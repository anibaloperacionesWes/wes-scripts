"""
Plan de regulación de agua — colegios Puente Alto con consumo nocturno.

Lee la muestra comparativa semanal y el comparativo LXM para proponer acciones
de optimización de corte parcial en horarios sin demanda hidráulica.

Uso:
  python plan_regulacion_nocturno_puente_alto.py
"""

from __future__ import annotations

import csv
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, OSError, ValueError):
            pass

ROOT = Path(__file__).resolve().parent
MUESTRA_DIR = ROOT / "reports" / "proyeccion ahorre puente 2025" / "muestra_semanal"
LXM_CSV = ROOT / "reports" / "puente_alto_lxm" / "comparativo_lxm_puente_alto_20260322_20260323_1723.csv"
OUT_DIR = MUESTRA_DIR

# Horarios de corte programados (Excel corporación / comparativo LXM)
HORARIOS_CORTE: Dict[str, str] = {
    "000010-01": "Corte nocturno 00:01–07:00 | Alto 07:00–15:00 | Medio 15:00–23:59",
    "000010-02": "Corte nocturno 23:00–06:00 | Medio 19:00–23:00 | Alto 06:00–19:00",
    "000010-03": "Corte nocturno 20:00–06:00 | Medio 19:00–20:00 | Alto 06:00–19:00",
    "000010-04": "Corte nocturno 23:00–06:00 | Medio 20:00–23:00 | Alto 06:00–20:00",
    "000010-05": "Config. atípica: banda única 00:01–23:59 (revisar ON/OFF)",
    "000010-06": "Corte nocturno 19:30–06:00 | Medio 18:00–19:30 | Alto 06:30–18:00",
    "000010-07": "Corte nocturno 20:00–06:30 | Medio 15:20–20:00 | Alto 06:30–15:20",
    "000010-08": "Config. atípica: banda «nocturno» 00:01–23:59 (revisar ON/OFF)",
    "000010-09": "Corte nocturno 23:00–06:30 | Medio 17:15–23:00 | Alto 06:30–17:15",
    "000010-10": "Corte nocturno 00:01–07:00 | Medio 21:00–23:59 | Alto 07:00–21:00",
    "000010-11": "Corte nocturno 23:00–06:30 | Medio 19:30–23:00 | Alto 06:30–19:30",
}


def _ultima_muestra() -> Path:
    files = sorted(MUESTRA_DIR.glob("muestra_con_sin_wes_pa_*.xlsx"))
    if not files:
        raise FileNotFoundError("Ejecute primero muestra_comparativa_con_sin_wes_puente_alto_semanas.py")
    return files[-1]


def _cargar_lxm_nocturno() -> Dict[str, str]:
    out: Dict[str, str] = {}
    if not LXM_CSV.is_file():
        return out
    with LXM_CSV.open(encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            if "nocturno" not in str(row.get("rango", "")).lower():
                continue
            nid = str(row.get("node_id", "")).strip()
            cumple = str(row.get("cumple", "")).strip().upper()
            nota = str(row.get("nota", "")).strip()
            out[nid] = "Cumple LXM" if cumple == "SI" else f"No cumple LXM ({nota})"
    return out


def _prioridad(
    noct_rec: float,
    pct_noct: float,
    dias_con_noct: int,
    delta_noct: float,
    lxm: str,
    total_rec: float,
) -> str:
    if noct_rec >= 100 or (pct_noct >= 28 and dias_con_noct >= 6):
        return "P1 — Crítica"
    if "No cumple" in lxm or delta_noct >= 30:
        return "P1 — Crítica"
    if noct_rec >= 40 or (dias_con_noct >= 6 and pct_noct >= 20):
        return "P2 — Alta"
    if noct_rec >= 15 or dias_con_noct >= 5:
        return "P3 — Media"
    if total_rec < 30:
        return "P4 — Baja actividad"
    return "P4 — Seguimiento"


def _acciones(
    nid: str,
    colegio: str,
    prioridad: str,
    noct_rec: float,
    pct_noct: float,
    dias_con_noct: int,
    delta_noct: float,
    lxm: str,
) -> List[str]:
    acts: List[str] = []
    hor = HORARIOS_CORTE.get(nid, "Verificar horario en Excel corporación")

    if dias_con_noct >= 6:
        acts.append(
            "Auditoría válvula WES: confirmar cierre efectivo 100% en ventana nocturna; "
            "revisar bypass manual o derivaciones paralelas a la red monitoreada."
        )

    if "No cumple LXM" in lxm or nid in ("000010-04", "000010-08"):
        acts.append(
            "Ajuste inmediato de horario ON/OFF: adelantar cierre 30–45 min y retrasar apertura matinal "
            "hasta validar caudal cero en madrugada."
        )

    if nid in ("000010-05", "000010-08"):
        acts.append(
            "Reconfigurar bandas horarias en plataforma WES: separar rangos Alto/Medio/Bajo; "
            "la banda «nocturno» no debe cubrir casi todo el día."
        )

    if delta_noct >= 20:
        acts.append(
            f"Consumo nocturno aumentó {delta_noct:+.1f} m³ vs mayo: inspección in situ de fugas "
            "(inodoros, llaves de jardín, calderas) en ventana 02:00–05:00."
        )
    elif delta_noct <= -15 and noct_rec >= 80:
        acts.append(
            "Aunque el nocturno bajó vs mayo, el volumen absoluto sigue alto: priorizar corte total "
            "en fin de semana y feriados."
        )

    if pct_noct >= 25:
        acts.append(
            "Evaluar corte parcial escalonado: red interior (aulas/baños) OFF nocturno; "
            "mantener solo estanques/riego si hay demanda justificada documentada."
        )

    if nid == "000010-02":
        acts.append(
            "Caso Villa Independencia: correlacionar pico nocturno con horario de aseo nocturno o "
            "riego automático; reprogramar o excluir circuito de riego del bypass."
        )

    if nid == "000010-09":
        acts.append(
            "Los Andes: consumo diurno reciente elevado con nocturno moderado — revisar uso fuera de "
            "horario escolar (talleres, riego) antes de endurecer solo el corte nocturno."
        )

    if nid == "000010-07":
        acts.append(
            "Juan Mackenna: caudal nocturno cercano al LXM (pico ~03:00); inspeccionar flotadores y "
            "cisternas que recirculan en madrugada."
        )

    if noct_rec < 5 and dias_con_noct <= 2:
        acts.append(
            "Establecimiento con baja actividad reciente: al reactivar clases, validar que el perfil "
            "ON/OFF se restaure antes del inicio de clases."
        )

    if not acts:
        acts.append("Monitoreo mensual; mantener horario actual y verificar cumplimiento LXM.")

    acts.append(f"Horario referencia: {hor}")
    return acts


def _hallazgos_globales(df: pd.DataFrame) -> List[str]:
    n = len(df)
    siete_dias = int((df["dias_con_noct"] >= 6).sum())
    alto_vol = int((df["noct_rec"] >= 80).sum())
    subio = int((df["delta_noct"] > 10).sum())
    return [
        f"De {n} establecimientos analizados, {siete_dias} presentan consumo nocturno en 6 o más días "
        f"de la semana reciente (9–15 jun 2026): el corte programado no está dejando la red en cero.",
        f"{alto_vol} colegios superan 80 m³ de consumo nocturno semanal medido — volumen incompatible "
        f"con recintos sin demanda hidráulica.",
        f"En {subio} colegios el nocturno reciente supera al de la última semana de mayo (con WES), "
        f"lo que sugiere regresión del corte o aparición de fugas.",
        "Colegio Maipo y Liceo Chiloé figuran con incumplimiento LXM en banda nocturna (marzo 2026); "
        "Luis Matte Larraín y Liceo Chiloé tienen bandas horarias mal parametrizadas en el Excel corporación.",
        "Esc. Gabriela y Esc. Padre Hurtado muestran actividad reciente muy baja: no usar su semana reciente "
        "como línea base de comparación hasta normalizar operación.",
        "Potencial de optimización municipal estimado: reducir 30–50% del nocturno actual en los 5 colegios "
        "P1/P2 (~400–500 m³/semana) mediante corte efectivo + reparación de fugas.",
    ]


def generar_plan() -> Path:
    muestra = _ultima_muestra()
    df_cmp = pd.read_excel(muestra, header=4)
    df_cmp = df_cmp[df_cmp["Node ID"].notna() & ~df_cmp["Establecimiento"].astype(str).str.contains("TOTAL", na=False)]

    analisis = MUESTRA_DIR / "analisis_nocturno_tmp.csv"
    if analisis.is_file():
        df = pd.read_csv(analisis, encoding="utf-8-sig")
    else:
        raise FileNotFoundError("Falta analisis_nocturno_tmp.csv; ejecute el análisis previo.")

    delta_map = {}
    for _, r in df_cmp.iterrows():
        nid = str(r["Node ID"]).strip()
        try:
            delta_map[nid] = float(r["Δ nocturno (rec − mayo)"])
        except (TypeError, ValueError, KeyError):
            delta_map[nid] = 0.0
        total_map = {}
    for _, r in df_cmp.iterrows():
        nid = str(r["Node ID"]).strip()
        total_map[nid] = float(r["Con WES reciente (m³)"])

    lxm = _cargar_lxm_nocturno()

    filas_plan: List[Dict[str, object]] = []
    for _, row in df.sort_values("noct_rec", ascending=False).iterrows():
        nid = str(row["node"]).strip()
        noct_rec = float(row["noct_rec"])
        pct = float(row["pct_noct_rec"])
        dcn = int(row["dias_con_noct"])
        delta = delta_map.get(nid, 0.0)
        total_rec = total_map.get(nid, 0.0)
        lx = lxm.get(nid, "Sin dato LXM")
        pri = _prioridad(noct_rec, pct, dcn, delta, lx, total_rec)
        acc = _acciones(nid, str(row["colegio"]), pri, noct_rec, pct, dcn, delta, lx)
        filas_plan.append(
            {
                "prioridad": pri,
                "colegio": row["colegio"],
                "node": nid,
                "noct_rec": noct_rec,
                "pct_noct": pct,
                "dias_con_noct": dcn,
                "delta_noct": delta,
                "lxm": lx,
                "acciones": acc,
            }
        )

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    out_docx = OUT_DIR / f"Plan_Regulacion_Agua_Nocturno_Puente_Alto_{ts}.docx"
    out_xlsx = OUT_DIR / f"Plan_Regulacion_Agua_Nocturno_Puente_Alto_{ts}.xlsx"

    doc = Document()
    t = doc.add_heading("Plan de regulación de agua — Corporación Puente Alto", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Optimización de detención parcial de agua en horarios sin demanda hidráulica "
        "(consumo nocturno con WES activo)"
    )
    doc.add_paragraph(f"Elaborado: {datetime.now():%d/%m/%Y %H:%M}")
    doc.add_paragraph(
        "Base: muestra comparativa 25–31 may 2026 vs 9–15 jun 2026; horarios de corte Excel corporación; "
        "validación LXM (mar 2026)."
    )

    doc.add_heading("1. Hallazgos que requieren acción", level=1)
    for h in _hallazgos_globales(df):
        doc.add_paragraph(h, style="List Bullet")

    doc.add_heading("2. Criterios de priorización", level=1)
    for txt in [
        "P1 — Crítica: ≥100 m³ nocturnos/semana, ≥28% del total, 6–7 días con madrugada activa, "
        "incumplimiento LXM o aumento >30 m³ vs mayo.",
        "P2 — Alta: ≥40 m³ nocturnos o patrón persistente 5+ días con >20% nocturno.",
        "P3 — Media: volumen moderado pero mejorable con ajuste fino de horarios.",
        "P4 — Seguimiento / baja actividad: recinto con operación mínima reciente.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_heading("3. Plan por establecimiento", level=1)

    for item in filas_plan:
        doc.add_heading(f"{item['colegio']} ({item['node']}) — {item['prioridad']}", level=2)
        p = doc.add_paragraph()
        p.add_run("Indicadores semana reciente (9–15 jun): ").bold = True
        p.add_run(
            f"Nocturno {item['noct_rec']:.1f} m³ ({item['pct_noct']:.1f}% del total); "
            f"{item['dias_con_noct']}/7 días con consumo en madrugada; "
            f"variación vs mayo {item['delta_noct']:+.1f} m³. {item['lxm']}."
        )
        doc.add_paragraph("Acciones recomendadas:", style="List Bullet")
        for a in item["acciones"]:
            doc.add_paragraph(a, style="List Bullet 2")

    doc.add_heading("4. Acciones transversales (municipal / WES)", level=1)
    transversales = [
        "Campaña de verificación física en los 5 establecimientos P1/P2 dentro de 15 días: "
        "válvula motorizada, estado ON/OFF en plataforma vs caudal real.",
        "Alerta automática WES si caudal > umbral (ej. 2 L/min) entre 00:00 y 06:00 por más de 2 horas consecutivas.",
        "Estandarizar bandas horarias en plataforma (Alto / Medio / Corte total) alineadas al Excel corporación; "
        "eliminar configuraciones «nocturno 24 h».",
        "Registro fotográfico de hallazgos en recorrida y cierre de plan con fecha; re-medición semanal 7 días post-intervención.",
        "Capacitación a conserjes: no abrir bypass manual salvo emergencia documentada.",
    ]
    for t in transversales:
        doc.add_paragraph(t, style="List Number")

    doc.add_heading("5. Seguimiento y KPI", level=1)
    doc.add_paragraph(
        "Meta municipal a 4 semanas: reducir consumo nocturno agregado de ~723 m³/semana a <450 m³/semana "
        "(−38%) en los colegios P1/P2, manteniendo consumo diurno para actividad escolar."
    )
    doc.add_paragraph("KPI por colegio: m³ nocturno semanal, % días con cero en 00:00–06:00, cumplimiento LXM, alertas WES.")

    doc.save(str(out_docx))

    # Excel operativo
    rows_x: List[Dict[str, object]] = []
    for item in filas_plan:
        rows_x.append(
            {
                "Prioridad": item["prioridad"],
                "Establecimiento": item["colegio"],
                "Node ID": item["node"],
                "Nocturno reciente (m³)": round(float(item["noct_rec"]), 1),
                "% nocturno": round(float(item["pct_noct"]), 1),
                "Días con nocturno /7": item["dias_con_noct"],
                "Δ noct vs mayo (m³)": round(float(item["delta_noct"]), 1),
                "LXM nocturno": item["lxm"],
                "Acciones (resumen)": " | ".join(str(a) for a in item["acciones"][:3]),
            }
        )
    pd.DataFrame(rows_x).to_excel(out_xlsx, index=False, sheet_name="Plan")

    print(f"[OK] Word: {out_docx}")
    print(f"[OK] Excel: {out_xlsx}")
    return out_docx


if __name__ == "__main__":
    generar_plan()
