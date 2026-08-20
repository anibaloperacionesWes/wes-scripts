"""
Renca — rendimiento hídrico con WES vs periodo sin control (ago-2026).

Contexto operativo: desde el lunes 17/08/2026 los equipos de ICCP, Esc. Lo Velásquez,
Gimnasio y Piscina Municipal quedaron sin control WES. Este script:

  1. Toma el periodo sin WES (lunes 17/08 hasta la última hora completa Chile de «ahora»).
  2. Replica la misma ventana horaria (lun + mar + mié hasta esa hora) en cada semana
     lunes–domingo desde mayo 2026 (semanas con control WES).
  3. Calcula rendimiento de ahorro: (Sin WES − Con WES) / Sin WES × 100.
  4. Identifica la semana de mayo en adelante con mayor rendimiento (m³ y %).
  5. Genera Excel, gráficos, Word y PDF.

Uso:
  python generar_comparacion_renca_sin_wes_ago17_vs_semanas_mayo.py
"""

from __future__ import annotations

import csv
import shutil
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple
from zoneinfo import ZoneInfo

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from generar_reporte_word import (
    acl_node_base_url,
    add_logo_to_header,
    fetch_json,
    flatten_measures,
    format_number_chilean,
    get_hourly_measures_for_day,
    normalize_measures_payload,
)

ROOT = Path(__file__).resolve().parent
TZ_CL = ZoneInfo("America/Santiago")
TARIFA_CLP_M3 = 1300.0
COLOR_WES = "#1F4788"
COLOR_SIN = "#C0504D"
COLOR_AHORRO = "#548235"
COLOR_HEAD = RGBColor(31, 71, 136)

PUNTOS: Tuple[Tuple[str, str], ...] = (
    ("000017-07", "ICCP (Cumbre de Cóndores pte.)"),
    ("000017-04", "Esc. Lo Velásquez"),
    ("000017-05", "Gimnasio municipal"),
    ("000017-06", "Piscina municipal"),
)

# Primera semana completa de mayo 2026 (lunes 4) hasta última semana con WES (lun 10–dom 16 ago).
PRIMER_LUNES_CON_WES = date(2026, 5, 4)
ULTIMO_LUNES_CON_WES = date(2026, 8, 10)
LUNES_SIN_WES = date(2026, 8, 17)


def _ahora_chile() -> datetime:
    return datetime.now(TZ_CL)


def _ultima_hora_completa(ahora: datetime) -> int:
    """Hora Chile 0–23 cuya franja [H:00, H+1:00) ya cerró. Si son las 13:31 → 12."""
    h = ahora.hour - 1
    return max(0, min(23, h))


def _lunes_semana(d: date) -> date:
    return d - timedelta(days=d.weekday())


def _iter_lunes(desde: date, hasta: date) -> List[date]:
    out: List[date] = []
    d = _lunes_semana(desde)
    if d < desde:
        d += timedelta(days=7)
    while d <= hasta:
        out.append(d)
        d += timedelta(days=7)
    return out


def _etiqueta_semana(lunes: date) -> str:
    domingo = lunes + timedelta(days=6)
    return f"{lunes:%d/%m}–{domingo:%d/%m}"


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shading = parse_xml(
        "<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
        f'w:val="clear" w:fill="{hex_fill}"/>'
    )
    tc_pr.append(shading)


def _fetch_diario(node_id: str, start: date, end: date) -> Dict[date, float]:
    raw = fetch_json(
        f"{acl_node_base_url()}/nodes/measures/dates",
        params=[
            ("id", node_id),
            ("start", start.strftime("%d%m%Y")),
            ("end", end.strftime("%d%m%Y")),
        ],
    )
    payload = normalize_measures_payload(raw, node_id)
    out: Dict[date, float] = {}
    for m in flatten_measures(payload):
        d = m.date.date()
        if start <= d <= end:
            out[d] = float(m.total_m3)
    return out


def _vector_24h(node_id: str, dia: date) -> List[float]:
    hourly = get_hourly_measures_for_day(node_id, datetime.combine(dia, datetime.min.time())) or []
    por_hora = {int(h): float(v) for h, v in hourly}
    return [float(por_hora.get(h, 0.0)) for h in range(24)]


def _suma_horas(vec: Sequence[float], h_ini: int, h_fin_incl: int) -> float:
    return float(sum(float(vec[h]) for h in range(h_ini, h_fin_incl + 1) if 0 <= h < len(vec)))


@dataclass
class VentanaPunto:
    node_id: str
    nombre: str
    lunes: date
    con_wes: bool
    m3_lunes: float
    m3_martes: float
    m3_miercoles_corte: float
    m3_jueves_corte: float
    m3_total: float
    m3_nocturno: float  # 00–06 del lun+mar + mié hasta min(6, corte)
    horas_equivalentes: int
    vec_miercoles: List[float] = field(default_factory=list)
    nota: str = ""


@dataclass
class SemanaPunto:
    node_id: str
    nombre: str
    lunes: date
    con_wes: bool
    m3_semana: float
    dias_con_dato: int
    m3_dia: float


def _consumo_ventana(
    node_id: str,
    nombre: str,
    lunes: date,
    diario: Dict[date, float],
    vec_parcial: Optional[List[float]],
    offsets_completos: Sequence[int],
    offset_parcial: Optional[int],
    hora_corte: int,
    con_wes: bool,
) -> VentanaPunto:
    m3_off = [0.0] * 7
    horas = 0
    for off in offsets_completos:
        if 0 <= off < 7:
            m3_off[off] = float(diario.get(lunes + timedelta(days=off), 0.0))
            horas += 24
    if offset_parcial is not None and vec_parcial is not None and 0 <= offset_parcial < 7:
        m3_off[offset_parcial] = _suma_horas(vec_parcial, 0, hora_corte)
        horas += hora_corte + 1
    noct = 0.0
    if offset_parcial is not None and vec_parcial is not None:
        noct = _suma_horas(vec_parcial, 0, min(6, hora_corte))
    return VentanaPunto(
        node_id=node_id,
        nombre=nombre,
        lunes=lunes,
        con_wes=con_wes,
        m3_lunes=m3_off[0],
        m3_martes=m3_off[1],
        m3_miercoles_corte=m3_off[2],
        m3_jueves_corte=m3_off[3],
        m3_total=float(sum(m3_off)),
        m3_nocturno=noct,
        horas_equivalentes=horas,
        vec_miercoles=list(vec_parcial or []),
    )


def _semana_completa(
    node_id: str,
    nombre: str,
    lunes: date,
    diario: Dict[date, float],
    con_wes: bool,
    hasta: Optional[date] = None,
) -> SemanaPunto:
    total = 0.0
    n = 0
    for i in range(7):
        d = lunes + timedelta(days=i)
        if hasta is not None and d > hasta:
            break
        if d in diario:
            total += float(diario[d])
            n += 1
        else:
            n += 1  # día sin fila API = 0 m³, cuenta para el promedio de 7 días
    dias = 7 if hasta is None else max(1, sum(1 for i in range(7) if lunes + timedelta(days=i) <= hasta))
    return SemanaPunto(
        node_id=node_id,
        nombre=nombre,
        lunes=lunes,
        con_wes=con_wes,
        m3_semana=total,
        dias_con_dato=n,
        m3_dia=total / dias if dias else 0.0,
    )


def _rendimiento(m3_sin: float, m3_con: float) -> Tuple[float, float]:
    ahorro = m3_sin - m3_con
    pct = (ahorro / m3_sin * 100.0) if m3_sin > 1e-9 else 0.0
    return ahorro, pct


def _mediana_positiva(vals: Sequence[float]) -> float:
    import statistics

    pos = [float(v) for v in vals if v > 1.0]
    if not pos:
        pos = [float(v) for v in vals]
    return float(statistics.median(pos)) if pos else 0.0


def _clasificar_ocupacion(m3: float, mediana: float) -> str:
    """receso / baja ocupación no se cuentan como el mejor rendimiento WES."""
    if mediana <= 1e-6:
        return "sin_referencia"
    if m3 < 0.20 * mediana:
        return "receso"
    if m3 < 0.50 * mediana:
        return "baja_ocupacion"
    if m3 > 2.5 * mediana:
        return "outlier_alto"
    return "operativa"


def _grafico_barras_semanas(
    etiquetas: List[str],
    valores_con: List[float],
    valor_sin: float,
    titulo: str,
    ylabel: str,
    out_png: Path,
    idx_mejor: Optional[int] = None,
    etiqueta_sin: str = "Sin WES hasta ahora",
) -> None:
    fig, ax = plt.subplots(figsize=(13.2, 5.4))
    x = list(range(len(etiquetas)))
    colors = []
    for i in range(len(etiquetas)):
        if idx_mejor is not None and i == idx_mejor:
            colors.append(COLOR_AHORRO)
        else:
            colors.append("#5B9BD5")
    ax.bar(x, valores_con, color=colors, zorder=2, label="Con WES (esa semana)")
    ax.axhline(
        valor_sin,
        color=COLOR_SIN,
        linestyle="--",
        linewidth=2.0,
        label=f"{etiqueta_sin} ({format_number_chilean(valor_sin, 1)} m³)",
        zorder=3,
    )
    ax.set_xticks(x)
    ax.set_xticklabels(etiquetas, rotation=55, ha="right", fontsize=8)
    ax.set_ylabel(ylabel)
    ax.set_title(titulo, fontweight="bold", fontsize=11, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=8, loc="upper right")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_pareado_sin_vs_semanas(
    etiquetas: List[str],
    valores_con: List[float],
    valor_sin: float,
    titulo: str,
    out_png: Path,
    etiqueta_sin: str = "Sin WES hasta ahora",
    figsize: Tuple[float, float] = (13.4, 5.6),
) -> None:
    """Barras pareadas: Con WES de cada semana calificada vs Sin WES acumulado."""
    import numpy as np

    n = len(etiquetas)
    x = np.arange(n)
    w = 0.38
    fig, ax = plt.subplots(figsize=figsize)
    bars_con = ax.bar(
        x - w / 2,
        valores_con,
        width=w,
        color=COLOR_WES,
        label="Con WES (esa semana)",
        zorder=2,
    )
    bars_sin = ax.bar(
        x + w / 2,
        [valor_sin] * n,
        width=w,
        color=COLOR_SIN,
        label=etiqueta_sin,
        zorder=2,
    )
    ymax = max([valor_sin] + list(valores_con) + [1.0]) * 1.18
    ax.set_ylim(0, ymax)
    ax.set_xticks(list(x))
    ax.set_xticklabels(etiquetas, rotation=50, ha="right", fontsize=8)
    ax.set_ylabel("Consumo ventana homóloga (m³)")
    ax.set_title(titulo, fontweight="bold", fontsize=11, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=8, loc="upper right")
    for bar in list(bars_con) + list(bars_sin):
        h = bar.get_height()
        if h <= 0:
            continue
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            h + ymax * 0.012,
            f"{h:.0f}" if h >= 10 else f"{h:.1f}",
            ha="center",
            va="bottom",
            fontsize=6.5,
            color="#333333",
        )
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_semanas_mas_barra_sin(
    etiquetas: List[str],
    valores_con: List[float],
    valor_sin: float,
    titulo: str,
    out_png: Path,
    etiqueta_sin: str = "Sin WES\nhasta ahora",
    idx_recomendada: Optional[int] = None,
) -> None:
    """Todas las semanas con WES + una barra roja final con el acumulado sin control."""
    from matplotlib.patches import Patch

    labs = list(etiquetas) + [etiqueta_sin]
    vals = list(valores_con) + [valor_sin]
    colors = ["#5B9BD5"] * len(etiquetas) + [COLOR_SIN]
    if idx_recomendada is not None and 0 <= idx_recomendada < len(etiquetas):
        colors[idx_recomendada] = COLOR_AHORRO
    fig, ax = plt.subplots(figsize=(13.6, 5.6))
    x = list(range(len(labs)))
    bars = ax.bar(x, vals, color=colors, zorder=2)
    ax.axhline(valor_sin, color=COLOR_SIN, linestyle=":", linewidth=1.2, alpha=0.8, zorder=1)
    ax.legend(
        handles=[
            Patch(facecolor="#5B9BD5", label="Semanas con WES"),
            Patch(facecolor=COLOR_AHORRO, label="Mejor para comparar (10–16 ago)"),
            Patch(facecolor=COLOR_SIN, label="Sin WES desde el 17/08"),
        ],
        fontsize=8,
        loc="upper right",
    )
    ax.set_xticks(x)
    ax.set_xticklabels(labs, rotation=50, ha="right", fontsize=8)
    ax.set_ylabel("Consumo ventana homóloga (m³)")
    ax.set_title(titulo, fontweight="bold", fontsize=11, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ymax = max(vals + [1.0]) * 1.16
    ax.set_ylim(0, ymax)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            v + ymax * 0.012,
            f"{v:.0f}" if v >= 10 else f"{v:.1f}",
            ha="center",
            va="bottom",
            fontsize=7,
            color="#333333",
        )
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_diario_sin_wes(
    dias: Sequence[date],
    valores: Sequence[float],
    titulo: str,
    out_png: Path,
    hora_corte_hoy: Optional[int] = None,
) -> None:
    """Barras diarias del periodo sin control (desde el lunes 17)."""
    nomb = ("lun", "mar", "mié", "jue", "vie", "sáb", "dom")
    labels = []
    for i, d in enumerate(dias):
        lab = f"{nomb[d.weekday()]} {d:%d/%m}"
        if i == len(dias) - 1 and hora_corte_hoy is not None:
            lab = f"{lab}\nhasta {hora_corte_hoy:02d}:59"
        labels.append(lab)
    fig, ax = plt.subplots(figsize=(8.8, 4.2))
    x = list(range(len(dias)))
    bars = ax.bar(x, list(valores), color=COLOR_SIN, zorder=2)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylabel("m³")
    ax.set_title(titulo, fontweight="bold", fontsize=11, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ymax = max(list(valores) + [0.1]) * 1.22
    ax.set_ylim(0, ymax)
    for bar, v in zip(bars, valores):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            v + ymax * 0.02,
            format_number_chilean(float(v), 1),
            ha="center",
            va="bottom",
            fontsize=9,
            color="#333333",
        )
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_ranking_pct(etiquetas: List[str], pcts: List[float], titulo: str, out_png: Path) -> None:
    fig, ax = plt.subplots(figsize=(12.5, 5.2))
    x = list(range(len(etiquetas)))
    colors = [COLOR_AHORRO if v > 0 else COLOR_SIN for v in pcts]
    ax.bar(x, pcts, color=colors, zorder=2)
    ax.axhline(0, color="#333333", linewidth=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels(etiquetas, rotation=55, ha="right", fontsize=8)
    ax.set_ylabel("Rendimiento de ahorro (%)")
    ax.set_title(titulo, fontweight="bold", fontsize=11, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_resumen_puntos(
    nombres: List[str],
    m3_sin: List[float],
    m3_mejor: List[float],
    m3_mediana: List[float],
    etiqueta_mejor: str,
    out_png: Path,
) -> None:
    import numpy as np

    x = np.arange(len(nombres))
    w = 0.26
    fig, ax = plt.subplots(figsize=(11, 5.2))
    ax.bar(x - w, m3_sin, width=w, color=COLOR_SIN, label="Sin WES hasta ahora", zorder=2)
    ax.bar(x, m3_mejor, width=w, color=COLOR_AHORRO, label=f"Última semana con WES ({etiqueta_mejor})", zorder=2)
    ax.bar(x + w, m3_mediana, width=w, color="#5B9BD5", label="Mediana semanas operativas con WES", zorder=2)
    ax.set_xticks(list(x))
    ax.set_xticklabels(nombres, rotation=18, ha="right", fontsize=9)
    ax.set_ylabel("Consumo ventana homóloga (m³)")
    ax.set_title(
        "Renca — ventana lun–mié: sin control vs última semana con WES y mediana operativa",
        fontweight="bold",
        fontsize=11,
        color="#1F4788",
    )
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_serie_diaria(
    diario_por_nodo: Dict[str, Dict[date, float]],
    nombres: Dict[str, str],
    d0: date,
    d1: date,
    out_png: Path,
) -> None:
    fig, axes = plt.subplots(2, 2, figsize=(12.8, 7.2), sharex=True)
    dias = []
    d = d0
    while d <= d1:
        dias.append(d)
        d += timedelta(days=1)
    xs = list(range(len(dias)))
    for ax, (nid, nom) in zip(axes.ravel(), PUNTOS):
        ys = [float(diario_por_nodo.get(nid, {}).get(dia, 0.0)) for dia in dias]
        ax.plot(xs, ys, color=COLOR_WES, linewidth=1.1)
        # sombrear periodo sin WES
        i_sin = next((i for i, dia in enumerate(dias) if dia >= LUNES_SIN_WES), None)
        if i_sin is not None:
            ax.axvspan(i_sin - 0.5, len(dias) - 0.5, color=COLOR_SIN, alpha=0.18, label="Sin WES")
        ax.set_title(nom, fontsize=10, fontweight="bold", color="#1F4788")
        ax.set_ylabel("m³/día")
        ax.grid(alpha=0.3)
        ticks = [i for i, dia in enumerate(dias) if dia.day == 1 or dia == d0]
        ax.set_xticks(ticks)
        ax.set_xticklabels([dias[i].strftime("%d/%m") for i in ticks], fontsize=8)
    fig.suptitle("Consumo diario mayo–agosto 2026 (sombra = sin control WES)", fontsize=12, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_perfil_24h(
    vec_sin: List[float],
    vec_con: List[float],
    titulo: str,
    out_png: Path,
    hora_corte: Optional[int] = None,
) -> None:
    fig, ax = plt.subplots(figsize=(10.5, 4.4))
    x = list(range(24))
    ax.bar([i - 0.18 for i in x], vec_sin, width=0.36, color=COLOR_SIN, label="Sin WES", zorder=2)
    ax.bar([i + 0.18 for i in x], vec_con, width=0.36, color=COLOR_WES, label="Con WES", zorder=2)
    ax.axvspan(-0.5, 6.5, color="#c41e1e", alpha=0.10, zorder=0)
    if hora_corte is not None:
        ax.axvline(hora_corte + 0.5, color="#666666", linestyle=":", linewidth=1.2)
    ax.set_xticks(x)
    ax.set_xticklabels([f"{h:02d}" for h in x], fontsize=7)
    ax.set_ylabel("m³/h")
    ax.set_title(titulo, fontweight="bold", fontsize=11, color="#1F4788")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3, zorder=0)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _style_header_row(ws, row: int, cols: int) -> None:
    fill = PatternFill("solid", fgColor="1F4788")
    font = Font(color="FFFFFF", bold=True, name="Calibri", size=10)
    thin = Border(
        left=Side(style="thin", color="B0B0B0"),
        right=Side(style="thin", color="B0B0B0"),
        top=Side(style="thin", color="B0B0B0"),
        bottom=Side(style="thin", color="B0B0B0"),
    )
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", wrap_text=True, vertical="center")
        cell.border = thin


def _excel(
    out_xlsx: Path,
    diario_por_nodo: Dict[str, Dict[date, float]],
    ventanas: Dict[str, List[VentanaPunto]],
    ventana_sin: Dict[str, VentanaPunto],
    semanas: Dict[str, List[SemanaPunto]],
    hora_corte: int,
    ahora: datetime,
) -> None:
    wb = Workbook()
    thin = Border(
        left=Side(style="thin", color="B0B0B0"),
        right=Side(style="thin", color="B0B0B0"),
        top=Side(style="thin", color="B0B0B0"),
        bottom=Side(style="thin", color="B0B0B0"),
    )

    # Diario
    ws = wb.active
    ws.title = "Consumo_diario"
    headers = ["Fecha", "Día"] + [n for _, n in PUNTOS] + ["Total 4 puntos (m³)"]
    ws.append(headers)
    _style_header_row(ws, 1, len(headers))
    d0 = date(2026, 5, 1)
    d1 = ahora.date()
    d = d0
    nomb_dia = ("lun", "mar", "mié", "jue", "vie", "sáb", "dom")
    while d <= d1:
        vals = [float(diario_por_nodo.get(nid, {}).get(d, 0.0)) for nid, _ in PUNTOS]
        ws.append([d.isoformat(), nomb_dia[d.weekday()]] + [round(v, 4) for v in vals] + [round(sum(vals), 4)])
        d += timedelta(days=1)
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=len(headers)):
        for c in row:
            c.border = thin
            c.font = Font(name="Calibri", size=10)

    # Ventana homóloga
    ws2 = wb.create_sheet("Ventana_homologa")
    h2 = [
        "Nodo",
        "Establecimiento",
        "Semana (lun–dom)",
        "Control",
        "m³ lunes",
        "m³ martes",
        "m³ miércoles",
        "m³ jueves (hasta hora corte)",
        "m³ ventana",
        "m³ Sin WES (misma ventana)",
        "Ahorro m³",
        "Rendimiento %",
        "Ahorro CLP",
        "Ocupación",
    ]
    ws2.append(h2)
    _style_header_row(ws2, 1, len(h2))
    for nid, nom in PUNTOS:
        sin = ventana_sin[nid]
        for v in ventanas[nid]:
            ahorro, pct = _rendimiento(sin.m3_total, v.m3_total)
            ws2.append(
                [
                    nid,
                    nom,
                    _etiqueta_semana(v.lunes),
                    "Con WES",
                    round(v.m3_lunes, 4),
                    round(v.m3_martes, 4),
                    round(v.m3_miercoles_corte, 4),
                    round(v.m3_jueves_corte, 4),
                    round(v.m3_total, 4),
                    round(sin.m3_total, 4),
                    round(ahorro, 4),
                    round(pct, 2),
                    round(ahorro * TARIFA_CLP_M3, 0),
                    v.nota or "operativa",
                ]
            )
        ws2.append(
            [
                nid,
                nom,
                _etiqueta_semana(sin.lunes),
                "Sin WES",
                round(sin.m3_lunes, 4),
                round(sin.m3_martes, 4),
                round(sin.m3_miercoles_corte, 4),
                round(sin.m3_jueves_corte, 4),
                round(sin.m3_total, 4),
                round(sin.m3_total, 4),
                0,
                0,
                0,
                "actual",
            ]
        )
    for row in ws2.iter_rows(min_row=2, max_row=ws2.max_row, max_col=len(h2)):
        for c in row:
            c.border = thin
            c.font = Font(name="Calibri", size=10)

    # Semanas completas
    ws3 = wb.create_sheet("Semanas_completas")
    h3 = [
        "Nodo",
        "Establecimiento",
        "Semana",
        "Control",
        "m³ semana (lun–dom)",
        "m³/día",
        "m³/día Sin WES (ritmo actual)",
        "Ahorro m³ vs ritmo sin WES × días",
        "Rendimiento % vs ritmo sin WES",
    ]
    ws3.append(h3)
    _style_header_row(ws3, 1, len(h3))
    horas_eq = next(iter(ventana_sin.values())).horas_equivalentes
    for nid, nom in PUNTOS:
        sin = ventana_sin[nid]
        ritmo = (sin.m3_total / horas_eq * 24.0) if horas_eq else 0.0
        for s in semanas[nid]:
            ahorro_est = ritmo * 7.0 - s.m3_semana
            pct = (ahorro_est / (ritmo * 7.0) * 100.0) if ritmo > 1e-9 else 0.0
            ws3.append(
                [
                    nid,
                    nom,
                    _etiqueta_semana(s.lunes),
                    "Con WES" if s.con_wes else "Sin WES (incompleta)",
                    round(s.m3_semana, 4),
                    round(s.m3_dia, 4),
                    round(ritmo, 4),
                    round(ahorro_est, 4),
                    round(pct, 2),
                ]
            )
    for row in ws3.iter_rows(min_row=2, max_row=ws3.max_row, max_col=len(h3)):
        for c in row:
            c.border = thin
            c.font = Font(name="Calibri", size=10)

    # Ranking
    ws4 = wb.create_sheet("Ranking_mejor_semana")
    h4 = [
        "Establecimiento",
        "Mejor semana (mayor %)",
        "m³ con WES",
        "m³ sin WES (hasta ahora)",
        "Ahorro m³",
        "Rendimiento %",
        "Ahorro CLP",
        "Mejor semana (mayor m³)",
        "Ahorro m³ (máximo volumen)",
        "Rendimiento % de esa semana",
    ]
    ws4.append(h4)
    _style_header_row(ws4, 1, len(h4))
    for nid, nom in PUNTOS:
        sin = ventana_sin[nid]
        ranked = []
        for v in ventanas[nid]:
            ahorro, pct = _rendimiento(sin.m3_total, v.m3_total)
            ranked.append((pct, ahorro, v))
        ops = [t for t in ranked if t[2].nota == "operativa"]
        pool = ops or ranked
        best_pct = max(pool, key=lambda t: t[0])
        best_m3 = max(pool, key=lambda t: t[1])
        ws4.append(
            [
                nom,
                _etiqueta_semana(best_pct[2].lunes),
                round(best_pct[2].m3_total, 4),
                round(sin.m3_total, 4),
                round(best_pct[1], 4),
                round(best_pct[0], 2),
                round(best_pct[1] * TARIFA_CLP_M3, 0),
                _etiqueta_semana(best_m3[2].lunes),
                round(best_m3[1], 4),
                round(best_m3[0], 2),
            ]
        )
    for row in ws4.iter_rows(min_row=2, max_row=ws4.max_row, max_col=len(h4)):
        for c in row:
            c.border = thin
            c.font = Font(name="Calibri", size=10)

    ws5 = wb.create_sheet("Metadatos")
    ws5["A1"] = "Periodo sin WES"
    ws5["B1"] = f"{LUNES_SIN_WES:%d/%m/%Y} 00:00 Chile hasta {ahora:%d/%m/%Y %H:%M} Chile"
    ws5["A2"] = "Última hora completa incluida (día de hoy)"
    ws5["B2"] = f"{hora_corte:02d}:00–{hora_corte:02d}:59"
    ws5["A3"] = "Horas equivalentes de la ventana"
    ws5["B3"] = next(iter(ventana_sin.values())).horas_equivalentes
    ws5["A4"] = "Semanas con WES"
    ws5["B4"] = f"{PRIMER_LUNES_CON_WES:%d/%m/%Y} a {ULTIMO_LUNES_CON_WES:%d/%m/%Y} (lunes de cada semana)"
    ws5["A5"] = "Tarifa CLP/m³"
    ws5["B5"] = TARIFA_CLP_M3
    ws5["A6"] = "Fuente"
    ws5["B6"] = "API WES totalM3 diario + serie horaria dates.measures.csv (hora Chile)"
    ws5["A7"] = "Rendimiento"
    ws5["B7"] = "(m³ Sin WES − m³ Con WES) / m³ Sin WES × 100"

    for wsx in wb.worksheets:
        for col in wsx.columns:
            letter = get_column_letter(col[0].column)
            maxlen = 12
            for cell in col[:40]:
                maxlen = max(maxlen, min(42, len(str(cell.value or ""))))
            wsx.column_dimensions[letter].width = maxlen + 2

    wb.save(out_xlsx)


def _word(
    out_docx: Path,
    pngs: Dict[str, Path],
    ventana_sin: Dict[str, VentanaPunto],
    ventanas: Dict[str, List[VentanaPunto]],
    hora_corte: int,
    ahora: datetime,
    ranking_agregado: List[Tuple[date, float, float, float]],
    mejor_por_punto: Dict[str, Tuple[VentanaPunto, float, float]],
    mediana_por_punto: Dict[str, float],
) -> None:
    """Un punto = días sin WES desde el 17 + barras por semanas + tabla de mejor comparación."""
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    try:
        add_logo_to_header(doc)
    except Exception:
        pass

    h = doc.add_heading("Renca — sin WES desde el 17/08 vs semanas con WES", level=0)
    if h.runs:
        h.runs[0].font.color.rgb = COLOR_HEAD

    horas_eq = next(iter(ventana_sin.values())).horas_equivalentes
    nomb_off = ("lun", "mar", "mié", "jue", "vie", "sáb", "dom")
    offset_hoy = (ahora.date() - LUNES_SIN_WES).days
    dias_completos = [nomb_off[i] for i in range(offset_hoy) if i < 7]
    dia_parcial = nomb_off[offset_hoy] if 0 <= offset_hoy < 7 else ""
    if dias_completos:
        ventana_txt = f"{', '.join(dias_completos)} completos + {dia_parcial} hasta {hora_corte:02d}:59"
    else:
        ventana_txt = f"{dia_parcial} hasta {hora_corte:02d}:59"
    p = doc.add_paragraph()
    p.add_run("Periodo sin control: ").bold = True
    p.add_run(
        f"lunes 17/08 a {ahora:%d/%m/%Y %H:%M} Chile "
        f"({ventana_txt}; {horas_eq} h). "
        "Cada semana con WES se mide en las mismas horas. "
        "Mejor para comparar = semana previa (10–16 ago, misma estación). "
        "Mayor ahorro = mejor % entre semanas operativas (sin receso)."
    )

    def _pinta_encabezado(tbl, headers: List[str]) -> None:
        for j, hd in enumerate(headers):
            cell = tbl.rows[0].cells[j]
            cell.text = hd
            _set_cell_shading(cell, "1F4788")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    for nid, nom in PUNTOS:
        sin = ventana_sin[nid]
        v_best, ahorro_best, pct_best = mejor_por_punto[nid]
        prev = next(v for v in ventanas[nid] if v.lunes == ULTIMO_LUNES_CON_WES)
        ahorro_prev, pct_prev = _rendimiento(sin.m3_total, prev.m3_total)

        doc.add_heading(nom, level=1)
        doc.add_paragraph(
            f"Sin WES hasta ahora: {format_number_chilean(sin.m3_total, 1)} m³. "
            f"Mejor para comparar: {_etiqueta_semana(prev.lunes)} "
            f"({format_number_chilean(prev.m3_total, 1)} m³ con WES → "
            f"{format_number_chilean(pct_prev, 1)} %). "
            f"Mayor ahorro operativo: {_etiqueta_semana(v_best.lunes)} "
            f"({format_number_chilean(pct_best, 1)} %; {format_number_chilean(ahorro_best, 1)} m³)."
        )

        if f"diario_{nid}" in pngs:
            doc.add_picture(str(pngs[f"diario_{nid}"]), width=Cm(15.6))
        if f"semanas_{nid}" in pngs:
            doc.add_paragraph("")
            doc.add_picture(str(pngs[f"semanas_{nid}"]), width=Cm(16.2))

        filas = [
            v
            for v in ventanas[nid]
            if v.nota == "operativa" or v.lunes == ULTIMO_LUNES_CON_WES
        ]
        if not filas:
            filas = [v for v in ventanas[nid] if v.nota != "receso"]
        tbl = doc.add_table(rows=1 + len(filas), cols=6)
        tbl.style = "Table Grid"
        _pinta_encabezado(
            tbl,
            ["Semana", "Con WES (m³)", "Sin WES ahora (m³)", "Ahorro (m³)", "%", "Selección"],
        )
        for i, v in enumerate(filas, start=1):
            ahorro, pct = _rendimiento(sin.m3_total, v.m3_total)
            marca = ""
            if v.lunes == ULTIMO_LUNES_CON_WES and v.lunes == v_best.lunes:
                marca = "Mejor para comparar y mayor ahorro"
            elif v.lunes == ULTIMO_LUNES_CON_WES:
                marca = "Mejor para comparar"
            elif v.lunes == v_best.lunes:
                marca = "Mayor ahorro"
            row = tbl.rows[i]
            row.cells[0].text = _etiqueta_semana(v.lunes)
            row.cells[1].text = format_number_chilean(v.m3_total, 1)
            row.cells[2].text = format_number_chilean(sin.m3_total, 1)
            row.cells[3].text = format_number_chilean(ahorro, 1)
            row.cells[4].text = format_number_chilean(pct, 1) + " %"
            row.cells[5].text = marca
            if marca:
                for c in row.cells:
                    _set_cell_shading(c, "E2EFDA")
                    for run in c.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
            else:
                for c in row.cells:
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(9)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)

def _convertir_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=180,
    )
    return pdf_path if pdf_path.is_file() else None


def main() -> int:
    ahora = _ahora_chile()
    offset_hoy = (ahora.date() - LUNES_SIN_WES).days
    if offset_hoy < 0:
        print("[ERROR] El periodo sin WES aún no comienza.")
        return 1
    offsets_completos = list(range(0, offset_hoy))
    offset_parcial = offset_hoy
    hora_corte = _ultima_hora_completa(ahora)
    nomb_off = ("lun", "mar", "mié", "jue", "vie", "sáb", "dom")
    print(
        f"Ahora Chile: {ahora:%Y-%m-%d %H:%M} | "
        f"completos: {[nomb_off[i] for i in offsets_completos] or 'ninguno'} | "
        f"parcial {nomb_off[offset_parcial]} 00–{hora_corte:02d}"
    )

    d_ini = date(2026, 5, 1)
    d_fin = ahora.date()
    lunes_con = _iter_lunes(PRIMER_LUNES_CON_WES, ULTIMO_LUNES_CON_WES)

    print("Descargando totalM3 diario (4 nodos, mayo–hoy)...")
    diario_por_nodo: Dict[str, Dict[date, float]] = {}
    for nid, nom in PUNTOS:
        diario_por_nodo[nid] = _fetch_diario(nid, d_ini, d_fin)
        print(f"  {nid} {nom}: {len(diario_por_nodo[nid])} días")

    parcial_ids: List[Tuple[str, date]] = []
    for lunes in lunes_con:
        dia_p = lunes + timedelta(days=offset_parcial)
        for nid, _ in PUNTOS:
            parcial_ids.append((nid, dia_p))
    dia_parcial_actual = LUNES_SIN_WES + timedelta(days=offset_parcial)
    for nid, _ in PUNTOS:
        parcial_ids.append((nid, dia_parcial_actual))

    print(
        f"Descargando serie horaria de {len(parcial_ids)} "
        f"{nomb_off[offset_parcial]}es (corte 00–{hora_corte:02d})..."
    )
    vecs: Dict[Tuple[str, date], List[float]] = {}

    def _job(item: Tuple[str, date]) -> Tuple[Tuple[str, date], List[float]]:
        nid, dia = item
        return item, _vector_24h(nid, dia)

    with ThreadPoolExecutor(max_workers=8) as ex:
        futs = [ex.submit(_job, it) for it in parcial_ids]
        done = 0
        for fut in as_completed(futs):
            key, vec = fut.result()
            vecs[key] = vec
            done += 1
            if done % 10 == 0 or done == len(parcial_ids):
                print(f"  horarios {done}/{len(parcial_ids)}")

    ventanas: Dict[str, List[VentanaPunto]] = {nid: [] for nid, _ in PUNTOS}
    ventana_sin: Dict[str, VentanaPunto] = {}
    semanas: Dict[str, List[SemanaPunto]] = {nid: [] for nid, _ in PUNTOS}

    for nid, nom in PUNTOS:
        for lunes in lunes_con:
            dia_p = lunes + timedelta(days=offset_parcial)
            ventanas[nid].append(
                _consumo_ventana(
                    nid,
                    nom,
                    lunes,
                    diario_por_nodo[nid],
                    vecs[(nid, dia_p)],
                    offsets_completos,
                    offset_parcial,
                    hora_corte,
                    True,
                )
            )
            semanas[nid].append(_semana_completa(nid, nom, lunes, diario_por_nodo[nid], True))
        ventana_sin[nid] = _consumo_ventana(
            nid,
            nom,
            LUNES_SIN_WES,
            diario_por_nodo[nid],
            vecs[(nid, dia_parcial_actual)],
            offsets_completos,
            offset_parcial,
            hora_corte,
            False,
        )
        semanas[nid].append(
            _semana_completa(nid, nom, LUNES_SIN_WES, diario_por_nodo[nid], False, hasta=d_fin)
        )

    # Clasificar ocupación (receso ≠ rendimiento WES)
    for nid, _ in PUNTOS:
        med = _mediana_positiva([v.m3_total for v in ventanas[nid]])
        for v in ventanas[nid]:
            v.nota = _clasificar_ocupacion(v.m3_total, med)

    # Ranking agregado (suma 4 puntos) por semana con WES
    ranking_agregado: List[Tuple[date, float, float, float]] = []
    ranking_operativo: List[Tuple[date, float, float, float]] = []
    m3_sin_total = sum(ventana_sin[nid].m3_total for nid, _ in PUNTOS)
    for i, lunes in enumerate(lunes_con):
        m3_con = sum(ventanas[nid][i].m3_total for nid, _ in PUNTOS)
        ahorro, pct = _rendimiento(m3_sin_total, m3_con)
        ranking_agregado.append((lunes, ahorro, pct, m3_con))
        # Semana operativa agregada: piscina abierta y escuela no en receso
        # (un receso escolar infla el % sin ser ahorro del equipo).
        pisc = ventanas["000017-06"][i].nota
        esc = ventanas["000017-04"][i].nota
        if pisc == "operativa" and esc != "receso":
            ranking_operativo.append((lunes, ahorro, pct, m3_con))

    import statistics

    mejor_por_punto: Dict[str, Tuple[VentanaPunto, float, float]] = {}
    mediana_por_punto: Dict[str, float] = {}
    for nid, _ in PUNTOS:
        sin = ventana_sin[nid]
        scored = []
        for v in ventanas[nid]:
            ahorro, pct = _rendimiento(sin.m3_total, v.m3_total)
            scored.append((pct, ahorro, v))
        ops = [t for t in scored if t[2].nota == "operativa"]
        pool = ops or [t for t in scored if t[2].nota != "receso"] or scored
        best = max(pool, key=lambda t: (t[0], t[1]))
        mejor_por_punto[nid] = (best[2], best[1], best[0])
        mediana_por_punto[nid] = statistics.median(
            [v.m3_total for v in ventanas[nid] if v.nota == "operativa"] or [v.m3_total for v in ventanas[nid]]
        )

    ts = ahora.strftime("%Y%m%d_%H%M")
    out_dir = ROOT / "reports" / "reporte de auditoria" / f"comparacion_renca_sin_wes_ago17_vs_semanas_mayo_{ts}"
    out_dir.mkdir(parents=True, exist_ok=True)
    png_dir = out_dir / "graficos"
    png_dir.mkdir(exist_ok=True)

    pngs: Dict[str, Path] = {}
    dias_sin = []
    d = LUNES_SIN_WES
    while d <= ahora.date():
        dias_sin.append(d)
        d += timedelta(days=1)

    for nid, nom in PUNTOS:
        sin = ventana_sin[nid]
        vals_dia = []
        for i, dia in enumerate(dias_sin):
            if i < offset_hoy:
                vals_dia.append(float(diario_por_nodo[nid].get(dia, 0.0)))
            else:
                vals_dia.append(max(0.0, float(sin.m3_total) - sum(vals_dia)))
        pdia = png_dir / f"diario_sin_{nid}.png"
        _grafico_diario_sin_wes(
            dias_sin,
            vals_dia,
            f"{nom} — días sin WES desde el 17/08",
            pdia,
            hora_corte_hoy=hora_corte,
        )
        pngs[f"diario_{nid}"] = pdia

        filas_op = [
            v
            for v in ventanas[nid]
            if v.nota == "operativa" or v.lunes == ULTIMO_LUNES_CON_WES
        ]
        if not filas_op:
            filas_op = [v for v in ventanas[nid] if v.nota != "receso"]
        etq_op = [_etiqueta_semana(v.lunes) for v in filas_op]
        vals_op = [v.m3_total for v in filas_op]
        idx_rec = next((i for i, v in enumerate(filas_op) if v.lunes == ULTIMO_LUNES_CON_WES), None)
        psem = png_dir / f"semanas_{nid}.png"
        _grafico_semanas_mas_barra_sin(
            etq_op,
            vals_op,
            sin.m3_total,
            f"{nom} — semanas con WES vs sin WES hasta ahora (rojo)",
            psem,
            etiqueta_sin=f"Sin WES\n{LUNES_SIN_WES:%d/%m}–{ahora:%d/%m}",
            idx_recomendada=idx_rec,
        )
        pngs[f"semanas_{nid}"] = psem

    xlsx = out_dir / "comparacion_renca_sin_wes_ago17_vs_semanas_mayo.xlsx"
    _excel(xlsx, diario_por_nodo, ventanas, ventana_sin, semanas, hora_corte, ahora)

    csv_path = out_dir / "ranking_agregado_ventana_homologa.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(
            [
                "lunes_semana",
                "etiqueta",
                "m3_con_wes",
                "m3_sin_wes",
                "ahorro_m3",
                "rendimiento_pct",
                "ahorro_clp",
                "ocupacion_piscina",
            ]
        )
        nota_piscina = {v.lunes: v.nota for v in ventanas["000017-06"]}
        for lunes, ahorro, pct, m3_con in ranking_agregado:
            w.writerow(
                [
                    lunes.isoformat(),
                    _etiqueta_semana(lunes),
                    f"{m3_con:.4f}",
                    f"{m3_sin_total:.4f}",
                    f"{ahorro:.4f}",
                    f"{pct:.2f}",
                    f"{ahorro * TARIFA_CLP_M3:.0f}",
                    nota_piscina.get(lunes, ""),
                ]
            )

    docx_path = out_dir / f"Rendimiento_hidrico_Renca_sinWES_17ago_vs_semanas_mayo_{ts}.docx"
    _word(
        docx_path,
        pngs,
        ventana_sin,
        ventanas,
        hora_corte,
        ahora,
        ranking_agregado,
        mejor_por_punto,
        mediana_por_punto,
    )

    pdf_path = _convertir_pdf(docx_path)

    # Consola
    horas_eq = next(iter(ventana_sin.values())).horas_equivalentes
    print("=" * 72)
    print("RENCA — SIN WES 17-19 AGO vs SEMANAS CON WES (MAYO→16 AGO)")
    print("=" * 72)
    print(f"Ventana: lun–mié completos + {nomb_off[offset_parcial]} 00–{hora_corte:02d} | {horas_eq} h equivalentes")
    print(f"Sin WES conjunto: {m3_sin_total:.2f} m³")
    for nid, nom in PUNTOS:
        v, ahorro, pct = mejor_por_punto[nid]
        print(
            f"  {nom}: sin {ventana_sin[nid].m3_total:.1f} m³ | "
            f"mejor {_etiqueta_semana(v.lunes)} {v.m3_total:.1f} m³ | "
            f"{pct:.1f}% ({ahorro:.1f} m³)"
        )
    best_agg = max(ranking_operativo or ranking_agregado, key=lambda t: t[2])
    prev = ranking_agregado[-1]
    print(
        f"Mejor semana operativa conjunta: {_etiqueta_semana(best_agg[0])} | "
        f"{best_agg[2]:.1f}% | {best_agg[1]:.1f} m³"
    )
    print(
        f"Última semana con WES (10–16 ago) vs actual: {prev[2]:.1f}% | {prev[1]:.1f} m³"
    )
    print(f"XLSX: {xlsx}")
    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path or '(no convertido)'}")
    print(f"DIR:  {out_dir}")
    return 0


if __name__ == "__main__":
    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
            sys.stderr.reconfigure(encoding="utf-8")
        except Exception:
            pass
    raise SystemExit(main())
