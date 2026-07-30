"""
Consolida los hallazgos más importantes por mall — Parque Arauco (000025).

Usa la misma metodología que los Word agregados con bloque Q3:
obtener_datos_agregados + _parrafo_hallazgos_q3 por mall, y de ahí extrae
síntesis ejecutiva, top aportes y línea de puntos en cero.

Salida: un solo .docx en reports/Parque_Arauco/Consolidados/

Uso:
  python consolidar_hallazgos_pa_malls.py
  python consolidar_hallazgos_pa_malls.py --desde 11/03/2026 --hasta 31/03/2026
  python consolidar_hallazgos_pa_malls.py --solo-mall Kennedy
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from generar_pa_agregado_todos_puntos import _parrafo_hallazgos_q3
from generar_pa_todos_malls_hallazgos_y_enviar import _nodos_por_mall
from generar_reporte_word import format_number_chilean, get_company_name
from generar_reportes_y_ppt_mall_maipu import obtener_datos_agregados

SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR / "reports" / "Parque_Arauco" / "Consolidados"


def _carpetas_escritorio() -> List[Path]:
    home = Path.home()
    out: List[Path] = []
    one = home / "OneDrive" / "Desktop"
    loc = home / "Desktop"
    if one.is_dir():
        out.append(one.resolve())
    if loc.is_dir() and (not out or loc.resolve() != out[0]):
        out.append(loc.resolve())
    if not out:
        out.append(loc.resolve())
    return out


def _copiar_escritorio(src: Path, nombre: str) -> None:
    for desk in _carpetas_escritorio():
        try:
            desk.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, desk / nombre)
            print(f"[OK] Copia Escritorio: {desk / nombre}")
        except OSError as e:
            print(f"[AVISO] No se pudo copiar a {desk}: {e}")


def extraer_hallazgos_clave(q3: str) -> Dict[str, Any]:
    """A partir del texto Q3, saca síntesis, ranking y alerta de ceros."""
    sintesis = ""
    if "Síntesis ejecutiva" in q3:
        rest = q3.split("Síntesis ejecutiva", 1)[1]
        sintesis = rest.split("Distribución por punto", 1)[0].strip()
    top_lines: List[str] = []
    if "Los cinco mayores aportes acumulados del período son:" in q3:
        chunk = q3.split("Los cinco mayores aportes acumulados del período son:", 1)[1]
        chunk = chunk.split("Los puntos con menor consumo", 1)[0]
        for line in chunk.split("\n"):
            s = line.strip()
            if re.match(r"^\d+\)", s):
                top_lines.append(s)
    cero_line = ""
    for line in q3.split("\n"):
        if "Puntos con consumo acumulado cero" in line:
            cero_line = line.strip()
            break
    return {
        "sintesis": sintesis,
        "top_lines": top_lines[:5],
        "cero_line": cero_line,
    }


def build_consolidado_word(
    out_path: Path,
    periodo: str,
    filas: List[Tuple[str, List[str], Dict[str, Any], float, int]],
) -> None:
    doc = Document()
    t = doc.add_heading(
        "Informe consolidado — Hallazgos por mall (Parque Arauco)",
        level=0,
    )
    for r in t.runs:
        r.font.size = Pt(18)

    intro = doc.add_paragraph()
    intro.add_run(
        f"Generado el {datetime.now().strftime('%d-%m-%Y %H:%M')}. "
        f"Período analizado: {periodo}. "
        "Cada sección resume el mall con la síntesis ejecutiva del análisis de consumo agregado, "
        "los tres mayores aportes por punto y el indicador de puntos sin consumo en el período. "
        "Los valores se calculan al ejecutar este script (API WES), alineado con los reportes Word agregados por mall."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in intro.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_heading("Resumen por mall (totales)", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Mall"
    hdr[1].text = "Puntos"
    hdr[2].text = "Total m³ (período)"
    hdr[3].text = "Prom. m³/día agreg."
    for mall, nids, ext, total_m3, dias in filas:
        row = tbl.add_row().cells
        row[0].text = mall
        row[1].text = str(len(nids))
        row[2].text = format_number_chilean(total_m3, 1)
        prom = (total_m3 / dias) if dias else 0.0
        row[3].text = format_number_chilean(prom, 1)

    doc.add_paragraph("")

    for mall, nids, ext, total_m3, dias in filas:
        doc.add_heading(f"Mall {mall}", level=1)
        p_ids = doc.add_paragraph()
        p_ids.add_run(f"Puntos incluidos ({len(nids)}): {', '.join(nids)}")
        p_ids.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.add_heading("Hallazgo principal (síntesis ejecutiva)", level=2)
        s = ext.get("sintesis") or "—"
        ps = doc.add_paragraph(s)
        ps.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in ps.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

        doc.add_heading("Mayores aportes de consumo (top 3)", level=2)
        tops = ext.get("top_lines") or []
        for line in tops[:3]:
            doc.add_paragraph(line, style="List Bullet")
        if not tops:
            doc.add_paragraph("— Sin ranking disponible.", style="List Bullet")

        doc.add_heading("Seguimiento operativo", level=2)
        cz = ext.get("cero_line") or "—"
        doc.add_paragraph(cz)
        doc.add_paragraph("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    ap = argparse.ArgumentParser(description="Consolidar hallazgos PA por mall en un Word")
    ap.add_argument("--desde", default="11/03/2026", help="DD/MM/YYYY inicio")
    ap.add_argument("--hasta", default=None, help="DD/MM/YYYY fin (default: hoy)")
    ap.add_argument("--solo-mall", default="", help='Solo un mall, ej. "Kennedy"')
    ap.add_argument("--no-copia-escritorio", action="store_true")
    ap.add_argument("--abrir-carpeta", action="store_true", help="Abrir carpeta Consolidados (Windows)")
    args = ap.parse_args()

    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass

    hasta = args.hasta or datetime.now().strftime("%d/%m/%Y")
    periodo = f"{args.desde} a {hasta}"

    company_id = "000025"
    cname = get_company_name(company_id)
    por_mall = _nodos_por_mall(company_id, cname)
    if args.solo_mall.strip():
        clave = args.solo_mall.strip()
        if clave not in por_mall:
            print(f"[ERROR] Mall no encontrado: {clave!r}. Opciones: {list(por_mall.keys())}")
            return 1
        por_mall = {clave: por_mall[clave]}

    if not por_mall:
        print("[ERROR] Sin malls con nodos.")
        return 1

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[INFO] Carpeta salida: {OUT_DIR.resolve()}")
    print(f"[INFO] Período: {periodo}")
    print(f"[INFO] Malls: {list(por_mall.keys())}")

    filas: List[Tuple[str, List[str], Dict[str, Any], float, int]] = []

    for mall in sorted(por_mall.keys(), key=lambda x: x.casefold()):
        nids = por_mall[mall]
        print(f"\n[INFO] Mall {mall} — {len(nids)} nodos")
        datos = obtener_datos_agregados(nids, args.desde, hasta)
        agg = datos.get("aggregate_summary") or {}
        total_m3 = float(agg.get("total") or 0)
        dias = int(agg.get("dias") or 0)
        sujeto = (
            f"Parque Arauco Mall {mall} (puntos de monitoreo del mall incluidos en este reporte)"
        )
        q3 = _parrafo_hallazgos_q3(datos, periodo, alcance_sujeto=sujeto)
        ext = extraer_hallazgos_clave(q3)
        filas.append((mall, nids, ext, total_m3, dias))
        print(f"  Total mall: {format_number_chilean(total_m3, 1)} m³ | días agreg.: {dias}")

    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    out = OUT_DIR / f"Hallazgos_consolidados_PA_{stamp}.docx"
    build_consolidado_word(out, periodo, filas)
    print(f"\n[OK] Word consolidado: {out.resolve()}")
    if not args.no_copia_escritorio:
        _copiar_escritorio(out, "Parque_Arauco_Hallazgos_Consolidados_ULTIMO.docx")
    if sys.platform == "win32" and args.abrir_carpeta:
        try:
            os.startfile(OUT_DIR)
        except OSError:
            subprocess.run(["explorer", str(OUT_DIR)], check=False)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
