"""
Genera el borrador Word del informe Puente Alto al estilo **resumen ejecutivo consolidado**
(tabla KPI + gráfica comparativa dos barras + Top 10), con opción de **prefijar portada y
metodología** desde un .docx previo.

- **Secciones 2–3 — Resumen ejecutivo y gráficas:** tabla KPI, rango CLP y nota; 3.1 barras
  comparativo mensual; 3.2 Top 10 y gráfico ranking. Van **antes** de la sección 4.
- **Sección 4 en adelante:** detalle mensual, agregado, ranking completo y anexo.

Por defecto se toman las **dos primeras páginas** de
``Informe_Auditoria_WES_Puente_Alto_2025_BORRADOR.docx`` (saltos de página explícitos en el Word)
como portada + metodología. Las tablas se enlazan con su gráfico (**mantener con siguiente**)
para reducir cortes feos entre tabla y figura.

Ejemplo::

  python generar_borrador_resumen_ejecutivo_puente_alto.py --year 2025

  python generar_borrador_resumen_ejecutivo_puente_alto.py --plantilla "R:\\solo_portada.docx"
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
OUT_REPORTS = ROOT / "reports" / "proyeccion ahorre puente 2025"
OUT_GRAFICOS = OUT_REPORTS / "graficos"
PLANTILLA_DEFAULT = OUT_REPORTS / "plantilla_portada_metodologia_Puente_Alto.docx"
INFORME_BORRADOR_REF = OUT_REPORTS / "Informe_Auditoria_WES_Puente_Alto_2025_BORRADOR.docx"

_RE_BR_PAGE = re.compile(r"w:type=['\"]page['\"]")


def _oxml_element_has_page_break(el: object) -> bool:
    from lxml import etree

    s = etree.tostring(el, encoding="unicode")
    return _RE_BR_PAGE.search(s) is not None


def _plain_text_w_body_child(ch: Any) -> str:
    """Texto del hijo ``w:p`` directo del body (vacío si no es párrafo)."""
    from docx.oxml.ns import qn

    if not ch.tag.endswith("p"):
        return ""
    texts = ch.findall(".//%s" % qn("w:t"))
    return "".join((t.text or "") for t in texts).replace("\xa0", " ").strip()


def _recortar_plantilla_si_arrastra_inicio_seccion_2(body: Any) -> int:
    """
    Si la extracción por saltos de página dejó colado el bloque que empieza en «2. Resumen…»
    del borrador fuente, elimínalo hasta antes de ``sectPr`` para no duplicarlo al fusionar el cuerpo.
    Devuelve cuántos nodos se quitaron.
    """
    children = list(body)
    idx_sect: Optional[int] = None
    for i, ch in enumerate(children):
        if ch.tag.endswith("sectPr"):
            idx_sect = i
            break
    if idx_sect is None:
        idx_sect = len(children)
    tit = "2. Resumen ejecutivo consolidado"
    for i, ch in enumerate(children[:idx_sect]):
        if ch.tag.endswith("p") and _plain_text_w_body_child(ch) == tit:
            to_rem = children[i:idx_sect]
            for el in reversed(to_rem):
                body.remove(el)
            return len(to_rem)
    return 0


def extraer_primeras_paginas_docx(src: Path, hasta_saltos_explicitos: int = 2) -> Path:
    """
    Copia bloques XML del ``Document`` hasta incluir el párrafo/tabla donde ocurre el salto
    ``w:br w:type='page'`` número ``hasta_saltos_explicitos`` (coincide con páginas del borrador WES).
    Guarda un .docx temporal; el llamador debe borrarlo tras fusionar.
    """
    src_doc = Document(str(src))
    dst_doc = Document()
    dst_body = dst_doc.element.body
    for child in list(dst_body):
        dst_body.remove(child)
    saltos = 0
    for child in list(src_doc.element.body):
        dst_body.append(deepcopy(child))
        if _oxml_element_has_page_break(child):
            saltos += 1
            if saltos >= hasta_saltos_explicitos:
                break
    _recortar_plantilla_si_arrastra_inicio_seccion_2(dst_body)
    fd, tmp = tempfile.mkstemp(suffix=".docx", prefix="plantilla_2pag_pa_")
    os.close(fd)
    out_p = Path(tmp)
    dst_doc.save(str(out_p))
    return out_p


def _para_keep_next(doc: Document, texto: str) -> None:
    p = doc.add_paragraph(texto)
    p.paragraph_format.keep_with_next = True


def _heading_keep_next(
    doc: Document,
    texto: str,
    nivel: int,
    *,
    page_break_before: bool = False,
) -> Any:
    h = doc.add_heading(texto, level=nivel)
    h.paragraph_format.keep_with_next = True
    if page_break_before:
        h.paragraph_format.page_break_before = True
    return h


def _puente_grafico(doc: Document) -> None:
    """Párrafo vacío que mantiene la tabla anterior junto al gráfico siguiente."""
    p = doc.add_paragraph("")
    p.paragraph_format.keep_with_next = True


TARIFA_PLANA_CLP_POR_M3 = 1300

# Ranking completo + tabla anexo: tipografía más compacta para favorecer una sola hoja.
TABLAS_FINALES_FONT_PT = 8.0

# Colores cercanos al informe de referencia (Sin WES naranja, Con WES verde)
COLOR_SIN_WES = "#D35400"
COLOR_CON_WES = "#1E8449"


def _elegir_csv_consolidado(out_dir: Path, explicit: Optional[Path]) -> Path:
    if explicit and explicit.is_file():
        return explicit.resolve()
    candidatos = sorted(
        out_dir.glob("consolidado_m3_mensual_colegios_puente_alto_*_desde_checkpoint.csv"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if candidatos:
        return candidatos[0]
    parcial = out_dir / "consolidado_m3_mensual_colegios_puente_alto_2025_PARCIAL.csv"
    if parcial.is_file():
        return parcial
    raise FileNotFoundError(
        "No hay CSV consolidado. Ejecute generar_consolidado_m3_mensual_puente_alto.py "
        "o indique --consolidado-csv."
    )


def _mes_clave(year: int, month: int) -> str:
    return f"{year}-{month:02d}"


def _columnas_mes(year: int, month: int) -> Tuple[str, str]:
    mk = _mes_clave(year, month)
    return mk, f"{mk}_sin_WES_est_m3"


def _fmt_clp(val: float) -> str:
    x = int(round(float(val)))
    neg = x < 0
    x = abs(x)
    s = f"{x:,}".replace(",", ".")
    return f"-${s}" if neg else f"${s}"


def _fmt_m3(v: float) -> str:
    return f"{float(v):,.2f}".replace(",", ".")


def _fmt_pct(v: float) -> str:
    return f"{float(v):.2f}".replace(".", ",") + " %"


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _compact_entire_table_font(table: Any, pt: float = TABLAS_FINALES_FONT_PT) -> None:
    """Reduce el tamaño de fuente en todas las celdas (últimas tablas densas del informe)."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(pt)


def _add_table_header_row(table, headers: Sequence[str]) -> None:
    row = table.rows[0].cells
    for i, h in enumerate(headers):
        row[i].text = ""
        p = row[i].paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(9)
        _set_cell_shading(row[i], "E7E6E6")


def _kpis_desde_consolidado(df: pd.DataFrame, year: int, tarifa: int) -> Dict[str, Any]:
    mask_t = df["node_id"].astype(str).str.strip().str.upper() == "TOTAL"
    n_escuelas = len(df[df["node_id"].astype(str).str.upper() != "TOTAL"])
    if not mask_t.any():
        raise ValueError("Consolidado sin fila TOTAL.")
    tr = df.loc[mask_t].iloc[0]
    total_con = float(str(tr["total_anio_m3"]).replace(",", "."))
    total_sin = float(str(tr["total_anio_sin_WES_est_m3"]).replace(",", "."))

    prom_sin_mes = total_sin / 12.0
    prom_con_mes = total_con / 12.0
    ahorro_m_m3 = prom_sin_mes - prom_con_mes
    ahorro_a_m3 = total_sin - total_con
    eff_global = (ahorro_m_m3 / prom_sin_mes * 100.0) if prom_sin_mes > 1e-9 else 0.0

    clp_m_base = ahorro_m_m3 * tarifa
    clp_a_base = ahorro_a_m3 * tarifa

    # Banda para línea narrativa (± referencia tarifa; no es escenario contractual)
    fac_lo, fac_hi = 0.82, 1.22
    rango_min = ahorro_a_m3 * tarifa * fac_lo
    rango_max = ahorro_a_m3 * tarifa * fac_hi

    return {
        "year": year,
        "n_escuelas": n_escuelas,
        "prom_sin_mes": prom_sin_mes,
        "prom_con_mes": prom_con_mes,
        "ahorro_m_m3": ahorro_m_m3,
        "ahorro_a_m3": ahorro_a_m3,
        "eff_global": eff_global,
        "clp_m_base": clp_m_base,
        "clp_a_base": clp_a_base,
        "rango_min_clp": rango_min,
        "rango_max_clp": rango_max,
        "tarifa": tarifa,
    }


def _tabla_kpi_resumen_consolidado(doc: Document, k: Dict[str, Any]) -> None:
    """Tabla estilo informe: indicador | valor."""
    tarifa = int(k["tarifa"])
    filas = [
        ("Puntos evaluados", str(int(k["n_escuelas"]))),
        ("Proyección mensual sin WES (m³)", _fmt_m3(k["prom_sin_mes"])),
        ("Proyección mensual con WES (m³)", _fmt_m3(k["prom_con_mes"])),
        ("Ahorro mensual (m³)", _fmt_m3(k["ahorro_m_m3"])),
        (f"Ahorro anual {k['year']} (m³)", _fmt_m3(k["ahorro_a_m3"])),
        ("Eficiencia global", _fmt_pct(k["eff_global"])),
        (f"Ahorro mensual base ({tarifa} CLP/m³)", _fmt_clp(k["clp_m_base"])),
        (f"Ahorro anual base ({tarifa} CLP/m³)", _fmt_clp(k["clp_a_base"])),
    ]
    table = doc.add_table(rows=1 + len(filas), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = ""
    p0 = hdr[0].paragraphs[0]
    r0 = p0.add_run("Indicador")
    r0.bold = True
    p1 = hdr[1].paragraphs[0]
    r1 = p1.add_run("Valor")
    r1.bold = True
    _set_cell_shading(hdr[0], "E7E6E6")
    _set_cell_shading(hdr[1], "E7E6E6")
    for i, (lab, val) in enumerate(filas, start=1):
        table.rows[i].cells[0].text = lab
        table.rows[i].cells[1].text = val

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        f"Rango anual estimado en CLP (banda referencia ~82–122 % sobre tarifa {tarifa} CLP/m³ "
        f"aplicada al ahorro proyectado): "
    )
    p.add_run(_fmt_clp(k["rango_min_clp"]) + " a " + _fmt_clp(k["rango_max_clp"])).bold = True
    p.add_run(".")

    doc.add_paragraph()
    n = doc.add_paragraph()
    n.add_run(
        "Nota metodológica: la proyección mensual consolidada se obtiene como total anual agregado "
        "(fila TOTAL del consolidado) dividido 12; la eficiencia global como ahorro mensual entre "
        "proyección mensual sin WES. La valorización económica usa una tarifa plana de referencia "
        f"({tarifa} CLP/m³), sin sustituir facturación real."
    ).italic = True


def _grafico_comparativo_mensual_consolidado(
    prom_sin: float,
    prom_con: float,
    year: int,
    out_png: Path,
) -> None:
    """Dos barras: Sin WES vs Con WES (promedios mensuales consolidados), estilo informe."""
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(6.0, 4.2), dpi=120)
    categorias = ["Sin WES", "Con WES"]
    valores = [prom_sin, prom_con]
    colores = [COLOR_SIN_WES, COLOR_CON_WES]
    bars = ax.bar(categorias, valores, color=colores, width=0.55, edgecolor="#333333", linewidth=0.6)
    ax.set_ylabel("m³/mes")
    ax.set_title(f"Proyección mensual consolidada ({year})")
    ax.yaxis.set_major_formatter(
        mticker.FuncFormatter(lambda v, _: f"{v:,.0f}".replace(",", "."))
    )
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_axisbelow(True)
    ymax = max(valores) * 1.18 if valores else 1
    ax.set_ylim(0, ymax)
    for bar, val in zip(bars, valores):
        ax.annotate(
            _fmt_m3(val),
            xy=(bar.get_x() + bar.get_width() / 2, bar.get_height()),
            xytext=(0, 4),
            textcoords="offset points",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(out_png, bbox_inches="tight")
    plt.close(fig)


def _tabla_anual_por_colegio(doc: Document, df: pd.DataFrame, year: int, tarifa: int) -> Any:
    sub = df[df["node_id"].astype(str).str.upper() != "TOTAL"].copy()
    sub["con"] = pd.to_numeric(sub["total_anio_m3"], errors="coerce").fillna(0)
    sub["sin"] = pd.to_numeric(sub["total_anio_sin_WES_est_m3"], errors="coerce").fillna(0)
    sub["$ con"] = sub["con"] * tarifa
    sub["$ sin"] = sub["sin"] * tarifa
    sub["$ ahorro (sin−con)"] = (sub["sin"] - sub["con"]) * tarifa
    sub = sub.sort_values("colegio")

    headers = [
        "Establecimiento",
        "m³ con WES (año)",
        "m³ sin WES est. (año)",
        f"$ con WES ({tarifa} CLP/m³)",
        "$ sin WES est.",
        "$ ahorro volumen",
    ]
    table = doc.add_table(rows=1 + len(sub), cols=len(headers))
    table.style = "Table Grid"
    _add_table_header_row(table, headers)
    for i, (_, row) in enumerate(sub.iterrows(), start=1):
        cells = table.rows[i].cells
        cells[0].text = str(row.get("colegio", ""))
        cells[1].text = _fmt_m3(float(row["con"]))
        cells[2].text = _fmt_m3(float(row["sin"]))
        cells[3].text = _fmt_clp(float(row["$ con"]))
        cells[4].text = _fmt_clp(float(row["$ sin"]))
        cells[5].text = _fmt_clp(float(row["$ ahorro (sin−con)"]))
    return table


def _tabla_mes_por_colegio(
    doc: Document,
    df: pd.DataFrame,
    year: int,
    month: int,
    tarifa: int,
    umbral_con: float = 1e-6,
) -> int:
    col_con, col_sin = _columnas_mes(year, month)
    if col_con not in df.columns:
        return 0
    sub = df[df["node_id"].astype(str).str.upper() != "TOTAL"].copy()
    rows_data: List[Tuple[str, float, float, float, float, float]] = []
    for _, row in sub.iterrows():
        try:
            c = float(str(row[col_con]).replace(",", "."))
            s = float(str(row[col_sin]).replace(",", "."))
        except (TypeError, ValueError):
            continue
        if c <= umbral_con:
            continue
        nombre = str(row.get("colegio", "")).strip()
        rows_data.append((nombre, c, s, c * tarifa, s * tarifa, (s - c) * tarifa))

    if not rows_data:
        doc.add_paragraph("(Sin filas con consumo con WES > 0 este mes.)", style="Intense Quote")
        return 0

    headers = [
        "Establecimiento",
        "m³ con WES",
        "m³ sin WES est.",
        f"$ con ({tarifa} CLP/m³)",
        "$ sin WES est.",
        "$ ahorro mes",
    ]
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = "Table Grid"
    _add_table_header_row(table, headers)
    for i, (nom, c, s, vc, vs, va) in enumerate(rows_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = nom
        cells[1].text = _fmt_m3(c)
        cells[2].text = _fmt_m3(s)
        cells[3].text = _fmt_clp(vc)
        cells[4].text = _fmt_clp(vs)
        cells[5].text = _fmt_clp(va)
    return len(rows_data)


def _tabla_agregado_mensual_total_row(doc: Document, df: pd.DataFrame, year: int, tarifa: int) -> None:
    mask = df["node_id"].astype(str).str.strip().str.upper() == "TOTAL"
    if not mask.any():
        doc.add_paragraph("(No hay fila TOTAL en el consolidado.)")
        return
    row_t = df.loc[mask].iloc[0]
    headers = [
        "Mes",
        "m³ con WES (total)",
        "m³ sin WES est. (total)",
        f"$ con ({tarifa} CLP/m³)",
        "$ sin WES est.",
        "$ diferencia (ahorro)",
    ]
    meses_nom = (
        "Enero",
        "Febrero",
        "Marzo",
        "Abril",
        "Mayo",
        "Junio",
        "Julio",
        "Agosto",
        "Septiembre",
        "Octubre",
        "Noviembre",
        "Diciembre",
    )
    filas: List[Tuple[str, float, float, float, float, float]] = []
    for m in range(1, 13):
        cc, cs = _columnas_mes(year, m)
        try:
            con_v = float(str(row_t[cc]).replace(",", "."))
            sin_v = float(str(row_t[cs]).replace(",", "."))
        except (TypeError, ValueError, KeyError):
            continue
        filas.append(
            (
                meses_nom[m - 1],
                con_v,
                sin_v,
                con_v * tarifa,
                sin_v * tarifa,
                (sin_v - con_v) * tarifa,
            )
        )

    table = doc.add_table(rows=1 + len(filas), cols=len(headers))
    table.style = "Table Grid"
    _add_table_header_row(table, headers)
    for i, (nombre, c, s, vc, vs, vd) in enumerate(filas, start=1):
        cells = table.rows[i].cells
        cells[0].text = nombre
        cells[1].text = _fmt_m3(c)
        cells[2].text = _fmt_m3(s)
        cells[3].text = _fmt_clp(vc)
        cells[4].text = _fmt_clp(vs)
        cells[5].text = _fmt_clp(vd)


def _add_picture_if_exists(doc: Document, path: Path, width_in: float = 6.3) -> bool:
    if not path.is_file():
        doc.add_paragraph(f"[Gráfico no encontrado: {path.name}]")
        return False
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def _find_body_child_index_containing(body: Any, substr: str) -> Optional[int]:
    from lxml import etree

    for i, child in enumerate(body):
        if substr in etree.tostring(child, encoding="unicode"):
            return i
    return None


def _wp_body_child_is_blank_gap(p_el: Any) -> bool:
    """True si es un ``w:p`` sin texto útil ni dibujos/objetos (hueco/salto típico entre secciones)."""
    from docx.oxml.ns import qn

    if not p_el.tag.endswith("p"):
        return False
    for nm in ("drawing", "pict", "object", "fldChar"):
        if p_el.findall(".//%s" % qn(f"w:{nm}")):
            return False
    texts = p_el.findall(".//%s" % qn("w:t"))
    joined = "".join((t.text or "") for t in texts).replace("\xa0", " ").strip()
    return len(joined) == 0


def _eliminar_huecos_en_blanco_antes_de_primer_bloque_cuerpo(
    doc: Document,
    marcador: str = "4. Detalle mensual por establecimiento",
) -> int:
    """
    Quita párrafos vacíos (o solo con saltos) **justo antes** del primer encabezado del cuerpo
    (por defecto sección 4), tras fusionar plantilla + cuerpo.
    """
    body = doc.element.body
    idx = _find_body_child_index_containing(body, marcador)
    if idx is None:
        return 0
    to_remove: List[Any] = []
    i = idx - 1
    while i >= 0:
        ch = body[i]
        if ch.tag.endswith("tbl"):
            break
        if ch.tag.endswith("p") and _wp_body_child_is_blank_gap(ch):
            to_remove.append(ch)
            i -= 1
        else:
            break
    for el in to_remove:
        body.remove(el)
    return len(to_remove)


def _fmt_tarifa_clp_chilena(tarifa_clp: int) -> str:
    return f"{tarifa_clp:,}".replace(",", ".")


def _actualizar_metodologia_punto_14_tarifa_plana(doc: Document, tarifa_clp: int) -> bool:
    """
    En el punto **1.4 Cálculo económico en pesos chilenos**, deja explícita la tarifa plana CLP/m³.
    Devuelve True si se reemplazó algún párrafo.
    """
    tarifa_txt = _fmt_tarifa_clp_chilena(tarifa_clp)
    nuevo = (
        f"1.4 Cálculo económico en pesos chilenos. La valorización monetaria utiliza una tarifa plana de "
        f"{tarifa_txt} pesos chilenos por metro cúbico (CLP/m³), como referencia orientativa y sin sustituir "
        "la facturación contractual real."
    )
    for p in doc.paragraphs:
        tx = (p.text or "").replace("\xa0", " ")
        if not re.search(r"1\s*\.\s*4", tx):
            continue
        if "económico" not in tx.lower() and "economico" not in tx.lower():
            continue
        if "chileno" not in tx.lower():
            continue
        p.text = nuevo
        return True
    return False


def _remover_bloque_sensibilidad_y_escenarios_tarifarios(doc: Document) -> int:
    """
    Quita de la metodología el texto de tarifa «efectiva», sensibilidad y escenarios
    (sustituido por la tarifa plana en 1.4). Devuelve cuántos párrafos se eliminaron.
    """
    to_remove: List[Any] = []
    for p in doc.paragraphs:
        t = (p.text or "").replace("\xa0", " ")
        if not t.strip():
            continue
        tl = t.lower()
        if "la valorización se obtiene multiplicando" in tl:
            to_remove.append(p)
            continue
        if "se incorpora" in tl and "sensibilidad" in tl and "escenario" in tl:
            to_remove.append(p)
            continue
        if "escenario conservador" in tl and "15%" in t and "tarifa" in tl:
            to_remove.append(p)
            continue
        if "escenario base" in tl and "promedio referencial" in tl:
            to_remove.append(p)
            continue
        if "escenario alto" in tl and "15%" in t and "tarifa" in tl:
            to_remove.append(p)
            continue
        if re.search(
            r"ahorro\s*clp\s*=\s*ahorro\s*m", tl, re.I
        ) and "tarifa" in tl and "clp" in tl:
            to_remove.append(p)
            continue
    for p in to_remove:
        p._element.getparent().remove(p._element)
    return len(to_remove)


def _body_child_plain_text_first_paragraph(ch: Any) -> str:
    """Texto visible del hijo ``w:p`` del body (vacío si no es párrafo)."""
    return _plain_text_w_body_child(ch)


def _eliminar_bloque_xml_secciones_2_y_3_cuerpo(doc: Document) -> int:
    """
    Informes antiguos llevaban secciones 2–3 en el cuerpo; el informe actual empieza en 4.
    Elimina nodos XML desde el párrafo cuyo texto es exactamente «2. Resumen ejecutivo consolidado»
    hasta antes del párrafo que empieza por «4. Detalle mensual por establecimiento».
    Coincidencia por texto plano evita borrar la plantilla (TOC u otras menciones a «2.»).
    """
    body = doc.element.body
    children = list(body)
    tit2 = "2. Resumen ejecutivo consolidado"
    idx2: Optional[int] = None
    idx4: Optional[int] = None
    for i, ch in enumerate(children):
        ht = _body_child_plain_text_first_paragraph(ch)
        if ht == tit2:
            idx2 = i
            break
    if idx2 is None:
        return 0
    from lxml import etree

    def _hijos_siguientes_sugieren_kpi_seccion2(i2: int) -> bool:
        for j in range(i2 + 1, min(i2 + 6, len(children))):
            s = etree.tostring(children[j], encoding="unicode")
            if "Puntos evaluados" in s or ("Indicador" in s and "Valor" in s):
                return True
        return False

    if not _hijos_siguientes_sugieren_kpi_seccion2(idx2):
        return 0
    for i, ch in enumerate(children):
        if i <= idx2:
            continue
        ht = _body_child_plain_text_first_paragraph(ch)
        if ht.startswith("4. Detalle mensual por establecimiento"):
            idx4 = i
            break
    if idx4 is None or idx4 <= idx2:
        return 0
    to_rem = children[idx2:idx4]
    for el in reversed(to_rem):
        body.remove(el)
    return len(to_rem)


def _remover_parrafos_banda_tarifaria_3_3(doc: Document) -> int:
    """Elimina restos de subapartado 3.3 / tres barras / bandas tarifarias (texto suelto)."""
    to_rm: List[Any] = []
    for p in doc.paragraphs:
        t = (p.text or "").replace("\xa0", " ")
        if not t.strip():
            continue
        tl = t.lower()
        if "3.3" in t and "escenario" in tl:
            to_rm.append(p)
            continue
        if "gráfico de 3 barras" in tl or ("3 barras" in tl and "banda" in tl):
            to_rm.append(p)
            continue
        if "barras conservadora" in tl and "alta" in tl:
            to_rm.append(p)
            continue
        if "tres bandas tarifarias" in tl:
            to_rm.append(p)
            continue
    for p in to_rm:
        p._element.getparent().remove(p._element)
    return len(to_rm)


def aplicar_post_procesado_informe_consolidado_pa(
    doc: Document,
    tarifa_clp: int = TARIFA_PLANA_CLP_POR_M3,
) -> Tuple[int, bool, int, int]:
    """
    Tras fusionar plantilla + cuerpo: huecos antes del 4; actualiza 1.4; quita sensibilidad
    en metodología; restos 3.3 (texto suelto de escenarios antiguos).
    Devuelve ``(n_gap, actualizó_1_4, n_tarifa, n_párr_3_3)``. No elimina secciones 2–3.
    """
    n_gap = _eliminar_huecos_en_blanco_antes_de_primer_bloque_cuerpo(doc)
    ok_14 = _actualizar_metodologia_punto_14_tarifa_plana(doc, tarifa_clp)
    n_tarifa = _remover_bloque_sensibilidad_y_escenarios_tarifarios(doc)
    n_33 = _remover_parrafos_banda_tarifaria_3_3(doc)
    return (n_gap, ok_14, n_tarifa, n_33)


def _fusionar_plantilla_y_cuerpo_ooxml(plantilla: Path, cuerpo_docx: Path, salida: Path) -> None:
    """
    Une plantilla + cuerpo con **docxcompose.Composer** para que las imágenes incrustadas
    (``word/media/``, relaciones) se copien al paquete final.

    Un merge solo por ``deepcopy`` del XML del ``body`` **rompe** las figuras: el XML copiado
    sigue apuntando a ``r:embed`` que no existen en el .docx destino.

    El post-proceso (huecos antes del bloque 4, etc.) compensa posibles saltos raros de página.
    """
    try:
        from docxcompose.composer import Composer
    except ImportError as e:
        raise RuntimeError(
            "La fusión con imágenes requiere docxcompose: pip install docxcompose"
        ) from e

    composer = Composer(Document(str(plantilla)))
    composer.append(Document(str(cuerpo_docx)))
    salida.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(salida))


def construir_cuerpo_documento(
    df: pd.DataFrame,
    year: int,
    tarifa: int,
    graficos_dir: Path,
) -> Document:
    k = _kpis_desde_consolidado(df, year, tarifa)

    doc = Document()
    doc.add_heading("2. Resumen ejecutivo consolidado", level=1)
    _tabla_kpi_resumen_consolidado(doc, k)

    h3 = doc.add_heading("3. Gráficas", level=1)
    h3.paragraph_format.keep_with_next = True

    _heading_keep_next(doc, "3.1 Comparativo mensual (Sin WES vs Con WES)", 2)
    _para_keep_next(
        doc,
        "Comparación de promedios mensuales consolidados (total anual ÷ 12 en la fila TOTAL). "
        f"Valorización económica en sección 2 con tarifa referencia {tarifa} CLP/m³.",
    )
    mini = doc.add_table(rows=2, cols=3)
    mini.style = "Table Grid"
    mini.rows[0].cells[0].text = "Concepto"
    mini.rows[0].cells[1].text = "Sin WES (m³/mes)"
    mini.rows[0].cells[2].text = "Con WES (m³/mes)"
    mini.rows[1].cells[0].text = "Total corporación"
    mini.rows[1].cells[1].text = _fmt_m3(k["prom_sin_mes"])
    mini.rows[1].cells[2].text = _fmt_m3(k["prom_con_mes"])

    png_comp = graficos_dir / f"puente_alto_comparativo_mensual_consolidado_{year}.png"
    _grafico_comparativo_mensual_consolidado(
        k["prom_sin_mes"], k["prom_con_mes"], year, png_comp
    )
    _puente_grafico(doc)
    _add_picture_if_exists(doc, png_comp, width_in=5.8)

    _heading_keep_next(doc, "3.2 Top 10 por ahorro anual (m³)", 2, page_break_before=True)
    _para_keep_next(
        doc,
        "Ranking por ahorro de volumen anual (m³ sin WES − m³ con WES). "
        "Se listan hasta 10 establecimientos.",
    )
    sub = df[df["node_id"].astype(str).str.upper() != "TOTAL"].copy()
    sub["con"] = pd.to_numeric(sub["total_anio_m3"], errors="coerce").fillna(0)
    sub["sin"] = pd.to_numeric(sub["total_anio_sin_WES_est_m3"], errors="coerce").fillna(0)
    sub["ahorro_m3"] = sub["sin"] - sub["con"]
    sub = sub.sort_values("ahorro_m3", ascending=False).head(10).reset_index(drop=True)
    sub.insert(0, "Ranking", range(1, len(sub) + 1))

    headers = [
        "Ranking",
        "Establecimiento",
        "Ahorro anual (m³)",
        f"Ahorro anual ({tarifa} CLP/m³)",
    ]
    table = doc.add_table(rows=1 + len(sub), cols=len(headers))
    table.style = "Table Grid"
    _add_table_header_row(table, headers)
    for i, (_, row) in enumerate(sub.iterrows(), start=1):
        cells = table.rows[i].cells
        cells[0].text = str(int(row["Ranking"]))
        cells[1].text = str(row.get("colegio", ""))
        cells[2].text = _fmt_m3(float(row["ahorro_m3"]))
        cells[3].text = _fmt_clp(float(row["ahorro_m3"]) * tarifa)
    _puente_grafico(doc)
    _add_picture_if_exists(doc, graficos_dir / f"ranking_ahorro_anual_pa_{year}.png", width_in=6.0)

    h4 = doc.add_heading("4. Detalle mensual por establecimiento", level=1)
    h4.paragraph_format.page_break_before = True
    _para_keep_next(
        doc,
        "Por cada mes: una página de trabajo con la tabla de establecimientos y el gráfico "
        "comparativo del mismo mes (solo colegios con medición con WES > 0), para revisar "
        "cifras de la tabla frente a las barras del gráfico.",
    )
    meses_txt = (
        "enero",
        "febrero",
        "marzo",
        "abril",
        "mayo",
        "junio",
        "julio",
        "agosto",
        "septiembre",
        "octubre",
        "noviembre",
        "diciembre",
    )
    for m in range(1, 13):
        _heading_keep_next(
            doc,
            f"4.{m} {meses_txt[m - 1].capitalize()} {year}",
            2,
            page_break_before=(m >= 2),
        )
        _para_keep_next(doc, "Tabla de datos seguida del gráfico comparativo del mismo mes.")
        nrows = _tabla_mes_por_colegio(doc, df, year, m, tarifa)
        if nrows > 0:
            _puente_grafico(doc)
            _add_picture_if_exists(doc, graficos_dir / f"puente_alto_con_vs_sin_wes_{year}_{m:02d}.png")
        else:
            doc.add_paragraph("(Sin gráfico: ningún colegio con con WES > 0 en el mes.)")

    h5 = doc.add_heading("5. Agregado mensual total y anexos", level=1)
    h5.paragraph_format.page_break_before = True
    h5.paragraph_format.keep_with_next = True

    _heading_keep_next(doc, "5.1 Totales mensuales corporación", 2)
    _para_keep_next(doc, "Totales mensuales de la fila TOTAL del consolidado y gráfico agregado asociado.")
    _tabla_agregado_mensual_total_row(doc, df, year, tarifa)
    _puente_grafico(doc)
    _add_picture_if_exists(doc, graficos_dir / f"puente_alto_AGREGADO_con_vs_sin_wes_{year}.png")

    _heading_keep_next(doc, "5.2 Ranking completo por ahorro anual", 2, page_break_before=True)
    sub_f = df[df["node_id"].astype(str).str.upper() != "TOTAL"].copy()
    sub_f["con"] = pd.to_numeric(sub_f["total_anio_m3"], errors="coerce").fillna(0)
    sub_f["sin"] = pd.to_numeric(sub_f["total_anio_sin_WES_est_m3"], errors="coerce").fillna(0)
    sub_f["ahorro_m3"] = sub_f["sin"] - sub_f["con"]
    sub_f["$ ahorro anual"] = sub_f["ahorro_m3"] * tarifa
    sub_f = sub_f.sort_values("ahorro_m3", ascending=False).reset_index(drop=True)
    sub_f.insert(0, "Ranking", range(1, len(sub_f) + 1))

    headers_f = [
        "Ranking",
        "Establecimiento",
        "m³ con WES",
        "m³ sin WES est.",
        f"$ con ({tarifa} CLP/m³)",
        "$ sin WES est.",
        "$ ahorro anual",
    ]
    table_f = doc.add_table(rows=1 + len(sub_f), cols=len(headers_f))
    table_f.style = "Table Grid"
    _add_table_header_row(table_f, headers_f)
    for i, (_, row) in enumerate(sub_f.iterrows(), start=1):
        cells = table_f.rows[i].cells
        cells[0].text = str(int(row["Ranking"]))
        cells[1].text = str(row.get("colegio", ""))
        cells[2].text = _fmt_m3(float(row["con"]))
        cells[3].text = _fmt_m3(float(row["sin"]))
        cells[4].text = _fmt_clp(float(row["con"]) * tarifa)
        cells[5].text = _fmt_clp(float(row["sin"]) * tarifa)
        cells[6].text = _fmt_clp(float(row["$ ahorro anual"]))

    _compact_entire_table_font(table_f, TABLAS_FINALES_FONT_PT)

    anexo_h = _heading_keep_next(doc, "Anexo — Tabla anual por establecimiento", 1, page_break_before=True)
    anexo_h.paragraph_format.space_before = Pt(6)
    anexo_h.paragraph_format.space_after = Pt(2)
    tab_anexo = _tabla_anual_por_colegio(doc, df, year, tarifa)
    _compact_entire_table_font(tab_anexo, TABLAS_FINALES_FONT_PT)

    return doc


def _generar_png_auxiliares_puente_alto(csv_path: Path, year: int, graficos_dir: Path) -> None:
    """
    Crea en ``graficos_dir`` los PNG que el Word incrusta por ruta (no van dentro del .docx hasta
    ``add_picture``): ranking anual, un PNG por mes y el agregado. El gráfico de dos barras de la
    sección 3.1 se genera más tarde dentro de ``construir_cuerpo_documento``.
    """
    graficos_dir.mkdir(parents=True, exist_ok=True)
    py = sys.executable
    script_graf = ROOT / "grafico_con_vs_sin_wes_puente_alto.py"
    script_rank = ROOT / "ranking_ahorro_anual_puente_alto.py"
    if script_graf.is_file():
        subprocess.run(
            [
                py,
                str(script_graf),
                "--year",
                str(year),
                "--out-dir",
                str(graficos_dir),
                "--consolidado-csv",
                str(csv_path),
                "--con-agregado",
            ],
            check=False,
        )
    else:
        print(f"[WARN] Falta {script_graf.name}; no habrá gráficos mensuales ni agregado.", flush=True)
    if script_rank.is_file():
        subprocess.run(
            [
                py,
                str(script_rank),
                "--year",
                str(year),
                "--consolidado-csv",
                str(csv_path),
            ],
            check=False,
        )
    else:
        print(f"[WARN] Falta {script_rank.name}; no habrá gráfico del Top 10.", flush=True)


def generar_borrador(
    year: int,
    csv_path: Path,
    out_docx: Path,
    tarifa: int = TARIFA_PLANA_CLP_POR_M3,
    graficos_dir: Path = OUT_GRAFICOS,
    plantilla: Optional[Path] = None,
    *,
    borrar_plantilla_temporal: bool = False,
    omitir_png_auxiliares: bool = False,
) -> None:
    df = pd.read_csv(csv_path, dtype=str)
    df.columns = [str(c).replace("\ufeff", "").strip() for c in df.columns]

    if not omitir_png_auxiliares:
        print("[INFO] Generando PNG auxiliares (ranking, meses, agregado)…", flush=True)
        _generar_png_auxiliares_puente_alto(csv_path, year, graficos_dir)

    doc_inner = construir_cuerpo_documento(df, year, tarifa, graficos_dir)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        inner_path = Path(tf.name)
    try:
        doc_inner.save(str(inner_path))
        if plantilla and plantilla.is_file():
            print(f"[INFO] Plantilla (portada/metodología): {plantilla}", flush=True)
            _fusionar_plantilla_y_cuerpo_ooxml(plantilla, inner_path, out_docx)
            doc_out = Document(str(out_docx))
            n_gap, ok14, n_tar, n_p33 = aplicar_post_procesado_informe_consolidado_pa(
                doc_out, tarifa_clp=tarifa
            )
            doc_out.save(str(out_docx))
            if n_gap:
                print(
                    f"[INFO] Post-proceso: eliminados {n_gap} párrafos hueco antes del bloque 4.",
                    flush=True,
                )
            if ok14:
                print("[INFO] Post-proceso: punto 1.4 metodología actualizado (tarifa plana CLP/m³).", flush=True)
            if n_tar:
                print(
                    f"[INFO] Post-proceso: eliminados {n_tar} párrafos (sensibilidad / escenarios tarifarios).",
                    flush=True,
                )
            if n_p33:
                print(
                    f"[INFO] Post-proceso: eliminados {n_p33} párrafos (restos apartado 3.3).",
                    flush=True,
                )
            if borrar_plantilla_temporal:
                try:
                    plantilla.unlink(missing_ok=True)
                except OSError:
                    pass
        else:
            shutil.copyfile(inner_path, out_docx)
    finally:
        if inner_path.is_file():
            inner_path.unlink(missing_ok=True)


def main() -> int:
    ap = argparse.ArgumentParser(description="Borrador Word informe Puente Alto (estilo KPI + gráficas)")
    ap.add_argument("--year", type=int, default=2025)
    ap.add_argument("--tarifa-clp-m3", type=int, default=TARIFA_PLANA_CLP_POR_M3)
    ap.add_argument("--consolidado-csv", type=Path, default=None)
    ap.add_argument(
        "--salida",
        type=Path,
        default=None,
        help="Ruta .docx de salida",
    )
    ap.add_argument("--graficos-dir", type=Path, default=None)
    ap.add_argument(
        "--plantilla",
        type=Path,
        default=None,
        help="Word completo con solo portada + metodología (opcional). "
        "Por defecto se extraen las 2 primeras páginas de Informe_Auditoria_WES_Puente_Alto_*_BORRADOR.docx.",
    )
    ap.add_argument(
        "--sin-plantilla",
        action="store_true",
        help="No anteponer portada/metodología.",
    )
    ap.add_argument(
        "--omitir-png-auxiliares",
        action="store_true",
        help="No ejecutar grafico_con_vs_sin_wes ni ranking (asume PNG ya en --graficos-dir).",
    )
    args = ap.parse_args()

    year = args.year
    csv_path = _elegir_csv_consolidado(OUT_REPORTS, args.consolidado_csv)
    out_docx = (
        Path(args.salida).expanduser().resolve()
        if args.salida
        else OUT_REPORTS / f"Borrador_resumen_ejecutivo_Puente_Alto_{year}.docx"
    )
    gdir = Path(args.graficos_dir).expanduser().resolve() if args.graficos_dir else OUT_GRAFICOS

    plantilla: Optional[Path] = None
    borrar_tmp = False
    if not args.sin_plantilla:
        if args.plantilla:
            plantilla = Path(args.plantilla).expanduser().resolve()
        elif INFORME_BORRADOR_REF.is_file():
            plantilla = extraer_primeras_paginas_docx(INFORME_BORRADOR_REF, hasta_saltos_explicitos=2)
            borrar_tmp = True
            print(
                f"[INFO] Plantilla automática: 2 primeras páginas desde {INFORME_BORRADOR_REF.name}",
                flush=True,
            )
        elif PLANTILLA_DEFAULT.is_file():
            plantilla = PLANTILLA_DEFAULT

    if plantilla and not plantilla.is_file():
        print(f"[WARN] Plantilla no encontrada: {plantilla}. Se genera solo el cuerpo.", flush=True)
        plantilla = None
        borrar_tmp = False

    print(f"[INFO] Consolidado: {csv_path}", flush=True)
    generar_borrador(
        year,
        csv_path,
        out_docx,
        tarifa=args.tarifa_clp_m3,
        graficos_dir=gdir,
        plantilla=plantilla,
        borrar_plantilla_temporal=borrar_tmp,
        omitir_png_auxiliares=args.omitir_png_auxiliares,
    )
    print(f"[OK] {out_docx}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
