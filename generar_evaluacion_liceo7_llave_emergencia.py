"""
Evaluación rápida Liceo 7 Providencia (000006-04):
control hídrico desactivado desde 7-ago + llave de emergencia abierta.
"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from statistics import mean

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generar_reporte_word import add_logo_to_header
from wes_estilo_graficos_app import guardar_grafico_horario_24h_app

NODE_ID = "000006-04"
NOMBRE = "Liceo 7 Luisa Saavedra (Providencia)"
HORARIO_CORTE = "23:00 a 06:00"
OUT_DIR = Path("reports/Providencia/Liceo7_evaluacion_llave_emergencia")
TARIFA_REF_CLP = 1400  # referencial Santiago / Aguas Andinas


def _fmt(n: float, dec: int = 1) -> str:
    s = f"{n:,.{dec}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt_clp(n: float) -> str:
    return f"${_fmt(n, 0)}"


def _set_run(run, *, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, bold=True, size=16 if level == 1 else 13, color=(30, 90, 150))
    return p


def _add_p(doc, text: str, *, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, bold=bold, size=size)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        _set_run(run, bold=True, size=9, color=(255, 255, 255))
        shading = hdr[i]._element.get_or_add_tcPr()
        from docx.oxml import parse_xml

        shading.append(
            parse_xml(
                r'<w:shd {} w:fill="1F5A96"/>'.format(
                    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                )
            )
        )
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_run(run, size=9)
    return table


def build_charts(rows: list[dict], out: Path) -> dict:
    charts = {}

    # 1) Barras diarias total + nocturno (julio–ago)
    sel = [r for r in rows if r["date"] >= "2026-07-01"]
    dates = [r["date"][5:] for r in sel]
    totals = [r["total"] for r in sel]
    nights = [r["night"] for r in sel]
    colors = []
    for r in sel:
        if r["date"] >= "2026-08-07":
            colors.append("#c41e1e")
        elif r["total"] > 50 and r["strict_1_5"] < 1:
            colors.append("#2e7ac8")
        else:
            colors.append("#4A8CB8")

    fig, ax = plt.subplots(figsize=(11, 4.2))
    ax.bar(range(len(dates)), totals, color=colors, width=0.75, label="Consumo diario")
    ax.plot(range(len(dates)), nights, color="#E67E22", marker="o", ms=3, lw=1.2, label="Nocturno 23–06 h")
    ax.axvline(dates.index("08-07") - 0.5, color="#c41e1e", ls="--", lw=1.2, label="Desactivación control (7-ago)")
    ax.set_ylabel("m³")
    ax.set_title("Liceo 7 — Consumo diario (jul–ago 2026)")
    ax.set_xticks(range(len(dates)))
    ax.set_xticklabels(dates, rotation=90, fontsize=6)
    ax.legend(fontsize=8, loc="upper left")
    ax.set_ylim(0, max(totals) * 1.15)
    fig.tight_layout()
    p1 = out / "chart_consumo_diario.png"
    fig.savefig(p1, dpi=150, bbox_inches="tight")
    plt.close(fig)
    charts["diario"] = p1

    # 2) Perfil con control (5-ago) vs sin (10-ago)
    con = next(r for r in rows if r["date"] == "2026-08-05")
    sin = next(r for r in rows if r["date"] == "2026-08-10")
    horas_con = {int(k): float(v) for k, v in con["profile"].items()}
    horas_sin = {int(k): float(v) for k, v in sin["profile"].items()}
    p2 = out / "chart_perfil_con_control_05ago.png"
    guardar_grafico_horario_24h_app(
        horas_con,
        p2,
        titulo="Con control activo — 05/08/2026 (mié)",
        subtitulo="Corte 23:00–06:00 · horas 01–05 ≈ 0 m³/h",
    )
    charts["con"] = p2
    p3 = out / "chart_perfil_sin_control_10ago.png"
    guardar_grafico_horario_24h_app(
        horas_sin,
        p3,
        titulo="Sin control / llave emergencia abierta — 10/08/2026 (lun)",
        subtitulo="Flujo continuo en madrugada · ~9–19 m³/h en 01–05",
    )
    charts["sin"] = p3

    # 3) Comparación proyección
    fig, ax = plt.subplots(figsize=(7.5, 4.2))
    labels = [
        "Con control\n(escenario)",
        "Sin control\n(llave abierta)",
        "Exceso\nmensual",
    ]
    vals = [8637, 15952, 7315]
    cols = ["#2e7ac8", "#c41e1e", "#E67E22"]
    bars = ax.bar(labels, vals, color=cols, width=0.55)
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2, v + 200, f"{_fmt(v, 0)} m³", ha="center", fontsize=9, fontweight="bold")
    ax.set_ylabel("m³ / mes")
    ax.set_title("Proyección mensual Liceo 7 (22 días hábiles + 8 fin de semana)")
    ax.set_ylim(0, 18500)
    fig.tight_layout()
    p4 = out / "chart_proyeccion_mensual.png"
    fig.savefig(p4, dpi=150, bbox_inches="tight")
    plt.close(fig)
    charts["proy"] = p4
    return charts


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = json.loads((OUT_DIR / "serie_diaria.json").read_text(encoding="utf-8"))
    charts = build_charts(rows, OUT_DIR)

    con = [r for r in rows if r["date"] in ["2026-08-03", "2026-08-04", "2026-08-05", "2026-08-06"]]
    sin_wd = [r for r in rows if r["date"] in ["2026-08-07", "2026-08-10", "2026-08-11"]]
    sin_all = [r for r in rows if "2026-08-07" <= r["date"] <= "2026-08-11"]
    sin_we = [r for r in rows if r["date"] in ["2026-08-08", "2026-08-09"]]
    jul = [r for r in rows if r["date"].startswith("2026-07")]

    avg_con = mean(r["total"] for r in con)
    avg_night_con = mean(r["night"] for r in con)
    avg_strict_con = mean(r["strict_1_5"] for r in con)
    avg_sin_wd = mean(r["total"] for r in sin_wd)
    avg_night_sin = mean(r["night"] for r in sin_all)
    avg_strict_sin = mean(r["strict_1_5"] for r in sin_all)
    avg_we_sin = mean(r["total"] for r in sin_we)

    proy_con = 22 * avg_con + 8 * 0
    proy_sin = 22 * avg_sin_wd + 8 * avg_we_sin
    delta = proy_sin - proy_con
    extra_dia = avg_sin_wd - avg_con
    extra_noche = avg_night_sin - avg_night_con
    costo_mes = delta * TARIFA_REF_CLP
    costo_dia_extra = extra_dia * TARIFA_REF_CLP

    # Acumulado 7–11 ago vs baseline con control
    acum_5d = sum(r["total"] for r in sin_all)
    baseline_5d = avg_con * 5  # aproximación (incluye 2 WE que con control ≈0)
    # Mejor: con control weekdays 393, weekends 0 → 3*393 + 2*0 for same mix? 
    # Days 7,10,11 = wd; 8,9 = we. Baseline = 3*avg_con + 0
    baseline_mismo_mix = 3 * avg_con + 2 * 0
    acum_real_wd_we = sum(r["total"] for r in sin_all)
    exceso_5d = acum_real_wd_we - baseline_mismo_mix

    summary = {
        "nodo": NODE_ID,
        "nombre": NOMBRE,
        "horario_corte": HORARIO_CORTE,
        "julio_total_m3": round(sum(r["total"] for r in jul), 1),
        "avg_con_control_m3_dia": round(avg_con, 1),
        "avg_sin_control_habil_m3_dia": round(avg_sin_wd, 1),
        "avg_sin_control_finde_m3_dia": round(avg_we_sin, 1),
        "avg_noche_con": round(avg_night_con, 1),
        "avg_noche_sin": round(avg_night_sin, 1),
        "avg_strict_sin_1_5": round(avg_strict_sin, 1),
        "extra_diario_habil_m3": round(extra_dia, 1),
        "extra_nocturno_m3": round(extra_noche, 1),
        "proy_mensual_con_m3": round(proy_con, 0),
        "proy_mensual_sin_m3": round(proy_sin, 0),
        "delta_mensual_m3": round(delta, 0),
        "tarifa_ref_clp": TARIFA_REF_CLP,
        "costo_mensual_exceso_clp": round(costo_mes, 0),
        "exceso_5dias_7_11_ago_m3": round(exceso_5d, 1),
        "generado": datetime.now().isoformat(timespec="seconds"),
    }
    (OUT_DIR / "resumen.json").write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

    doc = Document()
    add_logo_to_header(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Evaluación de consumo — Liceo 7 Providencia")
    _set_run(run, bold=True, size=18, color=(20, 70, 120))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Sala de bomba · Control hídrico desactivado desde 7/08/2026 · "
        f"Llave de emergencia abierta (visita técnica 12/08/2026)\n"
        f"Nodo WES {NODE_ID} · Horario de corte programado: {HORARIO_CORTE}"
    )
    _set_run(run, size=10, color=(80, 80, 80))

    _add_heading(doc, "1. Hallazgo operativo", 1)
    _add_p(
        doc,
        "Desde julio 2026 la sala de bomba del Liceo 7 opera con monitoreo WES. "
        "El equipo genera cada noche el control hídrico del colegio (ventana 23:00–06:00). "
        "A partir del 7 de agosto se desactivó el sistema de control. En la visita técnica "
        "del 12 de agosto se encontró la llave de emergencia abierta: todo el caudal "
        "programado para corte por horario pasa de largo.",
    )
    _add_p(
        doc,
        "Evidencia en serie horaria: con control activo (p. ej. 3–6 ago) las horas 01:00–05:00 "
        "quedan en ≈0 m³/h. Desde el 7 ago esas mismas horas marcan 10–17 m³/h de forma continua, "
        "también sábado y domingo.",
        bold=False,
    )

    _add_heading(doc, "2. Números clave", 1)
    _add_table(
        doc,
        ["Indicador", "Con control (3–6 ago)", "Sin control (7–11 ago)", "Diferencia"],
        [
            [
                "Consumo día hábil",
                f"{_fmt(avg_con)} m³/día",
                f"{_fmt(avg_sin_wd)} m³/día",
                f"+{_fmt(extra_dia)} m³/día (+{_fmt(100*extra_dia/avg_con,0)}%)",
            ],
            [
                "Consumo nocturno 23–06",
                f"{_fmt(avg_night_con)} m³/noche",
                f"{_fmt(avg_night_sin)} m³/noche",
                f"+{_fmt(extra_noche)} m³/noche",
            ],
            [
                "Ventana estricta 01–05",
                f"{_fmt(avg_strict_con)} m³",
                f"{_fmt(avg_strict_sin)} m³",
                f"+{_fmt(avg_strict_sin - avg_strict_con)} m³ (debería ser ~0)",
            ],
            [
                "Fin de semana",
                "≈ 0 m³/día",
                f"{_fmt(avg_we_sin)} m³/día",
                "Corte de finde anulado",
            ],
            [
                "Julio 2026 (mes completo)",
                f"{_fmt(sum(r['total'] for r in jul))} m³",
                "—",
                "Sala bomba en operación",
            ],
        ],
    )

    _add_heading(doc, "3. Qué significa para el colegio", 1)
    _add_p(
        doc,
        f"En solo 5 días (7–11 ago) el colegio consumió {_fmt(acum_real_wd_we)} m³. "
        f"Con el control funcionando, el mismo mix (3 hábiles + 2 finde) habría sido "
        f"≈ {_fmt(baseline_mismo_mix)} m³. Exceso acumulado: {_fmt(exceso_5d)} m³ "
        f"(≈ {_fmt_clp(exceso_5d * TARIFA_REF_CLP)} a tarifa referencial {_fmt_clp(TARIFA_REF_CLP)}/m³).",
    )
    _add_p(
        doc,
        f"En día hábil el exceso es ≈ {_fmt(extra_dia)} m³/día "
        f"(≈ {_fmt_clp(costo_dia_extra)}/día). Gran parte corresponde a agua que el "
        "horario de corte debería haber detenido en madrugada, más el arrastre de "
        "consumo en fines de semana que antes quedaban en cero.",
    )

    _add_heading(doc, "4. Proyección mensual", 1)
    _add_p(
        doc,
        "Escenario comercial a 30 días (22 hábiles + 8 fin de semana), extrapolando "
        "el comportamiento medido:",
    )
    _add_table(
        doc,
        ["Escenario", "m³ / mes", "Costo ref. (CLP)", "Comentario"],
        [
            [
                "Con control operativo",
                _fmt(proy_con, 0),
                _fmt_clp(proy_con * TARIFA_REF_CLP),
                "Hábiles ~393 m³; finde ~0",
            ],
            [
                "Sin control / llave abierta",
                _fmt(proy_sin, 0),
                _fmt_clp(proy_sin * TARIFA_REF_CLP),
                "Hábiles ~577 m³; finde ~407 m³",
            ],
            [
                "Exceso mensual",
                _fmt(delta, 0),
                _fmt_clp(costo_mes),
                "Agua que el control evitaría",
            ],
        ],
    )
    _add_p(
        doc,
        f"En síntesis: mantener la llave de emergencia abierta y el control desactivado "
        f"implica del orden de {_fmt(delta, 0)} m³/mes adicionales "
        f"(≈ {_fmt_clp(costo_mes)}/mes a tarifa referencial). "
        "Eso casi duplica la factura respecto al escenario con corte nocturno activo.",
    )

    _add_heading(doc, "5. Gráficos", 1)
    _add_p(doc, "Consumo diario julio–agosto (rojo = post desactivación 7/08):")
    doc.add_picture(str(charts["diario"]), width=Inches(6.3))
    _add_p(doc, "Perfil horario con control (05/08) vs sin control (10/08):")
    doc.add_picture(str(charts["con"]), width=Inches(6.0))
    doc.add_picture(str(charts["sin"]), width=Inches(6.0))
    _add_p(doc, "Proyección mensual comparada:")
    doc.add_picture(str(charts["proy"]), width=Inches(5.5))

    _add_heading(doc, "6. Detalle diario post-desactivación", 1)
    det_rows = []
    for r in rows:
        if "2026-08-03" <= r["date"] <= "2026-08-12":
            estado = "SIN CONTROL" if r["date"] >= "2026-08-07" else "con control"
            det_rows.append(
                [
                    r["date"],
                    r["weekday"],
                    _fmt(r["total"]),
                    _fmt(r["night"]),
                    _fmt(r["strict_1_5"]),
                    estado,
                ]
            )
    _add_table(
        doc,
        ["Fecha", "Día", "Total m³", "Noche 23–06", "01–05 m³", "Estado"],
        det_rows,
    )
    _add_p(
        doc,
        "Nota: 12/08 es día parcial (visita técnica en curso); no se usa para promedios ni proyección.",
        size=9,
    )

    _add_heading(doc, "7. Recomendación", 1)
    _add_p(
        doc,
        "1) Cerrar de inmediato la llave de emergencia y restablecer el control hídrico "
        "con el horario 23:00–06:00.",
    )
    _add_p(
        doc,
        "2) Verificar en la app/control nocturno WES que las horas 01–05 vuelvan a 0 m³/h "
        "la misma noche del cierre.",
    )
    _add_p(
        doc,
        "3) Mientras la llave permanezca abierta, el colegio proyecta ~16.000 m³/mes frente a "
        "~8.600 m³/mes con control: un sobrecosto del orden de 7.300 m³/mes.",
    )

    foot = doc.add_paragraph()
    run = foot.add_run(
        f"Fuente: API WES nodo {NODE_ID} · Generado {summary['generado']} · "
        f"Tarifa {_fmt_clp(TARIFA_REF_CLP)}/m³ solo referencial (no factura oficial)."
    )
    _set_run(run, size=8, color=(120, 120, 120))

    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    docx_path = OUT_DIR / f"Evaluacion_Liceo7_llave_emergencia_{stamp}.docx"
    doc.save(docx_path)
    print(f"[OK] DOCX {docx_path}")
    print(json.dumps(summary, indent=2, ensure_ascii=False))
    print(f"DOCX_PATH={docx_path}")


if __name__ == "__main__":
    main()
