"""
Informe de auditoría Renca — Con WES 10–13 ago vs Sin WES 17–20 ago 2026.

Puntos: ICCP, Esc. Lo Velásquez, Gimnasio municipal, Piscina municipal.
Días homólogos lun–jue. Si el 20/08 aún no cerró, ambos jueves se cortan a la
última hora completa Chile (misma ventana).

Uso:
  python generar_auditoria_renca_ago10_13_vs_ago17_20.py
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import load_workbook

import auditoria_cpa_icco_renca_grafico as graf
import generar_graficos_comparativos_desde_excel_consolidado as gxlsx
import generar_informe_auditoria_icco_renca_word as icco
from auditoria_cpa_icco_renca_grafico import Periodo
from generar_excel_auditoria_consolidado_dos_periodos import generar_excel_consolidado
from generar_graficos_comparativos_desde_excel_consolidado import (
    _limpiar_pngs_carpeta_graficos,
    generar_pngs,
    leer_matriz_consolidado,
)
from generar_reporte_word import add_logo_to_header, format_number_chilean

ROOT = Path(__file__).resolve().parent
TZ_CL = ZoneInfo("America/Santiago")
TARIFA_CLP_M3 = 1300.0
COLOR_WES = "#2a6fad"
COLOR_SIN = "#C0504D"
COLOR_HEAD = RGBColor(31, 71, 136)

PUNTOS: Tuple[Tuple[str, str], ...] = (
    ("000017-07", "ICCP (Cumbre de Cóndores pte.)"),
    ("000017-04", "Esc. Lo Velásquez"),
    ("000017-05", "Gimnasio municipal"),
    ("000017-06", "Piscina municipal"),
)

DIAS_CON = (date(2026, 8, 10), date(2026, 8, 11), date(2026, 8, 12), date(2026, 8, 13))
DIAS_SIN = (date(2026, 8, 17), date(2026, 8, 18), date(2026, 8, 19), date(2026, 8, 20))
PER_CON = Periodo("Con control (10-08 a 13-08-2026)", DIAS_CON)
PER_SIN = Periodo("Sin control (17-08 a 20-08-2026)", DIAS_SIN)


def _hora_corte_homologa(ahora: datetime) -> int:
    """Última hora Chile 0–23 a incluir en ambos jueves. 23 = día completo."""
    d20 = date(2026, 8, 20)
    if ahora.date() > d20:
        return 23
    if ahora.date() < d20:
        return max(0, min(23, ahora.hour - 1))
    return max(0, min(23, ahora.hour - 1))


def _aplicar_corte_jueves(xlsx: Path, hora_corte: int) -> None:
    """Horas posteriores al corte = 0 en ambos jueves (col E = 13/08, col I = 20/08)."""
    if hora_corte >= 23:
        return
    wb = load_workbook(xlsx)
    ws = wb["Consolidado"]
    for col in (5, 9):
        for h in range(hora_corte + 1, 24):
            ws.cell(row=4 + h, column=col, value=0.0)
    wb.save(xlsx)
    wb.close()


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shading = parse_xml(
        "<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
        f'w:val="clear" w:fill="{hex_fill}"/>'
    )
    tc_pr.append(shading)


def _rendimiento(m3_sin: float, m3_con: float) -> Tuple[float, float]:
    ahorro = m3_sin - m3_con
    pct = (ahorro / m3_sin * 100.0) if m3_sin > 1e-9 else 0.0
    return ahorro, pct


def _grafico_4_puntos(
    nombres: List[str],
    m3_con: List[float],
    m3_sin: List[float],
    titulo: str,
    out_png: Path,
) -> None:
    x = np.arange(len(nombres))
    w = 0.36
    fig, ax = plt.subplots(figsize=(11.4, 5.6))
    b1 = ax.bar(x - w / 2, m3_con, width=w, color=COLOR_WES, label="Con WES 10–13 ago", zorder=2)
    b2 = ax.bar(x + w / 2, m3_sin, width=w, color=COLOR_SIN, label="Sin WES 17–20 ago", zorder=2)
    ymax = max(list(m3_con) + list(m3_sin) + [1.0]) * 1.22
    ax.set_ylim(0, ymax)
    ax.set_xticks(list(x))
    ax.set_xticklabels(nombres, fontsize=10)
    ax.set_ylabel("Consumo (m³)")
    ax.set_title(titulo, fontweight="bold", fontsize=12, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=9, loc="upper left")
    for bars in (b1, b2):
        for bar in bars:
            h = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                h + ymax * 0.015,
                format_number_chilean(float(h), 1),
                ha="center",
                va="bottom",
                fontsize=8,
                color="#333333",
            )
    for i, (a, b) in enumerate(zip(m3_sin, m3_con)):
        _, pct = _rendimiento(float(a), float(b))
        ax.text(
            x[i],
            ymax * 0.96,
            f"{pct:+.0f} %",
            ha="center",
            va="top",
            fontsize=9,
            fontweight="bold",
            color="#548235" if pct > 0 else COLOR_SIN,
        )
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_dias_homologos(
    nombres_dia: List[str],
    series_con: Dict[str, List[float]],
    series_sin: Dict[str, List[float]],
    out_png: Path,
) -> None:
    fig, axes = plt.subplots(2, 2, figsize=(12.8, 7.6))
    for ax, (nid, nom) in zip(axes.ravel(), PUNTOS):
        x = np.arange(len(nombres_dia))
        w = 0.36
        ax.bar(x - w / 2, series_con[nid], width=w, color=COLOR_WES, label="Con WES", zorder=2)
        ax.bar(x + w / 2, series_sin[nid], width=w, color=COLOR_SIN, label="Sin WES", zorder=2)
        ax.set_xticks(list(x))
        ax.set_xticklabels(nombres_dia, fontsize=8)
        ax.set_ylabel("m³")
        ax.set_title(nom, fontsize=10, fontweight="bold", color="#1F4788")
        ax.grid(axis="y", alpha=0.3, zorder=0)
        ax.legend(fontsize=7, loc="upper left")
    fig.suptitle(
        "Días homólogos lun–jue: Con WES 10–13 ago vs Sin WES 17–20 ago",
        fontsize=13,
        fontweight="bold",
        color="#1F4788",
    )
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _convertir_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=180,
    )
    return pdf_path if pdf_path.is_file() else None


def _auditoria_un_punto(
    node_id: str,
    nombre: str,
    cliente_dir: Path,
    hora_corte: int,
) -> Tuple[Path, Optional[Path], float, float, List[float], List[float]]:
    cliente_dir.mkdir(parents=True, exist_ok=True)
    xlsx_out = cliente_dir / "consumo_consolidado_conWES_10-13ago_sinWES_17-20ago.xlsx"
    print(f"  CSV + Excel {node_id} {nombre}…")
    generar_excel_consolidado(
        xlsx_out,
        node_id,
        DIAS_CON,
        DIAS_SIN,
        titulo_p1="Con WES: 10-08-2026 al 13-08-2026",
        titulo_p2="Sin WES: 17-08-2026 al 20-08-2026",
        solo_desde_csv=False,
        csv_dir=cliente_dir / "csv_descarga_api",
    )
    _aplicar_corte_jueves(xlsx_out, hora_corte)

    gdir = cliente_dir / "graficos_comparativos"
    _limpiar_pngs_carpeta_graficos(gdir)
    fechas, mats = leer_matriz_consolidado(xlsx_out)
    generar_pngs(fechas, mats, gdir)

    mid = len(fechas) // 2
    m3_con = float(sum(sum(c) for c in mats[:mid]))
    m3_sin = float(sum(sum(c) for c in mats[mid:]))
    dias_con = [float(sum(c)) for c in mats[:mid]]
    dias_sin = [float(sum(c)) for c in mats[mid:]]

    stem = f"Auditoria_{nombre.replace(' ', '_')}_{node_id}"
    docx_path = cliente_dir / f"{stem}.docx"
    graf.NOMBRE_PUNTO = nombre
    icco.NOMBRE_PUNTO = nombre
    icco.PORTADA_TITULO = "Informe de Auditoría"
    icco.PORTADA_REFERENCIA_BORRADOR = stem
    icco.PORTADA_ESTABLECIMIENTO_LINEA1 = nombre
    icco.PORTADA_ESTABLECIMIENTO_LINEA2 = node_id
    icco.CLP_POR_M3_REF = TARIFA_CLP_M3
    icco._ETIQUETA_COLEGIO_CUADRO_RESUMEN = nombre.upper()[:22]

    print(f"  Word {node_id}…")
    p = icco.generar_informe_word(
        node_id=node_id,
        out_dir=cliente_dir,
        output_docx=docx_path,
        mantener_borrador_manual=False,
        solo_consolidado=False,
        periodo_ref=PER_CON,
        periodo_aud=PER_SIN,
        figuras_desde_xlsx=xlsx_out,
    )
    pdf = _convertir_pdf(p)
    return p, pdf, m3_con, m3_sin, dias_con, dias_sin


def _word_conjunto(
    out_docx: Path,
    hora_corte: int,
    ahora: datetime,
    pngs: Dict[str, Path],
    m3_con: Dict[str, float],
    m3_sin: Dict[str, float],
    dias_con: Dict[str, List[float]],
    dias_sin: Dict[str, List[float]],
    dirs: Dict[str, Path],
) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    try:
        add_logo_to_header(doc)
    except Exception:
        pass

    h = doc.add_heading("Informe de Auditoría — Renca", level=0)
    if h.runs:
        h.runs[0].font.color.rgb = COLOR_HEAD
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("Con WES 10 al 13 de agosto de 2026  vs  Sin WES 17 al 20 de agosto de 2026")
    r.bold = True
    r.font.color.rgb = COLOR_HEAD

    horas = "00:00–23:59" if hora_corte >= 23 else f"00:00–{hora_corte:02d}:59"
    p = doc.add_paragraph()
    p.add_run("Ventana homóloga: ").bold = True
    p.add_run(
        f"lunes a miércoles completos + jueves {horas} Chile "
        f"(corte {ahora:%d/%m/%Y %H:%M} Chile). "
        "Cada día sin control se compara con el mismo día de la semana anterior con WES "
        "(lun 17 vs lun 10, mar 18 vs mar 11, mié 19 vs mié 12, jue 20 vs jue 13). "
        "Ahorro = (Sin WES − Con WES) / Sin WES × 100. Tarifa de referencia 1.300 CLP/m³."
    )

    doc.add_heading("1. Los 4 puntos", level=1)
    if "todos" in pngs:
        doc.add_picture(str(pngs["todos"]), width=Cm(16.2))
    if "dias4" in pngs:
        doc.add_paragraph("")
        doc.add_picture(str(pngs["dias4"]), width=Cm(16.2))

    tot_con = sum(m3_con.values())
    tot_sin = sum(m3_sin.values())
    ahorro_t, pct_t = _rendimiento(tot_sin, tot_con)
    doc.add_paragraph(
        f"Conjunto: Con WES {format_number_chilean(tot_con, 1)} m³; "
        f"Sin WES {format_number_chilean(tot_sin, 1)} m³; "
        f"diferencia {format_number_chilean(ahorro_t, 1)} m³ "
        f"({format_number_chilean(pct_t, 1)} % sobre el periodo sin control; "
        f"~{format_number_chilean(max(0.0, ahorro_t) * TARIFA_CLP_M3, 0)} CLP)."
    )

    tbl = doc.add_table(rows=1 + len(PUNTOS) + 1, cols=6)
    tbl.style = "Table Grid"
    headers = ["Punto", "Con WES (m³)", "Sin WES (m³)", "Ahorro (m³)", "%", "Hallazgo"]
    for j, hd in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = hd
        _set_cell_shading(cell, "1F4788")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)

    for i, (nid, nom) in enumerate(PUNTOS, start=1):
        ahorro, pct = _rendimiento(m3_sin[nid], m3_con[nid])
        if pct >= 15:
            hall = "WES baja el consumo"
        elif pct > 0:
            hall = "Ahorro menor"
        else:
            hall = "Sin WES no gasta más"
        row = tbl.rows[i]
        vals = [
            nom,
            format_number_chilean(m3_con[nid], 1),
            format_number_chilean(m3_sin[nid], 1),
            format_number_chilean(ahorro, 1),
            format_number_chilean(pct, 1) + " %",
            hall,
        ]
        for j, v in enumerate(vals):
            row.cells[j].text = v
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
        if pct >= 15:
            for c in row.cells:
                _set_cell_shading(c, "E2EFDA")

    row = tbl.rows[1 + len(PUNTOS)]
    vals = [
        "Total 4 puntos",
        format_number_chilean(tot_con, 1),
        format_number_chilean(tot_sin, 1),
        format_number_chilean(ahorro_t, 1),
        format_number_chilean(pct_t, 1) + " %",
        "",
    ]
    for j, v in enumerate(vals):
        row.cells[j].text = v
        _set_cell_shading(row.cells[j], "D6DCE4")
        for run in row.cells[j].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    nomb_d = ["Lun 10 vs 17", "Mar 11 vs 18", "Mié 12 vs 19", f"Jue 13 vs 20 ({horas})"]
    for nid, nom in PUNTOS:
        doc.add_heading(nom, level=1)
        ahorro, pct = _rendimiento(m3_sin[nid], m3_con[nid])
        doc.add_paragraph(
            f"Con WES {format_number_chilean(m3_con[nid], 1)} m³ vs "
            f"Sin WES {format_number_chilean(m3_sin[nid], 1)} m³ → "
            f"{format_number_chilean(pct, 1)} % "
            f"({format_number_chilean(ahorro, 1)} m³)."
        )
        gdir = dirs[nid] / "graficos_comparativos"
        barras = gdir / "02_barras_total_rejilla_por_periodo.png"
        if barras.is_file():
            doc.add_picture(str(barras), width=Cm(12.5))
        for fname in ("04_area_Lunes.png", "05_area_Martes.png", "06_area_Miercoles.png", "07_area_Jueves.png"):
            pth = gdir / fname
            if pth.is_file():
                doc.add_picture(str(pth), width=Cm(16.0))
        prom = gdir / "03_area_promedio_24h.png"
        if prom.is_file():
            doc.add_picture(str(prom), width=Cm(16.0))

        t2 = doc.add_table(rows=5, cols=5)
        t2.style = "Table Grid"
        for j, hd in enumerate(["Día homólogo", "Con WES (m³)", "Sin WES (m³)", "Ahorro (m³)", "%"]):
            cell = t2.rows[0].cells[j]
            cell.text = hd
            _set_cell_shading(cell, "1F4788")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
        for i, lab in enumerate(nomb_d):
            a, pc = _rendimiento(dias_sin[nid][i], dias_con[nid][i])
            vals = [
                lab,
                format_number_chilean(dias_con[nid][i], 1),
                format_number_chilean(dias_sin[nid][i], 1),
                format_number_chilean(a, 1),
                format_number_chilean(pc, 1) + " %",
            ]
            for j, v in enumerate(vals):
                t2.rows[i + 1].cells[j].text = v
                for run in t2.rows[i + 1].cells[j].paragraphs[0].runs:
                    run.font.size = Pt(9)

    doc.add_heading("Conclusiones", level=1)
    doc.add_paragraph(
        "La comparación es la más limpia disponible: misma estación, mismos días de la semana "
        "y la última semana completa con control WES (10–13 ago) contra la primera semana sin "
        "control (17–20 ago). El aumento se concentra en ICCP y piscina municipal; escuela y "
        "gimnasio no explican el salto de volumen."
    )
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def main() -> int:
    ahora = datetime.now(TZ_CL)
    hora_corte = _hora_corte_homologa(ahora)
    ts = ahora.strftime("%Y%m%d_%H%M")
    out_root = (
        ROOT
        / "reports"
        / "reporte de auditoria"
        / f"auditoria_renca_conWES_10-13ago_sinWES_17-20ago_{ts}"
    )
    out_root.mkdir(parents=True, exist_ok=True)
    print(
        f"Auditoría Renca | Con WES 10–13 ago vs Sin WES 17–20 ago | "
        f"jueves hasta {hora_corte:02d}:59 Chile | {ahora:%Y-%m-%d %H:%M}"
    )

    gxlsx.LABEL_P1 = "Con WES 10–13 ago 2026"
    gxlsx.LABEL_P2 = "Sin WES 17–20 ago 2026"

    m3_con: Dict[str, float] = {}
    m3_sin: Dict[str, float] = {}
    dias_con: Dict[str, List[float]] = {}
    dias_sin: Dict[str, List[float]] = {}
    dirs: Dict[str, Path] = {}

    for nid, nom in PUNTOS:
        print("=" * 64)
        print(f"{nid} {nom}")
        cdir = out_root / f"Auditoria_{nom.split('(')[0].strip().replace(' ', '_')}_{nid}"
        dirs[nid] = cdir
        _docx, _pdf, c, s, dc, ds = _auditoria_un_punto(nid, nom, cdir, hora_corte)
        m3_con[nid] = c
        m3_sin[nid] = s
        dias_con[nid] = dc
        dias_sin[nid] = ds
        print(f"  Con {c:.1f} m³ | Sin {s:.1f} m³ | Word {_docx}")
        if _pdf:
            print(f"  PDF {_pdf}")

    png_dir = out_root / "graficos"
    png_dir.mkdir(exist_ok=True)
    p_all = png_dir / "todos_casos_con_vs_sin.png"
    _grafico_4_puntos(
        ["ICCP", "Lo Velásquez", "Gimnasio", "Piscina"],
        [m3_con[n] for n, _ in PUNTOS],
        [m3_sin[n] for n, _ in PUNTOS],
        "Renca — auditoría: Con WES 10–13 ago vs Sin WES 17–20 ago",
        p_all,
    )
    etiqueta_jue = "Jueves" if hora_corte >= 23 else f"Jueves\nhasta {hora_corte:02d}:59"
    p_dias = png_dir / "dias_homologos_4_puntos.png"
    _grafico_dias_homologos(
        ["Lunes", "Martes", "Miércoles", etiqueta_jue],
        dias_con,
        dias_sin,
        p_dias,
    )

    docx_path = out_root / f"Informe_Auditoria_Renca_conWES_10-13_sinWES_17-20_{ts}.docx"
    _word_conjunto(
        docx_path,
        hora_corte,
        ahora,
        {"todos": p_all, "dias4": p_dias},
        m3_con,
        m3_sin,
        dias_con,
        dias_sin,
        dirs,
    )
    pdf_path = _convertir_pdf(docx_path)

    tot_c = sum(m3_con.values())
    tot_s = sum(m3_sin.values())
    ahorro, pct = _rendimiento(tot_s, tot_c)
    print("=" * 64)
    print(f"TOTAL Con WES {tot_c:.1f} m³ | Sin WES {tot_s:.1f} m³ | {pct:.1f}% ({ahorro:.1f} m³)")
    print(f"DOCX conjunto: {docx_path}")
    print(f"PDF  conjunto: {pdf_path or '(no convertido)'}")
    print(f"DIR: {out_root}")
    return 0


if __name__ == "__main__":
    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
            sys.stderr.reconfigure(encoding="utf-8")
        except Exception:
            pass
    raise SystemExit(main())
