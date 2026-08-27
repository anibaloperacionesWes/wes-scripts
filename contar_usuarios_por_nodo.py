# -*- coding: utf-8 -*-
"""
Cuenta usuarios WES por nodeId (allowedNodes) y los ordena de mayor a menor.

La API acl-entities no lista todos los usuarios: solo GET /users?email= y
GET /users/{userId}. Este script consulta un universo de correos conocidos
(archivos del repo + contactos WES) y cuenta cuántos de esos usuarios tienen
cada nodo en allowedNodes.

Salida (reports/Usuarios/usuarios_por_nodo/):
  - CSV ranking (solo puntos activos)
  - XLSX (ranking + detalle de usuarios)
  - DOCX resumen
  - PDF

Por defecto se omiten puntos dados de baja o fuera de operación
(exclusiones_reportes, registro_puntos_deshabilitados y bajas PA).

Uso:
  python contar_usuarios_por_nodo.py
  python contar_usuarios_por_nodo.py --emails extra.txt
  python contar_usuarios_por_nodo.py --incluir-inactivos
"""

from __future__ import annotations

import argparse
import csv
import re
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Set, Tuple
from zoneinfo import ZoneInfo

import requests

ENTITY_BASE = "http://104.248.53.141:7001/wes/api/acl-entities/v1"
ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Usuarios" / "usuarios_por_nodo"
TZ_CL = ZoneInfo("America/Santiago")

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
JUNK_DOMAINS = {
    "empresa.com",
    "ejemplo.com",
    "email.com",
    "fu.my",
    "w.rt",
    "e.mqc",
    "p.ne",
    "r.lt",
    "o.lgp",
    "u7o.ex",
    "f.ya",
}
SKIP_DIR_PARTS = {".git", "__pycache__", "node_modules", "gmail_oauth"}
SKIP_SUFFIX = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".pdf",
    ".docx",
    ".xlsx",
    ".zip",
    ".pyc",
    ".mp4",
    ".pkl",
    ".pickle",
}

WES_SPA_ID = "000000"

# La API no tiene PUT de nombre (solo POST crear / DELETE). Nombres reales
# cuando el registro WES quedó desactualizado.
NOMBRE_DISPLAY_POR_EMAIL: Dict[str, Tuple[str, str]] = {
    "an_ambiental_pae@linkes.cl": ("Sergio", "Fuenzalida"),
}


def nombre_usuario(user: dict) -> str:
    email = str(user.get("username") or "").strip().lower()
    over = NOMBRE_DISPLAY_POR_EMAIL.get(email)
    if over:
        return f"{over[0]} {over[1]}".strip()
    return f"{user.get('name', '')} {user.get('lastName', '')}".strip()


# Parque Arauco fuera de operación (mismo set que LISTADO_PA_IDS_EXCLUIDOS;
# se copia acá para no importar listado_pa_que_esta_instalado → matplotlib).
_PA_IDS_FUERA_DE_OPERACION = frozenset(
    {
        "000025-02",
        "000025-03",
        "000025-05",
        "000025-06",
    }
)


def _ids_pa_inactivos() -> Set[str]:
    ids: Set[str] = set(_PA_IDS_FUERA_DE_OPERACION)
    try:
        from pa_nodos_inactivos_por_mall import NODOS_INACTIVOS_POR_MALL

        for grupo in NODOS_INACTIVOS_POR_MALL.values():
            ids.update(grupo)
    except ImportError:
        pass
    return ids


def es_nodo_activo(
    node_id: str,
    company_id: str = "",
    company_name: str = "",
) -> bool:
    """True si el punto está operativo (no dado de baja / no excluido de reportes)."""
    from exclusiones_reportes import (
        EXCLUDED_NODE_IDS_PUNTOS_EN_CERO,
        is_node_excluded,
    )

    nid = (node_id or "").strip()
    if not nid:
        return False
    if is_node_excluded(nid, company_id or None, company_name or None):
        return False
    if nid in EXCLUDED_NODE_IDS_PUNTOS_EN_CERO:
        return False
    if nid in _ids_pa_inactivos():
        return False
    return True


def _session() -> requests.Session:
    s = requests.Session()
    s.headers.update({"Accept": "application/json"})
    return s


def _email_valido(email: str) -> bool:
    email = email.strip().lower()
    if not email or "@" not in email:
        return False
    local, _, domain = email.partition("@")
    if not local or not domain or "." not in domain.replace(",", "."):
        return False
    domain = domain.replace(",", ".")
    if domain in JUNK_DOMAINS:
        return False
    if email.endswith((".png", ".jpg", ".css", ".js")):
        return False
    if len(local) < 2:
        return False
    tld = domain.rsplit(".", 1)[-1]
    if len(tld) < 2 or not tld.isalpha():
        return False
    if any(ch in local for ch in " \t"):
        return False
    # Evitar restos de regex sobre hashes / basura.
    if local.startswith((".", "_")) and len(local) <= 3:
        return False
    return True


def _normalizar_email(raw: str) -> str:
    email = raw.strip().lower()
    if "@" not in email:
        return email
    local, _, domain = email.partition("@")
    # Coma en dominio suele ser typo (externos,parauco.com). En el local se respeta
    # porque hay usuarios WES con coma real (p. ej. supervisor,fundozapallar@gmail.com).
    return f"{local}@{domain.replace(',', '.')}"


def _nodos_y_empresas(session: requests.Session) -> Tuple[Dict[str, str], Dict[str, str], Dict[str, str]]:
    """
    Returns:
      node_id -> nombre
      node_id -> company_id (prioriza empresa distinta de Wes Spa)
      company_id -> company_name
    """
    r = session.get(f"{ENTITY_BASE}/configuration/companies", timeout=30)
    r.raise_for_status()
    empresas = r.json()
    nombres: Dict[str, str] = {}
    company_de: Dict[str, str] = {}
    companies: Dict[str, str] = {}

    for emp in empresas or []:
        cid = str(emp.get("companyId") or "").strip()
        cname = str(emp.get("name") or cid).strip()
        if cid:
            companies[cid] = cname
        for node in emp.get("nodes") or []:
            nid = str(node.get("nodeId") or "").strip()
            if not nid:
                continue
            nname = str(node.get("name") or "").strip()
            if nid not in nombres or nname:
                nombres[nid] = nname or nombres.get(nid, "")
            prev = company_de.get(nid)
            if prev is None:
                company_de[nid] = cid
            elif prev == WES_SPA_ID and cid and cid != WES_SPA_ID:
                company_de[nid] = cid

    return nombres, company_de, companies


def _emails_contactos() -> Set[str]:
    found: Set[str] = set()
    try:
        from lista_contactos_reportes import CONTACTOS_REPORTES, CORREOS_AUTORIZADOS

        for c in CONTACTOS_REPORTES.values():
            e = _normalizar_email(c.get("email") or "")
            if _email_valido(e):
                found.add(e)
        for row in CORREOS_AUTORIZADOS:
            e = _normalizar_email(row.get("email") or "")
            if _email_valido(e):
                found.add(e)
    except ImportError:
        pass
    return found


def _emails_desde_archivo(path: Path) -> Set[str]:
    if not path.is_file():
        return set()
    found: Set[str] = set()
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return set()
    for m in EMAIL_RE.finditer(text):
        e = _normalizar_email(m.group(0))
        if _email_valido(e):
            found.add(e)
    return found


def _emails_desde_repo() -> Set[str]:
    found: Set[str] = set()
    for p in ROOT.rglob("*"):
        if not p.is_file():
            continue
        if any(part in SKIP_DIR_PARTS for part in p.parts):
            continue
        if p.suffix.lower() in SKIP_SUFFIX:
            continue
        try:
            if p.stat().st_size > 2_000_000:
                continue
        except OSError:
            continue
        found.update(_emails_desde_archivo(p))
    return found


def _recolectar_emails(extra: Iterable[Path]) -> List[str]:
    pool: Set[str] = set()
    pool.update(_emails_contactos())
    pool.update(_emails_desde_repo())
    for fp in extra:
        pool.update(_emails_desde_archivo(fp))
    return sorted(pool)


def _emails_desde_alertas(
    session: requests.Session,
    node_ids: Iterable[str],
    workers: int = 10,
) -> Set[str]:
    """Correos de receivers FILTRATION (clientes que el ranking por repo no ve)."""
    found: Set[str] = set()

    def _uno(nid: str) -> List[str]:
        nid = (nid or "").strip()
        if "-" not in nid:
            return []
        cid = nid.split("-", 1)[0]
        try:
            r = session.get(
                f"{ENTITY_BASE}/companies/{cid}/node/{nid}/alert/FILTRATION/information",
                timeout=15,
            )
        except requests.RequestException:
            return []
        if r.status_code != 200:
            return []
        try:
            data = r.json()
        except ValueError:
            return []
        out: List[str] = []
        for rec in data.get("receiverList") or []:
            if not isinstance(rec, dict):
                continue
            e = _normalizar_email(rec.get("email") or "")
            if e:
                out.append(e)
        return out

    nids = [n for n in node_ids if n]
    with ThreadPoolExecutor(max_workers=max(1, min(workers, 16))) as ex:
        futs = [ex.submit(_uno, nid) for nid in nids]
        for fut in as_completed(futs):
            for e in fut.result():
                found.add(e)
    return found


def _fetch_user(session: requests.Session, email: str) -> Tuple[str, Optional[dict], str]:
    candidatos: List[str] = []
    base = email.strip()
    if base:
        candidatos.append(base)
    # Typo frecuente: coma en vez de punto en el dominio. Si el local tiene coma,
    # también probar la variante con punto (por si el API la guarda distinta).
    if "," in base:
        candidatos.append(base.replace(",", "."))
    vistos: Set[str] = set()
    last_status = ""
    try:
        for cand in candidatos:
            key = cand.lower()
            if key in vistos:
                continue
            vistos.add(key)
            r = session.get(f"{ENTITY_BASE}/users", params={"email": cand}, timeout=25)
            last_status = str(r.status_code)
            if r.status_code == 200:
                u = r.json()
                if isinstance(u, dict):
                    u.pop("password", None)
                obs = "" if cand == base else f"email variante: {cand}"
                return email, u, obs
        if last_status == "404":
            return email, None, "no existe en API"
        return email, None, f"HTTP {last_status or 'error'}"
    except requests.RequestException as exc:
        return email, None, str(exc)


def _allowed_nodes(user: dict) -> List[str]:
    allowed = user.get("allowedNodes") or []
    if not isinstance(allowed, list):
        return []
    out: List[str] = []
    seen: Set[str] = set()
    for nid in allowed:
        if not isinstance(nid, str):
            continue
        nid = nid.strip()
        if not nid or nid in seen:
            continue
        seen.add(nid)
        out.append(nid)
    return out


def xlsx_a_pdf(
    xlsx_path: Path,
    pdf_path: Optional[Path] = None,
    *,
    solo_ranking: bool = True,
    max_paginas: int = 5,
) -> Path:
    """Convierte el Excel de ranking a PDF. Por defecto solo la hoja de ranking, págs. 1-5."""
    from openpyxl import load_workbook
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    xlsx_path = Path(xlsx_path)
    if pdf_path is None:
        pdf_path = xlsx_path.with_suffix(".pdf")
    else:
        pdf_path = Path(pdf_path)

    wb = load_workbook(xlsx_path, data_only=True)
    page = landscape(A4)
    margin = 10 * mm
    usable_w = page[0] - 2 * margin

    cell_style = ParagraphStyle(
        "celda",
        fontName="Helvetica",
        fontSize=6.5,
        leading=8,
        alignment=TA_LEFT,
        wordWrap="CJK",
    )
    header_style = ParagraphStyle(
        "encabezado",
        fontName="Helvetica-Bold",
        fontSize=7,
        leading=8.5,
        textColor=colors.white,
        alignment=TA_CENTER,
    )
    title_style = ParagraphStyle(
        "titulo",
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#1F4E79"),
        alignment=TA_LEFT,
    )
    sub_style = ParagraphStyle(
        "subtitulo",
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#555555"),
    )

    def _cell(val, header: bool = False) -> Paragraph:
        txt = "" if val is None else str(val)
        txt = (
            txt.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        return Paragraph(txt, header_style if header else cell_style)

    story = []
    story.append(Paragraph("Usuarios por Node ID — ranking mayor a menor", title_style))
    story.append(
        Paragraph(
            f"PDF generado desde {xlsx_path.name}"
            + (
                " · ranking compacto (págs. 1-5; correos completos en el Excel)"
                if solo_ranking and max_paginas == 5
                else ""
            ),
            sub_style,
        )
    )
    story.append(Spacer(1, 4 * mm))

    header_bg = colors.HexColor("#1F4E79")
    zebra = colors.HexColor("#F4F8FB")
    green = colors.HexColor("#C6EFCE")

    hojas = list(wb.worksheets)
    if solo_ranking:
        hojas = [ws for ws in hojas if ws.title == "Ranking_por_nodo"] or hojas[:1]

    for si, ws in enumerate(hojas):
        if si:
            story.append(PageBreak())
            story.append(Paragraph("Usuarios por Node ID — ranking mayor a menor", title_style))
            story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(f"Hoja: {ws.title}", sub_style))
        story.append(Spacer(1, 2 * mm))

        rows_raw = list(ws.iter_rows(values_only=True))
        if not rows_raw:
            continue
        headers = ["" if h is None else str(h) for h in rows_raw[0]]
        compact = bool(solo_ranking and max_paginas and max_paginas <= 5)
        drop = {"company_id", "emails"} if compact else set()
        keep_idx = [i for i, h in enumerate(headers) if h not in drop]
        if not keep_idx:
            keep_idx = list(range(len(headers)))
        headers = [headers[i] for i in keep_idx]
        n_cols = len(headers)
        data = [[_cell(h, header=True) for h in headers]]
        qty_col = None
        for i, h in enumerate(headers):
            if h.strip().lower() in {"cantidad_usuarios", "nodos_count"}:
                qty_col = i
                break
        qty_by_data_row: List[int] = [0]

        for row in rows_raw[1:]:
            if row is None or all(v is None or str(v).strip() == "" for v in row):
                continue
            vals = [(row[i] if i < len(row) else "") for i in keep_idx]
            data.append([_cell(v) for v in vals])
            qty = 0
            if qty_col is not None:
                try:
                    qty = int(vals[qty_col] or 0)
                except (TypeError, ValueError):
                    qty = 0
            qty_by_data_row.append(qty)

        # Anchos según cantidad de columnas.
        if n_cols >= 7:
            weights = [0.05, 0.08, 0.09, 0.18, 0.08, 0.14, 0.38]
        elif n_cols == 5:
            weights = [0.08, 0.14, 0.14, 0.40, 0.24]
        elif n_cols == 6:
            weights = [0.18, 0.14, 0.16, 0.10, 0.08, 0.34]
        else:
            weights = [1.0 / n_cols] * n_cols
        col_w = [usable_w * w for w in weights[:n_cols]]
        if len(col_w) < n_cols:
            rest = usable_w - sum(col_w)
            col_w.extend([rest / (n_cols - len(col_w))] * (n_cols - len(col_w)))

        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), header_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B0B0B0")),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]
        for r_i in range(1, len(data)):
            if r_i % 2 == 0:
                style_cmds.append(("BACKGROUND", (0, r_i), (-1, r_i), zebra))
            if qty_col is not None and qty_by_data_row[r_i] > 0:
                style_cmds.append(("BACKGROUND", (qty_col, r_i), (qty_col, r_i), green))

        table = Table(data, colWidths=col_w, repeatRows=1)
        table.setStyle(TableStyle(style_cmds))
        story.append(KeepTogether([table]) if len(data) <= 8 else table)

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=page,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
        title="Usuarios por Node ID",
        author="WES",
    )
    doc.build(story)
    if max_paginas and max_paginas > 0:
        _limitar_pdf_paginas(pdf_path, max_paginas)
    return pdf_path


def _limitar_pdf_paginas(pdf_path: Path, max_paginas: int) -> None:
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        return
    reader = PdfReader(str(pdf_path))
    if len(reader.pages) <= max_paginas:
        return
    writer = PdfWriter()
    for i in range(max_paginas):
        writer.add_page(reader.pages[i])
    with open(pdf_path, "wb") as f:
        writer.write(f)


def _escribir_xlsx(
    path: Path,
    ranking: List[dict],
    usuarios: List[dict],
) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    top_fill = PatternFill("solid", fgColor="C6EFCE")

    ws = wb.active
    ws.title = "Ranking_por_nodo"
    cols = [
        "ranking",
        "cantidad_usuarios",
        "node_id",
        "nombre_nodo",
        "company_id",
        "empresa",
        "emails",
    ]
    for c, h in enumerate(cols, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.fill = header_fill
        cell.font = header_font
    for r, row in enumerate(ranking, 2):
        for c, h in enumerate(cols, 1):
            cell = ws.cell(row=r, column=c, value=row.get(h, ""))
            if h == "cantidad_usuarios" and int(row.get("cantidad_usuarios") or 0) > 0:
                cell.fill = top_fill
    widths = (10, 18, 14, 42, 12, 28, 80)
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.auto_filter.ref = f"A1:{get_column_letter(len(cols))}{max(1, len(ranking)+1)}"
    ws.freeze_panes = "A2"

    ws2 = wb.create_sheet("Solo_con_usuarios")
    con_users = [row for row in ranking if int(row["cantidad_usuarios"]) > 0]
    for c, h in enumerate(cols, 1):
        cell = ws2.cell(row=1, column=c, value=h)
        cell.fill = header_fill
        cell.font = header_font
    for r, row in enumerate(con_users, 2):
        for c, h in enumerate(cols, 1):
            ws2.cell(row=r, column=c, value=row.get(h, ""))
    for i, w in enumerate(widths, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"

    ws3 = wb.create_sheet("Por_usuario")
    cols_u = ["email", "nombre", "user_id", "company_id_usuario", "nodos_count", "nodos"]
    for c, h in enumerate(cols_u, 1):
        cell = ws3.cell(row=1, column=c, value=h)
        cell.fill = header_fill
        cell.font = header_font
    for r, row in enumerate(usuarios, 2):
        for c, h in enumerate(cols_u, 1):
            ws3.cell(row=r, column=c, value=row.get(h, ""))
            ws3.cell(row=r, column=c).alignment = Alignment(wrap_text=True)
    for i, w in enumerate((32, 28, 26, 18, 14, 80), 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = "A2"

    wb.save(path)


def _escribir_docx(
    path: Path,
    ranking: List[dict],
    n_emails: int,
    n_users: int,
    generado: str,
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Usuarios por Node ID")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Ranking de mayor a menor · generado {generado} (hora Chile)")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    con_users = [row for row in ranking if int(row["cantidad_usuarios"]) > 0]
    sin_users = len(ranking) - len(con_users)
    max_c = max((int(r["cantidad_usuarios"]) for r in ranking), default=0)
    doc.add_paragraph(
        f"Nodos en API: {len(ranking)} · Nodos con al menos 1 usuario "
        f"(en la muestra): {len(con_users)} · Nodos sin usuarios en la muestra: {sin_users}. "
        f"Correos consultados: {n_emails} · Usuarios encontrados en API: {n_users}. "
        f"Máximo de usuarios en un nodo: {max_c}."
    )
    nota = doc.add_paragraph()
    nr = nota.add_run(
        "Nota: la API WES no expone un listado completo de usuarios. "
        "El conteo cubre los correos conocidos en este proyecto (WES, Parque Arauco, "
        "Linkes y otros rastreados en el repositorio). "
        "Solo se incluyen puntos activos (se omiten dados de baja, sin instalación "
        "o fuera de operación según exclusiones_reportes y bajas PA)."
    )
    nr.italic = True
    nr.font.size = Pt(9)

    doc.add_heading("Ranking (nodos con usuarios)", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Usuarios", "Node ID", "Punto", "Empresa"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True

    for row in con_users:
        cells = table.add_row().cells
        cells[0].text = str(row["ranking"])
        cells[1].text = str(row["cantidad_usuarios"])
        cells[2].text = str(row["node_id"])
        cells[3].text = str(row["nombre_nodo"])
        cells[4].text = str(row["empresa"])

    if sin_users:
        doc.add_heading("Nodos sin usuarios en la muestra", level=1)
        doc.add_paragraph(
            ", ".join(r["node_id"] for r in ranking if int(r["cantidad_usuarios"]) == 0)
        )

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Cuenta usuarios WES por nodeId y ordena de mayor a menor"
    )
    parser.add_argument(
        "--emails",
        type=Path,
        action="append",
        default=[],
        help="Archivo adicional con correos (uno por línea)",
    )
    parser.add_argument("--salida", type=Path, default=OUT_DIR, help="Carpeta de salida")
    parser.add_argument("--workers", type=int, default=8, help="Consultas HTTP en paralelo")
    parser.add_argument(
        "--incluir-inactivos",
        action="store_true",
        help="Incluir puntos dados de baja / fuera de operación (por defecto se omiten)",
    )
    parser.add_argument(
        "--desde-xlsx",
        type=Path,
        default=None,
        help="Solo convertir un Excel existente a PDF (sin consultar la API)",
    )
    parser.add_argument(
        "--pdf-paginas",
        type=int,
        default=5,
        help="Máximo de páginas del PDF (default 5). 0 = sin recorte",
    )
    args = parser.parse_args()

    if args.desde_xlsx:
        src = args.desde_xlsx
        if not src.is_file():
            print(f"[ERROR] No existe el Excel: {src}")
            return 1
        pdf = xlsx_a_pdf(src, max_paginas=args.pdf_paginas)
        print(f"[OK] PDF: {pdf}")
        return 0

    session = _session()
    print("[INFO] Obteniendo nodos y empresas desde API...")
    nombres, company_de, companies = _nodos_y_empresas(session)
    print(f"[OK] {len(nombres)} nodeId únicos · {len(companies)} empresas")

    emails = _recolectar_emails(args.emails)
    workers = max(1, min(args.workers, 16))
    extra_alertas = _emails_desde_alertas(session, nombres.keys(), workers=workers)
    if extra_alertas:
        n_antes = len(emails)
        emails = sorted(set(emails) | extra_alertas)
        print(
            f"[INFO] Correos extra desde alertas FILTRATION: "
            f"{len(emails) - n_antes} (total {len(emails)})"
        )
    print(f"[INFO] Consultando {len(emails)} correo(s) candidatos...")

    usuarios_por_email: Dict[str, dict] = {}
    sin_usuario: List[str] = []
    with ThreadPoolExecutor(max_workers=workers) as ex:
        futs = [ex.submit(_fetch_user, session, e) for e in emails]
        for fut in as_completed(futs):
            email, user, obs = fut.result()
            if not user:
                sin_usuario.append(f"{email} ({obs})" if obs else email)
                continue
            email_api = str(user.get("username") or email).strip().lower()
            usuarios_por_email[email_api] = user
            n_nodes = len(_allowed_nodes(user))
            print(f"[OK] {email_api} — {n_nodes} nodo(s)")

    por_nodo: Dict[str, Set[str]] = defaultdict(set)
    filas_usuario: List[dict] = []
    solo_activos = not args.incluir_inactivos
    n_inactivos = 0
    for email_api, user in sorted(usuarios_por_email.items()):
        nodos = _allowed_nodes(user)
        nombre = nombre_usuario(user)
        cid_user = str(user.get("companyId") or "").strip()
        nodos_visibles: List[str] = []
        for nid in nodos:
            if nid not in nombres:
                nombres[nid] = ""
            if nid not in company_de and cid_user:
                company_de[nid] = cid_user
            cid = company_de.get(nid, cid_user)
            cname = companies.get(cid, "")
            if solo_activos and not es_nodo_activo(nid, cid, cname):
                continue
            por_nodo[nid].add(email_api)
            nodos_visibles.append(nid)
        if not nodos_visibles:
            continue
        filas_usuario.append(
            {
                "email": email_api,
                "nombre": nombre,
                "user_id": str(user.get("userId") or "").strip(),
                "company_id_usuario": cid_user,
                "nodos_count": len(nodos_visibles),
                "nodos": ", ".join(nodos_visibles),
            }
        )

    ranking: List[dict] = []
    for nid in nombres:
        cid = company_de.get(nid, "")
        cname = companies.get(cid, "")
        if solo_activos and not es_nodo_activo(nid, cid, cname):
            n_inactivos += 1
            continue
        users = sorted(por_nodo.get(nid, set()), key=str.casefold)
        ranking.append(
            {
                "cantidad_usuarios": len(users),
                "node_id": nid,
                "nombre_nodo": nombres.get(nid, ""),
                "company_id": cid,
                "empresa": cname or cid,
                "emails": ", ".join(users),
            }
        )
    ranking.sort(key=lambda r: (-int(r["cantidad_usuarios"]), r["node_id"]))
    for i, row in enumerate(ranking, 1):
        row["ranking"] = i
    if solo_activos:
        print(
            f"[INFO] Puntos inactivos omitidos: {n_inactivos} "
            f"(quedan {len(ranking)} activos)"
        )

    args.salida.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now(TZ_CL).strftime("%Y%m%d_%H%M")
    generado = datetime.now(TZ_CL).strftime("%d/%m/%Y %H:%M")
    csv_path = args.salida / f"usuarios_por_nodo_ranking_{stamp}.csv"
    xlsx_path = args.salida / f"usuarios_por_nodo_ranking_{stamp}.xlsx"
    docx_path = args.salida / f"usuarios_por_nodo_ranking_{stamp}.docx"
    csv_latest = args.salida / "usuarios_por_nodo_ranking.csv"
    xlsx_latest = args.salida / "usuarios_por_nodo_ranking.xlsx"
    docx_latest = args.salida / "usuarios_por_nodo_ranking.docx"
    pdf_latest = args.salida / "usuarios_por_nodo_ranking.pdf"
    pdf_path = None

    cols = [
        "ranking",
        "cantidad_usuarios",
        "node_id",
        "nombre_nodo",
        "company_id",
        "empresa",
        "emails",
    ]
    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=cols, delimiter=";", extrasaction="ignore")
        w.writeheader()
        w.writerows(ranking)
    csv_latest.write_bytes(csv_path.read_bytes())

    try:
        _escribir_xlsx(xlsx_path, ranking, filas_usuario)
        xlsx_latest.write_bytes(xlsx_path.read_bytes())
        try:
            pdf_stamp = xlsx_path.with_suffix(".pdf")
            pdf_path = xlsx_a_pdf(
                xlsx_path, pdf_stamp, solo_ranking=True, max_paginas=args.pdf_paginas
            )
            pdf_latest.write_bytes(pdf_path.read_bytes())
        except Exception as exc:
            pdf_path = None
            print(f"[WARN] No se pudo generar PDF: {exc}")
    except ImportError:
        xlsx_path = None
        print("[WARN] openpyxl no instalado; se omite XLSX")

    try:
        _escribir_docx(docx_path, ranking, len(emails), len(usuarios_por_email), generado)
        docx_latest.write_bytes(docx_path.read_bytes())
    except ImportError:
        docx_path = None
        print("[WARN] python-docx no instalado; se omite DOCX")

    con_users = sum(1 for r in ranking if int(r["cantidad_usuarios"]) > 0)
    print("\n=== Ranking (mayor a menor) ===")
    print(f"{'#':>4}  {'Users':>5}  {'Node ID':<12}  Punto")
    for row in ranking:
        if int(row["cantidad_usuarios"]) == 0:
            continue
        print(
            f"{row['ranking']:>4}  {row['cantidad_usuarios']:>5}  "
            f"{row['node_id']:<12}  {row['nombre_nodo']}"
        )
    print(f"\n[OK] Nodos con usuarios: {con_users} / {len(ranking)}")
    print(f"[OK] CSV:  {csv_path}")
    if xlsx_path:
        print(f"[OK] XLSX: {xlsx_path}")
    if pdf_path:
        print(f"[OK] PDF:  {pdf_path}")
    if docx_path:
        print(f"[OK] DOCX: {docx_path}")
    print(f"[INFO] Correos sin usuario API: {len(sin_usuario)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
