"""
Secciones del reporte agregado — formato extendido (Club Providencia, UDD, etc.).

- Evolución diaria del periodo por punto.
- Día de mayor consumo diario (perfil horario + conclusión).
- Análisis nocturno: suma real del periodo, costo CLP, sin proyección a 30 días.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Dict, List

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

AGREGADO_EXTENDIDO_COMPANY_IDS = frozenset({
    "000031",  # Club Providencia
    "000026",  # UDD
    "000009",  # COPEC
    "000007",  # Nido de Águilas
    "000020",  # AGUNSA (Lampa e Intermodal)
    "000021",  # CDUC
    "000027",  # Fundo Zapallar
    "000028",  # La Florida
    "000006",  # Colegios Providencia
    "000012",  # DERCO
    "000024",  # La Reina
    "000002",  # Lo Valledor
})

_CLIENTE: Dict[str, dict] = {
    "000031": {
        "prefijo": "cp",
        "sujeto": "El club",
        "sujeto_min": "el club",
        "total_label": "Total club (periodo):",
    },
    "000026": {
        "prefijo": "udd",
        "sujeto": "La UDD",
        "sujeto_min": "la UDD",
        "total_label": "Total UDD (periodo):",
    },
    "000009": {
        "prefijo": "copec",
        "sujeto": "COPEC",
        "sujeto_min": "COPEC",
        "total_label": "Total COPEC (periodo):",
    },
    "000007": {
        "prefijo": "aguilas",
        "sujeto": "Nido de Águilas",
        "sujeto_min": "Nido de Águilas",
        "total_label": "Total Nido de Águilas (periodo):",
    },
    "000020": {
        "prefijo": "agunsa",
        "sujeto": "AGUNSA",
        "sujeto_min": "AGUNSA",
        "total_label": "Total AGUNSA (periodo):",
    },
    "000021": {
        "prefijo": "cduc",
        "sujeto": "El CDUC",
        "sujeto_min": "el CDUC",
        "total_label": "Total CDUC (periodo):",
    },
    "000027": {
        "prefijo": "zapallar",
        "sujeto": "Fundo Zapallar",
        "sujeto_min": "Fundo Zapallar",
        "total_label": "Total Fundo Zapallar (periodo):",
    },
    "000028": {
        "prefijo": "florida",
        "sujeto": "La Florida",
        "sujeto_min": "La Florida",
        "total_label": "Total La Florida (periodo):",
    },
    "000006": {
        "prefijo": "providencia",
        "sujeto": "Los colegios de Providencia",
        "sujeto_min": "los colegios de Providencia",
        "total_label": "Total Providencia (periodo):",
    },
    "000012": {
        "prefijo": "inchcape",
        "sujeto": "Inchcape",
        "sujeto_min": "Inchcape",
        "total_label": "Total Inchcape (periodo):",
    },
    "000024": {
        "prefijo": "reina",
        "sujeto": "La Reina",
        "sujeto_min": "La Reina",
        "total_label": "Total La Reina (periodo):",
    },
    "000002": {
        "prefijo": "valledor",
        "sujeto": "Lo Valledor",
        "sujeto_min": "Lo Valledor",
        "total_label": "Total Lo Valledor (periodo):",
    },
}

HORARIO_NOCTURNO_TEXTO = (
    "Se define como consumo nocturno el volumen medido entre las 00:00 y las 06:59 "
    "(hora Chile), es decir 7 horas por cada día del periodo. El total nocturno es la "
    "suma de esas horas en todos los días con datos — no se extrapola ni proyecta a 30 días."
)


def es_agregado_extendido(company_id: str) -> bool:
    return company_id in AGREGADO_EXTENDIDO_COMPANY_IDS


def _cfg(company_id: str) -> dict:
    return _CLIENTE[company_id]


def _conclusion_dia_mayor(
    node_name: str,
    fecha: datetime,
    total_dia: float,
    promedio_diario: float,
) -> str:
    if promedio_diario <= 0:
        ratio_txt = "sin referencia de promedio diario en el periodo"
    else:
        ratio = (total_dia / promedio_diario - 1.0) * 100.0
        if ratio > 15:
            ratio_txt = (
                f"un {ratio:.0f} % por sobre el promedio diario del periodo "
                f"({promedio_diario:.1f} m³/día)"
            )
        elif ratio < -15:
            ratio_txt = (
                f"un {abs(ratio):.0f} % por debajo del promedio diario del periodo "
                f"({promedio_diario:.1f} m³/día)"
            )
        else:
            ratio_txt = (
                f"alineado con el promedio diario del periodo ({promedio_diario:.1f} m³/día)"
            )
    return (
        f"El día de mayor consumo diario en {node_name} fue el {fecha.strftime('%d/%m/%Y')} "
        f"con {total_dia:.1f} m³ registrados, {ratio_txt}. Conviene revisar en terreno qué "
        f"actividades o equipos explican el pico horario de ese día y si corresponde a "
        f"operación programada o a un uso extraordinario de agua."
    )


def _conclusion_nocturno_agregado(
    company_id: str,
    total_m3: float,
    total_clp: float,
    nodos: List[dict],
    num_dias: int,
) -> str:
    from generar_reporte_word import format_currency_chilean, format_number_chilean

    c = _cfg(company_id)
    if total_m3 <= 0:
        return (
            "En el periodo analizado no se registró consumo en horario nocturno (00:00–06:59) "
            "en los puntos monitoreados, o no hubo datos horarios suficientes."
        )
    partes = [f"{n['nombre']} ({format_number_chilean(n['m3'], 1)} m³)" for n in nodos]
    if len(partes) > 3:
        top = sorted(nodos, key=lambda x: x["m3"], reverse=True)[:3]
        otros_m3 = total_m3 - sum(t["m3"] for t in top)
        detalle_top = ", ".join(
            f"{t['nombre']} ({format_number_chilean(t['m3'], 1)} m³)" for t in top
        )
        detalle = (
            f"{detalle_top}"
            + (f" y {format_number_chilean(otros_m3, 1)} m³ en los demás puntos" if otros_m3 > 0 else "")
        )
    else:
        detalle = " y ".join(partes)
    return (
        f"Durante los {num_dias} días del periodo, {c['sujeto_min']} registró un total de "
        f"{format_number_chilean(total_m3, 1)} m³ en horario nocturno ({detalle}), equivalente a "
        f"{format_currency_chilean(total_clp)} al precio de referencia del servicio. Ese volumen "
        f"corresponde a agua consumida en madrugada cuando la operación debería ser mínima; un "
        f"sistema de monitoreo y regulación WES permite detectar y reducir esos caudales, "
        f"mejorando el control del recurso y el costo asociado."
    )


def narrativa_consumo_total_extendido(company_id: str, nodes_data: List[dict]) -> str:
    from generar_reporte_word import format_number_chilean

    c = _cfg(company_id)
    if not nodes_data:
        return ""
    ordenados = sorted(nodes_data, key=lambda d: d["summary"]["total"], reverse=True)
    total = sum(d["summary"]["total"] for d in ordenados)
    if total <= 0:
        return "No se registró consumo en el periodo analizado."

    if len(ordenados) == 1:
        n = ordenados[0]
        nombre = n["node_name"].replace("\n", " ").strip()
        return (
            f"En el periodo, el único punto monitoreado ({nombre}) registró "
            f"{format_number_chilean(n['summary']['total'], 1)} m³."
        )

    if len(ordenados) == 2:
        mayor = ordenados[0]
        menor = ordenados[-1]
        nom_mayor = mayor["node_name"].replace("\n", " ").strip()
        nom_menor = menor["node_name"].replace("\n", " ").strip()
        cons_mayor = float(mayor["summary"]["total"])
        cons_menor = float(menor["summary"]["total"])
        pct_mayor = cons_mayor / total * 100.0
        pct_menor = cons_menor / total * 100.0
        diferencia = cons_mayor - cons_menor
        return (
            f"{c['sujeto']} consumió en total {format_number_chilean(total, 1)} m³ en el periodo. "
            f"De ese volumen, {nom_mayor} aportó {format_number_chilean(cons_mayor, 1)} m³ "
            f"({format_number_chilean(pct_mayor, 1)} %) y {nom_menor} "
            f"{format_number_chilean(cons_menor, 1)} m³ ({format_number_chilean(pct_menor, 1)} %). "
            f"{nom_mayor} fue el punto de mayor demanda, con "
            f"{format_number_chilean(diferencia, 1)} m³ más que {nom_menor}."
        )

    top3 = ordenados[:3]
    menor = ordenados[-1]
    nom_menor = menor["node_name"].replace("\n", " ").strip()
    cons_menor = float(menor["summary"]["total"])
    pct_menor = cons_menor / total * 100.0 if total else 0.0
    tops_txt = ", ".join(
        f"{n['node_name'].replace('\n', ' ').strip()} "
        f"({format_number_chilean(n['summary']['total'], 1)} m³, "
        f"{format_number_chilean(float(n['summary']['total']) / total * 100.0, 1)} %)"
        for n in top3
    )
    return (
        f"{c['sujeto']} registró un consumo total de {format_number_chilean(total, 1)} m³ en el periodo "
        f"entre {len(ordenados)} puntos de monitoreo. Los de mayor demanda fueron: {tops_txt}. "
        f"El menor consumo correspondió a {nom_menor} "
        f"({format_number_chilean(cons_menor, 1)} m³, {format_number_chilean(pct_menor, 1)} % del total)."
    )


COPEC_NODE_ESTANQUE_REUTILIZACION = "000009-02"
AGUILAS_NODE_PISCINA = "000007-05"
AGUILAS_NODE_ELEMENTARY = "000007-04"
INCHCAPE_NODE_MATRIZ_PRINCIPAL = "000012-06"

# Sin sección «Día de mayor consumo» ni marcadores/tabla de alertas en rojo.
OMITIR_DIA_MAYOR_Y_ALERTAS_ROJAS = frozenset({
    "000027",  # Fundo Zapallar
    "000012",  # Inchcape (ex DERCO)
    "000002",  # Lo Valledor
    "000026",  # UDD
    "000031",  # Club Providencia
    "000020",  # AGUNSA (Lampa e Intermodal)
})


def _omitir_grafico_dia_mayor(company_id: str, data: dict) -> bool:
    if company_id in OMITIR_DIA_MAYOR_Y_ALERTAS_ROJAS:
        return True
    summary = data.get("summary") or {}
    total = float(summary.get("total") or 0.0)
    if total <= 0:
        return True
    if company_id == "000007" and data.get("node_id") == AGUILAS_NODE_PISCINA:
        return True
    # Elementary con consumo residual (turbina pendiente): no tiene sentido el pico diario.
    if company_id == "000007" and data.get("node_id") == AGUILAS_NODE_ELEMENTARY and total < 1.0:
        return True
    max_m = summary.get("max")
    if max_m and float(max_m.total_m3 or 0) <= 0:
        return True
    return False


def _nota_estanque_reutilizacion_copec() -> str:
    return (
        "Nota — Estanque Reutilización: este punto corresponde a un estanque que almacena "
        "agua reutilizada y que, cuando el aporte de reutilización no es suficiente, puede "
        "completarse con agua potable desde la red. En el periodo analizado el medidor registra "
        "consumo cero (0 m³), lo que se explica porque el estanque ha sido rellenado con agua "
        "potable desde otros puntos de monitoreo cuando fue necesario mantener su nivel operativo. "
        "Esto no indica una falla del sistema: la reutilización ha operado de forma normal en el periodo."
    )


def _nota_turbina_elementary() -> str:
    return (
        "Nota — Elementary (turbina): se planificó la limpieza de la turbina de este punto; "
        "sin embargo, por condiciones climáticas y por necesidades operativas del colegio, "
        "el trabajo aún no se ha ejecutado. Por eso el medidor marca consumo cercano a cero "
        "en el periodo: no corresponde a una falla del sistema WES, sino a que la intervención "
        "programada quedó pendiente de realizar."
    )


def _agregar_nota_estanque_reutilizacion_copec(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run("Detalle operativo").bold = True
    nota = doc.add_paragraph(_nota_estanque_reutilizacion_copec())
    nota.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in nota.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _agregar_nota_turbina_elementary(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run("Detalle operativo").bold = True
    nota = doc.add_paragraph(_nota_turbina_elementary())
    nota.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in nota.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _agregar_destacado_matriz_principal_inchcape(
    doc: Document,
    nodes_data: List[dict],
    nodos_noct: List[dict],
) -> None:
    """Resalta control nocturno sobre Quilicura Matriz Principal."""
    from generar_reporte_word import add_formatted_title, format_number_chilean

    matriz = next(
        (d for d in nodes_data if d.get("node_id") == INCHCAPE_NODE_MATRIZ_PRINCIPAL),
        None,
    )
    noct = next(
        (
            n
            for n in nodos_noct
            if "matriz principal" in str(n.get("nombre", "")).lower()
        ),
        None,
    )
    if matriz is None and noct is None:
        return

    add_formatted_title(doc, "Control destacado — Matriz Principal")
    total_periodo = float((matriz.get("summary") or {}).get("total") or 0.0) if matriz else 0.0
    m3_noct = float(noct["m3"]) if noct else 0.0
    dias_noct = int(noct["dias_con"]) if noct else 0
    pct = (m3_noct / total_periodo * 100.0) if total_periodo > 0 else 0.0
    texto = (
        "El nodo Quilicura Matriz Principal es el punto de control prioritario del sitio: "
        f"concentra {format_number_chilean(total_periodo, 1)} m³ del periodo y "
        f"{format_number_chilean(m3_noct, 1)} m³ en horario nocturno (00:00–06:59), "
        f"equivalente al {format_number_chilean(pct, 1)} % de su consumo del periodo "
        f"({dias_noct} días con caudal en madrugada). "
        "Conviene priorizar la revisión y regulación WES sobre esta matriz para reducir "
        "consumos fuera de operación y reforzar el control del recurso en el predio."
    )
    p = doc.add_paragraph(texto)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def agregar_secciones_consumo_diario_y_max_dia(
    company_id: str,
    doc: Document,
    nodes_data: List[dict],
    start_dt: datetime,
    end_dt: datetime,
    output_dir: Path,
) -> None:
    from generar_reporte_word import (
        add_formatted_heading,
        add_formatted_title,
        add_picture_with_pagination,
        agregar_tabla_alertas_grafico_diario,
        alertas_marcadas_grafico_diario,
        build_consumption_chart,
        build_hourly_consumption_line_chart,
        filtrar_alertas_informativas,
        get_hourly_measures_for_day,
    )

    pref = _cfg(company_id)["prefijo"]

    add_formatted_heading(doc, "Evolución del consumo diario por punto", level=1)
    intro = doc.add_paragraph(
        "Gráficos de consumo total diario (m³) de cada día del periodo analizado, "
        "por punto de monitoreo."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for data in nodes_data:
        node_id = data["node_id"]
        node_name = data["node_name"]
        measures = data.get("measures") or []
        if not measures:
            continue
        chart_path = output_dir / f"{pref}_diario_{node_id.replace('-', '_')}.png"
        alerts = filtrar_alertas_informativas(data.get("alerts"))
        # Zapallar / Inchcape: solo curva de consumo (sin marcadores ni tabla de alertas en rojo).
        alerts_para_grafico = None if company_id in OMITIR_DIA_MAYOR_Y_ALERTAS_ROJAS else alerts
        built = build_consumption_chart(
            measures, chart_path, start_dt, end_dt, alerts_para_grafico
        )
        if not built or not chart_path.is_file():
            continue
        doc.add_paragraph("")
        add_formatted_title(doc, node_name.upper())
        add_picture_with_pagination(doc, str(chart_path), Inches(6), keep_with_next=True)
        if (
            es_agregado_extendido(company_id)
            and company_id not in OMITIR_DIA_MAYOR_Y_ALERTAS_ROJAS
        ):
            alerts_marcadas = alertas_marcadas_grafico_diario(alerts, measures, start_dt, end_dt)
            if alerts_marcadas:
                agregar_tabla_alertas_grafico_diario(doc, alerts_marcadas, wes_style=True)
        if company_id == "000009" and node_id == COPEC_NODE_ESTANQUE_REUTILIZACION:
            total_periodo = float((data.get("summary") or {}).get("total") or 0.0)
            if total_periodo <= 0:
                _agregar_nota_estanque_reutilizacion_copec(doc)
        if company_id == "000007" and node_id == AGUILAS_NODE_ELEMENTARY:
            total_periodo = float((data.get("summary") or {}).get("total") or 0.0)
            if total_periodo < 1.0:
                _agregar_nota_turbina_elementary(doc)

    # Zapallar / Inchcape: el cliente pide omitir «día de mayor consumo» por punto.
    if company_id in OMITIR_DIA_MAYOR_Y_ALERTAS_ROJAS:
        return

    add_formatted_heading(doc, "Día de mayor consumo diario por punto", level=1)
    intro2 = doc.add_paragraph(
        "Para cada punto se identifica el día con mayor consumo acumulado del periodo "
        "y se muestra el perfil horario (m³/h) de ese día."
    )
    intro2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for data in nodes_data:
        node_id = data["node_id"]
        node_name = data["node_name"]
        summary = data.get("summary") or {}
        if _omitir_grafico_dia_mayor(company_id, data):
            continue
        max_m = summary.get("max")
        if not max_m:
            continue
        max_day_dt = datetime.combine(max_m.date.date(), datetime.min.time())
        total_dia = float(max_m.total_m3)
        promedio = float(summary.get("promedio_diario") or 0.0)
        hourly = get_hourly_measures_for_day(node_id, max_day_dt)
        if not hourly:
            continue
        chart_path = output_dir / f"{pref}_max_dia_{node_id.replace('-', '_')}.png"
        built = build_hourly_consumption_line_chart(
            hourly,
            chart_path,
            max_day_dt,
            f"Mayor consumo diario — {max_day_dt.strftime('%d-%m-%Y')}",
        )
        if not built or not chart_path.is_file():
            continue
        doc.add_paragraph("")
        add_formatted_title(doc, f"{node_name.upper()} — {max_day_dt.strftime('%d/%m/%Y')}")
        add_picture_with_pagination(doc, str(chart_path), Inches(6), keep_with_next=True)
        conc = doc.add_paragraph(_conclusion_dia_mayor(node_name, max_day_dt, total_dia, promedio))
        conc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in conc.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)


def agregar_analisis_nocturno_extendido(
    company_id: str,
    doc: Document,
    nodes_data: List[dict],
    start_dt: datetime,
    end_dt: datetime,
    output_dir: Path,
    price_per_m3: float,
) -> None:
    from generar_reporte_word import (
        add_formatted_heading,
        add_formatted_title,
        apply_keep_with_next,
        calculate_nocturnal_metrics,
        estilizar_tabla_wes,
        format_currency_chilean,
        format_number_chilean,
    )

    pref = _cfg(company_id)["prefijo"]
    total_label = _cfg(company_id)["total_label"]
    num_dias = (end_dt.date() - start_dt.date()).days + 1
    nodos_noct: List[dict] = []
    names: List[str] = []
    values: List[float] = []

    for data in nodes_data:
        node_id = data["node_id"]
        node_name = data["node_name"].replace("\n", " ").strip()
        nm = calculate_nocturnal_metrics(node_id, start_dt, end_dt, company_id=company_id)
        c_noche = float(nm["consumo_nocturno_total"])
        nodos_noct.append(
            {
                "nombre": node_name,
                "m3": c_noche,
                "dias_con": int(nm["dias_con_consumo_nocturno"]),
                "clp": c_noche * price_per_m3,
            }
        )
        names.append(node_name)
        values.append(c_noche)

    total_m3 = sum(values)
    total_clp = total_m3 * price_per_m3

    add_formatted_heading(doc, "Análisis de consumos nocturnos", level=1, page_break_before=True)

    p_hor = doc.add_paragraph()
    lbl = p_hor.add_run("Horario nocturno considerado: ")
    lbl.bold = True
    lbl.font.color.rgb = RGBColor(0, 51, 102)
    txt = p_hor.add_run(HORARIO_NOCTURNO_TEXTO)
    txt.font.size = Pt(9)
    p_hor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_hor.paragraph_format.space_after = Pt(4)
    apply_keep_with_next(p_hor)

    chart_path = None
    if names and any(v > 0 for v in values):
        sorted_pairs = sorted(zip(names, values), key=lambda x: x[1], reverse=True)
        names = [n for n, _ in sorted_pairs]
        values = [v for _, v in sorted_pairs]
        n_pts = len(names)
        fig_w = max(7.0, min(9.5, n_pts * 0.85))
        fig_h = 2.5 if n_pts > 5 else 3.2
        rot = 48 if n_pts > 4 else 28
        fs = 7 if n_pts > 6 else 8
        fig, ax = plt.subplots(figsize=(fig_w, fig_h))
        bars = ax.bar(names, values, color="#FF8C00", alpha=0.88, edgecolor="#CC7000", linewidth=1.0)
        ax.set_ylabel("m³ (periodo)", fontsize=10, fontweight="bold")
        ax.set_title("Consumo nocturno por punto", fontsize=11, fontweight="bold", pad=8)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.set_ylim(bottom=0)
        plt.setp(ax.xaxis.get_majorticklabels(), rotation=rot, ha="right", fontsize=fs)
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height(),
                    f"{format_number_chilean(val, 1)}",
                    ha="center",
                    va="bottom",
                    fontsize=7,
                    fontweight="bold",
                )
        plt.tight_layout()
        chart_path = output_dir / f"{pref}_consumo_nocturno_periodo.png"
        plt.savefig(chart_path, dpi=140, bbox_inches="tight")
        plt.close()
        img_w = Inches(5.2) if n_pts > 5 else Inches(5.8)
        img_h = Inches(2.0) if n_pts > 5 else Inches(2.4)
        pic_para = doc.add_paragraph()
        pic_para.add_run().add_picture(str(chart_path), width=img_w, height=img_h)
        pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic_para.paragraph_format.space_before = Pt(2)
        pic_para.paragraph_format.space_after = Pt(4)
        apply_keep_with_next(pic_para)

    add_formatted_title(doc, "Detalle por punto")

    table_rows = [
        ["Punto", "m³ nocturnos", "Días 00–06:59", "Costo (CLP)"],
    ]
    for n in sorted(nodos_noct, key=lambda x: x["m3"], reverse=True):
        table_rows.append(
            [
                n["nombre"],
                format_number_chilean(n["m3"], 1),
                str(n["dias_con"]),
                format_currency_chilean(n["clp"]),
            ]
        )
    table_rows.append(
        [
            total_label,
            format_number_chilean(total_m3, 1),
            "",
            format_currency_chilean(total_clp),
        ]
    )

    table = doc.add_table(rows=len(table_rows), cols=4)
    table.style = "Table Grid"
    for i, row_vals in enumerate(table_rows):
        for j, val in enumerate(row_vals):
            cell = table.rows[i].cells[j]
            cell.text = str(val)
    estilizar_tabla_wes(table)
    if table.rows:
        apply_keep_with_next(table.rows[0].cells[0].paragraphs[0])

    add_formatted_title(doc, "Conclusión — consumo nocturno")
    conc = doc.add_paragraph(
        _conclusion_nocturno_agregado(company_id, total_m3, total_clp, nodos_noct, num_dias)
    )
    conc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conc.paragraph_format.space_before = Pt(2)
    for run in conc.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(10)

    if company_id == "000012":
        _agregar_destacado_matriz_principal_inchcape(doc, nodes_data, nodos_noct)
    elif company_id == "000027":
        # Destacar el volumen nocturno real del periodo (sin día de máximo consumo).
        from generar_reporte_word import add_formatted_title, format_currency_chilean, format_number_chilean

        add_formatted_title(doc, "Énfasis operativo — periodo nocturno")
        esval_noct = next(
            (n for n in nodos_noct if "esval" in str(n.get("nombre", "")).lower() or "matriz" in str(n.get("nombre", "")).lower()),
            nodos_noct[0] if nodos_noct else None,
        )
        if esval_noct and total_m3 > 0:
            enf = doc.add_paragraph(
                "En Fundo Zapallar el foco de control queda en el caudal nocturno del periodo "
                f"({format_number_chilean(total_m3, 1)} m³; {format_currency_chilean(total_clp)}). "
                f"La Matriz ESVAL concentra {format_number_chilean(esval_noct['m3'], 1)} m³ en madrugada "
                f"({esval_noct['dias_con']} días con consumo 00:00–06:59): es el punto de entrada "
                "donde conviene priorizar la regulación y el seguimiento WES."
            )
            enf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in enf.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
