"""
Informe UDD — comparación pre/referencia vs monitoreo post-reparación (Honduras).

Ventanas:
  Sin reparación / referencia: 29/06/2026 – 05/07/2026 (fija)
  Monitoreo post-reparación:   últimos 7 días completos hasta la fecha
  Checkpoint previo:           27/07/2026 – 02/08/2026 (informe ago2026 original)
  Reparación de fugas:         12/07/2026

Uso:
  python generar_informe_udd_comparacion_post_reparacion.py
"""

from __future__ import annotations

import csv
import shutil
import subprocess
import sys
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

if sys.platform == "win32":
    for s in (sys.stdout, sys.stderr):
        try:
            s.reconfigure(encoding="utf-8", line_buffering=True)
        except Exception:
            pass

from generar_reporte_word import (
    add_logo_to_header,
    convertir_word_a_pdf,
    estilizar_tabla_wes,
    format_currency_chilean,
    format_number_chilean,
    get_hourly_measures_for_day,
)

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "UDD" / "COMPARACION_POST_REPARACION_ago2026"
GRAFICOS = OUT_DIR / "graficos"

NODE_ID = "000026-01"
NODE_NAME = "Sala impulsión Honduras"
CLIENTE = "Universidad del Desarrollo (UDD)"
DIA_REPARACION = date(2026, 7, 12)

SIN_INI, SIN_FIN = date(2026, 6, 29), date(2026, 7, 5)
# Checkpoint del informe original (ago2026, generado el 03/08)
POST_PREV_INI, POST_PREV_FIN = date(2026, 7, 27), date(2026, 8, 2)
# Se rellenan en generar() con los últimos 7 días completos
POST_INI = date(2026, 7, 27)
POST_FIN = date(2026, 8, 2)

COLOR_SIN = "#C0392B"
COLOR_POST = "#0050b3"
COLOR_NOCT = "#FFD700"
COLOR_PREV = "#7F8C8D"
COLOR_OK = "#1E8449"
PRECIO = 1200.0
DIAS_ES = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]

Fila = Tuple[date, float, float, Dict[int, float]]
_CACHE: Dict[date, Fila] = {}


def _rango(ini: date, fin: date) -> List[date]:
    out = []
    d = ini
    while d <= fin:
        out.append(d)
        d += timedelta(days=1)
    return out


def _consumo_dia(d: date) -> Fila:
    """Retorna (fecha, total_m3, nocturno_00_06, horas)."""
    if d in _CACHE:
        return _CACHE[d]
    hourly = get_hourly_measures_for_day(NODE_ID, datetime.combine(d, datetime.min.time()))
    horas: Dict[int, float] = {h: 0.0 for h in range(24)}
    if hourly:
        for h, v in hourly:
            horas[int(h)] = horas.get(int(h), 0.0) + float(v)
    total = sum(horas.values())
    noct = sum(horas[h] for h in range(0, 7))
    fila: Fila = (d, total, noct, horas)
    _CACHE[d] = fila
    return fila


def _serie(ini: date, fin: date) -> List[Fila]:
    rows = []
    for d in _rango(ini, fin):
        fila = _consumo_dia(d)
        rows.append(fila)
        print(f"  {d}: total={fila[1]:.1f} noct={fila[2]:.1f}", flush=True)
    return rows


def _horas_con_dato(horas: Dict[int, float]) -> int:
    return sum(1 for v in horas.values() if v > 1e-9)


def _ventana_ultimos_7_completos() -> Tuple[date, date, bool]:
    """Últimos 7 días civiles con dato horario suficiente. Incluye hoy si tiene ≥20 h."""
    hoy = date.today()
    fila = _consumo_dia(hoy)
    incluye_hoy = _horas_con_dato(fila[3]) >= 20
    fin = hoy if incluye_hoy else hoy - timedelta(days=1)
    ini = fin - timedelta(days=6)
    return ini, fin, incluye_hoy


def _semanas_lun_dom(ini: date, fin: date) -> List[Tuple[date, date]]:
    d = ini
    while d.weekday() != 0:
        d += timedelta(days=1)
    weeks: List[Tuple[date, date]] = []
    while True:
        w_end = d + timedelta(days=6)
        if w_end > fin:
            break
        weeks.append((d, w_end))
        d += timedelta(days=7)
    return weeks


def _totales(serie: List[Fila]) -> Tuple[float, float, float]:
    n = len(serie) or 1
    tot = sum(r[1] for r in serie)
    noct = sum(r[2] for r in serie)
    return tot, noct, tot / n


def _pct(delta: float, base: float) -> float:
    return (100.0 * delta / base) if base > 0 else 0.0


def _plot_comparativo_diario(
    sin: List[Fila],
    post: List[Fila],
    out: Path,
) -> None:
    fig, ax = plt.subplots(figsize=(10, 5.2))
    x = list(range(1, max(len(sin), len(post)) + 1))
    v_sin = [r[1] for r in sin]
    v_post = [r[1] for r in post]
    labels = []
    for i in range(max(len(sin), len(post))):
        a = sin[i][0].strftime("%d/%m") if i < len(sin) else "—"
        b = post[i][0].strftime("%d/%m") if i < len(post) else "—"
        labels.append(f"{a}\nvs\n{b}")

    w = 0.38
    xs = [i - w / 2 for i in x[: len(v_sin)]]
    xp = [i + w / 2 for i in x[: len(v_post)]]
    ax.bar(xs, v_sin, width=w, color=COLOR_SIN, label=f"Sin / ref. {SIN_INI:%d/%m}–{SIN_FIN:%d/%m}")
    ax.bar(xp, v_post, width=w, color=COLOR_POST, label=f"Post {POST_INI:%d/%m}–{POST_FIN:%d/%m}")
    ax.set_xticks(x[: len(labels)])
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_ylabel("Consumo diario (m³)", fontweight="bold")
    ax.set_title("Comparativo diario — UDD Honduras (pareado por orden de día)", fontweight="bold")
    ax.legend(fontsize=9)
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def _plot_totales(tot_sin: float, tot_post: float, out: Path) -> None:
    fig, ax = plt.subplots(figsize=(7, 4.8))
    labels = [
        f"Sin reparación\n{SIN_INI:%d/%m}–{SIN_FIN:%d/%m}",
        f"Post-reparación\n{POST_INI:%d/%m}–{POST_FIN:%d/%m}",
    ]
    vals = [tot_sin, tot_post]
    bars = ax.bar(labels, vals, color=[COLOR_SIN, COLOR_POST], width=0.55)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 1)} m³",
            ha="center",
            va="bottom",
            fontweight="bold",
        )
    delta = tot_sin - tot_post
    pct = _pct(delta, tot_sin)
    signo = "↓" if delta >= 0 else "↑"
    ax.set_title(
        f"Consumo acumulado — variación {signo} {format_number_chilean(abs(delta), 1)} m³ "
        f"({format_number_chilean(abs(pct), 1)} %)",
        fontsize=11,
        fontweight="bold",
    )
    ax.set_ylabel("m³ acumulados")
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def _plot_detalle_post(post: List[Fila], out: Path) -> None:
    labels = [f"{DIAS_ES[d.weekday()][:3]} {d.strftime('%d/%m')}" for d, *_ in post]
    totals = [t for _, t, _, _ in post]
    nocts = [n for _, _, n, _ in post]
    fig, ax = plt.subplots(figsize=(10, 4.8))
    x = range(len(labels))
    ax.bar([i - 0.2 for i in x], totals, width=0.4, color=COLOR_POST, label="Total día")
    ax.bar(
        [i + 0.2 for i in x],
        nocts,
        width=0.4,
        color=COLOR_NOCT,
        edgecolor="#DAA520",
        label="Nocturno 00–06",
    )
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=20, ha="right", fontsize=9)
    ax.set_ylabel("m³")
    ax.set_title(
        f"Detalle monitoreo post-reparación ({POST_INI:%d/%m}–{POST_FIN:%d/%m})",
        fontweight="bold",
    )
    ax.legend()
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_ylim(bottom=0)
    for i, (t, n) in enumerate(zip(totals, nocts)):
        ax.text(i - 0.2, t, f"{t:.0f}", ha="center", va="bottom", fontsize=8)
        if n > 0:
            ax.text(i + 0.2, n, f"{n:.0f}", ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def _plot_perfil_dia(horas: Dict[int, float], dia: date, out: Path, titulo: str) -> None:
    hs = list(range(24))
    vals = [float(horas.get(h, 0.0)) for h in hs]
    fig, ax = plt.subplots(figsize=(9.5, 4.2))
    ax.plot(hs, vals, color=COLOR_POST, marker="o", markersize=4, linewidth=2)
    ax.fill_between(hs, vals, alpha=0.28, color=COLOR_POST)
    ax.axvspan(-0.5, 6.5, color=COLOR_NOCT, alpha=0.18, label="00:00–06:59")
    ax.set_xticks(hs)
    ax.set_xticklabels([f"{h:02d}" for h in hs], fontsize=8)
    ax.set_xlabel("Hora")
    ax.set_ylabel("m³/h")
    ax.set_title(
        f"{titulo} — {DIAS_ES[dia.weekday()]} {dia.strftime('%d/%m/%Y')}",
        fontweight="bold",
    )
    ax.set_ylim(bottom=0)
    ax.legend(fontsize=8)
    ax.grid(True, linestyle="--", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _plot_serie_diaria(
    serie: List[Fila],
    prom_ref: float,
    prom_noct_ref: float,
    out: Path,
) -> None:
    fechas = [datetime.combine(r[0], datetime.min.time()) for r in serie]
    totales = [r[1] for r in serie]
    nocts = [r[2] for r in serie]
    fig, ax = plt.subplots(figsize=(11.2, 5.4))
    ax.plot(fechas, totales, color=COLOR_POST, linewidth=1.8, marker="o", markersize=3.2, label="Total diario")
    ax.fill_between(fechas, totales, alpha=0.18, color=COLOR_POST)
    ax.plot(
        fechas,
        nocts,
        color="#DAA520",
        linewidth=1.5,
        marker="s",
        markersize=3,
        label="Nocturno 00–06",
    )
    ax.axhline(prom_ref, color=COLOR_SIN, linestyle="--", linewidth=1.4, label=f"Prom. ref. {prom_ref:.1f} m³/día")
    ax.axhline(
        prom_noct_ref,
        color="#B9770E",
        linestyle=":",
        linewidth=1.3,
        label=f"Nocturno ref. {prom_noct_ref:.1f} m³/día",
    )
    ax.axvline(
        datetime.combine(DIA_REPARACION, datetime.min.time()),
        color=COLOR_PREV,
        linestyle=":",
        linewidth=1.5,
        label=f"Reparación {DIA_REPARACION:%d/%m}",
    )
    ax.axvspan(
        datetime.combine(SIN_INI, datetime.min.time()),
        datetime.combine(SIN_FIN, datetime.min.time()),
        color=COLOR_SIN,
        alpha=0.08,
        label="Ventana referencia",
    )
    ax.axvspan(
        datetime.combine(POST_INI, datetime.min.time()),
        datetime.combine(POST_FIN, datetime.min.time()),
        color=COLOR_POST,
        alpha=0.08,
        label="Últimos 7 días",
    )
    ax.set_title("Serie diaria — ¿se sostiene el ahorro post-reparación?", fontweight="bold")
    ax.set_ylabel("m³")
    ax.set_ylim(bottom=0)
    ax.xaxis.set_major_locator(mdates.WeekdayLocator(byweekday=mdates.MO))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m"))
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=35, ha="right")
    ax.grid(True, linestyle="--", alpha=0.3)
    ax.legend(loc="upper right", fontsize=8, ncol=2)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def _plot_semanas(
    semanas: List[Tuple[str, float, float]],
    out: Path,
) -> None:
    labels = [s[0] for s in semanas]
    tots = [s[1] for s in semanas]
    nocts = [s[2] for s in semanas]
    colors = []
    for lab in labels:
        if lab.startswith("Ref"):
            colors.append(COLOR_SIN)
        else:
            colors.append(COLOR_POST)
    fig, ax = plt.subplots(figsize=(10.5, 5.0))
    x = list(range(len(labels)))
    ax.bar([i - 0.18 for i in x], tots, width=0.36, color=colors, label="Total semana")
    ax.bar(
        [i + 0.18 for i in x],
        nocts,
        width=0.36,
        color=COLOR_NOCT,
        edgecolor="#DAA520",
        label="Nocturno 00–06",
    )
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=8, rotation=15, ha="right")
    ax.set_ylabel("m³ / 7 días")
    ax.set_title("Evolución semanal (lunes–domingo) vs referencia", fontweight="bold")
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_ylim(bottom=0)
    ax.legend(fontsize=8)
    for i, (t, n) in enumerate(zip(tots, nocts)):
        ax.text(i - 0.18, t, f"{t:.0f}", ha="center", va="bottom", fontsize=7.5)
        ax.text(i + 0.18, n, f"{n:.0f}", ha="center", va="bottom", fontsize=7.5)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def _plot_tres_ventanas(tot_sin: float, tot_prev: float, tot_post: float, out: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    labels = [
        f"Referencia\n{SIN_INI:%d/%m}–{SIN_FIN:%d/%m}",
        f"Checkpoint ago\n{POST_PREV_INI:%d/%m}–{POST_PREV_FIN:%d/%m}",
        f"Últimos 7 días\n{POST_INI:%d/%m}–{POST_FIN:%d/%m}",
    ]
    vals = [tot_sin, tot_prev, tot_post]
    bars = ax.bar(labels, vals, color=[COLOR_SIN, COLOR_PREV, COLOR_POST], width=0.55)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 1)} m³",
            ha="center",
            va="bottom",
            fontweight="bold",
            fontsize=9,
        )
    ax.set_title("Persistencia del ahorro — 3 ventanas de 7 días", fontweight="bold")
    ax.set_ylabel("m³ acumulados")
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def _plot_tres_nocturno(noct_sin: float, noct_prev: float, noct_post: float, out: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    labels = [
        f"Referencia\n{SIN_INI:%d/%m}–{SIN_FIN:%d/%m}",
        f"Checkpoint ago\n{POST_PREV_INI:%d/%m}–{POST_PREV_FIN:%d/%m}",
        f"Últimos 7 días\n{POST_INI:%d/%m}–{POST_FIN:%d/%m}",
    ]
    vals = [noct_sin, noct_prev, noct_post]
    bars = ax.bar(labels, vals, color=[COLOR_SIN, COLOR_PREV, COLOR_NOCT], width=0.55, edgecolor="#DAA520")
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 1)} m³",
            ha="center",
            va="bottom",
            fontweight="bold",
            fontsize=9,
        )
    ax.set_title("Indicador de fugas — nocturno 00:00–06:59 (7 días)", fontweight="bold")
    ax.set_ylabel("m³ nocturnos acumulados")
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def _add_table(doc: Document, headers: List[str], rows: List[List[str]], *, total=False) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            table.rows[r_i + 1].cells[c_i].text = str(val)
    estilizar_tabla_wes(table, has_total_row=total)
    doc.add_paragraph("")


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>')
    )


def _set_cell_margins(cell, *, top=40, bottom=40, left=80, right=80) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is not None:
        tc_pr.remove(tc_mar)
    tc_mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _set_table_borders(table, *, color="0050B3", sz="8") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge in ("insideH", "insideV"):
            el.set(qn("w:val"), "nil")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _clear_para(para) -> None:
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    para.text = ""


def _fill_cell_runs(cell, parts: List[Tuple[str, dict]]) -> None:
    para = cell.paragraphs[0]
    _clear_para(para)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.05
    for text, style in parts:
        run = para.add_run(text)
        run.bold = bool(style.get("bold"))
        run.font.size = Pt(style.get("size", 10))
        run.font.name = style.get("name", "Calibri")
        color = style.get("color")
        if color is not None:
            run.font.color.rgb = color


def add_ficha_informe(doc: Document, campos: List[Tuple[str, str]]) -> None:
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = True
    c0 = banner.cell(0, 0)
    _set_cell_shading(c0, "0050B3")
    _set_cell_margins(c0, top=50, bottom=50, left=100, right=100)
    _fill_cell_runs(
        c0,
        [("FICHA DEL INFORME", {"bold": True, "size": 10, "color": RGBColor(255, 255, 255)})],
    )
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_table_borders(banner, color="0050B3", sz="4")

    table = doc.add_table(rows=len(campos), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table, color="0050B3", sz="8")
    try:
        table.columns[0].width = Cm(4.8)
        table.columns[1].width = Cm(11.2)
    except Exception:
        pass

    for i, (label, value) in enumerate(campos):
        left, right = table.cell(i, 0), table.cell(i, 1)
        fill_l = "003366" if i % 2 == 0 else "1F4788"
        fill_r = "F2F6FC" if i % 2 == 0 else "FFFFFF"
        _set_cell_shading(left, fill_l)
        _set_cell_shading(right, fill_r)
        _set_cell_margins(left, top=45, bottom=45, left=90, right=70)
        _set_cell_margins(right, top=45, bottom=45, left=90, right=90)
        _fill_cell_runs(left, [(label, {"bold": True, "size": 9, "color": RGBColor(255, 255, 255)})])
        _fill_cell_runs(right, [(value, {"bold": False, "size": 10, "color": RGBColor(33, 37, 41)})])

    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(4)


def aplicar_margenes_informe(doc: Document) -> None:
    for sec in doc.sections:
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)


def _veredicto(
    pct_total: float,
    pct_noct: float,
    pct_vs_prev_noct: float,
) -> Tuple[str, str]:
    """Título + párrafo. El nocturno manda: es el indicador de fugas, no el total diurno."""
    if pct_noct >= 50:
        if pct_vs_prev_noct >= -50:
            extra = (
                "El caudal 00:00–06:59 sigue muy por debajo de la referencia previa "
                "y en el mismo orden de magnitud que el checkpoint de 27/07–02/08. "
                "Eso indica que las fugas de matriz no reaparecieron."
            )
            if pct_vs_prev_noct < 0:
                extra += (
                    " Hay una alza menor del nocturno frente a inicios de agosto; "
                    "no alcanza para hablar de recaída."
                )
            extra += (
                " Un total diario más alto se interpreta como mayor demanda del campus "
                "(ocupación, riego u otros usos diurnos), no como recaída de fugas."
            )
            return (
                "Las reparaciones continúan buenas: el nocturno se mantiene bajo",
                extra,
            )
        return (
            "El nocturno bajó fuerte vs la referencia, con alza frente al checkpoint",
            "Sigue habiendo evidencia de reparación, pero el nocturno de la última "
            "semana supera claramente el de 27/07–02/08. Conviene seguir la serie "
            "00:00–06:59 las próximas noches.",
        )
    if pct_noct >= 25:
        return (
            "Las fugas siguen controladas de forma moderada",
            "El nocturno continúa bajo la referencia, aunque el recorte es menor que "
            "en las primeras semanas post-reparación.",
        )
    if pct_noct >= 8:
        extra = ""
        if pct_total < 0:
            extra = (
                " El total semanal está por encima de la referencia: separar demanda "
                "diurna de caudal base nocturno antes de concluir una recaída."
            )
        return (
            "Hay una baja nocturna leve; conviene seguir monitoreando",
            "El indicador de fugas no volvió al nivel previo, pero el margen es "
            "estrecho." + extra,
        )
    return (
        "El nocturno ya no muestra el recorte de la reparación",
        "Revisar caudal mínimo 00:00–06:59 y ocupación del campus. Un alza de total "
        "diurno sola no prueba recaída de fugas; el nocturno sí es la señal a chequear.",
    )


def _convertir_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                str(docx_path.parent),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        return pdf_path if pdf_path.is_file() else None
    try:
        pdf_out = convertir_word_a_pdf(docx_path)
        if pdf_out and Path(pdf_out).is_file():
            return Path(pdf_out)
    except Exception as exc:
        print(f"[WARN] PDF (docx2pdf): {exc}")
    return None


def _export_csv(serie: List[Fila], path: Path) -> None:
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["fecha", "dia", "total_m3", "nocturno_00_06_m3", "diurno_m3", "horas_con_dato"])
        for d, tot, noct, horas in serie:
            w.writerow(
                [
                    d.isoformat(),
                    DIAS_ES[d.weekday()],
                    f"{tot:.3f}",
                    f"{noct:.3f}",
                    f"{tot - noct:.3f}",
                    _horas_con_dato(horas),
                ]
            )


def generar() -> Path:
    global POST_INI, POST_FIN
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    GRAFICOS.mkdir(parents=True, exist_ok=True)

    print("Determinando últimos 7 días completos...", flush=True)
    POST_INI, POST_FIN, incluye_hoy = _ventana_ultimos_7_completos()
    print(
        f"Ventana post: {POST_INI} – {POST_FIN} (incluye_hoy={incluye_hoy})",
        flush=True,
    )

    print("Cargando ventana SIN / referencia...", flush=True)
    serie_sin = _serie(SIN_INI, SIN_FIN)
    print("Cargando checkpoint 27/07–02/08...", flush=True)
    serie_prev = _serie(POST_PREV_INI, POST_PREV_FIN)
    print("Cargando ventana POST (últimos 7 días)...", flush=True)
    serie_post = _serie(POST_INI, POST_FIN)

    print("Cargando serie diaria continua (referencia → último día)...", flush=True)
    serie_larga = _serie(SIN_INI, POST_FIN)

    tot_sin, noct_sin, prom_sin = _totales(serie_sin)
    tot_prev, noct_prev, prom_prev = _totales(serie_prev)
    tot_post, noct_post, prom_post = _totales(serie_post)
    n_sin, n_post = len(serie_sin), len(serie_post)
    delta = tot_sin - tot_post
    pct = _pct(delta, tot_sin)
    delta_noct = noct_sin - noct_post
    pct_noct = _pct(delta_noct, noct_sin)
    delta_vs_prev = tot_prev - tot_post
    pct_vs_prev = _pct(delta_vs_prev, tot_prev)
    delta_vs_prev_noct = noct_prev - noct_post
    pct_vs_prev_noct = _pct(delta_vs_prev_noct, noct_prev)
    prom_noct_sin = noct_sin / n_sin if n_sin else 0.0
    prom_noct_prev = noct_prev / len(serie_prev) if serie_prev else 0.0
    prom_noct_post = noct_post / n_post if n_post else 0.0

    semanas_plot: List[Tuple[str, float, float]] = [
        (f"Ref.\n{SIN_INI:%d/%m}–{SIN_FIN:%d/%m}", tot_sin, noct_sin)
    ]
    for w_ini, w_fin in _semanas_lun_dom(DIA_REPARACION + timedelta(days=1), POST_FIN):
        filas = [r for r in serie_larga if w_ini <= r[0] <= w_fin]
        t, n, _ = _totales(filas)
        semanas_plot.append((f"{w_ini:%d/%m}–{w_fin:%d/%m}", t, n))

    img_comp = GRAFICOS / "comparativo_diario.png"
    img_tot = GRAFICOS / "totales.png"
    img_det = GRAFICOS / "detalle_post.png"
    img_serie = GRAFICOS / "serie_diaria_persistencia.png"
    img_sem = GRAFICOS / "evolucion_semanal.png"
    img_tres = GRAFICOS / "tres_ventanas.png"
    _plot_comparativo_diario(serie_sin, serie_post, img_comp)
    _plot_totales(tot_sin, tot_post, img_tot)
    _plot_detalle_post(serie_post, img_det)
    _plot_serie_diaria(serie_larga, prom_sin, prom_noct_sin, img_serie)
    _plot_semanas(semanas_plot, img_sem)
    _plot_tres_ventanas(tot_sin, tot_prev, tot_post, img_tres)
    img_noct3 = GRAFICOS / "tres_ventanas_nocturno.png"
    _plot_tres_nocturno(noct_sin, noct_prev, noct_post, img_noct3)

    dia_max = max(serie_post, key=lambda r: r[1])
    img_perfil = GRAFICOS / "perfil_dia_max_post.png"
    _plot_perfil_dia(
        dia_max[3],
        dia_max[0],
        img_perfil,
        "Detalle horario — día de mayor consumo (últimos 7 días)",
    )

    csv_path = OUT_DIR / "serie_diaria_honduras.csv"
    _export_csv(serie_larga, csv_path)

    titulo_ver, extra_ver = _veredicto(pct, pct_noct, pct_vs_prev_noct)

    doc = Document()
    add_logo_to_header(doc)
    aplicar_margenes_informe(doc)

    title = doc.add_paragraph("Informe comparativo UDD — post-reparación Honduras")
    title.style = "Title"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 80, 179)
        run.font.name = "Calibri"
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph("Actualización a la fecha — ¿continúa el ahorro?")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 71, 136)
        run.italic = True
    sub.paragraph_format.space_after = Pt(8)

    nota_hoy = (
        f"Hoy {date.today().strftime('%d/%m/%Y')} aún no completa 20 h de dato; "
        f"la ventana cierra el {POST_FIN.strftime('%d/%m/%Y')}."
        if not incluye_hoy
        else f"Incluye {date.today().strftime('%d/%m/%Y')} (día con dato horario completo)."
    )
    add_ficha_informe(
        doc,
        [
            ("Cliente", CLIENTE),
            ("Punto WES", f"{NODE_NAME} ({NODE_ID})"),
            ("Reparación", f"Fugas en matriz — {DIA_REPARACION.strftime('%d/%m/%Y')}"),
            (
                "Sin reparación",
                f"{SIN_INI.strftime('%d/%m/%Y')} – {SIN_FIN.strftime('%d/%m/%Y')}",
            ),
            (
                "Checkpoint ago",
                f"{POST_PREV_INI.strftime('%d/%m/%Y')} – {POST_PREV_FIN.strftime('%d/%m/%Y')}",
            ),
            (
                "Últimos 7 días",
                f"{POST_INI.strftime('%d/%m/%Y')} – {POST_FIN.strftime('%d/%m/%Y')}",
            ),
            ("Cobertura", nota_hoy),
            ("Veredicto", titulo_ver),
            ("Generado", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ],
    )

    doc.add_heading("1. Resumen comparativo (referencia vs últimos 7 días)", level=1)
    p = doc.add_paragraph(
        f"Se compara el consumo de {NODE_NAME} en la ventana de referencia "
        f"({SIN_INI.strftime('%d/%m')}–{SIN_FIN.strftime('%d/%m')}, {n_sin} días) con el "
        f"monitoreo más reciente posterior a la reparación "
        f"({POST_INI.strftime('%d/%m')}–{POST_FIN.strftime('%d/%m')}, {n_post} días). "
        f"La referencia es la misma del informe de inicios de agosto."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    rows = [
        [
            f"Sin / ref. ({SIN_INI:%d/%m}–{SIN_FIN:%d/%m})",
            format_number_chilean(tot_sin, 1),
            format_number_chilean(prom_sin, 1),
            format_number_chilean(noct_sin, 1),
            format_currency_chilean(tot_sin * PRECIO),
        ],
        [
            f"Últimos 7 d ({POST_INI:%d/%m}–{POST_FIN:%d/%m})",
            format_number_chilean(tot_post, 1),
            format_number_chilean(prom_post, 1),
            format_number_chilean(noct_post, 1),
            format_currency_chilean(tot_post * PRECIO),
        ],
    ]
    _add_table(
        doc,
        ["Ventana", "Total (m³)", "Prom. diario (m³)", "Nocturno (m³)", "Costo (CLP)"],
        rows,
    )

    var_txt = (
        f"Variación del total: {format_number_chilean(delta, 1)} m³ "
        f"({'baja' if delta >= 0 else 'alza'} de {format_number_chilean(abs(pct), 1)} % "
        f"respecto de la ventana sin reparación / referencia)."
    )
    noct_txt = (
        f"Nocturno 00:00–06:59: {format_number_chilean(noct_sin, 1)} m³ → "
        f"{format_number_chilean(noct_post, 1)} m³ "
        f"({'baja' if delta_noct >= 0 else 'alza'} de "
        f"{format_number_chilean(abs(pct_noct), 1)} %)."
    )
    doc.add_paragraph(var_txt).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph(noct_txt).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lec = doc.add_paragraph()
    lec.add_run("Lectura: ").bold = True
    lec.add_run(
        "El nocturno (00:00–06:59) es el indicador de fugas. El total diario mezcla "
        "ocupación del campus, riego y otros usos. Un total más alto con nocturno "
        "aún bajo no significa que las reparaciones hayan fallado."
    )
    lec.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_picture(str(img_tot), width=Inches(5.5))
    doc.add_picture(str(img_comp), width=Inches(6.2))

    doc.add_heading("2. Persistencia: ¿sigue igual de bien que a inicios de agosto?", level=1)
    p2 = doc.add_paragraph(
        "El informe original de esta carpeta cerró el monitoreo en 27/07–02/08 "
        f"(total {format_number_chilean(tot_prev, 1)} m³; nocturno "
        f"{format_number_chilean(noct_prev, 1)} m³). Esa semana tuvo menos demanda "
        "diurna que la actual. El contraste correcto para fugas es el nocturno, "
        "no el total."
    )
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _add_table(
        doc,
        ["Ventana", "Total (m³)", "Prom. diario (m³)", "Nocturno (m³)", "Prom. noct. (m³/día)"],
        [
            [
                f"Referencia ({SIN_INI:%d/%m}–{SIN_FIN:%d/%m})",
                format_number_chilean(tot_sin, 1),
                format_number_chilean(prom_sin, 1),
                format_number_chilean(noct_sin, 1),
                format_number_chilean(prom_noct_sin, 1),
            ],
            [
                f"Checkpoint ({POST_PREV_INI:%d/%m}–{POST_PREV_FIN:%d/%m})",
                format_number_chilean(tot_prev, 1),
                format_number_chilean(prom_prev, 1),
                format_number_chilean(noct_prev, 1),
                format_number_chilean(prom_noct_prev, 1),
            ],
            [
                f"Últimos 7 d ({POST_INI:%d/%m}–{POST_FIN:%d/%m})",
                format_number_chilean(tot_post, 1),
                format_number_chilean(prom_post, 1),
                format_number_chilean(noct_post, 1),
                format_number_chilean(prom_noct_post, 1),
            ],
        ],
    )
    vs_prev = (
        f"Total últimos 7 días vs checkpoint: {format_number_chilean(delta_vs_prev, 1)} m³ "
        f"({'más bajo' if delta_vs_prev >= 0 else 'más alto'} en "
        f"{format_number_chilean(abs(pct_vs_prev), 1)} %). "
        f"Nocturno vs checkpoint: {format_number_chilean(delta_vs_prev_noct, 1)} m³ "
        f"({'más bajo' if delta_vs_prev_noct >= 0 else 'más alto'} en "
        f"{format_number_chilean(abs(pct_vs_prev_noct), 1)} %)."
    )
    doc.add_paragraph(vs_prev).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_picture(str(img_noct3), width=Inches(5.8))
    doc.add_picture(str(img_tres), width=Inches(5.8))

    doc.add_heading("3. Evolución semanal post-reparación", level=1)
    p3 = doc.add_paragraph(
        "Semanas lunes–domingo desde el día siguiente a la reparación. "
        "La barra roja es la semana de referencia previa. Un nocturno que se "
        "mantenga bajo es la mejor evidencia de que las fugas no reaparecieron."
    )
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_picture(str(img_sem), width=Inches(6.2))
    sem_rows = []
    for lab, t, n in semanas_plot:
        lab_flat = lab.replace("\n", " ")
        sem_rows.append(
            [
                lab_flat,
                format_number_chilean(t, 1),
                format_number_chilean(t / 7.0, 1),
                format_number_chilean(n, 1),
            ]
        )
    _add_table(
        doc,
        ["Semana", "Total (m³)", "Prom. diario (m³)", "Nocturno (m³)"],
        sem_rows,
    )

    doc.add_heading("4. Serie diaria hasta la fecha", level=1)
    p4 = doc.add_paragraph(
        f"Consumo diario desde {SIN_INI.strftime('%d/%m/%Y')} hasta {POST_FIN.strftime('%d/%m/%Y')}. "
        "Línea roja punteada = promedio diario de la referencia; línea amarilla = nocturno. "
        "Si el total oscila por uso del campus pero el nocturno se queda bajo, la reparación sigue vigente."
    )
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_picture(str(img_serie), width=Inches(6.3))

    doc.add_heading(
        f"5. Detalle últimos 7 días ({POST_INI:%d/%m}–{POST_FIN:%d/%m})",
        level=1,
    )
    p5 = doc.add_paragraph(
        "Detalle diario del consumo total y nocturno (00:00–06:59) en la semana más reciente."
    )
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_picture(str(img_det), width=Inches(6.2))

    det_rows = []
    for d, tot, noct, _ in serie_post:
        det_rows.append(
            [
                d.strftime("%d/%m/%Y"),
                DIAS_ES[d.weekday()],
                format_number_chilean(tot, 1),
                format_number_chilean(noct, 1),
                format_number_chilean(tot - noct, 1),
            ]
        )
    det_rows.append(
        [
            "",
            "Total",
            format_number_chilean(tot_post, 1),
            format_number_chilean(noct_post, 1),
            format_number_chilean(tot_post - noct_post, 1),
        ]
    )
    _add_table(
        doc,
        ["Fecha", "Día", "Total (m³)", "Nocturno (m³)", "Diurno (m³)"],
        det_rows,
        total=True,
    )

    doc.add_heading("5.1 Perfil horario — día de mayor consumo (últimos 7 días)", level=2)
    doc.add_paragraph(
        f"Día seleccionado: {DIAS_ES[dia_max[0].weekday()]} {dia_max[0].strftime('%d/%m/%Y')} "
        f"({format_number_chilean(dia_max[1], 1)} m³ totales; "
        f"nocturno {format_number_chilean(dia_max[2], 1)} m³)."
    )
    doc.add_picture(str(img_perfil), width=Inches(6.0))

    doc.add_heading("6. Conclusión", level=1)
    vtitle = doc.add_paragraph()
    run_v = vtitle.add_run(titulo_ver)
    run_v.bold = True
    run_v.font.size = Pt(12)
    run_v.font.color.rgb = RGBColor(0, 80, 179)

    concl = doc.add_paragraph(
        f"En el monitoreo del {POST_INI.strftime('%d/%m')} al {POST_FIN.strftime('%d/%m')}, "
        f"{NODE_NAME} registró {format_number_chilean(tot_post, 1)} m³ "
        f"(promedio {format_number_chilean(prom_post, 1)} m³/día), frente a "
        f"{format_number_chilean(tot_sin, 1)} m³ en la ventana de referencia "
        f"{SIN_INI.strftime('%d/%m')}–{SIN_FIN.strftime('%d/%m')} "
        f"(promedio {format_number_chilean(prom_sin, 1)} m³/día). "
        f"{var_txt} {noct_txt} "
        f"Respecto del checkpoint 27/07–02/08 ({format_number_chilean(tot_prev, 1)} m³; "
        f"nocturno {format_number_chilean(noct_prev, 1)} m³): {vs_prev} {extra_ver}"
    )
    concl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for t in [
        "La comparación de totales semanales se ve afectada por ocupación del campus, riego u otras demandas diurnas.",
        "El caudal entre 00:00 y 06:59 es el indicador más estable de fugas residuales.",
        "Entre el 13/07 y el 23/07 hay varios días sin dato horario (post-corte); no se usan como evidencia de fugas.",
        "Valores alineados con Sala impulsión Honduras, nodo 000026-01.",
        f"Serie diaria exportada en {csv_path.name}.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    pie = doc.add_paragraph()
    pie.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = pie.add_run("WES — Water Efficiency System")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    out = OUT_DIR / "Informe_UDD_Comparacion_Post_Reparacion_Ago2026.docx"
    doc.save(out)
    print(f"[OK] DOCX: {out}")
    print(f"[OK] CSV:  {csv_path}")

    pdf_path = _convertir_pdf(out)
    if pdf_path:
        print(f"[OK] PDF:  {pdf_path}")
    else:
        print("[WARN] No se pudo convertir a PDF")

    print(f"SIN:  {tot_sin:.1f} m3 (noct {noct_sin:.1f})")
    print(f"PREV: {tot_prev:.1f} m3 (noct {noct_prev:.1f})")
    print(f"POST: {tot_post:.1f} m3 (noct {noct_post:.1f})")
    print(f"Delta vs ref: {delta:.1f} m3 ({pct:.1f}%) | noct {delta_noct:.1f} m3 ({pct_noct:.1f}%)")
    print(f"Delta vs prev: {delta_vs_prev:.1f} m3 ({pct_vs_prev:.1f}%) | noct {delta_vs_prev_noct:.1f} m3 ({pct_vs_prev_noct:.1f}%)")
    print(f"Veredicto: {titulo_ver}")
    return out


def main() -> int:
    print("=" * 70)
    print("UDD — comparación post-reparación actualizada hasta la fecha")
    print(f"Referencia fija: {SIN_INI:%d/%m/%Y} – {SIN_FIN:%d/%m/%Y}")
    print("=" * 70)
    generar()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
