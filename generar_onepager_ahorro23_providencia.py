"""
One-pager Providencia: consumo oct→fecha + 23% ahorro supuesto.

Uso:
  python generar_onepager_ahorro23_providencia.py
"""

from __future__ import annotations

import subprocess
import sys
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

if sys.platform == "win32":
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Providencia" / "onepager_ahorro23"
COMPANY_ID = "000006"
DESDE = "01/10/2025"
AHORRO_PCT = 0.23

# Control WES detectado (informe activación)
CONTROL = {
    "000006-01": "Desde 01/10/2025 (hasta ~abr 2026)",
    "000006-02": "26/01–01/03/2026",
    "000006-04": "Solo 29/06–05/07/2026",
    "000006-05": "No detectado sostenido",
}

NODOS = [
    "000006-01",
    "000006-02",
    "000006-04",
    "000006-05",
]


def _fmt(v: float, d: int = 1) -> str:
    from generar_reporte_word import format_number_chilean

    return format_number_chilean(v, d)


def _fmt_clp(v: float) -> str:
    from generar_reporte_word import format_currency_chilean

    return format_currency_chilean(v)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc.append(shd)


def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_narrow_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def _grafico(filas, tot_ahorro, out: Path) -> None:
    labels = [f["corto"] for f in filas] + ["TOTAL"]
    cons = [f["tot"] for f in filas] + [sum(f["tot"] for f in filas)]
    aho = [f["ahorro_m3"] for f in filas] + [tot_ahorro]
    x = range(len(labels))
    w = 0.38
    fig, ax = plt.subplots(figsize=(9.2, 3.4))
    ax.bar([i - w / 2 for i in x], cons, width=w, label="Consumo medido", color="#1f4788")
    ax.bar([i + w / 2 for i in x], aho, width=w, label="Ahorro 23 %", color="#27ae60")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_ylabel("m³")
    ax.set_title("Consumo oct→fecha vs ahorro supuesto 23 %")
    ax.legend(fontsize=8, loc="upper right")
    ax.grid(True, axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out, dpi=140)
    plt.close(fig)


def generar() -> Path:
    from generar_reporte_word import (
        calculate_nocturnal_metrics,
        get_node_name,
        get_water_price_per_m3,
        parse_date,
    )
    from listado_pa_que_esta_instalado import _medidas_rango

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    hasta = datetime.now().strftime("%d/%m/%Y")
    start_dt = parse_date(DESDE)
    end_dt = parse_date(hasta, end_of_day=True)
    s = datetime(2025, 10, 1)
    e = end_dt.replace(tzinfo=None) if getattr(end_dt, "tzinfo", None) else end_dt
    precio = get_water_price_per_m3(COMPANY_ID, NODOS[0], None) or 1274.0

    filas = []
    print(f"[INFO] One-pager 23% | {DESDE} → {hasta} | {_fmt_clp(precio)}/m³")
    for nid in NODOS:
        nombre = get_node_name(nid) or nid
        pts = _medidas_rango(nid, s, e)
        tot = sum(float(m.total_m3 or 0) for m in pts)
        dias = sum(1 for m in pts if float(m.total_m3 or 0) > 1e-9)
        prom = tot / dias if dias else 0.0
        noct = calculate_nocturnal_metrics(nid, start_dt, end_dt)
        cn = float(noct["consumo_nocturno_total"])
        dd = int(noct["dias_con_datos_horarios"])
        prom_n = cn / dd if dd else 0.0
        ahorro = tot * AHORRO_PCT
        filas.append(
            {
                "node_id": nid,
                "nombre": nombre,
                "corto": nombre.replace("Liceo ", "")
                .replace("Luisa Saavedra", "7")
                .replace("Juan Pablo Duarte", "Duarte")[:18],
                "tot": tot,
                "dias": dias,
                "prom": prom,
                "noct": cn,
                "prom_n": prom_n,
                "ahorro_m3": ahorro,
                "ahorro_clp": ahorro * precio,
                "control": CONTROL.get(nid, "—"),
            }
        )
        print(f"  {nid}: tot={tot:.1f} ahorro23={ahorro:.1f}")

    tot_cons = sum(f["tot"] for f in filas)
    tot_ahorro = tot_cons * AHORRO_PCT
    tot_clp = tot_ahorro * precio
    tot_noct = sum(f["noct"] for f in filas)

    png = OUT_DIR / f"onepager_barras_{ts}.png"
    _grafico(filas, tot_ahorro, png)

    doc = Document()
    _set_narrow_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    h = doc.add_heading("Colegios Providencia — One page ahorro 23 %", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].font.size = Pt(16)
        h.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(
        f"Periodo: {DESDE} al {hasta}  |  4 liceos operativos  |  "
        f"Supuesto de ahorro: {AHORRO_PCT*100:.0f} % sobre consumo medido  |  "
        f"Tarifa ref. {_fmt_clp(precio)}/m³"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(80, 80, 80)

    # KPI boxes as compact paragraph
    kpi = doc.add_paragraph()
    kpi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kpi.paragraph_format.space_before = Pt(2)
    kpi.paragraph_format.space_after = Pt(6)
    k = kpi.add_run(
        f"Consumo total: {_fmt(tot_cons, 0)} m³   ·   "
        f"Ahorro 23 %: {_fmt(tot_ahorro, 0)} m³   ·   "
        f"Ahorro $: {_fmt_clp(tot_clp)}   ·   "
        f"Nocturno medido: {_fmt(tot_noct, 0)} m³"
    )
    k.bold = True
    k.font.size = Pt(10)
    k.font.color.rgb = RGBColor(39, 174, 96)

    # Main table
    headers = [
        "Liceo",
        "Total m³",
        "Prom. diario",
        "Prom. noct. 00–06",
        "Ahorro 23 % m³",
        "Ahorro 23 % $",
        "Control WES",
    ]
    tbl = doc.add_table(rows=1 + len(filas) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hd in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = hd
        _shade(cell, "1F4788")
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)

    for i, f in enumerate(filas, 1):
        vals = [
            f["nombre"],
            _fmt(f["tot"], 0),
            f"{_fmt(f['prom'], 1)} m³/d",
            f"{_fmt(f['prom_n'], 2)} m³/d",
            _fmt(f["ahorro_m3"], 0),
            _fmt_clp(f["ahorro_clp"]),
            f["control"],
        ]
        for j, v in enumerate(vals):
            cell = tbl.rows[i].cells[j]
            cell.text = v
            if i % 2 == 0:
                _shade(cell, "F2F6FA")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(8)

    tot_vals = [
        "TOTAL",
        _fmt(tot_cons, 0),
        "—",
        "—",
        _fmt(tot_ahorro, 0),
        _fmt_clp(tot_clp),
        "—",
    ]
    for j, v in enumerate(tot_vals):
        cell = tbl.rows[len(filas) + 1].cells[j]
        cell.text = v
        _shade(cell, "D5F5E3")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    if png.is_file():
        pimg = doc.add_paragraph()
        pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pimg.paragraph_format.space_before = Pt(6)
        pimg.paragraph_format.space_after = Pt(2)
        run = pimg.add_run()
        run.add_picture(str(png), width=Inches(6.3))

    # Notas compactas
    n1 = doc.add_paragraph()
    n1.paragraph_format.space_before = Pt(2)
    n1.paragraph_format.space_after = Pt(1)
    t = n1.add_run("Supuesto: ")
    t.bold = True
    t.font.size = Pt(8)
    t2 = n1.add_run(
        f"el {AHORRO_PCT*100:.0f} % se aplica sobre el consumo total medido WES "
        f"desde octubre a la fecha (no es baseline pre-WES). "
        f"Equivalente a {_fmt(tot_ahorro, 0)} m³ / {_fmt_clp(tot_clp)} en el periodo."
    )
    t2.font.size = Pt(8)

    n2 = doc.add_paragraph()
    n2.paragraph_format.space_before = Pt(0)
    n2.paragraph_format.space_after = Pt(1)
    u = n2.add_run("Control WES: ")
    u.bold = True
    u.font.size = Pt(8)
    u2 = n2.add_run(
        "detectado cuando noche 00–06 → ~0 con diurno hábil activo. "
        "Lastarria: control fuerte oct–abr; desde mayo el nocturno vuelve. "
        "Carmela: tramo ene–mar 2026. Liceo 7/Duarte: sin control sostenido."
    )
    u2.font.size = Pt(8)

    n3 = doc.add_paragraph()
    n3.paragraph_format.space_before = Pt(0)
    v = n3.add_run("Notas: ")
    v.bold = True
    v.font.size = Pt(7)
    v2 = n3.add_run(
        "Prom. diario = total ÷ días con consumo > 0. "
        "Prom. nocturno = nocturno ÷ días con datos horarios. "
        "Liceo 7 con jul–ago anómalo (revisar). Alessandri excluido (sin data reciente). "
        f"Generado {datetime.now():%d-%m-%Y %H:%M}."
    )
    v2.font.size = Pt(7)
    v2.font.color.rgb = RGBColor(100, 100, 100)

    out_docx = OUT_DIR / f"OnePager_Ahorro23_Providencia_{ts}.docx"
    doc.save(out_docx)

    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUT_DIR),
                str(out_docx),
            ],
            check=False,
            capture_output=True,
            timeout=120,
        )
    except Exception as ex:
        print(f"[WARN] PDF: {ex}")

    pdf = out_docx.with_suffix(".pdf")
    print(f"[OK] Word: {out_docx}")
    print(f"[OK] PDF:  {pdf if pdf.is_file() else 'no'}")
    print(f"[OK] TOTAL ahorro 23%: {_fmt(tot_ahorro, 1)} m³ = {_fmt_clp(tot_clp)}")
    return out_docx


def main() -> int:
    print("=" * 64)
    print("ONE-PAGER AHORRO 23% — PROVIDENCIA")
    print("=" * 64)
    generar()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
