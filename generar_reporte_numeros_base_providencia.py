"""
Reporte operativo — números base colegios Providencia.

Tablas + gráficos:
  1) Consumo total por liceo (octubre → fecha)
  2) Consumo promedio diario
  3) Consumo promedio nocturno 00:00–06:59

Uso:
  python generar_reporte_numeros_base_providencia.py
"""

from __future__ import annotations

import csv
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

ROOT = Path(__file__).resolve().parent
COMPANY_ID = "000006"
DESDE = "01/10/2025"

NODOS = [
    ("000006-01", "Liceo Lastarria", True),
    ("000006-02", "Carmela Carvajal", True),
    ("000006-04", "Liceo 7 Luisa Saavedra", True),
    ("000006-05", "Liceo Juan Pablo Duarte", True),
    ("000006-03", "Arturo Alessandri Palma", False),
]


def _fmt(v: float, d: int = 1) -> str:
    from generar_reporte_word import format_number_chilean

    return format_number_chilean(v, d)


def _fmt_clp(v: float) -> str:
    from generar_reporte_word import format_currency_chilean

    return format_currency_chilean(v)


def _corto(nombre: str) -> str:
    return (
        nombre.replace("Liceo ", "")
        .replace("Arturo Alessandri Palma", "Alessandri")
        .replace("Luisa Saavedra", "7")
        .replace("Juan Pablo Duarte", "Duarte")
    )


def _naive(dt: datetime) -> datetime:
    if getattr(dt, "tzinfo", None) is not None:
        return dt.replace(tzinfo=None)
    return dt


def _barra(
    labels: List[str],
    vals: List[float],
    out: Path,
    *,
    ylabel: str,
    titulo: str,
    color: str = "#1f4788",
    decimals: int = 1,
) -> None:
    fig, ax = plt.subplots(figsize=(9.2, 4.8))
    bars = ax.bar(labels, vals, color=color, edgecolor="white", width=0.65)
    ymax = max(vals) if vals else 1
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + ymax * 0.02,
            _fmt(v, decimals),
            ha="center",
            va="bottom",
            fontsize=9,
        )
    ax.set_ylabel(ylabel)
    ax.set_title(titulo)
    ax.grid(True, axis="y", alpha=0.3)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: _fmt(x, 0)))
    fig.tight_layout()
    fig.savefig(out, dpi=150)
    plt.close(fig)


def _comparativo_promedios(filas: List[dict], out: Path) -> None:
    labels = [f["corto"] for f in filas]
    x = range(len(labels))
    w = 0.38
    dia = [f["prom_dia"] for f in filas]
    noc = [f["prom_noct"] for f in filas]
    fig, ax = plt.subplots(figsize=(9.5, 5.0))
    b1 = ax.bar([i - w / 2 for i in x], dia, width=w, label="Prom. diario total", color="#1f4788")
    b2 = ax.bar([i + w / 2 for i in x], noc, width=w, label="Prom. nocturno 00–06 h", color="#c0392b")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_ylabel("m³/día")
    ax.set_title("Promedio diario total vs promedio nocturno (00:00–06:59)")
    ax.legend()
    ax.grid(True, axis="y", alpha=0.3)
    for bars in (b1, b2):
        for bar in bars:
            h = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                h + 0.15,
                _fmt(h, 2),
                ha="center",
                va="bottom",
                fontsize=8,
            )
    fig.tight_layout()
    fig.savefig(out, dpi=150)
    plt.close(fig)


def generar(output_dir: Optional[Path] = None) -> Path:
    from generar_reporte_word import (
        calculate_nocturnal_metrics,
        get_water_price_per_m3,
        parse_date,
    )
    from listado_pa_que_esta_instalado import _medidas_rango

    out_dir = output_dir or (
        ROOT / "reports" / "Providencia" / "numeros_base_octubre_fecha"
    )
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")

    start_dt = parse_date(DESDE)
    hasta_txt = datetime.now().strftime("%d/%m/%Y")
    end_dt = parse_date(hasta_txt, end_of_day=True)
    precio = get_water_price_per_m3(COMPANY_ID, NODOS[0][0], None) or 1274.0

    s_naive = datetime(_naive(start_dt).year, _naive(start_dt).month, _naive(start_dt).day)
    e_naive = _naive(end_dt)

    filas: List[dict] = []
    print(f"[INFO] Periodo {DESDE} → {hasta_txt} | precio {_fmt_clp(precio)}/m³")

    for nid, nombre, activo_flag in NODOS:
        print(f"  {nid} {nombre}...")
        pts = _medidas_rango(nid, s_naive, e_naive)
        tot = sum(float(m.total_m3 or 0) for m in pts)
        dias_cons = sum(1 for m in pts if float(m.total_m3 or 0) > 1e-9)
        prom_dia = (tot / dias_cons) if dias_cons else 0.0

        noct = calculate_nocturnal_metrics(nid, start_dt, end_dt)
        noct_tot = float(noct["consumo_nocturno_total"])
        dias_h = int(noct["dias_con_datos_horarios"])
        dias_con_n = int(noct["dias_con_consumo_nocturno"])
        prom_noct = (noct_tot / dias_h) if dias_h else 0.0
        pct_noct = (100.0 * noct_tot / tot) if tot > 0 else 0.0

        nota = ""
        if not activo_flag:
            nota = "Sin data reciente (~ene 2026)."
        elif nid == "000006-04" and tot > 8000:
            nota = "Revisar jul–ago 2026 (consumo anómalo)."

        filas.append(
            {
                "node_id": nid,
                "nombre": nombre,
                "corto": _corto(nombre),
                "activo": activo_flag and dias_cons > 30,
                "tot_m3": tot,
                "tot_clp": tot * precio,
                "dias_cons": dias_cons,
                "prom_dia": prom_dia,
                "noct_m3": noct_tot,
                "noct_clp": noct_tot * precio,
                "dias_h": dias_h,
                "dias_con_n": dias_con_n,
                "prom_noct": prom_noct,
                "pct_noct": pct_noct,
                "nota": nota,
            }
        )
        print(
            f"    tot={tot:.1f} prom_dia={prom_dia:.2f} "
            f"noct={noct_tot:.1f} prom_noct={prom_noct:.2f}"
        )

    # Priorizar activos en gráficos; Alessandri al final si tiene data
    graf = [f for f in filas if f["activo"]] + [f for f in filas if not f["activo"] and f["tot_m3"] > 0]

    png_tot = out_dir / f"total_consumo_{ts}.png"
    _barra(
        [f["corto"] for f in graf],
        [f["tot_m3"] for f in graf],
        png_tot,
        ylabel="m³",
        titulo=f"Consumo total por liceo — {DESDE} al {hasta_txt}",
        color="#1f4788",
        decimals=0,
    )

    png_dia = out_dir / f"promedio_diario_{ts}.png"
    _barra(
        [f["corto"] for f in graf],
        [f["prom_dia"] for f in graf],
        png_dia,
        ylabel="m³/día",
        titulo="Consumo promedio diario (días con consumo > 0)",
        color="#16a085",
        decimals=2,
    )

    png_noct = out_dir / f"promedio_nocturno_{ts}.png"
    _barra(
        [f["corto"] for f in graf],
        [f["prom_noct"] for f in graf],
        png_noct,
        ylabel="m³/día (00–06 h)",
        titulo="Consumo promedio nocturno 00:00–06:59",
        color="#c0392b",
        decimals=2,
    )

    png_cmp = out_dir / f"comparativo_promedios_{ts}.png"
    _comparativo_promedios(graf, png_cmp)

    csv_path = out_dir / f"numeros_base_providencia_{ts}.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(
            [
                "node_id",
                "liceo",
                "desde",
                "hasta",
                "consumo_total_m3",
                "dias_con_consumo",
                "promedio_diario_m3",
                "consumo_nocturno_total_m3",
                "dias_con_datos_horarios",
                "promedio_nocturno_00_06_m3_dia",
                "pct_nocturno_sobre_total",
                "nota",
            ]
        )
        for f in filas:
            w.writerow(
                [
                    f["node_id"],
                    f["nombre"],
                    DESDE,
                    hasta_txt,
                    f"{f['tot_m3']:.2f}".replace(".", ","),
                    f["dias_cons"],
                    f"{f['prom_dia']:.3f}".replace(".", ","),
                    f"{f['noct_m3']:.2f}".replace(".", ","),
                    f["dias_h"],
                    f"{f['prom_noct']:.3f}".replace(".", ","),
                    f"{f['pct_noct']:.1f}".replace(".", ","),
                    f["nota"],
                ]
            )

    # Word
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading(
        "Números base — Colegios Providencia (octubre → fecha)", 0
    )
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    p = doc.add_paragraph()
    p.add_run("Periodo: ").bold = True
    p.add_run(f"{DESDE} al {hasta_txt}")
    p2 = doc.add_paragraph()
    p2.add_run("Qué incluye: ").bold = True
    p2.add_run(
        "consumo total por liceo, promedio diario y promedio nocturno "
        "00:00–06:59 (hora Chile). Tarifa ref. "
        f"{_fmt_clp(precio)}/m³. Generado {datetime.now():%d-%m-%Y %H:%M}."
    )

    # Tabla principal
    doc.add_heading("1. Tabla resumen", level=1)
    cols = [
        "Liceo",
        "Total (m³)",
        "Días c/ cons.",
        "Prom. diario (m³)",
        "Nocturno total (m³)",
        "Prom. noct. 00–06 (m³/día)",
        "% nocturno",
    ]
    tbl = doc.add_table(rows=1 + len(filas) + 1, cols=len(cols))
    tbl.style = "Table Grid"
    for j, hd in enumerate(cols):
        tbl.rows[0].cells[j].text = hd
        for run in tbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    tot_t = tot_n = 0.0
    for i, f in enumerate(filas, 1):
        tot_t += f["tot_m3"]
        tot_n += f["noct_m3"]
        row = tbl.rows[i].cells
        row[0].text = f["nombre"]
        row[1].text = _fmt(f["tot_m3"], 1)
        row[2].text = str(f["dias_cons"])
        row[3].text = _fmt(f["prom_dia"], 2)
        row[4].text = _fmt(f["noct_m3"], 1)
        row[5].text = _fmt(f["prom_noct"], 2)
        row[6].text = f"{_fmt(f['pct_noct'], 1)} %"

    activos = [f for f in filas if f["activo"]]
    tot_act = sum(f["tot_m3"] for f in activos)
    noct_act = sum(f["noct_m3"] for f in activos)
    tr = tbl.rows[len(filas) + 1].cells
    tr[0].text = "TOTAL 4 activos"
    tr[1].text = _fmt(tot_act, 1)
    tr[2].text = "—"
    tr[3].text = "—"
    tr[4].text = _fmt(noct_act, 1)
    tr[5].text = "—"
    tr[6].text = f"{_fmt(100.0 * noct_act / tot_act, 1)} %" if tot_act else "—"
    for c in tr:
        for run in c.paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph(
        "Promedio diario = consumo total ÷ días con consumo > 0. "
        "Promedio nocturno = consumo 00:00–06:59 ÷ días con datos horarios."
    )
    for f in filas:
        if f["nota"]:
            n = doc.add_paragraph()
            n.add_run(f"{f['nombre']}: ").bold = True
            n.add_run(f["nota"])

    doc.add_heading("2. Consumo total por liceo", level=1)
    doc.add_paragraph(
        f"Volumen acumulado desde octubre 2025 hasta {hasta_txt}."
    )
    if png_tot.is_file():
        doc.add_picture(str(png_tot), width=Inches(6.0))

    doc.add_heading("3. Consumo promedio diario", level=1)
    doc.add_paragraph(
        "Ritmo medio de consumo en los días con medición > 0."
    )
    if png_dia.is_file():
        doc.add_picture(str(png_dia), width=Inches(6.0))

    doc.add_heading("4. Consumo promedio nocturno 00:00–06:00", level=1)
    doc.add_paragraph(
        "Promedio diario del volumen entre 00:00 y 06:59. Es la base para "
        "estimar el ahorro si la gestión hídrica CPA lleva ese tramo a cero."
    )
    if png_noct.is_file():
        doc.add_picture(str(png_noct), width=Inches(6.0))

    doc.add_heading("5. Comparativo promedios", level=1)
    if png_cmp.is_file():
        doc.add_picture(str(png_cmp), width=Inches(6.0))

    # Ficha corta por liceo
    doc.add_heading("6. Ficha por liceo", level=1)
    for f in filas:
        doc.add_heading(f"{f['nombre']} ({f['node_id']})", level=2)
        if not f["activo"] and f["tot_m3"] <= 0:
            doc.add_paragraph("Sin datos en el periodo.")
            continue
        doc.add_paragraph(
            f"Total desde octubre: {_fmt(f['tot_m3'], 1)} m³ "
            f"({_fmt_clp(f['tot_clp'])})."
        )
        doc.add_paragraph(
            f"Promedio diario: {_fmt(f['prom_dia'], 2)} m³/día "
            f"({f['dias_cons']} días con consumo)."
        )
        doc.add_paragraph(
            f"Promedio nocturno 00–06 h: {_fmt(f['prom_noct'], 2)} m³/día "
            f"(total nocturno {_fmt(f['noct_m3'], 1)} m³ en {f['dias_h']} días "
            f"con datos; {_fmt(f['pct_noct'], 1)} % del total)."
        )
        if f["nota"]:
            doc.add_paragraph(f"Nota: {f['nota']}")

    out_docx = out_dir / f"Numeros_base_Providencia_{ts}.docx"
    doc.save(out_docx)

    pdf_path = out_docx.with_suffix(".pdf")
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(out_docx),
            ],
            check=False,
            capture_output=True,
            timeout=120,
        )
    except Exception as ex:
        print(f"[WARN] PDF: {ex}")

    print(f"[OK] Word: {out_docx}")
    print(f"[OK] PDF:  {pdf_path if pdf_path.is_file() else 'no'}")
    print(f"[OK] CSV:  {csv_path}")
    return out_docx


def main() -> int:
    print("=" * 72)
    print("NÚMEROS BASE — PROVIDENCIA (tablas + gráficos)")
    print("=" * 72)
    generar()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
