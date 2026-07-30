"""
Genera auditorías CPA (Con control vs Sin control) para todos los puntos de Renca,
usando las mismas fechas del informe ICCO Renca:

- Con control: 23-03-2026 a 26-03-2026
- Sin control: 06-04-2026 a 09-04-2026

Salida: un DOCX + PDF por nodo en:
  reports/reporte de auditoria/auditoria_cpa_renca_todos_puntos_20260323_20260409/

Uso:
  python generar_auditorias_cpa_renca_todos_puntos.py
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from auditoria_cpa_icco_renca_grafico import PERIODO_AUDITORIA, PERIODO_REFERENCIA, computar_auditoria_cpa
from generar_reporte_word import ENTITY_BASE_URL, convertir_word_a_pdf, fetch_json, format_number_chilean, get_node_name


ROOT = Path(__file__).resolve().parent
OUT_DIR = (
    ROOT
    / "reports"
    / "reporte de auditoria"
    / "auditoria_cpa_renca_todos_puntos_20260323_20260409"
)

COMPANIES_RENCA = ("000016", "000017")

_HEADING_COLOR = RGBColor(31, 71, 136)


def _safe_filename(s: str) -> str:
    out = "".join(ch for ch in s if ch.isalnum() or ch in (" ", "-", "_")).strip()
    out = out.replace(" ", "_")
    return out or "reporte"


def _company_nodes(company_id: str) -> list[dict]:
    # /companies/{id} devuelve nodes en este backend; fallback a /companies/{id}/nodes
    data = fetch_json(f"{ENTITY_BASE_URL}/companies/{company_id}")
    nodes = []
    if isinstance(data, dict) and isinstance(data.get("nodes"), list):
        nodes = data["nodes"]
    if not nodes:
        data2 = fetch_json(f"{ENTITY_BASE_URL}/companies/{company_id}/nodes")
        if isinstance(data2, list):
            nodes = data2
        elif isinstance(data2, dict) and isinstance(data2.get("nodes"), list):
            nodes = data2["nodes"]
    # normalizar
    out = []
    for n in nodes or []:
        node_id = (n.get("nodeId") or n.get("node_id") or "").strip()
        name = (n.get("name") or n.get("nodeName") or "").strip()
        if node_id:
            out.append({"nodeId": node_id, "name": name})
    return out


def _add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.bold = True
        r.font.name = "Calibri"
        r.font.color.rgb = _HEADING_COLOR
        r.font.size = Pt(14) if level == 1 else Pt(13)


def _p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 0, 0)


def _build_docx(node_id: str, node_name: str) -> Path:
    res = computar_auditoria_cpa(node_id, ref=PERIODO_REFERENCIA, aud=PERIODO_AUDITORIA)

    con_m3 = float(res.total_ref_m3)
    sin_m3 = float(res.total_aud_m3)
    ahorro_m3 = max(0.0, sin_m3 - con_m3)
    pct_ahorro = (100.0 * ahorro_m3 / sin_m3) if sin_m3 > 0 else 0.0

    doc = Document()
    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_heading(doc, "Informe de Auditoría CPA", level=1)
    _add_heading(doc, f"Renca — {node_name} ({node_id})", level=2)
    doc.add_paragraph("")

    _p(
        doc,
        "Auditoría comparativa de consumos en jornada completa (hora Chile), con el objetivo de "
        "contrastar un periodo con control activo frente a un periodo sin control (línea base).",
    )
    _p(
        doc,
        f"Periodo con control: {PERIODO_REFERENCIA.dias[0]:%d-%m-%Y} al {PERIODO_REFERENCIA.dias[-1]:%d-%m-%Y}. "
        f"Periodo sin control: {PERIODO_AUDITORIA.dias[0]:%d-%m-%Y} al {PERIODO_AUDITORIA.dias[-1]:%d-%m-%Y}.",
    )
    doc.add_paragraph("")

    _add_heading(doc, "Resultados", level=2)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    filas = [
        ("Indicador", "Valor"),
        ("Σ Con control (m³)", format_number_chilean(con_m3, 2)),
        ("Σ Sin control (m³)", format_number_chilean(sin_m3, 2)),
        ("Volumen evitado vs Sin control (m³)", format_number_chilean(ahorro_m3, 2)),
        ("Ahorro porcentual vs Sin control", f"{format_number_chilean(pct_ahorro, 2)} %"),
        ("Nota", "Los totales se calculan desde la rejilla horaria WES (0–23 h, hora Chile)."),
    ]
    for i, (a, b) in enumerate(filas):
        c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        c0.text = a
        c1.text = b
        for c in (c0, c1):
            for par in c.paragraphs:
                par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i == 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in par.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True

    doc.add_paragraph("")
    _p(doc, f"Generado: {datetime.now().strftime('%d-%m-%Y %H:%M')}")

    fname = f"Auditoria_CPA_Renca_{_safe_filename(node_name)}_{node_id}.docx"
    out = OUT_DIR / fname
    doc.save(out)
    return out


def main() -> int:
    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print("=" * 72)
    print("AUDITORÍAS CPA — RENCA (TODOS LOS PUNTOS)")
    print("=" * 72)
    print(f"Salida: {OUT_DIR.resolve()}")

    nodes: list[tuple[str, str]] = []
    for cid in COMPANIES_RENCA:
        for n in _company_nodes(cid):
            nid = n["nodeId"]
            nm = (n.get("name") or "").strip() or get_node_name(nid)
            nodes.append((nid, nm))

    # de-dup por nodeId
    seen: set[str] = set()
    uniq: list[tuple[str, str]] = []
    for nid, nm in nodes:
        if nid in seen:
            continue
        seen.add(nid)
        uniq.append((nid, nm))

    ok = 0
    fail = 0
    for i, (nid, nm) in enumerate(sorted(uniq), start=1):
        try:
            print(f"[{i}/{len(uniq)}] {nid} — {nm} ...", end=" ", flush=True)
            docx = _build_docx(nid, nm)
            pdf = convertir_word_a_pdf(docx)
            ok += 1
            print(f"[OK] {docx.name}" + (f" + {pdf.name}" if pdf else " (sin PDF)"))
        except Exception as e:
            fail += 1
            print(f"[ERROR] {e}")

    print("-" * 72)
    print(f"Completado. OK: {ok} | ERROR: {fail}")
    return 0 if fail == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())

