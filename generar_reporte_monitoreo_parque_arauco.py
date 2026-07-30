"""Generar reporte de monitoreo de Parque Arauco Kennedy con tabla, estadísticas y gráficas."""

from pathlib import Path
from datetime import datetime
from typing import List, Dict
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn

from generar_reporte_word import format_number_chilean, format_currency_chilean, add_logo_to_header


# Datos de la tabla de monitoreo
DATOS_MONITOREO = [
    {
        "nombre": "Impulsión Ander3-4 Matriz Principal",
        "medicion_8dias": 1731.9,
        "promedio_dia": 216.5,
        "alerta": 69.6,
        "porc_nocturno": 32,
        "valor_nocturno": 2603736,
        "es_matriz": True,
        "abastece": ["Impulsión Ander3-4 Locales Gast.", "Impulsión Anden 3-4 Restaurante"]
    },
    {
        "nombre": "Impulsión Ander3-4 Locales Gast.",
        "medicion_8dias": 1140.9,
        "promedio_dia": 142.6,
        "alerta": 64.8,
        "porc_nocturno": 45,
        "valor_nocturno": 2424168,
        "es_matriz": False,
        "abastecido_por": "Impulsión Ander3-4 Matriz Principal"
    },
    {
        "nombre": "Impulsión Anden 3-4 Restaurante",
        "medicion_8dias": 204.8,
        "promedio_dia": 25.6,
        "alerta": 0,
        "porc_nocturno": 0,
        "valor_nocturno": 0,
        "es_matriz": False,
        "abastecido_por": "Impulsión Ander3-4 Matriz Principal"
    },
    {
        "nombre": "Impulsión Sandia Mall 1 Piso-4",
        "medicion_8dias": 1562.0,
        "promedio_dia": 195.2,
        "alerta": 55.2,
        "porc_nocturno": 28,
        "valor_nocturno": 2065032,
        "es_independiente": True
    },
    {
        "nombre": "Impulsión Sandia Baños 2-3-6-7 Fredo",
        "medicion_8dias": 2619.6,
        "promedio_dia": 327.4,
        "alerta": 0,
        "porc_nocturno": 0,
        "valor_nocturno": 0,
        "es_independiente": True
    },
    {
        "nombre": "Llenado Pileta",
        "medicion_8dias": 22.5,
        "promedio_dia": 2.8,
        "alerta": 0,
        "porc_nocturno": 0,
        "valor_nocturno": 0,
        "es_independiente": True
    },
    {
        "nombre": "Llenado Pileta Cascada",
        "medicion_8dias": 22.5,
        "promedio_dia": 2.8,
        "alerta": 0.96,
        "porc_nocturno": 34,
        "valor_nocturno": 35913,
        "es_independiente": True
    },
    {
        "nombre": "Baño N°5 Damas",
        "medicion_8dias": 1.4,
        "promedio_dia": 0.2,
        "alerta": 0,
        "porc_nocturno": 0,
        "valor_nocturno": 0,
        "es_independiente": True
    },
    {
        "nombre": "Baño N°6 Varones",
        "medicion_8dias": 0.8,
        "promedio_dia": 0.1,
        "alerta": 0,
        "porc_nocturno": 0,
        "valor_nocturno": 0,
        "es_independiente": True
    },
    {
        "nombre": "Distrito de lujo DL",
        "medicion_8dias": 2910.2,
        "promedio_dia": 363.8,
        "alerta": 0,
        "porc_nocturno": 0,
        "valor_nocturno": 0,
        "es_independiente": True
    }
]


def calcular_precio_m3() -> float:
    """Calcula el precio promedio del m³ basándose en los valores de consumo nocturno."""
    total_m3_nocturno = 0
    total_valor = 0
    
    for punto in DATOS_MONITOREO:
        if punto["porc_nocturno"] > 0 and punto["valor_nocturno"] > 0:
            # Calcular m³ nocturno en 8 días
            m3_nocturno_8dias = (punto["promedio_dia"] * 8) * (punto["porc_nocturno"] / 100)
            total_m3_nocturno += m3_nocturno_8dias
            total_valor += punto["valor_nocturno"]
    
    if total_m3_nocturno > 0:
        return total_valor / total_m3_nocturno
    # Precio estándar si no se puede calcular
    return 1200.0  # CLP por m³


def calcular_consumo_proyectado_mall() -> Dict:
    """Calcula el consumo proyectado total del mall."""
    # Matriz principal (incluye gastronomía y restaurante)
    matriz = next(p for p in DATOS_MONITOREO if p["es_matriz"])
    consumo_matriz = matriz["promedio_dia"]
    alerta_matriz = matriz["alerta"]
    
    # Puntos independientes (excluyendo gastronomía y restaurante)
    puntos_independientes = [
        p for p in DATOS_MONITOREO 
        if p.get("es_independiente", False) and not p.get("es_matriz", False)
    ]
    
    consumo_independientes = sum(p["promedio_dia"] for p in puntos_independientes)
    alerta_independientes = sum(p["alerta"] for p in puntos_independientes)
    
    # Consumo total proyectado
    consumo_total_dia = consumo_matriz + consumo_independientes
    alerta_total_dia = alerta_matriz + alerta_independientes
    
    # Proyección mensual (30 días)
    consumo_mensual_m3 = consumo_total_dia * 30
    alerta_mensual_m3 = alerta_total_dia * 30
    
    # Calcular precio
    precio_m3 = calcular_precio_m3()
    
    # Valor en pesos
    consumo_mensual_clp = consumo_mensual_m3 * precio_m3
    alerta_mensual_clp = alerta_mensual_m3 * precio_m3
    
    return {
        "consumo_diario_m3": consumo_total_dia,
        "alerta_diaria_m3": alerta_total_dia,
        "consumo_mensual_m3": consumo_mensual_m3,
        "alerta_mensual_m3": alerta_mensual_m3,
        "consumo_mensual_clp": consumo_mensual_clp,
        "alerta_mensual_clp": alerta_mensual_clp,
        "precio_m3": precio_m3,
        "puntos_independientes": len(puntos_independientes)
    }


def crear_grafica_consumo_promedio(output_path: Path) -> None:
    """Crea gráfica de barras verticales apiladas: consumo neto (consumo - filtración) y filtración."""
    fig, ax = plt.subplots(figsize=(16, 10))
    
    # Ordenar de mayor a menor consumo
    datos_ordenados = sorted(DATOS_MONITOREO, key=lambda x: x["promedio_dia"], reverse=True)
    
    nombres = [p["nombre"] for p in datos_ordenados]
    consumos = [p["promedio_dia"] for p in datos_ordenados]
    filtraciones = [p["alerta"] for p in datos_ordenados]  # Alerta = filtración
    
    # Calcular consumo neto (consumo - filtración)
    consumos_netos = [max(0, consumo - filtracion) for consumo, filtracion in zip(consumos, filtraciones)]
    
    # Posiciones de las barras
    x = np.arange(len(nombres))
    width = 0.8  # Ancho de las barras (más ancho)
    
    # Crear barras apiladas: primero consumo neto (azul), luego consumo nocturno (rojo) encima
    bars_consumo_neto = ax.bar(x, consumos_netos, width, label='Consumo Neto (m³/día)', color='#1f77b4', alpha=0.8)  # Azul
    bars_filtracion = ax.bar(x, filtraciones, width, bottom=consumos_netos, label='Consumo nocturno (m³/día)', color='#d62728', alpha=0.8)  # Rojo
    
    # Configurar ejes
    ax.set_xlabel("", fontsize=11, fontweight="bold")  # Eliminado "Puntos de Monitoreo"
    ax.set_ylabel("m³/día", fontsize=11, fontweight="bold")
    ax.set_title("Consumo promedio diario vs consumo promedio nocturno", fontsize=18, fontweight="bold")  # Título agrandado y cambiado
    ax.set_xticks(x)
    ax.set_xticklabels(nombres, rotation=45, ha='right', fontsize=12)  # Reducido de 18 a 12
    ax.grid(axis="y", alpha=0.3)
    ax.set_axisbelow(True)
    
    # Calcular umbral para determinar barras pequeñas (menos del 5% del máximo)
    umbral_altura = max(consumos) * 0.05
    
    # Agregar valores en las barras de consumo neto (tamaño de fuente reducido)
    for bar, consumo_neto, consumo_total in zip(bars_consumo_neto, consumos_netos, consumos):
        height = bar.get_height()
        # Si la barra es muy pequeña, mostrar valor por encima
        if height <= umbral_altura and consumo_total > 0:
            # Mostrar valor total por encima de la barra
            ax.text(bar.get_x() + bar.get_width()/2., consumo_total,
                    f"{format_number_chilean(consumo_total, 1)}",
                    ha="center", va="bottom", fontsize=14, fontweight="bold", color='black')  # Reducido de 20 a 14
        elif height > umbral_altura and consumo_neto > 0:
            # Mostrar valor dentro de la barra si es suficientemente alta
            ax.text(bar.get_x() + bar.get_width()/2., height/2,
                    f"{format_number_chilean(consumo_neto, 1)}",
                    ha="center", va="center", fontsize=14, fontweight="bold", color='white')  # Reducido de 20 a 14
    
    # Agregar valores en las barras de filtración (tamaño de fuente reducido)
    for bar_filt, filtracion, consumo_neto, consumo_total in zip(bars_filtracion, filtraciones, consumos_netos, consumos):
        if filtracion > 0:
            height_filt = bar_filt.get_height()
            bottom = consumo_neto
            # Si la barra total es muy pequeña, el valor ya se mostró arriba
            if consumo_total > umbral_altura:
                # Solo mostrar valor si la barra es suficientemente alta
                if height_filt > umbral_altura:
                    ax.text(bar_filt.get_x() + bar_filt.get_width()/2., bottom + height_filt/2,
                            f"{format_number_chilean(filtracion, 1)}",
                            ha="center", va="center", fontsize=14, fontweight="bold", color='white')  # Reducido de 20 a 14
    
    # Agregar leyenda al lado derecho de la gráfica
    ax.legend(loc='center left', bbox_to_anchor=(1.02, 0.5), fontsize=10, frameon=True)
    
    # Ajustar layout para dar espacio a la leyenda a la derecha
    plt.tight_layout(rect=[0, 0, 0.85, 1])
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()


def crear_grafica_alertas(output_path: Path) -> None:
    """Crea gráfica de barras con alertas por punto, ordenada de menor a mayor."""
    fig, ax = plt.subplots(figsize=(14, 6))
    
    puntos_con_alerta = [p for p in DATOS_MONITOREO if p["alerta"] > 0]
    
    if not puntos_con_alerta:
        # Si no hay alertas, crear gráfica vacía con mensaje
        ax.text(0.5, 0.5, "No se registraron alertas en el periodo", 
                ha="center", va="center", fontsize=14, transform=ax.transAxes)
        ax.set_title("Consumo Nocturno promedio m3/dia", fontsize=18, fontweight="bold")
        plt.tight_layout()
        plt.savefig(output_path, dpi=300, bbox_inches="tight")
        plt.close()
        return
    
    # Ordenar de menor a mayor alerta
    puntos_con_alerta.sort(key=lambda x: x["alerta"], reverse=False)
    
    nombres = [p["nombre"] for p in puntos_con_alerta]
    alertas = [p["alerta"] for p in puntos_con_alerta]
    
    bars = ax.barh(range(len(nombres)), alertas, color="#d62728")
    ax.set_yticks(range(len(nombres)))
    ax.set_yticklabels(nombres, fontsize=16)  # Agrandado de 9 a 16
    ax.set_xlabel("m³/día", fontsize=11, fontweight="bold")  # Eliminado "Alerta"
    ax.set_title("Consumo Nocturno promedio m3/dia", fontsize=18, fontweight="bold")  # Título cambiado y agrandado
    ax.grid(axis="x", alpha=0.3)
    
    # Agregar valores (agrandados y sin m³/día)
    for i, (bar, alerta) in enumerate(zip(bars, alertas)):
        width = bar.get_width()
        ax.text(width, bar.get_y() + bar.get_height()/2, 
                f"{format_number_chilean(alerta, 1)}",
                ha="left", va="center", fontsize=14, fontweight="bold")  # Agrandado de 8 a 14, eliminado " m³/día"
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()


def crear_grafica_consumo_nocturno(output_path: Path) -> None:
    """Crea gráfica de consumo nocturno valorizado, ordenada de menor a mayor."""
    fig, ax = plt.subplots(figsize=(14, 6))
    
    puntos_con_nocturno = [p for p in DATOS_MONITOREO if p["valor_nocturno"] > 0]
    
    if not puntos_con_nocturno:
        plt.close()
        return
    
    # Ordenar de menor a mayor valor
    puntos_con_nocturno.sort(key=lambda x: x["valor_nocturno"], reverse=False)
    
    nombres = [p["nombre"] for p in puntos_con_nocturno]
    valores = [p["valor_nocturno"] / 1_000_000 for p in puntos_con_nocturno]  # En millones
    
    bars = ax.barh(range(len(nombres)), valores, color="#9467bd")
    ax.set_yticks(range(len(nombres)))
    ax.set_yticklabels(nombres, fontsize=9)
    ax.set_xlabel("Proyección Mensual Consumo Nocturno (Millones de CLP)", fontsize=11, fontweight="bold")
    ax.set_title("Proyección Mensual de Consumo Nocturno (Presunta Fuga) por Punto", fontsize=13, fontweight="bold")
    ax.grid(axis="x", alpha=0.3)
    
    # Agregar valores
    for i, (bar, valor) in enumerate(zip(bars, valores)):
        width = bar.get_width()
        ax.text(width, bar.get_y() + bar.get_height()/2, 
                f"${format_number_chilean(valor * 1_000_000, 0)}",
                ha="left", va="center", fontsize=8, fontweight="bold")
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()


def crear_grafica_comparacion_matriz_independientes(output_path: Path) -> None:
    """Crea gráfica comparativa entre matriz principal y puntos independientes."""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    matriz = next(p for p in DATOS_MONITOREO if p["es_matriz"])
    puntos_independientes = [p for p in DATOS_MONITOREO if p.get("es_independiente", False)]
    
    consumo_matriz = matriz["promedio_dia"]
    consumo_independientes = sum(p["promedio_dia"] for p in puntos_independientes)
    
    categorias = ["Matriz Principal\n(Anden 3-4)", "Puntos Independientes"]
    consumos = [consumo_matriz, consumo_independientes]
    colores = ["#1f77b4", "#2ca02c"]
    
    bars = ax.bar(categorias, consumos, color=colores, width=0.6)
    ax.set_ylabel("Consumo Promedio (m³/día)", fontsize=11, fontweight="bold")
    ax.set_title("Comparación: Matriz Principal vs Puntos Independientes", 
                 fontsize=13, fontweight="bold")
    ax.grid(axis="y", alpha=0.3)
    
    # Agregar valores
    for bar, consumo in zip(bars, consumos):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, height,
                f"{format_number_chilean(consumo, 1)} m³/día",
                ha="center", va="bottom", fontsize=10, fontweight="bold")
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()


def agregar_tabla_monitoreo(doc: Document) -> None:
    """Agrega la tabla de monitoreo al documento."""
    doc.add_heading("Tabla de Monitoreo", level=1)
    
    # Crear tabla
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = "Light Grid Accent 1"
    
    # Encabezados
    headers = ["Monitoreo", "Medición (m³/8 días)", "Promedio (m³/día)", 
               "Alerta (m³/día)", "% consumo nocturno / Proyección mensual (presunta fuga) en $"]
    header_cells = tabla.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Agregar datos
    for punto in DATOS_MONITOREO:
        row = tabla.add_row()
        row.cells[0].text = punto["nombre"]
        row.cells[1].text = format_number_chilean(punto["medicion_8dias"], 1)
        row.cells[2].text = format_number_chilean(punto["promedio_dia"], 1)
        row.cells[3].text = format_number_chilean(punto["alerta"], 1) if punto["alerta"] > 0 else "0"
        if punto["porc_nocturno"] > 0:
            row.cells[4].text = f"{punto['porc_nocturno']}% ({format_currency_chilean(punto['valor_nocturno'])})"
        else:
            row.cells[4].text = "0,0% ($0)"
        
        # Alinear celdas numéricas
        for i in [1, 2, 3, 4]:
            row.cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Agregar total al final
    total_valor = sum(p["valor_nocturno"] for p in DATOS_MONITOREO)
    row_total = tabla.add_row()
    row_total.cells[0].text = "TOTAL"
    row_total.cells[1].text = ""
    row_total.cells[2].text = ""
    row_total.cells[3].text = ""
    row_total.cells[4].text = format_currency_chilean(total_valor)
    
    # Formatear fila total
    for cell in row_total.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            if cell != row_total.cells[0]:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def generar_reporte() -> Path:
    """Genera el reporte completo de monitoreo."""
    # Crear directorio de salida
    output_dir = Path("reports") / "Parque_Arauco" / "MONITOREO"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Crear subdirectorio con timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    folder_name = f"Monitoreo_{timestamp}"
    report_folder = output_dir / folder_name
    report_folder.mkdir(exist_ok=True)
    
    # Crear documento
    doc = Document()
    
    # Agregar logo
    add_logo_to_header(doc)
    
    # Título
    title = doc.add_heading("Reporte de Monitoreo de Consumo de Agua", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph("Parque Arauco Kennedy")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_paragraph()  # Espacio
    
    # Fecha del reporte
    fecha_reporte = doc.add_paragraph(f"Fecha del reporte: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    fecha_reporte.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in fecha_reporte.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # 1. INTRODUCCIÓN Y NARRACIÓN
    doc.add_heading("1. Introducción", level=1)
    
    intro_text = """
Este reporte presenta un análisis detallado del consumo de agua en el Mall Parque Arauco Kennedy, 
basado en los datos de monitoreo de 10 puntos de medición durante un periodo de 8 días.

El sistema de abastecimiento del mall está estructurado de la siguiente manera:

• **Matriz Principal (Anden 3-4)**: Este punto abastece directamente a los locales de gastronomía 
  y al restaurante del anden 3-4. Su consumo incluye tanto el consumo propio como el de los puntos 
  que abastece.

• **Puntos Independientes**: El resto de los puntos de monitoreo representan fuentes de agua 
  independientes que alimentan diferentes sectores del mall, incluyendo:
  - Impulsión Sandia Mall (1° piso a 4° piso)
  - Impulsión Sandia Baños
  - Sistemas de llenado de piletas
  - Baños públicos
  - Distrito de lujo

Para el cálculo del consumo total proyectado del mall, se considera únicamente la matriz principal 
(que ya incluye gastronomía y restaurante) más todos los puntos independientes, evitando así la 
doble contabilización del consumo de los locales de gastronomía y restaurante.
"""
    
    for paragraph_text in intro_text.strip().split("\n\n"):
        para = doc.add_paragraph(paragraph_text.strip())
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in para.runs:
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # 2. TABLA DE MONITOREO
    doc.add_paragraph()
    nota_tabla = doc.add_paragraph(
        "Nota: Los valores monetarios ($) en la columna '% consumo nocturno / Proyección mensual' "
        "representan la proyección mensual del consumo nocturno, que puede indicar presunta fuga de agua."
    )
    for run in nota_tabla.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    agregar_tabla_monitoreo(doc)
    
    doc.add_page_break()
    
    # 3. ESTADÍSTICAS
    doc.add_heading("2. Estadísticas Generales", level=1)
    
    proyeccion = calcular_consumo_proyectado_mall()
    
    stats_text = f"""
**Consumo Total del Mall (Proyectado)**

• **Consumo Diario Total**: {format_number_chilean(proyeccion["consumo_diario_m3"], 1)} m³/día
• **Consumo Mensual Total**: {format_number_chilean(proyeccion["consumo_mensual_m3"], 1)} m³/mes
• **Valor Mensual Proyectado**: {format_currency_chilean(proyeccion["consumo_mensual_clp"])}

**Alertas de Fuga (Proyectado)**

• **Alerta Diaria Total**: {format_number_chilean(proyeccion["alerta_diaria_m3"], 1)} m³/día
• **Alerta Mensual Total**: {format_number_chilean(proyeccion["alerta_mensual_m3"], 1)} m³/mes
• **Valor Mensual de Alertas**: {format_currency_chilean(proyeccion["alerta_mensual_clp"])}

**Estructura del Sistema**

• **Punto Matriz Principal**: 1 (Anden 3-4, incluye gastronomía y restaurante)
• **Puntos Independientes**: {proyeccion["puntos_independientes"]}
• **Total de Puntos Monitoreados**: {len(DATOS_MONITOREO)}

**Proyección Mensual de Consumo Nocturno (Presunta Fuga)**

• **Valor Total Proyectado**: {format_currency_chilean(sum(p["valor_nocturno"] for p in DATOS_MONITOREO))}
• **Nota**: Este valor representa la proyección mensual del consumo nocturno, que puede indicar presunta fuga de agua.
"""
    
    for line in stats_text.strip().split("\n"):
        if line.strip():
            if line.startswith("**"):
                # Es un encabezado
                para = doc.add_paragraph()
                run = para.add_run(line.replace("**", ""))
                run.font.bold = True
                run.font.size = Pt(11)
            else:
                para = doc.add_paragraph(line.strip())
                for run in para.runs:
                    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # 4. GRÁFICAS
    doc.add_heading("3. Análisis Gráfico", level=1)
    
    # Gráfica 1: Consumo promedio
    doc.add_heading("3.1. Consumo Promedio Diario por Punto", level=2)
    graph1_path = report_folder / "grafica_consumo_promedio.png"
    crear_grafica_consumo_promedio(graph1_path)
    doc.add_picture(str(graph1_path), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    legend_text = doc.add_paragraph(
        "Leyenda: Azul = Matriz Principal | Verde = Puntos Independientes | Naranja = Puntos Abastecidos por Matriz"
    )
    for run in legend_text.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
    legend_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Gráfica 2: Alertas
    doc.add_heading("3.2. Alertas de Fuga por Punto", level=2)
    graph2_path = report_folder / "grafica_alertas.png"
    crear_grafica_alertas(graph2_path)
    doc.add_picture(str(graph2_path), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Gráfica 3: Consumo nocturno
    doc.add_heading("3.3. Proyección Mensual de Consumo Nocturno (Presunta Fuga)", level=2)
    graph3_path = report_folder / "grafica_consumo_nocturno.png"
    crear_grafica_consumo_nocturno(graph3_path)
    if graph3_path.exists():
        doc.add_picture(str(graph3_path), width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Gráfica 4: Comparación matriz vs independientes
    doc.add_heading("3.4. Comparación: Matriz Principal vs Puntos Independientes", level=2)
    graph4_path = report_folder / "grafica_comparacion.png"
    crear_grafica_comparacion_matriz_independientes(graph4_path)
    doc.add_picture(str(graph4_path), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # 5. CONCLUSIONES Y PROYECCIONES
    doc.add_heading("4. Conclusiones y Proyecciones", level=1)
    
    conclusiones_text = f"""
**Consumo Total Proyectado del Mall**

El consumo total proyectado del Mall Parque Arauco Kennedy, considerando la matriz principal 
(que incluye gastronomía y restaurante) más todos los puntos independientes, es:

• **Consumo Diario**: {format_number_chilean(proyeccion["consumo_diario_m3"], 1)} m³/día
• **Consumo Mensual**: {format_number_chilean(proyeccion["consumo_mensual_m3"], 1)} m³/mes
• **Valor Mensual**: {format_currency_chilean(proyeccion["consumo_mensual_clp"])}

**Proyección de Alertas**

Las alertas de fuga detectadas proyectan un consumo adicional de:

• **Alerta Diaria**: {format_number_chilean(proyeccion["alerta_diaria_m3"], 1)} m³/día
• **Alerta Mensual**: {format_number_chilean(proyeccion["alerta_mensual_m3"], 1)} m³/mes
• **Valor Mensual de Alertas**: {format_currency_chilean(proyeccion["alerta_mensual_clp"])}

**Recomendaciones**

1. Se recomienda revisar los puntos con alertas de fuga activas para identificar y corregir 
   posibles pérdidas de agua.

2. El consumo nocturno representa un indicador importante de posibles fugas. Los puntos con 
   mayor porcentaje de consumo nocturno deben ser monitoreados con mayor frecuencia.

3. El punto "Distrito de lujo DL" presenta el mayor consumo promedio diario ({format_number_chilean(next(p["promedio_dia"] for p in DATOS_MONITOREO if "Distrito" in p["nombre"]), 1)} m³/día), 
   por lo que se recomienda un análisis detallado de sus patrones de consumo.

4. La matriz principal (Anden 3-4) muestra una alerta de {format_number_chilean(next(p["alerta"] for p in DATOS_MONITOREO if p["es_matriz"]), 1)} m³/día, 
   lo que requiere atención inmediata para evitar pérdidas significativas.
"""
    
    for paragraph_text in conclusiones_text.strip().split("\n\n"):
        para = doc.add_paragraph(paragraph_text.strip())
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in para.runs:
            run.font.size = Pt(11)
            if run.text.startswith("•") or run.text.startswith("**"):
                run.font.bold = True
    
    # Guardar documento
    doc_path = report_folder / "Reporte_Monitoreo_Parque_Arauco_Kennedy.docx"
    doc.save(str(doc_path))
    
    print(f"[OK] Reporte generado: {doc_path}")
    print(f"[INFO] Consumo total proyectado del mall:")
    print(f"  - Diario: {format_number_chilean(proyeccion['consumo_diario_m3'], 1)} m³/día")
    print(f"  - Mensual: {format_number_chilean(proyeccion['consumo_mensual_m3'], 1)} m³/mes")
    print(f"  - Valor mensual: {format_currency_chilean(proyeccion['consumo_mensual_clp'])}")
    
    return doc_path


if __name__ == "__main__":
    generar_reporte()

