"""
Seguimiento Bupa Antofagasta — proyección WES vs data actualizada.

Compara día a día el Medidor Principal Sanitaria (y salas) contra la proyección
del reporte 23–27/07/2026 (~250 m³/día → ~7.499 m³ / 30 días) y recalcula
proyección con días civiles completos hasta ayer.

Uso:
  python3 generar_seguimiento_bupa_antofagasta.py
"""

from __future__ import annotations

import sys
from datetime import datetime, timedelta
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from generar_reporte_word import (
    add_formatted_title,
    add_picture_with_pagination,
    add_table,
    acl_node_base_url,
    build_hourly_consumption_line_chart,
    convertir_word_a_pdf,
    fetch_json,
    find_max_nocturnal_consumption_day,
    flatten_measures,
    format_currency_chilean,
    format_number_chilean,
    normalize_measures_payload,
    parse_date,
    summarize_consumption,
)

# --- Baseline del último reporte (AGREGADO_20260728_1950) ---
BASELINE_START = "23/07/2026"
BASELINE_END = "27/07/2026"
BASELINE_DIAS = 5
BASELINE_SANIT_TOTAL = 1249.8
BASELINE_SANIT_PROM = 250.0  # redondeo del informe
BASELINE_SANIT_PROY = 7498.7
BASELINE_SANIT_PROY_CLP = 20_761_364.0
BASELINE_SALAS_PROY = 1435.6
BASELINE_GAP_PROY = 6063.2
BASELINE_PCT_SALAS = 19.1
BASELINE_PCT_GAP = 80.9

FACTURA_M3 = 6696.0
FACTURA_CLP = 18_538_860.0
PRECIO_M3 = FACTURA_CLP / FACTURA_M3  # ~$2.769 CLP/m³ (ticket factura julio)

# Día en que se consolida el descenso post-mejoras / pérdidas detectadas
MEJORAS_DESDE = "2026-07-29"

NODOS_DEF = [
    {"id": "000029-07", "nombre": "Sala de Bomba Principal", "tipo": "bomba"},
    {"id": "000029-08", "nombre": "Sala de Bomba Sexto Piso", "tipo": "bomba"},
    {"id": "000029-09", "nombre": "Medidor Principal Sanitaria", "tipo": "cuenta"},
    {"id": "000029-10", "nombre": "Sala de Bomba N°2", "tipo": "bomba"},
]

FOLDER = Path("reports/Bupa_Antofagasta/ABREGADO")


def _dias_completos_hasta_ayer() -> tuple[datetime, datetime, int]:
    hoy = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    start = parse_date(BASELINE_START)
    end = hoy - timedelta(days=1)
    end = end.replace(hour=23, minute=59, second=59)
    n = (end.date() - start.date()).days + 1
    return start, end, n


def _fetch_by_day(node_id: str, start: datetime, end: datetime) -> dict[str, float]:
    raw = fetch_json(
        f"{acl_node_base_url()}/nodes/measures/dates",
        params=[
            ("id", node_id),
            ("start", start.strftime("%d%m%Y")),
            ("end", end.strftime("%d%m%Y")),
        ],
    )
    measures = flatten_measures(normalize_measures_payload(raw, node_id))
    return {m.date.strftime("%Y-%m-%d"): float(m.total_m3) for m in measures}


def _variacion_txt(delta: float, umbral: float = 1.0) -> str:
    if abs(delta) < umbral:
        return "estable"
    return "aumentó" if delta > 0 else "bajó"


def _chart_diario_sanitaria(
    path: Path,
    days: list[str],
    values: list[float],
    baseline: float,
) -> Path:
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    colors = ["#c0392b" if v < baseline * 0.85 else "#0050b3" for v in values]
    bars = ax.bar(days, values, color=colors, edgecolor="#003a80", linewidth=0.8, alpha=0.88)
    ax.axhline(
        baseline,
        color="#e67e22",
        linestyle="--",
        linewidth=2,
        label=f"Proyección previa ({format_number_chilean(baseline, 0)} m³/día)",
    )
    ax.set_ylabel("m³ / día", fontsize=11, fontweight="bold")
    ax.set_title(
        "Medidor Principal Sanitaria — consumo diario vs proyección 23–27/07",
        fontsize=12,
        fontweight="bold",
    )
    ax.set_ylim(bottom=0)
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=40, ha="right", fontsize=9)
    for bar, v in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{v:.0f}",
            ha="center",
            va="bottom",
            fontsize=8,
            fontweight="bold",
        )
    ax.legend(loc="upper right", fontsize=9)
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _chart_proyeccion(path: Path, rows: list[dict], n_dias: int) -> Path:
    ordered = sorted(rows, key=lambda r: r["proy_m3"], reverse=True)
    fig, ax = plt.subplots(figsize=(8, 4.8))
    bars = ax.bar(
        [r["nombre"] for r in ordered],
        [r["proy_m3"] for r in ordered],
        color="#0050b3",
        alpha=0.85,
        edgecolor="#003a80",
        linewidth=1.1,
    )
    ax.axhline(
        BASELINE_SANIT_PROY,
        color="#e67e22",
        linestyle="--",
        linewidth=1.6,
        label=f"Proy. previa Sanitaria ({format_number_chilean(BASELINE_SANIT_PROY, 0)} m³)",
    )
    ax.set_ylabel("m³ / 30 días", fontsize=12, fontweight="bold")
    ax.set_title(
        f"Proyección mensual actualizada (promedio {n_dias} días × 30)",
        fontsize=13,
        fontweight="bold",
    )
    ax.set_ylim(bottom=0)
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=25, ha="right", fontsize=10)
    for bar, r in zip(bars, ordered):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(r['proy_m3'], 1)} m³",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )
    ax.legend(loc="upper right", fontsize=8)
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _chart_comparacion_periodos(
    path: Path,
    prom_base: float,
    proy_base: float,
    prom_nuevo: float,
    proy_nuevo: float,
    prom_reciente: float,
    proy_reciente: float,
) -> Path:
    labels = [
        "Baseline\n23–27/07",
        "Actualizado\n23/07–ayer",
        "Post-mejoras\n29/07–ayer",
    ]
    proms = [prom_base, prom_nuevo, prom_reciente]
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    bars = ax.bar(labels, proms, color=["#7f8c8d", "#0050b3", "#27ae60"], edgecolor="#222", alpha=0.9)
    ax.set_ylabel("m³ / día (Sanitaria)", fontsize=11, fontweight="bold")
    ax.set_title("¿Continúa la proyección? — promedio diario Sanitaria", fontsize=12, fontweight="bold")
    for bar, p, proy in zip(bars, proms, [proy_base, proy_nuevo, proy_reciente]):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{p:.1f} m³/d\n→ {proy:.0f} m³/30d",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    ax.set_ylim(0, max(proms) * 1.35)
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _chart_factura_fin_mes(
    path: Path,
    factura_clp: float,
    proy_previa_clp: float,
    proy_mejoras_clp: float,
    cierre_agosto_clp: float,
) -> Path:
    labels = [
        "Factura julio\n(ticket)",
        "Proy. previa\n(con pérdidas)",
        "Proy. 30 días\npost-mejoras",
        "Estimación\ncierre agosto",
    ]
    values = [factura_clp, proy_previa_clp, proy_mejoras_clp, cierre_agosto_clp]
    colors = ["#7f8c8d", "#c0392b", "#27ae60", "#0050b3"]
    fig, ax = plt.subplots(figsize=(9.0, 4.8))
    bars = ax.bar(labels, values, color=colors, edgecolor="#222", alpha=0.9)
    ax.set_ylabel("CLP", fontsize=11, fontweight="bold")
    ax.set_title(
        "Proyección de pago (tarifa ticket julio)",
        fontsize=12,
        fontweight="bold",
    )
    ax.yaxis.set_major_formatter(
        plt.FuncFormatter(lambda x, _: f"${x/1e6:.1f} M" if x >= 1e6 else f"${x:,.0f}")
    )
    for bar, v in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            format_currency_chilean(v),
            ha="center",
            va="bottom",
            fontsize=8,
            fontweight="bold",
        )
    ax.set_ylim(0, max(values) * 1.22)
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _convert_with_libreoffice(docx_path: Path) -> Path | None:
    """Fallback Linux: LibreOffice headless."""
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    out_dir = docx_path.parent
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=180,
        )
    except Exception as exc:
        print(f"[ADVERTENCIA] LibreOffice falló: {exc}")
        return None
    pdf = docx_path.with_suffix(".pdf")
    return pdf if pdf.is_file() else None


def _set_black(paragraph) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def build_report() -> tuple[Path, Path, dict]:
    start, end, n_dias = _dias_completos_hasta_ayer()
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    out_dir = FOLDER / f"AGREGADO_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    # Serie diaria por nodo
    series: dict[str, dict[str, float]] = {}
    for n in NODOS_DEF:
        series[n["id"]] = _fetch_by_day(n["id"], start, end)

    sanit_id = "000029-09"
    sanit_days = series[sanit_id]
    all_dates = sorted(sanit_days.keys())

    # Totales / proyección periodo completo
    rows = []
    for n in NODOS_DEF:
        by_day = series[n["id"]]
        total = sum(by_day.get(d, 0.0) for d in all_dates)
        prom = total / n_dias if n_dias else 0.0
        proy = prom * 30.0
        rows.append(
            {
                **n,
                "m3_periodo": total,
                "prom_dia": prom,
                "proy_m3": proy,
                "proy_clp": proy * PRECIO_M3,
                "by_day": by_day,
            }
        )

    sanit = next(r for r in rows if r["tipo"] == "cuenta")
    bombas = [r for r in rows if r["tipo"] == "bomba"]
    salas_m3 = sum(b["proy_m3"] for b in bombas)
    salas_clp = sum(b["proy_clp"] for b in bombas)
    gap_m3 = max(0.0, sanit["proy_m3"] - salas_m3)
    gap_clp = gap_m3 * PRECIO_M3
    pct_salas = salas_m3 / sanit["proy_m3"] * 100.0 if sanit["proy_m3"] else 0.0
    pct_gap = 100.0 - pct_salas

    # Periodo reciente (desde 28/07 = día siguiente al baseline)
    reciente_start = (parse_date(BASELINE_END) + timedelta(days=1)).strftime("%Y-%m-%d")
    reciente_dates = [d for d in all_dates if d >= reciente_start]
    n_reciente = len(reciente_dates)
    sanit_reciente_total = sum(sanit_days.get(d, 0.0) for d in reciente_dates)
    sanit_reciente_prom = sanit_reciente_total / n_reciente if n_reciente else 0.0
    sanit_reciente_proy = sanit_reciente_prom * 30.0

    # Post-mejoras / pérdidas corregidas (consolidado desde 29/07)
    mejoras_dates = [d for d in all_dates if d >= MEJORAS_DESDE]
    n_mejoras = len(mejoras_dates)
    sanit_mej_total = sum(sanit_days.get(d, 0.0) for d in mejoras_dates)
    sanit_mej_prom = sanit_mej_total / n_mejoras if n_mejoras else 0.0
    sanit_mej_proy_m3 = sanit_mej_prom * 30.0
    sanit_mej_proy_clp = sanit_mej_proy_m3 * PRECIO_M3

    # Estimación cierre mes calendario (agosto 2026)
    hoy = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    mes_ini = hoy.replace(day=1)
    if hoy.month == 12:
        mes_fin = hoy.replace(year=hoy.year + 1, month=1, day=1) - timedelta(days=1)
    else:
        mes_fin = hoy.replace(month=hoy.month + 1, day=1) - timedelta(days=1)
    dias_mes = mes_fin.day
    ago_dates = [
        d for d in all_dates
        if d.startswith(hoy.strftime("%Y-%m"))
    ]
    consumo_ago_real = sum(sanit_days.get(d, 0.0) for d in ago_dates)
    dias_ago_completos = len(ago_dates)
    dias_restantes = max(0, dias_mes - dias_ago_completos)
    # Ritmo proyectado restante = post-mejoras (más representativo del nuevo estado)
    ritmo_cierre = sanit_mej_prom if sanit_mej_prom > 0 else sanit_reciente_prom
    cierre_ago_m3 = consumo_ago_real + ritmo_cierre * dias_restantes
    cierre_ago_clp = cierre_ago_m3 * PRECIO_M3

    ahorro_vs_factura = FACTURA_CLP - sanit_mej_proy_clp
    ahorro_vs_proy_prev = BASELINE_SANIT_PROY_CLP - sanit_mej_proy_clp
    ahorro_cierre_vs_factura = FACTURA_CLP - cierre_ago_clp
    ahorro_cierre_vs_proy_prev = BASELINE_SANIT_PROY_CLP - cierre_ago_clp

    # Día a día vs baseline + vs día anterior
    dia_rows = []
    prev_val = None
    for d in all_dates:
        v = sanit_days.get(d, 0.0)
        vs_base = v - BASELINE_SANIT_PROM
        vs_prev = None if prev_val is None else v - prev_val
        salas_d = sum(series[b["id"]].get(d, 0.0) for b in bombas)
        dia_rows.append(
            {
                "fecha": d,
                "sanitaria": v,
                "salas": salas_d,
                "vs_base": vs_base,
                "vs_prev": vs_prev,
                "var_base": _variacion_txt(vs_base, umbral=5.0),
                "var_prev": "—" if vs_prev is None else _variacion_txt(vs_prev, umbral=1.0),
            }
        )
        prev_val = v

    # Charts
    labels_short = [datetime.strptime(d, "%Y-%m-%d").strftime("%d/%m") for d in all_dates]
    chart_diario = _chart_diario_sanitaria(
        out_dir / "chart_diario_sanitaria_vs_proyeccion.png",
        labels_short,
        [sanit_days[d] for d in all_dates],
        BASELINE_SANIT_PROM,
    )
    chart_proy = _chart_proyeccion(
        out_dir / "chart_proyeccion_mensual_actualizada.png", rows, n_dias
    )
    chart_cmp = _chart_comparacion_periodos(
        out_dir / "chart_comparacion_periodos_proyeccion.png",
        BASELINE_SANIT_PROM,
        BASELINE_SANIT_PROY,
        sanit["prom_dia"],
        sanit["proy_m3"],
        sanit_mej_prom,
        sanit_mej_proy_m3,
    )
    chart_factura = _chart_factura_fin_mes(
        out_dir / "chart_proyeccion_factura_fin_mes.png",
        FACTURA_CLP,
        BASELINE_SANIT_PROY_CLP,
        sanit_mej_proy_clp,
        cierre_ago_clp,
    )

    result_noct = find_max_nocturnal_consumption_day(sanit_id, None, start, end)
    chart_noct = None
    target_dt = None
    if result_noct:
        target_dt, hourly = result_noct
        chart_noct = out_dir / "chart_max_nocturnal_sanitaria.png"
        build_hourly_consumption_line_chart(
            hourly,
            chart_noct,
            target_dt,
            f"Día con mayor consumo nocturno ({target_dt.strftime('%d-%m-%y')})",
        )

    # --- Word ---
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Clínica Bupa Antofagasta — Seguimiento de proyección WES")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 80, 179)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph(
        f"Periodo actualizado: {start.strftime('%d/%m/%Y')} – {end.strftime('%d/%m/%Y')} "
        f"({n_dias} días civiles completos). "
        f"Referencia: reporte previo 23–27/07/2026 (proyección ~"
        f"{format_number_chilean(BASELINE_SANIT_PROY, 1)} m³ / 30 días). "
        f"Contexto operativo: el encargado de operaciones indicó que se realizaron "
        f"mejoras y se detectaron/corrigieron pérdidas; el descenso de consumo desde "
        f"fines de julio es coherente con ese trabajo."
    )
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_black(sub)

    # Veredicto
    add_formatted_title(doc, "Veredicto: ¿continúa la proyección?")
    delta_prom = sanit["prom_dia"] - BASELINE_SANIT_PROM
    delta_proy = sanit["proy_m3"] - BASELINE_SANIT_PROY
    delta_rec = sanit_reciente_prom - BASELINE_SANIT_PROM
    delta_mej = sanit_mej_prom - BASELINE_SANIT_PROM
    continua = abs(delta_prom) / BASELINE_SANIT_PROM < 0.10  # ±10%

    if continua:
        veredicto = (
            f"La proyección previa (~{format_number_chilean(BASELINE_SANIT_PROM, 0)} m³/día) "
            f"se mantiene en el orden esperado: el promedio actualizado es "
            f"{format_number_chilean(sanit['prom_dia'], 1)} m³/día "
            f"(proyección 30 días: {format_number_chilean(sanit['proy_m3'], 1)} m³)."
        )
    else:
        veredicto = (
            f"La proyección previa (ritmo con pérdidas, ~"
            f"{format_number_chilean(BASELINE_SANIT_PROM, 0)} m³/día → "
            f"{format_number_chilean(BASELINE_SANIT_PROY, 1)} m³ / "
            f"{format_currency_chilean(BASELINE_SANIT_PROY_CLP)}) "
            f"ya no aplica: tras las mejoras y la corrección de pérdidas, "
            f"el Medidor Principal Sanitaria bajó a "
            f"{format_number_chilean(sanit_mej_prom, 1)} m³/día "
            f"en el tramo post-mejoras ({MEJORAS_DESDE} a {mejoras_dates[-1] if mejoras_dates else '—'}, "
            f"{n_mejoras} días). "
            f"Eso implica una proyección de cuenta a 30 días de "
            f"{format_number_chilean(sanit_mej_proy_m3, 1)} m³ "
            f"({format_currency_chilean(sanit_mej_proy_clp)}), "
            f"{format_number_chilean(abs(delta_mej), 1)} m³/día bajo el baseline."
        )
    p = doc.add_paragraph(veredicto)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_black(p)

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_cmp), Inches(6.0), keep_with_next=True)

    # --- PROYECCIÓN FACTURA FIN DE MES (pedido explícito) ---
    add_formatted_title(doc, "Proyección de pago a fin de mes (tarifa ticket julio)")
    intro_fac = doc.add_paragraph(
        f"Valorización con la tarifa efectiva del ticket/factura de julio usado en el "
        f"reporte anterior: {format_currency_chilean(FACTURA_CLP)} ÷ "
        f"{format_number_chilean(FACTURA_M3, 0)} m³ = "
        f"${format_number_chilean(PRECIO_M3, 0)} CLP/m³. "
        f"Se proyecta cuánto pagarían si el ritmo post-mejoras se mantiene."
    )
    intro_fac.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_black(intro_fac)

    fac_table = [
        (
            "Escenario",
            "Base de cálculo",
            "m³",
            "CLP (tarifa ticket)",
            "Vs factura julio",
        ),
        (
            "Factura julio (ticket de referencia)",
            "Histórica",
            format_number_chilean(FACTURA_M3, 0),
            format_currency_chilean(FACTURA_CLP),
            "—",
        ),
        (
            "Proyección previa (23–27/07, con pérdidas)",
            f"{format_number_chilean(BASELINE_SANIT_PROM, 0)} m³/día × 30",
            format_number_chilean(BASELINE_SANIT_PROY, 1),
            format_currency_chilean(BASELINE_SANIT_PROY_CLP),
            f"+{format_currency_chilean(BASELINE_SANIT_PROY_CLP - FACTURA_CLP)}",
        ),
        (
            "Proyección 30 días post-mejoras (recomendado)",
            f"{format_number_chilean(sanit_mej_prom, 1)} m³/día × 30 "
            f"({n_mejoras} días desde {datetime.strptime(MEJORAS_DESDE, '%Y-%m-%d').strftime('%d/%m')})",
            format_number_chilean(sanit_mej_proy_m3, 1),
            format_currency_chilean(sanit_mej_proy_clp),
            f"{format_currency_chilean(ahorro_vs_factura)} "
            f"({'ahorro' if ahorro_vs_factura > 0 else 'alza'})",
        ),
        (
            f"Estimación cierre {hoy.strftime('%B %Y').replace('August', 'agosto')}",
            f"{format_number_chilean(consumo_ago_real, 1)} m³ reales "
            f"({dias_ago_completos} días) + "
            f"{format_number_chilean(ritmo_cierre, 1)} m³/día × {dias_restantes} días restantes",
            format_number_chilean(cierre_ago_m3, 1),
            format_currency_chilean(cierre_ago_clp),
            f"{format_currency_chilean(ahorro_cierre_vs_factura)} "
            f"({'ahorro' if ahorro_cierre_vs_factura > 0 else 'alza'})",
        ),
    ]
    add_table(doc, "Cuánto pagarían a fin de mes", fac_table, wes_style=True)

    lectura_fac = doc.add_paragraph()
    lectura_fac.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lectura_fac.add_run("Lectura para el cliente: ").bold = True
    lectura_fac.add_run(
        f"si se mantiene el ritmo post-mejoras ("
        f"{format_number_chilean(sanit_mej_prom, 1)} m³/día), la cuenta mensual "
        f"se acerca a {format_number_chilean(sanit_mej_proy_m3, 1)} m³ / "
        f"{format_currency_chilean(sanit_mej_proy_clp)}. "
        f"Respecto al ticket de julio ({format_currency_chilean(FACTURA_CLP)}) "
        f"eso implica un ahorro proyectado de "
        f"{format_currency_chilean(ahorro_vs_factura)} "
        f"({format_number_chilean(ahorro_vs_factura / FACTURA_CLP * 100, 1)}%). "
        f"Frente a la proyección previa con pérdidas "
        f"({format_currency_chilean(BASELINE_SANIT_PROY_CLP)}), el ahorro sería de "
        f"{format_currency_chilean(ahorro_vs_proy_prev)}. "
        f"Para el cierre de agosto {hoy.year}: "
        f"{format_number_chilean(cierre_ago_m3, 1)} m³ / "
        f"{format_currency_chilean(cierre_ago_clp)} "
        f"(ahorro vs ticket julio: {format_currency_chilean(ahorro_cierre_vs_factura)}; "
        f"vs proyección previa: {format_currency_chilean(ahorro_cierre_vs_proy_prev)})."
    )
    _set_black(lectura_fac)

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_factura), Inches(6.2), keep_with_next=True)

    # Tabla comparación proyecciones
    add_formatted_title(doc, "Comparación de proyecciones (Medidor Principal Sanitaria)")
    cmp_table = [
        (
            "Escenario",
            "Días",
            "Promedio diario (m³)",
            "Proyección 30 días (m³)",
            "Proyección 30 días (CLP)",
            "Vs baseline",
        ),
        (
            "Baseline (reporte 23–27/07)",
            str(BASELINE_DIAS),
            format_number_chilean(BASELINE_SANIT_PROM, 1),
            format_number_chilean(BASELINE_SANIT_PROY, 1),
            format_currency_chilean(BASELINE_SANIT_PROY_CLP),
            "—",
        ),
        (
            f"Actualizado ({start.strftime('%d/%m')}–{end.strftime('%d/%m')})",
            str(n_dias),
            format_number_chilean(sanit["prom_dia"], 1),
            format_number_chilean(sanit["proy_m3"], 1),
            format_currency_chilean(sanit["proy_clp"]),
            f"{format_number_chilean(delta_prom, 1)} m³/día",
        ),
        (
            f"Post-mejoras (desde {datetime.strptime(MEJORAS_DESDE, '%Y-%m-%d').strftime('%d/%m')})",
            str(n_mejoras),
            format_number_chilean(sanit_mej_prom, 1),
            format_number_chilean(sanit_mej_proy_m3, 1),
            format_currency_chilean(sanit_mej_proy_clp),
            f"{format_number_chilean(delta_mej, 1)} m³/día",
        ),
        (
            "Referencia: factura julio (ticket)",
            "—",
            "—",
            format_number_chilean(FACTURA_M3, 0),
            format_currency_chilean(FACTURA_CLP),
            "histórica",
        ),
    ]
    add_table(doc, "Proyecciones Sanitaria", cmp_table, wes_style=True)

    # Día a día
    add_formatted_title(doc, "Consumo día a día — ¿bajó o aumentó?")
    intro_d = doc.add_paragraph(
        f"Referencia de la proyección previa: "
        f"{format_number_chilean(BASELINE_SANIT_PROM, 0)} m³/día en Sanitaria. "
        f"Se indica si cada día bajó o aumentó respecto a ese ritmo y respecto al día anterior. "
        f"También se muestra la suma de las tres salas de bomba del mismo día."
    )
    intro_d.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_black(intro_d)

    table_dias = [
        (
            "Fecha",
            "Sanitaria (m³)",
            "Vs proyección",
            "Δ vs proy. (m³)",
            "Vs día ant.",
            "Salas bomba (m³)",
        )
    ]
    for r in dia_rows:
        fecha_fmt = datetime.strptime(r["fecha"], "%Y-%m-%d").strftime("%d/%m/%Y")
        vs_prev_txt = r["var_prev"]
        if r["vs_prev"] is not None:
            vs_prev_txt = f"{r['var_prev']} ({format_number_chilean(r['vs_prev'], 1)})"
        table_dias.append(
            (
                fecha_fmt,
                format_number_chilean(r["sanitaria"], 1),
                r["var_base"],
                format_number_chilean(r["vs_base"], 1),
                vs_prev_txt,
                format_number_chilean(r["salas"], 1),
            )
        )
    add_table(doc, "Serie diaria Sanitaria", table_dias, wes_style=True)

    # Resumen conteo
    n_bajo = sum(1 for r in dia_rows if r["vs_base"] < -5)
    n_alto = sum(1 for r in dia_rows if r["vs_base"] > 5)
    n_est = n_dias - n_bajo - n_alto
    # Desde 28/07
    post = [r for r in dia_rows if r["fecha"] >= reciente_start]
    resumen_d = doc.add_paragraph()
    resumen_d.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    resumen_d.add_run("Lectura: ").bold = True
    resumen_d.add_run(
        f"de {n_dias} días completos, {n_bajo} quedaron bajo la proyección previa, "
        f"{n_alto} por encima y {n_est} estables (±5 m³). "
        f"El quiebre coincide con las mejoras/pérdidas corregidas: desde el 29/07 "
        f"({n_mejoras} días) Sanitaria promedió {format_number_chilean(sanit_mej_prom, 1)} m³/día "
        f"(frente a {format_number_chilean(BASELINE_SANIT_PROM, 0)} m³/día del baseline). "
        f"Mínimo del periodo: {format_number_chilean(min(r['sanitaria'] for r in dia_rows), 1)} m³ "
        f"({min(dia_rows, key=lambda x: x['sanitaria'])['fecha']}); "
        f"máximo: {format_number_chilean(max(r['sanitaria'] for r in dia_rows), 1)} m³ "
        f"({max(dia_rows, key=lambda x: x['sanitaria'])['fecha']})."
    )
    _set_black(resumen_d)

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_diario), Inches(6.2), keep_with_next=True)

    # Proyección actualizada por punto
    add_formatted_title(doc, "Proyección mensual actualizada (WES · días completos)")
    intro_p = doc.add_paragraph(
        f"Promedio diario de {n_dias} días completos "
        f"({start.strftime('%d/%m/%Y')}–{end.strftime('%d/%m/%Y')}) × 30. "
        f"Tarifa efectiva factura julio: ${format_number_chilean(PRECIO_M3, 0)} CLP/m³. "
        f"La factura ({format_number_chilean(FACTURA_M3, 0)} m³ / "
        f"{format_currency_chilean(FACTURA_CLP)}) es solo referencia histórica."
    )
    intro_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_black(intro_p)

    table_proy = [
        (
            "Punto",
            f"Consumo {n_dias} días (m³)",
            "Promedio diario (m³)",
            "Proyección 30 días (m³)",
            "Proyección 30 días (CLP)",
        )
    ]
    for r in sorted(rows, key=lambda x: x["proy_m3"], reverse=True):
        table_proy.append(
            (
                f"{r['nombre']} ({r['id']})",
                format_number_chilean(r["m3_periodo"], 1),
                format_number_chilean(r["prom_dia"], 1),
                format_number_chilean(r["proy_m3"], 1),
                format_currency_chilean(r["proy_clp"]),
            )
        )
    table_proy.append(
        (
            "Referencia: factura julio (histórica)",
            "—",
            "—",
            format_number_chilean(FACTURA_M3, 0),
            format_currency_chilean(FACTURA_CLP),
        )
    )
    add_table(doc, "Proyección / cuenta mensual actualizada", table_proy, wes_style=True)

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_proy), Inches(6.0), keep_with_next=True)

    # Participación
    add_formatted_title(doc, "Participación salas vs cuenta (proyección actualizada)")
    expl = doc.add_paragraph(
        "Sanitaria = 100% de la cuenta. Las salas son submediciones. "
        "El diferencial es volumen no monitoreado punto a punto. "
        f"En el baseline era ~{format_number_chilean(BASELINE_PCT_GAP, 1)}% no monitoreado; "
        f"con data actualizada queda en {format_number_chilean(pct_gap, 1)}%."
    )
    expl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_black(expl)

    part_rows = [
        ("Concepto", "Proy. 30 días (m³)", "Proy. 30 días (CLP)", "% sobre Sanitaria"),
        (
            "Medidor Principal Sanitaria (cuenta / 100%)",
            format_number_chilean(sanit["proy_m3"], 1),
            format_currency_chilean(sanit["proy_clp"]),
            "100,0%",
        ),
    ]
    for b in sorted(bombas, key=lambda x: x["proy_m3"], reverse=True):
        part_rows.append(
            (
                f"{b['nombre']} ({b['id']})",
                format_number_chilean(b["proy_m3"], 1),
                format_currency_chilean(b["proy_clp"]),
                f"{format_number_chilean(b['proy_m3'] / sanit['proy_m3'] * 100, 1)}%",
            )
        )
    part_rows.append(
        (
            "TOTAL salas de bomba (07+08+10)",
            format_number_chilean(salas_m3, 1),
            format_currency_chilean(salas_clp),
            f"{format_number_chilean(pct_salas, 1)}%",
        )
    )
    part_rows.append(
        (
            "NO MONITOREADO (Sanitaria − salas)",
            format_number_chilean(gap_m3, 1),
            format_currency_chilean(gap_clp),
            f"{format_number_chilean(pct_gap, 1)}%",
        )
    )
    add_table(doc, "Participación actualizada", part_rows, wes_style=True)

    nota_anom = doc.add_paragraph()
    nota_anom.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    nota_anom.add_run("Contexto operativo: ").bold = True
    # Check if salas rose while sanitaria fell
    base_dates = [d for d in all_dates if d <= "2026-07-27"]
    post_dates = reciente_dates
    salas_base = sum(
        sum(series[b["id"]].get(d, 0.0) for b in bombas) for d in base_dates
    ) / max(len(base_dates), 1)
    salas_post = sum(
        sum(series[b["id"]].get(d, 0.0) for b in bombas) for d in post_dates
    ) / max(len(post_dates), 1)
    nota_anom.add_run(
        f"según el encargado de operaciones, se ejecutaron mejoras y se "
        f"encontraron/corrigieron pérdidas. Eso explica la caída de Sanitaria "
        f"de ~{format_number_chilean(BASELINE_SANIT_PROM, 0)} a "
        f"~{format_number_chilean(sanit_mej_prom, 0)} m³/día. "
        f"Las salas de bomba internas pasaron de ~{format_number_chilean(salas_base, 1)} a "
        f"~{format_number_chilean(salas_post, 1)} m³/día (no acompañan la misma magnitud "
        f"de baja, lo esperable si la pérdida corregida estaba fuera de esas submediciones)."
    )
    _set_black(nota_anom)

    if chart_noct and target_dt:
        add_formatted_title(
            doc,
            f"Día con mayor consumo nocturno — Sanitaria ({target_dt.strftime('%d-%m-%y')})",
        )
        add_picture_with_pagination(doc, str(chart_noct), Inches(6.0), keep_with_next=True)

    # Conclusiones
    add_formatted_title(doc, "Conclusiones")
    bullets = [
        (
            f"La proyección del reporte 23–27/07 (~{format_number_chilean(BASELINE_SANIT_PROY, 0)} m³ "
            f"/ ~{format_currency_chilean(BASELINE_SANIT_PROY_CLP)}) "
            f"correspondía al ritmo con pérdidas y ya no aplica tras las mejoras."
        ),
        (
            f"Proyección de pago post-mejoras (tarifa ticket julio "
            f"${format_number_chilean(PRECIO_M3, 0)}/m³): "
            f"{format_number_chilean(sanit_mej_proy_m3, 1)} m³ / "
            f"{format_currency_chilean(sanit_mej_proy_clp)} a 30 días "
            f"(ahorro vs ticket julio: {format_currency_chilean(ahorro_vs_factura)}; "
            f"vs proyección previa: {format_currency_chilean(ahorro_vs_proy_prev)})."
        ),
        (
            f"Estimación cierre agosto: {format_number_chilean(cierre_ago_m3, 1)} m³ / "
            f"{format_currency_chilean(cierre_ago_clp)} "
            f"({dias_ago_completos} días reales + {dias_restantes} proyectados al ritmo post-mejoras)."
        ),
        (
            f"Día a día: quiebre desde el 28–29/07; tramo post-mejoras "
            f"{format_number_chilean(sanit_mej_prom, 1)} m³/día."
        ),
        (
            f"Participación salas: de {format_number_chilean(BASELINE_PCT_SALAS, 1)}% "
            f"a {format_number_chilean(pct_salas, 1)}% de la cuenta proyectada; "
            f"no monitoreado de {format_number_chilean(BASELINE_PCT_GAP, 1)}% a "
            f"{format_number_chilean(pct_gap, 1)}%."
        ),
    ]
    for b in bullets:
        bp = doc.add_paragraph(b, style="List Bullet")
        bp.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_black(bp)

    docx_name = (
        f"Seguimiento_Proyeccion_Bupa_Antofagasta_"
        f"{start.strftime('%Y%m%d')}_{end.strftime('%Y%m%d')}.docx"
    )
    docx_path = out_dir / docx_name
    doc.save(str(docx_path))

    pdf_path = convertir_word_a_pdf(docx_path)
    if not pdf_path or not Path(pdf_path).is_file():
        pdf_path = _convert_with_libreoffice(docx_path)
    if not pdf_path or not Path(pdf_path).is_file():
        pdf_path = None
        print("[ADVERTENCIA] PDF no generado (conversión no disponible en este entorno).")

    meta = {
        "out_dir": str(out_dir),
        "n_dias": n_dias,
        "sanit_prom": sanit["prom_dia"],
        "sanit_proy": sanit["proy_m3"],
        "sanit_proy_clp": sanit["proy_clp"],
        "reciente_prom": sanit_reciente_prom,
        "reciente_proy": sanit_reciente_proy,
        "mejoras_prom": sanit_mej_prom,
        "mejoras_proy_m3": sanit_mej_proy_m3,
        "mejoras_proy_clp": sanit_mej_proy_clp,
        "cierre_agosto_m3": cierre_ago_m3,
        "cierre_agosto_clp": cierre_ago_clp,
        "ahorro_vs_factura": ahorro_vs_factura,
        "ahorro_vs_proy_prev": ahorro_vs_proy_prev,
        "precio_m3": PRECIO_M3,
        "continua": continua,
        "pct_salas": pct_salas,
        "pct_gap": pct_gap,
        "n_bajo": n_bajo,
        "n_alto": n_alto,
        "start": start.strftime("%d/%m/%Y"),
        "end": end.strftime("%d/%m/%Y"),
        "dia_rows": dia_rows,
    }
    return docx_path, Path(pdf_path) if pdf_path else Path(), meta


def main() -> int:
    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass

    print("=" * 70)
    print("[INFO] Bupa Antofagasta — seguimiento proyección vs data actualizada")
    print("=" * 70)
    docx_path, pdf_path, meta = build_report()
    print(f"[OK] Word: {docx_path}")
    if pdf_path and pdf_path.is_file():
        print(f"[OK] PDF:  {pdf_path}")
    print(
        f"     Periodo: {meta['start']} – {meta['end']} ({meta['n_dias']} días)"
    )
    print(
        f"     Sanitaria prom: {meta['sanit_prom']:.1f} m3/dia -> "
        f"proy {meta['sanit_proy']:.1f} m3 ({format_currency_chilean(meta['sanit_proy_clp'])})"
    )
    print(
        f"     Reciente prom: {meta['reciente_prom']:.1f} m3/dia -> "
        f"proy {meta['reciente_proy']:.1f} m3"
    )
    print(
        f"     Post-mejoras: {meta['mejoras_prom']:.1f} m3/dia -> "
        f"proy {meta['mejoras_proy_m3']:.1f} m3 "
        f"({format_currency_chilean(meta['mejoras_proy_clp'])})"
    )
    print(
        f"     Cierre agosto: {meta['cierre_agosto_m3']:.1f} m3 "
        f"({format_currency_chilean(meta['cierre_agosto_clp'])})"
    )
    print(
        f"     Ahorro vs factura julio: {format_currency_chilean(meta['ahorro_vs_factura'])} | "
        f"vs proy previa: {format_currency_chilean(meta['ahorro_vs_proy_prev'])}"
    )
    print(f"     Continua proyeccion previa: {'SI' if meta['continua'] else 'NO'}")
    print(f"     Dias bajo baseline: {meta['n_bajo']} | sobre: {meta['n_alto']}")
    print(
        f"     Participacion salas {meta['pct_salas']:.1f}% | "
        f"gap {meta['pct_gap']:.1f}%"
    )

    # Subir a Drive
    try:
        from wes_google_drive import credenciales_configuradas, subir_a_drive

        if credenciales_configuradas():
            sub = "Bupa_Antofagasta/ABREGADO"
            for path in (docx_path, pdf_path):
                if path and path.is_file():
                    info = subir_a_drive(path, subcarpeta=sub)
                    print(f"[DRIVE] {info['name']}")
                    print(f"        {info['web_view_link']}")
                    meta.setdefault("drive_links", []).append(info)
        else:
            print("[ADVERTENCIA] Sin credenciales GOOGLE_DRIVE_*")
    except Exception as exc:
        print(f"[ADVERTENCIA] Drive: {exc}")

    # Persist meta for summary
    meta_path = Path(meta["out_dir"]) / "meta_seguimiento.json"
    import json

    dump = {k: v for k, v in meta.items() if k != "dia_rows"}
    dump["drive_links"] = meta.get("drive_links", [])
    dump["dias"] = [
        {
            "fecha": r["fecha"],
            "sanitaria": round(r["sanitaria"], 2),
            "salas": round(r["salas"], 2),
            "var_base": r["var_base"],
            "vs_base": round(r["vs_base"], 2),
            "var_prev": r["var_prev"],
            "vs_prev": None if r["vs_prev"] is None else round(r["vs_prev"], 2),
        }
        for r in meta["dia_rows"]
    ]
    meta_path.write_text(json.dumps(dump, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"[OK] Meta: {meta_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
