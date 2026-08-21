"""
Regulación Renca desde el lunes 24/08/2026 (vuelve el periodo con control).

Objetivo: ~10 % de rendimiento vs el perfil sin WES (17–20 ago).
Programación WES en tres niveles sobre un caudal máximo (100 %):
  100 % = tope m³/h a cargar en el horario
   60 % = 0,60 × ese tope
   30 % = 0,30 × ese tope

Semáforo (vs el pico del día de línea base):
  ≥ 74 % del pico → 100 %   |  40–74 % → 60 %   |  < 40 % → 30 %

Uso:
  python generar_regulacion_renca_lunes24_rendimiento10.py
"""

from __future__ import annotations

import argparse
import math
import statistics
import subprocess
import shutil
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple
from zoneinfo import ZoneInfo

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from auditoria_cpa_icco_renca_grafico import _vector_m3h_24_desde_api
from generar_reporte_word import add_logo_to_header, format_number_chilean

ROOT = Path(__file__).resolve().parent
TZ_CL = ZoneInfo("America/Santiago")
OBJETIVO_PCT = 10.0  # rendimiento = (sin − con) / sin
FACTOR_OBJ = 1.0 - OBJETIVO_PCT / 100.0
RATIO_100 = 0.74
RATIO_60 = 0.40
COLOR_HEAD = RGBColor(31, 71, 136)
FILL_100 = PatternFill("solid", fgColor="C6EFCE")
FILL_60 = PatternFill("solid", fgColor="FFF2CC")
FILL_30 = PatternFill("solid", fgColor="FFC7CE")
COLOR_100 = "#548235"
COLOR_60 = "#BF8F00"
COLOR_30 = "#C0504D"

PUNTOS_TODOS: Tuple[Tuple[str, str], ...] = (
    ("000017-07", "ICCP (Cumbre de Cóndores pte.)"),
    ("000017-04", "Esc. Lo Velásquez"),
    ("000017-05", "Gimnasio municipal"),
    ("000017-06", "Piscina municipal"),
)
PUNTOS_ESCUELA_GIMNASIO: Tuple[Tuple[str, str], ...] = (
    ("000017-04", "Esc. Lo Velásquez"),
    ("000017-05", "Gimnasio municipal"),
)
DIAS_SIN = (date(2026, 8, 17), date(2026, 8, 18), date(2026, 8, 19), date(2026, 8, 20))
LUNES_CONTROL = date(2026, 8, 24)


MARGEN_ABAST = 1.10  # 100 % = pico medido × 1,10 para que el abastecimiento no falle


def _m3_a_lmin(m3h: float) -> float:
    return float(m3h) * 1000.0 / 60.0


def _lmin_a_m3(lmin: float) -> float:
    return float(lmin) * 60.0 / 1000.0


def _ceil_lmin(x: float) -> float:
    """Redondeo operativo a 0,5 L/min hacia arriba."""
    if x <= 0:
        return 0.0
    return math.ceil(x * 2.0 - 1e-9) / 2.0


def _fmt_rango(h0: int, h1: int) -> str:
    return f"{h0:02d}:00–{h1:02d}:59"


def _rangos_nivel(niveles: Sequence[int]) -> List[Tuple[int, int, int]]:
    """Bloques consecutivos (hora_ini, hora_fin inclusive, % )."""
    out: List[Tuple[int, int, int]] = []
    i = 0
    while i < len(niveles):
        j = i
        while j + 1 < len(niveles) and niveles[j + 1] == niveles[i]:
            j += 1
        out.append((i, j, int(niveles[i])))
        i = j + 1
    return out


def _suavizar_niveles(niveles: Sequence[int], sin_h: Sequence[float], lmin_de: Dict[int, float]) -> List[int]:
    """Bloques operativos: no dejar un agujero de 30/60 entre horas 100 %, ni un tope justo al caudal medido."""
    n = [int(x) for x in niveles]
    for h in range(24):
        dem = _m3_a_lmin(sin_h[h])
        while n[h] < 100 and dem * MARGEN_ABAST > lmin_de[n[h]] + 1e-9:
            n[h] = 60 if n[h] == 30 else 100
    # Valles de 1–2 h entre niveles más altos → subir (p. ej. 10 h a 100, 11 h a 30, 12 h a 100).
    changed = True
    while changed:
        changed = False
        i = 0
        while i < 24:
            j = i
            while j + 1 < 24 and n[j + 1] == n[i]:
                j += 1
            if i > 0 and j < 23 and n[i] < n[i - 1] and n[i] < n[j + 1] and (j - i + 1) <= 2:
                nuevo = min(n[i - 1], n[j + 1])
                if nuevo > n[i]:
                    for k in range(i, j + 1):
                        n[k] = nuevo
                    changed = True
            i = j + 1
    # Entre dos tramos 100 % separados por ≤ 3 h en 60 %, unificar a 100 % (programable).
    picos = [h for h, v in enumerate(n) if v == 100]
    for a, b in zip(picos, picos[1:]):
        gap = list(range(a + 1, b))
        if 0 < len(gap) <= 3 and all(n[k] >= 60 for k in gap):
            for k in range(a, b + 1):
                n[k] = 100
    return n


def _nivel(v: float, pico: float) -> int:
    if pico <= 1e-9 or v <= 1e-9:
        return 30
    r = v / pico
    if r >= RATIO_100:
        return 100
    if r >= RATIO_60:
        return 60
    return 30


def _consumo_esperado(q100: float, sin_h: Sequence[float], fracs: Sequence[float]) -> float:
    return float(sum(min(float(sin_h[h]), q100 * fracs[h]) for h in range(24)))


def _q100_para_objetivo(sin_h: Sequence[float], fracs: Sequence[float], factor: float) -> float:
    """Busca el 100 % (m³/h) tal que el tope 30/60/100 deje factor × Σ sin WES."""
    base = float(sum(sin_h))
    if base <= 1e-9:
        return 0.0
    target = factor * base
    lo, hi = 0.0, max(max(sin_h), 0.01) * 8.0
    # Si ni con tope enorme se llega al target (el patrón 30/60 recorta de más), devolver hi.
    if _consumo_esperado(hi, sin_h, fracs) + 1e-9 < target:
        return hi
    for _ in range(48):
        mid = (lo + hi) / 2.0
        got = _consumo_esperado(mid, sin_h, fracs)
        if got >= target:
            hi = mid
        else:
            lo = mid
    return hi


def _vector_mediana(vecs: Sequence[Sequence[float]]) -> List[float]:
    out: List[float] = []
    for h in range(24):
        vals = [float(v[h]) for v in vecs]
        out.append(float(statistics.median(vals)))
    return out


def _set_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(
        parse_xml(
            "<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
            f'w:val="clear" w:fill="{hex_fill}"/>'
        )
    )


def _grafico_dia(
    sin_h: Sequence[float],
    cap_h: Sequence[float],
    niveles: Sequence[int],
    titulo: str,
    out_png: Path,
) -> None:
    x = list(range(24))
    colors = [{100: COLOR_100, 60: COLOR_60, 30: COLOR_30}[n] for n in niveles]
    fig, ax = plt.subplots(figsize=(12.2, 4.8))
    ax.bar(
        [i - 0.18 for i in x],
        [_m3_a_lmin(v) for v in sin_h],
        width=0.36,
        color="#C0504D",
        label="Sin WES (L/min medido)",
        zorder=2,
    )
    ax.bar(
        [i + 0.18 for i in x],
        [_m3_a_lmin(v) for v in cap_h],
        width=0.36,
        color=colors,
        label="Tope 30 / 60 / 100 % (L/min)",
        zorder=2,
    )
    ax.set_xticks(x)
    ax.set_xticklabels([f"{h:02d}" for h in x], fontsize=7)
    ax.set_ylabel("L/min")
    ax.set_title(titulo, fontweight="bold", fontsize=11, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=8, loc="upper right")
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _pdf(docx: Path) -> Optional[Path]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx.parent),
            str(docx),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=180,
    )
    p = docx.with_suffix(".pdf")
    return p if p.is_file() else None


def _plan_para_perfil(sin_h: List[float]) -> dict:
    """100 % = pico × margen (L/min). 30/60 solo si el caudal medido ya cabe: no falla el abastecimiento."""
    pico = max(sin_h) if sin_h else 0.0
    q100_lmin = _ceil_lmin(_m3_a_lmin(pico) * MARGEN_ABAST)
    q60_lmin = _ceil_lmin(q100_lmin * 0.60)
    q30_lmin = _ceil_lmin(q100_lmin * 0.30)
    q100 = _lmin_a_m3(q100_lmin)
    lmin_de = {100: q100_lmin, 60: q60_lmin, 30: q30_lmin}
    niveles: List[int] = []
    for h in range(24):
        n = _nivel(sin_h[h], pico)
        dem = _m3_a_lmin(sin_h[h])
        while n < 100 and dem > lmin_de[n] + 1e-9:
            n = 60 if n == 30 else 100
        niveles.append(n)
    niveles = _suavizar_niveles(niveles, sin_h, lmin_de)
    cap = [_lmin_a_m3(lmin_de[n]) for n in niveles]
    fracs = [n / 100.0 for n in niveles]
    esperado = float(sum(min(float(sin_h[h]), cap[h]) for h in range(24)))
    base = float(sum(sin_h))
    rend = ((base - esperado) / base * 100.0) if base > 1e-9 else 0.0
    rangos = _rangos_nivel(niveles)
    return {
        "sin": sin_h,
        "pico": pico,
        "niveles": niveles,
        "q100": q100,
        "q60": _lmin_a_m3(q60_lmin),
        "q30": _lmin_a_m3(q30_lmin),
        "q100_lmin": q100_lmin,
        "q60_lmin": q60_lmin,
        "q30_lmin": q30_lmin,
        "cap": cap,
        "base": base,
        "esperado": esperado,
        "rend": rend,
        "rangos": rangos,
    }


def _excel(
    path: Path,
    puntos: Sequence[Tuple[str, str]],
    planes_lun: Dict[str, dict],
    planes_tipo: Dict[str, dict],
) -> None:
    wb = Workbook()
    thin = Border(
        left=Side(style="thin", color="B0B0B0"),
        right=Side(style="thin", color="B0B0B0"),
        top=Side(style="thin", color="B0B0B0"),
        bottom=Side(style="thin", color="B0B0B0"),
    )
    head = Font(bold=True, color="FFFFFF", name="Calibri")
    fill_h = PatternFill("solid", fgColor="1F4788")

    def _hdr(ws, row, vals):
        for i, v in enumerate(vals, 1):
            c = ws.cell(row, i, v)
            c.font = head
            c.fill = fill_h
            c.alignment = Alignment(wrap_text=True, horizontal="center")
            c.border = thin

    ws = wb.active
    ws.title = "Resumen_100_60_30"
    _hdr(
        ws,
        1,
        [
            "Punto",
            "Perfil",
            "Σ sin WES (m³)",
            "Σ con tope (m³)",
            "Rendimiento %",
            "100 % (m³/h)",
            "60 % (m³/h)",
            "30 % (m³/h)",
            "100 % (L/min)",
            "60 % (L/min)",
            "30 % (L/min)",
            "Horas 100",
            "Horas 60",
            "Horas 30",
        ],
    )
    r = 2
    for nid, nom in puntos:
        for etiqueta, plan in (("Lunes 24 (base lun 17)", planes_lun[nid]), ("Día tipo lun–jue (mediana)", planes_tipo[nid])):
            n = plan["niveles"]
            ws.cell(r, 1, nom)
            ws.cell(r, 2, etiqueta)
            ws.cell(r, 3, round(plan["base"], 2))
            ws.cell(r, 4, round(plan["esperado"], 2))
            ws.cell(r, 5, round(plan["rend"], 1))
            ws.cell(r, 6, round(plan["q100"], 3))
            ws.cell(r, 7, round(plan["q60"], 3))
            ws.cell(r, 8, round(plan["q30"], 3))
            ws.cell(r, 9, plan["q100_lmin"])
            ws.cell(r, 10, plan["q60_lmin"])
            ws.cell(r, 11, plan["q30_lmin"])
            ws.cell(r, 12, n.count(100))
            ws.cell(r, 13, n.count(60))
            ws.cell(r, 14, n.count(30))
            for c in range(1, 15):
                ws.cell(r, c).border = thin
            r += 1

    ws2 = wb.create_sheet("Horario_lunes_24")
    _hdr(
        ws2,
        1,
        ["Hora"]
        + [f"{nom} % / m³/h / L·min" for _, nom in puntos],
    )
    # Better: one block per point
    ws2.delete_rows(1)
    col = 1
    for nid, nom in puntos:
        plan = planes_lun[nid]
        ws2.cell(1, col, nom).font = head
        ws2.cell(1, col).fill = fill_h
        ws2.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + 3)
        for i, h in enumerate(["Hora", "% WES", "Tope m³/h", "Tope L/min"]):
            c = ws2.cell(2, col + i, h)
            c.font = head
            c.fill = fill_h
            c.border = thin
        for h in range(24):
            rr = 3 + h
            n = plan["niveles"][h]
            fill = {100: FILL_100, 60: FILL_60, 30: FILL_30}[n]
            vals = [f"{h:02d}:00–{h:02d}:59", n, round(plan["cap"][h], 3), round(_m3_a_lmin(plan["cap"][h]), 2)]
            for i, v in enumerate(vals):
                c = ws2.cell(rr, col + i, v)
                c.border = thin
                c.fill = fill
        col += 5

    ws3 = wb.create_sheet("Horario_dia_tipo")
    col = 1
    for nid, nom in puntos:
        plan = planes_tipo[nid]
        ws3.cell(1, col, nom).font = head
        ws3.cell(1, col).fill = fill_h
        ws3.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + 3)
        for i, h in enumerate(["Hora", "% WES", "Tope m³/h", "Tope L/min"]):
            c = ws3.cell(2, col + i, h)
            c.font = head
            c.fill = fill_h
            c.border = thin
        for h in range(24):
            rr = 3 + h
            n = plan["niveles"][h]
            fill = {100: FILL_100, 60: FILL_60, 30: FILL_30}[n]
            vals = [f"{h:02d}:00–{h:02d}:59", n, round(plan["cap"][h], 3), round(_m3_a_lmin(plan["cap"][h]), 2)]
            for i, v in enumerate(vals):
                c = ws3.cell(rr, col + i, v)
                c.border = thin
                c.fill = fill
        col += 5

    ws_r = wb.create_sheet("Rangos_Lmin")
    _hdr(ws_r, 1, ["Punto", "Horario", "% WES", "L/min a cargar", "Cubre el pico del bloque"])
    rr = 2
    for nid, nom in puntos:
        plan = planes_lun[nid]
        lmap = {100: plan["q100_lmin"], 60: plan["q60_lmin"], 30: plan["q30_lmin"]}
        for h0, h1, niv in plan["rangos"]:
            pico_b = max(_m3_a_lmin(plan["sin"][h]) for h in range(h0, h1 + 1))
            cubre = "sí" if lmap[niv] + 1e-9 >= pico_b else "revisar"
            fill = {100: FILL_100, 60: FILL_60, 30: FILL_30}[niv]
            vals = [nom, _fmt_rango(h0, h1), niv, lmap[niv], cubre]
            for i, v in enumerate(vals, 1):
                c = ws_r.cell(rr, i, v)
                c.border = thin
                c.fill = fill
            rr += 1

    ws4 = wb.create_sheet("Criterio")
    lines = [
        ("Cifras", "Todo el plan se carga en L/min (no m³/h)."),
        ("Inicio control", "Lunes 24/08/2026"),
        ("Línea base", "Sin WES 17–20/08/2026 (hora Chile). Lunes 24 usa el perfil del lunes 17."),
        ("100 %", "Pico medido del lunes 17 × 1,10, redondeado a 0,5 L/min. El abastecimiento del pico no falla."),
        ("60 % / 30 %", "0,60 y 0,30 de ese 100 %. Solo en tramos donde el lunes 17 ya cabía, con 10 % de holgura."),
        ("Rangos", "Horas consecutivas del mismo L/min. Valles cortos entre picos se suben para no cortar el servicio."),
        ("Rendimiento 10 %", "No se recorta el caudal del lunes 17. El 10 % aparece si el lunes 24 se pasa de estos topes (fuga o uso extra)."),
    ]
    for i, (a, b) in enumerate(lines, 1):
        ws4.cell(i, 1, a).font = Font(bold=True)
        ws4.cell(i, 2, b)
    ws4.column_dimensions["A"].width = 18
    ws4.column_dimensions["B"].width = 92

    for wsx in wb.worksheets:
        for col in wsx.columns:
            letter = get_column_letter(col[0].column)
            wsx.column_dimensions[letter].width = max(12, min(28, wsx.column_dimensions[letter].width or 12))
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def _word(
    path: Path,
    puntos: Sequence[Tuple[str, str]],
    hora_corte: int,
    pngs: Dict[str, Path],
    planes_lun: Dict[str, dict],
    planes_tipo: Dict[str, dict],
) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    try:
        add_logo_to_header(doc)
    except Exception:
        pass
    tit = "Lo Velásquez y Gimnasio — L/min por horario (lunes 24)"
    if len(puntos) != 2:
        tit = "Renca — L/min por horario (lunes 24)"
    h = doc.add_heading(tit, level=0)
    if h.runs:
        h.runs[0].font.color.rgb = COLOR_HEAD
    p = doc.add_paragraph()
    p.add_run("Cifras en litros por minuto. ").bold = True
    p.add_run(
        "El 100 % es el pico del lunes 17 sin WES con 10 % de holgura, para que no falle el abastecimiento. "
        "60 % y 30 % solo en tramos donde ese día el caudal ya era menor. "
        f"Jueves 20 de la línea base hasta {hora_corte:02d}:59 Chile."
    )

    doc.add_heading("Caudales a cargar en WES", level=1)
    tbl = doc.add_table(rows=1 + len(puntos), cols=4)
    tbl.style = "Table Grid"
    headers = ["Punto", "100 % (L/min)", "60 % (L/min)", "30 % (L/min)"]
    for j, hd in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = hd
        _set_shading(cell, "1F4788")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    for i, (nid, nom) in enumerate(puntos, start=1):
        pl = planes_lun[nid]
        vals = [
            nom,
            format_number_chilean(pl["q100_lmin"], 1),
            format_number_chilean(pl["q60_lmin"], 1),
            format_number_chilean(pl["q30_lmin"], 1),
        ]
        for j, v in enumerate(vals):
            tbl.rows[i].cells[j].text = v
            for run in tbl.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(12)
        _set_shading(tbl.rows[i].cells[1], "C6EFCE")
        _set_shading(tbl.rows[i].cells[2], "FFF2CC")
        _set_shading(tbl.rows[i].cells[3], "FFC7CE")

    for nid, nom in puntos:
        pl = planes_lun[nid]
        doc.add_heading(nom, level=1)
        doc.add_paragraph(
            f"Pico lun 17: {format_number_chilean(_m3_a_lmin(pl['pico']), 1)} L/min. "
            f"Cargar estos L/min por tramo. El lunes 17 cabe entero en estos topes "
            f"(recorte vs ese día: {format_number_chilean(pl['rend'], 1)} %)."
        )
        if f"lun_{nid}" in pngs:
            doc.add_picture(str(pngs[f"lun_{nid}"]), width=Cm(16.0))
        rangos = pl["rangos"]
        t = doc.add_table(rows=1 + len(rangos), cols=3)
        t.style = "Table Grid"
        for j, hd in enumerate(["Horario", "L/min", "% WES"]):
            cell = t.rows[0].cells[j]
            cell.text = hd
            _set_shading(cell, "1F4788")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        lmap = {100: pl["q100_lmin"], 60: pl["q60_lmin"], 30: pl["q30_lmin"]}
        hexf = {100: "C6EFCE", 60: "FFF2CC", 30: "FFC7CE"}
        for i, (h0, h1, niv) in enumerate(rangos, start=1):
            vals = [_fmt_rango(h0, h1), format_number_chilean(lmap[niv], 1), str(niv)]
            for j, v in enumerate(vals):
                t.rows[i].cells[j].text = v
                _set_shading(t.rows[i].cells[j], hexf[niv])
                for run in t.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(12)

    doc.add_heading("Notas", level=1)
    doc.add_paragraph(
        "En cada tramo el L/min cubre el caudal medido del lunes 17 más 10 % de holgura: no se corta el servicio de ese día. "
        "El 10 % de rendimiento no sale recortando el pico. Sale si el lunes 24 se pasa de estos L/min (fuga o uso extra). "
        "Si un tramo 30 o 60 se queda corto en el día, subir ese tramo a 100 %."
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    ap = argparse.ArgumentParser(description="Regulación Renca lunes 24 — 10 % de rendimiento")
    ap.add_argument(
        "--solo-escuela-gimnasio",
        action="store_true",
        help="Solo Esc. Lo Velásquez y Gimnasio municipal.",
    )
    args = ap.parse_args()
    puntos: Sequence[Tuple[str, str]] = (
        PUNTOS_ESCUELA_GIMNASIO if args.solo_escuela_gimnasio else PUNTOS_TODOS
    )

    ahora = datetime.now(TZ_CL)
    hora_corte = max(0, min(23, ahora.hour - 1)) if ahora.date() == date(2026, 8, 20) else 23
    if ahora.date() > date(2026, 8, 20):
        hora_corte = 23
    print(f"Chile {ahora:%Y-%m-%d %H:%M} | jueves 20 hasta {hora_corte:02d}:59")
    print("Puntos:", ", ".join(n for _, n in puntos))

    jobs = [(nid, d) for nid, _ in puntos for d in DIAS_SIN]
    vecs: Dict[Tuple[str, date], List[float]] = {}

    def _job(item: Tuple[str, date]) -> Tuple[Tuple[str, date], List[float]]:
        nid, d = item
        v = _vector_m3h_24_desde_api(nid, d)
        if d == date(2026, 8, 20) and hora_corte < 23:
            v = [v[h] if h <= hora_corte else 0.0 for h in range(24)]
        return item, v

    print(f"Descargando {len(jobs)} series horarias…")
    with ThreadPoolExecutor(max_workers=8) as ex:
        futs = [ex.submit(_job, it) for it in jobs]
        for i, fut in enumerate(as_completed(futs), 1):
            k, v = fut.result()
            vecs[k] = v
            if i % 4 == 0 or i == len(jobs):
                print(f"  {i}/{len(jobs)}")

    planes_lun: Dict[str, dict] = {}
    planes_tipo: Dict[str, dict] = {}
    for nid, nom in puntos:
        lun = vecs[(nid, date(2026, 8, 17))]
        tipo_src = [vecs[(nid, d)] for d in DIAS_SIN]
        tipo = _vector_mediana(tipo_src)
        planes_lun[nid] = _plan_para_perfil(lun)
        planes_tipo[nid] = _plan_para_perfil(tipo)
        pl = planes_lun[nid]
        print(
            f"  {nom}: 100%={pl['q100_lmin']:.1f} L/min | 60%={pl['q60_lmin']:.1f} | "
            f"30%={pl['q30_lmin']:.1f} | pico lun17={_m3_a_lmin(pl['pico']):.1f} L/min"
        )
        for h0, h1, niv in pl["rangos"]:
            lmin = {100: pl["q100_lmin"], 60: pl["q60_lmin"], 30: pl["q30_lmin"]}[niv]
            print(f"     {_fmt_rango(h0, h1)}  →  {lmin:.1f} L/min  ({niv} %)")

    ts = ahora.strftime("%Y%m%d_%H%M")
    suf = "escuela_gimnasio_" if args.solo_escuela_gimnasio else ""
    out_dir = ROOT / "reports" / "reporte de auditoria" / f"regulacion_renca_{suf}lunes24_rend10_{ts}"
    png_dir = out_dir / "graficos"
    png_dir.mkdir(parents=True, exist_ok=True)
    pngs: Dict[str, Path] = {}
    for nid, nom in puntos:
        p = png_dir / f"lunes24_{nid}.png"
        pl = planes_lun[nid]
        _grafico_dia(
            pl["sin"],
            pl["cap"],
            pl["niveles"],
            f"{nom} — lunes 24: L/min medido lun 17 vs tope 30/60/100",
            p,
        )
        pngs[f"lun_{nid}"] = p

    stem = "Regulacion_LoVelasquez_Gimnasio_lunes24_rendimiento10" if args.solo_escuela_gimnasio else "Regulacion_Renca_lunes24_rendimiento10"
    xlsx = out_dir / f"{stem}.xlsx"
    _excel(xlsx, puntos, planes_lun, planes_tipo)
    docx = out_dir / f"{stem}_{ts}.docx"
    _word(docx, puntos, hora_corte, pngs, planes_lun, planes_tipo)
    pdf = _pdf(docx)
    print(f"XLSX {xlsx}")
    print(f"DOCX {docx}")
    print(f"PDF  {pdf or '(no convertido)'}")
    return 0


if __name__ == "__main__":
    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
            sys.stderr.reconfigure(encoding="utf-8")
        except Exception:
            pass
    raise SystemExit(main())
