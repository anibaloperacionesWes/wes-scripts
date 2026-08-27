"""
Informe de cumplimiento hídrico según planilla de habilitación / corte.

Lee el Google Sheet «Horarios de control hídrico clientes WES», cruza con la
serie horaria de la API WES y clasifica cada punto (cumple / no cumple).

Reglas especiales:
- Derco matriz (000012-06): las ventanas listadas son mínimo nocturno de
  guardias; el resto del día es habilitación. Se alerta si el mínimo nocturno
  supera el histórico en más de 25 %.
- GYM Renca (000017-05): en corte nocturno hay un mínimo de guardias; misma
  regla del 25 % sobre el histórico.
- Tiempo de cerrado / retraso de corte: horas desde el inicio programado de
  corte hasta que el caudal cae a ~0 (p. ej. ICCO).
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple
from zoneinfo import ZoneInfo

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook

from control_nocturno import _compactar_horas_a_rangos, obtener_datos_horarios_dia
from generar_reporte_word import (
    add_formatted_heading,
    add_logo_to_header,
    estilizar_tabla_wes,
)
from wes_estilo_graficos_app import COLOR_BARRA_WES, COLOR_CONSUMO, COLOR_NOCHE

CHILE = ZoneInfo("America/Santiago")
ROOT = Path(__file__).resolve().parent
SHEET_ID = "1eM03xh4Pmqx5YTTKWiOku8_NT0115ausiWOKNaLc8Lk"
UMBRAL_CERO = 0.05  # m³/h: ruido de medidor
MULT_HIST = 1.25
DIAS_HISTORICO = 14
MAX_WORKERS = 12

DIAS_SEMANA = {
    "LUNES": 0,
    "MARTES": 1,
    "MIERCOLES": 2,
    "MIÉRCOLES": 2,
    "JUEVES": 3,
    "VIERNES": 4,
    "SABADO": 5,
    "SÁBADO": 5,
    "DOMINGO": 6,
}

# (subcadena normalizada del nombre en la planilla) -> (nodeId, cliente, nombre oficial)
SITE_MAP: List[Tuple[str, str, str, str]] = [
    ("antonio hermidas", "000008-01", "CORMUP", "Antonio Hermida Fabres"),
    ("carlos fernandes", "000008-03", "CORMUP", "Carlos Fernández Peña"),
    ("tobalaba", "000008-04", "CORMUP", "Tobalaba"),
    ("santa maria", "000008-05", "CORMUP", "Santa María"),
    ("luis arrieta", "000008-06", "CORMUP", "Luis Arrieta Cañas"),
    ("erasmo", "000008-07", "CORMUP", "Erasmo Escala"),
    ("juan bautista", "000008-09", "CORMUP", "Juan Bautista Pastene"),
    ("matilde", "000008-10", "CORMUP", "Matilde Huici Navas"),
    ("valle germoso", "000008-11", "CORMUP", "CE Valle Hermoso"),
    ("valle hermoso", "000008-11", "CORMUP", "CE Valle Hermoso"),
    ("union nacional", "000008-12", "CORMUP", "Unión Nacional Árabe"),
    ("juan pablo segundo", "000008-14", "CORMUP", "Juan Pablo II"),
    ("pae estanquer", "000025-01", "PARQUE ARAUCO", "Estanque Norte Locales"),
    ("mae", "000025-19", "PARQUE ARAUCO", "Sala de Bomba Estanque Sur"),
    ("pizza hut", "000025-07", "PARQUE ARAUCO", "Pizza Hut"),
    ("alto cordillera", "000028-01", "LA FLORIDA", "Liceo Alto Cordillera"),
    ("san ignacio 500", "000025-18", "PARQUE ARAUCO", "San Ignacio 500"),
    ("alexander fleming", "000022-00", "LAS CONDES", "Alexander Fleming"),
    ("derco matriz", "000012-06", "DERCO", "Quilicura Matriz Principal"),
    ("las tarrias", "000006-01", "PROVIDENCIA", "Liceo Lastarria"),
    ("carmela", "000006-02", "PROVIDENCIA", "Carmela Carvajal"),
    ("liceo 7", "000006-04", "PROVIDENCIA", "Liceo 7 Luisa Saavedra"),
    ("juan pablo duarte", "000006-05", "PROVIDENCIA", "Liceo Juan Pablo Duarte"),
    ("lo valledor", "000002-01", "LO VALLEDOR", "Lo Valledor P1"),
    ("club hause", "000021-01", "CDUC", "Club House CDUC"),
    ("club house", "000021-01", "CDUC", "Club House CDUC"),
    ("tupper", "000021-03", "CDUC", "Raimundo Tupper"),
    ("agunsa", "000020-02", "AGUNSA", "Módulo D"),
    ("lo velazques", "000017-04", "RENCA", "Escuela Lo Velázquez"),
    ("gym renca", "000017-05", "RENCA", "Gimnasio Renca"),
    ("picina", "000017-06", "RENCA", "Piscina Municipal Renca"),
    ("piscina", "000017-06", "RENCA", "Piscina Municipal Renca"),
    ("iccp", "000017-07", "RENCA", "ICCP (Cumbre de Cóndores Pte.)"),
    ("icco", "000017-08", "RENCA", "Colegio ICCO Renca"),
    ("hostos", "000024-01", "LA REINA", "Eugenio María de Hostos"),
]

NODOS_MIN_NOCTURNO = {"000012-06", "000017-05"}  # Derco matriz, GYM Renca
NODO_DERCO = "000012-06"
NODO_GYM = "000017-05"

RE_DIA = re.compile(
    r"#{2,}\s*(LUNES|MARTES|MIERCOLES|MIÉRCOLES|JUEVES|VIERNES|SABADO|SÁBADO|DOMINGO)\s*#{2,}",
    re.IGNORECASE,
)
RE_EVENTO = re.compile(
    r"(\d+)\s*=\s*(\d{1,2})\s*:\s*(\d{2})\s*-*>\s*(?:\d+\s+)?(\d{1,2})\s*:\s*(\d{2})\s*-\s*(ALTA|BAJA)",
    re.IGNORECASE,
)
RE_EV_HEADER = re.compile(r"EV\.\s*INICIAL", re.IGNORECASE)


def _norm(text: object) -> str:
    s = str(text or "").strip().lower()
    repl = (
        ("á", "a"),
        ("é", "e"),
        ("í", "i"),
        ("ó", "o"),
        ("ú", "u"),
        ("ñ", "n"),
        ("ü", "u"),
    )
    for a, b in repl:
        s = s.replace(a, b)
    return re.sub(r"\s+", " ", s)


def mapear_sitio(nombre: str) -> Optional[Tuple[str, str, str]]:
    n = _norm(nombre)
    if not n or n.startswith("#") or RE_EV_HEADER.search(n) or RE_EVENTO.search(n):
        return None
    for key, node_id, cliente, oficial in SITE_MAP:
        if key in n:
            return node_id, cliente, oficial
    return None


def _hhmm_a_min(h: int, m: int) -> int:
    return max(0, min(h, 23)) * 60 + max(0, min(m, 59))


@dataclass
class Ventana:
    start_min: int  # [0, 1440)
    end_min: int  # exclusive; 1440 = fin de día
    estado: str  # ALTA / BAJA
    valvula: str

    def cubre_minuto(self, minuto: int) -> bool:
        if self.start_min < self.end_min:
            return self.start_min <= minuto < self.end_min
        # cruza medianoche: no debería ocurrir si partimos ventanas, pero por si acaso
        return minuto >= self.start_min or minuto < self.end_min


@dataclass
class Sitio:
    nombre_planilla: str
    node_id: str
    cliente: str
    nombre: str
    por_dia: Dict[int, List[Ventana]] = field(default_factory=dict)  # 0=lun … 6=dom
    especial: str = ""  # derco | gym | ""


def descargar_planilla(destino: Path) -> Path:
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    import io

    creds = Credentials(
        token=None,
        refresh_token=os.environ["GOOGLE_DRIVE_REFRESH_TOKEN"],
        token_uri="https://oauth2.googleapis.com/token",
        client_id=os.environ["GOOGLE_DRIVE_CLIENT_ID"],
        client_secret=os.environ["GOOGLE_DRIVE_CLIENT_SECRET"],
        scopes=["https://www.googleapis.com/auth/drive"],
    )
    creds.refresh(Request())
    service = build("drive", "v3", credentials=creds, cache_discovery=False)
    request = service.files().export_media(
        fileId=SHEET_ID,
        mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    destino.parent.mkdir(parents=True, exist_ok=True)
    with destino.open("wb") as fh:
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
    return destino


def _partir_ventana(h1: int, m1: int, h2: int, m2: int) -> List[Tuple[int, int]]:
    """Devuelve tramos [start, end) en minutos del día, partiendo si cruza 24:00."""
    a = _hhmm_a_min(h1, m1)
    b = _hhmm_a_min(h2, m2)
    if h2 == 23 and m2 >= 50:
        b = 1440
    if a == b:
        return []
    if a < b:
        return [(a, b)]
    return [(a, 1440), (0, b)]


def parsear_planilla(path: Path) -> List[Sitio]:
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    max_r, max_c = ws.max_row, ws.max_column

    # Celdas (row, col) -> texto
    celdas: Dict[Tuple[int, int], str] = {}
    for r in range(1, max_r + 1):
        for c in range(1, max_c + 1):
            v = ws.cell(r, c).value
            if v is not None and str(v).strip():
                celdas[(r, c)] = str(v).strip()

    # Cabeceras de sitio: celdas que mapean a un nodo y no son eventos/días
    cabeceras: List[Tuple[int, int, Sitio]] = []
    vistos: set[str] = set()
    for (r, c), txt in sorted(celdas.items()):
        mapped = mapear_sitio(txt)
        if not mapped:
            continue
        node_id, cliente, oficial = mapped
        if node_id in vistos:
            continue
        # Evitar que un evento accidental coincida (no debería)
        if RE_EVENTO.search(txt) or RE_DIA.search(txt) or RE_EV_HEADER.search(txt):
            continue
        vistos.add(node_id)
        especial = ""
        if node_id == NODO_DERCO:
            especial = "derco"
        elif node_id == NODO_GYM:
            especial = "gym"
        cabeceras.append(
            (
                r,
                c,
                Sitio(
                    nombre_planilla=txt.strip(),
                    node_id=node_id,
                    cliente=cliente,
                    nombre=oficial,
                    especial=especial,
                ),
            )
        )

    # Para cada sitio, leer hacia abajo la misma columna hasta el siguiente
    # encabezado de esa columna (o fin), recolectando 7 días.
    cab_por_col: Dict[int, List[Tuple[int, Sitio]]] = {}
    for r, c, sitio in cabeceras:
        cab_por_col.setdefault(c, []).append((r, sitio))
    for c in cab_por_col:
        cab_por_col[c].sort()

    sitios: List[Sitio] = []
    for c, lista in cab_por_col.items():
        for i, (r0, sitio) in enumerate(lista):
            r_fin = lista[i + 1][0] if i + 1 < len(lista) else max_r + 1
            dia_actual: Optional[int] = None
            r = r0 + 1
            # No cruzar el siguiente encabezado de la misma columna (evita
            # mezclar Alto Cordillera con Agunsa, MAE con Club House, etc.).
            while r < r_fin and r <= max_r:
                txt = celdas.get((r, c), "")
                m_dia = RE_DIA.search(txt) if txt else None
                if m_dia:
                    key = m_dia.group(1).upper().replace("É", "E").replace("Á", "A")
                    dia_actual = DIAS_SEMANA[key]
                    sitio.por_dia.setdefault(dia_actual, [])
                    r += 1
                    continue
                m_ev = RE_EVENTO.search(txt) if txt else None
                if m_ev and dia_actual is not None:
                    valvula, h1, m1, h2, m2, estado = m_ev.groups()
                    for a, b in _partir_ventana(int(h1), int(m1), int(h2), int(m2)):
                        sitio.por_dia.setdefault(dia_actual, []).append(
                            Ventana(a, b, estado.upper(), valvula)
                        )
                r += 1
            sitios.append(sitio)
    sitios.sort(key=lambda s: (s.cliente, s.nombre))
    return sitios


def ventanas_habilitacion(sitio: Sitio, dia: date) -> List[Ventana]:
    return list(sitio.por_dia.get(dia.weekday(), []))


def minutos_habilitados(sitio: Sitio, dia: date) -> List[bool]:
    """Vector de 1440 minutos: True = habilitación (o mínimo Derco)."""
    flags = [False] * 1440
    for v in ventanas_habilitacion(sitio, dia):
        for m in range(v.start_min, v.end_min):
            if 0 <= m < 1440:
                flags[m] = True
    if sitio.especial == "derco":
        # Las ventanas listadas son el mínimo nocturno; el resto es habilitación.
        # Se invierte: True = agua habilitada (día) o mínimo (noche listada).
        # Para corte "estricto" Derco no tiene corte: todo el día hay agua
        # (mínimo o habilitación). Se marca todo True.
        return [True] * 1440
    return flags


def horas_en_corte(sitio: Sitio, dia: date) -> List[int]:
    """Horas 0-23 cuyo bloque [h:00, h+1:00) está mayoritariamente en corte."""
    flags = minutos_habilitados(sitio, dia)
    horas: List[int] = []
    for h in range(24):
        chunk = flags[h * 60 : (h + 1) * 60]
        habilitados = sum(1 for x in chunk if x)
        if habilitados < 30:
            horas.append(h)
    return horas


def horas_minimo_nocturno(sitio: Sitio, dia: date) -> List[int]:
    if sitio.especial == "gym":
        return horas_en_corte(sitio, dia)
    if sitio.especial != "derco":
        return []
    flags = [False] * 1440
    for v in ventanas_habilitacion(sitio, dia):
        for m in range(v.start_min, min(v.end_min, 1440)):
            flags[m] = True
    return [h for h in range(24) if sum(flags[h * 60 : (h + 1) * 60]) >= 30]


def texto_habilitacion(sitio: Sitio, dia: date) -> str:
    vents = ventanas_habilitacion(sitio, dia)
    if sitio.especial == "derco":
        if not vents:
            return "Habilitación 24 h (mín. nocturno no declarado)"
        partes = []
        for v in vents:
            partes.append(
                f"{v.start_min // 60:02d}:{v.start_min % 60:02d}-"
                f"{v.end_min // 60:02d}:{v.end_min % 60:02d} mín. {v.estado}"
            )
        return "Día habilitado; noche " + "; ".join(partes)
    if not vents:
        return "Sin habilitación (corte 24 h)"
    partes = []
    for v in vents:
        partes.append(
            f"{v.start_min // 60:02d}:{v.start_min % 60:02d}-"
            f"{(v.end_min // 60) if v.end_min < 1440 else 24:02d}:"
            f"{v.end_min % 60 if v.end_min < 1440 else 0:02d} {v.estado}"
        )
    return "; ".join(partes)


def texto_corte(sitio: Sitio, dia: date) -> str:
    if sitio.especial == "derco":
        mins = horas_minimo_nocturno(sitio, dia)
        if not mins:
            return "Sin corte (mín. nocturno guardias)"
        return "Sin corte; mín. guardias " + _compactar_horas_a_rangos(mins)
    horas = horas_en_corte(sitio, dia)
    if not horas:
        return "Sin corte (habilitado 24 h)"
    extra = " (mín. guardias)" if sitio.especial == "gym" else ""
    return _compactar_horas_a_rangos(horas) + extra


def _inicio_corte_minutos(sitio: Sitio, dia: date) -> Optional[int]:
    """Primer minuto de un tramo de corte (el más cercano a la tarde/noche)."""
    if sitio.especial == "derco":
        return None
    flags = minutos_habilitados(sitio, dia)
    # Preferir el tramo de corte que empieza después de las 12:00
    for start in range(12 * 60, 1440):
        if not flags[start] and (start == 0 or flags[start - 1]):
            return start
    for start in range(0, 12 * 60):
        if not flags[start] and (start == 0 or flags[start - 1]):
            return start
    return None


@dataclass
class ResultadoSitio:
    sitio: Sitio
    dia: date
    horas_corte: List[int]
    hourly_ayer: Dict[int, float]
    hourly_hoy: Dict[int, float]
    max_corte: float
    horas_con_consumo: List[int]
    cumple: bool
    estado_actual: str
    tiempo_cerrado: str
    retraso: str
    observacion: str
    hist_mediana: Optional[float] = None
    hist_actual: Optional[float] = None
    hist_alerta: bool = False


def _hora_chile_ahora() -> datetime:
    return datetime.now(CHILE)


def clasificar_cerrado(
    sitio: Sitio,
    dia: date,
    hourly_ayer: Dict[int, float],
    hourly_hoy: Dict[int, float],
    ahora: datetime,
) -> Tuple[str, str, str]:
    """
    Retorna (estado_actual, tiempo_cerrado, retraso).

    Recorre horas desde el inicio de corte de ayer/hoy hasta ahora.
    """
    hora_now = ahora.hour if ahora.date() == dia else 23
    # Serie (día_offset, hora) -> m3  con offset 0=hoy, -1=ayer
    serie: List[Tuple[int, int, float]] = []
    for h in range(24):
        serie.append((-1, h, float(hourly_ayer.get(h, 0.0))))
    for h in range(hora_now + 1):
        serie.append((0, h, float(hourly_hoy.get(h, 0.0))))

    def en_corte(offset: int, h: int) -> bool:
        d = dia + timedelta(days=offset)
        if sitio.especial == "derco":
            return False
        return h in set(horas_en_corte(sitio, d))

    # Estado actual
    h_act = ahora.hour
    cons_act = float(hourly_hoy.get(h_act, 0.0))
    corte_act = en_corte(0, h_act)
    if sitio.especial == "derco":
        if h_act in set(horas_minimo_nocturno(sitio, dia)):
            estado = "Mínimo nocturno (guardias)"
        else:
            estado = "Habilitado"
    elif corte_act:
        estado = "Cortado (programado)" if cons_act <= UMBRAL_CERO else "En corte, aún con caudal"
    else:
        estado = "Habilitado" if cons_act > UMBRAL_CERO else "Habilitado, caudal ~0 (posible válvula cerrada)"

    # Retraso: desde inicio de corte de AYER (ventana nocturna) hasta primera hora ~0
    ini = _inicio_corte_minutos(sitio, dia - timedelta(days=1))
    retraso_txt = "—"
    cerrado_at: Optional[Tuple[int, int]] = None  # (offset, hour)
    h_ini: Optional[int] = None
    if ini is not None:
        h_ini = ini // 60
        for off, h, v in serie:
            if off == -1 and h < h_ini:
                continue
            if en_corte(off, h) and v <= UMBRAL_CERO:
                cerrado_at = (off, h)
                break
        if cerrado_at is not None and h_ini is not None:
            off_c, h_c = cerrado_at
            delta_h = h_c - h_ini
            if off_c == 0:
                delta_h += 24
            if delta_h < 0:
                delta_h += 24
            if delta_h <= 0:
                retraso_txt = f"Cerró a las {h_c:02d}:00 (sin retraso)"
            else:
                retraso_txt = f"{delta_h} h (cerró {h_c:02d}:00; programado {h_ini:02d}:00)"
        elif horas_en_corte(sitio, dia - timedelta(days=1)):
            retraso_txt = "No cerró (caudal en toda la ventana de corte)"

    # Tiempo cerrado: racha ~0 de la última ventana de corte ya transcurrida.
    racha = 0
    for off, h, v in reversed(serie):
        if v <= UMBRAL_CERO:
            racha += 1
        else:
            break
    if corte_act and racha > 0:
        idx = len(serie) - racha
        off0, h0, _ = serie[idx]
        etiqueta_dia = "hoy" if off0 == 0 else "ayer"
        cerrado_txt = f"{racha} h (desde {h0:02d}:00 {etiqueta_dia})"
    else:
        i = len(serie) - 1
        while i >= 0 and not en_corte(serie[i][0], serie[i][1]):
            i -= 1
        fin_i = i
        while i >= 0 and en_corte(serie[i][0], serie[i][1]) and serie[i][2] <= UMBRAL_CERO:
            i -= 1
        start_i = i + 1
        if fin_i >= 0 and start_i <= fin_i and serie[start_i][2] <= UMBRAL_CERO:
            dur = fin_i - start_i + 1
            h0 = serie[start_i][1]
            h1 = (serie[fin_i][1] + 1) % 24
            if h1 == 0:
                h1_lab = "24:00"
            else:
                h1_lab = f"{h1:02d}:00"
            cerrado_txt = f"Cerrado {dur} h ({h0:02d}:00–{h1_lab}); ahora habilitado"
        elif racha == 0:
            cerrado_txt = "Abierto (hay caudal)"
        else:
            idx = len(serie) - racha
            off0, h0, _ = serie[idx]
            etiqueta_dia = "hoy" if off0 == 0 else "ayer"
            cerrado_txt = f"{racha} h (desde {h0:02d}:00 {etiqueta_dia})"

    return estado, cerrado_txt, retraso_txt


def mediana(vals: Sequence[float]) -> float:
    if not vals:
        return 0.0
    a = sorted(float(x) for x in vals)
    n = len(a)
    if n % 2:
        return a[n // 2]
    return 0.5 * (a[n // 2 - 1] + a[n // 2])


def evaluar_sitio(
    sitio: Sitio,
    dia: date,
    hourly_ayer: Dict[int, float],
    hourly_hoy: Dict[int, float],
    hist_horas: Optional[List[float]] = None,
) -> ResultadoSitio:
    ahora = _hora_chile_ahora()
    horas_c = horas_en_corte(sitio, dia)
    # Ventana de corte que cruza medianoche: unir corte de ayer (noche) + hoy (madrugada)
    horas_ayer_noche = [h for h in horas_en_corte(sitio, dia - timedelta(days=1)) if h >= 18]
    max_corte = 0.0
    horas_con: List[int] = []

    def considera_corte(h: int, v: float, es_hoy: bool) -> None:
        nonlocal max_corte
        if v > max_corte:
            max_corte = v
        if v > UMBRAL_CERO:
            horas_con.append(h if es_hoy else h)  # se compacta después

    if sitio.especial in {"derco", "gym"}:
        horas_min = horas_minimo_nocturno(sitio, dia)
        vals_noche = [float(hourly_hoy.get(h, 0.0)) for h in horas_min]
        # si la ventana cruza medianoche, sumar ayer
        if sitio.especial == "derco":
            horas_min_ayer = horas_minimo_nocturno(sitio, dia - timedelta(days=1))
            vals_noche += [
                float(hourly_ayer.get(h, 0.0)) for h in horas_min_ayer if h >= 22
            ]
        actual = mediana(vals_noche) if vals_noche else 0.0
        hist_med = mediana(hist_horas or []) if hist_horas else 0.0
        alerta = False
        obs_parts: List[str] = []
        if hist_med > UMBRAL_CERO:
            limite = hist_med * MULT_HIST
            alerta = actual > limite + 1e-9
            obs_parts.append(
                f"Mín. nocturno actual {actual:.2f} m³/h vs histórico {hist_med:.2f} "
                f"(umbral +25 % = {limite:.2f})"
            )
            if alerta:
                obs_parts.append("SUPERA 25 % del histórico")
            else:
                obs_parts.append("Dentro del histórico +25 %")
        else:
            # Histórico ~0: tratar como corte clásico
            alerta = actual > UMBRAL_CERO
            obs_parts.append(
                f"Histórico nocturno ~0; actual {actual:.2f} m³/h"
            )
        cumple = not alerta
        max_corte = max(vals_noche) if vals_noche else 0.0
        horas_con = [
            h
            for h, v in ((h, float(hourly_hoy.get(h, 0.0))) for h in horas_min)
            if v > (hist_med * MULT_HIST if hist_med > UMBRAL_CERO else UMBRAL_CERO)
        ]
        estado, cerrado, retraso = clasificar_cerrado(
            sitio, dia, hourly_ayer, hourly_hoy, ahora
        )
        return ResultadoSitio(
            sitio=sitio,
            dia=dia,
            horas_corte=horas_min,
            hourly_ayer=hourly_ayer,
            hourly_hoy=hourly_hoy,
            max_corte=max_corte,
            horas_con_consumo=horas_con,
            cumple=cumple,
            estado_actual=estado,
            tiempo_cerrado=cerrado,
            retraso=retraso,
            observacion="; ".join(obs_parts),
            hist_mediana=hist_med,
            hist_actual=actual,
            hist_alerta=alerta,
        )

    sin_prog = not any(sitio.por_dia.values())
    if sin_prog:
        estado, cerrado, retraso = clasificar_cerrado(
            sitio, dia, hourly_ayer, hourly_hoy, ahora
        )
        return ResultadoSitio(
            sitio=sitio,
            dia=dia,
            horas_corte=[],
            hourly_ayer=hourly_ayer,
            hourly_hoy=hourly_hoy,
            max_corte=0.0,
            horas_con_consumo=[],
            cumple=True,
            estado_actual="Sin programación en planilla",
            tiempo_cerrado=cerrado,
            retraso="—",
            observacion="La planilla no trae tramos de habilitación para este punto; no se evalúa corte.",
        )

    # Puntos estándar: consumo en horas de corte. Se tolera un retraso de cierre
    # de hasta GRACE_HORAS (la válvula tarda en cortar; caso típico ICCO).
    GRACE_HORAS = 2
    ini_min = _inicio_corte_minutos(sitio, dia - timedelta(days=1))
    h_ini_ayer = (ini_min // 60) if ini_min is not None else None

    def _dentro_gracia(offset: int, h: int) -> bool:
        if h_ini_ayer is None:
            return False
        if offset == -1:
            return 0 <= (h - h_ini_ayer) < GRACE_HORAS
        if h_ini_ayer == 0:
            return h < GRACE_HORAS
        return (24 - h_ini_ayer) + h < GRACE_HORAS

    for h in horas_ayer_noche:
        v = float(hourly_ayer.get(h, 0.0))
        if v > max_corte:
            max_corte = v
        if v > UMBRAL_CERO and not _dentro_gracia(-1, h):
            horas_con.append(h)
    for h in horas_c:
        v = float(hourly_hoy.get(h, 0.0))
        if v > max_corte:
            max_corte = v
        if v > UMBRAL_CERO and not _dentro_gracia(0, h):
            horas_con.append(h)

    cumple = len(horas_con) == 0
    estado, cerrado, retraso = clasificar_cerrado(
        sitio, dia, hourly_ayer, hourly_hoy, ahora
    )
    if cumple and "h (cerró" in retraso:
        obs = f"Cumple corte, con retraso de cierre. {retraso}"
    elif cumple:
        obs = "Cumple corte (caudal ~0)"
    else:
        obs = "Consumo en ventana de corte: " + _compactar_horas_a_rangos(sorted(set(horas_con)))
    if not horas_c and not horas_ayer_noche:
        obs = "Habilitado 24 h (no aplica corte)"
        cumple = True
    return ResultadoSitio(
        sitio=sitio,
        dia=dia,
        horas_corte=horas_c,
        hourly_ayer=hourly_ayer,
        hourly_hoy=hourly_hoy,
        max_corte=max_corte,
        horas_con_consumo=sorted(set(horas_con)),
        cumple=cumple,
        estado_actual=estado,
        tiempo_cerrado=cerrado,
        retraso=retraso,
        observacion=obs,
    )


def fetch_hourly(node_id: str, dia: date) -> Dict[int, float]:
    dt = datetime.combine(dia, datetime.min.time())
    try:
        return obtener_datos_horarios_dia(node_id, dt)
    except Exception:
        return {h: 0.0 for h in range(24)}


def recolectar_historico_nocturno(
    sitio: Sitio, dia: date, n_dias: int = DIAS_HISTORICO
) -> List[float]:
    vals: List[float] = []
    for i in range(1, n_dias + 1):
        d = dia - timedelta(days=i)
        hourly = fetch_hourly(sitio.node_id, d)
        for h in horas_minimo_nocturno(sitio, d):
            vals.append(float(hourly.get(h, 0.0)))
    return vals


def grafico_anillo(total_ok: int, total_no: int, path: Path) -> Path:
    fig, ax = plt.subplots(figsize=(5.4, 4.4))
    vals = [total_no, total_ok]
    colores = ["#E74C3C", "#2ECC71"]
    etiquetas = ["No cumplen", "Cumplen"]
    total = total_ok + total_no
    ax.pie(
        vals,
        labels=etiquetas,
        colors=colores,
        autopct=lambda pct: f"{pct:.0f}%\n({int(round(pct * total / 100.0))})" if total else "0",
        startangle=90,
        pctdistance=0.75,
        textprops={"fontsize": 9},
    )
    fig.gca().add_artist(plt.Circle((0, 0), 0.55, fc="white"))
    ax.set_title("Cumplimiento de corte / mínimo nocturno", fontsize=11)
    ax.text(0, 0, f"Total\n{total}", ha="center", va="center", fontsize=11, weight="bold")
    ax.axis("equal")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def grafico_24h(
    hourly: Dict[int, float],
    horas_corte: List[int],
    titulo: str,
    path: Path,
) -> Path:
    x = np.arange(24, dtype=float)
    y = np.array([float(hourly.get(h, 0.0)) for h in range(24)], dtype=float)
    corte = set(horas_corte)
    colors = [COLOR_NOCHE if h in corte else COLOR_BARRA_WES for h in range(24)]
    fig, ax = plt.subplots(figsize=(10.2, 3.8))
    ax.bar(x, y, width=0.78, color=colors, edgecolor="white", linewidth=0.4, zorder=2)
    for h in sorted(corte):
        ax.axvspan(h - 0.5, h + 0.5, alpha=0.10, color=COLOR_NOCHE, zorder=0)
    ax.plot(x, y, color=COLOR_CONSUMO, linewidth=1.1, marker="o", markersize=3, zorder=4)
    ax.set_xlim(-0.5, 23.5)
    ax.set_xticks(x)
    ax.set_xticklabels([f"{h:02d}" for h in range(24)], fontsize=7)
    ax.set_ylabel("m³/h")
    ax.set_title(titulo, fontsize=11, fontweight="bold")
    ax.grid(axis="y", linestyle="-", alpha=0.3)
    ax.set_axisbelow(True)
    fig.text(
        0.5,
        0.02,
        "Rojo = ventana de corte (o mínimo nocturno). Azul = habilitación.",
        ha="center",
        fontsize=7.5,
        color="#444",
    )
    fig.subplots_adjust(bottom=0.16, left=0.08, right=0.98, top=0.88)
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=160)
    plt.close(fig)
    return path


def _set_cell(cell, text: str, *, bold: bool = False, size: int = 8, color: Optional[RGBColor] = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def llenar_tabla(table, rows: List[List[str]], highlight: Optional[set[int]] = None) -> None:
    highlight = highlight or set()
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            if i == 0:
                _set_cell(cell, val, bold=True, size=8, color=RGBColor(255, 255, 255))
            else:
                col = RGBColor(180, 0, 0) if i in highlight else None
                _set_cell(cell, val, bold=i in highlight, size=7, color=col)
    estilizar_tabla_wes(table, highlight_rows=sorted(highlight), has_total_row=False)


def convertir_pdf(docx_path: Path) -> Optional[Path]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                str(docx_path.parent),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return None
    return pdf_path if pdf_path.exists() else None


def crear_docx(
    resultados: List[ResultadoSitio],
    out_path: Path,
    charts_dir: Path,
    dia: date,
    ahora: datetime,
) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.5)
    add_logo_to_header(doc)

    title = doc.add_heading("INFORME DE CUMPLIMIENTO HÍDRICO", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = sub.add_run(
        "Habilitación programada vs corte · planilla de control hídrico WES"
    )
    r.italic = True
    r.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.add_run(
        f"Generado: {ahora.strftime('%d-%m-%Y %H:%M')} hora Chile  ·  "
        f"Día evaluado: {dia.strftime('%A %d-%m-%Y')}  ·  "
        f"Fuente: Horarios de control hídrico clientes WES"
    ).font.size = Pt(9)

    analizados = [r for r in resultados if r.sitio.node_id]
    ok = [r for r in analizados if r.cumple]
    no = [r for r in analizados if not r.cumple]
    add_formatted_heading(doc, "Resumen ejecutivo", 1)
    p = doc.add_paragraph()
    p.add_run(
        f"Se contrastó el caudal horario de {len(analizados)} puntos con la programación "
        f"de habilitación de agua de la planilla (el resto del día se considera corte). "
        f"Umbral de «cerrado»: ≤ {UMBRAL_CERO:.2f} m³/h. Se tolera un retraso de cierre "
        f"de hasta 2 horas (electrovalvula).\n"
    )
    p.add_run(f"Puntos que cumplen: {len(ok)}\n").bold = True
    p.add_run(f"Puntos que no cumplen: {len(no)}\n").bold = True
    p.add_run(
        "Derco Matriz Principal y Gimnasio Renca se evalúan por mínimo nocturno de "
        "guardias: se alerta si el caudal nocturno supera en más de 25 % la mediana "
        f"histórica de {DIAS_HISTORICO} días.\n"
    )
    p.add_run(
        "El «tiempo de cerrado» es la racha de horas consecutivas con caudal ~0 hasta "
        "ahora. El «retraso de corte» es la diferencia entre la hora programada de "
        "cierre (inicio de corte de la noche anterior) y la primera hora efectivamente "
        "en cero (caso típico: ICCO)."
    )

    anillo = grafico_anillo(len(ok), len(no), charts_dir / "anillo_cumplimiento.png")
    doc.add_picture(str(anillo), width=Inches(4.6))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Tabla principal
    add_formatted_heading(doc, "Tabla de cumplimiento", 1)
    headers = [
        "Cliente",
        "Punto",
        "ID WES",
        "Habilitación hoy",
        "Ventana de corte",
        "Estado actual",
        "¿Cumple?",
        "Máx. en corte (m³/h)",
        "Tiempo cerrado",
        "Retraso de corte",
        "Observación",
    ]
    table_rows = [headers]
    highlight: set[int] = set()
    for idx, res in enumerate(sorted(analizados, key=lambda x: (not x.cumple, x.sitio.cliente, x.sitio.nombre)), start=1):
        table_rows.append(
            [
                res.sitio.cliente,
                res.sitio.nombre,
                res.sitio.node_id,
                texto_habilitacion(res.sitio, dia),
                texto_corte(res.sitio, dia),
                res.estado_actual,
                "SÍ" if res.cumple else "NO",
                f"{res.max_corte:.2f}".replace(".", ","),
                res.tiempo_cerrado,
                res.retraso,
                res.observacion,
            ]
        )
        if not res.cumple:
            highlight.add(idx)

    table = doc.add_table(rows=len(table_rows), cols=len(headers))
    table.style = "Table Grid"
    llenar_tabla(table, table_rows, highlight)

    nota = doc.add_paragraph()
    nota.add_run(
        "Nota: en Derco la noche listada (23:00–04:30 ALTA) es el mínimo de guardias, "
        "no un corte a cero. En GYM Renca el corte nocturno (fuera de 08:00–22:00) "
        "conserva un mínimo de guardias; ambos se comparan contra el histórico +25 %."
    ).italic = True
    nota.runs[0].font.size = Pt(8)
    nota.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    # Sección especiales
    especiales = [r for r in analizados if r.sitio.node_id in NODOS_MIN_NOCTURNO]
    if especiales:
        add_formatted_heading(doc, "Mínimo nocturno Derco y GYM Renca", 1, page_break_before=True)
        doc.add_paragraph(
            "Se calculó la mediana de las horas de mínimo nocturno de los últimos "
            f"{DIAS_HISTORICO} días (excluyendo hoy) y se comparó con la mediana de "
            "la última noche. Alerta si actual > histórico × 1,25."
        )
        h2 = ["Punto", "ID", "Mediana histórica (m³/h)", "Mediana última noche", "Límite +25 %", "¿Alerta?"]
        rows2 = [h2]
        hi2: set[int] = set()
        for i, res in enumerate(especiales, start=1):
            hist = res.hist_mediana or 0.0
            act = res.hist_actual or 0.0
            rows2.append(
                [
                    res.sitio.nombre,
                    res.sitio.node_id,
                    f"{hist:.3f}".replace(".", ","),
                    f"{act:.3f}".replace(".", ","),
                    f"{hist * MULT_HIST:.3f}".replace(".", ","),
                    "SÍ, supera 25 %" if res.hist_alerta else "No",
                ]
            )
            if res.hist_alerta:
                hi2.add(i)
        t2 = doc.add_table(rows=len(rows2), cols=len(h2))
        t2.style = "Table Grid"
        llenar_tabla(t2, rows2, hi2)

        for res in especiales:
            png = grafico_24h(
                res.hourly_hoy,
                res.horas_corte,
                f"{res.sitio.nombre} ({res.sitio.node_id}) — {dia.strftime('%d-%m-%Y')}",
                charts_dir / f"perfil_{res.sitio.node_id.replace('-', '_')}.png",
            )
            doc.add_paragraph("")
            doc.add_picture(str(png), width=Inches(9.2))

    # No cumplen detalle
    if no:
        add_formatted_heading(doc, "Detalle de puntos que no cumplen", 1, page_break_before=True)
        for res in no:
            add_formatted_heading(doc, f"{res.sitio.nombre} ({res.sitio.node_id})", 2)
            doc.add_paragraph(
                f"Cliente: {res.sitio.cliente}. Corte hoy: {texto_corte(res.sitio, dia)}. "
                f"Máximo en corte: {res.max_corte:.2f} m³/h. "
                f"Retraso: {res.retraso}. Tiempo cerrado: {res.tiempo_cerrado}."
            )
            png = grafico_24h(
                res.hourly_hoy,
                res.horas_corte,
                f"{res.sitio.nombre} — perfil {dia.strftime('%d-%m-%Y')}",
                charts_dir / f"nc_{res.sitio.node_id.replace('-', '_')}.png",
            )
            doc.add_picture(str(png), width=Inches(9.2))

    # ICCO destacado
    icco = next((r for r in analizados if r.sitio.node_id == "000017-08"), None)
    if icco:
        add_formatted_heading(doc, "ICCO — tiempo de cerrado", 1)
        doc.add_paragraph(
            f"Programación de hoy: {texto_habilitacion(icco.sitio, dia)}. "
            f"Ventana de corte: {texto_corte(icco.sitio, dia)}. "
            f"Estado actual: {icco.estado_actual}. "
            f"Tiempo cerrado: {icco.tiempo_cerrado}. "
            f"Retraso de corte (noche anterior): {icco.retraso}."
        )
        png = grafico_24h(
            icco.hourly_hoy,
            icco.horas_corte,
            f"ICCO Renca — {dia.strftime('%d-%m-%Y')}",
            charts_dir / "icco_perfil.png",
        )
        doc.add_picture(str(png), width=Inches(9.2))

    add_formatted_heading(doc, "Criterios", 1)
    doc.add_paragraph(
        "1) Habilitación = tramos ALTA/BAJA de electroválvula en la planilla para el "
        "día de la semana. Fuera de esos tramos = corte.\n"
        "2) Un punto «cumple» si, fuera de las primeras 2 h de retraso de válvula, "
        "el caudal en corte es ≤ 0,05 m³/h. El retraso y el tiempo cerrado se informan "
        "igual (caso ICCO: programado 20:00, cierra ~22:00).\n"
        "3) Derco: no hay corte a cero; la noche 23:00–04:30 es mínimo de guardias y el "
        "resto del día es habilitación. Alerta solo si el mínimo supera el histórico +25 %.\n"
        "4) GYM Renca: habilitación 08:00–22:00 (BAJA); de noche hay mínimo de guardias "
        "con la misma regla del 25 %.\n"
        "5) Retraso de corte: horas desde el inicio programado de corte de la noche "
        "anterior hasta la primera hora con caudal ~0."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Informe de cumplimiento hídrico")
    parser.add_argument("--fecha", default="", help="YYYY-MM-DD (default: hoy Chile)")
    parser.add_argument("--excel", default="", help="Ruta local al xlsx (salta Drive)")
    args = parser.parse_args()

    ahora = _hora_chile_ahora()
    dia = (
        datetime.strptime(args.fecha, "%Y-%m-%d").date()
        if args.fecha.strip()
        else ahora.date()
    )
    print(f"[INFO] Fecha Chile: {dia}  ahora={ahora.strftime('%Y-%m-%d %H:%M')}")

    reports = ROOT / "reports" / "control_nocturno"
    reports.mkdir(parents=True, exist_ok=True)
    xlsx = Path(args.excel) if args.excel.strip() else reports / "HORARIOS_CONTROL_HIDRICO.xlsx"
    if args.excel.strip():
        print(f"[INFO] Excel local: {xlsx}")
    else:
        print("[INFO] Descargando planilla de Google Drive…")
        descargar_planilla(xlsx)
        print(f"[OK] Planilla: {xlsx} ({xlsx.stat().st_size} bytes)")

    sitios = parsear_planilla(xlsx)
    print(f"[INFO] Sitios mapeados: {len(sitios)}")
    for s in sitios:
        dias = ",".join(str(d) for d in sorted(s.por_dia))
        print(f"  {s.node_id:10} {s.cliente:16} {s.nombre:32} días={dias} esp={s.especial or '-'}")

    ayer = dia - timedelta(days=1)
    hourly_cache: Dict[Tuple[str, date], Dict[int, float]] = {}

    def need(node_id: str, d: date) -> None:
        hourly_cache[(node_id, d)] = {}  # placeholder

    jobs: List[Tuple[str, date]] = []
    for s in sitios:
        jobs.append((s.node_id, dia))
        jobs.append((s.node_id, ayer))
        if s.node_id in NODOS_MIN_NOCTURNO:
            for i in range(1, DIAS_HISTORICO + 1):
                jobs.append((s.node_id, dia - timedelta(days=i)))
    jobs = list(dict.fromkeys(jobs))
    print(f"[INFO] Descargando {len(jobs)} series horarias…")

    def _job(item: Tuple[str, date]) -> Tuple[Tuple[str, date], Dict[int, float]]:
        nid, d = item
        return item, fetch_hourly(nid, d)

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futs = [ex.submit(_job, j) for j in jobs]
        for i, fut in enumerate(as_completed(futs), 1):
            key, data = fut.result()
            hourly_cache[key] = data
            if i % 20 == 0 or i == len(jobs):
                print(f"  {i}/{len(jobs)}")

    resultados: List[ResultadoSitio] = []
    for s in sitios:
        hoy_h = hourly_cache.get((s.node_id, dia), {h: 0.0 for h in range(24)})
        ayer_h = hourly_cache.get((s.node_id, ayer), {h: 0.0 for h in range(24)})
        hist_vals: Optional[List[float]] = None
        if s.node_id in NODOS_MIN_NOCTURNO:
            hist_vals = []
            for i in range(1, DIAS_HISTORICO + 1):
                d = dia - timedelta(days=i)
                hourly = hourly_cache.get((s.node_id, d), {})
                for h in horas_minimo_nocturno(s, d):
                    hist_vals.append(float(hourly.get(h, 0.0)))
        res = evaluar_sitio(s, dia, ayer_h, hoy_h, hist_vals)
        resultados.append(res)
        flag = "OK" if res.cumple else "NO"
        print(
            f"  [{flag}] {s.node_id} {s.nombre}: max_corte={res.max_corte:.2f} "
            f"estado={res.estado_actual} cerrado={res.tiempo_cerrado} retraso={res.retraso}"
        )

    ts = ahora.strftime("%Y%m%d_%H%M")
    out_dir = reports / f"cumplimiento_hidrico_{dia.strftime('%Y%m%d')}_{ts}"
    charts = out_dir / "charts"
    charts.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f"Informe_Cumplimiento_Hidrico_{dia.strftime('%Y%m%d')}.docx"
    crear_docx(resultados, docx_path, charts, dia, ahora)
    print(f"[OK] DOCX {docx_path}")
    pdf = convertir_pdf(docx_path)
    if pdf:
        print(f"[OK] PDF  {pdf}")
    else:
        print("[AVISO] No se pudo convertir a PDF")

    # CSV resumen
    import csv

    csv_path = out_dir / f"cumplimiento_hidrico_{dia.strftime('%Y%m%d')}.csv"
    with csv_path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(
            [
                "cliente",
                "punto",
                "node_id",
                "habilitacion_hoy",
                "ventana_corte",
                "estado_actual",
                "cumple",
                "max_corte_m3h",
                "tiempo_cerrado",
                "retraso_corte",
                "observacion",
                "hist_mediana",
                "hist_actual",
                "hist_alerta",
            ]
        )
        for res in resultados:
            w.writerow(
                [
                    res.sitio.cliente,
                    res.sitio.nombre,
                    res.sitio.node_id,
                    texto_habilitacion(res.sitio, dia),
                    texto_corte(res.sitio, dia),
                    res.estado_actual,
                    "SI" if res.cumple else "NO",
                    f"{res.max_corte:.3f}",
                    res.tiempo_cerrado,
                    res.retraso,
                    res.observacion,
                    "" if res.hist_mediana is None else f"{res.hist_mediana:.3f}",
                    "" if res.hist_actual is None else f"{res.hist_actual:.3f}",
                    "" if res.sitio.node_id not in NODOS_MIN_NOCTURNO else ("SI" if res.hist_alerta else "NO"),
                ]
            )
    print(f"[OK] CSV  {csv_path}")
    print(f"DOCX={docx_path}")
    print(f"PDF={pdf or ''}")
    print(f"CSV={csv_path}")
    print(f"XLSX={xlsx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
