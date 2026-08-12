"""
Detección de activación control WES en colegios Providencia.

Criterio: consumo nocturno (00–06) y/o fin de semana → ~0
MIENTRAS el consumo diurno hábil sigue activo (no es desconexión total).

Uso:
  python generar_reporte_activacion_control_wes_providencia.py
"""

from __future__ import annotations

import csv
import subprocess
import sys
from collections import defaultdict
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Providencia" / "activacion_control_wes"
UMBRAL = 0.05
START = date(2025, 10, 1)
END = date(2026, 8, 12)

NODOS = [
    ("000006-01", "Liceo Lastarria"),
    ("000006-02", "Carmela Carvajal"),
    ("000006-04", "Liceo 7 Luisa Saavedra"),
    ("000006-05", "Liceo Juan Pablo Duarte"),
]


def _fmt(v: float, d: int = 1) -> str:
    from generar_reporte_word import format_number_chilean

    return format_number_chilean(v, d)


def _load_days(nid: str) -> List[dict]:
    from control_nocturno import obtener_datos_horarios_dia

    rows = []
    d = START
    while d <= END:
        dt = datetime.combine(d, datetime.min.time())
        try:
            h = obtener_datos_horarios_dia(nid, dt)
        except Exception:
            h = {}
        h = {i: float(h.get(i, 0) or 0) for i in range(24)}
        noct = sum(h[i] for i in range(7))
        diurno = sum(h[i] for i in range(7, 24))
        total = noct + diurno
        wd = d.weekday()
        rows.append(
            {
                "d": d,
                "wd": wd,
                "noct": noct,
                "diurno": diurno,
                "total": total,
                "noct_cero": noct <= UMBRAL,
                "diurno_ok": diurno > 0.5,
                "finde": wd >= 5,
                "total_cero": total <= UMBRAL,
            }
        )
        d += timedelta(days=1)
    return rows


def _semanas(rows: List[dict]) -> List[dict]:
    by: Dict[Tuple[int, int], List[dict]] = defaultdict(list)
    for r in rows:
        by[r["d"].isocalendar()[:2]].append(r)
    out = []
    for w in sorted(by):
        dias = by[w]
        lv = [r for r in dias if r["wd"] < 5]
        fs = [r for r in dias if r["wd"] >= 5]
        lv_act = [r for r in lv if r["diurno_ok"]]
        ini = min(r["d"] for r in dias)
        fin = max(r["d"] for r in dias)
        if len(lv_act) < 2:
            out.append(
                {
                    "w": w,
                    "ini": ini,
                    "fin": fin,
                    "tipo": "sin_operacion",
                    "noct_lv": sum(r["noct"] for r in lv),
                    "diurno_lv": sum(r["diurno"] for r in lv),
                    "fs_tot": sum(r["total"] for r in fs),
                    "pct_noct_cero": None,
                    "fs_cero": None,
                    "control_noche": False,
                    "control_finde": False,
                }
            )
            continue
        pct = sum(1 for r in lv_act if r["noct_cero"]) / len(lv_act)
        fs_cero = bool(fs) and all(r["total_cero"] for r in fs)
        control_noche = pct >= 0.8
        control_finde = control_noche and fs_cero
        out.append(
            {
                "w": w,
                "ini": ini,
                "fin": fin,
                "tipo": "operativo",
                "noct_lv": sum(r["noct"] for r in lv_act),
                "diurno_lv": sum(r["diurno"] for r in lv_act),
                "fs_tot": sum(r["total"] for r in fs),
                "pct_noct_cero": pct,
                "fs_cero": fs_cero,
                "control_noche": control_noche,
                "control_finde": control_finde,
            }
        )
    return out


def _tramos(sems: List[dict], key: str) -> List[Tuple[dict, dict]]:
    tramos = []
    cur = None
    for s in sems:
        on = s["tipo"] == "operativo" and s.get(key)
        if on and cur is None:
            cur = [s, s]
        elif on and cur is not None:
            cur[1] = s
        elif not on and cur is not None:
            tramos.append((cur[0], cur[1]))
            cur = None
    if cur:
        tramos.append((cur[0], cur[1]))
    return tramos


def _grafico_timeline(nid: str, nombre: str, rows: List[dict], sems: List[dict], out: Path) -> None:
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(11, 6.2), sharex=True)

    xs = [r["d"] for r in rows]
    ax1.bar(xs, [r["noct"] for r in rows], color="#c0392b", width=0.8, label="Nocturno 00–06")
    ax1.set_ylabel("m³ noche")
    ax1.set_title(f"{nombre} — nocturno diario y marcas de control WES")
    ax1.grid(True, axis="y", alpha=0.3)
    ax1.legend(loc="upper right", fontsize=8)

    # sombrear tramos control
    for a, b in _tramos(sems, "control_noche"):
        ax1.axvspan(a["ini"], b["fin"], color="#27ae60", alpha=0.22)
        ax2.axvspan(a["ini"], b["fin"], color="#27ae60", alpha=0.22)

    # finde total
    fs_x = [r["d"] for r in rows if r["finde"]]
    fs_y = [r["total"] for r in rows if r["finde"]]
    ax2.bar(fs_x, fs_y, color="#2980b9", width=0.9, label="Consumo total sáb/dom")
    ax2.set_ylabel("m³ finde")
    ax2.set_xlabel("Fecha")
    ax2.grid(True, axis="y", alpha=0.3)
    ax2.legend(loc="upper right", fontsize=8)
    ax2.xaxis.set_major_formatter(mdates.DateFormatter("%b-%y"))
    ax2.xaxis.set_major_locator(mdates.MonthLocator())
    fig.autofmt_xdate()
    fig.text(
        0.5,
        0.01,
        "Franja verde = semanas con control nocturno (noche ~0 y diurno hábil > 0)",
        ha="center",
        fontsize=8,
        color="#555",
    )
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out, dpi=150)
    plt.close(fig)


def _grafico_resumen(hallazgos: List[dict], out: Path) -> None:
    labels = [h["corto"] for h in hallazgos]
    # días en control (aprox)
    dias = []
    for h in hallazgos:
        n = 0
        for a, b in h["tramos_noche"]:
            n += (b["fin"] - a["ini"]).days + 1
        dias.append(n)
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    bars = ax.bar(labels, dias, color="#27ae60", edgecolor="white")
    for bar, v in zip(bars, dias):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 1,
            str(v),
            ha="center",
            va="bottom",
            fontsize=10,
        )
    ax.set_ylabel("Días en tramos con control nocturno")
    ax.set_title("Cobertura del control WES detectado (oct 2025 – ago 2026)")
    ax.grid(True, axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out, dpi=150)
    plt.close(fig)


def generar() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    hallazgos = []

    print(f"[INFO] Detectando control WES {START} → {END}")
    for nid, nombre in NODOS:
        print(f"  {nid} {nombre}...")
        rows = _load_days(nid)
        sems = _semanas(rows)
        tramos_n = _tramos(sems, "control_noche")
        tramos_f = _tramos(sems, "control_finde")
        png = OUT_DIR / f"timeline_{nid}_{ts}.png"
        _grafico_timeline(nid, nombre, rows, sems, png)

        # primera activación: primer tramo con >= 7 días
        primera = None
        for a, b in tramos_n:
            if (b["fin"] - a["ini"]).days + 1 >= 7:
                primera = a["ini"]
                break
        if primera is None and tramos_n:
            primera = tramos_n[0][0]["ini"]

        # fin del último tramo / estado actual
        ultimo_fin = tramos_n[-1][1]["fin"] if tramos_n else None
        activo_hoy = bool(tramos_n) and (END - tramos_n[-1][1]["fin"]).days <= 10

        corto = (
            nombre.replace("Liceo ", "")
            .replace("Luisa Saavedra", "7")
            .replace("Juan Pablo Duarte", "Duarte")
        )
        hallazgos.append(
            {
                "node_id": nid,
                "nombre": nombre,
                "corto": corto,
                "tramos_noche": tramos_n,
                "tramos_finde": tramos_f,
                "primera": primera,
                "ultimo_fin": ultimo_fin,
                "activo_hoy": activo_hoy,
                "png": png,
                "sems": sems,
                "rows": rows,
            }
        )
        print(f"    primera={primera} tramos_noche={len(tramos_n)} tramos_finde={len(tramos_f)}")

    png_res = OUT_DIR / f"resumen_dias_control_{ts}.png"
    _grafico_resumen(hallazgos, png_res)

    # CSV tramos
    csv_path = OUT_DIR / f"tramos_control_wes_providencia_{ts}.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(
            [
                "node_id",
                "liceo",
                "tipo_tramo",
                "desde",
                "hasta",
                "dias",
                "firma",
            ]
        )
        for h in hallazgos:
            for a, b in h["tramos_noche"]:
                firma = "noche+finde" if any(
                    x[0]["ini"] == a["ini"] for x in h["tramos_finde"]
                ) else "solo_noche"
                w.writerow(
                    [
                        h["node_id"],
                        h["nombre"],
                        "control_wes",
                        a["ini"].strftime("%d/%m/%Y"),
                        b["fin"].strftime("%d/%m/%Y"),
                        (b["fin"] - a["ini"]).days + 1,
                        firma,
                    ]
                )
            if not h["tramos_noche"]:
                w.writerow(
                    [h["node_id"], h["nombre"], "sin_control_detectado", "", "", "", ""]
                )

    # Word
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    title = doc.add_heading(
        "Activación control WES — Colegios Providencia", 0
    )
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    p = doc.add_paragraph()
    p.add_run("Método: ").bold = True
    p.add_run(
        "se revisa día a día el consumo 00:00–06:59 y el de sábados/domingos. "
        "Se marca control WES cuando la noche (lun–vie) cae a ~0 m³ "
        f"(≤ {UMBRAL} m³) en ≥80 % de los días hábiles con actividad diurna, "
        "es decir el colegio sigue consumiendo de día pero la madrugada queda en cero. "
        "Si además el fin de semana completo queda en cero, se etiqueta noche+finde. "
        "Si el día entero (incluida la mañana) está en cero, se clasifica como "
        "sin operación (vacaciones / corte / sin data), no como control."
    )
    meta = doc.add_paragraph()
    meta.add_run("Periodo: ").bold = True
    meta.add_run(f"{START.strftime('%d/%m/%Y')} al {END.strftime('%d/%m/%Y')}. ")
    meta.add_run("Generado: ").bold = True
    meta.add_run(datetime.now().strftime("%d-%m-%Y %H:%M"))

    doc.add_heading("1. Resumen — cuándo se activó el control", level=1)
    tbl = doc.add_table(rows=1 + len(hallazgos), cols=5)
    tbl.style = "Table Grid"
    for j, hd in enumerate(
        ["Liceo", "Primera activación", "Tramos detectados", "Último tramo hasta", "¿Activo hoy?"]
    ):
        tbl.rows[0].cells[j].text = hd
        for run in tbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, h in enumerate(hallazgos, 1):
        row = tbl.rows[i].cells
        row[0].text = h["nombre"]
        row[1].text = h["primera"].strftime("%d/%m/%Y") if h["primera"] else "No detectado"
        row[2].text = str(len(h["tramos_noche"]))
        row[3].text = (
            h["ultimo_fin"].strftime("%d/%m/%Y") if h["ultimo_fin"] else "—"
        )
        row[4].text = "Sí" if h["activo_hoy"] else "No"

    if png_res.is_file():
        doc.add_paragraph()
        doc.add_picture(str(png_res), width=Inches(5.5))

    # Lectura ejecutiva
    doc.add_heading("2. Lectura por colegio", level=1)

    textos = {
        "000006-01": (
            "Lastarria muestra control nocturno casi desde el inicio de la data "
            "(octubre 2025): noches en cero con consumo diurno activo, y varios "
            "tramos con finde también en cero (firma alineada a la regulación "
            "lun–vie madrugada + sáb–dom). Ese patrón se mantiene con fuerza hasta "
            "mediados/fines de abril 2026. Desde ~27/04–mayo 2026 el nocturno "
            "vuelve a niveles altos: el control dejó de verse en la serie."
        ),
        "000006-02": (
            "Carmela tiene un tramo claro de control nocturno del 26/01/2026 al "
            "01/03/2026: las noches lun–vie caen a cero mientras el diurno sigue. "
            "Los fines de semana no quedan en cero (coherente con regulación solo "
            "madrugada, no corte de finde completo). Antes y después de ese tramo "
            "hay consumo nocturno habitual."
        ),
        "000006-04": (
            "Liceo 7 casi no muestra control sostenido. Hay muchos periodos "
            "sin_operacion (día completo en cero: corte/sin data), distintos del "
            "control. Solo aparece un tramo corto 29/06–05/07/2026. El alza "
            "fuerte de jul–ago no corresponde a control activo."
        ),
        "000006-05": (
            "Duarte no presenta tramos sostenidos de noche en cero con diurno "
            "activo. Hay fines de semana aislados en cero, pero las madrugadas "
            "hábiles siguen con consumo: no se ve activación clara de control WES "
            "en la serie analizada."
        ),
    }

    for h in hallazgos:
        doc.add_heading(f"{h['nombre']} ({h['node_id']})", level=2)
        doc.add_paragraph(textos.get(h["node_id"], ""))
        if h["tramos_noche"]:
            doc.add_paragraph("Tramos con control nocturno detectado:")
            for a, b in h["tramos_noche"]:
                dias = (b["fin"] - a["ini"]).days + 1
                con_f = any(
                    x[0]["ini"] == a["ini"] and x[1]["fin"] == b["fin"]
                    for x in h["tramos_finde"]
                )
                tag = "noche + finde → 0" if con_f else "solo noche → 0"
                doc.add_paragraph(
                    f"• {a['ini'].strftime('%d/%m/%Y')} al {b['fin'].strftime('%d/%m/%Y')} "
                    f"({dias} días) — {tag}",
                    style="List Bullet",
                )
        else:
            doc.add_paragraph("Sin tramos de control nocturno sostenido.")
        if h["png"].is_file():
            doc.add_picture(str(h["png"]), width=Inches(6.1))

    doc.add_heading("3. Cómo leer el gráfico", level=1)
    doc.add_paragraph(
        "Barras rojas = m³ nocturnos por día. Barras azules = m³ totales de "
        "sábado/domingo. Franja verde = semanas donde el algoritmo marca control "
        "(noche ~0 con colegio operando de día). Huecos sin franja verde y con "
        "barras en cero todo el día = sin operación, no control."
    )

    doc.add_heading("4. Implicancia para el ahorro", level=1)
    doc.add_paragraph(
        "Los periodos verdes son evidencia de que el control WES sí llevó el "
        "nocturno (y a veces el finde) a cero. Los periodos sin franja, con "
        "nocturno alto, son la oportunidad de ahorro si se reactiva/extiende "
        "el control. En Lastarria el contraste pre/mayo vs post/mayo es el más "
        "claro para cocinar el mensaje comercial."
    )

    out_docx = OUT_DIR / f"Activacion_control_WES_Providencia_{ts}.docx"
    doc.save(out_docx)
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUT_DIR),
                str(out_docx),
            ],
            check=False,
            capture_output=True,
            timeout=120,
        )
    except Exception as ex:
        print(f"[WARN] PDF: {ex}")

    pdf = out_docx.with_suffix(".pdf")
    print(f"[OK] Word: {out_docx}")
    print(f"[OK] PDF:  {pdf if pdf.is_file() else 'no'}")
    print(f"[OK] CSV:  {csv_path}")
    return out_docx


def main() -> int:
    print("=" * 72)
    print("ACTIVACIÓN CONTROL WES — PROVIDENCIA")
    print("=" * 72)
    generar()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
