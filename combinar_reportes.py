"""
Script para combinar el reporte de José con el reporte de puntos en cero.
Genera un documento unificado con resumen ejecutivo mejorado.
"""

import sys
import copy
import io
from pathlib import Path
from datetime import datetime, timezone
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement

# Configurar codificación UTF-8
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')


def extraer_texto_documento(doc_path: Path) -> str:
    """
    Extrae todo el texto de un documento Word.
    """
    try:
        doc = Document(doc_path)
        texto_completo = []
        
        for paragraph in doc.paragraphs:
            texto = paragraph.text.strip()
            if texto:
                texto_completo.append(texto)
        
        # También extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                fila_texto = []
                for cell in row.cells:
                    texto_celda = cell.text.strip()
                    if texto_celda:
                        fila_texto.append(texto_celda)
                if fila_texto:
                    texto_completo.append(" | ".join(fila_texto))
        
        return "\n".join(texto_completo)
    except Exception as e:
        print(f"[ERROR] No se pudo leer el documento {doc_path}: {e}")
        return ""


def extraer_tablas_documento(doc_path: Path) -> list:
    """
    Extrae las tablas de un documento Word.
    Retorna lista de tablas con sus datos.
    """
    try:
        doc = Document(doc_path)
        tablas_extraidas = []
        
        for table in doc.tables:
            datos_tabla = []
            for row in table.rows:
                fila = []
                for cell in row.cells:
                    fila.append(cell.text.strip())
                datos_tabla.append(fila)
            if datos_tabla:
                tablas_extraidas.append(datos_tabla)
        
        return tablas_extraidas
    except Exception as e:
        print(f"[ERROR] No se pudo extraer tablas de {doc_path}: {e}")
        return []


def copiar_imagen_con_relacion(origen_doc: Document, destino_doc: Document, rId_original: str):
    """
    Copia una imagen del documento origen al destino y retorna el nuevo rId.
    """
    try:
        # Obtener la relación original
        rel_original = origen_doc.part.rels.get(rId_original)
        if not rel_original or not hasattr(rel_original, 'target_part'):
            return None
        
        # Obtener el blob de la imagen
        imagen_blob = rel_original.target_part.blob
        content_type = rel_original.target_part.content_type
        
        # Verificar si ya existe una imagen idéntica en el destino
        for rId, rel in destino_doc.part.rels.items():
            if (hasattr(rel, 'target_part') and 
                hasattr(rel.target_part, 'blob') and
                rel.target_part.content_type == content_type and
                rel.target_part.blob == imagen_blob):
                return rId
        
        # Convertir blob a BytesIO para que sea file-like
        imagen_stream = io.BytesIO(imagen_blob)
        imagen_stream.seek(0)  # Asegurar que esté al inicio
        
        # Usar el método get_or_add_image_part del package
        try:
            nueva_parte = destino_doc.part.package.get_or_add_image_part(imagen_stream)
        except Exception:
            # Si falla, intentar método alternativo
            from docx.opc.packuri import PackURI
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.opc.part import ImagePart
            
            # Obtener extensión del content type
            ext_map = {
                'image/png': 'png',
                'image/jpeg': 'jpg',
                'image/jpg': 'jpg',
                'image/gif': 'gif',
                'image/bmp': 'bmp'
            }
            ext = ext_map.get(content_type, 'png')
            
            # Crear partname único
            import uuid
            filename = f"image{uuid.uuid4().hex[:8]}.{ext}"
            partname = PackURI(f"/word/media/{filename}")
            
            # Crear ImagePart directamente
            nueva_parte = ImagePart(partname, content_type, imagen_blob, destino_doc.part.package)
            destino_doc.part.package.parts.append(nueva_parte)
        
        # Agregar relación si no existe
        rId_existente = None
        for rId, rel in destino_doc.part.rels.items():
            if hasattr(rel, 'target_part') and rel.target_part == nueva_parte:
                rId_existente = rId
                break
        
        if not rId_existente:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            nueva_rel = destino_doc.part.rels.add_relationship(
                RT.IMAGE,
                nueva_parte,
                nueva_parte.partname
            )
            return nueva_rel.rId
        else:
            return rId_existente
    except Exception as e:
        print(f"[ADVERTENCIA] Error al copiar imagen: {e}")
        import traceback
        traceback.print_exc()
        return None


def copiar_elementos_documento(origen: Document, destino: Document, saltar_titulo: bool = False):
    """
    Copia párrafos, imágenes y tablas de un documento a otro, manteniendo el formato.
    """
    saltado_titulo = False
    
    # Copiar párrafos (incluyendo imágenes)
    for paragraph in origen.paragraphs:
        texto = paragraph.text.strip()
        tiene_imagen = any(run._element.xpath('.//a:blip') or run._element.xpath('.//w:drawing') for run in paragraph.runs)
        
        # Saltar título principal si se solicita
        if saltar_titulo and not saltado_titulo:
            if texto and (len(texto) > 30 or paragraph.style.name.startswith('Heading 0')):
                saltado_titulo = True
                continue
        
        # Si el párrafo tiene contenido (texto o imagen)
        if texto or tiene_imagen:
            # Copiar estilo si es un heading
            if paragraph.style.name.startswith('Heading') and texto:
                nivel = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                nuevo_parrafo = destino.add_heading(texto, min(nivel, 9))
            else:
                nuevo_parrafo = destino.add_paragraph()
                if texto:
                    nuevo_parrafo.style = paragraph.style
            
            nuevo_parrafo.alignment = paragraph.alignment
            
            # Copiar runs con formato e imágenes
            for run in paragraph.runs:
                # Verificar si el run tiene imagen
                if run._element.xpath('.//a:blip') or run._element.xpath('.//w:drawing'):
                    # Copiar imagen con sus relaciones - método simplificado
                    try:
                        # Obtener el rId de la imagen
                        r_embed = None
                        for blip in run._element.xpath('.//a:blip'):
                            r_embed = blip.get(qn('r:embed'))
                            if r_embed:
                                break
                        
                        if r_embed:
                            # Copiar la imagen y obtener el nuevo rId
                            nuevo_rId = copiar_imagen_con_relacion(origen, destino, r_embed)
                            if nuevo_rId:
                                # Crear nuevo run y agregar imagen usando add_picture indirectamente
                                # Copiar el XML del run pero actualizando el rId
                                run_xml = copy.deepcopy(run._element)
                                
                                # Actualizar todas las referencias r:embed en el XML
                                for blip in run_xml.xpath('.//a:blip'):
                                    if blip.get(qn('r:embed')):
                                        blip.set(qn('r:embed'), nuevo_rId)
                                
                                # Agregar el run clonado al párrafo
                                nuevo_parrafo._element.append(run_xml)
                            else:
                                # Si no se pudo copiar la imagen, agregar texto alternativo
                                if run.text:
                                    nuevo_run = nuevo_parrafo.add_run(f"[Imagen: {run.text}]")
                        else:
                            # Si no hay rId, copiar el run completo sin modificar
                            nuevo_parrafo._element.append(copy.deepcopy(run._element))
                    except Exception as e:
                        print(f"[ADVERTENCIA] No se pudo copiar una imagen: {e}")
                        # Si falla, intentar copiar solo el texto
                        if run.text:
                            nuevo_run = nuevo_parrafo.add_run(run.text)
                            nuevo_run.bold = run.bold
                            nuevo_run.italic = run.italic
                else:
                    # Copiar texto y formato
                    nuevo_run = nuevo_parrafo.add_run(run.text)
                    nuevo_run.bold = run.bold
                    nuevo_run.italic = run.italic
                    if run.font.size:
                        nuevo_run.font.size = run.font.size
                    if run.font.color and run.font.color.rgb:
                        nuevo_run.font.color.rgb = run.font.color.rgb
    
    # Copiar tablas (incluyendo imágenes en celdas)
    for table in origen.tables:
        num_rows = len(table.rows)
        num_cols = len(table.rows[0].cells) if num_rows > 0 else 0
        
        if num_rows > 0 and num_cols > 0:
            nueva_tabla = destino.add_table(rows=num_rows, cols=num_cols)
            nueva_tabla.style = table.style
            
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    nueva_celda = nueva_tabla.rows[i].cells[j]
                    
                    # Copiar contenido de la celda (texto e imágenes)
                    if cell.paragraphs:
                        # Limpiar párrafos por defecto de la celda nueva
                        nueva_celda._element.clear_content()
                        
                        for paragraph in cell.paragraphs:
                            nuevo_par_celda = nueva_celda.add_paragraph()
                            nuevo_par_celda.alignment = paragraph.alignment
                            
                            for run in paragraph.runs:
                                # Verificar si el run tiene imagen
                                if run._element.xpath('.//a:blip') or run._element.xpath('.//w:drawing'):
                                    # Copiar imagen con sus relaciones - método simplificado
                                    try:
                                        # Obtener el rId de la imagen
                                        r_embed = None
                                        for blip in run._element.xpath('.//a:blip'):
                                            r_embed = blip.get(qn('r:embed'))
                                            if r_embed:
                                                break
                                        
                                        if r_embed:
                                            # Copiar la imagen y obtener el nuevo rId
                                            nuevo_rId = copiar_imagen_con_relacion(origen, destino, r_embed)
                                            if nuevo_rId:
                                                # Copiar el XML del run pero actualizando el rId
                                                run_xml = copy.deepcopy(run._element)
                                                
                                                # Actualizar todas las referencias r:embed en el XML
                                                for blip in run_xml.xpath('.//a:blip'):
                                                    if blip.get(qn('r:embed')):
                                                        blip.set(qn('r:embed'), nuevo_rId)
                                                
                                                # Agregar el run clonado al párrafo de la celda
                                                nuevo_par_celda._element.append(run_xml)
                                            else:
                                                # Si no se pudo copiar la imagen, agregar texto alternativo
                                                if run.text:
                                                    nuevo_run = nuevo_par_celda.add_run(f"[Imagen: {run.text}]")
                                        else:
                                            # Si no hay rId, copiar el run completo sin modificar
                                            nuevo_par_celda._element.append(copy.deepcopy(run._element))
                                    except Exception as e:
                                        # Si falla, copiar solo el texto
                                        if run.text:
                                            nuevo_run = nuevo_par_celda.add_run(run.text)
                                            nuevo_run.bold = run.bold
                                            nuevo_run.italic = run.italic
                                else:
                                    # Copiar texto y formato
                                    nuevo_run = nuevo_par_celda.add_run(run.text)
                                    nuevo_run.bold = run.bold
                                    nuevo_run.italic = run.italic
                                    if run.font.size:
                                        nuevo_run.font.size = run.font.size
                                    if run.font.color and run.font.color.rgb:
                                        nuevo_run.font.color.rgb = run.font.color.rgb
                    else:
                        # Si no hay párrafos, copiar texto directo
                        nueva_celda.text = cell.text
                    
                    # Copiar formato de la celda (ancho, etc.)
                    try:
                        # Copiar propiedades de la celda
                        tc_pr_original = cell._element.tcPr
                        if tc_pr_original is not None:
                            tc_pr_nuevo = nueva_celda._element.get_or_add_tcPr()
                            for child in tc_pr_original:
                                tc_pr_nuevo.append(copy.deepcopy(child))
                    except:
                        pass


def crear_resumen_ejecutivo(
    doc: Document,
    texto_reporte_jose: str,
    texto_reporte_cero: str,
    num_puntos_cero: int,
    num_puntos_sin_datos: int,
    total_puntos: int
) -> None:
    """
    Crea un resumen ejecutivo mejorado al inicio del documento.
    """
    # Título principal
    title = doc.add_heading("REPORTE CONSOLIDADO DE MONITOREO", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
    title_run.bold = True
    title_run.font.size = Pt(22)
    
    # Fecha de generación
    fecha_generacion = datetime.now(timezone.utc).strftime("%d de %B de %Y, %H:%M:%S UTC")
    gen_para = doc.add_paragraph(f"Fecha de generación: {fecha_generacion}")
    gen_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    gen_para.runs[0].font.color.rgb = RGBColor(64, 64, 64)
    gen_para.runs[0].italic = True
    
    doc.add_paragraph("")  # Espacio
    
    # Resumen ejecutivo
    doc.add_heading("RESUMEN EJECUTIVO", 1)
    
    puntos_con_datos = total_puntos - num_puntos_sin_datos
    porcentaje_con_datos = ((puntos_con_datos / total_puntos * 100) if total_puntos > 0 else 0)
    porcentaje_cero = ((num_puntos_cero / puntos_con_datos * 100) if puntos_con_datos > 0 else 0)
    porcentaje_sin_datos = ((num_puntos_sin_datos / total_puntos * 100) if total_puntos > 0 else 0)
    
    resumen_texto = f"""Este documento presenta un análisis consolidado del estado operativo del sistema de monitoreo WES, integrando la evaluación de puntos de monitoreo con información complementaria del sistema.

ESTADO OPERATIVO DEL SISTEMA

Durante el análisis de los últimos 3 días, se evaluaron {total_puntos} puntos de monitoreo distribuidos en el sistema, arrojando los siguientes resultados:

• Puntos operativos con datos: {puntos_con_datos} puntos ({porcentaje_con_datos:.1f}% del total)
• Puntos con consumo cero: {num_puntos_cero} puntos ({porcentaje_cero:.1f}% de los puntos con datos)
• Puntos sin datos disponibles: {num_puntos_sin_datos} puntos ({porcentaje_sin_datos:.1f}% del total)

ANÁLISIS Y DIAGNÓSTICO

Los {num_puntos_sin_datos} puntos sin datos disponibles ({porcentaje_sin_datos:.1f}% del total) representan una interrupción crítica en la capacidad de monitoreo. Esta situación puede originarse por fallas de conectividad, sensores desconectados, problemas en la transmisión de datos o configuraciones incorrectas. Se requiere una revisión técnica inmediata para restaurar la visibilidad completa del sistema.

Por su parte, los {num_puntos_cero} puntos que registran consumo cero ({porcentaje_cero:.1f}% de los puntos operativos) necesitan verificación para determinar si corresponde a un consumo real nulo (situación esperada) o a un problema técnico que requiere intervención.

ESTRUCTURA DEL DOCUMENTO

Este reporte consolidado está organizado en las siguientes secciones:

1. Análisis detallado de puntos en cero y sin datos: Incluye tablas completas con identificación de cada punto, su ubicación y empresa asociada.

2. Información adicional del sistema: Contiene gráficos, métricas y datos complementarios del sistema de monitoreo.

3. Datos específicos y métricas relevantes: Proporciona información detallada para la toma de decisiones.

RECOMENDACIONES PRIORITARIAS

1. Revisión inmediata: Priorizar la atención de los {num_puntos_sin_datos} puntos sin datos disponibles para restaurar la cobertura completa del sistema.

2. Verificación de puntos en cero: Validar que los {num_puntos_cero} puntos con consumo cero correspondan a situaciones esperadas y no a fallas técnicas.

3. Monitoreo continuo: Establecer un seguimiento periódico para detectar tempranamente nuevas interrupciones en el sistema."""
    
    resumen_para = doc.add_paragraph(resumen_texto)
    resumen_para.runs[0].font.size = Pt(11)
    resumen_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph("")  # Espacio
    doc.add_paragraph("")  # Espacio
    
    # Línea separadora
    separador = doc.add_paragraph("_" * 80)
    separador.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    separador.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph("")  # Espacio


def crear_documento_combinado(
    ruta_reporte_jose: Path,
    ruta_reporte_cero: Path,
    carpeta_salida: Path
) -> Path:
    """
    Crea un documento combinado con ambos reportes.
    """
    print("=" * 70)
    print("COMBINANDO REPORTES")
    print("=" * 70)
    print()
    
    # Verificar que existan los archivos
    if not ruta_reporte_jose.exists():
        print(f"[ERROR] No se encontró el archivo: {ruta_reporte_jose}")
        return None
    
    if not ruta_reporte_cero.exists():
        print(f"[ERROR] No se encontró el archivo: {ruta_reporte_cero}")
        return None
    
    print(f"[OK] Leyendo: {ruta_reporte_jose.name}")
    print(f"[OK] Leyendo: {ruta_reporte_cero.name}")
    print()
    
    # Leer documentos
    doc_jose = Document(ruta_reporte_jose)
    doc_cero = Document(ruta_reporte_cero)
    
    # Extraer información del reporte de puntos en cero
    texto_cero = extraer_texto_documento(ruta_reporte_cero)
    
    # Contar puntos en cero y sin datos desde el texto
    num_puntos_cero = texto_cero.count("EN CERO") or texto_cero.count("marcando cero")
    num_puntos_sin_datos = texto_cero.count("Sin datos disponibles")
    
    # Buscar total de puntos en el texto
    total_puntos = 0
    for line in texto_cero.split("\n"):
        if "Total de puntos analizados:" in line or "Total nodos encontrados:" in line:
            try:
                total_puntos = int([s for s in line.split() if s.isdigit()][-1])
                break
            except:
                pass
    
    # Si no se encontró, usar valores por defecto
    if total_puntos == 0:
        total_puntos = 354  # Valor aproximado
    
    # Crear nuevo documento
    doc_combinado = Document()
    
    # Crear resumen ejecutivo
    texto_jose = extraer_texto_documento(ruta_reporte_jose)
    crear_resumen_ejecutivo(
        doc_combinado,
        texto_jose,
        texto_cero,
        num_puntos_cero,
        num_puntos_sin_datos,
        total_puntos
    )
    
    # Agregar sección del reporte de puntos en cero
    doc_combinado.add_heading("REPORTE DE PUNTOS EN CERO Y SIN DATOS", 1)
    doc_combinado.add_paragraph(
        "A continuación se presenta el análisis detallado de los puntos de monitoreo que están marcando cero o no tienen datos disponibles."
    )
    doc_combinado.add_paragraph("")  # Espacio
    
    # Copiar contenido del reporte de puntos en cero (sin el título inicial)
    copiar_elementos_documento(doc_cero, doc_combinado, saltar_titulo=True)
    
    doc_combinado.add_paragraph("")  # Espacio
    doc_combinado.add_paragraph("")  # Espacio
    
    # Línea separadora
    separador = doc_combinado.add_paragraph("_" * 80)
    separador.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    separador.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc_combinado.add_paragraph("")  # Espacio
    
    # Agregar sección del reporte de José
    doc_combinado.add_heading("INFORMACIÓN ADICIONAL DE MONITOREO", 1)
    doc_combinado.add_paragraph(
        "A continuación se presenta información adicional del sistema de monitoreo."
    )
    doc_combinado.add_paragraph("")  # Espacio
    
    # Copiar contenido del reporte de José
    copiar_elementos_documento(doc_jose, doc_combinado, saltar_titulo=False)
    
    # Guardar documento con nombre fijo
    filename = "Reporte_Consolidado.docx"
    output_path = carpeta_salida / filename
    
    # Eliminar archivo anterior si existe (puede estar abierto)
    if output_path.exists():
        try:
            output_path.unlink()
            print(f"[INFO] Archivo anterior eliminado: {filename}")
        except PermissionError:
            # Si está abierto, crear con timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Reporte_Consolidado_{timestamp}.docx"
            output_path = carpeta_salida / filename
            print(f"[INFO] Archivo anterior está abierto, creando nuevo: {filename}")
    
    doc_combinado.save(str(output_path))
    
    print(f"[OK] Documento combinado generado/actualizado:")
    print(f"  {output_path}")
    print()
    
    return output_path


def main():
    """Función principal."""
    carpeta = Path("reporte en cero")
    
    # Buscar archivos
    reporte_jose = None
    reporte_cero = None
    
    # Buscar reporte de José
    for archivo in carpeta.glob("*jose*.docx"):
        if not archivo.name.startswith("~$"):  # Ignorar archivos temporales
            reporte_jose = archivo
            break
    
    # Buscar reporte de puntos en cero más reciente
    reportes_cero = sorted(
        [f for f in carpeta.glob("*Puntos_En_Cero*.docx") if not f.name.startswith("~$")],
        key=lambda x: x.stat().st_mtime,
        reverse=True
    )
    
    if reportes_cero:
        reporte_cero = reportes_cero[0]
    
    if not reporte_jose:
        print("[ERROR] No se encontró el archivo 'reporte jose' en la carpeta.")
        return
    
    if not reporte_cero:
        print("[ERROR] No se encontró el archivo de reporte de puntos en cero.")
        return
    
    print(f"Archivo de José: {reporte_jose.name}")
    print(f"Archivo de puntos en cero: {reporte_cero.name}")
    print()
    
    # Crear documento combinado
    resultado = crear_documento_combinado(reporte_jose, reporte_cero, carpeta)
    
    if resultado:
        print("=" * 70)
        print("PROCESO COMPLETADO")
        print("=" * 70)
        print(f"Reporte consolidado guardado en: {resultado.absolute()}")
    else:
        print("[ERROR] No se pudo generar el documento combinado.")


if __name__ == "__main__":
    main()

