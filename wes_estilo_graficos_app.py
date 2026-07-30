"""
Gráficos y textos alineados a la app WES / informes de referencia (estilo Diego Carrasco).

- Perfil horario 24 h (barras + sombreado madrugada), datos API como la app.
- Proyección mensual desde consumo nocturno horario (sin depender de alertas MyAlert).
"""

from __future__ import annotations

from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.ticker import MultipleLocator

# Paleta cercana a app WES / informes comparación
COLOR_CONSUMO = "#2e7ac8"
COLOR_CONSUMO_FILL = (0.62, 0.82, 0.98)
COLOR_NOCHE = "#c41e1e"
COLOR_BARRA_WES = "#4A8CB8"
COLOR_BARRA_FACT = "#E67E22"


def horas_api_chile(node_id: str, dia: datetime) -> Dict[int, float]:
    """m³/h por hora 0–23, alineado visualización app (CSV hora local)."""
    from control_nocturno import obtener_datos_horarios_dia

    h = obtener_datos_horarios_dia(node_id, dia)
    return {i: float(h.get(i, 0.0)) for i in range(24)}


def dia_mayor_consumo_nocturno(
    node_id: str, start_dt: datetime, end_dt: datetime
) -> Tuple[datetime, Dict[int, float], float]:
    best_day = start_dt
    best_h: Dict[int, float] = {i: 0.0 for i in range(24)}
    best_n = -1.0
    d = start_dt.date()
    while d <= end_dt.date():
        dt = datetime.combine(d, datetime.min.time())
        horas = horas_api_chile(node_id, dt)
        noche = sum(horas.get(h, 0.0) for h in range(7))
        if noche > best_n:
            best_n = noche
            best_day = dt
            best_h = horas
        d += timedelta(days=1)
    return best_day, best_h, best_n


def guardar_grafico_horario_24h_app(
    horas: Dict[int, float],
    out_path: Path,
    *,
    titulo: str,
    subtitulo: str = "Consumo (m³/h) — hora Chile, serie API WES",
) -> Path:
    """Barras 24 h estilo pantalla app (referencia informes comparación)."""
    x = np.arange(24, dtype=float)
    y = np.round(np.array([horas.get(h, 0.0) for h in range(24)], dtype=float), 2)

    fig, ax = plt.subplots(figsize=(10.2, 4.6))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    colors = [COLOR_NOCHE if h <= 6 else COLOR_BARRA_WES for h in range(24)]
    ax.bar(x, y, width=0.78, color=colors, edgecolor="white", linewidth=0.5, zorder=2)
    ax.axvspan(-0.5, 6.5, alpha=0.12, color=COLOR_NOCHE, zorder=0)
    ax.plot(x, y, color=COLOR_CONSUMO, linewidth=1.2, marker="o", markersize=3, zorder=4)

    ymax = float(max(y.max() if len(y) else 0, 0.05)) * 1.18
    ax.set_xlim(-0.5, 23.5)
    ax.set_ylim(0, ymax)
    ax.set_xticks(x)
    ax.set_xticklabels([f"{int(h):02d}" for h in x], fontsize=7)
    ax.set_xlabel("Hora del día (Chile)", fontsize=9)
    ax.set_ylabel("m³/h", fontsize=10)
    ax.set_title(titulo, fontsize=11, fontweight="bold", pad=8)
    if ymax <= 15:
        ax.yaxis.set_major_locator(MultipleLocator(max(0.5, round(ymax / 8, 1))))
    ax.grid(axis="y", linestyle="-", alpha=0.35, zorder=0)
    ax.set_axisbelow(True)
    fig.text(0.5, 0.02, subtitulo, ha="center", fontsize=7.5, color="#444444")
    fig.subplots_adjust(bottom=0.14, left=0.08, right=0.98, top=0.88)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return out_path


def proyeccion_mensual_desde_nocturno(
    consumo_nocturno_total: float,
    num_dias_periodo: int,
    dias_con: int,
    dias_sin: int,
    *,
    dias_mes: int = 30,
    forzar: bool = False,
) -> Tuple[float, float, float, bool]:
    """
    Returns: (m³/mes proyectado, m³/día, m³/h promedio noche, cumple_umbral_75pct).
    """
    from generar_reporte_word import (
        HORAS_NOCTURNAS_POR_DIA,
        UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION,
        proyeccion_filtracion_desde_consumo_nocturno,
    )

    proy_per, proy_dia, prom_h = proyeccion_filtracion_desde_consumo_nocturno(
        consumo_nocturno_total, num_dias_periodo, dias_con, dias_sin
    )
    total_d = dias_con + dias_sin
    pct = (100.0 * dias_con / total_d) if total_d else 0.0
    cumple = pct >= UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION

    if (proy_dia <= 0 or proy_per <= 0) and consumo_nocturno_total > 0 and num_dias_periodo > 0:
        if forzar or consumo_nocturno_total > 0:
            total_h = float(num_dias_periodo) * float(HORAS_NOCTURNAS_POR_DIA)
            prom_h = consumo_nocturno_total / total_h
            proy_dia = prom_h * 24.0
            proy_per = proy_dia * float(num_dias_periodo)

    proy_mes = proy_dia * float(dias_mes)
    return proy_mes, proy_dia, prom_h, cumple


def agregar_perfiles_nocturnos_agregado_doc(
    doc,
    nodes_data: List[dict],
    start_dt: datetime,
    end_dt: datetime,
    output_dir: Path,
    *,
    nodos_ya_graficados: Optional[Set[str]] = None,
    price_per_m3: float = 1200.0,
) -> None:
    """Inserta en el Word agregado perfiles horarios + proyección (sin requerir alertas)."""
    from docx.shared import Inches
    from generar_reporte_word import (
        add_formatted_title,
        add_picture_with_pagination,
        calculate_nocturnal_metrics,
        format_currency_chilean,
        format_number_chilean,
    )

    ya = nodos_ya_graficados or set()
    num_dias = (end_dt.date() - start_dt.date()).days + 1
    if num_dias < 1:
        return

    pendientes: List[dict] = []
    for nd in nodes_data:
        nid = nd["node_id"]
        if nid in ya:
            continue
        nm = calculate_nocturnal_metrics(nid, start_dt, end_dt)
        if float(nm["consumo_nocturno_total"]) <= 0:
            continue
        pendientes.append({**nd, "nocturnal": nm})

    if not pendientes:
        return

    doc.add_paragraph("")
    add_formatted_title(
        doc,
        "Perfiles horarios y proyección (serie API — estilo app WES)",
    )
    intro = doc.add_paragraph(
        "Los siguientes gráficos replican el perfil horario de la app WES (sin captura de pantalla). "
        "La proyección mensual usa consumo en madrugada (00:00–06:59) extrapolado a 24 h × 30 días, "
        "útil para estimar ahorro con equipo de control cuando no hay alertas MyAlert recientes."
    )
    intro.alignment = 1  # JUSTIFY - use WD if available
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    except Exception:
        pass

    for item in pendientes:
        nid = item["node_id"]
        name = item["node_name"]
        nm = item["nocturnal"]
        c_noche = float(nm["consumo_nocturno_total"])
        d_con = int(nm["dias_con_consumo_nocturno"])
        d_sin = int(nm["dias_sin_consumo_nocturno"])

        proy_mes, proy_dia, prom_h, cumple = proyeccion_mensual_desde_nocturno(
            c_noche, num_dias, d_con, d_sin, forzar=False
        )
        if proy_mes <= 0 and c_noche > 0:
            proy_mes, proy_dia, prom_h, cumple = proyeccion_mensual_desde_nocturno(
                c_noche, num_dias, d_con, d_sin, forzar=True
            )

        dia_rep, horas_rep, noche_dia = dia_mayor_consumo_nocturno(nid, start_dt, end_dt)
        png = output_dir / f"app_horario_{nid.replace('-', '_')}.png"
        guardar_grafico_horario_24h_app(
            horas_rep,
            png,
            titulo=f"{name} — {dia_rep:%d/%m/%Y} (día con mayor consumo 00–06 h)",
        )

        doc.add_paragraph("")
        add_formatted_title(doc, f"{name.upper()} ({nid})")
        p = doc.add_paragraph()
        p.add_run("Consumo nocturno en el periodo: ").bold = True
        p.add_run(f"{format_number_chilean(c_noche, 1)} m³. ")
        p.add_run("Proyección a 30 días (24 h/día): ").bold = True
        p.add_run(
            f"{format_number_chilean(proy_mes, 1)} m³ "
            f"({format_currency_chilean(proy_mes * price_per_m3)} referencial)."
        )
        if not cumple:
            doc.add_paragraph(
                "Nota: el patrón nocturno no supera el umbral del 75 % de días con consumo en madrugada; "
                "la proyección se muestra como estimación técnica a partir de la serie horaria."
            )
        add_picture_with_pagination(doc, str(png), Inches(6), keep_with_next=True)
