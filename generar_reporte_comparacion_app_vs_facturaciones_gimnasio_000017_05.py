"""
Genera un reporte tipo informe (Word + PDF) que compara:
- "App" = consumo WES (API medidas) del nodo 000017-05 (Gimnasio municipal, Renca)
- Facturaciones Aguas Andinas = consumo total m³ informado en cada factura PDF (o filas CSV)

Fuentes (carpeta ``Gimnasio Facturaciones``), en este orden:
  1) ``facturacion_aguas_andinas.csv`` si existe (útil cuando el PDF es solo imagen escaneada).
  2) Si no hay CSV: PDF de Aguas Andinas (boleta electrónica con texto, o **historial de consumo** renderizado
     como imagen; en ese caso se usa OCR con PyMuPDF + rapidocr-onnxruntime — ``pip install pymupdf rapidocr-onnxruntime``).

CSV manual (UTF-8, separador coma), primera fila encabezado:
  boleta,fecha_emision,lectura_anterior,lectura_actual,m3_cuenta[,cuenta,medidor]
  Fechas: DD-MM-YYYY o DD-MMM-YYYY (mes en inglés o abreviatura de 3 letras, ej. 11-NOV-2025).

Cuadros comparativos (tablas A–C, gráficos y total acumulado): solo boletas con lectura anterior y lectura actual
≥ ``INFORME_COMPARATIVO_DESDE`` (06-06-2025). En este sitio ``MONITOREO_WES_DESDE`` coincide con esa fecha (referencia narrativa).

Salida (todo en ``reporte_comparacion_Gimnasio/``):
  - Comparacion_Facturacion_vs_WES_Gimnasio_Renca.xlsx (paso 1 del comando principal)
  - Comparacion_App_vs_Facturaciones_Gimnasio_Renca.docx y .pdf (tablas y gráficos usan el Excel si está disponible)
  - Gráficos ``chart_barras_gimnasio_*.png``

Uso:
  python generar_reporte_comparacion_app_vs_facturaciones_gimnasio_000017_05.py
"""

from __future__ import annotations

import csv
import os
import re
import shutil
import subprocess
import sys
import traceback
import tempfile
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from facturacion_aguas_andinas_pdf import listar_periodos_desde_pdf
from generar_reporte_word import convertir_word_a_pdf, format_number_chilean

ROOT = Path(__file__).resolve().parent
IN_DIR = (
    ROOT
    / "reports"
    / "Renca"
    / "Coparacion App con Aguas Andinas"
    / "Gimnasio Facturaciones"
)
OUT_DIR = (
    ROOT
    / "reports"
    / "Renca"
    / "Coparacion App con Aguas Andinas"
    / "reporte_comparacion_Gimnasio"
)

XLSX_INFORME = "Comparacion_Facturacion_vs_WES_Gimnasio_Renca.xlsx"
NODE_ID = "000017-05"
NODE_NOMBRE = "Gimnasio municipal (000017-05)"
EMPRESA = "Renca"

# Si existe en IN_DIR, tiene prioridad sobre PDF (p. ej. PDF escaneado sin capa de texto).
CSV_FACTURACION_MANUAL = IN_DIR / "facturacion_aguas_andinas.csv"

# Ventana de transición (calendario facturación, mismo criterio que otros informes Renca)
TRANSICION_INICIO = datetime(2025, 12, 11).date()
TRANSICION_FIN = datetime(2026, 1, 8).date()

# Narrativa operacional (Gimnasio); cuadros comparativos usan INFORME_COMPARATIVO_DESDE (misma fecha aquí).
MONITOREO_WES_DESDE = datetime(2025, 6, 6)
INFORME_COMPARATIVO_DESDE = date(2025, 6, 6)


def _ahora_informe_local() -> datetime:
    """Momento de generación en la zona horaria local del equipo (metadatos Word coherentes con el reloj)."""
    return datetime.now(timezone.utc).astimezone()


def _doc_establecer_metadatos(doc: Document, generado_en: datetime) -> None:
    cp = doc.core_properties
    cp.created = generado_en
    cp.modified = generado_en
    cp.title = "Comparación App WES vs Facturaciones Aguas Andinas"
    cp.subject = f"{NODE_ID} — {NODE_NOMBRE} — {EMPRESA}"
    cp.keywords = "WES; Aguas Andinas; Gimnasio municipal; Renca; facturación; consumo"
    cp.category = "Informe comparativo"
    cp.comments = (
        "Informe generado automáticamente: facturas Aguas Andinas (PDF/CSV/OCR) vs consumo API nodo WES."
    )
    cp.author = "Script informe Gimnasio municipal Renca"
    cp.last_modified_by = cp.author
    try:
        cp.revision = 1
    except Exception:
        pass


def _docx_parchear_core_hora_local(out_docx: Path, generado_en: datetime) -> None:
    """python-docx guarda created/modified como UTC; Word en Chile resta el huso. Reescribimos core.xml con offset local."""
    if generado_en.tzinfo is None:
        generado_en = generado_en.astimezone()
    marca = generado_en.isoformat(timespec="seconds")

    with ZipFile(out_docx, "r") as zin:
        nombres = zin.namelist()
        if "docProps/core.xml" not in nombres:
            return
        core = zin.read("docProps/core.xml").decode("utf-8")

    def _reemplaza(tag: str, xml: str) -> str:
        pat = re.compile(
            rf"(<dcterms:{tag}\s[^>]*>)([^<]*)(</dcterms:{tag}>)",
            re.DOTALL,
        )
        if not pat.search(xml):
            return xml
        return pat.sub(rf"\g<1>{marca}\g<3>", xml, count=1)

    core2 = _reemplaza("created", core)
    core2 = _reemplaza("modified", core2)
    if core2 == core:
        return

    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    tmp_path = Path(tmp)
    try:
        with ZipFile(out_docx, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "docProps/core.xml":
                    data = core2.encode("utf-8")
                zout.writestr(info, data)
        shutil.move(str(tmp_path), str(out_docx))
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass


def _fmt_num(n: float, dec: int = 1) -> str:
    return format_number_chilean(n, dec)


_MESES = {
    "ENE": 1,
    "JAN": 1,
    "FEB": 2,
    "MAR": 3,
    "ABR": 4,
    "APR": 4,
    "MAY": 5,
    "JUN": 6,
    "JUL": 7,
    "AGO": 8,
    "AUG": 8,
    "SEP": 9,
    "OCT": 10,
    "NOV": 11,
    "DIC": 12,
    "DEC": 12,
}


def _parse_fecha_es_dd_mmm_yyyy(s: str) -> datetime:
    p = s.strip().upper().split("-")
    if len(p) != 3:
        raise ValueError(f"Fecha inválida: {s!r}")
    dd = int(p[0])
    mm = _MESES[p[1][:3]]
    yy = int(p[2])
    return datetime(yy, mm, dd)


def _parse_fecha_flexible(s: str) -> datetime:
    s = s.strip()
    if not s:
        raise ValueError("Fecha vacía")
    p = s.split("-")
    if len(p) == 3 and p[1].isdigit() and p[0].isdigit() and p[2].isdigit():
        dd, mm, yy = int(p[0]), int(p[1]), int(p[2])
        return datetime(yy, mm, dd)
    return _parse_fecha_es_dd_mmm_yyyy(s)


def _fmt_periodo(dt: datetime) -> str:
    meses = {
        1: "ene",
        2: "feb",
        3: "mar",
        4: "abr",
        5: "may",
        6: "jun",
        7: "jul",
        8: "ago",
        9: "sep",
        10: "oct",
        11: "nov",
        12: "dic",
    }
    return f"{dt.day:02d}-{meses[dt.month]}-{dt.year}"


def _to_ddmmyyyy(dt: datetime) -> str:
    return dt.strftime("%d%m%Y")


@dataclass(frozen=True)
class FacturaAA:
    pdf: Path
    boleta: str
    cuenta: str | None
    medidor: str | None
    emision: datetime
    lectura_anterior: datetime
    lectura_actual: datetime
    m3_cuenta: int

    @property
    def api_start(self) -> str:
        return _to_ddmmyyyy(self.lectura_anterior)

    @property
    def api_end(self) -> str:
        return _to_ddmmyyyy(self.lectura_actual)

    @property
    def periodo_txt(self) -> str:
        return f"{_fmt_periodo(self.lectura_anterior)} → {_fmt_periodo(self.lectura_actual)}"


def _norm_csv_key(k: str) -> str:
    return re.sub(r"\s+", "_", (k or "").strip().lower())


def _cargar_facturas_desde_csv(path: Path) -> list[FacturaAA]:
    """Filas desde CSV con columnas boleta, fecha_emision, lectura_anterior, lectura_actual, m3_cuenta."""
    out: list[FacturaAA] = []
    with path.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        if not reader.fieldnames:
            return out
        remap = {_norm_csv_key(k): k for k in reader.fieldnames if k}
        req = ["boleta", "fecha_emision", "lectura_anterior", "lectura_actual", "m3_cuenta"]
        for col in req:
            if col not in remap:
                raise ValueError(
                    f"CSV {path.name}: falta columna «{col}». Encabezados: {reader.fieldnames}"
                )
        for row in reader:
            raw_boleta = (row.get(remap["boleta"]) or "").strip()
            if not raw_boleta or raw_boleta.startswith("#"):
                continue
            boleta = raw_boleta
            emision = _parse_fecha_flexible((row.get(remap["fecha_emision"]) or "").strip())
            la = _parse_fecha_flexible((row.get(remap["lectura_anterior"]) or "").strip())
            lb = _parse_fecha_flexible((row.get(remap["lectura_actual"]) or "").strip())
            m3_raw = (row.get(remap["m3_cuenta"]) or "").strip().replace(".", "").replace(",", ".")
            m3_cuenta = int(float(m3_raw))
            cuenta = None
            medidor = None
            if "cuenta" in remap:
                v = (row.get(remap["cuenta"]) or "").strip()
                cuenta = v or None
            if "medidor" in remap:
                v = (row.get(remap["medidor"]) or "").strip()
                medidor = v or None
            out.append(
                FacturaAA(
                    pdf=path,
                    boleta=boleta,
                    cuenta=cuenta,
                    medidor=medidor,
                    emision=emision,
                    lectura_anterior=la,
                    lectura_actual=lb,
                    m3_cuenta=m3_cuenta,
                )
            )
    out.sort(key=lambda r: r.lectura_actual)
    return out


def _cargar_facturas() -> list[FacturaAA]:
    if CSV_FACTURACION_MANUAL.is_file():
        rows = _cargar_facturas_desde_csv(CSV_FACTURACION_MANUAL)
        if not rows:
            raise ValueError(
                f"El archivo {CSV_FACTURACION_MANUAL.name} no tiene filas de datos válidas "
                "(revise encabezados y que haya al menos una boleta)."
            )
        return rows

    pdfs = sorted(IN_DIR.glob("*.pdf"))
    if not pdfs:
        guia = IN_DIR / "COMO_LLENAR_ESTA_CARPETA.txt"
        extra = f"\n\nLea: {guia.name} (en esta misma carpeta)." if guia.is_file() else ""
        raise FileNotFoundError(
            f"No hay {CSV_FACTURACION_MANUAL.name} ni archivos .pdf en:\n{IN_DIR}\n\n"
            "Coloque boletas PDF descargadas de Aguas Andinas (con texto), o bien cree "
            f"{CSV_FACTURACION_MANUAL.name} —puede partir de plantilla_facturacion_aguas_andinas.csv— "
            "con una fila por factura (ver docstring del script generar_reporte_comparacion_app_vs_facturaciones_gimnasio_000017_05.py)."
            + extra
        )
    errores: list[str] = []
    rows: list[FacturaAA] = []
    for p in pdfs:
        try:
            for per in listar_periodos_desde_pdf(p):
                rows.append(
                    FacturaAA(
                        pdf=per.pdf,
                        boleta=per.boleta,
                        cuenta=per.cuenta,
                        medidor=per.medidor,
                        emision=per.emision,
                        lectura_anterior=per.lectura_anterior,
                        lectura_actual=per.lectura_actual,
                        m3_cuenta=per.m3_cuenta,
                    )
                )
        except ValueError as e:
            errores.append(f"  - {p.name}: {e}")
    if not rows:
        detalle = "\n".join(errores) if errores else "(sin detalle)"
        raise ValueError(
            "No se pudieron extraer períodos de facturación desde ningún PDF.\n"
            f"{detalle}\n\n"
            f"Si el archivo no es una boleta AA ni un historial de consumo, use:\n  {CSV_FACTURACION_MANUAL}\n"
            "con columnas: boleta,fecha_emision,lectura_anterior,lectura_actual,m3_cuenta"
        )
    rows.sort(key=lambda r: r.lectura_actual)
    return rows


def _particionar_por_corte_informe(
    facturas: list[FacturaAA], desde: date
) -> tuple[list[FacturaAA], list[FacturaAA]]:
    """Incluye solo boletas con lectura anterior y lectura actual ≥ ``desde`` (cuadro comparable App vs cuenta)."""
    incl: list[FacturaAA] = []
    excl: list[FacturaAA] = []
    for f in facturas:
        d0 = f.lectura_anterior.date()
        d1 = f.lectura_actual.date()
        if d0 >= desde and d1 >= desde:
            incl.append(f)
        else:
            excl.append(f)
    return incl, excl


def _parse_celda_fecha_excel(val) -> datetime | None:
    if val is None or val == "":
        return None
    if isinstance(val, datetime):
        return val
    s = str(val).strip().replace("/", "-")[:10]
    for fmt in ("%d-%m-%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            continue
    return None


def _emision_tabla_desde_excel(emision_raw, lec_fin: datetime) -> str:
    dt: datetime | None = None
    if isinstance(emision_raw, datetime):
        dt = emision_raw
    elif emision_raw is not None and str(emision_raw).strip():
        dt = _parse_celda_fecha_excel(emision_raw)
    if dt is None:
        dt = lec_fin
    s = dt.strftime("%d-%b-%Y").lower()
    return s.replace("apr", "abr").replace("aug", "ago").replace("dec", "dic")


def _leer_filas_desde_excel_comparacion(path: Path) -> list[dict] | None:
    """Lee la hoja del Excel exportado (mismas columnas que exportar_excel...). Retorna None si falla."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return None
    if not path.is_file():
        return None
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        out: list[dict] = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row or row[0] is None:
                continue
            t = tuple(row) + (None,) * 20
            c0 = str(t[0] or "").strip()
            # Layout nuevo (export 2025+): Nodo, PDF, cuenta, medidor, periodo, emisión, boleta, …
            if c0.startswith("000017-"):
                periodo = str(t[4] or "").strip()
                lec_ini = _parse_celda_fecha_excel(t[7])
                lec_fin = _parse_celda_fecha_excel(t[9])
                if lec_fin is None:
                    continue
                boleta = str(t[6] or "").strip() if t[6] is not None else ""
                emision_raw = t[5]
                m3_cuenta = int(round(float(t[12])))
                m3_wes = float(t[13])
                dias_wes = int(t[14])
                diff = float(t[15])
                medidor_c = (str(t[3]).strip() if t[3] is not None and str(t[3]).strip() else "-")
                pct_raw = t[16]
                pct_diff = (
                    float(pct_raw)
                    if pct_raw is not None and str(pct_raw).strip() != ""
                    else ((100.0 * float(diff) / float(m3_cuenta)) if m3_cuenta else 0.0)
                )
            else:
                periodo = str(t[0]).strip()
                lec_ini = _parse_celda_fecha_excel(t[3])
                lec_fin = _parse_celda_fecha_excel(t[5])
                if lec_fin is None:
                    continue
                boleta = str(t[2]).strip() if t[2] is not None else ""
                emision_raw = t[1]
                m3_cuenta = int(round(float(t[8])))
                m3_wes = float(t[9])
                dias_wes = int(t[10])
                diff = float(t[11])
                medidor_c = "-"
                pct_diff = (100.0 * float(diff) / float(m3_cuenta)) if m3_cuenta else 0.0
            out.append(
                {
                    "periodo": periodo,
                    "periodo_short": _fmt_periodo(lec_fin),
                    "lectura_fin": lec_fin,
                    "lectura_ini": lec_ini if lec_ini is not None else lec_fin,
                    "emision": _emision_tabla_desde_excel(emision_raw, lec_fin),
                    "boleta": boleta or "-",
                    "medidor": medidor_c,
                    "m3_cuenta": m3_cuenta,
                    "m3_wes": float(m3_wes),
                    "dias_wes": int(dias_wes),
                    "diff": float(diff),
                    "pct": float(pct_diff),
                }
            )
        wb.close()
        return out if out else None
    except Exception:
        return None


def _split_filas_dict_por_transicion(filas: list[dict]) -> tuple[list[dict], list[dict], list[dict]]:
    """Misma regla A/B/C que con objetos FacturaAA (fechas de período)."""
    filas_bypass: list[dict] = []
    filas_transicion: list[dict] = []
    filas_alineado: list[dict] = []
    for r in filas:
        lec_fin = r["lectura_fin"]
        if not isinstance(lec_fin, datetime):
            continue
        li = r.get("lectura_ini")
        d0 = li.date() if isinstance(li, datetime) else lec_fin.date()
        d1 = lec_fin.date()
        overlap_trans = not (d1 < TRANSICION_INICIO or d0 > TRANSICION_FIN)
        if overlap_trans:
            filas_transicion.append(r)
        elif d1 < TRANSICION_INICIO:
            filas_bypass.append(r)
        else:
            filas_alineado.append(r)
    return filas_bypass, filas_transicion, filas_alineado


def _fetch_wes_total(f: FacturaAA) -> tuple[float, int]:
    from generar_reporte_word import (
        acl_node_base_url,
        fetch_json,
        normalize_measures_payload,
        flatten_measures,
        summarize_consumption,
    )

    start = f.api_start
    end_api = _to_ddmmyyyy(f.lectura_actual + timedelta(days=1))
    base = acl_node_base_url()
    raw = fetch_json(
        f"{base}/nodes/measures/dates",
        params=[("id", NODE_ID), ("start", start), ("end", end_api)],
    )
    norm = normalize_measures_payload(raw, NODE_ID)
    meas = flatten_measures(norm)
    d0, d1 = f.lectura_anterior.date(), f.lectura_actual.date()
    meas = [m for m in meas if d0 <= m.date.date() <= d1]
    s = summarize_consumption(meas)
    return float(s.get("total", 0.0)), int(s.get("dias", 0))


def _generar_grafico_barras(out_png: Path, filas: list[dict]) -> None:
    """Barras agrupadas con ``ax.bar``; dos azules. Izq = app WES, der = facturación."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib import patheffects as pe

    rc = {
        "font.family": "sans-serif",
        "font.sans-serif": ["Segoe UI", "Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 10.5,
        "axes.labelsize": 11,
        "axes.edgecolor": "#94a3b8",
        "axes.linewidth": 0.8,
        "xtick.labelsize": 9.5,
        "ytick.labelsize": 10,
        "legend.fontsize": 10,
        "figure.facecolor": "#f8fafc",
        "axes.facecolor": "#ffffff",
    }

    if not filas:
        return

    cuenta = np.array([float(f["m3_cuenta"]) for f in filas], dtype=float)
    wes = np.array([float(f["m3_wes"]) for f in filas], dtype=float)

    labels_x = [f'N° {f["boleta"]}\n{f["periodo_short"]}\n(izq.: app WES · der.: fact.)' for f in filas]

    col_wes = "#38bdf8"
    col_fact = "#1e3a8f"
    edge = "#ffffff"

    with matplotlib.rc_context(rc):
        fig, ax = plt.subplots(figsize=(13.5, 6.8))
        n = len(filas)
        xg = np.arange(n, dtype=float)
        bw = 0.36
        ymax = float(max(float(cuenta.max()), float(wes.max()), 1.0)) * 1.2

        bar_w = ax.bar(
            xg - bw / 2,
            wes,
            bw,
            color=col_wes,
            edgecolor=edge,
            linewidth=1.15,
            label="Consumo según app WES (m³)",
            zorder=3,
            clip_on=False,
        )
        bar_f = ax.bar(
            xg + bw / 2,
            cuenta,
            bw,
            color=col_fact,
            edgecolor=edge,
            linewidth=1.15,
            label="Consumo facturado Aguas Andinas (m³)",
            zorder=2,
            clip_on=False,
        )

        def _etiqueta_arriba(bar, vals: np.ndarray, *, zeros: bool = False) -> None:
            rects = getattr(bar, "patches", None) or tuple(bar)
            for rect, val in zip(rects, vals):
                hgt = rect.get_height()
                if val <= 0 and zeros:
                    zm = ax.annotate(
                        _fmt_num(float(val), 0 if abs(val) >= 10 else 1),
                        xy=(rect.get_x() + rect.get_width() / 2, 0),
                        xytext=(0, 11),
                        textcoords="offset points",
                        ha="center",
                        va="bottom",
                        fontsize=8.5,
                        color="#64748b",
                        zorder=5,
                    )
                    zm.set_path_effects([pe.withStroke(linewidth=3, foreground="white")])
                    continue
                if hgt <= 0:
                    continue
                t = ax.annotate(
                    _fmt_num(float(val), 0 if val >= 10 else 1),
                    xy=(rect.get_x() + rect.get_width() / 2, hgt),
                    xytext=(0, 6),
                    textcoords="offset points",
                    ha="center",
                    va="bottom",
                    fontsize=9.5,
                    color="#0f172a",
                    fontweight="600",
                    zorder=5,
                )
                t.set_path_effects([pe.withStroke(linewidth=3.5, foreground="white")])

        _etiqueta_arriba(bar_w, wes, zeros=True)
        _etiqueta_arriba(bar_f, cuenta, zeros=False)

        ax.set_ylim(0, ymax)
        ax.set_xticks(xg)
        ax.set_xticklabels(labels_x, ha="center", linespacing=1.35)
        ax.set_xlabel("Factura y fin de período (lectura actual)", color="#475569", labelpad=12)
        ax.set_ylabel("Consumo total del período (m³)", color="#475569", labelpad=8)

        ax.set_axisbelow(True)
        ax.yaxis.grid(True, linestyle="--", linewidth=0.7, color="#cbd5e1", alpha=0.9, zorder=0)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.legend(loc="upper right", framealpha=0.97, fancybox=True, edgecolor="#e2e8f0")

        fig.suptitle(
            "Por factura: app WES (izquierda) vs facturación (derecha)",
            fontsize=14,
            fontweight="600",
            color="#0f172a",
            y=0.97,
        )
        fig.text(
            0.5,
            0.935,
            f"Gimnasio municipal · nodo {NODE_ID} — lecturas anterior y actual ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}",
            ha="center",
            fontsize=10,
            color="#64748b",
        )

        fig.subplots_adjust(bottom=0.26, left=0.07, right=0.98, top=0.82)
        fig.savefig(out_png, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor(), edgecolor="none")
        plt.close(fig)


def _generar_grafico_totales(
    out_png: Path,
    *,
    total_wes: float,
    total_cuenta: float,
    n_boletas: int,
) -> None:
    """Dos barras: total App WES vs total facturación (suma de boletas incluidas)."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib import patheffects as pe

    rc = {
        "font.family": "sans-serif",
        "font.sans-serif": ["Segoe UI", "Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 11,
        "axes.labelsize": 11,
        "axes.edgecolor": "#94a3b8",
        "figure.facecolor": "#f8fafc",
        "axes.facecolor": "#ffffff",
    }
    col_wes = "#38bdf8"
    col_fact = "#1e3a8f"
    labels = ["Total App WES\n(suma m³)", "Total facturación\ncuenta (suma m³)"]
    vals = np.array([float(total_wes), float(total_cuenta)], dtype=float)
    x = np.arange(2, dtype=float)
    ymax = float(max(vals.max(), 1.0)) * 1.22

    with matplotlib.rc_context(rc):
        fig, ax = plt.subplots(figsize=(8.2, 5.4))
        bars = ax.bar(x, vals, width=0.52, color=[col_wes, col_fact], edgecolor="#ffffff", linewidth=1.2, zorder=2)
        for rect, val in zip(bars, vals):
            if val <= 0:
                continue
            t = ax.annotate(
                _fmt_num(float(val), 1),
                xy=(rect.get_x() + rect.get_width() / 2, float(val)),
                xytext=(0, 6),
                textcoords="offset points",
                ha="center",
                va="bottom",
                fontsize=11,
                fontweight="600",
                color="#0f172a",
                zorder=5,
            )
            t.set_path_effects([pe.withStroke(linewidth=3.5, foreground="white")])

        ax.set_xticks(x)
        ax.set_xticklabels(labels, ha="center", linespacing=1.25)
        ax.set_ylabel("m³ acumulados", color="#475569", labelpad=10)
        ax.set_ylim(0, ymax)
        ax.set_axisbelow(True)
        ax.yaxis.grid(True, linestyle="--", linewidth=0.7, color="#cbd5e1", alpha=0.9, zorder=0)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        fig.suptitle(
            "Total acumulado: App WES vs facturación en cuenta — Gimnasio municipal",
            fontsize=14,
            fontweight="600",
            color="#0f172a",
            y=0.96,
        )
        fig.text(
            0.5,
            0.90,
            (
                f"Nodo {NODE_ID} — {n_boletas} boleta(s) — lecturas ≥ "
                f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}"
            ),
            ha="center",
            fontsize=10,
            color="#64748b",
        )
        fig.subplots_adjust(bottom=0.18, left=0.12, right=0.97, top=0.82)
        fig.savefig(out_png, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor(), edgecolor="none")
        plt.close(fig)


def generar_reporte() -> tuple[Path, Path | None]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ahora = _ahora_informe_local()
    run_id = ahora.strftime("%Y%m%d_%H%M%S")

    facturas_todas = _cargar_facturas()
    facturas, facturas_excl = _particionar_por_corte_informe(facturas_todas, INFORME_COMPARATIVO_DESDE)
    if not facturas:
        raise ValueError(
            f"No hay facturas con lectura anterior y lectura actual ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}."
        )

    xlsx_comp = OUT_DIR / XLSX_INFORME
    filas_precalc = _leer_filas_desde_excel_comparacion(xlsx_comp)

    doc = Document()
    _doc_establecer_metadatos(doc, ahora)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Comparación App WES vs Facturaciones Aguas Andinas", level=1)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    p = doc.add_paragraph()
    p.add_run("Empresa: ").bold = True
    p.add_run(EMPRESA)
    p2 = doc.add_paragraph()
    p2.add_run("Nodo: ").bold = True
    p2.add_run(f"{NODE_ID} — {NODE_NOMBRE}")
    p3 = doc.add_paragraph()
    p3.add_run("Fuente facturas: ").bold = True
    p3.add_run(str(IN_DIR))
    p4 = doc.add_paragraph()
    p4.add_run("Generado: ").bold = True
    p4.add_run(ahora.strftime("%d-%m-%Y %H:%M"))

    doc.add_paragraph()
    doc.add_heading("1. Alcance y criterio de comparación", level=1)
    doc.add_paragraph(
        "Para cada boleta, se compara el consumo total facturado (m³) con el consumo total registrado por WES "
        "en el nodo indicado, para el período entre «Lectura anterior» y «Lectura actual» informados en la factura. "
        "Las diferencias pueden explicarse por alcances hidráulicos distintos (medidor principal vs subred), "
        "pérdidas/derivaciones no medidas por el punto, y diferencias de sincronización/calendario del período."
    )
    doc.add_paragraph(
        "En el Gimnasio municipal (nodo 000017-05) es frecuente que el consumo facturado por Aguas Andinas "
        "no coincida de forma directa con el volumen que muestra la app WES. La explicación principal es de "
        "alcance de medición: el medidor asociado a la cuenta de Aguas Andinas suele alimentar redes o grupos "
        "colectivos (otros usuarios o tramos de red compartida) antes de que el agua llegue al edificio o sector "
        "donde está instalado el punto monitoreado por WES. Por tanto, la facturación puede reflejar un caudal "
        "acumulado o repartido de forma distinta al flujo que atraviesa el ámbito del sensor en el gimnasio; ello "
        "se traduce en diferencias de m³ y en desfases temporales, sin que por sí solo implique error en la boleta "
        "ni falla del equipo WES en su propia zona de medición. Los consumos de cuenta se obtienen de boletas "
        "electrónicas o del historial de consumo en PDF (incluido OCR si el archivo es escaneado), o de un CSV "
        "manual en la carpeta de facturas cuando el PDF no es parseable. Para la evaluación con sentido operacional "
        f"se considera el monitoreo de referencia a partir del {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}. "
        "Existe un período de transición (11-dic-2025 a 08-ene-2026) donde se observan ajustes y estabilización en "
        "el calendario de facturación; desde el fin de esa transición en adelante, la comparación se lee como "
        "referencia principal (aun así, pueden existir diferencias por el alcance hidráulico descrito). "
        "Si se dispone del Excel comparativo exportado para este sitio, las cifras del informe pueden tomarse de "
        "esa hoja para homogeneizar números con operaciones."
    )
    doc.add_paragraph(
        f"Las tablas, gráficos por período y el total acumulado de la sección 2 incluyen únicamente boletas cuya "
        f"«Lectura anterior» y «Lectura actual» son el {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')} o posteriores. "
        "En ciclos anteriores o con alguna lectura previa a esa fecha, la cuenta suele reflejar liquidación por "
        "promedio estimado u homologación sin contraste de punta a punta con el punto WES en ese tramo."
    )

    doc.add_heading(
        f"1.1 Boletas fuera del cuadro comparativo (lecturas anteriores al {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})",
        level=2,
    )
    if facturas_excl:
        doc.add_paragraph(
            "Las siguientes boletas no entran en las tablas 2 ni en los gráficos ni en el total acumulado; "
            "suelen corresponder a promedio estimado o períodos no homologables con el monitoreo WES en el tramo indicado."
        )
        for f in sorted(facturas_excl, key=lambda x: x.lectura_actual):
            doc.add_paragraph(
                f"N° {f.boleta} — {f.periodo_txt} (m³ facturados en boleta: {f.m3_cuenta})",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No hay boletas excluidas por este criterio.")

    doc.add_heading(
        f"2. Resumen por período de factura (lecturas ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})",
        level=1,
    )
    headers = [
        "Período (lecturas)",
        "Emisión",
        "N° factura",
        "Medidor",
        "m³ facturado",
        "m³ App (WES)",
        "Días (WES)",
        "Dif (fact - WES)",
        "% Dif vs Fact",
    ]
    MAX_FILAS_TABLA_POR_PAGINA = 10

    filas_bypass: list[FacturaAA] = []
    filas_transicion: list[FacturaAA] = []
    filas_alineado: list[FacturaAA] = []
    for f in facturas:
        d0 = f.lectura_anterior.date()
        d1 = f.lectura_actual.date()
        overlap_trans = not (d1 < TRANSICION_INICIO or d0 > TRANSICION_FIN)
        if overlap_trans:
            filas_transicion.append(f)
        elif d1 < TRANSICION_INICIO:
            filas_bypass.append(f)
        else:
            filas_alineado.append(f)

    def _filas_dict(fs: list[FacturaAA]) -> list[dict]:
        out: list[dict] = []
        for f in fs:
            m3_wes, dias = _fetch_wes_total(f)
            diff = float(f.m3_cuenta) - float(m3_wes)
            pct_diff = (100.0 * float(diff) / float(f.m3_cuenta)) if f.m3_cuenta else 0.0
            out.append(
                {
                    "periodo": f.periodo_txt,
                    "periodo_short": _fmt_periodo(f.lectura_actual),
                    "lectura_fin": f.lectura_actual,
                    "emision": f.emision.strftime("%d-%b-%Y")
                    .lower()
                    .replace("apr", "abr")
                    .replace("aug", "ago")
                    .replace("dec", "dic"),
                    "boleta": f.boleta,
                    "medidor": f.medidor or "-",
                    "m3_cuenta": int(f.m3_cuenta),
                    "m3_wes": float(m3_wes),
                    "dias_wes": int(dias),
                    "diff": float(diff),
                    "pct": float(pct_diff),
                }
            )
        return out

    if filas_precalc:
        filas_bypass_d, filas_transicion_d, filas_alineado_d = _split_filas_dict_por_transicion(filas_precalc)
    else:
        filas_bypass_d = _filas_dict(filas_bypass)
        filas_transicion_d = _filas_dict(filas_transicion)
        filas_alineado_d = _filas_dict(filas_alineado)

    def _tbl_pr_set_full_width_fixed(table) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        for tag_nm in ("tblW", "tblLayout"):
            tq = qn(f"w:{tag_nm}")
            for el in list(tbl_pr):
                if el.tag == tq:
                    tbl_pr.remove(el)
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), "5000")
        tbl_w.set(qn("w:type"), "pct")
        tbl_pr.append(tbl_w)
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), "fixed")
        tbl_pr.append(tbl_layout)

    def _tbl_set_repeat_header(table) -> None:
        try:
            tr = table.rows[0]._tr
        except Exception:
            return
        tr_pr = tr.get_or_add_trPr()
        for el in list(tr_pr):
            if el.tag == qn("w:tblHeader"):
                tr_pr.remove(el)
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        tr_pr.append(hdr)

    def _add_table(title_txt: str, rows: list[dict]) -> None:
        p_tit = doc.add_paragraph()
        p_tit.add_run(title_txt).bold = True
        p_tit.paragraph_format.keep_with_next = True
        if not rows:
            doc.add_paragraph("Sin facturas en este rango.")
            return

        chunks: list[list[dict]] = [
            rows[i : i + MAX_FILAS_TABLA_POR_PAGINA]
            for i in range(0, len(rows), MAX_FILAS_TABLA_POR_PAGINA)
        ]
        for chunk_idx, chunk_rows in enumerate(chunks):
            if chunk_idx > 0:
                doc.add_page_break()
                p_ct = doc.add_paragraph()
                p_ct.add_run(f"{title_txt} (continuación)").bold = True
                p_ct.paragraph_format.keep_with_next = True

            table = doc.add_table(rows=1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            _tbl_pr_set_full_width_fixed(table)
            _tbl_set_repeat_header(table)
            hdr = table.rows[0].cells
            for i, h in enumerate(headers):
                rr = hdr[i].paragraphs[0].add_run(h)
                rr.bold = True
                rr.font.size = Pt(9)
            for row in chunk_rows:
                cells = table.add_row().cells
                cells[0].text = row["periodo"]
                cells[1].text = row["emision"]
                cells[2].text = str(row["boleta"])
                cells[3].text = str(row["medidor"])
                cells[4].text = _fmt_num(float(row["m3_cuenta"]), 0)
                cells[5].text = _fmt_num(float(row["m3_wes"]), 1)
                cells[6].text = str(row["dias_wes"])
                cells[7].text = _fmt_num(float(row["diff"]), 1)
                cells[8].text = _fmt_num(float(row["pct"]), 1) + "%"

            table.allow_autofit = False
            for r in table.rows:
                try:
                    r.allow_break_across_pages = False
                except Exception:
                    pass
                for c in r.cells:
                    for pp in c.paragraphs:
                        pp.paragraph_format.space_before = Pt(0)
                        pp.paragraph_format.space_after = Pt(0)
                        for run in pp.runs:
                            run.font.size = Pt(9)

    p_a = doc.add_paragraph(
        "Tabla A — Períodos con lectura final anterior al 11-dic-2025 (calendario de facturación), solo lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}: etapa previa a la ventana de transición; interpretar con "
        "cuidado si hubo arranque de monitoreo o lecturas estimadas. El alcance del medidor de cuenta puede incluir "
        "consumo de tramos compartidos antes del edificio del gimnasio (véase apartado 1)."
    )
    p_a.paragraph_format.keep_with_next = True
    _add_table("Períodos pre-transición (lectura final antes del 11-dic-2025)", filas_bypass_d)

    doc.add_paragraph()
    p_b = doc.add_paragraph(
        "Tabla B — Períodos de transición (11-dic-2025 a 08-ene-2026), solo lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}: ciclos con ajustes y estabilización."
    )
    p_b.paragraph_format.keep_with_next = True
    _add_table("Períodos de transición (11-dic-2025 a 08-ene-2026)", filas_transicion_d)

    doc.add_paragraph()
    p_c = doc.add_paragraph(
        "Tabla C — Períodos alineados (posterior a 08-ene-2026), solo lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}: referencia principal; «% Dif vs Fact» respecto a la facturación."
    )
    p_c.paragraph_format.keep_with_next = True
    _add_table("Períodos alineados (posterior a 08-ene-2026)", filas_alineado_d)

    doc.add_paragraph()
    filas_graf = [dict(r) for r in filas_bypass_d + filas_transicion_d + filas_alineado_d]
    filas_graf.sort(key=lambda r: r["lectura_fin"])
    for r in filas_graf:
        r.pop("lectura_fin", None)
    img_barras = OUT_DIR / f"chart_barras_gimnasio_todos_{run_id}.png"
    _generar_grafico_barras(img_barras, filas_graf)
    doc.add_paragraph(
        "Gráfico consolidado — todas las boletas del cuadro, orden cronológico por lectura actual; "
        "izquierda App WES, derecha facturación. "
        f"Solo lecturas ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}."
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if filas_graf:
        doc.add_picture(str(img_barras), width=Inches(6.7))

    def _grafico_categoria(titulo: str, filas_d: list[dict], sufijo: str) -> None:
        if not filas_d:
            return
        copia = [dict(r) for r in filas_d]
        copia.sort(key=lambda r: r["lectura_fin"])
        for r in copia:
            r.pop("lectura_fin", None)
        img = OUT_DIR / f"chart_barras_gimnasio_{sufijo}_{run_id}.png"
        _generar_grafico_barras(img, copia)
        doc.add_paragraph()
        pg = doc.add_paragraph()
        pg.add_run(titulo).bold = True
        doc.add_picture(str(img), width=Inches(6.7))

    _grafico_categoria("Gráfico — Tabla A (bypass)", filas_bypass_d, "bypass")
    _grafico_categoria("Gráfico — Tabla B (transición)", filas_transicion_d, "transicion")
    _grafico_categoria("Gráfico — Tabla C (alineados)", filas_alineado_d, "alineados")

    doc.add_paragraph()
    doc.add_heading(
        f"2.3 Total acumulado — cuenta (facturación) vs App WES (≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})",
        level=2,
    )
    txt_acum = (
        f"Suma de todos los períodos incluidos en el cuadro (lecturas anterior y actual ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}). "
        "Total en cuenta: suma del m³ facturado en esas boletas; total App: suma del consumo WES entre las mismas "
        "lecturas de cada período."
    )
    if filas_precalc:
        txt_acum += (
            f" En esta corrida, esos totales y las filas del cuadro provienen del archivo «{XLSX_INFORME}» "
            "(mismos números que la hoja exportada), sin volver a consultar la API al armar el Word."
        )
    doc.add_paragraph(txt_acum)
    rows_acum = filas_bypass_d + filas_transicion_d + filas_alineado_d
    n_boletas_acum = len(rows_acum)
    tot_cuenta = sum(float(r["m3_cuenta"]) for r in rows_acum)
    tot_wes = sum(float(r["m3_wes"]) for r in rows_acum)
    tot_diff = tot_cuenta - tot_wes
    pct_vs_fact = (100.0 * tot_diff / tot_cuenta) if tot_cuenta else 0.0

    ttot = doc.add_table(rows=1, cols=2)
    ttot.style = "Table Grid"
    ttot.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_pr_set_full_width_fixed(ttot)
    h0 = ttot.rows[0].cells
    h0[0].paragraphs[0].add_run("Indicador").bold = True
    h0[1].paragraphs[0].add_run("Valor").bold = True
    for lab, val in (
        ("Boletas incluidas en el acumulado (N°)", str(n_boletas_acum)),
        ("Total m³ facturación (cuenta, suma boletas)", _fmt_num(tot_cuenta, 1)),
        ("Total m³ App WES (suma períodos)", _fmt_num(tot_wes, 1)),
        ("Diferencia (facturación − App WES)", _fmt_num(tot_diff, 1)),
        ("% diferencia vs facturación", _fmt_num(pct_vs_fact, 1) + "%"),
    ):
        row_cells = ttot.add_row().cells
        row_cells[0].text = lab
        row_cells[1].text = val
    for r in ttot.rows:
        try:
            r.allow_break_across_pages = False
        except Exception:
            pass
        for c in r.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after = Pt(0)
                for run in pp.runs:
                    run.font.size = Pt(9)

    img_tot = OUT_DIR / f"chart_barras_gimnasio_total_acumulado_{run_id}.png"
    _generar_grafico_totales(
        img_tot,
        total_wes=tot_wes,
        total_cuenta=tot_cuenta,
        n_boletas=n_boletas_acum,
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Gráfico — total acumulado: barra izquierda suma App WES, barra derecha suma facturación en cuenta "
        f"(misma base que la tabla; solo facturaciones con ambas lecturas ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})."
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if n_boletas_acum > 0:
        doc.add_picture(str(img_tot), width=Inches(5.8))

    doc.add_heading("3. Referencias", level=1)
    doc.add_paragraph(
        "Facturas incluidas en tablas, gráficos y total acumulado (lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}), por lectura actual:\n- "
        + "\n- ".join([str(f.boleta) for f in sorted(facturas, key=lambda x: x.lectura_actual)])
    )

    out_docx = OUT_DIR / "Comparacion_App_vs_Facturaciones_Gimnasio_Renca.docx"
    if out_docx.exists():
        try:
            out_docx.unlink()
        except OSError:
            pass
    doc.save(out_docx)
    _docx_parchear_core_hora_local(out_docx, ahora)

    out_pdf = out_docx.with_suffix(".pdf")
    if out_pdf.exists():
        try:
            out_pdf.unlink()
        except OSError:
            pass
    out_pdf = convertir_word_a_pdf(out_docx)
    return out_docx, out_pdf


def main() -> None:
    try:
        OUT_DIR.mkdir(parents=True, exist_ok=True)
        print("[1/2] Generando Excel comparativo (base de datos para tablas y gráficos del informe)…")
        r = subprocess.run(
            [
                sys.executable,
                str(ROOT / "exportar_excel_comparacion_facturacion_vs_wes.py"),
                "--site",
                "gimnasio",
            ],
            cwd=str(ROOT),
        )
        if r.returncode != 0:
            print(
                "[ADVERTENCIA] No se generó el Excel. El informe usará consumos WES desde la API fila a fila si puede.\n"
                "  Revise facturas/CSV y ejecute: python exportar_excel_comparacion_facturacion_vs_wes.py --site gimnasio",
                file=sys.stderr,
            )
        else:
            xlsx = OUT_DIR / XLSX_INFORME
            if xlsx.is_file():
                print(f"[OK] XLSX: {xlsx}")

        print("[2/2] Generando informe Word, PDF y gráficos…")
        docx, pdf = generar_reporte()
    except FileNotFoundError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"[ERROR] {type(e).__name__}: {e}", file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)
    print(f"[OK] DOCX: {docx}")
    print(f"[OK] PDF: {pdf}" if pdf else "[ADVERTENCIA] No se generó PDF.")


if __name__ == "__main__":
    main()

