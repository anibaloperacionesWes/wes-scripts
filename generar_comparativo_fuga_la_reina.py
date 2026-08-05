"""
Comparativo fuga La Reina (Eugenio María de Hostos, 000024-01):
referencia del reporte agregado julio 2026 vs ventana reciente (p. ej. ayer 12:00 → ahora).

Uso:
  python generar_comparativo_fuga_la_reina.py
  python generar_comparativo_fuga_la_reina.py --desde-hora 12
"""

from __future__ import annotations

import argparse
import csv
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages

from generar_reporte_word import get_hourly_measures_for_day
from wes_estilo_graficos_app import guardar_grafico_horario_24h_app

CHILE = ZoneInfo("America/Santiago")
NODE_ID = "000024-01"
NODE_NAME = "Eugenio María de Hostos"
EMPRESA = "La Reina"
COMPANY_ID = "000024"
BASE_URL = "http://104.248.53.141:7003/wes/api/acl-node/v1"

# Referencia del reporte enviado (agregado julio 2026)
REF_JULIO_TOTAL_M3 = 21010.1
REF_JULIO_NOCTURNO_M3 = 5945.6
REF_JULIO_NOCTURNO_PCT = 28.3
REF_CAMBIO_NOCTURNO = "14/07/2026"
REF_CONTROL_OFF = "27/07/2026"
REF_BASELINE_DIA_M3 = 20.0  # aprox. 1–13 jul / junio


def _set_run_font(run, size: int = 11, bold: bool = False, color: Optional[RGBColor] = None) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if color is not None:
        run.font.color.rgb = color


def obtener_live() -> Dict:
    r = requests.get(f"{BASE_URL}/nodes/{NODE_ID}", timeout=30)
    r.raise_for_status()
    return r.json()


def obtener_diarios(inicio: datetime, fin: datetime) -> List[Dict]:
    r = requests.get(
        f"{BASE_URL}/nodes/measures/dates",
        params={
            "id": NODE_ID,
            "start": inicio.strftime("%d%m%Y"),
            "end": fin.strftime("%d%m%Y"),
        },
        timeout=60,
    )
    r.raise_for_status()
    data = r.json()
    if isinstance(data, list) and data:
        return data[0].get("month") or []
    if isinstance(data, dict):
        return data.get("month") or []
    return []


def obtener_alertas(inicio: datetime, fin: datetime) -> List[Dict]:
    r = requests.get(
        f"{BASE_URL}/nodes/myalert/alerts",
        params={
            "id": NODE_ID,
            "start": inicio.strftime("%d%m%Y"),
            "end": fin.strftime("%d%m%Y"),
        },
        timeout=60,
    )
    if r.status_code != 200:
        return []
    data = r.json()
    return data if isinstance(data, list) else []


def horario_dia(dia: datetime) -> Dict[int, float]:
    pairs = get_hourly_measures_for_day(NODE_ID, dia)
    return {int(h): float(v) for h, v in pairs}


def serie_desde_hasta(
    desde: datetime, hasta: datetime
) -> List[Tuple[datetime, float]]:
    """Serie horaria Chile desde ``desde`` inclusive hasta ``hasta`` (hora actual)."""
    out: List[Tuple[datetime, float]] = []
    d = desde.replace(minute=0, second=0, microsecond=0)
    fin_hora = hasta.replace(minute=0, second=0, microsecond=0)
    cache: Dict[str, Dict[int, float]] = {}
    while d <= fin_hora:
        key = d.strftime("%Y-%m-%d")
        if key not in cache:
            cache[key] = horario_dia(d)
        val = cache[key].get(d.hour)
        if val is not None:
            out.append((d, float(val)))
        d += timedelta(hours=1)
    return out


def grafico_ventana(
    serie: List[Tuple[datetime, float]], out_path: Path, titulo: str
) -> Path:
    fig, ax = plt.subplots(figsize=(11, 4.8))
    fig.patch.set_facecolor("white")
    xs = list(range(len(serie)))
    ys = [v for _, v in serie]
    labels = [t.strftime("%d/%m\n%H:00") for t, _ in serie]
    colors = ["#c41e1e" if t.hour <= 6 else "#4A8CB8" for t, _ in serie]
    ax.bar(xs, ys, color=colors, width=0.85, edgecolor="white", linewidth=0.3)
    ax.plot(xs, ys, color="#2e7ac8", linewidth=1.1, marker="o", markersize=3)
    ax.set_title(titulo, fontsize=12, fontweight="bold")
    ax.set_ylabel("m³/h (aprox. volumen de la hora)")
    step = max(1, len(xs) // 12)
    ax.set_xticks(xs[::step])
    ax.set_xticklabels([labels[i] for i in xs[::step]], fontsize=7)
    ax.set_ylim(0, max(ys + [0.1]) * 1.15)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=140)
    plt.close(fig)
    return out_path


def grafico_diarios_contexto(
    diarios: List[Dict], out_path: Path, titulo: str
) -> Path:
    fechas = [d["date"][5:] for d in diarios]  # MM-DD
    vals = [float(d["totalM3"]) for d in diarios]
    fig, ax = plt.subplots(figsize=(11, 4.2))
    fig.patch.set_facecolor("white")
    ax.bar(range(len(vals)), vals, color="#4A8CB8", edgecolor="white")
    ax.axhline(
        REF_BASELINE_DIA_M3,
        color="#2e7d32",
        linestyle="--",
        linewidth=1.4,
        label=f"Baseline pre-fuga ≈ {REF_BASELINE_DIA_M3:.0f} m³/día",
    )
    ax.set_title(titulo, fontsize=12, fontweight="bold")
    ax.set_ylabel("m³/día")
    step = max(1, len(fechas) // 10)
    ax.set_xticks(range(0, len(fechas), step))
    ax.set_xticklabels([fechas[i] for i in range(0, len(fechas), step)], fontsize=8)
    ax.legend(loc="upper left", fontsize=8)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out_path, dpi=140)
    plt.close(fig)
    return out_path


def fmt_m3(v: float) -> str:
    return f"{v:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")


def construir_docx(
    out_docx: Path,
    desde: datetime,
    hasta: datetime,
    serie: List[Tuple[datetime, float]],
    diarios: List[Dict],
    alertas: List[Dict],
    live: Dict,
    chart_ventana: Path,
    chart_contexto: Path,
    chart_ayer: Path,
    chart_hoy: Path,
) -> Path:
    vol_ventana = sum(v for _, v in serie)
    horas_con_dato = len(serie)
    # desde 12 ayer: subtramos tramos
    ayer = (hasta - timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)
    hoy = hasta.replace(hour=0, minute=0, second=0, microsecond=0)
    vol_ayer_desde = sum(v for t, v in serie if t.date() == ayer.date())
    vol_hoy = sum(v for t, v in serie if t.date() == hoy.date())
    max_h = max(serie, key=lambda x: x[1]) if serie else None
    noct_ventana = sum(v for t, v in serie if t.hour <= 6)

    mch = live.get("mch")
    lpm = live.get("lpm")
    wes_status = live.get("wesStatus")
    last_upd = live.get("lastUpdate")

    # Diarios recientes vs baseline
    ultimos = diarios[-7:] if len(diarios) >= 7 else diarios
    promedio_reciente = (
        sum(float(d["totalM3"]) for d in ultimos) / len(ultimos) if ultimos else 0.0
    )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading(
        "Comparativo fuga — La Reina (Eugenio María de Hostos)", level=0
    )
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    _set_run_font(p.add_run("MONITOREO WES — seguimiento operativo"), 12, True)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Punto: {NODE_NAME} ({NODE_ID}) · Cliente: {EMPRESA} ({COMPANY_ID})\n"
        f"Ventana analizada: {desde:%d-%m-%Y %H:%M} → {hasta:%d-%m-%Y %H:%M} (hora Chile)\n"
        f"Generado: {hasta:%d-%m-%Y %H:%M}"
    )

    doc.add_heading("1. Qué decía el reporte enviado (julio 2026)", level=1)
    doc.add_paragraph(
        "El reporte agregado de La Reina (01–31/07/2026) documentó un evento de fuga / "
        "pérdida anómala en Eugenio María de Hostos:"
    )
    bullets = [
        f"Consumo total del mes: {fmt_m3(REF_JULIO_TOTAL_M3)} m³ (un solo colegio).",
        f"Consumo nocturno (00:00–06:59): {fmt_m3(REF_JULIO_NOCTURNO_M3)} m³ "
        f"({REF_JULIO_NOCTURNO_PCT:.1f} % del total).",
        f"Baseline pre-fuga (1–13/07 y junio): ~{REF_BASELINE_DIA_M3:.0f} m³/día.",
        f"Cambio marcado en nocturno desde el {REF_CAMBIO_NOCTURNO}.",
        f"El {REF_CONTROL_OFF} se desactivó el control, lo que explica el alza posterior.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("2. Qué está pasando ahora (ventana pedida)", level=1)
    doc.add_paragraph(
        f"Desde el {desde:%d-%m-%Y %H:%M} hasta {hasta:%d-%m-%Y %H:%M} el punto acumula "
        f"aproximadamente {fmt_m3(vol_ventana)} m³ en {horas_con_dato} horas con dato "
        f"(suma de caudales horarios m³/h ≈ volumen de cada hora)."
    )
    detalle = [
        f"Volumen el {ayer:%d-%m} desde las {desde.hour:02d}:00: {fmt_m3(vol_ayer_desde)} m³.",
        f"Volumen el {hoy:%d-%m} hasta ahora: {fmt_m3(vol_hoy)} m³.",
        f"Nocturno dentro de la ventana (horas 00–06): {fmt_m3(noct_ventana)} m³.",
    ]
    if max_h:
        detalle.append(
            f"Máximo horario en la ventana: {fmt_m3(max_h[1])} m³/h a las "
            f"{max_h[0]:%d-%m %H:00}."
        )
    detalle.append(
        f"Lectura en vivo al generar: caudal {mch} m³/h · {lpm} L/min · "
        f"estado {wes_status} · última actualización {last_upd}."
    )
    for b in detalle:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3. Comparación con la referencia de fuga", level=1)
    factor = (promedio_reciente / REF_BASELINE_DIA_M3) if REF_BASELINE_DIA_M3 else 0
    doc.add_paragraph(
        f"Promedio de los últimos {len(ultimos)} días con dato en API: "
        f"{fmt_m3(promedio_reciente)} m³/día, frente a ~{REF_BASELINE_DIA_M3:.0f} m³/día "
        f"antes de la fuga (≈ {factor:.0f}× el baseline)."
    )
    doc.add_paragraph(
        "Lectura operativa de la ventana reciente: el {0:%d-%m} aún muestra caudales "
        "muy altos en la mañana (mismo orden de magnitud del evento de julio). "
        "Desde las 13:00 del {0:%d-%m} el caudal cae a ~1–3 m³/h y se mantiene "
        "continuo; la madrugada de hoy aparece en cero en la serie horaria, pero "
        "al momento de este informe el punto vuelve a registrar caudal (~{1} m³/h).".format(
            ayer, mch
        )
    )

    if alertas:
        doc.add_paragraph("Alertas MyAlert en el entorno de la ventana / días previos:")
        for a in alertas[-8:]:
            doc.add_paragraph(
                f"{a.get('creationDate', '')} · medida {a.get('measure')} · "
                f"stream {a.get('stream')}",
                style="List Bullet",
            )

    doc.add_heading("4. Gráficos", level=1)
    doc.add_paragraph("Perfil horario — ventana desde ayer 12:00 hasta ahora:")
    if chart_ventana.exists():
        doc.add_picture(str(chart_ventana), width=Inches(6.3))
    doc.add_paragraph("Contexto diario (vs baseline pre-fuga):")
    if chart_contexto.exists():
        doc.add_picture(str(chart_contexto), width=Inches(6.3))
    doc.add_paragraph(f"Perfil 24 h — {ayer:%d-%m-%Y} (día completo API):")
    if chart_ayer.exists():
        doc.add_picture(str(chart_ayer), width=Inches(6.3))
    doc.add_paragraph(f"Perfil 24 h — {hoy:%d-%m-%Y} (parcial hasta ahora):")
    if chart_hoy.exists():
        doc.add_picture(str(chart_hoy), width=Inches(6.3))

    doc.add_heading("5. Conclusión", level=1)
    doc.add_paragraph(
        "Respecto del reporte de julio (fuga / pérdida anómala en Hostos): el evento "
        "no está cerrado. Los días recientes siguen órdenes de magnitud por sobre el "
        "baseline (~20 m³/día). En la ventana desde ayer 12:00 se ve una baja brusca "
        "del caudal extremo al mediodía, residual continuo en la tarde-noche, madrugada "
        "de hoy en cero en la serie, y reaparición de caudal esta mañana. "
        "Se recomienda verificar en terreno si hubo cierre parcial de válvula / "
        "reactivación de control y confirmar si la fuga residual (~1–2 m³/h) sigue activa."
    )

    doc.add_paragraph(
        "Nota: volumen horario ≈ valor m³/h de la API × 1 h. Totales diarios vienen de "
        "/nodes/measures/dates (totalM3)."
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    return out_docx


def guardar_pdf_resumen(
    out_pdf: Path,
    desde: datetime,
    hasta: datetime,
    serie: List[Tuple[datetime, float]],
    live: Dict,
    chart_ventana: Path,
    chart_contexto: Path,
) -> Path:
    vol = sum(v for _, v in serie)
    with PdfPages(out_pdf) as pdf:
        fig = plt.figure(figsize=(11.69, 8.27))
        ax = fig.add_subplot(111)
        ax.axis("off")
        lines = [
            "COMPARATIVO FUGA — LA REINA",
            f"{NODE_NAME} ({NODE_ID})",
            "",
            f"Ventana: {desde:%d-%m-%Y %H:%M} → {hasta:%d-%m-%Y %H:%M} (Chile)",
            f"Volumen en ventana: {fmt_m3(vol)} m³",
            f"Live: {live.get('mch')} m³/h · {live.get('lpm')} L/min · {live.get('wesStatus')}",
            "",
            "Referencia reporte julio 2026:",
            f"  Total mes: {fmt_m3(REF_JULIO_TOTAL_M3)} m³ · Nocturno: {fmt_m3(REF_JULIO_NOCTURNO_M3)} m³",
            f"  Baseline pre-fuga ≈ {REF_BASELINE_DIA_M3:.0f} m³/día",
            f"  Cambio nocturno desde {REF_CAMBIO_NOCTURNO}; control off {REF_CONTROL_OFF}",
            "",
            "Conclusión: la fuga reportada en julio sigue impactando; la ventana reciente",
            "muestra baja brusca ~13:00 de ayer, residual continuo, y caudal activo hoy.",
        ]
        y = 0.92
        for i, line in enumerate(lines):
            weight = "bold" if i < 2 else "normal"
            size = 16 if i < 2 else 11
            ax.text(0.05, y, line, fontsize=size, fontweight=weight, va="top", family="DejaVu Sans")
            y -= 0.055 if i > 1 else 0.07
        pdf.savefig(fig)
        plt.close(fig)

        for chart, title in (
            (chart_ventana, "Ventana horaria"),
            (chart_contexto, "Contexto diario"),
        ):
            if not chart.exists():
                continue
            fig = plt.figure(figsize=(11.69, 8.27))
            ax = fig.add_subplot(111)
            ax.axis("off")
            ax.set_title(title)
            img = plt.imread(str(chart))
            ax.imshow(img)
            pdf.savefig(fig)
            plt.close(fig)
    return out_pdf


def main() -> int:
    ap = argparse.ArgumentParser(description="Comparativo fuga La Reina vs ventana actual")
    ap.add_argument("--desde-hora", type=int, default=12, help="Hora inicio ayer (default 12)")
    args = ap.parse_args()

    ahora = datetime.now(CHILE)
    ayer = ahora - timedelta(days=1)
    desde = ayer.replace(hour=args.desde_hora, minute=0, second=0, microsecond=0)

    ts = ahora.strftime("%Y%m%d_%H%M")
    out_dir = (
        Path(__file__).resolve().parent
        / "reports"
        / "La_Reina"
        / "COMPARATIVO_FUGA"
        / f"comparativo_{ts}"
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"Ventana: {desde} → {ahora}")
    live = obtener_live()
    print("Live:", live.get("mch"), "m3/h", live.get("lpm"), "L/min")

    serie = serie_desde_hasta(desde, ahora)
    print(f"Horas con dato: {len(serie)} · volumen ≈ {sum(v for _, v in serie):.1f} m³")

    diarios = obtener_diarios(
        datetime(2026, 7, 1, tzinfo=CHILE),
        ahora,
    )
    alertas = obtener_alertas(desde - timedelta(days=5), ahora)

    chart_ventana = grafico_ventana(
        serie,
        out_dir / "chart_ventana_desde_ayer12.png",
        f"Hostos — desde {desde:%d-%m %H:%M} hasta {ahora:%d-%m %H:%M}",
    )
    chart_contexto = grafico_diarios_contexto(
        diarios,
        out_dir / "chart_contexto_diario.png",
        "Hostos — consumo diario vs baseline pre-fuga (~20 m³/día)",
    )

    horas_ayer = {i: 0.0 for i in range(24)}
    horas_ayer.update(horario_dia(desde.replace(hour=0)))
    chart_ayer = guardar_grafico_horario_24h_app(
        horas_ayer,
        out_dir / "chart_24h_ayer.png",
        titulo=f"{NODE_NAME} — {desde:%d-%m-%Y}",
    )
    horas_hoy = {i: 0.0 for i in range(24)}
    horas_hoy.update(horario_dia(ahora.replace(hour=0)))
    chart_hoy = guardar_grafico_horario_24h_app(
        horas_hoy,
        out_dir / "chart_24h_hoy.png",
        titulo=f"{NODE_NAME} — {ahora:%d-%m-%Y} (parcial)",
    )

    out_docx = out_dir / f"Comparativo_Fuga_La_Reina_Hostos_{ts}.docx"
    construir_docx(
        out_docx,
        desde,
        ahora,
        serie,
        diarios,
        alertas,
        live,
        chart_ventana,
        chart_contexto,
        chart_ayer,
        chart_hoy,
    )
    print("DOCX:", out_docx)

    out_pdf = out_dir / f"Comparativo_Fuga_La_Reina_Hostos_{ts}.pdf"
    guardar_pdf_resumen(
        out_pdf, desde, ahora, serie, live, chart_ventana, chart_contexto
    )
    print("PDF:", out_pdf)

    # CSV detalle
    csv_path = out_dir / f"serie_horaria_{ts}.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["fecha_hora_chile", "m3_h"])
        for t, v in serie:
            w.writerow([t.strftime("%Y-%m-%d %H:%M"), f"{v:.3f}"])
    print("CSV:", csv_path)

    # resumen texto
    resumen = out_dir / "resumen.txt"
    resumen.write_text(
        "\n".join(
            [
                f"Ventana: {desde:%Y-%m-%d %H:%M} → {ahora:%Y-%m-%d %H:%M}",
                f"Volumen ventana m3: {sum(v for _, v in serie):.1f}",
                f"Live mch: {live.get('mch')} lpm: {live.get('lpm')}",
                f"DOCX: {out_docx}",
                f"PDF: {out_pdf}",
            ]
        ),
        encoding="utf-8",
    )
    print(resumen.read_text(encoding="utf-8"))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
