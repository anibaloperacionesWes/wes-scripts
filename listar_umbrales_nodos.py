# -*- coding: utf-8 -*-
"""
Listado de umbrales de alerta configurados en la API WES para todos los nodos
de una compañía (por defecto 000025 — Parque Arauco).

Fuente: GET /wes/api/acl-entities/v1/companies/{companyId}
        → nodes[].configuration.threshold

El umbral de la API es un único valor diario (m³/día): cuando el consumo
acumulado del día lo supera se dispara alerta. Valor 0 / vacío = sin umbral.

Uso:
  python listar_umbrales_nodos.py
  python listar_umbrales_nodos.py --company-id 000025
"""

from __future__ import annotations

import argparse
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Optional
from zoneinfo import ZoneInfo

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
import requests

from generar_reporte_word import (
    add_formatted_heading,
    add_logo_to_header,
    estilizar_tabla_wes,
    format_number_chilean,
    get_mall_name_for_parque_arauco,
)

if sys.platform == "win32":
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace", line_buffering=True)
        except Exception:
            pass

ENTITY_BASE = "http://104.248.53.141:7001/wes/api/acl-entities/v1"
TZ_CHILE = ZoneInfo("America/Santiago")
COMPANY_DEFAULT = "000025"

ESTADO_SIN = "Sin umbral"
ESTADO_OK = "Configurado"
ESTADO_INV = "Inválido"


def _parse_threshold(raw: Any) -> tuple[Optional[float], str, str]:
    """Retorna (valor_float_o_None, texto_crudo, estado)."""
    if raw is None:
        return None, "", ESTADO_SIN
    crudo = str(raw).strip()
    if crudo == "":
        return None, crudo, ESTADO_SIN
    norm = crudo.replace(",", ".")
    try:
        val = float(norm)
    except ValueError:
        return None, crudo, ESTADO_INV
    if val <= 0:
        return 0.0, crudo, ESTADO_SIN
    return val, crudo, ESTADO_OK


def listar_nodos(company_id: str) -> tuple[str, list[dict]]:
    url = f"{ENTITY_BASE}/companies/{company_id}"
    r = requests.get(url, timeout=60)
    r.raise_for_status()
    data = r.json()
    company_name = str(data.get("name") or company_id).strip() or company_id
    rows: list[dict] = []
    for node in data.get("nodes") or []:
        nid = str(node.get("nodeId") or "").strip()
        if not nid:
            continue
        name = str(node.get("name") or "").strip() or nid
        cfg = node.get("configuration") if isinstance(node.get("configuration"), dict) else {}
        raw = cfg.get("threshold")
        valor, crudo, estado = _parse_threshold(raw)
        mall = ""
        if company_id == "000025":
            mall = get_mall_name_for_parque_arauco(nid, name).strip()
        rows.append(
            {
                "nodeId": nid,
                "nodeName": name,
                "mall": mall or "—",
                "threshold_raw": crudo,
                "threshold": valor,
                "estado": estado,
            }
        )
    rows.sort(key=lambda x: (x["mall"] if x["mall"] != "—" else "zzz", x["nodeId"]))
    return company_name, rows


def _fmt_umbral(valor: Optional[float], estado: str) -> str:
    if estado == ESTADO_INV:
        return "N/D"
    if valor is None or (estado == ESTADO_SIN and (valor is None or valor <= 0)):
        return "0"
    return format_number_chilean(valor, 1)


def escribir_excel(rows: list[dict], path: Path, *, company_id: str, company_name: str, consultado: datetime) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = f"Umbrales {company_id}"
    headers = [
        "Node ID",
        "Punto",
        "Mall / recinto",
        "Umbral API (crudo)",
        "Umbral (m³/día)",
        "Estado",
    ]
    ws.append(headers)
    header_fill = PatternFill("solid", fgColor="003366")
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = header_fill
        c.alignment = Alignment(wrap_text=True, vertical="center")

    fill_sin = PatternFill("solid", fgColor="FFF2CC")
    fill_inv = PatternFill("solid", fgColor="FCE4D6")
    fill_ok = PatternFill("solid", fgColor="C6EFCE")
    alt = PatternFill("solid", fgColor="F2F6FC")

    for i, r in enumerate(rows, start=2):
        umbral_txt = _fmt_umbral(r["threshold"], r["estado"])
        ws.append(
            [
                r["nodeId"],
                r["nodeName"],
                r["mall"],
                r["threshold_raw"] if r["threshold_raw"] != "" else "(vacío)",
                umbral_txt,
                r["estado"],
            ]
        )
        fill = None
        if r["estado"] == ESTADO_INV:
            fill = fill_inv
        elif r["estado"] == ESTADO_SIN:
            fill = fill_sin
        elif r["estado"] == ESTADO_OK:
            fill = fill_ok
        elif i % 2 == 0:
            fill = alt
        if fill:
            for cell in ws[i]:
                cell.fill = fill
                cell.alignment = Alignment(vertical="center")

    ws.append([])
    ws.append(["Parámetros"])
    ws.append(["Compañía", f"{company_id} — {company_name}"])
    ws.append(["Consulta API", consultado.strftime("%d/%m/%Y %H:%M") + " (hora Chile)"])
    ws.append(["Endpoint", f"{ENTITY_BASE}/companies/{company_id}"])
    ws.append(["Campo", "nodes[].configuration.threshold"])
    ws.append(
        [
            "Interpretación",
            "Umbral diario de alerta de consumo máximo (m³/día). "
            "0 / vacío = sin umbral (no dispara alerta por este campo). "
            "Valor no numérico = inválido.",
        ]
    )
    counts = Counter(r["estado"] for r in rows)
    ws.append(["Nodos", len(rows)])
    ws.append([ESTADO_OK, counts.get(ESTADO_OK, 0)])
    ws.append([ESTADO_SIN, counts.get(ESTADO_SIN, 0)])
    ws.append([ESTADO_INV, counts.get(ESTADO_INV, 0)])

    widths = [14, 42, 16, 18, 16, 14]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.auto_filter.ref = f"A1:F{1 + len(rows)}"
    ws.freeze_panes = "A2"
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def escribir_word(rows: list[dict], path: Path, *, company_id: str, company_name: str, consultado: datetime) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    add_logo_to_header(doc)

    add_formatted_heading(
        doc,
        f"Listado de umbrales configurados — {company_name} ({company_id})",
        level=0,
    )
    p = doc.add_paragraph()
    p.add_run(f"Fecha de consulta: {consultado.strftime('%d/%m/%Y %H:%M')} (hora Chile)\n").bold = True
    p.add_run(
        f"Fuente: API WES acl-entities, GET /companies/{company_id} "
        f"(campo configuration.threshold de cada nodo).\n"
        f"Nodos listados: {len(rows)}."
    )

    add_formatted_heading(doc, "1. Qué representa este umbral", level=1)
    p1 = doc.add_paragraph(
        "Cada nodo tiene un umbral diario único (m³/día) en la configuración de la API. "
        "Cuando el consumo acumulado del día supera ese valor —independiente de la hora— "
        "se dispara alerta por correo. Un valor 0 o vacío significa que el punto no tiene "
        "umbral de máximo cargado: esa alerta no se activa."
    )
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    counts = Counter(r["estado"] for r in rows)
    add_formatted_heading(doc, "2. Resumen", level=1)
    table_res = doc.add_table(rows=5, cols=2)
    table_res.style = "Table Grid"
    res_rows = [
        ["Concepto", "Cantidad"],
        ["Nodos de la compañía", str(len(rows))],
        ["Con umbral configurado (> 0)", str(counts.get(ESTADO_OK, 0))],
        ["Sin umbral (0 o vacío)", str(counts.get(ESTADO_SIN, 0))],
        ["Valor inválido (no numérico)", str(counts.get(ESTADO_INV, 0))],
    ]
    for i, pair in enumerate(res_rows):
        table_res.rows[i].cells[0].text = pair[0]
        table_res.rows[i].cells[1].text = pair[1]
    estilizar_tabla_wes(table_res, has_total_row=False)
    doc.add_paragraph("")

    if counts.get(ESTADO_OK, 0) == 0:
        nota = doc.add_paragraph()
        nota.add_run(
            "Hallazgo: ningún nodo de esta compañía tiene un umbral de alerta diario "
            "mayor a cero en la API. Los puntos con valor inválido aparecen al final de la tabla."
        ).italic = True

    add_formatted_heading(doc, "3. Tabla por punto", level=1)
    headers = ["Node ID", "Punto", "Mall", "Umbral (m³/día)", "Estado"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    highlight: list[int] = []
    for r_i, r in enumerate(rows):
        vals = [
            r["nodeId"],
            r["nodeName"],
            r["mall"],
            _fmt_umbral(r["threshold"], r["estado"]),
            r["estado"],
        ]
        for c_i, val in enumerate(vals):
            table.rows[r_i + 1].cells[c_i].text = val
        if r["estado"] == ESTADO_INV:
            highlight.append(r_i + 1)
    estilizar_tabla_wes(table, highlight_rows=highlight, has_total_row=False)
    doc.add_paragraph("")

    add_formatted_heading(doc, "4. Notas", level=1)
    for b in [
        "Este listado refleja el umbral cargado hoy en la API, no una propuesta de recálculo (p. ej. promedio 90 días × 1,25).",
        "No se incluye el umbral de control nocturno del Excel operativo (HORARIOS CONTROL NOCTURNO.xlsx); ese es otro criterio, por horario.",
        "Puntos con valor crudo no numérico (p. ej. «string» o campo vacío distinto de 0) se marcan como inválidos y conviene corregirlos en configuración.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def escribir_pdf(rows: list[dict], path: Path, *, company_id: str, company_name: str, consultado: datetime) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    counts = Counter(r["estado"] for r in rows)
    headers = ["Node ID", "Punto", "Mall", "Umbral (m³/día)", "Estado"]
    table_data = [headers]
    for r in rows:
        table_data.append(
            [
                r["nodeId"],
                r["nodeName"],
                r["mall"],
                _fmt_umbral(r["threshold"], r["estado"]),
                r["estado"],
            ]
        )

    with PdfPages(path) as pdf:
        resumen = (
            f"Consulta: {consultado.strftime('%d/%m/%Y %H:%M')} hora Chile  |  "
            f"Nodos: {len(rows)}  |  Configurados: {counts.get(ESTADO_OK, 0)}  |  "
            f"Sin umbral: {counts.get(ESTADO_SIN, 0)}  |  Inválidos: {counts.get(ESTADO_INV, 0)}\n"
            "Fuente: API configuration.threshold (m³/día). Valor 0 / vacío = no dispara alerta por umbral."
        )
        page_size = 22
        body = table_data[1:]
        pages = [body[i : i + page_size] for i in range(0, len(body), page_size)] or [[]]

        for page_i, chunk in enumerate(pages):
            fig = plt.figure(figsize=(11.69, 8.27))
            ax = fig.add_subplot(111)
            ax.axis("off")
            if page_i == 0:
                ax.set_title(
                    f"Listado de umbrales — {company_name} ({company_id})",
                    fontsize=14,
                    fontweight="bold",
                    color="#003366",
                    pad=18,
                    loc="left",
                )
                ax.text(
                    0,
                    1.04,
                    resumen,
                    transform=ax.transAxes,
                    fontsize=8,
                    va="bottom",
                    ha="left",
                )
            else:
                ax.set_title(
                    f"Listado de umbrales — {company_name} ({company_id})  (cont.)",
                    fontsize=12,
                    fontweight="bold",
                    color="#003366",
                    pad=10,
                    loc="left",
                )
            tbl = ax.table(
                cellText=[headers] + chunk,
                loc="upper center",
                cellLoc="left",
                colWidths=[0.12, 0.38, 0.16, 0.16, 0.14],
            )
            tbl.auto_set_font_size(False)
            tbl.set_fontsize(7.5)
            tbl.scale(1, 1.35)
            for (row, col), cell in tbl.get_celld().items():
                cell.set_edgecolor("#CCCCCC")
                if row == 0:
                    cell.set_facecolor("#003366")
                    cell.set_text_props(color="white", fontweight="bold")
                else:
                    estado = chunk[row - 1][4] if row - 1 < len(chunk) else ""
                    if estado == ESTADO_INV:
                        cell.set_facecolor("#FCE4D6")
                    elif estado == ESTADO_SIN:
                        cell.set_facecolor("#FFF2CC")
                    elif estado == ESTADO_OK:
                        cell.set_facecolor("#C6EFCE")
                    elif row % 2 == 0:
                        cell.set_facecolor("#F2F6FC")
            fig.text(
                0.99,
                0.02,
                f"Página {page_i + 1} de {len(pages)}  |  WES — umbrales API {company_id}",
                ha="right",
                fontsize=7,
                color="#666666",
            )
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)


def generar(company_id: str) -> tuple[list[dict], Path, Path, Path]:
    consultado = datetime.now(TZ_CHILE)
    company_name, rows = listar_nodos(company_id)
    stamp = consultado.strftime("%Y%m%d_%H%M")
    if company_id == "000025":
        out_dir = Path("reports") / "Parque_Arauco" / "umbrales_consumo"
    else:
        out_dir = Path("reports") / company_id / "umbrales_consumo"
    out_dir.mkdir(parents=True, exist_ok=True)
    base = f"Listado_umbrales_API_{company_id}_{stamp}"
    xlsx = out_dir / f"{base}.xlsx"
    docx = out_dir / f"{base}.docx"
    pdf = out_dir / f"{base}.pdf"
    escribir_excel(rows, xlsx, company_id=company_id, company_name=company_name, consultado=consultado)
    escribir_word(rows, docx, company_id=company_id, company_name=company_name, consultado=consultado)
    escribir_pdf(rows, pdf, company_id=company_id, company_name=company_name, consultado=consultado)
    return rows, xlsx, docx, pdf


def main() -> int:
    parser = argparse.ArgumentParser(description="Listado de umbrales API por compañía")
    parser.add_argument("--company-id", default=COMPANY_DEFAULT, help="ID de compañía (default: 000025)")
    args = parser.parse_args()
    company_id = str(args.company_id).strip() or COMPANY_DEFAULT
    rows, xlsx, docx, pdf = generar(company_id)
    counts = Counter(r["estado"] for r in rows)
    print("=" * 72)
    print(f"LISTADO UMBRALES API — {company_id}")
    print("=" * 72)
    print(f"Nodos: {len(rows)}")
    print(f"  Configurados: {counts.get(ESTADO_OK, 0)}")
    print(f"  Sin umbral:   {counts.get(ESTADO_SIN, 0)}")
    print(f"  Inválidos:    {counts.get(ESTADO_INV, 0)}")
    print(f"Excel: {xlsx}")
    print(f"Word:  {docx}")
    print(f"PDF:   {pdf}")
    print("-" * 72)
    print(f"{'Node ID':<12} {'Mall':<14} {'Umbral':>10}  {'Estado':<14} Punto")
    for r in rows:
        print(
            f"{r['nodeId']:<12} {r['mall']:<14} {_fmt_umbral(r['threshold'], r['estado']):>10}  "
            f"{r['estado']:<14} {r['nodeName']}"
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
