"""
Edita el último Word agregado Fundo Zapallar:
1) Conserva contenido hasta el gráfico Riego Llenado ESVAL.
2) Elimina la sección «Día de mayor consumo diario» (páginas intermedias).
3) Inserta esquema hidráulico FZ y mantiene métricas / nocturno (cierre ejecutivo).
"""

from __future__ import annotations

import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", line_buffering=True)
    except Exception:
        pass

from generar_reporte_word import add_formatted_title, add_picture_with_pagination


def _latest_docx() -> Path:
    base = Path("reports/Fundo_Zapallar/ABREGADO")
    docs = sorted(
        base.rglob("Reporte_Agregado_Fundo_Zapallar_20260701_20260728.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not docs:
        raise FileNotFoundError("No hay Word de Fundo Zapallar julio 2026")
    return docs[0]


def _para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def _build_esquema(path: Path) -> Path:
    fig, ax = plt.subplots(figsize=(9.2, 5.2))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_title(
        "Esquema hidráulico Fundo Zapallar",
        fontsize=14,
        fontweight="bold",
        pad=12,
    )

    def box(x, y, w, h, text, color="#0050b3"):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.05,rounding_size=0.15",
            linewidth=1.4,
            edgecolor=color,
            facecolor=color,
            alpha=0.92,
        )
        ax.add_patch(patch)
        ax.text(
            x + w / 2,
            y + h / 2,
            text,
            ha="center",
            va="center",
            color="white",
            fontsize=9,
            fontweight="bold",
            wrap=True,
        )

    def arrow(x1, y1, x2, y2):
        ax.add_patch(
            FancyArrowPatch(
                (x1, y1),
                (x2, y2),
                arrowstyle="-|>",
                mutation_scale=14,
                linewidth=1.8,
                color="#333333",
            )
        )

    # Flujo: ESVAL -> Inferior -> Superior -> (Etapa 1-4 | Etapa 5)
    box(0.3, 2.9, 2.4, 1.2, "Matriz ESVAL\n(000027-01)")
    box(3.5, 2.9, 2.5, 1.2, "Estanque Inferior\n(000027-02)")
    box(6.8, 2.9, 2.5, 1.2, "Estanque Superior\n(carga desde\nestanque inferior)")

    box(6.2, 5.2, 2.4, 1.1, "Etapa N°1 al 4\n(000027-04)\n+ Etapas 1/2/3")
    box(9.2, 5.2, 2.3, 1.1, "Etapa N°5\n(000027-03)")

    arrow(2.7, 3.5, 3.5, 3.5)  # ESVAL -> Inferior
    arrow(6.0, 3.5, 6.8, 3.5)  # Inferior -> Superior
    # Superior split to two pipes
    arrow(8.05, 4.1, 7.4, 5.2)  # a Etapa 1-4
    arrow(8.05, 4.1, 10.2, 5.2)  # a Etapa 5

    ax.text(7.2, 4.55, "tubería", fontsize=8, color="#555555", ha="center")
    ax.text(9.3, 4.55, "tubería", fontsize=8, color="#555555", ha="center")

    ax.text(
        6.0,
        0.55,
        "Flujo: Matriz ESVAL carga Estanque Inferior → Estanque Superior → "
        "alimenta por una tubería Etapas 1–4 y por otra tubería Etapa 5.",
        ha="center",
        va="center",
        fontsize=9,
        color="#222222",
    )

    plt.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def main() -> None:
    docx_path = _latest_docx()
    print(f"[INFO] Editando: {docx_path}")
    doc = Document(str(docx_path))
    body = doc.element.body
    children = list(body)

    start_idx = end_idx = None
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] != "p":
            continue
        t = _para_text(child).upper()
        if start_idx is None and t.startswith("DÍA DE MAYOR CONSUMO DIARIO"):
            start_idx = i
        if start_idx is not None and (
            t.startswith("RESUMEN POR PUNTO") or t.startswith("MÉTRICAS POR PUNTO")
        ):
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        raise RuntimeError(f"No se encontró rango a eliminar ({start_idx}, {end_idx})")

    # También quitar párrafos vacíos previos al bloque si quedan sueltos
    n = 0
    for child in children[start_idx:end_idx]:
        body.remove(child)
        n += 1
    print(f"[OK] Eliminados {n} elementos (día de mayor consumo)")

    # Insertar esquema antes de RESUMEN POR PUNTO
    esquema_path = docx_path.parent / "esquema_hidraulico_fundo_zapallar.png"
    _build_esquema(esquema_path)

    # Agregar al final y mover antes de RESUMEN
    resumen = None
    for child in list(body):
        if child.tag.split("}")[-1] == "p":
            tu = _para_text(child).upper()
            if tu.startswith("RESUMEN POR PUNTO") or tu.startswith("MÉTRICAS POR PUNTO"):
                resumen = child
                break
    if resumen is None:
        raise RuntimeError("RESUMEN POR PUNTO no encontrado tras el recorte")

    before = len([c for c in body if c.tag.split("}")[-1] != "sectPr"])

    add_formatted_title(doc, "Esquema hidráulico Fundo Zapallar")
    expl = doc.add_paragraph(
        "La Matriz ESVAL carga el Estanque Inferior; desde ahí se carga el Estanque Superior, "
        "que alimenta el sistema en dos tuberías: una hacia Etapa N°1 al 4 (y submediciones "
        "Etapa 1, 2 y 3) y otra hacia Etapa N°5."
    )
    expl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in expl.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(11)
    add_picture_with_pagination(doc, str(esquema_path), Inches(6.2), keep_with_next=True)
    doc.add_paragraph("")

    after = len([c for c in body if c.tag.split("}")[-1] != "sectPr"])
    new_els = [c for c in list(body) if c.tag.split("}")[-1] != "sectPr"][-(after - before) :]
    for el in new_els:
        resumen.addprevious(el)

    try:
        doc.save(str(docx_path))
        out = docx_path
    except PermissionError:
        out = docx_path.with_name(docx_path.stem + "_ejecutivo.docx")
        doc.save(str(out))
        print(f"[ADVERTENCIA] Guardado como {out.name}")

    print(f"[OK] {out}")


if __name__ == "__main__":
    main()
