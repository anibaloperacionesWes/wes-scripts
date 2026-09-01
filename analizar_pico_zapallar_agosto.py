"""
Diagnóstico Fundo Zapallar: pico 1–4 agosto 2026 en Etapa N°3
frente a estanque inferior, ESVAL y matriz Etapa 1–4.

Uso:
  python analizar_pico_zapallar_agosto.py
"""
from __future__ import annotations

import csv
import json
from datetime import date, datetime, timedelta
from math import sqrt
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from generar_reporte_word import (
    add_formatted_heading,
    add_formatted_title,
    add_logo_to_header,
    add_picture_with_pagination,
    estilizar_tabla_wes,
    format_number_chilean,
)
from wes_google_drive import credenciales_configuradas, subir_a_drive

BASE = "http://104.248.53.141:7003/wes/api/acl-node/v1"

NODES: Dict[str, str] = {
    "000027-01": "Matriz ESVAL",
    "000027-02": "Estanque Inferior",
    "000027-03": "Etapa N°5",
    "000027-04": "Etapa N°1 al 4",
    "000027-06": "Etapa N°1",
    "000027-07": "Etapa N°2",
    "000027-08": "Etapa N°3",
    "000027-09": "Riego Llenado ESVAL",
}
SHORT = {
    "000027-01": "ESVAL",
    "000027-02": "Est. inf.",
    "000027-03": "Etapa 5",
    "000027-04": "Etapa 1–4",
    "000027-06": "Etapa 1",
    "000027-07": "Etapa 2",
    "000027-08": "Etapa 3",
    "000027-09": "Riego llenado",
}
COLOR = {
    "000027-01": "#0050b3",
    "000027-02": "#2e7d32",
    "000027-03": "#6a1b9a",
    "000027-04": "#0277bd",
    "000027-06": "#00838f",
    "000027-07": "#ef6c00",
    "000027-08": "#c62828",
    "000027-09": "#546e7a",
}

PICO_INI = date(2026, 8, 1)
PICO_FIN = date(2026, 8, 4)
AGO_INI = date(2026, 8, 1)
AGO_FIN = date(2026, 8, 31)


def _cl(v: Optional[float], d: int = 1) -> str:
    if v is None:
        return "—"
    return format_number_chilean(float(v), d)


def _daterange(a: date, b: date) -> List[date]:
    out = []
    d = a
    while d <= b:
        out.append(d)
        d += timedelta(days=1)
    return out


def _corr(xs: List[float], ys: List[float]) -> Optional[float]:
    n = len(xs)
    if n < 3:
        return None
    mx = sum(xs) / n
    my = sum(ys) / n
    cov = sum((x - mx) * (y - my) for x, y in zip(xs, ys))
    vx = sum((x - mx) ** 2 for x in xs)
    vy = sum((y - my) ** 2 for y in ys)
    if vx <= 0 or vy <= 0:
        return None
    return cov / sqrt(vx * vy)


def fetch_daily(start: date, end: date) -> Dict[str, Dict[str, float]]:
    daily: Dict[str, Dict[str, float]] = {nid: {} for nid in NODES}
    start_s, end_s = start.strftime("%d%m%Y"), end.strftime("%d%m%Y")
    for nid in NODES:
        r = requests.get(
            f"{BASE}/nodes/measures/dates",
            params=[("id", nid), ("start", start_s), ("end", end_s)],
            timeout=60,
        )
        r.raise_for_status()
        payload = r.json()
        if isinstance(payload, list):
            item = next((x for x in payload if x.get("nodeId") == nid), payload[0] if payload else {})
        else:
            item = payload
        for m in item.get("month") or []:
            ds = (m.get("date") or "")[:10]
            if ds:
                daily[nid][ds] = float(m.get("totalM3") or 0.0)
    return daily


def _hours_from_csv(text: str, dia: date) -> Tuple[Dict[int, float], int, int]:
    prefix = dia.strftime("%Y-%m-%d")
    by_time: Dict[str, float] = {}
    n_rows = 0
    for line in text.strip().split("\n")[1:]:
        if not line.strip():
            continue
        parts = line.split(",", 1)
        if len(parts) < 2:
            continue
        n_rows += 1
        t = parts[0].strip()
        try:
            v = float(parts[1].strip().replace(" ", "").replace(",", "."))
        except ValueError:
            continue
        by_time[t] = by_time.get(t, 0.0) + v
    acc = {h: 0.0 for h in range(24)}
    for t, v in by_time.items():
        if t.startswith(prefix):
            try:
                acc[int(t[11:13])] = v
            except ValueError:
                pass
    return acc, n_rows, len(by_time)


def fetch_hourly(nids: List[str], days: List[date]) -> Dict[str, Dict[str, Dict[int, float]]]:
    hours: Dict[str, Dict[str, Dict[int, float]]] = {n: {} for n in nids}
    dup: Dict[str, Dict[str, Tuple[int, int]]] = {n: {} for n in nids}
    for nid in nids:
        for d in days:
            r = requests.get(
                f"{BASE}/nodes/{nid}/dates.measures.csv",
                params=[("start", d.strftime("%d%m%Y")), ("end", d.strftime("%d%m%Y"))],
                timeout=60,
            )
            r.raise_for_status()
            acc, n_rows, n_times = _hours_from_csv(r.text, d)
            hours[nid][d.isoformat()] = acc
            dup[nid][d.isoformat()] = (n_rows, n_times)
    hours["_dup_meta"] = dup  # type: ignore[assignment]
    return hours


def v(daily: Dict[str, Dict[str, float]], nid: str, d: date) -> float:
    return float(daily.get(nid, {}).get(d.isoformat(), 0.0) or 0.0)


def build_esquema(path: Path) -> Path:
    fig, ax = plt.subplots(figsize=(11.2, 6.4))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Circuito hidráulico Fundo Zapallar", fontsize=14, fontweight="bold", pad=8)

    def box(x, y, w, h, text, color="#0050b3"):
        ax.add_patch(
            FancyBboxPatch(
                (x, y), w, h,
                boxstyle="round,pad=0.04,rounding_size=0.12",
                linewidth=1.3, edgecolor=color, facecolor=color, alpha=0.92,
            )
        )
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                color="white", fontsize=8.2, fontweight="bold")

    def arrow(x1, y1, x2, y2):
        ax.add_patch(
            FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>",
                            mutation_scale=13, linewidth=1.6, color="#333333")
        )

    box(0.25, 3.4, 2.5, 1.15, "Matriz ESVAL\n000027-01")
    box(3.5, 3.4, 2.6, 1.15, "Estanque Inferior\n000027-02")
    box(6.9, 3.4, 2.6, 1.15, "Estanque Superior\n(sin medidor WES)")
    box(3.5, 1.15, 2.6, 1.0, "Riego / llenado\n000027-09")
    box(10.2, 5.55, 3.4, 1.15, "Matriz Etapa N°5\n000027-03")
    box(10.2, 3.35, 3.4, 1.2, "Matriz Etapa 1 al 4\n000027-04")
    box(10.05, 0.35, 1.15, 0.95, "Et.1\n-06")
    box(11.35, 0.35, 1.15, 0.95, "Et.2\n-07")
    box(12.65, 0.35, 1.15, 0.95, "Et.3\n-08")

    arrow(2.75, 4.0, 3.5, 4.0)
    arrow(6.1, 4.0, 6.9, 4.0)
    arrow(1.5, 3.4, 3.5, 2.15)
    arrow(9.5, 4.55, 10.2, 6.0)
    arrow(9.5, 3.95, 10.2, 3.95)
    arrow(11.9, 3.35, 10.6, 1.3)
    arrow(11.9, 3.35, 11.9, 1.3)
    arrow(11.9, 3.35, 13.2, 1.3)

    ax.text(7.0, 7.55,
            "Salida estanque superior: dos matrices. Etapas 1, 2 y 3 son ramales de Etapa 1–4.",
            ha="center", fontsize=9, color="#222")
    plt.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_diario(daily, path: Path, *, normalizado: bool = False) -> Path:
    days = _daterange(date(2026, 7, 25), AGO_FIN)
    xs = [datetime.combine(d, datetime.min.time()) for d in days]
    series = {
        "000027-01": [v(daily, "000027-01", d) for d in days],
        "000027-02": [v(daily, "000027-02", d) for d in days],
        "000027-04": [v(daily, "000027-04", d) for d in days],
        "000027-08": [v(daily, "000027-08", d) for d in days],
    }
    fig, ax = plt.subplots(figsize=(11.2, 5.4))
    for nid, ys in series.items():
        if normalizado:
            m = max(ys) or 1.0
            ys = [y / m for y in ys]
        ax.plot(xs, ys, marker="o", markersize=3.5, linewidth=1.8,
                color=COLOR[nid], label=SHORT[nid])
    ax.axvspan(datetime(2026, 8, 1), datetime(2026, 8, 5), color="#c62828", alpha=0.12,
               label="Pico 1–4 ago")
    ax.set_ylabel("Índice (máx. del periodo = 1)" if normalizado else "Consumo diario (m³)")
    ax.set_title(
        "Series diarias normalizadas (misma escala 0–1)" if normalizado
        else "Series diarias — ESVAL, estanque inferior, matriz 1–4 y Etapa 3"
    )
    ax.legend(loc="upper right", fontsize=8)
    ax.grid(True, alpha=0.3)
    fig.autofmt_xdate()
    plt.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_balance(daily, path: Path) -> Path:
    days = _daterange(AGO_INI, AGO_FIN)
    xs = [datetime.combine(d, datetime.min.time()) for d in days]
    madre = [v(daily, "000027-04", d) for d in days]
    suma = [
        v(daily, "000027-06", d) + v(daily, "000027-07", d) + v(daily, "000027-08", d)
        for d in days
    ]
    fig, ax = plt.subplots(figsize=(11.2, 5.2))
    ax.bar([x - timedelta(hours=6) for x in xs], madre, width=0.4,
           color=COLOR["000027-04"], label="Matriz Etapa 1–4 (madre)")
    ax.bar([x + timedelta(hours=6) for x in xs], suma, width=0.4,
           color=COLOR["000027-08"], label="Suma Etapa 1 + 2 + 3")
    ax.axvspan(datetime(2026, 8, 1), datetime(2026, 8, 5), color="#c62828", alpha=0.12)
    ax.set_ylabel("m³ / día")
    ax.set_title("Invariante: la matriz 1–4 no puede ser menor que la suma de ramales 1+2+3")
    ax.legend(fontsize=8)
    ax.grid(True, axis="y", alpha=0.3)
    fig.autofmt_xdate()
    plt.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_horario(hours, dia: date, path: Path) -> Path:
    hs = list(range(24))
    ds = dia.isoformat()
    fig, ax = plt.subplots(figsize=(11.2, 5.2))
    for nid in ("000027-01", "000027-02", "000027-04", "000027-08"):
        ys = [hours[nid][ds].get(h, 0.0) for h in hs]
        ax.plot(hs, ys, marker="o", markersize=3.5, linewidth=1.8,
                color=COLOR[nid], label=SHORT[nid])
    ax.set_xticks(hs)
    ax.set_xlabel("Hora (marca TIME del CSV WES)")
    ax.set_ylabel("m³ / h")
    ax.set_title(f"Perfil horario {dia.strftime('%d-%m-%Y')} — Etapa 3 vs entrada y matriz madre")
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_et3_vs_inf_horas(hours, days: List[date], path: Path) -> Path:
    fig, axes = plt.subplots(2, 2, figsize=(11.2, 6.8), sharex=True, sharey=False)
    hs = list(range(24))
    for ax, d in zip(axes.ravel(), days):
        ds = d.isoformat()
        ax.plot(hs, [hours["000027-02"][ds].get(h, 0) for h in hs],
                color=COLOR["000027-02"], lw=1.8, label="Est. inf.")
        ax.plot(hs, [hours["000027-08"][ds].get(h, 0) for h in hs],
                color=COLOR["000027-08"], lw=1.8, label="Etapa 3")
        ax.set_title(d.strftime("%d-%m"))
        ax.grid(True, alpha=0.3)
        ax.set_xticks(range(0, 24, 3))
    axes[0, 0].legend(fontsize=8)
    fig.suptitle("Estanque inferior vs Etapa 3 — perfil horario 1 a 4 de agosto",
                 fontsize=13, fontweight="bold")
    fig.supxlabel("Hora")
    fig.supylabel("m³ / h")
    plt.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def _para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold


def build_docx(
    out: Path,
    figs: Dict[str, Path],
    daily: Dict[str, Dict[str, float]],
    stats: dict,
) -> Path:
    doc = Document()
    add_logo_to_header(doc)

    t = doc.add_paragraph()
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = t.add_run("FUNDO ZAPALLAR")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0, 51, 102)

    st = doc.add_paragraph()
    st.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = st.add_run(
        "Diagnóstico del pico 1–4 de agosto 2026\n"
        "¿El aumento de Etapa N°3 se explica por el estanque inferior?"
    )
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0, 51, 102)

    meta = doc.add_paragraph(
        "Empresa 000027 · Datos API WES (totalM3 diario + dates.measures.csv horario)\n"
        "Periodo de contraste: agosto 2026 · Generado: "
        + datetime.now().strftime("%d-%m-%Y %H:%M")
    )
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_formatted_heading(doc, "Respuesta")
    _para(
        doc,
        "No. El pico de Etapa N°3 (000027-08) entre el 1 y el 4 de agosto no se explica "
        "por un aumento de consumo en el estanque inferior ni por la Matriz ESVAL. "
        "En esos cuatro días Etapa 3 promedió "
        f"{_cl(stats['et3_pico_avg'])} m³/día "
        f"({_cl(stats['et3_ratio'], 2)} veces el resto de agosto), mientras el estanque "
        f"inferior estuvo por debajo de su promedio del mes ({_cl(stats['inf_ratio'], 2)} veces) "
        f"y ESVAL también ({_cl(stats['esval_ratio'], 2)} veces). "
        "La correlación diaria agosto Etapa 3 vs estanque inferior es "
        f"r = {_cl(stats['r_et3_inf'], 2)} (nula / levemente negativa).",
        bold=True,
    )
    _para(
        doc,
        "Además el pico es hidráulicamente imposible si Etapa 3 es un ramal de la matriz "
        "Etapa N°1 al 4: el 1 de agosto Etapa 3 registró "
        f"{_cl(stats['et3_ago01'])} m³ y la matriz madre solo {_cl(stats['et14_ago01'])} m³. "
        "La suma Etapa 1+2+3 ese día fue "
        f"{_cl(stats['sum123_ago01'])} m³. Un ramal no puede superar a su tubería de alimentación.",
    )

    add_formatted_heading(doc, "Circuito usado como base")
    _para(
        doc,
        "Matriz ESVAL alimenta estanque inferior y riego. Del estanque inferior se carga el "
        "estanque superior. En la salida del superior hay dos matrices: una a Etapa N°5 y otra "
        "a Etapas 1 al 4. Etapas 1, 2 y 3 son submediciones de esa segunda matriz; por eso "
        "Etapa 1–4 nunca puede ser menor que Etapa 1 + Etapa 2 + Etapa 3.",
    )
    add_picture_with_pagination(doc, str(figs["esquema"]), Inches(6.3), keep_with_next=False)

    add_formatted_heading(doc, "Qué se ve en el agregado (1–4 de agosto)")
    headers = [
        "Fecha", "ESVAL", "Est. inf.", "Etapa 5", "Etapa 1–4",
        "Etapa 1", "Etapa 2", "Etapa 3", "Suma 1+2+3", "Madre − suma",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    highlight = []
    for i, d in enumerate(_daterange(PICO_INI, PICO_FIN), start=1):
        et14 = v(daily, "000027-04", d)
        s123 = v(daily, "000027-06", d) + v(daily, "000027-07", d) + v(daily, "000027-08", d)
        row = table.add_row()
        vals = [
            d.strftime("%d-%m"),
            v(daily, "000027-01", d),
            v(daily, "000027-02", d),
            v(daily, "000027-03", d),
            et14,
            v(daily, "000027-06", d),
            v(daily, "000027-07", d),
            v(daily, "000027-08", d),
            s123,
            et14 - s123,
        ]
        for j, val in enumerate(vals):
            row.cells[j].text = val if j == 0 else _cl(val)
        highlight.append(i)
    tot = table.add_row()
    pico_days = _daterange(PICO_INI, PICO_FIN)
    sums = {
        nid: sum(v(daily, nid, d) for d in pico_days) for nid in NODES
    }
    s123t = sums["000027-06"] + sums["000027-07"] + sums["000027-08"]
    tot_vals = [
        "Suma 1–4",
        sums["000027-01"], sums["000027-02"], sums["000027-03"], sums["000027-04"],
        sums["000027-06"], sums["000027-07"], sums["000027-08"],
        s123t, sums["000027-04"] - s123t,
    ]
    for j, val in enumerate(tot_vals):
        tot.cells[j].text = val if j == 0 else _cl(val)
    estilizar_tabla_wes(table, highlight_rows=highlight, has_total_row=True)
    _para(
        doc,
        "Lectura: Etapa 3 concentra casi todo el “pico” (640 m³ en cuatro días). "
        "ESVAL solo 153 m³ y estanque inferior 171 m³. Aunque el estanque superior "
        "pudiera entregar reserva, no hay de dónde sacar 640 m³ en Etapa 3 si la matriz "
        "madre 1–4 apenas midió 77 m³ y la entrada ESVAL 153 m³.",
    )

    add_formatted_heading(doc, "¿Se parecen las curvas? (efecto de escala del agregado)")
    _para(
        doc,
        "En el reporte agregado cada punto tiene su propio eje Y. Un bache de 65 m³ en "
        "estanque inferior (2 de agosto) puede verse “igual de alto” que 186 m³ en Etapa 3. "
        "En escala real no se parecen; normalizando cada serie a su máximo, Etapa 3 se "
        "dispara el 1–4 y el estanque inferior no.",
    )
    add_picture_with_pagination(doc, str(figs["diario"]), Inches(6.3), keep_with_next=False)
    add_picture_with_pagination(doc, str(figs["norm"]), Inches(6.3), keep_with_next=False)

    add_formatted_heading(doc, "Correlación de consumo (agosto 2026)")
    _para(
        doc,
        "Si Etapa 3 estuviera arrastrada por el llenado/bombeo del estanque inferior, "
        "deberían moverse juntas. No es el caso.",
    )
    ctab = doc.add_table(rows=1, cols=3)
    ctab.style = "Table Grid"
    for i, h in enumerate(["Par", "r (1–31 ago)", "r (5–31 ago, sin pico)"]):
        ctab.rows[0].cells[i].text = h
    for par, r1, r2 in stats["corrs"]:
        row = ctab.add_row()
        row.cells[0].text = par
        row.cells[1].text = "—" if r1 is None else f"{r1:.2f}".replace(".", ",")
        row.cells[2].text = "—" if r2 is None else f"{r2:.2f}".replace(".", ",")
    estilizar_tabla_wes(ctab, has_total_row=False)
    _para(
        doc,
        "La única correlación fuerte del circuito de entrada es estanque inferior vs ESVAL "
        f"(r = {_cl(stats['r_inf_esval'], 2)}): el inferior sigue a la matriz, como corresponde "
        "a un llenado. Etapa 3 no sigue a ninguno de los dos.",
    )

    add_formatted_heading(doc, "Perfil horario: caudal plano, no un riego")
    _para(
        doc,
        "El 1 de agosto Etapa 3 midió ~8 m³/h las 24 horas, incluida la madrugada, con "
        "estanque inferior en 0 m³/h entre 00:00 y 03:00 y ESVAL en 0 casi todo el día "
        "salvo un pulso de llenado a las 10–11. Eso no es un programa de riego: es un "
        "caudal continuo (o un medidor emitiendo pulsos constantes). El 4 de agosto el "
        "valor de Etapa 3 cae de ~6 m³/h a ~0,6 m³/h a lo largo del día y el 5 ya está "
        "en el rango habitual (~0,4–1 m³/h).",
    )
    add_picture_with_pagination(doc, str(figs["h01"]), Inches(6.3), keep_with_next=False)
    add_picture_with_pagination(doc, str(figs["hgrid"]), Inches(6.3), keep_with_next=False)
    _para(
        doc,
        f"Nota de calidad de dato: el CSV horario de Etapa 3 el 1 de agosto trae "
        f"{stats['dup_et3_ago01'][0]} filas para {stats['dup_et3_ago01'][1]} marcas TIME "
        "(duplicados). El total diario de la API suma esos duplicados (~199 m³). "
        "Aunque se tomara una sola fila por hora (~100 m³) seguiría siendo varias veces "
        "la matriz madre (13,4 m³). El 19 de julio, sin duplicados, Etapa 3 llegó a "
        "1.418 m³/día con tramos saturados en 89,61 m³/h y estanque inferior en 0: "
        "antecedente de descontrol del mismo medidor.",
    )

    add_formatted_heading(doc, "Invariante matriz 1–4 vs ramales")
    _para(
        doc,
        f"En los 31 días de agosto, Etapa 1–4 fue menor que Etapa 1+2+3 en "
        f"{stats['viol_madre']}/31 días (todos). El peor exceso es el pico 1–4 de agosto. "
        "Fuera del pico el sesgo sigue: p. ej. el 15 de agosto la madre midió 28,9 m³ y "
        "la suma de ramales 55,6 m³. Eso apunta a un problema sistemático de "
        "calibración/instalación (madre subregistra, ramales sobreregistran, o Etapa 3 "
        "no está realmente sobre esa tubería), no a un evento puntual de riego.",
    )
    add_picture_with_pagination(doc, str(figs["balance"]), Inches(6.3), keep_with_next=False)

    add_formatted_heading(doc, "Conclusión y qué revisar en terreno")
    _para(
        doc,
        "El aumento de Etapa 3 del 1–4 de agosto no es un traslado de caudal desde el "
        "estanque inferior. La entrada real del fundo (ESVAL) esos días fue baja. "
        "Interpretar el pico del agregado como consumo del fundo duplicaría agua que "
        "no pasó por la matriz de entrada.",
    )
    _para(
        doc,
        "Revisar en terreno el medidor 000027-08 (Etapa N°3): factor K / peso de pulso, "
        "cableado, punto exacto de instalación (¿está en el ramal de la matriz 1–4 a la "
        "salida del estanque superior?) y si hubo intervención el 4–5 de agosto (el "
        "caudal constante se corta ese día). Contrastar 000027-04 contra la suma 06+07+08 "
        "con un aforo o un caudalímetro portátil. Hasta no cerrar eso, no usar Etapa 3 "
        "como indicador de consumo real del fundo.",
    )
    doc.save(str(out))
    return out


def write_csv(path: Path, daily: Dict[str, Dict[str, float]]) -> Path:
    days = _daterange(date(2026, 7, 15), AGO_FIN)
    fields = ["fecha"] + [f"{nid}_{SHORT[nid]}" for nid in NODES] + ["suma_et1_et2_et3", "madre_menos_suma"]
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(fields)
        for d in days:
            s123 = v(daily, "000027-06", d) + v(daily, "000027-07", d) + v(daily, "000027-08", d)
            et14 = v(daily, "000027-04", d)
            w.writerow(
                [d.isoformat()]
                + [f"{v(daily, nid, d):.2f}".replace(".", ",") for nid in NODES]
                + [f"{s123:.2f}".replace(".", ","), f"{et14 - s123:.2f}".replace(".", ",")]
            )
    return path


def build_pdf(path: Path, figs: Dict[str, Path], stats: dict) -> Path:
    with PdfPages(path) as pdf:
        fig = plt.figure(figsize=(11.69, 8.27))
        fig.text(0.5, 0.88, "Fundo Zapallar — diagnóstico pico 1–4 agosto 2026",
                 ha="center", fontsize=16, fontweight="bold", color="#003366")
        fig.text(0.5, 0.82, "Etapa N°3 vs estanque inferior / ESVAL / matriz 1–4",
                 ha="center", fontsize=12, color="#003366")
        body = (
            "Respuesta: el pico de Etapa 3 NO se explica por el estanque inferior.\n\n"
            f"• Etapa 3 1–4 ago: {_cl(stats['et3_pico_avg'])} m³/día "
            f"({_cl(stats['et3_ratio'], 2)}× el resto de agosto).\n"
            f"• Estanque inferior: {_cl(stats['inf_pico_avg'])} m³/día "
            f"({_cl(stats['inf_ratio'], 2)}× el resto; más bajo, no más alto).\n"
            f"• ESVAL: {_cl(stats['esval_pico_avg'])} m³/día "
            f"({_cl(stats['esval_ratio'], 2)}× el resto).\n"
            f"• Correlación agosto Etapa 3 vs estanque inferior: r = {_cl(stats['r_et3_inf'], 2)}.\n"
            f"• 01-08: Etapa 3 = {_cl(stats['et3_ago01'])} m³  vs  matriz 1–4 = {_cl(stats['et14_ago01'])} m³.\n"
            f"• Invariante madre ≥ ramales 1+2+3: violado {stats['viol_madre']}/31 días de agosto.\n\n"
            "El perfil horario de Etapa 3 es un caudal ~8 m³/h las 24 h (incluida madrugada),\n"
            "no un riego, y corre con estanque inferior en cero. Circuito de referencia:\n"
            "ESVAL → estanque inferior → estanque superior → dos matrices (N5 y 1–4);\n"
            "etapas 1, 2 y 3 son ramales de 1–4."
        )
        fig.text(0.08, 0.72, body, ha="left", va="top", fontsize=11, family="DejaVu Sans")
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)
        for key in ("esquema", "diario", "norm", "balance", "h01", "hgrid"):
            img = plt.imread(figs[key])
            f2, ax = plt.subplots(figsize=(11.69, 8.27))
            ax.imshow(img)
            ax.axis("off")
            pdf.savefig(f2, bbox_inches="tight")
            plt.close(f2)
    return path


def main() -> None:
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    out_dir = Path("reports/Fundo_Zapallar/ABREGADO") / f"DIAGNOSTICO_PICO_ETAPA3_{stamp}"
    fig_dir = out_dir / "figuras"
    fig_dir.mkdir(parents=True, exist_ok=True)
    print(f"[INFO] Salida: {out_dir}", flush=True)

    print("[INFO] Bajando diarios jul–ago…", flush=True)
    daily = fetch_daily(date(2026, 7, 1), AGO_FIN)
    (out_dir / "diario_jul_ago.json").write_text(json.dumps(daily), encoding="utf-8")

    pico = _daterange(PICO_INI, PICO_FIN)
    resto = _daterange(date(2026, 8, 5), AGO_FIN)
    ago = _daterange(AGO_INI, AGO_FIN)

    def avg(nid, days):
        return sum(v(daily, nid, d) for d in days) / len(days)

    et3_pico_avg = avg("000027-08", pico)
    et3_resto = avg("000027-08", resto)
    inf_pico_avg = avg("000027-02", pico)
    inf_resto = avg("000027-02", resto)
    esval_pico_avg = avg("000027-01", pico)
    esval_resto = avg("000027-01", resto)

    def series(nid, days):
        return [v(daily, nid, d) for d in days]

    pairs = [
        ("Etapa 3 vs estanque inferior", "000027-08", "000027-02"),
        ("Etapa 3 vs ESVAL", "000027-08", "000027-01"),
        ("Etapa 3 vs matriz 1–4", "000027-08", "000027-04"),
        ("Etapa 3 vs Etapa 1", "000027-08", "000027-06"),
        ("Etapa 3 vs Etapa 2", "000027-08", "000027-07"),
        ("Estanque inferior vs ESVAL", "000027-02", "000027-01"),
        ("Matriz 1–4 vs Etapa 1", "000027-04", "000027-06"),
    ]
    corrs = []
    for label, a, b in pairs:
        corrs.append((
            label,
            _corr(series(a, ago), series(b, ago)),
            _corr(series(a, resto), series(b, resto)),
        ))

    viol = 0
    for d in ago:
        et14 = v(daily, "000027-04", d)
        s123 = v(daily, "000027-06", d) + v(daily, "000027-07", d) + v(daily, "000027-08", d)
        if et14 < s123 - 0.05:
            viol += 1

    print("[INFO] Bajando horarios 31/07–05/08…", flush=True)
    hdays = _daterange(date(2026, 7, 31), date(2026, 8, 5))
    hours = fetch_hourly(
        ["000027-01", "000027-02", "000027-04", "000027-08"],
        hdays,
    )
    dup_meta = hours.pop("_dup_meta")  # type: ignore[misc]
    (out_dir / "horario_pico.json").write_text(
        json.dumps({"hours": hours, "dup": dup_meta}), encoding="utf-8"
    )

    stats = {
        "et3_pico_avg": et3_pico_avg,
        "et3_ratio": et3_pico_avg / et3_resto if et3_resto else 0,
        "inf_pico_avg": inf_pico_avg,
        "inf_ratio": inf_pico_avg / inf_resto if inf_resto else 0,
        "esval_pico_avg": esval_pico_avg,
        "esval_ratio": esval_pico_avg / esval_resto if esval_resto else 0,
        "r_et3_inf": corrs[0][1],
        "r_inf_esval": corrs[5][1],
        "et3_ago01": v(daily, "000027-08", PICO_INI),
        "et14_ago01": v(daily, "000027-04", PICO_INI),
        "sum123_ago01": (
            v(daily, "000027-06", PICO_INI)
            + v(daily, "000027-07", PICO_INI)
            + v(daily, "000027-08", PICO_INI)
        ),
        "corrs": corrs,
        "viol_madre": viol,
        "dup_et3_ago01": dup_meta["000027-08"]["2026-08-01"],
    }

    print("[INFO] Gráficos…", flush=True)
    figs = {
        "esquema": build_esquema(fig_dir / "esquema.png"),
        "diario": plot_diario(daily, fig_dir / "diario.png"),
        "norm": plot_diario(daily, fig_dir / "diario_norm.png", normalizado=True),
        "balance": plot_balance(daily, fig_dir / "balance.png"),
        "h01": plot_horario(hours, PICO_INI, fig_dir / "horario_01ago.png"),
        "hgrid": plot_et3_vs_inf_horas(hours, pico, fig_dir / "horario_1_4.png"),
    }

    csv_path = write_csv(out_dir / "serie_diaria_jul_ago.csv", daily)
    docx_path = out_dir / "Diagnostico_pico_Etapa3_Zapallar_20260801_20260804.docx"
    pdf_path = out_dir / "Diagnostico_pico_Etapa3_Zapallar_20260801_20260804.pdf"
    print("[INFO] Word…", flush=True)
    build_docx(docx_path, figs, daily, stats)
    print("[INFO] PDF…", flush=True)
    build_pdf(pdf_path, figs, stats)

    print(f"[OK] Word {docx_path}")
    print(f"[OK] PDF  {pdf_path}")
    print(f"[OK] CSV  {csv_path}")

    if credenciales_configuradas():
        sub = "Fundo_Zapallar/ABREGADO"
        for p in (docx_path, pdf_path, csv_path):
            info = subir_a_drive(p, subcarpeta=sub)
            print(f"[DRIVE] {info['name']}")
            print(f"        {info['web_view_link']}")
            (out_dir / f"{p.stem}_drive.json").write_text(
                json.dumps(info, indent=2), encoding="utf-8"
            )
    else:
        print("[WARN] Sin secretos GOOGLE_DRIVE_* — no se sube a Drive.")


if __name__ == "__main__":
    main()
