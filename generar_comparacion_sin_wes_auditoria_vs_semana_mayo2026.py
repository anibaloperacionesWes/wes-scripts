"""
Compara el consumo «Sin WES» de las auditorías Renca (abril 2026) con la semana
25/05/2026–31/05/2026, solo para puntos con rendimiento positivo en la auditoría.

Rendimiento auditoría (positivo = WES redujo consumo vs sin control):
  % = (Semana sin WES − Semana con WES) / Semana sin WES × 100

Semana mayo: suma horaria API (m³/h → m³/día), mismos 7 días calendario.

Uso:
  python generar_comparacion_sin_wes_auditoria_vs_semana_mayo2026.py
"""

from __future__ import annotations

import csv
import sys
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import List, Optional, Sequence, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
BASE_AUDIT = ROOT / "reports" / "reporte de auditoria" / "auditoria_puntos_renca_abril_2026"
XLSX_NOMBRE = "consumo_consolidado_parseo_filas_abr06-12_abr13-19_2026.xlsx"

AUDITORIAS: Tuple[Tuple[str, str], ...] = (
    ("000017-08", "Auditoria ICCO Renca 000017-08"),
    ("000017-04", "Auditoria Escuela Lo Velazquez 000017-04"),
    ("000017-06", "Auditoria Piscina Municipal 000017-06"),
    ("000017-05", "Auditoria Gimnasio 000017-05"),
    ("000017-07", "Auditoria Cumbre de condores 000017-07"),
)

MAYO_DIAS: Tuple[date, ...] = tuple(date(2026, 5, d) for d in range(25, 32))
AUDIT_SIN = "6–12 abril 2026"
AUDIT_CON = "13–19 abril 2026"
MAYO_LABEL = "25–31 mayo 2026"


@dataclass
class FilaComparacion:
    node_id: str
    nombre: str
    m3_con_wes_audit: float
    m3_sin_wes_audit: float
    pct_rendimiento_audit: float
    m3_semana_mayo: float
    diff_mayo_vs_sin: float
    pct_mayo_vs_sin: float
    diff_mayo_vs_con: float
    pct_mayo_vs_con: float


def _pct_rendimiento_audit(m3_con: float, m3_sin: float) -> Optional[float]:
    if m3_sin <= 1e-9:
        return None
    return (float(m3_sin) - float(m3_con)) / float(m3_sin) * 100.0


def _total_m3_semana(node_id: str, dias: Sequence[date]) -> float:
    from auditoria_cpa_icco_renca_grafico import _flatten_grilla_desde_vectores, _vectores_m3h_por_dias

    vecs = _vectores_m3h_por_dias(node_id, dias)
    _, _, total = _flatten_grilla_desde_vectores(dias, vecs)
    return float(total)


def _cargar_filas() -> List[FilaComparacion]:
    from generar_graficos_comparativos_desde_excel_consolidado import totales_rejilla_desde_excel_consolidado
    from generar_reporte_word import format_number_chilean, get_node_name

    out: List[FilaComparacion] = []
    for node_id, sub in AUDITORIAS:
        p_xlsx = (BASE_AUDIT / sub / XLSX_NOMBRE).resolve()
        if not p_xlsx.is_file():
            raise FileNotFoundError(f"No está el Excel de auditoría: {p_xlsx}")

        t_con, t_sin, _n = totales_rejilla_desde_excel_consolidado(p_xlsx)
        pct_aud = _pct_rendimiento_audit(t_con, t_sin)
        if pct_aud is None or pct_aud <= 0:
            continue

        m3_mayo = _total_m3_semana(node_id, MAYO_DIAS)
        nombre = (get_node_name(node_id) or "").strip() or node_id

        diff_sin = m3_mayo - t_sin
        pct_vs_sin = (100.0 * (t_sin - m3_mayo) / t_sin) if t_sin > 0 else 0.0
        diff_con = m3_mayo - t_con
        pct_vs_con = (100.0 * (t_con - m3_mayo) / t_con) if t_con > 0 else 0.0

        out.append(
            FilaComparacion(
                node_id=node_id,
                nombre=nombre,
                m3_con_wes_audit=t_con,
                m3_sin_wes_audit=t_sin,
                pct_rendimiento_audit=pct_aud,
                m3_semana_mayo=m3_mayo,
                diff_mayo_vs_sin=diff_sin,
                pct_mayo_vs_sin=pct_vs_sin,
                diff_mayo_vs_con=diff_con,
                pct_mayo_vs_con=pct_vs_con,
            )
        )
    return out


def _grafico_barras(filas: List[FilaComparacion], out_png: Path) -> None:
    labels = [f"{r.nombre}\n({r.node_id})" for r in filas]
    sin_vals = [r.m3_sin_wes_audit for r in filas]
    mayo_vals = [r.m3_semana_mayo for r in filas]
    x = range(len(filas))
    w = 0.35
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar([i - w / 2 for i in x], sin_vals, width=w, label=f"Sin WES auditoría ({AUDIT_SIN})", color="#C0504D")
    ax.bar([i + w / 2 for i in x], mayo_vals, width=w, label=f"Semana {MAYO_LABEL}", color="#4F81BD")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=28, ha="right", fontsize=8)
    ax.set_ylabel("Consumo acumulado (m³)")
    ax.set_title(
        "Renca — Sin WES (auditoría) vs semana 25–31 mayo 2026\n(solo puntos con rendimiento positivo en auditoría)",
        fontsize=11,
        fontweight="bold",
    )
    ax.legend(fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _word(filas: List[FilaComparacion], out_docx: Path, png: Path) -> None:
    from generar_reporte_word import format_number_chilean

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading("Comparación Sin WES (auditoría) vs semana 25–31 mayo 2026", level=0)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    doc.add_paragraph(
        f"Se incluyen únicamente los establecimientos Renca cuya auditoría de abril 2026 "
        f"mostró rendimiento positivo (% ahorro con WES respecto de la semana sin control). "
        f"Referencia sin WES: {AUDIT_SIN}. Semana con WES en auditoría: {AUDIT_CON}. "
        f"Periodo operativo comparado: {MAYO_LABEL} (serie horaria API WES, 24 h/día)."
    )

    p = doc.add_paragraph()
    p.add_run("Generado: ").bold = True
    p.add_run(datetime.now().strftime("%d-%m-%Y %H:%M"))

    doc.add_heading("Metodología", level=1)
    doc.add_paragraph(
        "• Rendimiento auditoría = (Sin WES − Con WES) / Sin WES × 100. Positivo: menor consumo con control WES.\n"
        "• Consumo Sin WES auditoría: Excel consolidado de cada auditoría (rejilla m³/h).\n"
        "• Semana mayo 2026: suma de consumo horario 25/05–31/05 desde API (misma lógica CPA).\n"
        "• % mayo vs Sin WES = (Sin WES audit. − Mayo) / Sin WES audit. × 100. Positivo: mayo por debajo del sin control histórico."
    )

    doc.add_heading("Resultados", level=1)
    tbl = doc.add_table(rows=1 + len(filas), cols=8)
    tbl.style = "Table Grid"
    headers = [
        "Nodo",
        "Establecimiento",
        "Sin WES audit. (m³)",
        "Con WES audit. (m³)",
        "% rend. audit.",
        f"Mayo 25–31 (m³)",
        "Δ vs Sin WES (m³)",
        "% vs Sin WES",
    ]
    for j, hd in enumerate(headers):
        tbl.rows[0].cells[j].text = hd
        for run in tbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, r in enumerate(filas, start=1):
        cells = tbl.rows[i].cells
        vals = [
            r.node_id,
            r.nombre,
            format_number_chilean(r.m3_sin_wes_audit, 1),
            format_number_chilean(r.m3_con_wes_audit, 1),
            format_number_chilean(r.pct_rendimiento_audit, 1) + " %",
            format_number_chilean(r.m3_semana_mayo, 1),
            format_number_chilean(r.diff_mayo_vs_sin, 1),
            format_number_chilean(r.pct_mayo_vs_sin, 1) + " %",
        ]
        for j, v in enumerate(vals):
            cells[j].text = v

    doc.add_paragraph("")
    doc.add_picture(str(png), width=Cm(16))

    doc.add_heading("Conclusiones", level=1)
    bajo_sin = [r for r in filas if r.pct_mayo_vs_sin > 0]
    sobre_sin = [r for r in filas if r.pct_mayo_vs_sin <= 0]
    doc.add_paragraph(
        f"De {len(filas)} punto(s) con rendimiento positivo en auditoría, "
        f"{len(bajo_sin)} registraron en mayo un consumo inferior al Sin WES de referencia "
        f"y {len(sobre_sin)} igual o superior."
    )
    for r in filas:
        doc.add_paragraph(
            f"• {r.nombre}: mayo {format_number_chilean(r.m3_semana_mayo, 1)} m³ vs "
            f"Sin WES {format_number_chilean(r.m3_sin_wes_audit, 1)} m³ "
            f"({format_number_chilean(r.pct_mayo_vs_sin, 1)} % vs referencia sin control)."
        )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def main() -> int:
    from generar_reporte_word import convertir_word_a_pdf

    filas = _cargar_filas()
    if not filas:
        print("[AVISO] Ningún punto con rendimiento positivo en auditoría abril 2026.")
        return 1

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    out_dir = ROOT / "reports" / "reporte de auditoria" / f"comparacion_sin_wes_vs_mayo25-31_{ts}"
    out_dir.mkdir(parents=True, exist_ok=True)

    csv_path = out_dir / "comparacion_sin_wes_vs_mayo25-31.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(
            [
                "node_id",
                "establecimiento",
                "m3_sin_wes_auditoria",
                "m3_con_wes_auditoria",
                "pct_rendimiento_auditoria",
                "m3_semana_25-31_mayo_2026",
                "diff_mayo_vs_sin_m3",
                "pct_mayo_vs_sin_wes",
                "diff_mayo_vs_con_m3",
                "pct_mayo_vs_con_wes",
            ]
        )
        for r in filas:
            w.writerow(
                [
                    r.node_id,
                    r.nombre,
                    f"{r.m3_sin_wes_audit:.4f}",
                    f"{r.m3_con_wes_audit:.4f}",
                    f"{r.pct_rendimiento_audit:.2f}",
                    f"{r.m3_semana_mayo:.4f}",
                    f"{r.diff_mayo_vs_sin:.4f}",
                    f"{r.pct_mayo_vs_sin:.2f}",
                    f"{r.diff_mayo_vs_con:.4f}",
                    f"{r.pct_mayo_vs_con:.2f}",
                ]
            )

    png_path = out_dir / "barras_sin_wes_vs_mayo.png"
    _grafico_barras(filas, png_path)

    docx_path = out_dir / f"Comparacion_SinWES_auditoria_vs_mayo25-31_{ts}.docx"
    _word(filas, docx_path, png_path)

    pdf_path = docx_path.with_suffix(".pdf")
    try:
        pdf_out = convertir_word_a_pdf(docx_path)
        if pdf_out and Path(pdf_out).is_file():
            pdf_path = Path(pdf_out)
    except Exception:
        pass

    print("=" * 72)
    print("COMPARACIÓN SIN WES AUDITORÍA vs 25–31 MAYO 2026 (rendimiento +)")
    print("=" * 72)
    print(f"Puntos incluidos: {len(filas)}")
    for r in filas:
        print(
            f"  {r.node_id} {r.nombre}: Sin WES {r.m3_sin_wes_audit:.1f} m³ | "
            f"Mayo {r.m3_semana_mayo:.1f} m³ | {r.pct_mayo_vs_sin:.1f} % vs Sin WES"
        )
    print(f"CSV:  {csv_path}")
    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
