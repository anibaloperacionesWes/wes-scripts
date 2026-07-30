"""
Ajustes al Word Bupa Antofagasta:
- Quita MÉTRICAS POR PUNTO (título + tabla + gráfico nocturno por nodo).
- Reemplaza día mayor consumo nocturno (Bomba N°2) por Medidor Principal Sanitaria.
- PARTICIPACIÓN SOBRE CUENTA con proyección mensual (promedio diario × 30).
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

from generar_reporte_word import (
    add_formatted_title,
    add_picture_with_pagination,
    add_table,
    build_hourly_consumption_line_chart,
    find_max_nocturnal_consumption_day,
    format_currency_chilean,
    format_number_chilean,
    parse_date,
)

DOC_PATH = Path(
    "reports/Bupa_Antofagasta/ABREGADO/AGREGADO_20260728_1725/"
    "Reporte_Agregado_BUPA_20260723_20260728.docx"
)
OUT_DIR = DOC_PATH.parent

NUM_DIAS = 6
PRECIO_M3 = 1300.0
START = parse_date("23/07/2026")
END = parse_date("28/07/2026", end_of_day=True)
SANITARIA_ID = "000029-09"
SANITARIA_NAME = "Medidor Principal Sanitaria"

NODOS = [
    {"id": "000029-07", "nombre": "Sala de Bomba Principal", "tipo": "bomba", "m3": 227.9},
    {"id": "000029-08", "nombre": "Sala de Bomba Sexto Piso", "tipo": "bomba", "m3": 36.1},
    {"id": "000029-09", "nombre": "Medidor Principal Sanitaria", "tipo": "cuenta", "m3": 1400.0},
    {"id": "000029-10", "nombre": "Sala de Bomba N°2", "tipo": "bomba", "m3": 6.2},
]


def _para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def _remove_range_by_markers(doc: Document, start_pred, end_pred, *, include_end: bool = False) -> int:
    body = doc.element.body
    children = list(body)
    start_idx = end_idx = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1]
        if tag != "p":
            continue
        t = _para_text(child)
        if start_idx is None and start_pred(t):
            start_idx = i
            continue
        if start_idx is not None and end_pred(t):
            end_idx = i if include_end else i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Rango no encontrado start={start_idx} end={end_idx}")
    # incluir tablas/imágenes entre medio: end_idx es el primer párrafo del siguiente bloque
    n = 0
    for child in children[start_idx:end_idx]:
        body.remove(child)
        n += 1
    return n


def _find_conclusiones(doc: Document):
    for child in list(doc.element.body):
        if child.tag.split("}")[-1] == "p" and _para_text(child).upper().startswith("CONCLUSIONES"):
            return child
    raise RuntimeError("CONCLUSIONES no encontrado")


def _build_sanitaria_nocturnal_chart() -> tuple[Path, datetime, str]:
    result = find_max_nocturnal_consumption_day(SANITARIA_ID, None, START, END)
    if not result:
        raise RuntimeError("Sin día de mayor consumo nocturno para Sanitaria")
    target_dt, hourly = result
    chart_path = OUT_DIR / "chart_max_nocturnal_000029_09.png"
    build_hourly_consumption_line_chart(
        hourly,
        chart_path,
        target_dt,
        f"Día con mayor consumo nocturno ({target_dt.strftime('%d-%m-%y')})",
    )
    texto = (
        f"El gráfico corresponde al día con mayor consumo nocturno (00:00–06:59) "
        f"del Medidor Principal Sanitaria en el periodo. Este punto representa el total "
        f"de la cuenta de agua de la clínica."
    )
    return chart_path, target_dt, texto


def _build_participacion_chart(path: Path, salas_m3: float, sanitaria_m3: float) -> Path:
    resto = max(0.0, sanitaria_m3 - salas_m3)
    labels = ["Salas de bomba\n(07+08+10)", "Resto de la cuenta\n(Sanitaria − salas)"]
    sizes = [salas_m3, resto]
    colors = ["#c0392b", "#0050b3"]
    fig, ax = plt.subplots(figsize=(6.2, 4.5))
    wedges, *_ = ax.pie(
        sizes,
        labels=labels,
        colors=colors,
        autopct=lambda p: f"{p:.1f}%",
        startangle=90,
        textprops={"fontsize": 10},
    )
    for w in wedges:
        w.set_edgecolor("white")
        w.set_linewidth(1.5)
    ax.set_title(
        "Participación proyectada a 30 días sobre Medidor Principal Sanitaria",
        fontsize=12,
        fontweight="bold",
    )
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _make_sections(chart_sanitaria: Path, fecha: datetime, texto_alerta: str, rows: list[dict], chart_part: Path):
    tmp = Document()

    # Día mayor nocturno — Sanitaria
    add_formatted_title(
        tmp,
        f"Día con mayor consumo nocturno - {SANITARIA_NAME.upper()} ({fecha.strftime('%d-%m-%y')}):",
    )
    add_picture_with_pagination(tmp, str(chart_sanitaria), Inches(6), keep_with_next=True)
    p = tmp.add_paragraph(texto_alerta)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Proyección mensual (se mantiene; se reescribe el bloque desde proyección en adelante)
    add_formatted_title(tmp, "Proyección de consumo mensual total por punto")
    intro = tmp.add_paragraph(
        f"Para cada punto se calcula el promedio diario del periodo ({NUM_DIAS} días: "
        f"23/07/2026–28/07/2026) y se proyecta a un mes de 30 días "
        f"(promedio diario × 30). La valorización usa ${format_number_chilean(PRECIO_M3, 0)} CLP/m³."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in intro.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    table_rows = [
        (
            "Punto",
            "Consumo periodo (m³)",
            "Promedio diario (m³)",
            "Proyección 30 días (m³)",
            "Proyección 30 días (CLP)",
        )
    ]
    for r in rows:
        table_rows.append(
            (
                f"{r['nombre']} ({r['id']})",
                format_number_chilean(r["m3"], 1),
                format_number_chilean(r["prom_dia"], 1),
                format_number_chilean(r["proy_30"], 1),
                format_currency_chilean(r["proy_clp"]),
            )
        )
    add_table(tmp, "Proyección mensual total", table_rows, wes_style=True)
    tmp.add_paragraph("")
    proy_chart = OUT_DIR / "chart_proyeccion_mensual_total_por_punto.png"
    if proy_chart.exists():
        add_picture_with_pagination(tmp, str(proy_chart), Inches(6), keep_with_next=True)

    # Participación con proyección mensual
    sanitaria = next(r for r in rows if r["id"] == SANITARIA_ID)
    bombas = [r for r in rows if r["tipo"] == "bomba"]
    salas_m3 = sum(b["proy_30"] for b in bombas)
    salas_clp = sum(b["proy_clp"] for b in bombas)
    sanitaria_m3 = sanitaria["proy_30"]
    sanitaria_clp = sanitaria["proy_clp"]
    pct_m3 = (salas_m3 / sanitaria_m3 * 100.0) if sanitaria_m3 > 0 else 0.0
    pct_clp = (salas_clp / sanitaria_clp * 100.0) if sanitaria_clp > 0 else 0.0

    tmp.add_paragraph("")
    add_formatted_title(tmp, "Participación de salas de bomba vs Medidor Principal Sanitaria")
    expl = tmp.add_paragraph(
        "El Medidor Principal Sanitaria (000029-09) representa el total de la cuenta de agua. "
        "Las cifras de esta sección usan la proyección de consumo mensual (promedio diario × 30). "
        "Las salas de bomba (000029-07, 000029-08 y 000029-10) se expresan como porcentaje del "
        "volumen y del costo proyectado de la cuenta (Sanitaria)."
    )
    expl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in expl.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    part_rows = [
        ("Concepto", "Proy. 30 días (m³)", "Proy. 30 días (CLP)", "% sobre Sanitaria"),
        (
            "Medidor Principal Sanitaria (cuenta total)",
            format_number_chilean(sanitaria_m3, 1),
            format_currency_chilean(sanitaria_clp),
            "100,0%",
        ),
    ]
    for b in bombas:
        pct = (b["proy_30"] / sanitaria_m3 * 100.0) if sanitaria_m3 > 0 else 0.0
        part_rows.append(
            (
                f"{b['nombre']} ({b['id']})",
                format_number_chilean(b["proy_30"], 1),
                format_currency_chilean(b["proy_clp"]),
                f"{format_number_chilean(pct, 1)}%",
            )
        )
    part_rows.append(
        (
            "TOTAL salas de bomba (07+08+10)",
            format_number_chilean(salas_m3, 1),
            format_currency_chilean(salas_clp),
            f"{format_number_chilean(pct_m3, 1)}%",
        )
    )
    add_table(tmp, "Participación sobre cuenta (proyección mensual)", part_rows, wes_style=True)

    resumen = tmp.add_paragraph()
    resumen.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    resumen.add_run("Resumen: ").bold = True
    resumen.add_run(
        f"proyectadas a 30 días, las tres salas de bomba suman "
        f"{format_number_chilean(salas_m3, 1)} m³ ({format_currency_chilean(salas_clp)}), "
        f"equivalentes al {format_number_chilean(pct_m3, 1)}% del volumen y al "
        f"{format_number_chilean(pct_clp, 1)}% del costo proyectado de la cuenta "
        f"(Medidor Principal Sanitaria: {format_number_chilean(sanitaria_m3, 1)} m³ / "
        f"{format_currency_chilean(sanitaria_clp)})."
    )
    for run in resumen.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    tmp.add_paragraph("")
    add_picture_with_pagination(tmp, str(chart_part), Inches(5.2), keep_with_next=True)
    tmp.add_paragraph("")

    return [
        child
        for child in list(tmp.element.body)
        if child.tag.split("}")[-1] != "sectPr"
    ]


def main() -> None:
    rows = []
    for n in NODOS:
        prom = n["m3"] / NUM_DIAS
        proy = prom * 30.0
        rows.append({**n, "prom_dia": prom, "proy_30": proy, "proy_clp": proy * PRECIO_M3})

    chart_sanitaria, fecha, texto = _build_sanitaria_nocturnal_chart()
    salas_proy = sum(r["proy_30"] for r in rows if r["tipo"] == "bomba")
    sanitaria_proy = next(r["proy_30"] for r in rows if r["id"] == SANITARIA_ID)
    chart_part = _build_participacion_chart(
        OUT_DIR / "chart_participacion_salas_vs_sanitaria.png",
        salas_proy,
        sanitaria_proy,
    )

    doc = Document(str(DOC_PATH))

    # Quitar desde MÉTRICAS / RESUMEN POR PUNTO hasta antes de PROYECCIÓN DE CONSUMO MENSUAL
    # (incluye tabla métricas, img nocturna nodos, día bomba N°2 y texto de alertas)
    n1 = _remove_range_by_markers(
        doc,
        lambda t: t.upper().startswith("RESUMEN POR PUNTO")
        or t.upper().startswith("MÉTRICAS POR PUNTO")
        or t.upper().startswith("METRICAS POR PUNTO"),
        lambda t: t.upper().startswith("PROYECCIÓN DE CONSUMO MENSUAL")
        or t.upper().startswith("PROYECCION DE CONSUMO MENSUAL"),
    )
    print(f"[OK] Eliminados {n1} elementos (métricas + día bomba N°2)")

    # Quitar bloque viejo de proyección + participación hasta CONCLUSIONES
    n2 = _remove_range_by_markers(
        doc,
        lambda t: t.upper().startswith("PROYECCIÓN DE CONSUMO MENSUAL")
        or t.upper().startswith("PROYECCION DE CONSUMO MENSUAL"),
        lambda t: t.upper().startswith("CONCLUSIONES"),
    )
    print(f"[OK] Eliminados {n2} elementos (proyección/participación anteriores)")

    elements = _make_sections(chart_sanitaria, fecha, texto, rows, chart_part)
    conclus = _find_conclusiones(doc)
    for el in elements:
        conclus.addprevious(el)

    try:
        doc.save(str(DOC_PATH))
        out = DOC_PATH
    except PermissionError:
        out = DOC_PATH.with_name(DOC_PATH.stem + "_v2.docx")
        doc.save(str(out))
        print(f"[ADVERTENCIA] Guardado como {out.name}")

    print("[OK]", out)
    print(f"     Día mayor nocturno Sanitaria: {fecha.strftime('%d-%m-%Y')}")
    print(
        f"     Participación proy. 30d: {format_number_chilean(salas_proy, 1)} / "
        f"{format_number_chilean(sanitaria_proy, 1)} m³ = "
        f"{format_number_chilean(salas_proy / sanitaria_proy * 100, 1)}%"
    )


if __name__ == "__main__":
    main()
