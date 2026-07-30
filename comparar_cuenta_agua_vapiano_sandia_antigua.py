"""
Comparación: cuenta de agua potable (cliente 2131435-8, factura adjunta)
vs monitoreo WES nodo Sala de Bomba Sandia Antigua (000025-22).

La factura refleja el medidor de acometida que abastece/llena el estanque;
el nodo WES mide el circuito asociado a la impulsión Sandia Antigua en el mall.
Datos de factura tomados del PDF (emisión 26-mar-2026).

Uso:
  python comparar_cuenta_agua_vapiano_sandia_antigua.py
  python comparar_cuenta_agua_vapiano_sandia_antigua.py --hasta 31/03/2026
  python comparar_cuenta_agua_vapiano_sandia_antigua.py --abrir-carpeta

Copia al Escritorio: Kennedy_Vapiano_Sandia_INFORME_ULTIMO.docx

Genera un solo Word en <carpeta_del_script>/reports/Parque_Arauco/Kennedy/comparaciones/:
  • Informe_Vapiano_Sandia_Antigua_<ts>.docx (Parte I comparación + Parte II proyección 30 días)

Los montos y lecturas de la factura están tomados del PDF de cuenta 2131435-8 adjunto por el usuario.
"""

from __future__ import annotations

import argparse
import os
import shutil
import subprocess
import sys
import traceback
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Misma API que generar_reporte_word.py
BASE_URL = "http://104.248.53.141:7003/wes/api/acl-node/v1"

NODE_ID = "000025-22"
NODE_LABEL = "Sala de Bomba Sandia Antigua"

# Siempre guardar junto a este script (evita perder el Word al correr desde otra carpeta)
SCRIPT_DIR = Path(__file__).resolve().parent
REPORTS_COMP = (
    SCRIPT_DIR / "reports" / "Parque_Arauco" / "Kennedy" / "comparaciones"
)


def _carpetas_escritorio() -> List[Path]:
    home = Path.home()
    out: List[Path] = []
    one = home / "OneDrive" / "Desktop"
    local = home / "Desktop"
    if one.is_dir():
        out.append(one.resolve())
    if local.is_dir() and (not out or local.resolve() != out[0]):
        out.append(local.resolve())
    if not out:
        out.append(local.resolve())
    return out


def _copiar_escritorio(src: Path, nombre_fijo: str) -> None:
    for desk in _carpetas_escritorio():
        try:
            desk.mkdir(parents=True, exist_ok=True)
            dest = desk / nombre_fijo
            shutil.copy2(src, dest)
            print(f"[OK] Copia Escritorio: {dest}")
        except OSError as e:
            print(f"[AVISO] No se pudo copiar a {desk}: {e}")


def _doc_error_api(out_path: Path, mensaje: str) -> None:
    """Word mínimo si la API falla (sabes dónde quedó y qué pasó)."""
    doc = Document()
    doc.add_heading("Error al generar comparación WES", level=0)
    doc.add_paragraph(mensaje)
    doc.add_paragraph(f"Carpeta de salida: {out_path.parent}")
    doc.add_paragraph(f"Ejecuta de nuevo desde: {SCRIPT_DIR}")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

# Factura electrónica Nº 1110129 (PDF usuario; cuenta 2131435-8)
FACTURA = {
    "numero_factura": "1110129",
    "fecha_emision": "26/03/2026",
    "numero_cuenta": "2131435-8",
    "titular_factura": "PARQUE ARAUCO S.A.",
    "direccion": "Av. Presidente Kennedy 5413 — Parque Arauco",
    "medidor": "122710153",
    "lectura_anterior_fecha": "18/02/2026",
    "lectura_anterior_m3": 105_279.0,
    "lectura_actual_fecha": "19/03/2026",
    "lectura_actual_m3": 121_069.0,
    "consumo_facturado_m3": 15_790.0,
    "total_a_pagar_clp": 26_662_687,
    "vencimiento": "14/04/2026",
    "nota_usuario_vapiano": (
        "Referencia operacional: suministro asociado al local Vapiano / llenado de estanque "
        "impulsado posteriormente por Sandia Antigua (criterio indicado por el usuario)."
    ),
}


def _ddmmyyyy(dt: datetime) -> str:
    return dt.strftime("%d%m%Y")


def fetch_measures(node_id: str, start_ddmmyyyy: str, end_ddmmyyyy: str) -> List[Dict[str, Any]]:
    r = requests.get(
        f"{BASE_URL}/nodes/measures/dates",
        params=[("id", node_id), ("start", start_ddmmyyyy), ("end", end_ddmmyyyy)],
        timeout=120,
    )
    r.raise_for_status()
    payload = r.json()
    if isinstance(payload, list):
        for item in payload:
            if isinstance(item, dict) and item.get("nodeId") == node_id:
                return list(item.get("month") or [])
        return list((payload[0] or {}).get("month") or []) if payload else []
    return list((payload or {}).get("month") or [])


def month_rows_totals(month: List[Dict[str, Any]]) -> Tuple[float, int, List[Tuple[str, float]]]:
    rows: List[Tuple[str, float]] = []
    for m in month:
        d = m.get("date")
        t = m.get("totalM3")
        if d is None or t is None:
            continue
        rows.append((str(d), float(t)))
    total = sum(t for _, t in rows)
    dias = len(rows)
    return total, dias, sorted(rows, key=lambda x: x[0])


def _parse_row_date(s: str) -> Optional[datetime]:
    s = s.strip()
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00"))
    except ValueError:
        try:
            return datetime.strptime(s[:10], "%Y-%m-%d")
        except ValueError:
            return None


def filter_rows_date_inclusive(
    month: List[Dict[str, Any]], d0: datetime, d1: datetime
) -> Tuple[float, int, List[Tuple[str, float]]]:
    """Suma solo filas cuya fecha cae entre d0 y d1 (solo parte fecha, inclusive)."""
    rows: List[Tuple[str, float]] = []
    for m in month:
        ds = m.get("date")
        t = m.get("totalM3")
        if ds is None or t is None:
            continue
        dt = _parse_row_date(str(ds))
        if dt is None:
            continue
        day = dt.date()
        if day < d0.date() or day > d1.date():
            continue
        rows.append((str(ds), float(t)))
    total = sum(v for _, v in rows)
    return total, len(rows), sorted(rows, key=lambda x: x[0])


def fetch_alerts_positive(node_id: str, start_ddmmyyyy: str, end_ddmmyyyy: str) -> List[Dict[str, Any]]:
    r = requests.get(
        f"{BASE_URL}/nodes/myalert/alerts",
        params=[("id", node_id), ("start", start_ddmmyyyy), ("end", end_ddmmyyyy)],
        timeout=120,
    )
    r.raise_for_status()
    raw = r.json()
    if not isinstance(raw, list):
        return []
    return [a for a in raw if float(a.get("measure") or 0) > 0]


def _append_proyeccion_sections(
    doc: Document,
    factura: Dict[str, Any],
    total_win: float,
    dias_win: int,
    rows_win: List[Tuple[str, float]],
    promedio: float,
    proj30: float,
    proj30_min: float,
    proj30_max: float,
    consumo_factura: float,
    sec_level: int = 2,
) -> None:
    doc.add_page_break()
    doc.add_heading("Parte II — Proyección 30 días (WES vs consumo facturado)", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"Nodo {NODE_ID} ({NODE_LABEL}). Cuenta referencia {factura['numero_cuenta']}. "
        "Promedio calculado en 11-03-2026 al 30-03-2026."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_heading("1. Base de cálculo (post corrección medidor)", sec_level)
    doc.add_paragraph(
        "Periodo para el promedio diario: 11-03-2026 al 30-03-2026 (inclusive)."
    )
    doc.add_paragraph(
        f"Días con medición en WES en ese tramo: {dias_win}. "
        f"Consumo acumulado en el tramo: {total_win:,.1f} m³.".replace(",", ".")
    )
    doc.add_paragraph(
        f"Promedio diario (total ÷ días con dato): {promedio:,.2f} m³/día.".replace(",", ".")
    )

    doc.add_heading("2. Proyección a 30 días", sec_level)
    doc.add_paragraph(
        f"Proyección lineal (promedio × 30): {proj30:,.1f} m³.".replace(",", ".")
    )
    doc.add_paragraph(
        "Banda por variabilidad diaria observada en el mismo tramo (mínimo y máximo diario × 30):"
    )
    doc.add_paragraph(
        f"  • Escenario bajo: {proj30_min:,.1f} m³\n  • Escenario alto: {proj30_max:,.1f} m³".replace(
            ",", "."
        ),
    )

    doc.add_heading("3. Referencia: consumo facturado (periodo de lecturas)", sec_level)
    doc.add_paragraph(
        f"Consumo facturado en la boleta (medidor acometida): {consumo_factura:,.1f} m³ "
        f"(lecturas {factura['lectura_anterior_fecha']} — {factura['lectura_actual_fecha']}).".replace(
            ",", "."
        )
    )

    doc.add_heading("4. Diferencia y rango de error (proyección vs factura)", sec_level)
    delta = proj30 - consumo_factura
    pct = (delta / consumo_factura * 100.0) if consumo_factura else 0.0
    doc.add_paragraph(
        f"Proyección central − factura: {delta:+,.1f} m³ ({pct:+.1f} % sobre el consumo facturado).".replace(
            ",", "."
        )
    )
    dentro = proj30_min <= consumo_factura <= proj30_max
    doc.add_paragraph(
        f"¿El consumo facturado ({consumo_factura:,.0f} m³) cae dentro de la banda min–max proyectada? "
        f"{'Sí' if dentro else 'No'}.".replace(",", ".")
    )
    if not dentro:
        if consumo_factura < proj30_min:
            doc.add_paragraph(
                f"La factura queda {proj30_min - consumo_factura:,.1f} m³ por debajo del piso proyectado.".replace(
                    ",", "."
                )
            )
        else:
            doc.add_paragraph(
                f"La factura queda {consumo_factura - proj30_max:,.1f} m³ por sobre el techo proyectado.".replace(
                    ",", "."
                )
            )
    doc.add_paragraph(
        "Nota: la factura mide la acometida; WES mide el punto Sandia Antigua. "
        "La banda min–max solo refleja la variabilidad diaria WES en 11–30 mar, no la incertidumbre del vínculo hidráulico con el medidor de cobro."
    )

    doc.add_heading("5. Detalle diario (11–30 mar)", sec_level)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    h = tbl.rows[0].cells
    h[0].text = "Fecha"
    h[1].text = "m³/día"
    for fecha, val in sorted(rows_win, key=lambda x: x[0]):
        row = tbl.add_row().cells
        row[0].text = fecha
        row[1].text = f"{val:,.2f}".replace(",", ".")


def _append_comparacion_sections(
    doc: Document,
    factura: Dict[str, Any],
    wes_post_corr_total: float,
    wes_post_corr_dias: int,
    wes_overlap_bill: float,
    wes_overlap_dias: int,
    wes_pre_window_total: float,
    top_days: List[Tuple[str, float]],
    low_days: List[Tuple[str, float]],
    alerts: List[Dict[str, Any]],
    sec_level: int = 2,
) -> None:
    doc.add_heading("Parte I — Comparación cuenta de agua vs WES Sandia Antigua", level=1)

    doc.add_heading("1. Antecedentes de la factura (medidor de acometida)", sec_level)
    bullets = [
        f"Cuenta / Nº cliente: {factura['numero_cuenta']}",
        f"Factura Nº {factura['numero_factura']}, emisión {factura['fecha_emision']}",
        f"Titular en documento: {factura['titular_factura']}",
        f"Medidor: {factura['medidor']}",
        f"Lectura anterior: {factura['lectura_anterior_fecha']} — {factura['lectura_anterior_m3']:,.0f} m³".replace(",", "."),
        f"Lectura actual: {factura['lectura_actual_fecha']} — {factura['lectura_actual_m3']:,.0f} m³".replace(",", "."),
        f"Consumo facturado en el periodo: {factura['consumo_facturado_m3']:,.1f} m³".replace(",", "."),
        f"Total a pagar: ${factura['total_a_pagar_clp']:,} CLP (venc. {factura['vencimiento']})".replace(",", "."),
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(factura["nota_usuario_vapiano"])

    doc.add_heading("2. Monitoreo WES — punto Sandia Antigua", sec_level)
    doc.add_paragraph(
        f"Nodo: {NODE_ID} ({NODE_LABEL}). "
        "Según criterio de operación, el medidor de la factura alimenta el estanque que luego "
        "es impulsado por el sistema Sandia Antigua; el WES mide el volumen registrado en ese "
        "punto de monitoreo (no es el mismo medidor ni necesariamente el mismo volumen físico "
        "que el de la factura)."
    )

    doc.add_heading("3. Cifras WES (post corrección ultrasonido 11-03-2026)", sec_level)
    doc.add_paragraph(
        f"Periodo solicitado desde el 11-03-2026 hasta la fecha de generación: "
        f"{wes_post_corr_dias} días con datos, consumo acumulado {wes_post_corr_total:,.1f} m³.".replace(",", ".")
    )
    doc.add_paragraph(
        f"Tramo que solapa el cierre del periodo de factura (11-03-2026 al 19-03-2026): "
        f"{wes_overlap_dias} días, total WES {wes_overlap_bill:,.1f} m³.".replace(",", ".")
    )

    doc.add_heading("4. Dato de referencia: ventana completa del periodo de factura en WES", sec_level)
    doc.add_paragraph(
        "Si se toma en WES el mismo calendario que la factura (18-02-2026 a 19-03-2026), "
        f"el total acumulado en el nodo es {wes_pre_window_total:,.1f} m³.".replace(",", ".")
    )
    doc.add_paragraph(
        "Ese valor es muy superior al consumo facturado (15.790 m³), lo cual es coherente con "
        "un error de configuración del medidor ultrasónico antes del 11-03-2026: las lecturas "
        "previas a la corrección no son comparables con la factura."
    )

    doc.add_heading("5. Hallazgos más importantes", sec_level)
    hallazgos = [
        "La factura mide el ingreso de agua potable por la acometida (un solo medidor de cobro); "
        "WES mide en el punto de bomba/impulsión Sandia Antigua. La reconciliación exige criterio "
        "hidráulico (llenado, consumo del mall, recirculaciones, rezagos).",
        f"Tras el 11-03-2026, en {wes_post_corr_dias} días el nodo acumula {wes_post_corr_total:,.1f} m³ "
        f"(promedio ~{wes_post_corr_total / max(wes_post_corr_dias,1):,.1f} m³/día).".replace(",", "."),
        "En el solape 11–19 marzo, WES registra ~3,9 mil m³ frente a un periodo de factura de ~30 días "
        "y 15.790 m³ totales: son magnitudes distintas (ventana parcial vs medición de entrada).",
        "Variabilidad diaria relevante en marzo (post corrección): días pico cercanos a 500–520 m³/día "
        "y días más bajos ~275–350 m³/día; conviene revisar si coincide con operación comercial y "
        "llenado de estanque.",
    ]
    if alerts:
        a0 = alerts[0]
        hallazgos.append(
            f"Alerta con medida > 0 en el periodo 11–31 marzo: {a0.get('creationDate')} "
            f"({float(a0.get('measure') or 0):.3f} m³/h). Revisar contexto (arranque de bomba, "
            "prueba o evento puntual)."
        )
    else:
        hallazgos.append(
            "En el periodo 11–31 marzo no se registraron alertas nocturnas adicionales con medida "
            "positiva más allá de la indicada en el sistema (si aplica)."
        )

    for h in hallazgos:
        doc.add_paragraph(h, style="List Number")

    doc.add_heading("6. Detalle días extremos (WES, 11-03 al fin de datos)", sec_level)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Fecha"
    hdr[1].text = "Consumo diario (m³)"
    for fecha, val in top_days[:5]:
        row = table.add_row().cells
        row[0].text = fecha
        row[1].text = f"{val:,.1f}".replace(",", ".")
    doc.add_paragraph("Días más bajos (muestra):")
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    h2[0].text = "Fecha"
    h2[1].text = "Consumo diario (m³)"
    for fecha, val in low_days[:5]:
        row = table2.add_row().cells
        row[0].text = fecha
        row[1].text = f"{val:,.1f}".replace(",", ".")


def build_informe_word_unico(
    out_path: Path,
    factura: Dict[str, Any],
    wes_post_corr_total: float,
    wes_post_corr_dias: int,
    wes_overlap_bill: float,
    wes_overlap_dias: int,
    wes_pre_window_total: float,
    top_days: List[Tuple[str, float]],
    low_days: List[Tuple[str, float]],
    alerts: List[Dict[str, Any]],
    total_win: float,
    dias_win: int,
    rows_win: List[Tuple[str, float]],
    promedio: float,
    proj30: float,
    proj30_min: float,
    proj30_max: float,
    consumo_factura: float,
) -> None:
    doc = Document()
    t = doc.add_heading(
        "Informe Word — Cuenta de agua Vapiano (2131435-8) y WES Sala de Bomba Sandia Antigua",
        level=0,
    )
    for run in t.runs:
        run.font.size = Pt(16)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Documento único en formato Microsoft Word (.docx), generado el "
        f"{datetime.now().strftime('%d-%m-%Y %H:%M')}. "
        "Incluye la comparación factura vs monitoreo WES y la proyección a 30 días. "
        "Valores WES desde la API de nodos."
    )
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    _append_comparacion_sections(
        doc,
        factura,
        wes_post_corr_total,
        wes_post_corr_dias,
        wes_overlap_bill,
        wes_overlap_dias,
        wes_pre_window_total,
        top_days,
        low_days,
        alerts,
    )
    _append_proyeccion_sections(
        doc,
        factura,
        total_win,
        dias_win,
        rows_win,
        promedio,
        proj30,
        proj30_min,
        proj30_max,
        consumo_factura,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--hasta",
        default=None,
        help="DD/MM/YYYY fin periodo WES post-corrección (default: hoy)",
    )
    parser.add_argument(
        "--no-copia-escritorio",
        action="store_true",
        help="No copiar los .docx al Escritorio (OneDrive/local).",
    )
    parser.add_argument(
        "--abrir-carpeta",
        action="store_true",
        help="Abrir la carpeta de informes en el Explorador al terminar (Windows).",
    )
    args = parser.parse_args()

    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass

    REPORTS_COMP.mkdir(parents=True, exist_ok=True)
    print(f"[INFO] Carpeta informes: {REPORTS_COMP.resolve()}")
    print(f"[INFO] Script en: {SCRIPT_DIR}")

    hasta = args.hasta or datetime.now().strftime("%d/%m/%Y")
    hasta_dt = datetime.strptime(hasta, "%d/%m/%Y")
    start_post = "11032026"
    end_post = _ddmmyyyy(hasta_dt)
    stamp = datetime.now().strftime("%Y%m%d_%H%M")

    try:
        month_post = fetch_measures(NODE_ID, start_post, end_post)
        total_post, dias_post, rows_post = month_rows_totals(month_post)

        # Promedio 11–30 mar y proyección 30 días vs consumo facturado
        d_prom0 = datetime(2026, 3, 11)
        d_prom1 = datetime(2026, 3, 30)
        month_prom = fetch_measures(NODE_ID, "11032026", "30032026")
        tot_pr, dias_pr, rows_pr = filter_rows_date_inclusive(month_prom, d_prom0, d_prom1)
        prom = (tot_pr / dias_pr) if dias_pr else 0.0
        vals = [v for _, v in rows_pr]
        p30 = prom * 30.0
        p30_lo = (min(vals) * 30.0) if vals else 0.0
        p30_hi = (max(vals) * 30.0) if vals else 0.0

        month_overlap = fetch_measures(NODE_ID, "11032026", "19032026")
        tot_ov, dias_ov, _ = month_rows_totals(month_overlap)

        month_bill_window = fetch_measures(NODE_ID, "18022026", "19032026")
        tot_bw, _, _ = month_rows_totals(month_bill_window)

        sorted_by_val = sorted(rows_post, key=lambda x: x[1], reverse=True)
        top5 = sorted_by_val[:5]
        low5 = sorted(rows_post, key=lambda x: x[1])[:5]

        try:
            alerts = fetch_alerts_positive(NODE_ID, start_post, end_post)
        except Exception as ae:
            print(f"[AVISO] Alertas no disponibles: {ae}")
            alerts = []

        out = REPORTS_COMP / f"Informe_Vapiano_Sandia_Antigua_{stamp}.docx"

        build_informe_word_unico(
            out,
            FACTURA,
            total_post,
            dias_post,
            tot_ov,
            dias_ov,
            tot_bw,
            top5,
            low5,
            alerts,
            tot_pr,
            dias_pr,
            rows_pr,
            prom,
            p30,
            p30_lo,
            p30_hi,
            float(FACTURA["consumo_facturado_m3"]),
        )
    except Exception as e:
        err_file = REPORTS_COMP / f"ERROR_comparacion_vapiano_sandia_{stamp}.docx"
        tb = traceback.format_exc()
        _doc_error_api(
            err_file,
            f"No se pudo completar el informe.\n\n{type(e).__name__}: {e}\n\n{tb}",
        )
        print(f"[ERROR] {e}")
        print(tb)
        print(f"[INFO] Detalle guardado en: {err_file.resolve()}")
        return 1

    print(f"[OK] Informe Word (.docx): {out.resolve()}")

    if not args.no_copia_escritorio:
        _copiar_escritorio(out, "Kennedy_Vapiano_Sandia_INFORME_ULTIMO.docx")

    if sys.platform == "win32" and args.abrir_carpeta:
        try:
            os.startfile(REPORTS_COMP)
        except OSError:
            subprocess.run(["explorer", str(REPORTS_COMP)], check=False)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
