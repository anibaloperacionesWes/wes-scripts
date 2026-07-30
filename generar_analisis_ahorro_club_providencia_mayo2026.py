"""
Análisis de ahorro potencial — Club Providencia (mayo 2026).

Cuando hay consumo nocturno en datos horarios pero pocas o ninguna alerta MyAlert,
la proyección del agregado (basada en alertas) queda en cero. Este informe complementa
con:
  - Perfil horario de un día representativo por punto (mayor consumo 00:00–06:59).
  - Proyección: promedio horario nocturno × 24 h × 30 días (m³/mes evitables con control).

Uso:
  python generar_analisis_ahorro_club_providencia_mayo2026.py
"""

from __future__ import annotations

import sys
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Club_Providencia" / "analisis_ahorro_mayo2026"

COMPANY_ID = "000031"
NODES = ["000031-01", "000031-02"]
START = "01/05/2026"
END = "31/05/2026"
DIAS_MES_REF = 30


def _parse_periodo() -> Tuple[datetime, datetime, int]:
    from generar_reporte_word import parse_date

    start_dt = parse_date(START)
    end_dt = parse_date(END, end_of_day=True)
    num_dias = (end_dt.date() - start_dt.date()).days + 1
    return start_dt, end_dt, num_dias


def _grafico_barras_ahorro(
    nodos: List[dict],
    out_path: Path,
) -> None:
    labels = [n["nombre_corto"] for n in nodos]
    vals = [n["proyeccion_mensual_m3"] for n in nodos]
    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(labels, vals, color=["#2980b9", "#16a085"], edgecolor="white")
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + max(vals) * 0.02 if vals else 0,
            f"{v:.1f}",
            ha="center",
            va="bottom",
            fontsize=10,
        )
    ax.set_ylabel("m³/mes (proyección)")
    ax.set_title("Ahorro potencial mensual por punto (control de consumo nocturno)")
    ax.grid(True, axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def generar() -> Path:
    from generar_reporte_word import (
        calculate_nocturnal_metrics,
        format_currency_chilean,
        format_number_chilean,
        get_node_name,
        get_water_price_per_m3,
        UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION,
    )
    from wes_estilo_graficos_app import (
        dia_mayor_consumo_nocturno,
        guardar_grafico_horario_24h_app,
        proyeccion_mensual_desde_nocturno,
    )

    start_dt, end_dt, num_dias = _parse_periodo()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    out_docx = OUT_DIR / f"Analisis_ahorro_Club_Providencia_mayo2026_{ts}.docx"

    precio = get_water_price_per_m3(COMPANY_ID, NODES[0], None)
    resultados: List[dict] = []

    for node_id in NODES:
        node_name = get_node_name(node_id) or node_id
        m = calculate_nocturnal_metrics(node_id, start_dt, end_dt)
        dia_rep, horas_rep, noche_rep = dia_mayor_consumo_nocturno(node_id, start_dt, end_dt)

        c_noche = float(m["consumo_nocturno_total"])
        d_con = int(m["dias_con_consumo_nocturno"])
        d_sin = int(m["dias_sin_consumo_nocturno"])
        d_datos = int(m["dias_con_datos_horarios"])
        pct = (100.0 * d_con / d_datos) if d_datos else 0.0

        proy_mes_30, proy_dia, prom_h, cumple_umbral = proyeccion_mensual_desde_nocturno(
            c_noche, num_dias, d_con, d_sin, dias_mes=DIAS_MES_REF, forzar=True
        )

        png_dia = OUT_DIR / f"perfil_{node_id}_{ts}.png"
        guardar_grafico_horario_24h_app(
            horas_rep,
            png_dia,
            titulo=f"{node_name} — {dia_rep:%d/%m/%Y} (día con mayor consumo 00–06 h)",
        )

        resultados.append(
            {
                "node_id": node_id,
                "nombre": node_name,
                "nombre_corto": node_name.replace("Matriz ", ""),
                "dia_rep": dia_rep,
                "noche_dia_rep": noche_rep,
                "consumo_nocturno_mes": c_noche,
                "dias_con_noche": d_con,
                "pct_dias": pct,
                "prom_h": prom_h,
                "proy_dia": proy_dia,
                "proyeccion_mensual_m3": proy_mes_30,
                "valor_clp": proy_mes_30 * precio,
                "cumple_umbral": cumple_umbral,
                "png_dia": png_dia,
            }
        )

    png_barras = OUT_DIR / f"ahorro_mensual_{ts}.png"
    _grafico_barras_ahorro(resultados, png_barras)

    total_mes = sum(r["proyeccion_mensual_m3"] for r in resultados)
    total_clp = total_mes * precio

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading("Análisis de ahorro potencial — Club Providencia", 0)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    doc.add_paragraph()
    for lab, txt in [
        ("Cliente", "Club Providencia"),
        ("Periodo", f"{START} al {END}"),
        ("Puntos", ", ".join(NODES)),
        ("Generado", datetime.now().strftime("%d-%m-%Y %H:%M")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{lab}: ").bold = True
        p.add_run(txt)

    doc.add_heading("1. Contexto", level=1)
    doc.add_paragraph(
        "El reporte agregado de mayo muestra consumo nocturno (00:00–06:59) en los datos horarios, "
        "pero la proyección basada en alertas MyAlert puede ser cero si no hay alertas recientes "
        "en horario 22:00–07:00. Este documento usa únicamente la serie horaria para estimar el "
        "volumen que se podría evitar instalando un equipo de control que corte o reduzca el flujo "
        "cuando el club no opera."
    )

    doc.add_heading("2. Metodología", level=1)
    doc.add_paragraph(
        "Por cada punto:\n"
        "1) Se suma el consumo entre 00:00 y 06:59 de todos los días de mayo.\n"
        "2) Se obtiene un promedio horario nocturno = consumo nocturno total ÷ (días del periodo × 7 h).\n"
        "3) Se proyecta un escenario de fuga/control deficiente 24 h/día: promedio × 24 h/día × 30 días.\n"
        "4) Se grafica el día con mayor consumo nocturno del mes (perfil 0–23 h).\n\n"
        f"Umbral WES habitual para proyección en agregado: ≥{UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION:.0f} % "
        "de los días con consumo en madrugada. Si no se cumple, igual se muestra la proyección técnica "
        "para evaluar el beneficio de un sistema de control."
    )

    doc.add_heading("3. Resultados por punto", level=1)
    tbl = doc.add_table(rows=1 + len(resultados), cols=7)
    tbl.style = "Table Grid"
    headers = [
        "Punto",
        "Cons. nocturno mayo (m³)",
        "Días con cons. 00–06h",
        "% días",
        "Prom. hora noct. (m³/h)",
        "Proy. mensual 30 d (m³)",
        "Valor ref. (CLP)",
    ]
    for j, hd in enumerate(headers):
        tbl.rows[0].cells[j].text = hd
        for r in tbl.rows[0].cells[j].paragraphs:
            for run in r.runs:
                run.bold = True
    for i, r in enumerate(resultados, 1):
        row = tbl.rows[i].cells
        vals = [
            r["nombre"],
            format_number_chilean(r["consumo_nocturno_mes"], 1),
            str(r["dias_con_noche"]),
            format_number_chilean(r["pct_dias"], 1) + " %",
            format_number_chilean(r["prom_h"], 3),
            format_number_chilean(r["proyeccion_mensual_m3"], 1),
            format_currency_chilean(r["valor_clp"]),
        ]
        for j, v in enumerate(vals):
            row[j].text = v

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Total club (proyección 30 días): ").bold = True
    p.add_run(
        f"{format_number_chilean(total_mes, 1)} m³ — {format_currency_chilean(total_clp)} "
        f"(precio ref. {format_currency_chilean(precio)}/m³)."
    )

    for r in resultados:
        doc.add_heading(f"3.{resultados.index(r)+1} {r['nombre']}", level=2)
        doc.add_paragraph(
            f"Día representativo (mayor consumo 00:00–06:59): {r['dia_rep'].strftime('%d/%m/%Y')} "
            f"({format_number_chilean(r['noche_dia_rep'], 2)} m³ en madrugada ese día)."
        )
        if not r["cumple_umbral"]:
            doc.add_paragraph(
                f"Nota: el punto no alcanza el umbral del {UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION:.0f} % "
                "de días con consumo nocturno; la proyección mensual mostrada es estimación técnica "
                "a partir del consumo horario observado."
            )
        doc.add_picture(str(r["png_dia"]), width=Inches(5.8))

    doc.add_heading("4. Resumen gráfico mensual", level=1)
    doc.add_picture(str(png_barras), width=Inches(5.5))

    doc.add_heading("5. Interpretación", level=1)
    doc.add_paragraph(
        "Si el consumo en horario de cierre (madrugada) corresponde a fugas, válvulas abiertas o "
        "equipos sin control, un sistema de monitoreo y corte podría acercar ese caudal a cero en "
        "esas horas. El ahorro real dependerá de la intervención en terreno; la proyección mensual "
        "cuantifica el orden de magnitud del agua que hoy se registra fuera del horario operativo "
        "o que se extrapolaría si ese patrón se mantuviera las 24 horas."
    )

    doc.save(out_docx)
    pdf = out_docx.with_suffix(".pdf")
    try:
        from generar_reporte_word import convertir_word_a_pdf

        p = convertir_word_a_pdf(out_docx)
        if p and p.is_file():
            pdf = p
    except Exception:
        pass
    print(f"[OK] Word: {out_docx}")
    print(f"[OK] PDF:  {pdf}")
    print(f"[OK] Total proyección mensual: {total_mes:.1f} m³")
    return out_docx


def main() -> int:
    print("=" * 72)
    print("ANÁLISIS AHORRO — CLUB PROVIDENCIA MAYO 2026")
    print("=" * 72)
    generar()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
