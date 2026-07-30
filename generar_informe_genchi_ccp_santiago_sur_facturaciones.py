"""
Informe comparativo facturaciones Aguas Andinas — CCP Santiago Sur (Genchi / Gendarmería).

Estructura:
  reports/Genchi/CCP Santiago Sur/Facturaciones/
  reports/Genchi/CCP Santiago Sur/informe/

Ventana de análisis: ene-2025 a la fecha.
Excluye factura 8427552 (feb-2025, equipos off por licitación) y mes 2025-02 del desglose.
Retiro WES: 30-06-2025.

Uso:
  python generar_informe_genchi_ccp_santiago_sur_facturaciones.py
"""

from __future__ import annotations

import csv
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
SITE_DIR = ROOT / "reports" / "Genchi" / "CCP Santiago Sur"
FACTURAS_DIR = next(
    (SITE_DIR / sub for sub in ("Facturaciones", "facturaciones") if (SITE_DIR / sub).is_dir()),
    SITE_DIR / "Facturaciones",
)
OUT_BASE = SITE_DIR / "informe"

FECHA_DESDE_ANALISIS = date(2025, 1, 1)
FECHA_RETIRO_WES = date(2025, 6, 30)
FECHA_FIN_CON_WES = date(2025, 6, 29)
FACTURA_EXCLUIDA = "8427552"
MESES_EXCLUIDOS = {"2025-02"}
MESES_PROYECCION_ANUAL = 12

CUENTA = "1007968-3"
MEDIDOR = "723461"
CLIENTE = "Gendarmería de Chile — CCP Santiago Sur (Unidad Genchi)"

_MESES_NOM = {
    1: "ene", 2: "feb", 3: "mar", 4: "abr", 5: "may", 6: "jun",
    7: "jul", 8: "ago", 9: "sep", 10: "oct", 11: "nov", 12: "dic",
}

_HEADING = RGBColor(31, 71, 136)


def _keep_with_next(paragraph) -> None:
    try:
        p_pr = paragraph._p.get_or_add_pPr()
        if p_pr.find(qn("w:keepNext")) is None:
            p_pr.append(OxmlElement("w:keepNext"))
    except Exception:
        pass


def _configurar_pagina_compacta(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0


def _heading_compacto(doc: Document, texto: str, level: int = 1):
    h = doc.add_heading(texto, level)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    _keep_with_next(h)
    return h


def _tabla_compacta(table, font_pt: float = 9.0) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(font_pt)
                    run.font.name = "Calibri"


def _encabezado_tabla_repite(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def _add_picture_compact(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    cap = doc.add_paragraph(caption)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(2)
    _keep_with_next(cap)
    pic = doc.add_paragraph()
    pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pic.add_run().add_picture(str(path), width=Cm(width_cm))
    pic.paragraph_format.space_after = Pt(4)


def _fmt_clp(n: int) -> str:
    s = f"{abs(int(n)):,}".replace(",", ".")
    return f"{s} CLP" if n >= 0 else f"-{s} CLP"


def _fmt_m3(n: float) -> str:
    return f"{n:,.0f} m³".replace(",", ".")


def _fmt_mes_corto(d: date) -> str:
    return f"{_MESES_NOM[d.month]}-{d.year}"


def _fmt_mes(key: str) -> str:
    y, m = key.split("-")
    return f"{_MESES_NOM[int(m)]}-{y}"


@dataclass
class Periodo:
    boleta: str
    lectura_desde: date
    lectura_hasta: date
    m3: int
    clp: int
    pdf: str
    clp_esperado: int = 0
    sobrecosto_clp: int = 0

    @property
    def dias(self) -> int:
        return (self.lectura_hasta - self.lectura_desde).days

    @property
    def m3_dia(self) -> float:
        return self.m3 / self.dias if self.dias else 0.0

    @property
    def etiqueta(self) -> str:
        if self.lectura_hasta <= FECHA_FIN_CON_WES:
            return "Con monitoreo WES"
        return "Sin monitoreo WES"


def _extraer_clp_pdf(pdf: Path) -> int:
    from facturacion_aguas_andinas_pdf import extraer_texto_pdf

    txt = extraer_texto_pdf(pdf)
    m = re.search(r"TOTAL\s+A\s+PAGAR\s*\$\s*([\d\.]+)", txt, flags=re.IGNORECASE)
    if not m:
        raise ValueError(f"No se encontró TOTAL A PAGAR en {pdf.name}")
    return int(m.group(1).replace(".", ""))


def _cargar_periodos() -> List[Periodo]:
    from facturacion_aguas_andinas_pdf import listar_periodos_desde_pdf

    raw: List[Periodo] = []
    for pdf in sorted(FACTURAS_DIR.glob("*.pdf")):
        clp = _extraer_clp_pdf(pdf)
        for p in listar_periodos_desde_pdf(pdf):
            raw.append(
                Periodo(
                    boleta=p.boleta,
                    lectura_desde=p.lectura_anterior.date(),
                    lectura_hasta=p.lectura_actual.date(),
                    m3=p.m3_cuenta,
                    clp=clp,
                    pdf=pdf.name,
                )
            )
    raw.sort(key=lambda x: x.lectura_hasta)
    return raw


def _filtrar_analisis(todos: List[Periodo]) -> List[Periodo]:
    """Ene-2025 a la fecha; sin factura anómala feb-2025."""
    return [
        p
        for p in todos
        if p.boleta != FACTURA_EXCLUIDA and p.lectura_hasta >= FECHA_DESDE_ANALISIS
    ]


def _aplicar_sobrecosto(periodos: List[Periodo], con_wes: List[Periodo]) -> float:
    m3_base = sum(p.m3 for p in con_wes)
    clp_base = sum(p.clp for p in con_wes)
    clp_m3 = clp_base / m3_base if m3_base else 0.0
    for p in periodos:
        p.clp_esperado = int(round(p.m3 * clp_m3))
        p.sobrecosto_clp = p.clp - p.clp_esperado
    return clp_m3


def _estado_mes(key: str) -> str:
    y, m = map(int, key.split("-"))
    d = date(y, m, 1)
    if d >= date(2025, 7, 1):
        return "Sin WES"
    return "Con WES"


def _prorrateo_mensual(periodos: List[Periodo]) -> List[dict]:
    agg: Dict[str, dict] = defaultdict(lambda: {"m3": 0.0, "clp": 0, "clp_esp": 0, "dias": 0})

    for p in periodos:
        if p.dias <= 0:
            continue
        d = max(p.lectura_desde, FECHA_DESDE_ANALISIS)
        fin = p.lectura_hasta
        while d < fin:
            if d.month == 12:
                fin_mes = date(d.year + 1, 1, 1)
            else:
                fin_mes = date(d.year, d.month + 1, 1)
            tramo_fin = min(fin, fin_mes)
            dias = (tramo_fin - d).days
            if dias <= 0:
                d = tramo_fin
                continue
            key = f"{d.year}-{d.month:02d}"
            if key in MESES_EXCLUIDOS or date(d.year, d.month, 1) < FECHA_DESDE_ANALISIS:
                d = tramo_fin
                continue
            frac = dias / p.dias
            agg[key]["m3"] += p.m3 * frac
            agg[key]["clp"] += int(round(p.clp * frac))
            agg[key]["clp_esp"] += int(round(p.clp_esperado * frac))
            agg[key]["dias"] += dias
            d = tramo_fin

    rows = []
    for mes in sorted(agg.keys()):
        a = agg[mes]
        rows.append(
            {
                "mes": mes,
                "mes_label": _fmt_mes(mes),
                "estado": _estado_mes(mes),
                "dias": a["dias"],
                "m3": a["m3"],
                "clp": a["clp"],
                "clp_esperado": a["clp_esp"],
                "sobrecosto_clp": a["clp"] - a["clp_esp"],
            }
        )
    return rows


def _meses_sin_ultimo_incompleto(meses: List[dict]) -> List[dict]:
    if len(meses) > 1:
        return meses[:-1]
    return meses


def _promedios_con_sin_wes(meses: List[dict]) -> dict:
    """Promedio mensual prorrateado por estado (excluye último mes en curso)."""
    plot = _meses_sin_ultimo_incompleto(meses)
    con = [m for m in plot if m["estado"] == "Con WES"]
    sin = [m for m in plot if m["estado"] == "Sin WES"]
    n_con, n_sin = len(con), len(sin)

    def _avg(arr: List[dict], key: str) -> float:
        return sum(m[key] for m in arr) / len(arr) if arr else 0.0

    prom_m3_con = _avg(con, "m3")
    prom_m3_sin = _avg(sin, "m3")
    prom_clp_con = _avg(con, "clp")
    prom_clp_sin = _avg(sin, "clp")
    clp_m3_con = prom_clp_con / prom_m3_con if prom_m3_con else 0.0
    clp_m3_sin = prom_clp_sin / prom_m3_sin if prom_m3_sin else 0.0

    def _pct(a: float, b: float) -> float:
        return ((b / a) - 1) * 100 if a else 0.0

    proj_m3_con = prom_m3_con * MESES_PROYECCION_ANUAL
    proj_m3_sin = prom_m3_sin * MESES_PROYECCION_ANUAL
    proj_clp_con = int(round(prom_clp_con * MESES_PROYECCION_ANUAL))
    proj_clp_sin = int(round(prom_clp_sin * MESES_PROYECCION_ANUAL))

    return {
        "n_meses_con": n_con,
        "n_meses_sin": n_sin,
        "prom_m3_con": prom_m3_con,
        "prom_m3_sin": prom_m3_sin,
        "prom_clp_con": prom_clp_con,
        "prom_clp_sin": prom_clp_sin,
        "clp_m3_con": clp_m3_con,
        "clp_m3_sin": clp_m3_sin,
        "pct_m3": _pct(prom_m3_con, prom_m3_sin),
        "pct_clp": _pct(prom_clp_con, prom_clp_sin),
        "pct_clp_m3": _pct(clp_m3_con, clp_m3_sin),
        "proj_m3_con": proj_m3_con,
        "proj_m3_sin": proj_m3_sin,
        "proj_clp_con": proj_clp_con,
        "proj_clp_sin": proj_clp_sin,
        "sobrecosto_anual_clp": proj_clp_sin - proj_clp_con,
    }


def _grafico_proyeccion_anual(prom: dict, png: Path) -> None:
    """Barras = promedio mensual × 12 meses (proyección anual comparable)."""
    labels = [
        f"Con WES\n(prom. × {MESES_PROYECCION_ANUAL} meses)",
        f"Sin WES\n(prom. × {MESES_PROYECCION_ANUAL} meses)",
    ]
    clp_anual = [prom["proj_clp_con"] / 1_000_000, prom["proj_clp_sin"] / 1_000_000]
    colors = ["#3498DB", "#E74C3C"]
    fig, ax = plt.subplots(figsize=(7.5, 5))
    bars = ax.bar(labels, clp_anual, color=colors, width=0.45, edgecolor="white")
    ax.set_ylabel("Proyección anual (millones CLP)")
    pct = prom["pct_clp"]
    ax.set_title(
        f"Proyección anual — promedio mensual × {MESES_PROYECCION_ANUAL} meses\n"
        f"(Sin WES vs Con WES: {pct:+.1f}% CLP; {prom['pct_m3']:+.1f}% m³)",
        fontsize=10,
    )
    for bar, m3, clp in zip(
        bars,
        [prom["proj_m3_con"], prom["proj_m3_sin"]],
        [prom["proj_clp_con"], prom["proj_clp_sin"]],
    ):
        h = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            h + max(clp_anual) * 0.03,
            f"{_fmt_m3(m3)}\n{_fmt_clp(clp)}",
            ha="center",
            va="bottom",
            fontsize=9,
        )
    ax.set_ylim(0, max(clp_anual) * 1.22 if clp_anual else 1)
    fig.tight_layout()
    fig.savefig(png, dpi=180)
    plt.close(fig)


def _meses_grafico(meses: List[dict], excluir_ultimo: bool = True) -> List[dict]:
    """Meses para gráfico lineal; omite el último si el periodo aún no cierra."""
    if excluir_ultimo and len(meses) > 1:
        return meses[:-1]
    return meses


def _grafico_clp_mensual(meses: List[dict], png: Path) -> None:
    plot = _meses_grafico(meses)
    labels = [m["mes_label"] for m in plot]
    clp = [m["clp"] / 1_000_000 for m in plot]
    esp = [m["clp_esperado"] / 1_000_000 for m in plot]
    colors = ["#3498DB" if m["estado"] == "Con WES" else "#E74C3C" for m in plot]
    x = range(len(labels))
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(x, clp, color=colors, alpha=0.85, label="Facturado")
    ax.plot(x, esp, color="#2C3E50", marker="o", linewidth=1.5, label="Esperado línea base WES")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=45, ha="right")
    ax.set_title("Costo mensual prorrateado (millones CLP) — ene-2025 a la fecha", fontsize=11)
    ax.set_ylabel("M CLP")
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(png, dpi=180)
    plt.close(fig)


def _word(
    path: Path,
    periodos: List[Periodo],
    meses: List[dict],
    png_mes: Path,
    png_totales: Path,
    stats: dict,
    prom: dict,
) -> None:
    doc = Document()
    _configurar_pagina_compacta(doc)
    t = doc.add_heading("INFORME COMPARATIVO DE FACTURACIÓN", 0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    t.runs[0].font.color.rgb = _HEADING

    sub = doc.add_paragraph(CLIENTE)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].font.size = Pt(11)

    doc.add_paragraph(f"Cuenta Aguas Andinas: {CUENTA}  |  Medidor: {MEDIDOR}")
    doc.add_paragraph(f"Ventana de análisis: ene-2025 a la fecha  |  Generado: {datetime.now():%d-%m-%Y %H:%M}")
    doc.add_paragraph("")

    doc.add_heading("1. Antecedentes", 1)
    doc.add_paragraph(
        "Este informe compara el costo facturado por Aguas Andinas (CLP) con y sin monitoreo WES. "
        "La evaluación considera únicamente desde enero de 2025 hasta la última factura disponible."
    )
    doc.add_paragraph(
        f"Se excluye del análisis la factura {FACTURA_EXCLUIDA} (periodo dic-2024 / feb-2025) y el mes "
        "feb-2025, porque los equipos WES estuvieron deshabilitados durante la licitación y ese consumo "
        "no representa operación normal."
    )
    doc.add_paragraph(
        f"Retiro definitivo de equipos WES: {FECHA_RETIRO_WES.strftime('%d-%m-%Y')}. "
        "A partir de jul-2025 el recinto quedó sin monitoreo WES."
    )

    doc.add_heading("2. Resumen económico (ene-2025 a la fecha)", 1)
    p = doc.add_paragraph()
    p.add_run("Línea base con WES (facturas normales ene–jun 2025):\n").bold = True
    p.add_run(
        f"  • {stats['n_con_wes']} facturas | {stats['clp_m3_base']:.0f} CLP/m³ | "
        f"{_fmt_clp(int(stats['prom_clp_con']))}/periodo\n"
    )
    p.add_run("Con monitoreo WES (mar–jun 2025):\n").bold = True
    p.add_run(f"  • {_fmt_m3(stats['total_m3_con'])} | {_fmt_clp(stats['total_clp_con'])}\n")
    p.add_run(f"Tras retiro del monitoreo WES (jul-2025 a {stats.get('label_fin_sin', '—')}):\n").bold = True
    p.add_run(
        f"  • {_fmt_m3(stats['total_m3_sin'])} | {_fmt_clp(stats['total_clp_sin'])}\n"
        f"  • Sobrecosto vs línea base WES: {_fmt_clp(stats['sobrecosto_sin'])}\n"
    )
    p.add_run("Totales en ventana (excl. factura feb-2025):\n").bold = True
    p.add_run(
        f"  • Pagado a Aguas Andinas: {_fmt_clp(stats['total_clp_ventana'])}\n"
        f"  • Costo esperado si todo hubiera sido como con WES: {_fmt_clp(stats['total_clp_esperado'])}\n"
        f"  • Sobrecosto acumulado sin WES: {_fmt_clp(stats['sobrecosto_sin'])}\n"
    )
    doc.add_paragraph(
        f"(Factura {FACTURA_EXCLUIDA} excluida: {_fmt_clp(stats['clp_excluida'])} — no suma en totales.)"
    )

    _heading_compacto(doc, "3. Promedio mensual y proyección anual — Con WES vs Sin WES", 1)
    p3 = doc.add_paragraph(
        f"Promedios sobre el desglose mensual prorrateado "
        f"({prom['n_meses_con']} meses con WES, {prom['n_meses_sin']} meses sin WES; "
        f"sin el último mes en curso). Proyección anual = promedio mensual × {MESES_PROYECCION_ANUAL} meses."
    )
    _keep_with_next(p3)
    rows_p = [
        ("Indicador", "Con WES", "Sin WES", "Variación Sin vs Con"),
        (
            "Promedio mensual — m³",
            f"{prom['prom_m3_con']:.1f}",
            f"{prom['prom_m3_sin']:.1f}",
            f"{prom['pct_m3']:+.1f}%",
        ),
        (
            "Promedio mensual — CLP",
            _fmt_clp(int(round(prom["prom_clp_con"]))),
            _fmt_clp(int(round(prom["prom_clp_sin"]))),
            f"{prom['pct_clp']:+.1f}%",
        ),
        (
            "CLP/m³ (prom.)",
            f"{prom['clp_m3_con']:.0f}",
            f"{prom['clp_m3_sin']:.0f}",
            f"{prom['pct_clp_m3']:+.1f}%",
        ),
        (
            f"Proyección anual — m³ (×{MESES_PROYECCION_ANUAL})",
            _fmt_m3(prom["proj_m3_con"]),
            _fmt_m3(prom["proj_m3_sin"]),
            f"{prom['pct_m3']:+.1f}%",
        ),
        (
            f"Proyección anual — CLP (×{MESES_PROYECCION_ANUAL})",
            _fmt_clp(prom["proj_clp_con"]),
            _fmt_clp(prom["proj_clp_sin"]),
            f"{prom['pct_clp']:+.1f}%",
        ),
    ]
    tp = doc.add_table(rows=len(rows_p), cols=len(rows_p[0]))
    tp.style = "Light Grid Accent 1"
    for i, hdr in enumerate(rows_p[0]):
        tp.rows[0].cells[i].text = hdr
        tp.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for ri, row in enumerate(rows_p[1:], 1):
        for ci, val in enumerate(row):
            tp.rows[ri].cells[ci].text = val
    _tabla_compacta(tp, 9)
    _encabezado_tabla_repite(tp)

    _heading_compacto(doc, "4. Desglose mes a mes", 1)
    p4 = doc.add_paragraph(
        "Prorrateo por días del mes. «Esperado» = línea base CLP/m³ con WES (mar–jun 2025). "
        "Feb-2025 excluido."
    )
    _keep_with_next(p4)
    rows_m = [
        ("Mes", "Estado", "Días", "m³", "Facturado (CLP)", "Esperado (CLP)", "Sobrecosto (CLP)"),
    ]
    tot_clp = tot_esp = tot_sob = tot_m3 = 0.0
    for m in meses:
        tot_clp += m["clp"]
        tot_esp += m["clp_esperado"]
        tot_sob += m["sobrecosto_clp"]
        tot_m3 += m["m3"]
        rows_m.append(
            (
                m["mes_label"],
                m["estado"],
                str(m["dias"]),
                f"{m['m3']:.1f}",
                _fmt_clp(m["clp"]),
                _fmt_clp(m["clp_esperado"]),
                _fmt_clp(m["sobrecosto_clp"]),
            )
        )
    rows_m.append(
        ("TOTAL", "", "", _fmt_m3(tot_m3), _fmt_clp(tot_clp), _fmt_clp(tot_esp), _fmt_clp(tot_sob))
    )
    tm = doc.add_table(rows=len(rows_m), cols=len(rows_m[0]))
    tm.style = "Light Grid Accent 1"
    for i, hdr in enumerate(rows_m[0]):
        c = tm.rows[0].cells[i]
        c.text = hdr
        c.paragraphs[0].runs[0].font.bold = True
    for ri, row in enumerate(rows_m[1:], 1):
        for ci, val in enumerate(row):
            cell = tm.rows[ri].cells[ci]
            cell.text = val
            if ri == len(rows_m) - 1:
                cell.paragraphs[0].runs[0].font.bold = True
    _tabla_compacta(tm, 8.5)
    _encabezado_tabla_repite(tm)

    if png_totales.exists():
        _add_picture_compact(
            doc,
            png_totales,
            12.5,
            f"Gráfico — proyección anual (promedio mensual × {MESES_PROYECCION_ANUAL} meses):",
        )
    if png_mes.exists():
        _add_picture_compact(
            doc,
            png_mes,
            13.5,
            "Gráfico — costo mensual prorrateado (CLP; sin el último mes en curso):",
        )

    doc.add_page_break()
    _heading_compacto(doc, "5. Referencia por factura (anexo)", 1)
    doc.add_paragraph("Detalle compacto de las boletas incluidas en la ventana ene-2025 a la fecha.")
    rows = [("Factura", "Periodo lectura", "m³", "Total (CLP)", "Estado")]
    for p in periodos:
        rows.append(
            (
                p.boleta,
                f"{p.lectura_desde:%d-%m-%Y} → {p.lectura_hasta:%d-%m-%Y}",
                f"{p.m3:,}".replace(",", "."),
                _fmt_clp(p.clp),
                p.etiqueta,
            )
        )
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, hdr in enumerate(rows[0]):
        table.rows[0].cells[i].text = hdr
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for ri, row in enumerate(rows[1:], 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    _tabla_compacta(table, 9)

    doc.add_heading("6. Conclusiones", 1)
    for line in stats["conclusiones"]:
        doc.add_paragraph(line, style="List Bullet")

    nota = doc.add_paragraph(
        "Nota: Montos = TOTAL A PAGAR (Aguas Andinas). Ventana ene-2025+. Excluye factura "
        f"{FACTURA_EXCLUIDA} y mes feb-2025. Línea base CLP/m³ calculada solo con facturas con WES "
        "en esa ventana (mar–jun 2025)."
    )
    nota.runs[0].font.italic = True
    nota.runs[0].font.size = Pt(9)
    doc.save(str(path))


def main() -> int:
    FACTURAS_DIR.mkdir(parents=True, exist_ok=True)
    OUT_BASE.mkdir(parents=True, exist_ok=True)

    if not any(FACTURAS_DIR.glob("*.pdf")):
        print(f"[ERROR] No hay PDFs en {FACTURAS_DIR}")
        return 1

    todos = _cargar_periodos()
    excluida = next((p for p in todos if p.boleta == FACTURA_EXCLUIDA), None)
    periodos = _filtrar_analisis(todos)

    con_wes = [p for p in periodos if p.lectura_hasta <= FECHA_FIN_CON_WES]
    sin_wes = [p for p in periodos if p.lectura_hasta > FECHA_FIN_CON_WES]

    clp_m3_base = _aplicar_sobrecosto(periodos, con_wes)
    meses = _prorrateo_mensual(periodos)
    prom = _promedios_con_sin_wes(meses)

    prom_clp_con = sum(p.clp for p in con_wes) / len(con_wes) if con_wes else 0
    total_m3_con = sum(p.m3 for p in con_wes)
    total_m3_sin = sum(p.m3 for p in sin_wes)
    total_clp_con = sum(p.clp for p in con_wes)
    total_clp_sin = sum(p.clp for p in sin_wes)
    total_clp_ventana = sum(p.clp for p in periodos)
    total_m3_ventana = sum(p.m3 for p in periodos)
    total_clp_esperado = sum(p.clp_esperado for p in periodos)
    sobrecosto_sin = sum(p.sobrecosto_clp for p in sin_wes)
    clp_excluida = excluida.clp if excluida else 0

    def _m3_fmt(n: float) -> str:
        return _fmt_m3(n)

    ultima_sin = max((p.lectura_hasta for p in sin_wes), default=None)
    ultima_ventana = max((p.lectura_hasta for p in periodos), default=None)
    fin_sin = _fmt_mes_corto(ultima_sin) if ultima_sin else "—"
    fin_ventana = _fmt_mes_corto(ultima_ventana) if ultima_ventana else "—"

    conclusiones = [
        f"Con monitoreo WES activo (mar–jun 2025) el recinto facturó {_m3_fmt(total_m3_con)} "
        f"por {_fmt_clp(total_clp_con)} (línea base {clp_m3_base:.0f} CLP/m³).",
        f"La factura {FACTURA_EXCLUIDA} ({_fmt_clp(clp_excluida)}) quedó excluida por la deshabilitación "
        "de equipos durante la licitación; feb-2025 no entra en totales ni en el desglose mensual.",
        f"Tras el retiro del monitoreo WES (jul-2025 a {fin_sin}) se facturaron {_m3_fmt(total_m3_sin)} "
        f"por {_fmt_clp(total_clp_sin)}; el sobrecosto respecto a la línea base WES fue "
        f"{_fmt_clp(sobrecosto_sin)}.",
        f"En la ventana ene-2025 a {fin_ventana} (sin feb anómalo) el total fue {_m3_fmt(total_m3_ventana)} "
        f"por {_fmt_clp(total_clp_ventana)}; al ritmo con WES habría sido {_fmt_clp(total_clp_esperado)} "
        f"(diferencia {_fmt_clp(total_clp_ventana - total_clp_esperado)}).",
        f"El promedio mensual sin WES ({_fmt_clp(int(round(prom['prom_clp_sin'])))} / "
        f"{prom['prom_m3_sin']:.0f} m³) supera en {prom['pct_clp']:+.1f}% el costo y en {prom['pct_m3']:+.1f}% "
        f"el volumen del promedio con WES ({_fmt_clp(int(round(prom['prom_clp_con'])))} / "
        f"{prom['prom_m3_con']:.0f} m³).",
        f"Proyectado a 12 meses: con WES {_fmt_m3(prom['proj_m3_con'])} por {_fmt_clp(prom['proj_clp_con'])}; "
        f"sin WES {_fmt_m3(prom['proj_m3_sin'])} por {_fmt_clp(prom['proj_clp_sin'])} "
        f"(diferencia anual {_fmt_clp(prom['sobrecosto_anual_clp'])}).",
    ]

    stats = {
        "n_con_wes": len(con_wes),
        "clp_m3_base": clp_m3_base,
        "prom_clp_con": prom_clp_con,
        "total_m3_con": total_m3_con,
        "total_m3_sin": total_m3_sin,
        "total_clp_con": total_clp_con,
        "total_clp_sin": total_clp_sin,
        "total_clp_ventana": total_clp_ventana,
        "total_clp_esperado": total_clp_esperado,
        "sobrecosto_sin": sobrecosto_sin,
        "clp_excluida": clp_excluida,
        "label_fin_sin": fin_sin,
        "label_fin_ventana": fin_ventana,
        "conclusiones": conclusiones,
    }

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    csv_mes = OUT_BASE / f"facturaciones_genchi_ccp_mensual_{ts}.csv"
    with csv_mes.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(["mes", "estado", "dias", "m3", "clp", "clp_esperado", "sobrecosto_clp"])
        for m in meses:
            w.writerow(
                [
                    m["mes_label"],
                    m["estado"],
                    m["dias"],
                    round(m["m3"], 2),
                    m["clp"],
                    m["clp_esperado"],
                    m["sobrecosto_clp"],
                ]
            )

    csv_path = OUT_BASE / f"facturaciones_genchi_ccp_{ts}.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(["factura", "desde", "hasta", "m3", "total_clp", "clp_esperado", "sobrecosto_clp", "estado"])
        for p in periodos:
            w.writerow(
                [p.boleta, p.lectura_desde, p.lectura_hasta, p.m3, p.clp, p.clp_esperado, p.sobrecosto_clp, p.etiqueta]
            )

    png_mes = OUT_BASE / f"grafico_clp_mensual_{ts}.png"
    png_tot = OUT_BASE / f"grafico_proyeccion_anual_{ts}.png"
    _grafico_clp_mensual(meses, png_mes)
    _grafico_proyeccion_anual(prom, png_tot)

    docx = OUT_BASE / f"Informe_Genchi_CCP_Santiago_Sur_facturaciones_{ts}.docx"
    _word(docx, periodos, meses, png_mes, png_tot, stats, prom)

    pdf = docx.with_suffix(".pdf")
    try:
        from generar_reporte_word import convertir_word_a_pdf

        p = convertir_word_a_pdf(docx)
        if p and Path(p).is_file():
            pdf = Path(p)
    except Exception:
        pass

    print("=" * 72)
    print(f"Ventana: ene-2025+ | Facturas: {len(periodos)} | Excluida: {FACTURA_EXCLUIDA}")
    print(f"Total ventana: {_fmt_clp(total_clp_ventana)} | Sobrecosto sin WES: {_fmt_clp(sobrecosto_sin)}")
    print(f"DOCX: {docx}")
    print(f"PDF:  {pdf}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
