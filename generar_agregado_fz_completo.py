"""
Reporte agregado Fundo Zapallar (FZ) — formato completo solicitado:

  - Consumo total (referencia Matriz ESVAL)
  - Consumo nocturno 00:00–06:59 (hora Chile) + gráfico de barras amarillo
  - Umbral recomendado para alertas de máximo (baseline diario × 1,25)
  - Comparativo del mes (nocturno vs efectivo)
  - Comparativo de los últimos 6 meses (entrada ESVAL)

Uso:
  python generar_agregado_fz_completo.py
  python generar_agregado_fz_completo.py --start-date 01/07/2026 --end-date 30/07/2026
"""

from __future__ import annotations

import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

if sys.platform == "win32":
    for s in (sys.stdout, sys.stderr):
        try:
            s.reconfigure(encoding="utf-8", line_buffering=True)
        except Exception:
            pass

from exclusiones_reportes import FUNDO_ZAPALLAR_NODE_IDS
from generar_reporte_word import (
    add_formatted_heading,
    add_formatted_title,
    add_logo_to_header,
    add_picture_with_pagination,
    build_consumption_chart,
    build_monthly_comparison_chart,
    calculate_nocturnal_metrics,
    estilizar_tabla_wes,
    fetch_json,
    flatten_measures,
    format_currency_chilean,
    format_number_chilean,
    get_company_name,
    get_node_name,
    get_water_price_per_m3,
    normalize_measures_payload,
    parse_date,
    summarize_consumption,
    acl_node_base_url,
)
from generar_consolidado_m3_mensual_puente_alto import consumo_mes_un_nodo

COMPANY_ID = "000027"
ESVAL_ID = "000027-01"
NODE_IDS = list(FUNDO_ZAPALLAR_NODE_IDS)
MULT_UMBRAL = 1.25
PRECIO_DEFAULT = 1200.0
COLOR_BARRA = "#0050b3"
COLOR_NOCTURNO = "#FFD700"
COLOR_NOCTURNO_EDGE = "#DAA520"


def _fmt_api(dt: datetime) -> str:
    return dt.strftime("%d%m%Y")


def _fetch_node(node_id: str, start_dt: datetime, end_dt: datetime) -> dict:
    name = get_node_name(node_id)
    print(f"  [data] {node_id} {name}...", flush=True)
    try:
        raw = fetch_json(
            f"{acl_node_base_url()}/nodes/measures/dates",
            params=[
                ("id", node_id),
                ("start", _fmt_api(start_dt)),
                ("end", _fmt_api(end_dt)),
            ],
        )
        payload = normalize_measures_payload(raw, node_id)
        measures = flatten_measures(payload)
        summary = summarize_consumption(measures)
        noct = calculate_nocturnal_metrics(node_id, start_dt, end_dt)
        price = get_water_price_per_m3(COMPANY_ID, node_id, payload)
        # Baseline diario y máximo diario del periodo
        daily: Dict[date, float] = {}
        for m in measures:
            d0 = m.date.date() if hasattr(m.date, "date") else m.date
            daily[d0] = daily.get(d0, 0.0) + float(m.total_m3 or 0.0)
        vals = [v for v in daily.values() if v > 0]
        baseline = (sum(vals) / len(vals)) if vals else 0.0
        max_diario = max(vals) if vals else 0.0
        return {
            "node_id": node_id,
            "node_name": name,
            "summary": summary,
            "measures": measures,
            "nocturno_m3": float(noct.get("consumo_nocturno_total") or 0.0),
            "diurno_m3": float(noct.get("consumo_diurno_efectivo") or 0.0),
            "dias_nocturno": int(noct.get("dias_con_consumo_nocturno") or 0),
            "baseline_diario": baseline,
            "max_diario": max_diario,
            "umbral": round(baseline * MULT_UMBRAL, 1) if baseline > 0 else None,
            "price": price if price else PRECIO_DEFAULT,
            "dias_con_data": len(vals),
        }
    except Exception as e:
        print(f"  [WARN] {node_id}: {e}", flush=True)
        return {
            "node_id": node_id,
            "node_name": name,
            "summary": summarize_consumption([]),
            "measures": [],
            "nocturno_m3": 0.0,
            "diurno_m3": 0.0,
            "dias_nocturno": 0,
            "baseline_diario": 0.0,
            "max_diario": 0.0,
            "umbral": None,
            "price": PRECIO_DEFAULT,
            "dias_con_data": 0,
            "error": str(e),
        }


def _plot_barras_periodo(nodes: List[dict], out: Path) -> Path:
    pairs = sorted(
        ((n["node_name"], float(n["summary"]["total"])) for n in nodes),
        key=lambda x: -x[1],
    )
    labels = [p[0] for p in pairs]
    vals = [p[1] for p in pairs]
    fig, ax = plt.subplots(figsize=(12, 6))
    bars = ax.bar(labels, vals, color=COLOR_BARRA)
    ax.set_ylabel("Consumo (m³)", fontsize=12, fontweight="bold")
    ax.set_title(
        "Consumo del periodo por punto (FZ: ESVAL = entrada real)",
        fontsize=13,
        fontweight="bold",
    )
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=35, ha="right", fontsize=9)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 1)}",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _plot_nocturno_amarillo(nodes: List[dict], out: Path) -> Path:
    """Barras amarillas de consumo nocturno por punto (estilo reportes agregados WES)."""
    pairs = sorted(
        ((n["node_name"], float(n["nocturno_m3"])) for n in nodes),
        key=lambda x: -x[1],
    )
    labels = [p[0] for p in pairs]
    vals = [p[1] for p in pairs]
    n_pts = max(1, len(labels))
    fig_w = max(8.0, min(12.0, n_pts * 1.1))
    fig, ax = plt.subplots(figsize=(fig_w, 5.5))
    bars = ax.bar(
        labels,
        vals,
        color=COLOR_NOCTURNO,
        alpha=0.9,
        edgecolor=COLOR_NOCTURNO_EDGE,
        linewidth=1.2,
    )
    ax.set_ylabel("Consumo nocturno (m³)", fontsize=12, fontweight="bold")
    ax.set_title(
        "Consumo nocturno por punto de monitoreo (00:00–06:59)",
        fontsize=13,
        fontweight="bold",
    )
    ax.set_ylim(bottom=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=35, ha="right", fontsize=9)
    for bar, v in zip(bars, vals):
        if v > 0:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height(),
                f"{format_number_chilean(v, 1)} m³",
                ha="center",
                va="bottom",
                fontsize=9,
                fontweight="bold",
            )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _meses_ultimos_6(end_d: date) -> List[Tuple[int, int]]:
    """Últimos 6 meses civiles incluyendo el mes de end_d."""
    y, m = end_d.year, end_d.month
    out: List[Tuple[int, int]] = []
    for _ in range(6):
        out.append((y, m))
        m -= 1
        if m == 0:
            m = 12
            y -= 1
    out.reverse()
    return out


def _plot_6_meses(series: List[Tuple[str, float]], out: Path) -> Path:
    labels = [s[0] for s in series]
    vals = [s[1] for s in series]
    fig, ax = plt.subplots(figsize=(10, 5.5))
    bars = ax.bar(labels, vals, color=COLOR_BARRA, width=0.65)
    ax.set_ylabel("Consumo mensual ESVAL (m³)", fontsize=11, fontweight="bold")
    ax.set_xlabel("Mes", fontsize=11, fontweight="bold")
    ax.set_title(
        "Comparativo últimos 6 meses — Matriz ESVAL (entrada al fundo)",
        fontsize=12,
        fontweight="bold",
    )
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 0)}",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _add_table(
    doc: Document,
    headers: List[str],
    rows: List[List[str]],
    *,
    has_total_row: bool = False,
) -> None:
    """Tabla con estilo WES (encabezado azul, filas alternadas)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            table.rows[r_i + 1].cells[c_i].text = str(val)
    estilizar_tabla_wes(table, has_total_row=has_total_row)
    doc.add_paragraph("")


def generar(start_date: str, end_date: str) -> Path:
    start_dt = parse_date(start_date)
    end_dt = parse_date(end_date, end_of_day=True)
    company_name = get_company_name(COMPANY_ID)

    print("=" * 70)
    print(f"AGREGADO FZ COMPLETO — {company_name}")
    print(f"Periodo: {start_date} → {end_date}")
    print(f"Nodos: {', '.join(NODE_IDS)}")
    print("=" * 70)

    t0 = time.perf_counter()
    nodes: List[dict] = []
    with ThreadPoolExecutor(max_workers=min(8, len(NODE_IDS))) as ex:
        futs = {ex.submit(_fetch_node, nid, start_dt, end_dt): nid for nid in NODE_IDS}
        for fut in as_completed(futs):
            nodes.append(fut.result())
    nodes.sort(key=lambda n: n["node_id"])

    esval = next((n for n in nodes if n["node_id"] == ESVAL_ID), None)
    esval_m3 = float(esval["summary"]["total"]) if esval else 0.0
    price = float(esval["price"]) if esval else PRECIO_DEFAULT

    # Salida
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    out_dir = Path("reports") / "Fundo_Zapallar" / "ABREGADO" / f"AGREGADO_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    chart_barras = _plot_barras_periodo(nodes, out_dir / "chart_consumo_periodo.png")
    chart_nocturno = _plot_nocturno_amarillo(
        nodes, out_dir / "chart_consumo_nocturno_nodos.png"
    )

    # Nocturno agregado (suma puntos; para FZ el total de entrada es ESVAL)
    noct_esval = float(esval["nocturno_m3"]) if esval else 0.0
    diurno_esval = float(esval["diurno_m3"]) if esval else 0.0
    num_dias = (end_dt.date() - start_dt.date()).days + 1
    # Proyección a 30 días para comparativo del mes
    factor_30 = 30.0 / max(num_dias, 1)
    leak_monthly = noct_esval * factor_30
    efectivo_monthly = max(0.0, (esval_m3 - noct_esval) * factor_30)
    chart_mes = out_dir / "chart_comparacion_mensual.png"
    built_mes = build_monthly_comparison_chart(
        leak_monthly, efectivo_monthly, price, chart_mes
    )

    # Últimos 6 meses ESVAL
    print("[INFO] Descargando últimos 6 meses ESVAL...", flush=True)
    sess = requests.Session()
    series_6: List[Tuple[str, float]] = []
    for y, m in _meses_ultimos_6(end_dt.date()):
        # Mes en curso: acotar al end_date si es julio incompleto... usamos mes civil completo
        # salvo el mes actual del periodo (hasta end_dt).
        if y == end_dt.year and m == end_dt.month:
            # suma diaria del periodo ya cargado
            m3 = esval_m3
            label = f"{y}-{m:02d}*"
        else:
            m3, _, _ = consumo_mes_un_nodo(sess, ESVAL_ID, y, m)
            label = f"{y}-{m:02d}"
        series_6.append((label, float(m3)))
        print(f"  {label}: {m3:.1f} m³", flush=True)
    chart_6m = _plot_6_meses(series_6, out_dir / "chart_ultimos_6_meses_esval.png")

    # Word
    doc = Document()
    add_logo_to_header(doc)
    title = doc.add_paragraph(f"Reporte Agregado — {company_name}")
    title.style = "Title"
    for run in title.runs:
        run.font.size = Pt(22)
    sub = doc.add_paragraph(
        f"MONITOREO WES\n"
        f"Análisis consolidado de {len(nodes)} puntos\n"
        f"Rango: {start_dt.strftime('%d-%m-%y')} – {end_dt.strftime('%d-%m-%y')}\n"
        f"Generado: {datetime.now().strftime('%d-%m-%y')}"
    )
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 1) Consumo total
    add_formatted_heading(doc, "1. Consumo total", level=1)
    p = doc.add_paragraph(
        f"El consumo real del fundo en el periodo corresponde a la Matriz ESVAL "
        f"(entrada de agua): {format_number_chilean(esval_m3, 1)} m³. "
        f"Estanques y etapas miden caudales aguas abajo y no se suman al total "
        f"(evitar doble conteo)."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_barras), Inches(6), keep_with_next=True)

    # 1b) Consumo diario del mes por punto (sin marcadores de alerta)
    add_formatted_heading(doc, "1.1 Consumo diario por punto (mes)", level=1)
    p_d = doc.add_paragraph(
        "Evolución del consumo diario (m³) de cada punto en el periodo. "
        "Solo la curva de consumo; sin marcadores de alertas."
    )
    p_d.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for n in nodes:
        measures = n.get("measures") or []
        if not measures:
            continue
        chart_path = out_dir / f"zapallar_diario_{n['node_id'].replace('-', '_')}.png"
        built = build_consumption_chart(
            measures, chart_path, start_dt, end_dt, alerts=None
        )
        if not built or not chart_path.is_file():
            continue
        doc.add_paragraph("")
        add_formatted_title(doc, n["node_name"].upper())
        add_picture_with_pagination(doc, str(chart_path), Inches(6), keep_with_next=True)

    # 2) Nocturno + barras amarillas
    add_formatted_heading(doc, "2. Consumo nocturno (00:00 a 06:59)", level=1)
    p2 = doc.add_paragraph(
        "Se define como consumo nocturno el volumen medido entre las 00:00 y las 06:59 "
        "(hora Chile). El total nocturno de referencia del fundo es el de Matriz ESVAL."
    )
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2b = doc.add_paragraph(
        f"Nocturno ESVAL: {format_number_chilean(noct_esval, 1)} m³ "
        f"({format_currency_chilean(noct_esval * price)}) en {num_dias} días. "
        f"Diurno ESVAL: {format_number_chilean(diurno_esval, 1)} m³."
    )
    p2b.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_nocturno), Inches(6), keep_with_next=True)

    noct_total_pts = sum(float(n["nocturno_m3"]) for n in nodes)
    rows_noct = []
    for n in sorted(nodes, key=lambda x: -x["nocturno_m3"]):
        rows_noct.append(
            [
                n["node_id"],
                n["node_name"],
                format_number_chilean(n["nocturno_m3"], 1),
                str(n["dias_nocturno"]),
                format_currency_chilean(n["nocturno_m3"] * price),
            ]
        )
    rows_noct.append(
        [
            "",
            "Total puntos (suma)",
            format_number_chilean(noct_total_pts, 1),
            "",
            format_currency_chilean(noct_total_pts * price),
        ]
    )
    _add_table(
        doc,
        ["Nodo", "Punto", "Nocturno (m³)", "Días con nocturno", "Costo (CLP)"],
        rows_noct,
        has_total_row=True,
    )

    # 3) Umbral alertas máximo
    add_formatted_heading(doc, "3. Umbral para alertas de máximo", level=1)
    p3 = doc.add_paragraph(
        f"Baseline = promedio diario del periodo (días con consumo > 0). "
        f"Umbral recomendado = baseline × {MULT_UMBRAL:.2f} (+25 %). "
        f"Sirve como referencia para configurar alertas de consumo máximo diario."
    )
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    rows_u = []
    for n in nodes:
        umbral = n["umbral"]
        rows_u.append(
            [
                n["node_id"],
                n["node_name"],
                str(n["dias_con_data"]),
                format_number_chilean(n["baseline_diario"], 1) if n["baseline_diario"] else "—",
                format_number_chilean(n["max_diario"], 1) if n["max_diario"] else "—",
                format_number_chilean(umbral, 1) if umbral is not None else "—",
            ]
        )
    _add_table(
        doc,
        [
            "Nodo",
            "Punto",
            "Días con data",
            "Baseline diario (m³)",
            "Máx. diario (m³)",
            f"Umbral ×{MULT_UMBRAL:.2f} (m³/día)",
        ],
        rows_u,
        has_total_row=False,
    )

    # 4) Comparativo del mes
    add_formatted_heading(doc, "4. Comparativo del mes (nocturno vs efectivo)", level=1)
    p5 = doc.add_paragraph(
        f"Proyección a 30 días a partir del periodo ({num_dias} días), usando Matriz ESVAL: "
        f"nocturno proyectado {format_number_chilean(leak_monthly, 1)} m³ "
        f"({format_currency_chilean(leak_monthly * price)}) y consumo efectivo "
        f"{format_number_chilean(efectivo_monthly, 1)} m³ "
        f"({format_currency_chilean(efectivo_monthly * price)})."
    )
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if built_mes and Path(built_mes).exists():
        add_picture_with_pagination(doc, str(chart_mes), Inches(4.5), keep_with_next=True)
    else:
        doc.add_paragraph("No fue posible generar la gráfica de comparación mensual.")

    # 5) Últimos 6 meses
    add_formatted_heading(doc, "5. Comparativo últimos 6 meses", level=1)
    p6 = doc.add_paragraph(
        "Consumo mensual de Matriz ESVAL (entrada real al fundo). "
        "El mes marcado con * corresponde al periodo de este reporte "
        f"(hasta {end_dt.strftime('%d/%m/%Y')}), no necesariamente al mes civil completo."
    )
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_6m), Inches(6), keep_with_next=True)
    rows_6 = [[lab, format_number_chilean(v, 1)] for lab, v in series_6]
    _add_table(doc, ["Mes", "ESVAL (m³)"], rows_6, has_total_row=False)

    # Cierre
    add_formatted_heading(doc, "Conclusión", level=1)
    concl = doc.add_paragraph(
        f"En el periodo, Fundo Zapallar registró {format_number_chilean(esval_m3, 1)} m³ "
        f"de entrada (ESVAL), de los cuales {format_number_chilean(noct_esval, 1)} m³ "
        f"correspondieron a horario nocturno. Se recomienda validar umbrales de alerta "
        f"de máximo en terreno y seguir la evolución mensual de ESVAL para detectar desvíos."
    )
    concl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    fname = (
        f"Reporte_Agregado_Fundo_Zapallar_{start_dt:%Y%m%d}_{end_dt:%Y%m%d}.docx"
    )
    out_path = out_dir / fname
    doc.save(out_path)
    print(f"[OK] {out_path}")
    print(f"[INFO] Tiempo: {time.perf_counter() - t0:.1f} s")
    return out_path


def main() -> int:
    import argparse

    ap = argparse.ArgumentParser(description="Agregado FZ completo")
    ap.add_argument("--start-date", default="01/07/2026")
    ap.add_argument("--end-date", default="30/07/2026")
    args = ap.parse_args()
    generar(args.start_date, args.end_date)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
