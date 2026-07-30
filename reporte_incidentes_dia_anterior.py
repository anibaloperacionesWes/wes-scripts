"""
Script para generar reporte de incidentes del día anterior.
Analiza todos los puntos y reporta:
- Puntos con consumo cero
- Puntos sin respuesta (sin data)
- Puntos con data incompleta (menos de 24 horas, indicando horas faltantes)

Uso:
    python reporte_incidentes_dia_anterior.py
"""

import argparse
import sys
import requests
import smtplib
from pathlib import Path
from datetime import datetime, timedelta, timezone
from typing import Dict, List, Tuple, Optional
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configurar codificación UTF-8
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

# URLs base de las APIs
BASE_URL = "http://104.248.53.141:7003/wes/api/acl-node/v1"
ENTITY_BASE_URL = "http://104.248.53.141:7001/wes/api/acl-entities/v1"


def obtener_todos_los_nodos() -> List[Dict[str, str]]:
    """
    Obtiene todos los nodos del sistema iterando por empresas.
    Retorna lista de diccionarios con nodeId, nodeName, companyId, companyName
    """
    all_nodes = []
    print("Obteniendo todos los nodos del sistema...")
    print("=" * 60)
    
    # Iterar por empresas desde 000000 hasta 000100
    for i in range(101):  # 000000 a 000100
        company_id = f"{i:06d}"
        url = f"{ENTITY_BASE_URL}/companies/{company_id}"
        
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                data = response.json()
                company_name = data.get("name", "").strip()
                
                if company_name:
                    # Obtener nodos de la empresa
                    nodes = data.get("nodes", [])
                    for node in nodes:
                        node_id = node.get("nodeId", "")
                        node_name = node.get("name", "").strip()
                        if node_id and node_name:
                            all_nodes.append({
                                "nodeId": node_id,
                                "nodeName": node_name,
                                "companyId": company_id,
                                "companyName": company_name
                            })
                    if nodes:
                        print(f"[OK] {company_id} ({company_name}): {len(nodes)} nodos")
        except requests.RequestException:
            # Empresa no existe o error, continuar
            pass
        except Exception as e:
            print(f"[ERROR] {company_id}: {e}")
    
    print("=" * 60)
    print(f"Total nodos encontrados: {len(all_nodes)}")
    return all_nodes


def obtener_datos_horarios_dia(node_id: str, fecha: datetime) -> Tuple[Optional[Dict], str]:
    """
    Obtiene datos horarios para un día específico.
    
    Returns:
        Tuple[dict, str]: (datos_horarios, estado)
        - datos_horarios: dict con {hora: valor} o None si hay error
        - estado: "ok", "sin_respuesta", "error"
    """
    date_str = fecha.strftime("%d%m%Y")
    url = f"{BASE_URL}/nodes/{node_id}/dates.measures.csv"
    params = [("start", date_str), ("end", date_str)]
    
    try:
        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        
        # Parsear CSV
        csv_content = response.text
        hourly_data = {}
        lines = csv_content.strip().split('\n')
        
        # Saltar encabezado
        for line in lines[1:]:
            if not line.strip():
                continue
            
            parts = line.split(',')
            if len(parts) >= 2:
                try:
                    time_str = parts[0].strip()
                    value_str = parts[1].strip()
                    
                    # Extraer hora del formato ISO
                    if 'T' in time_str:
                        hour_part = time_str.split('T')[1]
                        hour = int(hour_part.split(':')[0])
                    else:
                        dt = datetime.fromisoformat(time_str.replace('Z', '+00:00'))
                        hour = dt.hour
                    
                    value = float(value_str)
                    hourly_data[hour] = value
                except (ValueError, TypeError, IndexError):
                    continue
        
        return hourly_data, "ok"
    
    except requests.exceptions.RequestException:
        return None, "sin_respuesta"
    except Exception:
        return None, "error"


def analizar_datos_horarios(hourly_data: Dict[int, float]) -> Tuple[str, List[int]]:
    """
    Analiza los datos horarios y determina el estado.
    
    Returns:
        Tuple[str, List[int]]: (estado, horas_faltantes)
        - estado: "consumo_cero", "datos_incompletos", "ok"
        - horas_faltantes: lista de horas que faltan (0-23)
    """
    if not hourly_data:
        return "sin_datos", list(range(24))
    
    # Verificar horas faltantes
    horas_presentes = set(hourly_data.keys())
    horas_esperadas = set(range(24))
    horas_faltantes = sorted(list(horas_esperadas - horas_presentes))
    
    # Verificar si todos los valores son cero
    valores = list(hourly_data.values())
    consumo_total = sum(valores)
    
    if consumo_total == 0:
        return "consumo_cero", horas_faltantes
    
    if len(horas_presentes) < 24:
        return "datos_incompletos", horas_faltantes
    
    return "ok", horas_faltantes


def crear_reporte_word(
    fecha_analisis: datetime,
    puntos_cero: List[Dict],
    puntos_sin_respuesta: List[Dict],
    puntos_incompletos: List[Dict],
    output_dir: Path
) -> Path:
    """
    Crea un documento Word con el reporte de incidentes.
    """
    doc = Document()
    
    # Título principal
    title = doc.add_heading("REPORTE DE INCIDENTES - CONSUMO DÍA ANTERIOR", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
    title_run.bold = True
    
    # Fecha de análisis
    fecha_str = fecha_analisis.strftime("%d-%m-%Y")
    fecha_para = doc.add_paragraph(f"Fecha analizada: {fecha_str}")
    fecha_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fecha_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    fecha_generacion = datetime.now(timezone.utc).strftime("%d-%m-%Y %H:%M")
    gen_para = doc.add_paragraph(f"Reporte generado: {fecha_generacion}")
    gen_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    gen_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph("")  # Espacio
    
    # Resumen ejecutivo
    doc.add_heading("RESUMEN EJECUTIVO", 1)
    summary_para = doc.add_paragraph(
        f"Total de puntos analizados: {len(puntos_cero) + len(puntos_sin_respuesta) + len(puntos_incompletos)}\n"
        f"Puntos con consumo cero: {len(puntos_cero)}\n"
        f"Puntos sin respuesta: {len(puntos_sin_respuesta)}\n"
        f"Puntos con datos incompletos: {len(puntos_incompletos)}"
    )
    summary_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph("")  # Espacio
    
    # Sección: Puntos con consumo cero
    if puntos_cero:
        doc.add_heading("PUNTOS CON CONSUMO CERO", 1)
        doc.add_paragraph(
            "Los siguientes puntos registraron consumo cero durante todo el día analizado:"
        )
        
        rows = [("Nodo ID", "Nombre del Punto", "Empresa")]
        for punto in puntos_cero:
            rows.append((
                punto["nodeId"],
                punto["nodeName"],
                punto["companyName"]
            ))
        
        # Crear tabla
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        header_cells = table.rows[0].cells
        for i, header in enumerate(rows[0]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Datos
        for row_idx, row_data in enumerate(rows[1:], start=1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = str(value)
                table.rows[row_idx].cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        doc.add_paragraph("")  # Espacio
    
    # Sección: Puntos sin respuesta
    if puntos_sin_respuesta:
        doc.add_heading("PUNTOS SIN RESPUESTA", 1)
        doc.add_paragraph(
            "Los siguientes puntos no respondieron a la consulta de datos (error de API o sin conexión):"
        )
        
        rows = [("Nodo ID", "Nombre del Punto", "Empresa")]
        for punto in puntos_sin_respuesta:
            rows.append((
                punto["nodeId"],
                punto["nodeName"],
                punto["companyName"]
            ))
        
        # Crear tabla
        table = doc.add_table(rows=len(rows), cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        header_cells = table.rows[0].cells
        for i, header in enumerate(rows[0]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Datos
        for row_idx, row_data in enumerate(rows[1:], start=1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = str(value)
                table.rows[row_idx].cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        doc.add_paragraph("")  # Espacio
    
    # Sección: Puntos con datos incompletos
    if puntos_incompletos:
        doc.add_heading("PUNTOS CON DATOS INCOMPLETOS", 1)
        doc.add_paragraph(
            "Los siguientes puntos tienen menos de 24 horas de datos. Se indican las horas faltantes:"
        )
        
        rows = [("Nodo ID", "Nombre del Punto", "Empresa", "Horas Faltantes")]
        for punto in puntos_incompletos:
            horas_faltantes_str = ", ".join([f"{h:02d}:00" for h in punto["horasFaltantes"]])
            if len(horas_faltantes_str) > 100:
                horas_faltantes_str = horas_faltantes_str[:100] + "..."
            rows.append((
                punto["nodeId"],
                punto["nodeName"],
                punto["companyName"],
                horas_faltantes_str
            ))
        
        # Crear tabla
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Ajustar ancho de columnas
        table.columns[0].width = Inches(1.5)  # Nodo ID
        table.columns[1].width = Inches(2.5)  # Nombre
        table.columns[2].width = Inches(2.0)  # Empresa
        table.columns[3].width = Inches(2.0)  # Horas faltantes
        
        # Encabezados
        header_cells = table.rows[0].cells
        for i, header in enumerate(rows[0]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Datos
        for row_idx, row_data in enumerate(rows[1:], start=1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = str(value)
                table.rows[row_idx].cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        doc.add_paragraph("")  # Espacio
    
    # Guardar documento
    fecha_str_file = fecha_analisis.strftime("%Y%m%d")
    filename = f"Reporte_Incidentes_{fecha_str_file}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    
    return filepath


def convertir_word_a_pdf(word_path: Path) -> Optional[Path]:
    """
    Convierte un archivo Word (.docx) a PDF temporalmente.
    Retorna la ruta del PDF temporal o None si falla.
    El PDF debe ser eliminado después de usarlo.
    """
    try:
        # Intentar usar docx2pdf (requiere Microsoft Word instalado)
        try:
            import docx2pdf
            pdf_path = word_path.with_suffix('.pdf')
            docx2pdf.convert(str(word_path), str(pdf_path))
            if pdf_path.exists():
                return pdf_path
        except ImportError:
            pass
        except Exception as e:
            print(f"[DEBUG] docx2pdf falló: {e}")
        
        # Intentar usar win32com (Windows COM automation)
        try:
            import win32com.client
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            doc = word_app.Documents.Open(str(word_path.absolute()))
            pdf_path = word_path.with_suffix('.pdf')
            doc.SaveAs(str(pdf_path.absolute()), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word_app.Quit()
            
            if pdf_path.exists():
                return pdf_path
        except ImportError:
            pass
        except Exception as e:
            print(f"[DEBUG] win32com falló: {e}")
        
        # Intentar usar comtypes (alternativa a win32com)
        try:
            import comtypes.client
            word_app = comtypes.client.CreateObject('Word.Application')
            word_app.Visible = False
            
            doc = word_app.Documents.Open(str(word_path.absolute()))
            pdf_path = word_path.with_suffix('.pdf')
            doc.SaveAs(str(pdf_path.absolute()), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word_app.Quit()
            
            if pdf_path.exists():
                return pdf_path
        except ImportError:
            pass
        except Exception as e:
            print(f"[DEBUG] comtypes falló: {e}")
        
        # Si ninguna librería está disponible, retornar None
        print("[ADVERTENCIA] No se encontró ninguna librería para convertir Word a PDF.")
        print("[INFO] Instala una de estas opciones:")
        print("  - pip install docx2pdf (requiere Microsoft Word)")
        print("  - pip install pywin32 (para win32com)")
        print("  - pip install comtypes (alternativa)")
        
        return None
        
    except Exception as e:
        print(f"[ERROR] Error al convertir Word a PDF: {e}")
        return None


def enviar_reporte_por_correo(
    reporte_path: Path,
    destinatarios: List[str],
    smtp_servidor: str = "smtp.gmail.com",
    smtp_puerto: int = 587,
    smtp_usuario: Optional[str] = None,
    smtp_password: Optional[str] = None,
    fecha_analisis: Optional[datetime] = None,
) -> bool:
    """
    Envía el reporte de incidentes por correo electrónico a múltiples destinatarios.
    
    Args:
        reporte_path: Ruta al archivo del reporte Word
        destinatarios: Lista de correos electrónicos de los destinatarios
        smtp_servidor: Servidor SMTP (default: smtp.gmail.com)
        smtp_puerto: Puerto SMTP (default: 587)
        smtp_usuario: Usuario SMTP (correo del remitente)
        smtp_password: Contraseña SMTP o contraseña de aplicación
        fecha_analisis: Fecha analizada (opcional, para el asunto)
    
    Returns:
        True si el correo se envió exitosamente, False en caso contrario
    """
    if not smtp_usuario or not smtp_password:
        print("[ERROR] Se requiere --smtp-usuario y --smtp-password para enviar correo")
        return False
    
    if not reporte_path.exists():
        print(f"[ERROR] El archivo del reporte no existe: {reporte_path}")
        return False
    
    try:
        # Crear mensaje
        msg = MIMEMultipart()
        msg["From"] = smtp_usuario
        msg["To"] = ", ".join(destinatarios)  # Múltiples destinatarios
        
        # Crear asunto
        asunto = "Reporte de Incidentes - Consumo Día Anterior"
        if fecha_analisis:
            fecha_str = fecha_analisis.strftime("%d-%m-%Y")
            asunto += f" ({fecha_str})"
        
        msg["Subject"] = asunto
        
        # Crear cuerpo del mensaje
        cuerpo = f"""
Estimados/as,

Se adjunta el reporte de incidentes del día anterior con el análisis de consumo de todos los puntos del sistema.

"""
        if fecha_analisis:
            cuerpo += f"Fecha analizada: {fecha_analisis.strftime('%d-%m-%Y')}\n\n"
        
        cuerpo += """El reporte incluye:
- Puntos con consumo cero
- Puntos sin respuesta (error de API)
- Puntos con datos incompletos (menos de 24 horas de datos)

PROCESO DE GENERACIÓN DEL REPORTE:

Este reporte fue generado automáticamente por un agente de inteligencia artificial creado por José Luis Otarola. El proceso ejecutado fue el siguiente:

1. Obtención de nodos: Se consultaron todas las empresas del sistema mediante el endpoint de consulta de empresas para obtener la lista completa de puntos de monitoreo disponibles.

2. Consulta de datos horarios: Para cada nodo identificado, se consultaron los datos de consumo horario del día anterior utilizando el endpoint de medidas horarias en formato CSV, solicitando los datos para el día específico analizado.

3. Análisis y clasificación: Cada punto fue analizado y clasificado según su estado:
   - Consumo cero: Puntos que registraron consumo cero durante todo el día
   - Sin respuesta: Puntos que no respondieron a la consulta (error de API o sin conexión)
   - Datos incompletos: Puntos con menos de 24 horas de datos, identificando las horas faltantes
   - OK: Puntos con datos completos y consumo normal

4. Generación del reporte: Se generó un documento Word con tablas detalladas para cada categoría de incidente, incluyendo información del nodo, nombre del punto, empresa y detalles específicos (como horas faltantes).

5. Generación de log: Se creó un archivo de log completo con todos los nodos revisados y su estado para referencia y auditoría.

Este reporte se genera automáticamente para identificar posibles problemas en el sistema de monitoreo y facilitar la detección temprana de anomalías.

Saludos cordiales,
Agente IA - Sistema WES
Creado por José Luis Otarola
"""
        
        msg.attach(MIMEText(cuerpo, "plain", "utf-8"))
        
        # Convertir Word a PDF temporalmente para el envío
        pdf_path = None
        try:
            pdf_path = convertir_word_a_pdf(reporte_path)
            if pdf_path and pdf_path.exists():
                # Adjuntar PDF
                with open(pdf_path, "rb") as f:
                    adjunto = MIMEApplication(f.read(), _subtype="pdf")
                    adjunto.add_header(
                        "Content-Disposition",
                        "attachment",
                        filename=reporte_path.stem + ".pdf"
                    )
                    msg.attach(adjunto)
                print(f"[INFO] Reporte convertido a PDF temporalmente para envío")
            else:
                # Si falla la conversión, adjuntar Word original
                print(f"[ADVERTENCIA] No se pudo convertir a PDF, adjuntando Word original")
                with open(reporte_path, "rb") as f:
                    adjunto = MIMEApplication(f.read(), _subtype="docx")
                    adjunto.add_header(
                        "Content-Disposition",
                        "attachment",
                        filename=reporte_path.name
                    )
                    msg.attach(adjunto)
        except Exception as e:
            print(f"[ADVERTENCIA] Error al convertir a PDF: {e}. Adjuntando Word original.")
            # Si falla la conversión, adjuntar Word original
            with open(reporte_path, "rb") as f:
                adjunto = MIMEApplication(f.read(), _subtype="docx")
                adjunto.add_header(
                    "Content-Disposition",
                    "attachment",
                    filename=reporte_path.name
                )
                msg.attach(adjunto)
        
        # Enviar correo
        print(f"[INFO] Conectando al servidor SMTP {smtp_servidor}:{smtp_puerto}...")
        with smtplib.SMTP(smtp_servidor, smtp_puerto) as server:
            server.starttls()
            print(f"[INFO] Autenticando como {smtp_usuario}...")
            server.login(smtp_usuario, smtp_password)
            print(f"[INFO] Enviando correo a {len(destinatarios)} destinatario(s)...")
            server.send_message(msg)
        
        print(f"[OK] Correo enviado exitosamente a: {', '.join(destinatarios)}")
        
        # Eliminar el PDF temporal si existe
        if pdf_path and pdf_path.exists():
            try:
                pdf_path.unlink()
                print(f"[INFO] Archivo PDF temporal eliminado")
            except Exception as e:
                print(f"[ADVERTENCIA] No se pudo eliminar el PDF temporal: {e}")
        
        return True
        
    except smtplib.SMTPAuthenticationError:
        print("[ERROR] Error de autenticación SMTP. Verifica usuario y contraseña.")
        return False
    except smtplib.SMTPException as e:
        print(f"[ERROR] Error SMTP: {e}")
        return False
    except Exception as e:
        print(f"[ERROR] Error al enviar correo: {e}")
        return False


def main():
    """Función principal."""
    parser = argparse.ArgumentParser(
        description="Generar reporte de incidentes del día anterior y enviar por correo",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  # Generar reporte sin enviar correo
  python reporte_incidentes_dia_anterior.py
  
  # Generar reporte y enviar por correo
  python reporte_incidentes_dia_anterior.py --enviar-correo --smtp-usuario tu@email.com --smtp-password tu_password
        """
    )
    
    parser.add_argument(
        "--enviar-correo",
        action="store_true",
        help="Enviar el reporte por correo electrónico"
    )
    
    parser.add_argument(
        "--smtp-servidor",
        default="smtp.gmail.com",
        help="Servidor SMTP (default: smtp.gmail.com)"
    )
    
    parser.add_argument(
        "--smtp-puerto",
        type=int,
        default=587,
        help="Puerto SMTP (default: 587)"
    )
    
    parser.add_argument(
        "--smtp-usuario",
        help="Usuario SMTP (correo del remitente)"
    )
    
    parser.add_argument(
        "--smtp-password",
        help="Contraseña SMTP (o contraseña de aplicación)"
    )
    
    args = parser.parse_args()
    
    # Destinatarios por defecto
    destinatarios_default = [
        "diegocarrasco@wes.cl",
        "anibal.aoperaciones@wes.cl",
        "benjamingumucio@wes.cl",
        "juanlopez@wes.cl",
        "agente.ia@wes.cl"
    ]
    
    # Contraseña de aplicación SMTP (desde scripts de reportes)
    SMTP_PASSWORD = "gsptzgakauqasnfx"
    
    print("=" * 60)
    print("REPORTE DE INCIDENTES - CONSUMO DÍA ANTERIOR")
    print("=" * 60)
    print()
    
    # Fecha del día anterior
    fecha_analisis = datetime.now(timezone.utc) - timedelta(days=1)
    fecha_analisis = fecha_analisis.replace(hour=0, minute=0, second=0, microsecond=0)
    
    print(f"Analizando datos del día: {fecha_analisis.strftime('%d-%m-%Y')}")
    print()
    
    # Obtener todos los nodos
    todos_nodos = obtener_todos_los_nodos()
    
    if not todos_nodos:
        print("[ERROR] No se encontraron nodos en el sistema")
        return
    
    print()
    print("Analizando consumo de cada punto...")
    print("=" * 60)
    
    # Clasificar puntos
    puntos_cero = []
    puntos_sin_respuesta = []
    puntos_incompletos = []
    puntos_ok = []
    
    # Lista para el log completo
    log_completo = []
    
    total = len(todos_nodos)
    for idx, nodo in enumerate(todos_nodos, 1):
        node_id = nodo["nodeId"]
        node_name = nodo["nodeName"]
        company_name = nodo["companyName"]
        
        print(f"[{idx}/{total}] Analizando {node_id} ({node_name})...", end=" ")
        
        # Obtener datos horarios
        hourly_data, estado_api = obtener_datos_horarios_dia(node_id, fecha_analisis)
        
        estado_final = ""
        horas_faltantes_str = ""
        
        if estado_api == "sin_respuesta" or estado_api == "error":
            puntos_sin_respuesta.append(nodo)
            estado_final = "SIN RESPUESTA"
            print("✗ Sin respuesta")
        elif hourly_data is None:
            puntos_sin_respuesta.append(nodo)
            estado_final = "SIN DATOS"
            print("✗ Sin datos")
        else:
            estado_datos, horas_faltantes = analizar_datos_horarios(hourly_data)
            
            if estado_datos == "consumo_cero":
                puntos_cero.append(nodo)
                estado_final = "CONSUMO CERO"
                print("⚠ Consumo cero")
            elif estado_datos == "datos_incompletos":
                nodo["horasFaltantes"] = horas_faltantes
                puntos_incompletos.append(nodo)
                estado_final = "DATOS INCOMPLETOS"
                horas_faltantes_str = ", ".join([f"{h:02d}:00" for h in horas_faltantes])
                horas_str = ", ".join([f"{h:02d}:00" for h in horas_faltantes[:5]])
                if len(horas_faltantes) > 5:
                    horas_str += f" ... ({len(horas_faltantes)} horas faltantes)"
                print(f"⚠ Datos incompletos ({len(horas_faltantes)} horas faltantes)")
            else:
                puntos_ok.append(nodo)
                estado_final = "OK"
                print("✓ OK")
        
        # Agregar al log completo
        log_completo.append({
            "indice": idx,
            "nodeId": node_id,
            "nombre": node_name,
            "empresa": company_name,
            "estado": estado_final,
            "horas_faltantes": horas_faltantes_str
        })
    
    print()
    print("=" * 60)
    print("RESUMEN DEL ANÁLISIS")
    print("=" * 60)
    print(f"Total puntos analizados: {total}")
    print(f"Puntos OK: {len(puntos_ok)}")
    print(f"Puntos con consumo cero: {len(puntos_cero)}")
    print(f"Puntos sin respuesta: {len(puntos_sin_respuesta)}")
    print(f"Puntos con datos incompletos: {len(puntos_incompletos)}")
    print()
    
    # Crear carpeta de incidentes
    base_dir = Path.home() / "Desktop" / "wes-scripts" / "reports" / "incidentes"
    fecha_creacion = datetime.now(timezone.utc)
    folder_name = f"REPORTE_DIA_ANTERIOR_{fecha_creacion.strftime('%Y%m%d_%H%M')}"
    output_dir = base_dir / folder_name
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generar archivo de log completo
    print(f"Generando archivo de log...")
    log_path = output_dir / f"LOG_NODOS_REVISADOS_{fecha_creacion.strftime('%Y%m%d_%H%M')}.txt"
    
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("LOG COMPLETO DE NODOS REVISADOS\n")
        f.write("=" * 80 + "\n")
        f.write(f"Fecha analizada: {fecha_analisis.strftime('%d-%m-%Y')}\n")
        f.write(f"Fecha de generación: {fecha_creacion.strftime('%d-%m-%Y %H:%M:%S')}\n")
        f.write(f"Total de nodos analizados: {total}\n")
        f.write("=" * 80 + "\n\n")
        
        f.write(f"{'#':<5} {'Nodo ID':<15} {'Estado':<20} {'Nombre':<40} {'Empresa':<30}\n")
        f.write("-" * 80 + "\n")
        
        for item in log_completo:
            estado = item["estado"]
            nombre = item["nombre"][:38] if len(item["nombre"]) > 38 else item["nombre"]
            empresa = item["empresa"][:28] if len(item["empresa"]) > 28 else item["empresa"]
            
            f.write(f"{item['indice']:<5} {item['nodeId']:<15} {estado:<20} {nombre:<40} {empresa:<30}\n")
            
            # Si tiene horas faltantes, agregarlas en la siguiente línea
            if item["horas_faltantes"]:
                f.write(f"{'':5} {'':15} {'':20} Horas faltantes: {item['horas_faltantes']}\n")
        
        f.write("\n" + "=" * 80 + "\n")
        f.write("RESUMEN POR CATEGORÍA\n")
        f.write("=" * 80 + "\n")
        f.write(f"Puntos OK: {len(puntos_ok)}\n")
        f.write(f"Puntos con consumo cero: {len(puntos_cero)}\n")
        f.write(f"Puntos sin respuesta: {len(puntos_sin_respuesta)}\n")
        f.write(f"Puntos con datos incompletos: {len(puntos_incompletos)}\n")
        f.write("=" * 80 + "\n\n")
        
        # Detalle por categoría
        if puntos_cero:
            f.write("\nPUNTOS CON CONSUMO CERO:\n")
            f.write("-" * 80 + "\n")
            for punto in puntos_cero:
                f.write(f"  {punto['nodeId']:<15} {punto['nodeName']:<40} {punto['companyName']}\n")
        
        if puntos_sin_respuesta:
            f.write("\nPUNTOS SIN RESPUESTA:\n")
            f.write("-" * 80 + "\n")
            for punto in puntos_sin_respuesta:
                f.write(f"  {punto['nodeId']:<15} {punto['nodeName']:<40} {punto['companyName']}\n")
        
        if puntos_incompletos:
            f.write("\nPUNTOS CON DATOS INCOMPLETOS:\n")
            f.write("-" * 80 + "\n")
            for punto in puntos_incompletos:
                horas_str = ", ".join([f"{h:02d}:00" for h in punto.get("horasFaltantes", [])])
                f.write(f"  {punto['nodeId']:<15} {punto['nodeName']:<40} {punto['companyName']}\n")
                f.write(f"  {'':15} Horas faltantes: {horas_str}\n")
    
    print(f"✓ Log guardado en: {log_path}")
    
    print(f"Generando reporte Word...")
    
    # Crear reporte Word
    reporte_path = crear_reporte_word(
        fecha_analisis,
        puntos_cero,
        puntos_sin_respuesta,
        puntos_incompletos,
        output_dir
    )
    
    print(f"✓ Reporte guardado en: {reporte_path}")
    print()
    
    # Enviar correo si se solicita
    if args.enviar_correo:
        print("=" * 60)
        print("ENVIANDO CORREO")
        print("=" * 60)
        
        # Usar contraseña de aplicación por defecto si no se proporciona
        smtp_password = args.smtp_password if args.smtp_password else SMTP_PASSWORD
        smtp_usuario = args.smtp_usuario if args.smtp_usuario else "agente.ia@wes.cl"
        
        enviado = enviar_reporte_por_correo(
            reporte_path=reporte_path,
            destinatarios=destinatarios_default,
            smtp_servidor=args.smtp_servidor,
            smtp_puerto=args.smtp_puerto,
            smtp_usuario=smtp_usuario,
            smtp_password=smtp_password,
            fecha_analisis=fecha_analisis
        )
        
        if enviado:
            print(f"[OK] Correo enviado exitosamente a {len(destinatarios_default)} destinatario(s)")
        else:
            print(f"[ERROR] No se pudo enviar el correo")
    
    print()
    print("=" * 60)
    print("PROCESO COMPLETADO")
    print("=" * 60)


if __name__ == "__main__":
    main()

