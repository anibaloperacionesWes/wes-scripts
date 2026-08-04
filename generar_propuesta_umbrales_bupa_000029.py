# -*- coding: utf-8 -*-
"""
Estudio baseline + umbral +25% — solo Clínica Bupa Antofagasta
(nodos 000029-07, 000029-08, 000029-09, 000029-10).

Regla estándar WES (como Parque Arauco):
  umbral = promedio_diario × 1,25  (ventana objetivo 90 días, dejando ceros fuera)

Para esta clínica se evalúa desde el 22/06/2026 (sin esperar los 90 días completos),
para poder activar ya la alerta, que es importante para la operación.

No incluye puntos de matriz (000029-01..06).

Genera Excel/Word técnicos + borrador de correo al cliente para aprobar activación.
"""

from __future__ import annotations

import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime
from pathlib import Path
from statistics import median

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from generar_reporte_word import format_number_chilean as fn
from generar_reportes_y_ppt_mall_maipu import obtener_datos_agregados

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Bupa_Antofagasta" / "umbrales_consumo"
ENTITY_BASE = "http://104.248.53.141:7001/wes/api/acl-entities/v1"

COMPANY_ID = "000029"
COMPANY_LABEL = "Clínica Bupa Antofagasta"
REF = date(2026, 8, 4)
# Regla estándar WES: ventana objetivo 90 días. En esta clínica se evalúa
# desde el 22/06/2026 para poder mostrar/activar ya la alerta.
DIAS_BASELINE_OBJETIVO = 90
FECHA_INICIO_EVAL = date(2026, 6, 22)
MULT = 1.25
DESDE = FECHA_INICIO_EVAL.strftime("%d/%m/%Y")
HASTA = REF.strftime("%d/%m/%Y")
DIAS_VENTANA_EVAL = (REF - FECHA_INICIO_EVAL).days + 1

# Solo Antofagasta — no incluir matriz ni otros puntos Bupa
NODOS_ANTOFAGASTA = (
    "000029-07",
    "000029-08",
    "000029-09",
    "000029-10",
)

NOMBRES_FALLBACK = {
    "000029-07": "Sala de Bomba Principal",
    "000029-08": "Sala de Bomba Sexto Piso",
    "000029-09": "Medidor Principal Sanitaria",
    "000029-10": "Sala de Bomba N°2",
}


def listar_nodos_antofagasta() -> list[dict]:
    import requests

    url = f"{ENTITY_BASE}/companies/{COMPANY_ID}"
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    data = r.json()
    company_name = data.get("name", "BUPA")
    by_id = {
        (n.get("nodeId") or "").strip(): (n.get("name") or "").strip()
        for n in (data.get("nodes") or [])
    }
    rows = []
    for nid in NODOS_ANTOFAGASTA:
        nm = by_id.get(nid) or NOMBRES_FALLBACK.get(nid) or nid
        rows.append(
            {
                "nodeId": nid,
                "nodeName": nm,
                "companyName": company_name,
                "sede": COMPANY_LABEL,
            }
        )
    return rows


def _serie_diaria(node_id: str) -> dict[date, float]:
    d = obtener_datos_agregados([node_id], DESDE, HASTA)
    ns = (d.get("nodes_summary") or [None])[0]
    if not ns:
        return {}
    daily: dict[date, float] = {}
    for m in ns.get("measures") or []:
        d0 = m.date.date()
        daily[d0] = daily.get(d0, 0.0) + float(m.total_m3)
    return daily


def _baseline_operativo(daily: dict[date, float]) -> tuple[float | None, int, float | None, int, str]:
    """
    Retorna (promedio, dias_usados, max_diario, dias_brutos, nota).

    - Deja ceros fuera: parte desde el primer día con consumo > 0,05 m³.
    - Excluye outliers claros (p. ej. lecturas absurdas de calibración en Sala Bomba N°2).
    - Ventana de evaluación de esta clínica: desde FECHA_INICIO_EVAL (22/06/2026).
    """
    if not daily:
        return None, 0, None, 0, "Sin datos en la ventana desde 22/06/2026"

    ordered = sorted(daily.items())
    dias_brutos = len(ordered)
    first_nz = next((d for d, v in ordered if v > 0.05), None)
    if first_nz is None:
        return None, 0, None, dias_brutos, "Solo ceros / sin consumo medible (ceros dejados fuera)"

    series = [(d, v) for d, v in ordered if d >= first_nz]
    raw = [v for _, v in series]
    med = median(raw) if raw else 0.0
    lim = max(med * 5.0, 50.0)
    if med > 80:
        lim = max(med * 3.0, 400.0)

    clean = [(d, v) for d, v in series if v <= lim]
    n_out = len(series) - len(clean)
    if not clean:
        return None, 0, None, dias_brutos, "Todos los días filtrados como outlier"

    vals = [v for _, v in clean]
    mean = sum(vals) / len(vals)
    nota = (
        f"Evaluación desde {DESDE} (no se espera ventana completa de "
        f"{DIAS_BASELINE_OBJETIVO} días). Ceros dejados fuera; data operativa "
        f"desde {first_nz.strftime('%d/%m/%Y')}: {len(vals)} día(s)"
    )
    if n_out:
        nota += f"; se excluyeron {n_out} día(s) outlier de puesta en marcha"
    return mean, len(vals), max(vals), dias_brutos, nota


def calcular_filas(nodos: list[dict]) -> list[dict]:
    out: list[dict] = []

    def _one(n: dict) -> dict:
        nid = n["nodeId"]
        print(f"[data] {nid} {n['nodeName']}...", flush=True)
        try:
            daily = _serie_diaria(nid)
            mean, dias, mx, dias_brutos, nota = _baseline_operativo(daily)
        except Exception as e:
            print(f"[WARN] {nid}: {e}")
            mean, dias, mx, dias_brutos, nota = None, 0, None, 0, f"Error: {e}"
        umbral = round(mean * MULT, 1) if mean is not None else None
        return {
            "nodeId": nid,
            "nodeName": n["nodeName"],
            "companyName": n["companyName"],
            "sede": n["sede"],
            "dias": dias,
            "dias_brutos": dias_brutos,
            "baseline": round(mean, 1) if mean is not None else None,
            "max_periodo": round(mx, 1) if mx is not None else None,
            "umbral": umbral,
            "mult": MULT,
            "nota": nota,
        }

    with ThreadPoolExecutor(max_workers=6) as ex:
        futs = {ex.submit(_one, n): n for n in nodos}
        for fut in as_completed(futs):
            out.append(fut.result())
    out.sort(key=lambda r: r["nodeId"])
    return out


def escribir_excel(rows: list[dict], path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Umbrales 000029"
    headers = [
        "Node ID",
        "Sede",
        "Punto",
        "Días operativos usados",
        "Días brutos en ventana",
        "Baseline promedio diario (m³)",
        "Máximo diario usado (m³)",
        "Multiplicador",
        "Umbral recomendado (m³/día)",
        "Nota",
    ]
    ws.append(headers)
    header_fill = PatternFill("solid", fgColor="1F4E79")
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = header_fill
        c.alignment = Alignment(wrap_text=True, vertical="center")

    for r in rows:
        ws.append(
            [
                r["nodeId"],
                r["sede"],
                r["nodeName"],
                r["dias"],
                r["dias_brutos"],
                r["baseline"] if r["baseline"] is not None else "Sin datos",
                r["max_periodo"] if r["max_periodo"] is not None else "Sin datos",
                f"× {MULT:.2f} (+{(MULT - 1) * 100:.0f}%)",
                r["umbral"] if r["umbral"] is not None else "N/D",
                r["nota"],
            ]
        )

    ws.append([])
    ws.append(["Parámetros"])
    ws.append(["Compañía", COMPANY_ID])
    ws.append(["Cliente / sede", COMPANY_LABEL])
    ws.append(["Nodos incluidos", ", ".join(NODOS_ANTOFAGASTA)])
    ws.append(
        [
            "Regla estándar WES",
            f"umbral = promedio_diario × 1,25 (ventana objetivo {DIAS_BASELINE_OBJETIVO} días, dejando ceros fuera)",
        ]
    )
    ws.append(
        [
            "Evaluación en esta clínica",
            f"Desde {DESDE} hasta {HASTA} ({DIAS_VENTANA_EVAL} días de calendario) — "
            "no se espera completar los 90 días para poder activar ya la alerta",
        ]
    )
    ws.append(["Corte", REF.strftime("%d/%m/%Y")])
    ws.append(["Ceros", "Dejados fuera del baseline (solo días con consumo operativo)"])
    ws.append(
        [
            "Comportamiento API",
            "Cuando el acumulado del día supera el umbral (a cualquier hora) → alerta por correo",
        ]
    )

    for i, w in enumerate([14, 28, 32, 14, 14, 22, 18, 14, 20, 55], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def escribir_word(rows: list[dict], path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(
        f"Propuesta de umbrales de consumo diario — {COMPANY_LABEL}",
        0,
    )
    p = doc.add_paragraph()
    p.add_run(f"Fecha de corte: {REF.strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(
        f"Puntos incluidos: {', '.join(NODOS_ANTOFAGASTA)}\n"
        f"Regla estándar WES: {DIAS_BASELINE_OBJETIVO} días × 1,25, dejando ceros fuera.\n"
        f"Evaluación en esta clínica: desde el {DESDE} hasta el {HASTA} "
        f"({DIAS_VENTANA_EVAL} días de calendario) — sin esperar los 90 días "
        "completos, para poder activar ya esta alerta, que es importante para la clínica."
    )

    doc.add_heading("1. Objetivo", 1)
    doc.add_paragraph(
        f"Definir umbrales diarios (m³/día) para los 4 puntos de monitoreo de "
        f"{COMPANY_LABEL} ({', '.join(NODOS_ANTOFAGASTA)}), compatibles con la alerta "
        "por umbral de la API WES: cuando el consumo acumulado del día supera ese "
        "número —independiente de la hora— se dispara aviso por correo."
    )

    doc.add_heading("2. Alcance", 1)
    doc.add_paragraph(
        "Solo se incluyen los nodos de Clínica Bupa Antofagasta:\n"
        "• 000029-07 — Sala de Bomba Principal\n"
        "• 000029-08 — Sala de Bomba Sexto Piso\n"
        "• 000029-09 — Medidor Principal Sanitaria\n"
        "• 000029-10 — Sala de Bomba N°2\n\n"
        "No se incluyen otros puntos de la compañía 000029 (matriz u otras sedes)."
    )

    doc.add_heading("3. Criterio técnico (qué se declara y qué se evalúa)", 1)
    for b in [
        f"Regla estándar WES (igual que Parque Arauco): umbral = promedio diario × 1,25, con ventana objetivo de {DIAS_BASELINE_OBJETIVO} días y dejando ceros fuera del cálculo.",
        f"Evaluación de esta clínica: se calcula el baseline solo desde el {DESDE} (no desde 90 días atrás). Motivo: poder mostrar/activar ya la alerta, que es importante para la operación de la clínica, sin esperar a completar la ventana de 90 días.",
        "Ceros dejados fuera: el promedio usa solo días con consumo operativo (se omiten ceros previos a la puesta en marcha).",
        "También se excluyen outliers claros de calibración (caso Sala de Bomba N°2).",
        "Alerta API: un número por punto; se evalúa sobre el acumulado del día al cruzar el umbral.",
        f"Cuando se complete una ventana de {DIAS_BASELINE_OBJETIVO} días, se recalculará el baseline manteniendo la misma fórmula × 1,25.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    con_umbral = [r for r in rows if r["umbral"] is not None]
    doc.add_heading("4. Tabla recomendada", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Node ID"
    hdr[1].text = "Punto"
    hdr[2].text = "Días usados"
    hdr[3].text = "Baseline (m³/d)"
    hdr[4].text = "Umbral +25% (m³/d)"
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["nodeId"]
        cells[1].text = r["nodeName"]
        cells[2].text = str(r["dias"])
        cells[3].text = fn(r["baseline"], 1) if r["baseline"] is not None else "Sin datos"
        cells[4].text = fn(r["umbral"], 1) if r["umbral"] is not None else "N/D"

    if con_umbral:
        ej = max(con_umbral, key=lambda r: r["baseline"] or 0)
        doc.add_paragraph(
            f"Ejemplo ({ej['nodeName']} / {ej['nodeId']}): "
            f"baseline ≈ {fn(ej['baseline'], 1)} m³/día → "
            f"umbral ≈ {fn(ej['umbral'], 1)} m³/día."
        )

    doc.add_heading("5. Notas por punto", 1)
    for r in rows:
        doc.add_paragraph(
            f"{r['nodeId']} — {r['nodeName']}: {r['nota']}",
            style="List Bullet",
        )

    doc.add_heading("6. Pedido al cliente", 1)
    doc.add_paragraph(
        "1) Validar la regla estándar (90 días × 1,25, dejando ceros fuera) y la "
        f"evaluación desde el {DESDE} para activar ya la alerta en Clínica Bupa Antofagasta.\n"
        "2) Aprobar la carga de los umbrales en la API de alertas WES.\n"
        "3) Confirmar destinatarios de las alertas.\n"
        "4) Indicar exclusiones puntuales si algún medidor aún no está estabilizado.\n"
        "5) Recalcular cuando se complete una ventana de 90 días (revisión trimestral)."
    )

    foot = doc.add_paragraph()
    foot.add_run(
        "Documento generado por Sistema WES — propuesta de umbrales Clínica Bupa Antofagasta."
    ).italic = True

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _tabla_texto(rows: list[dict]) -> str:
    lines = [
        f"{'ID':<12} {'Punto':<32} {'Días':>5} {'Baseline':>10} {'Umbral':>10}",
        "-" * 74,
    ]
    for r in rows:
        b = fn(r["baseline"], 1) if r["baseline"] is not None else "N/D"
        u = fn(r["umbral"], 1) if r["umbral"] is not None else "N/D"
        lines.append(
            f"{r['nodeId']:<12} {r['nodeName'][:31]:<32} {r['dias']:>5} {b:>10} {u:>10}"
        )
    return "\n".join(lines)


def _tabla_html(rows: list[dict]) -> str:
    trs = []
    for r in rows:
        trs.append(
            "<tr>"
            f"<td>{r['nodeId']}</td>"
            f"<td>{r['nodeName']}</td>"
            f"<td style='text-align:right'>{r['dias']}</td>"
            f"<td style='text-align:right'>{fn(r['baseline'], 1)}</td>"
            f"<td style='text-align:right'>{fn(r['umbral'], 1)}</td>"
            "</tr>"
        )
    return (
        "<table border='1' cellpadding='6' cellspacing='0' "
        "style='border-collapse:collapse;font-family:Calibri,Arial,sans-serif;font-size:13px'>"
        "<thead><tr style='background:#1F4E79;color:#fff'>"
        "<th>Node ID</th><th>Punto</th><th>Días usados</th>"
        "<th>Baseline (m³/día)</th><th>Umbral +25% (m³/día)</th>"
        "</tr></thead><tbody>"
        + "".join(trs)
        + "</tbody></table>"
    )


def escribir_correo_cliente(rows: list[dict], path_txt: Path, path_html: Path) -> None:
    propuesta = [r for r in rows if r["umbral"] is not None]
    ejemplo = max(propuesta, key=lambda r: r["baseline"] or 0) if propuesta else None
    ejemplo_txt = ""
    if ejemplo:
        ejemplo_txt = (
            f"Por ejemplo, en {ejemplo['nodeName']} ({ejemplo['nodeId']}) el consumo "
            f"promedio operativo es de aproximadamente {fn(ejemplo['baseline'], 1)} m³/día; "
            f"el umbral propuesto sería {fn(ejemplo['umbral'], 1)} m³/día (+25%).\n\n"
        )

    asunto = (
        "WES — Propuesta de umbrales de alerta de consumo de agua — "
        "Clínica Bupa Antofagasta | solicitud de aprobación"
    )

    cuerpo = f"""Estimados,

Junto con saludar, les escribimos desde WES para presentar una propuesta de umbrales de alerta de consumo diario de agua en los puntos de monitoreo de Clínica Bupa Antofagasta (nodos 000029-07, 000029-08, 000029-09 y 000029-10), y solicitar su aprobación para activarlos en la plataforma.

¿Qué queremos hacer?
Activar, en cada uno de estos 4 puntos, un umbral máximo diario (m³/día). Cuando el consumo acumulado del día supere ese valor —a cualquier hora—, el sistema enviará automáticamente un correo de alerta al equipo que ustedes indiquen. Así pueden reaccionar el mismo día, sin esperar el cierre del período ni la boleta. Esta alerta es importante para la clínica y por eso proponemos activarla ahora.

¿Cómo se calcularon los umbrales?
La regla estándar de WES (la misma que usamos en otros clientes) es:
• Ventana objetivo: 90 días.
• Baseline = promedio del consumo diario, dejando ceros fuera.
• Umbral = baseline × 1,25 (+25% sobre el promedio).

Para Clínica Bupa Antofagasta evaluamos desde el {DESDE} hasta el {HASTA} (sin esperar a completar los 90 días), dejando ceros fuera del cálculo. El motivo es poder mostrar y activar ya esta alerta, que es importante para la operación de la clínica. Cuando se complete una ventana de 90 días, recalcularemos el baseline con la misma fórmula.

{ejemplo_txt}Tabla propuesta — Clínica Bupa Antofagasta (valores en m³/día):

{_tabla_texto(rows)}

¿Qué pedimos de su parte?
1) Revisar la tabla y confirmar si están de acuerdo con la regla (90 días × 1,25, dejando ceros fuera) y con evaluar desde el {DESDE} para activar ya la alerta.
2) Aprobar la aplicación (carga) de estos umbrales en la configuración de alertas WES.
3) Indicar a qué correos deben llegar las alertas (si difieren de los contactos actuales).
4) Señalar si algún punto debe excluirse por estar aún en estabilización o fuera de servicio.

Quedamos atentos a su OK para proceder con la activación.

Saludos cordiales,
Equipo WES — Monitoreo y Control
"""

    path_txt.parent.mkdir(parents=True, exist_ok=True)
    path_txt.write_text(f"ASUNTO:\n{asunto}\n\nCUERPO:\n{cuerpo}\n", encoding="utf-8")

    html = f"""<!DOCTYPE html>
<html lang="es">
<head><meta charset="utf-8"><title>{asunto}</title></head>
<body style="font-family:Calibri,Arial,sans-serif;font-size:14px;color:#222;line-height:1.45">
<p><strong>Asunto:</strong> {asunto}</p>
<hr>
<p>Estimados,</p>
<p>Junto con saludar, les escribimos desde WES para presentar una propuesta de
<strong>umbrales de alerta de consumo diario de agua</strong> en los puntos de
monitoreo de <strong>Clínica Bupa Antofagasta</strong>
(nodos 000029-07, 000029-08, 000029-09 y 000029-10), y solicitar su
<strong>aprobación para activarlos</strong> en la plataforma.</p>
<p><strong>¿Qué queremos hacer?</strong><br>
Activar, en cada uno de estos 4 puntos, un umbral máximo diario (m³/día).
Cuando el consumo acumulado del día supere ese valor —a cualquier hora—, el sistema
enviará automáticamente un correo de alerta al equipo que ustedes indiquen. Así
pueden reaccionar el mismo día. <strong>Esta alerta es importante para la clínica</strong>
y por eso proponemos activarla ahora.</p>
<p><strong>¿Cómo se calcularon los umbrales?</strong><br>
La <strong>regla estándar de WES</strong> (igual que en otros clientes) es:</p>
<ul>
<li>Ventana objetivo: <strong>90 días</strong>.</li>
<li>Baseline = promedio del consumo diario, <strong>dejando ceros fuera</strong>.</li>
<li>Umbral = baseline × 1,25 (+25% sobre el promedio).</li>
</ul>
<p>Para Clínica Bupa Antofagasta <strong>evaluamos desde el {DESDE} hasta el {HASTA}</strong>
(sin esperar a completar los 90 días), dejando ceros fuera del cálculo. El motivo es
<strong>poder mostrar y activar ya esta alerta</strong>, que es importante para la
operación de la clínica. Cuando se complete una ventana de 90 días, recalcularemos
el baseline con la misma fórmula.</p>
{f"<p>Por ejemplo, en <strong>{ejemplo['nodeName']}</strong> ({ejemplo['nodeId']}) el consumo promedio operativo es de aproximadamente <strong>{fn(ejemplo['baseline'], 1)} m³/día</strong>; el umbral propuesto sería <strong>{fn(ejemplo['umbral'], 1)} m³/día</strong> (+25%).</p>" if ejemplo else ""}
<p><strong>Tabla propuesta — Clínica Bupa Antofagasta (m³/día):</strong></p>
{_tabla_html(rows)}
<p><strong>¿Qué pedimos de su parte?</strong></p>
<ol>
<li>Revisar la tabla y confirmar si están de acuerdo con la regla (90 días × 1,25, dejando ceros fuera) y con evaluar desde el {DESDE} para activar ya la alerta.</li>
<li>Aprobar la aplicación (carga) de estos umbrales en la configuración de alertas WES.</li>
<li>Indicar a qué correos deben llegar las alertas (si difieren de los contactos actuales).</li>
<li>Señalar si algún punto debe excluirse por estar aún en estabilización o fuera de servicio.</li>
</ol>
<p>Quedamos atentos a su OK para proceder con la activación.</p>
<p>Saludos cordiales,<br>
<strong>Equipo WES — Monitoreo y Control</strong></p>
</body>
</html>
"""
    path_html.write_text(html, encoding="utf-8")


def escribir_word_correo_cliente(rows: list[dict], path: Path) -> None:
    propuesta = [r for r in rows if r["umbral"] is not None]
    ejemplo = max(propuesta, key=lambda r: r["baseline"] or 0) if propuesta else None

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(
        "Borrador de correo al cliente — aprobación umbrales Clínica Bupa Antofagasta",
        0,
    )
    doc.add_paragraph(
        "Usar este texto para enviar al cliente (completar destinatarios). "
        "Adjuntos sugeridos: Excel y Word de la propuesta técnica. "
        "Alcance: solo nodos 000029-07, 000029-08, 000029-09 y 000029-10."
    )

    asunto = (
        "WES — Propuesta de umbrales de alerta de consumo de agua — "
        "Clínica Bupa Antofagasta | solicitud de aprobación"
    )
    p = doc.add_paragraph()
    p.add_run("Asunto: ").bold = True
    p.add_run(asunto)

    doc.add_heading("Cuerpo", 1)
    doc.add_paragraph("Estimados,")
    doc.add_paragraph(
        "Junto con saludar, les escribimos desde WES para presentar una propuesta de "
        "umbrales de alerta de consumo diario de agua en los puntos de monitoreo de "
        "Clínica Bupa Antofagasta (nodos 000029-07, 000029-08, 000029-09 y 000029-10), "
        "y solicitar su aprobación para activarlos en la plataforma."
    )
    doc.add_paragraph("¿Qué queremos hacer?").runs[0].bold = True
    doc.add_paragraph(
        "Activar, en cada uno de estos 4 puntos, un umbral máximo diario (m³/día). "
        "Cuando el consumo acumulado del día supere ese valor —a cualquier hora—, el "
        "sistema enviará automáticamente un correo de alerta al equipo que ustedes "
        "indiquen. Así pueden reaccionar el mismo día. Esta alerta es importante para "
        "la clínica y por eso proponemos activarla ahora."
    )
    doc.add_paragraph("¿Cómo se calcularon los umbrales?").runs[0].bold = True
    doc.add_paragraph(
        "La regla estándar de WES (la misma que usamos en otros clientes) es:\n"
        "• Ventana objetivo: 90 días.\n"
        "• Baseline = promedio del consumo diario, dejando ceros fuera.\n"
        "• Umbral = baseline × 1,25 (+25% sobre el promedio).\n\n"
        f"Para Clínica Bupa Antofagasta evaluamos desde el {DESDE} hasta el {HASTA} "
        "(sin esperar a completar los 90 días), dejando ceros fuera del cálculo. "
        "El motivo es poder mostrar y activar ya esta alerta, que es importante para "
        "la operación de la clínica. Cuando se complete una ventana de 90 días, "
        "recalcularemos el baseline con la misma fórmula."
    )
    if ejemplo:
        doc.add_paragraph(
            f"Por ejemplo, en {ejemplo['nodeName']} ({ejemplo['nodeId']}) el consumo "
            f"promedio operativo es de aproximadamente {fn(ejemplo['baseline'], 1)} m³/día; "
            f"el umbral propuesto sería {fn(ejemplo['umbral'], 1)} m³/día (+25%)."
        )

    doc.add_paragraph(
        "Tabla propuesta — Clínica Bupa Antofagasta (valores en m³/día):"
    ).runs[0].bold = True
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Node ID"
    hdr[1].text = "Punto"
    hdr[2].text = "Días usados"
    hdr[3].text = "Baseline"
    hdr[4].text = "Umbral +25%"
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["nodeId"]
        cells[1].text = r["nodeName"]
        cells[2].text = str(r["dias"])
        cells[3].text = fn(r["baseline"], 1) if r["baseline"] is not None else "Sin datos"
        cells[4].text = fn(r["umbral"], 1) if r["umbral"] is not None else "N/D"

    doc.add_paragraph("¿Qué pedimos de su parte?").runs[0].bold = True
    for item in [
        f"Revisar la tabla y confirmar si están de acuerdo con la regla (90 días × 1,25, dejando ceros fuera) y con evaluar desde el {DESDE} para activar ya la alerta.",
        "Aprobar la aplicación (carga) de estos umbrales en la configuración de alertas WES.",
        "Indicar a qué correos deben llegar las alertas (si difieren de los contactos actuales).",
        "Señalar si algún punto debe excluirse por estar aún en estabilización o fuera de servicio.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph(
        "Quedamos atentos a su OK para proceder con la activación.\n\n"
        "Saludos cordiales,\n"
        "Equipo WES — Monitoreo y Control"
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Listando nodos Antofagasta ({', '.join(NODOS_ANTOFAGASTA)})...")
    nodos = listar_nodos_antofagasta()
    print(f"Nodos: {len(nodos)}")
    assert all(n["nodeId"] in NODOS_ANTOFAGASTA for n in nodos)
    assert len(nodos) == 4
    rows = calcular_filas(nodos)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    xlsx = OUT_DIR / f"Umbrales_Antofagasta_000029_07_10_x125_{stamp}.xlsx"
    docx = OUT_DIR / f"Propuesta_umbrales_Antofagasta_000029_07_10_x125_{stamp}.docx"
    mail_txt = OUT_DIR / f"Correo_cliente_aprobacion_umbrales_Antofagasta_{stamp}.txt"
    mail_html = OUT_DIR / f"Correo_cliente_aprobacion_umbrales_Antofagasta_{stamp}.html"
    mail_docx = OUT_DIR / f"Correo_cliente_aprobacion_umbrales_Antofagasta_{stamp}.docx"

    escribir_excel(rows, xlsx)
    escribir_word(rows, docx)
    escribir_correo_cliente(rows, mail_txt, mail_html)
    escribir_word_correo_cliente(rows, mail_docx)

    print(f"[OK] Excel: {xlsx}")
    print(f"[OK] Word:  {docx}")
    print(f"[OK] Correo txt:  {mail_txt}")
    print(f"[OK] Correo html: {mail_html}")
    print(f"[OK] Correo docx: {mail_docx}")

    ok = sum(1 for r in rows if r["umbral"] is not None)
    print(f"Con umbral calculado: {ok}/{len(rows)} (solo Antofagasta)")
    for r in rows:
        print(
            f"  {r['nodeId']} | {r['nodeName'][:28]:<28} | "
            f"d={r['dias']:3d} | base={r['baseline']} | umbral={r['umbral']}"
        )

    (OUT_DIR / "_last_paths.txt").write_text(
        f"{xlsx}\n{docx}\n{mail_txt}\n{mail_html}\n{mail_docx}\n",
        encoding="utf-8",
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
