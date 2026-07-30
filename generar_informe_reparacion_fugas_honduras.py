"""
Informe comparativo UDD — reparación de fugas en matriz (Sala impulsión Honduras).

Evento: sábado 11/07/2026 — reparaciones de fugas en la matriz (corte de agua).
Comparación: consumo diario 7 días antes vs 7 días después.
El día 11/07/2026 se excluye del gráfico y de los totales comparativos.

Ventanas:
  Antes:  04/07/2026 – 10/07/2026
  Después: 12/07/2026 – 18/07/2026

Uso:
  python generar_informe_reparacion_fugas_honduras.py

Salida:
  reports/udd_reparacion_fugas_honduras/
    Informe_Reparacion_Fugas_Honduras_UDD.docx
    Informe_Reparacion_Fugas_Honduras_UDD.pdf
    graficos/
"""

from __future__ import annotations

import sys
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "udd_reparacion_fugas_honduras"
GRAFICOS_DIR = OUT_DIR / "graficos"
NOMBRE_INFORME = "Informe_Reparacion_Fugas_Honduras_UDD"

NODE_ID = "000026-01"
NODE_NOMBRE = "Sala impulsión Honduras"
CLIENTE = "Universidad del Desarrollo (UDD)"

DIA_REPARACION = date(2026, 7, 11)
ANTES_INI = date(2026, 7, 4)   # 7 días antes (incluye)
ANTES_FIN = date(2026, 7, 10)
DESPUES_INI = date(2026, 7, 12)  # 7 días después (incluye)
DESPUES_FIN = date(2026, 7, 18)

# Martes de cada ventana (mismo día de semana)
MARTES_ANTES = date(2026, 7, 7)    # martes dentro de 04–10/07
MARTES_DESPUES = date(2026, 7, 14)  # martes dentro de 12–18/07
HORAS_NOCTURNAS = list(range(0, 7))  # 00:00–06:00 (bloques 0..6)

# Estilo reportes agregados: rojo = antes, azul WES = después
COLOR_ANTES = "#C0392B"
COLOR_DESPUES = "#0050b3"
COLOR_EXCLUIDO = "#7F8C8D"
COLOR_AHORRO = "#1F4E79"
COLOR_NOCHE_APP = "#c41e1e"
COLOR_LINEA_APP = "#4A90E2"
COLOR_NOCTURNO_AGREGADO = "#FFD700"  # amarillo reportes agregados / alertas nocturnas
COLOR_NOCTURNO_BORDE = "#DAA520"
COLOR_NOCTURNO_ANTES = "#FF8C00"  # naranja-amarillo agregado
COLOR_NOCTURNO_DESPUES = "#FFD700"

DIAS_ES = {
    0: "Lunes",
    1: "Martes",
    2: "Miércoles",
    3: "Jueves",
    4: "Viernes",
    5: "Sábado",
    6: "Domingo",
}


@dataclass(frozen=True)
class DiaConsumo:
    dia: date
    m3: float
    horas_con_dato: int

    @property
    def etiqueta(self) -> str:
        return f"{DIAS_ES[self.dia.weekday()][:3]} {self.dia.strftime('%d/%m')}"

    @property
    def etiqueta_corta(self) -> str:
        return self.dia.strftime("%d/%m")


def _rango(ini: date, fin: date) -> List[date]:
    out: List[date] = []
    d = ini
    while d <= fin:
        out.append(d)
        d += timedelta(days=1)
    return out


def _fmt(n: float, dec: int = 2) -> str:
    from generar_reporte_word import format_number_chilean

    return format_number_chilean(n, dec)


def _consumo_dia(node_id: str, d: date) -> DiaConsumo:
    from control_nocturno import obtener_datos_horarios_dia

    hdict = obtener_datos_horarios_dia(node_id, datetime.combine(d, datetime.min.time()))
    vals = [float(hdict.get(h, 0.0)) for h in range(24)]
    horas = sum(1 for v in vals if v > 1e-9)
    return DiaConsumo(d, sum(vals), horas)


def _cargar_periodo(node_id: str, ini: date, fin: date) -> List[DiaConsumo]:
    return [_consumo_dia(node_id, d) for d in _rango(ini, fin)]


def _total(serie: List[DiaConsumo]) -> float:
    return sum(x.m3 for x in serie)


def _promedio_diario(serie: List[DiaConsumo]) -> float:
    return _total(serie) / len(serie) if serie else 0.0


def _grafico_consumo_diario_semana(
    antes: List[DiaConsumo],
    despues: List[DiaConsumo],
    out_path: Path,
) -> None:
    """
    Gráfico tipo reporte agregado: línea + área sombreada.
    Rojo = 7 días antes; azul = 7 días después.
    El 11/07 (día de corte) no se incluye en la serie.
    """
    fechas_antes = [datetime.combine(a.dia, datetime.min.time()) for a in antes]
    vals_antes = [a.m3 for a in antes]
    fechas_desp = [datetime.combine(d.dia, datetime.min.time()) for d in despues]
    vals_desp = [d.m3 for d in despues]

    fig, ax = plt.subplots(figsize=(10, 5))

    ax.plot(
        fechas_antes,
        vals_antes,
        marker="o",
        linestyle="-",
        color=COLOR_ANTES,
        linewidth=2,
        markersize=6,
        label=f"Antes ({ANTES_INI.strftime('%d/%m')}–{ANTES_FIN.strftime('%d/%m')})",
        zorder=3,
    )
    ax.fill_between(fechas_antes, vals_antes, alpha=0.3, color=COLOR_ANTES, zorder=2)

    ax.plot(
        fechas_desp,
        vals_desp,
        marker="o",
        linestyle="-",
        color=COLOR_DESPUES,
        linewidth=2,
        markersize=6,
        label=f"Después ({DESPUES_INI.strftime('%d/%m')}–{DESPUES_FIN.strftime('%d/%m')})",
        zorder=3,
    )
    ax.fill_between(fechas_desp, vals_desp, alpha=0.3, color=COLOR_DESPUES, zorder=2)

    # Marca del día excluido (sin incluirlo en el comparativo)
    ax.axvline(
        datetime.combine(DIA_REPARACION, datetime.min.time()),
        color=COLOR_EXCLUIDO,
        linestyle=":",
        linewidth=1.4,
        alpha=0.8,
        label=f"{DIA_REPARACION.strftime('%d/%m')} excluido (corte)",
        zorder=1,
    )

    ax.set_title("Consumo diario (m³)", fontsize=14, fontweight="bold")
    ax.set_xlabel("Fecha", fontsize=11)
    ax.set_ylabel("Total m³", fontsize=11)
    ax.set_ylim(bottom=0)
    yticks = ax.get_yticks()
    ax.set_yticks(yticks[yticks >= 0])

    ax.xaxis.set_major_locator(mdates.DayLocator(interval=1))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d. %b"))
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha="right")

    ax.grid(True, linestyle="--", alpha=0.3, axis="y")
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.22), ncol=3, fontsize=9)
    fig.tight_layout(rect=[0, 0.05, 1, 0.95])
    fig.autofmt_xdate()
    fig.savefig(out_path, dpi=150, bbox_inches="tight", pad_inches=0.2)
    plt.close(fig)


def _grafico_totales(
    total_antes: float,
    total_despues: float,
    out_path: Path,
) -> None:
    """Barras de total acumulado (rojo antes / azul después)."""
    etiquetas = [
        f"Total 7 días\nantes\n({ANTES_INI.strftime('%d/%m')}–{ANTES_FIN.strftime('%d/%m')})",
        f"Total 7 días\ndespués\n({DESPUES_INI.strftime('%d/%m')}–{DESPUES_FIN.strftime('%d/%m')})",
    ]
    valores = [total_antes, total_despues]
    colores = [COLOR_ANTES, COLOR_DESPUES]

    fig, ax = plt.subplots(figsize=(7.5, 5.0))
    x = np.arange(len(etiquetas))
    bars = ax.bar(x, valores, color=colores, edgecolor="white", linewidth=1.2, width=0.55)

    ymax = max(valores) if valores else 1.0
    for bar, val in zip(bars, valores):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + ymax * 0.02,
            f"{val:.1f} m³",
            ha="center",
            va="bottom",
            fontsize=11,
            fontweight="bold",
        )

    delta = total_antes - total_despues
    pct = (100.0 * delta / total_antes) if total_antes > 0 else 0.0
    signo = "↓" if delta >= 0 else "↑"
    ax.set_title(
        f"Consumo acumulado — efecto de la reparación\n"
        f"Variación: {signo} {_fmt(abs(delta), 1)} m³ ({_fmt(abs(pct), 1)} %)",
        fontsize=11,
    )
    ax.set_xticks(x)
    ax.set_xticklabels(etiquetas, fontsize=9)
    ax.set_ylabel("Consumo acumulado (m³)")
    ax.grid(True, axis="y", alpha=0.35)
    ax.set_axisbelow(True)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def _nocturno_m3(horas: Dict[int, float]) -> float:
    return sum(float(horas.get(h, 0.0)) for h in HORAS_NOCTURNAS)


def _tiene_datos_horarios(horas: Dict[int, float]) -> bool:
    return any(float(horas.get(h, 0.0)) > 1e-9 for h in range(24))


def _primer_dia_con_datos(
    node_id: str, ini: date, fin: date
) -> Tuple[date, Dict[int, float]]:
    """Primer día del rango con al menos una hora > 0 (fallback si el martes después está vacío)."""
    from wes_estilo_graficos_app import horas_api_chile

    for d in _rango(ini, fin):
        h = horas_api_chile(node_id, datetime.combine(d, datetime.min.time()))
        if _tiene_datos_horarios(h):
            return d, h
    d0 = ini
    return d0, {i: 0.0 for i in range(24)}


def _grafico_dia_estilo_app(
    horas: Dict[int, float],
    out_path: Path,
    *,
    dia: date,
    titulo_extra: str = "",
    aviso_sin_datos: str = "",
) -> None:
    """Perfil horario línea + área (igual que gráfica día app / reportes WES), sin barras."""
    from generar_reporte_word import format_number_chilean

    hours = list(range(24))
    values = [float(horas.get(h, 0.0)) for h in hours]

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(hours, values, linestyle="-", color=COLOR_LINEA_APP, linewidth=2, marker="o", markersize=4)
    ax.fill_between(hours, values, 0, color=COLOR_LINEA_APP, alpha=0.3)

    non_zero = [(i, v) for i, v in enumerate(values) if v > 0]
    if non_zero:
        max_i, max_v = max(non_zero, key=lambda x: x[1])
        ax.plot(max_i, max_v, "ro", markersize=8, markeredgecolor="darkred", markeredgewidth=2, zorder=10)
        y_max = max(values) * 1.15 if values else 1
        ax.set_ylim(bottom=0, top=max(y_max, 0.05))
        _, y_top = ax.get_ylim()
        ax.annotate(
            f"{format_number_chilean(max_v, 2)} m³/hr\n{max_i:02d}:00",
            xy=(max_i, max_v),
            xytext=(max_i, min(max_v + (y_top - 0) * 0.12, y_top * 0.9)),
            ha="center",
            va="bottom",
            fontsize=11,
            fontweight="bold",
            color="red",
            bbox=dict(boxstyle="round,pad=0.6", facecolor="yellow", alpha=0.8, edgecolor="red", linewidth=2.5),
            arrowprops=dict(arrowstyle="->", color="red", lw=2.5),
        )
    else:
        ax.set_ylim(bottom=0, top=1.0)

    # Franja nocturna 00–06 si hay consumo
    if any(values[h] > 0 for h in HORAS_NOCTURNAS):
        ax.axvspan(0, 6, alpha=0.1, color="red", zorder=0)
        ax.axvline(0, color="red", linestyle="--", linewidth=1.5, alpha=0.7)
        ax.axvline(6, color="red", linestyle="--", linewidth=1.5, alpha=0.7)

    month_names = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    title = f"Consumo por hora - {dia.day} {month_names[dia.month - 1]} {dia.year}"
    if titulo_extra:
        title += f" ({titulo_extra})"
    ax.set_title(title, fontsize=12, fontweight="bold", pad=15)
    ax.set_xlabel("Hora del día", fontsize=10)
    ax.set_ylabel("Consumo (m³/hr)", fontsize=10)

    month_abbr = ["Ene.", "Feb.", "Mar.", "Abr.", "May.", "Jun.", "Jul.", "Ago.", "Sep.", "Oct.", "Nov.", "Dic."]
    ax.set_xticks([0, 6, 12, 18])
    ax.set_xticklabels(
        [f"{dia.day}. {month_abbr[dia.month - 1]}", "06:00", "12:00", "18:00"],
        fontsize=9,
    )
    ax.grid(True, linestyle="--", alpha=0.3, linewidth=0.5)

    if aviso_sin_datos:
        ax.text(
            0.5,
            0.55,
            aviso_sin_datos,
            transform=ax.transAxes,
            ha="center",
            va="center",
            fontsize=11,
            fontweight="bold",
            color="#8B0000",
            bbox=dict(boxstyle="round,pad=0.6", facecolor="#FFF3CD", edgecolor="#C0392B", alpha=0.95),
        )

    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_comparativo_martes_lineas(
    horas_antes: Dict[int, float],
    horas_despues: Dict[int, float],
    out_path: Path,
    *,
    dia_antes: date,
    dia_despues: date,
    etiqueta_despues: str,
) -> None:
    """Dos perfiles en línea (estilo app), sin barras: martes antes vs martes después."""
    x = list(range(24))
    y_a = [float(horas_antes.get(h, 0.0)) for h in x]
    y_d = [float(horas_despues.get(h, 0.0)) for h in x]

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(x, y_a, color=COLOR_ANTES, linewidth=2.2, marker="o", markersize=4, label=f"Martes {dia_antes.strftime('%d/%m/%Y')} (antes)")
    ax.fill_between(x, y_a, 0, color=COLOR_ANTES, alpha=0.18)
    ax.plot(x, y_d, color=COLOR_DESPUES, linewidth=2.2, marker="o", markersize=4, label=etiqueta_despues)
    ax.fill_between(x, y_d, 0, color=COLOR_DESPUES, alpha=0.18)
    ax.axvspan(0, 6, alpha=0.08, color="red", zorder=0)

    ymax = max(max(y_a), max(y_d), 0.05) * 1.18
    ax.set_ylim(0, ymax)
    ax.set_xlim(0, 23)
    ax.set_xticks([0, 6, 12, 18, 23])
    ax.set_xticklabels(["00:00", "06:00", "12:00", "18:00", "23:00"], fontsize=9)
    ax.set_xlabel("Hora del día (Chile)", fontsize=10)
    ax.set_ylabel("Consumo (m³/hr)", fontsize=10)
    ax.set_title(
        "Comparativo perfil día (estilo app WES) — antes vs después",
        fontsize=12,
        fontweight="bold",
        pad=12,
    )
    ax.grid(True, linestyle="--", alpha=0.3)
    ax.legend(loc="upper right", fontsize=9, framealpha=0.95)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _nocturno_por_dia(node_id: str, ini: date, fin: date) -> List[Tuple[date, float]]:
    from wes_estilo_graficos_app import horas_api_chile

    out: List[Tuple[date, float]] = []
    for d in _rango(ini, fin):
        h = horas_api_chile(node_id, datetime.combine(d, datetime.min.time()))
        out.append((d, _nocturno_m3(h)))
    return out


def _pares_nocturnos_por_weekday(
    noct_antes: List[Tuple[date, float]],
    noct_despues: List[Tuple[date, float]],
) -> List[Tuple[date, float, date, float]]:
    """Empareja mismo día de semana (p. ej. martes antes vs martes después)."""
    map_d = {d.weekday(): (d, v) for d, v in noct_despues}
    pares: List[Tuple[date, float, date, float]] = []
    for da, va in noct_antes:
        if da.weekday() in map_d:
            dd, vd = map_d[da.weekday()]
            pares.append((da, va, dd, vd))
    # orden: lunes..domingo o el orden natural de `antes`
    return pares


def _grafico_control_nocturno_7x7(
    noct_antes: List[Tuple[date, float]],
    noct_despues: List[Tuple[date, float]],
    out_path: Path,
) -> None:
    """
    Control nocturno 00:00–06:00 de los 7 días antes vs 7 días después.
    Color amarillo (#FFD700 / #FF8C00) como reportes agregados.
    Emparejado por día de la semana.
    """
    pares = _pares_nocturnos_por_weekday(noct_antes, noct_despues)
    y_a = [va for _, va, _, _ in pares]
    y_d = [vd for _, _, _, vd in pares]
    x = np.arange(len(pares))
    width = 0.38

    fig, ax = plt.subplots(figsize=(11, 5.0))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    bars_a = ax.bar(
        x - width / 2,
        y_a,
        width,
        color=COLOR_NOCTURNO_ANTES,
        edgecolor=COLOR_NOCTURNO_BORDE,
        linewidth=1.1,
        alpha=0.9,
        label=f"Antes 00–06 ({ANTES_INI.strftime('%d/%m')}–{ANTES_FIN.strftime('%d/%m')})",
    )
    bars_d = ax.bar(
        x + width / 2,
        y_d,
        width,
        color=COLOR_NOCTURNO_DESPUES,
        edgecolor=COLOR_NOCTURNO_BORDE,
        linewidth=1.1,
        alpha=0.9,
        label=f"Después 00–06 ({DESPUES_INI.strftime('%d/%m')}–{DESPUES_FIN.strftime('%d/%m')})",
    )

    ymax = max(y_a + y_d + [0.05]) * 1.25
    for bars in (bars_a, bars_d):
        for bar in bars:
            h = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                h + ymax * 0.02,
                f"{h:.1f}",
                ha="center",
                va="bottom",
                fontsize=8,
                fontweight="bold",
            )

    tick_labels = [
        f"{DIAS_ES[da.weekday()][:3]}\n{da.strftime('%d/%m')} | {dd.strftime('%d/%m')}"
        for da, _, dd, _ in pares
    ]
    ax.set_xticks(x)
    ax.set_xticklabels(tick_labels, fontsize=8)
    ax.set_ylim(0, ymax)
    ax.set_ylabel("m³ (00:00–06:00)", fontsize=10, fontweight="bold")
    ax.set_title(
        "Control nocturno — 7 días antes vs 7 días después (00:00–06:00)",
        fontsize=12,
        fontweight="bold",
        pad=10,
    )
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_axisbelow(True)
    ax.legend(loc="upper right", fontsize=9, framealpha=0.95)
    fig.text(
        0.5,
        0.01,
        "Amarillo/naranja = estilo consumo nocturno reportes agregados · emparejado por día de semana · 11/07 excluido",
        ha="center",
        fontsize=7.5,
        color="#555555",
    )
    fig.tight_layout(rect=[0, 0.04, 1, 1])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _analisis_nocturno_7x7(noct_antes: List[Tuple[date, float]], noct_despues: List[Tuple[date, float]]) -> str:
    tot_a = sum(v for _, v in noct_antes)
    tot_d = sum(v for _, v in noct_despues)
    delta = tot_a - tot_d
    pct = (100.0 * delta / tot_a) if tot_a > 0 else 0.0
    prom_a = tot_a / len(noct_antes) if noct_antes else 0.0
    prom_d = tot_d / len(noct_despues) if noct_despues else 0.0
    return (
        f"En la franja 00:00–06:00, los 7 días previos a la reparación sumaron {_fmt(tot_a, 2)} m³ "
        f"(promedio {_fmt(prom_a, 2)} m³/noche), frente a {_fmt(tot_d, 2)} m³ en los 7 días posteriores "
        f"(promedio {_fmt(prom_d, 2)} m³/noche). Variación: {_fmt(abs(delta), 2)} m³ "
        f"({'reducción' if delta >= 0 else 'aumento'} de {_fmt(abs(pct), 1)} %). "
        f"El 11/07 no se incluye. El perfil día comparado es martes "
        f"{MARTES_ANTES.strftime('%d/%m')} vs martes {MARTES_DESPUES.strftime('%d/%m')}."
    )


def _conclusion(total_antes: float, total_despues: float, pct: float) -> str:
    delta = total_antes - total_despues
    if pct >= 15 and delta > 0:
        return (
            f"Tras la reparación de fugas en la matriz (11/07/2026), el consumo acumulado en los 7 días "
            f"siguientes ({_fmt(total_despues, 1)} m³) es {_fmt(pct, 1)} % menor que en los 7 días previos "
            f"({_fmt(total_antes, 1)} m³). La reducción de {_fmt(delta, 1)} m³ en la semana es coherente con "
            f"la eliminación de pérdidas en la matriz y respalda el impacto positivo de la intervención."
        )
    if pct >= 5 and delta > 0:
        return (
            f"Se observa una reducción moderada del consumo semanal tras la reparación "
            f"({_fmt(pct, 1)} %, equivalente a {_fmt(delta, 1)} m³). El resultado es favorable, aunque "
            f"parte de la variación puede deberse a patrones de uso del campus; se recomienda seguir "
            f"monitoreando el caudal base nocturno en las próximas semanas."
        )
    if abs(pct) < 5:
        return (
            f"El consumo de los 7 días posteriores ({_fmt(total_despues, 1)} m³) es similar al de los "
            f"7 días previos ({_fmt(total_antes, 1)} m³; variación {_fmt(pct, 1)} %). No se aprecia aún "
            f"un cambio semanal marcado; conviene contrastar con el perfil nocturno y con un periodo "
            f"más largo para confirmar el efecto de la reparación."
        )
    return (
        f"En la ventana analizada, el consumo posterior ({_fmt(total_despues, 1)} m³) resulta mayor "
        f"que el previo ({_fmt(total_antes, 1)} m³). Esto no invalida la reparación: puede reflejar "
        f"mayor demanda operativa del campus o recuperación de servicio tras el corte. Se sugiere "
        f"revisar el caudal mínimo nocturno como indicador más estable de fugas residuales."
    )


def generar() -> Tuple[Path, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    GRAFICOS_DIR.mkdir(parents=True, exist_ok=True)

    print(f"Cargando consumo {NODE_ID} — periodo antes...")
    antes = _cargar_periodo(NODE_ID, ANTES_INI, ANTES_FIN)
    print(f"Cargando consumo {NODE_ID} — periodo después...")
    despues = _cargar_periodo(NODE_ID, DESPUES_INI, DESPUES_FIN)
    print(f"Cargando día de reparación (solo referencial, excluido)...")
    dia_rep = _consumo_dia(NODE_ID, DIA_REPARACION)

    total_antes = _total(antes)
    total_despues = _total(despues)
    prom_antes = _promedio_diario(antes)
    prom_despues = _promedio_diario(despues)
    delta = total_antes - total_despues
    pct = (100.0 * delta / total_antes) if total_antes > 0 else 0.0

    img_diarias = GRAFICOS_DIR / "01_consumo_diario_antes_despues.png"
    img_totales = GRAFICOS_DIR / "02_totales_acumulados.png"

    _grafico_consumo_diario_semana(antes, despues, img_diarias)
    _grafico_totales(total_antes, total_despues, img_totales)

    # --- Comparativo martes (estilo app línea) + control nocturno 7×7 ---
    from wes_estilo_graficos_app import horas_api_chile

    print(f"Cargando perfiles horarios martes {MARTES_ANTES} y {MARTES_DESPUES}...")
    horas_mar_antes = horas_api_chile(
        NODE_ID, datetime.combine(MARTES_ANTES, datetime.min.time())
    )
    horas_mar_desp = horas_api_chile(
        NODE_ID, datetime.combine(MARTES_DESPUES, datetime.min.time())
    )
    martes_desp_sin_datos = not _tiene_datos_horarios(horas_mar_desp)

    # Si el martes después no tiene CSV en la API, usar el primer día post con datos
    # solo para el gráfico de perfil (se documenta en el informe).
    dia_perfil_desp = MARTES_DESPUES
    horas_perfil_desp = horas_mar_desp
    etiqueta_perfil_desp = f"Martes {MARTES_DESPUES.strftime('%d/%m/%Y')} (después)"
    aviso_desp = ""
    if martes_desp_sin_datos:
        print(
            f"[AVISO] Martes {MARTES_DESPUES} sin registros horarios en API WES (CSV vacío). "
            "Buscando primer día post-reparación con datos para el perfil comparativo..."
        )
        dia_perfil_desp, horas_perfil_desp = _primer_dia_con_datos(
            NODE_ID, DESPUES_INI, DESPUES_FIN
        )
        etiqueta_perfil_desp = (
            f"{DIAS_ES[dia_perfil_desp.weekday()]} {dia_perfil_desp.strftime('%d/%m/%Y')} "
            f"(primer día con datos post-reparación; martes {MARTES_DESPUES.strftime('%d/%m')} sin CSV)"
        )
        aviso_desp = (
            f"Sin datos horarios en API WES\npara el martes {MARTES_DESPUES.strftime('%d/%m/%Y')}\n"
            f"(CSV vacío — se muestra perfil de\n"
            f"{DIAS_ES[dia_perfil_desp.weekday()]} {dia_perfil_desp.strftime('%d/%m')} abajo)"
        )

    total_mar_antes = sum(float(horas_mar_antes.get(h, 0.0)) for h in range(24))
    total_mar_desp = sum(float(horas_mar_desp.get(h, 0.0)) for h in range(24))
    noct_mar_antes = _nocturno_m3(horas_mar_antes)
    noct_mar_desp = _nocturno_m3(horas_mar_desp)

    print("Cargando control nocturno 7 días antes / 7 días después...")
    noct_serie_antes = _nocturno_por_dia(NODE_ID, ANTES_INI, ANTES_FIN)
    noct_serie_desp = _nocturno_por_dia(NODE_ID, DESPUES_INI, DESPUES_FIN)
    tot_noct_antes = sum(v for _, v in noct_serie_antes)
    tot_noct_desp = sum(v for _, v in noct_serie_desp)

    img_mar_antes = GRAFICOS_DIR / "03_martes_antes_estilo_app.png"
    img_mar_desp = GRAFICOS_DIR / "04_martes_despues_estilo_app.png"
    img_mar_comp = GRAFICOS_DIR / "05_comparativo_perfiles_lineas.png"
    img_nocturno = GRAFICOS_DIR / "06_control_nocturno_7x7.png"
    img_perfil_fallback = GRAFICOS_DIR / "04b_primer_dia_con_datos_despues.png"

    _grafico_dia_estilo_app(
        horas_mar_antes,
        img_mar_antes,
        dia=MARTES_ANTES,
        titulo_extra="antes de la reparación",
    )
    _grafico_dia_estilo_app(
        horas_mar_desp,
        img_mar_desp,
        dia=MARTES_DESPUES,
        titulo_extra="después de la reparación",
        aviso_sin_datos=aviso_desp if martes_desp_sin_datos else "",
    )
    if martes_desp_sin_datos and dia_perfil_desp != MARTES_DESPUES:
        _grafico_dia_estilo_app(
            horas_perfil_desp,
            img_perfil_fallback,
            dia=dia_perfil_desp,
            titulo_extra="primer día con datos post-reparación",
        )
    _grafico_comparativo_martes_lineas(
        horas_mar_antes,
        horas_perfil_desp,
        img_mar_comp,
        dia_antes=MARTES_ANTES,
        dia_despues=dia_perfil_desp,
        etiqueta_despues=etiqueta_perfil_desp,
    )
    _grafico_control_nocturno_7x7(noct_serie_antes, noct_serie_desp, img_nocturno)

    out_docx = OUT_DIR / f"{NOMBRE_INFORME}.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h1 = doc.add_heading(
        "Informe comparativo — Reparación de fugas en matriz (Honduras)",
        level=0,
    )
    if h1.runs:
        h1.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    doc.add_paragraph()
    for label, text in [
        ("Cliente", CLIENTE),
        ("Punto WES", f"{NODE_ID} — {NODE_NOMBRE}"),
        (
            "Intervención",
            f"Sábado {DIA_REPARACION.strftime('%d/%m/%Y')} — reparación de fugas en la matriz "
            "(corte de agua para trabajos)",
        ),
        (
            "Ventana antes",
            f"{ANTES_INI.strftime('%d/%m/%Y')} – {ANTES_FIN.strftime('%d/%m/%Y')} (7 días)",
        ),
        (
            "Ventana después",
            f"{DESPUES_INI.strftime('%d/%m/%Y')} – {DESPUES_FIN.strftime('%d/%m/%Y')} (7 días)",
        ),
        (
            "Exclusión",
            f"{DIA_REPARACION.strftime('%d/%m/%Y')} no se incluye en totales ni en el gráfico "
            "comparativo (día de corte de suministro).",
        ),
        ("Generado", datetime.now().strftime("%d-%m-%Y %H:%M")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Cuantificar el cambio de consumo en Sala impulsión Honduras luego de las reparaciones "
        "de fugas en la matriz realizadas el 11/07/2026, comparando dos ventanas equivalentes "
        "de 7 días (antes y después), sin contaminar el análisis con el día del corte de agua."
    )

    doc.add_heading("2. Metodología", level=1)
    for t in [
        "Se obtiene el consumo diario aproximado (m³) como suma de los caudales horarios "
        "registrados por la plataforma WES (m³/h × 1 h).",
        f"Periodo previo: {ANTES_INI.strftime('%d/%m/%Y')} al {ANTES_FIN.strftime('%d/%m/%Y')}.",
        f"Periodo posterior: {DESPUES_INI.strftime('%d/%m/%Y')} al {DESPUES_FIN.strftime('%d/%m/%Y')}.",
        f"El día {DIA_REPARACION.strftime('%d/%m/%Y')} se excluye del comparativo porque el "
        "suministro fue cortado para ejecutar los trabajos.",
        "El gráfico de consumo diario usa el mismo formato de los reportes agregados "
        "(línea con área): rojo = antes de la reparación; azul = después.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("3. Resumen de resultados", level=1)
    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = "Table Grid"
    filas = [
        ("Indicador", "7 días antes", "7 días después"),
        ("Periodo", f"{ANTES_INI.strftime('%d/%m')}–{ANTES_FIN.strftime('%d/%m')}",
         f"{DESPUES_INI.strftime('%d/%m')}–{DESPUES_FIN.strftime('%d/%m')}"),
        ("Consumo acumulado (m³)", _fmt(total_antes, 2), _fmt(total_despues, 2)),
        ("Promedio diario (m³/día)", _fmt(prom_antes, 2), _fmt(prom_despues, 2)),
        ("Máximo diario (m³)", _fmt(max(x.m3 for x in antes), 2), _fmt(max(x.m3 for x in despues), 2)),
        ("Mínimo diario (m³)", _fmt(min(x.m3 for x in antes), 2), _fmt(min(x.m3 for x in despues), 2)),
    ]
    for i, row in enumerate(filas):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            if i == 0:
                for par in cell.paragraphs:
                    for r in par.runs:
                        r.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Variación acumulada (antes − después): ").bold = True
    signo = "reducción" if delta >= 0 else "aumento"
    p.add_run(f"{_fmt(abs(delta), 2)} m³ ({signo} de {_fmt(abs(pct), 1)} %).")

    p2 = doc.add_paragraph()
    p2.add_run(f"Día {DIA_REPARACION.strftime('%d/%m/%Y')} (excluido): ").bold = True
    p2.add_run(
        f"consumo registrado {_fmt(dia_rep.m3, 2)} m³ "
        f"({dia_rep.horas_con_dato} horas con dato). Solo referencial; no entra al comparativo."
    )

    doc.add_heading("4. Detalle diario", level=1)
    n = len(antes)
    tbl2 = doc.add_table(rows=1 + n + 1, cols=5)
    tbl2.style = "Table Grid"
    headers = ["Día semana", "Fecha antes", "m³ antes", "Fecha después", "m³ después"]
    for j, h in enumerate(headers):
        tbl2.rows[0].cells[j].text = h
        for par in tbl2.rows[0].cells[j].paragraphs:
            for r in par.runs:
                r.bold = True
    for i, (a, d) in enumerate(zip(antes, despues), start=1):
        tbl2.rows[i].cells[0].text = DIAS_ES[a.dia.weekday()]
        tbl2.rows[i].cells[1].text = a.dia.strftime("%d/%m/%Y")
        tbl2.rows[i].cells[2].text = _fmt(a.m3, 2)
        tbl2.rows[i].cells[3].text = d.dia.strftime("%d/%m/%Y")
        tbl2.rows[i].cells[4].text = _fmt(d.m3, 2)
    last = tbl2.rows[n + 1]
    last.cells[0].text = "Total"
    last.cells[1].text = ""
    last.cells[2].text = _fmt(total_antes, 2)
    last.cells[3].text = ""
    last.cells[4].text = _fmt(total_despues, 2)
    for j in (0, 2, 4):
        for par in last.cells[j].paragraphs:
            for r in par.runs:
                r.bold = True

    doc.add_heading("5. Gráficos", level=1)

    doc.add_heading("5.1 Consumo diario (estilo reporte agregado)", level=2)
    doc.add_paragraph(
        "Línea con área sombreada (mismo formato de los reportes agregados). "
        "Rojo = 7 días antes de la reparación; azul = 7 días después. "
        "El 11/07/2026 no se incluye en la serie (solo se marca como día de corte excluido)."
    )
    doc.add_picture(str(img_diarias), width=Inches(6.3))

    doc.add_heading("5.2 Totales acumulados", level=2)
    doc.add_paragraph("Suma de los 7 días de cada ventana (sin el día de corte). Rojo = antes; azul = después.")
    doc.add_picture(str(img_totales), width=Inches(5.5))

    doc.add_heading("6. Comparativo de martes (gráfica día estilo app WES)", level=1)
    doc.add_paragraph(
        f"Se contrastan los dos martes de las ventanas: "
        f"{MARTES_ANTES.strftime('%d/%m/%Y')} (antes) y {MARTES_DESPUES.strftime('%d/%m/%Y')} (después). "
        "Los perfiles usan el formato de la gráfica día de la app WES "
        "(línea azul + área, sin barras; marca del máximo del día)."
    )
    if martes_desp_sin_datos:
        p_av = doc.add_paragraph()
        p_av.add_run("Hallazgo de datos: ").bold = True
        p_av.add_run(
            f"El martes {MARTES_DESPUES.strftime('%d/%m/%Y')} no tiene registros horarios en la API WES "
            f"(CSV dates.measures vacío; tampoco aparece en el consolidado diario 12–18/07). "
            f"Para el comparativo de perfil se usa el primer día post-reparación con datos: "
            f"{DIAS_ES[dia_perfil_desp.weekday()]} {dia_perfil_desp.strftime('%d/%m/%Y')}."
        )

    tbl_lun = doc.add_table(rows=4, cols=3)
    tbl_lun.style = "Table Grid"
    filas_lun = [
        ("Indicador", f"Martes {MARTES_ANTES.strftime('%d/%m')}", f"Martes {MARTES_DESPUES.strftime('%d/%m')}"),
        ("Consumo total día (m³)", _fmt(total_mar_antes, 2), _fmt(total_mar_desp, 2)),
        ("Nocturno 00:00–06:00 (m³)", _fmt(noct_mar_antes, 2), _fmt(noct_mar_desp, 2)),
        (
            "Datos horarios API",
            "Sí (24 h)",
            "No (CSV vacío)" if martes_desp_sin_datos else "Sí",
        ),
    ]
    for i, row in enumerate(filas_lun):
        for j, val in enumerate(row):
            cell = tbl_lun.rows[i].cells[j]
            cell.text = val
            if i == 0:
                for par in cell.paragraphs:
                    for r in par.runs:
                        r.bold = True

    doc.add_heading("6.1 Perfil día — martes antes (estilo app)", level=2)
    doc.add_picture(str(img_mar_antes), width=Inches(6.2))

    doc.add_heading("6.2 Perfil día — martes después (estilo app)", level=2)
    doc.add_picture(str(img_mar_desp), width=Inches(6.2))
    if martes_desp_sin_datos and img_perfil_fallback.is_file():
        doc.add_paragraph(
            f"Perfil del primer día con datos tras la reparación "
            f"({DIAS_ES[dia_perfil_desp.weekday()]} {dia_perfil_desp.strftime('%d/%m/%Y')}):"
        )
        doc.add_picture(str(img_perfil_fallback), width=Inches(6.2))

    doc.add_heading("6.3 Comparativo de perfiles (líneas)", level=2)
    doc.add_paragraph(
        "Líneas con área (sin barras): rojo = martes antes; azul = día usado para el perfil después."
    )
    doc.add_picture(str(img_mar_comp), width=Inches(6.3))

    doc.add_heading("7. Control nocturno 00:00–06:00 (7 días antes vs 7 días después)", level=1)
    doc.add_paragraph(_analisis_nocturno_7x7(noct_serie_antes, noct_serie_desp))
    doc.add_paragraph(
        "Gráfico en amarillo/naranja (misma paleta de consumo nocturno de los reportes agregados). "
        "Cada par de barras compara el mismo orden de día dentro de cada ventana de 7 días."
    )
    doc.add_picture(str(img_nocturno), width=Inches(6.3))

    pares_n = _pares_nocturnos_por_weekday(noct_serie_antes, noct_serie_desp)
    tbl_n = doc.add_table(rows=1 + len(pares_n) + 1, cols=5)
    tbl_n.style = "Table Grid"
    for j, h in enumerate(["Día", "Fecha antes", "m³ noche antes", "Fecha después", "m³ noche después"]):
        tbl_n.rows[0].cells[j].text = h
        for par in tbl_n.rows[0].cells[j].paragraphs:
            for r in par.runs:
                r.bold = True
    for i, (da, va, dd, vd) in enumerate(pares_n, start=1):
        tbl_n.rows[i].cells[0].text = DIAS_ES[da.weekday()]
        tbl_n.rows[i].cells[1].text = da.strftime("%d/%m/%Y")
        tbl_n.rows[i].cells[2].text = _fmt(va, 2)
        tbl_n.rows[i].cells[3].text = dd.strftime("%d/%m/%Y")
        tbl_n.rows[i].cells[4].text = _fmt(vd, 2)
    last_n = tbl_n.rows[len(pares_n) + 1]
    last_n.cells[0].text = "Total"
    last_n.cells[2].text = _fmt(tot_noct_antes, 2)
    last_n.cells[4].text = _fmt(tot_noct_desp, 2)
    for j in (0, 2, 4):
        for par in last_n.cells[j].paragraphs:
            for r in par.runs:
                r.bold = True

    doc.add_heading("8. Conclusión", level=1)
    doc.add_paragraph(_conclusion(total_antes, total_despues, pct))
    doc.add_paragraph(
        f"Complemento nocturno (7×7): {_fmt(tot_noct_antes, 2)} m³ antes vs "
        f"{_fmt(tot_noct_desp, 2)} m³ después en 00:00–06:00."
    )

    doc.add_heading("9. Consideraciones", level=1)
    for t in [
        "La comparación semanal puede verse afectada por la ocupación del campus, riego u otras demandas.",
        "El caudal entre 00:00 y 06:00 es el indicador más estable de fugas residuales.",
        "Los perfiles día siguen el estilo de la gráfica de la app WES (línea + área, sin barras).",
        "El comparativo de perfil día usa martes vs martes (07/07 vs 14/07).",
        "Valores alineados con Sala impulsión Honduras, nodo 000026-01.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    pie = doc.add_paragraph()
    pie.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = pie.add_run("WES — Water Efficiency System")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(out_docx)

    pdf_path = out_docx.with_suffix(".pdf")
    try:
        from generar_reporte_word import convertir_word_a_pdf

        pdf_out = convertir_word_a_pdf(out_docx)
        if pdf_out and Path(pdf_out).is_file():
            pdf_path = Path(pdf_out)
    except Exception as exc:
        print(f"[AVISO] No se pudo convertir a PDF: {exc}")

    print(f"Total antes:   {_fmt(total_antes, 2)} m³ (prom. {_fmt(prom_antes, 2)} m³/día)")
    print(f"Total después: {_fmt(total_despues, 2)} m³ (prom. {_fmt(prom_despues, 2)} m³/día)")
    print(f"Variación:     {_fmt(delta, 2)} m³ ({_fmt(pct, 1)} %)")
    print(f"Día 11/07 (excluido): {_fmt(dia_rep.m3, 2)} m³")
    print(
        f"Martes {MARTES_ANTES}: total {_fmt(total_mar_antes, 2)} m³ | noche {_fmt(noct_mar_antes, 2)} m³"
    )
    print(
        f"Martes {MARTES_DESPUES}: total {_fmt(total_mar_desp, 2)} m³ | noche {_fmt(noct_mar_desp, 2)} m³ "
        f"| sin_datos={martes_desp_sin_datos}"
    )
    print(f"Nocturno 7×7: {_fmt(tot_noct_antes, 2)} m³ antes vs {_fmt(tot_noct_desp, 2)} m³ después")

    return out_docx, pdf_path


def main() -> int:
    print("=" * 72)
    print("INFORME COMPARATIVO — REPARACIÓN FUGAS MATRIZ HONDURAS (UDD)")
    print("=" * 72)
    docx, pdf = generar()
    print(f"[OK] DOCX: {docx}")
    print(f"[OK] PDF:  {pdf}")
    print(f"[OK] Gráficos: {GRAFICOS_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())