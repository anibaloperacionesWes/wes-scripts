"""
Reporte agregado de las 5 auditorías Renca abril 2026 (incl. ICCO 000017-08): totales por semana
(Con WES 13–19 abr / Sin WES 6–12 abr), % de rendimiento, dos Pareto y el gráfico #8 de cada
auditoría (``08_area_<día>.png`` en ``graficos_comparativos/``).

Nombres de establecimiento: API WES (``get_node_name``).

Uso:
  python reporte_agregado_auditorias_renca_pareto.py
"""
from __future__ import annotations

import argparse
import csv
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Sequence, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

from generar_graficos_comparativos_desde_excel_consolidado import totales_rejilla_desde_excel_consolidado
from generar_reporte_word import get_node_name

ROOT = Path(__file__).resolve().parent
BASE_AUDIT = (
    ROOT
    / "reports"
    / "reporte de auditoria"
    / "auditoria_puntos_renca_abril_2026"
)
XLSX_NOMBRE = "consumo_consolidado_parseo_filas_abr06-12_abr13-19_2026.xlsx"

# (node_id, subcarpeta bajo BASE_AUDIT)
AUDITORIAS: Tuple[Tuple[str, str], ...] = (
    ("000017-08", "Auditoria ICCO Renca 000017-08"),
    ("000017-04", "Auditoria Escuela Lo Velazquez 000017-04"),
    ("000017-06", "Auditoria Piscina Municipal 000017-06"),
    ("000017-05", "Auditoria Gimnasio 000017-05"),
    ("000017-07", "Auditoria Cumbre de condores 000017-07"),
)

COLOR_BARRA = "#2980b9"
COLOR_LINEA = "#c0392b"


@dataclass
class FilaAuditoria:
    node_id: str
    nombre_app: str
    path_xlsx: Path
    m3_con_wes: float
    m3_sin_wes: float
    pct_rendimiento: Optional[float]
    """(Semana_sin − Semana_con) / Semana_sin × 100; positivo = menor consumo con WES."""
    path_grafico_08_origen: Optional[Path]
    path_grafico_08_copia: Optional[Path]


def _pct_rendimiento(m3_con: float, m3_sin: float) -> Optional[float]:
    if m3_sin <= 1e-9:
        return None
    return (float(m3_sin) - float(m3_con)) / float(m3_sin) * 100.0


def _find_png_comparativo_08(graficos_dir: Path) -> Optional[Path]:
    """
    Octavo gráfico del lote de ``generar_pngs``: ``08_area_<día>.png``
    (comparativo área + líneas, día homólogo entre periodos).
    """
    if not graficos_dir.is_dir():
        return None
    found = sorted(graficos_dir.glob("08_area_*.png"))
    return found[0] if found else None


def _cargar_filas(out_dir: Path) -> List[FilaAuditoria]:
    filas: List[FilaAuditoria] = []
    for node_id, sub in AUDITORIAS:
        base = (BASE_AUDIT / sub).resolve()
        p_xlsx = (base / XLSX_NOMBRE).resolve()
        if not p_xlsx.is_file():
            raise FileNotFoundError(f"No está el Excel de auditoría: {p_xlsx}")
        nombre = (get_node_name(node_id) or "").strip() or node_id
        t_con, t_sin, _n = totales_rejilla_desde_excel_consolidado(p_xlsx)
        tc, ts = float(t_con), float(t_sin)
        pct = _pct_rendimiento(tc, ts)
        src_08 = _find_png_comparativo_08(base / "graficos_comparativos")
        copia_08: Optional[Path] = None
        if src_08 is not None and src_08.is_file():
            copia_08 = out_dir / f"comparativo08_{node_id.replace('-', '_')}_{src_08.name}"
            shutil.copy2(src_08, copia_08)
        filas.append(
            FilaAuditoria(
                node_id=node_id,
                nombre_app=nombre,
                path_xlsx=p_xlsx,
                m3_con_wes=tc,
                m3_sin_wes=ts,
                pct_rendimiento=pct,
                path_grafico_08_origen=src_08,
                path_grafico_08_copia=copia_08,
            )
        )
    return filas


def _plot_pareto(
    etiquetas: Sequence[str],
    valores: Sequence[float],
    titulo: str,
    subtitulo_eje_y: str,
    out_png: Path,
) -> None:
    pairs = sorted(zip(etiquetas, valores), key=lambda x: -x[1])
    labels = [p[0] for p in pairs]
    vals = [p[1] for p in pairs]
    total = sum(vals) or 1.0
    cum: List[float] = []
    r = 0.0
    for v in vals:
        r += v
        cum.append(r / total * 100.0)

    fig, ax1 = plt.subplots(figsize=(11, 6))
    x = range(len(vals))
    ax1.bar(x, vals, color=COLOR_BARRA, width=0.65)
    ax1.set_ylabel(subtitulo_eje_y, fontsize=10)
    ax1.set_xlabel("Establecimiento (API WES), orden por consumo", fontsize=10)
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(labels, rotation=42, ha="right", fontsize=9)
    ax1.set_title(titulo, fontsize=12, fontweight="bold", pad=12)

    ax2 = ax1.twinx()
    ax2.plot(x, cum, color=COLOR_LINEA, marker="o", linewidth=2, markersize=5)
    ax2.set_ylabel("% acumulado", fontsize=10)
    ax2.set_ylim(0, 105)
    ax2.axhline(80, color=COLOR_LINEA, linestyle="--", linewidth=1, alpha=0.5)

    for i, v in enumerate(vals):
        ax1.text(i, v, f"{v:.0f}", ha="center", va="bottom", fontsize=8, rotation=0)

    fig.tight_layout()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_png, dpi=180, bbox_inches="tight")
    plt.close(fig)


def _escribir_csv(path: Path, filas: List[FilaAuditoria]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(
            [
                "node_id",
                "establecimiento_app",
                "ruta_excel",
                "m3_semana_CON_WES",
                "m3_semana_SIN_WES_referencia",
                "pct_rendimiento_ahorro_vs_ref",
                "grafico_08_origen",
            ]
        )
        for row in filas:
            pct_s = (
                f"{row.pct_rendimiento:.2f}"
                if row.pct_rendimiento is not None
                else ""
            )
            g8 = str(row.path_grafico_08_origen) if row.path_grafico_08_origen else ""
            w.writerow(
                [
                    row.node_id,
                    row.nombre_app,
                    str(row.path_xlsx),
                    f"{row.m3_con_wes:.4f}",
                    f"{row.m3_sin_wes:.4f}",
                    pct_s,
                    g8,
                ]
            )


def _word_agregado(
    out_docx: Path,
    filas: List[FilaAuditoria],
    png_con: Path,
    png_sin: Path,
) -> None:
    doc = Document()
    t = doc.add_heading("Reporte agregado — 5 auditorías Renca (abril 2026)", level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "Se consolidan los totales de la hoja «Consolidado» de cada auditoría "
        "(suma de m³/h en la rejilla del Excel). Los nombres de establecimiento provienen de la API WES. "
        "Periodo con control WES: 13–19 abril 2026. Periodo de referencia (sin WES): 6–12 abril 2026."
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("Fecha de generación: ").bold = True
    p.add_run(datetime.now().strftime("%d-%m-%Y %H:%M"))

    doc.add_heading("Totales y % rendimiento (ahorro vs semana referencia)", level=1)
    doc.add_paragraph(
        "% rendimiento = (Semana sin WES − Semana con WES) / Semana sin WES × 100. "
        "Valor positivo: menor volumen en la semana con control WES respecto a la semana de referencia."
    )

    table = doc.add_table(rows=1 + len(filas), cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Nodo"
    hdr[1].text = "Establecimiento (app)"
    hdr[2].text = "Semana Con WES (m³)"
    hdr[3].text = "Semana Sin WES ref. (m³)"
    hdr[4].text = "% rendimiento"
    for i, row in enumerate(filas, start=1):
        c = table.rows[i].cells
        c[0].text = row.node_id
        c[1].text = row.nombre_app
        c[2].text = f"{row.m3_con_wes:.2f}"
        c[3].text = f"{row.m3_sin_wes:.2f}"
        if row.pct_rendimiento is None:
            c[4].text = "—"
        else:
            c[4].text = f"{row.pct_rendimiento:.1f} %"

    doc.add_heading("Pareto — semana con control WES (13–19 abril 2026)", level=1)
    doc.add_picture(str(png_con), width=Cm(16))
    doc.add_heading("Pareto — semana de referencia sin WES (6–12 abril 2026)", level=1)
    doc.add_picture(str(png_sin), width=Cm(16))

    doc.add_heading(
        "Gráfico comparativo #8 por establecimiento (día homólogo, área + líneas)",
        level=1,
    )
    doc.add_paragraph(
        "Corresponde al archivo ``08_area_<día>.png`` generado en cada carpeta "
        "``graficos_comparativos/`` (mismo orden que en la auditoría individual)."
    )

    for row in filas:
        doc.add_heading(f"{row.nombre_app} ({row.node_id})", level=2)
        pic = row.path_grafico_08_copia or row.path_grafico_08_origen
        if pic is not None and Path(pic).is_file():
            doc.add_picture(str(pic), width=Cm(15))
        else:
            doc.add_paragraph(
                f"[No se encontró el gráfico #8 en {row.path_xlsx.parent / 'graficos_comparativos'}]"
            )

    pie = doc.add_paragraph(
        "Pareto: barras por consumo semanal; curva roja = % acumulado (línea de referencia 80 %)."
    )
    pie.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def main() -> int:
    ap = argparse.ArgumentParser(description="Reporte agregado 5 auditorías + Pareto + gráfico #8.")
    ap.add_argument(
        "--salida",
        type=Path,
        default=None,
        help="Carpeta de salida (default: …/Reporte_agregado_5_auditorias_<timestamp>)",
    )
    args = ap.parse_args()

    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    out_dir = (
        args.salida.resolve()
        if args.salida
        else (BASE_AUDIT / f"Reporte_agregado_5_auditorias_{stamp}").resolve()
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    filas = _cargar_filas(out_dir)
    etiquetas = [f"{f.nombre_app}\n({f.node_id})" for f in filas]
    v_con = [f.m3_con_wes for f in filas]
    v_sin = [f.m3_sin_wes for f in filas]

    png_con = out_dir / "pareto_semana_CON_WES_13_19_abr_2026.png"
    png_sin = out_dir / "pareto_semana_SIN_WES_06_12_abr_2026.png"

    _plot_pareto(
        etiquetas,
        v_con,
        "Pareto — consumo agregado semanal (Con WES, 13–19 abr 2026)",
        "Σ m³ (rejilla horaria, 7 días)",
        png_con,
    )
    _plot_pareto(
        etiquetas,
        v_sin,
        "Pareto — consumo agregado semanal (referencia sin WES, 6–12 abr 2026)",
        "Σ m³ (rejilla horaria, 7 días)",
        png_sin,
    )

    csv_path = out_dir / "resumen_totales_5_auditorias.csv"
    _escribir_csv(csv_path, filas)

    docx_path = out_dir / "Reporte_agregado_5_auditorias.docx"
    _word_agregado(docx_path, filas, png_con, png_sin)

    print(out_dir.resolve())
    print(f"  CSV:   {csv_path.name}")
    print(f"  PNG:   {png_con.name}")
    print(f"  PNG:   {png_sin.name}")
    print(f"  Word:  {docx_path.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
