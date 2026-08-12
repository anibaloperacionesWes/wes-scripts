"""
Reporte: colegios Providencia desde entrega de datos + ahorros entregables.

Inventario de los 5 nodos (company 000006), primera fecha con consumo > 0,
consumo acumulado y análisis de qué métricas de ahorro se pueden entregar
(consumo nocturno medido, proyección control, regulación diurna).

Uso:
  python generar_reporte_ahorro_colegios_providencia.py
  python generar_reporte_ahorro_colegios_providencia.py --desde 01/10/2025 --hasta 12/08/2026
"""

from __future__ import annotations

import argparse
import csv
import sys
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

ROOT = Path(__file__).resolve().parent
COMPANY_ID = "000006"
# Todos los colegios API; 000006-03 suele excluirse de agregados operativos.
ALL_NODES = [
    "000006-01",
    "000006-02",
    "000006-03",
    "000006-04",
    "000006-05",
]
# Reportados habitualmente en agregados mensuales
NODOS_AGREGADO = ["000006-01", "000006-02", "000006-04", "000006-05"]
MES_REF_START = "01/07/2026"
MES_REF_END = "31/07/2026"
MES_REF_LABEL = "julio 2026"


def _fmt_m3(v: float, d: int = 1) -> str:
    from generar_reporte_word import format_number_chilean

    return format_number_chilean(v, d)


def _fmt_clp(v: float) -> str:
    from generar_reporte_word import format_currency_chilean

    return format_currency_chilean(v)


def _naive(dt: datetime) -> datetime:
    if getattr(dt, "tzinfo", None) is not None:
        return dt.replace(tzinfo=None)
    return dt


def _primera_fecha(
    node_id: str, start_dt: datetime, end_dt: datetime
) -> Tuple[Optional[datetime], str]:
    from listado_pa_que_esta_instalado import _primera_fecha_consumo

    return _primera_fecha_consumo(node_id, start_dt, end_dt)


def _consumo_mensual(
    node_id: str, start_dt: datetime, end_dt: datetime
) -> Tuple[float, int, Dict[str, float]]:
    from listado_pa_que_esta_instalado import _medidas_rango

    s = _naive(start_dt)
    e = _naive(end_dt)
    s = datetime(s.year, s.month, s.day)
    pts = _medidas_rango(node_id, s, e)
    tot = 0.0
    dias = 0
    by_month: Dict[str, float] = {}
    for m in pts:
        d = m.date.date() if hasattr(m.date, "date") else m.date
        v = float(m.total_m3 or 0)
        tot += v
        if v > 1e-9:
            dias += 1
        key = f"{d.year}-{d.month:02d}"
        by_month[key] = by_month.get(key, 0.0) + v
    return tot, dias, by_month


def _grafico_consumo_mensual(filas: List[dict], out_path: Path) -> None:
    # Solo nodos con datos recientes en agregado
    meses = sorted({m for f in filas for m in f["by_month"]})
    if not meses:
        return
    fig, ax = plt.subplots(figsize=(10.5, 5.2))
    colors = ["#1f4788", "#16a085", "#c0392b", "#8e44ad", "#e67e22"]
    for i, f in enumerate(filas):
        ys = [f["by_month"].get(m, 0.0) for m in meses]
        ax.plot(
            meses,
            ys,
            marker="o",
            linewidth=1.8,
            markersize=4,
            label=f["nombre_corto"],
            color=colors[i % len(colors)],
        )
    ax.set_ylabel("Consumo mensual (m³)")
    ax.set_title("Consumo mensual desde entrega de datos — Colegios Providencia")
    ax.legend(loc="upper left", fontsize=8)
    ax.grid(True, axis="y", alpha=0.35)
    ax.tick_params(axis="x", rotation=45, labelsize=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: _fmt_m3(v, 0)))
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def _grafico_ahorro_proy(filas_proy: List[dict], out_path: Path) -> None:
    labels = [f["nombre_corto"] for f in filas_proy]
    vals = [f["proy_mes_m3"] for f in filas_proy]
    fig, ax = plt.subplots(figsize=(8.5, 4.6))
    bars = ax.bar(labels, vals, color="#2980b9", edgecolor="white")
    ymax = max(vals) if vals else 1
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + ymax * 0.02,
            _fmt_m3(v, 0),
            ha="center",
            va="bottom",
            fontsize=9,
        )
    ax.set_ylabel("m³/mes (proyección 30 d)")
    ax.set_title(
        f"Ahorro potencial mensual — control consumo nocturno ({MES_REF_LABEL})"
    )
    ax.grid(True, axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def generar(
    desde: str = "01/10/2025",
    hasta: str = "",
    output_dir: Optional[Path] = None,
) -> Path:
    from generar_reporte_word import (
        UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION,
        calculate_nocturnal_metrics,
        get_company_name,
        get_node_name,
        get_water_price_per_m3,
        parse_date,
    )
    from wes_estilo_graficos_app import (
        dia_mayor_consumo_nocturno,
        guardar_grafico_horario_24h_app,
        proyeccion_mensual_desde_nocturno,
    )

    start_dt = parse_date(desde)
    if hasta.strip():
        end_dt = parse_date(hasta, end_of_day=True)
    else:
        end_dt = parse_date(datetime.now().strftime("%d/%m/%Y"), end_of_day=True)

    company = get_company_name(COMPANY_ID) or "Providencia"
    precio = get_water_price_per_m3(COMPANY_ID, ALL_NODES[0], None)

    out_dir = output_dir or (
        ROOT / "reports" / "Providencia" / "analisis_ahorro_desde_entrega"
    )
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")

    filas: List[dict] = []
    print(f"[INFO] {company} | rango búsqueda primera data: {desde} → {_naive(end_dt).date()}")

    for nid in ALL_NODES:
        nombre = get_node_name(nid) or nid
        primera, err = _primera_fecha(nid, start_dt, end_dt)
        print(f"  {nid} {nombre}: primera={primera.date() if primera else None} {err}")
        if primera is None:
            filas.append(
                {
                    "node_id": nid,
                    "nombre": nombre,
                    "nombre_corto": nombre.replace("Liceo ", "").replace("Instituto ", "")[:28],
                    "primera": None,
                    "total_m3": 0.0,
                    "dias_con": 0,
                    "by_month": {},
                    "en_agregado": nid in NODOS_AGREGADO,
                    "nota": err or "Sin consumo > 0 en el rango",
                }
            )
            continue
        p0 = _naive(primera)
        p0 = datetime(p0.year, p0.month, p0.day)
        tot, dias, by_month = _consumo_mensual(nid, p0, end_dt)
        nota = ""
        if nid == "000006-03":
            nota = "Excluido habitualmente de agregados; sin datos recientes (corte ~ene 2026)."
        elif nid == "000006-04" and by_month.get("2026-07", 0) > 2000:
            nota = "Consumo jul–ago 2026 anómalo (revisar medidor / fugas / lecturas)."
        filas.append(
            {
                "node_id": nid,
                "nombre": nombre,
                "nombre_corto": nombre.replace("Liceo ", "")[:28],
                "primera": p0,
                "total_m3": tot,
                "dias_con": dias,
                "by_month": by_month,
                "en_agregado": nid in NODOS_AGREGADO,
                "nota": nota,
            }
        )

    # Proyección mes de referencia (julio 2026) — ahorros entregables
    mes_s = parse_date(MES_REF_START)
    mes_e = parse_date(MES_REF_END, end_of_day=True)
    proyecciones: List[dict] = []
    print(f"[INFO] Proyección nocturna {MES_REF_LABEL}...")
    for f in filas:
        nid = f["node_id"]
        m = calculate_nocturnal_metrics(nid, mes_s, mes_e)
        c_noche = float(m["consumo_nocturno_total"])
        d_con = int(m["dias_con_consumo_nocturno"])
        d_sin = int(m["dias_sin_consumo_nocturno"])
        d_datos = int(m["dias_con_datos_horarios"])
        pct = (100.0 * d_con / d_datos) if d_datos else 0.0
        proy_mes, proy_dia, prom_h, cumple = proyeccion_mensual_desde_nocturno(
            c_noche, 31, d_con, d_sin, dias_mes=30, forzar=True
        )
        png_dia = None
        if c_noche > 0 and d_datos > 0:
            try:
                dia_rep, horas_rep, noche_rep = dia_mayor_consumo_nocturno(
                    nid, mes_s, mes_e
                )
                png_dia = out_dir / f"perfil_{nid}_{ts}.png"
                guardar_grafico_horario_24h_app(
                    horas_rep,
                    png_dia,
                    titulo=(
                        f"{f['nombre']} — {dia_rep:%d/%m/%Y} "
                        f"(mayor consumo 00–06 h, {MES_REF_LABEL})"
                    ),
                )
            except Exception as ex:
                print(f"  [WARN] perfil {nid}: {ex}")
                noche_rep = 0.0
                dia_rep = None
        else:
            dia_rep = None
            noche_rep = 0.0

        proyecciones.append(
            {
                **f,
                "noct_mes": c_noche,
                "dias_con_noche": d_con,
                "dias_datos": d_datos,
                "pct_dias": pct,
                "prom_h": prom_h,
                "proy_mes_m3": proy_mes,
                "proy_clp": proy_mes * precio,
                "cumple_umbral": cumple,
                "png_dia": png_dia,
                "dia_rep": dia_rep,
                "noche_dia_rep": noche_rep,
            }
        )

    png_mensual = out_dir / f"consumo_mensual_{ts}.png"
    _grafico_consumo_mensual(filas, png_mensual)
    proy_con_datos = [p for p in proyecciones if p["noct_mes"] > 0]
    png_ahorro = out_dir / f"ahorro_proyeccion_{ts}.png"
    if proy_con_datos:
        _grafico_ahorro_proy(proy_con_datos, png_ahorro)

    # CSV resumen
    csv_path = out_dir / f"resumen_providencia_ahorro_{ts}.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(
            [
                "node_id",
                "colegio",
                "primera_fecha_datos",
                "consumo_desde_entrega_m3",
                "dias_con_consumo",
                "en_agregado_habitual",
                f"nocturno_{MES_REF_LABEL.replace(' ', '_')}_m3",
                "pct_dias_con_nocturno",
                "proyeccion_mensual_control_m3",
                "proyeccion_clp",
                "cumple_umbral_75pct",
                "nota",
            ]
        )
        for p in proyecciones:
            w.writerow(
                [
                    p["node_id"],
                    p["nombre"],
                    p["primera"].strftime("%d/%m/%Y") if p["primera"] else "",
                    f"{p['total_m3']:.2f}".replace(".", ","),
                    p["dias_con"],
                    "si" if p["en_agregado"] else "no",
                    f"{p['noct_mes']:.2f}".replace(".", ","),
                    f"{p['pct_dias']:.1f}".replace(".", ","),
                    f"{p['proy_mes_m3']:.2f}".replace(".", ","),
                    f"{p['proy_clp']:.0f}",
                    "si" if p["cumple_umbral"] else "no",
                    p["nota"],
                ]
            )

    # Word
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading(
        "Colegios Providencia — datos desde entrega y ahorros entregables", 0
    )
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    for lab, txt in [
        ("Cliente", f"{company} (companyId {COMPANY_ID})"),
        (
            "Periodo inventario",
            f"{desde} al {_naive(end_dt).strftime('%d/%m/%Y')}",
        ),
        ("Mes referencia ahorro", MES_REF_LABEL),
        ("Precio ref. agua", f"{_fmt_clp(precio)} / m³"),
        ("Generado", datetime.now().strftime("%d-%m-%Y %H:%M")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{lab}: ").bold = True
        p.add_run(txt)

    # 1. Inventario
    doc.add_heading("1. Inventario de colegios y fecha de entrega de datos", level=1)
    doc.add_paragraph(
        "Se busca la primera fecha con consumo diario > 0 desde el inicio del rango "
        "(proxy de “datos entregados / medidor operativo”). Hay 5 nodos en la API; "
        "los agregados mensuales habituales usan 4 (excluyen Arturo Alessandri Palma)."
    )
    tbl = doc.add_table(rows=1 + len(filas), cols=6)
    tbl.style = "Table Grid"
    for j, hd in enumerate(
        [
            "Node ID",
            "Colegio",
            "Primera data",
            "Consumo desde entrega (m³)",
            "Días c/ consumo",
            "En agregado",
        ]
    ):
        tbl.rows[0].cells[j].text = hd
        for run in tbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, f in enumerate(filas, 1):
        cells = tbl.rows[i].cells
        cells[0].text = f["node_id"]
        cells[1].text = f["nombre"]
        cells[2].text = (
            f["primera"].strftime("%d/%m/%Y") if f["primera"] else "—"
        )
        cells[3].text = _fmt_m3(f["total_m3"], 1)
        cells[4].text = str(f["dias_con"])
        cells[5].text = "Sí" if f["en_agregado"] else "No*"

    total_m3 = sum(f["total_m3"] for f in filas)
    p = doc.add_paragraph()
    p.add_run("Total consumo medido desde entrega (5 colegios): ").bold = True
    p.add_run(f"{_fmt_m3(total_m3, 1)} m³.")
    doc.add_paragraph(
        "* 000006-03 Arturo Alessandri Palma: excluido por criterio operativo de "
        "reportes agregados; deja de registrar datos hacia enero 2026."
    )

    if png_mensual.is_file():
        doc.add_paragraph()
        doc.add_picture(str(png_mensual), width=Inches(6.2))

    # 2. Qué ahorros podemos entregar
    doc.add_heading("2. Qué datos de ahorro podemos entregar", level=1)
    doc.add_paragraph(
        "Con la serie horaria/diaria WES de Providencia se pueden armar, de forma "
        "recurrente, estos productos de ahorro (sin depender de un baseline “sin WES” "
        "previo a la instalación, porque la data parte con el medidor ya entregado):"
    )

    doc.add_heading("2.1 Consumo nocturno medido (00:00–06:59)", level=2)
    doc.add_paragraph(
        "Volumen real en madrugada por colegio y por periodo. Es el dato más sólido "
        "para el cliente: no es proyección. Sirve para reportes mensuales, ranking "
        "entre liceos y seguimiento de fugas / válvulas abiertas fuera de horario."
    )

    doc.add_heading("2.2 Proyección de ahorro con equipo de control (filtración)", level=2)
    doc.add_paragraph(
        "A partir del consumo nocturno: promedio horario nocturno × 24 h × 30 días. "
        f"Umbral habitual WES: ≥{UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION:.0f} % de "
        "días con consumo en madrugada. Si no se cumple el umbral, igual se puede "
        "mostrar como estimación técnica. Valor en CLP con tarifa de referencia del "
        "nodo. Ideal para dimensionar beneficio de corte/regulación nocturna."
    )

    doc.add_heading("2.3 Propuestas de regulación diurna", level=2)
    doc.add_paragraph(
        "Ya existen análisis de regulación para Liceo Lastarria (000006-01) y "
        "Carmela Carvajal (000006-02) en `calculo de regulaciones/` "
        "(Excel detalle + PDF propuestas). Muestran escenarios de caudal reducido "
        "en horario hábil / inhábil. Se puede replicar el mismo formato para "
        "Liceo 7 y Juan Pablo Duarte cuando se priorice."
    )

    doc.add_heading("2.4 Reportes agregados mensuales", level=2)
    doc.add_paragraph(
        "Word/PDF agregado de los 4 colegios operativos (consumo total, nocturno, "
        "alertas, perfiles horarios). Ya se generan de forma recurrente bajo "
        "reports/Providencia/ABREGADO."
    )

    doc.add_heading("2.5 Lo que NO tenemos (o es débil) hoy", level=2)
    doc.add_paragraph(
        "• Baseline “antes de WES” facturado vs medido: la serie API arranca con la "
        "entrega del medidor (oct 2025); no hay comparación con/sin WES como en "
        "Puente Alto (cortes ON/OFF).\n"
        "• Ahorro auditado tipo ICCO (semana referencia vs semana intervención) "
        "salvo que se programe una auditoría en terreno.\n"
        "• Arturo Alessandri Palma: sin datos recientes → no hay ahorro entregable "
        "hasta reactivar el punto.\n"
        "• Liceo 7 (jul–ago 2026): consumo anormalmente alto; cualquier cifra de "
        "ahorro ahí debe validarse antes de presentarla al cliente."
    )

    # 3. Muestra julio
    doc.add_heading(
        f"3. Muestra de ahorro entregable — {MES_REF_LABEL}", level=1
    )
    doc.add_paragraph(
        f"Cálculo con serie horaria de {MES_REF_LABEL}. La proyección mensual "
        "cuantifica el orden de magnitud si el patrón nocturno se controlara "
        "(escenario comercial de equipo de control)."
    )
    tbl2 = doc.add_table(rows=1 + len(proyecciones), cols=7)
    tbl2.style = "Table Grid"
    for j, hd in enumerate(
        [
            "Colegio",
            "Nocturno (m³)",
            "Días c/ noche",
            "% días",
            "Prom. h noct.",
            "Proy. 30 d (m³)",
            "Valor CLP",
        ]
    ):
        tbl2.rows[0].cells[j].text = hd
        for run in tbl2.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, p in enumerate(proyecciones, 1):
        row = tbl2.rows[i].cells
        row[0].text = p["nombre"]
        row[1].text = _fmt_m3(p["noct_mes"], 1)
        row[2].text = f"{p['dias_con_noche']}/{p['dias_datos']}"
        row[3].text = f"{_fmt_m3(p['pct_dias'], 0)} %"
        row[4].text = _fmt_m3(p["prom_h"], 3)
        row[5].text = _fmt_m3(p["proy_mes_m3"], 1)
        row[6].text = _fmt_clp(p["proy_clp"])

    tot_proy = sum(p["proy_mes_m3"] for p in proyecciones)
    tot_clp = tot_proy * precio
    tot_proy_ok = sum(
        p["proy_mes_m3"]
        for p in proyecciones
        if p["en_agregado"] and p["cumple_umbral"]
    )
    p = doc.add_paragraph()
    p.add_run(f"Total proyección 5 colegios ({MES_REF_LABEL}): ").bold = True
    p.add_run(f"{_fmt_m3(tot_proy, 1)} m³ — {_fmt_clp(tot_clp)}.")
    p2 = doc.add_paragraph()
    p2.add_run(
        "Solo nodos en agregado que cumplen umbral 75 % días con nocturno: "
    ).bold = True
    p2.add_run(
        f"{_fmt_m3(tot_proy_ok, 1)} m³ — {_fmt_clp(tot_proy_ok * precio)} "
        "(cifra más defendible para entregar al cliente)."
    )

    if png_ahorro.is_file() and proy_con_datos:
        doc.add_paragraph()
        doc.add_picture(str(png_ahorro), width=Inches(5.8))

    for p in proyecciones:
        if not p.get("png_dia") or not Path(p["png_dia"]).is_file():
            continue
        doc.add_heading(p["nombre"], level=2)
        if p["dia_rep"]:
            doc.add_paragraph(
                f"Día con mayor consumo 00–06 h en {MES_REF_LABEL}: "
                f"{p['dia_rep'].strftime('%d/%m/%Y')} "
                f"({_fmt_m3(p['noche_dia_rep'], 2)} m³ en madrugada)."
            )
        if not p["cumple_umbral"] and p["noct_mes"] > 0:
            doc.add_paragraph(
                f"Nota: no alcanza el umbral del "
                f"{UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION:.0f} % de días; "
                "la proyección es estimación técnica."
            )
        if p["nota"]:
            doc.add_paragraph(f"Observación: {p['nota']}")
        doc.add_picture(str(p["png_dia"]), width=Inches(5.8))

    # 4. Recomendación de entregables
    doc.add_heading("4. Recomendación: paquete de ahorros a entregar", level=1)
    doc.add_paragraph(
        "1) Informe mensual agregado (4 colegios) con tabla de consumo nocturno "
        "medido + proyección de control donde cumpla umbral.\n"
        "2) Ficha de ahorro potencial (este formato) actualizada cada mes.\n"
        "3) PDF de regulación diurna para Lastarria y Carmela (ya disponibles); "
        "extender a Duarte y Liceo 7 tras validar datos.\n"
        "4) Alerta operativa sobre Liceo 7 (consumo jul–ago) antes de citar "
        "ahorros en CLP a la corporación.\n"
        "5) Reactivar / diagnosticar 000006-03 si se quiere cobertura de “todos” "
        "los colegios en el mensaje comercial."
    )

    doc.add_heading("5. Conclusión", level=1)
    doc.add_paragraph(
        "Desde la entrega de data (1–6 oct 2025) hay serie continua en 4 colegios "
        "y parcial en Arturo Alessandri. Los ahorros que sí se pueden entregar con "
        "confianza son: (a) nocturno medido, (b) proyección de control nocturno "
        "donde el patrón es recurrente (≥75 % días), (c) propuestas de regulación "
        "diurna ya armadas para Lastarria y Carmela. No hay baseline pre-WES; el "
        "valor a comunicar es potencial de reducción sobre el consumo actual "
        "fuera de horario / con regulación, no un “ahorro histórico auditado”."
    )

    out_docx = out_dir / f"Reporte_Ahorro_Colegios_Providencia_{ts}.docx"
    doc.save(out_docx)

    pdf_path = out_docx.with_suffix(".pdf")
    try:
        from generar_reporte_word import convertir_word_a_pdf

        converted = convertir_word_a_pdf(out_docx)
        if converted and Path(converted).is_file():
            pdf_path = Path(converted)
    except Exception as ex:
        print(f"[WARN] PDF: {ex}")

    print(f"[OK] Word: {out_docx}")
    print(f"[OK] CSV:  {csv_path}")
    print(f"[OK] PDF:  {pdf_path if pdf_path.is_file() else 'no generado'}")
    print(f"[OK] Proyección total {MES_REF_LABEL}: {tot_proy:.1f} m³")
    return out_docx


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Reporte ahorros colegios Providencia desde entrega de datos"
    )
    ap.add_argument("--desde", default="01/10/2025")
    ap.add_argument("--hasta", default="")
    ap.add_argument("--output-dir", default="")
    args = ap.parse_args()
    out = Path(args.output_dir) if args.output_dir else None
    print("=" * 72)
    print("REPORTE AHORRO — COLEGIOS PROVIDENCIA DESDE ENTREGA DE DATOS")
    print("=" * 72)
    generar(desde=args.desde, hasta=args.hasta, output_dir=out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
