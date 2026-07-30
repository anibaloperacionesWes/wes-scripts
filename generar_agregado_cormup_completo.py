"""
Reporte agregado CORMUP (Peñalolén) — formato solicitado:

  - Consumo total ordenado de mayor a menor (por colegio)
  - Desagregado por colegio (curva diaria)
  - Consumo nocturno 00:00–06:00 sumando todos los colegios
  - Comparativo últimos 6 meses (total = suma de todos los colegios)

Uso:
  python generar_agregado_cormup_completo.py
  python generar_agregado_cormup_completo.py --start-date 01/07/2026 --end-date 30/07/2026
"""

from __future__ import annotations

import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

if sys.platform == "win32":
    for s in (sys.stdout, sys.stderr):
        try:
            s.reconfigure(encoding="utf-8", line_buffering=True)
        except Exception:
            pass

from generar_consolidado_m3_mensual_puente_alto import consumo_mes_un_nodo
from generar_reporte_word import (
    _dt_to_chile,
    _parse_alert_creation_date,
    acl_node_base_url,
    add_formatted_heading,
    add_formatted_title,
    add_logo_to_header,
    add_picture_with_pagination,
    alerta_medida_informativa,
    build_consumption_chart,
    calculate_nocturnal_metrics,
    estilizar_tabla_wes,
    fetch_json,
    flatten_measures,
    format_currency_chilean,
    format_number_chilean,
    get_company_name,
    get_node_name,
    get_water_price_per_m3,
    normalize_measures_payload,
    parse_date,
    summarize_consumption,
)

COMPANY_ID = "000008"
COMPANY_FOLDER = "CORMUP"
NODE_IDS = [f"000008-{i:02d}" for i in range(1, 15)]
PRECIO_DEFAULT = 1200.0
COLOR_BARRA = "#0050b3"
COLOR_NOCTURNO = "#FFD700"
COLOR_NOCTURNO_EDGE = "#DAA520"
HORA_ALERTA_INI = 0
HORA_ALERTA_FIN = 6  # inclusive → 00:00–06:59


def _fmt_api(dt: datetime) -> str:
    return dt.strftime("%d%m%Y")


def _alerta_en_ventana_nocturna(alert: dict) -> bool:
    """True si la alerta cayó entre 00:00 y 06:59 hora Chile."""
    dt = _parse_alert_creation_date(alert)
    if not dt:
        return False
    h = _dt_to_chile(dt).hour
    return HORA_ALERTA_INI <= h <= HORA_ALERTA_FIN


def _fetch_alerts(node_id: str, start_dt: datetime, end_dt: datetime) -> List[dict]:
    try:
        raw = fetch_json(
            f"{acl_node_base_url()}/nodes/myalert/alerts",
            params=[
                ("id", node_id),
                ("start", _fmt_api(start_dt)),
                ("end", _fmt_api(end_dt)),
            ],
        )
        if not isinstance(raw, list):
            return []
        return [a for a in raw if isinstance(a, dict) and alerta_medida_informativa(a)]
    except Exception as e:
        print(f"  [WARN] alertas {node_id}: {e}", flush=True)
        return []


def _fetch_node(node_id: str, start_dt: datetime, end_dt: datetime) -> dict:
    name = get_node_name(node_id)
    print(f"  [data] {node_id} {name}...", flush=True)
    try:
        raw = fetch_json(
            f"{acl_node_base_url()}/nodes/measures/dates",
            params=[
                ("id", node_id),
                ("start", _fmt_api(start_dt)),
                ("end", _fmt_api(end_dt)),
            ],
        )
        payload = normalize_measures_payload(raw, node_id)
        measures = flatten_measures(payload)
        summary = summarize_consumption(measures)
        noct = calculate_nocturnal_metrics(
            node_id, start_dt, end_dt, company_id=COMPANY_ID
        )
        price = get_water_price_per_m3(COMPANY_ID, node_id, payload)
        alerts = _fetch_alerts(node_id, start_dt, end_dt)
        alerts_noct = [a for a in alerts if _alerta_en_ventana_nocturna(a)]
        return {
            "node_id": node_id,
            "node_name": name,
            "summary": summary,
            "measures": measures,
            "nocturno_m3": float(noct.get("consumo_nocturno_total") or 0.0),
            "diurno_m3": float(noct.get("consumo_diurno_efectivo") or 0.0),
            "dias_nocturno": int(noct.get("dias_con_consumo_nocturno") or 0),
            "price": price if price else PRECIO_DEFAULT,
            "alerts_noct": alerts_noct,
            "alerts_total": len(alerts),
        }
    except Exception as e:
        print(f"  [WARN] {node_id}: {e}", flush=True)
        return {
            "node_id": node_id,
            "node_name": name,
            "summary": summarize_consumption([]),
            "measures": [],
            "nocturno_m3": 0.0,
            "diurno_m3": 0.0,
            "dias_nocturno": 0,
            "price": PRECIO_DEFAULT,
            "alerts_noct": [],
            "alerts_total": 0,
            "error": str(e),
        }


def _plot_barras_periodo(nodes: List[dict], out: Path) -> Path:
    pairs = sorted(
        ((n["node_name"], float(n["summary"]["total"])) for n in nodes),
        key=lambda x: -x[1],
    )
    labels = [p[0] for p in pairs]
    vals = [p[1] for p in pairs]
    fig, ax = plt.subplots(figsize=(13, 6.5))
    bars = ax.bar(labels, vals, color=COLOR_BARRA)
    ax.set_ylabel("Consumo (m³)", fontsize=12, fontweight="bold")
    ax.set_title(
        "Consumo total del periodo por colegio — CORMUP (mayor → menor)",
        fontsize=13,
        fontweight="bold",
    )
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=40, ha="right", fontsize=8)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 0)}",
            ha="center",
            va="bottom",
            fontsize=8,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _plot_nocturno_amarillo(nodes: List[dict], out: Path) -> Path:
    pairs = sorted(
        ((n["node_name"], float(n["nocturno_m3"])) for n in nodes),
        key=lambda x: -x[1],
    )
    labels = [p[0] for p in pairs]
    vals = [p[1] for p in pairs]
    fig, ax = plt.subplots(figsize=(13, 6))
    bars = ax.bar(
        labels,
        vals,
        color=COLOR_NOCTURNO,
        alpha=0.9,
        edgecolor=COLOR_NOCTURNO_EDGE,
        linewidth=1.2,
    )
    ax.set_ylabel("Consumo nocturno (m³)", fontsize=12, fontweight="bold")
    ax.set_title(
        "Consumo nocturno por colegio (00:00–06:00) — orden mayor → menor",
        fontsize=13,
        fontweight="bold",
    )
    ax.set_ylim(bottom=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=40, ha="right", fontsize=8)
    for bar, v in zip(bars, vals):
        if v > 0:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height(),
                f"{format_number_chilean(v, 0)}",
                ha="center",
                va="bottom",
                fontsize=8,
                fontweight="bold",
            )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _meses_ultimos_6(end_d: date) -> List[Tuple[int, int]]:
    y, m = end_d.year, end_d.month
    out: List[Tuple[int, int]] = []
    for _ in range(6):
        out.append((y, m))
        m -= 1
        if m == 0:
            m = 12
            y -= 1
    out.reverse()
    return out


def _plot_6_meses(series: List[Tuple[str, float]], out: Path) -> Path:
    labels = [s[0] for s in series]
    vals = [s[1] for s in series]
    fig, ax = plt.subplots(figsize=(10, 5.5))
    bars = ax.bar(labels, vals, color=COLOR_BARRA, width=0.65)
    ax.set_ylabel("Consumo mensual total (m³)", fontsize=11, fontweight="bold")
    ax.set_xlabel("Mes", fontsize=11, fontweight="bold")
    ax.set_title(
        "Comparativo últimos 6 meses — CORMUP (suma todos los colegios)",
        fontsize=12,
        fontweight="bold",
    )
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 0)}",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _add_table(
    doc: Document,
    headers: List[str],
    rows: List[List[str]],
    *,
    has_total_row: bool = False,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            table.rows[r_i + 1].cells[c_i].text = str(val)
    estilizar_tabla_wes(table, has_total_row=has_total_row)
    doc.add_paragraph("")


def _mes_total_colegios(
    node_ids: List[str], year: int, month: int, workers: int = 8
) -> float:
    total = 0.0
    with ThreadPoolExecutor(max_workers=min(workers, len(node_ids))) as ex:
        futs = {
            ex.submit(consumo_mes_un_nodo, requests.Session(), nid, year, month): nid
            for nid in node_ids
        }
        for fut in as_completed(futs):
            v, _, _ = fut.result()
            total += float(v)
    return total


def generar(start_date: str, end_date: str, workers: int = 10) -> Path:
    start_dt = parse_date(start_date)
    end_dt = parse_date(end_date, end_of_day=True)
    company_name = get_company_name(COMPANY_ID)
    node_ids = list(NODE_IDS)

    print("=" * 70)
    print(f"AGREGADO CORMUP COMPLETO — {company_name}")
    print(f"Periodo: {start_date} → {end_date}")
    print(f"Colegios ({len(node_ids)}): {', '.join(node_ids)}")
    print("=" * 70)

    t0 = time.perf_counter()
    nodes: List[dict] = []
    w = max(1, min(workers, len(node_ids)))
    with ThreadPoolExecutor(max_workers=w) as ex:
        futs = {ex.submit(_fetch_node, nid, start_dt, end_dt): nid for nid in node_ids}
        for fut in as_completed(futs):
            nodes.append(fut.result())

    # Orden mayor → menor consumo
    nodes_by_cons = sorted(nodes, key=lambda n: -float(n["summary"]["total"]))
    total_m3 = sum(float(n["summary"]["total"]) for n in nodes)
    noct_total = sum(float(n["nocturno_m3"]) for n in nodes)
    diurno_total = sum(float(n["diurno_m3"]) for n in nodes)
    prices = [float(n["price"]) for n in nodes if n.get("price")]
    price = (sum(prices) / len(prices)) if prices else PRECIO_DEFAULT
    num_dias = (end_dt.date() - start_dt.date()).days + 1

    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    out_dir = Path("reports") / COMPANY_FOLDER / "ABREGADO" / f"AGREGADO_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    chart_barras = _plot_barras_periodo(nodes, out_dir / "chart_consumo_periodo.png")
    chart_nocturno = _plot_nocturno_amarillo(
        nodes, out_dir / "chart_consumo_nocturno_colegios.png"
    )

    print("[INFO] Descargando últimos 6 meses (suma todos los colegios)...", flush=True)
    series_6: List[Tuple[str, float]] = []
    for y, m in _meses_ultimos_6(end_dt.date()):
        if y == end_dt.year and m == end_dt.month:
            m3 = total_m3
            label = f"{y}-{m:02d}*"
        else:
            m3 = _mes_total_colegios(node_ids, y, m, workers=w)
            label = f"{y}-{m:02d}"
        series_6.append((label, float(m3)))
        print(f"  {label}: {m3:.1f} m³", flush=True)
    chart_6m = _plot_6_meses(series_6, out_dir / "chart_ultimos_6_meses.png")

    doc = Document()
    add_logo_to_header(doc)
    title = doc.add_paragraph(f"Reporte Agregado — {company_name} (Peñalolén)")
    title.style = "Title"
    for run in title.runs:
        run.font.size = Pt(22)
    sub = doc.add_paragraph(
        f"MONITOREO WES\n"
        f"Análisis consolidado de {len(nodes)} colegios\n"
        f"Rango: {start_dt.strftime('%d-%m-%y')} – {end_dt.strftime('%d-%m-%y')}\n"
        f"Generado: {datetime.now().strftime('%d-%m-%y')}"
    )
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 1. Consumo total (ordenado mayor → menor)
    add_formatted_heading(doc, "1. Consumo total (ordenado de mayor a menor)", level=1)
    p = doc.add_paragraph(
        f"CORMUP registró un consumo total de {format_number_chilean(total_m3, 1)} m³ "
        f"en el periodo, sumando los {len(nodes)} colegios. "
        f"A continuación se presenta el ranking de mayor a menor consumo."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_barras), Inches(6.2), keep_with_next=True)

    rows_tot = []
    for i, n in enumerate(nodes_by_cons, start=1):
        m3 = float(n["summary"]["total"])
        pct = (m3 / total_m3 * 100.0) if total_m3 > 0 else 0.0
        rows_tot.append(
            [
                str(i),
                n["node_id"],
                n["node_name"],
                format_number_chilean(m3, 1),
                format_number_chilean(pct, 1) + " %",
                format_currency_chilean(m3 * price),
            ]
        )
    rows_tot.append(
        [
            "",
            "",
            "Total CORMUP",
            format_number_chilean(total_m3, 1),
            "100 %",
            format_currency_chilean(total_m3 * price),
        ]
    )
    _add_table(
        doc,
        ["#", "Nodo", "Colegio", "Consumo (m³)", "%", "Costo (CLP)"],
        rows_tot,
        has_total_row=True,
    )

    # 2. Desagregado por colegio
    add_formatted_heading(doc, "2. Desagregado por colegio", level=1)
    p_d = doc.add_paragraph(
        "Evolución del consumo diario (m³) de cada colegio en el periodo. "
        "En el gráfico solo se marcan alertas ocurridas entre las 00:00 y las 06:00 "
        "(hora Chile)."
    )
    p_d.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for n in nodes_by_cons:
        measures = n.get("measures") or []
        if not measures:
            continue
        chart_path = out_dir / f"cormup_diario_{n['node_id'].replace('-', '_')}.png"
        built = build_consumption_chart(
            measures,
            chart_path,
            start_dt,
            end_dt,
            alerts=n.get("alerts_noct") or [],
        )
        if not built or not chart_path.is_file():
            continue
        doc.add_paragraph("")
        add_formatted_title(doc, n["node_name"].upper())
        add_picture_with_pagination(doc, str(chart_path), Inches(6), keep_with_next=True)

    # 3. Nocturno sumado
    add_formatted_heading(doc, "3. Consumo nocturno (00:00 a 06:00) — suma", level=1)
    p2 = doc.add_paragraph(
        "Se considera consumo nocturno el volumen en ventana de madrugada "
        "(criterio colegios CORMUP: CSV UTC 00:00–07:00, equivalente operativo "
        "a 00:00–06:00). Se presenta el total sumado de todos los colegios y el "
        "detalle ordenado de mayor a menor."
    )
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2b = doc.add_paragraph(
        f"Nocturno total (suma colegios): {format_number_chilean(noct_total, 1)} m³ "
        f"({format_currency_chilean(noct_total * price)}) en {num_dias} días. "
        f"Diurno total: {format_number_chilean(diurno_total, 1)} m³."
    )
    p2b.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_nocturno), Inches(6.2), keep_with_next=True)

    rows_noct = []
    for n in sorted(nodes, key=lambda x: -x["nocturno_m3"]):
        rows_noct.append(
            [
                n["node_id"],
                n["node_name"],
                format_number_chilean(n["nocturno_m3"], 1),
                str(n["dias_nocturno"]),
                format_currency_chilean(n["nocturno_m3"] * price),
            ]
        )
    rows_noct.append(
        [
            "",
            "Total CORMUP",
            format_number_chilean(noct_total, 1),
            "",
            format_currency_chilean(noct_total * price),
        ]
    )
    _add_table(
        doc,
        ["Nodo", "Colegio", "Nocturno (m³)", "Días con nocturno", "Costo (CLP)"],
        rows_noct,
        has_total_row=True,
    )

    # 4. Alertas solo 00:00–06:00
    add_formatted_heading(
        doc, "4. Alertas de consumo (solo 00:00 a 06:00)", level=1
    )
    alert_rows: List[List[str]] = []
    for n in nodes_by_cons:
        for a in n.get("alerts_noct") or []:
            dt = _parse_alert_creation_date(a)
            if not dt:
                continue
            dt_cl = _dt_to_chile(dt)
            alert_rows.append(
                [
                    n["node_id"],
                    n["node_name"],
                    dt_cl.strftime("%d/%m/%Y"),
                    dt_cl.strftime("%H:%M"),
                    format_number_chilean(float(a.get("measure", 0) or 0), 2),
                ]
            )
    p4 = doc.add_paragraph(
        "Se listan únicamente las alertas cuya hora de registro (Chile) está entre "
        "las 00:00 y las 06:59. Las alertas diurnas se omiten en este informe."
    )
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if alert_rows:
        p4b = doc.add_paragraph(
            f"Total alertas nocturnas en el periodo: {len(alert_rows)} "
            f"(de {sum(int(n.get('alerts_total') or 0) for n in nodes)} alertas "
            f"informativas en todos los horarios)."
        )
        p4b.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        alert_rows.sort(key=lambda r: (r[2], r[3]), reverse=True)
        _add_table(
            doc,
            ["Nodo", "Colegio", "Día", "Hora", "Medida (m³/h)"],
            alert_rows,
            has_total_row=False,
        )
    else:
        doc.add_paragraph(
            "No se registraron alertas informativas entre las 00:00 y las 06:59 "
            "en el periodo analizado."
        )

    # 5. Últimos 6 meses
    add_formatted_heading(doc, "5. Comparativo últimos 6 meses (total)", level=1)
    p6 = doc.add_paragraph(
        "Consumo mensual total sumando todos los colegios CORMUP. "
        "El mes marcado con * corresponde al periodo de este reporte "
        f"(hasta {end_dt.strftime('%d/%m/%Y')})."
    )
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_6m), Inches(6), keep_with_next=True)
    rows_6 = [[lab, format_number_chilean(v, 1)] for lab, v in series_6]
    _add_table(doc, ["Mes", "Total colegios (m³)"], rows_6, has_total_row=False)

    add_formatted_heading(doc, "Conclusión", level=1)
    top = nodes_by_cons[0] if nodes_by_cons else None
    top_txt = (
        f"El colegio de mayor consumo fue {top['node_name']} "
        f"({format_number_chilean(float(top['summary']['total']), 1)} m³). "
        if top
        else ""
    )
    concl = doc.add_paragraph(
        f"En el periodo, CORMUP registró {format_number_chilean(total_m3, 1)} m³. "
        f"{top_txt}"
        f"El consumo nocturno agregado fue {format_number_chilean(noct_total, 1)} m³. "
        f"Las alertas consideradas en este informe se limitan a la ventana "
        f"00:00–06:00."
    )
    concl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    fname = f"Reporte_Agregado_CORMUP_{start_dt:%Y%m%d}_{end_dt:%Y%m%d}.docx"
    out_path = out_dir / fname
    doc.save(out_path)
    print(f"[OK] {out_path}")
    print(f"[INFO] Tiempo: {time.perf_counter() - t0:.1f} s")
    return out_path


def main() -> int:
    import argparse

    ap = argparse.ArgumentParser(description="Agregado CORMUP completo (julio)")
    ap.add_argument("--start-date", default="01/07/2026")
    ap.add_argument("--end-date", default="30/07/2026")
    ap.add_argument("--workers", type=int, default=10)
    args = ap.parse_args()
    generar(args.start_date, args.end_date, workers=args.workers)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
