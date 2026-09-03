"""
Reporte de puntos con horas perdidas (huecos en la serie horaria),
ordenado de mayor a menor, con gráficos por semana ISO.

Periodo por defecto: lunes de la semana 32 de 2026 hasta el día de hoy (Chile).

Hora perdida = hora Chile esperada sin registro en dates.measures.csv.
Una hora con valor 0 SÍ cuenta como dato (no es hora perdida).

El ranking final lista puntos que hoy pierden data por desconexión
(lastUpdate de la app ≥ 2 h).
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from matplotlib.patches import Patch

from reporte_puntos_en_cero import (
    HORAS_UMBRAL_CONEXION_APP,
    _fmt_antiguedad,
    obtener_estado_conexion_nodo,
    obtener_todos_los_nodos,
)
from wes_paths import wes_scripts_root

try:
    from zoneinfo import ZoneInfo

    CHILE_TZ = ZoneInfo("America/Santiago")
except Exception:
    CHILE_TZ = timezone(timedelta(hours=-4))

BASE_URL = os.environ.get(
    "WES_API_BASE_URL", "http://104.248.53.141:7003/wes/api/acl-node/v1"
).rstrip("/")

MAX_WORKERS = max(4, int(os.environ.get("WES_HORAS_PERDIDAS_WORKERS", "24")))
REQUEST_TIMEOUT = 25

# Semanas ISO desde el lunes 03-08-2026 (S32) hasta hoy (se arma en configurar_periodo).
INICIO_PERIODO = date(2026, 8, 3)
_MESES_ES = {
    1: "ene",
    2: "feb",
    3: "mar",
    4: "abr",
    5: "may",
    6: "jun",
    7: "jul",
    8: "ago",
    9: "sep",
    10: "oct",
    11: "nov",
    12: "dic",
}
COLORES_SEMANA = ["#1F4E79", "#2E86AB", "#E67E22", "#C0392B", "#7D3C98", "#16A085", "#8E44AD"]
COLOR_S32 = COLORES_SEMANA[0]
COLOR_S33 = COLORES_SEMANA[1]
COLOR_S34 = COLORES_SEMANA[2]
COLOR_S35 = COLORES_SEMANA[3]
COLOR_DESC = "#C0392B"
COLOR_HUECOS = "#2E86AB"

SEMANA_DEFS: List[Tuple[int, date, date, str]] = []
TODOS_LOS_DIAS: List[date] = []
SEMANA_POR_DIA: Dict[date, int] = {}


def _iter_days(inicio: date, fin: date) -> List[date]:
    out: List[date] = []
    cur = inicio
    while cur <= fin:
        out.append(cur)
        cur += timedelta(days=1)
    return out


def _fmt_rango_corto(ini: date, fin: date) -> str:
    if ini.month == fin.month:
        return f"{ini.day:02d}–{fin.day:02d} {_MESES_ES[fin.month]}"
    return f"{ini.day:02d} {_MESES_ES[ini.month]}–{fin.day:02d} {_MESES_ES[fin.month]}"


def configurar_periodo(hoy: date) -> None:
    """Arma semanas ISO desde INICIO_PERIODO hasta ``hoy`` (última semana parcial)."""
    global SEMANA_DEFS, TODOS_LOS_DIAS, SEMANA_POR_DIA
    defs: List[Tuple[int, date, date, str]] = []
    lun = INICIO_PERIODO - timedelta(days=INICIO_PERIODO.weekday())
    while lun <= hoy:
        week = lun.isocalendar()[1]
        domingo = lun + timedelta(days=6)
        fin = min(domingo, hoy)
        parcial = fin < domingo
        extra = ", parcial" if parcial else ""
        label = f"Semana {week} ({_fmt_rango_corto(lun, fin)}{extra})"
        defs.append((week, lun, fin, label))
        lun += timedelta(days=7)
    SEMANA_DEFS = defs
    TODOS_LOS_DIAS = _iter_days(INICIO_PERIODO, hoy)
    SEMANA_POR_DIA = {}
    for num, ini, fin, _lbl in defs:
        for d in _iter_days(ini, fin):
            SEMANA_POR_DIA[d] = num


def _semanas_completas() -> List[Tuple[int, date, date, str]]:
    return [t for t in SEMANA_DEFS if (t[2] - t[1]).days == 6]


try:
    configurar_periodo(datetime.now(CHILE_TZ).date())
except Exception:
    configurar_periodo(date.today())
_SESSION = requests.Session()
_ADAPTER = requests.adapters.HTTPAdapter(pool_connections=32, pool_maxsize=32, max_retries=2)
_SESSION.mount("http://", _ADAPTER)
_SESSION.mount("https://", _ADAPTER)


def _fmt_int(n: float | int) -> str:
    return f"{int(round(n)):,}".replace(",", ".")


def _fmt_pct(part: float, total: float) -> str:
    if total <= 0:
        return "—"
    return f"{(part / total) * 100:.1f}".replace(".", ",") + "%"


def _trunc(texto: str, n: int = 36) -> str:
    t = (texto or "").strip()
    return t if len(t) <= n else t[: n - 1] + "…"


def horas_esperadas_dia(dia: date, ahora_chile: datetime) -> List[int]:
    """Horas 0–23 que ya deberían tener registro a la hora del reporte."""
    hoy = ahora_chile.date()
    if dia < hoy:
        return list(range(24))
    if dia > hoy:
        return []
    # La marca HH:00 del día en curso aparece cerca de esa hora; no exigir la hora actual.
    corte = max(0, ahora_chile.hour)
    return list(range(corte))


def _parse_horas_csv(csv_content: str, dia: date) -> Dict[int, float]:
    """Hora civil del TIME (convención app WES) para un día."""
    target = dia.strftime("%Y-%m-%d")
    acc: Dict[int, float] = {}
    lines = (csv_content or "").strip().splitlines()
    if len(lines) <= 1:
        return acc
    for line in lines[1:]:
        line = line.strip()
        if not line or "," not in line:
            continue
        time_str, value_str = line.split(",", 1)
        ts = time_str.strip()
        if not ts.startswith(target) or "T" not in ts:
            continue
        try:
            hi = int(ts[11:13])
            val = float(value_str.strip().replace(" ", "").replace(",", "."))
        except (ValueError, TypeError, IndexError):
            continue
        if 0 <= hi < 24:
            acc[hi] = acc.get(hi, 0.0) + val
    return acc


def _parse_dias_rango_csv(csv_content: str) -> Dict[date, float]:
    """CSV de rango (una fila por día, TIME=YYYY-MM-DD)."""
    out: Dict[date, float] = {}
    lines = (csv_content or "").strip().splitlines()
    if len(lines) <= 1:
        return out
    for line in lines[1:]:
        line = line.strip()
        if not line or "," not in line:
            continue
        time_str, value_str = line.split(",", 1)
        ts = time_str.strip()
        if "T" in ts:
            ts = ts.split("T", 1)[0]
        try:
            dia = date.fromisoformat(ts[:10])
            val = float(value_str.strip().replace(" ", "").replace(",", "."))
        except (ValueError, TypeError):
            continue
        out[dia] = out.get(dia, 0.0) + val
    return out


def _get_csv(node_id: str, start: date, end: date) -> Tuple[str, bool]:
    url = f"{BASE_URL}/nodes/{node_id}/dates.measures.csv"
    params = [("start", start.strftime("%d%m%Y")), ("end", end.strftime("%d%m%Y"))]
    try:
        r = _SESSION.get(url, params=params, timeout=REQUEST_TIMEOUT)
        if r.status_code != 200:
            return "", False
        return r.text, True
    except requests.RequestException:
        return "", False


@dataclass
class ResultadoNodo:
    node_id: str
    node_name: str
    company_id: str
    company_name: str
    perdidas_por_dia: Dict[str, int] = field(default_factory=dict)
    esperadas_por_dia: Dict[str, int] = field(default_factory=dict)
    error: str = ""
    last_update: str = ""
    wes_status: str = ""
    last_update_dt: Optional[datetime] = None
    horas_sin_conexion: Optional[float] = None
    desconectado: bool = False

    def perdidas_dia(self, dia: date) -> int:
        return int(self.perdidas_por_dia.get(dia.isoformat(), 0))

    def esperadas_dia(self, dia: date) -> int:
        return int(self.esperadas_por_dia.get(dia.isoformat(), 0))

    def perdidas_semana(self, semana: int) -> int:
        return sum(
            self.perdidas_dia(d) for d, s in SEMANA_POR_DIA.items() if s == semana
        )

    def esperadas_semana(self, semana: int) -> int:
        return sum(
            self.esperadas_dia(d) for d, s in SEMANA_POR_DIA.items() if s == semana
        )

    def perdidas_total(self) -> int:
        return sum(int(v) for v in self.perdidas_por_dia.values())

    def esperadas_total(self) -> int:
        return sum(int(v) for v in self.esperadas_por_dia.values())

    def etiqueta_estado(self) -> str:
        if not self.desconectado:
            return "Conectado (huecos)"
        if self.horas_sin_conexion is None:
            return "Desconectado (sin lastUpdate)"
        return f"Desconectado ({_fmt_antiguedad(self.horas_sin_conexion)})"


def analizar_nodo(nodo: Dict[str, str], ahora_chile: datetime) -> ResultadoNodo:
    res = ResultadoNodo(
        node_id=nodo["nodeId"],
        node_name=nodo.get("nodeName", ""),
        company_id=nodo.get("companyId", ""),
        company_name=nodo.get("companyName", ""),
    )
    texto, ok = _get_csv(res.node_id, TODOS_LOS_DIAS[0], TODOS_LOS_DIAS[-1])
    dias_con_total: Dict[date, float] = _parse_dias_rango_csv(texto) if ok else {}

    for dia in TODOS_LOS_DIAS:
        esperadas = horas_esperadas_dia(dia, ahora_chile)
        res.esperadas_por_dia[dia.isoformat()] = len(esperadas)
        if not esperadas:
            res.perdidas_por_dia[dia.isoformat()] = 0
            continue

        if not ok:
            # Sin respuesta del rango: intentar el día puntual.
            csv_dia, ok_dia = _get_csv(res.node_id, dia, dia)
            if not ok_dia:
                res.perdidas_por_dia[dia.isoformat()] = len(esperadas)
                res.error = "sin respuesta API"
                continue
            presentes = set(_parse_horas_csv(csv_dia, dia).keys())
            res.perdidas_por_dia[dia.isoformat()] = sum(
                1 for h in esperadas if h not in presentes
            )
            continue

        if dia not in dias_con_total:
            res.perdidas_por_dia[dia.isoformat()] = len(esperadas)
            continue

        csv_dia, ok_dia = _get_csv(res.node_id, dia, dia)
        if not ok_dia:
            # El total diario existe; no pudimos ver huecos intra-día → 0 perdidas
            # (no inflar el ranking por un fallo puntual).
            res.perdidas_por_dia[dia.isoformat()] = 0
            continue
        presentes = set(_parse_horas_csv(csv_dia, dia).keys())
        res.perdidas_por_dia[dia.isoformat()] = sum(
            1 for h in esperadas if h not in presentes
        )
    return res


def enriquecer_conexion(resultados: Sequence[ResultadoNodo], ahora_chile: datetime) -> None:
    """Marca desconectado si lastUpdate de la app tiene ≥ umbral (default 2 h)."""

    def _one(r: ResultadoNodo) -> None:
        try:
            st = obtener_estado_conexion_nodo(r.node_id)
        except Exception:
            r.desconectado = True
            r.last_update = "—"
            r.wes_status = "—"
            return
        lu = st.get("lastUpdate")
        r.wes_status = str(st.get("wesStatus") or "—")
        if hasattr(lu, "strftime"):
            r.last_update_dt = lu  # type: ignore[assignment]
            r.last_update = lu.strftime("%d-%m-%Y %H:%M")
        else:
            r.last_update_dt = None
            r.last_update = str(st.get("lastUpdateRaw") or "—")
        if r.last_update_dt is None:
            r.desconectado = True
            r.horas_sin_conexion = None
            return
        dt = r.last_update_dt
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=CHILE_TZ)
        r.horas_sin_conexion = (ahora_chile - dt.astimezone(CHILE_TZ)).total_seconds() / 3600.0
        r.desconectado = r.horas_sin_conexion >= float(HORAS_UMBRAL_CONEXION_APP)

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        list(ex.map(_one, resultados))


def particionar_perdidas(
    items: Sequence[ResultadoNodo],
) -> Tuple[List[ResultadoNodo], List[ResultadoNodo]]:
    """(desconectados, conectados-con-huecos), ambos mayor a menor por total."""
    desc = [r for r in items if r.desconectado]
    huec = [r for r in items if not r.desconectado]
    desc.sort(key=lambda r: r.perdidas_total(), reverse=True)
    huec.sort(key=lambda r: r.perdidas_total(), reverse=True)
    return desc, huec


def cargar_desde_json(path: Path) -> List[ResultadoNodo]:
    raw = json.loads(path.read_text(encoding="utf-8"))
    out: List[ResultadoNodo] = []
    for d in raw:
        r = ResultadoNodo(
            node_id=str(d.get("nodeId", "")),
            node_name=str(d.get("nodeName", "")),
            company_id=str(d.get("companyId", "")),
            company_name=str(d.get("companyName", "")),
            perdidas_por_dia={k: int(v) for k, v in (d.get("perdidasPorDia") or {}).items()},
            esperadas_por_dia={k: int(v) for k, v in (d.get("esperadasPorDia") or {}).items()},
            error=str(d.get("error") or ""),
        )
        out.append(r)
    return out


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:fill="{hex_color}"/>'
    )
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(shading)


def _set_run_white_bold(cell, size: int = 8) -> None:
    for p in cell.paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(size)


def _fill_table(
    table,
    rows: List[Sequence[str]],
    header_fill: str = "1F4E79",
    highlight_col: Optional[int] = None,
) -> None:
    for j, header in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        cell.text = str(header)
        _set_cell_shading(cell, header_fill)
        _set_run_white_bold(cell, 8)
    for i, row in enumerate(rows[1:], start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER if j != 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
                )
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            if i % 2 == 0:
                _set_cell_shading(cell, "F2F2F2")
            if highlight_col is not None and j == highlight_col and i > 0:
                try:
                    n = int(str(val).split()[0].replace(".", ""))
                except ValueError:
                    n = 0
                if n >= 100:
                    _set_cell_shading(cell, "F4CCCC")
                elif n >= 24:
                    _set_cell_shading(cell, "FCE4D6")


def _add_picture(doc: Document, path: Path, width: float = 6.4) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def grafico_totales_semana(
    resultados: Sequence[ResultadoNodo],
    esperados: Dict[int, int],
    out: Path,
) -> Path:
    labels = [f"S{n}\n{_fmt_rango_corto(ini, fin)}" for n, ini, fin, _ in SEMANA_DEFS]
    vals = [sum(r.perdidas_semana(n) for r in resultados) for n, *_ in SEMANA_DEFS]
    colors = [COLORES_SEMANA[i % len(COLORES_SEMANA)] for i in range(len(SEMANA_DEFS))]
    fig, ax = plt.subplots(figsize=(max(8.8, 1.7 * len(labels)), 4.6))
    fig.patch.set_facecolor("white")
    x = range(len(labels))
    bars = ax.bar(x, vals, color=colors, width=0.62, zorder=3)
    ax.set_xticks(list(x), labels)
    ax.set_ylabel("Horas perdidas (flota)")
    nsem = len(SEMANA_DEFS)
    ax.set_title(f"Comparativa de {nsem} semanas — horas sin dato")
    ax.grid(axis="y", linestyle=":", alpha=0.5, zorder=0)
    ax.set_axisbelow(True)
    for bar, (n, *_rest) in zip(bars, SEMANA_DEFS):
        h = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            h,
            f"{_fmt_int(h)}\n({_fmt_pct(h, esperados[n])})",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    ymax = max((max(vals) if vals else 0) * 1.22, 1.0)
    ax.set_ylim(0, ymax)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def grafico_ranking(
    items: Sequence[ResultadoNodo],
    value_fn,
    titulo: str,
    out: Path,
    color: Optional[str] = None,
    top: int = 18,
    leyenda_estado: bool = True,
) -> Path:
    subset = [r for r in items if value_fn(r) > 0][:top]
    if not subset:
        fig, ax = plt.subplots(figsize=(8.8, 3.2))
        ax.text(0.5, 0.5, "Sin horas perdidas en este recorte", ha="center", va="center")
        ax.axis("off")
        fig.savefig(out, dpi=160, bbox_inches="tight")
        plt.close(fig)
        return out
    orden = list(reversed(subset))
    labels = [f"{_trunc(r.node_name, 28)} ({r.node_id})" for r in orden]
    vals = [value_fn(r) for r in orden]
    if color:
        colors = [color] * len(orden)
    else:
        colors = [COLOR_DESC if r.desconectado else COLOR_HUECOS for r in orden]
    fig, ax = plt.subplots(figsize=(8.8, max(3.4, 0.32 * len(subset) + 1.4)))
    fig.patch.set_facecolor("white")
    ax.barh(labels, vals, color=colors, zorder=3)
    ax.set_xlabel("Horas perdidas")
    ax.set_title(titulo)
    ax.grid(axis="x", linestyle=":", alpha=0.5, zorder=0)
    ax.set_axisbelow(True)
    if leyenda_estado and color is None:
        ax.legend(
            handles=[
                Patch(facecolor=COLOR_DESC, label="Desconectado"),
                Patch(facecolor=COLOR_HUECOS, label="Conectado (huecos)"),
            ],
            loc="lower right",
            fontsize=8,
        )
    for y, v in enumerate(vals):
        ax.text(v, y, f" {_fmt_int(v)}", va="center", fontsize=7)
    ax.set_xlim(0, max((max(vals) if vals else 0) * 1.18, 1.0))
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def grafico_apilado_top(
    items: Sequence[ResultadoNodo],
    out: Path,
    top: int = 15,
    titulo: str = "Top puntos — horas perdidas apiladas por semana",
) -> Path:
    subset = list(items)[:top]
    if not subset:
        fig, ax = plt.subplots(figsize=(8.8, 3.2))
        ax.axis("off")
        fig.savefig(out, dpi=160, bbox_inches="tight")
        plt.close(fig)
        return out
    labels = [_trunc(r.node_name, 22) for r in reversed(subset)]
    fig, ax = plt.subplots(figsize=(8.8, max(3.8, 0.34 * len(subset) + 1.6)))
    fig.patch.set_facecolor("white")
    left = [0] * len(subset)
    rev = list(reversed(subset))
    for i, (num, _ini, _fin, _lbl) in enumerate(SEMANA_DEFS):
        vals = [r.perdidas_semana(num) for r in rev]
        ax.barh(
            labels,
            vals,
            left=left,
            color=COLORES_SEMANA[i % len(COLORES_SEMANA)],
            label=f"S{num}",
            zorder=3,
        )
        left = [a + b for a, b in zip(left, vals)]
    ax.set_xlabel("Horas perdidas")
    ax.set_title(titulo)
    ax.legend(loc="lower right", fontsize=8)
    ax.grid(axis="x", linestyle=":", alpha=0.5, zorder=0)
    ax.set_axisbelow(True)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def grafico_dias_ultima_semana(
    items: Sequence[ResultadoNodo],
    out: Path,
    top: int = 15,
) -> Path:
    if not SEMANA_DEFS:
        return out
    num, ini, fin, label = SEMANA_DEFS[-1]
    dias = _iter_days(ini, fin)
    subset = [r for r in items if r.perdidas_semana(num) > 0][:top]
    if not subset:
        fig, ax = plt.subplots(figsize=(8.8, 3.2))
        ax.text(0.5, 0.5, f"Sin horas perdidas en {label}", ha="center", va="center")
        ax.axis("off")
        fig.savefig(out, dpi=160, bbox_inches="tight")
        plt.close(fig)
        return out
    labels = [_trunc(r.node_name, 22) for r in reversed(subset)]
    rev = list(reversed(subset))
    fig, ax = plt.subplots(figsize=(8.8, max(3.8, 0.38 * len(subset) + 1.8)))
    fig.patch.set_facecolor("white")
    n = len(dias)
    h = 0.72 / max(n, 1)
    y = list(range(len(subset)))
    for i, dia in enumerate(dias):
        vals = [r.perdidas_dia(dia) for r in rev]
        offset = (i - (n - 1) / 2) * h
        ax.barh(
            [yi + offset for yi in y],
            vals,
            height=h * 0.9,
            color=COLORES_SEMANA[i % len(COLORES_SEMANA)],
            label=dia.strftime("%d-%m"),
            zorder=3,
        )
    ax.set_yticks(y, labels)
    ax.set_xlabel("Horas perdidas")
    ax.set_title(f"{label} — por día")
    ax.legend(loc="lower right", fontsize=8)
    ax.grid(axis="x", linestyle=":", alpha=0.5, zorder=0)
    ax.set_axisbelow(True)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def convertir_a_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        print("[ADVERTENCIA] No hay LibreOffice; se deja solo el Word.")
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    return pdf_path if pdf_path.exists() else None


def _filas_ranking_periodo(items: Sequence[ResultadoNodo]) -> List[Sequence[str]]:
    rows: List[Sequence[str]] = [
        ("#", "Punto", "Empresa", "S32", "S33", "S34", "S35*", "Total", "%", "Estado", "Última conexión"),
    ]
    for i, r in enumerate(items, start=1):
        rows.append(
            (
                str(i),
                f"{r.node_name}\n{r.node_id}",
                r.company_name,
                _fmt_int(r.perdidas_semana(32)),
                _fmt_int(r.perdidas_semana(33)),
                _fmt_int(r.perdidas_semana(34)),
                _fmt_int(r.perdidas_semana(35)),
                _fmt_int(r.perdidas_total()),
                _fmt_pct(r.perdidas_total(), r.esperadas_total()),
                r.etiqueta_estado(),
                r.last_update or "—",
            )
        )
    return rows


def _agregar_tabla_ranking(
    doc: Document,
    items: Sequence[ResultadoNodo],
    header_fill: str,
    highlight_col: int = 7,
) -> None:
    rows = _filas_ranking_periodo(items)
    if len(rows) == 1:
        doc.add_paragraph("Ningún punto en este grupo.")
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    _fill_table(t, rows, header_fill=header_fill, highlight_col=highlight_col)


def crear_reporte_word(
    resultados: List[ResultadoNodo],
    ahora_chile: datetime,
    output_dir: Path,
    charts: Dict[str, Path],
    top_chart: int,
) -> Path:
    con_perdidas = [r for r in resultados if r.perdidas_total() > 0]
    con_perdidas.sort(key=lambda r: r.perdidas_total(), reverse=True)
    desc, huec = particionar_perdidas(con_perdidas)
    n_total = len(resultados)
    n_con = len(con_perdidas)
    h_desc = sum(r.perdidas_total() for r in desc)
    h_huec = sum(r.perdidas_total() for r in huec)

    tot_sem = {n: sum(r.perdidas_semana(n) for r in resultados) for n, *_ in SEMANA_DEFS}
    esp_sem = {n: sum(r.esperadas_semana(n) for r in resultados) for n, *_ in SEMANA_DEFS}
    tot_all = sum(tot_sem.values())
    esp_all = sum(esp_sem.values())
    nsem = len(SEMANA_DEFS)
    ultima = SEMANA_DEFS[-1]
    ultima_parcial = (ultima[2] - ultima[1]).days < 6
    rango_txt = (
        f"{INICIO_PERIODO.strftime('%d-%m-%Y')} a {TODOS_LOS_DIAS[-1].strftime('%d-%m-%Y')}"
    )
    semanas_txt = ", ".join(
        f"S{n} ({_fmt_rango_corto(ini, fin)}{' parcial' if (fin - ini).days < 6 else ''})"
        for n, ini, fin, _ in SEMANA_DEFS
    )

    doc = Document()
    title = doc.add_heading("REPORTE DE PUNTOS CON HORAS PERDIDAS", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.color.rgb = RGBColor(192, 0, 0)
    title.runs[0].bold = True

    sub = doc.add_paragraph(
        f"Comparativa de {nsem} semanas · un gráfico por semana · "
        "al final, puntos con pérdida de data por desconexión"
    )
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in sub.runs:
        run.font.size = Pt(11)
        run.italic = True

    gen = doc.add_paragraph(
        f"Generado: {ahora_chile.strftime('%d-%m-%Y %H:%M')} hora Chile · "
        f"Serie horaria API dates.measures.csv · Universo: {n_total} puntos "
        f"(mismas exclusiones que puntos en cero)"
    )
    gen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    gen.runs[0].font.size = Pt(9)
    gen.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_heading("1. Criterio", 1)
    doc.add_paragraph(
        "Hora perdida = hora Chile esperada sin registro en la serie horaria "
        "(una hora en 0 no cuenta: el punto sí reportó). "
        f"Periodo: {rango_txt} ({semanas_txt}). "
        f"Desconexión: lastUpdate de la app ≥ {HORAS_UMBRAL_CONEXION_APP:.0f} h."
    )

    doc.add_heading(f"2. Comparativa de las {nsem} semanas", 1)
    nota_parcial = (
        f" S{ultima[0]} es parcial ({_fmt_rango_corto(ultima[1], ultima[2])}); "
        "no se compara en magnitud con una semana de 168 h."
        if ultima_parcial
        else ""
    )
    doc.add_paragraph(
        f"Flota: {_fmt_int(tot_all)} h perdidas de {_fmt_int(esp_all)} esperadas "
        f"({_fmt_pct(tot_all, esp_all)}). "
        f"{n_con} de {n_total} puntos con al menos 1 h faltante.{nota_parcial}"
    )
    _add_picture(doc, charts["totales"], 6.4)
    completas = _semanas_completas()
    if len(completas) >= 3:
        a, b, c = completas[-3], completas[-2], completas[-1]
        va, vb, vc = tot_sem[a[0]], tot_sem[b[0]], tot_sem[c[0]]
        if vc > va and vc > vb:
            tendencia = (
                f"S{c[0]} sube respecto de S{a[0]} y S{b[0]}: "
                "más huecos de telemetría en la última semana completa."
            )
        elif vc < va and vc < vb:
            tendencia = (
                f"S{c[0]} cierra con menos horas perdidas que S{a[0]} y S{b[0]}: "
                "mejora relativa de cobertura."
            )
        else:
            tendencia = (
                f"S{a[0]}, S{b[0]} y S{c[0]} se mantienen en un rango similar, "
                "sin un quiebre fuerte."
            )
        doc.add_paragraph(tendencia)

    doc.add_heading("3. Un gráfico por semana", 1)
    doc.add_paragraph(
        "Ranking de puntos con horas perdidas esa semana (mayor a menor). "
        "Rojo = desconectado ahora; azul = conectado con huecos."
    )
    primera = True
    for num, _ini, _fin, label in SEMANA_DEFS:
        if not primera:
            doc.add_page_break()
        primera = False
        con = [r for r in resultados if r.perdidas_semana(num) > 0]
        n_desc = sum(1 for r in con if r.desconectado)
        doc.add_heading(label, 2)
        doc.add_paragraph(
            f"{len(con)} puntos · {_fmt_int(tot_sem[num])} h de flota "
            f"({_fmt_pct(tot_sem[num], esp_sem[num])} de las esperadas) · "
            f"{n_desc} de ellos están desconectados ahora."
        )
        _add_picture(doc, charts[f"s{num}"], 6.4)
        if num == ultima[0] and ultima_parcial:
            doc.add_paragraph(
                f"Detalle diario de S{num} (hasta {ahora_chile.strftime('%d-%m %H:%M')} Chile)."
            )
            _add_picture(doc, charts["ultima_dias"], 6.4)

    doc.add_page_break()
    doc.add_heading("4. Pérdida de data por desconexión", 1)
    doc.add_paragraph(
        f"{len(desc)} puntos caídos ahora (lastUpdate ≥ {HORAS_UMBRAL_CONEXION_APP:.0f} h). "
        f"Suman {_fmt_int(h_desc)} h del periodo ({_fmt_pct(h_desc, tot_all)} del total perdido). "
        "Este es el ranking de prioridad de reconexión, mayor a menor."
    )
    _add_picture(doc, charts["ranking_desconexion"], 6.4)
    _add_picture(doc, charts["apilado_desconexion"], 6.4)
    if desc:
        lista_d = "; ".join(
            f"{r.node_name} ({r.node_id}, {_fmt_int(r.perdidas_total())} h, {r.etiqueta_estado()})"
            for r in desc
        )
        prio = doc.add_paragraph()
        prio.add_run("Atacar: ").bold = True
        prio.add_run(lista_d + ".")

    doc.add_heading("5. Conclusión", 1)
    doc.add_paragraph(
        f"La comparativa de las {nsem} semanas muestra el volumen de horas sin dato. "
        "Los gráficos semanales ordenan los puntos de mayor a menor. "
        "La prioridad operativa está en la sección 4: puntos que hoy pierden data por desconexión."
    )
    nota = doc.add_paragraph(
        "Nota: exclusiones iguales al reporte de puntos en cero. "
        f"Fuente: GET /nodes/{{id}}/dates.measures.csv. Periodo {rango_txt}."
    )
    nota.runs[0].font.size = Pt(8)
    nota.runs[0].italic = True
    nota.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    stamp = ahora_chile.strftime("%Y%m%d_%H%M")
    first_w = SEMANA_DEFS[0][0] if SEMANA_DEFS else 32
    last_w = SEMANA_DEFS[-1][0] if SEMANA_DEFS else 36
    out = output_dir / f"Reporte_Horas_Perdidas_S{first_w}_S{last_w}_{stamp}.docx"
    doc.save(str(out))
    return out


def serializar(resultados: Iterable[ResultadoNodo]) -> List[Dict]:
    out = []
    for r in resultados:
        out.append(
            {
                "nodeId": r.node_id,
                "nodeName": r.node_name,
                "companyId": r.company_id,
                "companyName": r.company_name,
                "perdidasPorDia": r.perdidas_por_dia,
                "esperadasPorDia": r.esperadas_por_dia,
                "semanas": {n: r.perdidas_semana(n) for n, *_ in SEMANA_DEFS},
                "s32": r.perdidas_semana(32) if 32 in SEMANA_POR_DIA.values() else 0,
                "s33": r.perdidas_semana(33) if 33 in SEMANA_POR_DIA.values() else 0,
                "s34": r.perdidas_semana(34) if 34 in SEMANA_POR_DIA.values() else 0,
                "s35": r.perdidas_semana(35) if 35 in SEMANA_POR_DIA.values() else 0,
                "s36": r.perdidas_semana(36) if 36 in SEMANA_POR_DIA.values() else 0,
                "total": r.perdidas_total(),
                "lastUpdate": r.last_update,
                "wesStatus": r.wes_status,
                "desconectado": r.desconectado,
                "horasSinConexion": (
                    round(r.horas_sin_conexion, 1) if r.horas_sin_conexion is not None else None
                ),
                "estado": r.etiqueta_estado() if r.perdidas_total() > 0 else "ok",
                "error": r.error,
            }
        )
    return out


def main() -> int:
    if sys.platform == "win32":
        for stream in (sys.stdout, sys.stderr):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass

    parser = argparse.ArgumentParser(description="Reporte semanal de horas perdidas")
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="Carpeta de salida (default: reports/Puntos_En_Cero)",
    )
    parser.add_argument("--top", type=int, default=25, help="Puntos en gráficos de ranking")
    parser.add_argument("--max-nodos", type=int, default=0, help="Limitar nodos (prueba)")
    parser.add_argument("--workers", type=int, default=MAX_WORKERS)
    parser.add_argument(
        "--desde-json",
        type=Path,
        default=None,
        help="Reusar horas ya calculadas (solo refresca estado de conexión)",
    )
    args = parser.parse_args()

    workers = max(4, int(args.workers))
    ahora = datetime.now(CHILE_TZ)
    configurar_periodo(ahora.date())
    output_dir = args.output_dir
    if output_dir is None:
        output_dir = Path(wes_scripts_root()) / "reports" / "Puntos_En_Cero"
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    charts_dir = output_dir / "_charts_horas_perdidas"
    charts_dir.mkdir(parents=True, exist_ok=True)

    semanas_txt = " / ".join(f"S{n}" for n, *_ in SEMANA_DEFS)
    print("=" * 70)
    print(f"REPORTE DE HORAS PERDIDAS — {semanas_txt}")
    print(f"Corte Chile: {ahora.strftime('%d-%m-%Y %H:%M')}")
    print(
        f"Periodo: {TODOS_LOS_DIAS[0].isoformat()} a {TODOS_LOS_DIAS[-1].isoformat()} "
        f"({len(TODOS_LOS_DIAS)} días)"
    )
    print("=" * 70)

    if args.desde_json:
        json_in = args.desde_json.resolve()
        if not json_in.is_file():
            print(f"[ERROR] No existe --desde-json: {json_in}")
            return 1
        print(f"[INFO] Reusando horas desde {json_in}")
        resultados = cargar_desde_json(json_in)
    else:
        nodos = obtener_todos_los_nodos()
        if args.max_nodos and args.max_nodos > 0:
            nodos = nodos[: args.max_nodos]
            print(f"[INFO] --max-nodos={args.max_nodos}")
        if not nodos:
            print("[ERROR] No hay nodos para analizar.")
            return 1

        print(f"[INFO] Analizando {len(nodos)} puntos con {workers} workers...")
        resultados = []
        with ThreadPoolExecutor(max_workers=workers) as ex:
            futs = {ex.submit(analizar_nodo, n, ahora): n for n in nodos}
            done = 0
            for fut in as_completed(futs):
                resultados.append(fut.result())
                done += 1
                if done % 20 == 0 or done == len(nodos):
                    print(f"  ... {done}/{len(nodos)} puntos")

    resultados.sort(key=lambda r: r.perdidas_total(), reverse=True)
    con = [r for r in resultados if r.perdidas_total() > 0]
    print(f"[INFO] Puntos con horas perdidas: {len(con)} / {len(resultados)}")
    print("[INFO] Clasificando desconexión (lastUpdate app)...")
    enriquecer_conexion(con, ahora)
    desc, huec = particionar_perdidas(con)
    print(
        f"[INFO] Desconectados: {len(desc)} ({sum(r.perdidas_total() for r in desc)} h) · "
        f"Conectados con huecos: {len(huec)} ({sum(r.perdidas_total() for r in huec)} h)"
    )

    top = max(5, int(args.top))
    first_w = SEMANA_DEFS[0][0]
    last_w = SEMANA_DEFS[-1][0]
    charts = {
        "totales": charts_dir / "totales_por_semana.png",
        "ranking_desconexion": charts_dir / "ranking_desconexion.png",
        "apilado_desconexion": charts_dir / "apilado_desconexion.png",
        "ultima_dias": charts_dir / f"ranking_s{last_w}_dias.png",
    }
    for num, *_rest in SEMANA_DEFS:
        charts[f"s{num}"] = charts_dir / f"ranking_s{num}.png"
    esp_sem = {n: sum(r.esperadas_semana(n) for r in resultados) for n, *_ in SEMANA_DEFS}
    grafico_totales_semana(resultados, esp_sem, charts["totales"])
    for num, _ini, _fin, label in SEMANA_DEFS:
        orden = sorted(resultados, key=lambda r: r.perdidas_semana(num), reverse=True)
        grafico_ranking(
            orden,
            lambda r, n=num: r.perdidas_semana(n),
            f"{label} — horas perdidas (mayor a menor)",
            charts[f"s{num}"],
            None,
            top,
        )
    grafico_ranking(
        desc,
        lambda r: r.perdidas_total(),
        "Pérdida de data por desconexión — ranking",
        charts["ranking_desconexion"],
        COLOR_DESC,
        top,
        leyenda_estado=False,
    )
    grafico_apilado_top(
        desc,
        charts["apilado_desconexion"],
        min(top, 15),
        f"Desconectados — horas por semana (S{first_w} a S{last_w})",
    )
    grafico_dias_ultima_semana(
        sorted(resultados, key=lambda r: r.perdidas_semana(last_w), reverse=True),
        charts["ultima_dias"],
        min(top, 15),
    )

    json_path = output_dir / f"horas_perdidas_s{first_w}_s{last_w}.json"
    json_path.write_text(
        json.dumps(serializar(resultados), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"[OK] JSON: {json_path}")

    docx_path = crear_reporte_word(resultados, ahora, output_dir, charts, top)
    print(f"[OK] DOCX: {docx_path}")
    pdf_path = convertir_a_pdf(docx_path)
    if pdf_path:
        print(f"[OK] PDF: {pdf_path}")

    print()
    print("Ranking desconexión:")
    for i, r in enumerate(desc[:10], start=1):
        print(
            f"  {i:2d}. {r.node_id}  {r.node_name}  T={r.perdidas_total()}  "
            f"{r.etiqueta_estado()}  lu={r.last_update}"
        )
    print("Ranking conectados con huecos:")
    for i, r in enumerate(huec[:10], start=1):
        print(
            f"  {i:2d}. {r.node_id}  {r.node_name}  T={r.perdidas_total()}  lu={r.last_update}"
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
