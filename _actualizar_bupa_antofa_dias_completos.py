"""
Personaliza el agregado Bupa Antofagasta (días civiles completos):
- Quita métricas / proyección nocturna / comparación mensual nocturna.
- Día mayor consumo nocturno del Medidor Principal Sanitaria.
- Proyección mensual = promedio diario WES × 30 (todos los puntos, incl. Sanitaria).
- Factura julio solo como referencia histórica.
- Participación salas vs Sanitaria + gap no monitoreado.

Uso:
  python _actualizar_bupa_antofa_dias_completos.py
  python _actualizar_bupa_antofa_dias_completos.py --doc PATH.docx
  python _actualizar_bupa_antofa_dias_completos.py --start 23/07/2026 --end 11/08/2026
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

from generar_reporte_word import (
    add_formatted_title,
    add_picture_with_pagination,
    add_table,
    acl_node_base_url,
    build_hourly_consumption_line_chart,
    fetch_json,
    find_max_nocturnal_consumption_day,
    flatten_measures,
    format_currency_chilean,
    format_number_chilean,
    normalize_measures_payload,
    parse_date,
    summarize_consumption,
)

# Defaults alineados con generar_agregado_bupa_antofagasta.py
START_STR = "23/07/2026"
END_STR = "11/08/2026"
START = parse_date(START_STR)
END = parse_date(END_STR, end_of_day=True)
NUM_DIAS = (END.date() - START.date()).days + 1

FACTURA_M3 = 6696.0
FACTURA_CLP = 18_538_860.0
PRECIO_M3 = FACTURA_CLP / FACTURA_M3

NODOS_DEF = [
    {"id": "000029-07", "nombre": "Sala de Bomba Principal", "tipo": "bomba"},
    {"id": "000029-08", "nombre": "Sala de Bomba Sexto Piso", "tipo": "bomba"},
    {"id": "000029-09", "nombre": "Medidor Principal Sanitaria", "tipo": "cuenta"},
    {"id": "000029-10", "nombre": "Sala de Bomba N°2", "tipo": "bomba"},
]

ABREGADO_ROOT = Path("reports/Bupa_Antofagasta/ABREGADO")


def _find_latest_doc() -> Path:
    candidates = sorted(
        ABREGADO_ROOT.glob("AGREGADO_*/Reporte_Agregado_BUPA_*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    # Preferir el docx "base" (sin _ordenado / _actualizado)
    for p in candidates:
        name = p.name.lower()
        if "_ordenado" in name or "_actualizado" in name:
            continue
        return p
    if candidates:
        return candidates[0]
    raise FileNotFoundError(f"No hay Word agregado en {ABREGADO_ROOT}")


DOC_PATH = Path(".")  # se fija en main()
OUT_DIR = Path(".")


def _para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def _find_conclusiones(doc: Document):
    for child in list(doc.element.body):
        if child.tag.split("}")[-1] == "p" and _para_text(child).upper().startswith("CONCLUSIONES"):
            return child
    raise RuntimeError("CONCLUSIONES no encontrado")


def _remove_from_until(doc: Document, start_pred, end_pred) -> int:
    body = doc.element.body
    children = list(body)
    start_idx = end_idx = None
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] != "p":
            continue
        t = _para_text(child)
        if start_idx is None and start_pred(t):
            start_idx = i
        elif start_idx is not None and end_pred(t):
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Rango no encontrado start={start_idx} end={end_idx}")
    n = 0
    for child in children[start_idx:end_idx]:
        body.remove(child)
        n += 1
    return n


def _move_tail_before(doc: Document, marker, n_new: int) -> None:
    children = [c for c in list(doc.element.body) if c.tag.split("}")[-1] != "sectPr"]
    for el in children[-n_new:]:
        marker.addprevious(el)


def _fetch_rows() -> list[dict]:
    rows = []
    for n in NODOS_DEF:
        raw = fetch_json(
            f"{acl_node_base_url()}/nodes/measures/dates",
            params=[
                ("id", n["id"]),
                ("start", START.strftime("%d%m%Y")),
                ("end", END.strftime("%d%m%Y")),
            ],
        )
        total = summarize_consumption(
            flatten_measures(normalize_measures_payload(raw, n["id"]))
        )["total"]
        prom = total / NUM_DIAS
        proy = prom * 30.0
        rows.append(
            {
                **n,
                "m3_periodo": total,
                "prom_dia": prom,
                "proy_m3": proy,
                "proy_clp": proy * PRECIO_M3,
            }
        )
    return rows


def _chart_proyeccion(path: Path, rows: list[dict]) -> Path:
    ordered = sorted(rows, key=lambda r: r["proy_m3"], reverse=True)
    fig, ax = plt.subplots(figsize=(8, 4.8))
    bars = ax.bar(
        [r["nombre"] for r in ordered],
        [r["proy_m3"] for r in ordered],
        color="#0050b3",
        alpha=0.85,
        edgecolor="#003a80",
        linewidth=1.1,
    )
    ax.set_ylabel("m³ / 30 días", fontsize=12, fontweight="bold")
    ax.set_title(
        f"Proyección mensual WES (promedio {NUM_DIAS} días completos × 30)",
        fontsize=13,
        fontweight="bold",
    )
    ax.set_ylim(bottom=0)
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=25, ha="right", fontsize=10)
    for bar, r in zip(bars, ordered):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(r['proy_m3'], 1)} m³",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _chart_participacion(path: Path, salas: float, sanit: float, gap: float) -> Path:
    fig, ax = plt.subplots(figsize=(6.5, 4.6))
    wedges, *_ = ax.pie(
        [salas, gap],
        labels=[
            f"Salas de bomba\n({format_number_chilean(salas / sanit * 100, 1)}%)",
            f"No monitoreado\n({format_number_chilean(gap / sanit * 100, 1)}%)",
        ],
        colors=["#0050b3", "#c0392b"],
        autopct=lambda p: f"{p:.1f}%",
        startangle=90,
        textprops={"fontsize": 10},
    )
    for w in wedges:
        w.set_edgecolor("white")
        w.set_linewidth(1.5)
    ax.set_title(
        "Participación proyectada sobre Medidor Principal Sanitaria",
        fontsize=12,
        fontweight="bold",
    )
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def main(argv: list[str] | None = None) -> None:
    global DOC_PATH, OUT_DIR, START, END, NUM_DIAS, START_STR, END_STR

    parser = argparse.ArgumentParser(description="Personaliza agregado Bupa Antofagasta")
    parser.add_argument("--doc", type=Path, default=None, help="Word agregado a personalizar")
    parser.add_argument("--start", default=START_STR, help="Inicio DD/MM/YYYY")
    parser.add_argument("--end", default=END_STR, help="Fin DD/MM/YYYY (día completo)")
    args = parser.parse_args(argv)

    START_STR = args.start
    END_STR = args.end
    START = parse_date(START_STR)
    END = parse_date(END_STR, end_of_day=True)
    NUM_DIAS = (END.date() - START.date()).days + 1

    DOC_PATH = args.doc if args.doc else _find_latest_doc()
    OUT_DIR = DOC_PATH.parent
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    print(f"[INFO] Doc: {DOC_PATH}")
    print(f"[INFO] Periodo: {START_STR} – {END_STR} ({NUM_DIAS} días completos)")

    rows = _fetch_rows()
    sanit = next(r for r in rows if r["tipo"] == "cuenta")
    bombas = [r for r in rows if r["tipo"] == "bomba"]
    salas_m3 = sum(b["proy_m3"] for b in bombas)
    salas_clp = sum(b["proy_clp"] for b in bombas)
    gap_m3 = max(0.0, sanit["proy_m3"] - salas_m3)
    gap_clp = gap_m3 * PRECIO_M3
    pct_salas = salas_m3 / sanit["proy_m3"] * 100.0 if sanit["proy_m3"] else 0.0
    pct_gap = 100.0 - pct_salas

    # Gráfico nocturno Sanitaria
    result = find_max_nocturnal_consumption_day("000029-09", None, START, END)
    if not result:
        raise RuntimeError("Sin día mayor nocturno Sanitaria")
    target_dt, hourly = result
    chart_sanitaria = OUT_DIR / "chart_max_nocturnal_000029_09.png"
    build_hourly_consumption_line_chart(
        hourly,
        chart_sanitaria,
        target_dt,
        f"Día con mayor consumo nocturno ({target_dt.strftime('%d-%m-%y')})",
    )

    chart_proy = _chart_proyeccion(OUT_DIR / "chart_proyeccion_mensual_total_por_punto.png", rows)
    chart_part = _chart_participacion(
        OUT_DIR / "chart_participacion_salas_vs_sanitaria.png",
        salas_m3,
        sanit["proy_m3"],
        gap_m3,
    )

    doc = Document(str(DOC_PATH))

    # Quitar desde RESUMEN/MÉTRICAS (o perfiles) hasta CONCLUSIONES
    # El formato estándar trae métricas + alertas + perfiles + comparación mensual
    n_del = _remove_from_until(
        doc,
        lambda t: (
            t.upper().startswith("RESUMEN POR PUNTO")
            or t.upper().startswith("MÉTRICAS POR PUNTO")
            or t.upper().startswith("METRICAS POR PUNTO")
            or ("PERFILES HORARIOS" in t.upper())
        ),
        lambda t: t.upper().startswith("CONCLUSIONES"),
    )
    print(f"[OK] Eliminados {n_del} elementos del bloque estándar a reemplazar")

    conclus = _find_conclusiones(doc)
    before = len([c for c in doc.element.body if c.tag.split("}")[-1] != "sectPr"])

    # --- Día mayor nocturno Sanitaria ---
    add_formatted_title(
        doc,
        f"Día con mayor consumo nocturno - MEDIDOR PRINCIPAL SANITARIA "
        f"({target_dt.strftime('%d-%m-%y')}):",
    )
    add_picture_with_pagination(doc, str(chart_sanitaria), Inches(6), keep_with_next=True)
    p = doc.add_paragraph(
        "El gráfico muestra el día con mayor consumo nocturno (00:00–06:59) del "
        "Medidor Principal Sanitaria en el periodo de días completos analizado. "
        "Este medidor representa el total de la cuenta de agua de la clínica."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # --- Proyección WES ---
    add_formatted_title(doc, "Proyección de consumo mensual (WES · días completos)")
    intro = doc.add_paragraph(
        f"Se consideran solo días civiles completos ({NUM_DIAS} días: {START_STR}–{END_STR}), "
        f"sin incluir el día en curso a fin de no distorsionar el promedio diario. "
        f"Para cada punto: promedio diario del periodo × 30. "
        f"La valorización usa la tarifa efectiva de la última factura "
        f"(${format_number_chilean(PRECIO_M3, 0)} CLP/m³ = "
        f"{format_currency_chilean(FACTURA_CLP)} ÷ {format_number_chilean(FACTURA_M3, 0)} m³). "
        f"La factura de julio ({format_number_chilean(FACTURA_M3, 0)} m³ / "
        f"{format_currency_chilean(FACTURA_CLP)}) se muestra solo como referencia histórica; "
        f"no reemplaza la proyección WES."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in intro.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    table_proy = [
        (
            "Punto",
            f"Consumo {NUM_DIAS} días (m³)",
            "Promedio diario (m³)",
            "Proyección 30 días (m³)",
            "Proyección 30 días (CLP)",
        )
    ]
    for r in sorted(rows, key=lambda x: x["proy_m3"], reverse=True):
        table_proy.append(
            (
                f"{r['nombre']} ({r['id']})",
                format_number_chilean(r["m3_periodo"], 1),
                format_number_chilean(r["prom_dia"], 1),
                format_number_chilean(r["proy_m3"], 1),
                format_currency_chilean(r["proy_clp"]),
            )
        )
    # Fila referencia factura
    table_proy.append(
        (
            "Referencia: factura julio (histórica)",
            "—",
            "—",
            format_number_chilean(FACTURA_M3, 0),
            format_currency_chilean(FACTURA_CLP),
        )
    )
    add_table(doc, "Proyección / cuenta mensual", table_proy, wes_style=True)

    nota = doc.add_paragraph()
    nota.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    nota.add_run("Lectura para el cliente: ").bold = True
    nota.add_run(
        f"en estos {NUM_DIAS} días el Medidor Principal Sanitaria registró "
        f"{format_number_chilean(sanit['m3_periodo'], 1)} m³ "
        f"({format_number_chilean(sanit['prom_dia'], 1)} m³/día). "
        f"Si ese ritmo se mantiene, la cuenta a 30 días se acerca a "
        f"{format_number_chilean(sanit['proy_m3'], 1)} m³ "
        f"({format_currency_chilean(sanit['proy_clp'])}), "
        f"frente a la última factura de {format_number_chilean(FACTURA_M3, 0)} m³ "
        f"({format_currency_chilean(FACTURA_CLP)})."
    )
    for run in nota.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_proy), Inches(6), keep_with_next=True)

    # --- Participación + gap ---
    add_formatted_title(doc, "Participación de salas de bomba vs cuenta (proyección WES)")
    expl = doc.add_paragraph(
        "El Medidor Principal Sanitaria representa el 100% de la cuenta (distribuye todo el consumo). "
        "Las salas de bomba son submediciones internas. La diferencia entre el 100% y la suma de las "
        "salas es volumen (y costo) que hoy no está siendo monitoreado punto a punto."
    )
    expl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in expl.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    part_rows = [
        ("Concepto", "Proy. 30 días (m³)", "Proy. 30 días (CLP)", "% sobre Sanitaria"),
        (
            "Medidor Principal Sanitaria (cuenta / 100%)",
            format_number_chilean(sanit["proy_m3"], 1),
            format_currency_chilean(sanit["proy_clp"]),
            "100,0%",
        ),
    ]
    for b in sorted(bombas, key=lambda x: x["proy_m3"], reverse=True):
        part_rows.append(
            (
                f"{b['nombre']} ({b['id']})",
                format_number_chilean(b["proy_m3"], 1),
                format_currency_chilean(b["proy_clp"]),
                f"{format_number_chilean(b['proy_m3'] / sanit['proy_m3'] * 100, 1)}%",
            )
        )
    part_rows.append(
        (
            "TOTAL salas de bomba (07+08+10)",
            format_number_chilean(salas_m3, 1),
            format_currency_chilean(salas_clp),
            f"{format_number_chilean(pct_salas, 1)}%",
        )
    )
    part_rows.append(
        (
            "NO MONITOREADO (Sanitaria − salas)",
            format_number_chilean(gap_m3, 1),
            format_currency_chilean(gap_clp),
            f"{format_number_chilean(pct_gap, 1)}%",
        )
    )
    add_table(doc, "Participación sobre cuenta (proyección mensual)", part_rows, wes_style=True)

    resumen = doc.add_paragraph()
    resumen.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    resumen.add_run("Resumen: ").bold = True
    resumen.add_run(
        f"proyectadas a 30 días, las tres salas de bomba suman "
        f"{format_number_chilean(salas_m3, 1)} m³ ({format_currency_chilean(salas_clp)}), "
        f"es decir el {format_number_chilean(pct_salas, 1)}% de la cuenta proyectada en Sanitaria. "
        f"Queda un diferencial no monitoreado de "
        f"{format_number_chilean(gap_m3, 1)} m³ ({format_currency_chilean(gap_clp)}), "
        f"equivalente al {format_number_chilean(pct_gap, 1)}% de la cuenta."
    )
    for run in resumen.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_part), Inches(5.2), keep_with_next=True)
    doc.add_paragraph("")

    after = len([c for c in doc.element.body if c.tag.split("}")[-1] != "sectPr"])
    _move_tail_before(doc, conclus, after - before)

    try:
        doc.save(str(DOC_PATH))
        out = DOC_PATH
    except PermissionError:
        out = DOC_PATH.with_name(DOC_PATH.stem + "_actualizado.docx")
        doc.save(str(out))
        print(f"[ADVERTENCIA] Guardado como {out.name}")

    print("[OK]", out)
    print(f"     Periodo: {NUM_DIAS} días completos {START_STR}–{END_STR}")
    print(
        f"     Sanitaria: {sanit['m3_periodo']:.1f} m3 periodo -> "
        f"{sanit['prom_dia']:.1f}/dia -> proy {sanit['proy_m3']:.1f} m3 "
        f"({format_currency_chilean(sanit['proy_clp'])})"
    )
    print(
        f"     Salas {salas_m3:.1f} m3 ({pct_salas:.1f}%) | "
        f"Gap no monitoreado {gap_m3:.1f} m3 ({pct_gap:.1f}%)"
    )
    print(f"     Factura ref: {FACTURA_M3:.0f} m3 / {format_currency_chilean(FACTURA_CLP)}")
    print(f"     Dia nocturno Sanitaria: {target_dt.strftime('%d-%m-%Y')}")


if __name__ == "__main__":
    main()
