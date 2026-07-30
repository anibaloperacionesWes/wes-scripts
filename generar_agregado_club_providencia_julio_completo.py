"""
Club Providencia — mismo formato agregado de junio (extendido), datos julio 2026,
más comparativo del mes y últimos 6 meses.

Uso:
  python generar_agregado_club_providencia_julio_completo.py
"""

from __future__ import annotations

import sys
import time
from datetime import date, datetime
from pathlib import Path
from typing import List, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

if sys.platform == "win32":
    for s in (sys.stdout, sys.stderr):
        try:
            s.reconfigure(encoding="utf-8", line_buffering=True)
        except Exception:
            pass

from generar_reporte_word import (
    add_formatted_heading,
    add_picture_with_pagination,
    build_monthly_comparison_chart,
    calculate_nocturnal_metrics,
    estilizar_tabla_wes,
    format_currency_chilean,
    format_number_chilean,
    generate_aggregated_report,
    get_water_price_per_m3,
    parse_date,
    acl_node_base_url,
    fetch_json,
    normalize_measures_payload,
)
from generar_consolidado_m3_mensual_puente_alto import consumo_mes_un_nodo

COMPANY_ID = "000031"
NODE_IDS = ["000031-01", "000031-02"]
START = "01/07/2026"
END = "30/07/2026"
COLOR_BARRA = "#0050b3"
PRECIO_DEFAULT = 1200.0


def _meses_ultimos_6(end_d: date) -> List[Tuple[int, int]]:
    y, m = end_d.year, end_d.month
    out: List[Tuple[int, int]] = []
    for _ in range(6):
        out.append((y, m))
        m -= 1
        if m == 0:
            m = 12
            y -= 1
    out.reverse()
    return out


def _plot_6_meses(series: List[Tuple[str, float]], out: Path) -> Path:
    labels = [s[0] for s in series]
    vals = [s[1] for s in series]
    fig, ax = plt.subplots(figsize=(10, 5.5))
    bars = ax.bar(labels, vals, color=COLOR_BARRA, width=0.65)
    ax.set_ylabel("Consumo mensual total (m³)", fontsize=11, fontweight="bold")
    ax.set_xlabel("Mes", fontsize=11, fontweight="bold")
    ax.set_title(
        "Comparativo últimos 6 meses — Club Providencia",
        fontsize=12,
        fontweight="bold",
    )
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 0)}",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _totales_periodo(start_dt: datetime, end_dt: datetime) -> Tuple[float, float, float]:
    """Retorna (total_m3, nocturno_m3, price)."""
    total = 0.0
    noct = 0.0
    prices: List[float] = []
    for nid in NODE_IDS:
        raw = fetch_json(
            f"{acl_node_base_url()}/nodes/measures/dates",
            params=[
                ("id", nid),
                ("start", start_dt.strftime("%d%m%Y")),
                ("end", end_dt.strftime("%d%m%Y")),
            ],
        )
        payload = normalize_measures_payload(raw, nid)
        from generar_reporte_word import flatten_measures, summarize_consumption

        measures = flatten_measures(payload)
        summary = summarize_consumption(measures)
        total += float(summary["total"])
        nm = calculate_nocturnal_metrics(nid, start_dt, end_dt)
        noct += float(nm.get("consumo_nocturno_total") or 0.0)
        prices.append(get_water_price_per_m3(COMPANY_ID, nid, payload) or PRECIO_DEFAULT)
    price = sum(prices) / len(prices) if prices else PRECIO_DEFAULT
    return total, noct, price


def _append_comparativos(docx_path: Path) -> Path:
    start_dt = parse_date(START)
    end_dt = parse_date(END, end_of_day=True)
    out_dir = docx_path.parent

    print("[INFO] Calculando totales periodo para comparativos...", flush=True)
    total_m3, noct_total, price = _totales_periodo(start_dt, end_dt)
    num_dias = (end_dt.date() - start_dt.date()).days + 1
    factor_30 = 30.0 / max(num_dias, 1)
    leak_monthly = noct_total * factor_30
    efectivo_monthly = max(0.0, (total_m3 - noct_total) * factor_30)

    chart_mes = out_dir / "chart_comparacion_mensual_extra.png"
    built_mes = build_monthly_comparison_chart(
        leak_monthly, efectivo_monthly, price, chart_mes
    )

    print("[INFO] Descargando últimos 6 meses...", flush=True)
    sess = requests.Session()
    series_6: List[Tuple[str, float]] = []
    for y, m in _meses_ultimos_6(end_dt.date()):
        if y == end_dt.year and m == end_dt.month:
            m3 = total_m3
            label = f"{y}-{m:02d}*"
        else:
            m3 = 0.0
            for nid in NODE_IDS:
                v, _, _ = consumo_mes_un_nodo(sess, nid, y, m)
                m3 += float(v)
            label = f"{y}-{m:02d}"
        series_6.append((label, float(m3)))
        print(f"  {label}: {m3:.1f} m³", flush=True)
    chart_6m = _plot_6_meses(series_6, out_dir / "chart_ultimos_6_meses.png")

    doc = Document(str(docx_path))
    add_formatted_heading(doc, "Comparativo del mes (nocturno vs efectivo)", level=1)
    p5 = doc.add_paragraph(
        f"Proyección a 30 días a partir del periodo ({num_dias} días): "
        f"nocturno proyectado {format_number_chilean(leak_monthly, 1)} m³ "
        f"({format_currency_chilean(leak_monthly * price)}) y consumo efectivo "
        f"{format_number_chilean(efectivo_monthly, 1)} m³ "
        f"({format_currency_chilean(efectivo_monthly * price)})."
    )
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if built_mes and Path(built_mes).exists():
        add_picture_with_pagination(doc, str(chart_mes), Inches(4.5), keep_with_next=True)

    add_formatted_heading(doc, "Comparativo últimos 6 meses", level=1)
    p6 = doc.add_paragraph(
        "Consumo mensual total Club Providencia (Matriz Fitness + Matriz Piscina). "
        "El mes marcado con * corresponde al periodo de este reporte "
        f"(hasta {end_dt.strftime('%d/%m/%Y')})."
    )
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_6m), Inches(6), keep_with_next=True)

    table = doc.add_table(rows=1 + len(series_6), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Mes"
    table.rows[0].cells[1].text = "Total (m³)"
    for i, (lab, v) in enumerate(series_6):
        table.rows[i + 1].cells[0].text = lab
        table.rows[i + 1].cells[1].text = format_number_chilean(v, 1)
    estilizar_tabla_wes(table, has_total_row=False)

    doc.save(str(docx_path))
    print(f"[OK] Comparativos agregados a {docx_path}")
    return docx_path


def main() -> int:
    print("=" * 70)
    print("CLUB PROVIDENCIA — agregado julio (formato junio) + comparativos")
    print(f"Periodo: {START} → {END}")
    print("=" * 70)
    t0 = time.perf_counter()
    out = generate_aggregated_report(
        company_id=COMPANY_ID,
        node_ids=list(NODE_IDS),
        start_date=START,
        end_date=END,
        output_dir="reports",
        apply_exclusions=False,
        generate_ppt=False,
        parallel_node_fetch=True,
        max_parallel_workers=4,
    )
    docx = Path(out)
    print(f"[OK] Base: {docx}")
    _append_comparativos(docx)
    print(f"[INFO] Tiempo total: {time.perf_counter() - t0:.1f} s")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
