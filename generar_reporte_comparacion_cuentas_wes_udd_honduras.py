"""
Genera informe Word + PDF: comparación boletas de agua (cuenta 1653854-K UDD)
vs consumo WES nodo 000026-01 (Sala impulsión Honduras) por período de lectura.

Uso:
  python generar_reporte_comparacion_cuentas_wes_udd_honduras.py
"""

from __future__ import annotations

import sys
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from generar_reporte_word import format_number_chilean
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "comparacion_udd_cuentas_vs_wes"
OUT_DOCX = OUT_DIR / "Comparacion_cuentas_vs_WES_Honduras_UDD.docx"
FACTURAS_DIR = OUT_DIR / "Facturaciones UDD H"

NODE_ID = "000026-01"
NODE_NOMBRE = "UDD — Sala impulsión Honduras"
CUENTA = "1653854-K — Universidad del Desarrollo, Av. La Plaza 700, Las Condes"

# Períodos extraídos de los PDF (lectura anterior -> lectura actual); consumo total medidor; emisión; N° boleta aprox del extracto
FILAS_INFORME: list[dict] = [
    {
        "periodo_txt": "03-oct-2025 → 04-nov-2025",
        "emision": "11-nov-2025",
        "boleta": "32709178",
        "medidor": "120.710.011",
        "m3_cuenta": 7732,
        "api_start": "03102025",
        "api_end": "04112025",
    },
    {
        "periodo_txt": "04-nov-2025 → 05-dic-2025",
        "emision": "11-dic-2025",
        "boleta": "32869189",
        "medidor": "120.710.011",
        "m3_cuenta": 9364,
        "api_start": "04112025",
        "api_end": "05122025",
    },
    {
        "periodo_txt": "05-dic-2025 → 06-ene-2026",
        "emision": "12-ene-2026",
        "boleta": "33194523",
        "medidor": "120.710.011",
        "m3_cuenta": 8319,
        "api_start": "05122025",
        "api_end": "06012026",
    },
    {
        "periodo_txt": "06-ene-2026 → 06-feb-2026",
        "emision": "10-feb-2026",
        "boleta": "33364353",
        "medidor": "2.025.759.723",
        "m3_cuenta": 8258,
        "nota_medidor": "Cambio de medidor: boleta indica 1.395 m³ (medidor antiguo) + 6.863 m³ (medidor nuevo).",
        "api_start": "06012026",
        "api_end": "06022026",
    },
    {
        "periodo_txt": "06-feb-2026 → 10-mar-2026",
        "emision": "11-mar-2026",
        "boleta": "33522229",
        "medidor": "2.025.759.723",
        "m3_cuenta": 8593,
        "api_start": "06022026",
        "api_end": "10032026",
    },
]

NOTA_ESTANQUE = (
    "Los totales no deben cuadrar de forma exacta: el medidor que alimenta el estanque registra el "
    "volumen bruto de ingreso a ese embalse. Desde el estanque existen dos salidas: una alimenta el "
    "riego y la otra el uso de agua de la universidad (baños, duchas, casino, etc.). El punto WES "
    f"({NODE_ID}) monitorea un tramo de la red asociado a la «Sala impulsión Honduras» y no reproduce "
    "el 100 % del volumen facturado en el medidor de la cuenta, ni el reparto entre riego y uso interior."
)

_MESES = {
    "ENE": 1,
    "FEB": 2,
    "MAR": 3,
    "ABR": 4,
    "MAY": 5,
    "JUN": 6,
    "JUL": 7,
    "AGO": 8,
    "SEP": 9,
    "OCT": 10,
    "NOV": 11,
    "DIC": 12,
}


def _parse_fecha_es(fecha_txt: str) -> datetime:
    # dd-MMM-YYYY, ej: 07-ABR-2026
    p = fecha_txt.strip().upper().split("-")
    dd = int(p[0])
    mm = _MESES[p[1][:3]]
    yy = int(p[2])
    return datetime(yy, mm, dd)


def _fmt_periodo(dt: datetime) -> str:
    meses = {
        1: "ene",
        2: "feb",
        3: "mar",
        4: "abr",
        5: "may",
        6: "jun",
        7: "jul",
        8: "ago",
        9: "sep",
        10: "oct",
        11: "nov",
        12: "dic",
    }
    return f"{dt.day:02d}-{meses[dt.month]}-{dt.year}"


def _to_ddmmyyyy(dt: datetime) -> str:
    return dt.strftime("%d%m%Y")


def _extraer_texto_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for pg in reader.pages:
        parts.append(pg.extract_text() or "")
    return "\n".join(parts)


def _parse_factura_pdf(path: Path) -> dict:
    txt = _extraer_texto_pdf(path)

    boleta_m = re.search(r"N[º°]\s*([0-9]{6,})", txt, flags=re.IGNORECASE)
    emision_m = re.search(r"FECHA\s*EMISI[ÓO]N[:\s]*([0-9]{2}-[A-Z]{3}-[0-9]{4})", txt, flags=re.IGNORECASE)
    medidor_m = re.search(r"N[úu]mero\s+de\s+Medidor\s+([0-9\.]+)", txt, flags=re.IGNORECASE)
    consumo_m = re.search(r"CONSUMO\s+TOTAL\s+([0-9\.\,]+)\s*m3", txt, flags=re.IGNORECASE)
    la_m = re.search(r"LECTURA\s+ACTUAL\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})\s+[0-9\.\,]+\s*m3", txt, flags=re.IGNORECASE)
    lan_m = re.search(r"LECTURA\s+ANTERIOR\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})\s+[0-9\.\,]+\s*m3", txt, flags=re.IGNORECASE)

    if not (boleta_m and emision_m and medidor_m and consumo_m and la_m and lan_m):
        raise ValueError(f"No se pudieron extraer todos los campos desde {path.name}")

    dt_actual = _parse_fecha_es(la_m.group(1))
    dt_anterior = _parse_fecha_es(lan_m.group(1))
    dt_emision = _parse_fecha_es(emision_m.group(1))
    m3_cuenta = int(float(consumo_m.group(1).replace(".", "").replace(",", ".")))

    out = {
        "periodo_txt": f"{_fmt_periodo(dt_anterior)} → {_fmt_periodo(dt_actual)}",
        "emision": dt_emision.strftime("%d-%b-%Y").lower().replace("apr", "abr").replace("aug", "ago").replace("dec", "dic"),
        "boleta": boleta_m.group(1),
        "medidor": medidor_m.group(1),
        "m3_cuenta": m3_cuenta,
        "api_start": _to_ddmmyyyy(dt_anterior),
        "api_end": _to_ddmmyyyy(dt_actual),
        "_sort_dt": dt_actual,
    }
    return out


def _cargar_facturas() -> list[dict]:
    pdfs = sorted(FACTURAS_DIR.glob("*.pdf"))
    if not pdfs:
        return FILAS_INFORME
    rows: list[dict] = []
    for p in pdfs:
        rows.append(_parse_factura_pdf(p))
    rows.sort(key=lambda r: r["_sort_dt"])
    for r in rows:
        r.pop("_sort_dt", None)
    return rows


def _fetch_wes_total(start: str, end: str) -> tuple[float, int]:
    from generar_reporte_word import (
        acl_node_base_url,
        fetch_json,
        normalize_measures_payload,
        flatten_measures,
        summarize_consumption,
    )

    base = acl_node_base_url()
    raw = fetch_json(
        f"{base}/nodes/measures/dates",
        params=[("id", NODE_ID), ("start", start), ("end", end)],
    )
    norm = normalize_measures_payload(raw, NODE_ID)
    meas = flatten_measures(norm)
    s = summarize_consumption(meas)
    return float(s.get("total", 0.0)), int(s.get("dias", 0))


def _fmt_num(n: float, dec: int = 2) -> str:
    return format_number_chilean(n, dec)


def generar_word() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_heading("Comparación cuentas de agua vs registro WES", level=1)
    if t.runs:
        t.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Cliente: ").bold = True
    p.add_run(CUENTA)
    p2 = doc.add_paragraph()
    p2.add_run("Punto WES: ").bold = True
    p2.add_run(f"{NODE_ID} — {NODE_NOMBRE}")
    p3 = doc.add_paragraph()
    p3.add_run("Fuente API: ").bold = True
    from generar_reporte_word import acl_node_base_url

    p3.add_run(acl_node_base_url())
    p4 = doc.add_paragraph()
    p4.add_run("Generado: ").bold = True
    p4.add_run(datetime.now().strftime("%d-%m-%Y %H:%M"))

    doc.add_paragraph()
    doc.add_heading("1. Contexto y limitación de la comparación", level=1)
    doc.add_paragraph(NOTA_ESTANQUE)

    doc.add_heading("2. Resumen por período de boleta", level=1)

    # Precompute WES
    filas: list[dict] = []
    filas_input = _cargar_facturas()
    for row in filas_input:
        total, dias = _fetch_wes_total(row["api_start"], row["api_end"])
        cuenta = row["m3_cuenta"]
        diff = cuenta - total
        pct = (100.0 * total / cuenta) if cuenta else 0.0
        filas.append({**row, "m3_wes": total, "dias_wes": dias, "diff": diff, "pct_wes_sobre_cuenta": pct})

    headers = [
        "Período (lecturas)",
        "Emisión",
        "N° boleta",
        "Medidor",
        "Consumo cuenta (m³)",
        "Total WES (m³)",
        "Días datos API",
        "Diferencia (m³)",
        "% volumen WES / cuenta",
    ]
    tbl = doc.add_table(rows=1 + len(filas), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for par in hdr[j].paragraphs:
            par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in par.runs:
                r.bold = True
                r.font.size = Pt(9)

    for i, r in enumerate(filas, start=1):
        cells = tbl.rows[i].cells
        vals = [
            r["periodo_txt"],
            r["emision"],
            r["boleta"],
            r["medidor"],
            _fmt_num(r["m3_cuenta"], 0),
            _fmt_num(r["m3_wes"], 2),
            str(r["dias_wes"]),
            _fmt_num(r["diff"], 2),
            _fmt_num(r["pct_wes_sobre_cuenta"], 1) + " %",
        ]
        for j, v in enumerate(vals):
            cells[j].text = v
            cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading("3. Notas sobre medidores", level=1)
    doc.add_paragraph(
        "Entre febrero y marzo de 2026 la boleta registra medidor 2.025.759.723 (diámetro 50 mm en extracto). "
        "La boleta del 10-feb-2026 documenta el reemplazo del medidor anterior (120.710.011) e incluye el texto "
        "sobre consumo combinado medidor antiguo y nuevo."
    )
    for r in filas:
        if r.get("nota_medidor"):
            doc.add_paragraph(r["nota_medidor"], style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("4. Metodología WES", level=1)
    doc.add_paragraph(
        "El total «Total WES» es la suma de los consumos diarios reportados por la API "
        "GET /nodes/measures/dates para el rango de fechas indicado (formato de fechas ddMMyyyy alineado "
        "a los extremos del período de lectura de cada boleta). Puede haber pequeñas diferencias de "
        "calendario (número de días) respecto del ciclo exacto de facturación."
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX.resolve()


def main() -> int:
    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass
    p = generar_word()
    print(f"[OK] Word: {p}")
    try:
        from generar_reporte_word import convertir_word_a_pdf

        pdf = convertir_word_a_pdf(p)
        if pdf:
            print(f"[OK] PDF:  {pdf.resolve()}")
        else:
            print("[AVISO] No se pudo generar PDF (instala docx2pdf o Word COM).")
    except Exception as e:
        print(f"[AVISO] PDF: {e}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
