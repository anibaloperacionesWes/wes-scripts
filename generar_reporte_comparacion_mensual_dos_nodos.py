"""
Comparación mensual de consumo entre dos nodos WES (totalM3 diario API).

Uso:
  python generar_reporte_comparacion_mensual_dos_nodos.py \\
    --node-a 000025-19 --node-b 000025-02 --desde 2025-04 --hasta 2026-04
"""

from __future__ import annotations

import argparse
import calendar
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from generar_reporte_diferencias_mensuales_nodo import (
    MESES_ES,
    MesResumen,
    _agregar_por_mes,
    _etiqueta_mes,
    _fetch_medidas_diarias,
    _iter_meses,
    _nombre_nodo,
    _parse_yyyy_mm,
)
from generar_reporte_word import acl_node_base_url, add_logo_to_header, format_number_chilean


@dataclass
class FilaComparacion:
    year: int
    month: int
    etiqueta: str
    m3_a: float
    dias_a: int
    m3_b: float
    dias_b: int
    diff_b_menos_a: float
    pct_b_vs_a: Optional[float]


def _filas_comparacion(
    filas_a: List[MesResumen],
    filas_b: List[MesResumen],
) -> List[FilaComparacion]:
    map_b = {(r.year, r.month): r for r in filas_b}
    out: List[FilaComparacion] = []
    for ra in filas_a:
        rb = map_b.get((ra.year, ra.month))
        mb = rb.m3 if rb else 0.0
        db = rb.dias_con_dato if rb else 0
        diff = round(mb - ra.m3, 2)
        pct = None
        if abs(ra.m3) > 1e-9:
            pct = round((diff / ra.m3) * 100.0, 1)
        elif abs(mb) > 1e-9:
            pct = None
        else:
            pct = 0.0
        out.append(
            FilaComparacion(
                year=ra.year,
                month=ra.month,
                etiqueta=ra.etiqueta,
                m3_a=ra.m3,
                dias_a=ra.dias_con_dato,
                m3_b=mb,
                dias_b=db,
                diff_b_menos_a=diff,
                pct_b_vs_a=pct,
            )
        )
    return out


def _generar_grafico(
    filas: List[FilaComparacion],
    node_a: str,
    name_a: str,
    node_b: str,
    name_b: str,
    out_path: Path,
) -> Path:
    labels = [f"{MESES_ES[r.month][:3].capitalize()}\n{r.year % 100:02d}" for r in filas]
    vals_a = [r.m3_a for r in filas]
    vals_b = [r.m3_b for r in filas]
    diffs = [r.diff_b_menos_a for r in filas]

    x = np.arange(len(labels))
    width = 0.38

    fig, ax1 = plt.subplots(figsize=(13, 5.5))
    ax1.bar(x - width / 2, vals_a, width, label=f"{name_a} ({node_a})", color="#3498DB", alpha=0.9)
    ax1.bar(x + width / 2, vals_b, width, label=f"{name_b} ({node_b})", color="#2ECC71", alpha=0.9)
    ax1.set_ylabel("Consumo mensual (m³)")
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels, fontsize=9)
    ax1.set_title(f"Comparación mensual — {node_a} vs {node_b}")
    ax1.legend(loc="upper left", fontsize=8)

    ax2 = ax1.twinx()
    ax2.plot(x, diffs, color="#E74C3C", marker="o", linewidth=2, label=f"Diferencia ({node_b} − {node_a})")
    ax2.axhline(0, color="#7F8C8D", linewidth=0.8, linestyle="--")
    ax2.set_ylabel(f"Δ m³ ({node_b} − {node_a})")
    ax2.legend(loc="upper right", fontsize=8)

    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=200)
    plt.close(fig)
    return out_path


def _generar_excel(
    filas: List[FilaComparacion],
    node_a: str,
    name_a: str,
    node_b: str,
    name_b: str,
    desde: Tuple[int, int],
    hasta: Tuple[int, int],
    out_path: Path,
) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Comparación"

    ws["A1"] = "Comparación mensual de consumo"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = f"Nodo A: {name_a} ({node_a})"
    ws["A3"] = f"Nodo B: {name_b} ({node_b})"
    ws["A4"] = f"Periodo: {_etiqueta_mes(desde[0], desde[1])} a {_etiqueta_mes(hasta[0], hasta[1])}"
    ws["A5"] = f"Fuente: API WES totalM3 diario ({acl_node_base_url()})"

    headers = [
        "Mes",
        "Año",
        "Etiqueta",
        f"m³ {node_a}",
        f"Días {node_a}",
        f"m³ {node_b}",
        f"Días {node_b}",
        f"Δ m³ ({node_b} − {node_a})",
        f"Δ % ({node_b} vs {node_a})",
    ]
    hr = 7
    header_fill = PatternFill("solid", fgColor="1F4788")
    header_font = Font(color="FFFFFF", bold=True)
    for col, h in enumerate(headers, start=1):
        c = ws.cell(row=hr, column=col, value=h)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="center", wrap_text=True)

    for i, r in enumerate(filas, start=hr + 1):
        ws.cell(row=i, column=1, value=MESES_ES[r.month].capitalize())
        ws.cell(row=i, column=2, value=r.year)
        ws.cell(row=i, column=3, value=r.etiqueta)
        ws.cell(row=i, column=4, value=r.m3_a)
        ws.cell(row=i, column=5, value=r.dias_a)
        ws.cell(row=i, column=6, value=r.m3_b)
        ws.cell(row=i, column=7, value=r.dias_b)
        ws.cell(row=i, column=8, value=r.diff_b_menos_a)
        ws.cell(
            row=i,
            column=9,
            value=f"{r.pct_b_vs_a:.1f}%" if r.pct_b_vs_a is not None else "—",
        )

    tr = hr + len(filas) + 2
    ws.cell(row=tr, column=3, value="TOTAL periodo").font = Font(bold=True)
    ws.cell(row=tr, column=4, value=round(sum(r.m3_a for r in filas), 2)).font = Font(bold=True)
    ws.cell(row=tr, column=6, value=round(sum(r.m3_b for r in filas), 2)).font = Font(bold=True)
    ws.cell(row=tr, column=8, value=round(sum(r.diff_b_menos_a for r in filas), 2)).font = Font(bold=True)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 15
    ws.column_dimensions["C"].width = 20

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def _generar_word(
    filas: List[FilaComparacion],
    node_a: str,
    name_a: str,
    node_b: str,
    name_b: str,
    desde: Tuple[int, int],
    hasta: Tuple[int, int],
    chart_path: Optional[Path],
    out_path: Path,
) -> None:
    doc = Document()
    add_logo_to_header(doc)

    t = doc.add_heading("Comparación mensual entre nodos", level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Nodo A: {name_a} ({node_a})")
    doc.add_paragraph(f"Nodo B: {name_b} ({node_b})")
    doc.add_paragraph(
        f"Periodo: {_etiqueta_mes(desde[0], desde[1])} — {_etiqueta_mes(hasta[0], hasta[1])}"
    )
    doc.add_paragraph(
        "Consumo mensual = suma de totalM3 diarios (API WES). "
        f"La diferencia indica cuántos m³ consume más el nodo B ({node_b}) que el nodo A ({node_a}) en cada mes."
    )

    if chart_path and chart_path.exists():
        doc.add_paragraph("")
        doc.add_picture(str(chart_path), width=Inches(6.4))

    doc.add_paragraph("")
    doc.add_heading("Tabla comparativa", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Mes"
    hdr[1].text = f"{name_a}\n(m³)"
    hdr[2].text = f"{name_b}\n(m³)"
    hdr[3].text = f"Δ m³\n(B − A)"
    hdr[4].text = "Δ %"

    total_a = sum(r.m3_a for r in filas)
    total_b = sum(r.m3_b for r in filas)

    for r in filas:
        row = table.add_row().cells
        row[0].text = r.etiqueta
        row[1].text = format_number_chilean(r.m3_a, 1)
        row[2].text = format_number_chilean(r.m3_b, 1)
        row[3].text = format_number_chilean(r.diff_b_menos_a, 1)
        row[4].text = f"{r.pct_b_vs_a:.1f}%" if r.pct_b_vs_a is not None else "—"

    doc.add_paragraph("")
    doc.add_heading("Resumen del periodo", level=1)
    doc.add_paragraph(
        f"Acumulado {name_a}: {format_number_chilean(total_a, 1)} m³. "
        f"Acumulado {name_b}: {format_number_chilean(total_b, 1)} m³. "
        f"Diferencia total (B − A): {format_number_chilean(total_b - total_a, 1)} m³."
    )

    mayor_b = max(filas, key=lambda r: r.diff_b_menos_a)
    mayor_a = min(filas, key=lambda r: r.diff_b_menos_a)
    doc.add_paragraph(
        f"El mes con mayor ventaja de {name_b} fue {mayor_b.etiqueta} "
        f"(+{format_number_chilean(mayor_b.diff_b_menos_a, 1)} m³). "
        f"El mes con mayor ventaja de {name_a} fue {mayor_a.etiqueta} "
        f"({format_number_chilean(mayor_a.diff_b_menos_a, 1)} m³)."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def generar_comparacion_mensual_dos_nodos(
    node_a: str,
    node_b: str,
    desde: Tuple[int, int],
    hasta: Tuple[int, int],
    out_dir: Optional[Path] = None,
) -> Tuple[Path, Path, Path]:
    meses = _iter_meses(desde, hasta)
    start = date(desde[0], desde[1], 1)
    last_day = calendar.monthrange(hasta[0], hasta[1])[1]
    end = date(hasta[0], hasta[1], last_day)

    name_a = _nombre_nodo(node_a)
    name_b = _nombre_nodo(node_b)

    filas_a = _agregar_por_mes(_fetch_medidas_diarias(node_a, start, end), meses)
    filas_b = _agregar_por_mes(_fetch_medidas_diarias(node_b, start, end), meses)
    filas = _filas_comparacion(filas_a, filas_b)

    base = out_dir or (Path("reports") / "comparacion_mensual_dos_nodos" / f"{node_a}_vs_{node_b}")
    base.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    stem = (
        f"comparacion_mensual_{node_a}_vs_{node_b}_"
        f"{desde[0]}{desde[1]:02d}_{hasta[0]}{hasta[1]:02d}_{ts}"
    )

    xlsx = base / f"{stem}.xlsx"
    docx = base / f"{stem}.docx"
    png = base / f"{stem}.png"

    _generar_excel(filas, node_a, name_a, node_b, name_b, desde, hasta, xlsx)
    _generar_grafico(filas, node_a, name_a, node_b, name_b, png)
    _generar_word(filas, node_a, name_a, node_b, name_b, desde, hasta, png, docx)

    return xlsx, docx, png


def main() -> int:
    parser = argparse.ArgumentParser(description="Comparación mensual entre dos nodos WES")
    parser.add_argument("--node-a", default="000025-19")
    parser.add_argument("--node-b", default="000025-02")
    parser.add_argument("--desde", default="2025-04", help="Mes inicio YYYY-MM")
    parser.add_argument("--hasta", default="2026-04", help="Mes fin YYYY-MM")
    parser.add_argument("--out-dir", default="")
    args = parser.parse_args()

    desde = _parse_yyyy_mm(args.desde)
    hasta = _parse_yyyy_mm(args.hasta)
    if desde > hasta:
        print("[ERROR] --desde debe ser anterior o igual a --hasta")
        return 1

    out_dir = Path(args.out_dir) if args.out_dir.strip() else None
    xlsx, docx, png = generar_comparacion_mensual_dos_nodos(
        args.node_a.strip(),
        args.node_b.strip(),
        desde,
        hasta,
        out_dir,
    )

    print("=" * 72)
    print("COMPARACIÓN MENSUAL DOS NODOS")
    print("=" * 72)
    print(f"A: {args.node_a} ({_nombre_nodo(args.node_a)})")
    print(f"B: {args.node_b} ({_nombre_nodo(args.node_b)})")
    print(f"Periodo: {args.desde} a {args.hasta}")
    print(f"Excel:   {xlsx.resolve()}")
    print(f"Word:    {docx.resolve()}")
    print(f"Gráfico: {png.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
