"""
Reporte agregado Parque Arauco (000025): todos los puntos de monitoreo,
con sección inicial que responde:
  0) Datos más clave (KPI: totales, días, picos, por mall, top 3 puntos)
  1) Qué está instalado / fechas (según API disponible)
  2) Quién está habilitado (alcance según datos disponibles)
  3) Hallazgos de consumo (resumen + remisión al cuerpo del reporte)

Uso:
  python generar_pa_agregado_todos_puntos.py
  python generar_pa_agregado_todos_puntos.py --desde 11/03/2026 --hasta 31/03/2026
"""

from __future__ import annotations

import argparse
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import requests

ENTITY_BASE = "http://104.248.53.141:7001/wes/api/acl-entities/v1"

SCRIPT_DIR = Path(__file__).resolve().parent

from generar_reporte_word import (
    format_number_chilean,
    generate_aggregated_report,
    get_mall_name_for_parque_arauco,
)
from generar_reportes_y_ppt_mall_maipu import obtener_datos_agregados
from lista_contactos_reportes import CONTACTOS_REPORTES, CORREOS_AUTORIZADOS


def _kpi_fechas(configuration: Any) -> tuple[str, str]:
    if not isinstance(configuration, dict):
        return "—", "—"
    kpi = configuration.get("kpi")
    if not kpi or not isinstance(kpi, list) or not kpi:
        return "—", "—"
    first = kpi[0]
    if not isinstance(first, dict):
        return "—", "—"
    cre = str(first.get("creationDate", "") or "").strip() or "—"
    exp = str(first.get("expirationDate", "") or "").strip() or "—"
    return cre, exp


def _cargar_nodos_pa() -> Dict[str, Any]:
    r = requests.get(f"{ENTITY_BASE}/companies/000025", timeout=60)
    r.raise_for_status()
    return r.json()


def _parrafo_hallazgos_q3(
    datos: dict,
    periodo: str,
    alcance_sujeto: Optional[str] = None,
) -> str:
    """
    Narrativa larga para la pregunta 3: consumo agregado, día máx/mín global,
    ranking de puntos y remisión al cuerpo del reporte.

    alcance_sujeto: frase de sujeto para la síntesis (ej. Mall Maipú).
        Por defecto: todos los puntos PA del reporte.
    """
    agg = datos.get("aggregate_summary") or {}
    total = float(agg.get("total") or 0)
    dias = int(agg.get("dias") or 0)
    prom = float(agg.get("promedio_diario") or 0)
    max_p = agg.get("max")
    min_p = agg.get("min")

    max_txt = "—"
    min_txt = "—"
    if max_p is not None and hasattr(max_p, "date") and hasattr(max_p, "total_m3"):
        max_txt = (
            f"{max_p.date.strftime('%d-%m-%Y')} con "
            f"{format_number_chilean(float(max_p.total_m3), 1)} m³"
        )
    if min_p is not None and hasattr(min_p, "date") and hasattr(min_p, "total_m3"):
        min_txt = (
            f"{min_p.date.strftime('%d-%m-%Y')} con "
            f"{format_number_chilean(float(min_p.total_m3), 1)} m³"
        )

    nodes_summary: List[Dict[str, Any]] = list(datos.get("nodes_summary") or [])
    ranked: List[tuple[float, str, str]] = []
    for n in nodes_summary:
        t = float((n.get("summary") or {}).get("total") or 0)
        ranked.append((t, str(n.get("node_id") or ""), str(n.get("node_name") or "")))
    ranked.sort(key=lambda x: -x[0])

    n_puntos = len(ranked)
    if total > 0 and ranked:
        top_lines = []
        for i, (t, nid, nombre) in enumerate(ranked[:5], 1):
            pct = (t / total) * 100.0 if total > 0 else 0.0
            top_lines.append(
                f"{i}) {nombre} ({nid}): {format_number_chilean(t, 1)} m³ "
                f"({format_number_chilean(pct, 1)} % del total agregado)"
            )
        top_bloque = "\n".join(top_lines)
    else:
        top_bloque = "— (sin ranking: no hay consumo acumulado o no hay datos por punto)."

    menores = [
        f"{nombre} ({nid}): {format_number_chilean(t, 1)} m³"
        for t, nid, nombre in sorted(ranked, key=lambda x: x[0])[:5]
    ]

    sin_o_cero = sum(1 for t, _, _ in ranked if t <= 0)
    sujeto = (
        alcance_sujeto
        or "la red Parque Arauco (todos los puntos de monitoreo incluidos en este reporte)"
    )
    partes = [
        "3) ¿Qué hallazgos encontramos? (Análisis de consumo)",
        "",
        f"Período analizado: {periodo}.",
        "",
        "Síntesis ejecutiva",
        f"En el período considerado, {sujeto} acumuló un volumen total de {format_number_chilean(total, 1)} m³. "
        f"Contando días con medición a nivel agregado, se observan {dias} día(s) con datos; "
        f"el promedio diario agregado resulta de {format_number_chilean(prom, 1)} m³/día "
        f"(total de consumo diario sumado entre puntos, dividido por días con registro). "
        f"El día de mayor consumo global (suma de todos los puntos) fue {max_txt}; "
        f"el día de menor consumo global fue {min_txt}.",
        "",
        f"Distribución por punto de monitoreo (universo: {n_puntos} puntos). "
        f"Los cinco mayores aportes acumulados del período son:",
        top_bloque,
        "",
        "Los puntos con menor consumo acumulado en el período (referencia para priorizar revisión o "
        "confirmar operación en cero / sin datos) son:",
        "\n".join(menores) if menores else "—",
        "",
        f"Puntos con consumo acumulado cero o sin registros en el período (referencia): {sin_o_cero}.",
        "",
        "Profundización de hallazgos",
        "El detalle de métricas por punto (consumo diario y nocturno, alertas, proyecciones y gráficas) "
        "se desarrolla en las secciones siguientes de este mismo documento, incluyendo análisis de "
        "periodicidad de alertas y comparaciones cuando el modelo lo permite.",
    ]
    return "\n".join(partes)


def _bloque_datos_clave(datos: dict, periodo: str, n_puntos_reporte: int) -> str:
    """
    Resumen numérico destacado: totales, picos diarios, consumo por mall, top 3 nodos.
    """
    agg = datos.get("aggregate_summary") or {}
    total = float(agg.get("total") or 0)
    dias = int(agg.get("dias") or 0)
    prom = float(agg.get("promedio_diario") or 0)
    max_p = agg.get("max")
    min_p = agg.get("min")

    max_txt = "—"
    min_txt = "—"
    if max_p is not None and hasattr(max_p, "date") and hasattr(max_p, "total_m3"):
        max_txt = (
            f"{max_p.date.strftime('%d-%m-%Y')} — "
            f"{format_number_chilean(float(max_p.total_m3), 1)} m³ (suma todos los puntos)"
        )
    if min_p is not None and hasattr(min_p, "date") and hasattr(min_p, "total_m3"):
        min_txt = (
            f"{min_p.date.strftime('%d-%m-%Y')} — "
            f"{format_number_chilean(float(min_p.total_m3), 1)} m³ (suma todos los puntos)"
        )

    nodes_summary: List[Dict[str, Any]] = list(datos.get("nodes_summary") or [])
    ranked: List[tuple[float, str, str]] = []
    for n in nodes_summary:
        t = float((n.get("summary") or {}).get("total") or 0)
        ranked.append((t, str(n.get("node_id") or ""), str(n.get("node_name") or "")))
    ranked.sort(key=lambda x: -x[0])
    sin_o_cero = sum(1 for t, _, _ in ranked if t <= 0)

    by_mall: Dict[str, float] = defaultdict(float)
    for t, nid, nombre in ranked:
        mall = (get_mall_name_for_parque_arauco(nid, nombre) or "").strip() or "Sin mall"
        by_mall[mall] += t
    malls_sorted = sorted(by_mall.items(), key=lambda x: -x[1])

    lineas_mall: List[str] = []
    for mall, v in malls_sorted:
        if total > 0:
            pct = (v / total) * 100.0
            lineas_mall.append(
                f"  • {mall}: {format_number_chilean(v, 1)} m³ "
                f"({format_number_chilean(pct, 1)} % del total agregado)"
            )
        else:
            lineas_mall.append(f"  • {mall}: {format_number_chilean(v, 1)} m³")

    top3: List[str] = []
    for i, (t, nid, nombre) in enumerate(ranked[:3], 1):
        pct = (t / total) * 100.0 if total > 0 else 0.0
        top3.append(
            f"  {i}) {nombre} ({nid}): {format_number_chilean(t, 1)} m³ "
            f"({format_number_chilean(pct, 1)} % del total)"
        )

    partes = [
        "0) Datos más clave (resumen ejecutivo numérico)",
        "",
        f"Período analizado: {periodo}.",
        "",
        f"• Puntos de monitoreo incluidos en este informe: {n_puntos_reporte}.",
        f"• Consumo total agregado (suma de todos los puntos): {format_number_chilean(total, 1)} m³.",
        f"• Días con registro a nivel agregado (serie diaria combinada): {dias}.",
        f"• Promedio diario agregado: {format_number_chilean(prom, 1)} m³/día.",
        f"• Día de mayor consumo global: {max_txt}.",
        f"• Día de menor consumo global: {min_txt}.",
        f"• Puntos con consumo acumulado cero o sin registros en el período: {sin_o_cero}.",
        "",
        "Consumo por mall (agrupación operativa WES)",
        "\n".join(lineas_mall) if lineas_mall else "  —",
        "",
        "Tres mayores puntos por volumen en el período",
        "\n".join(top3) if top3 else "  —",
        "",
        "Nota: el cuerpo del documento incluye tablas, gráficos de comparación, alertas y métricas "
        "nocturnas por punto. Use este bloque como lectura rápida de cifras consolidadas.",
    ]
    return "\n".join(partes)


def _contactos_pa_texto() -> str:
    lineas: List[str] = []
    lineas.append(
        "Usuarios con permiso de acceso a la plataforma WES (cliente) no están disponibles vía la API "
        "usada por estos scripts; deben confirmarse en el administrador de usuarios de WES / backoffice."
    )
    lineas.append("")
    lineas.append("Como referencia operativa, contactos habituales en reportes WES con interés en Parque Arauco:")
    for clave, c in CONTACTOS_REPORTES.items():
        emp = c.get("empresas_interes") or []
        if not isinstance(emp, list) or "Parque Arauco" not in emp:
            continue
        lineas.append(f"  - {c.get('nombre_completo', clave)} <{c.get('email', '')}>")
    lineas.append("")
    lineas.append("Correos autorizados para solicitar reportes automáticos (lista interna del agente):")
    for row in CORREOS_AUTORIZADOS:
        lineas.append(f"  - {row.get('email', '')}")
    return "\n".join(lineas)


def _prepend_secciones(doc_path: Path, payload: dict, datos: dict, periodo: str) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import RGBColor

    doc = Document(str(doc_path))
    if not doc.paragraphs:
        doc.add_paragraph("")

    nodes = sorted(payload.get("nodes") or [], key=lambda x: str(x.get("nodeId", "")))

    q3 = _parrafo_hallazgos_q3(datos, periodo)

    q2_full = "2) ¿Quién está habilitado?\n\n" + _contactos_pa_texto()

    q1_lines = [
        "1) ¿Qué está instalado? Fecha de instalación / recepción",
        "Por cada punto: nombre en plataforma y fecha de inicio del contrato/tarifa KPI en API (creationDate del primer KPI). "
        "La instalación física y la recepción de equipamiento no figuran en esta API; completar con actas / inventario de campo.",
        "",
    ]
    for n in nodes:
        nid = str(n.get("nodeId", "")).strip()
        nombre = str(n.get("name", "")).strip()
        cfg = n.get("configuration") or {}
        cre, exp = _kpi_fechas(cfg)
        q1_lines.append(f"  • {nid} — {nombre}")
        q1_lines.append(f"    Referencia inicio servicio (KPI): {cre} | Fin vigencia tarifa: {exp}")
        q1_lines.append("    Instalación física / recepción: (completar en terreno — no en API)")
        q1_lines.append("")
    q1 = "\n".join(q1_lines)

    intro = (
        "PARQUE ARAUCO — Respuestas guía (todas las mediciones)\n"
        "Nota metodológica: fechas contractuales/KPI desde API de entidades; análisis de consumo del período en las "
        "secciones posteriores de este documento."
    )

    def _add_block(texto: str) -> None:
        p0 = doc.paragraphs[0]
        p = p0.insert_paragraph_before("")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        lineas = texto.split("\n")
        titulos_sec = {
            "Síntesis ejecutiva",
            "Profundización de hallazgos",
            "Consumo por mall (agrupación operativa WES)",
            "Tres mayores puntos por volumen en el período",
        }
        for i, line in enumerate(lineas):
            run = p.add_run(("\n" if i else "") + line)
            s = line.strip()
            run.bold = bool(
                (
                    i == 0
                    and (
                        s.startswith("PARQUE ARAUCO")
                        or s.startswith(("0)", "1)", "2)", "3)"))
                    )
                )
                or s in titulos_sec
            )
            run.font.color.rgb = RGBColor(0, 0, 0)

    datos_clave = _bloque_datos_clave(datos, periodo, len(nodes))

    # Orden inserción: el último insertado queda arriba del todo → datos clave al inicio.
    for bloque in (q3, q2_full, q1, intro, datos_clave):
        _add_block(bloque)

    doc.save(str(doc_path))


def main() -> int:
    parser = argparse.ArgumentParser(
        description="PA: reporte Word agregado de todos los puntos + datos clave y guías Q1–Q3"
    )
    parser.add_argument("--desde", default="11/03/2026", help="DD/MM/YYYY inicio período")
    parser.add_argument(
        "--hasta",
        default=None,
        help="DD/MM/YYYY fin período (default: hoy)",
    )
    args = parser.parse_args()

    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass

    hasta = args.hasta or datetime.now().strftime("%d/%m/%Y")

    payload = _cargar_nodos_pa()
    nodes = payload.get("nodes") or []
    node_ids = [str(n.get("nodeId", "")).strip() for n in nodes if n.get("nodeId")]
    node_ids = sorted(set(node_ids))
    if not node_ids:
        print("[ERROR] Sin nodos en empresa 000025")
        return 1

    periodo = f"{args.desde} a {hasta}"
    print(f"[INFO] Generando reporte agregado PA: {len(node_ids)} nodos, período {periodo}")
    print(f"[INFO] Salida bajo: {(SCRIPT_DIR / 'reports').resolve()}")

    out = generate_aggregated_report(
        company_id="000025",
        node_ids=node_ids,
        start_date=args.desde,
        end_date=hasta,
        output_dir=str(SCRIPT_DIR / "reports"),
        fuente_agua_id=None,
        mall_name="PA_Todos_Los_Puntos",
        apply_exclusions=False,
        generate_ppt=False,
    )
    out_path = Path(out).resolve()
    print(f"[OK] Reporte base: {out_path}")

    print("[INFO] Resumen consumo agregado (API)...")
    datos = obtener_datos_agregados(node_ids, args.desde, hasta)

    print("[INFO] Insertando datos clave + secciones guía (0–3) al inicio del Word...")
    _prepend_secciones(out_path, payload, datos, periodo)
    print(f"[OK] Documento completo: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
