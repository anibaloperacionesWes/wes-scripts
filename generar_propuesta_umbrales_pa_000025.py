# -*- coding: utf-8 -*-
"""
Tabla baseline 90d + umbral +25% para compañía 000025 (Parque Arauco).
Genera Excel/Word y envía a Aníbal.
"""

from __future__ import annotations

import os
import smtplib
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime, timedelta
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import requests
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from docx import Document
from docx.shared import Pt, Inches, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))

from generar_reporte_word import format_number_chilean as fn
from generar_reportes_y_ppt_mall_maipu import obtener_datos_agregados

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Parque_Arauco" / "umbrales_consumo"
ENTITY_BASE = "http://104.248.53.141:7001/wes/api/acl-entities/v1"

COMPANY_ID = "000025"
REF = date(2026, 7, 14)
DIAS_BASELINE = 90
MULT = 1.25
DESDE = (REF - timedelta(days=DIAS_BASELINE - 1)).strftime("%d/%m/%Y")
HASTA = REF.strftime("%d/%m/%Y")

DESTINATARIO = "anibal.aoperaciones@wes.cl"
SMTP_USUARIO = "agente.ia@wes.cl"
SMTP_SERVIDOR = "smtp.gmail.com"
SMTP_PUERTO = 587


def _smtp_password() -> str:
    p = (
        os.environ.get("WES_GMAIL_APP_PASSWORD", "").strip()
        or os.environ.get("WES_SMTP_PASSWORD", "").strip()
    )
    if p:
        return p.replace(" ", "").strip()
    f = ROOT / "gmail_oauth" / "app_password.txt"
    if f.is_file():
        for line in f.read_text(encoding="utf-8", errors="replace").splitlines():
            line = line.strip()
            if line and not line.startswith("#"):
                return line.replace(" ", "").strip()
    return ""


def listar_nodos_000025() -> list[dict]:
    url = f"{ENTITY_BASE}/companies/{COMPANY_ID}"
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    data = r.json()
    company_name = data.get("name", "Parque Arauco")
    rows = []
    for node in data.get("nodes") or []:
        nid = (node.get("nodeId") or "").strip()
        nm = (node.get("name") or "").strip()
        if nid:
            rows.append({"nodeId": nid, "nodeName": nm or nid, "companyName": company_name})
    rows.sort(key=lambda x: x["nodeId"])
    return rows


def promedio_diario(node_id: str) -> tuple[float | None, int, float | None]:
    """Retorna (promedio_diario, dias_con_data, max_diario) o (None, 0, None)."""
    try:
        d = obtener_datos_agregados([node_id], DESDE, HASTA)
        ns = (d.get("nodes_summary") or [None])[0]
        if not ns:
            return None, 0, None
        daily: dict[date, float] = {}
        for m in ns.get("measures") or []:
            d0 = m.date.date()
            daily[d0] = daily.get(d0, 0.0) + float(m.total_m3)
        if not daily:
            return None, 0, None
        vals = list(daily.values())
        return sum(vals) / len(vals), len(vals), max(vals)
    except Exception as e:
        print(f"[WARN] {node_id}: {e}")
        return None, 0, None


def calcular_filas(nodos: list[dict]) -> list[dict]:
    out: list[dict] = []

    def _one(n: dict) -> dict:
        nid = n["nodeId"]
        print(f"[data] {nid} {n['nodeName']}...", flush=True)
        mean, dias, mx = promedio_diario(nid)
        umbral = round(mean * MULT, 1) if mean is not None else None
        return {
            "nodeId": nid,
            "nodeName": n["nodeName"],
            "companyName": n["companyName"],
            "dias": dias,
            "baseline": round(mean, 1) if mean is not None else None,
            "max_periodo": round(mx, 1) if mx is not None else None,
            "umbral": umbral,
            "mult": MULT,
        }

    with ThreadPoolExecutor(max_workers=8) as ex:
        futs = {ex.submit(_one, n): n for n in nodos}
        for fut in as_completed(futs):
            out.append(fut.result())
    out.sort(key=lambda r: r["nodeId"])
    return out


def escribir_excel(rows: list[dict], path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Umbrales 000025"
    headers = [
        "Node ID",
        "Punto",
        "Días con data (90d)",
        "Baseline promedio diario (m³)",
        "Máximo diario en periodo (m³)",
        "Multiplicador",
        "Umbral recomendado (m³/día)",
        "Regla",
    ]
    ws.append(headers)
    header_fill = PatternFill("solid", fgColor="1F4E79")
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = header_fill
        c.alignment = Alignment(wrap_text=True, vertical="center")

    for r in rows:
        ws.append([
            r["nodeId"],
            r["nodeName"],
            r["dias"],
            r["baseline"] if r["baseline"] is not None else "Sin datos",
            r["max_periodo"] if r["max_periodo"] is not None else "Sin datos",
            f"× {MULT:.2f} (+{(MULT - 1) * 100:.0f}%)",
            r["umbral"] if r["umbral"] is not None else "N/D",
            "promedio_90d × 1,25" if r["umbral"] is not None else "Revisar / sin baseline",
        ])

    ws.append([])
    ws.append(["Parámetros"])
    ws.append(["Compañía", COMPANY_ID])
    ws.append(["Periodo baseline", f"{DESDE} a {HASTA} ({DIAS_BASELINE} días)"])
    ws.append(["Corte", REF.strftime("%d/%m/%Y")])
    ws.append(["Fórmula", "umbral = promedio_diario_90d × 1,25"])
    ws.append([
        "Comportamiento API",
        "Cuando el acumulado del día supera el umbral (a cualquier hora) → alerta por correo",
    ])

    widths = [14, 40, 16, 22, 22, 14, 22, 28]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i) if i <= 26 else "A"].width = w
    for i, w in enumerate(widths, 1):
        from openpyxl.utils import get_column_letter
        ws.column_dimensions[get_column_letter(i)].width = w

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def escribir_word(rows: list[dict], path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Propuesta de umbrales de consumo diario — Parque Arauco (000025)", 0)
    p = doc.add_paragraph()
    p.add_run(f"Fecha de corte: {REF.strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(
        f"Ventana baseline: últimos {DIAS_BASELINE} días ({DESDE} – {HASTA})\n"
        f"Multiplicador propuesto: × {MULT:.2f} (+{(MULT - 1) * 100:.0f}% sobre el promedio diario)\n"
        f"Destinatario interno: {DESTINATARIO}"
    )

    doc.add_heading("1. Objetivo", 1)
    doc.add_paragraph(
        "Definir, para cada punto de monitoreo de la compañía 000025 (Parque Arauco), "
        "un valor único de umbral diario (m³/día) compatible con la alerta por umbral "
        "ya implementada en la API WES: cuando el consumo acumulado del día supera ese "
        "número —independiente de la hora— se dispara aviso por correo al cliente."
    )

    doc.add_heading("2. Criterio técnico recomendado", 1)
    doc.add_paragraph(
        "Baseline = promedio del consumo diario total en los últimos 90 días.\n"
        "Umbral = baseline × 1,25 (+25%)."
    )
    doc.add_paragraph(
        "No se usa el promedio de máximos diarios/mensuales como baseline, porque ese "
        "valor ya representa días atípicos altos. Multiplicarlo otra vez (+25%) retrasa "
        "la alerta. El promedio diario normal + 25% permite avisar antes, alineado al "
        "objetivo de dar tiempo de reacción al cliente del mall."
    )

    doc.add_heading("3. Por qué 90 días y +25% (y no +50%)", 1)
    bullets = [
        "Ventana 90 días: práctica habitual en utilities AMI (p. ej. referencias SFPUC usan ~90 días). Es más estable que 30 días (menos sensible a un mes atípico) y más actualizada que 6 meses.",
        "Multiplicador +25%: más temprano que el +50% usado por algunas utilities estadounidenses. En malls se busca detectar alzas operativas/fugas antes de que el exceso sea extremo; +50% avisaría más tarde.",
        "Ejemplo Estanque Norte (000025-01): baseline 90d ≈ 26,2 m³/día → umbral ≈ 32,7 m³/día. Con +50% el umbral sería ≈ 39,3 m³/día (≈ 6,6 m³ más tarde).",
        "Comparación 30d vs 90d en el mismo punto: baseline 30d ≈ 21,1 (umbral 26,4) vs 90d ≈ 26,2 (umbral 32,7). Se elige 90d por estabilidad para una propuesta única a jefatura/clientes.",
        "Compatibilidad API: un solo número por punto; sin distinción lunes/finde; la alerta se evalúa sobre el acumulado del día al cruzar el umbral.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4. Tabla de valores recomendados", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Node ID"
    hdr[1].text = "Punto"
    hdr[2].text = "Baseline 90d (m³/d)"
    hdr[3].text = "Máx. periodo (m³/d)"
    hdr[4].text = "Umbral +25% (m³/d)"

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["nodeId"]
        cells[1].text = r["nodeName"]
        cells[2].text = fn(r["baseline"], 1) if r["baseline"] is not None else "Sin datos"
        cells[3].text = fn(r["max_periodo"], 1) if r["max_periodo"] is not None else "—"
        cells[4].text = fn(r["umbral"], 1) if r["umbral"] is not None else "N/D"

    doc.add_heading("5. Uso propuesto ante jefatura / clientes", 1)
    doc.add_paragraph(
        "1) Validar esta regla (90 días × 1,25) como estándar Parque Arauco.\n"
        "2) Cargar el umbral recomendado por punto en la configuración de alerta de la API.\n"
        "3) Comunicar al cliente que recibirá correo el mismo día en que el acumulado "
        "supere el umbral, sin esperar fin de día ni lectura de boleta.\n"
        "4) Revisar trimestralmente el baseline (recalcular 90d) para puntos con cambios "
        "operativos (obras, reconfiguración de red, nuevos medidores)."
    )

    doc.add_heading("6. Limitaciones", 1)
    doc.add_paragraph(
        "• Puntos sin data suficiente en la ventana aparecen como «Sin datos» / N/D.\n"
        "• Puntos dados de baja, en reparación o sin operación típica (p. ej. algunos "
        "baños retirados / impulsiones en OC) deben revisarse caso a caso antes de activar.\n"
        "• Un umbral único diario no distingue día hábil vs fin de semana; es la "
        "restricción actual del backend de alertas."
    )

    doc.add_paragraph("")
    foot = doc.add_paragraph()
    foot.add_run(
        "Documento generado por Sistema WES — uso interno para definición de umbrales."
    ).italic = True

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def enviar(files: list[Path]) -> None:
    pw = _smtp_password()
    if not pw:
        raise RuntimeError("Falta contraseña SMTP")

    fecha = datetime.now().strftime("%d-%m-%Y")
    msg = MIMEMultipart()
    msg["From"] = SMTP_USUARIO
    msg["To"] = DESTINATARIO
    msg["Subject"] = (
        f"WES — Propuesta umbrales consumo diario Parque Arauco 000025 "
        f"(90d × 1,25) — {fecha}"
    )
    body = (
        "Hola Aníbal,\n\n"
        "Adjunto la propuesta técnica de umbrales máximos diarios para todos los puntos "
        "de la compañía 000025 (Parque Arauco), lista para validar con jefatura y luego "
        "proponer a los clientes de los malls.\n\n"
        "Regla aplicada:\n"
        f"- Baseline = promedio consumo diario últimos {DIAS_BASELINE} días "
        f"({DESDE} a {HASTA})\n"
        f"- Umbral = baseline × {MULT:.2f} (+{(MULT - 1) * 100:.0f}%)\n"
        "- Alerta API: cuando el acumulado del día supera el umbral (cualquier hora) "
        "→ correo\n\n"
        "Archivos:\n"
        "- Excel con la tabla completa\n"
        "- Word con explicación técnica y motivos\n\n"
        "Ejemplo de referencia (Estanque Norte 000025-01): "
        "baseline ≈ 26,2 m³/día → umbral ≈ 32,7 m³/día.\n\n"
        "Saludos,\n"
        "Sistema WES\n"
    )
    msg.attach(MIMEText(body, "plain", "utf-8"))

    for f in files:
        with open(f, "rb") as fh:
            part = MIMEApplication(fh.read())
        part.add_header("Content-Disposition", "attachment", filename=f.name)
        msg.attach(part)

    with smtplib.SMTP(SMTP_SERVIDOR, SMTP_PUERTO) as server:
        server.starttls()
        server.login(SMTP_USUARIO, pw)
        server.send_message(msg, to_addrs=[DESTINATARIO])


def main() -> int:
    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument("--enviar", action="store_true", help="Enviar correo a Aníbal")
    args = ap.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Listando nodos {COMPANY_ID}...")
    nodos = listar_nodos_000025()
    print(f"Nodos: {len(nodos)}")
    rows = calcular_filas(nodos)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    xlsx = OUT_DIR / f"Umbrales_000025_90d_x125_{stamp}.xlsx"
    docx = OUT_DIR / f"Propuesta_umbrales_000025_90d_x125_{stamp}.docx"
    escribir_excel(rows, xlsx)
    escribir_word(rows, docx)
    print(f"[OK] Excel: {xlsx}")
    print(f"[OK] Word:  {docx}")

    ok = sum(1 for r in rows if r["umbral"] is not None)
    print(f"Con umbral calculado: {ok}/{len(rows)}")

    # Guardar paths para envío posterior
    (OUT_DIR / "_last_paths.txt").write_text(
        f"{xlsx}\n{docx}\n", encoding="utf-8"
    )

    if args.enviar:
        print(f"[INFO] Enviando a {DESTINATARIO}...")
        enviar([xlsx, docx])
        print("[OK] Correo enviado.")
    else:
        print("[INFO] Archivos listos. Usa --enviar para despachar correo.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
