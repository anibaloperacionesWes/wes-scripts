"""
Genera informe Word de auditoría WES (Colegio ICCO Renca).

Portada: párrafos centrados (título, establecimiento, marca) sin tabla ni franja de color; opcional:
``--portada-png`` o ``--portada-rebeca-media``.

Estilos: por defecto un Word en ``reports/auditoria_cpa_icco`` cuyo nombre indique **auditoría para
colegios** (p. ej. ``Auditoría para colegios.docx``); si no existe, se usa el informe Rebeca Matte allí;
opcional ``--plantilla-estilos`` o .docx en ``auditoria_para_colegios/``. Si no hay plantilla, Calibri y azul WES.
El borrador (``-o``) se sobrescribe en cada ejecución salvo que uses ``--mantener-borrador-manual`` con un
``.docx`` ya existente: entonces solo se actualizan PNG y CSV; tu Word editado a mano no se pisa (sustituye
tú las figuras en el documento si quieres los datos nuevos).

**Registros de consumos** (un solo título de sección): gráfico de **barras** con vs sin control, cuadro Σ,
**perfiles Con WES vs línea base** (un gráfico por día homólogo más promedio) con **texto comparativo por día**,
CSV (serie horaria en archivo). **Resultados y Conclusiones** al final.

Uso:
  python generar_informe_auditoria_icco_renca_word.py
  python generar_informe_auditoria_icco_renca_word.py -o reports/reporte de auditoria/auditoria_puntos_renca_abril_2026/Auditoria ICCO abril/Auditoria_ICCO.docx
  python generar_informe_auditoria_icco_renca_word.py -o .../Borrador_auditoria_ICCO_Renca_abril_2026.docx --solo-consolidado
  python generar_informe_auditoria_icco_renca_word.py --mantener-borrador-manual
  python generar_informe_auditoria_icco_renca_word.py --portada-png ruta/portada.png
  python generar_informe_auditoria_icco_renca_word.py --portada-rebeca-media
"""

from __future__ import annotations

import argparse
import unicodedata
import zipfile
from datetime import date, datetime, timezone
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from auditoria_cpa_icco_renca_grafico import (
    ComparacionDia24h,
    HORA_FIN_EXCL,
    HORA_INICIO,
    NOMBRE_PUNTO,
    NODE_DEFAULT,
    PERIODO_AUDITORIA,
    PERIODO_REFERENCIA,
    Periodo,
    ResultadoAuditoriaCpa,
    ejecutar_auditoria_cpa_icco,
    generar_png_barras_con_sin,
)
from generar_reporte_word import add_picture_with_pagination, format_number_chilean
from generar_graficos_comparativos_desde_excel_consolidado import (
    generar_png_barras_rejilla_totales,
    leer_matriz_consolidado,
    totales_rejilla_desde_excel_consolidado,
)

# Salida por defecto del informe (carpeta auditoría abril 2026).
_DEFAULT_SALIDA_INFORME_DOCX = (
    Path(__file__).resolve().parent
    / "reports"
    / "reporte de auditoria"
    / "auditoria_puntos_renca_abril_2026"
    / "Auditoria ICCO abril"
    / "Auditoria_ICCO.docx"
)

# Misma base que ``generar_excel_auditoria_consolidado_dos_periodos --toda-la-carpeta``.
_XLSX_CONSOLIDADO_CSV_API = (
    Path(__file__).resolve().parent
    / "reports"
    / "reporte de auditoria"
    / "auditoria_puntos_renca_abril_2026"
    / "Auditoria ICCO abril"
    / "consolidado_revision_todos_los_csv_descarga_api.xlsx"
)

# Estimación solo para texto tipo informe Renca (resultados esperados); ajustar si aplica.
CLP_POR_M3_REF = 1200.0

_MESES_ES = (
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)


def _rango_fechas_prosa_cl(d0: date, d1: date) -> str:
    """Rango inclusivo en prosa (Chile), p. ej. «el 23 y el 26 de marzo de 2026»."""
    if d0 == d1:
        return f"el {d0.day} de {_MESES_ES[d0.month - 1]} de {d0.year}"
    if d0.year != d1.year:
        return (
            f"el {d0.day} de {_MESES_ES[d0.month - 1]} de {d0.year} "
            f"y el {d1.day} de {_MESES_ES[d1.month - 1]} de {d1.year}"
        )
    if d0.month == d1.month:
        return f"el {d0.day} y el {d1.day} de {_MESES_ES[d0.month - 1]} de {d0.year}"
    return (
        f"el {d0.day} de {_MESES_ES[d0.month - 1]} y el {d1.day} de {_MESES_ES[d1.month - 1]} "
        f"de {d0.year}"
    )

# Tipografía sin plantilla: alineado a informes Word/WES del repositorio (Calibri, azul títulos).
_BODY_FONT = "Calibri"
_HEADING_FONT = "Calibri"
_HEADING_COLOR = RGBColor(31, 71, 136)

# --- Portada generada (editar aquí el “informe nuevo”) ---
PORTADA_TITULO = "Auditoría ICCO"
# Referencia del borrador (aparece en portada generada y coincide con el .docx por defecto).
PORTADA_REFERENCIA_BORRADOR = "Auditoria_ICCO"
# Una o dos líneas para el establecimiento (estilo referencia Renca).
PORTADA_ESTABLECIMIENTO_LINEA1 = "Colegio ICCO Renca"
PORTADA_ESTABLECIMIENTO_LINEA2 = ""  # ej. "Renca"; vacío = solo una línea

PORTADA_AUTOR_NOMBRE = "Aníbal Aranda Alvarado"
PORTADA_AUTOR_CARGO = "Jefe de mantenimiento y operación"
PORTADA_AUTOR_FONO = "9-75595835"

# Carpeta de recursos (antes ``assets``): logos, plantillas opcionales, caché portada Rebeca.
DIR_AUDITORIA_COLEGIOS = "auditoria_para_colegios"

# Logo opcional centrado (si existe): auditoria_para_colegios/wes_marca_centro.png — si no, texto WES.
WES_MARCA_ARCHIVO = "wes_marca_centro.png"

# Portada alternativa: recortes del .docx Rebeca Matte (--portada-rebeca-media).
PORTADA_IMAGENES = ("image14.png", "image2.png", "image11.png")

# Tabla estilo referencia adjunta: encabezado azul y tipografía homogénea.
_TABLA_AUDITORIA_FILL_ENCABEZADO = "D9E1F2"
_TABLA_AUDITORIA_FILL_COL_CONDICION = "FFFFFF"
_TABLA_CUANTITATIVOS_FILL_HEADER = "D9E1F2"
# Borde de cuadrícula (tema «Accent» / armoniza con celeste D9E1F2; no gris neutro).
_TABLA_BORDE_HEX = "8FAADC"


def _set_table_grid_borders_thin(
    table, *, sz_eighths_pt: str = "4", color_hex: str | None = None
) -> None:
    """
    Fuerza cuadrícula fina sobre la tabla. Los estilos tipo «Light Grid» suelen aplicar líneas gruesas;
    w:sz está en octavos de punto (p. ej. 4 ≈ 0,5 pt). Color por defecto: azul de acento acorde al encabezado.
    """
    ch = color_hex if color_hex is not None else _TABLA_BORDE_HEX
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz_eighths_pt)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), ch)
        borders.append(el)
    tbl_pr.append(borders)


def _apply_metricas_clave_table_style(tbl, doc: Document) -> None:
    """Replica estilo de tabla 'Métricas clave' de reportes estándar."""
    try:
        tbl.style = "Light Grid Accent 1"
    except KeyError:
        _apply_table_style_safe(tbl, doc)
    _set_table_grid_borders_thin(tbl)


def _set_cell_shading_hex(cell, hex_fill: str) -> None:
    """Relleno de celda (w:shd) sin #; hex 6 caracteres RRGGGBB."""
    shading_xml = (
        "<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
        f'w:val="clear" w:fill="{hex_fill}"/>'
    )
    shading = parse_xml(shading_xml)
    tc_pr = cell._element.get_or_add_tcPr()
    for el in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(el)
    tc_pr.append(shading)


def _portada_generada_wes(doc: Document, root: Path) -> None:
    """Portada clásica WES: título y establecimiento centrados (sin tabla ni franja), marca, pie derecho."""
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_tit.paragraph_format.space_before = Pt(24)
    p_tit.paragraph_format.space_after = Pt(6)
    rt = p_tit.add_run(PORTADA_TITULO)
    rt.font.name = _BODY_FONT
    rt.font.size = Pt(11)
    rt.font.color.rgb = RGBColor(64, 64, 64)

    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_ref.paragraph_format.space_before = Pt(2)
    p_ref.paragraph_format.space_after = Pt(8)
    rr = p_ref.add_run(PORTADA_REFERENCIA_BORRADOR)
    rr.font.name = _BODY_FONT
    rr.font.size = Pt(10)
    rr.font.italic = True
    rr.font.color.rgb = RGBColor(90, 90, 90)

    p_est = doc.add_paragraph()
    p_est.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_est.paragraph_format.space_after = Pt(4)
    r1 = p_est.add_run(PORTADA_ESTABLECIMIENTO_LINEA1)
    r1.bold = True
    r1.font.size = Pt(17)
    r1.font.name = _HEADING_FONT
    r1.font.color.rgb = _HEADING_COLOR
    if PORTADA_ESTABLECIMIENTO_LINEA2.strip():
        p_est.add_run().add_break(WD_BREAK.LINE)
        r2 = p_est.add_run(PORTADA_ESTABLECIMIENTO_LINEA2.strip())
        r2.bold = True
        r2.font.size = Pt(17)
        r2.font.name = _HEADING_FONT
        r2.font.color.rgb = _HEADING_COLOR

    _spacer_lines(doc, 2)

    marca = root / DIR_AUDITORIA_COLEGIOS / WES_MARCA_ARCHIVO
    if marca.is_file():
        _add_centered_picture(doc, marca, width_inches=4.8, keep_with_next=False)
    else:
        p_w = doc.add_paragraph()
        p_w.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        rw = p_w.add_run("WES")
        rw.bold = True
        rw.font.size = Pt(40)
        rw.font.name = _HEADING_FONT
        rw.font.color.rgb = _HEADING_COLOR
        p_w.paragraph_format.space_after = Pt(2)
        p_tag = doc.add_paragraph()
        p_tag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r_tag = p_tag.add_run("USO EFICIENTE DE AGUA")
        r_tag.font.name = _BODY_FONT
        r_tag.font.size = Pt(10)
        r_tag.font.color.rgb = RGBColor(100, 120, 140)
        r_tag.bold = True

    _spacer_lines(doc, 10)

    for texto, negrita in (
        (PORTADA_AUTOR_NOMBRE, True),
        (PORTADA_AUTOR_CARGO, False),
        (PORTADA_AUTOR_FONO, False),
    ):
        pr = doc.add_paragraph()
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        rr = pr.add_run(texto)
        rr.bold = negrita
        rr.font.name = _BODY_FONT
        rr.font.size = Pt(11)
        rr.font.color.rgb = RGBColor(0, 0, 0)

    br = doc.add_paragraph()
    br.add_run().add_break(WD_BREAK.PAGE)


def _norm_docx_basename(name: str) -> str:
    """Nombre de archivo sin tildes, minúsculas (para emparejar 'auditoría' / 'colegios')."""
    nfd = unicodedata.normalize("NFD", name)
    stripped = "".join(ch for ch in nfd if unicodedata.category(ch) != "Mn")
    return stripped.lower()


def _reports_auditoria_cpa_dir(root: Path) -> Path:
    return root / "reports" / "auditoria_cpa_icco"


def _find_plantilla_auditoria_colegios_docx(root: Path) -> Optional[Path]:
    """
    Plantilla de estilos «auditoría para colegios»: cualquier .docx en ``reports/auditoria_cpa_icco``
    cuyo nombre contenga *auditoria* y *colegio* (ej. ``Auditoría para colegios.docx``).
    """
    d = _reports_auditoria_cpa_dir(root)
    if not d.is_dir():
        return None
    for p in sorted(d.glob("*.docx")):
        if p.name.startswith("~$"):
            continue
        fn = _norm_docx_basename(p.stem)
        if "auditoria" in fn and "colegio" in fn:
            return p
    return None


def _find_rebeca_matte_docx(root: Path) -> Optional[Path]:
    """Informe Rebeca Matte en ``reports/auditoria_cpa_icco`` (respaldo de estilos y portada media)."""
    d = _reports_auditoria_cpa_dir(root)
    if not d.is_dir():
        return None
    for p in sorted(d.glob("*.docx")):
        ln = p.name.lower()
        if ln.startswith("~$"):
            continue
        if "rebeca" in ln and "matte" in ln:
            return p
    return None


def _find_reference_docx(root: Path) -> Optional[Path]:
    """Alias: solo Rebeca Matte (p. ej. ``--portada-rebeca-media``)."""
    return _find_rebeca_matte_docx(root)


def _find_assets_icco_plantilla(root: Path) -> Optional[Path]:
    """Plantilla opcional en ``auditoria_para_colegios/`` (p. ej. ``plantilla_informe_icco.docx``)."""
    ad = root / DIR_AUDITORIA_COLEGIOS
    if not ad.is_dir():
        return None
    exact = ad / "plantilla_informe_icco.docx"
    if exact.is_file():
        return exact.resolve()
    for pat in ("plantilla_informe_icco*.docx", "*plantilla*icco*.docx"):
        for p in sorted(ad.glob(pat)):
            if p.is_file() and not p.name.startswith("~$"):
                return p.resolve()
    return None


def _find_style_template_docx(root: Path, plantilla_explicita: Optional[Path]) -> Optional[Path]:
    """
    .docx del que copiar estilos (Normal, títulos, colores).
    Orden: ``--plantilla-estilos``, luego «auditoría para colegios» en ``reports/auditoria_cpa_icco``,
    luego Rebeca Matte, luego (opcional) ``auditoria_para_colegios/``.
    """
    for p in (
        plantilla_explicita,
        _find_plantilla_auditoria_colegios_docx(root),
        _find_rebeca_matte_docx(root),
        _find_assets_icco_plantilla(root),
    ):
        if p is not None and Path(p).is_file():
            return Path(p).resolve()
    return None


def _clear_docx_body_keep_sectpr(doc: Document) -> None:
    """Vacía el cuerpo conservando ``sectPr`` (márgenes, sección) para reutilizar estilos de la plantilla."""
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _open_icco_document(root: Path, plantilla_explicita: Optional[Path]) -> Document:
    """Documento en blanco con estilos de plantilla si existe; si no, documento nuevo."""
    tpl = _find_style_template_docx(root, plantilla_explicita)
    if tpl is None:
        doc = Document()
        setattr(doc, "_icco_uses_template", False)
        return doc
    try:
        doc = Document(str(tpl))
    except (OSError, ValueError, KeyError):
        doc = Document()
        setattr(doc, "_icco_uses_template", False)
        return doc
    _clear_docx_body_keep_sectpr(doc)
    setattr(doc, "_icco_uses_template", True)
    setattr(doc, "_icco_template_path", str(tpl))
    return doc


def _icco_from_template(doc: Document) -> bool:
    return bool(getattr(doc, "_icco_uses_template", False))


# Cuerpo / pie: primer nombre que exista en la plantilla (Word ES / EN).
_STYLES_BODY = (
    "Normal",
    "Body Text",
    "Texto sin sangría",
    "Texto independiente",
    "Plain Text",
)


def _paragraph_apply_style_from_template(paragraph, doc: Document, candidates: tuple[str, ...]) -> None:
    if not _icco_from_template(doc):
        return
    for name in candidates:
        try:
            paragraph.style = doc.styles[name]
            return
        except KeyError:
            continue


def _apply_icco_title_runs(
    para: Paragraph,
    *,
    tier: int,
) -> None:
    """Calibri, azul WES, negrita — corrige títulos en negro que vienen de estilos de plantilla."""
    t = 1 if int(tier) == 1 else 2
    pt = Pt(14) if t == 1 else Pt(13)
    for r in para.runs:
        try:
            r.font.name = _HEADING_FONT
            r.font.color.rgb = _HEADING_COLOR
            r.bold = True
            r.font.size = pt
        except Exception:
            pass


def _add_heading_like_template(
    doc: Document,
    text: str,
    *,
    prefer_level: int,
) -> Paragraph:
    """
    Con plantilla: usa Título 1/2/3 según existan en el .docx (más parecido al Word de muestra).
    Sin plantilla: Título 1 (principal) o Título 2 (secundario), Calibri azul WES — mismo criterio
    que «Perfil horario completo (24 h): Con WES vs línea base» (nivel 2).
    Tras crear el párrafo, se fuerza color/fuente WES (evita títulos en negro).
    """
    if not _icco_from_template(doc):
        lev = min(max(int(prefer_level), 1), 9)
        h = doc.add_heading(text, level=lev)
        _apply_icco_title_runs(h, tier=prefer_level)
        return h
    order = (prefer_level, 1, 2, 3, 0)
    seen: set[int] = set()
    for lev in order:
        if lev in seen:
            continue
        seen.add(lev)
        try:
            h = doc.add_heading(text, level=lev)
            _apply_icco_title_runs(h, tier=prefer_level)
            return h
        except KeyError:
            continue
    p = doc.add_paragraph(text)
    _paragraph_apply_style_from_template(p, doc, _STYLES_BODY)
    _apply_icco_title_runs(p, tier=prefer_level)
    return p


def _section_heading_icco(doc: Document, text: str, *, space_after_pt: int = 8) -> Paragraph:
    """Secciones principales: Índice, Metodología, Registros, Resultados."""
    h = _add_heading_like_template(doc, text, prefer_level=1)
    hf = h.paragraph_format
    hf.keep_with_next = True
    hf.widow_control = True
    hf.space_after = Pt(space_after_pt)
    return h


def _subsection_heading_icco(doc: Document, text: str) -> Paragraph:
    """Subtítulo de sección (mismo estilo que «Perfil horario completo (24 h): …»)."""
    h = _add_heading_like_template(doc, text, prefer_level=2)
    hf = h.paragraph_format
    hf.keep_with_next = True
    hf.widow_control = True
    hf.space_after = Pt(8)
    return h


# Etiqueta corta en el título del cuadro resumen (siempre mayúsculas «ICCO»).
_ETIQUETA_COLEGIO_CUADRO_RESUMEN = "ICCO"


def _apply_table_style_safe(tbl, doc: Document) -> None:
    """Estilo de tabla: nombres EN/ES/IT comunes en plantillas Word."""
    names = [
        "Light Grid Accent 1",
        "Light Grid",
        "Table Grid",
        "Cuadriculada",
        "Tabla con cuadrícula",
        "Tabella con griglia",
        "Normal Table",
        "Medium Shading 1 Accent 1",
    ]
    for name in names:
        try:
            tbl.style = name
            return
        except KeyError:
            continue


def _apply_heading_run_style(doc: Document, heading_para) -> None:
    """Colores/fuente de título: desde plantilla o Calibri + azul WES."""
    if _icco_from_template(doc):
        return
    for r in heading_para.runs:
        r.font.name = _HEADING_FONT
        r.font.color.rgb = _HEADING_COLOR


def _apply_index_line_runs(doc: Document, para) -> None:
    if _icco_from_template(doc):
        return
    for r in para.runs:
        r.font.name = _BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 0, 0)


def _sync_portada_media(reference_docx: Path, media_dir: Path) -> None:
    """Extrae solo las imágenes de portada desde la plantilla .docx."""
    media_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(reference_docx) as zf:
        for fname in PORTADA_IMAGENES:
            arc = f"word/media/{fname}"
            if arc in zf.namelist():
                (media_dir / fname).write_bytes(zf.read(arc))


def _add_centered_picture(
    doc: Document,
    image_path: Path,
    width_inches: float = 6.3,
    *,
    keep_with_next: bool = False,
) -> None:
    p = doc.add_paragraph()
    _paragraph_apply_style_from_template(p, doc, _STYLES_BODY)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(str(image_path.resolve()), width=Inches(width_inches))
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.widow_control = True
    pf.keep_with_next = keep_with_next


def _p_justify(doc: Document, text: str, *, keep_with_next: bool = False) -> None:
    para = doc.add_paragraph(text)
    _paragraph_apply_style_from_template(para, doc, _STYLES_BODY)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pf = para.paragraph_format
    pf.widow_control = True
    pf.keep_with_next = keep_with_next
    if not _icco_from_template(doc):
        for run in para.runs:
            run.font.name = _BODY_FONT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)


def _spacer_lines(doc: Document, n: int) -> None:
    for _ in range(n):
        doc.add_paragraph("")


def _portada_y_indice(
    doc: Document,
    root: Path,
    media_dir: Optional[Path],
    *,
    portada_completa: Optional[Path] = None,
) -> None:
    """
    Por defecto: portada generada (título + establecimiento + WES + pie; sin banner en tabla).
    Opcional: ``portada_completa`` (PNG) o ``media_dir`` con recortes Rebeca.
    Tras la portada (pág. 1), salto a pág. 2: Índice y luego Metodología van en la misma página
    (sin salto entre Índice y Metodología; el salto va antes de Resultados).
    """
    if portada_completa and portada_completa.is_file():
        _add_centered_picture(
            doc,
            portada_completa,
            width_inches=6.55,
            keep_with_next=True,
        )
        br = doc.add_paragraph()
        br.add_run().add_break(WD_BREAK.PAGE)
    elif media_dir and media_dir.is_dir():
        ok = all((media_dir / name).is_file() for name in PORTADA_IMAGENES)
        if ok:
            _spacer_lines(doc, 1)
            for i, name in enumerate(PORTADA_IMAGENES):
                last = i == len(PORTADA_IMAGENES) - 1
                _add_centered_picture(
                    doc,
                    media_dir / name,
                    width_inches=6.4,
                    keep_with_next=last,
                )
            _spacer_lines(doc, 1)
            br = doc.add_paragraph()
            br.add_run().add_break(WD_BREAK.PAGE)
        else:
            _portada_generada_wes(doc, root)
    else:
        _portada_generada_wes(doc, root)

    idx_h = _section_heading_icco(doc, "Índice", space_after_pt=10)
    # Números de página orientativos: p.1 portada; p.2 índice + metodología; luego registros; resultados al final.
    idx_lines = (
        "Metodología" + " " * 8 + "." * 42 + " 2",
        "Registros de consumos" + " " * 4 + "." * 34 + " 3",
        "Resultados y Conclusiones" + " " * 4 + "." * 28 + " 4",
    )
    for i, line in enumerate(idx_lines):
        ip = doc.add_paragraph(line)
        _paragraph_apply_style_from_template(ip, doc, _STYLES_BODY)
        _apply_index_line_runs(doc, ip)
        ip.paragraph_format.widow_control = True
        if i == len(idx_lines) - 1:
            ip.paragraph_format.keep_with_next = True


def _metodologia_icco(
    doc: Document, node_id: str, *, ref: Periodo, aud: Periodo
) -> None:
    _section_heading_icco(doc, "Metodología")
    nd = len(ref.dias)
    rc = _rango_fechas_prosa_cl(ref.dias[0], ref.dias[-1])
    rs = _rango_fechas_prosa_cl(aud.dias[0], aud.dias[-1])
    n_intervalos = len(ref.dias) * (HORA_FIN_EXCL - HORA_INICIO)
    _p_justify(
        doc,
        "Se generaron dos auditorías hídricas que permiten comprobar de forma empírica el efecto de "
        "ahorro asociado al servicio WES «Control Inteligente de Agua Potable» en las redes del "
        "establecimiento con tecnología de control. En cuanto a contexto y alcance, se "
        f"definieron dos periodos de {nd} días cada uno: con control hidráulico activo entre {rc} "
        f"y con control desactivado entre {rs}, de modo que ambas series son comparables y permiten "
        "dimensionar la diferencia de consumo atribuible al servicio WES.",
    )
    _p_justify(
        doc,
        "En la elaboración de los indicadores se consideran, para cada nodo, solo registros con fecha "
        "válida; el promedio diario de un periodo es la suma de los consumos diarios dividida por la "
        "cantidad de días del periodo. El ahorro estimado en volumen (m³/día) se expresa como "
        "max(0, promedio_desactivado − promedio_activo) y el ahorro porcentual como "
        "(ahorro_m3_dia / promedio_desactivado) × 100. La visualización gráfica —en especial la curva "
        "diaria— contrasta ambos periodos y facilita evidenciar el impacto del control WES.",
    )
    _p_justify(
        doc,
        f"Para {PORTADA_ESTABLECIMIENTO_LINEA1.strip() or 'el establecimiento'}, punto de medición WES "
        f"{node_id} ({NOMBRE_PUNTO}), las series se "
        "construyen además en rejilla horaria a partir de datos de la plataforma en zona horaria Chile, "
        f"para el periodo con control ({_formato_fechas_periodo_auditoria(ref)}) y el "
        f"periodo sin control ({_formato_fechas_periodo_auditoria(aud)}), en jornada "
        f"completa {HORA_INICIO:02d}:00 a {HORA_FIN_EXCL:02d}:00 (una lectura por hora; fin de intervalo "
        "exclusivo). Si una hora no tiene dato, se asume 0 m³/h para mantener una grilla comparable "
        f"({n_intervalos} intervalos: {len(ref.dias)} días × {HORA_FIN_EXCL - HORA_INICIO} horas). "
        "Como resultados esperados de este ejercicio se busca disponer de una línea de base actualizada "
        "para la estimación de eficiencia y proyectar el ahorro de agua (m³) y una estimación "
        "económica orientativa en dinero (CLP $).",
    )


def _tabla_resultados_cuantitativos(
    doc: Document,
    *,
    prom_con_m3_dia: float,
    prom_sin_m3_dia: float,
    texto_periodo_con: str,
    texto_periodo_sin: str,
) -> None:
    """
    Tabla de resultados cuantitativos bajo «Resultados y Conclusiones» (sin segundo título duplicado).
    Encabezado de tabla azul; Indicador en negrita; Valor en texto normal; bordes de cuadrícula.
    """
    # Indicadores coherentes con la metodología (sin valores negativos).
    dif_neta = max(0.0, prom_sin_m3_dia - prom_con_m3_dia)
    pct_ahorro = (100.0 * dif_neta / prom_sin_m3_dia) if prom_sin_m3_dia > 0 else 0.0
    filas: tuple[tuple[str, str], ...] = (
        ("Indicador", "Valor"),
        (
            f"Promedio diario con control activo ({texto_periodo_con})",
            f"{prom_con_m3_dia:.2f} m3/dia",
        ),
        (
            f"Promedio diario con control desactivado ({texto_periodo_sin})",
            f"{prom_sin_m3_dia:.2f} m3/dia",
        ),
        (
            "Ahorro promedio diario estimado por control WES",
            f"{dif_neta:.2f} m3/dia",
        ),
        ("Ahorro porcentual", f"{pct_ahorro:.2f}%"),
        (
            "Diferencia neta (desactivado - activo)",
            f"{dif_neta:.2f} m3/dia",
        ),
    )
    tbl = doc.add_table(rows=len(filas), cols=2)
    _apply_metricas_clave_table_style(tbl, doc)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        tbl.columns[0].width = Inches(5.2)
        tbl.columns[1].width = Inches(2.0)
    except Exception:
        pass

    for i, (col_a, col_b) in enumerate(filas):
        row = tbl.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = col_a
        c1.text = col_b
        if i == 0:
            for c in (c0, c1):
                _set_cell_shading_hex(c, _TABLA_CUANTITATIVOS_FILL_HEADER)
                for par in c.paragraphs:
                    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in par.runs:
                        run.bold = True
                        run.font.color.rgb = _HEADING_COLOR
                        if not _icco_from_template(doc):
                            run.font.name = _BODY_FONT
                            run.font.size = Pt(10)
        else:
            for par in c0.paragraphs:
                par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for par in c1.paragraphs:
                par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in c0.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                if not _icco_from_template(doc):
                    run.font.name = _BODY_FONT
                    run.font.size = Pt(10)
            for run in c1.paragraphs[0].runs:
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)
                if not _icco_from_template(doc):
                    run.font.name = _BODY_FONT
                    run.font.size = Pt(10)
    gap = doc.add_paragraph("")
    gap.paragraph_format.space_after = Pt(8)


def _resultados_y_conclusiones(
    doc: Document,
    total_ref: float,
    total_aud: float,
    ahorro_m3: float,
    rend_pct: float,
    dias: int,
    *,
    ref: Periodo,
    aud: Periodo,
) -> None:
    """total_ref = con control; total_aud = sin control (misma convención que auditoria_cpa_icco_renca_grafico)."""
    del rend_pct, ahorro_m3  # el texto usa solo totales; evita cifras negativas en el informe.
    total_con_m3 = total_ref
    total_sin_m3 = total_aud
    # Ahorro favorable: menor Σ Con que Sin (sin mostrar signos negativos en el informe).
    vol_ahorro_m3 = max(0.0, total_sin_m3 - total_con_m3)
    pct_ahorro_sobre_sin = (
        (100.0 * vol_ahorro_m3 / total_sin_m3) if total_sin_m3 > 0 else 0.0
    )
    ahorro_clp_aprox = vol_ahorro_m3 * CLP_POR_M3_REF
    exceso_con_sobre_sin = max(0.0, total_con_m3 - total_sin_m3)

    _section_heading_icco(doc, "Resultados y Conclusiones")
    dias_periodo = dias if dias else 1
    prom_con_m3_dia = total_con_m3 / dias_periodo
    prom_sin_m3_dia = total_sin_m3 / dias_periodo
    texto_periodo_con = _formato_fechas_periodo_compacto(ref)
    texto_periodo_sin = _formato_fechas_periodo_compacto(aud)
    _tabla_resultados_cuantitativos(
        doc,
        prom_con_m3_dia=prom_con_m3_dia,
        prom_sin_m3_dia=prom_sin_m3_dia,
        texto_periodo_con=texto_periodo_con,
        texto_periodo_sin=texto_periodo_sin,
    )
    if vol_ahorro_m3 > 0:
        _p_justify(
            doc,
            f"En la rejilla analizada ({dias} días, {HORA_INICIO:02d}:00–{HORA_FIN_EXCL:02d}:00), "
            f"Con WES {format_number_chilean(total_con_m3, 3)} m³ y Sin WES {format_number_chilean(total_sin_m3, 3)} m³; "
            f"volumen evitado respecto de Sin WES: {format_number_chilean(vol_ahorro_m3, 3)} m³ "
            f"(~{format_number_chilean(ahorro_clp_aprox, 0)} CLP a {format_number_chilean(CLP_POR_M3_REF, 0)} CLP/m³, orientativo). "
            "Cuadros y gráficos: Registros de consumos.",
        )
    elif exceso_con_sobre_sin > 0:
        _p_justify(
            doc,
            f"En la rejilla ({dias} días, {HORA_INICIO:02d}:00–{HORA_FIN_EXCL:02d}:00), Con WES "
            f"{format_number_chilean(total_con_m3, 3)} m³ y Sin WES {format_number_chilean(total_sin_m3, 3)} m³ "
            f"(mayor Σ en Con WES en {format_number_chilean(exceso_con_sobre_sin, 3)} m³). "
            "Detalle en Registros de consumos.",
        )
    else:
        _p_justify(
            doc,
            f"Con WES {format_number_chilean(total_con_m3, 3)} m³; Sin WES {format_number_chilean(total_sin_m3, 3)} m³ "
            f"({dias} días, {HORA_INICIO:02d}:00–{HORA_FIN_EXCL:02d}:00). Detalle en Registros de consumos.",
        )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    _paragraph_apply_style_from_template(p, doc, _STYLES_BODY)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if vol_ahorro_m3 > 0 and total_sin_m3 > 0:
        r0 = p.add_run("Ahorro sobre Sin WES: ")
        r1 = p.add_run(f"{format_number_chilean(pct_ahorro_sobre_sin, 2)} %")
        r1.bold = True
        for r in (r0, r1):
            if not _icco_from_template(doc):
                r.font.name = _BODY_FONT
        r0.font.size = Pt(11)
        r1.font.size = Pt(14)
    else:
        r0 = p.add_run("Ahorro sobre Sin WES: no aplica en este tramo.")
        if not _icco_from_template(doc):
            r0.font.name = _BODY_FONT
            r0.font.size = Pt(11)


def _formato_fechas_periodo_auditoria(periodo) -> str:
    """Rango calendario Chile de los días incluidos en el periodo (inicio–fin)."""
    d0 = periodo.dias[0]
    d1 = periodo.dias[-1]
    return f"{d0.strftime('%d-%m-%Y')} al {d1.strftime('%d-%m-%Y')}"


def _formato_rango_fechas_compacto(d0: date, d1: date) -> str:
    """
    Texto corto para celdas de tabla, p. ej. «23 al 26 marzo 2026» o «26 marzo al 6 abril 2026».
    """
    m = _MESES_ES
    if d0 == d1:
        return f"{d0.day:02d} de {m[d0.month - 1]} {d0.year}"
    if d0.year == d1.year and d0.month == d1.month:
        return f"{d0.day:02d} al {d1.day:02d} {m[d0.month - 1]} {d0.year}"
    if d0.year == d1.year:
        return f"{d0.day:02d} {m[d0.month - 1]} al {d1.day:02d} {m[d1.month - 1]} {d0.year}"
    return f"{d0.day:02d} {m[d0.month - 1]} {d0.year} al {d1.day:02d} {m[d1.month - 1]} {d1.year}"


def _formato_fechas_periodo_compacto(periodo) -> str:
    """Mismo periodo que ``_formato_fechas_periodo_auditoria``, formato compacto para tablas."""
    return _formato_rango_fechas_compacto(periodo.dias[0], periodo.dias[-1])


def _rango_desde_lista_fechas(fechas: List[date], i0: int, i1: int) -> str:
    """Rango inclusivo entre índices i0..i1 en ``fechas`` (mismo formato compacto que las tablas)."""
    return _formato_rango_fechas_compacto(fechas[i0], fechas[i1])


def _texto_comparacion_un_dia_24h(doc: Document, c: ComparacionDia24h) -> None:
    dif = c.total_con_m3 - c.total_sin_m3
    pct = (100.0 * dif / c.total_con_m3) if c.total_con_m3 > 0 else 0.0
    _p_justify(
        doc,
        f"{c.nombre_dia} — día homólogo con control ({c.fecha_con.strftime('%d-%m-%Y')}) "
        f"vs sin control ({c.fecha_sin.strftime('%d-%m-%Y')}): consumo diario aproximado "
        f"(suma de las 24 horas en m³/h) Con WES {format_number_chilean(c.total_con_m3, 3)} m³; "
        f"Sin WES (línea base) {format_number_chilean(c.total_sin_m3, 3)} m³; "
        f"diferencia (Con − Sin) {format_number_chilean(dif, 3)} m³ "
        f"({format_number_chilean(pct, 2)} % respecto del día con WES).",
        keep_with_next=True,
    )


def _tabla_consumo_sin_y_con_control(
    doc: Document,
    total_con_m3: float,
    total_sin_m3: float,
    dias_periodo: int,
    texto_periodo_con_wes: Optional[str] = None,
    texto_periodo_sin_wes: Optional[str] = None,
    *,
    periodo_con: Optional[Periodo] = None,
    periodo_sin: Optional[Periodo] = None,
) -> None:
    """
    Filas Con WES / Sin WES. Periodos por defecto API; si se pasan textos, deben coincidir con el consolidado.
    Promedio diario = Σ rejilla / dias_periodo.
    """
    _p_justify(
        doc,
        f"Jornada de medición (cada día): {HORA_INICIO:02d}:00 a {HORA_FIN_EXCL:02d}:00 (Chile). "
        "Σ rejilla = suma de todos los intervalos horarios (m³/h) del periodo; "
        "promedio diario = Σ rejilla ÷ número de días del periodo.",
        keep_with_next=True,
    )
    nd = dias_periodo if dias_periodo else 1
    prom_con = total_con_m3 / nd
    prom_sin = total_sin_m3 / nd
    pc = periodo_con or PERIODO_REFERENCIA
    ps = periodo_sin or PERIODO_AUDITORIA
    txt_con = texto_periodo_con_wes or _formato_fechas_periodo_compacto(pc)
    txt_sin = texto_periodo_sin_wes or _formato_fechas_periodo_compacto(ps)
    tbl = doc.add_table(rows=3, cols=4)
    _apply_metricas_clave_table_style(tbl, doc)
    filas = (
        (
            "Condición",
            "Periodo",
            "Σ rejilla (m³/h)",
            "Promedio diario (m³/h)",
        ),
        (
            "Con WES",
            txt_con,
            format_number_chilean(total_con_m3, 3),
            format_number_chilean(prom_con, 3),
        ),
        (
            "Sin WES",
            txt_sin,
            format_number_chilean(total_sin_m3, 3),
            format_number_chilean(prom_sin, 3),
        ),
    )
    for i, fila in enumerate(filas):
        for j, texto in enumerate(fila):
            celda = tbl.rows[i].cells[j]
            celda.text = texto
            if i == 0:
                _set_cell_shading_hex(celda, _TABLA_AUDITORIA_FILL_ENCABEZADO)
            else:
                _set_cell_shading_hex(celda, _TABLA_AUDITORIA_FILL_COL_CONDICION)
            for par in celda.paragraphs:
                par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i == 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in par.runs:
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = _HEADING_COLOR
                    elif j == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        run.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    if not _icco_from_template(doc):
                        run.font.name = _BODY_FONT
                        run.font.size = Pt(10)
    p_after = doc.add_paragraph("")
    p_after.paragraph_format.space_after = Pt(6)


def _registros_consumos(
    doc: Document,
    total_con_m3: float,
    total_sin_m3: float,
    dias_periodo: int,
    png_barras: Optional[Path],
    res_png: Optional[Path],
    csv_name: Optional[str],
    *,
    png_perfil_24h: Optional[list[Path]] = None,
    comparaciones_24h: Optional[list[ComparacionDia24h]] = None,
    texto_periodo_con_wes: Optional[str] = None,
    texto_periodo_sin_wes: Optional[str] = None,
    periodo_con: Optional[Periodo] = None,
    periodo_sin: Optional[Periodo] = None,
) -> None:
    _section_heading_icco(doc, "Registros de consumos", space_after_pt=10)
    # Un solo título de sección (evita duplicar con un segundo encabezado de cuadro).
    # Barras primero: así queda en la parte superior tras Metodología (p. 3 si índice+metodología caben en p. 2)
    if png_barras and Path(png_barras).is_file():
        _p_justify(
            doc,
            "Comparación del total acumulado en la rejilla horaria Σ (m³/h) entre Con WES y Sin WES "
            "(misma base que el Excel consolidado por CSV):",
            keep_with_next=True,
        )
        add_picture_with_pagination(
            doc,
            str(Path(png_barras).resolve()),
            width=Inches(5.2),
            keep_with_next=True,
        )
        gap = doc.add_paragraph("")
        gap.paragraph_format.space_before = Pt(14)
        gap.paragraph_format.space_after = Pt(6)

    _tabla_consumo_sin_y_con_control(
        doc,
        total_con_m3,
        total_sin_m3,
        dias_periodo,
        texto_periodo_con_wes=texto_periodo_con_wes,
        texto_periodo_sin_wes=texto_periodo_sin_wes,
        periodo_con=periodo_con,
        periodo_sin=periodo_sin,
    )

    if csv_name:
        _p_justify(
            doc,
            f"Métricas y serie horaria (jornada completa) exportadas: {csv_name}.",
        )

    paths_24 = png_perfil_24h or []
    comps = comparaciones_24h or []
    n_c = len(comps)
    if paths_24:
        doc.add_paragraph("")
        _subsection_heading_icco(doc, "Perfil horario completo (24 h): Con WES vs línea base")
        _p_justify(
            doc,
            "Los gráficos siguientes muestran el consumo horario (m³/h) en zona horaria Chile, "
            "para cada par de días homólogos: periodo con control (Con WES) frente al periodo sin control "
            "(línea base Sin WES). Las áreas semitransparentes permiten comparar visualmente ambas series.",
            keep_with_next=True,
        )
        if n_c >= 1 and len(paths_24) >= n_c + 1 and comps:
            for i in range(n_c):
                _texto_comparacion_un_dia_24h(doc, comps[i])
                add_picture_with_pagination(
                    doc,
                    str(paths_24[i].resolve()),
                    width=Inches(6.0),
                    keep_with_next=True,
                )
                gap = doc.add_paragraph("")
                gap.paragraph_format.space_after = Pt(8)
            prom_con = sum(c.total_con_m3 for c in comps) / float(n_c)
            prom_sin = sum(c.total_sin_m3 for c in comps) / float(n_c)
            difp = prom_con - prom_sin
            pctp = (100.0 * difp / prom_con) if prom_con > 0 else 0.0
            _p_justify(
                doc,
                f"Promedio horario de los {n_c} días homólogos (gráfico siguiente): "
                f"Con WES {format_number_chilean(prom_con, 3)} m³/día; "
                f"Sin WES {format_number_chilean(prom_sin, 3)} m³/día; diferencia media "
                f"{format_number_chilean(difp, 3)} m³/día ({format_number_chilean(pctp, 2)} % sobre el promedio con WES).",
                keep_with_next=True,
            )
            add_picture_with_pagination(
                doc,
                str(paths_24[n_c].resolve()),
                width=Inches(6.0),
                keep_with_next=False,
            )
        else:
            for i, pth in enumerate(paths_24, start=1):
                _subsection_heading_icco(doc, f"Gráfico comparativo diario #{i}.")
                add_picture_with_pagination(
                    doc,
                    str(Path(pth).resolve()),
                    width=Inches(6.0),
                    keep_with_next=True,
                )
                gap = doc.add_paragraph("")
                gap.paragraph_format.space_after = Pt(8)


def _datos_desde_excel_consolidado(
    xlsx_c: Path,
) -> tuple[
    float,
    float,
    int,
    Path,
    Optional[list[Path]],
    Optional[list[ComparacionDia24h]],
    Optional[str],
    Optional[str],
]:
    """
    Totales, barras y perfiles desde el Excel consolidado y ``graficos_comparativos/`` (sin API).
    """
    tr, ta, nper = totales_rejilla_desde_excel_consolidado(xlsx_c)
    fechas_x, _ = leer_matriz_consolidado(xlsx_c)
    periodo_con_txt: Optional[str] = None
    periodo_sin_txt: Optional[str] = None
    if len(fechas_x) >= 14:
        periodo_con_txt = _rango_desde_lista_fechas(fechas_x, 0, 6)
        periodo_sin_txt = _rango_desde_lista_fechas(fechas_x, 7, 13)
    elif len(fechas_x) == 8:
        # Misma lógica que auditoría CPA ICCO: 4 días Con + 4 días Sin.
        periodo_con_txt = _rango_desde_lista_fechas(fechas_x, 0, 3)
        periodo_sin_txt = _rango_desde_lista_fechas(fechas_x, 4, 7)
    gdir = xlsx_c.parent / "graficos_comparativos"
    png_barras_inf = generar_png_barras_rejilla_totales(gdir, tr, ta)
    if len(fechas_x) >= 14:
        diarios = [gdir / f"{n:02d}_area_{nm}.png" for n, nm in [
            (4, "Lunes"),
            (5, "Martes"),
            (6, "Miercoles"),
            (7, "Jueves"),
            (8, "Viernes"),
            (9, "Sabado"),
            (10, "Domingo"),
        ]]
    elif len(fechas_x) == 8:
        diarios = sorted(
            (p for p in gdir.glob("[0-9][0-9]_area_*.png") if p.name[:2].isdigit()),
            key=lambda p: int(p.name.split("_", 1)[0]),
        )
        diarios = [p for p in diarios if 4 <= int(p.name.split("_", 1)[0]) <= 7]
    else:
        diarios = []
    prom = gdir / "03_area_promedio_24h.png"
    diarios_ok = [p for p in diarios if p.is_file()]
    png_perfiles_inf: Optional[list[Path]] = None
    comps_inf: Optional[list[ComparacionDia24h]] = None
    if diarios_ok:
        png_perfiles_inf = diarios_ok + ([prom] if prom.is_file() else [])
        comps_inf = []
    return (
        tr,
        ta,
        nper,
        png_barras_inf,
        png_perfiles_inf,
        comps_inf,
        periodo_con_txt,
        periodo_sin_txt,
    )


def generar_informe_word(
    node_id: str = NODE_DEFAULT,
    out_dir: Path | None = None,
    output_docx: Path | None = None,
    portada_png: Optional[Path] = None,
    portada_rebeca_media: bool = False,
    plantilla_estilos: Optional[Path] = None,
    *,
    mantener_borrador_manual: bool = False,
    solo_consolidado: bool = False,
    periodo_ref: Optional[Periodo] = None,
    periodo_aud: Optional[Periodo] = None,
    figuras_desde_xlsx: Optional[Path] = None,
) -> Path:
    """
    Si ``output_docx`` es una ruta, el .docx se guarda ahí (sobrescribe) y los
    gráficos/CSV usan timestamp fijo ``borrador_icco`` para poder regenerar sin
    acumular archivos.

    Con ``mantener_borrador_manual=True`` y el .docx ya existente, no se reescribe el Word:
    solo se regeneran PNG y CSV (para no perder ediciones manuales en el borrador).

    Con ``solo_consolidado=True`` no se llama a la API WES: el informe usa solo el Excel
    ``consolidado_revision_todos_los_csv_descarga_api.xlsx`` y la carpeta ``graficos_comparativos/``
    (regeneración rápida del Word).

    Con ``figuras_desde_xlsx`` (ruta a un .xlsx con hoja ``Consolidado`` junto a ``graficos_comparativos/``):
    tras la auditoría por API, el Word usa esos PNG y totales de la rejilla (misma lógica que
    ``--solo-consolidado`` pero conservando CSV/PNG de auditoría en ``out_dir``).
    """
    root = Path(__file__).resolve().parent
    out_dir = out_dir or (root / "reports" / "auditoria_cpa_icco")
    out_dir.mkdir(parents=True, exist_ok=True)
    # Cuando se usa en batch (muchos nodos con -o), un timestamp fijo haría que CSV/PNG se pisen.
    # Mantener nombre estable por nodo permite regenerar sin acumular y sin colisiones.
    if output_docx is not None:
        safe_node = "".join(ch for ch in str(node_id) if ch.isalnum() or ch in ("-", "_"))
        ts = f"borrador_{safe_node}"
    else:
        ts = datetime.now().strftime("%Y%m%d_%H%M")

    word_path_resolved: Optional[Path] = None
    if output_docx is not None:
        word_path_resolved = output_docx.expanduser().resolve()
        word_path_resolved.parent.mkdir(parents=True, exist_ok=True)

    skip_regenerar_word = bool(
        mantener_borrador_manual and word_path_resolved is not None and word_path_resolved.is_file()
    )

    assets_dir = root / DIR_AUDITORIA_COLEGIOS
    portada_completa: Optional[Path] = None
    if not skip_regenerar_word:
        if portada_png and portada_png.is_file():
            portada_completa = portada_png.resolve()

        media_dir: Optional[Path] = None
        if portada_completa is None and portada_rebeca_media:
            ref_docx = _find_reference_docx(root)
            if ref_docx and ref_docx.is_file():
                media_dir = assets_dir / "plantilla_auditoria_wes_rebeca_media"
                _sync_portada_media(ref_docx, media_dir)
            else:
                print(
                    "[AVISO] --portada-rebeca-media: no se encontró Informe Rebeca Matte en "
                    "reports/auditoria_cpa_icco; se usará portada generada."
                )
    else:
        media_dir = None
        portada_completa = None

    xlsx_c = _XLSX_CONSOLIDADO_CSV_API.resolve()

    pref = periodo_ref or PERIODO_REFERENCIA
    paud = periodo_aud or PERIODO_AUDITORIA

    if solo_consolidado:
        if not xlsx_c.is_file():
            raise FileNotFoundError(
                f"--solo-consolidado requiere el Excel consolidado: {xlsx_c}"
            )
        print(
            "[INFO] Modo --solo-consolidado: sin API WES; totales y gráficos desde Excel local."
        )
        (
            tr,
            ta,
            nper,
            png_barras_inf,
            png_perfiles_inf,
            comps_inf,
            periodo_con_txt,
            periodo_sin_txt,
        ) = _datos_desde_excel_consolidado(xlsx_c)
        total_ref_inf = tr
        total_aud_inf = ta
        ahorro_inf = tr - ta
        rend_inf = (100.0 * ahorro_inf / tr) if tr > 0 else 0.0
        dias = nper
        res = ResultadoAuditoriaCpa(
            node_id=node_id,
            etiquetas=[],
            y_referencia=[],
            y_auditoria=[],
            total_ref_m3=tr,
            total_aud_m3=ta,
            ahorro_m3=ahorro_inf,
            rendimiento_pct=rend_inf,
            png_path=None,
            csv_path=xlsx_c,
            png_barras_path=png_barras_inf,
            png_paths_24h=png_perfiles_inf,
            comparaciones_diarias_24h=comps_inf,
        )
        print(f"  Informe: barras y cuadro desde consolidado ({xlsx_c.name}).")
    else:
        res = ejecutar_auditoria_cpa_icco(
            node_id=node_id,
            out_dir=out_dir,
            timestamp=ts,
            ref=pref,
            aud=paud,
            graficos_24h=figuras_desde_xlsx is None,
        )
        if figuras_desde_xlsx is not None:
            _safe_node = "".join(
                ch for ch in str(node_id) if ch.isalnum() or ch in ("-", "_")
            )
            for _old24 in out_dir.glob(
                f"cpa_icco_renca_wes_24h*_borrador_{_safe_node}.png"
            ):
                try:
                    _old24.unlink()
                except OSError:
                    pass
        if not res.csv_path or not res.csv_path.exists():
            raise FileNotFoundError("No se generó el CSV de auditoría.")
        if res.png_barras_path is None or not res.png_barras_path.is_file():
            res.png_barras_path = generar_png_barras_con_sin(
                res.total_ref_m3,
                res.total_aud_m3,
                out_dir,
                ts,
                ref=pref,
                aud=paud,
            )

        dias = len(pref.dias)
        total_ref_inf = res.total_ref_m3
        total_aud_inf = res.total_aud_m3
        ahorro_inf = res.ahorro_m3
        rend_inf = res.rendimiento_pct
        png_barras_inf: Optional[Path] = res.png_barras_path
        png_perfiles_inf: Optional[list[Path]] = res.png_paths_24h
        comps_inf: Optional[list[ComparacionDia24h]] = res.comparaciones_diarias_24h
        periodo_con_txt: Optional[str] = None
        periodo_sin_txt: Optional[str] = None

        if figuras_desde_xlsx is None:
            print(
                "  Informe: totales y gráficos desde la auditoría API (periodos 7+7 del módulo)."
            )

    if (
        not solo_consolidado
        and figuras_desde_xlsx is not None
        and Path(figuras_desde_xlsx).expanduser().resolve().is_file()
    ):
        xlsx_fig = Path(figuras_desde_xlsx).expanduser().resolve()
        try:
            (
                tr_fig,
                ta_fig,
                nper_fig,
                png_barras_inf,
                png_perfiles_inf,
                comps_inf,
                pc_fig,
                ps_fig,
            ) = _datos_desde_excel_consolidado(xlsx_fig)
            total_ref_inf = tr_fig
            total_aud_inf = ta_fig
            dias = nper_fig
            ahorro_inf = tr_fig - ta_fig
            rend_inf = (100.0 * ahorro_inf / tr_fig) if tr_fig > 0 else 0.0
            if periodo_con_txt is None and pc_fig:
                periodo_con_txt = pc_fig
            if periodo_sin_txt is None and ps_fig:
                periodo_sin_txt = ps_fig
            print(
                f"  Word: barras y perfiles desde {xlsx_fig.name} "
                f"(graficos_comparativos/ junto a ese Excel; totales alineados a la rejilla)."
            )
        except Exception as exc:
            print(f"  [AVISO] figuras_desde_xlsx no aplicado: {exc}")

    if skip_regenerar_word and word_path_resolved is not None:
        print(
            "[INFO] Mantener borrador manual: no se sobrescribe el .docx (ediciones conservadas)."
        )
        if res.png_path:
            print(f"  PNG lineas:  {res.png_path.resolve()}")
        if res.png_barras_path:
            print(f"  PNG barras:  {res.png_barras_path.resolve()}")
        if res.csv_path:
            print(f"  CSV:         {res.csv_path.resolve()}")
        print(
            "  Actualiza las imagenes en Word si quieres los datos nuevos (Insertar > Imagenes o reemplazar)."
        )
        return word_path_resolved

    doc = _open_icco_document(root, plantilla_estilos)
    _portada_y_indice(doc, root, media_dir, portada_completa=portada_completa)
    doc.add_paragraph("")
    _metodologia_icco(doc, res.node_id, ref=pref, aud=paud)

    p_br = doc.add_paragraph()
    run = p_br.add_run()
    run.add_break(WD_BREAK.PAGE)

    _registros_consumos(
        doc,
        total_ref_inf,
        total_aud_inf,
        dias,
        png_barras_inf,
        res.png_path,
        res.csv_path.name if res.csv_path else None,
        png_perfil_24h=png_perfiles_inf,
        comparaciones_24h=comps_inf,
        texto_periodo_con_wes=periodo_con_txt,
        texto_periodo_sin_wes=periodo_sin_txt,
        periodo_con=pref,
        periodo_sin=paud,
    )

    p_br2 = doc.add_paragraph()
    run2 = p_br2.add_run()
    run2.add_break(WD_BREAK.PAGE)

    _resultados_y_conclusiones(
        doc,
        total_ref_inf,
        total_aud_inf,
        ahorro_inf,
        rend_inf,
        dias,
        ref=pref,
        aud=paud,
    )

    doc.add_paragraph("")
    pie = doc.add_paragraph()
    _paragraph_apply_style_from_template(pie, doc, _STYLES_BODY)
    pie.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pr = pie.add_run(
        f"Documento generado automáticamente — {datetime.now(timezone.utc).strftime('%d-%m-%Y %H:%M')} UTC"
    )
    if not _icco_from_template(doc):
        pr.font.name = _BODY_FONT
        pr.font.size = Pt(9)
        pr.font.color.rgb = RGBColor(100, 100, 100)
    else:
        pr.font.size = Pt(9)
        pr.font.color.rgb = RGBColor(100, 100, 100)

    if output_docx is not None:
        word_path = output_docx.resolve()
        word_path.parent.mkdir(parents=True, exist_ok=True)
    else:
        word_path = out_dir / f"informe_auditoria_wes_icco_renca_{ts}.docx"
    doc.save(word_path)
    return word_path.resolve()


def main() -> int:
    ap = argparse.ArgumentParser(description="Informe Word auditoría CPA ICCO Renca")
    ap.add_argument("--node-id", default=NODE_DEFAULT, help="ID nodo WES")
    ap.add_argument(
        "-o",
        "--salida",
        type=Path,
        default=_DEFAULT_SALIDA_INFORME_DOCX,
        help="Ruta del .docx (default: Borrador_auditoria_ICCO_Renca_abril_2026.docx en reporte de auditoría). "
        "Sobrescribe salvo --mantener-borrador-manual. PNG/CSV usan timestamp borrador_icco.",
    )
    ap.add_argument(
        "--portada-png",
        type=Path,
        default=None,
        help="Usar este PNG como portada en lugar de la portada generada (texto editable).",
    )
    ap.add_argument(
        "--portada-rebeca-media",
        action="store_true",
        help="Usar recortes del Word Rebeca Matte (requiere ese .docx en reports/auditoria_cpa_icco).",
    )
    ap.add_argument(
        "--plantilla-estilos",
        type=Path,
        default=None,
        help="Plantilla .docx. Si no: auditoría para colegios en reports/, luego Rebeca Matte, luego auditoria_para_colegios/.",
    )
    ap.add_argument(
        "--mantener-borrador-manual",
        action="store_true",
        help="Con -o: si el .docx ya existe, no lo sobrescribe; solo regenera PNG y CSV (conserva tu Word editado).",
    )
    ap.add_argument(
        "--solo-consolidado",
        action="store_true",
        help="Sin llamadas API WES: totales y figuras solo desde consolidado_revision_todos_los_csv_descarga_api.xlsx "
        "y graficos_comparativos/ (regenerar el .docx en segundos). Requiere ese Excel.",
    )
    ap.add_argument(
        "--figuras-desde-xlsx",
        type=Path,
        default=None,
        help="Tras auditoría API: usar totales y PNG de este .xlsx y su carpeta graficos_comparativos/ en el Word.",
    )
    args = ap.parse_args()
    borrador_existia = bool(
        args.salida and Path(args.salida).expanduser().resolve().is_file()
    )
    out_dir_borrador: Path | None = None
    if args.salida is not None:
        out_dir_borrador = Path(args.salida).expanduser().resolve().parent
    p = generar_informe_word(
        node_id=args.node_id,
        out_dir=out_dir_borrador,
        output_docx=args.salida,
        portada_png=args.portada_png,
        portada_rebeca_media=args.portada_rebeca_media,
        plantilla_estilos=args.plantilla_estilos,
        mantener_borrador_manual=args.mantener_borrador_manual,
        solo_consolidado=args.solo_consolidado,
        figuras_desde_xlsx=args.figuras_desde_xlsx,
    )
    if args.mantener_borrador_manual and borrador_existia:
        print(f"PNG/CSV actualizados; borrador .docx sin cambios: {p.resolve()}")
    else:
        print(f"Informe generado: {p.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
