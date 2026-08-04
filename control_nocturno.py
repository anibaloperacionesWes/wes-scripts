"""
Genera reporte de control nocturno usando nodos WES y horarios desde Excel.

Incluye un gráfico de anillo (mismo estilo que el reporte de puntos en cero):
fuera de control vs cumpliendo corte, con total en el centro.

Objetivo de estandarizacion (segun requerimientos del usuario):
- La columna "Horas detectadas" se muestra en rangos (ej: "00:00-06:00") en vez
  de listar cada hora.
- No se muestra "total de consumo" en tablas (Word/PDF/CSV), solo:
  Fecha, Cliente, Establecimiento, Horas detectadas, Max por hora, ID WES.
"""

from __future__ import annotations

import argparse
import csv
import os
import re
import shutil
import tempfile
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import requests
import matplotlib

matplotlib.use("Agg")

# Serie horaria en hora Chile (misma base que auditoría / app); ver generar_reporte_word.
from generar_reporte_word import (
    _dt_to_chile,
    _utc_calendar_dates_for_chile_day,
    get_hourly_measures_for_day,
)
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_URL = "http://104.248.53.141:7003/wes/api/acl-node/v1"

# Excel de referencia en reports/ (nombre actualizado por el usuario).
EXCEL_CONTROL_NOCTURNO = "HORARIOS CONTROL NOCTURNO.xlsx"
EXCEL_CONTROL_NOCTURNO_LEGACY = "HORARIOS COLEGIOS.xlsx"


def default_excel_path() -> Path:
    """
    Ruta oficial del Excel de horarios para control nocturno:
    wes-scripts/reports/HORARIOS CONTROL NOCTURNO.xlsx
    Si aún existe el archivo antiguo HORARIOS COLEGIOS.xlsx, se usa como respaldo.
    """
    base = Path(__file__).resolve().parent
    reports_dir = base / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    primary = reports_dir / EXCEL_CONTROL_NOCTURNO
    legacy = reports_dir / EXCEL_CONTROL_NOCTURNO_LEGACY
    if primary.exists():
        return primary
    if legacy.exists():
        return legacy
    return primary


# Ajustes puntuales por nodo (desactivados por defecto).
# Para respetar la tabla Excel como fuente única, no se aplican salvo que se active:
#   WES_APLICAR_AJUSTES_HORARIO=1
AJUSTES_HORARIO_POR_NODO: Dict[str, str] = {}


def parse_horario_a_horas(horario: object) -> List[int]:
    """
    Convierte texto tipo '22:00 a 06:00' o '00:01 a 04:59' en horas enteras (0-23),
    respetando minutos con criterio de bloques horarios [h:00, h+1:00):
    - Se evalúan solo horas cuyo inicio cae dentro de la ventana [inicio, fin).
    - Si cruza medianoche, aplica partición en dos tramos.
    """
    if horario is None:
        return list(range(0, 7))
    text = str(horario).strip()
    if not text:
        return list(range(0, 7))

    m = re.search(
        r"(\d{1,2})\s*:\s*(\d{2})\s+a\s+(\d{1,2})\s*:\s*(\d{2})",
        text,
        flags=re.IGNORECASE,
    )
    if not m:
        return list(range(0, 7))

    start_h = int(m.group(1))
    start_m = int(m.group(2))
    end_h = int(m.group(3))
    end_m = int(m.group(4))

    start_min = start_h * 60 + start_m
    end_min = end_h * 60 + end_m

    horas: List[int] = []
    for h in range(24):
        h_start = h * 60
        if start_min < end_min:
            if start_min <= h_start < end_min:
                horas.append(h)
        elif start_min > end_min:  # cruza medianoche
            if h_start >= start_min or h_start < end_min:
                horas.append(h)
        else:
            # misma marca inicio/fin: sin ventana operativa
            continue
    return horas


def _fila_tiene_encabezado_cliente(row: Tuple[Any, ...]) -> bool:
    for cell in row:
        if cell is None:
            continue
        if str(cell).strip().upper() == "CLIENTE":
            return True
    return False


def _indice_columna_encabezado(hdr: Tuple[Any, ...], *keywords: str) -> Optional[int]:
    """
    Primera columna cuyo texto contiene alguna combinacion de keywords.
    """
    for idx, cell in enumerate(hdr):
        if cell is None:
            continue
        s = str(cell).strip().upper()
        if all(k in s for k in keywords):
            return idx
    # fallback: cualquiera de las keywords
    for idx, cell in enumerate(hdr):
        if cell is None:
            continue
        s = str(cell).strip().upper()
        for k in keywords:
            if k in s:
                return idx
    return None


def _parse_umbral_alerta(raw: object) -> float:
    """
    Acepta umbrales en formatos:
    - 3.6
    - 3,6
    - >3.6 / >= 3,6
    - texto con numero (toma el primer numero encontrado)
    """
    if raw is None:
        return 0.0
    s = str(raw).strip()
    if not s:
        return 0.0
    s_norm = s.replace(",", ".")
    try:
        return float(s_norm)
    except Exception:
        pass
    m = re.search(r"-?\d+(?:\.\d+)?", s_norm)
    if not m:
        return 0.0
    try:
        return float(m.group(0))
    except Exception:
        return 0.0


def cargar_targets_desde_excel(path: Path) -> Dict[str, Dict[str, Any]]:
    """
    Lee el Excel de control nocturno (reports/HORARIOS CONTROL NOCTURNO.xlsx) con columnas:
    CLIENTE, NOMBRE DEL COLEGIO O LICEO, ID, HORARIO DE CORTE, UMBRAL DE ALERTA
    """
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise RuntimeError("Instale openpyxl: pip install openpyxl") from e

    if not path.is_file():
        raise FileNotFoundError(f"No se encontro el Excel: {path}")

    # Copia temporal: openpyxl no abre el .xlsx original → evita bloqueo / "infracción de uso compartido"
    # al guardar desde Excel u OneDrive mientras corre un reporte.
    tmp_path: Optional[str] = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".xlsx", prefix="ctrl_nocturno_")
        os.close(fd)
        shutil.copy2(path, tmp_path)

        wb = load_workbook(Path(tmp_path), read_only=True, data_only=True)
        try:
            ws = wb[wb.sheetnames[0]]
            rows = [tuple(r) for r in ws.iter_rows(values_only=True)]

            header_idx: Optional[int] = None
            for i, row in enumerate(rows):
                if _fila_tiene_encabezado_cliente(row):
                    header_idx = i
                    break
            if header_idx is None:
                raise ValueError(f"No se encontro la fila de encabezados en {path}")

            hdr = rows[header_idx]
            i_cli = _indice_columna_encabezado(hdr, "CLIENTE")
            i_nom = (
                _indice_columna_encabezado(hdr, "NOMBRE")
                or _indice_columna_encabezado(hdr, "COLEGIO", "LICEO")
                or _indice_columna_encabezado(hdr, "NOMBRE", "COLEGIO")
            )
            i_id = None
            for j, cell in enumerate(hdr):
                if cell is not None and str(cell).strip().upper() == "ID":
                    i_id = j
                    break
            if i_id is None:
                i_id = _indice_columna_encabezado(hdr, "ID")

            i_hor = _indice_columna_encabezado(hdr, "HORARIO") or _indice_columna_encabezado(
                hdr, "CORTE"
            )
            i_umbral = _indice_columna_encabezado(hdr, "UMBRAL", "ALERTA") or _indice_columna_encabezado(
                hdr, "UMBRAL"
            )
            if i_cli is None or i_nom is None or i_id is None or i_hor is None:
                raise ValueError(
                    f"Encabezados incompletos en {path}: "
                    f"cliente={i_cli}, nombre={i_nom}, id={i_id}, horario={i_hor}"
                )

            targets: Dict[str, Dict[str, Any]] = {}
            for row in rows[header_idx + 1 :]:
                if not row or len(row) <= max(i_cli, i_nom, i_id, i_hor):
                    continue
                raw_id = row[i_id]
                if raw_id is None or str(raw_id).strip() == "":
                    continue

                node_id = str(raw_id).strip()
                cliente = str(row[i_cli]).strip() if row[i_cli] is not None else "—"
                nombre = str(row[i_nom]).strip() if row[i_nom] is not None else node_id
                horario_txt = str(row[i_hor]).strip() if row[i_hor] is not None else ""
                horas = parse_horario_a_horas(horario_txt if horario_txt else None)
                umbral_alerta = 0.0
                if i_umbral is not None and len(row) > i_umbral:
                    umbral_alerta = _parse_umbral_alerta(row[i_umbral])

                # Ajuste por nodo (opcional; desactivado por defecto para respetar Excel).
                aplicar_ajustes = os.environ.get("WES_APLICAR_AJUSTES_HORARIO", "").strip() == "1"
                if aplicar_ajustes and node_id in AJUSTES_HORARIO_POR_NODO:
                    horario_txt = AJUSTES_HORARIO_POR_NODO[node_id]
                    horas = parse_horario_a_horas(horario_txt)

                targets[node_id] = {
                    "nodeId": node_id,
                    "nodeName": nombre,
                    "cliente": cliente,
                    "establecimiento": nombre,
                    "horas_corte": horas,
                    "horario_texto": horario_txt or "—",
                    "umbral_alerta": umbral_alerta,
                }

            if not targets:
                raise ValueError(f"No se leyeron filas con ID valido en {path}")
            return targets
        finally:
            wb.close()
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _compactar_horas_a_rangos(horas: List[int]) -> str:
    """
    Convierte una lista de horas (0–23) a rangos compactos de bloques horarios.

    Cada hora representa el bloque [HH:00, HH+1:00). Por ejemplo:
    - [4] -> "04:00-05:00"
    - [4,5] -> "04:00-06:00"
    """
    if not horas:
        return "—"
    hs = sorted(set(int(h) for h in horas))
    rangos: List[Tuple[int, int]] = []
    start = prev = hs[0]
    for h in hs[1:]:
        if h == prev + 1:
            prev = h
            continue
        rangos.append((start, prev))
        start = prev = h
    rangos.append((start, prev))

    parts: List[str] = []
    for s, e in rangos:
        # Mostrar el rango como [inicio, fin) en horas.
        end_exclusive = e + 1
        # Permitir 24:00 como cierre de día (más claro que 00:00 en este contexto).
        end_label = f"{end_exclusive:02d}:00" if end_exclusive <= 24 else "24:00"
        parts.append(f"{s:02d}:00-{end_label}")
    return "; ".join(parts)


def _date_to_ddmmyyyy(d: datetime) -> str:
    return d.strftime("%d%m%Y")


def _value_by_time_last_row(csv_content: str) -> Dict[str, float]:
    """
  Último ``VALUE`` por marca ``TIME`` (no sumar duplicados).

  El CSV a veces repite la misma fila TIME dos veces; sumar inflaba el m³/h (p. ej. 1,07+1,07=2,14).
  Para control nocturno se compara caudal horario vs umbral, no el total diario.
    """
    out: Dict[str, float] = {}
    for line in csv_content.strip().split("\n")[1:]:
        if not line.strip():
            continue
        parts = line.split(",", 1)
        if len(parts) < 2:
            continue
        try:
            time_str = parts[0].strip()
            value_str = parts[1].strip().replace(" ", "").replace(",", ".")
            out[time_str] = float(value_str)
        except (ValueError, TypeError, IndexError):
            continue
    return out


def _dt_local_from_csv_time(time_str: str) -> Optional[datetime]:
    """
    Fecha/hora del CSV como la muestra la app WES (visualización nueva).

    Las marcas vienen como ``...T08:00:00.000Z`` pero la app las grafica en la
    hora **08:00 del día** sin aplicar offset UTC→Chile. Si se convierte a Chile,
    el mismo valor cae ~4 h antes y el control nocturno alerta en falso.
    """
    try:
        raw = time_str.strip()
        if "T" not in raw:
            return None
        ts = raw.replace("Z", "")
        if "+" in ts[10:]:
            ts = ts.split("+", 1)[0]
        else:
            tail = ts[10:]
            if "-" in tail:
                ts = ts[: ts.rfind("-")]
        return datetime.fromisoformat(ts)
    except (ValueError, TypeError):
        return None


def _chile_hours_from_csv_text(csv_content: str, dia_chile: date) -> Dict[int, float]:
    """``hora local (0-23) -> m³/h`` desde ``dates.measures.csv``, alineado con la app."""
    by_time = _value_by_time_last_row(csv_content)
    acc: Dict[int, float] = {}
    for time_str in sorted(by_time.keys()):
        try:
            dt = _dt_local_from_csv_time(time_str)
            if dt is None or dt.date() != dia_chile:
                continue
            hi = int(dt.hour)
            if 0 <= hi < 24:
                acc[hi] = float(by_time[time_str])
        except (ValueError, TypeError):
            continue
    return acc


def obtener_datos_horarios_dia(node_id: str, fecha: datetime) -> Dict[int, float]:
    """
    ``hora -> m³/h`` en hora local (0–23), alineado con la app WES.

    - Pide el CSV del día civil solicitado (``ddMMyyyy``).
    - Interpreta ``TIME`` como hora local del gráfico (sin UTC→Chile).
    - Usa el último valor por marca TIME (evita duplicados del CSV).
    - No escala al totalM3 diario (eso inflaba alertas vs umbral en algunos nodos).
    """
    dia = fecha.date()
    acc: Dict[int, float] = {}
    url = f"{BASE_URL}/nodes/{node_id}/dates.measures.csv"
    try:
        for ud in _utc_calendar_dates_for_chile_day(dia):
            date_str = ud.strftime("%d%m%Y")
            r = requests.get(
                url,
                params=[("start", date_str), ("end", date_str)],
                timeout=60,
            )
            r.raise_for_status()
            acc.update(_chile_hours_from_csv_text(r.text, dia))
        if acc:
            return {h: acc.get(h, 0.0) for h in range(24)}
    except Exception:
        pass
    # Fallback: serie reconciliada con total diario (menos fiel al umbral horario).
    os.environ["WES_HOURLY_SIN_MEASURES"] = "csv"
    hourly_list = get_hourly_measures_for_day(node_id, fecha) or []
    fallback: Dict[int, float] = {h: 0.0 for h in range(24)}
    for h, v in hourly_list:
        hi = int(h)
        if 0 <= hi < 24:
            fallback[hi] = float(v)
    return fallback


@dataclass
class AlertRow:
    fecha: str
    colegio: str  # nodeId (clave del Excel)
    node_id: str
    node_name: str
    horas_con_consumo: str
    horas_con_consumo_detalle: str
    consumo_max_hora_00_06: float


@dataclass
class AlertEvent:
    fecha: str
    colegio: str
    node_id: str
    node_name: str
    hour: int
    value: float


def analizar_control_nocturno(
    targets: Dict[str, Dict[str, Any]],
    start_date: datetime,
    end_date: datetime,
    umbral: float = 0.0,
) -> List[AlertRow]:
    rows: List[AlertRow] = []
    fecha = start_date
    while fecha <= end_date:
        for node_id, meta in targets.items():
            node_name = meta["nodeName"]
            horas_eval = meta.get("horas_corte") or list(range(0, 7))
            umbral_node = float(meta.get("umbral_alerta", umbral))
            try:
                hourly = obtener_datos_horarios_dia(node_id, fecha)
            except Exception:
                continue

            detalle_list: List[str] = []
            horas_detectadas: List[int] = []
            max_h = 0.0

            for h in horas_eval:
                v = float(hourly.get(h, 0.0))
                if v > max_h:
                    max_h = v
                if v > umbral_node:
                    detalle_list.append(f"{h:02d}:00({v:.1f})")
                    horas_detectadas.append(int(h))

            if not detalle_list:
                continue

            # Presentación: respetar el horario de corte evaluado para ese nodo.
            # Antes se recortaba a 00:00–06:00 y se ocultaba 23:00 en ventanas 23:00–06:00.
            horas_en_ventana = [h for h in horas_detectadas if h in set(horas_eval)]
            horas_resumen = _compactar_horas_a_rangos(horas_en_ventana or horas_detectadas)

            rows.append(
                AlertRow(
                    fecha=fecha.strftime("%Y-%m-%d"),
                    colegio=node_id,
                    node_id=node_id,
                    node_name=node_name,
                    horas_con_consumo=horas_resumen,
                    horas_con_consumo_detalle="; ".join(detalle_list),
                    consumo_max_hora_00_06=round(max_h, 3),
                )
            )
        fecha += timedelta(days=1)
    return rows


def _cliente_y_establecimiento(colegio: str, targets: Dict[str, Dict[str, Any]]) -> Tuple[str, str]:
    meta = targets.get(colegio) or {}
    cliente = str(meta.get("cliente", "—"))
    establecimiento = str(meta.get("establecimiento") or meta.get("nodeName") or colegio)
    return cliente, establecimiento


def guardar_csv(
    rows: List[AlertRow],
    out_path: Path,
    targets: Dict[str, Dict[str, Any]],
) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(
            [
                "fecha",
                "cliente",
                "establecimiento",
                "colegio",
                "nodeId",
                "nodeName",
                "horas_detectadas",
                "consumo_max_hora",
            ]
        )
        for r in rows:
            cli, est = _cliente_y_establecimiento(r.colegio, targets)
            w.writerow(
                [
                    r.fecha,
                    cli,
                    est,
                    r.colegio,
                    r.node_id,
                    r.node_name,
                    r.horas_con_consumo,
                    r.consumo_max_hora_00_06,
                ]
            )


def _agregar_grafico_anillo_control_nocturno(
    doc: Document,
    *,
    total_puntos: int,
    total_fuera: int,
    total_cumplen: int,
    chart_path: Path,
) -> None:
    """
    Gráfico de anillo al estilo del reporte de puntos en cero: fuera de control vs cumpliendo corte.
    """
    if total_puntos <= 0:
        return
    try:
        fuera = max(total_fuera, 0)
        cumplen = max(total_cumplen, 0)
        # Coherencia numérica (por si difiere en 1 por redondeo)
        if fuera + cumplen != total_puntos:
            cumplen = max(total_puntos - fuera, 0)

        valores = [fuera, cumplen]
        etiquetas = ["Fuera de control", "Cumpliendo corte"]
        colores = ["#E74C3C", "#2ECC71"]

        fig, ax = plt.subplots(figsize=(5.5, 4.5))
        ax.pie(
            valores,
            labels=etiquetas,
            colors=colores,
            autopct=lambda pct: (
                f"{pct:.1f}%\n({int(round(pct * total_puntos / 100.0))})"
                if total_puntos > 0
                else "0%"
            ),
            startangle=90,
            pctdistance=0.8,
            textprops={"fontsize": 9},
        )
        centre_circle = plt.Circle((0, 0), 0.55, fc="white")
        fig.gca().add_artist(centre_circle)
        ax.axis("equal")
        ax.set_title(
            "Distribución de puntos (fuera de control / cumpliendo corte)",
            fontsize=11,
        )
        ax.text(
            0,
            0,
            f"Total\n{total_puntos}",
            ha="center",
            va="center",
            fontsize=11,
            weight="bold",
        )

        chart_path.parent.mkdir(parents=True, exist_ok=True)
        fig.tight_layout()
        fig.savefig(chart_path, dpi=200)
        plt.close(fig)

        if chart_path.exists():
            doc.add_paragraph("Gráfica de distribución de puntos:")
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.add_paragraph("")
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass


def extraer_eventos(rows: List[AlertRow]) -> List[AlertEvent]:
    eventos: List[AlertEvent] = []
    patron = re.compile(r"(\d{2}):00\(([\d\.]+)\)")
    for r in rows:
        for hh, vv in patron.findall(r.horas_con_consumo_detalle):
            try:
                eventos.append(
                    AlertEvent(
                        fecha=r.fecha,
                        colegio=r.colegio,
                        node_id=r.node_id,
                        node_name=r.node_name,
                        hour=int(hh),
                        value=float(vv),
                    )
                )
            except Exception:
                continue
    return eventos


def crear_reporte_word_control(
    rows: List[AlertRow],
    output_docx: Path,
    desde: datetime,
    hasta: datetime,
    umbral: float,
    targets: Dict[str, Dict[str, Any]],
) -> Path:
    doc = Document()

    title = doc.add_heading("REPORTE DE CONTROL NOCTURNO (CORTES PROGRAMADOS)", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(204, 0, 0)
    run.bold = True
    run.font.size = Pt(20)

    gen_para = doc.add_paragraph(
        f"Reporte generado: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')} (hora local)"
    )
    gen_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")
    doc.add_heading("RESUMEN EJECUTIVO", 1)

    total_puntos = len(targets)
    puntos_fuera = {(r.colegio, r.node_id) for r in rows}
    total_fuera = len(puntos_fuera)
    total_cumplen = max(total_puntos - total_fuera, 0)

    umbrales = sorted(
        {
            float(meta.get("umbral_alerta", umbral))
            for meta in targets.values()
        }
    )
    if len(umbrales) == 1:
        umbral_txt = f"> {umbrales[0]:g} (mismo para todos)"
    elif umbrales:
        umbral_txt = (
            f"según Excel por establecimiento "
            f"(desde > {umbrales[0]:g} hasta > {umbrales[-1]:g})"
        )
    else:
        umbral_txt = f"> {umbral:g}"

    resumen = (
        f"Periodo evaluado: {desde:%d-%m-%Y} a {hasta:%d-%m-%Y}\n"
        f"Ventana de control: según horario de corte de cada establecimiento (Excel).\n"
        f"Umbral de alerta: {umbral_txt}\n"
        f"Cantidad de puntos analizados: {total_puntos}\n"
        f"Puntos cumpliendo corte: {total_cumplen}\n"
        f"Puntos fuera de control: {total_fuera}\n"
        f"Porcentaje fuera de control: "
        f"{(total_fuera / total_puntos * 100) if total_puntos else 0:.2f}%"
    )
    doc.add_paragraph(resumen)
    doc.add_paragraph("")

    chart_png = output_docx.parent / f"{output_docx.stem}_anillo.png"
    _agregar_grafico_anillo_control_nocturno(
        doc,
        total_puntos=total_puntos,
        total_fuera=total_fuera,
        total_cumplen=total_cumplen,
        chart_path=chart_png,
    )

    # Tabla principal
    doc.add_heading("TABLA DE PUNTOS FUERA DE CONTROL", 1)
    if not rows:
        doc.add_paragraph("No se detectaron puntos fuera de control en su ventana de corte programada.")
    else:
        doc.add_paragraph(
            "Se listan los puntos que registraron consumo durante la ventana de corte "
            "programada (ordenados por máximo horario, de mayor a menor):"
        )
        # Prioridad operativa: primero los mayores caudales (como en revisión de alertas).
        rows_sorted = sorted(
            rows,
            key=lambda r: (-float(r.consumo_max_hora_00_06), r.fecha, r.colegio, r.node_id),
        )
        table_rows = [
            ("Fecha", "Cliente", "Establecimiento", "Horas detectadas", "Máximo por hora", "ID WES")
        ]
        for r in rows_sorted:
            cli, est = _cliente_y_establecimiento(r.colegio, targets)
            table_rows.append(
                (
                    datetime.strptime(r.fecha, "%Y-%m-%d").strftime("%d-%m-%Y")
                    if r.fecha and len(r.fecha) == 10
                    else r.fecha,
                    cli,
                    est,
                    r.horas_con_consumo,
                    _fmt_m3h(r.consumo_max_hora_00_06),
                    r.node_id,
                )
            )

        table = doc.add_table(rows=len(table_rows), cols=6)
        table.style = "Light Grid Accent 1"
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(1.1)
        table.columns[2].width = Inches(1.45)
        table.columns[3].width = Inches(2.7)  # mejora compaginado
        table.columns[4].width = Inches(1.0)
        table.columns[5].width = Inches(0.9)

        header_cells = table.rows[0].cells
        for i, header in enumerate(table_rows[0]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            try:
                shading_xml = (
                    '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    'w:val="clear" w:fill="4472C4"/>'
                )
                shading = parse_xml(shading_xml)
                tc_pr = header_cells[i]._element.get_or_add_tcPr()
                if tc_pr.find(qn("w:shd")) is None:
                    tc_pr.append(shading)
            except Exception:
                pass

        _cols_centro = {0, 1, 2, 4, 5}  # Fecha, Cliente, Establecimiento, Max, ID
        for row_idx, row_data in enumerate(table_rows[1:], start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value)
                if col_idx in _cols_centro:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                if row_idx % 2 == 0:
                    try:
                        shading_xml = (
                            '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                            'w:val="clear" w:fill="F2F2F2"/>'
                        )
                        shading = parse_xml(shading_xml)
                        tc_pr = cell._element.get_or_add_tcPr()
                        if tc_pr.find(qn("w:shd")) is None:
                            tc_pr.append(shading)
                    except Exception:
                        pass

    doc.add_paragraph("")
    nota = doc.add_paragraph(
        "Nota: un punto se marca 'fuera de control' cuando presenta consumo mayor al umbral "
        "dentro de la ventana de corte asignada en el Excel."
    )
    nota.runs[0].font.italic = True
    nota.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


def _fmt_m3h(value: float) -> str:
    """Formato compacto de m³/h (evita 0.400; mantiene decimales útiles)."""
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)
    s = f"{v:.3f}".rstrip("0").rstrip(".")
    return s if s else "0"


def convertir_docx_a_pdf(docx_path: Path) -> Optional[Path]:
    """
    Convierte DOCX a PDF.

    Orden (cloud Linux y PC):
    1) LibreOffice / soffice (mismo resultado visual que los PDF antiguos del equipo)
    2) docx2pdf (Windows + Microsoft Word)
    """
    import subprocess

    pdf_path = docx_path.with_suffix(".pdf")
    docx_path = Path(docx_path).resolve()
    out_dir = docx_path.parent

    for bin_name in ("soffice", "libreoffice"):
        exe = shutil.which(bin_name)
        if not exe:
            continue
        try:
            subprocess.run(
                [
                    exe,
                    "--headless",
                    "--norestore",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(out_dir),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                timeout=180,
            )
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                return pdf_path
        except Exception as e:
            print(f"[DEBUG] {bin_name} falló al convertir a PDF: {e}")

    try:
        from docx2pdf import convert  # type: ignore

        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return pdf_path
    except Exception as e:
        print(f"[DEBUG] docx2pdf falló: {e}")

    return None


def guardar_pdf_simple(
    rows: List[AlertRow],
    out_path: Path,
    desde: datetime,
    hasta: datetime,
    umbral: float,
    targets: Dict[str, Dict[str, Any]],
) -> None:
    """
    PDF alternativo con matplotlib si docx2pdf no esta disponible.
    """
    out_path.parent.mkdir(parents=True, exist_ok=True)
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages

    resumen = [
        "REPORTE DE CONTROL NOCTURNO (CORTES PROGRAMADOS)",
        "Clientes monitoreo + control | Ventana según Excel",
        f"Periodo: {desde:%d-%m-%Y} a {hasta:%d-%m-%Y}",
        f"Umbral de alerta: según Excel por establecimiento (default > {umbral:g})",
        f"Puntos fuera de control: {len(rows)}",
    ]

    with PdfPages(out_path) as pdf:
        # Portada
        fig = plt.figure(figsize=(11.69, 8.27))
        fig.clf()
        ax = fig.add_subplot(111)
        ax.axis("off")
        y = 0.95
        for line in resumen:
            ax.text(0.02, y, line, fontsize=12, va="top")
            y -= 0.05
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)

        # Tabla (paginada)
        if not rows:
            fig = plt.figure(figsize=(11.69, 8.27))
            fig.clf()
            ax = fig.add_subplot(111)
            ax.axis("off")
            ax.text(0.02, 0.9, "Sin alertas.", fontsize=13, va="top")
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)
            return

        page_size = 22
        headers = ["Fecha", "Cliente", "Establecimiento", "Horas detectadas", "Máx. hora", "ID WES"]
        rows_sorted = sorted(
            rows,
            key=lambda r: (-float(r.consumo_max_hora_00_06), r.fecha, r.colegio, r.node_id),
        )
        for start in range(0, len(rows_sorted), page_size):
            chunk = rows_sorted[start : start + page_size]
            table_data = [headers]
            for r in chunk:
                cli, est = _cliente_y_establecimiento(r.colegio, targets)
                fecha_txt = (
                    datetime.strptime(r.fecha, "%Y-%m-%d").strftime("%d-%m-%Y")
                    if r.fecha and len(r.fecha) == 10
                    else r.fecha
                )
                table_data.append(
                    [
                        fecha_txt,
                        cli,
                        est,
                        r.horas_con_consumo,
                        _fmt_m3h(r.consumo_max_hora_00_06),
                        r.node_id,
                    ]
                )

            fig = plt.figure(figsize=(11.69, 8.27))
            fig.clf()
            ax = fig.add_subplot(111)
            ax.axis("off")
            ax.set_title("Detalle de alertas detectadas", fontsize=12, pad=10)
            table = ax.table(cellText=table_data, loc="center", cellLoc="left")
            table.auto_set_font_size(False)
            table.set_fontsize(7)
            table.scale(1, 1.02)
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)


def generar_reporte_control_nocturno(
    desde: datetime,
    hasta: datetime,
    umbral: float = 0.0,
    excel_path: Optional[Path] = None,
) -> Tuple[List[AlertRow], Path, Path, Path]:
    path_usado = excel_path or default_excel_path()
    targets = cargar_targets_desde_excel(path_usado)

    rows = analizar_control_nocturno(targets, desde, hasta, umbral=umbral)

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    base_name = f"control_nocturno_{desde:%Y%m%d}_{hasta:%Y%m%d}_{ts}"
    base_dir = Path("reports") / "control_nocturno"
    out_csv = base_dir / f"{base_name}.csv"
    out_docx = base_dir / f"{base_name}.docx"
    out_pdf = base_dir / f"{base_name}.pdf"

    guardar_csv(rows, out_csv, targets)
    crear_reporte_word_control(rows, out_docx, desde, hasta, umbral, targets)
    pdf_convertido = convertir_docx_a_pdf(out_docx)
    if pdf_convertido and pdf_convertido.exists():
        out_pdf = pdf_convertido
    else:
        guardar_pdf_simple(rows, out_pdf, desde, hasta, umbral, targets)

    return rows, out_csv, out_docx, out_pdf


def main() -> int:
    parser = argparse.ArgumentParser(description="Reporte de control nocturno (por Excel)")
    parser.add_argument("--desde", help="Fecha inicio YYYY-MM-DD")
    parser.add_argument("--hasta", help="Fecha fin YYYY-MM-DD")
    parser.add_argument("--umbral", type=float, default=0.0, help="Umbral minimo (default: 0)")
    parser.add_argument(
        "--excel",
        type=str,
        default="",
        help=(
            "Ruta al Excel de control nocturno (por defecto: "
            "wes-scripts/reports/HORARIOS CONTROL NOCTURNO.xlsx)."
        ),
    )
    args = parser.parse_args()

    hoy = datetime.now()
    # Requerimiento operativo: por defecto siempre la madrugada del dia solicitado (hoy).
    # Solo usar otro periodo cuando el usuario lo indique por --desde/--hasta.
    default_hasta = hoy.date()
    default_desde = hoy.date()

    desde = datetime.strptime(args.desde, "%Y-%m-%d") if args.desde else datetime.combine(default_desde, datetime.min.time())
    hasta = datetime.strptime(args.hasta, "%Y-%m-%d") if args.hasta else datetime.combine(default_hasta, datetime.min.time())

    excel_arg = Path(args.excel) if args.excel.strip() else None

    rows, out_csv, out_docx, out_pdf = generar_reporte_control_nocturno(desde, hasta, umbral=args.umbral, excel_path=excel_arg)

    print("=" * 72)
    print("REPORTE CONTROL NOCTURNO")
    print("=" * 72)
    print(f"Periodo: {desde:%Y-%m-%d} a {hasta:%Y-%m-%d}")
    print(f"Alertas detectadas: {len(rows)}")
    print(f"CSV: {out_csv}")
    print(f"DOCX: {out_docx}")
    print(f"PDF: {out_pdf}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

