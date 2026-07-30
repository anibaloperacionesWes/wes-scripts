"""
Reporte de consumo y diferencias mes a mes para un nodo WES (totalM3 diario API).

Uso:
  python generar_reporte_diferencias_mensuales_nodo.py --node-id 000025-01 --desde 2025-05 --hasta 2026-04
"""

from __future__ import annotations

import argparse
import calendar
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from generar_reporte_word import (
    NODE_NAMES,
    acl_node_base_url,
    add_logo_to_header,
    fetch_json,
    flatten_measures,
    format_number_chilean,
    normalize_measures_payload,
)

MESES_ES = (
    "",
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)


@dataclass
class MesResumen:
    year: int
    month: int
    etiqueta: str
    m3: float
    dias_con_dato: int
    promedio_diario: float
    diff_m3_vs_anterior: Optional[float]
    diff_pct_vs_anterior: Optional[float]


def _parse_yyyy_mm(text: str) -> Tuple[int, int]:
    parts = text.strip().split("-")
    if len(parts) != 2:
        raise ValueError(f"Formato esperado YYYY-MM, recibido: {text!r}")
    return int(parts[0]), int(parts[1])


def _mes_siguiente(year: int, month: int) -> Tuple[int, int]:
    if month == 12:
        return year + 1, 1
    return year, month + 1


def _iter_meses(desde: Tuple[int, int], hasta: Tuple[int, int]) -> List[Tuple[int, int]]:
    y, m = desde
    hy, hm = hasta
    out: List[Tuple[int, int]] = []
    while (y, m) <= (hy, hm):
        out.append((y, m))
        y, m = _mes_siguiente(y, m)
    return out


def _to_ddmmyyyy(d: date) -> str:
    return d.strftime("%d%m%Y")


def _etiqueta_mes(year: int, month: int) -> str:
    return f"{MESES_ES[month].capitalize()} {year}"


def _nombre_nodo(node_id: str) -> str:
    return NODE_NAMES.get(node_id, node_id)


def _fetch_medidas_diarias(node_id: str, start: date, end: date):
    """Medidas diarias (totalM3) entre start y end inclusive."""
    end_api = end
    raw = fetch_json(
        f"{acl_node_base_url()}/nodes/measures/dates",
        params=[
            ("id", node_id),
            ("start", _to_ddmmyyyy(start)),
            ("end", _to_ddmmyyyy(end_api)),
        ],
    )
    payload = normalize_measures_payload(raw, node_id)
    measures = flatten_measures(payload)
    return [m for m in measures if start <= m.date.date() <= end]


def _agregar_por_mes(
    measures,
    meses: List[Tuple[int, int]],
) -> List[MesResumen]:
    por_mes: Dict[Tuple[int, int], List[float]] = defaultdict(list)
    for m in measures:
        d = m.date.date()
        por_mes[(d.year, d.month)].append(float(m.total_m3))

    filas: List[MesResumen] = []
    anterior_m3: Optional[float] = None
    for year, month in meses:
        vals = por_mes.get((year, month), [])
        total = sum(vals)
        dias = len(vals)
        prom = total / dias if dias else 0.0
        diff_m3 = None
        diff_pct = None
        if anterior_m3 is not None:
            diff_m3 = total - anterior_m3
            if abs(anterior_m3) > 1e-9:
                diff_pct = (diff_m3 / anterior_m3) * 100.0
            elif abs(total) > 1e-9:
                diff_pct = None
            else:
                diff_pct = 0.0
        filas.append(
            MesResumen(
                year=year,
                month=month,
                etiqueta=_etiqueta_mes(year, month),
                m3=round(total, 2),
                dias_con_dato=dias,
                promedio_diario=round(prom, 2),
                diff_m3_vs_anterior=round(diff_m3, 2) if diff_m3 is not None else None,
                diff_pct_vs_anterior=round(diff_pct, 1) if diff_pct is not None else None,
            )
        )
        anterior_m3 = total
    return filas


def _generar_grafico(filas: List[MesResumen], node_id: str, node_name: str, out_path: Path) -> Path:
    labels = [f"{MESES_ES[r.month][:3].capitalize()}\n{r.year % 100:02d}" for r in filas]
    valores = [r.m3 for r in filas]
    diffs = [r.diff_m3_vs_anterior if r.diff_m3_vs_anterior is not None else 0.0 for r in filas]

    fig, ax1 = plt.subplots(figsize=(12, 5))
    x = range(len(labels))
    bars = ax1.bar(x, valores, color="#3498DB", alpha=0.85, label="Consumo mensual (m³)")
    ax1.set_ylabel("m³")
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(labels, fontsize=9)
    ax1.set_title(f"Consumo mensual — {node_name} ({node_id})")

    ax2 = ax1.twinx()
    ax2.plot(x[1:], diffs[1:], color="#E74C3C", marker="o", linewidth=2, label="Δ vs mes anterior")
    ax2.axhline(0, color="#7F8C8D", linewidth=0.8, linestyle="--")
    ax2.set_ylabel("Δ m³ vs mes anterior")

    for i, b in enumerate(bars):
        h = b.get_height()
        if h > 0:
            ax1.text(b.get_x() + b.get_width() / 2, h, f"{h:.0f}", ha="center", va="bottom", fontsize=7)

    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=200)
    plt.close(fig)
    return out_path


def _generar_excel(
    filas: List[MesResumen],
    node_id: str,
    node_name: str,
    desde: Tuple[int, int],
    hasta: Tuple[int, int],
    out_path: Path,
) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Diferencias mensuales"

    ws["A1"] = f"Reporte diferencias mensuales — {node_name}"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = f"Nodo: {node_id}"
    ws["A3"] = f"Periodo: {_etiqueta_mes(desde[0], desde[1])} a {_etiqueta_mes(hasta[0], hasta[1])}"
    ws["A4"] = f"Fuente: API WES totalM3 diario ({acl_node_base_url()})"
    ws["A5"] = f"Generado: {datetime.now():%Y-%m-%d %H:%M}"

    headers = [
        "Mes",
        "Año",
        "Etiqueta",
        "Días con dato",
        "Consumo m³",
        "Promedio diario m³",
        "Δ m³ vs mes anterior",
        "Δ % vs mes anterior",
    ]
    hr = 7
    header_fill = PatternFill("solid", fgColor="1F4788")
    header_font = Font(color="FFFFFF", bold=True)
    for col, h in enumerate(headers, start=1):
        c = ws.cell(row=hr, column=col, value=h)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="center", wrap_text=True)

    for i, r in enumerate(filas, start=hr + 1):
        ws.cell(row=i, column=1, value=MESES_ES[r.month].capitalize())
        ws.cell(row=i, column=2, value=r.year)
        ws.cell(row=i, column=3, value=r.etiqueta)
        ws.cell(row=i, column=4, value=r.dias_con_dato)
        ws.cell(row=i, column=5, value=r.m3)
        ws.cell(row=i, column=6, value=r.promedio_diario)
        ws.cell(row=i, column=7, value=r.diff_m3_vs_anterior if r.diff_m3_vs_anterior is not None else "—")
        pct = r.diff_pct_vs_anterior
        ws.cell(row=i, column=8, value=f"{pct:.1f}%" if pct is not None else "—")

    total_m3 = sum(r.m3 for r in filas)
    tr = hr + len(filas) + 2
    ws.cell(row=tr, column=4, value="TOTAL periodo").font = Font(bold=True)
    ws.cell(row=tr, column=5, value=round(total_m3, 2)).font = Font(bold=True)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 16
    ws.column_dimensions["C"].width = 22

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def _generar_word(
    filas: List[MesResumen],
    node_id: str,
    node_name: str,
    desde: Tuple[int, int],
    hasta: Tuple[int, int],
    chart_path: Optional[Path],
    out_path: Path,
) -> None:
    doc = Document()
    add_logo_to_header(doc)

    t = doc.add_heading("Diferencias de consumo por mes", level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Nodo: {node_name} ({node_id})")
    doc.add_paragraph(
        f"Periodo analizado: {_etiqueta_mes(desde[0], desde[1])} — {_etiqueta_mes(hasta[0], hasta[1])}"
    )
    doc.add_paragraph(
        "Los valores provienen de la suma del consumo diario (totalM3) registrado en la API WES "
        "para cada mes civil."
    )

    if chart_path and chart_path.exists():
        doc.add_paragraph("")
        doc.add_picture(str(chart_path), width=Inches(6.2))

    doc.add_paragraph("")
    doc.add_heading("Tabla resumen", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Mes"
    hdr[1].text = "Consumo (m³)"
    hdr[2].text = "Días con dato"
    hdr[3].text = "Δ m³ vs mes ant."
    hdr[4].text = "Δ % vs mes ant."

    max_mes = max(filas, key=lambda r: r.m3)
    min_mes = min(filas, key=lambda r: r.m3)
    total = sum(r.m3 for r in filas)

    for r in filas:
        row = table.add_row().cells
        row[0].text = r.etiqueta
        row[1].text = format_number_chilean(r.m3, 1)
        row[2].text = str(r.dias_con_dato)
        if r.diff_m3_vs_anterior is None:
            row[3].text = "—"
            row[4].text = "—"
        else:
            row[3].text = format_number_chilean(r.diff_m3_vs_anterior, 1)
            row[4].text = (
                f"{r.diff_pct_vs_anterior:.1f}%" if r.diff_pct_vs_anterior is not None else "—"
            )

    doc.add_paragraph("")
    doc.add_heading("Observaciones", level=1)
    p = doc.add_paragraph(
        f"En el periodo, el consumo acumulado fue {format_number_chilean(total, 1)} m³. "
        f"El mes de mayor consumo fue {max_mes.etiqueta} ({format_number_chilean(max_mes.m3, 1)} m³) "
        f"y el de menor consumo fue {min_mes.etiqueta} ({format_number_chilean(min_mes.m3, 1)} m³)."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    variaciones = [
        r for r in filas if r.diff_m3_vs_anterior is not None and r.diff_pct_vs_anterior is not None
    ]
    if variaciones:
        sube = max(variaciones, key=lambda r: r.diff_m3_vs_anterior or 0.0)
        baja = min(variaciones, key=lambda r: r.diff_m3_vs_anterior or 0.0)
        doc.add_paragraph(
            f"La mayor subida respecto al mes anterior se registró en {sube.etiqueta} "
            f"({format_number_chilean(sube.diff_m3_vs_anterior or 0.0, 1)} m³, "
            f"{sube.diff_pct_vs_anterior:.1f}%). "
            f"La mayor baja fue en {baja.etiqueta} "
            f"({format_number_chilean(baja.diff_m3_vs_anterior or 0.0, 1)} m³, "
            f"{baja.diff_pct_vs_anterior:.1f}%)."
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def generar_reporte_diferencias_mensuales(
    node_id: str,
    desde: Tuple[int, int],
    hasta: Tuple[int, int],
    out_dir: Optional[Path] = None,
) -> Tuple[Path, Path, Path]:
    meses = _iter_meses(desde, hasta)
    start = date(desde[0], desde[1], 1)
    last_day = calendar.monthrange(hasta[0], hasta[1])[1]
    end = date(hasta[0], hasta[1], last_day)

    node_name = _nombre_nodo(node_id)
    measures = _fetch_medidas_diarias(node_id, start, end)
    filas = _agregar_por_mes(measures, meses)

    base = out_dir or (Path("reports") / "diferencias_mensuales" / node_id)
    base.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    stem = f"diferencias_mensuales_{node_id}_{desde[0]}{desde[1]:02d}_{hasta[0]}{hasta[1]:02d}_{ts}"

    xlsx = base / f"{stem}.xlsx"
    docx = base / f"{stem}.docx"
    png = base / f"{stem}.png"

    _generar_excel(filas, node_id, node_name, desde, hasta, xlsx)
    _generar_grafico(filas, node_id, node_name, png)
    _generar_word(filas, node_id, node_name, desde, hasta, png, docx)

    return xlsx, docx, png


def main() -> int:
    parser = argparse.ArgumentParser(description="Diferencias de consumo mensual por nodo WES")
    parser.add_argument("--node-id", default="000025-01")
    parser.add_argument("--desde", default="2025-05", help="Mes inicio YYYY-MM")
    parser.add_argument("--hasta", default="2026-04", help="Mes fin YYYY-MM")
    parser.add_argument("--out-dir", default="", help="Carpeta de salida opcional")
    args = parser.parse_args()

    desde = _parse_yyyy_mm(args.desde)
    hasta = _parse_yyyy_mm(args.hasta)
    if desde > hasta:
        print("[ERROR] --desde debe ser anterior o igual a --hasta")
        return 1

    out_dir = Path(args.out_dir) if args.out_dir.strip() else None
    xlsx, docx, png = generar_reporte_diferencias_mensuales(
        args.node_id.strip(),
        desde,
        hasta,
        out_dir,
    )

    print("=" * 72)
    print("REPORTE DIFERENCIAS MENSUALES")
    print("=" * 72)
    print(f"Nodo: {args.node_id} ({_nombre_nodo(args.node_id)})")
    print(f"Periodo: {args.desde} a {args.hasta}")
    print(f"Excel: {xlsx.resolve()}")
    print(f"Word:  {docx.resolve()}")
    print(f"Gráfico: {png.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
