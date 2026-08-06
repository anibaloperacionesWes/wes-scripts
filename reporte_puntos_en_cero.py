"""
Script para generar reporte de todos los puntos de monitoreo que están marcando cero.
Conecta con todos los puntos del sistema y genera un reporte Word con los que están en cero.
"""

import os
import sys
import requests
from concurrent.futures import ThreadPoolExecutor
from exclusiones_reportes import (
    EXCLUDED_COMPANY_IDS_PUNTOS_EN_CERO,
    EXCLUDED_COMPANY_NAME_KEYWORDS,
    EXCLUDED_NODE_IDS_PUNTOS_EN_CERO,
)
from pathlib import Path
from datetime import datetime, timedelta, timezone
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from collections import defaultdict

from wes_paths import reporte_cero_dir

try:
    from zoneinfo import ZoneInfo

    _CHILE_TZ = ZoneInfo("America/Santiago")
except Exception:
    _CHILE_TZ = timezone(timedelta(hours=-4))

# Ventana extendida para diagnosticar «sin datos» / rango «en cero».
DIAS_BUSCAR_ULTIMA_MEDIDA = max(7, int(os.environ.get("WES_CERO_DIAS_ULTIMA_MEDIDA", "30")))
# Si la última medida es más reciente que esto → motivo «Desconectado» (solo listados).
HORAS_UMBRAL_DESCONECTADO = float(os.environ.get("WES_CERO_HORAS_DESCONECTADO", "48"))
# No notificar como «sin datos» si la última medida fue hace menos de N horas (<1 día).
HORAS_OMITIR_SIN_DATOS = float(os.environ.get("WES_CERO_HORAS_OMITIR_SIN_DATOS", "24"))

# Carpeta de salida por defecto: Google Drive (salvo WES_SCRIPTS_ROOT o si G: no existe).
_DEFAULT_REPORTE_CERO_DIR = reporte_cero_dir()

# Consola: UTF-8 + sustituir caracteres no representables (evita UnicodeEncodeError en PowerShell/cmd).
if sys.platform == "win32":
    try:
        import ctypes

        ctypes.windll.kernel32.SetConsoleOutputCP(65001)
    except Exception:
        pass
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, TypeError, OSError, ValueError):
        try:
            _stream.reconfigure(errors="replace")
        except (AttributeError, TypeError, OSError, ValueError):
            pass

# URLs base de las APIs
BASE_URL = "http://104.248.53.141:7003/wes/api/acl-node/v1"
ENTITY_BASE_URL = "http://104.248.53.141:7001/wes/api/acl-entities/v1"

# Fecha referencia UTC para verificación (últimos N días desde esa fecha en API); None = fecha/hora actual
FECHA_REFERENCIA_UTC: Optional[datetime] = None

# Concurrencia HTTP (reporte diario: muchos nodos × varias llamadas; antes todo secuencial).
MAX_WORKERS_CERO = max(4, int(os.environ.get("WES_REPORTE_CERO_WORKERS", "20")))


def _obtener_empresas_config() -> List[Dict[str, str]]:
    """
    Obtiene todas las empresas desde la API de configuración.
    Retorna una lista de dicts con companyId y name.
    """
    url = f"{ENTITY_BASE_URL}/configuration/companies"
    try:
        response = requests.get(url, timeout=15)
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, list):
                return data
    except Exception:
        pass
    return []


def obtener_todos_los_nodos() -> List[Dict[str, str]]:
    """
    Obtiene todos los nodos del sistema iterando por empresas.
    Retorna lista de diccionarios con nodeId, nodeName, companyId, companyName
    
    Excluye las siguientes empresas:
    - WES / Wes Spa (ID: 000000)
    - Ejército de Chile (ID: 000001)
    - Gendarmería de Chile (ID: 000004)
    """
    # IDs de empresas a excluir
    empresas_excluidas = EXCLUDED_COMPANY_IDS_PUNTOS_EN_CERO
    
    # Nombres de empresas a excluir (búsqueda case-insensitive)
    # Se busca cualquier coincidencia parcial en el nombre
    nombres_empresas_excluidas = EXCLUDED_COMPANY_NAME_KEYWORDS
    
    # IDs de nodos específicos a excluir
    nodos_excluidos = EXCLUDED_NODE_IDS_PUNTOS_EN_CERO
    
    all_nodes = []
    print("Obteniendo todos los nodos del sistema...")
    print("=" * 60)
    print("Empresas excluidas del análisis:")
    print("  - WES / Wes Spa (000000)")
    print("  - Ejército de Chile (000001)")
    print("  - Gendarmería de Chile (000004)")
    print("  - BUPA (000029) — puntos creados en app, pendiente instalación")
    print("  - Corporación Puente Alto (000010) — colegios fuera del reporte en cero")
    print("  - MOP / Ministerio de Obras Públicas (por nombre)")
    print("  - Lo Boza (por nombre)")
    print("  - Tres Montes Lucchetti (por nombre)")
    print("  - MADECO (por nombre)")
    print("=" * 60)
    print("Nodos específicos excluidos:")
    print("  - Plaza Boulevard Pajaros Sur (000013-01)")
    print("  - Rebeca Matte Bello (000016-01, 000017-01)")
    print("  - Poniente 7 (000025-03)")
    print("  - Locales de Comida (000025-05)")
    print("  - KFC (000025-06)")
    print("  - Casa Juan Lopez (000011-01)")
    print("  - Lo Boza Lavado de Vehículos (000012-01)")
    print("  - Lo Boza Pozo (000012-02)")
    print("  - Lo Boza Reutilización (000012-03)")
    print("  - Lo Boza Edificio Principal Casino (000012-04)")
    print("  - Lo Boza Matriz Principal (000012-05)")
    print("  - La Cabaña (000012-20)")
    print("  - Colegio Juan Pablo II / Las Condes (000022-01)")
    print("  - Oficina WES (000019-01)")
    print("  - Rugby CDUC (000021-08)")
    print("  - Arturo Alessandri Palma (000006-03)")
    print("  - Juana Atala de Hirmas (000017-02)")
    print("  - José Luis Araneda (000017-03)")
    print("  - Nido Cancha (000007-08)")
    print("  - Lo valledor - Pozo (000002-02)")
    print("  - Edificio Deportivo (000021-02)")
    print("  - Matriz principal 1°piso (000025-11)")
    print("  - Red de Incendio (000025-14)")
    print("  - Matriz A.A reubicado (000025-30) — solo puntos en cero")
    print("  - Riego Fundo Zapallar (000027-05)")
    print("  - Control Nido de Aguilas (000007-09)")
    print("  - Puntos dados de baja / cliente (registro auditable): registro_puntos_deshabilitados.txt")
    print("    (incl. 000025-09 Impulsión Falabella — sala bombas en reparación)")
    print("=" * 60)
    
    # Obtener empresas desde la API de configuración
    empresas_config = _obtener_empresas_config()
    if not empresas_config:
        print("[ADVERTENCIA] No se pudieron obtener empresas desde la API de configuración.")
        print("[INFO] Usando rango 000000-000100 como respaldo.")
        empresas_config = [{"companyId": f"{i:06d}", "name": ""} for i in range(101)]

    # Iterar por empresas obtenidas desde la API
    for empresa in empresas_config:
        company_id_raw = str(empresa.get("companyId", "")).strip()
        if not company_id_raw:
            continue
        company_id = company_id_raw.zfill(6)
        
        # Saltar empresas excluidas
        if company_id in empresas_excluidas:
            continue
        
        url = f"{ENTITY_BASE_URL}/companies/{company_id}"
        
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                data = response.json()
                company_name = data.get("name", "").strip() or str(empresa.get("name", "")).strip()
                
                # Verificar si la empresa debe ser excluida por nombre
                if company_name:
                    company_name_upper = company_name.upper().strip()
                    # Verificar si el nombre contiene alguna de las palabras clave excluidas
                    if any(nombre_excluido in company_name_upper for nombre_excluido in nombres_empresas_excluidas):
                        print(f"[EXCLUIDO] {company_id} ({company_name}) - Excluida por nombre")
                        continue
                
                if company_name:
                    # Obtener nodos de la empresa
                    nodes = data.get("nodes", [])
                    for node in nodes:
                        node_id = node.get("nodeId", "")
                        node_name = node.get("name", "").strip()
                        
                        # Excluir nodos específicos
                        if node_id in nodos_excluidos:
                            print(f"  [EXCLUIDO] Nodo {node_id} ({node_name}) - Excluido específicamente")
                            continue
                        
                        if node_id and node_name:
                            all_nodes.append({
                                "nodeId": node_id,
                                "nodeName": node_name,
                                "companyId": company_id,
                                "companyName": company_name
                            })
                    if nodes:
                        print(f"[OK] {company_id} ({company_name}): {len(nodes)} nodos")
        except requests.RequestException:
            # Empresa no existe o error, continuar
            pass
        except Exception as e:
            print(f"[ERROR] {company_id}: {e}")
    
    print("=" * 60)
    print(f"Total nodos encontrados: {len(all_nodes)}")
    return all_nodes


def _horas_desde_csv_medidas(csv_content: str) -> Tuple[Dict[int, float], bool]:
    """
    Parsea el cuerpo de /nodes/{id}/dates.measures.csv (encabezado TIME,VALUE + líneas de datos).
    Usa splitlines() (evita \\r colgando en Windows) y split(',', 1) por si el valor trae comas.
    """
    hourly_data: Dict[int, float] = {}
    lines = csv_content.strip().splitlines()
    if len(lines) <= 1:
        return hourly_data, False

    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
        parts = line.split(",", 1)
        if len(parts) < 2:
            continue
        time_str = parts[0].strip()
        value_str = parts[1].strip().replace(" ", "").replace(",", ".")
        try:
            if "T" in time_str:
                hour_part = time_str.split("T")[1]
                hour = int(hour_part.split(":")[0])
            else:
                hour = datetime.fromisoformat(time_str.replace("Z", "+00:00")).hour
            value = float(value_str)
        except (ValueError, TypeError, IndexError):
            continue
        hourly_data[hour] = value

    return hourly_data, True


def _parse_time_utc_medida(time_str: str) -> Optional[datetime]:
    try:
        return datetime.fromisoformat(time_str.strip().replace("Z", "+00:00"))
    except (ValueError, TypeError, AttributeError):
        return None


def _iter_medidas_csv(csv_content: str) -> List[Tuple[datetime, float]]:
    """Parsea TIME,VALUE del CSV de medidas a lista (datetime UTC aware, valor)."""
    out: List[Tuple[datetime, float]] = []
    lines = csv_content.strip().splitlines()
    if len(lines) <= 1:
        return out
    for line in lines[1:]:
        line = line.strip()
        if not line or "," not in line:
            continue
        time_str, value_str = line.split(",", 1)
        dt = _parse_time_utc_medida(time_str.strip())
        if dt is None:
            continue
        try:
            value = float(value_str.strip().replace(" ", "").replace(",", "."))
        except (ValueError, TypeError):
            continue
        out.append((dt, value))
    return out


def verificar_consumo_cero(node_id: str, dias_revisar: int = 3) -> Tuple[bool, Optional[str]]:
    """
    Verifica si un nodo está marcando cero en una ventana móvil.

    Ventana = últimas ``dias_revisar * 24`` horas hasta la referencia, en tiempo real
    (hora Chile para el criterio operativo; los timestamps de API se interpretan en UTC).

    - Si en esa ventana hay algún valor > 0 → NO está en cero (p. ej. ayer consumió
      y hoy de madrugada va en 0: no se notifica).
    - Si hay registros y todos son 0 → en cero.
    - Si no hay registros en la ventana → sin datos.

    Args:
        node_id: ID del nodo a verificar
        dias_revisar: días de ventana (1 = últimas 24 h, 3 = últimas 72 h)

    Returns:
        (esta_en_cero, mensaje_error)
    """
    global FECHA_REFERENCIA_UTC
    ref = FECHA_REFERENCIA_UTC if FECHA_REFERENCIA_UTC is not None else datetime.now(timezone.utc)
    if ref.tzinfo is None:
        ref = ref.replace(tzinfo=timezone.utc)

    dias_revisar = max(1, int(dias_revisar))
    horas_ventana = dias_revisar * 24
    inicio = ref - timedelta(hours=horas_ventana)

    # Días calendario UTC a pedir (margen +1 por cruce de huso / borde de ventana).
    dias_api = dias_revisar + 2
    url = f"{BASE_URL}/nodes/{node_id}/dates.measures.csv"
    medidas: List[Tuple[datetime, float]] = []
    errores = 0

    for dias_atras in range(dias_api):
        fecha = ref - timedelta(days=dias_atras)
        date_str = fecha.strftime("%d%m%Y")
        try:
            response = requests.get(url, params=[("start", date_str), ("end", date_str)], timeout=30)
            if response.status_code != 200:
                errores += 1
                continue
            medidas.extend(_iter_medidas_csv(response.text))
        except requests.exceptions.RequestException:
            errores += 1
            continue
        except Exception:
            continue

    en_ventana = [(dt, v) for dt, v in medidas if inicio <= dt <= ref]
    if en_ventana:
        if any(v > 0 for _, v in en_ventana):
            return False, None
        return True, None

    if errores >= dias_api:
        return False, "Sin datos disponibles"
    return False, "Sin datos disponibles"


def obtener_ultima_medida_positiva(
    node_id: str,
    dias_buscar: int = DIAS_BUSCAR_ULTIMA_MEDIDA,
) -> Tuple[Optional[datetime], Optional[float], str]:
    """
    Última marca con valor > 0 en dates.measures.csv (ventana extendida).
    Retorna (datetime UTC, valor, estado).
    """
    global FECHA_REFERENCIA_UTC
    ref = FECHA_REFERENCIA_UTC if FECHA_REFERENCIA_UTC is not None else datetime.now(timezone.utc)
    if ref.tzinfo is None:
        ref = ref.replace(tzinfo=timezone.utc)

    ultima: Optional[datetime] = None
    ultimo_val: Optional[float] = None
    errores = 0
    dias_buscar = max(1, int(dias_buscar))
    url = f"{BASE_URL}/nodes/{node_id}/dates.measures.csv"

    for dias_atras in range(dias_buscar):
        fecha = ref - timedelta(days=dias_atras)
        date_str = fecha.strftime("%d%m%Y")
        try:
            response = requests.get(
                url, params=[("start", date_str), ("end", date_str)], timeout=25
            )
            if response.status_code != 200:
                errores += 1
                continue
            for dt, val in _iter_medidas_csv(response.text):
                if val > 0 and (ultima is None or dt > ultima):
                    ultima = dt
                    ultimo_val = val
        except requests.RequestException:
            errores += 1
            continue
        except Exception:
            continue

    if ultima is not None:
        return ultima, ultimo_val, "OK"
    if errores >= dias_buscar:
        return None, None, "sin respuesta API"
    return None, None, "sin consumo >0 en ventana"


def _fmt_antiguedad(horas: float) -> str:
    if horas < 0:
        horas = 0.0
    if horas < 1:
        return "menos de 1 h"
    if horas < 48:
        return f"~{int(round(horas))} h"
    dias = max(1, int(horas // 24))
    if horas % 24 >= 12:
        dias = max(dias, int(round(horas / 24.0)))
    return f"~{dias} día{'s' if dias != 1 else ''}"


def clasificar_motivo_en_cero(
    ultima_positiva: Optional[datetime],
    estado_busqueda: str,
    *,
    dias_ventana_reporte: int,
    dias_busqueda: int = DIAS_BUSCAR_ULTIMA_MEDIDA,
    ref: Optional[datetime] = None,
) -> str:
    """
    Motivo para puntos en cero: rango desde que dejó de haber consumo > 0.
    """
    ref = ref if ref is not None else datetime.now(timezone.utc)
    if ref.tzinfo is None:
        ref = ref.replace(tzinfo=timezone.utc)

    horas_ventana = max(1, int(dias_ventana_reporte)) * 24
    if ultima_positiva is not None:
        # El tramo en cero empieza después de la última hora con consumo.
        inicio_cero = ultima_positiva + timedelta(hours=1)
        horas = _horas_desde(inicio_cero, ref) or 0.0
        return (
            f"En cero desde {_fmt_chile(inicio_cero)} Chile "
            f"({_fmt_antiguedad(horas)}; última con consumo: {_fmt_chile(ultima_positiva)} Chile). "
            f"Ventana chequeo: últimas {horas_ventana} h en 0"
        )

    if estado_busqueda == "sin respuesta API":
        return "En cero (no se pudo ampliar historial por API)"
    return (
        f"En cero ≥{dias_busqueda} días "
        f"(sin consumo >0 en historial revisado). "
        f"Ventana chequeo: últimas {horas_ventana} h en 0"
    )


def enriquecer_puntos_en_cero(
    puntos_en_cero: List[Dict],
    dias_ventana_reporte: int,
    dias_buscar: int = DIAS_BUSCAR_ULTIMA_MEDIDA,
) -> List[Dict]:
    """Agrega motivo / rango en cero a cada punto (en paralelo)."""
    if not puntos_en_cero:
        return puntos_en_cero

    ref = FECHA_REFERENCIA_UTC if FECHA_REFERENCIA_UTC is not None else datetime.now(timezone.utc)

    def _one(punto: Dict) -> Dict:
        node_id = str(punto.get("nodeId", "")).strip()
        ultima_pos, val, estado = obtener_ultima_medida_positiva(
            node_id, dias_buscar=dias_buscar
        )
        motivo = clasificar_motivo_en_cero(
            ultima_pos,
            estado,
            dias_ventana_reporte=dias_ventana_reporte,
            dias_busqueda=dias_buscar,
            ref=ref,
        )
        out = dict(punto)
        out["motivo"] = motivo
        out["ultimaPositivaChile"] = _fmt_chile(ultima_pos)
        out["ultimaPositivaValor"] = val
        out["estadoRangoCero"] = estado
        return out

    with ThreadPoolExecutor(max_workers=MAX_WORKERS_CERO) as ex:
        return list(ex.map(_one, puntos_en_cero))


def obtener_ultima_medida(
    node_id: str,
    dias_buscar: int = DIAS_BUSCAR_ULTIMA_MEDIDA,
) -> Tuple[Optional[datetime], str]:
    """
    Busca la marca TIME más reciente en dates.measures.csv (ventana extendida).
    Retorna (datetime UTC aware o None, estado).
    """
    global FECHA_REFERENCIA_UTC
    ref = FECHA_REFERENCIA_UTC if FECHA_REFERENCIA_UTC is not None else datetime.now(timezone.utc)
    ultima: Optional[datetime] = None
    errores = 0
    dias_buscar = max(1, int(dias_buscar))
    url = f"{BASE_URL}/nodes/{node_id}/dates.measures.csv"

    for dias_atras in range(dias_buscar):
        fecha = ref - timedelta(days=dias_atras)
        date_str = fecha.strftime("%d%m%Y")
        try:
            response = requests.get(url, params=[("start", date_str), ("end", date_str)], timeout=25)
            if response.status_code != 200:
                errores += 1
                continue
            for line in response.text.strip().splitlines()[1:]:
                line = line.strip()
                if not line or "," not in line:
                    continue
                time_str = line.split(",", 1)[0].strip()
                dt = _parse_time_utc_medida(time_str)
                if dt is not None and (ultima is None or dt > ultima):
                    ultima = dt
        except requests.RequestException:
            errores += 1
            continue
        except Exception:
            continue

    if ultima is not None:
        return ultima, "OK"
    if errores >= dias_buscar:
        return None, "sin respuesta API"
    return None, "sin medida en ventana"


def _fmt_chile(dt: Optional[datetime]) -> str:
    if dt is None:
        return "—"
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(_CHILE_TZ).strftime("%d-%m-%Y %H:%M")


def _horas_desde(dt: Optional[datetime], ref: Optional[datetime] = None) -> Optional[float]:
    if dt is None:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    ahora = ref if ref is not None else datetime.now(timezone.utc)
    if ahora.tzinfo is None:
        ahora = ahora.replace(tzinfo=timezone.utc)
    return (ahora - dt).total_seconds() / 3600.0


def clasificar_motivo_sin_datos(
    ultima: Optional[datetime],
    estado_busqueda: str,
    *,
    dias_ventana_reporte: int,
    dias_busqueda: int = DIAS_BUSCAR_ULTIMA_MEDIDA,
    umbral_horas_desconectado: float = HORAS_UMBRAL_DESCONECTADO,
    ref: Optional[datetime] = None,
) -> str:
    """
    Motivo para puntos **notificados** como sin datos (≥24 h sin medida).

    No usa «Desconectado» aquí: los cortes <24 h se omiten del informe.
    Los listados se tratan como equipo que sigue/volverá conectado, con hueco
    desde la última medida (al reconectar no recupera el histórico de ese periodo).
    """
    del umbral_horas_desconectado  # reservado; el filtro <24 h vive en enriquecer_*
    horas = _horas_desde(ultima, ref)
    ultima_txt = _fmt_chile(ultima)

    if ultima is not None and horas is not None:
        if horas < 0:
            horas = 0.0
        dias = max(1, int(horas // 24))
        if horas % 24 >= 12:
            dias = max(dias, int(round(horas / 24.0)))
        if horas < 48:
            antiguedad = f"~{int(round(horas))} h"
        else:
            antiguedad = f"{dias} día{'s' if dias != 1 else ''}"
        return (
            f"Equipo conectado — sin datos desde {ultima_txt} Chile "
            f"({antiguedad}; no recupera histórico)"
        )

    if estado_busqueda == "sin respuesta API":
        return "Sin respuesta de API de medidas"
    return (
        f"Equipo conectado — sin medida en ≥{dias_busqueda} días "
        f"(no recupera histórico; ventana reporte: "
        f"{dias_ventana_reporte} día{'s' if dias_ventana_reporte != 1 else ''})"
    )


def _set_cell_width(cell, width_inches: float) -> None:
    """Fija ancho de celda en DXA (twips*20) para que Word/PDF respeten el layout."""
    width = Inches(width_inches)
    cell.width = width
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = parse_xml(
            f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:w="{width.twips}" w:type="dxa"/>'
        )
        tcPr.append(tcW)
    else:
        tcW.set(qn("w:w"), str(width.twips))
        tcW.set(qn("w:type"), "dxa")


def _aplicar_anchos_tabla(table, widths_inches: List[float]) -> None:
    """Layout fijo mitad/mitad u otra proporción por columna."""
    try:
        table.autofit = False
    except Exception:
        pass
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(
            '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tblPr)
    # Preferir layout fijo
    if tblPr.find(qn("w:tblLayout")) is None:
        tblPr.append(
            parse_xml(
                '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'w:type="fixed"/>'
            )
        )
    total = sum(widths_inches)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblPr.append(
            parse_xml(
                f'<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                f'w:w="{Inches(total).twips}" w:type="dxa"/>'
            )
        )
    else:
        tblW.set(qn("w:w"), str(Inches(total).twips))
        tblW.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, w in enumerate(widths_inches):
            if idx < len(row.cells):
                _set_cell_width(row.cells[idx], w)


def enriquecer_puntos_sin_datos(
    puntos_sin_datos: List[Dict],
    dias_ventana_reporte: int,
    dias_buscar: int = DIAS_BUSCAR_ULTIMA_MEDIDA,
    horas_omitir: float = HORAS_OMITIR_SIN_DATOS,
) -> Tuple[List[Dict], List[Dict]]:
    """
    Agrega motivo / última medida a cada punto sin datos (en paralelo).

    Filtra (no notifica) puntos con última medida hace menos de ``horas_omitir``
    (por defecto 24 h): corte reciente / equipo desconectado <1 día.

    Retorna (a_reportar, omitidos_corte_reciente).
    """
    if not puntos_sin_datos:
        return [], []

    ref = FECHA_REFERENCIA_UTC if FECHA_REFERENCIA_UTC is not None else datetime.now(timezone.utc)

    def _one(punto: Dict) -> Dict:
        node_id = str(punto.get("nodeId", "")).strip()
        ultima, estado = obtener_ultima_medida(node_id, dias_buscar=dias_buscar)
        horas = _horas_desde(ultima, ref)
        motivo = clasificar_motivo_sin_datos(
            ultima,
            estado,
            dias_ventana_reporte=dias_ventana_reporte,
            dias_busqueda=dias_buscar,
            ref=ref,
        )
        out = dict(punto)
        out["motivo"] = motivo
        out["ultimaMedidaChile"] = _fmt_chile(ultima)
        out["haceHoras"] = round(horas, 1) if horas is not None else None
        out["estadoUltima"] = estado
        out["omitirPorCorteReciente"] = bool(
            horas is not None and horas < float(horas_omitir)
        )
        return out

    with ThreadPoolExecutor(max_workers=MAX_WORKERS_CERO) as ex:
        enriquecidos = list(ex.map(_one, puntos_sin_datos))

    a_reportar = [p for p in enriquecidos if not p.get("omitirPorCorteReciente")]
    omitidos = [p for p in enriquecidos if p.get("omitirPorCorteReciente")]
    return a_reportar, omitidos


def _parse_alert_hour(alert: Dict) -> str:
    """Extrae hora HH:MM desde un registro de alerta."""
    for key in ["date", "datetime", "time", "timestamp", "createdAt", "updatedAt"]:
        value = alert.get(key)
        if isinstance(value, str) and value.strip():
            try:
                dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
                return dt.strftime("%H:%M")
            except Exception:
                continue
    return ""


def _parse_alert_measure(alert: Dict) -> Optional[float]:
    """Extrae el volumen/medida desde un registro de alerta."""
    for key in ["measure", "value", "volume", "consumption", "m3", "m3_h"]:
        value = alert.get(key)
        try:
            if value is None:
                continue
            if isinstance(value, str):
                cleaned = value.replace(",", ".").replace("m3", "").replace("m³", "").strip()
                return float(cleaned)
            return float(value)
        except (ValueError, TypeError):
            continue
    return None


def obtener_alertas_dia(node_id: str, fecha: datetime) -> List[Dict]:
    """Obtiene alertas del nodo para un día específico."""
    date_str = fecha.strftime("%d%m%Y")
    url = f"{BASE_URL}/nodes/myalert/alerts"
    params = [("id", node_id), ("start", date_str), ("end", date_str)]
    try:
        response = requests.get(url, params=params, timeout=15)
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, list):
                return data
    except Exception:
        pass
    return []


def _resumen_alertas_un_nodo(nodo: Dict, fecha: datetime) -> Optional[Dict]:
    node_id = nodo.get("nodeId")
    if not node_id:
        return None
    alerts = obtener_alertas_dia(node_id, fecha)
    if not alerts:
        return None

    total_volume = 0.0
    max_alert = None
    max_measure = -1.0

    for alert in alerts:
        measure = _parse_alert_measure(alert)
        if measure is None:
            continue
        total_volume += measure
        if measure > max_measure:
            max_measure = measure
            max_alert = alert

    alert_hour = _parse_alert_hour(max_alert) if max_alert else ""
    return {
        "nodeId": node_id,
        "nodeName": nodo.get("nodeName", ""),
        "companyName": nodo.get("companyName", ""),
        "alertVolume": f"{total_volume:.1f} m3",
        "alertVolumeValue": total_volume,
        "alertHour": alert_hour or "-",
    }


def construir_resumen_alertas(nodos: List[Dict], fecha: datetime) -> List[Dict]:
    """Construye resumen de alertas por nodo para un día (peticiones en paralelo)."""
    if not nodos:
        return []

    def _one(n: Dict) -> Optional[Dict]:
        return _resumen_alertas_un_nodo(n, fecha)

    with ThreadPoolExecutor(max_workers=MAX_WORKERS_CERO) as ex:
        chunks = list(ex.map(_one, nodos))
    resumen = [c for c in chunks if c is not None]
    resumen.sort(key=lambda x: x.get("alertVolumeValue", 0.0), reverse=True)
    return resumen


def obtener_alertas_dia_anterior(puntos: List[Dict]) -> List[Dict]:
    """
    Obtiene alertas del día anterior para los puntos entregados y
    retorna un resumen ordenado por mayor volumen de alerta.
    """
    fecha_obj = datetime.now(timezone.utc) - timedelta(days=1)
    date_str = fecha_obj.strftime("%d%m%Y")

    resumen_por_nodo = {}

    for punto in puntos:
        node_id = punto.get("nodeId")
        if not node_id:
            continue

        url = f"{BASE_URL}/nodes/myalert/alerts"
        params = [("id", node_id), ("start", date_str), ("end", date_str)]
        try:
            response = requests.get(url, params=params, timeout=20)
            if response.status_code != 200:
                continue
            alerts = response.json()
            if not isinstance(alerts, list):
                continue

            # Tomar la alerta con mayor medida del día
            for alert in alerts:
                try:
                    measure = float(alert.get("measure", 0) or 0)
                except (TypeError, ValueError):
                    measure = 0
                if measure <= 0:
                    continue

                ts = alert.get("date") or alert.get("timestamp") or alert.get("datetime")
                hora = ""
                if ts:
                    try:
                        dt = datetime.fromisoformat(str(ts).replace("Z", "+00:00"))
                        hora = dt.strftime("%H:%M")
                    except Exception:
                        hora = str(ts)

                prev = resumen_por_nodo.get(node_id)
                if not prev or measure > prev["measure"]:
                    resumen_por_nodo[node_id] = {
                        "nodeId": node_id,
                        "nodeName": punto.get("nodeName", ""),
                        "companyName": punto.get("companyName", ""),
                        "measure": measure,
                        "volumen": f"{measure:.2f} m³",
                        "hora": hora,
                    }
        except requests.RequestException:
            continue

    resumen = sorted(resumen_por_nodo.values(), key=lambda x: x["measure"], reverse=True)
    return resumen


def crear_reporte_word(
    puntos_en_cero: List[Dict],
    puntos_sin_datos: List[Dict],
    total_puntos: int,
    output_dir: Path,
    alertas_resumen: Optional[List[Dict]] = None,
    alertas_fecha: Optional[str] = None,
    fecha_generacion_utc: Optional[datetime] = None,
    solo_sin_datos: bool = False,
    dias_revisar: int = 3,
    omitidos_corte_reciente: int = 0,
) -> Path:
    """
    Crea un documento Word con el reporte de puntos en cero y sin datos.
    Con ``solo_sin_datos=True`` solo incluye resumen y tabla de puntos sin datos (sin sección «en cero»).
    ``omitidos_corte_reciente``: puntos con última medida <24 h que no se notifican.
    """
    doc = Document()

    titulo_txt = (
        "REPORTE DE PUNTOS SIN DATOS DISPONIBLES"
        if solo_sin_datos
        else "REPORTE DE PUNTOS EN CERO Y SIN DATOS"
    )
    color_titulo = RGBColor(230, 126, 34) if solo_sin_datos else RGBColor(204, 0, 0)

    title = doc.add_heading(titulo_txt, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = color_titulo
    title_run.bold = True
    title_run.font.size = Pt(20)

    # Fecha de generación
    _gen = fecha_generacion_utc if fecha_generacion_utc is not None else datetime.now(timezone.utc)
    fecha_generacion = _gen.strftime("%d-%m-%Y %H:%M:%S")
    gen_para = doc.add_paragraph(f"Reporte generado: {fecha_generacion} UTC")
    gen_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    gen_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")  # Espacio

    # Resumen ejecutivo
    doc.add_heading("RESUMEN EJECUTIVO", 1)
    puntos_con_datos = total_puntos - len(puntos_sin_datos)
    if solo_sin_datos:
        summary_txt = (
            f"Total de puntos analizados: {total_puntos}\n"
            f"Puntos sin datos (notificados): {len(puntos_sin_datos)}\n"
            f"Puntos con al menos un registro en el periodo revisado: {puntos_con_datos}\n"
            f"Porcentaje sin datos (notificados): {(len(puntos_sin_datos) / total_puntos * 100) if total_puntos > 0 else 0:.2f}%"
        )
    else:
        summary_txt = (
            f"Total de puntos analizados: {total_puntos}\n"
            f"Puntos con datos disponibles: {puntos_con_datos}\n"
            f"Puntos marcando cero: {len(puntos_en_cero)}\n"
            f"Puntos sin datos (notificados): {len(puntos_sin_datos)}\n"
            f"Porcentaje en cero: {(len(puntos_en_cero) / puntos_con_datos * 100) if puntos_con_datos > 0 else 0:.2f}%\n"
            f"Porcentaje sin datos (notificados): {(len(puntos_sin_datos) / total_puntos * 100) if total_puntos > 0 else 0:.2f}%"
        )
    if omitidos_corte_reciente > 0:
        summary_txt += (
            f"\nCortes recientes omitidos (<{HORAS_OMITIR_SIN_DATOS:.0f} h desde última medida): "
            f"{omitidos_corte_reciente} (no se notifican como sin datos)"
        )
    summary_para = doc.add_paragraph(summary_txt)
    summary_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")  # Espacio

    # Gráfica de anillo
    try:
        en_cero = len(puntos_en_cero)
        sin_datos = len(puntos_sin_datos)
        otros = max(total_puntos - en_cero - sin_datos, 0)
        if solo_sin_datos:
            valores = [sin_datos, max(total_puntos - sin_datos, 0)]
            etiquetas = ["Sin datos", "Con datos"]
            colores = ["#E67E22", "#2ECC71"]
            titulo_graf = "Distribución: sin datos vs con datos"
        else:
            valores = [en_cero, sin_datos, otros]
            etiquetas = ["En cero", "Sin datos", "Con datos"]
            colores = ["#E74C3C", "#E67E22", "#2ECC71"]
            titulo_graf = "Distribución de puntos (en cero / sin datos / con datos)"

        fig, ax = plt.subplots(figsize=(5.5, 4.5))
        ax.pie(
            valores,
            labels=etiquetas,
            colors=colores,
            autopct=lambda pct: (
                f"{pct:.1f}%\n({int(round(pct * total_puntos / 100.0))})"
                if total_puntos > 0 else "0%"
            ),
            startangle=90,
            pctdistance=0.8,
            textprops={"fontsize": 9},
        )
        centre_circle = plt.Circle((0, 0), 0.55, fc="white")
        fig.gca().add_artist(centre_circle)
        ax.axis("equal")
        ax.set_title(titulo_graf, fontsize=11)

        ax.text(0, 0, f"Total\n{total_puntos}", ha="center", va="center", fontsize=11, weight="bold")

        chart_name = (
            "grafica_anillo_puntos_sin_datos.png" if solo_sin_datos else "grafica_anillo_puntos_en_cero.png"
        )
        chart_path = output_dir / chart_name
        fig.tight_layout()
        fig.savefig(chart_path, dpi=200)
        plt.close(fig)

        if chart_path.exists():
            doc.add_paragraph("Gráfica de distribución de puntos:")
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.add_paragraph("")  # Espacio
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass

    # Sección: Puntos en cero
    if not solo_sin_datos and puntos_en_cero:
        doc.add_heading("PUNTOS MARCANDO CERO", 1)
        doc.add_paragraph(
            "Los siguientes puntos registran consumo cero en toda la ventana revisada "
            "(móvil desde la hora del reporte). "
            "Motivo indica desde cuándo está en cero (última hora con consumo > 0)."
        )
        
        # Ordenar por empresa y luego por nombre del punto
        puntos_en_cero_ordenados = sorted(
            puntos_en_cero,
            key=lambda x: (x["companyName"], x["nodeName"])
        )
        
        rows = [("Nodo ID", "Nombre del Punto", "Empresa", "Motivo")]
        for punto in puntos_en_cero_ordenados:
            rows.append((
                punto["nodeId"],
                punto["nodeName"],
                punto["companyName"],
                punto.get("motivo") or "En cero (rango no determinado)",
            ))
        
        # Crear tabla — mitad izquierda / mitad Motivo
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = 'Light Grid Accent 1'
        anchos = [0.95, 1.35, 0.95, 3.25]
        
        # Encabezados
        header_cells = table.rows[0].cells
        for i, header in enumerate(rows[0]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            # Fondo azul para encabezados
            try:
                shading_xml = '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="4472C4"/>'
                shading = parse_xml(shading_xml)
                tc_pr = header_cells[i]._element.get_or_add_tcPr()
                if tc_pr.find(qn("w:shd")) is None:
                    tc_pr.append(shading)
            except Exception:
                pass
        
        # Datos
        for row_idx, row_data in enumerate(rows[1:], start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8 if col_idx == 3 else 9)
                # Alternar color de fondo para mejor legibilidad
                if row_idx % 2 == 0:
                    try:
                        shading_xml = '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="F2F2F2"/>'
                        shading = parse_xml(shading_xml)
                        tc_pr = cell._element.get_or_add_tcPr()
                        if tc_pr.find(qn("w:shd")) is None:
                            tc_pr.append(shading)
                    except Exception:
                        pass

        _aplicar_anchos_tabla(table, anchos)
        
        doc.add_paragraph("")  # Espacio
    elif not solo_sin_datos:
        doc.add_heading("PUNTOS MARCANDO CERO", 1)
        doc.add_paragraph("No se encontraron puntos marcando cero en el sistema.")
        doc.add_paragraph("")  # Espacio

    # Sección: Puntos sin datos
    if puntos_sin_datos:
        doc.add_heading("PUNTOS SIN DATOS DISPONIBLES", 1)
        doc.add_paragraph(
            "Los siguientes puntos no tienen datos en la ventana del reporte y llevan "
            f"≥{HORAS_OMITIR_SIN_DATOS:.0f} h sin medida (desconexiones de menos de 1 día no se notifican). "
            "Motivo: equipo conectado (o que volverá a conectar) con hueco desde la última medida; "
            "al reconectar no recupera el histórico de ese periodo."
        )
        
        # Ordenar por empresa y luego por nombre del punto
        puntos_sin_datos_ordenados = sorted(
            puntos_sin_datos,
            key=lambda x: (x["companyName"], x["nodeName"])
        )
        
        rows = [("Nodo ID", "Nombre del Punto", "Empresa", "Motivo")]
        for punto in puntos_sin_datos_ordenados:
            rows.append((
                punto["nodeId"],
                punto["nodeName"],
                punto["companyName"],
                punto.get("motivo") or "Sin datos (motivo no determinado)",
            ))
        
        # Crear tabla — mitad izquierda (ID/Nombre/Empresa) / mitad Motivo
        table = doc.add_table(rows=len(rows), cols=4)
        table.style = 'Light Grid Accent 1'
        # ~6.5" útiles: 50% izquierda (3.25") + 50% Motivo (3.25")
        anchos = [0.95, 1.35, 0.95, 3.25]
        
        # Encabezados
        header_cells = table.rows[0].cells
        for i, header in enumerate(rows[0]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            # Fondo naranja para encabezados
            try:
                shading_xml = '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="E67E22"/>'
                shading = parse_xml(shading_xml)
                tc_pr = header_cells[i]._element.get_or_add_tcPr()
                if tc_pr.find(qn("w:shd")) is None:
                    tc_pr.append(shading)
            except Exception:
                pass
        
        # Datos
        for row_idx, row_data in enumerate(rows[1:], start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8 if col_idx == 3 else 9)
                # Alternar color de fondo para mejor legibilidad
                if row_idx % 2 == 0:
                    try:
                        shading_xml = '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="FDF2E9"/>'
                        shading = parse_xml(shading_xml)
                        tc_pr = cell._element.get_or_add_tcPr()
                        if tc_pr.find(qn("w:shd")) is None:
                            tc_pr.append(shading)
                    except Exception:
                        pass

        _aplicar_anchos_tabla(table, anchos)
        
        doc.add_paragraph("")  # Espacio
    else:
        doc.add_heading("PUNTOS SIN DATOS DISPONIBLES", 1)
        doc.add_paragraph(
            "No hay puntos sin datos para notificar "
            f"(cortes con última medida <{HORAS_OMITIR_SIN_DATOS:.0f} h se omiten)."
            if omitidos_corte_reciente
            else "Todos los puntos tienen datos disponibles."
        )
        doc.add_paragraph("")  # Espacio

    # Nota al pie
    if dias_revisar <= 1:
        ventana_txt = "las últimas 24 horas (ventana móvil; todas las horas con registro)"
    else:
        ventana_txt = (
            f"las últimas {dias_revisar * 24} horas "
            f"({dias_revisar} días móviles; todas las horas con registro)"
        )
    if solo_sin_datos:
        texto_nota = (
            f"Nota: Este reporte verifica {ventana_txt}. "
            "Un punto se considera «sin datos» si no hay respuesta de la API o no hay registros en esa ventana. "
            f"Desconexiones <{HORAS_OMITIR_SIN_DATOS:.0f} h no se notifican. "
            "Los listados se marcan como «Equipo conectado — sin datos desde… (no recupera histórico)»."
        )
    else:
        texto_nota = (
            f"Nota: Este reporte verifica {ventana_txt}. "
            "Un punto se considera 'en cero' solo si TODOS los registros de esa ventana son cero "
            "(si hubo consumo ayer y hoy de madrugada va en 0, NO se notifica). "
            "Motivo en cero: desde cuándo (última hora con consumo > 0 hacia atrás, hasta 30 días). "
            "Un punto se considera 'sin datos' si no hay registros en esa ventana. "
            f"Filtro: desconexiones con última medida <{HORAS_OMITIR_SIN_DATOS:.0f} h no se listan. "
            "Motivo sin datos: equipo conectado (o que reconectará) con hueco desde la última medida; "
            "al reconectar no recupera el histórico de ese periodo."
        )
    nota = doc.add_paragraph(texto_nota)
    nota.runs[0].font.italic = True
    nota.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Guardar documento (fecha en nombre = día referencia si aplica + hora actual)
    _ref = fecha_generacion_utc if fecha_generacion_utc is not None else datetime.now(timezone.utc)
    # Usar la misma fecha/hora de generación para el nombre del archivo (evita depender del clock local)
    timestamp = _ref.strftime("%Y%m%d_%H%M%S")
    prefix = "Reporte_Puntos_Sin_Datos" if solo_sin_datos else "Reporte_Puntos_En_Cero"
    filename = f"{prefix}_{timestamp}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))
    
    return output_path


def main(
    output_dir: Optional[Path] = None,
    fecha_generacion_utc: Optional[datetime] = None,
    solo_sin_datos: bool = False,
    dias_revisar: int = 3,
) -> None:
    """Función principal."""
    if output_dir is None:
        output_dir = _DEFAULT_REPORTE_CERO_DIR
    else:
        output_dir = Path(output_dir).resolve()

    dias_revisar = max(1, int(dias_revisar))

    print("=" * 70)
    print(
        "REPORTE DE PUNTOS SIN DATOS DISPONIBLES"
        if solo_sin_datos
        else "REPORTE DE PUNTOS EN CERO"
    )
    if dias_revisar <= 1:
        print("Ventana en cero/sin datos: últimas 24 horas (móvil)")
    else:
        print(f"Ventana en cero/sin datos: últimas {dias_revisar * 24} horas ({dias_revisar} días móviles)")
    print("=" * 70)
    print()
    
    # Crear carpeta de salida (ruta fija junto al script salvo --output-dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"Carpeta de salida: {output_dir}")
    print()
    
    # 1. Obtener todos los nodos
    todos_nodos = obtener_todos_los_nodos()
    
    if not todos_nodos:
        print("[ERROR] No se encontraron nodos en el sistema.")
        return
    
    print()
    print("=" * 70)
    print("VERIFICANDO CONSUMO EN CERO")
    print("=" * 70)
    print(f"Verificando {len(todos_nodos)} puntos (hasta {MAX_WORKERS_CERO} en paralelo; ajustar WES_REPORTE_CERO_WORKERS)...")
    print()

    def _verificar_nodo(nodo: Dict) -> Tuple[Dict, bool, Optional[str], bool]:
        """Devuelve (nodo, esta_en_cero, error, ok_verificacion). ok_verificacion=False si hubo excepción."""
        try:
            node_id = nodo["nodeId"]
            esta_en_cero, error = verificar_consumo_cero(node_id, dias_revisar=dias_revisar)
            return nodo, esta_en_cero, error, True
        except Exception as e:
            return nodo, False, f"Error al verificar: {e}", False

    with ThreadPoolExecutor(max_workers=MAX_WORKERS_CERO) as ex:
        resultados = list(ex.map(_verificar_nodo, todos_nodos))

    puntos_en_cero = []
    puntos_sin_datos = []
    fallos_verificacion = 0

    for i, (nodo, esta_en_cero, error, ok_verificacion) in enumerate(resultados, 1):
        node_id = nodo["nodeId"]
        node_name = nodo["nodeName"]

        print(f"[{i}/{len(todos_nodos)}] {node_name} ({node_id})...", end=" ", flush=True)

        if not ok_verificacion:
            print(f"[ERROR] {error}")
            fallos_verificacion += 1
        elif esta_en_cero:
            print("[OK] EN CERO")
            puntos_en_cero.append(nodo)
        elif error and "Sin datos" in (error or ""):
            print(f"[SIN DATOS] {error}")
            puntos_sin_datos.append(nodo)
        else:
            print("[OK] Con consumo")

        if i % 10 == 0:
            print(
                f"\nProgreso (listado): {i}/{len(todos_nodos)} | "
                f"en cero: {len(puntos_en_cero)} | sin datos: {len(puntos_sin_datos)}\n"
            )

    if puntos_en_cero:
        print()
        print("=" * 70)
        print("DIAGNÓSTICO EN CERO (rango / motivo)")
        print("=" * 70)
        print(
            f"Buscando desde cuándo están en cero (hasta {DIAS_BUSCAR_ULTIMA_MEDIDA} días) "
            f"en {len(puntos_en_cero)} punto(s)..."
        )
        puntos_en_cero = enriquecer_puntos_en_cero(
            puntos_en_cero,
            dias_ventana_reporte=dias_revisar,
        )
        for p in sorted(puntos_en_cero, key=lambda x: (x["companyName"], x["nodeName"])):
            print(f"  {p['nodeId']} | {p['nodeName'][:28]:<28} | {p.get('motivo', '')}")

    omitidos_corte_reciente = 0
    if puntos_sin_datos:
        print()
        print("=" * 70)
        print("DIAGNÓSTICO SIN DATOS (última medida / motivo)")
        print("=" * 70)
        print(
            f"Buscando última medida (hasta {DIAS_BUSCAR_ULTIMA_MEDIDA} días) "
            f"en {len(puntos_sin_datos)} punto(s)..."
        )
        print(
            f"Filtro: no notificar si última medida < {HORAS_OMITIR_SIN_DATOS:.0f} h "
            "(desconexión / corte reciente)."
        )
        puntos_sin_datos, omitidos_recientes = enriquecer_puntos_sin_datos(
            puntos_sin_datos,
            dias_ventana_reporte=dias_revisar,
        )
        omitidos_corte_reciente = len(omitidos_recientes)
        if omitidos_recientes:
            print(
                f"\nOmitidos del informe ({len(omitidos_recientes)}; "
                f"última medida < {HORAS_OMITIR_SIN_DATOS:.0f} h):"
            )
            for p in sorted(omitidos_recientes, key=lambda x: (x["companyName"], x["nodeName"])):
                horas = p.get("haceHoras")
                horas_txt = f"~{horas} h" if horas is not None else "?"
                print(
                    f"  [OMITIDO] {p['nodeId']} | {p['nodeName'][:32]:<32} | "
                    f"última hace {horas_txt} ({p.get('ultimaMedidaChile', '—')} Chile)"
                )
        if puntos_sin_datos:
            print(f"\nA notificar como sin datos: {len(puntos_sin_datos)}")
            for p in sorted(puntos_sin_datos, key=lambda x: (x["companyName"], x["nodeName"])):
                print(f"  {p['nodeId']} | {p['nodeName'][:32]:<32} | {p.get('motivo', '')}")
        else:
            print("\nA notificar como sin datos: 0 (todos los cortes eran <1 día).")
    
    print()
    print("=" * 70)
    print("GENERANDO REPORTE")
    print("=" * 70)
    
    # 3. Construir análisis de alertas del día anterior
    ref = (
        fecha_generacion_utc
        if fecha_generacion_utc is not None
        else (FECHA_REFERENCIA_UTC if FECHA_REFERENCIA_UTC is not None else datetime.now(timezone.utc))
    )
    fecha_alertas = ref - timedelta(days=1)
    alertas_resumen = construir_resumen_alertas(todos_nodos, fecha_alertas)

    # 4. Generar reporte Word
    try:
        reporte_path = crear_reporte_word(
            puntos_en_cero,
            puntos_sin_datos,
            len(todos_nodos),
            output_dir,
            alertas_resumen=alertas_resumen,
            alertas_fecha=fecha_alertas.strftime("%d-%m-%Y"),
            fecha_generacion_utc=ref,
            solo_sin_datos=solo_sin_datos,
            dias_revisar=dias_revisar,
            omitidos_corte_reciente=omitidos_corte_reciente,
        )
        print(f"[OK] Reporte generado exitosamente:")
        print(f"  {reporte_path}")
        print()
        print("=" * 70)
        print("RESUMEN FINAL")
        print("=" * 70)
        puntos_con_datos = len(todos_nodos) - len(puntos_sin_datos)
        print(f"Total de puntos analizados: {len(todos_nodos)}")
        if solo_sin_datos:
            print(f"Puntos sin datos disponibles: {len(puntos_sin_datos)}")
            print(f"Puntos con al menos un registro en el periodo: {puntos_con_datos}")
        else:
            print(f"Puntos con datos disponibles: {puntos_con_datos}")
            print(f"Puntos marcando cero: {len(puntos_en_cero)}")
            print(f"Puntos sin datos disponibles: {len(puntos_sin_datos)}")
            print(f"Porcentaje en cero: {(len(puntos_en_cero) / puntos_con_datos * 100) if puntos_con_datos > 0 else 0:.2f}%")
        print(f"Porcentaje sin datos: {(len(puntos_sin_datos) / len(todos_nodos) * 100) if todos_nodos else 0:.2f}%")
        if omitidos_corte_reciente:
            print(
                f"Cortes <{HORAS_OMITIR_SIN_DATOS:.0f} h omitidos (no notificados): "
                f"{omitidos_corte_reciente}"
            )
        if fallos_verificacion:
            print(f"Advertencia: {fallos_verificacion} punto(s) con error al verificar (revisar mensajes [ERROR] arriba).")
        print()
        print(f"Reporte guardado en: {reporte_path.absolute()}")
    except Exception as e:
        print(f"[ERROR] Error al generar reporte: {e}")
        import traceback
        traceback.print_exc()


def regenerar_reporte_sin_api(output_dir: Optional[Path] = None) -> None:
    """
    Regenera el último reporte de puntos en cero sin conectarse a la API.
    Lee el último reporte generado y lo regenera con las tablas ordenadas por empresa.
    """
    from docx import Document as DocxDocument

    if output_dir is None:
        output_dir = _DEFAULT_REPORTE_CERO_DIR
    else:
        output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Buscar el último reporte generado
    reportes = list(output_dir.glob("Reporte_Puntos_En_Cero_*.docx"))
    if not reportes:
        print("[ERROR] No se encontró ningún reporte previo.")
        return
    
    # Ordenar por fecha de modificación (más reciente primero)
    reportes.sort(key=lambda x: x.stat().st_mtime, reverse=True)
    ultimo_reporte = reportes[0]
    
    print("=" * 70)
    print("REGENERANDO REPORTE SIN CONECTARSE A LA API")
    print("=" * 70)
    print(f"Leyendo reporte: {ultimo_reporte.name}")
    print()
    
    try:
        # Leer el documento Word existente
        doc = DocxDocument(str(ultimo_reporte))
        
        # Extraer datos de las tablas
        puntos_en_cero = []
        puntos_sin_datos = []
        total_puntos = 0
        
        # Buscar el resumen ejecutivo para obtener total_puntos
        for para in doc.paragraphs:
            text = para.text
            if "Total de puntos analizados:" in text:
                try:
                    total_str = text.split("Total de puntos analizados:")[1].split("\n")[0].strip()
                    total_puntos = int(total_str)
                except:
                    pass
        
        # Buscar las tablas en el documento
        tabla_indice = 0
        for table in doc.tables:
            if len(table.rows) > 0:
                header = table.rows[0].cells[0].text.strip()
                if "Nodo ID" in header:
                    # Primera tabla = puntos en cero, segunda tabla = puntos sin datos
                    for row in table.rows[1:]:  # Saltar encabezado
                        if len(row.cells) >= 4:
                            node_id = row.cells[0].text.strip()
                            node_name = row.cells[1].text.strip()
                            company_name = row.cells[2].text.strip()
                            company_id = row.cells[3].text.strip()
                            
                            if node_id and node_name and company_name:
                                punto = {
                                    "nodeId": node_id,
                                    "nodeName": node_name,
                                    "companyName": company_name,
                                    "companyId": company_id
                                }
                                
                                # Primera tabla encontrada = puntos en cero
                                if tabla_indice == 0:
                                    puntos_en_cero.append(punto)
                                else:
                                    puntos_sin_datos.append(punto)
                    tabla_indice += 1
        
        if not puntos_en_cero and not puntos_sin_datos:
            print("[ADVERTENCIA] No se pudieron extraer los datos del reporte anterior.")
            return
        
        print(f"[OK] Datos extraídos:")
        print(f"  - Puntos en cero: {len(puntos_en_cero)}")
        print(f"  - Puntos sin datos: {len(puntos_sin_datos)}")
        if total_puntos > 0:
            print(f"  - Total de puntos: {total_puntos}")
        print()
        
        # Regenerar el reporte con las tablas ordenadas
        print("Generando nuevo reporte con tablas ordenadas por empresa...")
        nuevo_reporte = crear_reporte_word(
            puntos_en_cero,
            puntos_sin_datos,
            total_puntos if total_puntos > 0 else len(puntos_en_cero) + len(puntos_sin_datos) + 100,
            output_dir
        )
        
        print()
        print("=" * 70)
        print("REPORTE REGENERADO EXITOSAMENTE")
        print("=" * 70)
        print(f"Reporte anterior: {ultimo_reporte.name}")
        print(f"Reporte nuevo: {nuevo_reporte.name}")
        print(f"Ubicación: {nuevo_reporte.absolute()}")
        
    except Exception as e:
        print(f"[ERROR] Error al regenerar reporte: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Reporte de puntos en cero")
    parser.add_argument("--regenerar", action="store_true", help="Regenerar desde último Word sin API")
    parser.add_argument(
        "--fecha",
        metavar="DD/MM/YYYY",
        help="Fecha referencia para verificar últimos días (UTC); ej. 06/04/2026",
    )
    parser.add_argument(
        "--fecha-generacion",
        metavar="DD/MM/YYYY HH:MM",
        help="Fecha/hora (UTC) para timbrar el reporte y el nombre del archivo; ej. 07/04/2026 09:00",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        metavar="DIR",
        help=f"Carpeta de salida (por defecto: {_DEFAULT_REPORTE_CERO_DIR})",
    )
    parser.add_argument(
        "--solo-sin-datos",
        action="store_true",
        help="Generar solo el informe de puntos sin datos (sin sección «en cero» ni tabla asociada).",
    )
    parser.add_argument(
        "--dias",
        type=int,
        default=3,
        metavar="N",
        help="Días a revisar hacia atrás desde la fecha de referencia (default: 3). Use 1 para solo hoy.",
    )
    args = parser.parse_args()
    out = args.output_dir.resolve() if args.output_dir else None
    if args.regenerar:
        regenerar_reporte_sin_api(out)
    else:
        forced_gen_utc: Optional[datetime] = None
        if args.fecha_generacion:
            forced_gen_utc = datetime.strptime(args.fecha_generacion.strip(), "%d/%m/%Y %H:%M").replace(
                tzinfo=timezone.utc
            )

        if args.fecha:
            globals()["FECHA_REFERENCIA_UTC"] = datetime.strptime(
                args.fecha.strip(), "%d/%m/%Y"
            ).replace(tzinfo=timezone.utc)
        elif forced_gen_utc is not None:
            # Si no se entrega --fecha, usar la generación como referencia para el "hoy" del análisis
            globals()["FECHA_REFERENCIA_UTC"] = forced_gen_utc

        main(
            out,
            fecha_generacion_utc=forced_gen_utc,
            solo_sin_datos=args.solo_sin_datos,
            dias_revisar=args.dias,
        )

