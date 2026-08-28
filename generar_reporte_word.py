"""Script para generar reportes Word de consumo y consumos nocturnos usando la API WES ACL Node."""

from __future__ import annotations

import argparse
import json
import os
import smtplib
import subprocess
import sys
import threading
import tempfile
import time
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
try:
    from zoneinfo import ZoneInfo
except ImportError:  # Python < 3.9
    from backports.zoneinfo import ZoneInfo  # type: ignore
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple, Union

import matplotlib
import numpy as np
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Usar backend no interactivo para generar gráficos en servidores/headless
matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
import matplotlib.dates as mdates  # noqa: E402

from exclusiones_reportes import (
    filter_node_ids,
    is_node_excluded,
)


_DEFAULT_ACL_NODE_BASE = "http://104.248.53.141:7003/wes/api/acl-node/v1"
ENTITY_BASE_URL = "http://104.248.53.141:7001/wes/api/acl-entities/v1"


def acl_node_base_url() -> str:
    """Base URL del API acl-node (nodos, medidas, alertas). Sobreescribe con env ``WES_API_BASE_URL``."""
    return os.environ.get("WES_API_BASE_URL", _DEFAULT_ACL_NODE_BASE).rstrip("/")


def __getattr__(name: str):
    if name == "BASE_URL":
        return acl_node_base_url()
    raise AttributeError(f"module {__name__!r} has no attribute {name!r}")


# CSV horario: lo **genera el servidor** WES al atender
# ``GET .../nodes/{id}/dates.measures.csv?start=ddMMyyyy&end=ddMMyyyy`` (ver ``acl_node_base_url()``).
# Este repositorio solo **consume** ese endpoint (igual que la app); no hay aquí código que arme el CSV.
# Otra serie que suelen usar los clientes es ``GET /nodes/measures/dates`` (JSON, campo ``measures``).

# Zona horaria de los reportes (Chile): la API entrega marcas en UTC; hay que convertir
# para ventanas nocturnas y para medidas horarias.
# En Windows: pip install tzdata (también en requirements.txt).
try:
    CHILE_TZ = ZoneInfo("America/Santiago")
except Exception:
    try:
        import tzdata  # noqa: F401 — datos IANA para zoneinfo en Windows

        CHILE_TZ = ZoneInfo("America/Santiago")
    except Exception:
        CHILE_TZ = timezone(timedelta(hours=-3))


def _dt_to_chile(dt: datetime) -> datetime:
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(CHILE_TZ)


def _hour_chile_from_iso_time(time_str: str) -> Optional[int]:
    """Hora local Chile desde una marca ISO del CSV/API (típicamente en UTC)."""
    try:
        if "T" not in time_str.strip():
            return None
        dt = datetime.fromisoformat(time_str.strip().replace("Z", "+00:00"))
        return _dt_to_chile(dt).hour
    except (ValueError, AttributeError, OSError):
        return None


def format_number_chilean(value: float, decimals: int = 2) -> str:
    """
    Formatea un número al estilo chileno: punto para miles, coma para decimales.
    Ejemplo: 1234.56 -> "1.234,56"
    """
    if value is None:
        return "0"
    
    # Formatear con el número de decimales especificado usando formato estándar
    formatted = f"{value:.{decimals}f}"
    
    # Separar parte entera y decimal (el formato estándar usa punto como separador decimal)
    parts = formatted.split(".")
    integer_part = parts[0]
    decimal_part = parts[1] if len(parts) > 1 else ""
    
    # Manejar signo negativo
    is_negative = integer_part.startswith("-")
    if is_negative:
        integer_part = integer_part[1:]
    
    # Agregar puntos como separadores de miles
    if len(integer_part) > 3:
        # Agregar puntos cada 3 dígitos desde la derecha
        formatted_integer = ""
        for i, digit in enumerate(reversed(integer_part)):
            if i > 0 and i % 3 == 0:
                formatted_integer = "." + formatted_integer
            formatted_integer = digit + formatted_integer
        integer_part = formatted_integer
    
    # Combinar con coma como separador decimal
    if decimal_part:
        # Eliminar ceros finales innecesarios
        decimal_part = decimal_part.rstrip("0")
        if decimal_part:
            result = f"{integer_part},{decimal_part}"
        else:
            result = integer_part
    else:
        result = integer_part
    
    # Agregar signo negativo si es necesario
    if is_negative:
        return f"-{result}"
    return result


def format_currency_chilean(value: float) -> str:
    """
    Formatea un valor monetario al estilo chileno.
    Ejemplo: 1234567.89 -> "$1.234.567,89"
    """
    return f"${format_number_chilean(value, 0)}"


def add_logo_to_header(doc: Document, logo_path: Optional[Path] = None) -> None:
    """
    Agrega un logo al encabezado de cada página del documento en el lado derecho superior.
    Si no se proporciona logo_path, busca 'logo.png' en el directorio actual.
    El logo se posiciona para no traslaparse con el contenido del reporte.
    """
    # Buscar el logo si no se proporciona
    if logo_path is None:
        current_dir = Path(__file__).parent
        possible_logos = [
            # Buscar "logo wes" con diferentes extensiones y variaciones
            current_dir / "logo wes.bmp",
            current_dir / "logo wes.png",
            current_dir / "logo wes.jpg",
            current_dir / "logo wes.jpeg",
            current_dir / "logo wes.BMP",
            current_dir / "logo wes.PNG",
            current_dir / "logo wes.JPG",
            current_dir / "logo wes.JPEG",
            current_dir / "logo_wes.bmp",
            current_dir / "logo_wes.png",
            current_dir / "logo_wes.jpg",
            current_dir / "logo_wes.jpeg",
            current_dir / "Logo WES.bmp",
            current_dir / "Logo WES.png",
            current_dir / "Logo WES.jpg",
            current_dir / "Logo WES.jpeg",
            # Buscar "logo" simple también
            current_dir / "logo.bmp",
            current_dir / "logo.png",
            current_dir / "logo.jpg",
            current_dir / "logo.jpeg",
            current_dir / "logo.BMP",
            current_dir / "logo.PNG",
            current_dir / "logo.JPG",
            current_dir / "logo.JPEG",
            Path("logo wes.bmp"),
            Path("logo wes.png"),
            Path("logo wes.jpg"),
            Path("logo wes.jpeg"),
            Path("logo.bmp"),
            Path("logo.png"),
            Path("logo.jpg"),
            Path("logo.jpeg"),
        ]
        for possible_logo in possible_logos:
            if possible_logo.exists():
                logo_path = possible_logo
                print(f"[INFO] Logo encontrado: {logo_path}")
                break
    
    if logo_path is None or not logo_path.exists():
        # Si no se encuentra el logo, mostrar advertencia pero continuar
        print(f"[ADVERTENCIA] No se encontró el archivo del logo. Buscado en:")
        current_dir = Path(__file__).parent
        print(f"  - {current_dir / 'logo wes.bmp'}")
        print(f"  - {current_dir / 'logo wes.png'}")
        print(f"  - {current_dir / 'logo wes.jpg'}")
        print(f"  - {current_dir / 'logo wes.jpeg'}")
        print(f"  - {current_dir / 'logo.bmp'}")
        print(f"  - {current_dir / 'logo.png'}")
        print(f"  - {current_dir / 'logo.jpg'}")
        print(f"  - {current_dir / 'logo.jpeg'}")
        print(f"  Coloca el archivo del logo en el directorio del script para que se agregue automáticamente.")
        return
    
    try:
        # Obtener la sección del documento
        section = doc.sections[0]
        header = section.header
        
        # Ajustar márgenes del encabezado para evitar traslape
        # Aumentar el margen superior del encabezado para dar espacio al logo
        section.header_distance = Inches(0.5)
        
        # Limpiar el encabezado existente
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        # Crear un párrafo para el logo
        logo_paragraph = header.paragraphs[0]
        logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Agregar el logo con tamaño aumentado
        # Tamaño más grande (1.5 pulgadas) para mejor visibilidad
        run = logo_paragraph.add_run()
        run.add_picture(str(logo_path), width=Inches(1.5))  # 1.5 pulgadas de ancho
        
        # Ajustar el espaciado del párrafo para que el logo esté bien posicionado
        logo_paragraph.paragraph_format.space_after = Pt(0)
        logo_paragraph.paragraph_format.space_before = Pt(0)
        
        print(f"[OK] Logo agregado al encabezado: {logo_path.name}")
    except Exception as e:
        print(f"[ERROR] No se pudo agregar el logo al encabezado: {e}")
        import traceback
        traceback.print_exc()

# Diccionario de mapeo de IDs de nodos a sus nombres
NODE_NAMES = {
    "000000-02": "Nodo isla Parcela",
    "000001-03": "Arrieta",
    "000001-05": "Acad. de Guerra",
    "000001-06": "Jef. Adm y Log (J.A.L)",
    "000001-08": "Regto. Tacna",
    "000001-09": "Regto. Limache",
    "000001-11": "Regto. Calama",
    "000001-12": "Regto. Copiapo",
    "000001-13": "Regto. Buin",
    "000001-14": "Regto. Maipo",
    "000001-15": "Regto. Chacabuco",
    "000002-01": "Lo Valledor - P1",
    "000002-02": "Lo valledor - Pozo",
    "000002-03": "Lo Valledor - Barrio Norte",
    "000004-01": "CDP Puente Alto",
    "000004-02": "CPF San Miguel",
    "000004-03": "CDP Colina 2 Red Sur",
    "000004-04": "CDP Colina 1",
    "000004-05": "Esforpen Simulación",
    "000004-06": "Esforpen Administración",
    "000004-07": "Colina 2 Red de Impulsión",
    "000004-08": "CDP STGO SUR Óvalo",
    "000004-09": "CDP STGO SUR Módulos",
    "000004-10": "CDP STGO SUR Ovalo Matriz Principal",
    "000004-11": "CDP STGO SUR Modulo Matriz Principal",
    "000004-12": "Esc Of. Genchi Carmen",
    "000004-13": "Esc Of. Genchi Artemio",
    "000004-14": "CPF San Miguel sala imp.",
    "000005-01": "M.O.P",
    "000006-01": "Liceo Lastarria",
    "000006-02": "Carmela Carvajal",
    "000006-03": "Arturo Alessandri Palma",
    "000006-04": "Liceo 7 Luisa Saavedra",
    "000006-05": "Liceo Juan Pablo Duarte",
    "000007-01": "Estanque B",
    "000007-02": "Teatro",
    "000007-03": "Nido High School",
    "000007-04": "Elementary",
    "000007-05": "Nido Piscina",
    "000007-06": "Nido Pozo Profundo",
    "000007-07": "Estanque C",
    "000007-08": "Nido Cancha",
    "000007-09": "Control Nido de Aguilas",
    "000008-01": "Lic. Antonio Hermida F",
    "000008-02": "Eduardo de la Barra",
    "000008-03": "Carlos Fernandez P.",
    "000008-04": "Tobalaba",
    "000008-05": "Santa Maria",
    "000008-06": "Luis Arrieta Caña",
    "000008-07": "Erasmo Escala",
    "000008-08": "Alicura",
    "000008-09": "Juan Bautista Pasten",
    "000008-10": "Matilde Huici Navas",
    "000008-11": "CE Valle Hermoso",
    "000008-12": "Unión Nacional Árabe",
    "000008-13": "Likankura",
    "000008-14": "Juan Pablo II",
    "000009-00": "Copec Costanera",
    "000009-01": "Oficina Admin.",
    "000009-02": "Estanque Reutilización",
    "000009-03": "Lavado Automático Norte",
    "000009-04": "Lavado Automático Sur",
    "000009-05": "Copec Riego",
    "000009-06": "Copec Matriz Principal",
    "000009-07": "Copec Costanera - ¿?",
    "000009-08": "Copec Pronto Baños",
    "000009-09": "Copec Lavado Auto servicio Norte",
    "000009-10": "Copec Lavado Auto servicio Sur",
    "000009-11": "Copec Pronto Tienda",
    "000010-01": "Escuela Andes del SUR",
    "000010-02": "Esc. Villa Independencia",
    "000010-03": "Esc. Padre Hurtado",
    "000010-04": "Colegio Maipo",
    "000010-05": "Esc. Luis Matte Larraín",
    "000010-06": "Esc. Gabriela",
    "000010-07": "Esc. Juan Mackenna",
    "000010-08": "Liceo Chiloé",
    "000010-09": "Esc. Los Andes",
    "000010-10": "Consolidada",
    "000010-11": "Esc. Nonato Coo",
    "000011-01": "Casa Juan Lopez",
    "000012-01": "Lo Boza Lavado de Vehículos",
    "000012-02": "Lo Boza Pozo",
    "000012-03": "Lo Boza Reutilización",
    "000012-04": "Lo Boza Edificio Principal Casino",
    "000012-05": "Lo Boza Matriz Principal",
    "000012-06": "Quilicura Matriz Principal",
    "000012-07": "Quilicura Dercomaq",
    "000012-08": "Quilicura Lav. de Maquinas",
    "000012-09": "Quilicura Casino",
    "000012-10": "Quilicura Proderco",
    "000012-11": "Quilicura Camarines",
    "000012-12": "Quilicura Edificio JCB",
    "000012-13": "Open Plaza Lavado de Vehiculos",
    "000012-14": "Open Plaza Matriz Principal",
    "000013-01": "Plaza Boulevard Pajaros Sur",
    "000014-01": "Agua de Rechazo Ablandador",
    "000015-01": "Molymet Red Riego",
    "000015-02": "Molymet Casa de Cambio",
    "000015-03": "Molymet Casa de Cambio",
    "000016-01": "SCL Rebeca Matte Bello",
    "000017-01": "Rebeca Matte Bello",
    "000017-02": "Juana Atala de Hirmas",
    "000017-03": "José Luis Araneda",
    "000017-04": "Esc. Lo Velásquez",
    "000017-05": "Gimnasio",
    "000017-06": "Piscina Municipal",
    "000017-07": "Cumbre de cóndores pte.",
    "000017-08": "Colegio ICCO Renca",
    "000018-01": "San Nicolas",
    "000018-02": "Ureta Cox",
    "000018-03": "Estanque Cisterna",
    "000018-04": "Casino",
    "000019-01": "Oficina WES",
    "000020-01": "Deposito",
    "000020-02": "Modulo D",
    "000020-03": "Módulo ABC",
    "000020-04": "Módulo E",
    "000020-05": "Intermodal-San Antonio",
    "000021-01": "Club House CDUC",
    "000021-02": "Edificio Deportivo",
    "000021-03": "Raimundo Tupper",
    "000021-04": "Equitación",
    "000021-05": "Calle de Servicio",
    "000021-07": "Canchas de Tenis",
    "000021-08": "Rugby CDUC",
    "000022-00": "Alexander Fleming",
    "000022-01": "Juan Pablo II",
    "000023-01": "San Nicolas",
    "000023-02": "Ureta Cox",
    "000023-03": "Estanque Cisterna",
    "000023-04": "Casino",
    "000024-01": "Eugenio María De Hostos",
    "000025-01": "Estanque Norte Locales Mall",
    "000025-02": "Abastecimiento Sur Terminal",
    "000025-03": "Poniente 7",
    "000025-04": "Baños Públicos",
    "000025-05": "Locales de Comida",
    "000025-06": "KFC",
    "000025-07": "PIZZA HUT",
    "000025-08": "Placa Bancaria",
    "000025-09": "Impulsión Falabella",
    "000025-10": "Impulsión Ripley",
    "000025-11": "Matriz principal 1°piso",
    "000025-12": "Anillo Plaza",
    "000025-13": "Matriz Principal",
    "000025-14": "Red de Incendio",
    "000025-15": "Matriz Principal",
    "000025-16": "Baños",
    "000025-17": "San Ignacio 300",
    "000025-18": "San Ignacio 500",
    "000025-19": "Sala de Bomba Estanque Sur",
    "000025-20": "Impulsión Anden 3-4 Matriz Principal",
    "000025-21": "Impulsión Anden 3-4 Locales Gast.",
    "000025-22": "Sala de Bomba Sandia Antigua",
    "000025-23": "Llenado Pileta",
    "000025-24": "Llenado Pileta Cascada",
    "000025-25": "Baño N°5 Damas",
    "000025-26": "Baño N°6 Varones",
    "000025-27": "Distrito de lujo DL",
    "000025-28": "Sala de Bomba Sandia Nueva",
    "000025-29": "Impulsión Anden 3-4 Restaurante",
    "000025-35": "PAK BAZAR GOURMET",
    "000025-36": "PAK DL KENNEDY",
    "000026-01": "UDD-Sala impulsión Honduras",
    "000026-02": "UDD-Edificio o Aula Magna",
    "000027-01": "Matriz ESVAL",
    "000027-02": "Estanque Inferior",
    "000027-03": "Etapa N°5",
    "000027-04": "Etapa N°1 al 4",
    "000027-05": "Riego",
    "000027-06": "Etapa N°1",
    "000027-07": "Etapa N°2",
    "000027-08": "Etapa N°3",
    "000027-09": "Riego Llenado de Estanque ESVAL",
    "000028-01": "Liceo Alto Cordillera",
    "000029-01": "Llenado de Estanques",
    "000029-02": "Torre A",
    "000029-03": "Torre B1",
    "000029-04": "Torre B2",
    "000029-05": "Torre C",
    "000029-06": "Central Térmica",
    "000030-01": "Matriz Av. Las Condes",
    "000030-02": "Matriz Chesterton",
    "000031-01": "Matriz Fitness",
    "000031-02": "Matriz Piscina",
    "999999-01": "Test wes - Gonza",
}

# Cache opcional con nombres de nodos obtenidos desde la API
NODE_NAMES_CACHE_FILE = Path("node_names_api_cache.json")

def _load_node_names_cache() -> dict:
    try:
        if NODE_NAMES_CACHE_FILE.exists():
            return json.loads(NODE_NAMES_CACHE_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}

NODE_NAMES_CACHE = _load_node_names_cache()
COMPANY_NAME_CACHE: dict[str, str] = {}
COMPANY_NODES_CACHE: dict[str, list[dict]] = {}

# Diccionario de mapeo de IDs de empresas a sus nombres
COMPANY_NAMES = {
    "000000": "Wes Spa",
    "000001": "Ejercito de Chile",
    "000002": "Lo valledor",
    "000004": "Gendarmeria de Chile",
    "000005": "Ministerio de O.P",
    "000006": "Providencia",
    "000007": "Nido de Aguilas",
    "000008": "CORMUP",
    "000009": "COPEC",
    "000010": "Corporación Puente Alto",
    "000011": "Sistemas Socios Wes",
    "000012": "Inchcape",
    "000013": "Lo Barnechea",
    "000014": "Tres Montes Lucchetti",
    "000016": "Renca",
    "000017": "Renca",
    "000018": "MADECO",
    "000019": "WESSPA",
    "000020": "AGUNSA",
    "000021": "CDUC",
    "000022": "Las Condes",
    "000023": "MADECO",
    "000024": "La Reina",
    "000025": "Parque Arauco",
    "000026": "UDD",
    "000027": "Fundo Zapallar",
    "000028": "La Florida",
    "000029": "BUPA",
    "000030": "Estadio Israelita Maccabi",
    "000031": "Club Providencia",
}


def get_company_name(company_id: str) -> str:
    """Obtiene el nombre de la empresa basado en su ID."""
    company_id = str(company_id).strip()
    # Primero intentar desde el diccionario estático
    if company_id in COMPANY_NAMES:
        return COMPANY_NAMES[company_id]

    # Luego intentar cache local en memoria
    if company_id in COMPANY_NAME_CACHE:
        return COMPANY_NAME_CACHE[company_id]

    # Si no está en el diccionario, consultar la API
    try:
        url = f"{ENTITY_BASE_URL}/companies/{company_id}"
        response = _requests_session().get(url, timeout=5)
        if response.status_code == 200:
            data = response.json()
            name = str(data.get('name', company_id))
            COMPANY_NAME_CACHE[company_id] = name
            return name
    except Exception:
        pass
    
    return company_id


def get_mall_name_for_parque_arauco(node_id: str, node_name: str) -> str:
    """
    Extrae el nombre del mall para nodos de Parque Arauco.
    Retorna el nombre del mall o cadena vacía si no se puede determinar.
    
    Usa el mapeo completo de nodos a malls basado en la configuración oficial.
    """
    # Mapeo completo de nodos a nombres de mall (basado en la configuración oficial)
    mall_mapping = {
        # Estación
        "000025-01": "Estación",
        "000025-19": "Estación",
        "000025-03": "Estación",
        "000025-05": "Estación",
        "000025-06": "Estación",
        "000025-04": "Estación",
        "000025-07": "Estación",
        # Maipú
        "000025-08": "Maipú",
        "000025-09": "Maipú",
        "000025-10": "Maipú",
        "000025-32": "Maipú",
        "000025-33": "Maipú",
        # El Bosque
        "000025-11": "El Bosque",
        "000025-12": "El Bosque",
        "000025-30": "El Bosque",
        # Quilicura
        "000025-13": "Quilicura",
        "000025-14": "Quilicura",
        "000025-34": "Quilicura",
        # Curauma
        "000025-15": "Curauma",
        "000025-16": "Curauma",
        # Buenaventura
        "000025-17": "Buenaventura",
        "000025-18": "Buenaventura",
        # Kennedy
        "000025-20": "Kennedy",
        "000025-21": "Kennedy",
        "000025-22": "Kennedy",
        "000025-23": "Kennedy",
        "000025-24": "Kennedy",
        "000025-25": "Kennedy",
        "000025-26": "Kennedy",
        "000025-27": "Kennedy",
        "000025-28": "Kennedy",
        "000025-29": "Kennedy",
        "000025-35": "Kennedy",
        "000025-36": "Kennedy",
    }
    
    # Si hay un mapeo directo, usarlo
    if node_id in mall_mapping:
        return mall_mapping[node_id]
    
    # Intentar extraer del nombre del nodo como respaldo
    node_name_lower = node_name.lower()
    
    # Buscar palabras clave que indiquen el mall
    if "maipú" in node_name_lower or "maipu" in node_name_lower:
        return "Maipú"
    if "buenaventura" in node_name_lower:
        return "Buenaventura"
    if "el bosque" in node_name_lower or "bosque" in node_name_lower:
        return "El Bosque"
    if "quilicura" in node_name_lower:
        return "Quilicura"
    if "curauma" in node_name_lower:
        return "Curauma"
    if "kennedy" in node_name_lower:
        return "Kennedy"
    if "estación" in node_name_lower or "estacion" in node_name_lower:
        return "Estación"
    
    # Si no se puede determinar, retornar cadena vacía en lugar de un valor por defecto incorrecto
    return ""


def get_node_name(node_id: str) -> str:
    """Obtiene el nombre del nodo basado en su ID."""
    node_id = str(node_id).strip()
    # Primero intentar desde el diccionario estático
    if node_id in NODE_NAMES:
        return NODE_NAMES[node_id]

    # Luego intentar desde el cache local
    if node_id in NODE_NAMES_CACHE:
        return NODE_NAMES_CACHE[node_id]
    
    # Si no está en el diccionario, intentar obtener desde la API
    # Extraer company_id del node_id (formato: company_id-node_number)
    try:
        if '-' in node_id:
            company_id = node_id.split('-')[0]
            nodes = COMPANY_NODES_CACHE.get(company_id)
            if nodes is None:
                url = f"{ENTITY_BASE_URL}/companies/{company_id}"
                response = _requests_session().get(url, timeout=5)
                if response.status_code == 200:
                    data = response.json()
                    nodes = data.get('nodes', []) or []
                    COMPANY_NODES_CACHE[company_id] = nodes
                    if company_id not in COMPANY_NAME_CACHE and data.get("name"):
                        COMPANY_NAME_CACHE[company_id] = str(data.get("name"))
                else:
                    nodes = []
            for node in nodes:
                if str(node.get('nodeId', '')).strip() == node_id:
                    name = str(node.get('name', node_id))
                    NODE_NAMES_CACHE[node_id] = name
                    return name
    except Exception:
        pass
    
    return node_id


def limpiar_nombre_archivo(nombre: str) -> str:
    """
    Limpia un nombre para usarlo como nombre de archivo.
    Elimina caracteres especiales y reemplaza espacios con guiones bajos.
    """
    # Reemplazar caracteres problemáticos con guiones bajos
    nombre_limpio = nombre.replace('/', '_').replace('\\', '_').replace(':', '_')
    nombre_limpio = nombre_limpio.replace('*', '_').replace('?', '_').replace('"', '_')
    nombre_limpio = nombre_limpio.replace('<', '_').replace('>', '_').replace('|', '_')
    nombre_limpio = nombre_limpio.replace('-', '_')  # Reemplazar guiones con guiones bajos
    
    # Mantener solo caracteres alfanuméricos, espacios y guiones bajos
    nombre_limpio = "".join(c for c in nombre_limpio if c.isalnum() or c in (" ", "_")).strip()
    
    # Reemplazar espacios con guiones bajos
    nombre_limpio = nombre_limpio.replace(' ', '_')
    
    # Eliminar guiones bajos múltiples (hacerlo varias veces para asegurar)
    import re
    nombre_limpio = re.sub(r'_+', '_', nombre_limpio)
    
    # Eliminar guiones bajos al inicio y final
    nombre_limpio = nombre_limpio.strip('_')
    
    # Limitar longitud (Windows tiene límite de 255 caracteres, pero mejor limitar a 80)
    if len(nombre_limpio) > 80:
        nombre_limpio = nombre_limpio[:80]
    
    return nombre_limpio


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Genera un reporte Word con gráficas de consumo y consumos nocturnos para un nodo PIC."
    )
    parser.add_argument("--company-id", required=True, help="ID de la empresa (CompanyId).")
    parser.add_argument("--node-id", required=True, help="ID del nodo/punto.")
    parser.add_argument("--start-date", required=True, help="Fecha inicio (ISO 8601 o dd/mm/aaaa).")
    parser.add_argument("--end-date", required=True, help="Fecha término (ISO 8601 o dd/mm/aaaa).")
    parser.add_argument(
        "--output-dir",
        default="reports",
        help="Directorio donde se guardará el Word generado (default: reports).",
    )
    # Parámetros opcionales para envío de correo
    parser.add_argument("--enviar-correo", action="store_true", help="Enviar el reporte por correo electrónico.")
    parser.add_argument("--destinatario", help="Correo electrónico del destinatario.")
    parser.add_argument("--smtp-servidor", default="smtp.gmail.com", help="Servidor SMTP (default: smtp.gmail.com).")
    parser.add_argument("--smtp-puerto", type=int, default=587, help="Puerto SMTP (default: 587).")
    parser.add_argument("--smtp-usuario", help="Usuario SMTP (correo del remitente).")
    parser.add_argument("--smtp-password", help="Contraseña SMTP (o contraseña de aplicación).")
    return parser.parse_args()


def parse_date(value: str, end_of_day: bool = False) -> datetime:
    """Acepta múltiples formatos y devuelve UTC sin timezone (ISO)."""
    value = value.strip()
    formats = [
        "%Y-%m-%d",
        "%d/%m/%Y",
        "%d%m%Y",
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M:%SZ",
    ]
    for fmt in formats:
        try:
            dt = datetime.strptime(value, fmt)
            break
        except ValueError:
            continue
    else:
        try:
            dt = datetime.fromisoformat(value.replace("Z", ""))
        except ValueError as exc:
            raise ValueError(f"No se pudo interpretar la fecha '{value}'.") from exc

    if end_of_day and dt.hour == 0 and dt.minute == 0:
        dt = dt.replace(hour=23, minute=59, second=59)
    return dt.replace(tzinfo=timezone.utc)


def isoformat(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


# Sesión por hilo: reutiliza conexiones TCP (keep-alive) y es segura con ThreadPoolExecutor.
_thread_local = threading.local()


def _requests_session() -> requests.Session:
    s = getattr(_thread_local, "session", None)
    if s is None:
        s = requests.Session()
        _thread_local.session = s
    return s


def fetch_json(url: str, params: Optional[Union[dict, Sequence[tuple]]] = None) -> Union[dict, list]:
    try:
        response = _requests_session().get(url, params=params, timeout=30)
        response.raise_for_status()
        if response.headers.get("Content-Type", "").startswith("application/json"):
            return response.json()
        raise ValueError(f"Respuesta inesperada (Content-Type={response.headers.get('Content-Type')}).")
    except requests.RequestException as exc:
        raise RuntimeError(f"Error al consultar {url}: {exc}") from exc


def get_water_price_per_m3(company_id: str, node_id: str, measures_payload: Optional[dict] = None) -> float:
    """
    Obtiene el precio del agua por m³ desde la API.
    Calcula el precio desde nodeKpi.expenses (costo total) dividido por totalM3 (consumo total).
    Retorna un valor por defecto si no se encuentra.
    """
    DEFAULT_PRICE_CLP = 1200.0  # Precio por defecto: $1,200 CLP por m³
    
    # Calcular el precio desde las medidas (nodeKpi.expenses / totalM3)
    prices_found = []
    if measures_payload:
        try:
            for node_measure in measures_payload.get("month", []):
                node_kpi = node_measure.get("nodeKpi")
                total_m3 = node_measure.get("totalM3", 0)
                
                if node_kpi and isinstance(node_kpi, dict) and total_m3 and total_m3 > 0:
                    expenses_str = node_kpi.get("expenses")
                    if expenses_str:
                        try:
                            # Limpiar el string de expenses (puede tener formato "$1,200" o "1200" o similar)
                            expenses_clean = str(expenses_str).replace("$", "").replace(",", "").replace(".", "").replace(" ", "").strip()
                            # Remover cualquier carácter no numérico al final
                            expenses_clean = ''.join(c for c in expenses_clean if c.isdigit() or c == '.')
                            if expenses_clean:
                                expenses_value = float(expenses_clean)
                                if expenses_value > 0:
                                    # Calcular precio por m³: expenses (costo total) / totalM3 (consumo)
                                    price_per_m3 = expenses_value / float(total_m3)
                                    # Validar que el precio esté en un rango razonable (100-10000 CLP/m³)
                                    if 100 <= price_per_m3 <= 10000:
                                        prices_found.append(price_per_m3)
                        except (ValueError, TypeError, ZeroDivisionError):
                            continue
            
            # Si encontramos precios válidos, retornar el promedio
            if prices_found:
                return sum(prices_found) / len(prices_found)
        except Exception:
            pass
    
    # Intentar obtener el precio desde la API de entities (campo "amount")
    # Extraer el ID del punto sin el guion y sin los números después del guion
    # Ejemplo: "000025-17" -> "000025"
    node_id_base = node_id.split("-")[0] if "-" in node_id else node_id
    
    entity_endpoints_to_try = [
        (f"{ENTITY_BASE_URL}/companies/{company_id}/nodes", None),  # GET /companies/{companyId}/nodes
        (f"{ENTITY_BASE_URL}/companies/{company_id}", None),  # GET /companies/{companyId}
        (f"{ENTITY_BASE_URL}/nodes/{node_id_base}", None),  # GET /nodes/{nodeIdBase} (sin guion)
    ]
    
    # Lista para almacenar todos los amounts encontrados con sus fechas
    amounts_with_dates = []
    
    for endpoint_url, params in entity_endpoints_to_try:
        try:
            data = fetch_json(endpoint_url, params=params)
            
            def find_amounts_with_dates(obj, parent_date=None, depth=0):
                """Busca recursivamente todos los campos 'amount' con sus fechas asociadas."""
                if depth > 10:  # Limitar profundidad
                    return
                
                if isinstance(obj, dict):
                    # Buscar campos de fecha comunes
                    date_fields = ["date", "createdDate", "updatedDate", "creationDate", "updateDate", 
                                  "startDate", "endDate", "effectiveDate", "validFrom", "validTo"]
                    current_date = parent_date
                    
                    for date_field in date_fields:
                        if date_field in obj:
                            try:
                                date_str = obj[date_field]
                                if isinstance(date_str, str):
                                    # Intentar parsear la fecha
                                    for fmt in ["%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%SZ", 
                                                "%d/%m/%Y", "%d-%m-%Y"]:
                                        try:
                                            current_date = datetime.strptime(date_str.split("T")[0], fmt)
                                            break
                                        except ValueError:
                                            continue
                            except Exception:
                                pass
                    
                    # Si encontramos amount, guardarlo con su fecha
                    if "amount" in obj:
                        value = obj["amount"]
                        amount_value = None
                        
                        if isinstance(value, (int, float)) and value > 0:
                            amount_value = float(value)
                        elif isinstance(value, str):
                            try:
                                value_clean = value.replace("$", "").replace(",", "").replace(" ", "").strip()
                                # Manejar formato con punto decimal
                                if "." in value_clean:
                                    value_clean = ''.join(c for c in value_clean if c.isdigit() or c == '.')
                                else:
                                    value_clean = ''.join(c for c in value_clean if c.isdigit())
                                if value_clean:
                                    amount_value = float(value_clean)
                            except (ValueError, TypeError):
                                pass
                        
                        if amount_value and 100 <= amount_value <= 10000:
                            amounts_with_dates.append({
                                "amount": amount_value,
                                "date": current_date if current_date else datetime.min,
                                "source": endpoint_url
                            })
                    
                    # Buscar recursivamente en valores del diccionario
                    for key, val in obj.items():
                        find_amounts_with_dates(val, current_date, depth + 1)
                        
                elif isinstance(obj, list):
                    # Buscar en cada elemento de la lista
                    for item in obj:
                        find_amounts_with_dates(item, parent_date, depth + 1)
            
            find_amounts_with_dates(data)
            
        except Exception:
            continue
    
    # Si encontramos amounts con fechas, usar el más reciente
    if amounts_with_dates:
        # Ordenar por fecha descendente (más reciente primero)
        amounts_with_dates.sort(key=lambda x: x["date"], reverse=True)
        most_recent = amounts_with_dates[0]
        return most_recent["amount"]
    
    # Si encontramos amounts pero sin fechas, usar el primero
    # (esto es un fallback por si no hay fechas en la respuesta)
    
    # Si no se encontró en entities, intentar desde diferentes endpoints del nodo
    endpoints_to_try = [
        (f"{acl_node_base_url()}/nodes", [("id", node_id)]),  # GET /nodes?id=...
        (f"{acl_node_base_url()}/nodes/{node_id}", None),  # GET /nodes/{id}
    ]
    
    for endpoint_url, params in endpoints_to_try:
        try:
            data = fetch_json(endpoint_url, params=params)
            if isinstance(data, dict):
                # Buscar campos comunes que puedan contener el precio
                price_fields = [
                    "pricePerM3", "pricePerM3CLP", "price", "costPerM3", 
                    "costPerM3CLP", "cost", "waterPrice", "waterCost",
                    "tarifa", "precio", "costo", "expenses", "amount"
                ]
                for field in price_fields:
                    if field in data:
                        value = data[field]
                        # Si es string, intentar convertirlo
                        if isinstance(value, str):
                            try:
                                value_clean = value.replace("$", "").replace(",", "").replace(".", "").replace(" ", "").strip()
                                value_clean = ''.join(c for c in value_clean if c.isdigit() or c == '.')
                                if value_clean:
                                    value = float(value_clean)
                                else:
                                    continue
                            except (ValueError, TypeError):
                                continue
                        if isinstance(value, (int, float)) and value > 0:
                            if 100 <= value <= 10000:  # Validar rango razonable
                                return float(value)
            elif isinstance(data, list) and len(data) > 0:
                # Si es una lista, buscar en el primer elemento
                node_data = data[0]
                if isinstance(node_data, dict):
                    for field in ["pricePerM3", "pricePerM3CLP", "price", "costPerM3", "cost", "expenses", "amount"]:
                        if field in node_data:
                            value = node_data[field]
                            # Si es string, intentar convertirlo
                            if isinstance(value, str):
                                try:
                                    value_clean = value.replace("$", "").replace(",", "").replace(".", "").replace(" ", "").strip()
                                    value_clean = ''.join(c for c in value_clean if c.isdigit() or c == '.')
                                    if value_clean:
                                        value = float(value_clean)
                                    else:
                                        continue
                                except (ValueError, TypeError):
                                    continue
                            if isinstance(value, (int, float)) and value > 0:
                                if 100 <= value <= 10000:
                                    return float(value)
        except Exception:
            continue
    
    # Si no se encontró el precio, retornar el valor por defecto
    return DEFAULT_PRICE_CLP


@dataclass
class MeasurePoint:
    date: datetime
    total_m3: float
    details: dict


def normalize_measures_payload(payload: Union[dict, list], node_id: str) -> dict:
    """El endpoint puede devolver un dict (/{id}/...) o una lista (/measures/dates)."""
    if isinstance(payload, dict):
        return payload
    if isinstance(payload, list):
        for item in payload:
            if isinstance(item, dict) and item.get("nodeId") == node_id:
                return item
        return {}
    return {}


def flatten_measures(payload: dict) -> List[MeasurePoint]:
    results: List[MeasurePoint] = []
    for node_measure in payload.get("month", []):
        date_str = node_measure.get("date")
        total = node_measure.get("totalM3")
        if not date_str or total is None:
            continue
        try:
            dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        except ValueError:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
        results.append(
            MeasurePoint(
                date=dt,
                total_m3=float(total),
                details=node_measure,
            )
        )
    return sorted(results, key=lambda m: m.date)


def _parse_alert_creation_date(alert: dict) -> Optional[datetime]:
    creation = alert.get("creationDate")
    if not creation:
        return None
    try:
        return datetime.fromisoformat(creation.replace("Z", "+00:00"))
    except (ValueError, AttributeError):
        return None


def _filter_measures_by_range(
    measures: List[MeasurePoint],
    start_date: Optional[datetime] = None,
    end_date: Optional[datetime] = None,
) -> List[MeasurePoint]:
    if not start_date and not end_date:
        return list(measures)
    filtered: List[MeasurePoint] = []
    for m in measures:
        m_date_only = m.date.date()
        if start_date and m_date_only < start_date.date():
            continue
        if end_date and m_date_only > end_date.date():
            continue
        filtered.append(m)
    return filtered


# Alertas con medida por debajo de este umbral (m³/h) no se marcan ni listan en informes extendidos.
ALERTA_MIN_MEDIDA_INFORME_M3H = 0.01


def alerta_medida_informativa(alert: dict, umbral: float = ALERTA_MIN_MEDIDA_INFORME_M3H) -> bool:
    return float(alert.get("measure", 0) or 0) >= umbral


def filtrar_alertas_informativas(
    alerts: Optional[List[dict]],
    umbral: float = ALERTA_MIN_MEDIDA_INFORME_M3H,
) -> List[dict]:
    if not alerts:
        return []
    return [a for a in alerts if alerta_medida_informativa(a, umbral)]


def alertas_marcadas_grafico_diario(
    alerts: Optional[List[dict]],
    measures: List[MeasurePoint],
    start_date: Optional[datetime] = None,
    end_date: Optional[datetime] = None,
) -> List[dict]:
    """Alertas cuya fecha coincide con un día del gráfico diario (marcador rojo)."""
    alerts = filtrar_alertas_informativas(alerts)
    filtered_measures = _filter_measures_by_range(measures, start_date, end_date)
    if not filtered_measures:
        return []
    measure_dates = {m.date.date() for m in filtered_measures}
    matched: List[dict] = []
    for alert in alerts:
        dt = _parse_alert_creation_date(alert)
        if dt and dt.date() in measure_dates:
            matched.append(alert)
    matched.sort(
        key=lambda a: _parse_alert_creation_date(a)
        or datetime.min.replace(tzinfo=timezone.utc)
    )
    return matched


def agregar_tabla_alertas_grafico_diario(doc: Document, alerts: List[dict], *, wes_style: bool = False) -> None:
    """Tabla día / hora / tamaño bajo el gráfico diario (hora Chile)."""
    rows = [["Día", "Hora", "Tamaño de la alerta (m³/h)"]]
    for alert in alerts:
        dt = _parse_alert_creation_date(alert)
        if not dt:
            continue
        dt_cl = _dt_to_chile(dt)
        measure = float(alert.get("measure", 0) or 0)
        rows.append(
            [
                dt_cl.strftime("%d/%m/%Y"),
                dt_cl.strftime("%H:%M"),
                format_number_chilean(measure, 2),
            ]
        )
    if len(rows) <= 1:
        return
    add_table(doc, "Alertas detectadas en el gráfico", rows, wes_style=wes_style, has_total_row=False)


def build_consumption_chart(measures: List[MeasurePoint], output: Path, start_date: Optional[datetime] = None, end_date: Optional[datetime] = None, alerts: Optional[List[dict]] = None) -> Optional[Path]:
    if not measures:
        return None
    
    filtered_measures = _filter_measures_by_range(measures, start_date, end_date)
    
    if not filtered_measures:
        return None
    
    dates = [m.date for m in filtered_measures]
    totals = [m.total_m3 for m in filtered_measures]
    
    # Crear mapeo de fechas con alertas (solo medida >= umbral informativo)
    alert_dates = set()
    if alerts:
        for alert in filtrar_alertas_informativas(alerts):
            dt = _parse_alert_creation_date(alert)
            if dt:
                alert_dates.add(dt.date())

    plt.figure(figsize=(10, 5))
    
    # Crear gráfica de línea con área sombreada
    plt.plot(dates, totals, marker="o", linestyle="-", color="#0050b3", linewidth=2, markersize=6, label="Consumo diario")
    
    # Agregar área sombreada debajo de la línea
    plt.fill_between(dates, totals, alpha=0.3, color="#0050b3")
    
    # Marcar con círculos rojos las fechas con alertas
    alert_label_added = False
    for i, date in enumerate(dates):
        if date.date() in alert_dates:
            label = "Alerta detectada" if not alert_label_added else ""
            plt.plot(date, totals[i], 'ro', markersize=12, markeredgewidth=2, markeredgecolor='darkred', zorder=5, label=label)
            alert_label_added = True
    
    plt.title("Consumo diario (m³)", fontsize=14, fontweight="bold")
    plt.xlabel("Fecha", fontsize=11)
    plt.ylabel("Total m³", fontsize=11)
    
    # Forzar que el eje Y empiece en 0 y no muestre números negativos
    plt.ylim(bottom=0)
    ax = plt.gca()
    ax.set_ylim(bottom=0)
    # Ocultar números negativos en el eje Y
    yticks = ax.get_yticks()
    yticks = yticks[yticks >= 0]
    ax.set_yticks(yticks)
    
    # Formatear fechas en el eje X (formato: "3. Nov.", "10. Nov.", etc.)
    # Mejorar la visualización de fechas: mostrar más fechas y rotarlas para mejor legibilidad
    ax.xaxis.set_major_locator(mdates.DayLocator(interval=max(1, len(dates) // 10)))  # Mostrar aproximadamente 10 fechas
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%d. %b'))
    ax.xaxis.set_minor_locator(mdates.DayLocator(interval=1))
    
    # Rotar las etiquetas de fecha para mejor legibilidad
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha='right')
    
    plt.grid(True, linestyle="--", alpha=0.3, axis="y")
    # Mover leyenda abajo de la gráfica, más abajo para que no tape el eje X
    # Ajustar bbox_to_anchor para dar más espacio al eje X
    if alert_dates:
        plt.legend(loc="upper center", bbox_to_anchor=(0.5, -0.20), ncol=2, fontsize=9)
    else:
        plt.legend(loc="upper center", bbox_to_anchor=(0.5, -0.20), ncol=1, fontsize=9)
    # Ajustar layout para dar más espacio en la parte inferior
    plt.tight_layout(rect=[0, 0.05, 1, 0.95])  # Dejar más espacio abajo para la leyenda
    plt.gcf().autofmt_xdate()
    plt.savefig(output, dpi=150, bbox_inches='tight', pad_inches=0.2)
    plt.close()
    return output


def calculate_weekly_averages(measures: List[MeasurePoint]) -> Optional[dict]:
    """
    Calcula promedios por día de la semana (lunes a domingo) considerando solo semanas completas.
    Una semana completa tiene todos los días de lunes a domingo.
    
    Retorna un diccionario con:
    - 'averages': dict con claves 'Lunes', 'Martes', etc. y valores promedio
    - 'weeks': lista de tuplas (start_date, end_date) representando las semanas completas
    - 'first_week_start': fecha de inicio de la primera semana completa
    - 'last_week_end': fecha de fin de la última semana completa
    """
    if not measures:
        return None
    
    # Crear diccionario por fecha
    consumption_by_date = {}
    for m in measures:
        date_key = m.date.date()
        consumption_by_date[date_key] = m.total_m3
    
    # Encontrar el rango de fechas
    all_dates = sorted(consumption_by_date.keys())
    if not all_dates:
        return None
    
    min_date = all_dates[0]
    max_date = all_dates[-1]
    
    # Encontrar el primer lunes en el rango
    first_monday = min_date
    while first_monday.weekday() != 0:  # 0 = lunes
        first_monday += timedelta(days=1)
    
    # Encontrar el último domingo en el rango
    last_sunday = max_date
    while last_sunday.weekday() != 6:  # 6 = domingo
        last_sunday -= timedelta(days=1)
    
    # Si no hay semanas completas, retornar None
    if first_monday > last_sunday:
        return None
    
    # Agrupar por semanas completas
    weeks_complete = []
    current_monday = first_monday
    
    while current_monday <= last_sunday:
        week_end = current_monday + timedelta(days=6)  # Domingo
        # Verificar que todos los días de la semana tengan datos
        week_complete = True
        for i in range(7):
            check_date = current_monday + timedelta(days=i)
            if check_date not in consumption_by_date:
                week_complete = False
                break
        
        if week_complete:
            weeks_complete.append((current_monday, week_end))
        
        current_monday += timedelta(days=7)
    
    # Si hay menos de 2 semanas completas, retornar None
    if len(weeks_complete) < 2:
        return None
    
    # Calcular promedios por día de la semana
    day_names = ['Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes', 'Sábado', 'Domingo']
    day_consumptions = {day: [] for day in day_names}
    
    # Agregar consumos por día de la semana para cada semana completa
    for week_start, week_end in weeks_complete:
        for i in range(7):
            day_date = week_start + timedelta(days=i)
            if day_date in consumption_by_date:
                day_name = day_names[i]
                day_consumptions[day_name].append(consumption_by_date[day_date])
    
    # Calcular promedios
    averages = {}
    for day_name in day_names:
        if day_consumptions[day_name]:
            averages[day_name] = sum(day_consumptions[day_name]) / len(day_consumptions[day_name])
        else:
            averages[day_name] = 0.0
    
    return {
        'averages': averages,
        'weeks': weeks_complete,
        'first_week_start': weeks_complete[0][0],
        'last_week_end': weeks_complete[-1][1]
    }


def build_weekly_averages_chart(weekly_data: dict, output: Path) -> Optional[Path]:
    """
    Genera un gráfico de barras verticales azules con los promedios por día de la semana.
    Muestra el valor promedio arriba de cada barra.
    """
    if not weekly_data or 'averages' not in weekly_data:
        return None
    
    averages = weekly_data['averages']
    day_names = ['Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes', 'Sábado', 'Domingo']
    
    # Obtener valores en el orden correcto
    values = [averages.get(day, 0.0) for day in day_names]
    
    # Crear gráfico
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Crear barras verticales azules
    bars = ax.bar(day_names, values, color='#0050b3', alpha=0.8, edgecolor='#003d8c', linewidth=1.5)
    
    # Agregar valores arriba de cada barra
    for bar, value in zip(bars, values):
        height = bar.get_height()
        # Mostrar el valor con 2 decimales
        ax.text(bar.get_x() + bar.get_width() / 2., height,
                f'{format_number_chilean(value, 2)}',
                ha='center', va='bottom', fontsize=11, fontweight='bold')
    
    # Configurar título y etiquetas
    ax.set_title('Promedio de consumo por día de la semana (m³)', fontsize=14, fontweight='bold', pad=15)
    ax.set_xlabel('Día de la semana', fontsize=11)
    ax.set_ylabel('Promedio (m³)', fontsize=11)
    
    # Forzar que el eje Y empiece en 0
    ax.set_ylim(bottom=0)
    yticks = ax.get_yticks()
    yticks = yticks[yticks >= 0]
    ax.set_yticks(yticks)
    
    # Rotar etiquetas del eje X si es necesario
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=0, ha='center')
    
    # Agregar grid
    ax.grid(True, linestyle='--', alpha=0.3, axis='y')
    
    plt.tight_layout()
    plt.savefig(output, dpi=150, bbox_inches='tight', pad_inches=0.2)
    plt.close()
    return output


def build_leak_chart(alerts: List[dict], output: Path) -> Optional[Path]:
    if not alerts:
        return None
    dates = []
    measures = []
    for alert in alerts:
        creation = alert.get("creationDate")
        measure = alert.get("measure")
        if not creation or measure is None:
            continue
        try:
            dt = datetime.fromisoformat(creation.replace("Z", "+00:00"))
        except ValueError:
            continue
        dates.append(dt)
        measures.append(float(measure))

    if not dates:
        return None

    # Crear figura más ancha para comprimir mejor las fechas
    fig, ax = plt.subplots(figsize=(14, 5))
    
    # Usar posiciones numéricas igualmente espaciadas en el eje X para evitar traslapes
    # Esto asegura que todas las barras tengan el mismo espaciado
    num_bars = len(dates)
    x_positions = list(range(num_bars))  # Posiciones: 0, 1, 2, 3, ...
    
    # Calcular ancho de barras (usar 60% del espaciado entre barras)
    bar_width = 0.6
    
    # Calcular valores máximos antes de crear las barras
    max_measure_value = max(measures) if measures else 0
    max_measure_idx = measures.index(max_measure_value)
    max_measure_date = dates[max_measure_idx]
    
    # Crear barras amarillas más gruesas (agrandadas) usando posiciones numéricas
    bars = ax.bar(x_positions, measures, width=bar_width, color='#FFD700', alpha=0.8, edgecolor='#DAA520', linewidth=1.2)
    
    # Establecer límites iniciales del eje Y
    ax.set_ylim(bottom=0, top=None)
    
    # Obtener el máximo actual para calcular espacio superior
    y_min_current, y_max_current = ax.get_ylim()
    
    # Agrupar barras por fecha para evitar traslapes en los números
    from collections import defaultdict
    bars_by_date = defaultdict(list)
    for i, (bar, measure, date, x_pos) in enumerate(zip(bars, measures, dates, x_positions)):
        date_key = date.strftime('%Y-%m-%d')
        bars_by_date[date_key].append((i, bar, measure, date, x_pos))
    
    # Calcular el espacio necesario para los números sobre las barras
    max_height_needed = 0
    for date_key, bar_group in bars_by_date.items():
        if len(bar_group) > 1:
            max_height_in_group = max(bar.get_height() for _, bar, _, _, _ in bar_group)
            # Calcular espacio necesario para múltiples números apilados
            spacing = max_height_in_group * 0.2
            total_space = max_height_in_group + spacing * len(bar_group)
            max_height_needed = max(max_height_needed, total_space)
        else:
            max_height_needed = max(max_height_needed, bar_group[0][1].get_height())
    
    # Aumentar el límite superior del eje Y para dar espacio a los números y al rectángulo (40% adicional)
    y_top_adjusted = max_height_needed * 1.40 if max_height_needed > 0 else y_max_current * 1.4
    ax.set_ylim(bottom=0, top=y_top_adjusted)
    
    # Establecer título en la parte superior
    ax.set_title("Alertas de consumo nocturno", fontsize=14, fontweight='bold', pad=20)
    ax.set_ylabel("Medida", fontsize=11, fontweight='bold')
    
    # Agregar etiquetas con los valores sobre cada barra (agrandadas)
    # Con lógica mejorada para evitar traslapes cuando hay múltiples barras en la misma fecha
    for date_key, bar_group in bars_by_date.items():
        if len(bar_group) == 1:
            # Una sola barra en esta fecha
            i, bar, measure, date, x_pos = bar_group[0]
            height = bar.get_height()
            label_text = f'{measure:.2f}'
            # Colocar el texto arriba de la barra con pequeño margen
            ax.text(bar.get_x() + bar.get_width() / 2., height + y_top_adjusted * 0.03,
                    label_text,
                    ha='center', va='bottom', fontsize=12, fontweight='bold')
        else:
            # Múltiples barras en la misma fecha - distribuir verticalmente sin traslapes
            bar_group_sorted = sorted(bar_group, key=lambda x: x[2], reverse=True)  # Ordenar por medida descendente
            max_height = max(bar.get_height() for _, bar, _, _, _ in bar_group_sorted)
            spacing = y_top_adjusted * 0.06  # Espaciado fijo basado en el rango del eje
            
            for idx, (i, bar, measure, date, x_pos) in enumerate(bar_group_sorted):
                height = bar.get_height()
                label_text = f'{measure:.2f}'
                # Calcular offset vertical para evitar traslapes
                vertical_offset = max_height + spacing * (idx + 1) + y_top_adjusted * 0.03
                # Asegurar que no exceda el límite superior
                if vertical_offset > y_top_adjusted * 0.92:
                    vertical_offset = y_top_adjusted * 0.92 - spacing * (len(bar_group_sorted) - idx - 1)
                # Colocar el texto arriba de la barra con offset
                ax.text(bar.get_x() + bar.get_width() / 2., vertical_offset,
                        label_text,
                        ha='center', va='bottom', fontsize=12, fontweight='bold')
    
    # Calcular espacio necesario para las fechas verticales
    # Obtener límites actuales del eje Y (ya ajustados con y_top_adjusted)
    y_min_current, y_max_current = ax.get_ylim()
    y_range_current = y_max_current - y_min_current
    
    # Calcular espacio necesario para fechas verticales (aproximadamente 12% del rango)
    space_for_dates = y_range_current * 0.12
    
    # Ajustar límite inferior del eje Y para dar espacio a las fechas
    y_bottom_new = -space_for_dates
    ax.set_ylim(bottom=y_bottom_new, top=y_top_adjusted)
    
    # Agregar fecha en cada barra, rotada verticalmente
    # Con espaciado uniforme, todas las fechas pueden estar en la misma posición vertical
    # Aumentar tamaño de fuente para mejor legibilidad
    date_y_position = y_bottom_new * 0.7  # Posición en el espacio negativo reservado
    
    # Agrupar barras por fecha para detectar fechas duplicadas
    from collections import defaultdict
    dates_by_position = defaultdict(list)
    for i, (bar, date, x_pos) in enumerate(zip(bars, dates, x_positions)):
        date_key = date.strftime('%Y-%m-%d')
        dates_by_position[date_key].append((i, bar, date, x_pos))
    
    for date_key, date_group in dates_by_position.items():
        if len(date_group) == 1:
            # Una sola barra para esta fecha
            i, bar, date, x_pos = date_group[0]
            date_text = date.strftime('%d-%m-%y')
            bar_center_x = bar.get_x() + bar.get_width() / 2.
            
            # Colocar la fecha debajo de la barra, rotada 90 grados (vertical)
            # Tamaño de fuente aumentado para mejor legibilidad
            ax.text(bar_center_x, date_y_position,
                    date_text,
                    ha='center', va='top', fontsize=11, fontweight='bold',
                    rotation=90)
        else:
            # Múltiples barras para la misma fecha - distribuir fechas verticalmente
            # Usar diferentes posiciones verticales para evitar traslapes
            for idx, (i, bar, date, x_pos) in enumerate(date_group):
                date_text = date.strftime('%d-%m-%y')
                bar_center_x = bar.get_x() + bar.get_width() / 2.
                
                # Distribuir fechas en diferentes niveles verticales
                offset_factor = (idx - len(date_group) / 2 + 0.5) * 0.15
                date_y_pos = date_y_position * (1 + offset_factor)
                
                # Colocar la fecha debajo de la barra, rotada 90 grados (vertical)
                # Tamaño de fuente aumentado para mejor legibilidad
                ax.text(bar_center_x, date_y_pos,
                        date_text,
                        ha='center', va='top', fontsize=10, fontweight='bold',
                        rotation=90)
    
    # Guardar información del punto máximo para marcarlo después de establecer límites
    max_bar = bars[max_measure_idx]
    max_bar_x = max_bar.get_x() + max_bar.get_width() / 2.
    max_bar_y = max_bar.get_height()
    max_point_info = {
        'x': max_bar_x,
        'y': max_bar_y,
        'text': f'{max_measure_value:.2f} m³/h',
        'label': f'Máximo: {max_measure_value:.2f} m³/h'
    }
    
    # Obtener límites actuales del eje Y para calcular posición de anotación
    # (ya ajustados para incluir espacio de fechas)
    y_min, y_max_axis = ax.get_ylim()
    y_range = y_max_axis - y_min
    
    # Dibujar un punto rojo más pequeño en la posición del máximo para evitar traslape con anotación
    ax.plot(max_point_info['x'], max_point_info['y'], 'ro', markersize=8, markeredgecolor='darkred', 
            markeredgewidth=1.5, label=max_point_info['label'])
    
    # Calcular posición de la anotación considerando los límites del eje Y
    # Posicionar la anotación en la parte superior del gráfico, muy lejos de las fechas
    # Usar 90-95% del límite superior para asegurar que no se traslape con números ni fechas
    annotation_y = y_max_axis * 0.92
    
    # Si el máximo está muy cerca de la anotación, ajustar aún más hacia arriba
    if max_point_info['y'] > y_max_axis * 0.70:
        annotation_y = y_max_axis * 0.95
    
    # Asegurar que la anotación esté siempre en la parte superior, lejos de cualquier elemento
    # Usar al menos el 92% del límite superior para garantizar separación con fechas
    if annotation_y < y_max_axis * 0.92:
        annotation_y = y_max_axis * 0.92
    
    # Agregar anotación con el valor del máximo consumo
    # Posicionar el rectángulo en la parte superior, centrado horizontalmente en el máximo
    # Usar va='top' para que el rectángulo quede completamente arriba, lejos de las fechas
    ax.annotate(max_point_info['text'],
                xy=(max_point_info['x'], max_point_info['y']),
                xytext=(max_point_info['x'], annotation_y),
                ha='center', va='top',
                fontsize=10, fontweight='bold', color='red',
                bbox=dict(boxstyle='round,pad=0.5', facecolor='yellow', alpha=0.7, edgecolor='red', linewidth=2),
                arrowprops=dict(arrowstyle='->', color='red', lw=2))
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    # Ocultar etiquetas del eje X ya que cada barra tiene su propia fecha
    ax.set_xticks([])
    ax.set_xlabel("")  # Remover etiqueta del eje X ya que las fechas están en cada barra
    
    # Ajustar layout para que el título esté en la parte superior y todo quede dentro del rectángulo
    # Dejar espacio superior para el título y espacio inferior para las fechas
    plt.tight_layout(rect=[0, 0.08, 1, 0.92])  # 8% inferior para fechas, 8% superior para título
    plt.savefig(output, dpi=300, bbox_inches='tight', pad_inches=0.2)
    plt.close()
    return output


def summarize_consumption(measures: List[MeasurePoint]) -> dict:
    if not measures:
        return {
            "total": 0.0,
            "promedio_diario": 0.0,
            "dias": 0,
            "max": None,
            "min": None,
        }
    total = sum(m.total_m3 for m in measures)
    dias = len({m.date.date() for m in measures})
    promedio = total / dias if dias else 0.0
    max_point = max(measures, key=lambda m: m.total_m3)
    min_point = min(measures, key=lambda m: m.total_m3)
    return {
        "total": total,
        "promedio_diario": promedio,
        "dias": dias,
        "max": max_point,
        "min": min_point,
    }


def summarize_alerts(alerts: List[dict], start_date: Optional[datetime] = None, end_date: Optional[datetime] = None) -> dict:
    """
    Calcula promedio de alerta y proyección diaria.
    La proyección diaria se calcula como el promedio de las 2 últimas alertas registradas con medida > 0
    que estén en los últimos 2 días y en horario nocturno (22:00 a 07:00).
    Siempre retorna la cantidad total de alertas (solo las que tienen medida > 0).
    """
    # Contar solo alertas con medida mayor a cero
    alerts_con_medida = [a for a in alerts if float(a.get("measure", 0) or 0) > 0] if alerts else []
    cantidad_total = len(alerts_con_medida)
    
    if not alerts or not end_date:
        return {
            "promedio_alerta": 0.0,
            "proyeccion_24h": 0.0,
            "cantidad": cantidad_total,
            "hay_alertas_ultimos_2_dias": False,
        }
    
    # Últimos 2 días calendario en Chile (fin de periodo del reporte), no en UTC crudo
    end_aware = end_date if end_date.tzinfo else end_date.replace(tzinfo=timezone.utc)
    end_chile_date = _dt_to_chile(end_aware).date()
    ultimos_2_dias = {end_chile_date - timedelta(days=i) for i in range(2)}
    
    # Filtrar alertas de los últimos 2 días, con medida > 0, y en horario nocturno (22:00 a 07:00) en Chile
    alerts_nocturnas_ultimos_2_dias = []
    for alert in alerts_con_medida:
        creation = alert.get("creationDate")
        if creation:
            try:
                dt = datetime.fromisoformat(creation.replace("Z", "+00:00"))
                dt_chile = _dt_to_chile(dt)
                if dt_chile.date() in ultimos_2_dias:
                    hora = dt_chile.hour
                    if hora >= 22 or hora <= 7:
                        alerts_nocturnas_ultimos_2_dias.append((dt_chile, alert))
            except (ValueError, AttributeError, OSError):
                continue
    
    # Ordenar por fecha descendente (más recientes primero)
    alerts_nocturnas_ultimos_2_dias.sort(key=lambda x: x[0], reverse=True)
    
    # Tomar las 2 últimas alertas nocturnas de los últimos 2 días
    ultimas_2_alertas = [alert for _, alert in alerts_nocturnas_ultimos_2_dias[:2]]
    
    hay_alertas_ultimos_2_dias = len(ultimas_2_alertas) > 0
    
    # Debug: imprimir información detallada sobre las alertas usadas
    print(f"DEBUG: Total alertas con medida > 0: {len(alerts_con_medida)}")
    print(f"DEBUG: Alertas nocturnas (22:00-07:00) en últimos 2 días encontradas: {len(alerts_nocturnas_ultimos_2_dias)}")
    
    # Si no hay alertas nocturnas en los últimos 2 días, retornar proyección cero
    if not ultimas_2_alertas:
        print(f"DEBUG: No hay alertas nocturnas (22:00-07:00) en los últimos 2 días. Proyección = 0")
        if alerts_nocturnas_ultimos_2_dias:
            print(f"DEBUG: Alertas nocturnas encontradas (pero no suficientes para cálculo):")
            for i, (dt, alert) in enumerate(alerts_nocturnas_ultimos_2_dias, 1):
                measure = float(alert.get("measure", 0) or 0)
                print(f"  {i}. Fecha={dt.strftime('%Y-%m-%d %H:%M')}, Medida={measure} m³/h")
        return {
            "promedio_alerta": 0.0,
            "proyeccion_24h": 0.0,
            "cantidad": cantidad_total,
            "hay_alertas_ultimos_2_dias": False,
        }
    
    # Calcular promedio de las 2 últimas alertas nocturnas
    medidas = [float(a.get("measure", 0) or 0) for a in ultimas_2_alertas]
    promedio = sum(medidas) / len(medidas) if medidas else 0.0
    proyeccion = promedio * 24.0
    
    # Debug: imprimir información detallada sobre las alertas usadas
    print(f"DEBUG: Proyección diaria calculada con {len(ultimas_2_alertas)} alertas nocturnas de los últimos 2 días:")
    print(f"DEBUG: Todas las alertas nocturnas (22:00-07:00) en últimos 2 días:")
    for i, (dt, alert) in enumerate(alerts_nocturnas_ultimos_2_dias, 1):
        measure = float(alert.get("measure", 0) or 0)
        print(f"  {i}. Fecha={dt.strftime('%Y-%m-%d %H:%M')}, Medida={measure} m³/h")
    print(f"DEBUG: Las 2 últimas alertas nocturnas usadas para el cálculo:")
    for i, alert in enumerate(ultimas_2_alertas, 1):
        creation = alert.get("creationDate", "")
        measure = float(alert.get("measure", 0) or 0)
        dt_raw = datetime.fromisoformat(creation.replace("Z", "+00:00"))
        dt = _dt_to_chile(dt_raw)
        print(f"  Alerta {i}: Fecha={dt.strftime('%Y-%m-%d %H:%M')} (Chile), Medida={measure} m³/h")
    print(f"  Suma de las 2 últimas: {sum(medidas)} m³/h")
    print(f"  Promedio (suma/2): {promedio} m³/h")
    print(f"  Proyección diaria (promedio * 24): {proyeccion} m³/día")
    print(f"  Proyección redondeada: {round(proyeccion, 1)} m³/día")
    
    return {
        "promedio_alerta": promedio,
        "proyeccion_24h": proyeccion,
        "cantidad": cantidad_total,
        "hay_alertas_ultimos_2_dias": hay_alertas_ultimos_2_dias,
    }


def add_formatted_heading(doc: Document, text: str, level: int = 1, *, page_break_before: bool = False) -> None:
    """Agrega un encabezado con formato: mayúsculas, negrita, azul oscuro.
    Configura el párrafo para mantenerlo junto con el siguiente elemento."""
    if page_break_before:
        doc.add_page_break()
    heading = doc.add_heading(text.upper(), level=level)
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
    # Mantener el encabezado junto con el siguiente párrafo
    heading.paragraph_format.keep_with_next = True
    # Evitar viudas y huérfanas
    heading.paragraph_format.widow_control = True
    # Espaciado mínimo después del encabezado
    heading.paragraph_format.space_after = Pt(3)


def add_formatted_title(doc: Document, text: str) -> None:
    """Agrega un título con formato: mayúsculas, negrita, azul oscuro.
    Configura el párrafo para mantenerlo junto con el siguiente elemento."""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(text.upper())
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
    # Mantener el título junto con el siguiente párrafo (gráfico)
    title_para.paragraph_format.keep_with_next = True
    # Evitar viudas y huérfanas
    title_para.paragraph_format.widow_control = True
    # Espaciado mínimo después del título
    title_para.paragraph_format.space_after = Pt(3)


def add_picture_with_pagination(doc: Document, image_path: str, width: Inches, keep_with_next: bool = True) -> None:
    """Agrega una imagen al documento con control de paginación mejorado.
    Mantiene la imagen junto con el título anterior si keep_with_next es True."""
    pic_para = doc.add_paragraph()
    pic_run = pic_para.add_run()
    pic_run.add_picture(image_path, width=width)
    pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if keep_with_next:
        # Mantener el gráfico junto con el título anterior
        pic_para.paragraph_format.keep_with_next = True
    # Evitar viudas y huérfanas
    pic_para.paragraph_format.widow_control = True
    # Espaciado mínimo después de la imagen
    pic_para.paragraph_format.space_after = Pt(3)


def apply_keep_with_next(paragraph) -> None:
    """Evita que el párrafo quede separado del siguiente (misma página)."""
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.widow_control = True
    paragraph.paragraph_format.space_after = Pt(2)


def add_summary_section(doc: Document, summary: dict, alerts: List[dict], alert_stats: dict, start_date: datetime, end_date: datetime) -> None:
    add_formatted_heading(doc, "Resumen ejecutivo", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Consumo total: {format_number_chilean(summary['total'], 1)} m³.\n")
    p.add_run(f"Promedio diario: {format_number_chilean(summary['promedio_diario'], 1)} m³.\n")
    # Calcular número de días del periodo
    num_dias = (end_date - start_date).days + 1
    p.add_run(f"Número de días del periodo del reporte: {num_dias}.\n")
    if summary["max"]:
        p.add_run(f"Día pico: {summary['max'].date.strftime('%d-%m-%y')} ({format_number_chilean(summary['max'].total_m3, 1)} m³).\n")
    # Contar solo alertas con medida mayor a cero
    alerts_con_medida = [a for a in alerts if float(a.get("measure", 0) or 0) > 0] if alerts else []
    p.add_run(f"Número de alertas de consumo nocturno: {len(alerts_con_medida)}.\n")


def _aplicar_shading_celda(cell, fill_hex: str) -> None:
    from docx.oxml import parse_xml

    try:
        shading = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:val="clear" w:fill="{fill_hex}"/>'
        )
        tc_pr = cell._element.get_or_add_tcPr()
        existing = tc_pr.find(qn("w:shd"))
        if existing is not None:
            tc_pr.remove(existing)
        tc_pr.append(shading)
    except Exception:
        pass


def estilizar_tabla_wes(
    table,
    *,
    highlight_rows: Optional[Iterable[int]] = None,
    has_total_row: bool = True,
) -> None:
    """Estilo compacto WES: encabezado azul, filas alternadas, fila total resaltada."""
    highlight_set = set(highlight_rows or [])
    n_rows = len(table.rows)
    total_idx = n_rows - 1 if has_total_row and n_rows > 1 else -1

    for i, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            if not para.runs and (cell.text or "").strip():
                para.add_run(cell.text.strip())
            if not para.runs:
                continue
            run = para.runs[0]
            run.font.size = Pt(8 if i > 0 else 9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                _aplicar_shading_celda(cell, "003366")
            elif i == total_idx:
                run.bold = True
                _aplicar_shading_celda(cell, "D9E1F2")
            elif i in highlight_set:
                run.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0)
                _aplicar_shading_celda(cell, "FFE6E6")
            elif i % 2 == 0:
                _aplicar_shading_celda(cell, "F2F6FC")


def add_table(
    doc: Document,
    title: str,
    rows: Iterable[Iterable[str]],
    highlight_rows: Optional[List[int]] = None,
    *,
    wes_style: bool = False,
    has_total_row: bool = True,
) -> None:
    doc.add_paragraph("")
    add_formatted_title(doc, title)
    rows = list(rows)
    if not rows:
        doc.add_paragraph("Sin datos disponibles.")
        return
    
    num_cols = len(rows[0])
    
    # Calcular el ancho máximo necesario para cada columna basándose en el contenido
    col_max_lengths = [0] * num_cols
    for row in rows:
        for j, value in enumerate(row):
            if j < num_cols:
                # Calcular longitud del texto (aproximación: 1 carácter ≈ 0.1 pulgadas)
                text_length = len(str(value))
                col_max_lengths[j] = max(col_max_lengths[j], text_length)
    
    # Calcular el ancho total necesario (suma de todos los anchos máximos)
    total_chars = sum(col_max_lengths)
    
    # Ancho total de la tabla
    # Si hay textos largos (más de 30 caracteres en alguna columna), usar ancho mayor
    max_col_length = max(col_max_lengths) if col_max_lengths else 0
    if max_col_length > 30:
        # Para tablas con textos largos (como "Proyección diaria de consumo nocturno"), usar ancho mayor
        table_total_width = Inches(6.5)
    else:
        # Para tablas normales, usar ancho estándar
        table_total_width = Inches(5)
    
    # Calcular anchos proporcionales para cada columna
    col_widths = []
    max_col_length = max(col_max_lengths) if col_max_lengths else 0
    for j, max_len in enumerate(col_max_lengths):
        if total_chars > 0:
            # Proporción basada en caracteres
            proportion = max_len / total_chars
            # Si la primera columna es "Dispositivo", darle un ancho mínimo más grande (1.5 pulgadas)
            if j == 0 and rows and len(rows[0]) > 0 and rows[0][0] == "Dispositivo":
                width_inches = max(1.5, table_total_width.inches * proportion)
            elif j == 0 and max_col_length > 30:
                # Para la primera columna con textos largos (etiquetas), darle más espacio (mínimo 2.5 pulgadas)
                width_inches = max(2.5, table_total_width.inches * proportion)
            elif max_col_length > 30:
                # Para otras columnas en tablas con textos largos, usar un mínimo mayor
                width_inches = max(1.0, table_total_width.inches * proportion)
            else:
                width_inches = max(0.6, table_total_width.inches * proportion)
            width = Inches(width_inches)
        else:
            # Si la primera columna es "Dispositivo", darle un ancho mínimo más grande
            if j == 0 and rows and len(rows[0]) > 0 and rows[0][0] == "Dispositivo":
                width = Inches(1.5)
            else:
                width = Inches(table_total_width.inches / num_cols)  # Distribución equitativa si no hay datos
        col_widths.append(width)
    
    # Ajustar si la suma excede el ancho total (normalizar)
    total_width_inches = sum(w.inches for w in col_widths)
    if total_width_inches > table_total_width.inches:
        factor = table_total_width.inches / total_width_inches
        col_widths = [Inches(w.inches * factor) for w in col_widths]
    
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid" if wes_style else "Light Grid Accent 1"
    
    # Aplicar anchos a las columnas
    # En python-docx, el ancho se especifica en EMU (English Metric Units)
    # 1 pulgada = 914400 EMU
    for j, width in enumerate(col_widths):
        if j < len(table.columns):
            # Convertir de Inches a EMU (1 inch = 914400 EMU)
            width_emu = int(width.inches * 914400)
            table.columns[j].width = width_emu
    
    highlight_set = set(highlight_rows or [])
    if wes_style:
        highlight_set = {i for i in highlight_set if not (has_total_row and i == len(rows) - 1)}

    if not wes_style:
        header_cells = table.rows[0].cells
        for cell in header_cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Inches(0.12)
            else:
                run = cell.paragraphs[0].add_run(cell.text)
                run.bold = True
                run.font.size = Inches(0.12)
            try:
                from docx.oxml import parse_xml
                shading_xml = '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="D9E1F2"/>'
                shading = parse_xml(shading_xml)
                tc_pr = cell._element.get_or_add_tcPr()
                if tc_pr.find(qn("w:shd")) is None:
                    tc_pr.append(shading)
            except Exception:
                pass

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            if j >= num_cols:
                continue
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            if not wes_style:
                try:
                    tc_pr = cell._element.get_or_add_tcPr()
                    no_wrap = OxmlElement('w:noWrap')
                    tc_pr.append(no_wrap)
                except Exception:
                    pass
            if i in highlight_set and not wes_style:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(value)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                try:
                    from docx.oxml import parse_xml
                    shading_xml = '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="FFE6E6"/>'
                    shading = parse_xml(shading_xml)
                    tc_pr = cell._element.get_or_add_tcPr()
                    if tc_pr.find(qn("w:shd")) is None:
                        tc_pr.append(shading)
                except Exception:
                    pass
            else:
                value_str = str(value)
                if "\n" in value_str or "\r" in value_str:
                    cell.paragraphs[0].clear()
                    lines = value_str.replace("\r", "").split("\n")
                    for line_idx, line in enumerate(lines):
                        line_clean = line.strip()
                        if line_clean:
                            if line_idx == 0:
                                cell.paragraphs[0].add_run(line_clean)
                            else:
                                new_para = cell.add_paragraph(line_clean)
                                new_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    clean_value = value_str.replace("\n", " ").replace("\r", " ").strip()
                    while "  " in clean_value:
                        clean_value = clean_value.replace("  ", " ")
                    cell.text = clean_value

    if wes_style:
        estilizar_tabla_wes(table, highlight_rows=highlight_set, has_total_row=has_total_row)


def generate_comparison_narrative(nodes_data: List[dict], avg_consumption_per_node: float) -> str:
    """
    Genera narrativa destacando los hallazgos de la gráfica comparativa por punto.
    Analiza el punto con mayor consumo, menor consumo y comparaciones con el promedio.
    """
    if not nodes_data or len(nodes_data) < 2:
        return "No hay suficientes datos para realizar una comparación entre puntos."
    
    narrative_parts = []
    
    # Ordenar por consumo total
    sorted_nodes = sorted(nodes_data, key=lambda d: d["summary"]["total"], reverse=True)
    
    # Punto con mayor consumo
    max_node = sorted_nodes[0]
    max_name = max_node["node_name"]
    max_consumption = max_node["summary"]["total"]
    
    # Punto con menor consumo
    min_node = sorted_nodes[-1]
    min_name = min_node["node_name"]
    min_consumption = min_node["summary"]["total"]
    
    # Calcular diferencia entre máximo y mínimo
    if min_consumption > 0:
        diff_ratio = (max_consumption / min_consumption) if min_consumption > 0 else 0
    else:
        diff_ratio = 0
    
    # Comparar máximo con promedio
    if avg_consumption_per_node > 0:
        max_vs_avg_pct = ((max_consumption - avg_consumption_per_node) / avg_consumption_per_node * 100)
    else:
        max_vs_avg_pct = 0
    
    # Comparar mínimo con promedio
    if avg_consumption_per_node > 0:
        min_vs_avg_pct = ((min_consumption - avg_consumption_per_node) / avg_consumption_per_node * 100)
    else:
        min_vs_avg_pct = 0
    
    # Construir narrativa
    narrative_parts.append(
        f"El análisis comparativo revela que el punto con mayor consumo es {max_name}, "
        f"con un total de {format_number_chilean(max_consumption, 1)} m³ durante el periodo analizado."
    )
    
    if max_vs_avg_pct > 20:
        narrative_parts.append(
            f"Este valor representa un {format_number_chilean(max_vs_avg_pct, 1)}% por encima del promedio "
            f"({format_number_chilean(avg_consumption_per_node, 1)} m³ por punto), "
            f"indicando una demanda significativamente superior a la media."
        )
    elif max_vs_avg_pct < -20:
        narrative_parts.append(
            f"Este valor se encuentra un {format_number_chilean(abs(max_vs_avg_pct), 1)}% por debajo del promedio "
            f"({format_number_chilean(avg_consumption_per_node, 1)} m³ por punto)."
        )
    else:
        narrative_parts.append(
            f"Este valor se encuentra cercano al promedio de {format_number_chilean(avg_consumption_per_node, 1)} m³ por punto."
        )
    
    narrative_parts.append(
        f"Por otro lado, el punto con menor consumo es {min_name}, "
        f"con un total de {format_number_chilean(min_consumption, 1)} m³."
    )
    
    if min_vs_avg_pct < -20:
        narrative_parts.append(
            f"Este valor representa un {format_number_chilean(abs(min_vs_avg_pct), 1)}% por debajo del promedio, "
            f"mostrando una demanda significativamente inferior."
        )
    elif min_vs_avg_pct > 20:
        narrative_parts.append(
            f"Este valor se encuentra un {format_number_chilean(min_vs_avg_pct, 1)}% por encima del promedio."
        )
    else:
        narrative_parts.append(
            f"Este valor también se encuentra cercano al promedio."
        )
    
    if diff_ratio > 2:
        narrative_parts.append(
            f"La diferencia entre el punto de mayor y menor consumo es considerable, "
            f"siendo el consumo máximo aproximadamente {format_number_chilean(diff_ratio, 1)} veces superior al mínimo, "
            f"lo que sugiere variaciones significativas en los patrones de uso entre los diferentes puntos de monitoreo."
        )
    elif diff_ratio > 1.5:
        narrative_parts.append(
            f"Existe una diferencia moderada entre los puntos, "
            f"con el consumo máximo siendo {format_number_chilean(diff_ratio, 1)} veces superior al mínimo."
        )
    else:
        narrative_parts.append(
            f"Los puntos muestran consumos relativamente homogéneos, "
            f"con una diferencia limitada entre el máximo y el mínimo."
        )
    
    return " ".join(narrative_parts)


def generate_consumption_narrative(summary: dict, measures: List[MeasurePoint], alerts: Optional[List[dict]] = None) -> str:
    """Genera narrativa destacando las 2 alzas más significativas de consumo."""
    if not measures or len(measures) < 2:
        return "No hay suficientes datos para analizar variaciones de consumo."
    
    narrative_parts = []
    
    # Calcular variaciones día a día
    variations = []
    for i in range(1, len(measures)):
        prev = measures[i-1].total_m3
        curr = measures[i].total_m3
        variation = ((curr - prev) / prev * 100) if prev > 0 else 0
        variations.append((measures[i].date, curr, prev, variation))
    
    # Encontrar alzas significativas (>20% de aumento) y ordenarlas por porcentaje descendente
    significant_increases = [v for v in variations if v[3] > 20]
    significant_increases.sort(key=lambda x: x[3], reverse=True)  # Ordenar por porcentaje descendente
    
    # Crear set de fechas con alertas para verificar si hubo alerta ese día
    alert_dates = set()
    if alerts:
        for alert in alerts:
            creation = alert.get("creationDate")
            if creation:
                try:
                    dt = datetime.fromisoformat(creation.replace("Z", "+00:00"))
                    alert_dates.add(dt.date())
                except (ValueError, AttributeError):
                    continue
    
    # Tomar solo las 2 alzas más significativas
    top_2_increases = significant_increases[:2]
    
    if top_2_increases:
        narrative_parts.append("Se detectaron alzas significativas de consumo en los siguientes días:")
        for date, curr, prev, var in top_2_increases:
            date_str = date.strftime('%d-%m-%y')
            alert_text = " El equipo de monitoreo WES alertó este evento." if date.date() in alert_dates else ""
            narrative_parts.append(
                f"• {date_str}: Aumento del {format_number_chilean(var, 1)}% "
                f"(de {format_number_chilean(prev, 1)} m³ a {format_number_chilean(curr, 1)} m³).{alert_text}"
            )
    else:
        narrative_parts.append("No se registraron alzas significativas de consumo durante el periodo analizado.")
    
    # Comparar con promedio
    if summary["max"] and summary["promedio_diario"] > 0:
        max_vs_avg = ((summary["max"].total_m3 - summary["promedio_diario"]) / summary["promedio_diario"] * 100)
        if max_vs_avg > 30:
            narrative_parts.append(
                f"El día pico ({summary['max'].date.strftime('%d-%m-%y')}) registró un consumo "
                f"de {format_number_chilean(summary['max'].total_m3, 1)} m³, lo que representa un {format_number_chilean(max_vs_avg, 1)}% "
                f"por encima del promedio diario ({format_number_chilean(summary['promedio_diario'], 1)} m³)."
            )
    
    return " ".join(narrative_parts)


def add_consumption_section(
    doc: Document,
    summary: dict,
    chart_path: Optional[Path],
    measures: List[MeasurePoint],
    alerts: Optional[List[dict]] = None,
    max_day_chart: Optional[Path] = None,
    min_day_chart: Optional[Path] = None,
    start_dt: Optional[datetime] = None,
    end_dt: Optional[datetime] = None,
    weekly_averages_chart: Optional[Path] = None,
    weekly_averages_data: Optional[dict] = None,
) -> None:
    add_formatted_heading(doc, "Consumo", level=1)
    
    # Inicializar description_text con un valor por defecto
    description_text = "Gráfica de consumo de agua."
    
    # Generar texto descriptivo simplificado
    if measures and start_dt and end_dt:
        # Obtener los días presentes en las medidas
        days_present = sorted([m.date.date() for m in measures])
        num_days = len(days_present)
        
        # Obtener todos los días del rango completo
        all_days_in_range = []
        current_date = start_dt.date()
        while current_date <= end_dt.date():
            all_days_in_range.append(current_date)
            current_date += timedelta(days=1)
        
        # Identificar días de la semana presentes y faltantes
        day_names = ["lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"]
        weekdays_present = set([d.weekday() for d in days_present])
        
        # Construir el texto descriptivo
        description_text = f"En la gráfica se representan {num_days} días."
        
        # Verificar si están todos los días de la semana
        all_weekdays_in_range = set([d.weekday() for d in all_days_in_range])
        missing_weekdays = sorted([wd for wd in all_weekdays_in_range if wd not in weekdays_present])
        
        if len(weekdays_present) == 7 and len(missing_weekdays) == 0:
            description_text += " El periodo incluye todos los días de la semana."
        elif missing_weekdays:
            missing_names = [day_names[wd] for wd in missing_weekdays]
            description_text += f" En el periodo falta día {', '.join(missing_names)}."
    
    desc_para = doc.add_paragraph(description_text)
    desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in desc_para.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
    if chart_path:
        add_picture_with_pagination(doc, str(chart_path), Inches(6), keep_with_next=True)
    
    # Agregar análisis del día con mayor consumo (justo después de la gráfica de consumo diario)
    if summary.get("max"):
        max_measure = summary["max"]
        max_date = max_measure.date.date()
        day_of_week = max_date.weekday()  # 0 = Lunes, 6 = Domingo
        
        day_names = ["lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"]
        day_name = day_names[day_of_week]
        
        # Determinar nivel de actividad hídrica basado en el día de la semana
        # Lunes-Viernes: alta actividad (días laborales)
        # Sábado-Domingo: baja actividad (fines de semana)
        if day_of_week < 5:  # Lunes a Viernes
            actividad = "alta actividad hídrica"
            contexto = "día laboral"
        else:  # Sábado o Domingo
            actividad = "baja actividad hídrica"
            contexto = "fin de semana"
        
        # Generar texto descriptivo
        max_consumption_text = (
            f"El mayor consumo se registró el {day_name} {max_date.strftime('%d-%m-%y')} "
            f"con {format_number_chilean(max_measure.total_m3, 1)} m³. "
            f"Este día corresponde a un {contexto}, lo que típicamente se asocia con "
            f"{actividad}. "
            f"El análisis horario detallado de este día permite identificar los períodos "
            f"de mayor demanda y los patrones de consumo específicos."
        )
        
        max_consumption_para = doc.add_paragraph(max_consumption_text)
        max_consumption_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in max_consumption_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
    
    # Agregar gráfica de promedios por día de la semana si hay 2+ semanas completas
    if weekly_averages_chart and weekly_averages_data:
        # Espacio mínimo antes de la gráfica
        para_spacing = doc.add_paragraph("")
        para_spacing.paragraph_format.space_after = Pt(3)
        
        add_picture_with_pagination(doc, str(weekly_averages_chart), Inches(6), keep_with_next=True)
        
        # Agregar narración explicativa con espacio mínimo
        first_week_start = weekly_averages_data['first_week_start']
        last_week_end = weekly_averages_data['last_week_end']
        num_weeks = len(weekly_averages_data['weeks'])
        
        month_names_es = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                          "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        
        first_week_start_str = f"{first_week_start.day} de {month_names_es[first_week_start.month - 1]} de {first_week_start.year}"
        last_week_end_str = f"{last_week_end.day} de {month_names_es[last_week_end.month - 1]} de {last_week_end.year}"
        
        explanation_text = (
            f"Esta gráfica muestra el promedio de consumo por día de la semana (lunes a domingo), "
            f"calculado exclusivamente con semanas completas del período analizado. "
            f"Para su elaboración, se identificaron {num_weeks} semana(s) completa(s) "
            f"(desde el lunes {first_week_start_str} hasta el domingo {last_week_end_str}), "
            f"excluyendo aquellas semanas que no contenían todos los días de lunes a domingo. "
            f"Los valores mostrados representan el promedio de consumo para cada día de la semana "
            f"considerando únicamente estas semanas completas."
        )
        
        explanation_para = doc.add_paragraph(explanation_text)
        explanation_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        explanation_para.paragraph_format.space_after = Pt(3)  # Espacio mínimo después
        explanation_para.paragraph_format.keep_with_next = True  # Mantener con el siguiente
        for run in explanation_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Agregar gráfica del día con mayor consumo justo después, sin espacio grande
        if max_day_chart:
            # Espacio mínimo antes del título (reducido para mantener todo junto)
            para_before_title = doc.add_paragraph("")
            para_before_title.paragraph_format.space_after = Pt(0)
            para_before_title.paragraph_format.keep_with_next = True
            
            add_formatted_title(doc, "Análisis detallado por hora de los días extremos:")
            add_formatted_title(doc, "Día con mayor consumo:")
            add_picture_with_pagination(doc, str(max_day_chart), Inches(6), keep_with_next=True)
    
    # Agregar gráfica del día con menor consumo (si existe y no se agregó antes)
    if min_day_chart and not (weekly_averages_chart and max_day_chart):
        # Espacio mínimo antes del título
        para_before_min = doc.add_paragraph("")
        para_before_min.paragraph_format.space_after = Pt(3)
        para_before_min.paragraph_format.keep_with_next = True
        
        add_formatted_title(doc, "Análisis detallado por hora de los días extremos:")
        add_formatted_title(doc, "Día con menor consumo:")
        add_picture_with_pagination(doc, str(min_day_chart), Inches(6), keep_with_next=True)
    elif min_day_chart and weekly_averages_chart and max_day_chart:
        # Si ya se agregó el título "Análisis detallado", solo agregar el subtítulo y gráfica
        para_before_min = doc.add_paragraph("")
        para_before_min.paragraph_format.space_after = Pt(3)
        para_before_min.paragraph_format.keep_with_next = True
        
        add_formatted_title(doc, "Día con menor consumo:")
        add_picture_with_pagination(doc, str(min_day_chart), Inches(6), keep_with_next=True)
    
    # Agregar narrativa destacando alzas
    narrative = generate_consumption_narrative(summary, measures, alerts)
    if narrative:
        para_narrative = doc.add_paragraph("")
        para_narrative.paragraph_format.space_after = Pt(3)
        narrative_para = doc.add_paragraph(narrative)
        narrative_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # Asegurar que el texto sea negro (no azul)
        for run in narrative_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
    
    # Calcular promedio mensual (promedio diario * 30 días)
    promedio_mensual = summary['promedio_diario'] * 30.0 if summary['promedio_diario'] > 0 else 0.0
    
    add_table(
        doc,
        "Métricas clave",
        [
            ("Total (m³)", format_number_chilean(summary['total'], 1)),
            ("Promedio diario (m³)", format_number_chilean(summary['promedio_diario'], 1)),
            ("Promedio mensual (m³)", format_number_chilean(promedio_mensual, 1)),
            (
                "Máximo diario",
                f"{format_number_chilean(summary['max'].total_m3, 1)} m³ ({summary['max'].date.strftime('%d-%m-%y')})"
                if summary["max"]
                else "N/D",
            ),
            (
                "Mínimo diario",
                f"{format_number_chilean(summary['min'].total_m3, 1)} m³ ({summary['min'].date.strftime('%d-%m-%y')})"
                if summary["min"]
                else "N/D",
            ),
        ],
    )
    # Tabla de detalle diario eliminada según solicitud del usuario


def find_max_alert_day(alerts: List[dict]) -> Optional[dict]:
    """Encuentra el día con la mayor alerta (solo entre alertas con medida > 0)."""
    if not alerts:
        return None
    # Filtrar solo alertas con medida mayor a cero
    alerts_con_medida = [a for a in alerts if float(a.get("measure", 0) or 0) > 0]
    if not alerts_con_medida:
        return None
    max_alert = max(alerts_con_medida, key=lambda a: float(a.get("measure", 0) or 0))
    return max_alert


def find_max_nocturnal_consumption_day(
    node_id: str,
    measures: Optional[List[MeasurePoint]],
    start_dt: datetime,
    end_dt: datetime,
) -> Optional[tuple]:
    """
    Encuentra el día con mayor consumo nocturno (00:00 a 06:00) en el periodo.
    
    Retorna (datetime, hourly_data) del día con mayor consumo nocturno, o None si no hay datos.
    """
    max_nocturnal_consumption = 0.0
    best_day = None
    best_hourly_data = None
    
    # Iterar sobre todos los días del periodo
    current_date = start_dt.date()
    end_date_only = end_dt.date()
    
    while current_date <= end_date_only:
        try:
            target_dt = datetime.combine(current_date, datetime.min.time())
            consumo_nocturno_dia = 0.0
            hourly_data: Optional[List[tuple]] = None

            if es_nodo_colegio(node_id):
                noct, _, ok = _nocturno_diurno_dia_colegios_utc(node_id, current_date)
                if ok:
                    consumo_nocturno_dia = noct
                    try:
                        by_time = _fetch_csv_colegio_dia(node_id, current_date)
                        rows: List[tuple] = []
                        for time_str in sorted(by_time.keys()):
                            ts_norm = time_str.strip().replace("Z", "+00:00")
                            dt_utc = datetime.fromisoformat(ts_norm)
                            if dt_utc.tzinfo is None:
                                dt_utc = dt_utc.replace(tzinfo=timezone.utc)
                            h = int(dt_utc.hour)
                            if 0 <= h <= 7:
                                rows.append((h, float(by_time[time_str])))
                        hourly_data = rows
                    except Exception:
                        hourly_data = []
            else:
                hourly_data = get_hourly_measures_for_day(node_id, target_dt)
                if hourly_data:
                    for hour, value in hourly_data:
                        if 0 <= hour <= 6:
                            consumo_nocturno_dia += value

            if hourly_data is not None:
                if consumo_nocturno_dia > max_nocturnal_consumption:
                    max_nocturnal_consumption = consumo_nocturno_dia
                    best_day = target_dt
                    best_hourly_data = hourly_data
        
        except Exception as e:
            # DEBUG deshabilitado para reducir ruido en logs
            # print(f"DEBUG: Error al obtener datos horarios para {current_date}: {e}")
            pass
        
        # Avanzar al siguiente día
        current_date += timedelta(days=1)
    
    if best_day and best_hourly_data:
        return (best_day, best_hourly_data)
    
    return None


def find_day_closest_to_alert_average(
    node_id: str,
    measures: List[MeasurePoint],
    alert_average: float,
    start_dt: datetime,
    end_dt: datetime,
    alerts: List[dict] = None,
) -> Optional[tuple]:
    """
    Encuentra el día cuyo consumo por hora promedio se acerca más al promedio de alerta.
    Si hay alertas disponibles, primero busca el día que tiene una alerta más cercana al promedio.
    Retorna (fecha, datos_horarios) o None si no se encuentra.
    
    El promedio de alerta está en m³/h, así que comparamos:
    - Promedio por hora del día = total diario (m³) / 24 horas
    - Con el promedio de alerta (m³/h)
    
    Estrategia: 
    1. Si hay alertas, buscar el día que tiene una alerta más cercana al promedio de alerta
    2. Si no hay alertas o no se encuentra, buscar el día cuyo consumo promedio se acerca más
    """
    if not measures or alert_average <= 0:
        return None
    
    # Primero, si hay alertas, buscar el día que tiene una alerta más cercana al promedio
    if alerts:
        alert_candidates = []
        for alert in alerts:
            raw_date = alert.get("creationDate", "") or ""
            measure_value = float(alert.get("measure", 0) or 0)
            
            if not raw_date or measure_value <= 0:
                continue
            
            try:
                dt = datetime.fromisoformat(raw_date.replace("Z", "+00:00"))
                alert_date = dt.date()
                
                # Verificar que la fecha esté en el rango
                if alert_date < start_dt.date() or alert_date > end_dt.date():
                    continue
                
                # Calcular la diferencia absoluta con el promedio de alerta
                diff = abs(measure_value - alert_average)
                alert_candidates.append((alert_date, diff, measure_value))
            except (ValueError, TypeError):
                continue
        
        if alert_candidates:
            # Ordenar por diferencia (menor diferencia = más cercano)
            alert_candidates.sort(key=lambda x: x[1])
            best_alert_date, best_alert_diff, best_alert_value = alert_candidates[0]
            
            # Obtener los datos horarios reales para el día seleccionado
            try:
                target_dt = datetime.combine(best_alert_date, datetime.min.time())
                hourly_data = get_hourly_measures_for_day(node_id, target_dt)
                
                if hourly_data and len(hourly_data) > 0:
                    return (target_dt, hourly_data)
                else:
                    # Si no hay datos horarios, buscar en las medidas diarias
                    for measure in measures:
                        if measure.date.date() == best_alert_date:
                            hourly_value = measure.total_m3 / 24.0
                            synthetic_data = [(h, hourly_value) for h in range(24)]
                            return (target_dt, synthetic_data)
            except Exception:
                pass
    
    # Si no se encontró usando alertas, buscar el día cuyo consumo promedio se acerca más
    candidates = []
    for measure in measures:
        measure_date = measure.date.date()
        
        # Verificar que la fecha esté en el rango
        if measure_date < start_dt.date() or measure_date > end_dt.date():
            continue
        
        # Intentar obtener datos horarios reales
        try:
            target_dt = datetime.combine(measure_date, datetime.min.time())
            hourly_data = get_hourly_measures_for_day(node_id, target_dt)
            
            if hourly_data and len(hourly_data) > 0:
                # Calcular el promedio por hora usando los datos horarios reales
                hourly_values = [v for _, v in hourly_data]
                daily_hourly_avg = sum(hourly_values) / len(hourly_values) if hourly_values else 0
            else:
                # Si no hay datos horarios, usar el promedio calculado del total diario
                daily_hourly_avg = measure.total_m3 / 24.0
        except Exception:
            # Si falla, usar el promedio calculado del total diario
            daily_hourly_avg = measure.total_m3 / 24.0
        
        # Calcular la diferencia absoluta con el promedio de alerta
        diff = abs(daily_hourly_avg - alert_average)
        
        candidates.append((measure_date, diff, measure, daily_hourly_avg))
    
    if not candidates:
        return None
    
    # Ordenar por diferencia (menor diferencia = más cercano)
    candidates.sort(key=lambda x: x[1])
    
    # Tomar el día con la menor diferencia
    best_date, _, best_measure, _ = candidates[0]
    
    # Ahora obtener los datos horarios reales para el día seleccionado
    try:
        target_dt = datetime.combine(best_date, datetime.min.time())
        hourly_data = get_hourly_measures_for_day(node_id, target_dt)
        
        if hourly_data and len(hourly_data) > 0:
            return (target_dt, hourly_data)
        else:
            # Si no hay datos horarios, crear datos sintéticos basados en el total diario
            hourly_value = best_measure.total_m3 / 24.0
            synthetic_data = [(h, hourly_value) for h in range(24)]
            return (target_dt, synthetic_data)
    except Exception:
        # Si falla, usar datos sintéticos
        target_dt = datetime.combine(best_date, datetime.min.time())
        hourly_value = best_measure.total_m3 / 24.0
        synthetic_data = [(h, hourly_value) for h in range(24)]
        return (target_dt, synthetic_data)


def _utc_calendar_dates_for_chile_day(dia_chile: date) -> List[date]:
    """
    El CSV ``dates.measures.csv`` con ``start=end=ddmmaaaa`` devuelve 24 filas por **día UTC**
    (marcas ``T00Z``…``T23Z`` de ese día). Para reconstruir las 24 h **locales Chile** del día
    civil ``dia_chile`` hay que fusionar todas las respuestas cuyas marcas UTC caen en
    [medianoche Chile, medianoche siguiente), normalmente **uno o dos** días UTC.
    """
    t0 = datetime.combine(dia_chile, datetime.min.time()).replace(tzinfo=CHILE_TZ)
    t1 = t0 + timedelta(days=1)
    u0 = t0.astimezone(timezone.utc).date()
    u1 = (t1 - timedelta(microseconds=1)).astimezone(timezone.utc).date()
    out: List[date] = []
    cur = u0
    while cur <= u1:
        out.append(cur)
        cur += timedelta(days=1)
    return out


def _value_by_time_sum_duplicate_rows(csv_content: str) -> Dict[str, float]:
    """
    Suma ``VALUE`` por cadena ``TIME`` única.

    El endpoint ``dates.measures.csv`` a veces devuelve **dos filas con la misma marca TIME**
    (mismo instante UTC) con el mismo u otro valor. Sumar antes de mapear a hora Chile hace que
    la suma de los caudales horarios coincida con ``totalM3`` del día; si solo se toma la última
    fila, el total queda ~la mitad (caso típico en medición por pulsos / duplicado de registro).
    """
    out: Dict[str, float] = {}
    for line in csv_content.strip().split("\n")[1:]:
        if not line.strip():
            continue
        parts = line.split(",", 1)
        if len(parts) < 2:
            continue
        try:
            time_str = parts[0].strip()
            value_str = parts[1].strip().replace(" ", "").replace(",", ".")
            out[time_str] = out.get(time_str, 0.0) + float(value_str)
        except (ValueError, TypeError, IndexError):
            continue
    return out


def _chile_hours_from_dates_measures_csv_text(csv_content: str, dia_chile: date) -> Dict[int, float]:
    """
    A partir del cuerpo de ``dates.measures.csv``, hora Chile 0–23 para ``dia_chile``.

    Las marcas ``TIME`` (p. ej. ``2026-07-03T07:00:00.000Z``) representan **hora civil Chile**
    del día en la fecha del timestamp; la app WES grafica esa hora sin convertir de UTC.
    Se suman ``VALUE`` duplicados por la misma marca ``TIME``; si hay varias filas para la misma
    hora, se conserva la **última** en orden cronológico del texto.
    """
    target_prefix = dia_chile.strftime("%Y-%m-%d")
    by_time = _value_by_time_sum_duplicate_rows(csv_content)
    acc: Dict[int, float] = {}
    for time_str in sorted(by_time.keys()):
        try:
            ts = time_str.strip()
            if not ts.startswith(target_prefix):
                continue
            hi = int(ts[11:13])
            if 0 <= hi < 24:
                acc[hi] = float(by_time[time_str])
        except (ValueError, TypeError, IndexError):
            continue
    return acc


def _total_m3_from_json_for_chile_day(node_id: str, dia_chile: date) -> Optional[float]:
    """``totalM3`` del día civil en ``/nodes/measures/dates`` (si existe)."""
    target = dia_chile.strftime("%Y-%m-%d")
    for ud in _utc_calendar_dates_for_chile_day(dia_chile):
        date_str = ud.strftime("%d%m%Y")
        try:
            payload_raw = fetch_json(
                f"{acl_node_base_url()}/nodes/measures/dates",
                params=[
                    ("id", node_id),
                    ("start", date_str),
                    ("end", date_str),
                ],
            )
            payload = normalize_measures_payload(payload_raw, node_id)
            for node_measure in payload.get("month", []):
                ds = str(node_measure.get("date", ""))
                if target in ds or ds.startswith(target[:10]):
                    t = node_measure.get("totalM3")
                    if t is not None:
                        return float(t)
        except Exception:
            continue
    return None


def _reconcile_chile_hours_with_total_m3(
    acc_chile: Dict[int, float],
    total_m3: Optional[float],
) -> Tuple[Dict[int, float], bool]:
    """
    Alinea la suma de las 24 horas Chile con ``totalM3`` del JSON (referencia habitual de la app).

    - **CSV casi vacío** frente al total: reparto uniforme ``totalM3/24``.
    - Si la suma CSV difiere de ``totalM3`` en más de **~2 %**, se **escala** cada hora por
      ``totalM3 / suma(CSV)`` (sube o baja), manteniendo la forma del día. Así se corrigen tanto
      sobre-conteos (p. ej. duplicados de ``TIME``) como sub-registros respecto al backoffice.
    - Diferencias ≤ 2 % se dejan (ruido de redondeo).

    Returns:
        (horas 0–23, True si se aplicó ajuste).
    """
    _REL_OK = 0.02  # 2 %: debajo de esto no se toca
    _SCALE_MAX = 15.0
    _SCALE_MIN = 1.0 / _SCALE_MAX

    if total_m3 is None or total_m3 <= 0:
        return acc_chile, False
    t = float(total_m3)
    s = sum(float(acc_chile.get(h, 0.0)) for h in range(24))
    hv = t / 24.0
    if s <= 1e-9:
        return {h: hv for h in range(24)}, True
    # Serie horaria casi nula pero el backoffice tiene volumen
    if t > 3.0 and s < 0.35 * t:
        return {h: hv for h in range(24)}, True
    rel = abs(s - t) / t if t > 1e-9 else 1.0
    if rel <= _REL_OK:
        return acc_chile, False
    scale = t / s
    if not (_SCALE_MIN < scale < _SCALE_MAX):
        return acc_chile, False
    return {h: float(acc_chile.get(h, 0.0)) * scale for h in range(24)}, True


def get_hourly_measures_for_day(node_id: str, target_date: datetime) -> List[tuple]:
    """
    Medidas por hora **Chile** (0–23) para el día civil de ``target_date``.

    Orden de preferencia:

    0. Archivo local ``WES_MEDIDAS_CSV_DIR/AAAA-MM-DD.csv`` si existe (export de la app).
    1. **JSON** ``/nodes/measures/dates``: si hay lista ``measures`` no vacía, se usa como serie horaria.
    2. **CSV** API (fusión UTC → hora Chile, suma por ``TIME`` duplicado). Se **reconcilia** con
       ``totalM3`` del JSON: reparto uniforme si el CSV está casi vacío; si no, escala proporcional
       cuando la suma difiere > ~2 % del total (sube o baja horas para coincidir con la app).
    3. **JSON** ``totalM3/24`` solo si no se pudo armar serie desde CSV ni había ``measures``.

    Env: ``WES_API_BASE_URL``, ``WES_MEDIDAS_CSV_DIR``.
    """
    try:
        dia = target_date.date()
        target_date_str = target_date.strftime("%Y-%m-%d")

        # --- 0) CSV guardado localmente (idéntico al de la app) ---
        override_dir = os.environ.get("WES_MEDIDAS_CSV_DIR", "").strip()
        if override_dir:
            local_csv = Path(override_dir) / f"{dia.isoformat()}.csv"
            if local_csv.is_file():
                body = local_csv.read_text(encoding="utf-8", errors="replace")
                acc0 = _chile_hours_from_dates_measures_csv_text(body, dia)
                tj0 = _total_m3_from_json_for_chile_day(node_id, dia)
                acc0, _ = _reconcile_chile_hours_with_total_m3(acc0, tj0)
                if acc0:
                    return sorted(acc0.items(), key=lambda x: x[0])

        # --- 1) JSON: lista ``measures`` con serie horaria (si existe) ---
        try:
            for ud in _utc_calendar_dates_for_chile_day(dia):
                date_str = ud.strftime("%d%m%Y")
                payload_raw = fetch_json(
                    f"{acl_node_base_url()}/nodes/measures/dates",
                    params=[
                        ("id", node_id),
                        ("start", date_str),
                        ("end", date_str),
                    ],
                )
                payload = normalize_measures_payload(payload_raw, node_id)
                for node_measure in payload.get("month", []):
                    date_str_measure = node_measure.get("date", "")
                    if not (
                        target_date_str in date_str_measure
                        or date_str_measure.startswith(target_date_str[:10])
                    ):
                        continue
                    measures_list = node_measure.get("measures") or []
                    if measures_list:
                        acc_j: Dict[int, float] = {}
                        for measure in measures_list:
                            hour_str = measure.get("hour", "")
                            measurement = measure.get("measurement", "0")
                            try:
                                hour = int(hour_str) if hour_str else 0
                                value = float(measurement)
                                hi = int(hour)
                                if 0 <= hi < 24:
                                    acc_j[hi] = float(value)
                            except (ValueError, TypeError):
                                continue
                        if acc_j:
                            return sorted(acc_j.items(), key=lambda x: x[0])
                    # No devolver totalM3/24 aquí: si ``measures`` viene vacío pero el CSV horario
                    # existe (caso frecuente en nodos como medidores de pulso), la app y los informes
                    # deben usar la curva real de ``dates.measures.csv`` (sección 2), no un reparto uniforme.
        except Exception:
            pass

        # --- 2) CSV por API: una fila por hora civil Chile (fecha del TIME = dia) ---
        acc: Dict[int, float] = {}
        url = f"{acl_node_base_url()}/nodes/{node_id}/dates.measures.csv"
        try:
            date_str = dia.strftime("%d%m%Y")
            response = _requests_session().get(
                url,
                params=[("start", date_str), ("end", date_str)],
                timeout=10,
            )
            response.raise_for_status()
            acc = _chile_hours_from_dates_measures_csv_text(response.text, dia)

            tj_csv = _total_m3_from_json_for_chile_day(node_id, dia)
            acc, _ = _reconcile_chile_hours_with_total_m3(acc, tj_csv)
            if acc:
                return sorted(acc.items(), key=lambda x: x[0])
        except requests.RequestException:
            pass

        # --- 3) JSON: solo total diario cuando no hay lista horaria (mismo criterio que antes) ---
        for ud in _utc_calendar_dates_for_chile_day(dia):
            date_str = ud.strftime("%d%m%Y")
            payload_raw = fetch_json(
                f"{acl_node_base_url()}/nodes/measures/dates",
                params=[
                    ("id", node_id),
                    ("start", date_str),
                    ("end", date_str),
                ],
            )
            payload = normalize_measures_payload(payload_raw, node_id)
            for node_measure in payload.get("month", []):
                date_str_measure = node_measure.get("date", "")
                if not (
                    target_date_str in date_str_measure
                    or date_str_measure.startswith(target_date_str[:10])
                ):
                    continue
                measures_list = node_measure.get("measures") or []
                if measures_list:
                    acc_j: Dict[int, float] = {}
                    for measure in measures_list:
                        hour_str = measure.get("hour", "")
                        measurement = measure.get("measurement", "0")
                        try:
                            hour = int(hour_str) if hour_str else 0
                            value = float(measurement)
                            hi = int(hour)
                            if 0 <= hi < 24:
                                acc_j[hi] = float(value)
                        except (ValueError, TypeError):
                            continue
                    if acc_j:
                        return sorted(acc_j.items(), key=lambda x: x[0])
                total_m3 = node_measure.get("totalM3", 0)
                if total_m3:
                    hourly_value = float(total_m3) / 24.0
                    return [(h, hourly_value) for h in range(24)]

        return []
    except Exception:
        return []


COLEGIOS_COMPANY_IDS = frozenset({"000008"})


def es_nodo_colegio(node_id: str, company_id: Optional[str] = None) -> bool:
    if company_id and company_id in COLEGIOS_COMPANY_IDS:
        return True
    return any(node_id.startswith(f"{cid}-") for cid in COLEGIOS_COMPANY_IDS)


def horas_nocturnas_por_dia_para_nodo(node_id: str, company_id: Optional[str] = None) -> int:
    """Colegios (CORMUP): ventana CSV UTC 00:00–07:00 (8 h). Resto: Chile 00:00–06:59 (7 h)."""
    return HORAS_NOCTURNAS_COLEGIOS_UTC if es_nodo_colegio(node_id, company_id) else HORAS_NOCTURNAS_POR_DIA


def _fetch_csv_colegio_dia(node_id: str, dia: date) -> Dict[str, float]:
    """CSV horario como en la app: ``dates.measures.csv`` con start=end=fecha del día."""
    url = f"{acl_node_base_url()}/nodes/{node_id}/dates.measures.csv"
    ds = dia.strftime("%d%m%Y")
    response = _requests_session().get(url, params=[("start", ds), ("end", ds)], timeout=10)
    response.raise_for_status()
    return _value_by_time_sum_duplicate_rows(response.text)


def _nocturno_diurno_dia_colegios_utc(node_id: str, dia: date) -> Tuple[float, float, bool]:
    """Suma filas TIME con hora UTC 0–7 (nocturno) y 8–23 (diurno) del CSV del día."""
    try:
        by_time = _fetch_csv_colegio_dia(node_id, dia)
    except Exception:
        return 0.0, 0.0, False
    if not by_time:
        return 0.0, 0.0, False
    nocturno = 0.0
    diurno = 0.0
    for time_str, val in by_time.items():
        try:
            ts_norm = time_str.strip().replace("Z", "+00:00")
            dt_utc = datetime.fromisoformat(ts_norm)
            if dt_utc.tzinfo is None:
                dt_utc = dt_utc.replace(tzinfo=timezone.utc)
            v = float(val)
            h = int(dt_utc.hour)
            if 0 <= h <= 7:
                nocturno += v
            elif 8 <= h <= 23:
                diurno += v
        except (ValueError, TypeError):
            continue
    return nocturno, diurno, True


def _calculate_nocturnal_metrics_chile(
    node_id: str,
    start_date: datetime,
    end_date: datetime,
) -> dict:
    dias_con_consumo_nocturno = 0
    dias_sin_consumo_nocturno = 0
    dias_sin_datos_horarios = 0
    consumo_nocturno_total = 0.0
    consumo_diurno_efectivo = 0.0
    current_date = start_date.date()
    end_date_only = end_date.date()
    while current_date <= end_date_only:
        try:
            hourly_data = get_hourly_measures_for_day(node_id, datetime.combine(current_date, datetime.min.time()))
            if hourly_data:
                consumo_nocturno_dia = 0.0
                tiene_consumo_nocturno = False
                consumo_diurno_dia = 0.0
                for hour, value in hourly_data:
                    if 0 <= hour <= 6:
                        consumo_nocturno_dia += value
                        if value > 0:
                            tiene_consumo_nocturno = True
                    elif 7 <= hour <= 23:
                        consumo_diurno_dia += value
                consumo_nocturno_total += consumo_nocturno_dia
                consumo_diurno_efectivo += consumo_diurno_dia
                if tiene_consumo_nocturno:
                    dias_con_consumo_nocturno += 1
                else:
                    dias_sin_consumo_nocturno += 1
            else:
                dias_sin_datos_horarios += 1
        except Exception:
            dias_sin_datos_horarios += 1
        current_date += timedelta(days=1)
    dias_con_datos_horarios = dias_con_consumo_nocturno + dias_sin_consumo_nocturno
    return {
        "dias_con_consumo_nocturno": dias_con_consumo_nocturno,
        "dias_sin_consumo_nocturno": dias_sin_consumo_nocturno,
        "dias_sin_datos_horarios": dias_sin_datos_horarios,
        "dias_con_datos_horarios": dias_con_datos_horarios,
        "consumo_nocturno_total": consumo_nocturno_total,
        "consumo_diurno_efectivo": consumo_diurno_efectivo,
    }


def _calculate_nocturnal_metrics_colegios_utc(
    node_id: str,
    start_date: datetime,
    end_date: datetime,
) -> dict:
    dias_con_consumo_nocturno = 0
    dias_sin_consumo_nocturno = 0
    dias_sin_datos_horarios = 0
    consumo_nocturno_total = 0.0
    consumo_diurno_efectivo = 0.0
    current_date = start_date.date()
    end_date_only = end_date.date()
    while current_date <= end_date_only:
        noct, diu, ok = _nocturno_diurno_dia_colegios_utc(node_id, current_date)
        if ok:
            consumo_nocturno_total += noct
            consumo_diurno_efectivo += diu
            if noct > 0:
                dias_con_consumo_nocturno += 1
            else:
                dias_sin_consumo_nocturno += 1
        else:
            dias_sin_datos_horarios += 1
        current_date += timedelta(days=1)
    dias_con_datos_horarios = dias_con_consumo_nocturno + dias_sin_consumo_nocturno
    return {
        "dias_con_consumo_nocturno": dias_con_consumo_nocturno,
        "dias_sin_consumo_nocturno": dias_sin_consumo_nocturno,
        "dias_sin_datos_horarios": dias_sin_datos_horarios,
        "dias_con_datos_horarios": dias_con_datos_horarios,
        "consumo_nocturno_total": consumo_nocturno_total,
        "consumo_diurno_efectivo": consumo_diurno_efectivo,
    }


def calculate_nocturnal_metrics(
    node_id: str,
    start_date: datetime,
    end_date: datetime,
    company_id: Optional[str] = None,
) -> dict:
    """
    Colegios (CORMUP / 000008): CSV ``dates.measures.csv`` por día, suma marcas TIME UTC 00:00–07:00.
    Demás clientes: hora Chile 00:00–06:59 vía ``get_hourly_measures_for_day``.
    """
    if es_nodo_colegio(node_id, company_id):
        return _calculate_nocturnal_metrics_colegios_utc(node_id, start_date, end_date)
    return _calculate_nocturnal_metrics_chile(node_id, start_date, end_date)


# Chile 0..6 inclusive (7 h). Colegios: UTC 0..7 inclusive (8 h) — ver HORAS_NOCTURNAS_COLEGIOS_UTC.
HORAS_NOCTURNAS_POR_DIA = 7
HORAS_NOCTURNAS_COLEGIOS_UTC = 8
# Mismo umbral que el análisis de filtración en reportes individuales (sección "Análisis de consumo nocturno...").
# ≥75 % de los días **con datos horarios** con consumo en madrugada (00:00–06:59 Chile).
UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION = 75.0


def proyeccion_filtracion_desde_consumo_nocturno(
    consumo_nocturno_total: float,
    num_dias_periodo: int,
    dias_con_consumo_nocturno: int,
    dias_sin_consumo_nocturno: int,
    umbral_pct: float = UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION,
    horas_nocturnas_por_dia: Optional[int] = None,
) -> Tuple[float, float, float]:
    """
    Proyección de filtración (m³) alineada al consumo nocturno de calculate_nocturnal_metrics.

    Promedio horario = consumo_nocturno_total / (días del periodo × horas nocturnas por día).
    Proyección día = promedio_hora × 24; proyección periodo = proyección día × días del periodo.

    Returns:
        (proyección periodo m³, proyección día m³, promedio hora m³/h)
    """
    if num_dias_periodo < 7:
        return 0.0, 0.0, 0.0
    total_dias = dias_con_consumo_nocturno + dias_sin_consumo_nocturno
    if total_dias <= 0:
        return 0.0, 0.0, 0.0
    pct = (dias_con_consumo_nocturno / total_dias) * 100.0
    if pct < umbral_pct:
        return 0.0, 0.0, 0.0
    h_noche = float(horas_nocturnas_por_dia if horas_nocturnas_por_dia is not None else HORAS_NOCTURNAS_POR_DIA)
    total_horas = float(num_dias_periodo) * h_noche
    if total_horas <= 0 or consumo_nocturno_total <= 0:
        return 0.0, 0.0, 0.0
    promedio_hora = consumo_nocturno_total / total_horas
    proyeccion_dia = promedio_hora * 24.0
    proyeccion_periodo = proyeccion_dia * float(num_dias_periodo)
    return proyeccion_periodo, proyeccion_dia, promedio_hora


def analyze_alert_periodicity(alerts: List[dict], start_date: datetime, end_date: datetime) -> dict:
    """
    Analiza la periodicidad de las alertas de un nodo.
    
    Retorna:
    - total_alertas: Número total de alertas con medida > 0
    - dias_con_alertas: Número de días únicos con alertas
    - periodicidad: Descripción de la periodicidad (diaria, intermitente, puntual, etc.)
    - dias_alertas: Lista de fechas con alertas
    """
    if not alerts:
        return {
            "total_alertas": 0,
            "dias_con_alertas": 0,
            "periodicidad": "Sin alertas",
            "dias_alertas": []
        }
    
    # Filtrar solo alertas con medida mayor a cero
    alerts_con_medida = [a for a in alerts if float(a.get("measure", 0) or 0) > 0]
    
    if not alerts_con_medida:
        return {
            "total_alertas": 0,
            "dias_con_alertas": 0,
            "periodicidad": "Sin alertas",
            "dias_alertas": []
        }
    
    # Extraer fechas únicas de las alertas
    dias_alertas = set()
    for alert in alerts_con_medida:
        raw_date = alert.get("creationDate", "") or ""
        if raw_date:
            try:
                dt = datetime.fromisoformat(raw_date.replace("Z", "+00:00"))
                dias_alertas.add(dt.date())
            except Exception:
                pass
    
    total_alertas = len(alerts_con_medida)
    dias_con_alertas = len(dias_alertas)
    
    # Calcular número total de días en el periodo
    num_dias_periodo = (end_date.date() - start_date.date()).days + 1
    
    # Determinar periodicidad
    if dias_con_alertas == 0:
        periodicidad = "Sin alertas"
    elif dias_con_alertas == num_dias_periodo:
        periodicidad = "Diaria (todos los días del periodo)"
    elif dias_con_alertas >= num_dias_periodo * 0.8:
        periodicidad = "Frecuente (más del 80% de los días)"
    elif dias_con_alertas >= num_dias_periodo * 0.5:
        periodicidad = "Regular (más del 50% de los días)"
    elif dias_con_alertas >= num_dias_periodo * 0.2:
        periodicidad = "Intermitente (más del 20% de los días)"
    else:
        periodicidad = "Puntual (menos del 20% de los días)"
    
    return {
        "total_alertas": total_alertas,
        "dias_con_alertas": dias_con_alertas,
        "periodicidad": periodicidad,
        "dias_alertas": sorted(list(dias_alertas))
    }


def build_nocturnal_pie_chart(
    consumo_nocturno: float,
    consumo_diurno: float,
    output: Path,
) -> Optional[Path]:
    """
    Genera una gráfica de torta mostrando consumo nocturno vs consumo diurno efectivo.
    """
    if consumo_nocturno == 0 and consumo_diurno == 0:
        return None
    
    try:
        import matplotlib.pyplot as plt
        
        # Calcular porcentajes
        total = consumo_nocturno + consumo_diurno
        if total == 0:
            return None
        
        pct_nocturno = (consumo_nocturno / total) * 100
        pct_diurno = (consumo_diurno / total) * 100
        
        # Crear la gráfica (figura amplia para que al insertar en Word no quede pequeña)
        fig, ax = plt.subplots(figsize=(10, 7.5))
        
        # Datos para la gráfica
        labels = [
            f"Consumo nocturno\n{format_number_chilean(consumo_nocturno, 1)} m³\n({pct_nocturno:.1f}%)",
            f"Consumo diurno efectivo\n{format_number_chilean(consumo_diurno, 1)} m³\n({pct_diurno:.1f}%)"
        ]
        sizes = [consumo_nocturno, consumo_diurno]
        colors = ['#FF8C00', '#0050b3']  # Naranja para nocturno, azul para diurno
        explode = (0.05, 0)  # Separar ligeramente la primera porción
        
        # Crear la gráfica de torta
        wedges, texts, autotexts = ax.pie(
            sizes,
            explode=explode,
            labels=labels,
            colors=colors,
            autopct='',
            shadow=True,
            startangle=90,
            textprops={'fontsize': 11, 'fontweight': 'bold'}
        )
        
        # Ajustar el formato de los textos
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(10)
        
        ax.set_title(
            "Distribución de consumo: Nocturno vs Diurno Efectivo",
            fontsize=14,
            fontweight='bold',
            pad=20
        )
        
        plt.tight_layout()
        plt.savefig(str(output), dpi=300, bbox_inches='tight', pad_inches=0.2)
        plt.close()
        
        return output
    except Exception as e:
        print(f"DEBUG: Error al generar gráfica de torta: {e}")
        import traceback
        traceback.print_exc()
        return None


def build_hourly_consumption_line_chart(
    hourly_data: List[tuple],
    output: Path,
    day_date: Optional[datetime] = None,
    title_suffix: str = "",
    alert_value: Optional[float] = None,
    alert_hour: Optional[int] = None,
    alert_datetime: Optional[datetime] = None,
) -> Optional[Path]:
    """Grafica consumo por hora como línea simple (sin área sombreada ni alertas).
    
    Si se proporciona alert_value y alert_hour, marca ese punto en lugar del máximo del día.
    """
    if not hourly_data:
        return None
    
    hours = [h for h, _ in hourly_data]
    values = [v for _, v in hourly_data]
    
    fig, ax = plt.subplots(figsize=(10, 5))
    
    # Graficar línea simple
    ax.plot(hours, values, linestyle="-", color="#4A90E2", linewidth=2, marker="o", markersize=4)
    
    # Rellenar el área bajo la línea con color azul claro (igual que la primera gráfica)
    ax.fill_between(hours, values, 0, color="#4A90E2", alpha=0.3)
    
    # SIEMPRE encontrar y marcar el valor máximo del día (sin condiciones)
    # Filtrar valores que sean mayores a 0 para encontrar el máximo real
    # Si todos son 0, no marcar ningún punto
    non_zero_values = [(i, v) for i, v in enumerate(values) if v > 0]
    
    if non_zero_values:
        # Encontrar el índice y valor del máximo entre los valores no cero
        max_index, max_value = max(non_zero_values, key=lambda x: x[1])
        max_hour = hours[max_index]
        marked_hour = max_hour
        marked_value = max_value
        
        # Guardar información del punto máximo para marcarlo después de establecer límites
        max_point_info = {
            'hour': marked_hour,
            'value': marked_value,
            'text': f"{format_number_chilean(marked_value, 2)} m³/hr\n{marked_hour:02d}:00"
        }
    else:
        max_point_info = None
    # Si todos los valores son 0, no marcar ningún punto
    
    # Título con fecha si está disponible
    if day_date:
        month_names_es = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                          "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        month_name = month_names_es[day_date.month - 1]
        title = f"Consumo por hora - {day_date.day} {month_name} {day_date.year}"
        if title_suffix:
            title += f" ({title_suffix})"
    else:
        title = "Consumo por hora"
        if title_suffix:
            title += f" ({title_suffix})"
    
    ax.set_title(title, fontsize=12, fontweight="bold", pad=15)
    ax.set_xlabel("Hora del día", fontsize=10)
    ax.set_ylabel("Consumo (m³/hr)", fontsize=10)
    
    # Verificar si hay consumo real (valores > 0) en el período nocturno (00:00 a 06:00)
    # Solo destacar si hay consumo real en ese período
    tiene_consumo_nocturno = False
    for hour, value in zip(hours, values):
        if 0 <= hour <= 6 and value > 0:
            tiene_consumo_nocturno = True
            break
    
    # Destacar período nocturno (00:00 a 06:00) solo si hay consumo real
    if tiene_consumo_nocturno:
        y_max = max(values) if values else 10
        ax.axvline(x=0, color='red', linestyle='--', linewidth=2, alpha=0.7, label='Período nocturno')
        ax.axvline(x=6, color='red', linestyle='--', linewidth=2, alpha=0.7)
        # Agregar área sombreada para el período nocturno
        ax.axvspan(0, 6, alpha=0.1, color='red', zorder=0)
    
    # Forzar que el eje Y empiece en 0 y no muestre números negativos
    ax.set_ylim(bottom=0)
    yticks = ax.get_yticks()
    yticks = yticks[yticks >= 0]
    ax.set_yticks(yticks)
    
    # Ahora marcar el punto máximo y agregar anotación (después de establecer límites)
    if max_point_info:
        # Marcar el punto con un círculo rojo más pequeño para evitar traslape con anotación
        ax.plot(max_point_info['hour'], max_point_info['value'], 'ro', markersize=8, markeredgecolor='darkred', 
                markeredgewidth=2, zorder=10)
        
        # Calcular posición de la anotación considerando los límites del eje Y
        y_min, y_max_axis = ax.get_ylim()
        y_range = y_max_axis - y_min
        
        # Calcular posición de la anotación (arriba del punto, pero dentro del área visible)
        # Usar un offset relativo al rango del eje Y, pero asegurarse de que no exceda el límite superior
        y_offset = min(y_range * 0.12, y_max_axis * 0.12)  # 12% del rango o del máximo, el menor
        annotation_y = max_point_info['value'] + y_offset
        
        # Asegurarse de que la anotación no exceda el 90% del límite superior del eje Y
        max_annotation_y = y_max_axis * 0.90
        if annotation_y > max_annotation_y:
            annotation_y = max_annotation_y
        
        # Agregar anotación con el valor del máximo consumo
        ax.annotate(
            max_point_info['text'],
            xy=(max_point_info['hour'], max_point_info['value']),
            xytext=(max_point_info['hour'], annotation_y),
            ha='center', va='bottom',
            fontsize=11, fontweight='bold', color='red',
            bbox=dict(boxstyle='round,pad=0.6', facecolor='yellow', alpha=0.8, edgecolor='red', linewidth=2.5),
            arrowprops=dict(arrowstyle='->', color='red', lw=2.5, connectionstyle='arc3,rad=0')
        )
    
    # Formatear el eje X
    if day_date:
        x_labels = []
        x_positions = []
        month_abbr = ["Ene.", "Feb.", "Mar.", "Abr.", "May.", "Jun.",
                     "Jul.", "Ago.", "Sep.", "Oct.", "Nov.", "Dic."]
        x_labels.append(f"{day_date.day}. {month_abbr[day_date.month - 1]}")
        x_positions.append(0)
        for hour in [6, 12, 18]:
            x_labels.append(f"{hour:02d}:00")
            x_positions.append(hour)
        ax.set_xticks(x_positions)
        ax.set_xticklabels(x_labels, fontsize=9)
    else:
        ax.set_xticks(range(0, 24, 6))
        ax.set_xticklabels([f"{h:02d}:00" for h in range(0, 24, 6)], fontsize=9)
    
    # Grid sutil
    ax.grid(True, linestyle="--", alpha=0.3, linewidth=0.5)
    
    plt.tight_layout()
    plt.savefig(output, dpi=150, bbox_inches='tight')
    plt.close()
    return output


def build_hourly_consumption_chart(
    hourly_data: List[tuple],
    alert_average: float,
    output: Path,
    day_date: Optional[datetime] = None,
    alerts: List[dict] = None,
) -> Optional[Path]:
    """Grafica consumo por hora con área sombreada, marcando las horas con alertas."""
    if not hourly_data:
        return None
    
    
    hours = [h for h, _ in hourly_data]
    values = [v for _, v in hourly_data]
    
    fig, ax = plt.subplots(figsize=(10, 5))
    
    # Graficar línea con marcadores y área sombreada (sin marcadores por defecto)
    ax.plot(hours, values, linestyle="-", color="#4A90E2", label="Consumo por hora", linewidth=2)
    
    # Rellenar el área bajo la línea con color azul claro
    ax.fill_between(hours, values, 0, color="#4A90E2", alpha=0.3)
    
    # Identificar horas con alertas en el día graficado
    alert_hours = set()
    if day_date and alerts:
        target_date = day_date.date()
        for alert in alerts:
            raw_date = alert.get("creationDate", "") or ""
            if raw_date:
                try:
                    dt = datetime.fromisoformat(raw_date.replace("Z", "+00:00"))
                    alert_date = dt.date()
                    if alert_date == target_date:
                        # Obtener la hora de la alerta
                        alert_hour = dt.hour
                        alert_hours.add(alert_hour)
                except (ValueError, TypeError):
                    continue
    
    # Marcar las horas con alertas con círculos rojos
    if alert_hours:
        alert_hours_list = sorted(list(alert_hours))
        alert_values = []
        alert_hours_plot = []
        for h in alert_hours_list:
            # Buscar el valor correspondiente a esa hora
            for hour, value in hourly_data:
                if hour == h:
                    alert_hours_plot.append(hour)
                    alert_values.append(value)
                    break
        
        if alert_hours_plot:
            ax.scatter(alert_hours_plot, alert_values, color="red", s=100, 
                      zorder=5, edgecolors="darkred", linewidths=2, label="Alertas")
    
    # Título con fecha si está disponible
    if day_date:
        # Formato: "26 noviembre 2025 | Escala representada en m3/hr"
        month_names_es = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                          "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        month_name = month_names_es[day_date.month - 1]
        title = f"{day_date.day} {month_name} {day_date.year} | Escala representada en m³/hr"
    else:
        title = "Consumo por hora | Escala representada en m³/hr"
    ax.set_title(title, fontsize=12, fontweight="bold", pad=15)
    ax.set_xlabel("Hora del día", fontsize=10)
    ax.set_ylabel("Consumo (m³/hr)", fontsize=10)
    
    # Verificar si hay consumo real (valores > 0) en el período nocturno (00:00 a 06:00)
    # Solo destacar si hay consumo real en ese período
    tiene_consumo_nocturno = False
    for hour, value in zip(hours, values):
        if 0 <= hour <= 6 and value > 0:
            tiene_consumo_nocturno = True
            break
    
    # Destacar período nocturno (00:00 a 06:00) solo si hay consumo real
    if tiene_consumo_nocturno:
        y_max = max(values) if values else 10
        ax.axvline(x=0, color='red', linestyle='--', linewidth=2, alpha=0.7, label='Período nocturno')
        ax.axvline(x=6, color='red', linestyle='--', linewidth=2, alpha=0.7)
        # Agregar área sombreada para el período nocturno
        ax.axvspan(0, 6, alpha=0.1, color='red', zorder=0)
    
    # Forzar que el eje Y empiece en 0 y no muestre números negativos
    ax.set_ylim(bottom=0)
    yticks = ax.get_yticks()
    yticks = yticks[yticks >= 0]
    ax.set_yticks(yticks)
    
    # Formatear el eje X con etiquetas mejor espaciadas
    # Mostrar: inicio del día (26. Nov.), 06:00, 12:00, 18:00
    if day_date:
        x_labels = []
        x_positions = []
        # Etiqueta inicial con fecha - formato "26. Nov."
        month_abbr = ["Ene.", "Feb.", "Mar.", "Abr.", "May.", "Jun.",
                     "Jul.", "Ago.", "Sep.", "Oct.", "Nov.", "Dic."]
        x_labels.append(f"{day_date.day}. {month_abbr[day_date.month - 1]}")
        x_positions.append(0)
        # Horas principales
        for hour in [6, 12, 18]:
            x_labels.append(f"{hour:02d}:00")
            x_positions.append(hour)
        ax.set_xticks(x_positions)
        ax.set_xticklabels(x_labels, fontsize=9)
    else:
        ax.set_xticks(range(0, 24, 6))
        ax.set_xticklabels([f"{h:02d}:00" for h in range(0, 24, 6)], fontsize=9)
    
    # Grid sutil
    ax.grid(True, linestyle="--", alpha=0.3, linewidth=0.5)
    
    # Mover leyenda abajo (solo si hay alertas marcadas)
    if alert_hours:
        ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.15), ncol=2, fontsize=9)
    else:
        ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.15), ncol=1, fontsize=9)
    
    plt.tight_layout()
    plt.savefig(output, dpi=150, bbox_inches='tight')
    plt.close()
    return output


def build_monthly_comparison_chart(
    leak_monthly: float,
    effective_consumption_monthly: float,
    price_per_m3_clp: float,
    output: Path,
) -> Optional[Path]:
    """
    Genera gráfica de anillo comparando consumo nocturno mensual (naranja) 
    vs consumo efectivo mensual (azul).
    Muestra valores en m³, $CLP y porcentajes en el centro del anillo.
    """
    # Permitir valores pequeños para casos con 3+ alertas
    if leak_monthly < 0 or effective_consumption_monthly < 0:
        return None

    leak = max(leak_monthly, 0.0)
    efectivo = max(effective_consumption_monthly, 0.0)

    if leak == 0 and efectivo == 0:
        return None

    # Calcular valores en CLP
    leak_value_clp = leak * price_per_m3_clp
    efectivo_value_clp = efectivo * price_per_m3_clp
    
    # Calcular porcentajes
    total = leak + efectivo
    leak_pct = (leak / total * 100) if total > 0 else 0
    efectivo_pct = (efectivo / total * 100) if total > 0 else 0

    fig, ax = plt.subplots(figsize=(6, 4))
    
    # Datos para el gráfico de anillo
    sizes = [efectivo, leak]
    colors = ["#0050b3", "#FF8C00"]  # Azul para consumo efectivo, naranja para consumo nocturno
    labels = ["Consumo efectivo mensual", "Consumo nocturno mensual"]
    
    # Crear gráfico de anillo (donut chart)
    # wedgeprops crea el agujero en el centro
    wedges, texts, autotexts = plt.pie(
        sizes,
        labels=None,  # No mostrar etiquetas alrededor del gráfico
        autopct='',  # No mostrar porcentajes automáticos
        startangle=90,
        colors=colors,
        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2),  # width=0.5 crea el efecto de anillo
        textprops={"fontsize": 10},
        pctdistance=0.85,  # Distancia para porcentajes (no usado pero necesario)
    )
    
    # Textos para cada segmento
    segment_texts = [
        (
            f"Consumo efectivo:\n"
            f"{format_number_chilean(efectivo, 1)} m³\n"
            f"{format_currency_chilean(efectivo_value_clp)}\n"
            f"({format_number_chilean(efectivo_pct, 1)}%)"
        ),
        (
            f"Consumo nocturno:\n"
            f"{format_number_chilean(leak, 1)} m³\n"
            f"{format_currency_chilean(leak_value_clp)}\n"
            f"({format_number_chilean(leak_pct, 1)}%)"
        ),
    ]
    
    # Calcular ángulos directamente desde los tamaños
    # startangle=90 significa que empezamos desde arriba (90°)
    # Los segmentos se dibujan en sentido horario
    total_size = sum(sizes)
    start_angle = 90  # startangle del pie chart
    
    # Agregar texto en el exterior de cada segmento
    current_angle = start_angle
    for i, (size, text) in enumerate(zip(sizes, segment_texts)):
        # Calcular el ángulo medio del segmento
        angle_span = (size / total_size) * 360  # Espacio angular del segmento en grados
        mid_angle_deg = current_angle - angle_span / 2  # Ángulo medio (sentido horario desde startangle)
        current_angle -= angle_span  # Mover al siguiente segmento
        
        # Convertir a coordenadas estándar
        # En matplotlib pie con startangle=90, los ángulos van en sentido horario desde arriba (90°)
        # En coordenadas estándar (trigonométricas), 0° está a la derecha y va en sentido antihorario
        # La conversión correcta:
        # - En pie: 90° = arriba, sentido horario
        # - En estándar: 90° = arriba, sentido antihorario
        # Si mid_angle_deg está medido desde 90° en sentido horario,
        # para convertir a estándar (desde 0° en sentido antihorario):
        # estándar = 90° - (mid_angle_deg - 90°) = 180° - mid_angle_deg
        # Pero esto da: si mid_angle_deg = 90°, estándar = 90° ✓ (correcto)
        # Si mid_angle_deg = 0°, estándar = 180° (izquierda) - esto puede ser correcto dependiendo del contexto
        # Probemos con: estándar = 90° - mid_angle_deg + 90° = 180° - mid_angle_deg
        standard_angle_deg = 180 - mid_angle_deg
        
        # Convertir a radianes
        mid_angle_rad = np.radians(standard_angle_deg)
        
        # Radio exterior del anillo (aumentado para alejar los indicadores)
        radius = 1.4
        
        # Calcular posición en el exterior del anillo
        x = radius * np.cos(mid_angle_rad)
        y = radius * np.sin(mid_angle_rad)
        
        # Si el indicador está en la parte superior (entre 60° y 120° o entre 240° y 300°),
        # agregar un offset vertical adicional para evitar el título
        # Convertir el ángulo estándar a un rango de 0-360
        angle_normalized = standard_angle_deg % 360
        if (angle_normalized >= 60 and angle_normalized <= 120) or (angle_normalized >= 240 and angle_normalized <= 300):
            # Está en la parte superior, mover más arriba
            y += 0.15
        
        # Agregar texto con fondo blanco para mejor legibilidad
        ax.text(x, y, text, ha='center', va='center',
                fontsize=9, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.5', facecolor='white', alpha=0.9, edgecolor='gray', linewidth=1))
    
    # Aumentar el padding superior del título para evitar superposición con indicadores
    ax.set_title("Comparación mensual: Consumo nocturno vs Consumo efectivo", fontsize=12, fontweight="bold", pad=25)
    
    # Mover leyenda abajo de la gráfica
    ax.legend(wedges, labels, loc="upper center", bbox_to_anchor=(0.5, -0.1), ncol=2, fontsize=10)
    
    # Ajustar límites del eje para dar más espacio a los indicadores
    ax.set_xlim(-1.8, 1.8)
    ax.set_ylim(-1.8, 1.8)
    
    # Asegurar que el layout tenga suficiente espacio superior
    plt.tight_layout(rect=[0, 0, 1, 0.95])  # Dejar 5% de espacio arriba
    plt.savefig(output, dpi=150, bbox_inches='tight', pad_inches=0.3)
    plt.close()
    return output


def build_filtracion_anillo_chart(
    proyeccion_filtracion: float,
    consumo_efectivo: float,
    output: Path,
) -> Optional[Path]:
    """
    Genera gráfica de anillo comparando proyección de consumo nocturno del periodo vs consumo efectivo del periodo.
    Muestra valores en m³ y porcentajes.
    """
    if proyeccion_filtracion < 0 or consumo_efectivo < 0:
        return None

    filtracion = max(proyeccion_filtracion, 0.0)
    efectivo = max(consumo_efectivo, 0.0)

    if filtracion == 0 and efectivo == 0:
        return None

    # Calcular porcentajes
    total = filtracion + efectivo
    filtracion_pct = (filtracion / total * 100) if total > 0 else 0
    efectivo_pct = (efectivo / total * 100) if total > 0 else 0

    fig, ax = plt.subplots(figsize=(10, 8))
    
    # Datos para el gráfico de anillo
    sizes = [efectivo, filtracion]
    colors = ["#0050b3", "#FF8C00"]  # Azul para consumo efectivo, naranja para consumo nocturno
    labels = ["Consumo efectivo del periodo", "Proyección consumo nocturno del periodo"]
    
    # Crear gráfico de anillo (donut chart)
    wedges, texts, autotexts = plt.pie(
        sizes,
        labels=None,
        autopct='',
        startangle=90,
        colors=colors,
        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2),
        textprops={"fontsize": 10},
        pctdistance=0.85,
    )
    
    # Textos para cada segmento
    segment_texts = [
        (
            f"Consumo efectivo:\n"
            f"{format_number_chilean(efectivo, 1)} m³\n"
            f"({format_number_chilean(efectivo_pct, 1)}%)"
        ),
        (
            f"Proyección consumo nocturno:\n"
            f"{format_number_chilean(filtracion, 1)} m³\n"
            f"({format_number_chilean(filtracion_pct, 1)}%)"
        ),
    ]
    
    # Calcular ángulos y posicionar textos
    total_size = sum(sizes)
    start_angle = 90
    current_angle = start_angle
    
    for i, (size, text) in enumerate(zip(sizes, segment_texts)):
        angle_span = (size / total_size) * 360
        mid_angle_deg = current_angle - angle_span / 2
        current_angle -= angle_span
        
        standard_angle_deg = 180 - mid_angle_deg
        mid_angle_rad = np.radians(standard_angle_deg)
        
        radius = 1.4
        x = radius * np.cos(mid_angle_rad)
        y = radius * np.sin(mid_angle_rad)
        
        angle_normalized = standard_angle_deg % 360
        if (angle_normalized >= 60 and angle_normalized <= 120) or (angle_normalized >= 240 and angle_normalized <= 300):
            y += 0.15
        
        ax.text(x, y, text, ha='center', va='center',
                fontsize=9, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.5', facecolor='white', alpha=0.9, edgecolor='gray', linewidth=1))
    
    # Título muy cerca de la gráfica
    ax.set_title("Distribución: Proyección filtración vs Consumo efectivo", fontsize=14, fontweight="bold", pad=5)
    
    # Leyenda muy cerca pero sin traslaparse
    ax.legend(wedges, labels, loc="upper center", bbox_to_anchor=(0.5, 0.02), ncol=1, fontsize=10, framealpha=0.9)
    
    ax.set_xlim(-1.8, 1.8)
    ax.set_ylim(-1.8, 1.8)
    
    # Ajustar layout para que todo quede más compacto y pegado
    plt.tight_layout(rect=[0, 0.08, 1, 0.99])
    plt.savefig(output, dpi=150, bbox_inches='tight', pad_inches=0.3)
    plt.close()
    return output


def build_alert_pie_chart(
    projected_leak_daily: float,
    effective_consumption: float,
    output: Path,
    price_per_m3_clp: float = 1200.0,
) -> Optional[Path]:
    """
    Genera gráfica de torta entre:
    - Proyección diaria de consumo nocturno (m³/día)
    - Consumo efectivo promedio (m³/día)
    
    Muestra valores en m³ y valorización en $.
    Genera la gráfica incluso si la proyección es pequeña (>= 0.1) para casos con 3+ alertas.
    """
    # No generar gráfica solo si la proyección de consumo nocturno es estrictamente cero o negativa
    # Permitir valores >= 0.01 para casos con 3+ alertas
    if projected_leak_daily <= 0:
        return None

    fuga_proj = max(projected_leak_daily, 0.0)
    efectivo = max(effective_consumption, 0.0)

    if fuga_proj == 0 and efectivo == 0:
        return None

    labels = [
        "Proyección diaria de consumo nocturno",
        "Consumo efectivo promedio",
    ]
    sizes = [fuga_proj, efectivo]
    # Colores: consumo nocturno en naranja, consumo efectivo en azul
    colors = ["#FF8C00", "#0050b3"]

    total = sum(sizes)
    
    # Calcular valores mensuales para mostrar en $
    fuga_mensual = fuga_proj * 30.0
    efectivo_mensual = efectivo * 30.0
    fuga_value_clp = fuga_mensual * price_per_m3_clp
    efectivo_value_clp = efectivo_mensual * price_per_m3_clp

    def autopct_fmt(pct: float) -> str:
        valor = total * pct / 100.0
        # Calcular valor en $ basado en el valor real (m³) del segmento
        # Proyectar a mensual y multiplicar por precio
        valor_mensual = valor * 30.0
        valor_clp = valor_mensual * price_per_m3_clp
        return f"{format_number_chilean(pct, 1)}%\n({format_number_chilean(valor, 1)} m³)\n{format_currency_chilean(valor_clp)}"

    plt.figure(figsize=(5, 4))
    wedges, _texts, autotexts = plt.pie(
        sizes,
        labels=None,  # evitamos que el texto se corte alrededor del gráfico
        autopct=autopct_fmt,
        startangle=90,
        textprops={"fontsize": 8},
        shadow=False,  # gráfica normal 2D
        colors=colors,
    )

    # Ajustar tamaño de las letras de valores
    for t in autotexts:
        t.set_fontsize(8)

    # Agregar leyenda debajo de la gráfica para que los textos no se corten
    plt.legend(
        wedges,
        labels,
        loc="upper center",
        bbox_to_anchor=(0.5, -0.05),
        ncol=2,
        fontsize=8,
    )

    # Aumentar el padding superior del título para evitar superposición con valores automáticos
    plt.title("Proyección de consumo nocturno vs consumo efectivo (m³)", fontsize=11, pad=25)
    
    # Asegurar que el layout tenga suficiente espacio superior
    plt.tight_layout(rect=[0, 0, 1, 0.95])  # Dejar 5% de espacio arriba
    plt.savefig(output, dpi=150, bbox_inches='tight', pad_inches=0.3)
    plt.close()
    return output


def add_alerts_section(
    doc: Document,
    chart_path: Optional[Path],
    pie_chart_path: Optional[Path],
    hourly_chart_path: Optional[Path],
    monthly_chart_path: Optional[Path],
    alerts: List[dict],
    alert_stats: dict,
    avg_daily_consumption: float,
    effective_consumption: float,
    leak_monthly: float,
    effective_consumption_monthly: float,
    leak_value_clp: float,
    node_id: str,
    output_dir: Optional[Path] = None,
    se_recalculo_proyeccion: bool = False,  # Parámetro obsoleto, mantenido por compatibilidad
    start_date: Optional[datetime] = None,
    end_date: Optional[datetime] = None,
    company_id: Optional[str] = None,
    price_per_m3_clp: float = 1200.0,
) -> None:
    # Calcular métricas de consumo nocturno y diurno SIEMPRE una sola vez, independientemente de si hay alertas
    nocturnal_metrics = None
    if start_date and end_date:
        print("Calculando métricas de consumo nocturno y diurno...")
        nocturnal_metrics = calculate_nocturnal_metrics(node_id, start_date, end_date, company_id=company_id)
        _lbl_dias_noct = (
            "Número de días con consumo nocturno (00:00–07:00)"
            if es_nodo_colegio(node_id, company_id)
            else "Número de días con consumo nocturno (00:00–06:59, hora Chile)"
        )
        metric_rows = [
            (_lbl_dias_noct, str(nocturnal_metrics["dias_con_consumo_nocturno"])),
            (
                "Número de días sin consumo nocturno (con serie horaria ese día)",
                str(nocturnal_metrics["dias_sin_consumo_nocturno"]),
            ),
        ]
        if int(nocturnal_metrics.get("dias_sin_datos_horarios", 0) or 0) > 0:
            metric_rows.append(
                ("Días sin datos horarios (sin serie ese día)", str(nocturnal_metrics["dias_sin_datos_horarios"]))
            )
        metric_rows.extend(
            [
                ("Consumo nocturno (m³)", format_number_chilean(nocturnal_metrics["consumo_nocturno_total"], 1)),
                ("Consumo diurno efectivo (m³)", format_number_chilean(nocturnal_metrics["consumo_diurno_efectivo"], 1)),
            ]
        )
        add_table(doc, "Métricas de consumo nocturno", metric_rows)
        
        # Gráfica de torta: omitir cuando el consumo nocturno es 0 (por control CPA/WES no agrega valor).
        consumo_nocturno_total = float(nocturnal_metrics.get("consumo_nocturno_total", 0.0) or 0.0)
        dias_con_nocturno = int(nocturnal_metrics.get("dias_con_consumo_nocturno", 0) or 0)

        if consumo_nocturno_total <= 0.0 or dias_con_nocturno == 0:
            doc.add_paragraph("")
            nota_para = doc.add_paragraph(
                "Nota: En este punto el consumo nocturno aparece en 0 debido al control operativo de la máquina WES (CPA). "
                "Por este motivo se omite la gráfica de torta (no aporta valor adicional en la interpretación)."
            )
            nota_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in nota_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            if output_dir:
                pie_chart_nocturnal_path = output_dir / "consumo_nocturno_vs_diurno.png"
                pie_chart_nocturnal = build_nocturnal_pie_chart(
                    nocturnal_metrics["consumo_nocturno_total"],
                    nocturnal_metrics["consumo_diurno_efectivo"],
                    pie_chart_nocturnal_path,
                )
                
                if pie_chart_nocturnal and pie_chart_nocturnal.exists():
                    doc.add_paragraph("")
                    add_formatted_title(doc, "Distribución de consumo: Nocturno vs Diurno Efectivo:")
                    add_picture_with_pagination(doc, str(pie_chart_nocturnal), Inches(6.5), keep_with_next=True)
                    
                    # Explicación de cómo se elaboró la gráfica de torta
                    explanation_para = doc.add_paragraph(
                        f"La gráfica de torta presentada muestra la distribución del consumo de agua entre el consumo nocturno y el consumo diurno efectivo. "
                        f"El consumo nocturno se calcula sumando todos los valores de consumo registrados entre las 00:00 y las 06:00 horas de todos los días del periodo analizado. "
                        f"El consumo diurno efectivo se calcula sumando todos los valores de consumo registrados entre las 06:00 y las 23:00 horas de todos los días del periodo analizado. "
                        f"Los porcentajes y valores mostrados en cada segmento de la gráfica representan la proporción y el volumen total (en m³) de cada tipo de consumo respecto al consumo total del periodo."
                    )
                    explanation_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    for run in explanation_para.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
    
    # Si no hay alertas, mostrar mensaje
    if not alerts:
        doc.add_paragraph("No se gatillaron alertas en el periodo analizado.")
    
    # Si hay alertas, mostrar información adicional sobre alertas
    node_name = get_node_name(node_id)
    if alerts:
        intro_para = doc.add_paragraph(
            f"Los eventos de consumo nocturno provienen del punto {node_name} "
            "y se utilizan para estimar el promedio diario de consumo nocturno y su proyección."
        )
        intro_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in intro_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Explicación del consumo efectivo
        efectivo_para = doc.add_paragraph(
            f"Según el algoritmo implementado en el sistema WES de monitoreo en terreno, los valores presentados son los siguientes:"
        )
        efectivo_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in efectivo_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        if chart_path:
            add_picture_with_pagination(doc, str(chart_path), Inches(6), keep_with_next=True)

    # NO generar gráfica de proyección basada en alertas
    # La proyección de filtración se basa en consumo nocturno (no en alertas)
    # Solo se mantiene la gráfica de torta "Distribución de consumo: Nocturno vs Diurno Efectivo"
    
    # Gráfica de comparación mensual eliminada según solicitud del usuario
    
    # Párrafo antiguo eliminado según solicitud del usuario
    
    # Tabla de eventos registrados eliminada según solicitud del usuario
    
    # Agregar gráfica del día con mayor consumo nocturno
    if start_date and end_date and output_dir:
        max_nocturnal_result = find_max_nocturnal_consumption_day(node_id, None, start_date, end_date)
        if max_nocturnal_result:
            try:
                target_dt, hourly_data = max_nocturnal_result
                dt = target_dt
                
                if hourly_data:
                    # Generar gráfica del día con mayor consumo nocturno
                    max_nocturnal_chart_path = output_dir / f"chart_max_nocturnal_day_{node_id}.png"
                    max_nocturnal_chart = build_hourly_consumption_line_chart(
                        hourly_data,
                        max_nocturnal_chart_path,
                        target_dt,
                        f"Día con mayor consumo nocturno ({dt.strftime('%d-%m-%y')})"
                    )
                    
                    if max_nocturnal_chart and max_nocturnal_chart.exists():
                        # No agregar salto de página, dejar que la paginación automática lo maneje
                        doc.add_paragraph("")  # Espacio antes del título
                        add_formatted_title(doc, f"Día con mayor consumo nocturno ({dt.strftime('%d-%m-%y')}):")
                        add_picture_with_pagination(doc, str(max_nocturnal_chart), Inches(6), keep_with_next=True)
            except Exception as e:
                print(f"DEBUG: Error generando gráfica día con mayor consumo nocturno: {e}")


def generate_report(args: argparse.Namespace) -> Path:
    company_name = get_company_name(args.company_id)
    if is_node_excluded(args.node_id, company_id=args.company_id, company_name=company_name):
        raise ValueError(f"Nodo excluido por configuración: {args.node_id}")

    start_dt = parse_date(args.start_date)
    end_dt = parse_date(args.end_date, end_of_day=True)
    if end_dt < start_dt:
        raise ValueError("La fecha término debe ser mayor o igual a la fecha inicio.")

    measures_error = None

    def _format_ddmmyyyy(dt: datetime) -> str:
        return dt.strftime("%d%m%Y")
    try:
        measures_payload_raw = fetch_json(
            f"{acl_node_base_url()}/nodes/measures/dates",
            params=[
                ("id", args.node_id),
                ("start", _format_ddmmyyyy(start_dt)),
                ("end", _format_ddmmyyyy(end_dt)),
            ],
        )
        measures_payload = normalize_measures_payload(measures_payload_raw, args.node_id)
        measures = flatten_measures(measures_payload)
    except Exception as exc:
        measures_error = f"No se pudo recuperar las medidas: {exc}"
        measures = []
        measures_payload = {}
    summary = summarize_consumption(measures)

    alerts_error = None
    alerts: List[dict] = []

    try:
        alerts_payload = fetch_json(
            f"{acl_node_base_url()}/nodes/myalert/alerts",
            params=[
                ("id", args.node_id),
                ("start", _format_ddmmyyyy(start_dt)),
                ("end", _format_ddmmyyyy(end_dt)),
            ],
        )
        if isinstance(alerts_payload, list):
            alerts = alerts_payload
            # Si el endpoint principal funciona, no intentar el fallback
    except Exception as exc:
        # El endpoint principal falló, intentar el fallback solo una vez
        alerts_error = (
            "No se pudieron recuperar las alertas desde /nodes/myalert/alerts. "
            "Es posible que no haya alertas registradas en el periodo solicitado."
        )
        try:
            # Intentar con formato ddMMyyyy (igual que myalert/alerts)
            fallback_payload = fetch_json(
                f"{acl_node_base_url()}/nodes/leak/alerts",
                params=[
                    ("id", args.node_id),
                    ("start", _format_ddmmyyyy(start_dt)),
                    ("end", _format_ddmmyyyy(end_dt)),
                ],
            )
            if isinstance(fallback_payload, list):
                alerts = fallback_payload
                alerts_error = None  # Si el fallback funciona, limpiar el error
        except Exception:
            # Si el fallback también falla, simplemente dejar alerts vacío
            # El error ya está documentado arriba
            alerts = []

    alert_stats = summarize_alerts(alerts, start_dt, end_dt)
    avg_daily = summary["promedio_diario"]
    # El consumo efectivo debe calcularse restando la proyección diaria (m³/día), no el promedio de alerta (m³/h)
    proyeccion_diaria = alert_stats["proyeccion_24h"]
    effective_consumption = max(avg_daily - proyeccion_diaria, 0.0)
    
    # Calcular proyecciones mensuales
    leak_monthly = alert_stats["proyeccion_24h"] * 30.0  # Proyección diaria * 30 días
    effective_consumption_monthly = effective_consumption * 30.0  # Consumo efectivo diario * 30 días
    
    # Obtener precio del agua desde la API
    price_per_m3_clp = get_water_price_per_m3(args.company_id, args.node_id, measures_payload)
    
    # Valorización en pesos chilenos
    leak_value_clp = leak_monthly * price_per_m3_clp
    
    # Crear estructura de carpetas: empresa/REPORTE/punto/fecha
    # Para Parque Arauco: empresa/mall/REPORTE/punto/fecha
    node_name = get_node_name(args.node_id)
    
    # Limpiar nombres para usarlos como nombres de carpeta (remover caracteres inválidos)
    safe_company_name = "".join(c for c in company_name if c.isalnum() or c in (" ", "-", "_")).strip()
    safe_company_name = safe_company_name.replace(" ", "_")  # Reemplazar espacios con guiones bajos
    
    safe_node_name = "".join(c for c in node_name if c.isalnum() or c in (" ", "-", "_")).strip()
    safe_node_name = safe_node_name.replace(" ", "_")  # Reemplazar espacios con guiones bajos
    
    # Detectar mall si es Parque Arauco
    mall_name = None
    if args.company_id == "000025":  # Parque Arauco
        # Intentar obtener mall_name del argumento si existe
        if hasattr(args, 'mall_name') and args.mall_name:
            mall_name = args.mall_name
        else:
            # Detectar automáticamente el mall del nodo
            mall_name = get_mall_name_for_parque_arauco(args.node_id, node_name)
    
    # Fecha de creación en formato YYYYMMDD_HHMM
    # Verificar si ya existe un reporte para este nodo y periodo
    base_output_dir = Path(args.output_dir)
    company_dir = base_output_dir / safe_company_name
    
    # Si es Parque Arauco y hay mall, crear estructura por mall
    if mall_name and args.company_id == "000025":
        # Limpiar nombre del mall
        safe_mall_name = "".join(c for c in mall_name if c.isalnum() or c in (" ", "-", "_")).strip()
        safe_mall_name = safe_mall_name.replace(" ", "_")
        company_dir = company_dir / safe_mall_name
    
    reporte_dir = company_dir / "REPORTE"
    
    # Formatear fechas para buscar reportes existentes
    start_str = start_dt.strftime("%Y%m%d")
    end_str = end_dt.strftime("%Y%m%d")
    pattern = f"Reporte_{args.company_id}_{args.node_id}_{start_str}_{end_str}.docx"
    
    # Buscar reportes existentes para este nodo y periodo
    existing_reports = []
    if reporte_dir.exists():
        for carpeta in reporte_dir.iterdir():
            if carpeta.is_dir() and carpeta.name.startswith(safe_node_name):
                reporte_file = carpeta / pattern
                if reporte_file.exists():
                    existing_reports.append(carpeta)
    
    # Eliminar reportes duplicados más antiguos y reutilizar el más reciente
    if existing_reports:
        # Ordenar por fecha de modificación (más reciente primero)
        existing_reports.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        # Eliminar todos excepto el más reciente
        for old_folder in existing_reports[1:]:
            try:
                import shutil
                shutil.rmtree(old_folder)
                print(f"Eliminado reporte duplicado antiguo: {old_folder.name}")
            except Exception as e:
                print(f"Advertencia: No se pudo eliminar {old_folder.name}: {e}")
        # Usar la carpeta más reciente (reutilizar)
        output_dir = existing_reports[0]
        folder_name = output_dir.name
        print(f"Reutilizando carpeta existente: {output_dir.name}")
    else:
        # Crear nueva carpeta solo si no existe ninguna
        creation_date = datetime.now(timezone.utc)
        folder_name = f"{safe_node_name}_{creation_date.strftime('%Y%m%d_%H%M')}"
        output_dir = reporte_dir / folder_name
        output_dir.mkdir(parents=True, exist_ok=True)
    
    # Crear también la carpeta ABREGADO para uso futuro (aunque esté vacía por ahora)
    # Si es Parque Arauco y hay mall, crear ABREGADO dentro de la carpeta del mall
    agregado_dir = company_dir / "ABREGADO"
    agregado_dir.mkdir(parents=True, exist_ok=True)

    consumption_chart = (
        build_consumption_chart(measures, output_dir / f"chart_consumo_{args.node_id}.png", start_dt, end_dt, alerts)
        if measures
        else None
    )
    
    # Generar gráfica de promedios por día de la semana (solo si hay 2+ semanas completas)
    weekly_averages_chart = None
    weekly_averages_data = None
    if measures:
        weekly_averages_data = calculate_weekly_averages(measures)
        if weekly_averages_data:
            weekly_averages_chart = build_weekly_averages_chart(
                weekly_averages_data,
                output_dir / f"chart_promedios_semanales_{args.node_id}.png"
    )
    
    # Generar gráficas de consumo por hora para día máximo y mínimo
    max_day_chart = None
    min_day_chart = None
    if measures and summary.get("max") and summary.get("min"):
        try:
            # Día con mayor consumo
            max_measure = summary["max"]
            max_day_dt = datetime.combine(max_measure.date.date(), datetime.min.time())
            max_hourly_data = get_hourly_measures_for_day(args.node_id, max_day_dt)
            if max_hourly_data:
                max_day_chart = build_hourly_consumption_line_chart(
                    max_hourly_data,
                    output_dir / f"chart_hourly_max_{args.node_id}.png",
                    max_day_dt,
                    "Mayor consumo"
                )
        except Exception as e:
            print(f"DEBUG: Error generando gráfica día máximo: {e}")
        
        try:
            # Día con menor consumo
            min_measure = summary["min"]
            min_day_dt = datetime.combine(min_measure.date.date(), datetime.min.time())
            min_hourly_data = get_hourly_measures_for_day(args.node_id, min_day_dt)
            if min_hourly_data:
                min_day_chart = build_hourly_consumption_line_chart(
                    min_hourly_data,
                    output_dir / f"chart_hourly_min_{args.node_id}.png",
                    min_day_dt,
                    "Menor consumo"
                )
        except Exception as e:
            print(f"DEBUG: Error generando gráfica día mínimo: {e}")
    
    alerts_chart = build_leak_chart(alerts, output_dir / f"chart_fugas_{args.node_id}.png") if alerts else None

    # Usar directamente los valores correctos de summarize_alerts
    # que solo considera alertas nocturnas (22:00-07:00) de los últimos 2 días
    hay_3_o_mas_alertas = alert_stats["cantidad"] >= 3
    proyeccion_para_graficas = alert_stats["proyeccion_24h"]
    promedio_alerta = alert_stats["promedio_alerta"]
    
    # Recalcular valores mensuales usando la proyección correcta
    leak_monthly_para_graficas = proyeccion_para_graficas * 30.0
    leak_value_clp_para_graficas = leak_monthly_para_graficas * price_per_m3_clp
    
    # Actualizar también los valores originales para que se muestren en la tabla
    leak_monthly = leak_monthly_para_graficas
    leak_value_clp = leak_value_clp_para_graficas
    
    # NO actualizar alert_stats porque ya tiene los valores correctos de summarize_alerts
    
    # Generar gráficas SOLO si hay 3 o más alertas
    # NO generar gráficas si hay 1 o 2 alertas (incluso si hay proyección)
    debe_generar_graficas = hay_3_o_mas_alertas

    # Gráfica de torta antigua eliminada según solicitud del usuario
    
    # Gráfica de barras mensual antigua eliminada según solicitud del usuario
    monthly_chart = None

    # Gráfica de consumo por hora eliminada según solicitud del usuario

    doc = Document()
    # Agregar logo al encabezado de cada página
    add_logo_to_header(doc)
    
    company_name = get_company_name(args.company_id)
    
    # Formato del título: "Análisis [Empresa] [Nodo]"
    title_base = f"Análisis {company_name} {node_name}"
    
    # Título con todas las letras del mismo tamaño
    title_text = doc.add_paragraph(title_base)
    title_text.style = "Title"
    # Asegurar que todo el texto tenga el mismo tamaño de fuente
    # Reducir tamaño si el título es muy largo para que quepa en una línea
    title_length = len(title_base)
    if title_length > 50:
        font_size = Pt(20)  # Tamaño más pequeño para títulos largos
    elif title_length > 40:
        font_size = Pt(22)  # Tamaño medio
    else:
        font_size = Pt(24)  # Tamaño normal
    
    for run in title_text.runs:
        run.font.size = font_size
    subtitle = doc.add_paragraph(
        f"MONITOREO WES\n"
        f"{node_name}\n"
        f"Rango: {start_dt.strftime('%d-%m-%y')} - {end_dt.strftime('%d-%m-%y')}\n"
        f"Generado: {datetime.now(timezone.utc).strftime('%d-%m-%y')}"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_summary_section(doc, summary, alerts, alert_stats, start_dt, end_dt)
    add_consumption_section(doc, summary, consumption_chart, measures, alerts, max_day_chart, min_day_chart, start_dt, end_dt, weekly_averages_chart, weekly_averages_data)
    # Ya no se usa el flag se_recalculo_proyeccion porque el cálculo correcto
    # ya viene de summarize_alerts (solo alertas nocturnas de últimos 2 días)
    
    add_alerts_section(
        doc, 
        alerts_chart, 
        None,  # pie_chart antigua eliminada
        None,  # hourly_chart eliminado
        monthly_chart,
        alerts, 
        alert_stats, 
        avg_daily, 
        effective_consumption,
        leak_monthly,
        effective_consumption_monthly,
        leak_value_clp,
        args.node_id,
        output_dir,
        se_recalculo_proyeccion=False,  # Ya no se recalcula
        start_date=start_dt,
        end_date=end_dt,
        company_id=args.company_id,
        price_per_m3_clp=price_per_m3_clp,
    )

    # Proyección de consumo nocturno si se cumplen las condiciones
    se_realizo_analisis_filtracion = False
    num_dias_periodo = (end_dt - start_dt).days + 1
    if num_dias_periodo >= 7:
        # Calcular métricas nocturnas para verificar el porcentaje
        nocturnal_metrics_analysis = calculate_nocturnal_metrics(
            args.node_id, start_dt, end_dt, company_id=args.company_id
        )
        dias_con_consumo = nocturnal_metrics_analysis["dias_con_consumo_nocturno"]
        dias_sin_consumo = nocturnal_metrics_analysis["dias_sin_consumo_nocturno"]
        total_dias_analisis = dias_con_consumo + dias_sin_consumo
        
        if total_dias_analisis > 0:
            porcentaje_consumo_nocturno = (dias_con_consumo / total_dias_analisis) * 100
            
            if porcentaje_consumo_nocturno >= UMBRAL_PCT_DIAS_CONSUMO_NOCTURNO_FILTRACION:
                se_realizo_analisis_filtracion = True
                consumo_nocturno_total = nocturnal_metrics_analysis["consumo_nocturno_total"]
                proyeccion_filtracion_periodo, proyeccion_dia_fuga, promedio_hora_consumo_nocturno = (
                    proyeccion_filtracion_desde_consumo_nocturno(
                        consumo_nocturno_total,
                        num_dias_periodo,
                        dias_con_consumo,
                        dias_sin_consumo,
                        horas_nocturnas_por_dia=horas_nocturnas_por_dia_para_nodo(
                            args.node_id, args.company_id
                        ),
                    )
                )
                
                # Obtener consumo total del resumen ejecutivo
                consumo_total_periodo = summary['total']
                consumo_efectivo_periodo = consumo_total_periodo - proyeccion_filtracion_periodo
                
                # Agregar sección de análisis de consumo nocturno y proyección de filtración
                add_formatted_heading(doc, "Análisis de consumo nocturno y proyección de filtración", level=1)
                
                # Agregar narración explicativa
                narracion_para = doc.add_paragraph(
                    f"Producto del análisis de los datos del periodo, se ha identificado que el porcentaje de días con datos horarios "
                    f"que presentan consumo nocturno es igual o superior al 75% ({format_number_chilean(porcentaje_consumo_nocturno, 1)}%) y el periodo del reporte "
                    f"es mayor o igual a 7 días ({num_dias_periodo} días). En base a estas condiciones, se realiza un análisis del consumo nocturno: "
                    f"el consumo nocturno registrado durante el periodo analizado presenta un patrón consistente que permite proyectar el volumen de agua "
                    f"asociado a la ventana nocturna (horas 00:00 a 06:00 inclusive, 7 intervalos horarios por día) a lo largo del periodo del reporte. "
                    f"Esta proyección se fundamenta en el cálculo del promedio horario de consumo nocturno, "
                    f"permitiendo estimar el volumen total de consumo nocturno durante todo el periodo del reporte."
                )
                narracion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in narracion_para.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
                
                doc.add_paragraph("")  # Espacio en blanco
                
                # Crear tabla con los parámetros de proyección de filtración
                add_table(
                    doc,
                    "Proyección de filtración basada en consumo nocturno",
                    [
                        ("Promedio hora consumo nocturno (m³/h)", format_number_chilean(promedio_hora_consumo_nocturno, 2)),
                        ("Proyección día filtración (m³/día)", format_number_chilean(proyeccion_dia_fuga, 1)),
                        ("Proyección filtración del periodo (m³)", format_number_chilean(proyeccion_filtracion_periodo, 1)),
                        ("Consumo efectivo del periodo (m³)", format_number_chilean(consumo_efectivo_periodo, 1)),
                    ],
                )
                
                # NO generar gráfica de anillo de "proyección vs consumo efectivo"
                # Solo se mantiene el análisis y la tabla, sin la gráfica comparativa

    add_formatted_heading(doc, "Conclusiones", level=1)
    
    if se_realizo_analisis_filtracion:
        # Conclusiones con mención de consumo nocturno
        concl_para = doc.add_paragraph(
            "Este reporte sintetiza los principales hallazgos de consumo y alertas detectados durante el periodo analizado. "
            "Se recomienda revisar los días con mayor consumo y atender las alertas registradas. "
        )
        concl_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # Mantener el primer párrafo junto con el título "Conclusiones"
        concl_para.paragraph_format.keep_with_next = True
        concl_para.paragraph_format.widow_control = True
        for run in concl_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Agregar recomendación específica sobre consumo nocturno
        recomendacion_para = doc.add_paragraph(
            "Adicionalmente, producto del análisis de consumo nocturno realizado, se ha identificado un patrón consistente de consumo "
            "durante las horas nocturnas (00:00 a 06:00) que representa una proporción significativa del consumo total del periodo. "
            "Se recomienda revisar el correcto funcionamiento de todos los artefactos conectados a la red de agua monitoreada por WES, "
            "verificando que no presenten fallas, desgaste o mal funcionamiento que puedan generar consumo durante las horas nocturnas. "
            "Asimismo, se sugiere realizar una inspección en el terreno para identificar posibles indicadores de consumo no esperado, "
            "tales como áreas húmedas, charcos persistentes, o cualquier otra señal que pueda indicar la presencia de consumo continuo. "
            "Estas acciones permitirán identificar y optimizar el consumo de agua durante las horas nocturnas."
        )
        recomendacion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in recomendacion_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
    else:
        # Conclusiones estándar sin análisis de filtración
        concl_para = doc.add_paragraph(
            "Este reporte sintetiza los principales hallazgos de consumo y fugas. "
            "Se recomienda revisar los días con mayor consumo y atender las alertas registradas."
        )
        concl_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # Mantener el primer párrafo junto con el título "Conclusiones"
        concl_para.paragraph_format.keep_with_next = True
        concl_para.paragraph_format.widow_control = True
        for run in concl_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro

    # Generar nombre descriptivo del archivo
    company_name_clean = limpiar_nombre_archivo(company_name)
    node_name_clean = limpiar_nombre_archivo(node_name)
    filename = f"Reporte_{company_name_clean}_{node_name_clean}_{start_dt:%Y%m%d}_{end_dt:%Y%m%d}.docx"
    output_path = output_dir / filename
    try:
        doc.save(output_path)
    except PermissionError:
        alt_name = f"{filename.rsplit('.', 1)[0]}_{int(time.time())}.docx"
        output_path = output_dir / alt_name
        doc.save(output_path)
    
    # NO generar PPT individual aquí - solo se genera presentación agregada
    # La presentación agregada se genera desde monitorear_correos_y_generar_reportes.py
    # y contiene todos los nodos seleccionados
    
    return output_path


def generate_aggregated_report(
    company_id: str,
    node_ids: List[str],
    start_date: str,
    end_date: str,
    output_dir: str = "reports",
    fuente_agua_id: Optional[str] = None,
    mall_name: Optional[str] = None,
    apply_exclusions: bool = True,
    generate_ppt: bool = True,
    nota_contexto_periodo: Optional[str] = None,
    parallel_node_fetch: bool = False,
    max_parallel_workers: int = 4,
    company_folder_override: Optional[str] = None,
) -> Path:
    """
    Genera un reporte agregado Word que sintetiza estadísticas de múltiples nodos.
    Guarda el reporte en la carpeta ABREGADO/ de la empresa.

    parallel_node_fetch: si True, descarga medidas/alertas de cada nodo en paralelo
        (acelera el agregado; si la API falla por carga, usar False o bajar max_parallel_workers).
    """
    start_dt = parse_date(start_date)
    end_dt = parse_date(end_date, end_of_day=True)
    
    def _format_ddmmyyyy(dt: datetime) -> str:
        return dt.strftime("%d%m%Y")
    
    company_name = get_company_name(company_id)
    if apply_exclusions:
        node_ids_filtrados = filter_node_ids(node_ids, company_id=company_id, company_name=company_name)
        if len(node_ids_filtrados) < len(node_ids):
            print(f"[INFO] Se excluyeron {len(node_ids) - len(node_ids_filtrados)} nodo(s) por configuración.")
        node_ids = node_ids_filtrados
    else:
        print("[INFO] Reporte agregado: sin aplicar exclusiones de nodos (todos los IDs solicitados).")

    # Parque Arauco: siempre quitar puntos dados de baja por mall (p. ej. 000025-14 Quilicura,
    # 000025-11 El Bosque). Importante: mall_name puede venir como slug "El_Bosque" desde scripts.
    if company_id == "000025" and node_ids:
        from pa_nodos_inactivos_por_mall import aplicar_bajas_mall_pa

        node_ids = aplicar_bajas_mall_pa(mall_name, node_ids)

    if not node_ids:
        raise ValueError("Todos los nodos fueron excluidos por configuración.")
    
    # Recopilar datos de todos los nodos
    nodes_data = []
    total_consumption = 0.0
    total_alerts = 0
    all_alerts = []
    all_measures = []

    def _aggregate_fetch_one_node(node_id: str) -> tuple:
        node_name = get_node_name(node_id)
        print(f"Procesando nodo {node_id} ({node_name})...")
        try:
            measures_payload_raw = fetch_json(
                f"{acl_node_base_url()}/nodes/measures/dates",
                params=[
                    ("id", node_id),
                    ("start", _format_ddmmyyyy(start_dt)),
                    ("end", _format_ddmmyyyy(end_dt)),
                ],
            )
            measures_payload = normalize_measures_payload(measures_payload_raw, node_id)
            measures = flatten_measures(measures_payload)
            summary = summarize_consumption(measures)
            alerts = []
            if company_id != "000027":
                try:
                    alerts_payload = fetch_json(
                        f"{acl_node_base_url()}/nodes/myalert/alerts",
                        params=[
                            ("id", node_id),
                            ("start", _format_ddmmyyyy(start_dt)),
                            ("end", _format_ddmmyyyy(end_dt)),
                        ],
                    )
                    if isinstance(alerts_payload, list):
                        alerts = alerts_payload
                except Exception:
                    pass
            alert_stats = summarize_alerts(alerts, start_dt, end_dt)
            for alert in alerts:
                alert["nodeId"] = node_id
            entry = {
                "node_id": node_id,
                "node_name": node_name,
                "summary": summary,
                "alerts": alerts,
                "alert_stats": alert_stats,
                "measures": measures,
            }
            alerts_con_medida = [a for a in alerts if float(a.get("measure", 0) or 0) > 0]
            return (
                entry,
                summary["total"],
                len(alerts_con_medida),
                alerts,
                measures,
            )
        except Exception as e:
            print(f"Error procesando nodo {node_id}: {e}")
            entry = {
                "node_id": node_id,
                "node_name": node_name,
                "summary": summarize_consumption([]),
                "alerts": [],
                "alert_stats": summarize_alerts([], start_dt, end_dt),
                "measures": [],
                "error": str(e),
            }
            return entry, 0.0, 0, [], []

    use_parallel = parallel_node_fetch and len(node_ids) > 1
    if use_parallel:
        from concurrent.futures import ThreadPoolExecutor, as_completed

        workers = max(1, min(max_parallel_workers, len(node_ids)))
        print(f"[INFO] Descarga paralela de nodos (workers={workers})")
        results_by_id: dict = {}
        with ThreadPoolExecutor(max_workers=workers) as ex:
            futures = {ex.submit(_aggregate_fetch_one_node, nid): nid for nid in node_ids}
            for fut in as_completed(futures):
                nid = futures[fut]
                results_by_id[nid] = fut.result()
        for nid in node_ids:
            entry, c_delta, a_delta, alerts_ext, measures_ext = results_by_id[nid]
            nodes_data.append(entry)
            total_consumption += c_delta
            total_alerts += a_delta
            all_alerts.extend(alerts_ext)
            all_measures.extend(measures_ext)
    else:
        for node_id in node_ids:
            entry, c_delta, a_delta, alerts_ext, measures_ext = _aggregate_fetch_one_node(node_id)
            nodes_data.append(entry)
            total_consumption += c_delta
            total_alerts += a_delta
            all_alerts.extend(alerts_ext)
            all_measures.extend(measures_ext)
    
    if not nodes_data:
        raise ValueError("No se pudieron obtener datos de ningún nodo.")
    
    # Calcular estadísticas agregadas
    total_nodes = len(nodes_data)
    avg_consumption_per_node = total_consumption / total_nodes if total_nodes > 0 else 0.0
    
    # Para reporte agregado: SUMAR los promedios de cada punto, no promediar todas las alertas
    # Si hay fuente de agua especificada, excluirla de ciertos cálculos
    # Detectar automáticamente si es Fundo Zapallar (compatibilidad hacia atrás)
    es_fundo_zapallar = company_id == "000027"
    from agregado_extendido_extra import (
        OMITIR_DIA_MAYOR_Y_ALERTAS_ROJAS,
        es_agregado_extendido,
    )

    es_agregado_fmt = es_agregado_extendido(company_id)
    omitir_dia_mayor = company_id in OMITIR_DIA_MAYOR_Y_ALERTAS_ROJAS
    nodo_estanque_inferior = "000027-02" if es_fundo_zapallar else None
    
    # Si no se especificó fuente_agua_id pero es Fundo Zapallar, usar ESVAL como fuente
    if fuente_agua_id is None and es_fundo_zapallar:
        fuente_agua_id = "000027-01"  # Matriz ESVAL
    
    # Calcular proyecciones usando la misma lógica que en reportes individuales (2 últimas alertas)
    # Aplicar recálculo a cada nodo antes de sumar
    sum_promedio_alerta = 0.0
    sum_proyeccion_24h = 0.0
    sum_effective_consumption_daily = 0.0
    
    for data in nodes_data:
        # Solo excluir Estanque Inferior si es Fundo Zapallar (no excluir fuente de agua)
        if es_fundo_zapallar and data["node_id"] == nodo_estanque_inferior:
            continue
            
        alerts = data["alerts"]
        alert_stats = data["alert_stats"]
        summary = data["summary"]
        
        # Usar directamente los valores correctos de summarize_alerts
        # que ya considera solo alertas nocturnas (22:00-07:00) de los últimos 2 días
        proyeccion_24h_para_suma = alert_stats["proyeccion_24h"]
        promedio_alerta_para_suma = alert_stats["promedio_alerta"]
        
        # Sumar proyecciones (ya correctas de summarize_alerts)
        sum_promedio_alerta += promedio_alerta_para_suma
        sum_proyeccion_24h += proyeccion_24h_para_suma
        
        # Calcular consumo efectivo diario usando la proyección recalculada
        avg_daily = summary["promedio_diario"]
        effective_daily = max(0.0, avg_daily - proyeccion_24h_para_suma)
        sum_effective_consumption_daily += effective_daily
    
    # Para la gráfica, usar la misma suma (ya excluye Estanque Inferior si es Fundo Zapallar)
    sum_proyeccion_24h_para_grafica = sum_proyeccion_24h
    
    # Crear diccionario con valores agregados (sumas)
    aggregated_alert_stats = {
        "promedio_alerta": sum_promedio_alerta,
        "proyeccion_24h": sum_proyeccion_24h,
        "cantidad": total_alerts,
    }
    
    # Para consumo: ya tenemos total_consumption que es la suma correcta
    aggregated_summary = {
        "total": total_consumption,
        "promedio_diario": sum(data["summary"]["promedio_diario"] for data in nodes_data),  # Suma de promedios diarios
    }
    
    # Obtener precio del agua (usar el del primer nodo)
    price_per_m3_clp = 1200.0
    if nodes_data:
        try:
            price_per_m3_clp = get_water_price_per_m3(company_id, nodes_data[0]["node_id"], {})
        except Exception:
            pass
    
    # Crear estructura de carpetas: empresa/ABREGADO/
    # Para Parque Arauco: empresa/mall/ABREGADO/
    safe_company_name = "".join(c for c in company_name if c.isalnum() or c in (" ", "-", "_")).strip()
    safe_company_name = safe_company_name.replace(" ", "_")
    if company_folder_override:
        safe_company_name = "".join(
            c for c in company_folder_override if c.isalnum() or c in (" ", "-", "_")
        ).strip().replace(" ", "_")
    
    # Formatear fechas para buscar reportes agregados existentes
    start_dt = parse_date(start_date)
    end_dt = parse_date(end_date, end_of_day=True)
    start_str = start_dt.strftime("%Y%m%d")
    end_str = end_dt.strftime("%Y%m%d")
    pattern = f"Reporte_Agregado_{company_id}_{start_str}_{end_str}.docx"
    
    base_output_dir = Path(output_dir)
    company_dir = base_output_dir / safe_company_name
    
    # Si es Parque Arauco y hay mall_name, crear estructura por mall
    if mall_name and company_id == "000025":
        # Limpiar nombre del mall
        safe_mall_name = "".join(c for c in mall_name if c.isalnum() or c in (" ", "-", "_")).strip()
        safe_mall_name = safe_mall_name.replace(" ", "_")
        company_dir = company_dir / safe_mall_name
    elif company_id == "000025" and not mall_name:
        # Si es Parque Arauco pero no se especificó mall, intentar detectarlo del primer nodo
        if node_ids:
            first_node_name = get_node_name(node_ids[0])
            detected_mall = get_mall_name_for_parque_arauco(node_ids[0], first_node_name)
            if detected_mall:
                safe_mall_name = "".join(c for c in detected_mall if c.isalnum() or c in (" ", "-", "_")).strip()
                safe_mall_name = safe_mall_name.replace(" ", "_")
                company_dir = company_dir / safe_mall_name
    
    agregado_dir = company_dir / "ABREGADO"
    
    # Buscar reportes agregados existentes para este periodo
    existing_reports = []
    if agregado_dir.exists():
        for carpeta in agregado_dir.iterdir():
            if carpeta.is_dir() and carpeta.name.startswith("AGREGADO_"):
                reporte_file = carpeta / pattern
                if reporte_file.exists():
                    existing_reports.append(carpeta)
    
    # Eliminar reportes agregados duplicados más antiguos
    if existing_reports:
        # Ordenar por fecha de modificación (más reciente primero)
        existing_reports.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        # Eliminar todos excepto el más reciente
        for old_folder in existing_reports[1:]:
            try:
                import shutil
                shutil.rmtree(old_folder)
                print(f"Eliminado reporte agregado duplicado antiguo: {old_folder.name}")
            except Exception as e:
                print(f"Advertencia: No se pudo eliminar {old_folder.name}: {e}")
        # Usar la carpeta más reciente (reutilizar)
        output_dir_path = existing_reports[0]
        folder_name = output_dir_path.name
        # Obtener la fecha de creación de la carpeta existente o usar la actual
        creation_date = datetime.fromtimestamp(output_dir_path.stat().st_mtime, tz=timezone.utc)
        print(f"Reutilizando carpeta agregada existente: {output_dir_path.name}")
    else:
        # Crear nueva carpeta solo si no existe ninguna
        creation_date = datetime.now(timezone.utc)
        folder_name = f"AGREGADO_{creation_date.strftime('%Y%m%d_%H%M')}"
        output_dir_path = agregado_dir / folder_name
        output_dir_path.mkdir(parents=True, exist_ok=True)
    
    # Asegurar que el directorio existe
    output_dir_path.mkdir(parents=True, exist_ok=True)
    
    # Generar gráficas agregadas
    # Gráfica comparativa de consumo total por nodo
    # Si hay fuente de agua, excluirla de la gráfica de comparación
    if nodes_data:
        fig, ax = plt.subplots(figsize=(16, 8))
        
        # Crear lista de tuplas (node_name, consumption, node_id) incluyendo todos los puntos.
        # Fundo Zapallar: se muestran todos los puntos en el gráfico, pero el "total" del
        # fundo NO es la suma (doble conteo): la referencia real es Matriz ESVAL (fuente).
        # Para BUPA, excluir "Llenado de Estanques" (000029-01) de la gráfica
        es_bupa = company_id == "000029"
        node_consumption_pairs = [
            (d["node_name"], d["summary"]["total"], d["node_id"])
            for d in nodes_data
            if not (es_bupa and d["node_id"] == "000029-01")
        ]
        node_consumption_pairs.sort(key=lambda x: x[1], reverse=True)  # Ordenar de mayor a menor
        
        # Extraer nombres y consumos ordenados
        node_names = [pair[0] for pair in node_consumption_pairs]
        consumptions = [pair[1] for pair in node_consumption_pairs]
        node_ids_chart = [pair[2] for pair in node_consumption_pairs]

        if es_fundo_zapallar:
            from agregado_extendido_extra import COLOR_AGUAS_ABAJO, COLOR_MATRIZ_ESVAL
            from matplotlib.patches import Patch

            bar_colors = [
                COLOR_MATRIZ_ESVAL if nid == "000027-01" else COLOR_AGUAS_ABAJO
                for nid in node_ids_chart
            ]
            bars = ax.bar(node_names, consumptions, color=bar_colors)
            ax.legend(
                handles=[
                    Patch(facecolor=COLOR_MATRIZ_ESVAL, label="Matriz ESVAL (entrada real)"),
                    Patch(facecolor=COLOR_AGUAS_ABAJO, label="Estanques y etapas (aguas abajo)"),
                ],
                loc="upper right",
                fontsize=12,
                frameon=False,
            )
        else:
            bars = ax.bar(node_names, consumptions, color="#0050b3")
        ax.set_ylabel("Consumo total (m³)", fontsize=18, fontweight='bold')
        ax.set_xlabel("Punto de monitoreo", fontsize=18, fontweight='bold')
        titulo_barras = (
            "Consumo del periodo por punto — Matriz ESVAL destacada"
            if es_fundo_zapallar
            else (
                "Consumo total del periodo por punto de monitoreo"
                if es_agregado_fmt
                else "Consumo total por punto de monitoreo"
            )
        )
        ax.set_title(titulo_barras, fontsize=20, fontweight="bold")
        ax.set_ylim(bottom=0)
        
        # Rotar las etiquetas del eje X para mejor legibilidad con fuente más grande
        plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha='right', fontsize=16)
        plt.setp(ax.yaxis.get_majorticklabels(), fontsize=16)
        
        # Agregar valores en las barras con fuente más grande
        for i, (bar, val) in enumerate(zip(bars, consumptions)):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2, height, 
                   f"{format_number_chilean(val, 1)} m³",
                   ha='center', va='bottom', fontsize=17, fontweight='bold')
        
        plt.tight_layout()
        comparison_chart_path = output_dir_path / "chart_comparacion_nodos.png"
        plt.savefig(comparison_chart_path, dpi=150, bbox_inches='tight')
        plt.close()
    
    # Crear documento Word
    doc = Document()
    add_logo_to_header(doc)
    
    # Título
    # Para Parque Arauco, determinar el mall basado en los nodos
    mall_name = ""
    if company_id == "000025" and node_ids:
        # Obtener el nombre del mall del primer nodo (asumiendo que todos los nodos son del mismo mall)
        # Si hay múltiples malls, usar el más común
        mall_names = []
        for node_id in node_ids:
            node_name = get_node_name(node_id)
            mall = get_mall_name_for_parque_arauco(node_id, node_name)
            if mall:
                mall_names.append(mall)
        
        # Si todos los nodos son del mismo mall, usar ese
        if mall_names:
            # Contar ocurrencias de cada mall
            mall_counter = Counter(mall_names)
            # Obtener el mall más común
            mall_name = mall_counter.most_common(1)[0][0]
    
    # Construir título
    if mall_name:
        title_text = doc.add_paragraph(f"Reporte Agregado - {company_name} Mall {mall_name}")
    else:
        title_text = doc.add_paragraph(f"Reporte Agregado - {company_name}")
    title_text.style = "Title"
    for run in title_text.runs:
        run.font.size = Pt(24)
    
    subtitle = doc.add_paragraph(
        f"MONITOREO WES\n"
        f"Análisis consolidado de {total_nodes} puntos de monitoreo\n"
        f"Rango: {start_dt.strftime('%d-%m-%y')} - {end_dt.strftime('%d-%m-%y')}\n"
        f"Generado: {creation_date.strftime('%d-%m-%y')}"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if nota_contexto_periodo:
        nota_para = doc.add_paragraph(nota_contexto_periodo.strip())
        nota_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in nota_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph("")
    
    # Resumen ejecutivo agregado
    add_formatted_heading(doc, "Resumen ejecutivo agregado", level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Puntos de monitoreo analizados: {total_nodes}.\n")
    if es_fundo_zapallar:
        esval = next((d for d in nodes_data if d["node_id"] == "000027-01"), None)
        esval_m3 = float(esval["summary"]["total"]) if esval else 0.0
        summary_para.add_run(
            f"Consumo del fundo (Matriz ESVAL / entrada): {format_number_chilean(esval_m3, 1)} m³.\n"
        )
        summary_para.add_run(
            "Nota: no se suma estanques ni etapas al total (mediciones aguas abajo de ESVAL).\n"
        )
    else:
        summary_para.add_run(f"Consumo total agregado: {format_number_chilean(total_consumption, 1)} m³.\n")
        summary_para.add_run(f"Consumo promedio por punto: {format_number_chilean(avg_consumption_per_node, 1)} m³.\n")
        summary_para.add_run(f"Total de alertas registradas: {total_alerts}.\n")
    if not es_agregado_fmt and sum_promedio_alerta > 0:
        summary_para.add_run(f"Promedio de alerta agregado: {format_number_chilean(sum_promedio_alerta, 1)} m³/h.\n")
        summary_para.add_run(f"Proyección diaria de consumo nocturno agregada: {format_number_chilean(sum_proyeccion_24h, 1)} m³/día.\n")
    
    # Gráfica comparativa
    if nodes_data and comparison_chart_path.exists():
        add_formatted_heading(doc, "Comparación de consumo por punto", level=1)
        if es_fundo_zapallar:
            comp_texto = (
                "La barra naranja es la Matriz ESVAL, la entrada de agua al fundo: es el consumo real. "
                "Las barras grises son estanques y etapas aguas abajo (el mismo caudal medido en cadena); "
                "no se suman al total."
            )
        elif es_agregado_fmt:
            comp_texto = (
                "Consumo acumulado de cada punto en el periodo (suma de todos los días analizados)."
            )
        else:
            comp_texto = (
                "Consumo total registrado en cada punto de monitoreo durante el periodo analizado."
            )
        comp_para = doc.add_paragraph(comp_texto)
        comp_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in comp_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        add_picture_with_pagination(doc, str(comparison_chart_path), Inches(6), keep_with_next=True)

        if es_fundo_zapallar:
            from agregado_extendido_extra import agregar_tabla_kpis_consumo

            esval_kpi = next((d for d in nodes_data if d["node_id"] == "000027-01"), None)
            if esval_kpi:
                agregar_tabla_kpis_consumo(
                    doc, esval_kpi.get("summary") or {}, destacar_matriz=True
                )
        
        consumer_nodes_for_narrative = [
            d for d in nodes_data 
            if not (es_bupa and d["node_id"] == "000029-01")
        ]
        if es_fundo_zapallar:
            esval = next((d for d in nodes_data if d["node_id"] == "000027-01"), None)
            esval_m3 = float(esval["summary"]["total"]) if esval else 0.0
            narrative = (
                f"Consumo real del fundo = Matriz ESVAL: {format_number_chilean(esval_m3, 1)} m³. "
                f"Estanques y etapas miden caudales aguas abajo de esa matriz; "
                f"el máximo del gráfico es ESVAL, no la suma de barras."
            )
        elif es_agregado_fmt:
            from agregado_extendido_extra import narrativa_consumo_total_extendido

            narrative = narrativa_consumo_total_extendido(company_id, consumer_nodes_for_narrative)
        else:
            narrative = generate_comparison_narrative(consumer_nodes_for_narrative, avg_consumption_per_node)
        if narrative:
            doc.add_paragraph("")
            narrative_para = doc.add_paragraph(narrative)
            narrative_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in narrative_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Negro

    if es_agregado_fmt:
        try:
            from agregado_extendido_extra import agregar_secciones_consumo_diario_y_max_dia

            agregar_secciones_consumo_diario_y_max_dia(
                company_id, doc, nodes_data, start_dt, end_dt, output_dir_path
            )
        except Exception as e:
            print(f"[ADVERTENCIA] Agregado extendido — secciones consumo diario: {e}")
    
    # Tabla resumen por nodo (ordenar de mayor a menor consumo)
    add_formatted_heading(doc, "Resumen por punto de monitoreo", level=1)
    col_ultima = "Costo nocturno (CLP)" if es_agregado_fmt else "Proyección de filtración"
    omitir_col_alertas = es_fundo_zapallar
    if omitir_col_alertas:
        table_rows = [
            ("Ranking", "Dispositivo", "Consumo total (m³)", "Consumo nocturno", col_ultima)
        ]
    else:
        table_rows = [
            ("Ranking", "Dispositivo", "Consumo total (m³)", "Número de alerta", "Consumo nocturno", col_ultima)
        ]
    
    # Ordenar nodes_data por consumo total de mayor a menor
    sorted_nodes_data = sorted(nodes_data, key=lambda d: d["summary"]["total"], reverse=True)
    
    # Calcular número total de días del periodo
    num_dias_periodo = (end_dt.date() - start_dt.date()).days + 1
    
    # Variables para acumular totales
    total_consumo_total = 0.0
    total_num_alertas = 0
    total_consumo_nocturno = 0.0
    total_proyeccion_filtracion = 0.0
    total_costo_nocturno_clp = 0.0
    
    # Lista para almacenar nodos con proyección de filtración
    nodos_con_filtracion = []
    
    # Listas para la gráfica de consumo nocturno
    node_names_for_chart = []
    consumo_nocturno_values = []
    
    for rank, data in enumerate(sorted_nodes_data, start=1):
        summary = data["summary"]
        alert_stats = data["alert_stats"]
        alerts = data["alerts"]
        node_id = data["node_id"]
        
        # Reemplazar saltos de línea y espacios múltiples por espacios simples para que se muestre en una línea
        node_name = data["node_name"].replace("\n", " ").replace("\r", " ").strip()
        # Reemplazar múltiples espacios por uno solo
        while "  " in node_name:
            node_name = node_name.replace("  ", " ")
        
        # Contar número de alertas con medida > 0
        if es_agregado_fmt:
            num_alertas = len(filtrar_alertas_informativas(alerts))
        else:
            num_alertas = len([a for a in alerts if float(a.get("measure", 0) or 0) > 0])
        
        # Calcular métricas nocturnas para este nodo
        try:
            nocturnal_metrics = calculate_nocturnal_metrics(node_id, start_dt, end_dt, company_id=company_id)
            consumo_nocturno = nocturnal_metrics["consumo_nocturno_total"]
            dias_con_consumo_nocturno = nocturnal_metrics["dias_con_consumo_nocturno"]
            dias_con_datos = int(nocturnal_metrics.get("dias_con_datos_horarios", 0) or 0)
            porcentaje_dias_nocturno = (
                (dias_con_consumo_nocturno / dias_con_datos * 100.0) if dias_con_datos > 0 else 0.0
            )
        except Exception as e:
            print(f"DEBUG: Error calculando métricas nocturnas para {node_id}: {e}")
            nocturnal_metrics = {}
            consumo_nocturno = 0.0
            dias_con_consumo_nocturno = 0
            porcentaje_dias_nocturno = 0.0
        
        # Formatear consumo nocturno con días y porcentaje
        consumo_nocturno_str = f"{format_number_chilean(consumo_nocturno, 1)} m³\n({dias_con_consumo_nocturno} días, {format_number_chilean(porcentaje_dias_nocturno, 1)}%)"
        
        dias_sin_noct = int(nocturnal_metrics.get("dias_sin_consumo_nocturno", 0) or 0)
        costo_nocturno_clp = consumo_nocturno * price_per_m3_clp
        if es_agregado_fmt:
            proyeccion_filtracion_periodo = 0.0
            proyeccion_filtracion_str = format_currency_chilean(costo_nocturno_clp)
            total_costo_nocturno_clp += costo_nocturno_clp
        else:
            proyeccion_filtracion_periodo, _, _ = proyeccion_filtracion_desde_consumo_nocturno(
                consumo_nocturno,
                num_dias_periodo,
                dias_con_consumo_nocturno,
                dias_sin_noct,
                horas_nocturnas_por_dia=horas_nocturnas_por_dia_para_nodo(node_id, company_id),
            )
            porcentaje_filtracion = (
                (proyeccion_filtracion_periodo / summary["total"] * 100.0) if summary["total"] > 0 else 0.0
            )
            if proyeccion_filtracion_periodo > 0:
                proyeccion_filtracion_str = (
                    f"{format_number_chilean(proyeccion_filtracion_periodo, 1)} m³\n"
                    f"({format_number_chilean(porcentaje_filtracion, 1)}%)"
                )
                nodos_con_filtracion.append({
                    "nombre": node_name,
                    "proyeccion": proyeccion_filtracion_periodo,
                    "porcentaje": porcentaje_filtracion,
                })
            else:
                proyeccion_filtracion_str = "0,0 m³\n(0,0%)"
        
        # Acumular totales
        total_consumo_total += summary["total"]
        total_num_alertas += num_alertas
        total_consumo_nocturno += consumo_nocturno
        total_proyeccion_filtracion += proyeccion_filtracion_periodo
        
        # Agregar datos para la gráfica
        node_names_for_chart.append(node_name)
        consumo_nocturno_values.append(consumo_nocturno)
        
        if omitir_col_alertas:
            table_rows.append((
                str(rank),
                node_name,
                format_number_chilean(summary["total"], 1),
                consumo_nocturno_str,
                proyeccion_filtracion_str,
            ))
        else:
            table_rows.append((
                str(rank),
                node_name,
                format_number_chilean(summary["total"], 1),
                str(num_alertas),
                consumo_nocturno_str,
                proyeccion_filtracion_str,
            ))
    
    # Agregar fila de totales
    ultima_col_total = (
        format_currency_chilean(total_costo_nocturno_clp)
        if es_agregado_fmt
        else format_number_chilean(total_proyeccion_filtracion, 1) + " m³"
    )
    if omitir_col_alertas:
        table_rows.append((
            "",
            "TOTAL",
            format_number_chilean(total_consumo_total, 1),
            format_number_chilean(total_consumo_nocturno, 1) + " m³",
            ultima_col_total,
        ))
    else:
        table_rows.append((
            "",
            "TOTAL",
            format_number_chilean(total_consumo_total, 1),
            str(total_num_alertas),
            format_number_chilean(total_consumo_nocturno, 1) + " m³",
            ultima_col_total,
        ))
    
    add_table(
        doc,
        "Métricas por punto",
        table_rows,
        highlight_rows=None if es_agregado_fmt else [len(table_rows) - 1],
        wes_style=es_agregado_fmt,
    )
    
    if es_agregado_fmt:
        try:
            from agregado_extendido_extra import agregar_analisis_nocturno_extendido

            agregar_analisis_nocturno_extendido(
                company_id, doc, nodes_data, start_dt, end_dt, output_dir_path, price_per_m3_clp
            )
        except Exception as e:
            print(f"[ADVERTENCIA] Agregado extendido — análisis nocturno: {e}")
    
    # Generar gráfica de consumo nocturno por nodo (orden: mayor a menor consumo nocturno)
    if (
        not es_agregado_fmt
        and node_names_for_chart
        and any(v > 0 for v in consumo_nocturno_values)
    ):
        try:
            _chart_sorted = sorted(
                zip(node_names_for_chart, consumo_nocturno_values),
                key=lambda x: x[1],
                reverse=True,
            )
            node_names_for_chart = [n for n, _ in _chart_sorted]
            consumo_nocturno_values = [v for _, v in _chart_sorted]
            fig, ax = plt.subplots(figsize=(8, 5))
            bars = ax.bar(node_names_for_chart, consumo_nocturno_values, color="#FF8C00", alpha=0.8, edgecolor='#DAA520', linewidth=1.2)
            ax.set_ylabel("(m³)", fontsize=14, fontweight='bold')
            ax.set_xlabel("", fontsize=14, fontweight='bold')
            ax.set_title("Consumo nocturno por punto de monitoreo", fontsize=16, fontweight="bold")
            ax.set_ylim(bottom=0)
            
            # Rotar las etiquetas del eje X para mejor legibilidad
            plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha='right', fontsize=11)
            plt.setp(ax.yaxis.get_majorticklabels(), fontsize=12)
            
            # Agregar valores en las barras
            for bar, val in zip(bars, consumo_nocturno_values):
                if val > 0:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2, height, 
                           f"{format_number_chilean(val, 1)} m³",
                           ha='center', va='bottom', fontsize=11, fontweight='bold')
            
            plt.tight_layout()
            nocturnal_chart_path = output_dir_path / "chart_consumo_nocturno_nodos.png"
            plt.savefig(nocturnal_chart_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            # Agregar gráfica al documento
            doc.add_paragraph("")
            add_picture_with_pagination(doc, str(nocturnal_chart_path), Inches(6), keep_with_next=True)
        except Exception as e:
            print(f"DEBUG: Error generando gráfica de consumo nocturno: {e}")
    
    # Estadísticas de alertas agregadas
    nodos_graficados_horario: set = set()
    if total_alerts > 0 and not es_agregado_fmt and not omitir_dia_mayor:
        # Sección "Análisis de alertas agregado" y tabla "Métricas agregadas de consumos nocturnos" eliminadas según solicitud del usuario
        
        # Tabla de eventos registrados eliminada según solicitud del usuario
        
        # Gráfica de torta antigua eliminada según solicitud del usuario (no se genera en reportes agregados)
        
        # Agregar gráficas y análisis de alertas por nodo
        # Procesar cada nodo que tenga alertas
        for node_data in nodes_data:
            node_id = node_data["node_id"]
            node_name = node_data["node_name"]
            alerts = node_data["alerts"]
            summary = node_data["summary"]
            alert_stats = node_data["alert_stats"]
            
            # Filtrar solo alertas con medida mayor a cero
            alerts_con_medida = [a for a in alerts if float(a.get("measure", 0) or 0) > 0]
            
            # Solo procesar nodos con alertas
            if not alerts_con_medida:
                continue
            
            try:
                # Encontrar el día con mayor consumo nocturno para este nodo
                max_nocturnal_result = find_max_nocturnal_consumption_day(node_id, None, start_dt, end_dt)
                if not max_nocturnal_result:
                    continue
                
                target_dt, hourly_data = max_nocturnal_result
                dt = target_dt
                
                if hourly_data:
                    # Generar gráfica del día con mayor consumo nocturno para este nodo
                    chart_filename = f"chart_max_nocturnal_{node_id.replace('-', '_')}.png"
                    max_nocturnal_chart_path = output_dir_path / chart_filename
                    max_nocturnal_chart = build_hourly_consumption_line_chart(
                        hourly_data,
                        max_nocturnal_chart_path,
                        target_dt,
                        f"Día con mayor consumo nocturno ({dt.strftime('%d-%m-%y')})"
                    )
                    
                    if max_nocturnal_chart and max_nocturnal_chart.exists():
                        nodos_graficados_horario.add(node_id)
                        # No agregar salto de página, dejar que la paginación automática lo maneje
                        doc.add_paragraph("")  # Espacio antes del título
                        add_formatted_title(doc, f"DÍA CON MAYOR CONSUMO NOCTURNO - {node_name.upper()} ({dt.strftime('%d-%m-%y')}):")
                        add_picture_with_pagination(doc, str(max_nocturnal_chart), Inches(6), keep_with_next=True)
                        
                        # Analizar periodicidad de las alertas
                        periodicidad_info = analyze_alert_periodicity(alerts, start_dt, end_dt)
                        
                        # Calcular proyección de filtración si existe
                        num_dias_periodo = (end_dt.date() - start_dt.date()).days + 1
                        proyeccion_filtracion_periodo = 0.0
                        proyeccion_filtracion_valor_clp = 0.0
                        
                        if num_dias_periodo >= 7:
                            nocturnal_metrics = calculate_nocturnal_metrics(
                                node_id, start_dt, end_dt, company_id=company_id
                            )
                            dias_con_consumo = nocturnal_metrics["dias_con_consumo_nocturno"]
                            dias_sin_consumo = nocturnal_metrics["dias_sin_consumo_nocturno"]
                            consumo_nocturno_total = nocturnal_metrics["consumo_nocturno_total"]
                            proyeccion_filtracion_periodo, _, _ = proyeccion_filtracion_desde_consumo_nocturno(
                                consumo_nocturno_total,
                                num_dias_periodo,
                                dias_con_consumo,
                                dias_sin_consumo,
                                horas_nocturnas_por_dia=horas_nocturnas_por_dia_para_nodo(node_id, company_id),
                            )
                            if proyeccion_filtracion_periodo > 0:
                                proyeccion_filtracion_valor_clp = proyeccion_filtracion_periodo * price_per_m3_clp
                        
                        # Generar párrafo con análisis
                        analisis_texto = (
                            f"El nodo {node_name} registró {periodicidad_info['total_alertas']} alerta(s) "
                            f"durante {periodicidad_info['dias_con_alertas']} día(s) del periodo analizado. "
                            f"La periodicidad de las alertas es {periodicidad_info['periodicidad'].lower()}."
                        )
                        
                        if proyeccion_filtracion_periodo > 0:
                            analisis_texto += (
                                f" Se identificó una proyección de filtración de "
                                f"{format_number_chilean(proyeccion_filtracion_periodo, 1)} m³ "
                                f"({format_currency_chilean(proyeccion_filtracion_valor_clp)}), "
                                f"lo que indica un patrón de fuga constante durante el periodo."
                            )
                        
                        analisis_para = doc.add_paragraph(analisis_texto)
                        analisis_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        for run in analisis_para.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
                            
            except Exception as e:
                print(f"DEBUG: Error generando gráfica y análisis para nodo {node_id}: {e}")
                import traceback
                traceback.print_exc()
                continue

    # Perfiles horarios tipo app + proyección (nodos con madrugada, con o sin alertas)
    if not es_agregado_fmt:
        try:
            from wes_estilo_graficos_app import agregar_perfiles_nocturnos_agregado_doc

            agregar_perfiles_nocturnos_agregado_doc(
                doc,
                nodes_data,
                start_dt,
                end_dt,
                output_dir_path,
                nodos_ya_graficados=nodos_graficados_horario,
                price_per_m3=price_per_m3_clp,
            )
        except Exception as e:
            print(f"[ADVERTENCIA] Perfiles horarios estilo app: {e}")

    if (total_alerts > 0 or total_proyeccion_filtracion > 0) and not es_agregado_fmt:
        # Calcular consumo efectivo mensual agregado
        # Consumo efectivo diario agregado = suma de (promedio diario - proyección diaria) de cada punto
        # (excluyendo ESVAL y Estanque Inferior para Fundo Zapallar)
        effective_consumption_daily_aggregated = sum_effective_consumption_daily
        effective_consumption_monthly_aggregated = effective_consumption_daily_aggregated * 30.0
        
        # Calcular leak_monthly para la gráfica (necesario para gráfica de comparación mensual)
        num_dias_periodo_chart = (end_dt.date() - start_dt.date()).days + 1
        leak_monthly = aggregated_alert_stats["proyeccion_24h"] * 30.0
        proyeccion_mensual_desde_horaria = False
        if leak_monthly <= 0 and total_proyeccion_filtracion > 0 and num_dias_periodo_chart >= 7:
            leak_monthly = (total_proyeccion_filtracion / num_dias_periodo_chart) * 30.0
            proyeccion_mensual_desde_horaria = True
            effective_consumption_monthly_aggregated = max(
                0.0,
                (total_consumption - total_proyeccion_filtracion) / num_dias_periodo_chart * 30.0,
            )

        # Para la gráfica, usar la proyección que solo excluye Estanque Inferior si es Fundo Zapallar
        # (no excluir fuente de agua, todos los puntos son consumidores)
        if es_fundo_zapallar and nodo_estanque_inferior:
            leak_monthly_para_grafica = sum_proyeccion_24h_para_grafica * 30.0
        else:
            leak_monthly_para_grafica = leak_monthly

        # Generar gráfica de comparación mensual agregada
        monthly_chart_path = output_dir_path / "chart_comparacion_mensual_agregada.png"
        monthly_chart = None
        if aggregated_alert_stats["proyeccion_24h"] > 0 or proyeccion_mensual_desde_horaria:
            monthly_chart = build_monthly_comparison_chart(
                leak_monthly_para_grafica,
                effective_consumption_monthly_aggregated,
                price_per_m3_clp,
                monthly_chart_path,
            )
        
        # Agregar sección de comparación mensual
        if monthly_chart and monthly_chart.exists():
            doc.add_paragraph("")
            add_formatted_title(doc, "Comparación mensual: Consumo nocturno vs Consumo efectivo:")
            
            # Agregar texto explicativo
            if proyeccion_mensual_desde_horaria:
                explanation_text = (
                    "La comparación mensual proyecta los valores a 30 días. "
                    "La proyección de consumo nocturno proviene de la serie horaria API (madrugada 00:00–06:59), "
                    "no de alertas MyAlert, cuando estas no están disponibles o son puntuales. "
                    "El consumo efectivo mensual es el total del periodo menos esa proyección, escalado a 30 días. "
                )
            else:
                explanation_text = (
                    f"La comparación mensual proyecta los valores diarios a un mes completo (30 días). "
                    f"La proyección mensual de consumo nocturno se obtiene multiplicando la proyección diaria de consumo nocturno agregada por 30 días, "
                    f"mientras que el consumo efectivo mensual se calcula multiplicando el consumo efectivo promedio diario agregado por 30 días. "
                )
            
            # Si es Fundo Zapallar, agregar nota sobre exclusión de Estanque Inferior
            if es_fundo_zapallar:
                estanque_inferior_name = get_node_name(nodo_estanque_inferior)
                explanation_text += (
                    f"NOTA: Para este análisis, se ha excluido el punto {estanque_inferior_name} "
                    f"del cálculo de consumo efectivo y proyección de consumo nocturno, ya que este punto no se considera en el análisis de balance hídrico. "
                )
            
            explanation_text += (
                f"Esta comparación permite visualizar el impacto económico de los consumos nocturnos en un periodo mensual, "
                f"valorizando el consumo nocturno en pesos chilenos según el precio por metro cúbico del punto."
            )
            
            explanation_para = doc.add_paragraph(explanation_text)
            explanation_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in explanation_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
            
            doc.add_paragraph("")
            add_picture_with_pagination(doc, str(monthly_chart_path), Inches(4), keep_with_next=True)
        elif (
            aggregated_alert_stats["proyeccion_24h"] == 0
            and total_alerts > 0
            and not proyeccion_mensual_desde_horaria
        ):
            # Si no hay proyección pero hay alertas, explicar que son puntuales
            doc.add_paragraph("")
            explanation_para = doc.add_paragraph(
                f"Las alertas registradas son puntuales y no se repiten todos los días durante el periodo analizado. "
                f"Por lo tanto, no se realiza una proyección de filtración continua ni se genera la gráfica comparativa "
                f"de consumo efectivo versus filtración, ya que las alertas no representan un patrón de fuga constante."
            )
            explanation_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in explanation_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
    
    se_realizo_analisis_filtracion_agregado = len(nodos_con_filtracion) > 0

    # Conclusiones (formato estándar; el extendido cierra con conclusión nocturna)
    if not es_agregado_fmt:
        add_formatted_heading(doc, "Conclusiones", level=1)
        if se_realizo_analisis_filtracion_agregado:
            concl_para = doc.add_paragraph(
                "Este reporte sintetiza los principales hallazgos de consumo y fugas detectados durante el periodo analizado. "
                "Se recomienda revisar los días con mayor consumo y atender las alertas registradas. "
            )
            concl_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            concl_para.paragraph_format.keep_with_next = True
            concl_para.paragraph_format.widow_control = True
            for run in concl_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

            if len(nodos_con_filtracion) == 1:
                nodo_info = nodos_con_filtracion[0]
                puntos_filtracion_texto = (
                    f"El análisis de filtración detectó una proyección de fuga en el punto de monitoreo "
                    f"{nodo_info['nombre']}, con un volumen proyectado de {format_number_chilean(nodo_info['proyeccion'], 1)} m³ "
                    f"({format_number_chilean(nodo_info['porcentaje'], 1)}% del consumo total del punto). "
                )
            else:
                nombres_nodos = ", ".join([nodo['nombre'] for nodo in nodos_con_filtracion[:-1]])
                nombres_nodos += f" y {nodos_con_filtracion[-1]['nombre']}"
                puntos_filtracion_texto = (
                    f"El análisis de filtración detectó proyecciones de fuga en los siguientes puntos de monitoreo: {nombres_nodos}. "
                    f"Estos puntos presentan consumos nocturnos significativos que sugieren la presencia de fugas constantes durante el periodo analizado. "
                )

            puntos_para = doc.add_paragraph(puntos_filtracion_texto)
            puntos_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in puntos_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

            recomendacion_para = doc.add_paragraph(
                "Se recomienda realizar una inspección exhaustiva en el terreno de los puntos mencionados para identificar posibles fugas. "
                "Específicamente, se sugiere revisar el correcto funcionamiento de todos los artefactos conectados a la red de agua monitoreada "
                "por WES en estos puntos, verificando que no presenten fallas, desgaste o mal funcionamiento que puedan generar pérdidas de agua. "
                "Asimismo, se recomienda buscar indicadores de humedad en superficies del terreno, tales como áreas húmedas, charcos persistentes, "
                "crecimiento anormal de vegetación, o cualquier otra señal que pueda indicar la presencia de una filtración subterránea. "
                "Estas acciones permitirán confirmar y localizar la fuente de la posible filtración identificada en el análisis."
            )
            recomendacion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in recomendacion_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            concl_para = doc.add_paragraph(
                "Este reporte sintetiza los principales hallazgos de consumo y fugas. "
                "Se recomienda revisar los días con mayor consumo y atender las alertas registradas."
            )
            concl_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            concl_para.paragraph_format.keep_with_next = True
            concl_para.paragraph_format.widow_control = True
            for run in concl_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Guardar documento con nombre descriptivo
    company_name_clean = limpiar_nombre_archivo(company_name)
    filename = f"Reporte_Agregado_{company_name_clean}_{start_dt:%Y%m%d}_{end_dt:%Y%m%d}.docx"
    output_path = output_dir_path / filename
    try:
        doc.save(output_path)
    except PermissionError:
        alt_name = f"{filename.rsplit('.', 1)[0]}_{int(time.time())}.docx"
        output_path = output_dir_path / alt_name
        doc.save(output_path)
    
    # Generar PPT agregado si tiene más de 7 días (opcional)
    if generate_ppt:
        try:
            from generar_reportes_y_ppt_mall_maipu import generar_ppt_desde_agregado
            ppt_path = generar_ppt_desde_agregado(
                company_id=company_id,
                node_ids=node_ids,
                start_date=start_date,
                end_date=end_date,
                aggregated_report_path=output_path,
                company_name=company_name
            )
            if ppt_path:
                print(f"[OK] Presentación PPT agregada generada: {ppt_path}")
        except Exception as e:
            print(f"[ADVERTENCIA] No se pudo generar PPT agregado: {e}")
            import traceback
            traceback.print_exc()
    
    return output_path


def convertir_word_a_pdf(word_path: Path) -> Optional[Path]:
    """
    Convierte un archivo Word (.docx) a PDF temporalmente.
    Retorna la ruta del PDF temporal o None si falla.
    El PDF debe ser eliminado después de usarlo.
    """
    try:
        # Intentar usar docx2pdf (requiere Microsoft Word instalado)
        try:
            import docx2pdf
            pdf_path = word_path.with_suffix('.pdf')
            docx2pdf.convert(str(word_path), str(pdf_path))
            if pdf_path.exists():
                return pdf_path
        except ImportError:
            pass
        except Exception as e:
            print(f"[DEBUG] docx2pdf falló: {e}")
        
        # Intentar usar win32com (Windows COM automation)
        try:
            import win32com.client
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            doc = word_app.Documents.Open(str(word_path.absolute()))
            pdf_path = word_path.with_suffix('.pdf')
            doc.SaveAs(str(pdf_path.absolute()), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word_app.Quit()
            
            if pdf_path.exists():
                return pdf_path
        except ImportError:
            pass
        except Exception as e:
            print(f"[DEBUG] win32com falló: {e}")
        
        # Intentar usar comtypes (alternativa a win32com)
        try:
            import comtypes.client
            word_app = comtypes.client.CreateObject('Word.Application')
            word_app.Visible = False
            
            doc = word_app.Documents.Open(str(word_path.absolute()))
            pdf_path = word_path.with_suffix('.pdf')
            doc.SaveAs(str(pdf_path.absolute()), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word_app.Quit()
            
            if pdf_path.exists():
                return pdf_path
        except ImportError:
            pass
        except Exception as e:
            print(f"[DEBUG] comtypes falló: {e}")
        
        # Si ninguna librería está disponible, retornar None
        print("[ADVERTENCIA] No se encontró ninguna librería para convertir Word a PDF.")
        print("[INFO] Instala una de estas opciones:")
        print("  - pip install docx2pdf (requiere Microsoft Word)")
        print("  - pip install pywin32 (para win32com)")
        print("  - pip install comtypes (alternativa)")
        
        return None
        
    except Exception as e:
        print(f"[ERROR] Error al convertir Word a PDF: {e}")
        return None


def enviar_reporte_por_correo(
    reporte_path: Path,
    destinatario: str,
    smtp_servidor: str = "smtp.gmail.com",
    smtp_puerto: int = 587,
    smtp_usuario: Optional[str] = None,
    smtp_password: Optional[str] = None,
    company_name: Optional[str] = None,
    node_name: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
) -> bool:
    """
    Envía el reporte generado por correo electrónico.
    
    Args:
        reporte_path: Ruta al archivo del reporte Word
        destinatario: Correo electrónico del destinatario
        smtp_servidor: Servidor SMTP (default: smtp.gmail.com)
        smtp_puerto: Puerto SMTP (default: 587)
        smtp_usuario: Usuario SMTP (correo del remitente)
        smtp_password: Contraseña SMTP o contraseña de aplicación
        company_name: Nombre de la empresa (opcional, para el asunto)
        node_name: Nombre del nodo (opcional, para el asunto)
        start_date: Fecha inicio (opcional, para el asunto)
        end_date: Fecha fin (opcional, para el asunto)
    
    Returns:
        True si el correo se envió exitosamente, False en caso contrario
    """
    if not smtp_usuario or not smtp_password:
        print("[ERROR] Se requiere --smtp-usuario y --smtp-password para enviar correo")
        return False
    
    if not reporte_path.exists():
        print(f"[ERROR] El archivo del reporte no existe: {reporte_path}")
        return False
    
    try:
        # Crear mensaje
        msg = MIMEMultipart()
        msg["From"] = smtp_usuario
        msg["To"] = destinatario
        
        # Crear asunto
        asunto_parts = ["Reporte de Consumo"]
        if company_name:
            asunto_parts.append(f"- {company_name}")
        if node_name:
            asunto_parts.append(f"({node_name})")
        if start_date and end_date:
            asunto_parts.append(f"- {start_date} a {end_date}")
        
        msg["Subject"] = " ".join(asunto_parts)
        
        # Crear cuerpo del mensaje
        cuerpo = f"""
Estimado/a,

Se adjunta el reporte de consumo y fugas generado.

"""
        if company_name:
            cuerpo += f"Empresa: {company_name}\n"
        if node_name:
            cuerpo += f"Punto de monitoreo: {node_name}\n"
        if start_date and end_date:
            cuerpo += f"Periodo: {start_date} a {end_date}\n"
        
        cuerpo += """
Este reporte contiene análisis detallado de consumo, alertas de consumo nocturno y métricas consolidadas.

Saludos cordiales,
Sistema WES
"""
        
        msg.attach(MIMEText(cuerpo, "plain", "utf-8"))
        
        # Convertir Word a PDF temporalmente para el envío
        pdf_path = None
        try:
            pdf_path = convertir_word_a_pdf(reporte_path)
            if pdf_path and pdf_path.exists():
                # Adjuntar PDF en lugar de Word
                with open(pdf_path, "rb") as f:
                    adjunto = MIMEApplication(f.read(), _subtype="pdf")
                    adjunto.add_header(
                        "Content-Disposition",
                        "attachment",
                        filename=reporte_path.stem + ".pdf"
                    )
                    msg.attach(adjunto)
                print(f"[INFO] Reporte convertido a PDF temporalmente para envío")
            else:
                # Si falla la conversión, adjuntar Word original
                print(f"[ADVERTENCIA] No se pudo convertir a PDF, adjuntando Word original")
                with open(reporte_path, "rb") as f:
                    adjunto = MIMEApplication(f.read(), _subtype="docx")
                    adjunto.add_header(
                        "Content-Disposition",
                        "attachment",
                        filename=reporte_path.name
                    )
                    msg.attach(adjunto)
        except Exception as e:
            print(f"[ADVERTENCIA] Error al convertir a PDF: {e}. Adjuntando Word original.")
            # Si falla la conversión, adjuntar Word original
            with open(reporte_path, "rb") as f:
                adjunto = MIMEApplication(f.read(), _subtype="docx")
                adjunto.add_header(
                    "Content-Disposition",
                    "attachment",
                    filename=reporte_path.name
                )
                msg.attach(adjunto)
        
        # Enviar correo
        print(f"[INFO] Conectando al servidor SMTP {smtp_servidor}:{smtp_puerto}...")
        with smtplib.SMTP(smtp_servidor, smtp_puerto) as server:
            server.starttls()
            print(f"[INFO] Autenticando como {smtp_usuario}...")
            server.login(smtp_usuario, smtp_password)
            print(f"[INFO] Enviando correo a {destinatario}...")
            server.send_message(msg)
        
        print(f"[OK] Correo enviado exitosamente a {destinatario}")
        
        # Eliminar el PDF temporal si existe
        if pdf_path and pdf_path.exists():
            try:
                pdf_path.unlink()
                print(f"[INFO] Archivo PDF temporal eliminado")
            except Exception as e:
                print(f"[ADVERTENCIA] No se pudo eliminar el PDF temporal: {e}")
        
        return True
        
    except smtplib.SMTPAuthenticationError as e:
        print(f"[ERROR] Error de autenticación SMTP: {e}")
        print("[INFO] Verifica que el usuario y contraseña sean correctos.")
        print("[INFO] Si usas Gmail, necesitas una 'Contraseña de aplicación' en lugar de tu contraseña normal.")
        return False
    except smtplib.SMTPException as e:
        print(f"[ERROR] Error SMTP: {e}")
        return False
    except Exception as e:
        print(f"[ERROR] Error al enviar correo: {e}")
        import traceback
        traceback.print_exc()
        return False


def main() -> None:
    args = parse_args()
    try:
        output = generate_report(args)
        print(f"Reporte generado en: {output}")
        
        # Enviar por correo si se solicita
        if args.enviar_correo:
            if not args.destinatario:
                print("[ERROR] Se requiere --destinatario para enviar correo")
                sys.exit(1)
            
            company_name = get_company_name(args.company_id)
            node_name = get_node_name(args.node_id)
            start_dt = parse_date(args.start_date)
            end_dt = parse_date(args.end_date)
            
            enviar_reporte_por_correo(
                reporte_path=output,
                destinatario=args.destinatario,
                smtp_servidor=args.smtp_servidor,
                smtp_puerto=args.smtp_puerto,
                smtp_usuario=args.smtp_usuario,
                smtp_password=args.smtp_password,
                company_name=company_name,
                node_name=node_name,
                start_date=start_dt.strftime('%d-%m-%y'),
                end_date=end_dt.strftime('%d-%m-%y'),
            )
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()

