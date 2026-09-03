# -*- coding: utf-8 -*-
"""
Reenvía propuesta umbrales 000025 a Juan y Diego (CC Aníbal),
excluyendo puntos sin operación / sin data acordados con operaciones.
"""

from __future__ import annotations

import os
import smtplib
import sys
from datetime import date, datetime, timedelta
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Parque_Arauco" / "umbrales_consumo"
SRC = OUT_DIR / "Umbrales_000025_90d_x125_20260714_110239.xlsx"

REF = date(2026, 7, 14)
DIAS = 90
MULT = 1.25
DESDE = (REF - timedelta(days=DIAS - 1)).strftime("%d/%m/%Y")
HASTA = REF.strftime("%d/%m/%Y")

# Puntos fuera de alcance operativo (pruebas / baja / no prioritarios Jef. OO.MM.)
EXCLUIDOS = {
    "000025-03",  # Poniente 7
    "000025-05",  # Locales de Comida
    "000025-06",  # KFC
    "000025-14",  # Red de Incendio
    "000025-25",  # Baño N°5 Damas
    "000025-26",  # Baño N°6 Varones
}

TO = ["juanlopez@wes.cl", "diegocarrasco@wes.cl"]
CC = ["anibal.aoperaciones@wes.cl"]
SMTP_USUARIO = "agente.ia@wes.cl"
SMTP_SERVIDOR = "smtp.gmail.com"
SMTP_PUERTO = 587


def _smtp_password() -> str:
    p = (
        os.environ.get("WES_GMAIL_APP_PASSWORD", "").strip()
        or os.environ.get("WES_SMTP_PASSWORD", "").strip()
    )
    if p:
        return p.replace(" ", "").strip()
    f = ROOT / "gmail_oauth" / "app_password.txt"
    if f.is_file():
        for line in f.read_text(encoding="utf-8", errors="replace").splitlines():
            line = line.strip()
            if line and not line.startswith("#"):
                return line.replace(" ", "").strip()
    return ""


def _fn(v) -> str:
    if v is None:
        return "—"
    try:
        x = float(v)
    except (TypeError, ValueError):
        return str(v)
    s = f"{x:.1f}".replace(".", ",")
    if s.endswith(",0"):
        s = s[:-2]
    return s


def load_rows() -> list[dict]:
    wb = load_workbook(SRC, data_only=True)
    ws = wb.active
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or not row[0] or not str(row[0]).startswith("000025"):
            continue
        nid = str(row[0]).strip()
        if nid in EXCLUIDOS:
            continue
        baseline = row[3]
        umbral = row[6]
        if baseline in (None, "Sin datos") or umbral in (None, "N/D", "Sin datos"):
            continue
        try:
            b = float(baseline)
            u = float(umbral)
        except (TypeError, ValueError):
            continue
        rows.append({
            "nodeId": nid,
            "nodeName": str(row[1] or nid),
            "dias": row[2],
            "baseline": b,
            "max_periodo": row[4] if isinstance(row[4], (int, float)) else None,
            "umbral": u,
        })
    rows.sort(key=lambda r: r["nodeId"])
    return rows


def write_excel(rows: list[dict], path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Umbrales operativos"
    headers = [
        "Node ID",
        "Punto",
        "Días con data (90d)",
        "Baseline promedio diario (m³)",
        "Máximo diario en periodo (m³)",
        "Multiplicador",
        "Umbral recomendado (m³/día)",
    ]
    ws.append(headers)
    fill = PatternFill("solid", fgColor="1F4E79")
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = fill
    for r in rows:
        ws.append([
            r["nodeId"],
            r["nodeName"],
            r["dias"],
            round(r["baseline"], 1),
            round(r["max_periodo"], 1) if r["max_periodo"] is not None else "—",
            "× 1,25 (+25%)",
            round(r["umbral"], 1),
        ])
    ws.append([])
    ws.append(["Excluidos (sin data / fuera de alcance operativo Jef. OO.MM.)"])
    for nid, label in [
        ("000025-03", "Poniente 7"),
        ("000025-05", "Locales de Comida"),
        ("000025-06", "KFC"),
        ("000025-14", "Red de Incendio"),
        ("000025-25", "Baño N°5 Damas"),
        ("000025-26", "Baño N°6 Varones"),
    ]:
        ws.append([nid, label, "Excluido — no se propone umbral"])
    ws.append([])
    ws.append(["Regla", "umbral = promedio_diario_90d × 1,25"])
    ws.append(["Periodo", f"{DESDE} a {HASTA}"])
    for i in range(1, 8):
        ws.column_dimensions[get_column_letter(i)].width = [14, 42, 14, 22, 22, 14, 22][i - 1]
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def write_word(rows: list[dict], path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(
        "Propuesta de umbrales de consumo diario — Parque Arauco (000025)", 0
    )
    p = doc.add_paragraph()
    p.add_run("Agente virtual WES · Monitoreo y Control\n").bold = True
    p.add_run(
        f"Corte: {REF.strftime('%d/%m/%Y')} · Baseline: {DIAS} días "
        f"({DESDE} – {HASTA}) · Umbral: promedio × 1,25 (+25%)\n"
        f"Puntos incluidos: {len(rows)} (solo operativos con data)"
    )

    doc.add_heading("1. Contexto", 1)
    doc.add_paragraph(
        "Junto con Aníbal revisamos los valores de línea de base de la compañía 000025 "
        "(Parque Arauco) para proponer umbrales de alerta diarios compatibles con la "
        "API actual: un número por punto (m³/día); cuando el acumulado del día lo supera "
        "—a cualquier hora— se dispara el correo al cliente."
    )
    doc.add_paragraph(
        "La idea de fondo es que el agente virtual forme parte del servicio de "
        "Monitoreo WES y Control WES: no solo reportar, sino orientar umbrales, "
        "oportunidades de control diurno (p. ej. regulación de % en estanques) y "
        "recomendaciones nocturnas (p. ej. corte en horario de mall cerrado), "
        "hablando cada vez más directo con la operación del cliente."
    )

    doc.add_heading("2. Criterio técnico (por qué así)", 1)
    for b in [
        "Baseline = promedio del consumo diario total de los últimos 90 días (práctica cercana a utilities AMI; p. ej. referencias SFPUC usan ventanas del orden de 90 días).",
        "Umbral = baseline × 1,25 (+25%). Más temprano que el +50% que usan algunos esquemas tipo SFPUC: en malls queremos avisar antes para dar tiempo de reacción.",
        "No usamos el promedio de máximos como baseline: ese valor ya es un techo alto; sumarle +25% retrasa la alerta.",
        "Ejemplo Estanque Norte (000025-01): baseline ≈ 26,2 m³/día → umbral ≈ 32,8 m³/día.",
        "Se excluyeron 6 puntos sin data útil / fuera de alcance operativo según jefaturas de OO.MM. de los malls (Poniente 7, Locales de Comida, KFC, Red de Incendio, Baño Damas, Baño Varones). Impulsión Falabella (000025-09) vuelve a incluirse: activo nuevamente.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3. Tabla recomendada (puntos operativos)", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Node ID"
    hdr[1].text = "Punto"
    hdr[2].text = "Baseline 90d (m³/d)"
    hdr[3].text = "Umbral +25% (m³/d)"
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["nodeId"]
        cells[1].text = r["nodeName"]
        cells[2].text = _fn(r["baseline"])
        cells[3].text = _fn(r["umbral"])

    doc.add_heading("4. Pedido a ustedes", 1)
    doc.add_paragraph(
        "Validar y revisar en conjunto esta regla y la tabla (90 días × 1,25) antes de "
        "cargar umbrales en la API y proponerlos a los clientes de los malls. "
        "Aníbal queda en copia para alinear operación."
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def enviar(files: list[Path]) -> None:
    pw = _smtp_password()
    if not pw:
        raise RuntimeError("Falta contraseña SMTP")

    fecha = datetime.now().strftime("%d-%m-%Y")
    msg = MIMEMultipart()
    msg["From"] = SMTP_USUARIO
    msg["To"] = ", ".join(TO)
    msg["Cc"] = ", ".join(CC)
    msg["Subject"] = (
        f"Agente WES — Umbrales diarios Parque Arauco (000025): "
        f"propuesta 90d × 1,25 para validar juntos ({fecha})"
    )
    body = (
        "Hola Juan y Diego,\n\n"
        "Les escribe el agente virtual WES (Monitoreo y Control).\n\n"
        "Junto con Aníbal revisamos los valores de línea de base de todos los puntos "
        "operativos de la compañía 000025 (Parque Arauco) para proponer umbrales de "
        "alerta de consumo diario, alineados a cómo funciona hoy la API: un umbral "
        "único por punto (m³/día); cuando el acumulado del día lo supera —a la hora "
        "que sea— se envía el correo al cliente.\n\n"
        "Lo que vimos en la práctica internacional (referencias tipo SFPUC / AMI): "
        "ventanas de baseline del orden de ~90 días y multiplicadores frecuentemente "
        "en +50% sobre el promedio. En nuestro caso recomendamos avisar antes: "
        "baseline = promedio diario 90 días, umbral = baseline × 1,25 (+25%). "
        "Así el cliente recibe la alerta con más margen de reacción, sin esperar "
        "un exceso tan extremo como el +50%.\n\n"
        "Ejemplo Estanque Norte (000025-01): baseline ≈ 26,2 m³/día → umbral ≈ 32,8 m³/día.\n\n"
        "Excluimos de esta propuesta 6 puntos que están en la compañía pero no son "
        "prioridad operativa / no tienen data útil (Poniente 7, Locales de Comida, KFC, "
        "Red de Incendio, Baño Damas y Baño Varones), según el "
        "criterio de las jefaturas de operaciones de los malls. "
        "Impulsión Falabella (000025-09) se incluye: volvió a operación.\n\n"
        "Adjunto Excel (tabla limpia) y Word (explicación técnica). La idea es "
        "validarlo y revisarlo juntos; Aníbal va en copia. "
        "Este tipo de propuestas (umbrales, y más adelante controles diurnos/nocturnos "
        "donde el punto lo permita) es el rol que buscamos consolidar para el agente "
        "como parte del servicio WES frente al cliente.\n\n"
        "Quedamos atentos a sus comentarios.\n\n"
        "Saludos,\n"
        "Agente virtual WES\n"
        "Monitoreo y Control\n"
    )
    msg.attach(MIMEText(body, "plain", "utf-8"))
    for f in files:
        with open(f, "rb") as fh:
            part = MIMEApplication(fh.read())
        part.add_header("Content-Disposition", "attachment", filename=f.name)
        msg.attach(part)

    with smtplib.SMTP(SMTP_SERVIDOR, SMTP_PUERTO) as server:
        server.starttls()
        server.login(SMTP_USUARIO, pw)
        server.send_message(msg, to_addrs=TO + CC)


def main() -> int:
    if not SRC.is_file():
        print(f"[ERROR] No existe fuente: {SRC}")
        return 1
    rows = load_rows()
    print(f"[OK] Filas operativas: {len(rows)}")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    xlsx = OUT_DIR / f"Umbrales_000025_operativos_90d_x125_{stamp}.xlsx"
    docx = OUT_DIR / f"Propuesta_umbrales_000025_Juan_Diego_{stamp}.docx"
    write_excel(rows, xlsx)
    write_word(rows, docx)
    print(f"[OK] {xlsx.name}")
    print(f"[OK] {docx.name}")
    print(f"[INFO] Enviando TO={TO} CC={CC}")
    enviar([xlsx, docx])
    print("[OK] Correo enviado.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
