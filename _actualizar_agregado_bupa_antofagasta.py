"""
Actualiza el Word agregado Bupa Antofagasta ya generado:
- Elimina proyección mensual por consumo nocturno (perfiles + comparación mensual).
- Agrega proyección de consumo mensual total (promedio diario × 30) por punto.
- Agrega participación de salas de bomba vs Medidor Principal Sanitaria (m³ y CLP).
"""

from __future__ import annotations

import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

from generar_reporte_word import (
    add_formatted_title,
    add_picture_with_pagination,
    add_table,
    format_currency_chilean,
    format_number_chilean,
)

DOC_PATH = Path(
    "reports/Bupa_Antofagasta/ABREGADO/AGREGADO_20260728_1725/"
    "Reporte_Agregado_BUPA_20260723_20260728.docx"
)
OUT_DIR = DOC_PATH.parent

# Datos del reporte generado (periodo 23–28/07/2026 = 6 días)
NUM_DIAS = 6
PRECIO_M3 = 1300.0

NODOS = [
    {"id": "000029-07", "nombre": "Sala de Bomba Principal", "tipo": "bomba", "m3": 227.9},
    {"id": "000029-08", "nombre": "Sala de Bomba Sexto Piso", "tipo": "bomba", "m3": 36.1},
    {"id": "000029-09", "nombre": "Medidor Principal Sanitaria", "tipo": "cuenta", "m3": 1400.0},
    {"id": "000029-10", "nombre": "Sala de Bomba N°2", "tipo": "bomba", "m3": 6.2},
]


def _para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def _remove_from_perfiles_until_conclusiones(doc: Document) -> None:
    body = doc.element.body
    children = list(body)
    start_idx = None
    end_idx = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1]
        if tag != "p":
            continue
        t = _para_text(child).upper()
        if start_idx is None and "PERFILES HORARIOS" in t and "PROYECCI" in t:
            start_idx = i
        if start_idx is not None and t.startswith("CONCLUSIONES"):
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError(
            f"No se encontró el bloque a eliminar (start={start_idx}, end={end_idx})"
        )
    # También quitar párrafo vacío previo al bloque de perfiles si existe
    if start_idx > 0 and _para_text(children[start_idx - 1]) == "":
        start_idx -= 1
    for child in children[start_idx:end_idx]:
        body.remove(child)
    print(f"[OK] Eliminados {end_idx - start_idx} elementos (proyección nocturna / mensual)")


def _insert_before_conclusiones(doc: Document, elements) -> None:
    """Inserta elementos XML justo antes del párrafo CONCLUSIONES."""
    body = doc.element.body
    conclus = None
    for child in list(body):
        if child.tag.split("}")[-1] == "p" and _para_text(child).upper().startswith("CONCLUSIONES"):
            conclus = child
            break
    if conclus is None:
        raise RuntimeError("No se encontró CONCLUSIONES para insertar contenido")
    for el in elements:
        conclus.addprevious(el)


def _build_proyeccion_chart(path: Path, rows: list[dict]) -> Path:
    nombres = [r["nombre"] for r in rows]
    valores = [r["proy_30"] for r in rows]
    fig, ax = plt.subplots(figsize=(8, 4.8))
    bars = ax.bar(nombres, valores, color="#0050b3", alpha=0.85, edgecolor="#003a80", linewidth=1.1)
    ax.set_ylabel("m³ / 30 días", fontsize=12, fontweight="bold")
    ax.set_title("Proyección de consumo mensual total por punto", fontsize=14, fontweight="bold")
    ax.set_ylim(bottom=0)
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=25, ha="right", fontsize=10)
    for bar, val in zip(bars, valores):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(val, 1)} m³",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _build_participacion_chart(path: Path, salas_m3: float, sanitaria_m3: float) -> Path:
    resto = max(0.0, sanitaria_m3 - salas_m3)
    labels = ["Salas de bomba\n(07+08+10)", "Resto de la cuenta\n(Sanitaria − salas)"]
    sizes = [salas_m3, resto]
    colors = ["#c0392b", "#0050b3"]
    fig, ax = plt.subplots(figsize=(6.2, 4.5))
    wedges, *_ = ax.pie(
        sizes,
        labels=labels,
        colors=colors,
        autopct=lambda p: f"{p:.1f}%",
        startangle=90,
        textprops={"fontsize": 10},
    )
    for w in wedges:
        w.set_edgecolor("white")
        w.set_linewidth(1.5)
    ax.set_title(
        "Participación de salas de bomba sobre Medidor Principal Sanitaria",
        fontsize=12,
        fontweight="bold",
    )
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _make_temp_doc_sections(rows: list[dict], chart_proy: Path, chart_part: Path) -> list:
    """Construye párrafos/tablas/imágenes en un doc temporal y devuelve sus XML elements."""
    from docx.oxml import OxmlElement

    tmp = Document()

    # --- Proyección mensual total ---
    add_formatted_title(tmp, "Proyección de consumo mensual total por punto")
    intro = tmp.add_paragraph(
        f"Para cada punto se calcula el promedio diario del periodo ({NUM_DIAS} días: "
        f"23/07/2026–28/07/2026) y se proyecta a un mes de 30 días "
        f"(promedio diario × 30). La valorización usa ${format_number_chilean(PRECIO_M3, 0)} CLP/m³."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in intro.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    table_rows = [
        (
            "Punto",
            "Consumo periodo (m³)",
            "Promedio diario (m³)",
            "Proyección 30 días (m³)",
            "Proyección 30 días (CLP)",
        )
    ]
    for r in rows:
        table_rows.append(
            (
                f"{r['nombre']} ({r['id']})",
                format_number_chilean(r["m3"], 1),
                format_number_chilean(r["prom_dia"], 1),
                format_number_chilean(r["proy_30"], 1),
                format_currency_chilean(r["proy_clp"]),
            )
        )
    add_table(tmp, "Proyección mensual total", table_rows, wes_style=True)
    tmp.add_paragraph("")
    add_picture_with_pagination(tmp, str(chart_proy), Inches(6), keep_with_next=True)

    # --- Participación salas vs sanitaria ---
    sanitaria = next(r for r in rows if r["id"] == "000029-09")
    bombas = [r for r in rows if r["tipo"] == "bomba"]
    salas_m3 = sum(b["m3"] for b in bombas)
    salas_clp = salas_m3 * PRECIO_M3
    sanitaria_m3 = sanitaria["m3"]
    sanitaria_clp = sanitaria_m3 * PRECIO_M3
    pct_m3 = (salas_m3 / sanitaria_m3 * 100.0) if sanitaria_m3 > 0 else 0.0
    pct_clp = (salas_clp / sanitaria_clp * 100.0) if sanitaria_clp > 0 else 0.0

    tmp.add_paragraph("")
    add_formatted_title(tmp, "Participación de salas de bomba vs Medidor Principal Sanitaria")
    expl = tmp.add_paragraph(
        "El Medidor Principal Sanitaria (000029-09) representa el total de la cuenta de agua. "
        "Las salas de bomba (000029-07, 000029-08 y 000029-10) son submediciones internas; "
        "su participación se expresa como porcentaje del volumen y del costo de la cuenta "
        "(Sanitaria) en el mismo periodo."
    )
    expl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in expl.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    part_rows = [
        ("Concepto", "Consumo (m³)", "Costo (CLP)", "% sobre Sanitaria"),
        (
            "Medidor Principal Sanitaria (cuenta total)",
            format_number_chilean(sanitaria_m3, 1),
            format_currency_chilean(sanitaria_clp),
            "100,0%",
        ),
    ]
    for b in bombas:
        pct = (b["m3"] / sanitaria_m3 * 100.0) if sanitaria_m3 > 0 else 0.0
        part_rows.append(
            (
                f"{b['nombre']} ({b['id']})",
                format_number_chilean(b["m3"], 1),
                format_currency_chilean(b["m3"] * PRECIO_M3),
                f"{format_number_chilean(pct, 1)}%",
            )
        )
    part_rows.append(
        (
            "TOTAL salas de bomba (07+08+10)",
            format_number_chilean(salas_m3, 1),
            format_currency_chilean(salas_clp),
            f"{format_number_chilean(pct_m3, 1)}%",
        )
    )
    add_table(tmp, "Participación sobre cuenta", part_rows, wes_style=True)

    resumen = tmp.add_paragraph()
    resumen.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    resumen.add_run("Resumen: ").bold = True
    resumen.add_run(
        f"las tres salas de bomba suman {format_number_chilean(salas_m3, 1)} m³ "
        f"({format_currency_chilean(salas_clp)}), equivalentes al "
        f"{format_number_chilean(pct_m3, 1)}% del volumen y al "
        f"{format_number_chilean(pct_clp, 1)}% del costo de la cuenta medido en "
        f"Medidor Principal Sanitaria ({format_number_chilean(sanitaria_m3, 1)} m³ / "
        f"{format_currency_chilean(sanitaria_clp)})."
    )
    for run in resumen.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    tmp.add_paragraph("")
    add_picture_with_pagination(tmp, str(chart_part), Inches(5.2), keep_with_next=True)
    tmp.add_paragraph("")

    # Devolver todos los elementos del body del tmp excepto sectPr
    elements = []
    for child in list(tmp.element.body):
        tag = child.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        elements.append(child)
    return elements


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    rows = []
    for n in NODOS:
        prom = n["m3"] / NUM_DIAS
        proy = prom * 30.0
        rows.append(
            {
                **n,
                "prom_dia": prom,
                "proy_30": proy,
                "proy_clp": proy * PRECIO_M3,
            }
        )

    chart_proy = _build_proyeccion_chart(
        OUT_DIR / "chart_proyeccion_mensual_total_por_punto.png", rows
    )
    sanitaria_m3 = next(r["m3"] for r in rows if r["id"] == "000029-09")
    salas_m3 = sum(r["m3"] for r in rows if r["tipo"] == "bomba")
    chart_part = _build_participacion_chart(
        OUT_DIR / "chart_participacion_salas_vs_sanitaria.png",
        salas_m3,
        sanitaria_m3,
    )

    doc = Document(str(DOC_PATH))
    _remove_from_perfiles_until_conclusiones(doc)
    elements = _make_temp_doc_sections(rows, chart_proy, chart_part)
    _insert_before_conclusiones(doc, elements)

    try:
        doc.save(str(DOC_PATH))
        out = DOC_PATH
    except PermissionError:
        alt = DOC_PATH.with_name(DOC_PATH.stem + "_actualizado.docx")
        doc.save(str(alt))
        out = alt
        print(f"[ADVERTENCIA] Archivo abierto; guardado como {alt.name}")

    print("[OK] Word actualizado:", out)
    print("     Proyecciones 30 días:")
    for r in rows:
        print(
            f"       {r['id']} {r['nombre']}: "
            f"{format_number_chilean(r['prom_dia'], 1)} m³/día → "
            f"{format_number_chilean(r['proy_30'], 1)} m³ "
            f"({format_currency_chilean(r['proy_clp'])})"
        )
    print(
        f"     Salas/Sanitaria: {format_number_chilean(salas_m3, 1)} / "
        f"{format_number_chilean(sanitaria_m3, 1)} m³ = "
        f"{format_number_chilean(salas_m3 / sanitaria_m3 * 100, 1)}%"
    )


if __name__ == "__main__":
    main()
