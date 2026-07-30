"""
Genera un reporte tipo informe (Word + PDF) que compara:
- "App" = consumo WES (API medidas) del nodo 000017-08
- Facturaciones Aguas Andinas = consumo total m³ informado en cada factura PDF

Fuente de facturas:
  reports/Renca/Coparacion App con Aguas Andinas/ICCO facturas/*.pdf

Uso:
  python generar_reporte_comparacion_app_vs_facturaciones_icco_000017_08.py

Cuadros comparativos (tablas A–C, gráficos y total acumulado): solo boletas con lectura anterior y lectura actual
≥ 28-09-2025 (``INFORME_COMPARATIVO_DESDE``).
"""

from __future__ import annotations

import os
import re
import shutil
import sys
import tempfile
import traceback
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

from generar_reporte_word import convertir_word_a_pdf, format_number_chilean


def _ahora_informe_local() -> datetime:
    """Momento de generación en la zona horaria **local** del equipo.

    Si se guardan metadatos como UTC (``...Z``), Word suele mostrar la hora corrida
    (p. ej. Chile: 21:59 real → 17:59 en propiedades). Con ``tzinfo`` local, la
    hora del documento coincide con el reloj que ves al ejecutar el script.
    """
    return datetime.now(timezone.utc).astimezone()


def _doc_establecer_metadatos(doc: Document, generado_en: datetime) -> None:
    """La plantilla por defecto de python-docx suele traer fechas heredadas (p. ej. 2013).
    Fijamos creación y modificación al momento real del informe para Word/Explorador."""
    cp = doc.core_properties
    cp.created = generado_en
    cp.modified = generado_en
    cp.title = "Comparación App WES vs Facturaciones Aguas Andinas"
    cp.subject = f"{NODE_ID} — {NODE_NOMBRE} — {EMPRESA}"
    cp.keywords = "WES; Aguas Andinas; ICCO; Renca; facturación; consumo"
    cp.category = "Informe comparativo"
    cp.comments = (
        "Informe generado automáticamente: facturas Aguas Andinas (PDF) vs consumo API nodo WES."
    )
    cp.author = "Script informe ICCO Renca"
    cp.last_modified_by = cp.author
    try:
        cp.revision = 1
    except Exception:
        pass


def _docx_parchear_core_hora_local(out_docx: Path, generado_en: datetime) -> None:
    """python-docx guarda ``created``/``modified`` como UTC (``...Z``); Word en Chile resta el huso
    y la hora de propiedades queda desfasada (p. ej. 21:59 → 17:59). Reescribimos ``core.xml`` con
    W3CDTF **con offset local** (``-04:00`` / ``-03:00``, etc.)."""
    if generado_en.tzinfo is None:
        generado_en = generado_en.astimezone()
    # ISO 8601 con offset; Word acepta el formato con dos puntos en el huso.
    marca = generado_en.isoformat(timespec="seconds")

    with ZipFile(out_docx, "r") as zin:
        nombres = zin.namelist()
        if "docProps/core.xml" not in nombres:
            return
        core = zin.read("docProps/core.xml").decode("utf-8")

    def _reemplaza(tag: str, xml: str) -> str:
        pat = re.compile(
            rf"(<dcterms:{tag}\s[^>]*>)([^<]*)(</dcterms:{tag}>)",
            re.DOTALL,
        )
        m = pat.search(xml)
        if not m:
            return xml
        return pat.sub(rf"\g<1>{marca}\g<3>", xml, count=1)

    core2 = _reemplaza("created", core)
    core2 = _reemplaza("modified", core2)
    if core2 == core:
        return

    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    tmp_path = Path(tmp)
    try:
        with ZipFile(out_docx, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "docProps/core.xml":
                    data = core2.encode("utf-8")
                zout.writestr(info, data)
        shutil.move(str(tmp_path), str(out_docx))
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass


ROOT = Path(__file__).resolve().parent
IN_DIR = (
    ROOT
    / "reports"
    / "Renca"
    / "Coparacion App con Aguas Andinas"
    / "ICCO facturas"
)
OUT_DIR = ROOT / "reports" / "Renca" / "Coparacion App con Aguas Andinas" / "reporte_comparacion_Icco"

NODE_ID = "000017-08"
NODE_NOMBRE = "ICCO Renca (000017-08)"
EMPRESA = "Renca"

# Referencia operaciones / cruce Excel: inicio monitoreo WES y único ciclo “cuadrado” previo a transición estable.
# Narrativa operación (ICCO); cuadros comparativos del informe usan INFORME_COMPARATIVO_DESDE.
MONITOREO_WES_DESDE = datetime(2025, 8, 26)
INFORME_COMPARATIVO_DESDE = date(2025, 9, 28)
PERIODO_ALINEADO_EXCEL_INI = datetime(2025, 10, 11)
PERIODO_ALINEADO_EXCEL_FIN = datetime(2025, 11, 11)

# Clasificación por número de factura (operación ICCO); prioridad sobre fechas.
BOLETA_SOLO_TABLA_B = "9019235"  # única fila en transición
BOLETAS_TABLA_C = frozenset({"9075418", "9131949", "9188945"})

_MESES = {
    "ENE": 1,
    "FEB": 2,
    "MAR": 3,
    "ABR": 4,
    "MAY": 5,
    "JUN": 6,
    "JUL": 7,
    "AGO": 8,
    "SEP": 9,
    "OCT": 10,
    "NOV": 11,
    "DIC": 12,
}


def _parse_fecha_es_dd_mmm_yyyy(s: str) -> datetime:
    # Ej: 16-ABR-2026
    p = s.strip().upper().split("-")
    if len(p) != 3:
        raise ValueError(f"Fecha inválida: {s!r}")
    dd = int(p[0])
    mm = _MESES[p[1][:3]]
    yy = int(p[2])
    return datetime(yy, mm, dd)


def _fmt_periodo(dt: datetime) -> str:
    meses = {
        1: "ene",
        2: "feb",
        3: "mar",
        4: "abr",
        5: "may",
        6: "jun",
        7: "jul",
        8: "ago",
        9: "sep",
        10: "oct",
        11: "nov",
        12: "dic",
    }
    return f"{dt.day:02d}-{meses[dt.month]}-{dt.year}"


def _to_ddmmyyyy(dt: datetime) -> str:
    return dt.strftime("%d%m%Y")


def _extraer_texto_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for pg in reader.pages:
        parts.append(pg.extract_text() or "")
    return "\n".join(parts)


@dataclass(frozen=True)
class FacturaAA:
    pdf: Path
    boleta: str
    cuenta: str | None
    medidor: str | None
    emision: datetime
    lectura_anterior: datetime
    lectura_actual: datetime
    m3_cuenta: int

    @property
    def api_start(self) -> str:
        return _to_ddmmyyyy(self.lectura_anterior)

    @property
    def api_end(self) -> str:
        return _to_ddmmyyyy(self.lectura_actual)

    @property
    def periodo_txt(self) -> str:
        return f"{_fmt_periodo(self.lectura_anterior)} → {_fmt_periodo(self.lectura_actual)}"


def _parse_factura_aguas_andinas_pdf(path: Path) -> FacturaAA:
    txt = _extraer_texto_pdf(path)

    boleta_m = re.search(
        r"(?:FACTURA|BOLETA)\s+ELECTR[ÓO]NICA\s*\n?\s*N[º°]\s*([0-9]{6,})",
        txt,
        flags=re.IGNORECASE,
    )
    if not boleta_m:
        boleta_m = re.search(r"\nN[º°]\s*([0-9]{6,})\n", txt, flags=re.IGNORECASE)
    emision_m = re.search(
        r"FECHA\s+EMISI[ÓO]N[:\s]*([0-9]{2}-[A-Z]{3}-[0-9]{4})",
        txt,
        flags=re.IGNORECASE,
    )
    cuenta_m = re.search(r"\n([0-9]{7,}-[0-9])\n", txt)  # ej: 1923771-0
    medidor_m = re.search(r"N[úu]mero\s+de\s+Medidor\s+([0-9\.]+)", txt, flags=re.IGNORECASE)
    consumo_m = re.search(r"CONSUMO\s+TOTAL\s+([0-9\.\,]+)\s*m3", txt, flags=re.IGNORECASE)
    # Formato estricto (m³ en la misma línea) o solo fecha (como en algunas boletas).
    la_m = re.search(
        r"LECTURA\s+ACTUAL\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})\s+[0-9\.\,]+\s*m3",
        txt,
        flags=re.IGNORECASE,
    ) or re.search(
        r"LECTURA\s+ACTUAL\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})",
        txt,
        flags=re.IGNORECASE,
    )
    lan_m = re.search(
        r"LECTURA\s+ANTERIOR\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})\s+[0-9\.\,]+\s*m3",
        txt,
        flags=re.IGNORECASE,
    ) or re.search(
        r"LECTURA\s+ANTERIOR\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})",
        txt,
        flags=re.IGNORECASE,
    )

    if not (emision_m and consumo_m and la_m and lan_m):
        raise ValueError(f"No se pudieron extraer campos desde {path.name}")

    dt_emision = _parse_fecha_es_dd_mmm_yyyy(emision_m.group(1))
    dt_actual = _parse_fecha_es_dd_mmm_yyyy(la_m.group(1))
    dt_anterior = _parse_fecha_es_dd_mmm_yyyy(lan_m.group(1))
    m3_cuenta = int(float(consumo_m.group(1).replace(".", "").replace(",", ".")))

    return FacturaAA(
        pdf=path,
        boleta=(boleta_m.group(1) if boleta_m else path.stem),
        cuenta=(cuenta_m.group(1) if cuenta_m else None),
        medidor=(medidor_m.group(1) if medidor_m else None),
        emision=dt_emision,
        lectura_anterior=dt_anterior,
        lectura_actual=dt_actual,
        m3_cuenta=m3_cuenta,
    )


def _cargar_facturas() -> list[FacturaAA]:
    pdfs = sorted(IN_DIR.glob("*.pdf"))
    if not pdfs:
        raise FileNotFoundError(f"No hay PDFs en {IN_DIR}")
    rows = [_parse_factura_aguas_andinas_pdf(p) for p in pdfs]
    rows.sort(key=lambda r: r.lectura_actual)
    return rows


def _particionar_por_corte_informe(
    facturas: list[FacturaAA], desde: date
) -> tuple[list[FacturaAA], list[FacturaAA]]:
    """Incluye solo boletas con lectura anterior y lectura actual ≥ ``desde`` (cuadro comparable App vs cuenta)."""
    incl: list[FacturaAA] = []
    excl: list[FacturaAA] = []
    for f in facturas:
        d0 = f.lectura_anterior.date()
        d1 = f.lectura_actual.date()
        if d0 >= desde and d1 >= desde:
            incl.append(f)
        else:
            excl.append(f)
    return incl, excl


def _fetch_wes_total(f: FacturaAA) -> tuple[float, int]:
    """Suma WES entre lectura anterior y lectura actual (inclusive).

    El backend suele interpretar ``end`` como **exclusivo**; se pide un día extra
    y se filtra contra el rango civil de la boleta para no perder el último día
    (caso típico: período que termina en enero sin barra App en el gráfico).
    """
    from generar_reporte_word import (
        acl_node_base_url,
        fetch_json,
        normalize_measures_payload,
        flatten_measures,
        summarize_consumption,
    )

    start = f.api_start
    end_api = _to_ddmmyyyy(f.lectura_actual + timedelta(days=1))
    base = acl_node_base_url()
    raw = fetch_json(
        f"{base}/nodes/measures/dates",
        params=[("id", NODE_ID), ("start", start), ("end", end_api)],
    )
    norm = normalize_measures_payload(raw, NODE_ID)
    meas = flatten_measures(norm)
    d0, d1 = f.lectura_anterior.date(), f.lectura_actual.date()
    meas = [m for m in meas if d0 <= m.date.date() <= d1]
    s = summarize_consumption(meas)
    return float(s.get("total", 0.0)), int(s.get("dias", 0))


def _fmt_num(n: float, dec: int = 1) -> str:
    return format_number_chilean(n, dec)


def _generar_grafico_barras(out_png: Path, filas: list[dict]) -> None:
    """Barras agrupadas con ``ax.bar`` (escala correcta); dos azules. Izq = app WES, der = facturación."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib import patheffects as pe

    rc = {
        "font.family": "sans-serif",
        "font.sans-serif": ["Segoe UI", "Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 10.5,
        "axes.labelsize": 11,
        "axes.edgecolor": "#94a3b8",
        "axes.linewidth": 0.8,
        "xtick.labelsize": 9.5,
        "ytick.labelsize": 10,
        "legend.fontsize": 10,
        "figure.facecolor": "#f8fafc",
        "axes.facecolor": "#ffffff",
    }

    if not filas:
        return

    cuenta = np.array([float(f["m3_cuenta"]) for f in filas], dtype=float)
    wes = np.array([float(f["m3_wes"]) for f in filas], dtype=float)

    labels_x = [
        f'N° {f["boleta"]}\n{f["periodo_short"]}\n(izq.: app WES · der.: fact.)'
        for f in filas
    ]

    col_wes = "#38bdf8"
    col_fact = "#1e3a8f"
    edge = "#ffffff"

    with matplotlib.rc_context(rc):
        fig, ax = plt.subplots(figsize=(13.5, 6.8))
        n = len(filas)
        xg = np.arange(n, dtype=float)
        bw = 0.36

        ymax = float(max(float(cuenta.max()), float(wes.max()), 1.0)) * 1.2

        bar_w = ax.bar(
            xg - bw / 2,
            wes,
            bw,
            color=col_wes,
            edgecolor=edge,
            linewidth=1.15,
            label="Consumo según app WES (m³)",
            zorder=3,
            clip_on=False,
        )
        bar_f = ax.bar(
            xg + bw / 2,
            cuenta,
            bw,
            color=col_fact,
            edgecolor=edge,
            linewidth=1.15,
            label="Consumo facturado Aguas Andinas (m³)",
            zorder=2,
            clip_on=False,
        )

        def _etiqueta_arriba(bar, vals: np.ndarray, *, zeros: bool = False) -> None:
            rects = getattr(bar, "patches", None) or tuple(bar)
            for rect, val in zip(rects, vals):
                hgt = rect.get_height()
                if val <= 0 and zeros:
                    zm = ax.annotate(
                        _fmt_num(float(val), 0 if abs(val) >= 10 else 1),
                        xy=(rect.get_x() + rect.get_width() / 2, 0),
                        xytext=(0, 11),
                        textcoords="offset points",
                        ha="center",
                        va="bottom",
                        fontsize=8.5,
                        color="#64748b",
                        zorder=5,
                    )
                    zm.set_path_effects([pe.withStroke(linewidth=3, foreground="white")])
                    continue
                if hgt <= 0:
                    continue
                t = ax.annotate(
                    _fmt_num(float(val), 0 if val >= 10 else 1),
                    xy=(rect.get_x() + rect.get_width() / 2, hgt),
                    xytext=(0, 6),
                    textcoords="offset points",
                    ha="center",
                    va="bottom",
                    fontsize=9.5,
                    color="#0f172a",
                    fontweight="600",
                    zorder=5,
                )
                t.set_path_effects([pe.withStroke(linewidth=3.5, foreground="white")])

        _etiqueta_arriba(bar_w, wes, zeros=True)
        _etiqueta_arriba(bar_f, cuenta, zeros=False)

        ax.set_ylim(0, ymax)
        ax.set_xticks(xg)
        ax.set_xticklabels(labels_x, ha="center", linespacing=1.35)
        ax.set_xlabel("Factura y fin de período (lectura actual)", color="#475569", labelpad=12)
        ax.set_ylabel("Consumo total del período (m³)", color="#475569", labelpad=8)

        ax.set_axisbelow(True)
        ax.yaxis.grid(True, linestyle="--", linewidth=0.7, color="#cbd5e1", alpha=0.9, zorder=0)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        ax.legend(loc="upper right", framealpha=0.97, fancybox=True, edgecolor="#e2e8f0")

        fig.suptitle(
            "Por factura: app WES (izquierda) vs facturación (derecha)",
            fontsize=14,
            fontweight="600",
            color="#0f172a",
            y=0.97,
        )
        fig.text(
            0.5,
            0.935,
            f"ICCO Renca · nodo {NODE_ID} — lecturas anterior y actual ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}",
            ha="center",
            fontsize=10,
            color="#64748b",
        )

        fig.subplots_adjust(bottom=0.26, left=0.07, right=0.98, top=0.82)
        fig.savefig(out_png, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor(), edgecolor="none")
        plt.close(fig)


def _generar_grafico_totales(
    out_png: Path,
    *,
    total_wes: float,
    total_cuenta: float,
    n_boletas: int,
) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib import patheffects as pe

    rc = {
        "font.family": "sans-serif",
        "font.sans-serif": ["Segoe UI", "Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 11,
        "axes.labelsize": 11,
        "axes.edgecolor": "#94a3b8",
        "figure.facecolor": "#f8fafc",
        "axes.facecolor": "#ffffff",
    }
    col_wes = "#38bdf8"
    col_fact = "#1e3a8f"
    labels = ["Total App WES\n(suma m³)", "Total facturación\ncuenta (suma m³)"]
    vals = np.array([float(total_wes), float(total_cuenta)], dtype=float)
    x = np.arange(2, dtype=float)
    ymax = float(max(vals.max(), 1.0)) * 1.22

    with matplotlib.rc_context(rc):
        fig, ax = plt.subplots(figsize=(8.2, 5.4))
        bars = ax.bar(x, vals, width=0.52, color=[col_wes, col_fact], edgecolor="#ffffff", linewidth=1.2, zorder=2)
        for rect, val in zip(bars, vals):
            if val <= 0:
                continue
            t = ax.annotate(
                _fmt_num(float(val), 1),
                xy=(rect.get_x() + rect.get_width() / 2, float(val)),
                xytext=(0, 6),
                textcoords="offset points",
                ha="center",
                va="bottom",
                fontsize=11,
                fontweight="600",
                color="#0f172a",
                zorder=5,
            )
            t.set_path_effects([pe.withStroke(linewidth=3.5, foreground="white")])

        ax.set_xticks(x)
        ax.set_xticklabels(labels, ha="center", linespacing=1.25)
        ax.set_ylabel("m³ acumulados", color="#475569", labelpad=10)
        ax.set_ylim(0, ymax)
        ax.set_axisbelow(True)
        ax.yaxis.grid(True, linestyle="--", linewidth=0.7, color="#cbd5e1", alpha=0.9, zorder=0)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        fig.suptitle(
            "Total acumulado: App WES vs facturación en cuenta — ICCO",
            fontsize=14,
            fontweight="600",
            color="#0f172a",
            y=0.96,
        )
        fig.text(
            0.5,
            0.90,
            (
                f"Nodo {NODE_ID} — {n_boletas} boleta(s) — lecturas ≥ "
                f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}"
            ),
            ha="center",
            fontsize=10,
            color="#64748b",
        )
        fig.subplots_adjust(bottom=0.18, left=0.12, right=0.97, top=0.82)
        fig.savefig(out_png, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor(), edgecolor="none")
        plt.close(fig)


def generar_reporte() -> tuple[Path, Path | None]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ahora = _ahora_informe_local()
    run_id = ahora.strftime("%Y%m%d_%H%M%S")

    facturas_todas = _cargar_facturas()
    facturas, facturas_excl = _particionar_por_corte_informe(facturas_todas, INFORME_COMPARATIVO_DESDE)
    if not facturas:
        raise ValueError(
            f"No hay facturas ICCO con lectura anterior y lectura actual ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}."
        )

    doc = Document()
    _doc_establecer_metadatos(doc, ahora)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Comparación App WES vs Facturaciones Aguas Andinas", level=1)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    p = doc.add_paragraph()
    p.add_run("Empresa: ").bold = True
    p.add_run(EMPRESA)
    p2 = doc.add_paragraph()
    p2.add_run("Nodo: ").bold = True
    p2.add_run(f"{NODE_ID} — {NODE_NOMBRE}")
    p3 = doc.add_paragraph()
    p3.add_run("Fuente facturas: ").bold = True
    p3.add_run(str(IN_DIR))
    p4 = doc.add_paragraph()
    p4.add_run("Generado: ").bold = True
    p4.add_run(ahora.strftime("%d-%m-%Y %H:%M"))

    doc.add_paragraph()
    doc.add_heading("1. Alcance y criterio de comparación", level=1)
    doc.add_paragraph(
        "Para cada boleta, se compara el consumo total facturado (m³) con el consumo total registrado por WES "
        "en el nodo indicado, para el período entre «Lectura anterior» y «Lectura actual» informados en la factura. "
        "Las diferencias pueden explicarse por alcances hidráulicos distintos (medidor principal vs subred), "
        "pérdidas/derivaciones no medidas por el punto, y diferencias de sincronización/calendario del período."
    )
    doc.add_paragraph(
        "Importante: el punto WES no siempre estuvo midiendo el mismo tramo hidráulico ni en condiciones "
        "de operación normal. Para la evaluación con sentido operacional se considera el monitoreo a partir del "
        f"{MONITOREO_WES_DESDE.strftime('%d-%m-%Y')} (inicio de monitoreo de referencia). Hasta antes de la transición "
        "posterior, el registro está afectado por operación con bypass y/o uso intermitente de la sala de bombas, "
        "por lo que la comparación App vs Facturación suele no ser representativa. En el cruce con el Excel comparativo "
        f"facturación vs app, el único ciclo mensual que aparece alineado («cuadrado») en ese tramo es el período entre "
        f"lecturas {_fmt_periodo(PERIODO_ALINEADO_EXCEL_INI)} y {_fmt_periodo(PERIODO_ALINEADO_EXCEL_FIN)}; el resto "
        "de ciclos en esa etapa muestra desalineación acorde a la configuración hidráulica descrita. "
        "Existe un período de transición (11-dic-2025 a 08-ene-2026) donde se observan ajustes y estabilización. "
        "Desde el fin de esa transición en adelante, el funcionamiento se considera más alineado con la medición "
        "correcta y la comparación es la referencia principal (aun así, pueden existir diferencias)."
    )
    doc.add_paragraph(
        f"Las tablas, gráficos por período y el total acumulado de la sección 2 incluyen únicamente boletas cuya "
        f"«Lectura anterior» y «Lectura actual» son el {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')} o posteriores. "
        "En ciclos anteriores o con alguna lectura previa a esa fecha, la cuenta suele reflejar liquidación por "
        "promedio estimado u homologación sin contraste de punta a punta con el punto WES en ese tramo."
    )

    doc.add_heading(
        f"1.1 Boletas fuera del cuadro comparativo (lecturas anteriores al {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})",
        level=2,
    )
    if facturas_excl:
        doc.add_paragraph(
            "Las siguientes boletas no entran en las tablas 2 ni en los gráficos ni en el total acumulado; "
            "suelen corresponder a promedio estimado o períodos no homologables con el monitoreo WES en el tramo indicado."
        )
        for f in sorted(facturas_excl, key=lambda x: x.lectura_actual):
            doc.add_paragraph(
                f"N° {f.boleta} — {f.periodo_txt} (m³ facturados en boleta: {f.m3_cuenta})",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No hay boletas excluidas por este criterio.")

    doc.add_heading(
        f"2. Resumen por período de factura (lecturas ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})",
        level=1,
    )
    headers = [
        "Período (lecturas)",
        "Emisión",
        "N° factura",
        "Medidor",
        "m³ facturado",
        "m³ App (WES)",
        "Días (WES)",
        "Dif (fact - WES)",
        "% Dif vs Fact",
    ]
    # Evitar tablas “cortadas”: si hay muchas filas, se dividen en sub-tablas por página.
    MAX_FILAS_TABLA_POR_PAGINA = 10  # filas de datos (sin contar encabezado)
    # Tablas según número de factura acordado con operaciones (prioridad sobre fechas del PDF).
    filas_bypass: list[FacturaAA] = []
    filas_alineado: list[FacturaAA] = []
    filas_transicion: list[FacturaAA] = []
    for f in facturas:
        if f.boleta == BOLETA_SOLO_TABLA_B:
            filas_transicion.append(f)
        elif f.boleta in BOLETAS_TABLA_C:
            filas_alineado.append(f)
        else:
            filas_bypass.append(f)

    def _filas_dict(fs: list[FacturaAA]) -> list[dict]:
        out: list[dict] = []
        for f in fs:
            m3_wes, dias = _fetch_wes_total(f)
            diff = float(f.m3_cuenta) - float(m3_wes)
            # Porcentaje de diferencia respecto a lo facturado:
            # (% "no explicado por WES" si diff>0). Usa base facturación para lectura intuitiva.
            pct_diff = (100.0 * float(diff) / float(f.m3_cuenta)) if f.m3_cuenta else 0.0
            out.append(
                {
                    "periodo": f.periodo_txt,
                    "periodo_short": _fmt_periodo(f.lectura_actual),
                    "emision": f.emision.strftime("%d-%b-%Y").lower().replace("apr", "abr").replace("aug", "ago").replace("dec", "dic"),
                    "boleta": f.boleta,
                    "cuenta": f.cuenta or "-",
                    "medidor": f.medidor or "-",
                    "m3_cuenta": int(f.m3_cuenta),
                    "m3_wes": float(m3_wes),
                    "dias_wes": int(dias),
                    "diff": float(diff),
                    "pct": float(pct_diff),
                    "pdf": f.pdf.name,
                    "lectura_fin": f.lectura_actual,
                }
            )
        return out

    filas_bypass_d = _filas_dict(filas_bypass)
    filas_transicion_d = _filas_dict(filas_transicion)
    filas_alineado_d = _filas_dict(filas_alineado)

    def _tbl_pr_set_full_width_fixed(table) -> None:
        """Ancho 100 % y diseño fijo: evita que Word reacomode columnas al abrir el archivo."""
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        for tag_nm in ("tblW", "tblLayout"):
            tq = qn(f"w:{tag_nm}")
            for el in list(tbl_pr):
                if el.tag == tq:
                    tbl_pr.remove(el)
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), "5000")
        tbl_w.set(qn("w:type"), "pct")
        tbl_pr.append(tbl_w)
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), "fixed")
        tbl_pr.append(tbl_layout)

    def _tbl_set_repeat_header(table) -> None:
        """Repetir encabezado cuando la tabla continúa en otra página."""
        try:
            tr = table.rows[0]._tr
        except Exception:
            return
        tr_pr = tr.get_or_add_trPr()
        # Quitar headers repetidos anteriores si existieran
        for el in list(tr_pr):
            if el.tag == qn("w:tblHeader"):
                tr_pr.remove(el)
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        tr_pr.append(hdr)

    def _add_table(title_txt: str, rows: list[dict]) -> None:
        p_tit = doc.add_paragraph()
        p_tit.add_run(title_txt).bold = True
        p_tit.paragraph_format.keep_with_next = True
        if not rows:
            doc.add_paragraph("Sin facturas en este rango.")
            return
        chunks: list[list[dict]] = [
            rows[i : i + MAX_FILAS_TABLA_POR_PAGINA]
            for i in range(0, len(rows), MAX_FILAS_TABLA_POR_PAGINA)
        ]

        for chunk_idx, chunk_rows in enumerate(chunks):
            if chunk_idx > 0:
                doc.add_page_break()
                p_ct = doc.add_paragraph()
                p_ct.add_run(f"{title_txt} (continuación)").bold = True
                p_ct.paragraph_format.keep_with_next = True

            table = doc.add_table(rows=1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            _tbl_pr_set_full_width_fixed(table)
            _tbl_set_repeat_header(table)
            hdr = table.rows[0].cells
            for i, h in enumerate(headers):
                r = hdr[i].paragraphs[0].add_run(h)
                r.bold = True
                r.font.size = Pt(9)
            for row in chunk_rows:
                cells = table.add_row().cells
                cells[0].text = row["periodo"]
                cells[1].text = row["emision"]
                cells[2].text = str(row["boleta"])
                cells[3].text = str(row["medidor"])
                cells[4].text = _fmt_num(float(row["m3_cuenta"]), 0)
                cells[5].text = _fmt_num(float(row["m3_wes"]), 1)
                cells[6].text = str(row["dias_wes"])
                cells[7].text = _fmt_num(float(row["diff"]), 1)
                cells[8].text = _fmt_num(float(row["pct"]), 1) + "%"

            table.allow_autofit = False
            for r in table.rows:
                try:
                    # No partir una fila en dos páginas.
                    r.allow_break_across_pages = False
                except Exception:
                    pass
                for c in r.cells:
                    for p in c.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        for run in p.runs:
                            run.font.size = Pt(9)

    p_a = doc.add_paragraph(
        "Tabla A — Períodos fuera de rango (bypass), solo lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}: operación no representativa (bypass / sala de bombas). "
        f"Referencia narrativa de monitoreo {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}; en Excel, ciclo alineado "
        f"entre {_fmt_periodo(PERIODO_ALINEADO_EXCEL_INI)} y {_fmt_periodo(PERIODO_ALINEADO_EXCEL_FIN)}."
    )
    p_a.paragraph_format.keep_with_next = True
    _add_table("Períodos fuera de rango (bypass)", filas_bypass_d)

    doc.add_paragraph()
    p_b = doc.add_paragraph(
        "Tabla B — Períodos de transición (11-dic-2025 a 08-ene-2026), solo lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}: ajustes y estabilización."
    )
    p_b.paragraph_format.keep_with_next = True
    _add_table("Períodos de transición (11-dic-2025 a 08-ene-2026)", filas_transicion_d)

    doc.add_paragraph()
    p_c = doc.add_paragraph(
        "Tabla C — Períodos alineados (posterior a 08-ene-2026), solo lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}: referencia principal; «% Dif vs Fact» respecto a la facturación."
    )
    p_c.paragraph_format.keep_with_next = True
    _add_table("Períodos alineados (posterior a 08-ene-2026)", filas_alineado_d)

    doc.add_paragraph()
    filas_graf = [dict(r) for r in filas_bypass_d + filas_transicion_d + filas_alineado_d]
    filas_graf.sort(key=lambda r: r["lectura_fin"])
    for r in filas_graf:
        r.pop("lectura_fin", None)
    img_barras = OUT_DIR / f"chart_barras_icco_todos_{run_id}.png"
    _generar_grafico_barras(img_barras, filas_graf)
    doc.add_paragraph(
        "Gráfico consolidado — todas las boletas del cuadro, orden cronológico por lectura actual; "
        "izquierda App WES, derecha facturación. "
        f"Solo lecturas ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}."
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if filas_graf:
        doc.add_picture(str(img_barras), width=Inches(6.7))

    def _grafico_categoria(titulo: str, filas_d: list[dict], sufijo: str) -> None:
        if not filas_d:
            return
        copia = [dict(r) for r in filas_d]
        copia.sort(key=lambda r: r["lectura_fin"])
        for r in copia:
            r.pop("lectura_fin", None)
        img = OUT_DIR / f"chart_barras_icco_{sufijo}_{run_id}.png"
        _generar_grafico_barras(img, copia)
        doc.add_paragraph()
        pg = doc.add_paragraph()
        pg.add_run(titulo).bold = True
        doc.add_picture(str(img), width=Inches(6.7))

    _grafico_categoria("Gráfico — Tabla A (bypass)", filas_bypass_d, "bypass")
    _grafico_categoria("Gráfico — Tabla B (transición)", filas_transicion_d, "transicion")
    _grafico_categoria("Gráfico — Tabla C (alineados)", filas_alineado_d, "alineados")

    doc.add_paragraph()
    doc.add_heading(
        f"2.3 Total acumulado — cuenta (facturación) vs App WES (≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})",
        level=2,
    )
    doc.add_paragraph(
        f"Suma de todos los períodos incluidos en el cuadro (lecturas anterior y actual ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}). "
        "Total en cuenta: suma del m³ facturado en esas boletas; total App: suma del consumo WES entre las mismas lecturas de cada período."
    )
    rows_acum = filas_bypass_d + filas_transicion_d + filas_alineado_d
    n_boletas_acum = len(rows_acum)
    tot_cuenta = sum(float(r["m3_cuenta"]) for r in rows_acum)
    tot_wes = sum(float(r["m3_wes"]) for r in rows_acum)
    tot_diff = tot_cuenta - tot_wes
    pct_vs_fact = (100.0 * tot_diff / tot_cuenta) if tot_cuenta else 0.0

    ttot = doc.add_table(rows=1, cols=2)
    ttot.style = "Table Grid"
    ttot.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_pr_set_full_width_fixed(ttot)
    h0 = ttot.rows[0].cells
    h0[0].paragraphs[0].add_run("Indicador").bold = True
    h0[1].paragraphs[0].add_run("Valor").bold = True
    for lab, val in (
        ("Boletas incluidas en el acumulado (N°)", str(n_boletas_acum)),
        ("Total m³ facturación (cuenta, suma boletas)", _fmt_num(tot_cuenta, 1)),
        ("Total m³ App WES (suma períodos)", _fmt_num(tot_wes, 1)),
        ("Diferencia (facturación − App WES)", _fmt_num(tot_diff, 1)),
        ("% diferencia vs facturación", _fmt_num(pct_vs_fact, 1) + "%"),
    ):
        row_cells = ttot.add_row().cells
        row_cells[0].text = lab
        row_cells[1].text = val
    for r in ttot.rows:
        try:
            r.allow_break_across_pages = False
        except Exception:
            pass
        for c in r.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after = Pt(0)
                for run in pp.runs:
                    run.font.size = Pt(9)

    img_tot = OUT_DIR / f"chart_barras_icco_total_acumulado_{run_id}.png"
    _generar_grafico_totales(
        img_tot,
        total_wes=tot_wes,
        total_cuenta=tot_cuenta,
        n_boletas=n_boletas_acum,
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Gráfico — total acumulado: barra izquierda suma App WES, barra derecha suma facturación en cuenta "
        f"(misma base que la tabla; solo facturaciones con ambas lecturas ≥ {INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')})."
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if n_boletas_acum > 0:
        doc.add_picture(str(img_tot), width=Inches(5.8))

    doc.add_heading("3. Referencias", level=1)
    doc.add_paragraph(
        "Facturas incluidas en tablas, gráficos y total acumulado (lecturas ≥ "
        f"{INFORME_COMPARATIVO_DESDE.strftime('%d-%m-%Y')}), por lectura actual:\n- "
        + "\n- ".join([str(f.boleta) for f in sorted(facturas, key=lambda x: x.lectura_actual)])
    )

    out_docx = OUT_DIR / "Comparacion_App_vs_Facturaciones_ICCO_Renca.docx"
    # En Windows, sobrescribir el mismo nombre conserva a veces la “fecha de creación” antigua
    # del archivo en el Explorador; borrar antes evita confundir con un informe viejo.
    if out_docx.exists():
        try:
            out_docx.unlink()
        except OSError:
            pass
    doc.save(out_docx)
    _docx_parchear_core_hora_local(out_docx, ahora)
    out_pdf = out_docx.with_suffix(".pdf")
    if out_pdf.exists():
        try:
            out_pdf.unlink()
        except OSError:
            pass
    out_pdf = convertir_word_a_pdf(out_docx)
    return out_docx, out_pdf


def main() -> None:
    try:
        docx, pdf = generar_reporte()
    except FileNotFoundError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"[ERROR] {type(e).__name__}: {e}", file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)
    print(f"[OK] DOCX: {docx}")
    print(f"[OK] PDF: {pdf}" if pdf else "[ADVERTENCIA] No se generó PDF.")


if __name__ == "__main__":
    main()

