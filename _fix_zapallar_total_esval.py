"""
Ajusta Word Fundo Zapallar:
- Quita esquema hidráulico y RESUMEN/MÉTRICAS POR PUNTO.
- Corrige total: no sumar nodos (doble conteo). El consumo real del fundo = Matriz ESVAL.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", line_buffering=True)
    except Exception:
        pass

from generar_reporte_word import format_number_chilean

DOC = Path(
    "reports/Fundo_Zapallar/ABREGADO/AGREGADO_20260729_1607/"
    "Reporte_Agregado_Fundo_Zapallar_20260701_20260728.docx"
)

# Valores del reporte actual (periodo 01–28/07)
ESVAL_M3 = 2973.1
N_PUNTOS = 8
ALERTAS = 37


def _para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def _remove_from_until(doc: Document, start_pred, end_pred) -> int:
    body = doc.element.body
    children = list(body)
    start_idx = end_idx = None
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] != "p":
            continue
        t = _para_text(child)
        if start_idx is None and start_pred(t):
            start_idx = i
        elif start_idx is not None and end_pred(t):
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Rango no encontrado ({start_idx}, {end_idx})")
    n = 0
    for child in children[start_idx:end_idx]:
        body.remove(child)
        n += 1
    return n


def _set_para_text(para_elem, text: str) -> None:
    """Reemplaza el texto de un párrafo w:p conservando el primer run si existe."""
    from docx.text.paragraph import Paragraph

    # Find corresponding paragraph object
    # Simpler: clear all w:t and set on first, or rebuild runs via python-docx Paragraph
    pass


def main() -> None:
    doc = Document(str(DOC))

    # 1) Quitar esquema + RESUMEN/MÉTRICAS hasta ANÁLISIS NOCTURNO
    n1 = _remove_from_until(
        doc,
        lambda t: t.upper().startswith("ESQUEMA HIDR")
        or t.upper().startswith("RESUMEN POR PUNTO"),
        lambda t: t.upper().startswith("ANÁLISIS DE CONSUMOS NOCTURNOS")
        or t.upper().startswith("ANALISIS DE CONSUMOS NOCTURNOS"),
    )
    print(f"[OK] Eliminados {n1} elementos (esquema + resumen/métricas)")

    # Si quedó esquema pero no se unió al resumen (esquema ya borrado junto a resumen)
    # También borrar si solo quedó RESUMEN por si el esquema ya no estaba al inicio del rango
    # (ya cubierto arriba si start es ESQUEMA)

    # 2) Corregir resumen ejecutivo y narrativa del gráfico comparativo
    for child in list(doc.element.body):
        if child.tag.split("}")[-1] != "p":
            continue
        t = _para_text(child)
        tu = t.upper()

        if t.startswith("Puntos de monitoreo analizados:"):
            # Reescribir runs del párrafo
            from docx.text.paragraph import Paragraph

            para = Paragraph(child, doc)
            para.clear()
            para.add_run(f"Puntos de monitoreo analizados: {N_PUNTOS}.\n")
            para.add_run(
                f"Consumo del fundo (Matriz ESVAL / entrada): "
                f"{format_number_chilean(ESVAL_M3, 1)} m³.\n"
            )
            para.add_run(
                "Nota: no se suma el consumo de estanques y etapas al total, "
                "porque son puntos aguas abajo de ESVAL (mismo caudal medido en cadena).\n"
            )
            para.add_run(f"Total de alertas registradas: {ALERTAS}.\n")
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            print("[OK] Resumen ejecutivo corregido (total = Matriz ESVAL)")

        if "registró un consumo total de" in t and "Fundo Zapallar" in t:
            from docx.text.paragraph import Paragraph

            para = Paragraph(child, doc)
            para.clear()
            para.add_run(
                f"En Fundo Zapallar el consumo real del periodo corresponde a la Matriz ESVAL "
                f"(entrada de agua al fundo): {format_number_chilean(ESVAL_M3, 1)} m³. "
                f"Los demás puntos (estanques y etapas) miden caudales aguas abajo de esa matriz; "
                f"por eso no se suman entre sí ni con ESVAL — el máximo del gráfico es Matriz ESVAL, "
                f"no la suma de barras (~8.000 m³), que no representa un consumo adicional."
            )
            para.alignment = 1  # justify if possible
            try:
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except Exception:
                pass
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            print("[OK] Narrativa del gráfico corregida")

        if tu.startswith("CONSUMO ACUMULADO DE CADA PUNTO"):
            from docx.text.paragraph import Paragraph

            para = Paragraph(child, doc)
            para.clear()
            para.add_run(
                "Consumo acumulado por punto en el periodo. "
                "Matriz ESVAL es la entrada al fundo (referencia de consumo real); "
                "estanques y etapas son mediciones aguas abajo y no deben sumarse al total."
            )
            try:
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except Exception:
                pass
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Subtítulo cabecera: a veces dice "consolidado" sin detallar — dejar
    try:
        doc.save(str(DOC))
        out = DOC
    except PermissionError:
        out = DOC.with_name(DOC.stem + "_esval.docx")
        doc.save(str(out))
        print(f"[ADVERTENCIA] Guardado como {out.name}")

    print(f"[OK] {out}")


if __name__ == "__main__":
    main()
