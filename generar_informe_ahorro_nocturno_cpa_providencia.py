"""
Informe concreto por colegio Providencia:

  El colegio X consumió Y m³ en horario nocturno desde la implementación.
  La gestión hídrica CPA permite llevar a cero ese consumo nocturno.
  Proyección 30 días = Z m³ / CLP.
  Esa proyección × meses de funcionamiento = ahorro en todo el periodo.

Uso:
  python generar_informe_ahorro_nocturno_cpa_providencia.py
"""

from __future__ import annotations

import csv
import sys
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

if sys.platform == "win32":
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

ROOT = Path(__file__).resolve().parent
COMPANY_ID = "000006"
PRECIO_DEFAULT = 1274.0

# Primera data (entrega / implementación) por nodo
NODOS = [
    ("000006-01", "Liceo Lastarria", "01/10/2025", True),
    ("000006-02", "Carmela Carvajal", "03/10/2025", True),
    ("000006-04", "Liceo 7 Luisa Saavedra", "03/10/2025", True),
    ("000006-05", "Liceo Juan Pablo Duarte", "06/10/2025", True),
    ("000006-03", "Arturo Alessandri Palma", "01/10/2025", False),  # sin data reciente
]

# Hasta último mes completo (proyección limpia)
HASTA = "31/07/2026"
HASTA_TXT = "31/07/2026"


def _fmt_m3(v: float, d: int = 1) -> str:
    from generar_reporte_word import format_number_chilean

    return format_number_chilean(v, d)


def _fmt_clp(v: float) -> str:
    from generar_reporte_word import format_currency_chilean

    return format_currency_chilean(v)


def _grafico_barras(
    filas: List[dict],
    out_path: Path,
    *,
    campo: str,
    ylabel: str,
    titulo: str,
) -> None:
    labels = [f["corto"] for f in filas if f["activo"]]
    vals = [f[campo] for f in filas if f["activo"]]
    fig, ax = plt.subplots(figsize=(8.8, 4.6))
    bars = ax.bar(labels, vals, color="#1f4788", edgecolor="white")
    ymax = max(vals) if vals else 1
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + ymax * 0.02,
            f"{_fmt_m3(v, 1)} m³",
            ha="center",
            va="bottom",
            fontsize=9,
        )
    ax.set_ylabel(ylabel)
    ax.set_title(titulo)
    ax.grid(True, axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def generar(output_dir: Optional[Path] = None) -> Path:
    from generar_reporte_word import (
        calculate_nocturnal_metrics,
        get_water_price_per_m3,
        parse_date,
    )

    out_dir = output_dir or (
        ROOT / "reports" / "Providencia" / "ahorro_nocturno_cpa"
    )
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")

    end_dt = parse_date(HASTA, end_of_day=True)
    precio = get_water_price_per_m3(COMPANY_ID, NODOS[0][0], None) or PRECIO_DEFAULT

    filas: List[dict] = []
    print(f"[INFO] Precio ref: {_fmt_clp(precio)}/m³ | hasta {HASTA_TXT}")

    for nid, nombre, desde, activo in NODOS:
        start_dt = parse_date(desde)
        print(f"  Calculando nocturno {nid} {nombre} desde {desde}...")
        m = calculate_nocturnal_metrics(nid, start_dt, end_dt)
        noct = float(m["consumo_nocturno_total"])
        d_datos = int(m["dias_con_datos_horarios"])
        d_con = int(m["dias_con_consumo_nocturno"])
        dias_cal = (end_dt.date() - start_dt.date()).days + 1

        # Ahorro = llevar nocturno a cero.
        # Proyección 30 d = promedio diario nocturno (sobre días con datos) × 30.
        # Periodo completo = esa proyección × meses de funcionamiento (días cal / 30).
        prom_dia = (noct / d_datos) if d_datos > 0 else 0.0
        proy_30 = prom_dia * 30.0
        proy_clp = proy_30 * precio
        acum_clp = noct * precio
        meses_func = dias_cal / 30.0
        proy_periodo_m3 = proy_30 * meses_func
        proy_periodo_clp = proy_periodo_m3 * precio

        corto = (
            nombre.replace("Liceo ", "")
            .replace("Arturo Alessandri Palma", "Alessandri")
            .replace("Luisa Saavedra", "7")
            .replace("Juan Pablo Duarte", "Duarte")
        )
        filas.append(
            {
                "node_id": nid,
                "nombre": nombre,
                "corto": corto,
                "desde": desde,
                "activo": activo and d_datos > 30,
                "noct_m3": noct,
                "noct_clp": acum_clp,
                "dias_datos": d_datos,
                "dias_cal": dias_cal,
                "dias_con_noche": d_con,
                "prom_dia": prom_dia,
                "proy_30_m3": proy_30,
                "proy_30_clp": proy_clp,
                "meses_func": meses_func,
                "proy_periodo_m3": proy_periodo_m3,
                "proy_periodo_clp": proy_periodo_clp,
            }
        )
        print(
            f"    noct={noct:.1f} m³ | proy30={proy_30:.1f} m³ | "
            f"meses={meses_func:.1f} | periodo={proy_periodo_m3:.1f} m³ | "
            f"{_fmt_clp(proy_periodo_clp)}"
        )

    activos = [f for f in filas if f["activo"]]
    png = out_dir / f"ahorro_proy_30d_{ts}.png"
    _grafico_barras(
        filas,
        png,
        campo="proy_30_m3",
        ylabel="Ahorro proyección 30 días (m³)",
        titulo="Ahorro mensual si la gestión CPA lleva a cero el consumo nocturno",
    )
    png_periodo = out_dir / f"ahorro_proy_periodo_{ts}.png"
    _grafico_barras(
        filas,
        png_periodo,
        campo="proy_periodo_m3",
        ylabel="Ahorro en todo el periodo (m³)",
        titulo="Proyección 30 d × meses de funcionamiento — ahorro total CPA",
    )

    # CSV
    csv_path = out_dir / f"ahorro_nocturno_cpa_providencia_{ts}.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh, delimiter=";")
        w.writerow(
            [
                "node_id",
                "colegio",
                "implementacion",
                "hasta",
                "consumo_nocturno_acumulado_m3",
                "valor_acumulado_clp",
                "dias_calendario",
                "meses_funcionamiento",
                "dias_con_datos",
                "promedio_diario_nocturno_m3",
                "ahorro_proyeccion_30d_m3",
                "ahorro_proyeccion_30d_clp",
                "ahorro_proyeccion_periodo_m3",
                "ahorro_proyeccion_periodo_clp",
            ]
        )
        for f in filas:
            w.writerow(
                [
                    f["node_id"],
                    f["nombre"],
                    f["desde"],
                    HASTA_TXT,
                    f"{f['noct_m3']:.2f}".replace(".", ","),
                    f"{f['noct_clp']:.0f}",
                    f["dias_cal"],
                    f"{f['meses_func']:.2f}".replace(".", ","),
                    f["dias_datos"],
                    f"{f['prom_dia']:.3f}".replace(".", ","),
                    f"{f['proy_30_m3']:.2f}".replace(".", ","),
                    f"{f['proy_30_clp']:.0f}",
                    f"{f['proy_periodo_m3']:.2f}".replace(".", ","),
                    f"{f['proy_periodo_clp']:.0f}",
                ]
            )

    # Word — mensaje concreto
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading(
        "Ahorro nocturno CPA — Colegios Providencia", 0
    )
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    intro = doc.add_paragraph()
    intro.add_run("Premisa: ").bold = True
    intro.add_run(
        "desde la implementación del medidor WES se registra consumo entre "
        "00:00 y 06:59 (hora Chile). La gestión hídrica CPA permite llevar ese "
        "consumo nocturno a cero. "
        "1) Proyección 30 días = promedio diario nocturno × 30. "
        "2) Ahorro en todo el periodo = esa proyección × meses de funcionamiento "
        f"(días calendario ÷ 30). Valorizado a {_fmt_clp(precio)}/m³."
    )

    meta = doc.add_paragraph()
    meta.add_run("Periodo de medición: ").bold = True
    meta.add_run(
        f"desde la fecha de implementación de cada colegio hasta el {HASTA_TXT} "
        "(último mes completo)."
    )
    gen = doc.add_paragraph()
    gen.add_run("Generado: ").bold = True
    gen.add_run(datetime.now().strftime("%d-%m-%Y %H:%M"))

    # Ficha por colegio (solo activos + Alessandri con nota)
    doc.add_heading("1. Ficha por colegio", level=1)

    for i, f in enumerate(filas, 1):
        doc.add_heading(f"{i}. {f['nombre']} ({f['node_id']})", level=2)

        if not f["activo"]:
            doc.add_paragraph(
                f"Implementación / primera data: {f['desde']}. "
                f"Consumo nocturno acumulado mientras hubo datos: "
                f"{_fmt_m3(f['noct_m3'], 1)} m³ "
                f"({_fmt_clp(f['noct_clp'])}). "
                "Hoy no hay serie reciente (corte ~enero 2026); "
                "no se entrega proyección hasta reactivar el punto."
            )
            continue

        # Bloque concreto pedido por el usuario
        p = doc.add_paragraph()
        p.add_run(
            f"Desde la implementación ({f['desde']}) hasta el {HASTA_TXT}, "
            f"{f['nombre']} consumió {_fmt_m3(f['noct_m3'], 1)} m³ en horario "
            f"nocturno (00:00–06:59), equivalentes a {_fmt_clp(f['noct_clp'])} "
            f"a tarifa de referencia."
        )

        p2 = doc.add_paragraph()
        p2.add_run(
            "La gestión hídrica CPA permite llevar a cero ese consumo nocturno "
            "(corte / control fuera de horario operativo)."
        )

        p3 = doc.add_paragraph()
        run = p3.add_run(
            f"Proyección 30 días — ahorro: {_fmt_m3(f['proy_30_m3'], 1)} m³ "
            f"= {_fmt_clp(f['proy_30_clp'])}."
        )
        run.bold = True

        p4 = doc.add_paragraph()
        run4 = p4.add_run(
            f"Llevado a {_fmt_m3(f['meses_func'], 1)} meses de funcionamiento "
            f"(proyección 30 d × meses): ahorro "
            f"{_fmt_m3(f['proy_periodo_m3'], 1)} m³ "
            f"= {_fmt_clp(f['proy_periodo_clp'])}."
        )
        run4.bold = True

        det = doc.add_paragraph()
        det.add_run("Detalle: ").bold = True
        det.add_run(
            f"{f['dias_cal']} días calendario = {_fmt_m3(f['meses_func'], 1)} meses; "
            f"{f['dias_datos']} días con datos horarios; "
            f"{f['dias_con_noche']} días con consumo nocturno > 0; "
            f"promedio {_fmt_m3(f['prom_dia'], 2)} m³/día nocturno × 30 = "
            f"proy. mensual; × {_fmt_m3(f['meses_func'], 1)} meses = periodo."
        )

    # Resumen tabla 30 d
    doc.add_heading("2. Resumen — proyección 30 días (colegios activos)", level=1)
    tbl = doc.add_table(rows=1 + len(activos) + 1, cols=5)
    tbl.style = "Table Grid"
    headers = [
        "Colegio",
        "Nocturno medido (m³)",
        "Meses func.",
        "Ahorro 30 d (m³)",
        "Ahorro 30 d (CLP)",
    ]
    for j, hd in enumerate(headers):
        tbl.rows[0].cells[j].text = hd
        for run in tbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    tot_noct = tot_proy = tot_periodo = 0.0
    for i, f in enumerate(activos, 1):
        tot_noct += f["noct_m3"]
        tot_proy += f["proy_30_m3"]
        tot_periodo += f["proy_periodo_m3"]
        row = tbl.rows[i].cells
        row[0].text = f["nombre"]
        row[1].text = _fmt_m3(f["noct_m3"], 1)
        row[2].text = _fmt_m3(f["meses_func"], 1)
        row[3].text = _fmt_m3(f["proy_30_m3"], 1)
        row[4].text = _fmt_clp(f["proy_30_clp"])

    tot_row = tbl.rows[len(activos) + 1].cells
    tot_row[0].text = "TOTAL (4 colegios)"
    tot_row[1].text = _fmt_m3(tot_noct, 1)
    tot_row[2].text = "—"
    tot_row[3].text = _fmt_m3(tot_proy, 1)
    tot_row[4].text = _fmt_clp(tot_proy * precio)
    for c in tot_row:
        for run in c.paragraphs[0].runs:
            run.bold = True

    if png.is_file():
        doc.add_paragraph()
        doc.add_picture(str(png), width=Inches(5.8))

    # Resumen periodo completo
    doc.add_heading(
        "3. Resumen — proyección llevada a meses de funcionamiento", level=1
    )
    doc.add_paragraph(
        "Fórmula: ahorro periodo = (proyección 30 días) × (días calendario ÷ 30). "
        "Es el ahorro total si la gestión CPA hubiera llevado a cero el nocturno "
        "durante todo el tiempo de operación del medidor."
    )
    tbl2 = doc.add_table(rows=1 + len(activos) + 1, cols=5)
    tbl2.style = "Table Grid"
    headers2 = [
        "Colegio",
        "Meses func.",
        "Ahorro 30 d (m³)",
        "Ahorro periodo (m³)",
        "Ahorro periodo (CLP)",
    ]
    for j, hd in enumerate(headers2):
        tbl2.rows[0].cells[j].text = hd
        for run in tbl2.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, f in enumerate(activos, 1):
        row = tbl2.rows[i].cells
        row[0].text = f["nombre"]
        row[1].text = _fmt_m3(f["meses_func"], 1)
        row[2].text = _fmt_m3(f["proy_30_m3"], 1)
        row[3].text = _fmt_m3(f["proy_periodo_m3"], 1)
        row[4].text = _fmt_clp(f["proy_periodo_clp"])

    tot2 = tbl2.rows[len(activos) + 1].cells
    tot2[0].text = "TOTAL (4 colegios)"
    tot2[1].text = "—"
    tot2[2].text = _fmt_m3(tot_proy, 1)
    tot2[3].text = _fmt_m3(tot_periodo, 1)
    tot2[4].text = _fmt_clp(tot_periodo * precio)
    for c in tot2:
        for run in c.paragraphs[0].runs:
            run.bold = True

    if png_periodo.is_file():
        doc.add_paragraph()
        doc.add_picture(str(png_periodo), width=Inches(5.8))

    doc.add_heading("4. Mensaje para el cliente", level=1)
    cierre = doc.add_paragraph()
    cierre.add_run(
        f"En los 4 colegios con data continua se midieron "
        f"{_fmt_m3(tot_noct, 1)} m³ nocturnos "
        f"({_fmt_clp(tot_noct * precio)}). "
        f"Al ritmo actual, el ahorro CPA a 30 días es "
        f"{_fmt_m3(tot_proy, 1)} m³ ({_fmt_clp(tot_proy * precio)}). "
        f"Llevado a los meses de funcionamiento de cada colegio, el ahorro "
        f"proyectado del periodo es {_fmt_m3(tot_periodo, 1)} m³ "
        f"({_fmt_clp(tot_periodo * precio)})."
    )

    nota = doc.add_paragraph()
    nota.add_run("Nota: ").bold = True
    nota.add_run(
        "el promedio diario incluye vacaciones y baja ocupación; "
        "la proyección es conservadora. "
        "Meses de funcionamiento = días calendario desde implementación ÷ 30."
    )

    out_docx = out_dir / f"Ahorro_nocturno_CPA_Providencia_{ts}.docx"
    doc.save(out_docx)

    # PDF
    pdf_path = out_docx.with_suffix(".pdf")
    try:
        import subprocess

        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(out_docx),
            ],
            check=False,
            capture_output=True,
            timeout=120,
        )
    except Exception as ex:
        print(f"[WARN] PDF: {ex}")

    print(f"[OK] Word: {out_docx}")
    print(f"[OK] PDF:  {pdf_path if pdf_path.is_file() else 'no'}")
    print(f"[OK] CSV:  {csv_path}")
    print(
        f"[OK] TOTAL proy 30d: {_fmt_m3(tot_proy, 1)} m³ = {_fmt_clp(tot_proy * precio)}"
    )
    print(
        f"[OK] TOTAL periodo: {_fmt_m3(tot_periodo, 1)} m³ = "
        f"{_fmt_clp(tot_periodo * precio)}"
    )
    return out_docx


def main() -> int:
    print("=" * 72)
    print("AHORRO NOCTURNO CPA — COLEGIOS PROVIDENCIA (concreto)")
    print("=" * 72)
    generar()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
