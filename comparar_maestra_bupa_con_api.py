"""
Script para comparar los datos del archivo Excel "Ficha maestra Bupa.xlsx"
con los datos de la API WES y generar un reporte Word con las inconsistencias encontradas.
"""

import sys
from pathlib import Path
from datetime import datetime, timedelta
import pandas as pd
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

# Agregar el directorio actual al path para importar módulos
sys.path.insert(0, str(Path(__file__).parent))

from generar_reporte_word import (
    get_company_name,
    get_node_name,
    format_number_chilean,
    add_logo_to_header,
    BASE_URL,
    ENTITY_BASE_URL,
    fetch_json,
    parse_date,
    normalize_measures_payload,
    flatten_measures,
    summarize_consumption,
)

# Configuración
COMPANY_ID = "000029"  # BUPA
EXCEL_PATH = Path("Analisis BUpa/Ficha maestra Bupa.xlsx")
OUTPUT_DIR = Path("Analisis BUpa")

def obtener_nodos_bupa_desde_api() -> dict:
    """
    Obtiene todos los nodos de BUPA desde la API WES.
    Retorna un diccionario {node_id: node_name}
    """
    url = f"{ENTITY_BASE_URL}/companies/{COMPANY_ID}"
    
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()
            nodes = data.get("nodes", [])
            return {
                node.get("nodeId", ""): node.get("name", "").strip()
                for node in nodes
                if node.get("nodeId") and node.get("name")
            }
        else:
            print(f"[ERROR] No se pudo obtener información de la empresa {COMPANY_ID}: {response.status_code}")
            return {}
    except Exception as e:
        print(f"[ERROR] Error al obtener nodos de la empresa {COMPANY_ID}: {e}")
        return {}

def normalizar_nombre_dispositivo(nombre_excel: str) -> str:
    """
    Normaliza el nombre del dispositivo del Excel para compararlo con la API.
    """
    if pd.isna(nombre_excel):
        return ""
    
    nombre = str(nombre_excel).strip()
    # Corregir errores comunes
    nombre = nombre.replace("martesatriz", "matriz")
    nombre = nombre.replace("Matriz Principal Bupa", "Matriz Principal BUPA")
    return nombre

def mapear_dispositivo_a_nodo(dispositivo_excel: str, nodos_api: dict) -> tuple:
    """
    Intenta mapear un dispositivo del Excel a un node_id de la API.
    Retorna (node_id, node_name) o (None, None) si no se encuentra.
    """
    dispositivo_normalizado = normalizar_nombre_dispositivo(dispositivo_excel)
    dispositivo_lower = dispositivo_normalizado.lower()
    
    # Mapeo manual conocido - "matriz Principal Bupa" generalmente se refiere al consumo total
    # o al nodo principal. En BUPA, puede ser "Llenado de Estanques" o necesitamos sumar todos los nodos
    # Por ahora, mapeamos a "Llenado de Estanques" como el nodo principal
    mapeos_manuales = {
        "matriz principal bupa": "000029-01",  # Llenado de Estanques
        "matriz principal": "000029-01",
        "principal bupa": "000029-01",
    }
    
    # Buscar en mapeos manuales primero
    for key, node_id in mapeos_manuales.items():
        if key in dispositivo_lower:
            if node_id in nodos_api:
                return (node_id, nodos_api[node_id])
    
    # Buscar coincidencias parciales en los nombres de nodos
    for node_id, node_name in nodos_api.items():
        node_name_lower = node_name.lower()
        
        # Buscar palabras clave comunes
        if "matriz" in dispositivo_lower:
            # Si el dispositivo tiene "matriz", puede ser el nodo principal
            # En BUPA, "Llenado de Estanques" es el nodo principal
            if "llenado" in node_name_lower or "estanque" in node_name_lower:
                return (node_id, node_name)
            if "matriz" in node_name_lower:
                return (node_id, node_name)
        
        if "torre" in dispositivo_lower and "torre" in node_name_lower:
            # Intentar mapear por letra
            if (" a" in dispositivo_lower or dispositivo_lower.endswith(" a")) and "torre a" in node_name_lower:
                return (node_id, node_name)
            if (" b1" in dispositivo_lower or " b 1" in dispositivo_lower or dispositivo_lower.endswith(" b1")) and "b1" in node_name_lower:
                return (node_id, node_name)
            if (" b2" in dispositivo_lower or " b 2" in dispositivo_lower or dispositivo_lower.endswith(" b2")) and "b2" in node_name_lower:
                return (node_id, node_name)
            if (" c" in dispositivo_lower or dispositivo_lower.endswith(" c")) and "torre c" in node_name_lower:
                return (node_id, node_name)
        
        if "central" in dispositivo_lower and "central" in node_name_lower:
            return (node_id, node_name)
    
    # Si no se encuentra mapeo específico y el dispositivo contiene "matriz" o "principal",
    # intentar usar el primer nodo (generalmente el principal)
    if "matriz" in dispositivo_lower or "principal" in dispositivo_lower:
        # Retornar el nodo "Llenado de Estanques" como principal
        if "000029-01" in nodos_api:
            return ("000029-01", nodos_api["000029-01"])
    
    return (None, None)

def obtener_consumo_total_todos_nodos(fecha: datetime, nodos_api: dict) -> dict:
    """
    Obtiene el consumo total sumando todos los nodos de BUPA para una fecha.
    Retorna un diccionario con los datos o None si hay error.
    """
    consumo_total = 0.0
    nodos_con_datos = []
    nodos_sin_datos = []
    
    for node_id in nodos_api.keys():
        datos = obtener_consumo_desde_api(node_id, fecha)
        if datos is not None:
            consumo_nodo = datos.get("total", 0.0)
            consumo_total += consumo_nodo
            nodos_con_datos.append({
                "node_id": node_id,
                "node_name": nodos_api[node_id],
                "consumo": consumo_nodo
            })
        else:
            nodos_sin_datos.append({
                "node_id": node_id,
                "node_name": nodos_api[node_id]
            })
    
    return {
        "total": consumo_total,
        "nodos_con_datos": nodos_con_datos,
        "nodos_sin_datos": nodos_sin_datos
    }

def obtener_consumo_desde_api(node_id: str, fecha: datetime) -> dict:
    """
    Obtiene el consumo de un nodo para una fecha específica desde la API.
    Retorna un diccionario con los datos o None si hay error.
    """
    try:
        fecha_str = fecha.strftime("%d%m%Y")
        
        measures_payload_raw = fetch_json(
            f"{BASE_URL}/nodes/measures/dates",
            params=[
                ("id", node_id),
                ("start", fecha_str),
                ("end", fecha_str),
            ],
        )
        
        measures_payload = normalize_measures_payload(measures_payload_raw, node_id)
        measures = flatten_measures(measures_payload)
        summary = summarize_consumption(measures)
        
        return {
            "total": summary.get("total", 0.0),
            "promedio_diario": summary.get("promedio_diario", 0.0),
            "medidas": measures,
        }
    except Exception as e:
        print(f"[ERROR] Error al obtener consumo para {node_id} en {fecha.strftime('%Y-%m-%d')}: {e}")
        return None

def leer_excel() -> pd.DataFrame:
    """
    Lee el archivo Excel y retorna un DataFrame limpio.
    """
    try:
        df = pd.read_excel(EXCEL_PATH)
        
        # Limpiar columnas vacías
        df = df.dropna(axis=1, how='all')
        
        # Limpiar filas completamente vacías
        df = df.dropna(how='all')
        
        # Convertir fecha a datetime si es necesario
        if 'fecha' in df.columns:
            df['fecha'] = pd.to_datetime(df['fecha'], errors='coerce')
        
        return df
    except Exception as e:
        print(f"[ERROR] No se pudo leer el archivo Excel: {e}")
        raise

def obtener_consumo_total_todos_nodos(fecha: datetime, nodos_api: dict) -> dict:
    """
    Obtiene el consumo total sumando todos los nodos de BUPA para una fecha.
    Retorna un diccionario con los datos o None si hay error.
    """
    consumo_total = 0.0
    nodos_con_datos = []
    nodos_sin_datos = []
    
    for node_id in nodos_api.keys():
        datos = obtener_consumo_desde_api(node_id, fecha)
        if datos is not None:
            consumo_nodo = datos.get("total", 0.0)
            consumo_total += consumo_nodo
            nodos_con_datos.append({
                "node_id": node_id,
                "node_name": nodos_api[node_id],
                "consumo": consumo_nodo
            })
        else:
            nodos_sin_datos.append({
                "node_id": node_id,
                "node_name": nodos_api[node_id]
            })
    
    return {
        "total": consumo_total,
        "nodos_con_datos": nodos_con_datos,
        "nodos_sin_datos": nodos_sin_datos
    }

def comparar_datos(excel_df: pd.DataFrame, nodos_api: dict) -> list:
    """
    Compara los datos del Excel con los de la API.
    Retorna una lista de inconsistencias encontradas.
    """
    inconsistencias = []
    
    # Agrupar por dispositivo y fecha
    for dispositivo in excel_df['dispositivo'].dropna().unique():
        dispositivo_normalizado = normalizar_nombre_dispositivo(dispositivo)
        dispositivo_lower = dispositivo_normalizado.lower()
        
        # Determinar si es consumo total (matriz/principal) o de un nodo específico
        es_consumo_total = "matriz" in dispositivo_lower or "principal" in dispositivo_lower
        
        # Mapear dispositivo a node_id (si no es consumo total)
        node_id, node_name = None, None
        if not es_consumo_total:
            node_id, node_name = mapear_dispositivo_a_nodo(dispositivo, nodos_api)
        
        # Obtener filas de este dispositivo
        filas_dispositivo = excel_df[excel_df['dispositivo'] == dispositivo]
        
        for idx, fila in filas_dispositivo.iterrows():
            fecha_excel = fila.get('fecha')
            consumo_excel = fila.get('M3 dia')
            
            # Saltar si no hay fecha o consumo válido
            if pd.isna(fecha_excel) or pd.isna(consumo_excel):
                continue
            
            try:
                fecha_dt = pd.to_datetime(fecha_excel)
                consumo_excel_float = float(consumo_excel)
                
                if es_consumo_total:
                    # Si es consumo total, sumar todos los nodos
                    datos_api = obtener_consumo_total_todos_nodos(fecha_dt, nodos_api)
                    
                    if datos_api is None or len(datos_api.get("nodos_con_datos", [])) == 0:
                        inconsistencias.append({
                            "tipo": "SIN_DATOS_API",
                            "dispositivo_excel": dispositivo,
                            "node_id": "TODOS",
                            "node_name": "Suma de todos los nodos",
                            "fecha": fecha_dt.strftime("%Y-%m-%d"),
                            "consumo_excel": consumo_excel_float,
                            "consumo_api": None,
                            "diferencia": None,
                            "mensaje": f"No se pudieron obtener datos de la API para ningún nodo en {fecha_dt.strftime('%Y-%m-%d')}"
                        })
                        continue
                    
                    consumo_api = datos_api.get("total", 0.0)
                    nodos_sin_datos = datos_api.get("nodos_sin_datos", [])
                    
                    if nodos_sin_datos:
                        inconsistencias.append({
                            "tipo": "NODOS_SIN_DATOS",
                            "dispositivo_excel": dispositivo,
                            "fecha": fecha_dt.strftime("%Y-%m-%d"),
                            "nodos_sin_datos": nodos_sin_datos,
                            "mensaje": f"Algunos nodos no tienen datos en {fecha_dt.strftime('%Y-%m-%d')}: {', '.join([n['node_name'] for n in nodos_sin_datos])}"
                        })
                    
                else:
                    # Si es un nodo específico
                    if node_id is None:
                        inconsistencias.append({
                            "tipo": "DISPOSITIVO_NO_ENCONTRADO",
                            "dispositivo_excel": dispositivo,
                            "dispositivo_normalizado": dispositivo_normalizado,
                            "fecha": fecha_dt.strftime("%Y-%m-%d"),
                            "consumo_excel": consumo_excel_float,
                            "consumo_api": None,
                            "diferencia": None,
                            "mensaje": f"El dispositivo '{dispositivo}' no se pudo mapear a ningún nodo de la API"
                        })
                        continue
                    
                    # Obtener consumo desde API
                    datos_api = obtener_consumo_desde_api(node_id, fecha_dt)
                    
                    if datos_api is None:
                        inconsistencias.append({
                            "tipo": "SIN_DATOS_API",
                            "dispositivo_excel": dispositivo,
                            "node_id": node_id,
                            "node_name": node_name,
                            "fecha": fecha_dt.strftime("%Y-%m-%d"),
                            "consumo_excel": consumo_excel_float,
                            "consumo_api": None,
                            "diferencia": None,
                            "mensaje": f"No se pudieron obtener datos de la API para {node_name} ({node_id}) en {fecha_dt.strftime('%Y-%m-%d')}"
                        })
                        continue
                    
                    consumo_api = datos_api.get("total", 0.0)
                
                # Calcular diferencia
                diferencia = abs(consumo_excel_float - consumo_api)
                
                # Tolerancia: considerar inconsistencia si la diferencia es mayor a 1 m³ o 5%
                tolerancia_absoluta = 1.0
                tolerancia_porcentual = 0.05
                
                es_inconsistente = False
                if consumo_excel_float > 0:
                    diferencia_porcentual = diferencia / consumo_excel_float
                    es_inconsistente = diferencia > tolerancia_absoluta or diferencia_porcentual > tolerancia_porcentual
                else:
                    es_inconsistente = diferencia > tolerancia_absoluta
                
                if es_inconsistente:
                    tipo_comparacion = "DIFERENCIA_CONSUMO_TOTAL" if es_consumo_total else "DIFERENCIA_CONSUMO"
                    inconsistencias.append({
                        "tipo": tipo_comparacion,
                        "dispositivo_excel": dispositivo,
                        "node_id": "TODOS" if es_consumo_total else node_id,
                        "node_name": "Suma de todos los nodos" if es_consumo_total else node_name,
                        "fecha": fecha_dt.strftime("%Y-%m-%d"),
                        "consumo_excel": consumo_excel_float,
                        "consumo_api": consumo_api,
                        "diferencia": diferencia,
                        "diferencia_porcentual": (diferencia / consumo_excel_float * 100) if consumo_excel_float > 0 else 0,
                        "mensaje": f"Diferencia de {format_number_chilean(diferencia, 2)} m³ ({format_number_chilean(diferencia / consumo_excel_float * 100, 1) if consumo_excel_float > 0 else 0}%)"
                    })
                
            except Exception as e:
                inconsistencias.append({
                    "tipo": "ERROR_PROCESAMIENTO",
                    "dispositivo_excel": dispositivo,
                    "fecha": str(fecha_excel) if not pd.isna(fecha_excel) else "N/A",
                    "error": str(e),
                    "mensaje": f"Error al procesar la fila: {e}"
                })
    
    return inconsistencias

def generar_reporte_word(inconsistencias: list, nodos_api: dict, excel_df: pd.DataFrame):
    """
    Genera un documento Word con las inconsistencias encontradas.
    """
    doc = Document()
    
    # Agregar logo
    add_logo_to_header(doc)
    
    # Título
    title = doc.add_heading("Reporte de Inconsistencias: Maestra BUPA vs API WES", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Información general
    doc.add_paragraph("")
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph(f"Empresa: {get_company_name(COMPANY_ID)} (ID: {COMPANY_ID})")
    doc.add_paragraph(f"Archivo Excel analizado: {EXCEL_PATH.name}")
    doc.add_paragraph(f"Total de inconsistencias encontradas: {len(inconsistencias)}")
    
    doc.add_page_break()
    
    # Resumen por tipo
    if inconsistencias:
        doc.add_heading("Resumen por Tipo de Inconsistencia", 1)
        
        tipos = {}
        for inc in inconsistencias:
            tipo = inc.get("tipo", "DESCONOCIDO")
            tipos[tipo] = tipos.get(tipo, 0) + 1
        
        tabla_resumen = doc.add_table(rows=1, cols=2)
        tabla_resumen.style = 'Light Grid Accent 1'
        
        # Encabezados
        header_cells = tabla_resumen.rows[0].cells
        header_cells[0].text = "Tipo de Inconsistencia"
        header_cells[1].text = "Cantidad"
        
        for tipo, cantidad in sorted(tipos.items()):
            row_cells = tabla_resumen.add_row().cells
            row_cells[0].text = tipo
            row_cells[1].text = str(cantidad)
        
        doc.add_page_break()
    
    # Dispositivos no encontrados
    dispositivos_no_encontrados = [inc for inc in inconsistencias if inc.get("tipo") == "DISPOSITIVO_NO_ENCONTRADO"]
    if dispositivos_no_encontrados:
        doc.add_heading("Dispositivos del Excel No Encontrados en la API", 1)
        doc.add_paragraph("Los siguientes dispositivos del Excel no se pudieron mapear a ningún nodo de la API WES:")
        
        for inc in dispositivos_no_encontrados:
            p = doc.add_paragraph(f"• {inc['dispositivo_excel']}", style='List Bullet')
            if inc.get('dispositivo_normalizado') and inc['dispositivo_normalizado'] != inc['dispositivo_excel']:
                p.add_run(f" (normalizado: {inc['dispositivo_normalizado']})")
        
        doc.add_paragraph("")
        doc.add_paragraph("Nodos disponibles en la API WES:")
        for node_id, node_name in sorted(nodos_api.items()):
            doc.add_paragraph(f"  - {node_id}: {node_name}", style='List Bullet')
        
        doc.add_page_break()
    
    # Diferencias de consumo total
    diferencias_consumo_total = [inc for inc in inconsistencias if inc.get("tipo") == "DIFERENCIA_CONSUMO_TOTAL"]
    if diferencias_consumo_total:
        doc.add_heading("Diferencias de Consumo Total (Matriz Principal)", 1)
        doc.add_paragraph("Se encontraron diferencias entre el consumo total del Excel y la suma de todos los nodos en la API:")
        
        tabla = doc.add_table(rows=1, cols=6)
        tabla.style = 'Light Grid Accent 1'
        
        # Encabezados
        header_cells = tabla.rows[0].cells
        header_cells[0].text = "Fecha"
        header_cells[1].text = "Dispositivo (Excel)"
        header_cells[2].text = "Consumo Excel (m³)"
        header_cells[3].text = "Consumo API Total (m³)"
        header_cells[4].text = "Diferencia"
        header_cells[5].text = "Diferencia %"
        
        for inc in sorted(diferencias_consumo_total, key=lambda x: x.get("fecha", "")):
            row_cells = tabla.add_row().cells
            row_cells[0].text = inc.get("fecha", "N/A")
            row_cells[1].text = inc.get("dispositivo_excel", "N/A")
            row_cells[2].text = format_number_chilean(inc.get("consumo_excel", 0), 2)
            row_cells[3].text = format_number_chilean(inc.get("consumo_api", 0), 2)
            row_cells[4].text = format_number_chilean(inc.get("diferencia", 0), 2)
            row_cells[5].text = f"{format_number_chilean(inc.get('diferencia_porcentual', 0), 1)}%"
        
        doc.add_page_break()
    
    # Diferencias de consumo
    diferencias_consumo = [inc for inc in inconsistencias if inc.get("tipo") == "DIFERENCIA_CONSUMO"]
    if diferencias_consumo:
        doc.add_heading("Diferencias de Consumo", 1)
        doc.add_paragraph("Se encontraron diferencias significativas entre los consumos del Excel y los de la API:")
        
        tabla = doc.add_table(rows=1, cols=6)
        tabla.style = 'Light Grid Accent 1'
        
        # Encabezados
        header_cells = tabla.rows[0].cells
        header_cells[0].text = "Fecha"
        header_cells[1].text = "Dispositivo (Excel)"
        header_cells[2].text = "Nodo (API)"
        header_cells[3].text = "Consumo Excel (m³)"
        header_cells[4].text = "Consumo API (m³)"
        header_cells[5].text = "Diferencia"
        
        for inc in sorted(diferencias_consumo, key=lambda x: x.get("fecha", "")):
            row_cells = tabla.add_row().cells
            row_cells[0].text = inc.get("fecha", "N/A")
            row_cells[1].text = inc.get("dispositivo_excel", "N/A")
            row_cells[2].text = f"{inc.get('node_name', 'N/A')} ({inc.get('node_id', 'N/A')})"
            row_cells[3].text = format_number_chilean(inc.get("consumo_excel", 0), 2)
            row_cells[4].text = format_number_chilean(inc.get("consumo_api", 0), 2)
            row_cells[5].text = inc.get("mensaje", "N/A")
        
        doc.add_page_break()
    
    # Nodos sin datos
    nodos_sin_datos = [inc for inc in inconsistencias if inc.get("tipo") == "NODOS_SIN_DATOS"]
    if nodos_sin_datos:
        doc.add_heading("Nodos Sin Datos en la API", 1)
        doc.add_paragraph("Las siguientes fechas tienen algunos nodos sin datos disponibles en la API WES:")
        
        for inc in nodos_sin_datos:
            doc.add_paragraph(
                f"• Fecha: {inc.get('fecha', 'N/A')} - Dispositivo: {inc.get('dispositivo_excel', 'N/A')}",
                style='List Bullet'
            )
            doc.add_paragraph(f"  Nodos sin datos: {inc.get('mensaje', 'N/A')}")
        
        doc.add_page_break()
    
    # Sin datos en API
    sin_datos = [inc for inc in inconsistencias if inc.get("tipo") == "SIN_DATOS_API"]
    if sin_datos:
        doc.add_heading("Fechas Sin Datos en la API", 1)
        doc.add_paragraph("Las siguientes fechas no tienen datos disponibles en la API WES:")
        
        for inc in sin_datos:
            doc.add_paragraph(
                f"• {inc.get('fecha', 'N/A')} - {inc.get('node_name', 'N/A')} ({inc.get('node_id', 'N/A')}): "
                f"Consumo Excel: {format_number_chilean(inc.get('consumo_excel', 0), 2)} m³",
                style='List Bullet'
            )
        
        doc.add_page_break()
    
    # Errores de procesamiento
    errores = [inc for inc in inconsistencias if inc.get("tipo") == "ERROR_PROCESAMIENTO"]
    if errores:
        doc.add_heading("Errores de Procesamiento", 1)
        doc.add_paragraph("Se encontraron errores al procesar algunas filas del Excel:")
        
        for inc in errores:
            doc.add_paragraph(
                f"• Dispositivo: {inc.get('dispositivo_excel', 'N/A')}, "
                f"Fecha: {inc.get('fecha', 'N/A')}, "
                f"Error: {inc.get('error', 'N/A')}",
                style='List Bullet'
            )
    
    # Guardar documento
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"Inconsistencias_BUPA_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(output_path))
    
    return output_path

def main():
    print("=" * 60)
    print("  COMPARACIÓN MAESTRA BUPA vs API WES")
    print("=" * 60)
    print()
    
    # Leer Excel
    print("[1/4] Leyendo archivo Excel...")
    try:
        excel_df = leer_excel()
        print(f"[OK] Excel leído: {len(excel_df)} filas")
        print(f"     Dispositivos únicos: {excel_df['dispositivo'].nunique()}")
        print(f"     Rango de fechas: {excel_df['fecha'].min()} a {excel_df['fecha'].max()}")
    except Exception as e:
        print(f"[ERROR] No se pudo leer el Excel: {e}")
        return
    print()
    
    # Obtener nodos de la API
    print("[2/4] Obteniendo nodos de BUPA desde la API...")
    nodos_api = obtener_nodos_bupa_desde_api()
    if not nodos_api:
        print("[ERROR] No se pudieron obtener los nodos de la API")
        return
    print(f"[OK] Se encontraron {len(nodos_api)} nodo(s) en la API:")
    for node_id, node_name in sorted(nodos_api.items()):
        print(f"  - {node_id}: {node_name}")
    print()
    
    # Comparar datos
    print("[3/4] Comparando datos del Excel con la API...")
    print("      (Esto puede tomar varios minutos dependiendo de la cantidad de datos)")
    inconsistencias = comparar_datos(excel_df, nodos_api)
    print(f"[OK] Se encontraron {len(inconsistencias)} inconsistencia(s)")
    print()
    
    # Generar reporte Word
    print("[4/4] Generando reporte Word...")
    try:
        output_path = generar_reporte_word(inconsistencias, nodos_api, excel_df)
        print(f"[OK] Reporte generado: {output_path}")
    except Exception as e:
        print(f"[ERROR] No se pudo generar el reporte Word: {e}")
        import traceback
        traceback.print_exc()
        return
    
    print()
    print("=" * 60)
    print("  PROCESO COMPLETADO EXITOSAMENTE")
    print("=" * 60)
    print()
    print(f"[INFO] Total inconsistencias: {len(inconsistencias)}")
    print(f"[INFO] Reporte guardado en: {output_path}")

if __name__ == "__main__":
    main()

