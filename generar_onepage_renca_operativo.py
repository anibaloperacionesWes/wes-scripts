#!/usr/bin/env python3
"""
OnePage operativo Renca — resumen de lo ocurrido en el último tiempo.

Incluye consumos (julio y agosto parcial), estado por punto, hitos
(llave de emergencia ICCO, bombas Celebridad) y % de ahorro vs línea
base sin WES (auditoría abril 2026 / comparativo mayo para Gimnasio).

Uso:
  python generar_onepage_renca_operativo.py
"""

from __future__ import annotations

import csv
import io
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List, Tuple

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    KeepTogether,
)

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "Renca" / "ONEPAGE"
TARIFA = 1300  # CLP/m³ referencial

# Ahorro auditoría abr 2026 (semana CON WES 13–19 vs SIN WES 6–12)
AHORRO_ABRIL = {
    "000017-08": {"nombre": "Colegio ICCO Renca", "pct": 37.2, "con": 214.80, "sin": 341.80},
    "000017-04": {"nombre": "Esc. Lo Velásquez", "pct": 10.1, "con": 33.00, "sin": 36.70},
    "000017-06": {"nombre": "Piscina Municipal", "pct": 57.4, "con": 297.80, "sin": 698.74},
    "000017-07": {"nombre": "Cumbre de cóndores pte.", "pct": 6.2, "con": 48.40, "sin": 51.60},
}
# Gimnasio: abril anómalo (medición); se usa comparativo mayo 25–31 vs línea sin WES abr
AHORRO_GIMNASIO_MAYO = {"nombre": "Gimnasio", "pct": 96.5, "con": 16.41, "sin": 470.18, "nota": "ref. mayo vs sin WES abr"}

NODES = [
    ("000017-08", "Colegio ICCO Renca"),
    ("000017-06", "Piscina Municipal"),
    ("000017-07", "Cumbre de cóndores pte."),
    ("000017-05", "Gimnasio"),
    ("000017-04", "Esc. Lo Velásquez"),
]

WES_BLUE = RGBColor(0x0B, 0x3D, 0x91)
WES_TEAL = RGBColor(0x0D, 0x7A, 0x6F)
WES_DARK = RGBColor(0x1A, 0x1A, 0x1A)
WES_GRAY = RGBColor(0x4A, 0x4A, 0x4A)


def _acl_base() -> str:
    import os

    return os.environ.get(
        "WES_API_BASE_URL", "http://104.248.53.141:7003/wes/api/acl-node/v1"
    ).rstrip("/")


def _fmt_m3(v: float, decimals: int = 1) -> str:
    s = f"{v:,.{decimals}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt_pct(v: float) -> str:
    return f"{v:.1f}".replace(".", ",") + " %"


def fetch_period(node_id: str, start: str, end: str) -> Tuple[float, List[Tuple[str, float]]]:
    url = f"{_acl_base()}/nodes/{node_id}/dates.measures.csv"
    r = requests.get(url, params={"start": start, "end": end}, timeout=60)
    r.raise_for_status()
    rows = list(csv.DictReader(io.StringIO(r.text)))
    days: List[Tuple[str, float]] = []
    total = 0.0
    for row in rows:
        vals = list(row.values())
        t = vals[0]
        raw = vals[1] if len(vals) > 1 else "0"
        try:
            m3 = float(str(raw).replace(",", ".") or 0)
        except ValueError:
            m3 = 0.0
        days.append((t, m3))
        total += m3
    return total, days


def collect_data() -> Dict:
    jul: Dict[str, float] = {}
    ago: Dict[str, float] = {}
    ago_days: Dict[str, List[Tuple[str, float]]] = {}
    for nid, _ in NODES:
        jul[nid], _ = fetch_period(nid, "01072026", "31072026")
        ago[nid], ago_days[nid] = fetch_period(nid, "01082026", "12082026")
    return {
        "jul": jul,
        "ago": ago,
        "ago_days": ago_days,
        "jul_total": sum(jul.values()),
        "ago_total": sum(ago.values()),
        "generado": datetime.now(),
        "corte_ago": date(2026, 8, 12),
    }


def _set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_narrow_margins(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Cm(1.1)
        sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.3)
        sec.right_margin = Cm(1.3)


def _p(doc: Document, text: str, *, size=9, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, space_before=0):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.05
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.font.color.rgb = color or WES_DARK
    return para


def build_docx(data: Dict, out_path: Path) -> Path:
    doc = Document()
    _set_narrow_margins(doc)

    logo = ROOT / "logo wes.bmp"
    if logo.is_file():
        try:
            from generar_reporte_word import add_logo_to_header

            add_logo_to_header(doc, logo)
        except Exception:
            pass

    _p(
        doc,
        "ONEPAGE OPERATIVO — RENCA",
        size=16,
        bold=True,
        color=WES_BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )
    _p(
        doc,
        "Monitoreo WES · Situación reciente, consumos y próximos pasos",
        size=10,
        color=WES_TEAL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    _p(
        doc,
        f"Corte de datos: 01-jul-2026 → {data['corte_ago'].strftime('%d-%b-%Y').lower()}  ·  "
        f"Generado: {data['generado'].strftime('%d-%m-%Y %H:%M')}",
        size=8,
        color=WES_GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )

    _p(doc, "1. Qué ha pasado en el último tiempo", size=11, bold=True, color=WES_BLUE, space_before=2, space_after=3)
    _p(
        doc,
        "Los colegios de Renca mantuvieron el régimen de ahorro con control nocturno WES "
        "durante el periodo escolar hasta fines de julio. Hace aproximadamente dos semanas, "
        "en el ICCO (Colegio ICCO Renca) se activó la llave de emergencia, dejando el recinto "
        "sin corte nocturno mientras se resolvía la situación hidráulica.",
        size=9,
        space_after=3,
    )
    _p(
        doc,
        "Esta semana el equipo WES dio soporte a la partida de las bombas del ICCO. Las bombas "
        "partieron correctamente gracias al trabajo del equipo de mantención de Celebridad, que "
        "dejó el sistema operativo. Con ello se cierra el episodio de contingencia hidráulica y "
        "se recupera la capacidad de alimentar el establecimiento de forma estable.",
        size=9,
        space_after=3,
    )
    _p(
        doc,
        "Próxima semana —a más tardar el lunes— se inicia la auditoría del mes de agosto, "
        "para consolidar consumos, control nocturno y rendimiento de cada punto frente a la "
        "línea base sin WES.",
        size=9,
        space_after=6,
    )

    _p(doc, "2. Consumos recientes", size=11, bold=True, color=WES_BLUE, space_after=3)
    _p(
        doc,
        f"Julio 2026 (mes completo): {_fmt_m3(data['jul_total'])} m³ agregados en 5 puntos. "
        f"Agosto 2026 (01–12): {_fmt_m3(data['ago_total'])} m³. "
        "El ICCO concentra el mayor volumen, coherente con la llave de emergencia activa "
        "en la segunda quincena de julio y lo que va de agosto.",
        size=9,
        space_after=4,
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Punto", "Julio (m³)", "Ago 1–12 (m³)", "Estado operativo"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "0B3D91")

    estados = {
        "000017-08": "Llave emergencia · bombas OK (Celebridad)",
        "000017-06": "Control nocturno activo",
        "000017-07": "Operativo sin novedades",
        "000017-05": "Funcionando bien",
        "000017-04": "Control nocturno · proceso auto. activo",
    }
    for nid, nom in NODES:
        row = table.add_row().cells
        vals = [
            nom,
            _fmt_m3(data["jul"][nid]),
            _fmt_m3(data["ago"][nid]),
            estados[nid],
        ]
        for i, v in enumerate(vals):
            row[i].text = v
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    # total row
    row = table.add_row().cells
    for i, v in enumerate(["TOTAL", _fmt_m3(data["jul_total"]), _fmt_m3(data["ago_total"]), ""]):
        row[i].text = v
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
                r.bold = True
        _shade(row[i], "E8EEF7")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _p(doc, "3. Estado por establecimiento", size=11, bold=True, color=WES_BLUE, space_before=4, space_after=3)
    bullets = [
        "ICCO: llave de emergencia ~hace dos semanas; esta semana bombas partieron con soporte WES + mantención Celebridad. Acta jun-2026: asesoramiento de corte hídrico manual.",
        "Gimnasio: funcionando bien (actas 2026: equipo 100% operativo tras mantenciones).",
        "Piscina Municipal: control nocturno activo; actas documentan programación de horario/presión y un evento de llave de emergencia abierta (jun-2026).",
        "Esc. Lo Velásquez: control nocturno; tiende a cero; días con consumo por llaves abiertas → proceso automático activado.",
        "Cumbre de cóndores: operativo sin problemas (acta: equipo operativo tras chequeo técnico).",
    ]
    for b in bullets:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(b)
        run.font.size = Pt(8)
        run.font.name = "Calibri"

    _p(doc, "4. Ahorro vs antes de operar (auditoría abril 2026)", size=11, bold=True, color=WES_BLUE, space_before=6, space_after=3)
    _p(
        doc,
        "Comparación de la semana con control WES (13–19 abr) frente a la semana de referencia "
        "sin WES (6–12 abr). Para el Gimnasio se reporta el comparativo de mayo (25–31) vs la "
        "misma línea base sin WES de abril, representativo de la operación normal.",
        size=8,
        color=WES_GRAY,
        space_after=3,
    )

    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    hdr2 = t2.rows[0].cells
    for i, h in enumerate(["Punto", "% ahorro", "Semana sin WES (m³)", "Semana con WES (m³)"]):
        hdr2[i].text = h
        for p in hdr2[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr2[i], "0D7A6F")

    rows_ahorro = [
        ("Colegio ICCO Renca", 37.2, 341.80, 214.80, ""),
        ("Esc. Lo Velásquez", 10.1, 36.70, 33.00, ""),
        ("Piscina Municipal", 57.4, 698.74, 297.80, ""),
        ("Gimnasio", 96.5, 470.18, 16.41, "mayo vs sin WES abr"),
        ("Cumbre de cóndores pte.", 6.2, 51.60, 48.40, ""),
    ]
    for nom, pct, sin, con, nota in rows_ahorro:
        row = t2.add_row().cells
        label = nom + (f" ({nota})" if nota else "")
        for i, v in enumerate([label, _fmt_pct(pct), _fmt_m3(sin, 1), _fmt_m3(con, 1)]):
            row[i].text = v
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    _p(
        doc,
        "Lectura: con el sistema operando, Piscina e ICCO concentran el mayor rendimiento de "
        "ahorro (57 % y 37 %). Lo Velásquez ahorra ~10 % y además suele caer a cero en noche. "
        "Cóndores aporta un ahorro moderado (~6 %) con operación estable. El Gimnasio, en "
        "régimen normal, muestra un ahorro muy alto respecto a la línea base previa.",
        size=8.5,
        space_before=4,
        space_after=4,
    )

    _p(
        doc,
        "5. Actas de mantención (Drive: Actas de Mantencion / RENCA / 2026 + Corporación)",
        size=11,
        bold=True,
        color=WES_BLUE,
        space_after=2,
    )
    _p(
        doc,
        "Revisión de formularios WES cargados en la carpeta indicada (escaneos). "
        "Hito reciente ICCO (llave emergencia ~2 semanas + bombas Celebridad esta semana): "
        "aún no figura acta formal de julio/agosto en esa carpeta; se reporta por seguimiento operativo.",
        size=7.5,
        color=WES_GRAY,
        space_after=2,
    )

    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    hdr3 = t3.rows[0].cells
    for i, h in enumerate(["Fecha / Acta", "Punto", "Hallazgo (observaciones)"]):
        hdr3[i].text = h
        for p in hdr3[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(7.5)
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr3[i], "0B3D91")

    actas_rows = [
        ("20-06-2026 · N°2221", "ICCO", "Asesoramiento para corte hídrico MANUAL."),
        (
            "25-06-2026 · N°2216",
            "Piscina",
            "Llave emergencia abierta desde 19-06; fuga ~2,45 m³/h (ruido eléctrico).",
        ),
        ("13-05-2026 · N°2183", "Gimnasio", "Automático caído; levantado. Equipo 100% operativo."),
        ("21-04-2026 · N°2149", "Gimnasio", "Reemplazo equipo + transductores. 100% operativo."),
        ("06-04-2026", "Gimnasio", "Ruido eléctrico → filtro instalado para asegurar horario."),
        ("30-01-2026", "Piscina", "Consumo nocturno ~1,3 m³/h; nuevo horario/presión. Operativo."),
        ("26-05-2026 · N°2198", "Cóndores Ote.", "Chequeo caudalímetro + micro. Equipo operativo."),
    ]
    for fecha, punto, hallazgo in actas_rows:
        row = t3.add_row().cells
        for i, v in enumerate([fecha, punto, hallazgo]):
            row[i].text = v
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)

    _p(
        doc,
        "6. Próximos pasos",
        size=11,
        bold=True,
        color=WES_BLUE,
        space_before=4,
        space_after=2,
    )
    next_steps = [
        "Lunes próximo (a más tardar): auditoría de agosto — consumos, nocturno y % ahorro por punto.",
        "ICCO: normalizar llave de emergencia, retomar control nocturno y cargar acta de bombas/Celebridad en Drive.",
        "Piscina / Lo Velásquez: mantener control nocturno; vigilar llaves abiertas / llave emergencia.",
        "Gimnasio y Cóndores: seguimiento estable (equipos operativos según actas).",
    ]
    for s in next_steps:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(s)
        run.font.size = Pt(8.5)

    _p(
        doc,
        "Fuente consumos: API WES. Ahorros: auditoría abr-2026 (Gimnasio: may-2026). "
        "Actas: Drive Actas de Mantencion/RENCA/2026 + Corporación Renca. Tarifa ref. $1.300/m³.",
        size=7,
        color=WES_GRAY,
        space_before=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_pdf(data: Dict, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=1.3 * cm,
        rightMargin=1.3 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "T",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=colors.HexColor("#0B3D91"),
        alignment=TA_CENTER,
        spaceAfter=2,
        leading=16,
    )
    sub = ParagraphStyle(
        "S",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        textColor=colors.HexColor("#0D7A6F"),
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    meta = ParagraphStyle(
        "M",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7.5,
        textColor=colors.HexColor("#666666"),
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    h = ParagraphStyle(
        "H",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10,
        textColor=colors.HexColor("#0B3D91"),
        spaceBefore=4,
        spaceAfter=2,
        leading=12,
    )
    body = ParagraphStyle(
        "B",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leading=10.5,
        alignment=TA_JUSTIFY,
        spaceAfter=3,
        textColor=colors.HexColor("#1A1A1A"),
    )
    bullet = ParagraphStyle(
        "Bu",
        parent=body,
        leftIndent=8,
        spaceAfter=1.5,
        alignment=TA_LEFT,
    )
    foot = ParagraphStyle(
        "F",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=6.5,
        textColor=colors.HexColor("#666666"),
        alignment=TA_CENTER,
        spaceBefore=4,
    )
    cell = ParagraphStyle(
        "C",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7,
        leading=9,
    )
    cellb = ParagraphStyle(
        "Cb",
        parent=cell,
        fontName="Helvetica-Bold",
        textColor=colors.white,
    )

    story = []
    story.append(Paragraph("ONEPAGE OPERATIVO — RENCA", title))
    story.append(Paragraph("Monitoreo WES · Situación reciente, consumos y próximos pasos", sub))
    story.append(
        Paragraph(
            f"Corte de datos: 01-jul-2026 → {data['corte_ago'].strftime('%d-%m-%Y')} · "
            f"Generado: {data['generado'].strftime('%d-%m-%Y %H:%M')}",
            meta,
        )
    )

    story.append(Paragraph("1. Qué ha pasado en el último tiempo", h))
    story.append(
        Paragraph(
            "Los colegios de Renca mantuvieron el régimen de <b>ahorro con control nocturno WES</b> "
            "durante el periodo escolar <b>hasta fines de julio</b>. Hace aproximadamente dos semanas, "
            "en el <b>ICCO</b> (Colegio ICCO Renca) se activó la <b>llave de emergencia</b>, dejando el "
            "recinto sin corte nocturno mientras se resolvía la situación hidráulica.",
            body,
        )
    )
    story.append(
        Paragraph(
            "Esta semana el equipo WES dio soporte a la <b>partida de las bombas</b> del ICCO. "
            "Las bombas partieron correctamente gracias al trabajo del equipo de mantención de "
            "<b>Celebridad</b>, que dejó el sistema operativo.",
            body,
        )
    )
    story.append(
        Paragraph(
            "La <b>próxima semana —a más tardar el lunes—</b> se inicia la <b>auditoría del mes de agosto</b>, "
            "para consolidar consumos, control nocturno y rendimiento de cada punto.",
            body,
        )
    )

    story.append(Paragraph("2. Consumos recientes", h))
    story.append(
        Paragraph(
            f"<b>Julio 2026:</b> {_fmt_m3(data['jul_total'])} m³ (5 puntos). "
            f"<b>Agosto 01–12:</b> {_fmt_m3(data['ago_total'])} m³. "
            "El ICCO concentra el mayor volumen, coherente con la llave de emergencia activa.",
            body,
        )
    )

    estados = {
        "000017-08": "Llave emergencia · bombas OK",
        "000017-06": "Control nocturno activo",
        "000017-07": "Operativo sin novedades",
        "000017-05": "Funcionando bien",
        "000017-04": "Control nocturno · auto. activo",
    }
    tdata = [[
        Paragraph("Punto", cellb),
        Paragraph("Julio m³", cellb),
        Paragraph("Ago 1–12 m³", cellb),
        Paragraph("Estado", cellb),
    ]]
    for nid, nom in NODES:
        tdata.append([
            Paragraph(nom, cell),
            Paragraph(_fmt_m3(data["jul"][nid]), cell),
            Paragraph(_fmt_m3(data["ago"][nid]), cell),
            Paragraph(estados[nid], cell),
        ])
    tdata.append([
        Paragraph("<b>TOTAL</b>", cell),
        Paragraph(f"<b>{_fmt_m3(data['jul_total'])}</b>", cell),
        Paragraph(f"<b>{_fmt_m3(data['ago_total'])}</b>", cell),
        Paragraph("", cell),
    ])
    t = Table(tdata, colWidths=[5.2 * cm, 2.4 * cm, 2.8 * cm, 6.5 * cm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D91")),
                ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#E8EEF7")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#99AACC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 3 * mm))

    story.append(Paragraph("3. Estado por establecimiento", h))
    for txt in [
        "<b>ICCO:</b> llave emergencia ~2 semanas; bombas OK (Celebridad). Acta jun: corte hídrico MANUAL.",
        "<b>Gimnasio:</b> funcionando bien (actas 2026: 100% operativo).",
        "<b>Piscina:</b> control nocturno; acta jun documenta llave emergencia abierta (19–25 jun).",
        "<b>Lo Velásquez:</b> control nocturno; tiende a cero; llaves abiertas → proceso automático.",
        "<b>Cóndores:</b> operativo sin problema (acta: equipo operativo).",
    ]:
        story.append(Paragraph("• " + txt, bullet))

    story.append(Paragraph("4. Ahorro vs antes de operar (auditoría abril 2026)", h))
    story.append(
        Paragraph(
            "Semana con WES (13–19 abr) vs sin WES (6–12 abr). Gimnasio: comparativo mayo 25–31 vs línea sin WES abr.",
            body,
        )
    )
    t2data = [[
        Paragraph("Punto", cellb),
        Paragraph("% ahorro", cellb),
        Paragraph("Sin WES m³", cellb),
        Paragraph("Con WES m³", cellb),
    ]]
    for nom, pct, sin, con, nota in [
        ("Colegio ICCO Renca", 37.2, 341.80, 214.80, ""),
        ("Esc. Lo Velásquez", 10.1, 36.70, 33.00, ""),
        ("Piscina Municipal", 57.4, 698.74, 297.80, ""),
        ("Gimnasio (mayo vs sin WES abr)", 96.5, 470.18, 16.41, ""),
        ("Cumbre de cóndores pte.", 6.2, 51.60, 48.40, ""),
    ]:
        t2data.append([
            Paragraph(nom, cell),
            Paragraph(_fmt_pct(pct), cell),
            Paragraph(_fmt_m3(sin, 1), cell),
            Paragraph(_fmt_m3(con, 1), cell),
        ])
    t2 = Table(t2data, colWidths=[7.0 * cm, 2.5 * cm, 3.5 * cm, 3.5 * cm])
    t2.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0D7A6F")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#88BBA8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(t2)
    story.append(Spacer(1, 1.5 * mm))
    story.append(
        Paragraph(
            "Piscina ~57 % · ICCO ~37 % · Lo Velásquez ~10 % · Cóndores ~6 % · Gimnasio ~96 % (régimen normal).",
            body,
        )
    )

    story.append(Paragraph("5. Actas Drive (RENCA/2026 + Corporación) y próximos pasos", h))
    story.append(
        Paragraph(
            "<b>ICCO N°2221 (20-06):</b> asesoramiento corte hídrico MANUAL. "
            "<b>Piscina N°2216 (25-06):</b> llave emergencia abierta desde 19-06; fuga ~2,45 m³/h. "
            "<b>Gimnasio:</b> abr–may operativos (filtro ruido eléctrico; reemplazo equipo; automático levantado). "
            "<b>Cóndores Ote. N°2198 (26-05):</b> equipo operativo. "
            "Hito reciente ICCO (llave emergencia + bombas Celebridad): sin acta jul/ago aún en carpeta. "
            "<b>Próximo lunes:</b> auditoría agosto; normalizar llave ICCO; cargar acta bombas; mantener nocturno.",
            body,
        )
    )
    story.append(
        Paragraph(
            "Fuente: API WES · Auditoría abr-2026 · Actas Drive Actas de Mantencion/RENCA · Tarifa ref. $1.300/m³",
            foot,
        )
    )

    doc.build(story)
    return out_path


def main() -> int:
    if sys.platform == "win32":
        for s in (sys.stdout, sys.stderr):
            try:
                s.reconfigure(encoding="utf-8", line_buffering=True)
            except Exception:
                pass

    print("[INFO] Descargando consumos julio / agosto desde API WES…")
    data = collect_data()
    stamp = data["generado"].strftime("%Y%m%d_%H%M")
    out_dir = OUT_DIR / f"ONEPAGE_{stamp}"
    docx_path = out_dir / f"OnePage_Renca_Operativo_{stamp}.docx"
    pdf_path = out_dir / f"OnePage_Renca_Operativo_{stamp}.pdf"

    print("[INFO] Generando DOCX…")
    build_docx(data, docx_path)
    print(f"[OK] {docx_path}")

    print("[INFO] Generando PDF…")
    build_pdf(data, pdf_path)
    print(f"[OK] {pdf_path}")

    print(f"[RESUMEN] Julio total: {_fmt_m3(data['jul_total'])} m³")
    print(f"[RESUMEN] Agosto 1–12 total: {_fmt_m3(data['ago_total'])} m³")
    for nid, nom in NODES:
        print(f"  {nom}: jul {_fmt_m3(data['jul'][nid])} | ago {_fmt_m3(data['ago'][nid])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
