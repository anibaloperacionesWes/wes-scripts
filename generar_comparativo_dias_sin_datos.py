"""
Comparativo de días sin datos entre dos periodos (p. ej. lun–vie vs lun–vie).

Un día civil es «sin datos» si en GET /nodes/measures/dates el día no aparece
en month[] o totalM3 es null. totalM3 = 0 cuenta como dato (punto en cero),
no como día sin data.

Uso:
  python generar_comparativo_dias_sin_datos.py \\
    --periodo-a 03/08/2026 07/08/2026 \\
    --periodo-b 10/08/2026 14/08/2026
"""

from __future__ import annotations

import argparse
import csv
import os
import shutil
import subprocess
import sys
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

from generar_reporte_word import acl_node_base_url, normalize_measures_payload
from reporte_puntos_en_cero import MAX_WORKERS_CERO, obtener_todos_los_nodos

DIAS_ES = {
    0: "Lunes",
    1: "Martes",
    2: "Miércoles",
    3: "Jueves",
    4: "Viernes",
    5: "Sábado",
    6: "Domingo",
}
COLOR_A = "#1F4E79"
COLOR_B = "#E67E22"
COLOR_OK = "#2ECC71"
COLOR_BAD = "#C0392B"


def _iter_days(inicio: date, fin: date) -> List[date]:
    out: List[date] = []
    cur = inicio
    while cur <= fin:
        out.append(cur)
        cur += timedelta(days=1)
    return out


def _parse_fecha(txt: str) -> date:
    return datetime.strptime(txt.strip(), "%d/%m/%Y").date()


def _fmt(d: date) -> str:
    return d.strftime("%d-%m-%Y")


def _label_periodo(inicio: date, fin: date) -> str:
    return f"{_fmt(inicio)} al {_fmt(fin)}"


def _parse_json_date(raw: object) -> Optional[date]:
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    try:
        if "T" in s:
            return datetime.fromisoformat(s.replace("Z", "+00:00")).date()
        return datetime.strptime(s[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return None


def _dias_con_total(payload: dict, dias: Sequence[date]) -> Dict[date, Optional[float]]:
    """Mapa día → totalM3 (None si el día no tiene total)."""
    found: Dict[date, Optional[float]] = {}
    for item in payload.get("month") or []:
        dia = _parse_json_date(item.get("date"))
        if dia is None:
            continue
        total = item.get("totalM3")
        if total is None:
            found[dia] = None
        else:
            try:
                found[dia] = float(total)
            except (TypeError, ValueError):
                found[dia] = None
    return {d: found.get(d) for d in dias}


def consultar_nodo(
    node_id: str,
    inicio: date,
    fin: date,
) -> Tuple[Dict[date, Optional[float]], Optional[str]]:
    url = f"{acl_node_base_url()}/nodes/measures/dates"
    dias = _iter_days(inicio, fin)
    try:
        resp = requests.get(
            url,
            params=[
                ("id", node_id),
                ("start", inicio.strftime("%d%m%Y")),
                ("end", fin.strftime("%d%m%Y")),
            ],
            timeout=45,
        )
        if resp.status_code != 200:
            return {d: None for d in dias}, f"HTTP {resp.status_code}"
        payload = normalize_measures_payload(resp.json(), node_id)
        return _dias_con_total(payload, dias), None
    except Exception as exc:
        return {d: None for d in dias}, str(exc)


def _set_shading(cell, hex_color: str) -> None:
    shading = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _style_header_row(row, fill: str = "1F4E79") -> None:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
        _set_shading(cell, fill)


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _style_header_row(hdr)
    for r_i, row_vals in enumerate(rows, start=1):
        for c_i, val in enumerate(row_vals):
            cell = table.rows[r_i].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            if r_i % 2 == 0:
                _set_shading(cell, "F2F2F2")


def _sign(n: int) -> str:
    return f"+{n}" if n > 0 else str(n)


def _pct(num: int, den: int) -> str:
    if den <= 0:
        return "0,00%"
    return f"{(num / den) * 100:.2f}%".replace(".", ",")


def _dias_txt(dias: Iterable[date]) -> str:
    xs = sorted(dias)
    return ", ".join(d.strftime("%d-%m") for d in xs) if xs else "—"


def grafico_por_dia(
    dias_a: Sequence[date],
    dias_b: Sequence[date],
    count_a: Sequence[int],
    count_b: Sequence[int],
    label_a: str,
    label_b: str,
    out_path: Path,
) -> Path:
    labels = []
    n = min(len(dias_a), len(dias_b), len(count_a), len(count_b))
    for i in range(n):
        da, db = dias_a[i], dias_b[i]
        labels.append(f"{DIAS_ES[da.weekday()][:3]}\n{_fmt(da)[:5]} / {_fmt(db)[:5]}")
    x = list(range(n))
    width = 0.38
    fig, ax = plt.subplots(figsize=(8.2, 4.2))
    ax.bar([i - width / 2 for i in x], count_a[:n], width, label=label_a, color=COLOR_A)
    ax.bar([i + width / 2 for i in x], count_b[:n], width, label=label_b, color=COLOR_B)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_ylabel("Puntos sin datos")
    ax.set_title("Puntos sin datos por día (misma jornada lun–vie)")
    ax.legend(fontsize=8)
    ax.grid(axis="y", linestyle="--", alpha=0.4)
    ymax = max(list(count_a[:n]) + list(count_b[:n]) + [1])
    ax.set_ylim(0, ymax * 1.18)
    for i in range(n):
        ax.text(i - width / 2, count_a[i] + ymax * 0.02, str(count_a[i]), ha="center", va="bottom", fontsize=8)
        ax.text(i + width / 2, count_b[i] + ymax * 0.02, str(count_b[i]), ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    fig.savefig(out_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out_path


def grafico_nodos_por_cantidad(
    hist_a: Dict[int, int],
    hist_b: Dict[int, int],
    max_dias: int,
    label_a: str,
    label_b: str,
    out_path: Path,
) -> Path:
    xs = list(range(0, max_dias + 1))
    ya = [hist_a.get(i, 0) for i in xs]
    yb = [hist_b.get(i, 0) for i in xs]
    width = 0.38
    fig, ax = plt.subplots(figsize=(8.2, 4.0))
    ax.bar([i - width / 2 for i in xs], ya, width, label=label_a, color=COLOR_A)
    ax.bar([i + width / 2 for i in xs], yb, width, label=label_b, color=COLOR_B)
    ax.set_xticks(xs)
    ax.set_xticklabels([str(i) for i in xs])
    ax.set_xlabel("Días sin datos en el periodo")
    ax.set_ylabel("Cantidad de puntos")
    ax.set_title("Distribución de puntos según días sin datos")
    ax.legend(fontsize=8)
    ax.grid(axis="y", linestyle="--", alpha=0.4)
    fig.tight_layout()
    fig.savefig(out_path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out_path


def convertir_a_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    return pdf_path if pdf_path.exists() else None


def crear_informe(
    nodos: List[Dict[str, str]],
    por_nodo: Dict[str, Dict[date, Optional[float]]],
    errores: Dict[str, str],
    dias_a: List[date],
    dias_b: List[date],
    output_dir: Path,
    hoy_chile: date,
) -> Tuple[Path, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    label_a = _label_periodo(dias_a[0], dias_a[-1])
    label_b = _label_periodo(dias_b[0], dias_b[-1])
    n_puntos = len(nodos)
    n_dias_a = len(dias_a)
    n_dias_b = len(dias_b)

    def _sin(nid: str, dias: Sequence[date]) -> List[date]:
        mapa = por_nodo.get(nid, {})
        return [d for d in dias if mapa.get(d) is None]

    filas = []
    for nodo in nodos:
        nid = nodo["nodeId"]
        sin_a = _sin(nid, dias_a)
        sin_b = _sin(nid, dias_b)
        filas.append(
            {
                "nodeId": nid,
                "nodeName": nodo.get("nodeName", ""),
                "companyName": nodo.get("companyName", ""),
                "companyId": nodo.get("companyId", ""),
                "sin_a": sin_a,
                "sin_b": sin_b,
                "n_a": len(sin_a),
                "n_b": len(sin_b),
                "delta": len(sin_b) - len(sin_a),
                "error": errores.get(nid, ""),
            }
        )

    count_a = [sum(1 for f in filas if d in f["sin_a"]) for d in dias_a]
    count_b = [sum(1 for f in filas if d in f["sin_b"]) for d in dias_b]
    node_days_a = sum(f["n_a"] for f in filas)
    node_days_b = sum(f["n_b"] for f in filas)
    max_a = n_puntos * n_dias_a
    max_b = n_puntos * n_dias_b
    unicos_a = sum(1 for f in filas if f["n_a"] > 0)
    unicos_b = sum(1 for f in filas if f["n_b"] > 0)
    delta_nd = node_days_b - node_days_a
    delta_u = unicos_b - unicos_a

    nuevos = [f for f in filas if f["n_a"] == 0 and f["n_b"] > 0]
    recuperados = [f for f in filas if f["n_a"] > 0 and f["n_b"] == 0]
    empeoran = [f for f in filas if f["delta"] > 0]
    mejoran = [f for f in filas if f["delta"] < 0]
    persisten = [f for f in filas if f["n_a"] > 0 and f["n_b"] > 0]
    full_a = [f for f in filas if f["n_a"] == n_dias_a]
    full_b = [f for f in filas if f["n_b"] == n_dias_b]

    hist_a: Dict[int, int] = defaultdict(int)
    hist_b: Dict[int, int] = defaultdict(int)
    for f in filas:
        hist_a[f["n_a"]] += 1
        hist_b[f["n_b"]] += 1

    png_dias = grafico_por_dia(
        dias_a,
        dias_b,
        count_a,
        count_b,
        label_a,
        label_b,
        output_dir / "chart_puntos_sin_datos_por_dia.png",
    )
    png_hist = grafico_nodos_por_cantidad(
        dict(hist_a),
        dict(hist_b),
        max(n_dias_a, n_dias_b),
        label_a,
        label_b,
        output_dir / "chart_distribucion_dias_sin_datos.png",
    )

    doc = Document()
    title = doc.add_heading("COMPARATIVO DÍAS SIN DATOS", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.color.rgb = RGBColor(230, 126, 34)

    sub = doc.add_paragraph(f"{label_a}  vs  {label_b}")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].bold = True

    gen = doc.add_paragraph(
        f"Informe generado: {datetime.now(timezone.utc).strftime('%d-%m-%Y %H:%M:%S')} UTC"
    )
    gen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    criterio = doc.add_paragraph(
        "Criterio: un día civil se cuenta como «sin datos» cuando el API "
        "(/nodes/measures/dates) no entrega totalM3 para ese día. "
        "Un totalM3 = 0 es punto en cero (sí hay dato), no día sin data. "
        "Universo: mismos puntos del reporte diario de puntos en cero "
        "(mismas exclusiones de empresa y nodo)."
    )
    criterio.runs[0].font.size = Pt(10)

    if hoy_chile in dias_a or hoy_chile in dias_b:
        nota = doc.add_paragraph(
            f"Nota: el {_fmt(hoy_chile)} está en curso a la hora de este informe. "
            "Un punto que aún no transmite hoy puede aparecer como sin datos "
            "aunque recupere más tarde."
        )
        nota.runs[0].italic = True
        nota.runs[0].font.size = Pt(10)

    if errores:
        av = doc.add_paragraph(
            f"Consultas con error de API: {len(errores)} punto(s). "
            "Esos puntos se cuentan como sin datos en todos los días del rango."
        )
        av.runs[0].font.size = Pt(10)

    doc.add_heading("RESUMEN EJECUTIVO", 1)
    resumen = doc.add_paragraph(
        f"Puntos analizados: {n_puntos}\n"
        f"Punto-días sin datos: {node_days_a} de {max_a} ({_pct(node_days_a, max_a)}) → "
        f"{node_days_b} de {max_b} ({_pct(node_days_b, max_b)}) ({_sign(delta_nd)})\n"
        f"Puntos con ≥1 día sin datos: {unicos_a} ({_pct(unicos_a, n_puntos)}) → "
        f"{unicos_b} ({_pct(unicos_b, n_puntos)}) ({_sign(delta_u)})\n"
        f"Promedio diario de puntos sin datos: "
        f"{(sum(count_a) / n_dias_a if n_dias_a else 0):.1f} → "
        f"{(sum(count_b) / n_dias_b if n_dias_b else 0):.1f}\n"
        f"Puntos sin datos los {n_dias_a}/{n_dias_a} días: {len(full_a)} → "
        f"los {n_dias_b}/{n_dias_b} días: {len(full_b)}"
    )
    resumen.runs[0].font.size = Pt(11)

    _add_table(
        doc,
        ["Indicador", label_a, label_b, "Δ"],
        [
            ["Puntos analizados", str(n_puntos), str(n_puntos), "0"],
            ["Punto-días sin datos", str(node_days_a), str(node_days_b), _sign(delta_nd)],
            ["% punto-días sin datos", _pct(node_days_a, max_a), _pct(node_days_b, max_b), ""],
            ["Puntos con ≥1 día sin datos", str(unicos_a), str(unicos_b), _sign(delta_u)],
            ["% puntos con ≥1 día sin datos", _pct(unicos_a, n_puntos), _pct(unicos_b, n_puntos), ""],
            [
                "Promedio diario (puntos sin datos)",
                f"{(sum(count_a) / n_dias_a if n_dias_a else 0):.1f}",
                f"{(sum(count_b) / n_dias_b if n_dias_b else 0):.1f}",
                "",
            ],
            [f"Sin datos los {n_dias_a} días", str(len(full_a)), str(len(full_b)), _sign(len(full_b) - len(full_a))],
        ],
    )

    doc.add_heading("PUNTOS SIN DATOS POR DÍA", 1)
    filas_dia = []
    n_par = min(len(dias_a), len(dias_b))
    for i in range(n_par):
        da, db = dias_a[i], dias_b[i]
        filas_dia.append(
            [
                DIAS_ES[da.weekday()],
                _fmt(da),
                str(count_a[i]),
                _fmt(db),
                str(count_b[i]),
                _sign(count_b[i] - count_a[i]),
            ]
        )
    _add_table(
        doc,
        ["Jornada", f"Fecha {label_a}", "Sin datos", f"Fecha {label_b}", "Sin datos", "Δ"],
        filas_dia,
    )
    pic = doc.add_paragraph()
    pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pic.add_run().add_picture(str(png_dias), width=Inches(6.3))

    pic2 = doc.add_paragraph()
    pic2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pic2.add_run().add_picture(str(png_hist), width=Inches(6.3))

    doc.add_heading("HALLAZGOS CLAVE", 1)
    hallazgos = doc.add_paragraph(
        f"• Nuevos (0 días en {label_a}, ≥1 en {label_b}): {len(nuevos)}\n"
        f"• Recuperados (≥1 en {label_a}, 0 en {label_b}): {len(recuperados)}\n"
        f"• Empeoran (más días sin datos en {label_b}): {len(empeoran)}\n"
        f"• Mejoran (menos días sin datos en {label_b}): {len(mejoran)}\n"
        f"• Persisten (≥1 día en ambos periodos): {len(persisten)}"
    )
    hallazgos.runs[0].font.size = Pt(11)

    headers_pts = [
        "Nodo ID",
        "Nombre",
        "Empresa",
        f"Días {label_a}",
        f"Días {label_b}",
        "Δ",
        f"Fechas {label_a}",
        f"Fechas {label_b}",
    ]

    def _rows_pts(items: List[dict]) -> List[List[str]]:
        ordered = sorted(items, key=lambda x: (-x["delta"], -x["n_b"], x["companyName"], x["nodeName"]))
        out = []
        for f in ordered:
            out.append(
                [
                    f["nodeId"],
                    f["nodeName"],
                    f["companyName"],
                    str(f["n_a"]),
                    str(f["n_b"]),
                    _sign(f["delta"]),
                    _dias_txt(f["sin_a"]),
                    _dias_txt(f["sin_b"]),
                ]
            )
        return out

    def _seccion(titulo: str, items: List[dict]) -> None:
        doc.add_heading(f"{titulo} ({len(items)})", 1)
        if items:
            _add_table(doc, headers_pts, _rows_pts(items))
        else:
            doc.add_paragraph("Ninguno.")

    _seccion("NUEVOS SIN DATOS", nuevos)
    _seccion("RECUPERADOS", recuperados)
    _seccion("EMPEORAN", empeoran)
    _seccion("MEJORAN", mejoran)
    _seccion("PERSISTEN EN AMBOS PERIODOS", persisten)

    afectados = [f for f in filas if f["n_a"] > 0 or f["n_b"] > 0]
    doc.add_heading(f"DETALLE — PUNTOS CON ALGÚN DÍA SIN DATOS ({len(afectados)})", 1)
    if afectados:
        _add_table(doc, headers_pts, _rows_pts(afectados))
    else:
        doc.add_paragraph("Ningún punto del universo tuvo días sin datos en estos periodos.")

    doc.add_heading("CONCLUSIÓN", 1)
    if delta_nd > 0:
        conclusion = (
            f"Respecto de {label_a}, los punto-días sin datos aumentaron en {delta_nd} "
            f"({node_days_a} → {node_days_b}; {_pct(node_days_b, max_b)} del universo). "
            f"Hay {len(nuevos)} puntos nuevos y {len(persisten)} que se mantienen con al menos un día sin datos."
        )
    elif delta_nd < 0:
        conclusion = (
            f"Respecto de {label_a}, los punto-días sin datos bajaron en {abs(delta_nd)} "
            f"({node_days_a} → {node_days_b}; {_pct(node_days_b, max_b)} del universo). "
            f"Se recuperaron {len(recuperados)} puntos y empeoraron {len(empeoran)}."
        )
    else:
        conclusion = (
            f"La cantidad de punto-días sin datos se mantiene igual que en {label_a} "
            f"({node_days_b}). Igual hay movimiento interno: {len(nuevos)} nuevos y "
            f"{len(recuperados)} recuperados."
        )
    doc.add_paragraph(conclusion)

    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    stem = (
        f"Comparativo_Dias_Sin_Datos_"
        f"{dias_a[0].strftime('%Y%m%d')}_{dias_a[-1].strftime('%Y%m%d')}_vs_"
        f"{dias_b[0].strftime('%Y%m%d')}_{dias_b[-1].strftime('%Y%m%d')}_{stamp}"
    )
    docx_path = output_dir / f"{stem}.docx"
    csv_path = output_dir / f"{stem}.csv"
    doc.save(str(docx_path))

    with csv_path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        w.writerow(
            [
                "nodeId",
                "nodeName",
                "companyName",
                "companyId",
                "dias_sin_periodo_a",
                "dias_sin_periodo_b",
                "delta",
                "fechas_a",
                "fechas_b",
                "error_api",
            ]
        )
        for f in sorted(filas, key=lambda x: (-x["delta"], -x["n_b"], x["companyName"], x["nodeName"])):
            w.writerow(
                [
                    f["nodeId"],
                    f["nodeName"],
                    f["companyName"],
                    f["companyId"],
                    f["n_a"],
                    f["n_b"],
                    f["delta"],
                    _dias_txt(f["sin_a"]),
                    _dias_txt(f["sin_b"]),
                    f["error"],
                ]
            )

    return docx_path, csv_path


def main() -> int:
    if sys.platform == "win32":
        for stream in (sys.stdout, sys.stderr):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass

    parser = argparse.ArgumentParser(description="Comparativo de días sin datos entre dos periodos")
    parser.add_argument("--periodo-a", nargs=2, metavar=("DD/MM/YYYY", "DD/MM/YYYY"), required=True)
    parser.add_argument("--periodo-b", nargs=2, metavar=("DD/MM/YYYY", "DD/MM/YYYY"), required=True)
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("reports/Puntos_En_Cero"),
    )
    args = parser.parse_args()

    ini_a, fin_a = _parse_fecha(args.periodo_a[0]), _parse_fecha(args.periodo_a[1])
    ini_b, fin_b = _parse_fecha(args.periodo_b[0]), _parse_fecha(args.periodo_b[1])
    if fin_a < ini_a or fin_b < ini_b:
        print("[ERROR] Cada periodo debe ir de fecha inicial a final.", flush=True)
        return 1

    dias_a = _iter_days(ini_a, fin_a)
    dias_b = _iter_days(ini_b, fin_b)
    rango_ini = min(ini_a, ini_b)
    rango_fin = max(fin_a, fin_b)

    print("=" * 70)
    print("COMPARATIVO DÍAS SIN DATOS")
    print(f"Periodo A: {_label_periodo(ini_a, fin_a)} ({len(dias_a)} días)")
    print(f"Periodo B: {_label_periodo(ini_b, fin_b)} ({len(dias_b)} días)")
    print("=" * 70)

    nodos = obtener_todos_los_nodos()
    if not nodos:
        print("[ERROR] No se encontraron nodos.")
        return 1

    por_nodo: Dict[str, Dict[date, Optional[float]]] = {}
    errores: Dict[str, str] = {}
    workers = max(4, int(os.environ.get("WES_REPORTE_CERO_WORKERS", str(MAX_WORKERS_CERO))))
    print(f"Consultando medidas de {len(nodos)} puntos ({workers} en paralelo)...")

    def _job(nodo: Dict[str, str]) -> Tuple[str, Dict[date, Optional[float]], Optional[str]]:
        mapa, err = consultar_nodo(nodo["nodeId"], rango_ini, rango_fin)
        return nodo["nodeId"], mapa, err

    listos = 0
    with ThreadPoolExecutor(max_workers=workers) as ex:
        futs = [ex.submit(_job, n) for n in nodos]
        for fut in as_completed(futs):
            nid, mapa, err = fut.result()
            por_nodo[nid] = mapa
            if err:
                errores[nid] = err
            listos += 1
            if listos % 20 == 0 or listos == len(nodos):
                print(f"  Progreso API: {listos}/{len(nodos)} | errores: {len(errores)}", flush=True)

    try:
        from zoneinfo import ZoneInfo

        hoy_chile = datetime.now(ZoneInfo("America/Santiago")).date()
    except Exception:
        hoy_chile = datetime.now(timezone.utc).date()

    out_dir = args.output_dir.resolve()
    docx_path, csv_path = crear_informe(
        nodos,
        por_nodo,
        errores,
        dias_a,
        dias_b,
        out_dir,
        hoy_chile,
    )
    print(f"[OK] DOCX: {docx_path}")
    print(f"[OK] CSV:  {csv_path}")
    pdf = convertir_a_pdf(docx_path)
    if not pdf:
        print("[ERROR] No se pudo convertir a PDF (falta LibreOffice/soffice).", flush=True)
        return 1
    print(f"[OK] PDF:  {pdf}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
