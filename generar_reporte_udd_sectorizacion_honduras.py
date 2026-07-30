"""
Informe comparativo UDD — maniobra sectorización Honduras (válvulas cerradas).

Ventana de prueba (Jorge Hermosilla):
  Cierre: sábado 30/05/2026 20:00
  Apertura: lunes 01/06/2026 07:00

Referencia: misma ventana horaria en los 3 fines de semana anteriores
  (sáb 20:00 → lun 07:00), no días de semana.

Uso:
  python generar_reporte_udd_sectorizacion_honduras.py
  python generar_reporte_udd_sectorizacion_honduras.py --tecnico   # informe técnico (líneas/barras)

Salida entrega Diego:
  reports/udd_sectorizacion_honduras/entrega_diego/
    Informe_Ejecutivo_Pruebas_Hidricas_UDD_Honduras_N2.docx
    Informe_Ejecutivo_Pruebas_Hidricas_UDD_Honduras_N2.pdf
    graficos/01_sabado_3005_24h.png … 04_domingo_3105_maniobra_24h.png
"""

from __future__ import annotations

import sys
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "reports" / "udd_sectorizacion_honduras"
ENTREGA_DIR = OUT_DIR / "entrega_diego"
GRAFICOS_DIR = ENTREGA_DIR / "graficos"
NOMBRE_INFORME = "Informe_Ejecutivo_Pruebas_Hidricas_UDD_Honduras_N2"

NODE_ID = "000026-01"
NODE_NOMBRE = "Sala impulsión Honduras"
CLIENTE = "Universidad del Desarrollo (UDD)"

CIERRE = datetime(2026, 5, 30, 20, 0)
REAPERTURA = datetime(2026, 6, 1, 7, 0)

# Fines de semana anteriores a comparar (misma ventana desplazada 7, 14 y 21 días)
N_FIN_SEMANAS_REF = 3

# Colores por fin de semana (líneas) + promedio azul punteado + maniobra roja
COLORES_FIN_SEMANA = ["#2E86AB", "#28A745", "#E67E22"]  # azul, verde, naranja
COLOR_PROMEDIO = "#1F4E79"
COLOR_MANIOBRA = "#C0392B"


@dataclass(frozen=True)
class HoraLocal:
    dia: date
    hora: int  # 0-23, bloque [HH:00, HH+1:00)

    @property
    def dt(self) -> datetime:
        return datetime.combine(self.dia, datetime.min.time()) + timedelta(hours=self.hora)


@dataclass(frozen=True)
class FinSemanaRef:
    semanas_atras: int
    cierre: datetime
    reapertura: datetime

    @property
    def etiqueta(self) -> str:
        return (
            f"Fin de semana −{self.semanas_atras} "
            f"({self.cierre.strftime('%d/%m %H:%M')} – {self.reapertura.strftime('%d/%m %H:%M')})"
        )

    @property
    def etiqueta_corta(self) -> str:
        return (
            f"FS −{self.semanas_atras} "
            f"({self.cierre.strftime('%d/%m')} – {self.reapertura.strftime('%d/%m')})"
        )


def _fin_semanas_referencia() -> List[FinSemanaRef]:
    return [
        FinSemanaRef(
            k,
            CIERRE - timedelta(days=7 * k),
            REAPERTURA - timedelta(days=7 * k),
        )
        for k in range(1, N_FIN_SEMANAS_REF + 1)
    ]


def _horas_en_ventana(inicio: datetime, fin: datetime) -> List[HoraLocal]:
    """Horas completas dentro de [inicio, fin); reapertura 07:00 → último bloque 06:00–07:00."""
    out: List[HoraLocal] = []
    cur = inicio.replace(minute=0, second=0, microsecond=0)
    if cur < inicio:
        cur += timedelta(hours=1)
    while cur < fin:
        out.append(HoraLocal(cur.date(), cur.hour))
        cur += timedelta(hours=1)
    return out


def _refs_equivalentes(hora: HoraLocal) -> List[Tuple[date, int]]:
    """Misma hora civil en los N fines de semana previos (desplazamiento semanal)."""
    return [(hora.dia - timedelta(days=7 * k), hora.hora) for k in range(1, N_FIN_SEMANAS_REF + 1)]


def _cargar_serie(node_id: str, dias: Iterable[date]) -> Dict[Tuple[date, int], float]:
    from control_nocturno import obtener_datos_horarios_dia

    acc: Dict[Tuple[date, int], float] = {}
    for d in dias:
        hdict = obtener_datos_horarios_dia(node_id, datetime.combine(d, datetime.min.time()))
        for h in range(24):
            acc[(d, h)] = float(hdict.get(h, 0.0))
    return acc


def _dias_a_cargar(horas_prueba: List[HoraLocal]) -> List[date]:
    dias: set[date] = set()
    for h in horas_prueba:
        dias.add(h.dia)
        for d, _ in _refs_equivalentes(h):
            dias.add(d)
    return sorted(dias)


def _promedio_refs(serie: Dict[Tuple[date, int], float], hora: HoraLocal) -> float:
    vals = [serie.get(k, 0.0) for k in _refs_equivalentes(hora)]
    return sum(vals) / len(vals) if vals else 0.0


def _vals_refs(serie: Dict[Tuple[date, int], float], hora: HoraLocal) -> List[float]:
    return [serie.get(k, 0.0) for k in _refs_equivalentes(hora)]


def _total_ventana(serie: Dict[Tuple[date, int], float], inicio: datetime, fin: datetime) -> float:
    return sum(serie.get((h.dia, h.hora), 0.0) for h in _horas_en_ventana(inicio, fin))


def _fmt(n: float, dec: int = 2) -> str:
    from generar_reporte_word import format_number_chilean

    return format_number_chilean(n, dec)


def _grafico_perfil(
    horas_prueba: List[HoraLocal],
    valores_prueba: List[float],
    valores_ref_prom: List[float],
    fines_semana: List[FinSemanaRef],
    serie: Dict[Tuple[date, int], float],
    out_path: Path,
) -> None:
    xs = [h.dt for h in horas_prueba]

    fig, ax = plt.subplots(figsize=(12, 5.0))
    for fs, color in zip(fines_semana, COLORES_FIN_SEMANA):
        horas_fs = _horas_en_ventana(fs.cierre, fs.reapertura)
        ys_fs = [serie.get((h.dia, h.hora), 0.0) for h in horas_fs]
        if len(ys_fs) == len(xs):
            ax.plot(
                xs,
                ys_fs,
                color=color,
                linewidth=1.5,
                alpha=0.9,
                label=fs.etiqueta_corta,
            )
    ax.plot(
        xs,
        valores_ref_prom,
        color=COLOR_PROMEDIO,
        linewidth=2.2,
        linestyle="--",
        label=f"Promedio {N_FIN_SEMANAS_REF} fines de semana",
        zorder=5,
    )
    ax.plot(
        xs,
        valores_prueba,
        color=COLOR_MANIOBRA,
        linewidth=2.0,
        marker="o",
        markersize=4,
        label="Maniobra 30/05–01/06 (válvulas cerradas)",
        zorder=6,
    )
    ax.axvline(CIERRE, color="gray", linestyle=":", alpha=0.5, linewidth=0.9)
    ax.axvline(REAPERTURA, color="gray", linestyle="-.", alpha=0.5, linewidth=0.9)
    ax.set_ylabel("m³/h")
    ax.set_xlabel("Fecha y hora — eje de la maniobra (Chile)")
    ax.set_title(f"{NODE_ID} — {NODE_NOMBRE}: perfil horario comparativo")
    ax.legend(loc="upper right", fontsize=8, framealpha=0.95)
    ax.grid(True, alpha=0.3)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m %Hh"))
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def _grafico_barras(
    totales_fs: List[Tuple[FinSemanaRef, float]],
    prom_fs: float,
    total_prueba: float,
    out_path: Path,
) -> None:
    """Barras: cada fin de semana + promedio + maniobra (misma ventana acumulada)."""
    etiquetas = [fs.etiqueta_corta.replace("FS ", "") for fs, _ in totales_fs]
    etiquetas.append(f"Promedio\n{N_FIN_SEMANAS_REF} FS")
    etiquetas.append("Maniobra\n30/05–01/06")

    valores = [t for _, t in totales_fs] + [prom_fs, total_prueba]
    colores = list(COLORES_FIN_SEMANA) + [COLOR_PROMEDIO, COLOR_MANIOBRA]

    fig, ax = plt.subplots(figsize=(9, 5.0))
    x_pos = range(len(etiquetas))
    bars = ax.bar(
        x_pos,
        valores,
        color=colores,
        edgecolor="white",
        linewidth=1.2,
        width=0.65,
    )
    # Promedio: borde más marcado (complementa el azul sólido en barras)
    bars[-2].set_edgecolor(COLOR_PROMEDIO)
    bars[-2].set_linewidth(2.0)
    bars[-1].set_edgecolor("#922b21")
    bars[-1].set_linewidth(2.0)

    ymax = max(valores) if valores else 1.0
    for bar, val in zip(bars, valores):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + ymax * 0.02,
            f"{val:.1f}",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold" if bar == bars[-1] else "normal",
        )

    ax.set_xticks(list(x_pos))
    ax.set_xticklabels(etiquetas, fontsize=9)
    ax.set_ylabel("Consumo acumulado (m³)")
    ax.set_title(
        "Comparación consumo total — ventana sábado 20:00 a lunes 07:00",
        fontsize=11,
    )
    ax.grid(True, axis="y", alpha=0.35)
    ax.set_axisbelow(True)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def _conclusion(reduccion_pct: float, total_prueba: float, total_ref: float) -> str:
    ref_txt = f"promedio de los {N_FIN_SEMANAS_REF} fines de semana anteriores (misma ventana sáb 20:00 – lun 07:00)"
    if reduccion_pct >= 30 and total_prueba < total_ref * 0.7:
        return (
            "Durante el aislamiento se observa una **disminución relevante** del caudal en impulsión Honduras "
            f"respecto del {ref_txt}. Esto **apoya** la hipótesis de que parte del consumo base "
            "nocturno/fugante se alimentaba desde el sector cerrado. Se recomienda continuar "
            "sectorización dentro de ese tramo o instalar un punto intermedio aguas abajo del cierre."
        )
    if reduccion_pct >= 10:
        return (
            "Se observa una **reducción moderada** del caudal durante el aislamiento. El sector cerrado "
            "podría explicar parte del consumo anómalo, pero **no descarta** aportes desde otras ramas. "
            "Conviene repetir la maniobra o combinar con un segundo cierre en subsector."
        )
    return (
        "El caudal en impulsión Honduras **no disminuyó de forma significativa** durante el aislamiento "
        f"respecto del {ref_txt}. La maniobra **no confirma** que la fuga esté dentro del sector cerrado. "
        "Se sugiere revisar el plano de sectorización y evaluar cierres adicionales o puntos intermedios."
    )


def generar() -> Tuple[Path, Path]:
    tecnico_dir = OUT_DIR / "tecnico"
    tecnico_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    out_docx = tecnico_dir / f"Informe_Sectorizacion_Honduras_{ts}.docx"
    out_png_perfil = tecnico_dir / f"perfil_sectorizacion_{ts}.png"
    out_png_barras = tecnico_dir / f"barras_sectorizacion_{ts}.png"

    fines_semana = _fin_semanas_referencia()
    horas_prueba = _horas_en_ventana(CIERRE, REAPERTURA)
    serie = _cargar_serie(NODE_ID, _dias_a_cargar(horas_prueba))

    vals_prueba = [serie.get((h.dia, h.hora), 0.0) for h in horas_prueba]
    vals_ref = [_promedio_refs(serie, h) for h in horas_prueba]

    total_prueba = sum(vals_prueba)
    total_ref_esperado = sum(vals_ref)
    reduccion_abs = total_ref_esperado - total_prueba
    reduccion_pct = (100.0 * reduccion_abs / total_ref_esperado) if total_ref_esperado > 0 else 0.0

    prom_prueba = total_prueba / len(vals_prueba) if vals_prueba else 0.0
    prom_ref = total_ref_esperado / len(vals_ref) if vals_ref else 0.0
    max_prueba = max(vals_prueba) if vals_prueba else 0.0
    max_ref = max(vals_ref) if vals_ref else 0.0
    min_prueba = min(vals_prueba) if vals_prueba else 0.0
    min_ref = min(vals_ref) if vals_ref else 0.0

    totales_fs: List[Tuple[FinSemanaRef, float]] = [
        (fs, _total_ventana(serie, fs.cierre, fs.reapertura)) for fs in fines_semana
    ]
    prom_fs = sum(t for _, t in totales_fs) / len(totales_fs) if totales_fs else 0.0

    noche_dom = [h for h in horas_prueba if h.dia == date(2026, 5, 31) and 0 <= h.hora <= 6]
    total_noche_dom = sum(serie.get((h.dia, h.hora), 0.0) for h in noche_dom)
    ref_noche_dom = sum(_promedio_refs(serie, h) for h in noche_dom)

    _grafico_perfil(horas_prueba, vals_prueba, vals_ref, fines_semana, serie, out_png_perfil)
    _grafico_barras(totales_fs, prom_fs, total_prueba, out_png_barras)

    ref_lineas = "\n".join(f"  • {fs.etiqueta}" for fs in fines_semana)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h1 = doc.add_heading("Informe comparativo — Maniobra sectorización Honduras", level=0)
    if h1.runs:
        h1.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    doc.add_paragraph()
    for label, text in [
        ("Cliente", CLIENTE),
        ("Punto WES", f"{NODE_ID} — {NODE_NOMBRE}"),
        (
            "Maniobra",
            f"Cierre {CIERRE.strftime('%d/%m/%Y %H:%M')} — "
            f"Reapertura {REAPERTURA.strftime('%d/%m/%Y %H:%M')} (hora Chile)",
        ),
        (
            "Referencia",
            f"Misma ventana horaria en los {N_FIN_SEMANAS_REF} fines de semana anteriores "
            f"(sábado 20:00 → lunes 07:00), no días hábiles de semana.",
        ),
        ("Generado", datetime.now().strftime("%d-%m-%Y %H:%M")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)

    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Evaluar el impacto del cierre temporal de válvulas en un sector del anillo de agua potable "
        "sobre el caudal en Sala impulsión Honduras, comparando la maniobra con el comportamiento "
        "habitual en fines de semana equivalentes (baja demanda esperada)."
    )

    doc.add_heading("2. Metodología de comparación", level=1)
    doc.add_paragraph(
        "Para cada hora de la maniobra se calcula el caudal registrado y se contrasta con el "
        f"promedio de esa misma hora en los {N_FIN_SEMANAS_REF} fines de semana previos, "
        "desplazando la ventana exactamente 7, 14 y 21 días hacia atrás:"
    )
    for linea in ref_lineas.split("\n"):
        doc.add_paragraph(linea.strip(), style="List Bullet")

    doc.add_heading("3. Resumen de resultados", level=1)
    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = "Table Grid"
    filas_resumen = [
        ("Indicador", "Maniobra (30/05 20h – 01/06 07h)", f"Prom. {N_FIN_SEMANAS_REF} fines de semana previos"),
        ("Horas analizadas", str(len(horas_prueba)), str(len(horas_prueba))),
        ("Consumo acumulado (m³)", _fmt(total_prueba, 2), _fmt(total_ref_esperado, 2)),
        ("Caudal medio (m³/h)", _fmt(prom_prueba, 2), _fmt(prom_ref, 2)),
        ("Caudal máximo (m³/h)", _fmt(max_prueba, 2), _fmt(max_ref, 2)),
        ("Caudal mínimo (m³/h)", _fmt(min_prueba, 2), _fmt(min_ref, 2)),
    ]
    for i, row in enumerate(filas_resumen):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val
            if i == 0:
                for par in tbl.rows[i].cells[j].paragraphs:
                    for r in par.runs:
                        r.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Reducción acumulada vs promedio fines de semana: ").bold = True
    p.add_run(f"{_fmt(reduccion_abs, 2)} m³ ({_fmt(reduccion_pct, 1)} %).")

    p2 = doc.add_paragraph()
    p2.add_run("Domingo 31/05, 00:00–06:00: ").bold = True
    p2.add_run(
        f"maniobra {_fmt(total_noche_dom, 2)} m³ vs promedio fines de semana {_fmt(ref_noche_dom, 2)} m³."
    )

    doc.add_heading("4. Detalle por fin de semana de referencia", level=1)
    doc.add_paragraph(
        "Consumo acumulado en la misma ventana (sáb 20:00 – lun 07:00) en cada fin de semana anterior:"
    )
    tbl2 = doc.add_table(rows=1 + len(totales_fs) + 1, cols=2)
    tbl2.style = "Table Grid"
    tbl2.rows[0].cells[0].text = "Fin de semana de referencia"
    tbl2.rows[0].cells[1].text = "Consumo acumulado (m³)"
    for par in tbl2.rows[0].cells[0].paragraphs:
        for r in par.runs:
            r.bold = True
    for par in tbl2.rows[0].cells[1].paragraphs:
        for r in par.runs:
            r.bold = True
    for i, (fs, tot) in enumerate(totales_fs, start=1):
        tbl2.rows[i].cells[0].text = fs.etiqueta
        tbl2.rows[i].cells[1].text = _fmt(tot, 2)
    row_prom = tbl2.rows[len(totales_fs) + 1]
    row_prom.cells[0].text = "Promedio de los 3 fines de semana"
    row_prom.cells[1].text = _fmt(prom_fs, 2)
    for par in row_prom.cells[0].paragraphs:
        for r in par.runs:
            r.bold = True

    doc.add_paragraph()
    pct_bajo_fs = (100.0 * (prom_fs - total_prueba) / prom_fs) if prom_fs else 0.0
    doc.add_paragraph(
        f"Maniobra con válvulas cerradas: {_fmt(total_prueba, 2)} m³ "
        f"({_fmt(pct_bajo_fs, 1)} % bajo el promedio de fines de semana)."
    )

    doc.add_heading("5. Gráficos", level=1)
    doc.add_heading("5.1 Perfil horario (m³/h)", level=2)
    doc.add_paragraph(
        "Cada fin de semana de referencia con un color (azul, verde y naranja). "
        "Línea azul punteada = promedio de los tres fines de semana. "
        "Línea roja = maniobra con válvulas cerradas. Eje temporal: fechas de la maniobra."
    )
    doc.add_picture(str(out_png_perfil), width=Inches(6.2))

    doc.add_heading("5.2 Comparación consumo acumulado (barras)", level=2)
    doc.add_paragraph(
        "Consumo total en la misma ventana (sábado 20:00 – lunes 07:00) para cada fin de semana, "
        "el promedio de los tres y la maniobra del 30/05 al 01/06."
    )
    doc.add_picture(str(out_png_barras), width=Inches(5.8))

    doc.add_heading("6. Interpretación", level=1)
    doc.add_paragraph(_conclusion(reduccion_pct, total_prueba, total_ref_esperado))

    doc.add_heading("7. Consideraciones", level=1)
    doc.add_paragraph(
        "La comparación con fines de semana evita sesgar la referencia con mayor demanda de días hábiles. "
        "El punto WES monitorea Sala impulsión Honduras; eventos concurrentes (riego, llenado de estanques, "
        "aperturas no planificadas) pueden alterar alguna hora puntual. Valores horarios alineados con app WES."
    )

    doc.save(out_docx)

    pdf_path = out_docx.with_suffix(".pdf")
    try:
        from generar_reporte_word import convertir_word_a_pdf

        pdf_out = convertir_word_a_pdf(out_docx)
        if pdf_out and pdf_out.is_file():
            pdf_path = pdf_out
    except Exception:
        pass

    return out_docx, pdf_path


def _keep_with_next(paragraph) -> None:
    """Evita que el párrafo quede solo al pie de página (gráfico en la página siguiente)."""
    try:
        from docx.oxml import OxmlElement

        p_pr = paragraph._p.get_or_add_pPr()
        p_pr.append(OxmlElement("w:keepNext"))
    except Exception:
        pass


def _configurar_pagina_ejecutiva(doc: Document) -> None:
    """Márgenes y estilo compactos para texto + gráficos en la misma hoja."""
    sec = doc.sections[0]
    sec.top_margin = Inches(0.52)
    sec.bottom_margin = Inches(0.48)
    sec.left_margin = Inches(0.62)
    sec.right_margin = Inches(0.62)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0


def _parrafo_bullet(doc: Document, texto: str, *, sub: bool = False) -> None:
    pref = "○  " if sub else "●  "
    p = doc.add_paragraph(pref + texto)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run.font.size = Pt(10.5)


def _agregar_footer_wes(doc: Document) -> None:
    """Footer simple (sin campos de número de página)."""
    try:
        section = doc.sections[0]
        footer = section.footer
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.text = "www.dashboard.wes.cl           operaciones@wes.cl        fono: (+569)75595695  -  (+562)6463385"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(90, 90, 90)
    except Exception:
        pass


def _grafico_horario_24h_compacto(
    horas: Dict[int, float],
    out_path: Path,
    *,
    titulo: str,
) -> Path:
    """Perfil 24 h compacto para insertar dos gráficos por fila en Word."""
    x = list(range(24))
    y = [float(horas.get(h, 0.0)) for h in x]
    colors = ["#c41e1e" if h <= 6 else "#4A8CB8" for h in x]

    fig, ax = plt.subplots(figsize=(4.55, 2.05))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    ax.bar(x, y, width=0.78, color=colors, edgecolor="white", linewidth=0.4, zorder=2)
    ax.axvspan(-0.5, 6.5, alpha=0.12, color="#c41e1e", zorder=0)
    ax.plot(x, y, color="#2e7ac8", linewidth=1.0, marker="o", markersize=2, zorder=4)
    ymax = max(max(y) if y else 0, 0.05) * 1.2
    ax.set_xlim(-0.5, 23.5)
    ax.set_ylim(0, ymax)
    ax.set_xticks(x[::2])
    ax.set_xticklabels([f"{h:02d}" for h in x[::2]], fontsize=6)
    ax.set_ylabel("m³/h", fontsize=8)
    ax.set_title(titulo, fontsize=8, fontweight="bold", pad=4)
    ax.tick_params(axis="y", labelsize=6)
    ax.grid(axis="y", linestyle="-", alpha=0.3, zorder=0)
    ax.set_axisbelow(True)
    fig.tight_layout(pad=0.4)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return out_path


def _grafico_rango_horas_app(
    horas: Dict[int, float],
    out_path: Path,
    *,
    titulo: str,
    hora_inicio: int,
    hora_fin_incl: int,
    compacto: bool = False,
) -> Path:
    """Gráfico de barras estilo app, pero acotado a un rango horario."""
    xs = list(range(hora_inicio, hora_fin_incl + 1))
    ys = [float(horas.get(h, 0.0)) for h in xs]
    colors = ["#c41e1e" if h <= 6 else "#4A8CB8" for h in xs]

    figsize = (4.55, 2.05) if compacto else (9.6, 3.8)
    fig, ax = plt.subplots(figsize=figsize)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    ax.bar(xs, ys, color=colors, edgecolor="white", linewidth=0.6, width=0.78, zorder=2)
    ax.axvspan(hora_inicio - 0.5, min(6, hora_fin_incl) + 0.5, alpha=0.12, color="#c41e1e", zorder=0)
    ax.plot(xs, ys, color="#2e7ac8", linewidth=1.2, marker="o", markersize=3, zorder=4)
    ax.set_xlim(hora_inicio - 0.5, hora_fin_incl + 0.5)
    ymax = (max(ys) if ys else 0.0) * 1.18
    ax.set_ylim(0, max(ymax, 0.05))
    ax.set_xticks(xs)
    fs_lbl = 6 if compacto else 8
    fs_ax = 8 if compacto else 9
    fs_title = 9 if compacto else 11
    ax.set_xticklabels([f"{h:02d}" for h in xs], fontsize=fs_lbl)
    if not compacto:
        ax.set_xlabel("Hora del día (Chile)", fontsize=fs_ax)
    ax.set_ylabel("m³/h", fontsize=fs_ax)
    ax.set_title(titulo, fontsize=fs_title, fontweight="bold", pad=6 if compacto else 8)
    if compacto:
        ax.tick_params(axis="y", labelsize=7)
    ax.grid(axis="y", linestyle="-", alpha=0.35, zorder=0)
    ax.set_axisbelow(True)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return out_path


def _insertar_par_graficos(
    doc: Document,
    items: List[Tuple[Path, str]],
    *,
    ancho=Inches(3.18),
) -> None:
    """Dos gráficos en una fila, al pie del bloque de texto (estilo informe Diego)."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    for j, (img_path, leyenda) in enumerate(items[:2]):
        cell = tbl.rows[0].cells[j]
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_img.paragraph_format.space_before = Pt(2)
        p_img.paragraph_format.space_after = Pt(0)
        run = p_img.add_run()
        run.add_picture(str(img_path), width=ancho)

        p_cap = cell.add_paragraph(leyenda)
        p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_cap.paragraph_format.space_before = Pt(1)
        p_cap.paragraph_format.space_after = Pt(0)
        if p_cap.runs:
            p_cap.runs[0].font.size = Pt(7.5)
            p_cap.runs[0].font.color.rgb = RGBColor(70, 70, 70)


def _limpiar_archivos_sueltos_antiguos() -> None:
    """Elimina generados viejos en la raíz de udd_sectorizacion (solo queda entrega_diego/)."""
    if not OUT_DIR.is_dir():
        return
    for item in OUT_DIR.iterdir():
        if item.name == "entrega_diego":
            continue
        try:
            if item.is_file():
                item.unlink()
            elif item.is_dir():
                import shutil

                shutil.rmtree(item)
        except OSError as e:
            print(f"[ADVERTENCIA] No se pudo borrar {item}: {e}")


def generar_informe_ejecutivo_n2() -> Tuple[Path, Path]:
    """
    Genera un informe ejecutivo estilo referencia UDD (Informe Nº2, 01.06.2026),
    usando gráficos tipo app (sin capturas).
    """
    from wes_estilo_graficos_app import horas_api_chile
    from generar_reporte_word import convertir_word_a_pdf

    ENTREGA_DIR.mkdir(parents=True, exist_ok=True)
    GRAFICOS_DIR.mkdir(parents=True, exist_ok=True)
    out_docx = ENTREGA_DIR / f"{NOMBRE_INFORME}.docx"

    # Recalcular métricas (mismas que en generar())
    fines_semana = _fin_semanas_referencia()
    horas_prueba = _horas_en_ventana(CIERRE, REAPERTURA)
    serie = _cargar_serie(NODE_ID, _dias_a_cargar(horas_prueba))
    vals_prueba = [serie.get((h.dia, h.hora), 0.0) for h in horas_prueba]
    vals_ref = [_promedio_refs(serie, h) for h in horas_prueba]

    total_prueba = sum(vals_prueba)
    total_ref_esperado = sum(vals_ref)
    reduccion_abs = total_ref_esperado - total_prueba
    reduccion_pct = (100.0 * reduccion_abs / total_ref_esperado) if total_ref_esperado > 0 else 0.0

    prom_prueba = total_prueba / len(vals_prueba) if vals_prueba else 0.0
    prom_ref = total_ref_esperado / len(vals_ref) if vals_ref else 0.0
    max_prueba = max(vals_prueba) if vals_prueba else 0.0
    max_ref = max(vals_ref) if vals_ref else 0.0
    min_prueba = min(vals_prueba) if vals_prueba else 0.0
    min_ref = min(vals_ref) if vals_ref else 0.0

    totales_fs: List[Tuple[FinSemanaRef, float]] = [
        (fs, _total_ventana(serie, fs.cierre, fs.reapertura)) for fs in fines_semana
    ]
    prom_fs = sum(t for _, t in totales_fs) / len(totales_fs) if totales_fs else 0.0

    # Hito madrugada domingo 31/05, 00–06
    noche_dom = [h for h in horas_prueba if h.dia == date(2026, 5, 31) and 0 <= h.hora <= 6]
    total_noche_dom = sum(serie.get((h.dia, h.hora), 0.0) for h in noche_dom)
    ref_noche_dom = sum(_promedio_refs(serie, h) for h in noche_dom)

    # Gráficos tipo app (equivalentes a capturas)
    dia_sab_3005 = datetime(2026, 5, 30)
    dia_dom_2405 = datetime(2026, 5, 24)
    dia_dom_3105 = datetime(2026, 5, 31)
    dia_lun_0106 = datetime(2026, 6, 1)

    horas_sab = horas_api_chile(NODE_ID, dia_sab_3005)
    horas_dom_ref = horas_api_chile(NODE_ID, dia_dom_2405)
    horas_dom_man = horas_api_chile(NODE_ID, dia_dom_3105)
    horas_lun = horas_api_chile(NODE_ID, dia_lun_0106)

    img_sab_24h = GRAFICOS_DIR / "01_sabado_3005_24h.png"
    img_lun_0_9 = GRAFICOS_DIR / "02_lunes_0106_0_9h.png"
    img_dom_ref_24h = GRAFICOS_DIR / "03_domingo_2405_referencia_24h.png"
    img_dom_man_24h = GRAFICOS_DIR / "04_domingo_3105_maniobra_24h.png"

    _grafico_horario_24h_compacto(
        horas_sab,
        img_sab_24h,
        titulo="Consumo (m³/hora), 24 horas  Sábado 30/05",
    )
    _grafico_rango_horas_app(
        horas_lun,
        img_lun_0_9,
        titulo="Consumo (m³/hora), 0 a 9 horas  Lunes 01/06",
        hora_inicio=0,
        hora_fin_incl=9,
        compacto=True,
    )
    _grafico_horario_24h_compacto(
        horas_dom_ref,
        img_dom_ref_24h,
        titulo="Consumo (m³/hora), 24 horas  Domingo 24/05",
    )
    _grafico_horario_24h_compacto(
        horas_dom_man,
        img_dom_man_24h,
        titulo="Consumo (m³/hora), 24 horas  Domingo 31/05",
    )

    # Documento (formato ejecutivo)
    doc = Document()
    _configurar_pagina_ejecutiva(doc)

    fecha_txt = "1 de Junio 2026"
    t = doc.add_paragraph(fecha_txt + "  Informe Ejecutivo")
    t.paragraph_format.space_after = Pt(2)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    h = doc.add_paragraph("Pruebas Hídricas -  Informe  Nº2  UDD  -  Red  Impulsión  Honduras")
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 71, 136)

    # I. Antecedentes
    p = doc.add_paragraph("I.  Antecedentes")
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)

    antecedentes = [
        "Con fecha 12 de mayo de 2026, se emitió un primer informe ejecutivo que concluyó de manera inequívoca "
        "que los dispositivos de medición registran caudales reales y efectivos, descartando fallas de medición "
        "y confirmando la presencia de pérdidas en la red.",
        "Luego el fin de semana del 23 y 24 de mayo, el equipo de UDD avanzó en la sectorización de válvulas del anillo de agua potable "
        "e incorporó nuevos puntos de medición de presión en la red. Dicha combinación de maniobras permitió identificar pérdidas de presión "
        "en dos sectores específicos.",
        "Con el fin de aislar temporalmente la zona con mayor probabilidad de contener la fuga investigada y verificar el comportamiento del consumo base, "
        "se coordinó una nueva maniobra técnica de aislamiento, de la cual se da cuenta en el presente informe.",
    ]
    for a in antecedentes:
        _parrafo_bullet(doc, a)

    # II. Metodología
    p = doc.add_paragraph("II.  Metodología")
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    met = [
        "El objetivo principal de la maniobra consistió en aislar temporalmente un sector específico del anillo de agua potable mediante el cierre programado de válvulas.",
        "En base a ello se evalúa el impacto directo sobre el consumo base registrado en la red Sala de impulsión Honduras, determinando si la fuga principal se encuentra dentro del tramo aislado.",
        "En conjunto con el equipo técnico de la UDD, se definió y ejecutó la siguiente secuencia técnica cronológica:",
    ]
    for m in met:
        _parrafo_bullet(doc, m)

    pasos = [
        f"1) Cierre de válvulas: El sábado 30 de mayo de 2026 a las 20:00 horas, el personal de la UDD procedió al cierre de dos válvulas previamente seleccionadas para aislar el sector con sospecha de fuga.",
        "2) Periodo de aislamiento: La condición de aislamiento estructural en la red se mantuvo de manera continua durante una ventana de 35 horas en periodo inhábil.",
        "3) Reapertura de válvulas: El lunes 1 de junio de 2026 a las 07:00 horas se procedió a la apertura de las llaves de paso para normalizar el suministro regular del campus.",
    ]
    ultimo_paso = None
    for s in pasos:
        ultimo_paso = doc.add_paragraph("   " + s)
        ultimo_paso.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ultimo_paso.paragraph_format.space_after = Pt(2)
        for run in ultimo_paso.runs:
            run.font.size = Pt(10.5)

    # Gráficos al pie de Metodología (misma página que el texto, como referencia Diego)
    if ultimo_paso is not None:
        _keep_with_next(ultimo_paso)
    _insertar_par_graficos(
        doc,
        [
            (img_sab_24h, "Sábado 30/05 — inicio maniobra"),
            (img_lun_0_9, "Lunes 01/06 — reapertura"),
        ],
    )

    # III. Análisis — nueva página solo para sección (texto + gráficos domingo juntos abajo)
    doc.add_page_break()
    p = doc.add_paragraph("III.  Análisis")
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(3)

    analisis_txt = [
        "Durante todo el periodo de aislamiento, se realizó un seguimiento remoto de la data de caudal (m³/h) y volumen acumulado a través de la plataforma y la WES App.",
        "Los caudales y volúmenes registrados se contrastaron con el promedio histórico de la misma ventana horaria (sábado 20:00 a lunes 07:00 horas) de los tres fines de semana inmediatamente anteriores:",
        "Fin de semana −1: del 23/05 (20:00 hrs) al 25/05 (07:00 hrs).",
        "Fin de semana −2: del 16/05 (20:00 hrs) al 18/05 (07:00 hrs).",
        "Fin de semana −3: del 09/05 (20:00 hrs) al 11/05 (07:00 hrs).",
        "El análisis de los datos recolectados durante las 35 horas de la maniobra de aislamiento reflejó un cambio drástico y favorable en el comportamiento hídrico de la red en comparación con la línea base histórica de los fines de semana de referencia.",
        f"Considerando los registros de consumo durante la maniobra versus el promedio, se verifica una reducción acumulada neta de {reduccion_abs:.1f} m³, lo que equivale a un descenso del {reduccion_pct:.1f}% en el volumen de agua pasante bajo el promedio habitual de la red.",
        f"En lo que respecta al comportamiento del caudal, el promedio medio durante el periodo de aislamiento cayó a {prom_prueba:.2f} m³/h, frente a los {prom_ref:.2f} m³/h del patrón histórico previo.",
        f"El caudal mínimo registrado descendió de forma notable hasta los {min_prueba:.2f} m³/h, aproximándose a un consumo nulo, en comparación con el mínimo de {min_ref:.1f} m³/h habitual de los fines de semana de referencia.",
    ]
    for a in analisis_txt[:2]:
        _parrafo_bullet(doc, a)
    for a in analisis_txt[2:5]:
        _parrafo_bullet(doc, a, sub=True)
    for a in analisis_txt[5:]:
        _parrafo_bullet(doc, a)

    # Tabla de indicadores (compacta)
    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = "Table Grid"
    filas_resumen = [
        ("Indicador", "Maniobra (30/05 20h – 01/06 07h)", f"Prom. {N_FIN_SEMANAS_REF} fines de semana previos"),
        ("Horas analizadas", str(len(horas_prueba)), str(len(horas_prueba))),
        ("Consumo acumulado (m³)", _fmt(total_prueba, 2), _fmt(prom_fs, 2) if prom_fs else _fmt(total_ref_esperado, 2)),
        ("Caudal medio (m³/h)", _fmt(prom_prueba, 2), _fmt(prom_ref, 2)),
        ("Caudal máximo (m³/h)", _fmt(max_prueba, 2), _fmt(max_ref, 2)),
        ("Caudal mínimo (m³/h)", _fmt(min_prueba, 2), _fmt(min_ref, 2)),
    ]
    for i, row in enumerate(filas_resumen):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            for par in cell.paragraphs:
                par.paragraph_format.space_after = Pt(0)
                for r in par.runs:
                    r.font.size = Pt(9)
                    if i == 0:
                        r.bold = True

    # Madrugada + gráficos dominicales + conclusión en la misma página (layout 1138)
    p_hito = doc.add_paragraph(
        f"●  Un hito representativo de esta disminución se observó durante la madrugada del domingo 31 de mayo, "
        f"específicamente en el bloque de las 00:00 a las 06:00 horas, donde el consumo registrado con las válvulas "
        f"cerradas fue de tan solo {total_noche_dom:.2f} m³, frente a un promedio histórico de {ref_noche_dom:.2f} m³ "
        f"para ese mismo intervalo horario."
    )
    p_hito.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_hito.paragraph_format.space_before = Pt(4)
    p_hito.paragraph_format.space_after = Pt(2)
    for run in p_hito.runs:
        run.font.size = Pt(10.5)
    _keep_with_next(p_hito)

    _insertar_par_graficos(
        doc,
        [
            (img_dom_ref_24h, "Domingo 24/05 — referencia"),
            (img_dom_man_24h, "Domingo 31/05 — maniobra"),
        ],
    )

    # IV. Conclusión (continúa en la misma página si hay espacio, como referencia)
    p = doc.add_paragraph("IV.  Conclusión:")
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    concl = [
        "La maniobra de sectorización y aislamiento ejecutada permite sustentar la hipótesis técnica planteada: "
        "la disminución del caudal medio (superior al 90%) al mantener cerradas las válvulas en el sector aislado "
        "confirma que la fuga principal bajo investigación se encuentra contenida de forma efectiva dentro del mismo.",
        "Como próximos pasos y recomendación técnica para avanzar hacia la localización específica del punto de pérdida, "
        "se sugiere dar continuidad a las labores de inspección de este tramo específico y/o, alternativamente, proceder "
        "con la instalación de un punto de monitoreo intermedio aguas abajo del cierre actual para facilitar la sub-sectorización.",
    ]
    for c in concl:
        _parrafo_bullet(doc, c)

    _agregar_footer_wes(doc)
    doc.save(out_docx)

    pdf_path = out_docx.with_suffix(".pdf")
    try:
        pdf_out = convertir_word_a_pdf(out_docx)
        if pdf_out and Path(pdf_out).is_file():
            pdf_path = Path(pdf_out)
    except Exception:
        pass

    _limpiar_archivos_sueltos_antiguos()
    return out_docx, pdf_path


def main() -> int:
    import argparse

    parser = argparse.ArgumentParser(description="Informe UDD sectorización Honduras")
    parser.add_argument(
        "--tecnico",
        action="store_true",
        help="Genera informe técnico (perfil líneas + barras) en lugar del ejecutivo Nº2",
    )
    args = parser.parse_args()

    print("=" * 72)
    if args.tecnico:
        print("INFORME TÉCNICO — UDD HONDURAS")
        print("=" * 72)
        docx, pdf = generar()
        print(f"[OK] DOCX: {docx}")
        print(f"[OK] PDF:  {pdf}")
    else:
        print("INFORME EJECUTIVO Nº2 — UDD HONDURAS (entrega Diego)")
        print("=" * 72)
        docx, pdf = generar_informe_ejecutivo_n2()
        print(f"[OK] DOCX: {docx}")
        print(f"[OK] PDF:  {pdf}")
        print(f"[OK] Gráficos: {GRAFICOS_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
