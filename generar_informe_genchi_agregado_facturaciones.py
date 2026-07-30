"""
Informe agregado — CDP Puente Alto, CPF San Miguel y CCP Santiago Sur (Genchi).

Uso:
  python generar_informe_genchi_agregado_facturaciones.py
"""

from __future__ import annotations

import csv
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
GENCHI_DIR = ROOT / "reports" / "Genchi"
OUT_BASE = GENCHI_DIR / "informe"
FECHA_DESDE = date(2025, 1, 1)
MESES_EXCLUIDOS = {"2025-02"}
MESES_PROYECCION = 12
_MESES_NOM = {1: "ene", 2: "feb", 3: "mar", 4: "abr", 5: "may", 6: "jun", 7: "jul", 8: "ago", 9: "sep", 10: "oct", 11: "nov", 12: "dic"}
_HEADING = RGBColor(31, 71, 136)


@dataclass
class SiteConfig:
    slug: str
    nombre: str
    cuenta: str
    medidor: str
    factura_excluida: str
    fecha_retiro_wes: date
    fecha_fin_con_wes: date
    site_dir: Path

    @property
    def facturas_dir(self) -> Path:
        for sub in ("Facturaciones", "facturaciones"):
            p = self.site_dir / sub
            if p.is_dir():
                return p
        return self.site_dir / "Facturaciones"


SITES = [
    SiteConfig("cdp", "CDP Detención Preventiva Puente Alto", "1008941-7", "120715107", "8427884", date(2025, 6, 28), date(2025, 6, 27), GENCHI_DIR / "CDP Puente Alto"),
    SiteConfig("cpf", "CPF San Miguel", "1008398-2", "119758364", "8428005", date(2025, 7, 2), date(2025, 7, 1), GENCHI_DIR / "CPF San Miguel"),
    SiteConfig("ccp", "CCP Santiago Sur", "1007968-3", "723461", "8427552", date(2025, 6, 30), date(2025, 6, 29), GENCHI_DIR / "CCP Santiago Sur"),
]


@dataclass
class Periodo:
    boleta: str
    lectura_desde: date
    lectura_hasta: date
    m3: int
    clp: int
    pdf: str
    cfg: SiteConfig
    clp_esperado: int = 0
    sobrecosto_clp: int = 0

    @property
    def dias(self) -> int:
        return (self.lectura_hasta - self.lectura_desde).days

    @property
    def etiqueta(self) -> str:
        return "Con monitoreo WES" if self.lectura_hasta <= self.cfg.fecha_fin_con_wes else "Sin monitoreo WES"


@dataclass
class ResultadoUnidad:
    cfg: SiteConfig
    periodos: List[Periodo]
    meses: List[dict]
    prom: dict
    stats: dict


def _fmt_clp(n: int) -> str:
    s = f"{abs(int(n)):,}".replace(",", ".")
    return f"{s} CLP" if n >= 0 else f"-{s} CLP"


def _fmt_m3(n: float) -> str:
    return f"{n:,.0f} m³".replace(",", ".")


def _fmt_mes(key: str) -> str:
    y, m = key.split("-")
    return f"{_MESES_NOM[int(m)]}-{y}"


def _fmt_mes_corto(d: date) -> str:
    return f"{_MESES_NOM[d.month]}-{d.year}"


def _extraer_clp_pdf(pdf: Path) -> int:
    from facturacion_aguas_andinas_pdf import extraer_texto_pdf
    txt = extraer_texto_pdf(pdf)
    m = re.search(r"TOTAL\s+A\s+PAGAR\s*\$\s*([\d\.]+)", txt, flags=re.IGNORECASE)
    if not m:
        raise ValueError(f"No se encontró TOTAL A PAGAR en {pdf.name}")
    return int(m.group(1).replace(".", ""))


def _cargar_periodos(cfg: SiteConfig) -> List[Periodo]:
    from facturacion_aguas_andinas_pdf import listar_periodos_desde_pdf
    raw: List[Periodo] = []
    for pdf in sorted(cfg.facturas_dir.glob("*.pdf")):
        clp = _extraer_clp_pdf(pdf)
        for p in listar_periodos_desde_pdf(pdf):
            raw.append(Periodo(p.boleta, p.lectura_anterior.date(), p.lectura_actual.date(), p.m3_cuenta, clp, pdf.name, cfg))
    raw.sort(key=lambda x: x.lectura_hasta)
    return raw


def _filtrar(periodos: List[Periodo], cfg: SiteConfig) -> List[Periodo]:
    return [p for p in periodos if p.boleta != cfg.factura_excluida and p.lectura_hasta >= FECHA_DESDE]


def _aplicar_sobrecosto(periodos: List[Periodo], con_wes: List[Periodo]) -> float:
    m3_base = sum(p.m3 for p in con_wes)
    clp_base = sum(p.clp for p in con_wes)
    clp_m3 = clp_base / m3_base if m3_base else 0.0
    for p in periodos:
        p.clp_esperado = int(round(p.m3 * clp_m3))
        p.sobrecosto_clp = p.clp - p.clp_esperado
    return clp_m3


def _estado_mes(key: str) -> str:
    y, m = map(int, key.split("-"))
    return "Sin WES" if date(y, m, 1) >= date(2025, 7, 1) else "Con WES"


def _prorrateo_mensual(periodos: List[Periodo]) -> List[dict]:
    agg: Dict[str, dict] = defaultdict(lambda: {"m3": 0.0, "clp": 0, "clp_esp": 0, "dias": 0})
    for p in periodos:
        if p.dias <= 0:
            continue
        d = max(p.lectura_desde, FECHA_DESDE)
        fin = p.lectura_hasta
        while d < fin:
            fin_mes = date(d.year + 1, 1, 1) if d.month == 12 else date(d.year, d.month + 1, 1)
            tramo_fin = min(fin, fin_mes)
            dias = (tramo_fin - d).days
            if dias <= 0:
                d = tramo_fin
                continue
            key = f"{d.year}-{d.month:02d}"
            if key in MESES_EXCLUIDOS or date(d.year, d.month, 1) < FECHA_DESDE:
                d = tramo_fin
                continue
            frac = dias / p.dias
            agg[key]["m3"] += p.m3 * frac
            agg[key]["clp"] += int(round(p.clp * frac))
            agg[key]["clp_esp"] += int(round(p.clp_esperado * frac))
            agg[key]["dias"] += dias
            d = tramo_fin
    return [{"mes": k, "mes_label": _fmt_mes(k), "estado": _estado_mes(k), "dias": a["dias"], "m3": a["m3"], "clp": a["clp"], "clp_esperado": a["clp_esp"], "sobrecosto_clp": a["clp"] - a["clp_esp"]} for k, a in sorted(agg.items())]


def _meses_plot(meses: List[dict]) -> List[dict]:
    return meses[:-1] if len(meses) > 1 else meses


def _promedios(meses: List[dict]) -> dict:
    plot = _meses_plot(meses)
    con = [m for m in plot if m["estado"] == "Con WES"]
    sin = [m for m in plot if m["estado"] == "Sin WES"]

    def avg(arr: List[dict], key: str) -> float:
        return sum(m[key] for m in arr) / len(arr) if arr else 0.0

    def pct(a: float, b: float) -> float:
        return ((b / a) - 1) * 100 if a else 0.0

    prom_m3_con, prom_m3_sin = avg(con, "m3"), avg(sin, "m3")
    prom_clp_con, prom_clp_sin = avg(con, "clp"), avg(sin, "clp")
    clp_m3_con = prom_clp_con / prom_m3_con if prom_m3_con else 0.0
    clp_m3_sin = prom_clp_sin / prom_m3_sin if prom_m3_sin else 0.0
    return {
        "n_meses_con": len(con), "n_meses_sin": len(sin),
        "prom_m3_con": prom_m3_con, "prom_m3_sin": prom_m3_sin,
        "prom_clp_con": prom_clp_con, "prom_clp_sin": prom_clp_sin,
        "clp_m3_con": clp_m3_con, "clp_m3_sin": clp_m3_sin,
        "pct_m3": pct(prom_m3_con, prom_m3_sin), "pct_clp": pct(prom_clp_con, prom_clp_sin),
        "pct_clp_m3": pct(clp_m3_con, clp_m3_sin),
        "proj_m3_con": prom_m3_con * MESES_PROYECCION, "proj_m3_sin": prom_m3_sin * MESES_PROYECCION,
        "proj_clp_con": int(round(prom_clp_con * MESES_PROYECCION)),
        "proj_clp_sin": int(round(prom_clp_sin * MESES_PROYECCION)),
        "sobrecosto_anual_clp": int(round(prom_clp_sin * MESES_PROYECCION)) - int(round(prom_clp_con * MESES_PROYECCION)),
    }


def _analizar_unidad(cfg: SiteConfig) -> ResultadoUnidad:
    todos = _cargar_periodos(cfg)
    excluida = next((p for p in todos if p.boleta == cfg.factura_excluida), None)
    periodos = _filtrar(todos, cfg)
    con_wes = [p for p in periodos if p.lectura_hasta <= cfg.fecha_fin_con_wes]
    sin_wes = [p for p in periodos if p.lectura_hasta > cfg.fecha_fin_con_wes]
    clp_m3_base = _aplicar_sobrecosto(periodos, con_wes)
    meses = _prorrateo_mensual(periodos)
    prom = _promedios(meses)
    stats = {
        "n_con_wes": len(con_wes), "clp_m3_base": clp_m3_base,
        "total_m3_con": sum(p.m3 for p in con_wes), "total_m3_sin": sum(p.m3 for p in sin_wes),
        "total_clp_con": sum(p.clp for p in con_wes), "total_clp_sin": sum(p.clp for p in sin_wes),
        "total_clp_ventana": sum(p.clp for p in periodos), "total_m3_ventana": sum(p.m3 for p in periodos),
        "total_clp_esperado": sum(p.clp_esperado for p in periodos),
        "sobrecosto_sin": sum(p.sobrecosto_clp for p in sin_wes),
        "clp_excluida": excluida.clp if excluida else 0,
        "label_fin_sin": _fmt_mes_corto(max((p.lectura_hasta for p in sin_wes), default=date.today())),
    }
    return ResultadoUnidad(cfg, periodos, meses, prom, stats)


def _agregar_meses(unidades: List[ResultadoUnidad]) -> List[dict]:
    agg: Dict[str, dict] = defaultdict(lambda: {"m3": 0.0, "clp": 0, "clp_esp": 0, "dias": 0})
    for u in unidades:
        for m in u.meses:
            k = m["mes"]
            agg[k]["m3"] += m["m3"]
            agg[k]["clp"] += m["clp"]
            agg[k]["clp_esp"] += m["clp_esperado"]
            agg[k]["dias"] += m["dias"]
    return [{"mes": k, "mes_label": _fmt_mes(k), "estado": _estado_mes(k), "dias": a["dias"], "m3": a["m3"], "clp": a["clp"], "clp_esperado": a["clp_esp"], "sobrecosto_clp": a["clp"] - a["clp_esp"]} for k, a in sorted(agg.items())]


def _stats_agregado(unidades: List[ResultadoUnidad], prom: dict) -> dict:
    s = {k: sum(u.stats[k] for u in unidades) for k in ("total_m3_con", "total_m3_sin", "total_clp_con", "total_clp_sin", "total_clp_ventana", "total_m3_ventana", "total_clp_esperado", "sobrecosto_sin", "clp_excluida")}
    s["clp_m3_base"] = s["total_clp_con"] / s["total_m3_con"] if s["total_m3_con"] else 0.0
    s["conclusiones"] = [
        f"Tres recintos Genchi (ene-2025+): {_fmt_m3(s['total_m3_ventana'])} por {_fmt_clp(s['total_clp_ventana'])}.",
        f"Con WES agregado (mar–jun 2025): {_fmt_m3(s['total_m3_con'])} por {_fmt_clp(s['total_clp_con'])}.",
        f"Sin WES agregado (jul-2025+): {_fmt_m3(s['total_m3_sin'])} por {_fmt_clp(s['total_clp_sin'])}; sobrecosto {_fmt_clp(s['sobrecosto_sin'])}.",
        f"Promedio mensual agregado sin WES vs con WES: {prom['pct_clp']:+.1f}% CLP, {prom['pct_m3']:+.1f}% m³.",
        f"Proyección anual agregada (×{MESES_PROYECCION}): con WES {_fmt_clp(prom['proj_clp_con'])}; sin WES {_fmt_clp(prom['proj_clp_sin'])} (dif. {_fmt_clp(prom['sobrecosto_anual_clp'])}).",
        "Línea base «esperado» calculada por recinto; feb-2025 excluido en cada unidad.",
    ]
    return s


def _keep_with_next(paragraph) -> None:
    try:
        p_pr = paragraph._p.get_or_add_pPr()
        if p_pr.find(qn("w:keepNext")) is None:
            p_pr.append(OxmlElement("w:keepNext"))
    except Exception:
        pass


def _tabla_compacta(table, font_pt: float = 9.0) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(font_pt)
                    run.font.name = "Calibri"


def _grafico_proyeccion(prom: dict, png: Path) -> None:
    labels = [f"Con WES\n(× {MESES_PROYECCION})", f"Sin WES\n(× {MESES_PROYECCION})"]
    clp = [prom["proj_clp_con"] / 1e6, prom["proj_clp_sin"] / 1e6]
    fig, ax = plt.subplots(figsize=(7.5, 5))
    bars = ax.bar(labels, clp, color=["#3498DB", "#E74C3C"], width=0.45)
    ax.set_ylabel("Millones CLP")
    ax.set_title(f"Agregado Genchi — proyección anual\n(Sin vs Con: {prom['pct_clp']:+.1f}% CLP; {prom['pct_m3']:+.1f}% m³)", fontsize=10)
    for bar, m3, c in zip(bars, [prom["proj_m3_con"], prom["proj_m3_sin"]], [prom["proj_clp_con"], prom["proj_clp_sin"]]):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(clp) * 0.03, f"{_fmt_m3(m3)}\n{_fmt_clp(c)}", ha="center", va="bottom", fontsize=9)
    ax.set_ylim(0, max(clp) * 1.22 if clp else 1)
    fig.tight_layout()
    fig.savefig(png, dpi=180)
    plt.close(fig)


def _grafico_mensual(meses: List[dict], png: Path) -> None:
    plot = _meses_plot(meses)
    labels = [m["mes_label"] for m in plot]
    clp = [m["clp"] / 1e6 for m in plot]
    esp = [m["clp_esperado"] / 1e6 for m in plot]
    colors = ["#3498DB" if m["estado"] == "Con WES" else "#E74C3C" for m in plot]
    x = range(len(labels))
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(x, clp, color=colors, alpha=0.85, label="Facturado")
    ax.plot(x, esp, color="#2C3E50", marker="o", linewidth=1.5, label="Esperado WES")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=45, ha="right")
    ax.set_title("Agregado Genchi — costo mensual (M CLP)", fontsize=11)
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(png, dpi=180)
    plt.close(fig)


def _grafico_unidades(unidades: List[ResultadoUnidad], png: Path) -> None:
    nombres = [u.cfg.nombre.replace("CDP Detención Preventiva ", "").replace("CPF ", "").replace("CCP ", "") for u in unidades]
    sob = [u.stats["sobrecosto_sin"] / 1e6 for u in unidades]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.barh(nombres, sob, color="#E74C3C", alpha=0.85)
    ax.set_xlabel("Sobrecosto sin WES (M CLP)")
    ax.set_title("Sobrecosto acumulado por unidad", fontsize=11)
    fig.tight_layout()
    fig.savefig(png, dpi=180)
    plt.close(fig)


def _word(path: Path, unidades: List[ResultadoUnidad], meses: List[dict], stats: dict, prom: dict, pngs: dict) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.5)
    t = doc.add_heading("INFORME AGREGADO DE FACTURACIÓN", 0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    t.runs[0].font.color.rgb = _HEADING
    sub = doc.add_paragraph("Gendarmería de Chile — Unidad Genchi (tres recintos)")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Ventana: ene-2025 a la fecha  |  Generado: {datetime.now():%d-%m-%Y %H:%M}")

    doc.add_heading("1. Antecedentes", 1)
    doc.add_paragraph("Consolidado CDP Puente Alto, CPF San Miguel y CCP Santiago Sur. Misma metodología que informes individuales.")
    for u in unidades:
        doc.add_paragraph(f"• {u.cfg.nombre}: cuenta {u.cfg.cuenta}, retiro WES {u.cfg.fecha_retiro_wes:%d-%m-%Y}, excl. factura {u.cfg.factura_excluida}.", style="List Bullet")

    doc.add_heading("2. Resumen económico agregado", 1)
    p = doc.add_paragraph()
    p.add_run("Totales tres unidades:\n").bold = True
    p.add_run(f"  • {_fmt_m3(stats['total_m3_ventana'])} | {_fmt_clp(stats['total_clp_ventana'])}\n  • Con WES: {_fmt_clp(stats['total_clp_con'])} | Sin WES: {_fmt_clp(stats['total_clp_sin'])}\n  • Sobrecosto sin WES: {_fmt_clp(stats['sobrecosto_sin'])}\n  • Esperado línea base: {_fmt_clp(stats['total_clp_esperado'])}\n")

    doc.add_heading("2.1 Por unidad", 2)
    rows_u = [("Unidad", "Cuenta", "Con WES", "Sin WES", "Total", "Sobrecosto")]
    for u in unidades:
        rows_u.append((u.cfg.nombre, u.cfg.cuenta, _fmt_clp(u.stats["total_clp_con"]), _fmt_clp(u.stats["total_clp_sin"]), _fmt_clp(u.stats["total_clp_ventana"]), _fmt_clp(u.stats["sobrecosto_sin"])))
    rows_u.append(("TOTAL", "", _fmt_clp(stats["total_clp_con"]), _fmt_clp(stats["total_clp_sin"]), _fmt_clp(stats["total_clp_ventana"]), _fmt_clp(stats["sobrecosto_sin"])))
    tu = doc.add_table(rows=len(rows_u), cols=6)
    tu.style = "Light Grid Accent 1"
    for i, row in enumerate(rows_u):
        for j, val in enumerate(row):
            tu.rows[i].cells[j].text = val
            if i == 0 or i == len(rows_u) - 1:
                tu.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    _tabla_compacta(tu, 8.5)

    doc.add_heading("3. Promedio mensual y proyección anual — agregado", 1)
    rows_p = [("Indicador", "Con WES", "Sin WES", "Variación"), ("Prom. mensual m³", f"{prom['prom_m3_con']:.1f}", f"{prom['prom_m3_sin']:.1f}", f"{prom['pct_m3']:+.1f}%"), ("Prom. mensual CLP", _fmt_clp(int(round(prom["prom_clp_con"]))), _fmt_clp(int(round(prom["prom_clp_sin"]))), f"{prom['pct_clp']:+.1f}%"), (f"Proy. anual CLP (×{MESES_PROYECCION})", _fmt_clp(prom["proj_clp_con"]), _fmt_clp(prom["proj_clp_sin"]), f"{prom['pct_clp']:+.1f}%")]
    tp = doc.add_table(rows=len(rows_p), cols=4)
    tp.style = "Light Grid Accent 1"
    for i, row in enumerate(rows_p):
        for j, val in enumerate(row):
            tp.rows[i].cells[j].text = val
            if i == 0:
                tp.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    _tabla_compacta(tp, 9)

    doc.add_heading("4. Desglose mes a mes — agregado", 1)
    rows_m = [("Mes", "Estado", "m³", "Facturado", "Esperado", "Sobrecosto")]
    tc = te = ts = tm = 0.0
    for m in meses:
        tc += m["clp"]; te += m["clp_esperado"]; ts += m["sobrecosto_clp"]; tm += m["m3"]
        rows_m.append((m["mes_label"], m["estado"], f"{m['m3']:.1f}", _fmt_clp(m["clp"]), _fmt_clp(m["clp_esperado"]), _fmt_clp(m["sobrecosto_clp"])))
    rows_m.append(("TOTAL", "", _fmt_m3(tm), _fmt_clp(tc), _fmt_clp(te), _fmt_clp(ts)))
    tm_t = doc.add_table(rows=len(rows_m), cols=6)
    tm_t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows_m):
        for j, val in enumerate(row):
            tm_t.rows[i].cells[j].text = val
            if i == 0 or i == len(rows_m) - 1:
                tm_t.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    _tabla_compacta(tm_t, 8.5)

    for key, cap, w in (("uni", "Gráfico — sobrecosto por unidad:", 14), ("proj", "Gráfico — proyección anual agregada:", 12.5), ("mes", "Gráfico — costo mensual agregado:", 13.5)):
        if pngs[key].exists():
            doc.add_paragraph(cap)
            pic = doc.add_paragraph()
            pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic.add_run().add_picture(str(pngs[key]), width=Cm(w))

    doc.add_page_break()
    doc.add_heading("5. Anexo por unidad", 1)
    for u in unidades:
        doc.add_heading(u.cfg.nombre, 2)
        doc.add_paragraph(f"Cuenta {u.cfg.cuenta} | Línea base {u.stats['clp_m3_base']:.0f} CLP/m³ | Sobrecosto {_fmt_clp(u.stats['sobrecosto_sin'])}")
        rows = [("Factura", "Periodo", "m³", "CLP", "Estado")]
        for p in u.periodos:
            rows.append((p.boleta, f"{p.lectura_desde:%d-%m-%Y} → {p.lectura_hasta:%d-%m-%Y}", f"{p.m3:,}".replace(",", "."), _fmt_clp(p.clp), p.etiqueta))
        tbl = doc.add_table(rows=len(rows), cols=5)
        tbl.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                tbl.rows[i].cells[j].text = val
                if i == 0:
                    tbl.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
        _tabla_compacta(tbl, 8)
        doc.add_paragraph("")

    doc.add_heading("6. Conclusiones", 1)
    for line in stats["conclusiones"]:
        doc.add_paragraph(line, style="List Bullet")
    doc.save(str(path))


def main() -> int:
    OUT_BASE.mkdir(parents=True, exist_ok=True)
    unidades = []
    for cfg in SITES:
        if not any(cfg.facturas_dir.glob("*.pdf")):
            print(f"[ERROR] Sin PDFs en {cfg.facturas_dir}")
            return 1
        print(f"[OK] {cfg.nombre}")
        unidades.append(_analizar_unidad(cfg))

    meses_agg = _agregar_meses(unidades)
    prom_agg = _promedios(meses_agg)
    stats = _stats_agregado(unidades, prom_agg)
    ts = datetime.now().strftime("%Y%m%d_%H%M")

    with (OUT_BASE / f"facturaciones_genchi_agregado_mensual_{ts}.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(["mes", "estado", "m3", "clp", "clp_esperado", "sobrecosto_clp"])
        for m in meses_agg:
            w.writerow([m["mes_label"], m["estado"], round(m["m3"], 2), m["clp"], m["clp_esperado"], m["sobrecosto_clp"]])

    pngs = {"proj": OUT_BASE / f"grafico_agregado_proyeccion_{ts}.png", "mes": OUT_BASE / f"grafico_agregado_mensual_{ts}.png", "uni": OUT_BASE / f"grafico_agregado_sobrecosto_{ts}.png"}
    _grafico_proyeccion(prom_agg, pngs["proj"])
    _grafico_mensual(meses_agg, pngs["mes"])
    _grafico_unidades(unidades, pngs["uni"])

    docx = OUT_BASE / f"Informe_Genchi_Agregado_facturaciones_{ts}.docx"
    _word(docx, unidades, meses_agg, stats, prom_agg, pngs)
    pdf = docx.with_suffix(".pdf")
    try:
        from generar_reporte_word import convertir_word_a_pdf
        p = convertir_word_a_pdf(docx)
        if p and Path(p).is_file():
            pdf = Path(p)
    except Exception:
        pass

    print("=" * 72)
    print(f"Total ventana: {_fmt_clp(stats['total_clp_ventana'])} | Sobrecosto: {_fmt_clp(stats['sobrecosto_sin'])}")
    print(f"DOCX: {docx}")
    print(f"PDF:  {pdf}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
