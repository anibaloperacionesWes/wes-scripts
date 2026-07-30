"""Utilidades compartidas: anteponer Q3 hallazgos en Word y envío SMTP a Aníbal."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable, List


def prepend_q3_hallazgos(doc_path: Path, texto_q3: str) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import RGBColor

    doc = Document(str(doc_path))
    if not doc.paragraphs:
        doc.add_paragraph("")

    p0 = doc.paragraphs[0]
    p = p0.insert_paragraph_before("")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    lineas = texto_q3.split("\n")
    titulos_sec = {"Síntesis ejecutiva", "Profundización de hallazgos"}
    for i, line in enumerate(lineas):
        run = p.add_run(("\n" if i else "") + line)
        s = line.strip()
        run.bold = bool((i == 0 and s.startswith("3)")) or s in titulos_sec)
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(str(doc_path))


def enviar_anibal_adjuntos(
    archivos: Iterable[Path],
    asunto: str,
    cuerpo: str,
    destinatario: str = "anibal.aoperaciones@wes.cl",
) -> None:
    import smtplib
    from email.mime.application import MIMEApplication
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    lista: List[Path] = [Path(p) for p in archivos if Path(p).exists()]
    if not lista:
        raise ValueError("No hay archivos existentes para adjuntar.")

    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    smtp_user = "agente.ia@wes.cl"
    smtp_pass = (os.environ.get("WES_GMAIL_APP_PASSWORD") or "vxbynfpoehbweelj").replace(" ", "").strip()

    msg = MIMEMultipart()
    msg["From"] = smtp_user
    msg["To"] = destinatario
    msg["Subject"] = asunto
    msg.attach(MIMEText(cuerpo, "plain", "utf-8"))

    for fp in lista:
        with open(fp, "rb") as f:
            part = MIMEApplication(f.read())
            part.add_header("Content-Disposition", "attachment", filename=fp.name)
            msg.attach(part)

    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_user, smtp_pass)
        server.sendmail(smtp_user, [destinatario], msg.as_string())
