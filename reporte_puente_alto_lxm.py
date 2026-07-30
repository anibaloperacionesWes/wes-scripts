"""
Compara LXM del Excel 'conusmos puente alto.xlsx' con el caudal horario de la API WES.

La API entrega valores por hora; se asumen en m³/h (caudal medio de esa hora).
Conversión a L/min (comparable con LXM del Excel):  L/min = m³/h × (1000/60).

En el Excel de salida: **Estado OK/NOK** según el **promedio** L/min en el horario (≤ LXM = OK).

Para que el .xlsx sea **idéntico** a tu plantilla (formatos, logos, anchos), en Windows se usa
**Microsoft Excel por COM** (pywin32). Sin Excel instalado se usa openpyxl y el archivo puede
verse distinto al original.

Solo Excel (sin Word/CSV): ``python reporte_puente_alto_lxm.py --no-word``

Salida por defecto: ``reports/puente_alto_lxm/`` (para automatizar o enviar por correo).
Copia en Escritorio solo al revisar: ``--escritorio``.
"""

from __future__ import annotations

import argparse
import csv
import re
import shutil
import sys
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_URL = "http://104.248.53.141:7003/wes/api/acl-node/v1"
ENTITY_BASE = "http://104.248.53.141:7001/wes/api/acl-entities/v1"

# m³/h -> L/min
M3H_TO_LMIN = 1000.0 / 60.0

DEFAULT_EXCEL = Path.home() / "OneDrive" / "Desktop" / "conusmos puente alto.xlsx"


def _norm(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", s.strip().lower())


def default_excel_path() -> Path:
    p = Path(__file__).resolve().parent / "conusmos puente alto.xlsx"
    if p.exists():
        return p
    return DEFAULT_EXCEL


def carpeta_escritorio_informes() -> Path:
    """
    Escritorio\\informes_lxm_puente_alto (solo si usas --escritorio para revisar a mano).
    """
    for candidate in (
        Path.home() / "OneDrive" / "Desktop",
        Path.home() / "Desktop",
    ):
        if candidate.is_dir():
            sub = candidate / "informes_lxm_puente_alto"
            sub.mkdir(parents=True, exist_ok=True)
            return sub
    sub = Path.home() / "informes_lxm_puente_alto"
    sub.mkdir(parents=True, exist_ok=True)
    return sub


def carpeta_salida_por_defecto() -> Path:
    """Salida habitual: reports/puente_alto_lxm (lista para adjuntar por correo o tareas)."""
    base = Path(__file__).resolve().parent / "reports" / "puente_alto_lxm"
    base.mkdir(parents=True, exist_ok=True)
    return base


def obtener_nodos_puente_alto() -> List[Dict[str, str]]:
    url = f"{ENTITY_BASE}/companies/000010"
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    data = r.json()
    out = []
    for n in data.get("nodes", []):
        nid = n.get("nodeId", "")
        name = (n.get("name") or "").strip()
        if nid and name:
            out.append({"nodeId": nid, "nodeName": name})
    return out


# Mapeo explícito Excel (nombre aproximado) -> nodeId Corporación Puente Alto (000010)
MAPA_EXCEL_A_NODE: Dict[str, str] = {
    _norm("Complejo Educacional Consolidada"): "000010-10",
    _norm("Escuela Villa Independencia"): "000010-02",
    _norm("Escuela Andes del Sur"): "000010-01",
    _norm("Escuela Padre Hurtado"): "000010-03",
    _norm("Colegio Maipo"): "000010-04",
    # Excel de proyección / auditoría dice "LICEO MAIPO"; en API es "Colegio Maipo".
    # Sin esta fila, el token "liceo" empareja antes con Liceo Chiloé (000010-08).
    _norm("Liceo Maipo"): "000010-04",
    _norm("Escuela Luis Matte Larraín Central"): "000010-05",
    _norm("Escuela Gabriela"): "000010-06",
    _norm("Escuela Juan Mackenna O'Reilly"): "000010-07",
    _norm("Liceo Chiloé"): "000010-08",
    _norm("Escuela Los Andes"): "000010-09",
    _norm("Escuela Nonato Coo"): "000010-11",
}


def mapear_establecimiento_a_nodo(nombre_excel: str, nodos: Sequence[Dict[str, str]]) -> Optional[str]:
    """Empareja nombre del Excel con nodeId (Puente Alto)."""
    if not nombre_excel or not str(nombre_excel).strip():
        return None
    ne = _norm(str(nombre_excel))
    if ne in MAPA_EXCEL_A_NODE:
        return MAPA_EXCEL_A_NODE[ne]

    # Variantes sin tilde / typo
    for k, nid in MAPA_EXCEL_A_NODE.items():
        if k in ne or ne in k:
            return nid

    # Por token largo en nombre API
    for n in nodos:
        nn = _norm(n["nodeName"])
        if ne in nn or nn in ne:
            return n["nodeId"]
    for n in nodos:
        nn = _norm(n["nodeName"])
        for parte in re.split(r"[^\w]+", ne):
            if len(parte) >= 5 and parte in nn:
                return n["nodeId"]
    return None


def _date_to_ddmmyyyy(d: datetime) -> str:
    return d.strftime("%d%m%Y")


def obtener_datos_horarios_dia(node_id: str, fecha: datetime) -> Dict[int, float]:
    date_str = _date_to_ddmmyyyy(fecha)
    url = f"{BASE_URL}/nodes/{node_id}/dates.measures.csv"
    params = [("start", date_str), ("end", date_str)]
    response = requests.get(url, params=params, timeout=30)
    response.raise_for_status()
    hourly_data: Dict[int, float] = {}
    lines = response.text.strip().split("\n")
    for line in lines[1:]:
        if not line.strip():
            continue
        parts = line.split(",")
        if len(parts) < 2:
            continue
        time_str = parts[0].strip()
        value_str = parts[1].strip()
        try:
            if "T" in time_str:
                hour = int(time_str.split("T")[1].split(":")[0])
            else:
                dt = datetime.fromisoformat(time_str.replace("Z", "+00:00"))
                hour = dt.hour
            hourly_data[hour] = float(value_str)
        except Exception:
            continue
    return hourly_data


def m3h_a_lmin(val_m3h: float) -> float:
    return val_m3h * M3H_TO_LMIN


def parse_rango_horario(txt: object) -> Optional[Tuple[int, int, bool]]:
    """
    Interpreta 'HH:MM - HH:MM'. Devuelve (hora_inicio, hora_fin_inclusive, cruza_medianoche).
    Si cruza medianoche, hora_fin_inclusive es la última hora del tramo que termina al día siguiente.
    """
    if txt is None:
        return None
    s = str(txt).strip()
    if not s or s.lower() in ("none", "-", ""):
        return None
    m = re.search(
        r"(\d{1,2})\s*:\s*(\d{2})\s*[-–]\s*(\d{1,2})\s*:\s*(\d{2})",
        s,
        re.I,
    )
    if not m:
        return None
    h1, m1, h2, m2 = int(m.group(1)), int(m.group(2)), int(m.group(3)), int(m.group(4))
    # Usamos la hora entera del reloj (bloque API): inicio en h1, fin en h2 si mismo día
    start_h = h1
    end_h = h2
    if start_h == end_h and m1 >= m2:
        return (start_h, start_h, False)
    if start_h < end_h or (start_h == end_h and m1 < m2):
        return (start_h, end_h, False)
    # Cruce medianoche: ej 21:00 - 23:59 o 23:00 - 06:00
    return (start_h, end_h, True)


def horas_incluidas_en_rango(
    start_h: int, end_h: int, cruza: bool
) -> List[int]:
    if not cruza:
        return list(range(start_h, end_h + 1))
    # overnight: start_h .. 23 y 0 .. end_h
    return list(range(start_h, 24)) + list(range(0, end_h + 1))


def max_lmin_en_horas(
    hourly_m3h: Dict[int, float], horas: List[int]
) -> Tuple[float, Optional[int]]:
    """Máximo L/min y hora donde ocurre."""
    best = -1.0
    best_h: Optional[int] = None
    for h in horas:
        v = float(hourly_m3h.get(h, 0.0))
        lmin = m3h_a_lmin(v)
        if lmin > best:
            best = lmin
            best_h = h
    if best < 0:
        return 0.0, None
    return best, best_h


def promedio_lmin_en_horas(
    hourly_m3h: Dict[int, float], horas: List[int]
) -> float:
    """Promedio L/min en las horas del rango (bloques horarios de la API)."""
    if not horas:
        return 0.0
    total = 0.0
    for h in horas:
        v = float(hourly_m3h.get(h, 0.0))
        total += m3h_a_lmin(v)
    return total / len(horas)


@dataclass
class FilaComparacion:
    item: str
    establecimiento: str
    node_id: str
    rango_nombre: str
    horario_txt: str
    lxm_excel: Optional[float]
    prom_lmin_app: float
    max_lmin_app: float
    hora_pico: Optional[int]
    cumple: bool
    nota: str


@dataclass
class LayoutPuenteAlto:
    """Columnas 1-based detectadas en tu Excel (no asumimos siempre B u H,K,N)."""

    sheet_index_0: int
    header_row_1based: int
    col_item: int
    col_est: int
    horario_cols: Tuple[int, int, int]
    col_max_pct: int = 4
    col_max_lxm: int = 5
    col_onoff: int = 15

    def horarios_ordenados(self) -> List[int]:
        return sorted(self.horario_cols)


def _header_match_horario(val: object) -> bool:
    if val is None:
        return False
    s = _norm(str(val))
    return "horario" in s


def _header_match_item(val: object) -> bool:
    if val is None:
        return False
    return str(val).strip().upper() == "ITEM"


def _header_match_onoff(val: object) -> bool:
    if val is None:
        return False
    s = str(val).strip().lower()
    return ("on" in s and "off" in s) or "on/off" in s.replace(" ", "")


def shift_col_despues_inserts(
    col_orig_1based: int, horario_cols: Sequence[int]
) -> int:
    """Tras insertar 2 columnas antes de (h+1) por cada h en horario (de derecha a izquierda)."""
    h = col_orig_1based
    for hi in sorted(horario_cols, reverse=True):
        ins = hi + 1
        if h >= ins:
            h += 2
    return h


def detectar_layout_puente_alto(path: Path) -> Tuple[LayoutPuenteAlto, List[str]]:
    """
    Detecta hoja, fila ITEM y las 3 columnas «Horario» leyendo los encabezados reales.
    """
    from openpyxl import load_workbook

    avisos: List[str] = []
    wb = load_workbook(path, read_only=True, data_only=True)

    for si, sname in enumerate(wb.sheetnames):
        ws = wb[sname]
        max_r = min(ws.max_row or 1, 250)
        max_c_scan = 30
        for r in range(1, max_r + 1):
            header_row_1based = r
            col_item: Optional[int] = None
            for c in range(1, max_c_scan + 1):
                v = ws.cell(r, c).value
                if _header_match_item(v):
                    col_item = c
                    break
            if col_item is None:
                continue

            col_est = col_item + 1
            horarios: List[int] = []
            for c in range(1, max_c_scan + 1):
                v = ws.cell(r, c).value
                if _header_match_horario(v):
                    horarios.append(c)
            horarios = sorted(set(horarios))

            if len(horarios) == 1:
                h0 = horarios[0]
                avisos.append(
                    f"Hoja '{sname}' fila {r}: una sola columna con texto 'Horario'; "
                    f"asumiendo 3 rangos en columnas {h0}, {h0 + 3}, {h0 + 6}."
                )
                horarios = [h0, h0 + 3, h0 + 6]
            elif len(horarios) < 3:
                avisos.append(
                    f"Hoja '{sname}' fila {r}: se encontraron {len(horarios)} columnas "
                    f"'Horario' (se esperan 3). Usando posiciones tipicas 8,11,14."
                )
                horarios = [8, 11, 14]
            elif len(horarios) > 3:
                avisos.append(
                    f"Hoja '{sname}' fila {r}: varias columnas 'Horario'; "
                    f"se usan las 3 primeras de izquierda a derecha."
                )
                horarios = horarios[:3]

            col_onoff = max(horarios) + 1
            for c in range(max(horarios) + 1, max_c_scan + 1):
                v = ws.cell(r, c).value
                if _header_match_onoff(v):
                    col_onoff = c
                    break

            col_max_pct, col_max_lxm = 4, 5
            if col_item <= 3:
                for c in range(col_item + 2, min(horarios[0], max_c_scan + 1)):
                    v = ws.cell(r, c).value
                    if v is not None and "%" in str(v):
                        col_max_pct = c
                        break
                for c in range(col_max_pct + 1, horarios[0]):
                    v = ws.cell(r, c).value
                    if v is not None and "lxm" in str(v).lower():
                        col_max_lxm = c
                        break

            wb.close()
            return (
                LayoutPuenteAlto(
                    sheet_index_0=si,
                    header_row_1based=header_row_1based,
                    col_item=col_item,
                    col_est=col_est,
                    horario_cols=(horarios[0], horarios[1], horarios[2]),
                    col_max_pct=col_max_pct,
                    col_max_lxm=col_max_lxm,
                    col_onoff=col_onoff,
                ),
                avisos,
            )

    try:
        wb.close()
    except Exception:
        pass
    raise ValueError(
        "No se encontro una fila con 'ITEM' y columnas 'Horario' en ninguna hoja."
    )


def _celda_fila(row: Optional[Sequence[object]], col_1based: int) -> object:
    if not row or col_1based < 1:
        return None
    i = col_1based - 1
    if i >= len(row):
        return None
    return row[i]


def leer_excel_y_comparar(
    path: Path,
    fecha: datetime,
    nodos: List[Dict[str, str]],
    layout: LayoutPuenteAlto,
) -> Tuple[List[FilaComparacion], List[str]]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[layout.sheet_index_0]]
    rows = list(ws.iter_rows(values_only=True))

    header_row = layout.header_row_1based - 1
    if header_row < 0 or header_row >= len(rows):
        wb.close()
        raise ValueError("Fila de encabezado invalida segun layout detectado")

    resultados: List[FilaComparacion] = []
    warnings: List[str] = []

    nombres_banda = ["Rango Alto", "Rango Medio", "Rango Bajo (nocturno)"]

    for row in rows[header_row + 1 :]:
        if not row:
            continue
        item = _celda_fila(row, layout.col_item)
        est = _celda_fila(row, layout.col_est)
        if item is None or est is None:
            continue
        try:
            item_s = str(int(item)) if isinstance(item, float) and item == int(item) else str(item)
        except Exception:
            item_s = str(item)
        est_s = str(est).strip()
        if not est_s:
            continue

        node_id = mapear_establecimiento_a_nodo(est_s, nodos)
        if not node_id:
            warnings.append(f"Sin mapeo API: {est_s}")
            continue

        try:
            hourly = obtener_datos_horarios_dia(node_id, fecha)
        except Exception as e:
            warnings.append(f"API error {node_id} ({est_s}): {e}")
            continue

        bandas = []
        for bi, h in enumerate(sorted(layout.horario_cols)):
            bandas.append(
                (
                    nombres_banda[bi],
                    _celda_fila(row, h - 2),
                    _celda_fila(row, h - 1),
                    _celda_fila(row, h),
                )
            )

        for nombre, pct, lxm_cell, horario_cell in bandas:
            lxm = None
            if lxm_cell is not None and str(lxm_cell).strip() != "":
                try:
                    lxm = float(lxm_cell)
                except (TypeError, ValueError):
                    lxm = None

            pr = parse_rango_horario(horario_cell)
            if pr is None:
                if lxm is not None:
                    resultados.append(
                        FilaComparacion(
                            item=item_s,
                            establecimiento=est_s,
                            node_id=node_id,
                            rango_nombre=nombre,
                            horario_txt=str(horario_cell or "—"),
                            lxm_excel=lxm,
                            prom_lmin_app=0.0,
                            max_lmin_app=0.0,
                            hora_pico=None,
                            cumple=True,
                            nota="Sin horario válido",
                        )
                    )
                continue

            start_h, end_h, cruza = pr
            horas = horas_incluidas_en_rango(start_h, end_h, cruza)
            max_lmin, h_pico = max_lmin_en_horas(hourly, horas)
            prom_lmin = promedio_lmin_en_horas(hourly, horas)

            if lxm is None:
                resultados.append(
                    FilaComparacion(
                        item=item_s,
                        establecimiento=est_s,
                        node_id=node_id,
                        rango_nombre=nombre,
                        horario_txt=str(horario_cell).strip(),
                        lxm_excel=None,
                        prom_lmin_app=round(prom_lmin, 3),
                        max_lmin_app=round(max_lmin, 3),
                        hora_pico=h_pico,
                        cumple=True,
                        nota="Sin LXM en Excel",
                    )
                )
                continue

            cumple = prom_lmin <= lxm + 1e-6
            resultados.append(
                FilaComparacion(
                    item=item_s,
                    establecimiento=est_s,
                    node_id=node_id,
                    rango_nombre=nombre,
                    horario_txt=str(horario_cell).strip(),
                    lxm_excel=round(lxm, 4),
                    prom_lmin_app=round(prom_lmin, 3),
                    max_lmin_app=round(max_lmin, 3),
                    hora_pico=h_pico,
                    cumple=cumple,
                    nota="OK" if cumple else "Supera LXM (promedio)",
                )
            )

    wb.close()
    return resultados, warnings


def escribir_word(
    filas: List[FilaComparacion],
    fecha: datetime,
    out_path: Path,
) -> None:
    doc = Document()
    t = doc.add_heading("COMPARATIVO LXM vs CAUDAL APP (Puente Alto)", 0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    t.runs[0].font.color.rgb = RGBColor(204, 0, 0)
    t.runs[0].bold = True

    doc.add_paragraph(
        f"Fecha medición: {fecha:%Y-%m-%d} | Unidad app: m³/h → L/min (×{M3H_TO_LMIN:.4f})"
    )
    doc.add_paragraph("")

    headers = [
        "Item",
        "Establecimiento",
        "ID",
        "Rango",
        "Horario Excel",
        "LXM (L/min)",
        "Prom L/min app",
        "Max L/min app",
        "Hora pico",
        "Cumple",
    ]
    table = doc.add_table(rows=1 + len(filas), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        try:
            shading_xml = (
                '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'w:val="clear" w:fill="4472C4"/>'
            )
            shading = parse_xml(shading_xml)
            tc_pr = cell._element.get_or_add_tcPr()
            if tc_pr.find(qn("w:shd")) is None:
                tc_pr.append(shading)
        except Exception:
            pass

    for r_idx, f in enumerate(filas, start=1):
        vals = [
            f.item,
            f.establecimiento,
            f.node_id,
            f.rango_nombre,
            f.horario_txt,
            f"{f.lxm_excel:.4f}" if f.lxm_excel is not None else "—",
            f"{f.prom_lmin_app:.3f}",
            f"{f.max_lmin_app:.3f}",
            f"{f.hora_pico:02d}:00" if f.hora_pico is not None else "—",
            "Sí" if f.cumple else "No",
        ]
        for c_idx, val in enumerate(vals):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            # Rojo si no cumple: Prom / Max / Cumple
            if not f.cumple and c_idx in (6, 7, 9):
                run.font.color.rgb = RGBColor(200, 0, 0)
                run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if c_idx != 1 else WD_PARAGRAPH_ALIGNMENT.LEFT

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def escribir_csv(filas: List[FilaComparacion], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as fp:
        w = csv.writer(fp)
        w.writerow(
            [
                "item",
                "establecimiento",
                "node_id",
                "rango",
                "horario_excel",
                "lxm_l_min",
                "prom_l_min_app",
                "max_l_min_app",
                "hora_pico",
                "cumple",
                "nota",
            ]
        )
        for f in filas:
            w.writerow(
                [
                    f.item,
                    f.establecimiento,
                    f.node_id,
                    f.rango_nombre,
                    f.horario_txt,
                    f.lxm_excel,
                    f.prom_lmin_app,
                    f.max_lmin_app,
                    f.hora_pico,
                    "SI" if f.cumple else "NO",
                    f.nota,
                ]
            )


def _excel_rgb(r: int, g: int, b: int) -> int:
    """Color interior Excel (BGR entero)."""
    return r + (g << 8) + (b << 16)


def _excel_com_disponible() -> bool:
    """Windows + pywin32: abre el archivo con Excel real (copia idéntica al original)."""
    if sys.platform != "win32":
        return False
    try:
        import win32com.client  # noqa: F401

        return True
    except ImportError:
        return False


def _excel_unmerge_fila(ws: object, row_1based: int, max_col: int = 15) -> None:
    """Quita fusiones en una fila (COM)."""
    seen: set[str] = set()
    for c in range(1, max_col + 1):
        cell = ws.Cells(row_1based, c)
        try:
            if cell.MergeCells:
                addr = str(cell.MergeArea.Address)
                if addr not in seen:
                    seen.add(addr)
                    cell.MergeArea.UnMerge()
        except Exception:
            pass


def escribir_excel_misma_tabla_excel_com(
    path_excel_origen: Path,
    fecha: datetime,
    out_path: Path,
    nodos: List[Dict[str, str]],
    layout: LayoutPuenteAlto,
) -> List[str]:
    """
    Usa Microsoft Excel instalado (COM): copia tu archivo e inserta columnas según el
    layout detectado (mismas columnas Horario que en tu Excel).
    """
    import win32com.client

    warnings: List[str] = []
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path_excel_origen, out_path)
    path_abs = str(out_path.resolve())

    horarios = layout.horarios_ordenados()

    excel = win32com.client.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    excel.ScreenUpdating = False
    wb = None
    try:
        wb = excel.Workbooks.Open(path_abs)
        ws = wb.Worksheets(layout.sheet_index_0 + 1)

        header_row_1based = layout.header_row_1based
        merge_row = header_row_1based - 1

        titles_band: List[object] = []
        if merge_row >= 1:
            t_max = ws.Cells(merge_row, layout.col_max_pct).Value
            for h in horarios:
                titles_band.append(ws.Cells(merge_row, h - 2).Value)
            t_onoff = ws.Cells(merge_row, layout.col_onoff).Value
            _excel_unmerge_fila(ws, merge_row, max_col=35)
        else:
            t_max = None
            titles_band = [None, None, None]
            t_onoff = None

        for h in sorted(horarios, reverse=True):
            for _ in range(2):
                ws.Columns(h + 1).Insert()

        if merge_row >= 1:
            ws.Range(
                ws.Cells(merge_row, layout.col_max_pct),
                ws.Cells(merge_row, layout.col_max_lxm),
            ).Merge()
            ws.Cells(merge_row, layout.col_max_pct).Value = t_max
            for i, h in enumerate(horarios):
                sa = shift_col_despues_inserts(h - 2, horarios)
                sb = shift_col_despues_inserts(h, horarios) + 2
                ws.Range(ws.Cells(merge_row, sa), ws.Cells(merge_row, sb)).Merge()
                ws.Cells(merge_row, sa).Value = titles_band[i]
            co = shift_col_despues_inserts(layout.col_onoff, horarios)
            ws.Cells(merge_row, co).Value = t_onoff

        subh = header_row_1based
        for h in horarios:
            pc = shift_col_despues_inserts(h, horarios) + 1
            ec = shift_col_despues_inserts(h, horarios) + 2
            ws.Cells(subh, pc).Value = "Prom L/min app"
            ws.Cells(subh, ec).Value = "Estado"

        xl_center = -4108
        bases_pct = [shift_col_despues_inserts(h - 2, horarios) for h in horarios]

        try:
            max_row = int(
                ws.Cells(ws.Rows.Count, layout.col_item).End(-4162).Row
            )
        except Exception:
            max_row = subh + 200

        for r in range(subh + 1, max_row + 1):
            item = ws.Cells(r, layout.col_item).Value
            est = ws.Cells(r, layout.col_est).Value
            if item is None and est is None:
                continue
            est_s = str(est).strip() if est is not None else ""
            if not est_s:
                continue

            node_id = mapear_establecimiento_a_nodo(est_s, nodos)
            hourly: Optional[Dict[int, float]] = None
            if not node_id:
                warnings.append(f"Sin mapeo API: {est_s}")
            else:
                try:
                    hourly = obtener_datos_horarios_dia(node_id, fecha)
                except Exception as e:
                    warnings.append(f"API error {node_id} ({est_s}): {e}")
                    hourly = None

            for base in bases_pct:
                lxm_cell = ws.Cells(r, base + 1).Value
                horario_cell = ws.Cells(r, base + 2).Value

                lxm: Optional[float] = None
                if lxm_cell is not None and str(lxm_cell).strip() != "":
                    try:
                        lxm = float(lxm_cell)
                    except (TypeError, ValueError):
                        lxm = None

                prom = 0.0
                estado = "—"
                if hourly is None:
                    estado = "—"
                else:
                    pr = parse_rango_horario(horario_cell)
                    if pr is None:
                        prom = 0.0
                        estado = "—"
                    else:
                        start_h, end_h, cruza = pr
                        horas = horas_incluidas_en_rango(start_h, end_h, cruza)
                        prom = round(promedio_lmin_en_horas(hourly, horas), 3)
                        if lxm is None:
                            estado = "—"
                        else:
                            estado = "OK" if prom <= lxm + 1e-6 else "NOK"

                c_prom = ws.Cells(r, base + 3)
                c_est = ws.Cells(r, base + 4)
                c_prom.Value = prom
                c_est.Value = estado
                c_prom.HorizontalAlignment = xl_center
                c_est.HorizontalAlignment = xl_center
                if estado == "OK":
                    c_est.Interior.Color = _excel_rgb(198, 239, 206)
                elif estado == "NOK":
                    c_est.Interior.Color = _excel_rgb(255, 199, 206)

        try:
            last = int(
                ws.Cells(ws.Rows.Count, layout.col_item).End(-4162).Row
            )
        except Exception:
            last = max_row
        note_row = last + 1
        ws.Cells(note_row, layout.col_item).Value = (
            f"Fecha medicion API: {fecha:%Y-%m-%d}"
        )

        wb.Save()
    finally:
        if wb is not None:
            wb.Close(SaveChanges=False)
        excel.Quit()

    return warnings


def _unmerge_rows(ws: object, rows: List[int]) -> None:
    """Quita fusiones que cruzan alguna de las filas indicadas (1-based)."""
    to_remove: List[object] = []
    for rng in list(ws.merged_cells.ranges):
        for r in rows:
            if rng.min_row <= r <= rng.max_row:
                to_remove.append(rng)
                break
    for rng in to_remove:
        try:
            ws.unmerge_cells(str(rng))
        except Exception:
            pass


def escribir_excel_misma_tabla_openpyxl(
    path_excel_origen: Path,
    fecha: datetime,
    out_path: Path,
    nodos: List[Dict[str, str]],
    layout: LayoutPuenteAlto,
) -> List[str]:
    """
    Fallback sin Excel instalado: openpyxl **reescribe** el .xlsx y puede verse distinto
    (imágenes/anchos). Preferir Excel COM (Windows + Office).
    """
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, PatternFill

    horarios = layout.horarios_ordenados()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path_excel_origen, out_path)

    wb = load_workbook(out_path)
    ws = wb[wb.sheetnames[layout.sheet_index_0]]

    header_row_1based = layout.header_row_1based
    merge_row = header_row_1based - 1
    warnings: List[str] = []

    titles_band: List[object] = []
    if merge_row >= 1:
        t_max = ws.cell(row=merge_row, column=layout.col_max_pct).value
        for h in horarios:
            titles_band.append(ws.cell(row=merge_row, column=h - 2).value)
        t_onoff = ws.cell(row=merge_row, column=layout.col_onoff).value
    else:
        t_max = t_onoff = None
        titles_band = [None, None, None]

    for h in sorted(horarios, reverse=True):
        ws.insert_cols(h + 1, amount=2)

    if merge_row >= 1:
        _unmerge_rows(ws, [merge_row])
        ws.merge_cells(
            start_row=merge_row,
            start_column=layout.col_max_pct,
            end_row=merge_row,
            end_column=layout.col_max_lxm,
        )
        ws.cell(row=merge_row, column=layout.col_max_pct, value=t_max)
        for i, h in enumerate(horarios):
            sa = shift_col_despues_inserts(h - 2, horarios)
            sb = shift_col_despues_inserts(h, horarios) + 2
            ws.merge_cells(
                start_row=merge_row,
                start_column=sa,
                end_row=merge_row,
                end_column=sb,
            )
            ws.cell(row=merge_row, column=sa, value=titles_band[i])
        co = shift_col_despues_inserts(layout.col_onoff, horarios)
        ws.cell(row=merge_row, column=co, value=t_onoff)

    subh = header_row_1based
    for h in horarios:
        pc = shift_col_despues_inserts(h, horarios) + 1
        ec = shift_col_despues_inserts(h, horarios) + 2
        ws.cell(row=subh, column=pc, value="Prom L/min app")
        ws.cell(row=subh, column=ec, value="Estado")

    fill_ok = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    fill_nok = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    center = Alignment(horizontal="center", vertical="center")

    bases_pct = [shift_col_despues_inserts(h - 2, horarios) for h in horarios]

    for r in range(subh + 1, ws.max_row + 1):
        item = ws.cell(row=r, column=layout.col_item).value
        est = ws.cell(row=r, column=layout.col_est).value
        if item is None and est is None:
            continue
        est_s = str(est).strip() if est is not None else ""
        if not est_s:
            continue

        node_id = mapear_establecimiento_a_nodo(est_s, nodos)
        hourly: Optional[Dict[int, float]] = None
        if not node_id:
            warnings.append(f"Sin mapeo API: {est_s}")
        else:
            try:
                hourly = obtener_datos_horarios_dia(node_id, fecha)
            except Exception as e:
                warnings.append(f"API error {node_id} ({est_s}): {e}")
                hourly = None

        for base in bases_pct:
            lxm_cell = ws.cell(row=r, column=base + 1).value
            horario_cell = ws.cell(row=r, column=base + 2).value

            lxm: Optional[float] = None
            if lxm_cell is not None and str(lxm_cell).strip() != "":
                try:
                    lxm = float(lxm_cell)
                except (TypeError, ValueError):
                    lxm = None

            prom = 0.0
            estado = "—"
            if hourly is None:
                estado = "—"
            else:
                pr = parse_rango_horario(horario_cell)
                if pr is None:
                    prom = 0.0
                    estado = "—"
                else:
                    start_h, end_h, cruza = pr
                    horas = horas_incluidas_en_rango(start_h, end_h, cruza)
                    prom = round(promedio_lmin_en_horas(hourly, horas), 3)
                    if lxm is None:
                        estado = "—"
                    else:
                        estado = "OK" if prom <= lxm + 1e-6 else "NOK"

            c_prom = ws.cell(row=r, column=base + 3, value=prom)
            c_est = ws.cell(row=r, column=base + 4, value=estado)
            c_prom.alignment = center
            c_est.alignment = center
            if estado == "OK":
                c_est.fill = fill_ok
            elif estado == "NOK":
                c_est.fill = fill_nok

    note_row = ws.max_row + 1
    ws.cell(
        row=note_row,
        column=layout.col_item,
        value=f"Fecha medicion API: {fecha:%Y-%m-%d}",
    )

    wb.save(str(out_path))
    return warnings


def escribir_excel_misma_tabla(
    path_excel_origen: Path,
    fecha: datetime,
    out_path: Path,
    nodos: List[Dict[str, str]],
    layout: LayoutPuenteAlto,
) -> Tuple[List[str], str]:
    """
    Genera el Excel ampliado. En Windows intenta **Excel real (COM)** = copia idéntica
    a tu archivo; si falla, usa openpyxl.

    Returns:
        (warnings, metodo_usado) con metodo_usado en ("excel_com", "openpyxl").
    """
    if _excel_com_disponible():
        try:
            w = escribir_excel_misma_tabla_excel_com(
                path_excel_origen, fecha, out_path, nodos, layout
            )
            return (w, "excel_com")
        except Exception as e:
            print(
                f"[AVISO] Excel (Office) no pudo procesar el archivo ({e}). "
                "Usando openpyxl (el resultado puede no ser identico al original)."
            )
            w = escribir_excel_misma_tabla_openpyxl(
                path_excel_origen, fecha, out_path, nodos, layout
            )
            return (w, "openpyxl")
    w = escribir_excel_misma_tabla_openpyxl(
        path_excel_origen, fecha, out_path, nodos, layout
    )
    return (w, "openpyxl")


def main() -> int:
    parser = argparse.ArgumentParser(description="Comparativo LXM Puente Alto vs API")
    parser.add_argument(
        "--fecha",
        type=str,
        default="",
        help="Fecha medicion YYYY-MM-DD (default: ayer)",
    )
    parser.add_argument(
        "--excel",
        type=str,
        default="",
        help="Ruta al Excel (default: escritorio conusmos puente alto.xlsx)",
    )
    parser.add_argument(
        "--no-word",
        action="store_true",
        help="No generar Word ni CSV; solo Excel (y datos de comparacion en memoria).",
    )
    parser.add_argument(
        "--escritorio",
        action="store_true",
        help="Guardar en Escritorio/informes_lxm_puente_alto (revisar a mano). "
        "Sin este flag: reports/puente_alto_lxm (recomendado para adjuntar por correo / tareas).",
    )
    args = parser.parse_args()

    if args.fecha.strip():
        fecha = datetime.strptime(args.fecha.strip(), "%Y-%m-%d")
    else:
        fecha = datetime.combine(
            (datetime.now() - timedelta(days=1)).date(), datetime.min.time()
        )

    excel_path = Path(args.excel) if args.excel.strip() else default_excel_path()
    if not excel_path.is_file():
        print(f"[ERROR] No existe Excel: {excel_path}")
        return 1

    try:
        layout, avisos_layout = detectar_layout_puente_alto(excel_path)
    except ValueError as e:
        print(f"[ERROR] No se pudo leer el formato del Excel: {e}")
        return 1

    print(
        f"[INFO] Layout detectado: hoja #{layout.sheet_index_0 + 1}, "
        f"fila ITEM={layout.header_row_1based}, "
        f"columnas Horario={layout.horario_cols}, ON/OFF col={layout.col_onoff}"
    )

    nodos = obtener_nodos_puente_alto()
    warnings: List[str] = list(avisos_layout)
    filas, w_leer = leer_excel_y_comparar(excel_path, fecha, nodos, layout)
    warnings.extend(w_leer)

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    base = (
        carpeta_escritorio_informes()
        if args.escritorio
        else carpeta_salida_por_defecto()
    )
    docx_path = base / f"comparativo_lxm_puente_alto_{fecha:%Y%m%d}_{ts}.docx"
    xlsx_out = base / f"comparativo_lxm_puente_alto_{fecha:%Y%m%d}_{ts}.xlsx"

    if not args.no_word:
        escribir_word(filas, fecha, docx_path)
        csv_path = docx_path.with_suffix(".csv")
        escribir_csv(filas, csv_path)
    else:
        docx_path = Path("")
        csv_path = Path("")

    try:
        w_excel, metodo_xlsx = escribir_excel_misma_tabla(
            excel_path, fecha, xlsx_out, nodos, layout
        )
        warnings.extend(w_excel)
        if metodo_xlsx == "excel_com":
            print(
                "[INFO] Excel generado con Microsoft Excel (misma apariencia que tu archivo)."
            )
        else:
            print(
                "[INFO] Excel generado con openpyxl (si no se ve igual, instala pywin32 y abre el script en Windows con Excel instalado)."
            )
    except Exception as e:
        print(f"[ERROR] No se pudo generar Excel con tu tabla: {e}")
        import traceback

        traceback.print_exc()

    print("=" * 60)
    print(f"Fecha: {fecha:%Y-%m-%d}")
    print(f"Excel origen: {excel_path}")
    print(f"Filas comparacion: {len(filas)}")
    print(f"Carpeta de salida: {base}")
    if not args.no_word:
        print(f"DOCX: {docx_path}")
        print(f"CSV:  {csv_path}")
    print(f"XLSX (tu tabla + Prom/Estado): {xlsx_out}")
    if warnings:
        print("\nAvisos:")
        for w in warnings[:30]:
            print(f"  - {w}")
        if len(warnings) > 30:
            print(f"  ... y {len(warnings)-30} mas")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
