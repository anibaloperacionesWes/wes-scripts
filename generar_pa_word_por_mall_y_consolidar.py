"""
Genera reportes Word agregados de Parque Arauco (000025) por mall y un consolidado.

Motivación: más rápido/estable que PPT para periodos largos.

Exclusiones requeridas:
- 000025-03 (no operativo desde 14-10-2025)
- 000025-05
- 000025-06
- Por mall: 000025-14 (Quilicura) y 000025-11 (El Bosque) dados de baja (pa_nodos_inactivos_por_mall)

Salida:
- reports/Parque_Arauco/<Mall>/ABREGADO/AGREGADO_<ts>/Reporte_Agregado_Parque_Arauco_<desde>_<hasta>.docx
- reports/Parque_Arauco/CONSOLIDADO/Reporte_Agregado_Parque_Arauco_CONSOLIDADO_<desde>_<hasta>.docx
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple

import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt

from generar_reporte_word import (
    format_number_chilean,
    get_mall_name_for_parque_arauco,
    generate_aggregated_report,
)
from generar_reportes_y_ppt_mall_maipu import obtener_datos_agregados
from pa_nodos_inactivos_por_mall import filtrar_nodos_activos_mall


ENTITY_BASE = "http://104.248.53.141:7001/wes/api/acl-entities/v1"

NODOS_EXCLUIDOS: Dict[str, str] = {
    "000025-03": "No operativo desde 14-10-2025",
    "000025-05": "Excluido",
    "000025-06": "Excluido",
}


def _cargar_empresa_pa() -> dict:
    r = requests.get(f"{ENTITY_BASE}/companies/000025", timeout=60)
    r.raise_for_status()
    return r.json()


def _slug(s: str) -> str:
    s2 = "".join(c for c in s if c.isalnum() or c in (" ", "-", "_")).strip()
    return s2.replace(" ", "_") or "SIN_NOMBRE"


def _resumen_mall(datos: dict) -> Dict[str, Any]:
    agg = datos.get("aggregate_summary") or {}
    total = float(agg.get("total") or 0)
    prom = float(agg.get("promedio_diario") or 0)
    dias = int(agg.get("dias") or 0)

    ranked: List[Tuple[float, str, str]] = []
    for n in (datos.get("nodes_summary") or []):
        t = float((n.get("summary") or {}).get("total") or 0)
        ranked.append((t, str(n.get("node_id") or ""), str(n.get("node_name") or "")))
    ranked.sort(key=lambda x: -x[0])
    top = ranked[0] if ranked else (0.0, "—", "—")

    return {
        "total": total,
        "promedio_diario": prom,
        "dias": dias,
        "top_total": float(top[0]),
        "top_id": top[1],
        "top_name": top[2],
        "num_puntos": len(ranked),
    }


def _crear_consolidado(
    out_path: Path,
    desde: str,
    hasta: str,
    resumenes: Dict[str, Dict[str, Any]],
    rutas_word: Dict[str, Path],
) -> Path:
    doc = Document()

    title = doc.add_paragraph("Parque Arauco — Consolidado de métricas (por mall)")
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph(f"Período: {desde} a {hasta}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    # Tabla resumen
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Mall"
    hdr[1].text = "Total (m³)"
    hdr[2].text = "Prom. diario (m³/día)"
    hdr[3].text = "Días"
    hdr[4].text = "Top punto (ID)"
    hdr[5].text = "Ruta Word"

    for mall in sorted(resumenes.keys()):
        r = resumenes[mall]
        row = table.add_row().cells
        row[0].text = mall
        row[1].text = format_number_chilean(r["total"], 1)
        row[2].text = format_number_chilean(r["promedio_diario"], 1)
        row[3].text = str(r["dias"])
        row[4].text = f"{r['top_name']} ({r['top_id']})"
        row[5].text = str(rutas_word.get(mall, "—"))

    doc.add_paragraph("")
    nota = doc.add_paragraph(
        "Notas: este consolidado se basa en reportes agregados por mall. "
        "No se consideran 000025-03, 000025-05, 000025-06; en Quilicura/El Bosque tampoco "
        "000025-14 / 000025-11 si están dados de baja (listado operativo PA)."
    )
    nota.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in nota.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description="PA Word por mall + consolidado")
    parser.add_argument("--desde", default="01/03/2026", help="DD/MM/YYYY")
    parser.add_argument("--hasta", default="27/03/2026", help="DD/MM/YYYY")
    args = parser.parse_args()

    empresa = _cargar_empresa_pa()
    nodes = empresa.get("nodes") or []

    mall_to_nodes: Dict[str, List[Tuple[str, str]]] = {}
    for n in nodes:
        nid = str(n.get("nodeId") or "").strip()
        name = str(n.get("name") or "").strip()
        if not nid:
            continue
        mall = get_mall_name_for_parque_arauco(nid, name) or "Parque Arauco"
        mall_to_nodes.setdefault(mall, []).append((nid, name))

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    rutas_word: Dict[str, Path] = {}
    resumenes: Dict[str, Dict[str, Any]] = {}

    for mall, pairs in sorted(mall_to_nodes.items()):
        node_ids = [nid for nid, _ in pairs if nid not in NODOS_EXCLUIDOS]
        node_ids = filtrar_nodos_activos_mall(mall, node_ids)
        if not node_ids:
            print(f"[INFO] {mall}: sin nodos tras exclusiones, se omite.")
            continue

        # Estructura por mall para evitar auto-detección con nodos mezclados
        mall_safe = _slug(mall)
        print(f"[INFO] Agregado Word mall {mall}: {len(node_ids)} punto(s) ...")

        out = generate_aggregated_report(
            company_id="000025",
            node_ids=node_ids,
            start_date=args.desde,
            end_date=args.hasta,
            output_dir="reports",
            mall_name=mall_safe,
            apply_exclusions=False,  # ya filtramos aquí
            generate_ppt=False,
        )
        rutas_word[mall] = Path(out)
        print(f"[OK] Word {mall}: {out}")

        datos = obtener_datos_agregados(node_ids, args.desde, args.hasta)
        resumenes[mall] = _resumen_mall(datos)

    cons_dir = Path("reports") / "Parque_Arauco" / "CONSOLIDADO"
    out_cons = cons_dir / f"Reporte_Agregado_Parque_Arauco_CONSOLIDADO_{args.desde.replace('/','')}-{args.hasta.replace('/','')}.docx"
    _crear_consolidado(out_cons, args.desde, args.hasta, resumenes, rutas_word)
    print(f"[OK] CONSOLIDADO: {out_cons}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

