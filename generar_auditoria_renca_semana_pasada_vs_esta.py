"""
Informe de auditoría Renca — 5 puntos en una tabla.

Lo Velásquez, gimnasio, piscina e ICCO: esta semana (desde lun 24) vs semana del 17.
ICCP (al final, no se mezcla en el % de los 4): lun 10–dom 16 con WES vs lun 17–dom 23 sin WES.

Uso:
  python generar_auditoria_renca_semana_pasada_vs_esta.py
  python generar_auditoria_renca_semana_pasada_vs_esta.py --hasta 2026-08-30
"""

from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple
from zoneinfo import ZoneInfo

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font as XlFont, PatternFill, Side
from openpyxl.utils import get_column_letter

import auditoria_cpa_icco_renca_grafico as graf
import generar_graficos_comparativos_desde_excel_consolidado as gxlsx
import generar_informe_auditoria_icco_renca_word as icco
from auditoria_cpa_icco_renca_grafico import Periodo, _vector_m3h_24_desde_api
from generar_excel_auditoria_consolidado_dos_periodos import generar_excel_consolidado
from generar_graficos_comparativos_desde_excel_consolidado import (
    _limpiar_pngs_carpeta_graficos,
    generar_pngs,
    leer_matriz_consolidado,
)
from generar_reporte_word import add_logo_to_header, format_number_chilean

ROOT = Path(__file__).resolve().parent
TZ_CL = ZoneInfo("America/Santiago")
TARIFA_CLP_M3 = 1300.0
COLOR_WES = "#2a6fad"
COLOR_SIN = "#C0504D"
COLOR_HEAD = RGBColor(31, 71, 136)

PUNTOS: Tuple[Tuple[str, str], ...] = (
    ("000017-07", "ICCP (Cumbre de Cóndores pte.)"),
    ("000017-04", "Esc. Lo Velásquez"),
    ("000017-05", "Gimnasio municipal"),
    ("000017-06", "Piscina municipal"),
    ("000017-08", "Colegio ICCO Renca"),
)
NODO_ICCP = "000017-07"
NODO_ICCO = "000017-08"
NODO_ESCUELA = "000017-04"
NODO_GIMNASIO = "000017-05"
PUNTOS_CONTROL: Tuple[Tuple[str, str], ...] = (
    ("000017-08", "Colegio ICCO Renca"),
    ("000017-04", "Esc. Lo Velásquez"),
    ("000017-05", "Gimnasio municipal"),
    ("000017-06", "Piscina municipal"),
)
LUNES_CON = date(2026, 8, 10)  # última semana con WES (ICCP: 10–16)
LUNES_SIN = date(2026, 8, 17)
LUNES_PROX = date(2026, 8, 24)  # vuelve el control en los 4 puntos (no ICCP: trabajos)
DIAS_ICCP_CON = tuple(LUNES_CON + timedelta(days=i) for i in range(7))  # 10–16
DIAS_ICCP_SIN = tuple(LUNES_SIN + timedelta(days=i) for i in range(7))  # 17–23
HORAS_REGULADAS = (11, 12, 13)  # 11:00–13:59 cubre la regulación 11:00–13:30
HORAS_CONTROL_ICCO = frozenset(range(0, 6))
DOMINGO_SIN_ICCO = date(2026, 8, 16)
NOMBRE_ICCO = "Colegio ICCO Renca"
NOMBRE_ICCP = "ICCP (Cumbre de Cóndores pte.)"
TEXTO_SALA_BOMBAS = ""  # texto viejo; el informe usa TEXTO_CONTROL
TEXTO_CONTROL = (
    "Lo Velásquez, gimnasio, piscina e ICCO se evalúan con lo que va de esta semana "
    "(desde lun 24) contra la semana del 17/08 sin WES. "
    "Ahorro = (Sin WES − Con WES) / Sin WES × 100. "
    "El lunes 24/08, entre 11:00 y 13:30 Chile, se reguló Escuela Lo Velásquez y el "
    "gimnasio municipal con las mejoras de horario (el gimnasio ya no va con el "
    "tope parejo de 0,54 m³/h que se quedaba corto en un evento). "
    "En Lo Velásquez abrieron el bypass el jueves 27 ~17:00: de noche volvió el caudal "
    "parejo (~0,55 m³/h; antes 0). Desde entonces el punto no está con control."
)
TEXTO_ICCP = (
    "ICCP se informa aparte, con semana completa: "
    "lun 10–dom 16/08 con WES vs lun 17–dom 23/08 sin WES. "
    "Esta semana el control está off por OT 2282, por eso no se usa 24–26. "
    "Ahorro = (Sin WES − Con WES) / Sin WES × 100."
)
WD_CORTO = ("Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom")
WD_LARGO = ("Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo")


def _ventana_homologa(ahora: datetime) -> Tuple[Tuple[date, ...], Tuple[date, ...], int, bool]:
    """Días homólogos lun…hoy (máx. 7) y hora de corte del último día (23 = completo)."""
    hoy = ahora.date()
    if hoy >= LUNES_PROX:
        n = min(7, (hoy - LUNES_PROX).days + 1)
        ultimo_incompleto = hoy == (LUNES_PROX + timedelta(days=n - 1))
        hora_corte = max(0, min(23, ahora.hour - 1)) if ultimo_incompleto else 23
        if hora_corte < 0:
            hora_corte = 0
        dias_con = tuple(LUNES_PROX + timedelta(days=i) for i in range(n))
        dias_sin = tuple(LUNES_SIN + timedelta(days=i) for i in range(n))
        return dias_con, dias_sin, hora_corte, ultimo_incompleto
    if hoy < LUNES_SIN:
        n = 1
        hora_corte = 23
        ultimo_incompleto = False
    else:
        n = min(7, (hoy - LUNES_SIN).days + 1)
        ultimo_incompleto = hoy == (LUNES_SIN + timedelta(days=n - 1))
        hora_corte = max(0, min(23, ahora.hour - 1)) if ultimo_incompleto else 23
        if hora_corte < 0:
            hora_corte = 0
    dias_con = tuple(LUNES_CON + timedelta(days=i) for i in range(n))
    dias_sin = tuple(LUNES_SIN + timedelta(days=i) for i in range(n))
    return dias_con, dias_sin, hora_corte, ultimo_incompleto


def _ventana_hasta(hasta: date) -> Tuple[Tuple[date, ...], Tuple[date, ...], int, bool]:
    """Ventana homologada con último día completo (00:00–23:59)."""
    if hasta >= LUNES_PROX:
        n = min(7, (hasta - LUNES_PROX).days + 1)
        dias_con = tuple(LUNES_PROX + timedelta(days=i) for i in range(n))
        dias_sin = tuple(LUNES_SIN + timedelta(days=i) for i in range(n))
        return dias_con, dias_sin, 23, False
    if hasta < LUNES_SIN:
        return (LUNES_CON,), (LUNES_SIN,), 23, False
    n = min(7, (hasta - LUNES_SIN).days + 1)
    dias_con = tuple(LUNES_CON + timedelta(days=i) for i in range(n))
    dias_sin = tuple(LUNES_SIN + timedelta(days=i) for i in range(n))
    return dias_con, dias_sin, 23, False


def _fmt_rango_dias(dias: Sequence[date]) -> str:
    if len(dias) == 1:
        return f"{dias[0]:%d/%m/%Y}"
    if len(dias) == 2:
        return f"{dias[0]:%d/%m} y {dias[1]:%d/%m/%Y}"
    a, b = dias[0], dias[-1]
    return f"{a:%d/%m} al {b:%d/%m/%Y}"


def _hora_corte_de_dia(i: int, n_dias: int, hora_corte: int) -> int:
    return hora_corte if i == n_dias - 1 else 23


def _etiqueta_par(d_con: date, d_sin: date, hora_corte_dia: int) -> str:
    extra = f" ({_etiqueta_horas(hora_corte_dia)})" if hora_corte_dia < 23 else ""
    return f"{WD_CORTO[d_con.weekday()]} {d_con:%d} vs {d_sin:%d}{extra}"


def _etiqueta_horas(hora_corte: int) -> str:
    return "00:00–23:59" if hora_corte >= 23 else f"00:00–{hora_corte:02d}:59"


def _aplicar_corte_ultimo_dia(xlsx: Path, n_dias: int, hora_corte: int) -> None:
    """Horas posteriores al corte = 0 en el último día de cada periodo."""
    if hora_corte >= 23 or n_dias < 1:
        return
    wb = load_workbook(xlsx)
    ws = wb["Consolidado"]
    col_p1 = 1 + n_dias  # A=hora, B… = días periodo 1
    col_p2 = 1 + n_dias + n_dias
    for col in (col_p1, col_p2):
        for h in range(hora_corte + 1, 24):
            ws.cell(row=4 + h, column=col, value=0.0)
    wb.save(xlsx)
    wb.close()


def _rellenar_encabezado(row, headers: Sequence[str], fill: str = "1F4788") -> None:
    for j, hd in enumerate(headers):
        cell = row.cells[j]
        cell.text = hd
        _set_cell_shading(cell, fill)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)


def _rellenar_fila(row, vals: Sequence[str], size: int = 9, shade: Optional[str] = None) -> None:
    for j, v in enumerate(vals):
        row.cells[j].text = str(v)
        if shade:
            _set_cell_shading(row.cells[j], shade)
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(size)


def _nota_en_celda(cell, titulo: str, nota: str, size: int = 9) -> None:
    """Nombre del punto + criterio de evaluación (para Diego, en la misma celda)."""
    cell.text = titulo
    for run in cell.paragraphs[0].runs:
        run.font.size = Pt(size)
        run.bold = True
    p = cell.add_paragraph()
    r = p.add_run(nota)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(89, 89, 89)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _eval_esta_semana(dias_con: Sequence[date], dias_sin: Sequence[date]) -> str:
    return (
        f"Eval.: esta semana {_fmt_rango_dias(dias_con)} vs "
        f"sem. del 17 {_fmt_rango_dias(dias_sin)}"
    )


def _eval_iccp(iccp_info: dict) -> str:
    dc = iccp_info["dias_con"]
    ds = iccp_info["dias_sin"]
    return (
        f"Eval.: {_fmt_rango_dias(dc)} con WES vs {_fmt_rango_dias(ds)} sin WES "
        "(semana completa; no 24–26)"
    )


def _set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(
        parse_xml(
            "<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
            f'w:val="clear" w:fill="{hex_fill}"/>'
        )
    )


def _rendimiento(m3_sin: float, m3_con: float) -> Tuple[float, float]:
    ahorro = m3_sin - m3_con
    pct = (ahorro / m3_sin * 100.0) if m3_sin > 1e-9 else 0.0
    return ahorro, pct


def _hallazgo(pct: float) -> str:
    if pct >= 15:
        return "WES baja el consumo"
    if pct > 0:
        return "Ahorro menor"
    return "Sin WES no gasta más"


def _grafico_4_puntos(
    nombres: List[str],
    m3_con: List[float],
    m3_sin: List[float],
    titulo: str,
    lab_con: str,
    lab_sin: str,
    out_png: Path,
    ylabel: str = "Consumo (m³)",
) -> None:
    x = np.arange(len(nombres))
    w = 0.36
    fig, ax = plt.subplots(figsize=(11.4, 5.6))
    b1 = ax.bar(x - w / 2, m3_con, width=w, color=COLOR_WES, label=lab_con, zorder=2)
    b2 = ax.bar(x + w / 2, m3_sin, width=w, color=COLOR_SIN, label=lab_sin, zorder=2)
    ymax = max(list(m3_con) + list(m3_sin) + [1.0]) * 1.22
    ax.set_ylim(0, ymax)
    ax.set_xticks(list(x))
    ax.set_xticklabels(nombres, fontsize=10)
    ax.set_ylabel(ylabel)
    ax.set_title(titulo, fontweight="bold", fontsize=12, color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=9, loc="upper left")
    for bars in (b1, b2):
        for bar in bars:
            h = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                h + ymax * 0.015,
                format_number_chilean(float(h), 1),
                ha="center",
                va="bottom",
                fontsize=8,
                color="#333333",
            )
    for i, (a, b) in enumerate(zip(m3_sin, m3_con)):
        _, pct = _rendimiento(float(a), float(b))
        ax.text(
            x[i],
            ymax * 0.96,
            f"{pct:+.0f} %",
            ha="center",
            va="top",
            fontsize=9,
            fontweight="bold",
            color="#548235" if pct > 0 else COLOR_SIN,
        )
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_perfil_24h(
    vec_con: Sequence[float],
    vec_sin: Sequence[float],
    hora_corte: int,
    titulo: str,
    lab_con: str,
    lab_sin: str,
    out_png: Path,
    sombrear: Sequence[int] = (),
) -> None:
    h1 = 23 if hora_corte >= 23 else hora_corte
    horas = list(range(h1 + 1))
    y_con = [float(vec_con[h]) for h in horas]
    y_sin = [float(vec_sin[h]) for h in horas]
    fig, ax = plt.subplots(figsize=(11.4, 5.0))
    ax.plot(horas, y_con, color=COLOR_WES, linewidth=2.2, marker="o", markersize=4, label=lab_con, zorder=3)
    ax.plot(horas, y_sin, color=COLOR_SIN, linewidth=2.2, marker="o", markersize=4, label=lab_sin, zorder=3)
    for h in sombrear:
        if h <= h1:
            ax.axvspan(h - 0.5, h + 0.5, color="#FFF2CC", alpha=0.55, zorder=0)
    ax.set_xticks(list(range(0, h1 + 1, 1 if h1 <= 16 else 2)))
    ax.set_xlabel("Hora Chile (sombreado = regulación 11:00–13:30)")
    ax.set_ylabel("m³/h")
    ax.set_title(titulo, fontweight="bold", color="#1F4788")
    ax.grid(axis="y", alpha=0.3)
    ax.legend(fontsize=8, loc="upper left")
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _grafico_perfil_pares(
    pares: Sequence[Tuple[str, Sequence[float], Sequence[float], int, Sequence[int]]],
    titulo: str,
    lab_con: str,
    lab_sin: str,
    out_png: Path,
) -> None:
    """Un panel por par homólogo (lun 24 vs 17, mar 25 vs 18, …)."""
    n = len(pares)
    fig, axes = plt.subplots(1, n, figsize=(6.0 * max(n, 1), 4.8), sharey=True)
    if n == 1:
        axes = [axes]
    for ax, (subtitulo, vc, vs, hc, sombrear) in zip(axes, pares):
        h1 = 23 if hc >= 23 else hc
        horas = list(range(h1 + 1))
        ax.plot(
            horas,
            [float(vc[h]) for h in horas],
            color=COLOR_WES,
            linewidth=2.2,
            marker="o",
            markersize=4,
            label=lab_con,
            zorder=3,
        )
        ax.plot(
            horas,
            [float(vs[h]) for h in horas],
            color=COLOR_SIN,
            linewidth=2.2,
            marker="o",
            markersize=4,
            label=lab_sin,
            zorder=3,
        )
        for h in sombrear:
            if h <= h1:
                ax.axvspan(h - 0.5, h + 0.5, color="#FFF2CC", alpha=0.55, zorder=0)
        ax.set_xticks(list(range(0, h1 + 1, 1 if h1 <= 16 else 2)))
        ax.set_title(subtitulo, fontsize=10, fontweight="bold", color="#1F4788")
        ax.set_xlabel("Hora Chile")
        ax.set_ylabel("m³/h")
        ax.grid(axis="y", alpha=0.3)
        ax.legend(fontsize=7, loc="upper left")
        ax.set_ylim(bottom=0)
    fig.suptitle(titulo, fontweight="bold", color="#1F4788", y=1.02)
    fig.tight_layout()
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _suma_hasta(vec: Sequence[float], hora_corte: int) -> float:
    h1 = 23 if hora_corte >= 23 else hora_corte
    return float(sum(float(vec[h]) for h in range(h1 + 1)))


def _horas_noche_icco(hora_corte: int) -> List[int]:
    """Horas 00–05 que ya cerraron (si el corte es antes de las 06:00, la noche va incompleta)."""
    h1 = 23 if hora_corte >= 23 else hora_corte
    return [h for h in sorted(HORAS_CONTROL_ICCO) if h <= h1]


def _suma_noche_icco(vec: Sequence[float], hora_corte: int = 23) -> float:
    return float(sum(float(vec[h]) for h in _horas_noche_icco(hora_corte)))


def _preparar_icco(
    dias_con: Sequence[date],
    dias_sin: Sequence[date],
    hora_corte: int,
    png_dir: Path,
    xlsx_path: Path,
) -> dict:
    """Homóloga lun–vie con WES vs sin WES. El KPI es la noche 00:01–06:00."""
    print(f"  ICCO: descargando {len(dias_con) + len(dias_sin) + 1} días (homóloga + dom 16)…")
    v16 = _vector_m3h_24_desde_api(NODO_ICCO, DOMINGO_SIN_ICCO)
    m3_16_full = float(sum(v16))
    testigo_noche_full = _suma_noche_icco(v16, 23)
    vecs: Dict[date, List[float]] = {DOMINGO_SIN_ICCO: v16}
    for d in list(dias_con) + list(dias_sin):
        vecs[d] = _vector_m3h_24_desde_api(NODO_ICCO, d)

    filas: List[dict] = []
    for i, (d_con, d_sin) in enumerate(zip(dias_con, dias_sin)):
        incompleto = i == len(dias_con) - 1 and hora_corte < 23
        hc = hora_corte if incompleto else 23
        vc, vs = vecs[d_con], vecs[d_sin]
        m3_con = _suma_hasta(vc, hc)
        m3_sin = _suma_hasta(vs, hc)
        noche_con = _suma_noche_icco(vc, 23 if hc >= 5 else hc)
        noche_sin = _suma_noche_icco(vs, 23 if hc >= 5 else hc)
        a_dia, p_dia = _rendimiento(m3_sin, m3_con)
        a_noc, p_noc = _rendimiento(noche_sin, noche_con)
        filas.append(
            {
                "d_con": d_con,
                "d_sin": d_sin,
                "wd": WD_LARGO[d_con.weekday()],
                "m3_con": m3_con,
                "m3_sin": m3_sin,
                "ahorro_dia": a_dia,
                "pct_dia": p_dia,
                "noche_con": noche_con,
                "noche_sin": noche_sin,
                "ahorro": a_noc,
                "pct": p_noc,
                "incompleto": incompleto,
                "hora_corte": hc,
                "vec_con": vc,
                "vec_sin": vs,
            }
        )

    tot_con = float(sum(f["m3_con"] for f in filas))
    tot_sin = float(sum(f["m3_sin"] for f in filas))
    ahorro_dia, pct_dia = _rendimiento(tot_sin, tot_con)
    tot_noche_con = float(sum(f["noche_con"] for f in filas))
    tot_noche_sin = float(sum(f["noche_sin"] for f in filas))
    ahorro_t, pct_t = _rendimiento(tot_noche_sin, tot_noche_con)
    n_cero = sum(1 for f in filas if f["noche_con"] < 1.0)
    noches_sin_parejas = all(f["noche_sin"] >= 7.0 for f in filas)
    inc_pct = (ahorro_t / tot_sin * 100.0) if tot_sin > 1e-9 else 0.0

    png_dir.mkdir(parents=True, exist_ok=True)
    labels = []
    for f in filas:
        suf = f"\nhasta {f['hora_corte']:02d}:59" if f["incompleto"] else ""
        labels.append(f"{WD_CORTO[f['d_con'].weekday()]}\n{f['d_con']:%d} vs {f['d_sin']:%d}{suf}")

    png_barras = png_dir / "icco_noche_homologa_00_06.png"
    fig, ax = plt.subplots(figsize=(11.4, 5.2))
    x = np.arange(len(labels))
    w = 0.36
    ax.bar(x - w / 2, [f["noche_con"] for f in filas], width=w, color=COLOR_WES, label="Noche con WES", zorder=2)
    ax.bar(x + w / 2, [f["noche_sin"] for f in filas], width=w, color=COLOR_SIN, label="Noche sin WES", zorder=2)
    ax.axhline(
        testigo_noche_full,
        color="#C0504D",
        linestyle="--",
        linewidth=1.3,
        label=f"Fuga dom 16 ({format_number_chilean(testigo_noche_full, 1)} m³)",
    )
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_ylabel("m³ en 00:01–06:00")
    ax.set_title("ICCO — noche homóloga (esta semana sin control todos los días)", fontweight="bold", color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=8, loc="upper left")
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    fig.savefig(png_barras, dpi=150, bbox_inches="tight")
    plt.close(fig)

    png_dias = png_dir / "icco_dia_homologo.png"
    fig, ax = plt.subplots(figsize=(11.4, 5.2))
    ax.bar(x - w / 2, [f["m3_con"] for f in filas], width=w, color=COLOR_WES, label="Día con WES", zorder=2)
    ax.bar(x + w / 2, [f["m3_sin"] for f in filas], width=w, color=COLOR_SIN, label="Día sin WES", zorder=2)
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_ylabel("m³ día completo")
    ax.set_title("ICCO — día homólogo (ocupación: esta semana gasta menos de día)", fontweight="bold", color="#1F4788")
    ax.grid(axis="y", alpha=0.3, zorder=0)
    ax.legend(fontsize=8, loc="upper left")
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    fig.savefig(png_dias, dpi=150, bbox_inches="tight")
    plt.close(fig)

    png_perfil = png_dir / "icco_perfil_homologo.png"
    f_cero = next((f for f in filas if f["noche_con"] < 1.0), filas[-2] if len(filas) > 1 else filas[0])
    fig, ax = plt.subplots(figsize=(12.2, 5.0))
    horas = list(range(24))
    ax.plot(horas, v16, color="#C0504D", linewidth=1.6, linestyle=":", label="Dom 16 sin WES (fuga pareja)")
    ax.plot(horas, f_cero["vec_con"], color=COLOR_WES, linewidth=2.2, label=f"{WD_CORTO[f_cero['d_con'].weekday()]} {f_cero['d_con']:%d/%m} con WES")
    ax.plot(horas, f_cero["vec_sin"], color=COLOR_SIN, linewidth=2.2, label=f"{WD_CORTO[f_cero['d_sin'].weekday()]} {f_cero['d_sin']:%d/%m} sin WES")
    ax.axvspan(-0.5, 5.5, color="#FFF2CC", alpha=0.5, zorder=0)
    ax.set_xticks(list(range(0, 24, 2)))
    ax.set_xlabel("Hora Chile (sombreado = 00:01–06:00)")
    ax.set_ylabel("m³/h")
    ax.set_title("ICCO — de noche esta semana iguala la fuga; de día hay menos ocupación", fontweight="bold", color="#1F4788")
    ax.grid(axis="y", alpha=0.3)
    ax.legend(fontsize=8, loc="upper left")
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    fig.savefig(png_perfil, dpi=150, bbox_inches="tight")
    plt.close(fig)

    thin = Border(
        left=Side(style="thin", color="B0B0B0"),
        right=Side(style="thin", color="B0B0B0"),
        top=Side(style="thin", color="B0B0B0"),
        bottom=Side(style="thin", color="B0B0B0"),
    )
    wb = Workbook()
    ws = wb.active
    ws.title = "ICCO_homologa"
    headers = [
        "Día",
        "Con WES",
        "Sin WES",
        "Día con (m³)",
        "Día sin (m³)",
        "Ahorro día (m³)",
        "% día",
        "Noche 00–06 con (m³)",
        "Noche 00–06 sin (m³)",
        "Ahorro noche (m³)",
        "% noche",
        "Notas",
    ]
    head_font = XlFont(bold=True, color="FFFFFF", name="Calibri")
    fill_h = PatternFill("solid", fgColor="1F4788")
    fill_ok = PatternFill("solid", fgColor="E2EFDA")
    fill_no = PatternFill("solid", fgColor="F8CBAD")
    for i, h in enumerate(headers, 1):
        c = ws.cell(1, i, h)
        c.font = head_font
        c.fill = fill_h
        c.alignment = Alignment(wrap_text=True, horizontal="center")
        c.border = thin
    for r, f in enumerate(filas, start=2):
        nota = "Noche sin WES ≈ fuga del domingo 16" if f["noche_sin"] >= 7 else ""
        if f["noche_con"] < 1:
            nota = (nota + "; " if nota else "") + "WES dejó la noche en 0"
        if f["incompleto"]:
            nota += f"; hasta {f['hora_corte']:02d}:59 Chile"
        vals = [
            f["wd"],
            f["d_con"].strftime("%d/%m/%Y"),
            f["d_sin"].strftime("%d/%m/%Y"),
            round(f["m3_con"], 2),
            round(f["m3_sin"], 2),
            round(f["ahorro_dia"], 2),
            round(f["pct_dia"], 1),
            round(f["noche_con"], 2),
            round(f["noche_sin"], 2),
            round(f["ahorro"], 2),
            round(f["pct"], 1),
            nota,
        ]
        fill = fill_ok if f["pct"] >= 15 else fill_no
        for i, v in enumerate(vals, 1):
            c = ws.cell(r, i, v)
            c.border = thin
            c.fill = fill
    rr = 2 + len(filas)
    ws.cell(rr, 1, "TOTAL").font = XlFont(bold=True)
    ws.cell(rr, 4, round(tot_con, 2)).font = XlFont(bold=True)
    ws.cell(rr, 5, round(tot_sin, 2)).font = XlFont(bold=True)
    ws.cell(rr, 6, round(ahorro_dia, 2)).font = XlFont(bold=True)
    ws.cell(rr, 7, round(pct_dia, 1)).font = XlFont(bold=True)
    ws.cell(rr, 8, round(tot_noche_con, 2)).font = XlFont(bold=True)
    ws.cell(rr, 9, round(tot_noche_sin, 2)).font = XlFont(bold=True)
    ws.cell(rr, 10, round(ahorro_t, 2)).font = XlFont(bold=True)
    ws.cell(rr, 11, round(pct_t, 1)).font = XlFont(bold=True)
    ws2 = wb.create_sheet("Criterio")
    ws2["A1"] = "Qué cambió"
    ws2["A1"].font = XlFont(bold=True)
    ws2["B1"] = TEXTO_SALA_BOMBAS
    ws2["A2"] = "Dato a mostrar"
    ws2["A2"].font = XlFont(bold=True)
    ws2["B2"] = (
        f"Rendimiento ICCO ahora: noche homóloga {pct_t:.1f} % ({ahorro_t:.1f} m³). "
        f"Día completo {pct_dia:.1f} % no se usa (ocupación). "
        f"Fuga domingo 16: {m3_16_full:.1f} m³/día, {testigo_noche_full:.1f} m³ de 00:01 a 06:00. "
        "Semana del 24/08: repetir contra esta semana."
    )
    ws2.column_dimensions["A"].width = 22
    ws2.column_dimensions["B"].width = 110
    for col in range(1, 13):
        ws.column_dimensions[get_column_letter(col)].width = 14
    ws.column_dimensions["L"].width = 48
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(xlsx_path)

    return {
        "filas": filas,
        "m3_16": m3_16_full,
        "m3_9": 0.0,
        "ahorro_dom": 0.0,
        "pct_dom": 0.0,
        "testigo_noche": testigo_noche_full,
        "tot_con": tot_con,
        "tot_sin": tot_sin,
        "ahorro_dia": ahorro_dia,
        "pct_dia": pct_dia,
        "tot_noche_con": tot_noche_con,
        "tot_noche_sin": tot_noche_sin,
        "ahorro_t": ahorro_t,
        "pct_t": pct_t,
        "n_noches": len(filas),
        "n_cero": n_cero,
        "vol_dias_wes": tot_con,
        "inc_pct": inc_pct,
        "sin_control_esta_semana": noches_sin_parejas,
        "png_barras": png_barras,
        "png_dias": png_dias,
        "png_perfil": png_perfil,
        "xlsx": xlsx_path,
    }


def _convertir_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=180,
    )
    return pdf_path if pdf_path.is_file() else None


def _auditoria_un_punto(
    node_id: str,
    nombre: str,
    cliente_dir: Path,
    dias_con: Sequence[date],
    dias_sin: Sequence[date],
    hora_corte: int,
    per_con: Periodo,
    per_sin: Periodo,
    titulo_p1: str,
    titulo_p2: str,
) -> Tuple[Path, Optional[Path], float, float, List[float], List[float]]:
    cliente_dir.mkdir(parents=True, exist_ok=True)
    a, b = dias_con[0], dias_con[-1]
    c, d = dias_sin[0], dias_sin[-1]
    xlsx_out = (
        cliente_dir
        / f"consumo_consolidado_conWES_{a:%d%m}_sinWES_{c:%d%m}.xlsx"
    )
    print(f"  CSV + Excel {node_id} {nombre}…")
    generar_excel_consolidado(
        xlsx_out,
        node_id,
        dias_con,
        dias_sin,
        titulo_p1=titulo_p1,
        titulo_p2=titulo_p2,
        solo_desde_csv=False,
        csv_dir=cliente_dir / "csv_descarga_api",
    )
    _aplicar_corte_ultimo_dia(xlsx_out, len(dias_con), hora_corte)
    fechas, mats = leer_matriz_consolidado(xlsx_out)
    mid = len(fechas) // 2
    m3_con = float(sum(sum(col) for col in mats[:mid]))
    m3_sin = float(sum(sum(col) for col in mats[mid:]))
    dias_c = [float(sum(col)) for col in mats[:mid]]
    dias_s = [float(sum(col)) for col in mats[mid:]]
    return xlsx_out, None, m3_con, m3_sin, dias_c, dias_s


def _word_conjunto(
    out_docx: Path,
    hora_corte: int,
    ahora: datetime,
    dias_con: Sequence[date],
    dias_sin: Sequence[date],
    pngs: Dict[str, Path],
    m3_con: Dict[str, float],
    m3_sin: Dict[str, float],
    dias_m3_con: Dict[str, List[float]],
    dias_m3_sin: Dict[str, List[float]],
    dirs: Dict[str, Path],
    icco_info: Optional[dict] = None,
    horas_por_dia: Optional[Dict[str, List[Tuple[List[float], List[float]]]]] = None,
    puntos: Sequence[Tuple[str, str]] = PUNTOS_CONTROL,
    iccp_info: Optional[dict] = None,
) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    try:
        add_logo_to_header(doc)
    except Exception:
        pass

    h = doc.add_heading("Renca — 5 puntos", level=0)
    if h.runs:
        h.runs[0].font.color.rgb = COLOR_HEAD

    horas = _etiqueta_horas(hora_corte)
    n_dias = len(dias_con)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(
        f"Hasta {dias_con[-1]:%d/%m/%Y} {horas} Chile · tarifa 1.300 CLP/m³"
    )
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_HEAD

    tper = doc.add_table(rows=3, cols=3)
    tper.style = "Table Grid"
    _rellenar_encabezado(tper.rows[0], ["Puntos", "Semana sin WES", "Semana con WES"])
    _rellenar_fila(
        tper.rows[1],
        [
            "ICCO · Lo Velásquez · gimnasio · piscina",
            _fmt_rango_dias(dias_sin),
            _fmt_rango_dias(dias_con),
        ],
    )
    if iccp_info:
        _rellenar_fila(
            tper.rows[2],
            [
                "ICCP (10–16 vs 17–23)",
                _fmt_rango_dias(iccp_info["dias_sin"]),
                _fmt_rango_dias(iccp_info["dias_con"]),
            ],
            shade="FFF2CC",
        )
    else:
        _rellenar_fila(tper.rows[2], ["ICCP", "—", "—"])

    if "todos" in pngs:
        doc.add_picture(str(pngs["todos"]), width=Cm(16.2))

    tot_con = sum(m3_con[nid] for nid, _ in puntos)
    tot_sin = sum(m3_sin[nid] for nid, _ in puntos)
    ahorro_t, pct_t = _rendimiento(tot_sin, tot_con)
    n_iccp = 1 if iccp_info else 0
    tbl = doc.add_table(rows=1 + len(puntos) + 1 + n_iccp, cols=5)
    tbl.style = "Table Grid"
    _rellenar_encabezado(
        tbl.rows[0],
        ["Punto", "Con WES (m³)", "Sin WES (m³)", "Ahorro (m³)", "%"],
    )

    for i, (nid, nom) in enumerate(puntos, start=1):
        ahorro, pct = _rendimiento(m3_sin[nid], m3_con[nid])
        _rellenar_fila(
            tbl.rows[i],
            [
                nom,
                format_number_chilean(m3_con[nid], 1),
                format_number_chilean(m3_sin[nid], 1),
                format_number_chilean(ahorro, 1),
                format_number_chilean(pct, 1) + " %",
            ],
            shade="E2EFDA" if pct >= 15 else None,
        )

    row = tbl.rows[1 + len(puntos)]
    _rellenar_fila(
        row,
        [
            "Subtotal 4 (ICCO, Lo Velásquez, gimnasio, piscina)",
            format_number_chilean(tot_con, 1),
            format_number_chilean(tot_sin, 1),
            format_number_chilean(ahorro_t, 1),
            format_number_chilean(pct_t, 1) + " %",
        ],
        shade="D6DCE4",
    )
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True

    if iccp_info:
        a_i, p_i = _rendimiento(iccp_info["m3_sin"], iccp_info["m3_con"])
        _rellenar_fila(
            tbl.rows[2 + len(puntos)],
            [
                "ICCP",
                format_number_chilean(iccp_info["m3_con"], 1),
                format_number_chilean(iccp_info["m3_sin"], 1),
                format_number_chilean(a_i, 1),
                format_number_chilean(p_i, 1) + " %",
            ],
            shade="FFF2CC",
        )

    nomb_d: List[str] = [
        _etiqueta_par(d1, d2, _hora_corte_de_dia(i, n_dias, hora_corte))
        for i, (d1, d2) in enumerate(zip(dias_con, dias_sin))
    ]

    tot_con_dia = [sum(dias_m3_con[nid][i] for nid, _ in puntos) for i in range(n_dias)]
    tot_sin_dia = [sum(dias_m3_sin[nid][i] for nid, _ in puntos) for i in range(n_dias)]
    if "pares" in pngs:
        doc.add_picture(str(pngs["pares"]), width=Cm(16.2))
    tpar = doc.add_table(rows=1 + n_dias, cols=5)
    tpar.style = "Table Grid"
    _rellenar_encabezado(tpar.rows[0], ["Par (4 puntos)", "Con WES (m³)", "Sin WES (m³)", "Ahorro (m³)", "%"])
    for i, lab in enumerate(nomb_d):
        a, pc = _rendimiento(tot_sin_dia[i], tot_con_dia[i])
        _rellenar_fila(
            tpar.rows[i + 1],
            [
                lab,
                format_number_chilean(tot_con_dia[i], 1),
                format_number_chilean(tot_sin_dia[i], 1),
                format_number_chilean(a, 1),
                format_number_chilean(pc, 1) + " %",
            ],
        )

    if iccp_info:
        if iccp_info.get("png_barras") and Path(iccp_info["png_barras"]).is_file():
            doc.add_picture(str(iccp_info["png_barras"]), width=Cm(16.2))
        nomb_i = [
            _etiqueta_par(d1, d2, _hora_corte_de_dia(i, len(iccp_info["dias_con"]), iccp_info["hora_corte"]))
            for i, (d1, d2) in enumerate(zip(iccp_info["dias_con"], iccp_info["dias_sin"]))
        ]
        t_i = doc.add_table(rows=1 + len(nomb_i), cols=5)
        t_i.style = "Table Grid"
        _rellenar_encabezado(
            t_i.rows[0],
            ["Par ICCP", "Con WES 10–16 (m³)", "Sin WES 17–23 (m³)", "Ahorro (m³)", "%"],
        )
        for i, lab in enumerate(nomb_i):
            a, pc = _rendimiento(iccp_info["dias_m3_sin"][i], iccp_info["dias_m3_con"][i])
            _rellenar_fila(
                t_i.rows[i + 1],
                [
                    lab,
                    format_number_chilean(iccp_info["dias_m3_con"][i], 1),
                    format_number_chilean(iccp_info["dias_m3_sin"][i], 1),
                    format_number_chilean(a, 1),
                    format_number_chilean(pc, 1) + " %",
                ],
            )

    if horas_por_dia:
        for key, nom in (
            ("escuela", "Lo Velásquez"),
            ("gimnasio", "Gimnasio"),
            ("icco", "ICCO"),
            ("iccp", "ICCP"),
        ):
            if key in pngs and Path(pngs[key]).is_file():
                doc.add_paragraph(nom).runs[0].bold = True
                doc.add_picture(str(pngs[key]), width=Cm(16.2))

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def main() -> int:
    ap = argparse.ArgumentParser(description="Auditoría Renca homologada (4 puntos + ICCP al final).")
    ap.add_argument(
        "--hasta",
        type=str,
        default=None,
        help="Último día Chile completo (YYYY-MM-DD). Sin esto, corta a la hora actual.",
    )
    args = ap.parse_args()
    ahora = datetime.now(TZ_CL)
    if args.hasta:
        hasta = date.fromisoformat(args.hasta)
        dias_con, dias_sin, hora_corte, _inc = _ventana_hasta(hasta)
        ts = f"{hasta:%Y%m%d}_2359"
        ahora = datetime(hasta.year, hasta.month, hasta.day, 23, 59, tzinfo=TZ_CL)
    else:
        dias_con, dias_sin, hora_corte, _inc = _ventana_homologa(ahora)
        ts = ahora.strftime("%Y%m%d_%H%M")
    lab_con = f"Con WES {_fmt_rango_dias(dias_con)}"
    lab_sin = f"Sin WES {_fmt_rango_dias(dias_sin)}"
    per_con = Periodo(f"Con control ({dias_con[0]:%d-%m} a {dias_con[-1]:%d-%m-%Y})", tuple(dias_con))
    per_sin = Periodo(f"Sin control ({dias_sin[0]:%d-%m} a {dias_sin[-1]:%d-%m-%Y})", tuple(dias_sin))
    titulo_p1 = f"Con WES: {dias_con[0]:%d-%m-%Y} al {dias_con[-1]:%d-%m-%Y}"
    titulo_p2 = f"Sin WES: {dias_sin[0]:%d-%m-%Y} al {dias_sin[-1]:%d-%m-%Y}"

    out_root = (
        ROOT
        / "reports"
        / "reporte de auditoria"
        / f"auditoria_renca_semana_pasada_vs_esta_{ts}"
    )
    out_root.mkdir(parents=True, exist_ok=True)
    (out_root / "hora_corte.txt").write_text(str(hora_corte), encoding="utf-8")
    (out_root / "run_meta.json").write_text(
        json.dumps(
            {
                "hasta": dias_con[-1].isoformat(),
                "hora_corte": hora_corte,
                "dias_con": [d.isoformat() for d in dias_con],
                "dias_sin": [d.isoformat() for d in dias_sin],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    print(
        f"Auditoría Renca | {lab_con} vs {lab_sin} | "
        f"último día hasta {hora_corte:02d}:59 Chile | {ahora:%Y-%m-%d %H:%M}"
    )
    print("4 puntos con control: esta semana vs semana del 17")
    print("ICCP (al final): lun 10–dom 16 con WES vs lun 17–dom 23 sin WES")

    gxlsx.LABEL_P1 = lab_con
    gxlsx.LABEL_P2 = lab_sin

    m3_con: Dict[str, float] = {}
    m3_sin: Dict[str, float] = {}
    dias_m3_con: Dict[str, List[float]] = {}
    dias_m3_sin: Dict[str, List[float]] = {}
    dirs: Dict[str, Path] = {}
    horas_por_dia: Dict[str, List[Tuple[List[float], List[float]]]] = {}

    for nid, nom in PUNTOS_CONTROL:
        print("=" * 64)
        print(f"{nid} {nom}")
        cdir = out_root / f"Auditoria_{nom.split('(')[0].strip().replace(' ', '_')}_{nid}"
        dirs[nid] = cdir
        xlsx, _pdf, c, s, dc, ds = _auditoria_un_punto(
            nid,
            nom,
            cdir,
            dias_con,
            dias_sin,
            hora_corte,
            per_con,
            per_sin,
            titulo_p1,
            titulo_p2,
        )
        m3_con[nid] = c
        m3_sin[nid] = s
        dias_m3_con[nid] = dc
        dias_m3_sin[nid] = ds
        _fechas, mats = leer_matriz_consolidado(xlsx)
        n = len(dias_con)
        horas_por_dia[nid] = [(list(mats[i]), list(mats[n + i])) for i in range(n)]
        print(f"  Con {c:.1f} m³ | Sin {s:.1f} m³")

    n = len(dias_con)
    dias_iccp_con = DIAS_ICCP_CON
    dias_iccp_sin = DIAS_ICCP_SIN
    n_iccp = len(dias_iccp_con)
    hora_corte_iccp = 23
    print("=" * 64)
    print(f"{NODO_ICCP} {NOMBRE_ICCP} (10–16 con WES vs 17–23 sin WES)")
    cdir_iccp = out_root / f"Auditoria_ICCP_{NODO_ICCP}"
    dirs[NODO_ICCP] = cdir_iccp
    lab_iccp_con = f"Con WES {_fmt_rango_dias(dias_iccp_con)}"
    lab_iccp_sin = f"Sin WES {_fmt_rango_dias(dias_iccp_sin)}"
    xlsx_i, _pdf, c_i, s_i, dc_i, ds_i = _auditoria_un_punto(
        NODO_ICCP,
        NOMBRE_ICCP,
        cdir_iccp,
        dias_iccp_con,
        dias_iccp_sin,
        hora_corte_iccp,
        Periodo(lab_iccp_con, dias_iccp_con),
        Periodo(lab_iccp_sin, dias_iccp_sin),
        f"Con WES: {dias_iccp_con[0]:%d-%m-%Y} al {dias_iccp_con[-1]:%d-%m-%Y}",
        f"Sin WES: {dias_iccp_sin[0]:%d-%m-%Y} al {dias_iccp_sin[-1]:%d-%m-%Y}",
    )
    _fechas_i, mats_i = leer_matriz_consolidado(xlsx_i)
    horas_iccp = [(list(mats_i[i]), list(mats_i[n_iccp + i])) for i in range(n_iccp)]
    print(f"  Con {c_i:.1f} m³ | Sin {s_i:.1f} m³")

    png_dir = out_root / "graficos"
    png_dir.mkdir(exist_ok=True)
    nombres_5 = ["ICCO", "Lo Velásquez", "Gimnasio", "Piscina", "ICCP"]
    h4 = (n - 1) * 24 + (hora_corte + 1 if hora_corte < 23 else 24)
    m3d_4c = [m3_con[nid] * 24.0 / h4 for nid, _ in PUNTOS_CONTROL]
    m3d_4s = [m3_sin[nid] * 24.0 / h4 for nid, _ in PUNTOS_CONTROL]
    m3d_ic = c_i / 7.0
    m3d_is = s_i / 7.0
    p_all = png_dir / "todos_casos_con_vs_sin.png"
    _grafico_4_puntos(
        nombres_5,
        m3d_4c + [m3d_ic],
        m3d_4s + [m3d_is],
        "Renca — 5 puntos (m³/día)  |  ICCP al final: 10–16 vs 17–23",
        "Con WES",
        "Sin WES",
        p_all,
        ylabel="m³/día",
    )
    pngs: Dict[str, Path] = {"todos": p_all}

    labels_pares = [
        _etiqueta_par(d1, d2, _hora_corte_de_dia(i, n, hora_corte))
        for i, (d1, d2) in enumerate(zip(dias_con, dias_sin))
    ]
    tot_con_dia = [sum(dias_m3_con[nid][i] for nid, _ in PUNTOS_CONTROL) for i in range(n)]
    tot_sin_dia = [sum(dias_m3_sin[nid][i] for nid, _ in PUNTOS_CONTROL) for i in range(n)]
    p_pares = png_dir / "pares_4puntos_esta_vs_17.png"
    _grafico_4_puntos(
        labels_pares,
        tot_con_dia,
        tot_sin_dia,
        "4 puntos — ICCO, Lo Velásquez, gimnasio, piscina",
        lab_con,
        lab_sin,
        p_pares,
    )
    pngs["pares"] = p_pares

    for nid, key, nom in (
        (NODO_ESCUELA, "escuela", "Lo Velásquez"),
        (NODO_GIMNASIO, "gimnasio", "Gimnasio"),
        (NODO_ICCO, "icco", "ICCO"),
    ):
        pares_plot: List[Tuple[str, Sequence[float], Sequence[float], int, Sequence[int]]] = []
        for i, (d1, d2) in enumerate(zip(dias_con, dias_sin)):
            hc = _hora_corte_de_dia(i, n, hora_corte)
            vc, vs = horas_por_dia[nid][i]
            sombrear = HORAS_REGULADAS if nid != NODO_ICCO else ()
            pares_plot.append(
                (
                    f"{WD_CORTO[d1.weekday()]} {d1:%d/%m} vs {d2:%d/%m}",
                    vc,
                    vs,
                    hc,
                    sombrear,
                )
            )
        png = png_dir / f"perfil_{key}_esta_vs_17.png"
        _grafico_perfil_pares(
            pares_plot,
            f"{nom} — esta semana (con WES) contra semana del 17 (sin WES)",
            lab_con,
            lab_sin,
            png,
        )
        pngs[key] = png

    p_iccp_bar = png_dir / "iccp_10_16_vs_17_23.png"
    labels_iccp = [
        _etiqueta_par(d1, d2, 23)
        for d1, d2 in zip(dias_iccp_con, dias_iccp_sin)
    ]
    _grafico_4_puntos(
        labels_iccp,
        dc_i,
        ds_i,
        "ICCP — lun 10–dom 16 con WES vs lun 17–dom 23 sin WES",
        lab_iccp_con,
        lab_iccp_sin,
        p_iccp_bar,
    )
    pares_iccp: List[Tuple[str, Sequence[float], Sequence[float], int, Sequence[int]]] = []
    for i in (0, 6):
        d1, d2 = dias_iccp_con[i], dias_iccp_sin[i]
        vc, vs = horas_iccp[i]
        pares_iccp.append(
            (
                f"{WD_CORTO[d1.weekday()]} {d1:%d/%m} vs {d2:%d/%m}",
                vc,
                vs,
                23,
                (),
            )
        )
    p_iccp_prf = png_dir / "perfil_iccp_10_16_vs_17_23.png"
    _grafico_perfil_pares(
        pares_iccp,
        "ICCP — lunes y domingo homólogos (10–16 vs 17–23)",
        lab_iccp_con,
        lab_iccp_sin,
        p_iccp_prf,
    )
    iccp_info = {
        "dias_con": dias_iccp_con,
        "dias_sin": dias_iccp_sin,
        "m3_con": c_i,
        "m3_sin": s_i,
        "dias_m3_con": dc_i,
        "dias_m3_sin": ds_i,
        "horas": horas_iccp,
        "hora_corte": hora_corte_iccp,
        "png": p_iccp_prf,
        "png_barras": p_iccp_bar,
    }
    pngs["iccp"] = p_iccp_prf

    docx_path = out_root / f"Informe_Auditoria_Renca_semana_pasada_vs_esta_{ts}.docx"
    _word_conjunto(
        docx_path,
        hora_corte,
        ahora,
        dias_con,
        dias_sin,
        pngs,
        m3_con,
        m3_sin,
        dias_m3_con,
        dias_m3_sin,
        dirs,
        None,
        horas_por_dia,
        PUNTOS_CONTROL,
        iccp_info,
    )
    pdf_path = _convertir_pdf(docx_path)

    tot_c = sum(m3_con[nid] for nid, _ in PUNTOS_CONTROL)
    tot_s = sum(m3_sin[nid] for nid, _ in PUNTOS_CONTROL)
    ahorro, pct = _rendimiento(tot_s, tot_c)
    print("=" * 64)
    print(f"4 puntos Con {tot_c:.1f} | Sin {tot_s:.1f} | {pct:.1f}% ({ahorro:.1f} m³)")
    for i, lab in enumerate(labels_pares):
        a_d, p_d = _rendimiento(tot_sin_dia[i], tot_con_dia[i])
        print(
            f"  {lab}: Con {tot_con_dia[i]:.1f} | Sin {tot_sin_dia[i]:.1f} | "
            f"{p_d:.1f}% ({a_d:.1f} m³)"
        )
    for nid, nom in PUNTOS_CONTROL:
        a, p = _rendimiento(m3_sin[nid], m3_con[nid])
        print(f"  {nom}: {p:.1f}% ({a:.1f} m³)")
    a_i, p_i = _rendimiento(s_i, c_i)
    print(f"  ICCP 10–16 vs 17–23: {p_i:.1f}% ({a_i:.1f} m³) | con {c_i:.1f} | sin {s_i:.1f}")
    print(f"DOCX conjunto: {docx_path}")
    print(f"PDF  conjunto: {pdf_path or '(no convertido)'}")
    print(f"DIR: {out_root}")
    return 0


if __name__ == "__main__":
    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
            sys.stderr.reconfigure(encoding="utf-8")
        except Exception:
            pass
    raise SystemExit(main())
