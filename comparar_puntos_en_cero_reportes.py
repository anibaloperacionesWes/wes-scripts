"""
Compara dos reportes Word de puntos en cero (p. ej. mayo vs hoy) y genera
un informe Word + PDF con resumen y diferencias por punto.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from wes_paths import reporte_cero_dir


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:fill="{hex_color}"/>'
    )
    cell._tePr = cell._tc.get_or_add_tcPr()
    cell._tc.get_or_add_tcPr().append(shading)


def _style_header_row(row, fill: str = "1F4E79") -> None:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        shading = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:fill="{fill}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)


def extraer_reporte(path: Path) -> Dict[str, Any]:
    doc = Document(str(path))
    meta: Dict[str, Any] = {
        "total": 0,
        "con_datos": 0,
        "cero": 0,
        "sin_datos": 0,
        "pct_cero": "0%",
        "pct_sin": "0%",
        "fecha": path.stem,
        "archivo": path.name,
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Reporte generado:"):
            meta["fecha"] = t.replace("Reporte generado:", "").strip()
        if "Total de puntos analizados:" in t:
            for line in t.splitlines():
                line = line.strip()
                if line.startswith("Total de puntos analizados:"):
                    meta["total"] = int(line.split(":", 1)[1].strip())
                elif line.startswith("Puntos con datos disponibles:"):
                    meta["con_datos"] = int(line.split(":", 1)[1].strip())
                elif line.startswith("Puntos marcando cero:"):
                    meta["cero"] = int(line.split(":", 1)[1].strip())
                elif line.startswith("Puntos sin datos disponibles:"):
                    meta["sin_datos"] = int(line.split(":", 1)[1].strip())
                elif line.startswith("Porcentaje en cero:"):
                    meta["pct_cero"] = line.split(":", 1)[1].strip()
                elif line.startswith("Porcentaje sin datos:"):
                    meta["pct_sin"] = line.split(":", 1)[1].strip()
            break

    tablas: List[List[Dict[str, str]]] = []
    for table in doc.tables:
        rows: List[Dict[str, str]] = []
        for i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if i == 0:
                continue
            if cells and cells[0] and "Nodo" not in cells[0]:
                rows.append(
                    {
                        "nodeId": cells[0],
                        "nodeName": cells[1] if len(cells) > 1 else "",
                        "companyName": cells[2] if len(cells) > 2 else "",
                        "companyId": cells[3] if len(cells) > 3 else "",
                    }
                )
        if rows:
            tablas.append(rows)

    meta["puntos_cero"] = tablas[0] if len(tablas) > 0 else []
    meta["puntos_sin_datos"] = tablas[1] if len(tablas) > 1 else []
    return meta


def _index(puntos: List[Dict[str, str]]) -> Dict[str, Dict[str, str]]:
    return {p["nodeId"]: p for p in puntos if p.get("nodeId")}


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _style_header_row(hdr)
    for r_i, row_vals in enumerate(rows, start=1):
        for c_i, val in enumerate(row_vals):
            table.rows[r_i].cells[c_i].text = str(val)
            for p in table.rows[r_i].cells[c_i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)


def crear_comparacion(
    base: Dict[str, Any],
    actual: Dict[str, Any],
    label_base: str,
    label_actual: str,
    output_dir: Path,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("COMPARACIÓN PUNTOS EN CERO", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph(f"{label_base}  vs  {label_actual}")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].bold = True

    gen = doc.add_paragraph(
        f"Informe generado: {datetime.now(timezone.utc).strftime('%d-%m-%Y %H:%M:%S')} UTC"
    )
    gen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        f"Base ({label_base}): {base['archivo']} — {base['fecha']}\n"
        f"Actual ({label_actual}): {actual['archivo']} — {actual['fecha']}"
    )

    doc.add_heading("RESUMEN COMPARATIVO", 1)
    delta_cero = actual["cero"] - base["cero"]
    delta_sin = actual["sin_datos"] - base["sin_datos"]
    delta_total = actual["total"] - base["total"]
    sign = lambda n: f"+{n}" if n > 0 else str(n)

    resumen = doc.add_paragraph(
        f"Total puntos: {base['total']} → {actual['total']} ({sign(delta_total)})\n"
        f"Puntos en cero: {base['cero']} ({base['pct_cero']}) → "
        f"{actual['cero']} ({actual['pct_cero']}) ({sign(delta_cero)})\n"
        f"Puntos sin datos: {base['sin_datos']} ({base['pct_sin']}) → "
        f"{actual['sin_datos']} ({actual['pct_sin']}) ({sign(delta_sin)})"
    )
    resumen.runs[0].font.size = Pt(11)

    _add_table(
        doc,
        ["Indicador", label_base, label_actual, "Δ"],
        [
            ["Total analizados", str(base["total"]), str(actual["total"]), sign(delta_total)],
            ["Con datos", str(base["con_datos"]), str(actual["con_datos"]), sign(actual["con_datos"] - base["con_datos"])],
            ["En cero", str(base["cero"]), str(actual["cero"]), sign(delta_cero)],
            ["% en cero", base["pct_cero"], actual["pct_cero"], ""],
            ["Sin datos", str(base["sin_datos"]), str(actual["sin_datos"]), sign(delta_sin)],
            ["% sin datos", base["pct_sin"], actual["pct_sin"], ""],
        ],
    )

    base_cero = _index(base["puntos_cero"])
    act_cero = _index(actual["puntos_cero"])
    base_sin = _index(base["puntos_sin_datos"])
    act_sin = _index(actual["puntos_sin_datos"])

    persistentes = sorted(set(base_cero) & set(act_cero))
    nuevos = sorted(set(act_cero) - set(base_cero))
    recuperados = sorted(set(base_cero) - set(act_cero))
    nuevos_sin = sorted(set(act_sin) - set(base_sin))
    recuperados_sin = sorted(set(base_sin) - set(act_sin))
    persistentes_sin = sorted(set(base_sin) & set(act_sin))

    doc.add_heading("HALLAZGOS CLAVE", 1)
    hallazgos = doc.add_paragraph(
        f"• Siguen en cero desde {label_base}: {len(persistentes)}\n"
        f"• Nuevos en cero (no estaban en {label_base}): {len(nuevos)}\n"
        f"• Recuperados (estaban en cero en {label_base} y ya no): {len(recuperados)}\n"
        f"• Nuevos sin datos: {len(nuevos_sin)}\n"
        f"• Dejaron de estar sin datos: {len(recuperados_sin)}"
    )
    hallazgos.runs[0].font.size = Pt(11)

    def _rows_from_ids(ids: List[str], catalog: Dict[str, Dict[str, str]]) -> List[List[str]]:
        out = []
        for nid in ids:
            p = catalog.get(nid) or base_cero.get(nid) or act_cero.get(nid) or base_sin.get(nid) or act_sin.get(nid) or {
                "nodeId": nid,
                "nodeName": "",
                "companyName": "",
            }
            out.append([p["nodeId"], p.get("nodeName", ""), p.get("companyName", "")])
        return out

    headers_pts = ["Nodo ID", "Nombre", "Empresa"]

    doc.add_heading(f"PERSISTEN EN CERO ({len(persistentes)})", 1)
    if persistentes:
        _add_table(doc, headers_pts, _rows_from_ids(persistentes, act_cero))
    else:
        doc.add_paragraph("Ninguno.")

    doc.add_heading(f"NUEVOS EN CERO ({len(nuevos)})", 1)
    if nuevos:
        _add_table(doc, headers_pts, _rows_from_ids(nuevos, act_cero))
    else:
        doc.add_paragraph("Ninguno.")

    doc.add_heading(f"RECUPERADOS DESDE {label_base.upper()} ({len(recuperados)})", 1)
    if recuperados:
        note = []
        for nid in recuperados:
            p = base_cero[nid]
            estado = "sin datos hoy" if nid in act_sin else "con consumo / fuera de cero"
            note.append([p["nodeId"], p["nodeName"], p["companyName"], estado])
        _add_table(doc, ["Nodo ID", "Nombre", "Empresa", "Estado actual"], note)
    else:
        doc.add_paragraph("Ninguno.")

    doc.add_heading(f"SIN DATOS — NUEVOS ({len(nuevos_sin)})", 1)
    if nuevos_sin:
        _add_table(doc, headers_pts, _rows_from_ids(nuevos_sin, act_sin))
    else:
        doc.add_paragraph("Ninguno.")

    doc.add_heading(f"SIN DATOS — PERSISTEN ({len(persistentes_sin)})", 1)
    if persistentes_sin:
        _add_table(doc, headers_pts, _rows_from_ids(persistentes_sin, act_sin))
    else:
        doc.add_paragraph("Ninguno.")

    doc.add_heading("CONCLUSIÓN", 1)
    if delta_cero > 0:
        conclusion = (
            f"Respecto de {label_base}, los puntos en cero aumentaron en {delta_cero} "
            f"({base['cero']} → {actual['cero']}; {actual['pct_cero']} del universo con datos). "
            f"Hay {len(nuevos)} puntos nuevos en cero y {len(persistentes)} que se mantienen desde la base."
        )
    elif delta_cero < 0:
        conclusion = (
            f"Respecto de {label_base}, los puntos en cero bajaron en {abs(delta_cero)} "
            f"({base['cero']} → {actual['cero']})."
        )
    else:
        conclusion = f"La cantidad de puntos en cero se mantiene igual que en {label_base} ({actual['cero']})."
    doc.add_paragraph(conclusion)

    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    out = output_dir / f"Comparacion_Puntos_En_Cero_{stamp}.docx"
    doc.save(str(out))
    return out


def convertir_a_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    return pdf_path if pdf_path.exists() else None


def main() -> int:
    parser = argparse.ArgumentParser(description="Comparar dos reportes de puntos en cero")
    parser.add_argument("--base", type=Path, required=True, help="DOCX base (p. ej. mayo)")
    parser.add_argument("--actual", type=Path, required=True, help="DOCX actual (p. ej. hoy 07:00)")
    parser.add_argument("--label-base", default="Mayo 2026")
    parser.add_argument("--label-actual", default="Hoy 07:00")
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("reports/Puntos_En_Cero"),
    )
    parser.add_argument("--pdf-only-keep", action="store_true", help="Elimina el DOCX tras convertir a PDF")
    args = parser.parse_args()

    base = extraer_reporte(args.base.resolve())
    actual = extraer_reporte(args.actual.resolve())
    docx = crear_comparacion(
        base,
        actual,
        args.label_base,
        args.label_actual,
        args.output_dir.resolve(),
    )
    print(f"[OK] DOCX: {docx}")
    pdf = convertir_a_pdf(docx)
    if not pdf:
        print("[ERROR] No se pudo convertir a PDF (falta LibreOffice/soffice).", flush=True)
        return 1
    print(f"[OK] PDF: {pdf}")
    if args.pdf_only_keep:
        docx.unlink(missing_ok=True)
        print("[INFO] DOCX eliminado (--pdf-only-keep).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
