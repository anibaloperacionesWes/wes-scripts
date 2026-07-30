"""
Corrige Word Bupa Antofagasta:
1) Reinserta el gráfico de día mayor nocturno de Sanitaria (relación de imagen rota).
2) Ajusta m³/CLP de la cuenta con la factura julio (6.696 m³ / $18.538.860).
"""

from __future__ import annotations

import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

from generar_reporte_word import (
    add_formatted_title,
    add_picture_with_pagination,
    add_table,
    build_hourly_consumption_line_chart,
    find_max_nocturnal_consumption_day,
    format_currency_chilean,
    format_number_chilean,
    parse_date,
)

DOC_PATH = Path(
    "reports/Bupa_Antofagasta/ABREGADO/AGREGADO_20260728_1725/"
    "Reporte_Agregado_BUPA_20260723_20260728.docx"
)
OUT_DIR = DOC_PATH.parent

# Factura julio — Centro Médico Antofagasta (N° 1460583, emisión 03-07-2026)
FACTURA_M3 = 6696.0
FACTURA_CLP = 18_538_860.0
PRECIO_M3 = FACTURA_CLP / FACTURA_M3  # ~2.768,65 CLP/m³

NUM_DIAS_WES = 6
START = parse_date("23/07/2026")
END = parse_date("28/07/2026", end_of_day=True)

# Consumos WES del periodo (salas); sanitaria cuenta = factura
NODOS = [
    {"id": "000029-07", "nombre": "Sala de Bomba Principal", "tipo": "bomba", "m3_periodo": 227.9},
    {"id": "000029-08", "nombre": "Sala de Bomba Sexto Piso", "tipo": "bomba", "m3_periodo": 36.1},
    {"id": "000029-09", "nombre": "Medidor Principal Sanitaria", "tipo": "cuenta", "m3_periodo": 1400.0},
    {"id": "000029-10", "nombre": "Sala de Bomba N°2", "tipo": "bomba", "m3_periodo": 6.2},
]


def _para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()


def _find_elem(doc: Document, pred):
    for child in list(doc.element.body):
        if child.tag.split("}")[-1] == "p" and pred(_para_text(child)):
            return child
    return None


def _remove_from_until(doc: Document, start_pred, end_pred) -> int:
    body = doc.element.body
    children = list(body)
    start_idx = end_idx = None
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] != "p":
            continue
        t = _para_text(child)
        if start_idx is None and start_pred(t):
            start_idx = i
        elif start_idx is not None and end_pred(t):
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Rango no encontrado ({start_idx}, {end_idx})")
    n = 0
    for child in children[start_idx:end_idx]:
        body.remove(child)
        n += 1
    return n


def _move_tail_before(doc: Document, marker, n_new: int) -> None:
    """Mueve los últimos n_new elementos del body (antes de sectPr) delante de marker."""
    body = doc.element.body
    children = [c for c in list(body) if c.tag.split("}")[-1] != "sectPr"]
    new_els = children[-n_new:]
    for el in new_els:
        marker.addprevious(el)


def _chart_participacion(path: Path, salas: float, cuenta: float) -> Path:
    resto = max(0.0, cuenta - salas)
    fig, ax = plt.subplots(figsize=(6.2, 4.5))
    wedges, *_ = ax.pie(
        [salas, resto],
        labels=["Salas de bomba\n(07+08+10)", "Resto de la cuenta\n(factura − salas)"],
        colors=["#c0392b", "#0050b3"],
        autopct=lambda p: f"{p:.1f}%",
        startangle=90,
        textprops={"fontsize": 10},
    )
    for w in wedges:
        w.set_edgecolor("white")
        w.set_linewidth(1.5)
    ax.set_title(
        "Participación proyectada sobre cuenta factura (6.696 m³)",
        fontsize=12,
        fontweight="bold",
    )
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _chart_proyeccion(path: Path, rows: list[dict]) -> Path:
    nombres = [r["nombre"] for r in rows]
    valores = [r["proy_m3"] for r in rows]
    fig, ax = plt.subplots(figsize=(8, 4.8))
    bars = ax.bar(nombres, valores, color="#0050b3", alpha=0.85, edgecolor="#003a80", linewidth=1.1)
    ax.set_ylabel("m³ / mes", fontsize=12, fontweight="bold")
    ax.set_title("Proyección / cuenta mensual por punto", fontsize=14, fontweight="bold")
    ax.set_ylim(bottom=0)
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=25, ha="right", fontsize=10)
    for bar, val in zip(bars, valores):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(val, 1)} m³",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )
    plt.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def main() -> None:
    rows = []
    for n in NODOS:
        if n["tipo"] == "cuenta":
            proy_m3 = FACTURA_M3
            proy_clp = FACTURA_CLP
            fuente = "Factura julio (consumo período)"
            prom = n["m3_periodo"] / NUM_DIAS_WES
        else:
            prom = n["m3_periodo"] / NUM_DIAS_WES
            proy_m3 = prom * 30.0
            proy_clp = proy_m3 * PRECIO_M3
            fuente = "WES · prom. diario × 30"
        rows.append(
            {
                **n,
                "prom_dia": prom,
                "proy_m3": proy_m3,
                "proy_clp": proy_clp,
                "fuente": fuente,
            }
        )

    # Gráfico Sanitaria
    result = find_max_nocturnal_consumption_day("000029-09", None, START, END)
    if not result:
        raise RuntimeError("Sin datos horarios para Sanitaria")
    target_dt, hourly = result
    chart_sanitaria = OUT_DIR / "chart_max_nocturnal_000029_09.png"
    build_hourly_consumption_line_chart(
        hourly,
        chart_sanitaria,
        target_dt,
        f"Día con mayor consumo nocturno ({target_dt.strftime('%d-%m-%y')})",
    )
    if not chart_sanitaria.exists() or chart_sanitaria.stat().st_size < 1000:
        raise RuntimeError("Gráfico Sanitaria no se generó correctamente")

    salas = [r for r in rows if r["tipo"] == "bomba"]
    salas_m3 = sum(r["proy_m3"] for r in salas)
    salas_clp = sum(r["proy_clp"] for r in salas)
    pct = salas_m3 / FACTURA_M3 * 100.0

    chart_proy = _chart_proyeccion(OUT_DIR / "chart_proyeccion_mensual_total_por_punto.png", rows)
    chart_part = _chart_participacion(
        OUT_DIR / "chart_participacion_salas_vs_sanitaria.png", salas_m3, FACTURA_M3
    )

    doc = Document(str(DOC_PATH))
    n_del = _remove_from_until(
        doc,
        lambda t: "MAYOR CONSUMO NOCTURNO" in t.upper() and "SANITARIA" in t.upper(),
        lambda t: t.upper().startswith("CONCLUSIONES"),
    )
    print(f"[OK] Eliminado bloque roto/anterior ({n_del} elementos)")

    conclus = _find_elem(doc, lambda t: t.upper().startswith("CONCLUSIONES"))
    if conclus is None:
        raise RuntimeError("CONCLUSIONES no encontrado")

    # Contar hijos actuales (sin sectPr) para saber cuántos agregamos
    before = len([c for c in doc.element.body if c.tag.split("}")[-1] != "sectPr"])

    # Agregar al final del doc (crea relaciones de imagen válidas en ESTE documento)
    add_formatted_title(
        doc,
        f"Día con mayor consumo nocturno - MEDIDOR PRINCIPAL SANITARIA "
        f"({target_dt.strftime('%d-%m-%y')}):",
    )
    add_picture_with_pagination(doc, str(chart_sanitaria), Inches(6), keep_with_next=True)
    p = doc.add_paragraph(
        "El gráfico corresponde al día con mayor consumo nocturno (00:00–06:59) del "
        "Medidor Principal Sanitaria en el periodo WES. Este punto representa el total "
        "de la cuenta de agua de la clínica."
    )
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    add_formatted_title(doc, "Proyección de consumo mensual total por punto")
    intro = doc.add_paragraph(
        f"Las salas de bomba se proyectan con promedio diario WES del periodo "
        f"({NUM_DIAS_WES} días: 23/07/2026–28/07/2026) × 30. "
        f"El Medidor Principal Sanitaria se corrige con la factura de agua de julio: "
        f"{format_number_chilean(FACTURA_M3, 0)} m³ (lectura {format_number_chilean(94664, 0)} → "
        f"{format_number_chilean(101360, 0)}) y total a pagar "
        f"{format_currency_chilean(FACTURA_CLP)} "
        f"(≈ ${format_number_chilean(PRECIO_M3, 0)} CLP/m³)."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in intro.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    table_proy = [
        (
            "Punto",
            "Consumo WES periodo (m³)",
            "Prom. diario (m³)",
            "Mensual (m³)",
            "Mensual (CLP)",
            "Fuente",
        )
    ]
    for r in rows:
        table_proy.append(
            (
                f"{r['nombre']} ({r['id']})",
                format_number_chilean(r["m3_periodo"], 1),
                format_number_chilean(r["prom_dia"], 1),
                format_number_chilean(r["proy_m3"], 1),
                format_currency_chilean(r["proy_clp"]),
                r["fuente"],
            )
        )
    add_table(doc, "Proyección / cuenta mensual", table_proy, wes_style=True)
    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_proy), Inches(6), keep_with_next=True)

    add_formatted_title(doc, "Participación de salas de bomba vs cuenta (factura)")
    expl = doc.add_paragraph(
        "La base de la cuenta es el consumo facturado de julio "
        f"({format_number_chilean(FACTURA_M3, 0)} m³ / {format_currency_chilean(FACTURA_CLP)}). "
        "Las salas de bomba se valorizan con la misma tarifa efectiva de la factura "
        f"(${format_number_chilean(PRECIO_M3, 0)} CLP/m³) sobre su proyección WES a 30 días."
    )
    expl.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in expl.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    part_rows = [
        ("Concepto", "Mensual (m³)", "Mensual (CLP)", "% sobre factura"),
        (
            "Medidor Principal Sanitaria / cuenta factura",
            format_number_chilean(FACTURA_M3, 0),
            format_currency_chilean(FACTURA_CLP),
            "100,0%",
        ),
    ]
    for b in salas:
        part_rows.append(
            (
                f"{b['nombre']} ({b['id']})",
                format_number_chilean(b["proy_m3"], 1),
                format_currency_chilean(b["proy_clp"]),
                f"{format_number_chilean(b['proy_m3'] / FACTURA_M3 * 100, 1)}%",
            )
        )
    part_rows.append(
        (
            "TOTAL salas de bomba (07+08+10)",
            format_number_chilean(salas_m3, 1),
            format_currency_chilean(salas_clp),
            f"{format_number_chilean(pct, 1)}%",
        )
    )
    add_table(doc, "Participación sobre cuenta (factura julio)", part_rows, wes_style=True)

    resumen = doc.add_paragraph()
    resumen.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    resumen.add_run("Resumen: ").bold = True
    resumen.add_run(
        f"sobre la cuenta facturada de {format_number_chilean(FACTURA_M3, 0)} m³ "
        f"({format_currency_chilean(FACTURA_CLP)}), las tres salas de bomba proyectan "
        f"{format_number_chilean(salas_m3, 1)} m³ ({format_currency_chilean(salas_clp)}), "
        f"equivalentes al {format_number_chilean(pct, 1)}% del volumen y del costo de la cuenta."
    )
    for run in resumen.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart_part), Inches(5.2), keep_with_next=True)
    doc.add_paragraph("")

    after = len([c for c in doc.element.body if c.tag.split("}")[-1] != "sectPr"])
    n_new = after - before
    _move_tail_before(doc, conclus, n_new)
    print(f"[OK] Insertados {n_new} elementos con imágenes válidas")

    # Verificar que el gráfico de Sanitaria quedó con relación existente
    ok_img = False
    for child in doc.element.body:
        if child.tag.split("}")[-1] != "p":
            continue
        if "MAYOR CONSUMO NOCTURNO" in _para_text(child).upper() and "SANITARIA" in _para_text(child).upper():
            # siguiente párrafo con imagen
            nxt = child.getnext()
            while nxt is not None and nxt.tag.split("}")[-1] == "p" and not list(nxt.iter(qn("a:blip"))):
                if _para_text(nxt):
                    break
                nxt = nxt.getnext()
            if nxt is not None:
                blips = list(nxt.iter(qn("a:blip")))
                if blips:
                    rid = blips[0].get(qn("r:embed"))
                    ok_img = rid in doc.part.rels
                    print(f"[CHECK] blip {rid} en rels: {ok_img}")
            break
    if not ok_img:
        print("[ADVERTENCIA] No se pudo verificar la relación de imagen")

    try:
        doc.save(str(DOC_PATH))
        out = DOC_PATH
    except PermissionError:
        out = DOC_PATH.with_name(DOC_PATH.stem + "_factura.docx")
        doc.save(str(out))
        print(f"[ADVERTENCIA] Guardado como {out.name}")

    print("[OK]", out)
    print(f"     Factura: {FACTURA_M3:.0f} m³ / {format_currency_chilean(FACTURA_CLP)}")
    print(f"     Precio efectivo: ${PRECIO_M3:,.0f} CLP/m³".replace(",", "."))
    print(f"     Salas: {salas_m3:.1f} m³ = {pct:.1f}% de la cuenta")
    print(f"     Día nocturno Sanitaria: {target_dt.strftime('%d-%m-%Y')}")


if __name__ == "__main__":
    main()
