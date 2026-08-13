"""
Bupa / UPA Antofagasta — reporte agregado formato clásico extendido
(mismo estilo Club Providencia / Fundo Zapallar de fin de mes).

Periodo por defecto: 23/07/2026 – 11/08/2026 (días civiles completos).

Uso:
  python generar_agregado_bupa_antofagasta_clasico.py
"""

from __future__ import annotations

import sys
import time
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

if sys.platform == "win32":
    for s in (sys.stdout, sys.stderr):
        try:
            s.reconfigure(encoding="utf-8", line_buffering=True)
        except Exception:
            pass

from generar_reporte_word import (
    add_formatted_heading,
    add_formatted_title,
    add_picture_with_pagination,
    add_table,
    build_monthly_comparison_chart,
    calculate_nocturnal_metrics,
    estilizar_tabla_wes,
    format_currency_chilean,
    format_number_chilean,
    generate_aggregated_report,
    get_water_price_per_m3,
    parse_date,
    acl_node_base_url,
    fetch_json,
    flatten_measures,
    normalize_measures_payload,
    summarize_consumption,
)

COMPANY_ID = "000029"
# Puntos UPA Antofagasta (clínica)
NODE_IDS = [
    "000029-07",  # Sala de Bomba Principal
    "000029-08",  # Sala de Bomba Sexto Piso
    "000029-09",  # Medidor Principal Sanitaria (cuenta)
    "000029-10",  # Sala de Bomba N°2
]
# Comparativos de cuenta (evita doble conteo con salas)
CUENTA_ID = "000029-09"
START = "23/07/2026"
END = "11/08/2026"
FOLDER = "Bupa_Antofagasta"
COLOR_BARRA = "#0050b3"
PRECIO_DEFAULT = 2768.65  # tarifa ref. factura julio

# Hito: bajada de consumo por mejora del equipo de mantención
MEJORA_DATE = date(2026, 7, 29)
BEFORE_START = "23/07/2026"
BEFORE_END = "28/07/2026"  # ritmo previo a la mejora
AFTER_START = "29/07/2026"
AFTER_END = "11/08/2026"  # ritmo post-mejora (días completos)
AUGUST_YEAR = 2026
AUGUST_DAYS = 31
AUGUST_OBS_END = "11/08/2026"  # último día observado de agosto
# 4 semanas de 5 días sobre el periodo monitoreado (20 días)
NUM_SEMANAS = 4
DIAS_POR_SEMANA = 5


def _semanas_periodo(start_d: date, end_d: date) -> List[Tuple[date, date]]:
    """Parte el periodo en NUM_SEMANAS bloques de DIAS_POR_SEMANA días."""
    out: List[Tuple[date, date]] = []
    d = start_d
    for _ in range(NUM_SEMANAS):
        if d > end_d:
            break
        w_end = min(d + timedelta(days=DIAS_POR_SEMANA - 1), end_d)
        out.append((d, w_end))
        d = w_end + timedelta(days=1)
    return out


def _daily_cuenta(start_dt: datetime, end_dt: datetime) -> Dict[date, float]:
    raw = fetch_json(
        f"{acl_node_base_url()}/nodes/measures/dates",
        params=[
            ("id", CUENTA_ID),
            ("start", start_dt.strftime("%d%m%Y")),
            ("end", end_dt.strftime("%d%m%Y")),
        ],
    )
    measures = flatten_measures(normalize_measures_payload(raw, CUENTA_ID))
    return {m.date.date(): float(m.total_m3) for m in measures}


def _plot_semanas(
    series: List[Tuple[str, float, float]],
    out: Path,
) -> Path:
    """series: (label, total_m3, promedio_diario)."""
    labels = [s[0] for s in series]
    vals = [s[1] for s in series]
    # Semana 1 (previa) en rojo; resto en azul WES — resalta la bajada
    colors = ["#c0392b" if i == 0 else COLOR_BARRA for i in range(len(vals))]
    fig, ax = plt.subplots(figsize=(9.5, 5.4))
    bars = ax.bar(labels, vals, color=colors, width=0.62, edgecolor="#333333", linewidth=0.7)
    ax.set_ylabel("Consumo semanal cuenta (m³)", fontsize=11, fontweight="bold")
    ax.set_xlabel("Semana del periodo monitoreado", fontsize=11, fontweight="bold")
    ax.set_title(
        "Comparativo semanal — Medidor Principal Sanitaria",
        fontsize=12,
        fontweight="bold",
    )
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    ax.set_ylim(bottom=0)
    for bar, (_lab, tot, prom) in zip(bars, series):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(tot, 0)} m³\n({format_number_chilean(prom, 0)} /día)",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return out


def _totales_periodo(start_dt: datetime, end_dt: datetime) -> Tuple[float, float, float]:
    """Retorna (total_m3 todos los puntos, nocturno_m3, price)."""
    total = 0.0
    noct = 0.0
    prices: List[float] = []
    for nid in NODE_IDS:
        raw = fetch_json(
            f"{acl_node_base_url()}/nodes/measures/dates",
            params=[
                ("id", nid),
                ("start", start_dt.strftime("%d%m%Y")),
                ("end", end_dt.strftime("%d%m%Y")),
            ],
        )
        payload = normalize_measures_payload(raw, nid)
        measures = flatten_measures(payload)
        summary = summarize_consumption(measures)
        total += float(summary["total"])
        nm = calculate_nocturnal_metrics(nid, start_dt, end_dt)
        noct += float(nm.get("consumo_nocturno_total") or 0.0)
        prices.append(get_water_price_per_m3(COMPANY_ID, nid, payload) or PRECIO_DEFAULT)
    price = sum(prices) / len(prices) if prices else PRECIO_DEFAULT
    return total, noct, price


def _consumo_cuenta_periodo(start_dt: datetime, end_dt: datetime) -> float:
    raw = fetch_json(
        f"{acl_node_base_url()}/nodes/measures/dates",
        params=[
            ("id", CUENTA_ID),
            ("start", start_dt.strftime("%d%m%Y")),
            ("end", end_dt.strftime("%d%m%Y")),
        ],
    )
    payload = normalize_measures_payload(raw, CUENTA_ID)
    return float(summarize_consumption(flatten_measures(payload))["total"])


def _promedio_diario_cuenta(start_s: str, end_s: str) -> Tuple[float, float, int]:
    """Retorna (total_m3, promedio_diario, num_dias)."""
    start_dt = parse_date(start_s)
    end_dt = parse_date(end_s, end_of_day=True)
    total = _consumo_cuenta_periodo(start_dt, end_dt)
    n = (end_dt.date() - start_dt.date()).days + 1
    return total, (total / n if n else 0.0), n


def _plot_proyeccion_agosto(
    path: Path,
    sin_mejora: float,
    con_mejora: float,
    observado: float,
) -> Path:
    labels = [
        "Agosto si se\nmantenía el ritmo\nprevio",
        "Agosto proyectado\ncon la mejora\n(cierre de mes)",
        "Ya consumido\nen agosto\n(01–11)",
    ]
    vals = [sin_mejora, con_mejora, observado]
    colors = ["#c0392b", "#2ecc71", "#0050b3"]
    fig, ax = plt.subplots(figsize=(8.5, 5.2))
    bars = ax.bar(labels, vals, color=colors, width=0.62, edgecolor="#333333", linewidth=0.8)
    ax.set_ylabel("m³ en agosto", fontsize=11, fontweight="bold")
    ax.set_title(
        "Proyección de cierre de agosto — Medidor Principal Sanitaria",
        fontsize=12,
        fontweight="bold",
    )
    ax.set_ylim(bottom=0)
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    for bar, v in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height(),
            f"{format_number_chilean(v, 0)} m³",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches="tight")
    plt.close(fig)
    return path


def _append_proyeccion_agosto_y_mejora(doc: Document, out_dir: Path, price: float) -> None:
    """Proyección de cierre de agosto + narrativa WES / mantención."""
    before_tot, before_avg, before_n = _promedio_diario_cuenta(BEFORE_START, BEFORE_END)
    after_tot, after_avg, after_n = _promedio_diario_cuenta(AFTER_START, AFTER_END)
    ago_obs = _consumo_cuenta_periodo(
        parse_date("01/08/2026"),
        parse_date(AUGUST_OBS_END, end_of_day=True),
    )
    dias_obs_ago = (parse_date(AUGUST_OBS_END).date() - date(AUGUST_YEAR, 8, 1)).days + 1
    dias_restantes = AUGUST_DAYS - dias_obs_ago
    proy_cierre = ago_obs + after_avg * dias_restantes
    proy_sin_mejora = before_avg * AUGUST_DAYS
    ahorro_m3 = max(0.0, proy_sin_mejora - proy_cierre)
    ahorro_clp = ahorro_m3 * price
    pct_baja = (1.0 - after_avg / before_avg) * 100.0 if before_avg > 0 else 0.0

    chart = _plot_proyeccion_agosto(
        out_dir / "chart_proyeccion_cierre_agosto.png",
        proy_sin_mejora,
        proy_cierre,
        ago_obs,
    )

    add_formatted_heading(doc, "Proyección de cierre de agosto y efecto de la mejora", level=1)

    hito = doc.add_paragraph()
    hito.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r0 = hito.add_run("Hito 29/07/2026: ")
    r0.bold = True
    hito.add_run(
        "en los gráficos diarios (punto verde) se marca la fecha en que el consumo de la "
        "cuenta bajó de forma nítida. Ese día el Medidor Principal Sanitaria pasó de un "
        f"ritmo cercano a {format_number_chilean(before_avg, 1)} m³/día "
        f"(promedio {BEFORE_START}–{BEFORE_END}) a unos "
        f"{format_number_chilean(after_avg, 1)} m³/día "
        f"(promedio {AFTER_START}–{AFTER_END}), una reducción del orden del "
        f"{format_number_chilean(pct_baja, 0)} %. Esta bajada se asocia a una mejora "
        "ejecutada por el equipo de mantención, a partir de la información de consumo "
        "que entrega WES."
    )

    add_formatted_title(doc, "Proyección mensual — cierre de agosto 2026")
    intro = doc.add_paragraph(
        "Para cerrar la cuenta de agosto se usa el consumo ya registrado "
        f"(01/08–{AUGUST_OBS_END}) y se proyectan los "
        f"{dias_restantes} días restantes al ritmo posterior a la mejora. "
        "En paralelo se muestra qué habría sido el mes si se hubiera mantenido "
        "el consumo previo a la intervención."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    rows = [
        (
            "Concepto",
            "m³ / día",
            "Base de cálculo",
            "Total agosto (m³)",
            "Valorizado (CLP)",
        ),
        (
            "Ritmo antes de la mejora",
            format_number_chilean(before_avg, 1),
            f"{BEFORE_START}–{BEFORE_END} ({before_n} días)",
            format_number_chilean(proy_sin_mejora, 0),
            format_currency_chilean(proy_sin_mejora * price),
        ),
        (
            "Ritmo después de la mejora",
            format_number_chilean(after_avg, 1),
            f"{AFTER_START}–{AFTER_END} ({after_n} días)",
            "—",
            "—",
        ),
        (
            "Ya consumido en agosto",
            format_number_chilean(ago_obs / dias_obs_ago, 1),
            f"01/08–{AUGUST_OBS_END} ({dias_obs_ago} días)",
            format_number_chilean(ago_obs, 0),
            format_currency_chilean(ago_obs * price),
        ),
        (
            "Proyección cierre agosto (con mejora)",
            format_number_chilean(proy_cierre / AUGUST_DAYS, 1),
            f"Observado + {dias_restantes} días × ritmo post-mejora",
            format_number_chilean(proy_cierre, 0),
            format_currency_chilean(proy_cierre * price),
        ),
        (
            "Ahorro proyectado vs ritmo previo",
            "—",
            "Diferencia de cierre de mes",
            format_number_chilean(ahorro_m3, 0),
            format_currency_chilean(ahorro_clp),
        ),
    ]
    add_table(doc, "Proyección / ahorro agosto", rows, wes_style=True)

    lectura = doc.add_paragraph()
    lectura.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r1 = lectura.add_run("Lectura para el cliente: ")
    r1.bold = True
    lectura.add_run(
        f"si el consumo se hubiera mantenido en ~{format_number_chilean(before_avg, 0)} m³/día, "
        f"agosto habría cerrado cerca de {format_number_chilean(proy_sin_mejora, 0)} m³ "
        f"({format_currency_chilean(proy_sin_mejora * price)}). Con el ritmo actual post-mejora, "
        f"el cierre proyectado es de ~{format_number_chilean(proy_cierre, 0)} m³ "
        f"({format_currency_chilean(proy_cierre * price)}), es decir un ahorro del orden de "
        f"{format_number_chilean(ahorro_m3, 0)} m³ "
        f"({format_currency_chilean(ahorro_clp)}) solo en este mes."
    )

    doc.add_paragraph("")
    add_picture_with_pagination(doc, str(chart), Inches(5.8), keep_with_next=True)

    add_formatted_title(doc, "Rol de WES y del equipo de mantención")
    wes = doc.add_paragraph()
    wes.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    wes.add_run(
        "La información que entrega WES es clave para el equipo de mantención: permite ver "
        "en qué momento el consumo se disparó o se mantuvo alto, actuar con precisión y "
        "verificar de inmediato el efecto de las mejoras. En este caso, la bajada marcada "
        f"el {MEJORA_DATE.strftime('%d/%m/%Y')} muestra que, cuando el equipo interviene "
        "apoyado en los datos de monitoreo, la clínica puede reducir de forma sustancial "
        "el volumen facturado y, con ello, la cuenta de agua. Mantener este seguimiento "
        "permite sostener el ahorro y detectar a tiempo cualquier nuevo desvío."
    )

    print(
        f"[OK] Proyección agosto: sin mejora {proy_sin_mejora:.0f} m³ | "
        f"cierre con mejora {proy_cierre:.0f} m³ | ahorro {ahorro_m3:.0f} m³ "
        f"({format_currency_chilean(ahorro_clp)})"
    )


def _append_comparativos(docx_path: Path) -> Path:
    start_dt = parse_date(START)
    end_dt = parse_date(END, end_of_day=True)
    out_dir = docx_path.parent

    print("[INFO] Calculando totales periodo para comparativos...", flush=True)
    total_m3, noct_total, price = _totales_periodo(start_dt, end_dt)
    num_dias = (end_dt.date() - start_dt.date()).days + 1
    factor_30 = 30.0 / max(num_dias, 1)
    leak_monthly = noct_total * factor_30
    efectivo_monthly = max(0.0, (total_m3 - noct_total) * factor_30)

    chart_mes = out_dir / "chart_comparacion_mensual_extra.png"
    built_mes = build_monthly_comparison_chart(
        leak_monthly, efectivo_monthly, price, chart_mes
    )

    print("[INFO] Armando comparativo semanal (cuenta Sanitaria)...", flush=True)
    daily = _daily_cuenta(start_dt, end_dt)
    cuenta_periodo = sum(daily.values())
    semanas = _semanas_periodo(start_dt.date(), end_dt.date())
    series_sem: List[Tuple[str, float, float, int, str]] = []
    for i, (w0, w1) in enumerate(semanas, start=1):
        days = [w0 + timedelta(days=k) for k in range((w1 - w0).days + 1)]
        tot = sum(daily.get(d, 0.0) for d in days)
        n = len(days)
        prom = tot / n if n else 0.0
        rango = f"{w0.strftime('%d/%m')}–{w1.strftime('%d/%m')}"
        label = f"Semana {i}\n{rango}"
        nota = "Antes de la mejora" if i == 1 else (
            "Transición (incluye 29/07)" if w0 <= MEJORA_DATE <= w1 else "Después de la mejora"
        )
        series_sem.append((label, tot, prom, n, nota))
        print(f"  Semana {i} {rango}: {tot:.1f} m³ ({prom:.1f}/día) — {nota}", flush=True)

    chart_sem = _plot_semanas(
        [(s[0], s[1], s[2]) for s in series_sem],
        out_dir / "chart_comparativo_semanal.png",
    )

    doc = Document(str(docx_path))
    add_formatted_heading(doc, "Comparativo del mes (nocturno vs efectivo)", level=1)
    p5 = doc.add_paragraph(
        f"Proyección a 30 días a partir del periodo ({num_dias} días): "
        f"nocturno proyectado {format_number_chilean(leak_monthly, 1)} m³ "
        f"({format_currency_chilean(leak_monthly * price)}) y consumo efectivo "
        f"{format_number_chilean(efectivo_monthly, 1)} m³ "
        f"({format_currency_chilean(efectivo_monthly * price)})."
    )
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if built_mes and Path(built_mes).exists():
        add_picture_with_pagination(doc, str(chart_mes), Inches(4.5), keep_with_next=True)

    # En vez de 6 meses (aún no hay historial suficiente), comparamos las 4 semanas monitoreadas.
    add_formatted_heading(doc, "Comparativo semanal del periodo monitoreado", level=1)
    p6 = doc.add_paragraph(
        "Todavía no hay historial de varios meses en este medidor, por eso el comparativo "
        "se hace sobre las 4 semanas del periodo con datos "
        f"({start_dt.strftime('%d/%m/%Y')} – {end_dt.strftime('%d/%m/%Y')}, "
        f"{format_number_chilean(cuenta_periodo, 1)} m³ en total en la cuenta). "
        "Cada barra agrupa 5 días. La Semana 1 (rojo) refleja el ritmo alto previo a la "
        f"mejora del {MEJORA_DATE.strftime('%d/%m/%Y')}; las siguientes muestran la bajada sostenida."
    )
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_picture_with_pagination(doc, str(chart_sem), Inches(6), keep_with_next=True)

    table_rows = [
        ("Semana", "Rango", "Días", "Total (m³)", "Promedio diario (m³)", "Observación"),
    ]
    for i, (lab, tot, prom, n, nota) in enumerate(series_sem, start=1):
        w0, w1 = semanas[i - 1]
        table_rows.append(
            (
                f"Semana {i}",
                f"{w0.strftime('%d/%m')}–{w1.strftime('%d/%m')}",
                str(n),
                format_number_chilean(tot, 1),
                format_number_chilean(prom, 1),
                nota,
            )
        )
    # Fila de contraste S1 vs promedio S3–S4
    if len(series_sem) >= 4:
        s1 = series_sem[0][2]
        post = (series_sem[2][2] + series_sem[3][2]) / 2.0
        baja = (1.0 - post / s1) * 100.0 if s1 else 0.0
        table_rows.append(
            (
                "Contraste",
                "S1 vs S3–S4",
                "—",
                "—",
                f"{format_number_chilean(s1, 0)} → {format_number_chilean(post, 0)}",
                f"Bajada ~{format_number_chilean(baja, 0)} % en m³/día",
            )
        )
    add_table(doc, "Consumo semanal — cuenta Sanitaria", table_rows, wes_style=True)

    # Proyección cierre agosto + narrativa mejora mantención / WES
    # Valorización con tarifa de factura julio (ref. cuenta), no el default API.
    _append_proyeccion_agosto_y_mejora(doc, out_dir, PRECIO_DEFAULT)

    doc.save(str(docx_path))
    print(f"[OK] Comparativos agregados a {docx_path}")
    return docx_path


def main() -> int:
    print("=" * 70)
    print("BUPA / UPA ANTOFAGASTA — agregado clásico extendido + comparativos")
    print(f"Periodo: {START} → {END}")
    print(f"Nodos: {', '.join(NODE_IDS)}")
    print("=" * 70)
    t0 = time.perf_counter()
    out = generate_aggregated_report(
        company_id=COMPANY_ID,
        node_ids=list(NODE_IDS),
        start_date=START,
        end_date=END,
        output_dir="reports",
        apply_exclusions=False,
        generate_ppt=False,
        parallel_node_fetch=True,
        max_parallel_workers=4,
        fuente_agua_id=None,
        company_folder_override=FOLDER,
    )
    docx = Path(out)
    print(f"[OK] Base: {docx}")
    _append_comparativos(docx)
    print(f"[INFO] Tiempo total: {time.perf_counter() - t0:.1f} s")
    print(f"[INFO] Word: {docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
