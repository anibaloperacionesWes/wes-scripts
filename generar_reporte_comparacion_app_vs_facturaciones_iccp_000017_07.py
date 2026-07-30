"""
Genera un reporte tipo informe (Word + PDF) que compara:
- "App" = consumo WES (API medidas) del nodo 000017-07 (ICCP Renca)
- Facturaciones Aguas Andinas = consumo total m³ informado en cada factura PDF

Fuente de facturas:
  reports/Renca/Coparacion App con Aguas Andinas/ICCP facturas/*.pdf

Salida:
  reports/Renca/Coparacion App con Aguas Andinas/reporte_comparacion_Iccp/

Uso:
  python generar_reporte_comparacion_app_vs_facturaciones_iccp_000017_07.py
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

from generar_reporte_word import convertir_word_a_pdf, format_number_chilean

ROOT = Path(__file__).resolve().parent
IN_DIR = (
    ROOT
    / "reports"
    / "Renca"
    / "Coparacion App con Aguas Andinas"
    / "ICCP facturas"
)
OUT_DIR = (
    ROOT
    / "reports"
    / "Renca"
    / "Coparacion App con Aguas Andinas"
    / "reporte_comparacion_Iccp"
)

NODE_ID = "000017-07"
NODE_NOMBRE = "ICCP Renca (000017-07)"
EMPRESA = "Renca"

# Ventana de transición (operación): 11-dic-2025 → 08-ene-2026
TRANSICION_INICIO = datetime(2025, 12, 11).date()
TRANSICION_FIN = datetime(2026, 1, 8).date()

# Referencia operaciones / cruce Excel (misma línea de tiempo que ICCO Renca cuando aplica).
MONITOREO_WES_DESDE = datetime(2025, 8, 26).date()
PERIODO_ALINEADO_EXCEL_INI = datetime(2025, 10, 11).date()
PERIODO_ALINEADO_EXCEL_FIN = datetime(2025, 11, 11).date()

# Orden del gráfico: por lectura actual (cronológico)
# (si necesitas boletas específicas como en ICCO, se puede fijar lista aquí)


def _ahora_informe_local() -> datetime:
    return datetime.now(timezone.utc).astimezone()


def _doc_establecer_metadatos(doc: Document, generado_en: datetime) -> None:
    cp = doc.core_properties
    cp.created = generado_en
    cp.modified = generado_en
    cp.title = "Comparación App WES vs Facturaciones Aguas Andinas"
    cp.subject = f"{NODE_ID} — {NODE_NOMBRE} — {EMPRESA}"
    cp.keywords = "WES; Aguas Andinas; ICCP; Renca; facturación; consumo"
    cp.category = "Informe comparativo"
    cp.comments = "Informe generado automáticamente: facturas Aguas Andinas (PDF) vs consumo API nodo WES."
    cp.author = "Script informe ICCP Renca"
    cp.last_modified_by = cp.author
    try:
        cp.revision = 1
    except Exception:
        pass


def _docx_parchear_core_hora_local(out_docx: Path, generado_en: datetime) -> None:
    # Reescribir core.xml con offset local (evita desfase de hora en propiedades Word)
    if generado_en.tzinfo is None:
        generado_en = generado_en.astimezone()
    marca = generado_en.isoformat(timespec="seconds")

    with ZipFile(out_docx, "r") as zin:
        if "docProps/core.xml" not in zin.namelist():
            return
        core = zin.read("docProps/core.xml").decode("utf-8")

    def _reemplaza(tag: str, xml: str) -> str:
        pat = re.compile(rf"(<dcterms:{tag}\s[^>]*>)([^<]*)(</dcterms:{tag}>)", re.DOTALL)
        if not pat.search(xml):
            return xml
        return pat.sub(rf"\g<1>{marca}\g<3>", xml, count=1)

    core2 = _reemplaza("created", core)
    core2 = _reemplaza("modified", core2)
    if core2 == core:
        return

    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    tmp_path = Path(tmp)
    try:
        with ZipFile(out_docx, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "docProps/core.xml":
                    data = core2.encode("utf-8")
                zout.writestr(info, data)
        shutil.move(str(tmp_path), str(out_docx))
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass


def _fmt_num(n: float, dec: int = 1) -> str:
    return format_number_chilean(n, dec)


_MESES = {
    "ENE": 1,
    "FEB": 2,
    "MAR": 3,
    "ABR": 4,
    "MAY": 5,
    "JUN": 6,
    "JUL": 7,
    "AGO": 8,
    "SEP": 9,
    "OCT": 10,
    "NOV": 11,
    "DIC": 12,
}


def _parse_fecha_es_dd_mmm_yyyy(s: str) -> datetime:
    p = s.strip().upper().split("-")
    if len(p) != 3:
        raise ValueError(f"Fecha inválida: {s!r}")
    dd = int(p[0])
    mm = _MESES[p[1][:3]]
    yy = int(p[2])
    return datetime(yy, mm, dd)


def _fmt_periodo(dt: datetime) -> str:
    meses = {
        1: "ene",
        2: "feb",
        3: "mar",
        4: "abr",
        5: "may",
        6: "jun",
        7: "jul",
        8: "ago",
        9: "sep",
        10: "oct",
        11: "nov",
        12: "dic",
    }
    return f"{dt.day:02d}-{meses[dt.month]}-{dt.year}"


def _to_ddmmyyyy(dt: datetime) -> str:
    return dt.strftime("%d%m%Y")


def _extraer_texto_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join([(pg.extract_text() or "") for pg in reader.pages])


@dataclass(frozen=True)
class FacturaAA:
    pdf: Path
    boleta: str
    cuenta: str | None
    medidor: str | None
    emision: datetime
    lectura_anterior: datetime
    lectura_actual: datetime
    m3_cuenta: int

    @property
    def api_start(self) -> str:
        return _to_ddmmyyyy(self.lectura_anterior)

    @property
    def api_end(self) -> str:
        return _to_ddmmyyyy(self.lectura_actual)

    @property
    def periodo_txt(self) -> str:
        return f"{_fmt_periodo(self.lectura_anterior)} → {_fmt_periodo(self.lectura_actual)}"


def _parse_factura_aguas_andinas_pdf(path: Path) -> FacturaAA:
    txt = _extraer_texto_pdf(path)

    boleta_m = re.search(
        r"(?:FACTURA|BOLETA)\s+ELECTR[ÓO]NICA\s*\n?\s*N[º°]\s*([0-9]{6,})",
        txt,
        flags=re.IGNORECASE,
    )
    if not boleta_m:
        boleta_m = re.search(r"\nN[º°]\s*([0-9]{6,})\n", txt, flags=re.IGNORECASE)
    emision_m = re.search(
        r"FECHA\s+EMISI[ÓO]N[:\s]*([0-9]{2}-[A-Z]{3}-[0-9]{4})",
        txt,
        flags=re.IGNORECASE,
    )
    cuenta_m = re.search(r"\n([0-9]{7,}-[0-9])\n", txt)
    medidor_m = re.search(r"N[úu]mero\s+de\s+Medidor\s+([0-9\.]+)", txt, flags=re.IGNORECASE)
    consumo_m = re.search(r"CONSUMO\s+TOTAL\s+([0-9\.\,]+)\s*m3", txt, flags=re.IGNORECASE)
    # Hay PDFs que traen solo la fecha y un guion (sin m3 a continuación).
    la_m = re.search(
        r"LECTURA\s+ACTUAL\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})",
        txt,
        flags=re.IGNORECASE,
    )
    lan_m = re.search(
        r"LECTURA\s+ANTERIOR\s+([0-9]{2}-[A-Z]{3}-[0-9]{4})",
        txt,
        flags=re.IGNORECASE,
    )

    if not (emision_m and consumo_m and la_m and lan_m):
        raise ValueError(f"No se pudieron extraer campos desde {path.name}")

    dt_emision = _parse_fecha_es_dd_mmm_yyyy(emision_m.group(1))
    dt_actual = _parse_fecha_es_dd_mmm_yyyy(la_m.group(1))
    dt_anterior = _parse_fecha_es_dd_mmm_yyyy(lan_m.group(1))
    m3_cuenta = int(float(consumo_m.group(1).replace(".", "").replace(",", ".")))

    return FacturaAA(
        pdf=path,
        boleta=(boleta_m.group(1) if boleta_m else path.stem),
        cuenta=(cuenta_m.group(1) if cuenta_m else None),
        medidor=(medidor_m.group(1) if medidor_m else None),
        emision=dt_emision,
        lectura_anterior=dt_anterior,
        lectura_actual=dt_actual,
        m3_cuenta=m3_cuenta,
    )


def _cargar_facturas() -> list[FacturaAA]:
    pdfs = sorted(IN_DIR.glob("*.pdf"))
    if not pdfs:
        raise FileNotFoundError(f"No hay PDFs en {IN_DIR}")
    rows = [_parse_factura_aguas_andinas_pdf(p) for p in pdfs]
    rows.sort(key=lambda r: r.lectura_actual)
    return rows


def _particionar_por_inicio_monitoreo(
    facturas: list[FacturaAA], desde: date
) -> tuple[list[FacturaAA], list[FacturaAA]]:
    """Incluye solo ciclos cuya lectura anterior y lectura actual son ``desde`` o posteriores.

    Quedan fuera períodos enteramente anteriores, o con alguna lectura antes de ``desde``
    (no comparable de punta a punta con el monitoreo WES desde esa fecha).
    """
    incl: list[FacturaAA] = []
    excl: list[FacturaAA] = []
    for f in facturas:
        d0 = f.lectura_anterior.date()
        d1 = f.lectura_actual.date()
        if d0 >= desde and d1 >= desde:
            incl.append(f)
        else:
            excl.append(f)
    return incl, excl


def _fetch_wes_total(f: FacturaAA) -> tuple[float, int]:
    from generar_reporte_word import (
        acl_node_base_url,
        fetch_json,
        normalize_measures_payload,
        flatten_measures,
        summarize_consumption,
    )

    start = f.api_start
    end_api = _to_ddmmyyyy(f.lectura_actual + timedelta(days=1))
    base = acl_node_base_url()
    raw = fetch_json(
        f"{base}/nodes/measures/dates",
        params=[("id", NODE_ID), ("start", start), ("end", end_api)],
    )
    norm = normalize_measures_payload(raw, NODE_ID)
    meas = flatten_measures(norm)
    d0, d1 = f.lectura_anterior.date(), f.lectura_actual.date()
    meas = [m for m in meas if d0 <= m.date.date() <= d1]
    s = summarize_consumption(meas)
    return float(s.get("total", 0.0)), int(s.get("dias", 0))


def _generar_grafico_barras(out_png: Path, filas: list[dict]) -> None:
    """Barras agrupadas con ``ax.bar``; dos azules. Izq = app WES, der = facturación."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib import patheffects as pe

    rc = {
        "font.family": "sans-serif",
        "font.sans-serif": ["Segoe UI", "Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 10.5,
        "axes.labelsize": 11,
        "axes.edgecolor": "#94a3b8",
        "axes.linewidth": 0.8,
        "xtick.labelsize": 9.5,
        "ytick.labelsize": 10,
        "legend.fontsize": 10,
        "figure.facecolor": "#f8fafc",
        "axes.facecolor": "#ffffff",
    }

    if not filas:
        return

    cuenta = np.array([float(f["m3_cuenta"]) for f in filas], dtype=float)
    wes = np.array([float(f["m3_wes"]) for f in filas], dtype=float)

    labels_x = [f'N° {f["boleta"]}\n{f["periodo_short"]}\n(izq.: app WES · der.: fact.)' for f in filas]

    col_wes = "#38bdf8"
    col_fact = "#1e3a8f"
    edge = "#ffffff"

    with matplotlib.rc_context(rc):
        fig, ax = plt.subplots(figsize=(13.5, 6.8))
        n = len(filas)
        xg = np.arange(n, dtype=float)
        bw = 0.36
        ymax = float(max(float(cuenta.max()), float(wes.max()), 1.0)) * 1.2

        bar_w = ax.bar(
            xg - bw / 2,
            wes,
            bw,
            color=col_wes,
            edgecolor=edge,
            linewidth=1.15,
            label="Consumo según app WES (m³)",
            zorder=3,
        )
        bar_f = ax.bar(
            xg + bw / 2,
            cuenta,
            bw,
            color=col_fact,
            edgecolor=edge,
            linewidth=1.15,
            label="Consumo facturado Aguas Andinas (m³)",
            zorder=2,
        )

        def _etiqueta_arriba(bar, vals: np.ndarray, *, zeros: bool = False) -> None:
            rects = getattr(bar, "patches", None) or tuple(bar)
            for rect, val in zip(rects, vals):
                hgt = rect.get_height()
                if val <= 0 and zeros:
                    t0 = ax.annotate(
                        _fmt_num(float(val), 0 if abs(val) >= 10 else 1),
                        xy=(rect.get_x() + rect.get_width() / 2, 0),
                        xytext=(0, 11),
                        textcoords="offset points",
                        ha="center",
                        va="bottom",
                        fontsize=8.5,
                        color="#64748b",
                        zorder=5,
                    )
                    t0.set_path_effects([pe.withStroke(linewidth=3, foreground="white")])
                    continue
                if hgt <= 0:
                    continue
                t = ax.annotate(
                    _fmt_num(float(val), 0 if val >= 10 else 1),
                    xy=(rect.get_x() + rect.get_width() / 2, hgt),
                    xytext=(0, 6),
                    textcoords="offset points",
                    ha="center",
                    va="bottom",
                    fontsize=9.5,
                    color="#0f172a",
                    fontweight="600",
                    zorder=5,
                )
                t.set_path_effects([pe.withStroke(linewidth=3.5, foreground="white")])

        _etiqueta_arriba(bar_w, wes, zeros=True)
        _etiqueta_arriba(bar_f, cuenta, zeros=False)

        ax.set_ylim(0, ymax)
        ax.set_xticks(xg)
        ax.set_xticklabels(labels_x, ha="center", linespacing=1.35)
        ax.set_xlabel("Factura y fin de período (lectura actual)", color="#475569", labelpad=12)
        ax.set_ylabel("Consumo total del período (m³)", color="#475569", labelpad=8)

        ax.set_axisbelow(True)
        ax.yaxis.grid(True, linestyle="--", linewidth=0.7, color="#cbd5e1", alpha=0.9, zorder=0)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.legend(loc="upper right", framealpha=0.97, fancybox=True, edgecolor="#e2e8f0")

        fig.suptitle(
            "Por factura: app WES (izquierda) vs facturación (derecha) — ICCP",
            fontsize=14,
            fontweight="600",
            color="#0f172a",
            y=0.97,
        )
        fig.text(
            0.5,
            0.935,
            f"Nodo {NODE_ID} — {NODE_NOMBRE} — solo lecturas ≥ {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}",
            ha="center",
            fontsize=10,
            color="#64748b",
        )

        fig.subplots_adjust(bottom=0.26, left=0.07, right=0.98, top=0.82)
        fig.savefig(out_png, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor(), edgecolor="none")
        plt.close(fig)


def _generar_grafico_totales(
    out_png: Path,
    *,
    total_wes: float,
    total_cuenta: float,
    n_boletas: int,
) -> None:
    """Dos barras: total App WES vs total facturación (suma de boletas incluidas)."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib import patheffects as pe

    rc = {
        "font.family": "sans-serif",
        "font.sans-serif": ["Segoe UI", "Arial", "Helvetica", "DejaVu Sans"],
        "font.size": 11,
        "axes.labelsize": 11,
        "axes.edgecolor": "#94a3b8",
        "figure.facecolor": "#f8fafc",
        "axes.facecolor": "#ffffff",
    }
    col_wes = "#38bdf8"
    col_fact = "#1e3a8f"
    labels = ["Total App WES\n(suma m³)", "Total facturación\ncuenta (suma m³)"]
    vals = np.array([float(total_wes), float(total_cuenta)], dtype=float)
    x = np.arange(2, dtype=float)
    ymax = float(max(vals.max(), 1.0)) * 1.22

    with matplotlib.rc_context(rc):
        fig, ax = plt.subplots(figsize=(8.2, 5.4))
        bars = ax.bar(x, vals, width=0.52, color=[col_wes, col_fact], edgecolor="#ffffff", linewidth=1.2, zorder=2)
        for rect, val in zip(bars, vals):
            if val <= 0:
                continue
            t = ax.annotate(
                _fmt_num(float(val), 1),
                xy=(rect.get_x() + rect.get_width() / 2, float(val)),
                xytext=(0, 6),
                textcoords="offset points",
                ha="center",
                va="bottom",
                fontsize=11,
                fontweight="600",
                color="#0f172a",
                zorder=5,
            )
            t.set_path_effects([pe.withStroke(linewidth=3.5, foreground="white")])

        ax.set_xticks(x)
        ax.set_xticklabels(labels, ha="center", linespacing=1.25)
        ax.set_ylabel("m³ acumulados", color="#475569", labelpad=10)
        ax.set_ylim(0, ymax)
        ax.set_axisbelow(True)
        ax.yaxis.grid(True, linestyle="--", linewidth=0.7, color="#cbd5e1", alpha=0.9, zorder=0)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        fig.suptitle(
            "Total acumulado: App WES vs facturación en cuenta",
            fontsize=14,
            fontweight="600",
            color="#0f172a",
            y=0.96,
        )
        fig.text(
            0.5,
            0.90,
            (
                f"Nodo {NODE_ID} — {n_boletas} boleta(s) — lecturas anterior y actual ≥ "
                f"{MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}"
            ),
            ha="center",
            fontsize=10,
            color="#64748b",
        )
        fig.subplots_adjust(bottom=0.18, left=0.12, right=0.97, top=0.82)
        fig.savefig(out_png, dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor(), edgecolor="none")
        plt.close(fig)


def generar_reporte() -> tuple[Path, Path | None]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ahora = _ahora_informe_local()
    run_id = ahora.strftime("%Y%m%d_%H%M%S")

    facturas_todas = _cargar_facturas()
    facturas, facturas_excl = _particionar_por_inicio_monitoreo(facturas_todas, MONITOREO_WES_DESDE)
    if not facturas:
        raise ValueError(
            f"No hay facturas con lectura anterior y lectura actual ≥ {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}."
        )

    doc = Document()
    _doc_establecer_metadatos(doc, ahora)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Comparación App WES vs Facturaciones Aguas Andinas", level=1)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(31, 71, 136)

    p = doc.add_paragraph()
    p.add_run("Empresa: ").bold = True
    p.add_run(EMPRESA)
    p2 = doc.add_paragraph()
    p2.add_run("Nodo: ").bold = True
    p2.add_run(f"{NODE_ID} — {NODE_NOMBRE}")
    p3 = doc.add_paragraph()
    p3.add_run("Fuente facturas: ").bold = True
    p3.add_run(str(IN_DIR))
    p4 = doc.add_paragraph()
    p4.add_run("Generado: ").bold = True
    p4.add_run(ahora.strftime("%d-%m-%Y %H:%M"))

    doc.add_paragraph()
    doc.add_heading("1. Alcance y criterio de comparación", level=1)
    doc.add_paragraph(
        "Para cada boleta, se compara el consumo total facturado (m³) con el consumo total registrado por WES "
        "en el nodo indicado, para el período entre «Lectura anterior» y «Lectura actual» informados en la factura."
    )
    doc.add_paragraph(
        "Para la lectura operacional del informe, la evaluación con sentido se ancla al inicio de monitoreo de "
        f"referencia ({MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}) en adelante. En el cruce con el Excel comparativo "
        "facturación vs app, en el histórico revisado el único ciclo mensual que aparece alineado («cuadrado») antes "
        "de la fase estable posterior a la transición es el período entre lecturas "
        f"{PERIODO_ALINEADO_EXCEL_INI.strftime('%d-%m-%Y')} y {PERIODO_ALINEADO_EXCEL_FIN.strftime('%d-%m-%Y')}; "
        "el resto de ciclos en la etapa «fuera de rango» muestra desalineación esperable según bypass e intermitencia "
        "de la sala de bombas (cuando aplica la misma lógica operativa que en ICCO Renca)."
    )
    doc.add_paragraph(
        f"Las tablas y el gráfico de este documento incluyen únicamente boletas cuya «Lectura anterior» y "
        f"«Lectura actual» son el {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')} o posteriores: el consumo facturado en "
        "ciclos anteriores o con lecturas previas a esa fecha suele corresponder a liquidación por promedio estimado "
        "(u homologación sin medición comparable en el punto WES en ese tramo), por lo que no se contrasta aquí con "
        "la app."
    )

    doc.add_heading("1.1 Boletas fuera del cuadro comparativo (antes del 26-08-2025 o sin ambas lecturas desde esa fecha)", level=2)
    if facturas_excl:
        doc.add_paragraph(
            "Las siguientes boletas no entran en las tablas ni en el gráfico. En esos períodos Aguas Andinas "
            "liquidó consumo según promedio estimado o lectura no asociada de punta a punta al monitoreo WES desde "
            f"el {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}."
        )
        for f in sorted(facturas_excl, key=lambda x: x.lectura_actual):
            doc.add_paragraph(f"N° {f.boleta} — {f.periodo_txt} (m³ facturados en boleta: {f.m3_cuenta})", style="List Bullet")
    else:
        doc.add_paragraph("No hay boletas excluidas por este criterio.")

    doc.add_heading("2. Resumen por período de factura (lecturas ≥ 26-08-2025)", level=1)
    headers = [
        "Período (lecturas)",
        "Emisión",
        "N° factura",
        "Medidor",
        "m³ facturado",
        "m³ App (WES)",
        "Días (WES)",
        "Dif (fact - WES)",
        "% Dif vs Fact",
    ]
    MAX_FILAS_TABLA_POR_PAGINA = 10

    filas_bypass: list[FacturaAA] = []
    filas_transicion: list[FacturaAA] = []
    filas_alineado: list[FacturaAA] = []
    for f in facturas:
        d0 = f.lectura_anterior.date()
        d1 = f.lectura_actual.date()
        overlap_trans = not (d1 < TRANSICION_INICIO or d0 > TRANSICION_FIN)
        if overlap_trans:
            filas_transicion.append(f)
        elif d1 < TRANSICION_INICIO:
            filas_bypass.append(f)
        else:
            filas_alineado.append(f)

    def _filas_dict(fs: list[FacturaAA]) -> list[dict]:
        out: list[dict] = []
        for f in fs:
            m3_wes, dias = _fetch_wes_total(f)
            diff = float(f.m3_cuenta) - float(m3_wes)
            pct_diff = (100.0 * float(diff) / float(f.m3_cuenta)) if f.m3_cuenta else 0.0
            out.append(
                {
                    "periodo": f.periodo_txt,
                    "periodo_short": _fmt_periodo(f.lectura_actual),
                    "lectura_fin": f.lectura_actual,
                    "emision": f.emision.strftime("%d-%b-%Y")
                    .lower()
                    .replace("apr", "abr")
                    .replace("aug", "ago")
                    .replace("dec", "dic"),
                    "boleta": f.boleta,
                    "medidor": f.medidor or "-",
                    "m3_cuenta": int(f.m3_cuenta),
                    "m3_wes": float(m3_wes),
                    "dias_wes": int(dias),
                    "diff": float(diff),
                    "pct": float(pct_diff),
                }
            )
        return out

    filas_bypass_d = _filas_dict(filas_bypass)
    filas_transicion_d = _filas_dict(filas_transicion)
    filas_alineado_d = _filas_dict(filas_alineado)

    def _tbl_pr_set_full_width_fixed(table) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        for tag_nm in ("tblW", "tblLayout"):
            tq = qn(f"w:{tag_nm}")
            for el in list(tbl_pr):
                if el.tag == tq:
                    tbl_pr.remove(el)
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), "5000")
        tbl_w.set(qn("w:type"), "pct")
        tbl_pr.append(tbl_w)
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), "fixed")
        tbl_pr.append(tbl_layout)

    def _tbl_set_repeat_header(table) -> None:
        try:
            tr = table.rows[0]._tr
        except Exception:
            return
        tr_pr = tr.get_or_add_trPr()
        for el in list(tr_pr):
            if el.tag == qn("w:tblHeader"):
                tr_pr.remove(el)
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        tr_pr.append(hdr)

    def _add_table(title_txt: str, rows: list[dict]) -> None:
        p_tit = doc.add_paragraph()
        p_tit.add_run(title_txt).bold = True
        p_tit.paragraph_format.keep_with_next = True
        if not rows:
            doc.add_paragraph("Sin facturas en este rango.")
            return

        chunks: list[list[dict]] = [
            rows[i : i + MAX_FILAS_TABLA_POR_PAGINA]
            for i in range(0, len(rows), MAX_FILAS_TABLA_POR_PAGINA)
        ]
        for chunk_idx, chunk_rows in enumerate(chunks):
            if chunk_idx > 0:
                doc.add_page_break()
                p_ct = doc.add_paragraph()
                p_ct.add_run(f"{title_txt} (continuación)").bold = True
                p_ct.paragraph_format.keep_with_next = True

            table = doc.add_table(rows=1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            _tbl_pr_set_full_width_fixed(table)
            _tbl_set_repeat_header(table)
            hdr = table.rows[0].cells
            for i, h in enumerate(headers):
                rr = hdr[i].paragraphs[0].add_run(h)
                rr.bold = True
                rr.font.size = Pt(9)
            for row in chunk_rows:
                cells = table.add_row().cells
                cells[0].text = row["periodo"]
                cells[1].text = row["emision"]
                cells[2].text = str(row["boleta"])
                cells[3].text = str(row["medidor"])
                cells[4].text = _fmt_num(float(row["m3_cuenta"]), 0)
                cells[5].text = _fmt_num(float(row["m3_wes"]), 1)
                cells[6].text = str(row["dias_wes"])
                cells[7].text = _fmt_num(float(row["diff"]), 1)
                cells[8].text = _fmt_num(float(row["pct"]), 1) + "%"

            table.allow_autofit = False
            for r in table.rows:
                try:
                    r.allow_break_across_pages = False
                except Exception:
                    pass
                for c in r.cells:
                    for pp in c.paragraphs:
                        pp.paragraph_format.space_before = Pt(0)
                        pp.paragraph_format.space_after = Pt(0)
                        for run in pp.runs:
                            run.font.size = Pt(9)

    doc.add_paragraph(
        "Tabla A — Períodos fuera de rango (bypass), solo lecturas ≥ "
        f"{MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}: comparación no representativa (operación con bypass/intermitencia). "
        f"En Excel, único ciclo alineado en esa etapa: lecturas {PERIODO_ALINEADO_EXCEL_INI.strftime('%d-%m-%Y')} → "
        f"{PERIODO_ALINEADO_EXCEL_FIN.strftime('%d-%m-%Y')}."
    ).paragraph_format.keep_with_next = True
    _add_table("Períodos fuera de rango (bypass)", filas_bypass_d)

    doc.add_paragraph()
    doc.add_paragraph(
        "Tabla B — Períodos de transición (11-dic-2025 a 08-ene-2026), solo lecturas ≥ "
        f"{MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}: ciclos con ajustes y estabilización."
    ).paragraph_format.keep_with_next = True
    _add_table("Períodos de transición (11-dic-2025 a 08-ene-2026)", filas_transicion_d)

    doc.add_paragraph()
    doc.add_paragraph(
        "Tabla C — Períodos alineados (posterior a 08-ene-2026), solo lecturas ≥ "
        f"{MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}: referencia principal para comparación."
    ).paragraph_format.keep_with_next = True
    _add_table("Períodos alineados (posterior a 08-ene-2026)", filas_alineado_d)

    doc.add_paragraph()
    filas_graf = [dict(r) for r in filas_bypass_d + filas_transicion_d + filas_alineado_d]
    filas_graf.sort(key=lambda r: r["lectura_fin"])
    for r in filas_graf:
        r.pop("lectura_fin", None)
    img_barras = OUT_DIR / f"chart_barras_iccp_todos_{run_id}.png"
    _generar_grafico_barras(img_barras, filas_graf)
    doc.add_paragraph(
        "Gráfico consolidado — todas las boletas incluidas en el informe, en orden cronológico por lectura actual; "
        "en cada grupo: izquierda app WES, derecha facturación (dos tonos azules). "
        f"Solo períodos con ambas lecturas ≥ {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}."
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if filas_graf:
        doc.add_picture(str(img_barras), width=Inches(6.7))

    def _grafico_categoria(titulo: str, filas_d: list[dict], sufijo: str) -> None:
        if not filas_d:
            return
        copia = [dict(r) for r in filas_d]
        copia.sort(key=lambda r: r["lectura_fin"])
        for r in copia:
            r.pop("lectura_fin", None)
        img = OUT_DIR / f"chart_barras_iccp_{sufijo}_{run_id}.png"
        _generar_grafico_barras(img, copia)
        doc.add_paragraph()
        pg = doc.add_paragraph()
        pg.add_run(titulo).bold = True
        doc.add_picture(str(img), width=Inches(6.7))

    _grafico_categoria("Gráfico — Tabla A (bypass)", filas_bypass_d, "bypass")
    _grafico_categoria("Gráfico — Tabla B (transición)", filas_transicion_d, "transicion")
    _grafico_categoria("Gráfico — Tabla C (alineados)", filas_alineado_d, "alineados")

    doc.add_paragraph()
    doc.add_heading(
        f"2.3 Total acumulado — cuenta (facturación) vs App WES (≥ {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')})",
        level=2,
    )
    doc.add_paragraph(
        f"Suma de todos los períodos incluidos en este informe: solo boletas cuya lectura anterior y lectura actual "
        f"son el {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')} o posteriores. "
        "La columna «cuenta» es la suma del consumo total facturado (m³) informado en esas boletas; «App WES» es la "
        "suma del consumo registrado por la API en cada período entre las mismas lecturas."
    )
    rows_acum = filas_bypass_d + filas_transicion_d + filas_alineado_d
    n_boletas_acum = len(rows_acum)
    tot_cuenta = sum(float(r["m3_cuenta"]) for r in rows_acum)
    tot_wes = sum(float(r["m3_wes"]) for r in rows_acum)
    tot_diff = tot_cuenta - tot_wes
    pct_vs_fact = (100.0 * tot_diff / tot_cuenta) if tot_cuenta else 0.0

    ttot = doc.add_table(rows=1, cols=2)
    ttot.style = "Table Grid"
    ttot.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_pr_set_full_width_fixed(ttot)
    h0 = ttot.rows[0].cells
    h0[0].paragraphs[0].add_run("Indicador").bold = True
    h0[1].paragraphs[0].add_run("Valor").bold = True
    for lab, val in (
        ("Boletas incluidas en el acumulado (N°)", str(n_boletas_acum)),
        ("Total m³ facturación (cuenta, suma boletas)", _fmt_num(tot_cuenta, 1)),
        ("Total m³ App WES (suma períodos)", _fmt_num(tot_wes, 1)),
        ("Diferencia (facturación − App WES)", _fmt_num(tot_diff, 1)),
        ("% diferencia vs facturación", _fmt_num(pct_vs_fact, 1) + "%"),
    ):
        row_cells = ttot.add_row().cells
        row_cells[0].text = lab
        row_cells[1].text = val
    for r in ttot.rows:
        try:
            r.allow_break_across_pages = False
        except Exception:
            pass
        for c in r.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after = Pt(0)
                for run in pp.runs:
                    run.font.size = Pt(9)

    img_tot = OUT_DIR / f"chart_barras_iccp_total_acumulado_{run_id}.png"
    _generar_grafico_totales(
        img_tot,
        total_wes=tot_wes,
        total_cuenta=tot_cuenta,
        n_boletas=n_boletas_acum,
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Gráfico — total acumulado: barra izquierda suma App WES, barra derecha suma facturación en cuenta; "
        f"misma base que la tabla (solo facturaciones posteriores al {MONITOREO_WES_DESDE.strftime('%d-%m-%Y')} "
        "en ambas lecturas del período)."
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_picture(str(img_tot), width=Inches(5.8))

    doc.add_heading("3. Referencias", level=1)
    doc.add_paragraph(
        "Facturas incluidas en tablas, gráficos por período y total acumulado (lecturas ≥ "
        f"{MONITOREO_WES_DESDE.strftime('%d-%m-%Y')}), por lectura actual:\n- "
        + "\n- ".join([str(f.boleta) for f in sorted(facturas, key=lambda x: x.lectura_actual)])
    )

    out_docx = OUT_DIR / "Comparacion_App_vs_Facturaciones_ICCP_Renca.docx"
    if out_docx.exists():
        try:
            out_docx.unlink()
        except OSError:
            pass
    doc.save(out_docx)
    _docx_parchear_core_hora_local(out_docx, ahora)

    out_pdf = out_docx.with_suffix(".pdf")
    if out_pdf.exists():
        try:
            out_pdf.unlink()
        except OSError:
            pass
    out_pdf = convertir_word_a_pdf(out_docx)
    return out_docx, out_pdf


def main() -> None:
    docx, pdf = generar_reporte()
    print(f"[OK] DOCX: {docx}")
    print(f"[OK] PDF: {pdf}" if pdf else "[ADVERTENCIA] No se generó PDF.")


if __name__ == "__main__":
    main()

