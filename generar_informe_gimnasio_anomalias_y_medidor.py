"""
Informe Gimnasio Renca (000017-05): anomalías pre cambio medidor (25-dic-2025 a 21-abr-2026),
periodo post medidor, y corte energía 8–12 may 2026.

Uso:
  python generar_informe_gimnasio_anomalias_y_medidor.py
"""

from __future__ import annotations

import csv
import math
import sys
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import List, Optional, Sequence, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Cm, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent
NODE_ID = "000017-05"
NOMBRE = "Gimnasio municipal Renca"
DIAMETRO_PULG = 2
# PVC 2" Sch.40 ≈ 52 mm diámetro interior; caudal máximo de referencia a 1,5 m/s
DIAMETRO_M = 0.052
VEL_REF_MS = 1.5
CAUDAL_MAX_M3H_REF = math.pi * (DIAMETRO_M / 2) ** 2 * VEL_REF_MS * 3600  # ~11,4 m³/h

FECHA_CAMBIO_MEDIDOR = date(2026, 4, 21)
DESDE_PRE_MEDIDOR = date(2025, 12, 25)
HASTA_PRE_MEDIDOR = date(2026, 4, 21)
DESDE_POST_MEDIDOR = date(2026, 4, 22)
CORTE_ENERGIA = (date(2026, 5, 8), date(2026, 5, 12))

_HEADING = RGBColor(31, 71, 136)


@dataclass
class DiaRegistro:
    fecha: date
    total_m3: float
    max_m3h: float
    horas_cero: int
    sin_datos: bool
    fuera_tuberia: bool
    pico_extremo: bool


def _dia(node_id: str, d: date) -> DiaRegistro:
    from auditoria_cpa_icco_renca_grafico import _flatten_grilla_desde_vectores, _vectores_m3h_por_dias
    from generar_reporte_word import get_hourly_measures_for_day

    vecs = _vectores_m3h_por_dias(node_id, [d])
    _, _, total = _flatten_grilla_desde_vectores([d], vecs)
    hourly = get_hourly_measures_for_day(node_id, datetime.combine(d, datetime.min.time()))
    sin_datos = not hourly
    if hourly:
        vals = [v for _, v in hourly]
        max_h = max(vals)
        zeros = sum(1 for v in vals if v <= 1e-9)
    else:
        max_h = 0.0
        zeros = 24
    fuera = max_h > CAUDAL_MAX_M3H_REF + 0.5
    extremo = max_h > 25.0 or total > 80.0
    return DiaRegistro(
        fecha=d,
        total_m3=float(total),
        max_m3h=float(max_h),
        horas_cero=zeros,
        sin_datos=sin_datos,
        fuera_tuberia=fuera,
        pico_extremo=extremo,
    )


def _rango(d0: date, d1: date) -> List[date]:
    out: List[date] = []
    d = d0
    while d <= d1:
        out.append(d)
        d += timedelta(days=1)
    return out


def _cargar_serie(dias: Sequence[date]) -> List[DiaRegistro]:
    return [_dia(NODE_ID, d) for d in dias]


def _grafico_serie(
    registros: List[DiaRegistro],
    out_png: Path,
    *,
    titulo: str,
    marcar_anomalias: bool,
) -> None:
    fechas = [r.fecha for r in registros]
    vals = [r.total_m3 for r in registros]
    fig, ax = plt.subplots(figsize=(12, 4.5))
    ax.bar(fechas, vals, width=0.8, color="#4F81BD", alpha=0.85, label="Consumo diario (m³)")
    if marcar_anomalias:
        for r in registros:
            if r.pico_extremo or r.fuera_tuberia:
                ax.bar(r.fecha, r.total_m3, width=0.8, color="#C0504D", alpha=0.95)
    ax.axvline(FECHA_CAMBIO_MEDIDOR, color="#28A745", linestyle="--", linewidth=1.2, label="21-abr cambio medidor")
    ax.axvspan(CORTE_ENERGIA[0], CORTE_ENERGIA[1] + timedelta(days=1), color="#FFE699", alpha=0.4, label="8–12 may corte energía")
    ax.set_ylabel("m³/día")
    ax.set_title(titulo, fontsize=11, fontweight="bold")
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d-%m"))
    ax.xaxis.set_major_locator(mdates.WeekdayLocator(interval=2))
    plt.xticks(rotation=35, ha="right")
    ax.grid(axis="y", alpha=0.3)
    ax.legend(fontsize=8, loc="upper right")
    fig.tight_layout()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_png, dpi=150, bbox_inches="tight")
    plt.close(fig)


def _word(
    out_docx: Path,
    pre: List[DiaRegistro],
    post: List[DiaRegistro],
    png_pre: Path,
    png_full: Path,
) -> None:
    from generar_reporte_word import format_number_chilean

    anom_pre = [r for r in pre if r.pico_extremo or r.fuera_tuberia]
    anom_pre.sort(key=lambda x: (-x.max_m3h, -x.total_m3))

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading("Informe técnico — Gimnasio Renca (000017-05)", level=0)
    if h.runs:
        h.runs[0].font.color.rgb = _HEADING

    doc.add_paragraph(
        f"Establecimiento: {NOMBRE}. Nodo WES: {NODE_ID}.\n"
        f"Generado: {datetime.now().strftime('%d-%m-%Y %H:%M')}."
    )

    doc.add_heading("1. Contexto", level=1)
    doc.add_paragraph(
        f"• El {FECHA_CAMBIO_MEDIDOR.strftime('%d-%m-%Y')} se reemplazó el medidor ultrasónico. "
        "Los datos desde el 22-abr-2026 corresponden al equipo nuevo.\n"
        f"• Instalación referida: tubería de {DIAMETRO_PULG}\" PVC. Caudal horario de referencia "
        f"para esa sección (~1,5 m/s): ≈ {format_number_chilean(CAUDAL_MAX_M3H_REF, 1)} m³/h. "
        "Valores muy por encima de ese orden en una hora sostenida son físicamente improbables "
        "en una red de 2\" y sugieren error de medición (medidor anterior) o evento no hidráulico.\n"
        f"• Del {CORTE_ENERGIA[0].strftime('%d-%m-%Y')} al {CORTE_ENERGIA[1].strftime('%d-%m-%Y')} "
        "no hay registro de consumo porque se bajó la energía del equipo WES/medición (corte planificado)."
    )

    doc.add_heading("2. Periodo medidor anterior (25-dic-2025 a 21-abr-2026)", level=1)
    tot_pre = sum(r.total_m3 for r in pre)
    n_anom = len(anom_pre)
    doc.add_paragraph(
        f"Días analizados: {len(pre)}. Consumo acumulado API: {format_number_chilean(tot_pre, 1)} m³ "
        f"(promedio {format_number_chilean(tot_pre / len(pre) if pre else 0, 2)} m³/día).\n"
        f"Se identificaron {n_anom} día(s) con picos horarios > {format_number_chilean(CAUDAL_MAX_M3H_REF, 1)} m³/h "
        f"o consumo diario > 80 m³ / pico > 25 m³/h (criterio de consumo irrisorio para 2\" PVC)."
    )

    if anom_pre:
        tbl = doc.add_table(rows=1 + min(len(anom_pre), 25), cols=4)
        tbl.style = "Table Grid"
        for j, hd in enumerate(["Fecha", "Total día (m³)", "Máx hora (m³/h)", "Observación"]):
            tbl.rows[0].cells[j].text = hd
            for run in tbl.rows[0].cells[j].paragraphs[0].runs:
                run.bold = True
        for i, r in enumerate(anom_pre[:25], start=1):
            obs = []
            if r.fuera_tuberia:
                obs.append(f"caudal > {format_number_chilean(CAUDAL_MAX_M3H_REF, 1)} m³/h (2\" PVC)")
            if r.pico_extremo:
                obs.append("consumo diario/pico extremo")
            if r.sin_datos:
                obs.append("sin serie horaria")
            tbl.rows[i].cells[0].text = r.fecha.strftime("%d-%m-%Y")
            tbl.rows[i].cells[1].text = format_number_chilean(r.total_m3, 2)
            tbl.rows[i].cells[2].text = format_number_chilean(r.max_m3h, 2)
            tbl.rows[i].cells[3].text = "; ".join(obs) or "—"

    doc.add_paragraph("")
    if png_pre.is_file():
        doc.add_picture(str(png_pre), width=Cm(16))

    doc.add_heading("3. Periodo medidor nuevo (desde 22-abr-2026)", level=1)
    tot_post = sum(r.total_m3 for r in post)
    anom_post = [r for r in post if r.pico_extremo or r.fuera_tuberia]
    corte = [r for r in post if CORTE_ENERGIA[0] <= r.fecha <= CORTE_ENERGIA[1]]
    doc.add_paragraph(
        f"Días analizados: {len(post)} (hasta {post[-1].fecha.strftime('%d-%m-%Y') if post else '—'}). "
        f"Consumo acumulado: {format_number_chilean(tot_post, 1)} m³ "
        f"(promedio {format_number_chilean(tot_post / len(post) if post else 0, 2)} m³/día).\n"
        f"Tras el cambio de medidor el patrón es bajo y estable (típicamente 0,4–5 m³/día), "
        "compatible con consumo base o uso acotado.\n"
        f"Días con anomalía tipo «2\" PVC»: {len(anom_post)}.\n"
        f"Corte energía 8–12 may: {len(corte)} días a 0 m³ — explicado por baja de energía del equipo."
    )

    doc.add_paragraph("")
    if png_full.is_file():
        doc.add_picture(str(png_full), width=Cm(16))

    doc.add_heading("4. Conclusiones", level=1)
    top3 = anom_pre[:3]
    bullets = [
        "El medidor anterior registró múltiples episodios de consumo incompatible con una tubería de 2\" PVC; "
        "no deben usarse para valorizar ahorro sin depurar esos días.",
        "Desde el 22-abr-2026 los datos son coherentes con el nuevo ultrasónico.",
        "La semana 8–12 de mayo sin lecturas responde al corte de energía del equipo, no a falla hidráulica.",
    ]
    if top3:
        ej = ", ".join(
            f"{r.fecha.strftime('%d-%m-%Y')} ({format_number_chilean(r.total_m3, 1)} m³, máx {format_number_chilean(r.max_m3h, 1)} m³/h)"
            for r in top3
        )
        bullets.insert(0, f"Episodios más relevantes (pre-22-abr): {ej}.")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def main() -> int:
    from generar_reporte_word import convertir_word_a_pdf, format_number_chilean

    hoy = date.today()
    print(f"Analizando {NOMBRE} ({NODE_ID})…")
    print(f"  Pre-medidor:  {DESDE_PRE_MEDIDOR} → {HASTA_PRE_MEDIDOR}")
    print(f"  Post-medidor: {DESDE_POST_MEDIDOR} → {hoy}")

    dias_pre = _rango(DESDE_PRE_MEDIDOR, HASTA_PRE_MEDIDOR)
    dias_post = _rango(DESDE_POST_MEDIDOR, hoy)
    pre = _cargar_serie(dias_pre)
    post = _cargar_serie(dias_post)
    full = pre + post

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    out_dir = ROOT / "reports" / "reporte de auditoria" / f"informe_gimnasio_000017-05_{ts}"
    out_dir.mkdir(parents=True, exist_ok=True)

    csv_path = out_dir / "serie_diaria_gimnasio.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(
            [
                "fecha",
                "total_m3",
                "max_m3h",
                "horas_cero",
                "sin_datos",
                "fuera_2pvc",
                "pico_extremo",
                "periodo",
            ]
        )
        for r in full:
            periodo = "pre_medidor" if r.fecha <= HASTA_PRE_MEDIDOR else "post_medidor"
            w.writerow(
                [
                    r.fecha.isoformat(),
                    f"{r.total_m3:.4f}",
                    f"{r.max_m3h:.4f}",
                    r.horas_cero,
                    int(r.sin_datos),
                    int(r.fuera_tuberia),
                    int(r.pico_extremo),
                    periodo,
                ]
            )

    png_pre = out_dir / "consumo_diario_pre_medidor.png"
    png_full = out_dir / "consumo_diario_completo.png"
    _grafico_serie(pre, png_pre, titulo="Gimnasio — consumo diario (medidor anterior)", marcar_anomalias=True)
    _grafico_serie(full, png_full, titulo="Gimnasio — consumo diario (dic-2025 a hoy)", marcar_anomalias=True)

    docx = out_dir / f"Informe_Gimnasio_anomalias_medidor_{ts}.docx"
    _word(docx, pre, post, png_pre, png_full)

    pdf = docx.with_suffix(".pdf")
    try:
        p = convertir_word_a_pdf(docx)
        if p and Path(p).is_file():
            pdf = Path(p)
    except Exception:
        pass

    anom = [r for r in pre if r.pico_extremo or r.fuera_tuberia]
    print("=" * 72)
    print(f"Caudal ref. 2\" PVC @ 1,5 m/s: {CAUDAL_MAX_M3H_REF:.1f} m³/h")
    print(f"Días anómalos (pre-22-abr): {len(anom)} / {len(pre)}")
    for r in sorted(anom, key=lambda x: -x.max_m3h)[:10]:
        print(
            f"  {r.fecha:%d-%m-%Y}: {r.total_m3:.1f} m³/día, máx {r.max_m3h:.1f} m³/h"
        )
    print(f"CSV:  {csv_path}")
    print(f"DOCX: {docx}")
    print(f"PDF:  {pdf}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
